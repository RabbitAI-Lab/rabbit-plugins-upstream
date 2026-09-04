#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
qq_hw.py —— QQ 群作业 → Word/PDF 文档 一站式工具

子命令:
  doctor           环境自检：CLI/daemon/扩展连通性/登录态/bkn 缓存/现存产物，一次说清
                   （定时任务每轮开头先跑它，避免整轮耗在错误方向上）
  bkn              探测 bkn（会打开作业页；需已登录）并写入配置。--gid 支持逗号分隔多群
  list             拉取作业列表（多群合并）-> hw_list.json
  day <YYYY-MM-DD> 拉取该天作业详情（含图片 URL，多群合并）-> hw_day_<date>.json
  docx <YYYY-MM-DD> [--scale N] [--courses 语文,数学] [--text-only] [--allow-multi]
                   生成 Word 文档；默认单页、全部科目、含图片
  pdf  <YYYY-MM-DD> docx -> PDF（Word COM 回退方案；优先用本地 pdf skill）
  pages <file.docx> 统计页数（需本机装 Word）

依赖: qqbrowser-skill CLI 在 PATH；pip install python-docx pillow
注意: 浏览器自动化必须在**沙箱外**运行（否则安全删除机制会 fail-closed）
"""
import argparse, io, json, os, re, subprocess, sys, urllib.request
from datetime import datetime

HERE = os.getcwd()
CFG = os.path.join(HERE, "qq_hw.json")
IMG_DIR = os.path.join(HERE, "hw_imgs")

CLI = os.environ.get("QQB_CLI", "qqbrowser-skill")
API_LIST = "https://qun.qq.com/cgi-bin/homework/hw/get_hw_list.fcg"
API_DETAIL = "https://qun.qq.com/cgi-bin/homework/hw/get_hw_detail.fcg"
INDEX_URL = "https://qun.qq.com/homework/features/index.html?gid={gid}"
LOGIN_URL = "https://qun.qq.com/#/login"

# 服务端会强制下线旧版 CLI：命中后所有浏览器子命令都返回这条提示，
# 与「扩展协议不兼容 / 未登录」症状相似，务必优先识别（2026-09-04 踩坑）
BLOCKED_HINT = "blocked until you upgrade"


class CliBlocked(Exception):
    """qqbrowser-skill 不可用（版本被服务端下线 / 可执行文件找不到）。"""
    def __init__(self, help_text):
        super().__init__("qqbrowser-skill unavailable")
        self.help_text = help_text


def _blocked_help():
    return "\n".join([
        "=" * 60,
        "❌ qqbrowser-skill 版本过旧，已被服务端强制下线（<=1.5.4）",
        "   症状：所有浏览器子命令都返回 'Browser skill commands are blocked until you upgrade'。",
        "   这不是扩展协议问题、也不是未登录 —— 别去查扩展 / 别让用户重新登录。",
        "",
        "   修复（三步，缺一不可）：",
        "     1) pip install --index-url https://pypi.org/simple --upgrade qqbrowser-skill",
        "        （必须显式指定官方源：本地镜像常滞后，装完还是旧版就白装）",
        "     2) qqbrowser-skill stop",
        "     3) qqbrowser-skill serve --daemon",
        "        （旧 daemon 仍在跑旧代码，不重启不会生效）",
        "",
        "   ⚠️ 升到 1.5.6+ 后 `status` 的 Connected clients 可能恒为 0，",
        "      但业务命令照常可用 —— 不要死等 clients≥1，",
        "      直接跑一次 `qq_hw.py doctor` 或 `browser_go_to_url` 验证即可。",
        "=" * 60,
    ])


def _cli_not_found_help():
    return "\n".join([
        "=" * 60,
        "❌ 找不到 qqbrowser-skill 可执行文件（当前解析为：%s）" % CLI,
        "   隔离 venv 安装时 CLI 通常不在 PATH，会抛 FileNotFoundError: [WinError 2]。",
        "",
        "   修复：用 QQB_CLI 指向可执行文件绝对路径，例如：",
        "     QQB_CLI=\"<venv>/Scripts/qqbrowser-skill.exe\" python scripts/qq_hw.py doctor",
        "   不知道路径就先跑：python -c \"import shutil;print(shutil.which('qqbrowser-skill'))\"",
        "=" * 60,
    ])

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                  "(KHTML, like Gecko) Chrome/120.0 Safari/537.36",
    "Referer": "https://qun.qq.com/",
}

# ---------------- 配置 ----------------

def load_cfg():
    if os.path.exists(CFG):
        with open(CFG, encoding="utf-8") as f:
            return json.load(f)
    return {"gid": "", "bkn": "", "gids": []}


def save_cfg(cfg):
    with open(CFG, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=1)


def parse_gids(arg, cfg):
    """解析群号列表：支持 --gid 逗号分隔多群，或从配置 gid/gids 读取。"""
    raw = arg or cfg.get("gids") or cfg.get("gid") or []
    if isinstance(raw, str):
        raw = [x.strip() for x in raw.split(",") if x.strip()]
    elif isinstance(raw, (int, float)):
        raw = [str(int(raw))]
    elif not isinstance(raw, (list, tuple)):
        raw = []
    return [str(x) for x in raw if str(x).strip()]


# ---------------- 浏览器交互 ----------------

def _run(args, timeout=300, check_blocked=True):
    try:
        p = subprocess.run([CLI] + args, capture_output=True, text=True,
                           encoding="utf-8", errors="replace", cwd=HERE, timeout=timeout)
    except FileNotFoundError:
        raise CliBlocked(_cli_not_found_help())
    out = (p.stdout or "") + (p.stderr or "")
    if check_blocked and BLOCKED_HINT in out:
        raise CliBlocked(_blocked_help())
    return out


def extract_result(raw):
    """从 qqbrowser-skill CLI 输出中取出 eval_content_js 的返回值。
    返回值夹在 'Result: ' 与 '>>>>>' 之间，且内部引号被转义成 \\" """
    i = raw.find("Result: ")
    if i < 0:
        raise RuntimeError("CLI 输出中没有 'Result:'，原始输出前 300 字符:\n" + raw[:300])
    i += len("Result: ")
    j = raw.find(">>>>>", i)
    chunk = (raw[i:j] if j > 0 else raw[i:]).strip()
    last = None
    for cand in (chunk, chunk.replace('\\"', '"').replace("\\\\", "\\")):
        for wrap in (True, False):
            try:
                return json.loads('"' + cand + '"') if wrap else json.loads(cand)
            except Exception as e:
                last = e
    raise RuntimeError("无法解析 Result 内容: %s | 前 200 字符: %s" % (last, chunk[:200]))


def eval_js(js, timeout=300):
    return json.loads(extract_result(_run(["browser_eval_content_js", "--script", js], timeout)))


def goto(url, timeout=180):
    return _run(["browser_go_to_url", "--url", url], timeout)


def _prompt_login():
    """登录态失效（ptlogin-ex verify fail / retcode 2001）时，自动打开登录页并提示用户。"""
    try:
        goto(LOGIN_URL)
    except Exception:
        pass
    print("!" * 56)
    print("⚠️ 登录态失效：接口返回 ptlogin-ex verify fail（未登录）")
    print("   根因：尚未在 https://qun.qq.com/#/login 完成登录。")
    print("已为你自动打开登录页：")
    print("   " + LOGIN_URL)
    print("请在弹出的 QQ 浏览器窗口中点击「登录」，登录成功后再重新执行本命令。")
    print("!" * 56)
    return 2


def _ensure_venv():
    """用 base 解释器运行时，自动切到「与 CLI 同目录」的 venv 解释器重跑自己。

    坑（2026-09-04）：base 解释器（.../python/versions/3.13.x/python.exe）通常**没装任何包**，
    而 python-docx / pillow 装在隔离 venv（.../python/envs/default/Scripts）里。
    用 base 跑 `docx` 会直接 ModuleNotFoundError —— 而包其实早就装过了。
    与其要求调用方记住该用哪个解释器（定时任务 prompt 里写死路径迟早会错），
    不如让脚本自己兜住。
    """
    if os.environ.get("QQHW_REEXEC"):
        return                                   # 已重跑过，别无限递归
    try:
        import docx, PIL                         # noqa: F401  够用就直接返回
        return
    except ImportError:
        pass

    import shutil
    exe = CLI if os.path.isabs(CLI) else shutil.which(CLI)
    if not exe:
        return
    cand = os.path.join(os.path.dirname(exe), "python.exe")
    if not os.path.isfile(cand):
        return
    if os.path.abspath(cand).lower() == os.path.abspath(sys.executable).lower():
        return
    try:                                          # 确认目标解释器真有依赖，避免无谓切换
        rc = subprocess.run([cand, "-c", "import docx,PIL"],
                            capture_output=True, timeout=60).returncode
    except Exception:
        return
    if rc != 0:
        return
    print("⚠️ 当前解释器（%s）缺少 python-docx/pillow，" % sys.executable)
    print("   自动切换到 venv 解释器重跑：%s" % cand)
    sys.stdout.flush()
    # 用子进程而不是 os.execve：Windows 上 exec 的进程替换语义不可靠（输出会丢、产物可能不落盘），
    # subprocess 透传 stdout 与退出码，行为可预期。QQHW_REEXEC 防递归。
    env = dict(os.environ, QQHW_REEXEC="1")
    try:
        r = subprocess.run([cand, os.path.abspath(__file__)] + sys.argv[1:], env=env)
        sys.exit(r.returncode)
    except Exception as e:                        # 起不来就退回原解释器，让真实错误暴露
        print("   （自动切换失败：%s，继续用当前解释器）" % e)


# ---------------- bkn 探测 ----------------

JS_BKN = (
    "var e=performance.getEntriesByType('resource').map(function(x){return x.name});"
    "var hit=null;for(var i=0;i<e.length;i++){var m=e[i].match(/[?&]bkn=(\\d+)/);if(m){hit=m[1];break;}}"
    "JSON.stringify({bkn:hit,cgiCount:e.filter(function(n){return n.indexOf('cgi-bin')>-1}).length})"
)

def cmd_bkn(args):
    cfg = load_cfg()
    gids = parse_gids(args.gid, cfg)
    if not gids:
        print("需要 --gid 群号（支持逗号分隔多群）"); return 1
    cfg["gids"] = gids
    cfg["gid"] = gids[0]
    goto(INDEX_URL.format(gid=gids[0]))
    r = eval_js(JS_BKN)
    if not r.get("bkn"):
        print("未探测到 bkn —— 多半是**未登录**。")
        print("请在弹出的 QQ 浏览器窗口中完成登录后，重新执行本命令。")
        save_cfg(cfg)
        return 2
    cfg["bkn"] = r["bkn"]
    save_cfg(cfg)
    print("bkn =", cfg["bkn"], "| 群:", ", ".join(gids), "| cgi 请求数:", r.get("cgiCount"))
    return 0


# ---------------- 列表 / 详情 ----------------

JS_LIST = (
    "var x=new XMLHttpRequest();x.open('POST','__API__',false);x.withCredentials=true;"
    "x.setRequestHeader('Content-Type','application/x-www-form-urlencoded');"
    "x.send('cmd=21&group_id=__GID__&num=1&page_size=__SIZE__&bkn=__BKN__');"
    "var j=JSON.parse(x.responseText);var hw=(j.data&&j.data.homework)||[];var out=[];"
    "for(var i=0;i<hw.length;i++){var h=hw[i];var cc=(h.content&&h.content.c)?h.content.c:[];var cs=[];"
    "for(var k=0;k<cc.length;k++){cs.push({t:cc[k].type,s:cc[k].text||''});}"
    "out.push({id:h.hw_id,title:h.hw_title,ts:h.ts_create,course:h.course_name,c:cs});}"
    "JSON.stringify({retcode:j.retcode,msg:j.msg||'',endFlag:j.data?j.data.end_flag:-1,total:hw.length,list:out})"
)

def cmd_list(args):
    cfg = load_cfg()
    gids = parse_gids(args.gid, cfg)
    bkn = args.bkn or cfg.get("bkn")
    if not (gids and bkn):
        print("缺少 gid/bkn，先执行: qq_hw.py bkn --gid <群号>"); return 1
    all_items, end_flag = [], -1
    for gid in gids:
        js = (JS_LIST.replace("__API__", API_LIST)
                     .replace("__GID__", str(gid))
                     .replace("__SIZE__", str(args.size))
                     .replace("__BKN__", str(bkn)))
        d = eval_js(js)
        if d.get("retcode") in (2001, 2004, 2007) or "verify fail" in str(d.get("msg", "")):
            return _prompt_login()
        if d.get("retcode") != 0:
            print("群 %s 接口返回 retcode=%s msg=%s（跳过）" % (gid, d.get("retcode"), d.get("msg", "")))
            continue
        for h in d["list"]:
            h["group_id"] = gid
        all_items.extend(d["list"])
        end_flag = d.get("endFlag", end_flag)
    # 去重（同群号重复拉取会产生重复条目）
    seen, dedup = set(), []
    for h in all_items:
        k = (h.get("group_id"), h["id"])
        if k in seen:
            continue
        seen.add(k); dedup.append(h)
    all_items = dedup
    with open(os.path.join(HERE, "hw_list.json"), "w", encoding="utf-8") as f:
        json.dump({"list": all_items, "endFlag": end_flag}, f, ensure_ascii=False, indent=1)
    print("作业 %d 条（%d 个群）-> hw_list.json (endFlag=%s)" % (len(all_items), len(gids), end_flag))
    return 0


JS_DETAIL = (
    "var ids=[__IDS__];var out=[];"
    "for(var k=0;k<ids.length;k++){"
    "var x=new XMLHttpRequest();"
    "x.open('POST','__API__',false);x.withCredentials=true;"
    "x.setRequestHeader('Content-Type','application/x-www-form-urlencoded');"
    "x.send('hw_id='+ids[k]+'&group_id=__GID__&bkn=__BKN__&puin=0&need_feedback=0');"
    "try{var j=JSON.parse(x.responseText);"
    "if(j.retcode!==0){out.push({id:ids[k],err:j.retcode});continue;}"
    "var d=j.data;var cc=(d.content&&d.content.c)?d.content.c:[];var cs=[];"
    "for(var m=0;m<cc.length;m++){cs.push({t:cc[m].type,s:cc[m].text||'',u:cc[m].url||'',"
    "w:cc[m].width,h:cc[m].height});}"
    "out.push({id:d.hw_id,title:d.hw_title,ts:d.ts_create,course:d.course_name,"
    "pub:d.pnick_name,pubuin:d.puin,fbname:(d.feedback&&d.feedback.nick_name)||'',c:cs});"
    "}catch(e){out.push({id:ids[k],err:String(e).slice(0,80)});}}"
    "JSON.stringify(out)"
)


def cmd_day(args):
    cfg = load_cfg()
    gids = parse_gids(args.gid, cfg)
    bkn = args.bkn or cfg.get("bkn")
    day = args.date
    lst_path = os.path.join(HERE, "hw_list.json")
    if not os.path.exists(lst_path) or args.refresh:
        if cmd_list(argparse.Namespace(gid=",".join(gids), bkn=bkn, size=100)) != 0:
            return 1
    data = json.load(open(lst_path, encoding="utf-8"))
    day_items = [h for h in data["list"]
                 if datetime.fromtimestamp(h["ts"]).strftime("%Y-%m-%d") == day]
    if not day_items:
        print("%s 没有作业" % day); return 1
    # 按群分组（兼容旧数据无 group_id：默认首个 gid）
    default_gid = gids[0]
    groups = {}
    for h in day_items:
        g = str(h.get("group_id") or default_gid)
        groups.setdefault(g, []).append(h)
    merged = []
    for gid, items in groups.items():
        ids = [h["id"] for h in items]
        try:
            js = (JS_DETAIL.replace("__IDS__", ",".join(str(i) for i in ids))
                           .replace("__API__", API_DETAIL)
                           .replace("__GID__", str(gid))
                           .replace("__BKN__", str(bkn)))
            details = eval_js(js)
        except Exception as e:
            print("群 %s 详情接口调用失败：%s" % (gid, e))
            return _prompt_login()
        bad = [d for d in details if not isinstance(d, dict) or "err" in d]
        if bad:
            print("群 %s 以下作业详情获取失败：%s"
                  % (gid, ", ".join(str(d.get("err", "?")) for d in bad if isinstance(d, dict))))
            return _prompt_login()
        for d in details:
            d["group_id"] = gid
            merged.append(d)
    # 去重（防御：同群号重复拉取会产生重复条目）
    seen, dedup = set(), []
    for d in merged:
        k = (str(d.get("group_id") or ""), d.get("id"))
        if k in seen:
            continue
        seen.add(k); dedup.append(d)
    merged = dedup
    merged.sort(key=lambda x: x.get("ts", 0))
    out = os.path.join(HERE, "hw_day_%s.json" % day)
    with open(out, "w", encoding="utf-8") as f:
        json.dump(merged, f, ensure_ascii=False, indent=1)
    n_img = sum(1 for r in merged for c in r.get("c", []) if c.get("t") == "img" and c.get("u"))
    print("%s: %d 条作业（%d 个群）, %d 张图片 -> %s" % (day, len(merged), len(groups), n_img, out))
    return 0


# ---------------- docx ----------------

def download(url, name):
    from PIL import Image
    Image.MAX_IMAGE_PIXELS = None
    os.makedirs(IMG_DIR, exist_ok=True)
    path = os.path.join(IMG_DIR, name)
    if os.path.exists(path) and os.path.getsize(path) > 1000:
        return path
    with urllib.request.urlopen(urllib.request.Request(url, headers=HEADERS), timeout=60) as r:
        data = r.read()
    im = Image.open(io.BytesIO(data))
    if im.mode not in ("RGB", "L"):
        im = im.convert("RGB")
    im.thumbnail((1400, 1400), Image.LANCZOS)
    im.save(path, "JPEG", quality=82, optimize=True)
    return path


def build_docx(day_items, date_str, out_path, scale, multi=False, text_only=False):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    HEAD_FONT, BODY_FONT = "微软雅黑", "宋体"

    def sf(run, name, size, bold=False):
        run.font.name = name; run.font.size = Pt(size); run.font.bold = bold
        rPr = run._element.get_or_add_rPr()
        rf = rPr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); rPr.append(rf)
        rf.set(qn("w:eastAsia"), name)

    def style_font(st, name, size, bold=False):
        st.font.name = name; st.font.size = Pt(size); st.font.bold = bold
        rPr = st.element.get_or_add_rPr()
        rf = rPr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); rPr.append(rf)
        rf.set(qn("w:eastAsia"), name)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(0.8)
    sec.left_margin = sec.right_margin = Cm(1.2)
    usable_w, usable_h = 21 - 2 * 1.2, 29.7 - 2 * 0.8

    # 关键：在**样式层**压掉自带间距，否则 Heading2/ListNumber 会撑到第二页
    for sname, fnt, sz, bd in (("Normal", BODY_FONT, 10.5, False),
                               ("Heading 2", HEAD_FONT, 13, True),
                               ("List Number", BODY_FONT, 10.5, False)):
        try:
            st = doc.styles[sname]
        except KeyError:
            continue
        style_font(st, fnt, sz, bd)
        pf = st.paragraph_format
        pf.space_before = pf.space_after = Pt(0)
        pf.line_spacing = 1.0

    clean = lambda t: re.sub(r"^\d{1,2}月\d{1,2}日", "", t).strip() or t

    def split_items(text):
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        if not lines:
            return []
        out = [re.sub(r"^\d{1,2}\s*[.、)．]\s*", "", l).strip() for l in lines]
        if len(out) == 1 and not re.match(r"^\d{1,2}\s*[.、)．]", lines[0]):
            return [("plain", out[0])]
        return [("num", x) for x in out]

    def group_rows(imgs):
        rows, i = [], 0
        while i < len(imgs):
            cur, nxt = imgs[i], imgs[i + 1] if i + 1 < len(imgs) else None
            rc = (cur.get("w") or 1) / (cur.get("h") or 1)
            rn = (nxt.get("w") or 1) / (nxt.get("h") or 1) if nxt else 99
            if rc < 0.95 and rn < 0.95:      # 两张竖图 → 横排
                rows.append([cur, nxt]); i += 2
            else:
                rows.append([cur]); i += 1
        return rows

    gids_present = {str(it.get("group_id") or "") for it in day_items}
    multi_group = len([g for g in gids_present if g]) > 1

    prepared, seq = [], 0
    for it in day_items:
        text = "".join(c.get("s", "") for c in it.get("c", []) if c.get("t") == "str").strip()
        imgs = [] if text_only else [c for c in it.get("c", []) if c.get("t") == "img" and c.get("u")]
        for c in imgs:
            seq += 1
            c["_file"] = f"{date_str}_{seq}.jpg"
        prepared.append({"title": clean(it.get("title", "")), "pub": it.get("pub") or "",
                         "group_id": str(it.get("group_id") or ""),
                         "items": split_items(text), "rows": group_rows(imgs)})

    fixed = 1.15
    for p in prepared:
        fixed += 0.72 + 0.46 * len(p["items"])
    n_rows = sum(len(p["rows"]) for p in prepared)
    if multi:
        per_row_h = 6.0          # 多页模式：图片用自然高度，允许分页
    else:
        per_row_h = max(2.0, ((usable_h - fixed) / n_rows if n_rows else 0) / scale)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    sf(p.add_run(f"{date_str} 作业"), HEAD_FONT, 14, True)

    for pr in prepared:
        h = doc.add_heading(level=2)
        h.paragraph_format.space_before = Pt(4)
        h.paragraph_format.space_after = Pt(2)
        for r in list(h.runs):
            r.text = ""
        sf(h.add_run(pr["title"]), HEAD_FONT, 13, True)
        if multi_group and pr["group_id"]:
            sf(h.add_run(f"　（群{pr['group_id']}）"), HEAD_FONT, 9, False)

        num_idx = 0
        for kind, txt in pr["items"]:
            if kind == "num":
                num_idx += 1
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.space_before = pf.space_after = Pt(0)
                pf.line_spacing = 1.0
                pf.left_indent = Cm(0.63)
                pf.first_line_indent = Cm(-0.63)
                sf(p.add_run(f"{num_idx}. "), BODY_FONT, 10.5)
                sf(p.add_run(txt), BODY_FONT, 10.5)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                sf(p.add_run(txt), BODY_FONT, 10.5)

        for row in pr["rows"]:
            n = len(row)
            hs = []
            for c in row:
                ratio = (c.get("w") or 1000) / (c.get("h") or 1000)
                hs.append(min(per_row_h, ((usable_w - 0.3 * (n - 1)) / n) / ratio))
            tgt_h = min(hs)
            tbl = doc.add_table(rows=1, cols=n)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = False
            mar = OxmlElement("w:tblCellMar")
            for tag in ("top", "start", "bottom", "end"):
                el = OxmlElement("w:" + tag)
                el.set(qn("w:w"), "0"); el.set(qn("w:type"), "dxa")
                mar.append(el)
            tbl._tbl.tblPr.append(mar)
            for j, c in enumerate(row):
                ratio = (c.get("w") or 1000) / (c.get("h") or 1000)
                w = tgt_h * ratio
                cell = tbl.cell(0, j)
                cell.width = Cm(w)
                cp = cell.paragraphs[0]
                cp.text = ""
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_before = cp.paragraph_format.space_after = Pt(0)
                try:
                    fp = download(c["u"], c["_file"])
                except Exception as e:
                    sf(cp.add_run(f"[图片下载失败: {e}]"), BODY_FONT, 8)
                    continue
                cp.add_run().add_picture(fp, width=Cm(w), height=Cm(tgt_h))

    doc.save(out_path)
    return out_path


def cmd_docx(args):
    _ensure_venv()          # base 解释器没装 docx/PIL 时自动切到 venv 重跑
    src = os.path.join(HERE, "hw_day_%s.json" % args.date)
    if not os.path.exists(src):
        print("缺少 %s，先执行: qq_hw.py day %s" % (src, args.date)); return 1
    data = json.load(open(src, encoding="utf-8"))
    # 科目筛选（默认全部；匹配 course 或 title 包含任一指定科目名）
    if args.courses:
        cs = [c.strip() for c in args.courses.split(",") if c.strip()]
        data = [it for it in data
                if any(k in (it.get("course") or "") or k in (it.get("title") or "") for k in cs)]
        if not data:
            print("筛选后无匹配科目作业（--courses %s）" % args.courses); return 1
    data.sort(key=lambda x: x.get("ts", 0))
    dt = datetime.strptime(args.date, "%Y-%m-%d")
    date_str = "%d年%d月%d日" % (dt.year, dt.month, dt.day)
    out = args.out
    if not out:
        base = os.path.join(HERE, "作业_%s" % args.date)
        out = base + ("_文字版.docx" if args.text_only else ".docx")
    build_docx(data, date_str, out, args.scale, multi=args.allow_multi, text_only=args.text_only)
    print("saved:", out, os.path.getsize(out), "bytes")
    return 0


def cmd_pdf(args):
    src = args.docx or os.path.join(HERE, "作业_%s.docx" % args.date)
    if not os.path.exists(src):
        print("缺少 %s，先执行: qq_hw.py docx %s" % (src, args.date)); return 1
    out = args.out or os.path.join(HERE, "作业_%s.pdf" % args.date)
    tmp = os.path.join(HERE, "_tmp_conv.pdf")
    # 中文路径经 -File 参数传递会被 PowerShell 按系统编码解码乱码，
    # 故把路径作为字面量写入 UTF-8 BOM 的临时 ps1。
    ps = os.path.join(HERE, "_conv_pdf_run.ps1")
    code = (
        "$tmp = '%s'\n" % tmp +
        "$src = '%s'\n" % src +
        "$dst = '%s'\n" % out +
        "try {\n"
        "  $w = New-Object -ComObject Word.Application\n"
        "  $w.Visible = $false\n"
        "  $d = $w.Documents.Open($src)\n"
        "  $d.SaveAs([ref]$tmp, [ref]17)\n"
        "  $d.Close()\n"
        "} catch {\n"
        "  Write-Error ('SAVE_FAIL: ' + $_.Exception.Message)\n"
        "  exit 3\n"
        "} finally { if ($w) { try { $w.Quit() } catch {} } }\n"
        "if (Test-Path $tmp) { Copy-Item $tmp $dst -Force; Remove-Item $tmp -Force -ErrorAction SilentlyContinue; exit 0 } else { Write-Error 'PDF_NOT_CREATED'; exit 4 }\n"
    )
    with open(ps, "w", encoding="utf-8-sig") as f:
        f.write(code)
    rc = subprocess.run(["powershell.exe", "-NoProfile", "-ExecutionPolicy", "Bypass",
                         "-File", ps], cwd=HERE).returncode
    if rc != 0 or not os.path.exists(out):
        print("PDF 转换失败（rc=%s）" % rc); return 1
    print("saved:", out, os.path.getsize(out), "bytes")
    return 0


def cmd_doctor(args):
    """无人值守前置自检：一次性回答「CLI 能用吗 / 扩展连着吗 / 登录了没 / 产物可复用吗」。

    定时任务每一轮开头跑一次，可避免把整轮耗在错误方向上
    （例如 CLI 被封锁却去让用户登录、clients=0 就死等、产物已生成却重做一遍）。
    """
    import shutil, glob

    print("== 环境自检 doctor ==")
    problems = []

    def ok(t): print("  [ OK ] " + t)
    def warn(t): print("  [WARN] " + t)
    def fail(t): print("  [FAIL] " + t)

    # 0. 先解析 CLI，并据此推导「同 venv 的解释器」
    #    常见坑：用 base 解释器（如 .../python/versions/3.13.x/python.exe）跑本脚本，
    #    而 python-docx/pillow 装在隔离 venv 里 → 报「依赖未安装」其实是解释器选错了
    resolved = CLI if (os.path.isabs(CLI) and os.path.exists(CLI)) else shutil.which(CLI)
    hint_py = ""
    if resolved:
        cand = os.path.join(os.path.dirname(resolved), "python.exe")
        if os.path.exists(cand):
            hint_py = cand

    # 1. Python 依赖
    for mod, pkg, required in (("docx", "python-docx", True),
                               ("PIL", "pillow", True),
                               ("pypdf", "pypdf", False)):
        try:
            __import__(mod)
            ok("依赖 %s" % pkg)
        except ImportError:
            msg = "依赖 %s 未安装 → pip install %s" % (pkg, pkg)
            (fail if required else warn)(msg)
            if required:
                if hint_py:
                    warn("      若明明装过 → 多半是**用错了解释器**，"
                         "改用与 CLI 同 venv 的：%s" % hint_py)
                problems.append("dep:" + pkg)

    # 2. CLI 可执行文件（隔离 venv 安装时常常不在 PATH）
    if resolved:
        ok("CLI = %s" % resolved)
    else:
        fail("找不到 CLI（QQB_CLI=%s）；隔离 venv 安装时需显式指定绝对路径" % CLI)
        problems.append("cli")

    # 3. daemon 状态 + 版本封锁
    if resolved:
        try:
            st = _run(["status"], timeout=60, check_blocked=False)
            if BLOCKED_HINT in st:
                print(_blocked_help())
                problems.append("blocked")
            elif re.search(r"running", st, re.I):
                ok("daemon 运行中")
            else:
                warn("daemon 未运行 → qqbrowser-skill serve --daemon")
                problems.append("daemon")
            m = re.search(r"Connected clients\D*(\d+)", st)
            if m:
                print("         Connected clients = %s"
                      "（1.5.6+ 恒为 0 属正常，以第 4 步实测为准，不要死等）" % m.group(1))
        except CliBlocked as e:
            print(e.help_text)
            problems.append("cli")
        except Exception as e:
            fail("status 调用失败：%s" % str(e)[:120])
            problems.append("cli")

    # 4. 真实连通性 + 登录态（打开登录页，一次调用两用）
    if not ({"cli", "blocked", "daemon"} & set(problems)):
        try:
            goto(LOGIN_URL, timeout=60)
            info = eval_js("JSON.stringify({t:document.body.innerText.slice(0,80),u:location.href})",
                           timeout=60)
            ok("扩展连通性实测通过（业务命令可用）")
            body = (info or {}).get("t", "") or ""
            if "登录" in body and "退出" not in body:
                warn("登录态：未登录（页面：%s）" % body.replace("\n", " ")[:30])
                warn("  → 请在浏览器点击登录；登录成功后【必须先重探 bkn】再拉数据，"
                     "否则仍会 verify fail")
                problems.append("login")
            else:
                ok("登录态：已登录（%s）" % (info or {}).get("u", ""))
        except CliBlocked as e:
            print(e.help_text)
            problems.append("blocked")
        except Exception as e:
            fail("连通性实测失败：%s" % str(e)[:120])
            problems.append("connect")

    # 5. bkn 缓存
    cfg = load_cfg()
    gids = parse_gids(getattr(args, "gid", ""), cfg)
    if cfg.get("bkn"):
        ok("已缓存 bkn = %s（群：%s）" % (cfg["bkn"], ",".join(gids) or "-"))
    else:
        warn("无 bkn 缓存 → 先跑：qq_hw.py bkn --gid <群号>")
        problems.append("bkn")

    # 6. 现存产物（帮助定时任务判断能否直接复用，避免重复生成）
    for pat in ("hw_day_*.json", "作业_*.docx", "作业_*.pdf", "hw_auto_sent_*.flag"):
        fs = sorted(glob.glob(os.path.join(HERE, pat)))
        if fs:
            print("  [info] %s：%d 个（最近：%s）"
                  % (pat, len(fs), ", ".join(os.path.basename(x) for x in fs[-3:])))

    if problems:
        print("== 自检结束：存在问题 → %s ==" % ", ".join(problems))
        return 1
    print("== 自检结束：全部通过 ==")
    return 0


def cmd_pages(args):
    ps = os.path.join(os.path.dirname(os.path.abspath(__file__)), "count_pages.ps1")
    tmp = os.path.join(HERE, "page_count.txt")
    subprocess.run(["powershell.exe", "-NoProfile", "-ExecutionPolicy", "Bypass",
                    "-File", ps, "-Docx", os.path.abspath(args.file)], cwd=HERE)
    if os.path.exists(tmp):
        print(open(tmp, encoding="utf-8").read().strip())
    return 0


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="QQ 群作业 → Word/PDF 文档")
    sub = ap.add_subparsers(dest="cmd", required=True)

    p = sub.add_parser("bkn"); p.add_argument("--gid", default="",
        help="群号，支持逗号分隔多群（如 123456789,987654321）"); p.set_defaults(fn=cmd_bkn)
    p = sub.add_parser("list"); p.add_argument("--gid", default="",
        help="群号，支持逗号分隔多群"); p.add_argument("--bkn", default="")
    p.add_argument("--size", type=int, default=100); p.set_defaults(fn=cmd_list)
    p = sub.add_parser("day"); p.add_argument("date"); p.add_argument("--gid", default="",
        help="群号，支持逗号分隔多群")
    p.add_argument("--bkn", default=""); p.add_argument("--refresh", action="store_true")
    p.set_defaults(fn=cmd_day)
    p = sub.add_parser("docx"); p.add_argument("date"); p.add_argument("--out", default="")
    p.add_argument("--scale", type=float, default=1.15)
    p.add_argument("--courses", default="", help="科目筛选，逗号分隔，如 语文,数学；默认全部")
    p.add_argument("--text-only", action="store_true", help="仅文字，不插入图片")
    p.add_argument("--allow-multi", action="store_true", help="允许分页（内容过多时，需经用户同意）")
    p.set_defaults(fn=cmd_docx)
    p = sub.add_parser("pdf"); p.add_argument("date"); p.add_argument("--docx", default="")
    p.add_argument("--out", default=""); p.set_defaults(fn=cmd_pdf)
    p = sub.add_parser("doctor"); p.add_argument("--gid", default="",
        help="群号（可选；给出时会顺带判定登录态）"); p.set_defaults(fn=cmd_doctor)
    p = sub.add_parser("pages"); p.add_argument("file"); p.set_defaults(fn=cmd_pages)

    a = ap.parse_args()
    try:
        sys.exit(a.fn(a) or 0)
    except CliBlocked as e:
        # CLI 被服务端下线 / 找不到可执行文件：给出确定性修复指引，避免误判成未登录
        print(e.help_text)
        sys.exit(3)
