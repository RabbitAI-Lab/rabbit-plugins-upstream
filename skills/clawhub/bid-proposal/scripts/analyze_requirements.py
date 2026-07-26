#!/usr/bin/env python3
"""
Analyze bidding requirement documents (text or .docx).

Extracts project type, scoring items, technical domains, and key requirements
from tender documents. Supports plain text input and .docx files.

Usage:
    python3 analyze_requirements.py <input_text_or_docx> [--format json]
"""

import json
import os
import re
import sys


# ---------------------------------------------------------------------------
# Load scene mapping
# ---------------------------------------------------------------------------

def _load_scene_mapping():
    """Load scene-mapping.json relative to this script's directory."""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    mapping_path = os.path.join(script_dir, '..', 'references', 'scene-mapping.json')
    with open(mapping_path, 'r', encoding='utf-8') as f:
        return json.load(f)


# ---------------------------------------------------------------------------
# Extract scoring items from text
# ---------------------------------------------------------------------------

def extract_scoring_items(text: str) -> list:
    """
    Extract scoring items from tender document text.

    Recognizes table-like patterns:
      | 评分项 | 分值 |
      | 总体设计方案 | 15分 |

    Also handles lines like: 评分项名称 X分
    """
    items = []
    # Pattern 1: Markdown table rows with scoring items
    # Matches: | 项目理解 | 5分 | or | 项目理解 | 5 | 要求说明 |
    # NOTE: Uses non-greedy match with explicit cell boundary at |
    # \s* between cells is intentional but limited by {2,30} on name
    table_pattern = re.compile(
        r'\|\s*([^|]{2,30}?)\s*\|\s*(\d+)\s*分?\s*\|',
        re.MULTILINE
    )
    for match in table_pattern.finditer(text):
        name = match.group(1).strip()
        score = int(match.group(2))
        # Filter out table headers
        if name in ('评分项', '评分标准', '评审项', '评审因素'):
            continue
        items.append({'name': name, 'max_score': score})

    # Pattern 2: Lines like "XX方案：XX分" or "XX方案（XX分）"
    loose_pattern = re.compile(
        r'(.{4,30}?)[：:]\s*(\d+)\s*分',
        re.MULTILINE
    )
    for match in loose_pattern.finditer(text):
        name = match.group(1).strip()
        # Skip if it looks like a section header rather than scoring item
        if len(name) > 20:
            continue
        if name in ('评分项', '评分标准', '评审因素'):
            continue
        score = int(match.group(2))
        # Avoid duplicates
        if not any(item['name'] == name for item in items):
            items.append({'name': name, 'max_score': score})

    return items


# ---------------------------------------------------------------------------
# Tech domain detection
# ---------------------------------------------------------------------------

DOMAIN_KEYWORDS = {
    '数据安全': ['分类分级', '脱敏', 'DLP', '数据防泄漏', '数据加密', '数据库审计', '数据水印', '数据销毁'],
    '网络安全': ['防火墙', 'WAF', 'IDS', 'IPS', '入侵检测', '入侵防御', 'VPN', '流量清洗', '抗DDoS', '蜜罐'],
    '终端安全': ['EDR', '终端安全', '防病毒', '桌面管理', '准入控制', '终端管控'],
    '应用安全': ['API安全', 'WAF', '漏洞扫描', '代码审计', 'RASP', 'IAST', 'SAST', 'DAST'],
    '安全运营': ['态势感知', 'SOC', 'SIEM', 'SOAR', '安全运营', 'MSS', '安全监控', '日志分析'],
    '安全服务': ['渗透测试', '重保', '应急响应', '等保测评', '安全巡检', '攻防演练', '红蓝对抗', '风险评估'],
    '身份安全': ['IAM', '零信任', '身份认证', '权限管理', '堡垒机', 'MFA', 'SSO', '统一认证'],
    '云安全': ['CWPP', '容器安全', '云安全', 'SASE', 'CASB', 'CSPM'],
    '信创': ['信创', '国产化', '自主可控', '国产适配'],
}


def _detect_tech_domains(text: str) -> list:
    """Detect technical domains from text keywords."""
    detected = []
    for domain, keywords in DOMAIN_KEYWORDS.items():
        for kw in keywords:
            if kw.lower() in text.lower():
                detected.append(domain)
                break
    return list(set(detected))


def _extract_key_requirements(text: str) -> list:
    """Extract key requirements as keywords from text."""
    requirements = []
    all_keywords = []
    for domain_kws in DOMAIN_KEYWORDS.values():
        all_keywords.extend(domain_kws)

    # Find specific product/tech mentions
    text_lower = text.lower()
    for kw in set(all_keywords):
        if kw.lower() in text_lower:
            # Only include specific tech terms, not broad domains
            requirements.append(kw)

    # Also look for numbered items that look like requirements
    req_pattern = re.compile(r'^[\d]+[\.\、)）]\s*(.{5,50})$', re.MULTILINE)
    for match in req_pattern.finditer(text):
        line = match.group(1).strip()
        # Filter out headings/non-requirement lines
        if any(h in line for h in ['一、', '二、', '三、', '四、', '五、', '（', '）']):
            continue
        if len(line) < 50:
            requirements.append(line)

    # Deduplicate and limit
    seen = set()
    unique = []
    for r in requirements:
        if r not in seen:
            seen.add(r)
            unique.append(r)
    return unique[:20]  # Limit to top 20


def _extract_project_name(text: str) -> str:
    """Attempt to extract project name from text."""
    # Pattern: XXX项目 or XXX项目 招标文件
    patterns = [
        r'^(.{3,30}项目)',
        r'(?:项目名称|项目][：:]\s*)(.{3,30})',
        r'^([^\n]{3,40})(?:招标|技术需求|技术规格书)',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.MULTILINE)
        if match:
            name = match.group(1).strip()
            # Clean up
            name = re.sub(r'^\d+[\.\、]?\s*', '', name)
            if len(name) >= 3:
                return name
    return "未命名项目"


# ---------------------------------------------------------------------------
# Scene matching
# ---------------------------------------------------------------------------

def match_scene(text: str) -> dict:
    """
    Match input text against known scenes in scene-mapping.json.

    Returns dict with 'scene' name, 'score', and 'matched_keywords'.
    Falls back to default scene ('综合类') if no clear match.
    """
    mapping = _load_scene_mapping()
    scenes = mapping.get('scenes', {})
    default_scene = mapping.get('default_scene', '综合类')
    threshold = mapping.get('matching_rules', {}).get('threshold', 2)

    text_lower = text.lower()
    scores = {}
    matched_kw = {}

    for scene_name, scene_config in scenes.items():
        keywords = scene_config.get('keywords', [])
        matches = [kw for kw in keywords if kw.lower() in text_lower]
        scores[scene_name] = len(matches)
        matched_kw[scene_name] = matches

    # Find best match
    if scores:
        best_scene = max(scores, key=scores.get)
        best_score = scores[best_scene]
        if best_score >= threshold:
            return {
                'scene': best_scene,
                'score': best_score,
                'matched_keywords': matched_kw[best_scene],
                'all_scores': scores
            }

    # Fallback
    return {
        'scene': default_scene,
        'score': 0,
        'matched_keywords': [],
        'all_scores': scores
    }


# ---------------------------------------------------------------------------
# Main analysis functions
# ---------------------------------------------------------------------------

def analyze_text(text: str) -> dict:
    """
    Analyze tender requirement text and return structured JSON.

    Args:
        text: Plain text content of the tender document.

    Returns:
        dict with keys: project_type, project_name, scoring_items,
                        tech_domains, key_requirements
    """
    if not text or not text.strip():
        return {
            'project_type': '综合类',
            'project_name': '',
            'scoring_items': [],
            'tech_domains': [],
            'key_requirements': [],
            'special_notes': ''
        }

    scene_result = match_scene(text)
    scoring_items = extract_scoring_items(text)
    tech_domains = _detect_tech_domains(text)
    key_requirements = _extract_key_requirements(text)
    project_name = _extract_project_name(text)

    # Detect special notes
    special_notes = ''
    if '信创' in text or '国产化' in text:
        special_notes = '需支持信创环境'

    return {
        'project_type': scene_result['scene'],
        'project_name': project_name,
        'scoring_items': scoring_items,
        'tech_domains': tech_domains,
        'key_requirements': key_requirements,
        'special_notes': special_notes,
        'scene_match_detail': {
            'score': scene_result['score'],
            'matched_keywords': scene_result['matched_keywords']
        }
    }


def analyze_docx(docx_path: str) -> dict:
    """
    Read a .docx tender document and analyze it.

    Args:
        docx_path: Path to the .docx file.

    Returns:
        Same dict structure as analyze_text().
        On python-docx not installed, returns dict with 'error' key
        and 'fallback_suggestion' for graceful degradation.
    """
    try:
        from docx import Document
    except ImportError:
        return {
            'project_type': '综合类',
            'project_name': '',
            'scoring_items': [],
            'tech_domains': [],
            'key_requirements': [],
            'special_notes': '',
            'error': 'python-docx not installed. Please paste the tender text directly.',
            'fallback_suggestion': '请复制招标文件中的技术需求文本，粘贴到对话框。'
        }

    try:
        doc = Document(docx_path)
    except Exception as e:
        return {
            'project_type': '综合类',
            'project_name': '',
            'scoring_items': [],
            'tech_domains': [],
            'key_requirements': [],
            'special_notes': '',
            'error': f'Failed to read docx: {e}',
            'fallback_suggestion': '文件读取失败，请确认文件格式正确或直接粘贴文本。'
        }

    paragraphs = [p.text for p in doc.paragraphs]

    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
            paragraphs.append(row_text)

    full_text = '\n'.join(paragraphs)
    return analyze_text(full_text)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print("Usage: python3 analyze_requirements.py <input_text_or_docx> [--format json|text]")
        sys.exit(1)

    input_arg = sys.argv[1]
    output_format = 'json'

    if '--format' in sys.argv:
        idx = sys.argv.index('--format')
        if idx + 1 < len(sys.argv):
            output_format = sys.argv[idx + 1]

    # Determine if input is a file path or inline text
    if os.path.isfile(input_arg):
        if input_arg.endswith('.docx'):
            result = analyze_docx(input_arg)
        else:
            with open(input_arg, 'r', encoding='utf-8') as f:
                result = analyze_text(f.read())
    else:
        result = analyze_text(input_arg)

    if output_format == 'json':
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        # Text format
        print(f"项目类型: {result['project_type']}")
        print(f"项目名称: {result['project_name']}")
        print(f"技术域: {', '.join(result['tech_domains'])}")
        print(f"关键词: {', '.join(result['key_requirements'][:10])}")
        print(f"评分项: {len(result['scoring_items'])} 项")
        for item in result['scoring_items']:
            print(f"  - {item['name']}: {item['max_score']}分")
        if result['special_notes']:
            print(f"特殊说明: {result['special_notes']}")


if __name__ == '__main__':
    main()
