#!/usr/bin/env python3
"""
Generate a bid proposal Word document (.docx) from a template.

Reuses styles from brief-proposal specifications:
- Normal: 10.5pt (五号)
- Heading 1: 16pt (三号), Bold
- Heading 2: 15pt (小三号), Bold
- Heading 3: 14pt (四号), Bold
- Page breaks before chapter titles (on preceding paragraph)
- SDT (structured document tag) table of contents preserved

Usage:
    python3 generate_docx.py <template.docx> <chapters.json> <output.docx>
"""

import json
import os
import re
import sys

HAS_DOCX = False
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    pass


# ---------------------------------------------------------------------------
# Sample chapters for testing / defaults
# ---------------------------------------------------------------------------

SAMPLE_CHAPTERS_SAFETY_SERVICE = [
    {
        "title": "项目背景与目标",
        "level": 1,
        "content": """## （一）项目背景

随着网络安全形势日益严峻，QG集团面临着持续增长的安全威胁。集团现有安全防护体系在安全监测、应急响应、渗透测试等方面存在能力缺口。

目前，QG集团已建设基础网络安全防护设施，但在安全运营和主动防御方面仍有不足。

## （二）建设目标

本项目旨在通过引入专业安全服务，实现以下目标：
1. 建立常态化安全运营机制，7×24小时安全监控
2. 每季度对核心系统进行渗透测试，及时发现安全漏洞
3. 重大活动期间提供重保服务，确保零事故
4. 安全事件发生后2小时内响应，降低安全风险"""
    },
    {
        "title": "需求分析",
        "level": 1,
        "content": """## （一）业务需求分析

QG集团核心业务涵盖生产、销售、供应链、财务等多个领域，核心业务系统超过20套。随着数字化转型推进，业务连续性要求日益提升。

## （二）安全需求分析

根据招标文件技术要求，梳理以下核心安全需求：

1. **安全运营需求**
   - 现状：缺乏7×24小时专业安全监控能力
   - 需求：建立安全运营中心（SOC），实现全天候安全事件监测和分析
   - 依据：招标文件第四章"安全运营服务要求"

2. **渗透测试需求**
   - 现状：每年仅开展一次外部渗透测试
   - 需求：每季度对核心业务系统进行深度渗透测试
   - 依据：招标文件第四章"渗透测试服务要求"

3. **应急响应需求**
   - 现状：安全事件发生后缺乏专业处置能力
   - 需求：建立应急响应机制，2小时内到达现场
   - 依据：招标文件第四章"应急响应服务要求"

## （三）合规需求分析

本项目需满足以下合规要求：
- 《网络安全法》第二十一条：采取监测、记录网络运行状态、网络安全事件的技术措施
- 《关键信息基础设施安全保护条例》：定期开展安全检测评估"""
    },
    {
        "title": "服务方案",
        "level": 1,
        "content": """## （一）安全运营服务

### 1. 服务概述

提供7×24小时安全运营服务，通过安全运营中心对QG集团信息系统进行全天候监控、分析和处置。

### 2. 服务内容

| 服务项 | 内容说明 | 交付频次 |
|--------|---------|---------|
| 安全监控 | 实时监控安全事件和告警 | 7×24小时 |
| 日志分析 | 对安全设备日志进行关联分析 | 每日 |
| 威胁情报 | 提供最新威胁情报和预警 | 实时 |
| 安全周报 | 安全态势分析和改进建议 | 每周 |
| 安全月报 | 安全运营总结和趋势分析 | 每月 |

### 3. 技术平台

采用自研安全运营平台，具备以下核心能力：
- 日志采集：支持Syslog、SNMP、API等多种接入方式
- 关联分析：内置500+条安全分析规则
- SOAR自动化：支持自动化告警处置和工单派发
- 可视化大屏：实时展示安全态势

### 4. 服务团队

配备以下专业人员：
| 角色 | 人数 | 资质要求 |
|------|------|---------|
| 安全运营经理 | 1 | CISSP/CISP，5年以上经验 |
| 安全分析师 | 3 | CISP/CISAW，3年以上经验 |
| 安全运维工程师 | 2 | 熟悉主流安全产品运维 |

## （二）渗透测试服务

### 1. 服务概述

每季度对QG集团核心业务系统进行深度渗透测试，及时发现并修复安全漏洞。

### 2. 测试范围

- 每季度覆盖不少于5套核心业务系统
- 包含Web应用、移动APP、API接口
- 测试内容：信息收集、漏洞扫描、手工渗透、权限提升、横向移动

### 3. 测试流程

1. 信息收集阶段（1天）
2. 漏洞扫描阶段（2天）
3. 手工渗透阶段（3天）
4. 漏洞验证阶段（2天）
5. 报告编写阶段（2天）

### 4. 交付物

- 《渗透测试报告》（含漏洞详情、风险等级、修复建议）
- 《漏洞修复验证报告》

## （三）应急响应服务

### 1. 服务概述

安全事件发生后，快速响应、有效处置，最大限度降低损失。

### 2. 响应时效

| 事件级别 | 远程响应 | 现场到达 | 问题解决 |
|---------|---------|---------|---------|
| P1-紧急 | 15分钟 | 2小时 | 4小时 |
| P2-严重 | 30分钟 | 4小时 | 8小时 |
| P3-一般 | 1小时 | 8小时 | 24小时 |

### 3. 响应流程

1. 事件上报 → 2. 初步研判 → 3. 应急处置 → 4. 根因分析 → 5. 整改加固 → 6. 复盘总结"""
    },
    {
        "title": "实施方案",
        "level": 1,
        "content": """## （一）项目组织

本项目组建如下项目团队：

| 角色 | 人数 | 职责 |
|------|------|------|
| 项目经理 | 1 | 项目整体管理、客户沟通、进度把控 |
| 安全运营负责人 | 1 | 安全运营服务技术把控 |
| 渗透测试负责人 | 1 | 渗透测试服务技术把控 |
| 安全运营工程师 | 3 | 日常安全监控和分析 |
| 渗透测试工程师 | 2 | 渗透测试执行 |

## （二）实施计划

项目分为以下阶段：

**第一阶段：服务启动（预计2周）**
- 完成项目启动会和服务方案确认
- 完成账户开通和平台部署
- 完成服务团队组建和培训

**第二阶段：试运行（预计1个月）**
- 开展安全运营监控试运行
- 完成首次渗透测试
- 优化告警规则和服务流程

**第三阶段：正式服务（10个月）**
- 全面开展各项安全服务
- 每月提交服务报告
- 每季度完成渗透测试

**第四阶段：服务总结（1周）**
- 完成年度服务总结报告
- 提供下一年度服务建议

## （三）风险应对

| 风险项 | 影响 | 应对措施 |
|--------|------|---------|
| 人员变动 | 高 | 建立备岗机制，核心岗位不少于2人 |
| 工具平台故障 | 中 | 部署冗余平台，实现自动切换 |
| 客户配合不足 | 中 | 明确接口人，建立定期沟通机制 |"""
    },
    {
        "title": "服务保障与SLA",
        "level": 1,
        "content": """## （一）服务内容

本项目服务期内提供以下服务保障：

| 服务项 | 服务内容 | 服务频次 |
|--------|---------|---------|
| 热线支持 | 7×24小时服务热线和技术支持 | 全天候 |
| 远程支持 | VPN远程接入提供技术支持 | 按需 |
| 现场服务 | 工程师现场处理故障和问题 | 按需 |
| 定期巡检 | 现场巡检设备和系统状态 | 每月 |

## （二）服务水平承诺

| 服务指标 | 承诺值 |
|---------|--------|
| 安全事件监测覆盖率 | ≥99.9% |
| 告警准确率 | ≥95% |
| 渗透测试按时完成率 | 100% |
| 应急响应到位率 | 100% |
| 客户满意度 | ≥90% |

## （三）服务保障措施

- 建立服务质量监督机制，每月进行服务质量评估
- 设立客户满意度回访制度，每季度进行满意度调查
- 建立投诉处理机制，投诉24小时内响应
- 每季度提交服务改进报告"""
    }
]


# ---------------------------------------------------------------------------
# Default template creation
# ---------------------------------------------------------------------------

def create_default_template(output_path: str) -> str:
    """
    Create a default .docx template with proper styles for bid proposals.

    The template follows brief-proposal style specifications:
    - Normal: 10.5pt
    - Heading 1-3: 16pt/15pt/14pt, Bold
    - Cover page with title and date
    - Placeholder for SDT table of contents
    - Page breaks between major sections

    Args:
        output_path: Where to save the generated .docx template.

    Returns:
        Path to the created template file.
    """
    if not HAS_DOCX:
        raise ImportError(
            "python-docx is required. Install with: pip install python-docx"
        )

    doc = Document()

    # ---- Configure styles ----
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)
    style.font.name = '宋体'
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    for level, size in [(1, 16), (2, 15), (3, 14)]:
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.size = Pt(size)
        heading_style.font.bold = True
        heading_style.font.name = '黑体'
        heading_style.font.color.rgb = None  # Default black
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)

    # ---- Cover page ----
    # Empty paragraphs for spacing
    for _ in range(6):
        doc.add_paragraph('', style='Normal')

    title_para = doc.add_paragraph('XX项目技术方案', style='Normal')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.size = Pt(22)
        run.font.bold = True

    doc.add_paragraph('', style='Normal')

    date_para = doc.add_paragraph('2026年6月', style='Normal')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date_para.runs:
        run.font.size = Pt(16)
        run.font.bold = True

    # ---- Page break after cover ----
    doc.add_page_break()

    # ---- SDT Table of Contents placeholder ----
    # Note: python-docx cannot natively create SDT (structured document tags)
    # for auto-TOC. We add a placeholder paragraph that users can replace
    # with Word's built-in TOC (Insert > Table of Contents).
    toc_para = doc.add_paragraph('', style='Normal')
    toc_run = toc_para.add_run('【请在Word中插入自动目录：插入 → 目录 → 自动目录】')
    toc_run.font.size = Pt(10.5)
    toc_run.font.color.rgb = None
    toc_run.font.italic = True

    doc.add_page_break()

    # ---- Content area ----
    content_placeholder = doc.add_paragraph(
        '（正文内容将由脚本自动填充。以下为生成时的内容区域。）',
        style='Normal'
    )
    # Mark this paragraph as the insertion point
    # We'll find it and add content after it

    # Add an empty paragraph after for clean insertion
    doc.add_paragraph('', style='Normal')

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Chapter content restructuring
# ---------------------------------------------------------------------------

def _markdown_to_paragraphs(doc, markdown_text: str) -> list:
    """
    Parse markdown text and return structured paragraph data.

    Converts markdown headings, lists, and paragraphs into structured data
    that can be inserted into a docx document.

    Args:
        doc: python-docx Document (used for style references).
        markdown_text: Markdown formatted chapter content.

    Returns:
        List of dicts with 'type' and 'text' keys, and optional 'level' for headings.
    """
    items = []
    lines = markdown_text.strip().split('\n')

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Heading detection: ## Title, ### Title
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            items.append({
                'type': 'heading',
                'text': heading_match.group(2).strip(),
                'level': level
            })
            continue

        # Table row detection: | col1 | col2 |
        if stripped.startswith('|') and stripped.endswith('|'):
            # Skip separator lines like |---|---|
            if re.match(r'^\|[\s\-:]+\|', stripped):
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            items.append({
                'type': 'table_row',
                'cells': cells
            })
            continue

        # List item detection: - item or 1. item
        list_match = re.match(r'^[\-\*] |^\d+[\.\、] ', stripped)
        if list_match:
            items.append({
                'type': 'list_item',
                'text': stripped
            })
            continue

        # Regular paragraph
        # Strip markdown bold (**text**) for display
        clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
        items.append({
            'type': 'paragraph',
            'text': clean_text
        })

    return items


# ---------------------------------------------------------------------------
# Main generation function
# ---------------------------------------------------------------------------

def generate_bid_docx(template_path: str, chapters: list, output_path: str) -> str:
    """
    Generate a bid proposal .docx from a template and chapters content.

    Args:
        template_path: Path to the template .docx file (will not be modified).
        chapters: List of chapter dicts, each with:
                  - title: Chapter title string
                  - level: Heading level (1, 2, or 3)
                  - content: Markdown-formatted chapter content
        output_path: Where to save the generated .docx file.

    Returns:
        Path to the generated output file.
    """
    if not HAS_DOCX:
        raise ImportError(
            "python-docx is required. Install with: pip install python-docx"
        )

    if not os.path.isfile(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")

    # Load template
    doc = Document(template_path)

    # Apply style adjustments (in case template has different styles)
    _ensure_styles(doc)

    # ---- Find content insertion point ----
    # Strategy: find the last paragraph before content area, or append to end
    # Look for placeholder text to determine insertion point
    body = doc.element.body
    insert_after_element = None

    for p in doc.paragraphs:
        if '正文内容将由脚本自动填充' in p.text:
            insert_after_element = p._element
            break

    # If no placeholder found, find page breaks and insert after the TOC area
    if insert_after_element is None:
        # Find the first paragraph after all page breaks that isn't a heading
        for p in doc.paragraphs:
            # Skip paragraphs that are likely cover/TOC
            if any(kw in p.text for kw in ['目录', '更新域', 'INSERT', 'TOC']):
                continue
            # Find a normal paragraph after the initial section
            if p.style.name == 'Normal' and len(p.text.strip()) > 0:
                insert_after_element = p._element
                break

    if insert_after_element is None:
        # Fallback: append to end of body
        insert_after_element = body[-1] if len(body) > 0 else None

    # ---- Insert chapter content ----
    prev_element = insert_after_element

    for ch in chapters:
        ch_title = ch.get('title', '')
        ch_level = ch.get('level', 1)
        ch_content = ch.get('content', '')

        # ---- Add page break before chapter (on preceding paragraph) ----
        if prev_element is not None and ch_level <= 2:
            _add_page_break_to_element(prev_element, doc)

        # ---- Insert chapter title as heading ----
        heading_para = _insert_element_after(prev_element, doc, 'paragraph')
        heading_para.style = doc.styles[f'Heading {ch_level}']
        heading_para.add_run(ch_title)
        prev_element = heading_para._element

        # ---- Insert chapter content ----
        content_items = _markdown_to_paragraphs(doc, ch_content)

        for item in content_items:
            item_type = item.get('type', 'paragraph')

            if item_type == 'heading':
                # Sub-headings within content
                sub_level = min(item.get('level', 2) + 1, 3)  # Offset by 1 level
                sub_para = _insert_element_after(prev_element, doc, 'paragraph')
                sub_para.style = doc.styles[f'Heading {sub_level}']
                sub_para.add_run(item['text'])
                prev_element = sub_para._element

            elif item_type == 'table_row':
                # Collect consecutive table rows
                table_rows = [item['cells']]
                # Peek next items for table rows
                # (We insert tables in batches here; single-row tables are fine too)
                table = _insert_table_after(prev_element, doc, table_rows)
                prev_element = table._element

            elif item_type == 'paragraph' or item_type == 'list_item':
                para = _insert_element_after(prev_element, doc, 'paragraph')
                para.style = doc.styles['Normal']
                para.add_run(item['text'])
                prev_element = para._element

    # Save output (never overwrite template)
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def _ensure_styles(doc):
    """Ensure document has required styles configured."""
    try:
        for level, size in [(1, 16), (2, 15), (3, 14)]:
            style_name = f'Heading {level}'
            if style_name in [s.name for s in doc.styles]:
                style = doc.styles[style_name]
                if style.font.size is None:
                    style.font.size = Pt(size)
                style.font.bold = True
    except Exception:
        pass  # Non-critical style fixes


def _insert_element_after(ref_element, doc, element_type='paragraph'):
    """Insert a new element (paragraph) after ref_element."""
    if element_type == 'paragraph':
        new_para = doc.add_paragraph('')
        ref_element.addnext(new_para._element)
        return new_para
    return None


def _add_page_break_to_element(element, doc):
    """Add a page break to an element's last run."""
    run = None
    # Find the element in doc paragraphs
    for p in doc.paragraphs:
        if p._element == element:
            if p.runs:
                run = p.runs[-1]
            else:
                run = p.add_run(' ')
            break

    if run is None:
        # If element is not a paragraph in paragraphs list, try direct
        try:
            # Access runs via lxml
            runs = element.findall(qn('w:r'))
            if runs:
                run_elem = runs[-1]
            else:
                run_elem = OxmlElement('w:r')
                element.append(run_elem)
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            run_elem.append(br)
            return
        except Exception:
            pass

    if run:
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run._element.append(br)


def _insert_table_after(ref_element, doc, rows: list) -> object:
    """Insert a simple table after ref_element."""
    if not rows:
        return None

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols, style='Table Grid')
    ref_element.addnext(table._element)

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                # Set cell font size
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

    return table


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    if sys.argv[1] == '--create-template':
        create_default_template(sys.argv[2])
        print(f"Template created: {sys.argv[2]}")
        return

    if len(sys.argv) < 4:
        print("Usage: python3 generate_docx.py <template.docx> <chapters.json> <output.docx>")
        print()
        print("  template.docx: Path to template .docx file")
        print("  chapters.json: Path to JSON file with chapters list")
        print("                 Format: [{\"title\": \"...\", \"level\": 1, \"content\": \"...\"}, ...]")
        print("  output.docx: Path where generated document will be saved")
        print()
        print("  If template.docx doesn't exist, a default template will be created first.")
        print("  Use: python3 generate_docx.py --create-template <output.docx>")
        sys.exit(1)
        create_default_template(sys.argv[2])
        print(f"Template created: {sys.argv[2]}")
        return

    template = sys.argv[1]
    chapters_file = sys.argv[2]
    output = sys.argv[3]

    with open(chapters_file, 'r', encoding='utf-8') as f:
        chapters = json.load(f)

    result = generate_bid_docx(template, chapters, output)
    print(f"Generated: {result}")


if __name__ == '__main__':
    main()
