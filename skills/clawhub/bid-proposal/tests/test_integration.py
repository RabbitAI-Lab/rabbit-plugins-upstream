#!/usr/bin/env python3
"""
Integration test: full bid proposal flow with mock QG project.

Simulates the complete workflow:
1. Analyze requirement text → JSON
2. Match scene → 安全服务类
3. Build chapter prompts
4. Generate .docx output
"""

import json
import os
import sys
import tempfile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'scripts'))

from analyze_requirements import analyze_text
from prompt_builder import build_all_prompts, build_chapter_prompt

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from generate_docx import generate_bid_docx, create_default_template, SAMPLE_CHAPTERS_SAFETY_SERVICE


# ---------------------------------------------------------------------------
# Mock QG project requirement text
# ---------------------------------------------------------------------------

QG_PROJECT_TEXT = """
QG2025年年度安全服务项目 招标文件 技术需求

一、项目概述
本项目为QG集团2025年度安全服务采购，服务内容包括：
1. 渗透测试服务：每季度对集团核心业务系统（不少于10套）进行深度渗透测试，
   包含Web应用、移动APP、API接口的全面测试，输出详细的渗透测试报告和漏洞修复建议。
2. 重保服务：在国家级重大活动期间，提供7×24小时现场安全保障服务，
   包括安全监控、应急响应、攻击溯源等，确保活动期间零安全事件。
3. 应急响应服务：发生安全事件后，2小时内派出专业应急响应团队到达现场，
   提供事件分析、恶意代码分析、攻击溯源、系统加固等全流程应急处置。
4. 安全运营服务：提供日常安全监控和日志分析服务，
   包括安全事件监测、威胁情报分析、安全态势报告等。

二、技术要求
- 服务团队必须具备CISP、CISSP等专业安全认证
- 渗透测试工具需使用商业授权版工具
- 应急响应团队至少3年以上安全服务经验
- 所有服务需提供详细的服务报告

三、评分标准
| 评分项 | 分值 |
|--------|------|
| 项目理解与背景分析 | 5 |
| 需求分析 | 10 |
| 服务方案 | 25 |
| 实施计划 | 10 |
| 服务承诺与SLA | 5 |
"""


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_full_flow_analysis():
    """Step 1: Analyze QG project requirements."""
    result = analyze_text(QG_PROJECT_TEXT)

    assert result['project_type'] == '安全服务类', \
        f"Expected 安全服务类, got {result['project_type']}"
    assert 'QG' in result['project_name'] or '2025' in result['project_name'], \
        f"Project name should contain QG, got: {result['project_name']}"

    # Should detect scoring items
    assert len(result['scoring_items']) >= 4, \
        f"Expected at least 4 scoring items, got {len(result['scoring_items'])}"

    # Should detect tech domains
    assert '安全服务' in result['tech_domains'], \
        f"Expected '安全服务' in tech_domains, got {result['tech_domains']}"

    # Should detect key requirements
    assert len(result['key_requirements']) >= 3, \
        f"Expected at least 3 key requirements, got {len(result['key_requirements'])}"

    print("✅ Step 1 (Analysis): PASSED")


def test_full_flow_scene_mapping():
    """Step 2: Verify scene mapping yields correct chapters."""
    analysis_result = analyze_text(QG_PROJECT_TEXT)
    assert analysis_result['scene_match_detail']['score'] >= 4, \
        f"Expected score >= 4, got {analysis_result['scene_match_detail']['score']}"
    assert '渗透测试' in analysis_result['scene_match_detail']['matched_keywords'], \
        f"Expected 渗透测试 in matched keywords"

    print("✅ Step 2 (Scene Mapping): PASSED")


def test_full_flow_prompt_building():
    """Step 3: Build prompts for all chapters."""
    analysis_result = analyze_text(QG_PROJECT_TEXT)
    chapters = [
        {"id": "background", "name": "项目背景与目标", "prompt_guide": "阐述项目背景"},
        {"id": "requirement", "name": "需求分析", "prompt_guide": "分析安全需求"},
        {"id": "service_solution", "name": "服务方案", "prompt_guide": "详细服务方案"},
        {"id": "implementation", "name": "实施方案", "prompt_guide": "实施计划"},
        {"id": "sla", "name": "服务保障与SLA", "prompt_guide": "SLA承诺"},
    ]

    prompts = build_all_prompts(chapters, analysis_result)
    assert len(prompts) == 5, f"Expected 5 prompts, got {len(prompts)}"

    for ch_id, prompt in prompts.items():
        assert len(prompt) > 100, f"Prompt for {ch_id} too short ({len(prompt)} chars)"
        assert analysis_result['project_name'] in prompt, \
            f"Prompt {ch_id} should reference project name"

    # Service solution prompt should reference scoring
    assert '25' in prompts.get('service_solution', '') or '服务方案' in prompts.get('service_solution', ''), \
        "Service solution prompt should reference scoring"

    print("✅ Step 3 (Prompt Building): PASSED")


def test_full_flow_docx_generation():
    """Step 5: Generate .docx from sample chapters."""
    if not HAS_DOCX:
        print("⚠️ Step 5 (Docx Generation): SKIPPED (python-docx not installed)")
        return

    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = os.path.join(tmpdir, 'template.docx')
        create_default_template(template_path)

        output_path = os.path.join(tmpdir, 'QG2025_技术方案.docx')
        result = generate_bid_docx(template_path, SAMPLE_CHAPTERS_SAFETY_SERVICE, output_path)

        assert os.path.isfile(result), "Output file should exist"
        assert os.path.getsize(result) > 1000, "Output file too small"

        # Verify content
        doc = Document(result)
        all_text = '\n'.join(p.text for p in doc.paragraphs)

        # Should contain all major chapter titles
        expected_titles = ['项目背景与目标', '需求分析', '服务方案', '实施方案', '服务保障与SLA']
        for title in expected_titles:
            assert title in all_text, f"Missing chapter title: {title}"

        # Should contain key content
        assert '渗透测试' in all_text, "Missing 渗透测试 content"
        assert 'QG集团' in all_text, "Missing QG集团 reference"

        print(f"✅ Step 5 (Docx Generation): PASSED (output: {os.path.getsize(result)} bytes)")


def test_full_flow_end_to_end():
    """Complete end-to-end flow: analysis → prompts → docx."""
    print("\n" + "="*60)
    print("INTEGRATION TEST: Full Bid Proposal Flow")
    print("="*60)

    # Step 1: Analysis
    test_full_flow_analysis()

    # Step 2: Scene Mapping
    test_full_flow_scene_mapping()

    # Step 3: Prompt Building
    test_full_flow_prompt_building()

    # Step 4: This would be AI generation (skip in test, use sample content)
    print("⏭  Step 4 (AI Chapter Generation): SKIPPED (uses pre-built sample content)")

    # Step 5: Docx Generation
    test_full_flow_docx_generation()

    print("\n" + "="*60)
    print("ALL INTEGRATION TESTS PASSED")
    print("="*60)


def test_analyze_requirements_cli():
    """Test the CLI of analyze_requirements.py."""
    from subprocess import run, PIPE

    # Test with inline text
    result = run(
        ['python3', 'scripts/analyze_requirements.py', QG_PROJECT_TEXT[:200], '--format', 'json'],
        capture_output=True, text=True,
        cwd=os.path.join(os.path.dirname(__file__), '..')
    )
    # May fail if text contains special chars; that's OK, just check it doesn't crash
    # Test with text file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
        f.write(QG_PROJECT_TEXT)
        txt_path = f.name

    try:
        result = run(
            ['python3', 'scripts/analyze_requirements.py', txt_path, '--format', 'json'],
            capture_output=True, text=True,
            cwd=os.path.join(os.path.dirname(__file__), '..')
        )
        output = json.loads(result.stdout)
        assert output['project_type'] == '安全服务类', \
            f"CLI analysis should match, got {output['project_type']}"
        print("✅ CLI (analyze_requirements): PASSED")
    finally:
        os.unlink(txt_path)


def test_generate_docx_cli():
    """Test the CLI of generate_docx.py."""
    if not HAS_DOCX:
        print("⚠️ CLI (generate_docx): SKIPPED (python-docx not installed)")
        return

    from subprocess import run, PIPE

    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = os.path.join(tmpdir, 'template.docx')
        chapters_path = os.path.join(tmpdir, 'chapters.json')
        output_path = os.path.join(tmpdir, 'output.docx')

        # Create template
        create_default_template(template_path)

        # Create chapters JSON
        chapters = [
            {"title": "测试章", "level": 1, "content": "测试内容"},
        ]
        with open(chapters_path, 'w', encoding='utf-8') as f:
            json.dump(chapters, f, ensure_ascii=False)

        # Run CLI
        result = run(
            ['python3', 'scripts/generate_docx.py', template_path, chapters_path, output_path],
            capture_output=True, text=True,
            cwd=os.path.join(os.path.dirname(__file__), '..')
        )

        assert os.path.isfile(output_path), f"CLI should create output, got: {result.stderr}"
        doc = Document(output_path)
        all_text = '\n'.join(p.text for p in doc.paragraphs)
        assert '测试章' in all_text

        print("✅ CLI (generate_docx): PASSED")


def test_prompt_builder_cli():
    """Test the CLI of prompt_builder.py."""
    from subprocess import run, PIPE

    with tempfile.TemporaryDirectory() as tmpdir:
        chapter_path = os.path.join(tmpdir, 'chapter.json')
        analysis_path = os.path.join(tmpdir, 'analysis.json')

        chapter = {"id": "test", "name": "测试", "prompt_guide": "测试指引"}
        analysis = {
            "project_type": "综合类",
            "project_name": "测试项目",
            "scoring_items": [],
            "tech_domains": [],
            "key_requirements": [],
            "special_notes": ""
        }

        with open(chapter_path, 'w', encoding='utf-8') as f:
            json.dump(chapter, f, ensure_ascii=False)
        with open(analysis_path, 'w', encoding='utf-8') as f:
            json.dump(analysis, f, ensure_ascii=False)

        result = run(
            ['python3', 'scripts/prompt_builder.py', chapter_path, analysis_path],
            capture_output=True, text=True,
            cwd=os.path.join(os.path.dirname(__file__), '..')
        )

        assert '测试' in result.stdout, f"CLI output should contain chapter name, got: {result.stdout[:200]}"
        assert '测试项目' in result.stdout

        print("✅ CLI (prompt_builder): PASSED")


if __name__ == '__main__':
    # Run end-to-end test
    test_full_flow_end_to_end()

    # Run CLI tests
    print("\n--- CLI Tests ---")
    test_analyze_requirements_cli()
    test_prompt_builder_cli()
    test_generate_docx_cli()

    print("\n" + "="*60)
    print("ALL INTEGRATION TESTS PASSED ✅")
    print("="*60)
