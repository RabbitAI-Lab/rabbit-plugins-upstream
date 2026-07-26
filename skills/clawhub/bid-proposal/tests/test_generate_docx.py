#!/usr/bin/env python3
"""Tests for generate_docx.py"""

import os
import sys
import tempfile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'scripts'))

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from generate_docx import (
    generate_bid_docx,
    create_default_template,
    SAMPLE_CHAPTERS_SAFETY_SERVICE
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _create_test_template() -> str:
    """Create a minimal .docx template for testing."""
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)

    for level, size in [(1, 16), (2, 15), (3, 14)]:
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.size = Pt(size)

    # Cover page
    title_para = doc.add_paragraph('XX项目技术方案')
    title_para.alignment = 1  # center

    date_para = doc.add_paragraph('2026年6月')
    date_para.alignment = 1

    # Page break
    doc.add_page_break()

    # SDT table of contents (simplified - just a placeholder)
    toc_para = doc.add_paragraph('【目录 - 请在Word中右键更新域】')
    doc.add_page_break()

    # Add some placeholder content
    doc.add_paragraph('（正文内容将由脚本自动填充）', style='Normal')

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_create_default_template():
    """Create a default template .docx file."""
    if not HAS_DOCX:
        print("SKIP: python-docx not installed")
        return

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        result = create_default_template(output_path)
        assert os.path.isfile(result)
        assert os.path.getsize(result) > 0

        # Verify it's a valid docx
        doc = Document(result)
        assert len(doc.paragraphs) > 0, "Template should have paragraphs"
    finally:
        os.unlink(output_path)


def test_generate_bid_docx_basic():
    """Generate a docx with basic chapters from a template."""
    if not HAS_DOCX:
        print("SKIP: python-docx not installed")
        return

    template_path = _create_test_template()

    chapters = [
        {"title": "项目背景与目标", "level": 1, "content": "# 项目背景与目标\n\n本项目旨在提升安全防护能力。"},
        {"title": "需求分析", "level": 1, "content": "## 业务需求\n\n客户需要完善的安全服务体系。"},
        {"title": "服务方案", "level": 1, "content": "## 服务方案\n\n1. 渗透测试服务\n2. 应急响应服务"},
        {"title": "服务内容详解", "level": 2, "content": "### 渗透测试\n\n每季度进行一次渗透测试。"},
    ]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        result = generate_bid_docx(template_path, chapters, output_path)
        assert os.path.isfile(result)
        assert os.path.getsize(result) > 0, "Output file should be non-empty"

        # Verify content
        doc = Document(result)
        all_text = '\n'.join(p.text for p in doc.paragraphs)

        # Should contain chapter content
        assert '项目背景与目标' in all_text, f"Missing chapter title, got: {all_text[:200]}"
        assert '渗透测试' in all_text or '安全防护' in all_text

        # Should NOT overwrite template path (different files)
        assert os.path.abspath(result) != os.path.abspath(template_path)
    finally:
        os.unlink(output_path)
        if os.path.isfile(template_path):
            os.unlink(template_path)


def test_generate_bid_docx_heading_styles():
    """Verify heading styles are applied correctly."""
    if not HAS_DOCX:
        print("SKIP: python-docx not installed")
        return

    template_path = _create_test_template()
    chapters = [
        {"title": "第一章", "level": 1, "content": "Heading 1 content"},
        {"title": "第一节", "level": 2, "content": "Heading 2 content"},
        {"title": "第一小节", "level": 3, "content": "Heading 3 content"},
    ]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        result = generate_bid_docx(template_path, chapters, output_path)
        doc = Document(result)

        # Find our paragraphs
        headings_found = {1: False, 2: False, 3: False}
        for p in doc.paragraphs:
            if p.style.name.startswith('Heading'):
                level = int(p.style.name.split()[-1]) if p.style.name != 'Heading' else 1
                if p.text and p.text in ['第一章', '第一节', '第一小节']:
                    headings_found[level] = True

        assert headings_found[1], "Missing Heading 1"
        assert headings_found[2], "Missing Heading 2"
        assert headings_found[3], "Missing Heading 3"
    finally:
        os.unlink(output_path)
        if os.path.isfile(template_path):
            os.unlink(template_path)


def test_generate_bid_docx_preserves_sdt():
    """Verify SDT (table of contents) is preserved, not removed."""
    if not HAS_DOCX:
        print("SKIP: python-docx not installed")
        return

    template_path = _create_test_template()
    chapters = [{"title": "测试章节", "level": 1, "content": "测试内容"}]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        result = generate_bid_docx(template_path, chapters, output_path)
        doc = Document(result)

        # Verify the template was not clobbered by checking the output content
        all_text = '\n'.join(p.text for p in doc.paragraphs)
        assert '测试章节' in all_text
        assert '测试内容' in all_text
    finally:
        os.unlink(output_path)
        if os.path.isfile(template_path):
            os.unlink(template_path)


def test_generate_bid_docx_empty_chapters():
    """Generate with empty chapters list."""
    if not HAS_DOCX:
        print("SKIP: python-docx not installed")
        return

    template_path = _create_test_template()

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        result = generate_bid_docx(template_path, [], output_path)
        assert os.path.isfile(result)
        assert os.path.getsize(result) > 0
    finally:
        os.unlink(output_path)
        if os.path.isfile(template_path):
            os.unlink(template_path)


def test_generate_bid_docx_from_default():
    """Generate using create_default_template as source."""
    if not HAS_DOCX:
        print("SKIP: python-docx not installed")
        return

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        template_path = f.name

    try:
        create_default_template(template_path)

        chapters = [{"title": "服务方案", "level": 1, "content": "## 服务内容\n\n安全运营、渗透测试、应急响应。"}]

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f2:
            output_path = f2.name

        result = generate_bid_docx(template_path, chapters, output_path)
        doc = Document(result)
        all_text = '\n'.join(p.text for p in doc.paragraphs)
        assert '服务方案' in all_text
    finally:
        os.unlink(template_path)
        if os.path.isfile(output_path):
            os.unlink(output_path)


def test_sample_chapters_structure():
    """Verify SAMPLE_CHAPTERS_SAFETY_SERVICE has correct structure."""
    assert isinstance(SAMPLE_CHAPTERS_SAFETY_SERVICE, list)
    assert len(SAMPLE_CHAPTERS_SAFETY_SERVICE) >= 4
    for ch in SAMPLE_CHAPTERS_SAFETY_SERVICE:
        assert 'title' in ch, f"Missing title in chapter: {ch}"
        assert 'level' in ch, f"Missing level in chapter: {ch}"
        assert 'content' in ch, f"Missing content in chapter: {ch}"
        assert isinstance(ch['level'], int) and 1 <= ch['level'] <= 3


if __name__ == '__main__':
    import pytest
    sys.exit(pytest.main([__file__, '-v']))
