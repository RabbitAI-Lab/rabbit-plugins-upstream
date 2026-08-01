#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""docx_gen.py — 从 Markdown 生成中文排版规范的 .docx。

支持的 Markdown 语法（按行解析，保持简单）：
    # 一级标题 / ## 二级标题 / ### 三级标题
    普通段落
    - 无序列表项
    > 引用

排版规范：A4 页面；正文宋体小四（12pt），标题黑体；西文字体 Times New Roman。
中文字体通过 rPr.rFonts 的 w:eastAsia 属性设置，确保在 Word/WPS 中正确显示。

用法：
    python scripts/docx_gen.py 输入.md -o 输出.docx

依赖：python-docx（pip install python-docx）。
"""

import argparse
import sys

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False


def set_run_fonts(run, east_asia, ascii_font="Times New Roman", size_pt=None, bold=None):
    """同时设置 run 的西文字体与中文（eastAsia）字体。"""
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def setup_page(doc):
    """页面设置：A4 纸张，常规页边距。"""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def add_heading(doc, text, level):
    """添加标题：黑体，字号按级别递减。"""
    sizes = {1: 16, 2: 14, 3: 13}
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_fonts(run, "黑体", size_pt=sizes.get(level, 12), bold=True)
    return paragraph


def add_body_paragraph(doc, text):
    """添加正文段落：宋体小四（12pt），首行缩进两字符。"""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 两个 12pt 字符宽
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_run_fonts(run, "宋体", size_pt=12)
    return paragraph


def add_list_item(doc, text):
    """添加无序列表项。"""
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_fonts(run, "宋体", size_pt=12)
    return paragraph


def add_quote(doc, text):
    """添加引用段落：楷体，左侧缩进。"""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_run_fonts(run, "楷体", size_pt=12)
    run.font.italic = True
    return paragraph


def parse_markdown(md_text):
    """把 Markdown 文本解析为 (类型, 内容, 级别) 元组列表。

    类型：heading / paragraph / list / quote
    """
    blocks = []
    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue  # 空行仅作为段落分隔，直接跳过
        if line.startswith("### "):
            blocks.append(("heading", line[4:].strip(), 3))
        elif line.startswith("## "):
            blocks.append(("heading", line[3:].strip(), 2))
        elif line.startswith("# "):
            blocks.append(("heading", line[2:].strip(), 1))
        elif line.startswith("> "):
            blocks.append(("quote", line[2:].strip(), 0))
        elif line.startswith("- "):
            blocks.append(("list", line[2:].strip(), 0))
        else:
            blocks.append(("paragraph", line.strip(), 0))
    return blocks


def generate(input_path, output_path):
    """读取 Markdown 文件并生成 docx，返回生成的块数量。"""
    with open(input_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    blocks = parse_markdown(md_text)
    if not blocks:
        print("[警告] 输入的 Markdown 文件没有可转换的内容。")

    doc = Document()
    setup_page(doc)
    for kind, text, level in blocks:
        if kind == "heading":
            add_heading(doc, text, level)
        elif kind == "list":
            add_list_item(doc, text)
        elif kind == "quote":
            add_quote(doc, text)
        else:
            add_body_paragraph(doc, text)

    doc.save(output_path)
    return len(blocks)


def main():
    parser = argparse.ArgumentParser(
        description="从 Markdown 生成中文排版规范的 .docx（A4、宋体正文、黑体标题、西文 Times New Roman）。",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "支持的语法：# / ## / ### 标题、普通段落、- 列表、> 引用。\n"
            "示例：\n"
            "  python scripts/docx_gen.py 输入.md -o 输出.docx"
        ),
    )
    parser.add_argument("input", help="输入的 Markdown 文件路径")
    parser.add_argument("-o", "--output", required=True, help="输出的 .docx 文件路径")
    args = parser.parse_args()

    if not _DOCX_OK:
        print("[错误] 缺少第三方库 python-docx，请先安装：")
        print("       pip install python-docx")
        sys.exit(2)

    try:
        count = generate(args.input, args.output)
    except FileNotFoundError:
        print(f"[错误] 找不到输入文件：{args.input}")
        sys.exit(1)
    except UnicodeDecodeError:
        print(f"[错误] 无法按 UTF-8 读取文件：{args.input}，请确认文件编码。")
        sys.exit(1)

    print(f"[完成] 已生成 {args.output}（共转换 {count} 个内容块）。")


if __name__ == "__main__":
    main()
