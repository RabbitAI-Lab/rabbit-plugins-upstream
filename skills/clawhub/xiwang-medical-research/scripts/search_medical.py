#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
烯旺医疗科研成果检索脚本
支持扫描、关键词检索、主题检索、疾病检索和文档列表功能。
可处理 .docx 和 .pdf 两种文件格式。
"""

import os
import sys
import json
import re
from pathlib import Path

# 尝试导入文档处理库
try:
    import docx as python_docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False


# ============================================================
# 主题分类体系
# ============================================================

TOPIC_KEYWORDS = {
    "tumor": {
        "label": "肿瘤/癌症治疗",
        "keywords": ["肿瘤", "癌症", "抗癌", "乳腺癌", "前列腺癌", "结肠癌",
                      "无创", "光热疗", "化疗", "纳米载体", "氧化石墨烯",
                      "MDA-MB-231", "PC3", "LoVo", "抑制", "凋亡", "转移"]
    },
    "pain": {
        "label": "骨科/疼痛",
        "keywords": ["疼痛", "关节炎", "腰腿痛", "颈椎", "膝骨", "骨性",
                      "腰椎", "椎间盘", "护腰", "护膝", "骨关节"]
    },
    "ent": {
        "label": "五官科",
        "keywords": ["耳鼻喉", "干眼症", "鼻炎", "扁桃体", "鼻罩", "眼罩",
                      "颈罩", "变应性", "尘螨", "泪液", "泪膜", "眼科"]
    },
    "gynecology": {
        "label": "妇科",
        "keywords": ["痛经", "乳腺增生", "甲状腺结节", "妇科", "寒湿",
                      "寒凝血瘀", "乳腺", "结节", "内分泌"]
    },
    "skin": {
        "label": "皮肤/美容",
        "keywords": ["黄褐斑", "面部皮肤", "毛发生长", "皮肤", "面膜",
                      "米诺地尔", "美容", "胶原蛋白"]
    },
    "sleep": {
        "label": "神经/睡眠",
        "keywords": ["失眠", "顽固性失眠", "脑卒中", "焦虑", "肥胖",
                      "神经炎症", "睡眠", "抑郁"]
    },
    "cardiovascular": {
        "label": "心血管",
        "keywords": ["心脑血管", "红细胞", "缗钱状", "心血管", "血液",
                      "微循环", "血压", "血管"]
    },
    "tcm": {
        "label": "中医/艾灸",
        "keywords": ["热敏灸", "艾灸", "中医", "经络", "阳气", "灸感",
                      "得气", "胃脘痛", "虚寒", "悬灸", "针灸", "好转反应",
                      "瞑眩反应", "治未病", "寒湿"]
    },
    "theory": {
        "label": "基础理论/机制",
        "keywords": ["非热效应", "远红外", "共振", "红外线", "生物效应",
                      "石墨烯", "6-14微米", "8.0μm", "吸收谱", "特征峰",
                      "碳纤维", "发热材料"]
    },
    "enterprise": {
        "label": "企业/品牌",
        "keywords": ["烯旺", "集团", "企业", "大事记", "医疗器械",
                      "认证", "冬奥会", "二类医疗器械", "军工"]
    },
    "product": {
        "label": "产品应用",
        "keywords": ["光波房", "能量房", "悬灸仪", "穿戴", "光波",
                      "能量", "汗蒸", "理疗", "护腰", "护膝", "面罩"]
    }
}

# 疾病名称到主题的映射
DISEASE_MAP = {
    "肿瘤": "tumor", "癌症": "tumor", "乳腺癌": "tumor", "前列腺癌": "tumor",
    "结肠癌": "tumor", "甲状腺结节": "gynecology",
    "膝骨性关节炎": "pain", "骨关节炎": "pain", "腰腿痛": "pain",
    "颈椎病": "pain", "腰椎间盘突出": "pain",
    "干眼症": "ent", "鼻炎": "ent", "变应性鼻炎": "ent", "扁桃体炎": "ent",
    "耳鼻喉": "ent",
    "痛经": "gynecology", "乳腺增生": "gynecology",
    "黄褐斑": "skin", "脱发": "skin", "毛发生长": "skin",
    "失眠": "sleep", "顽固性失眠": "sleep", "脑卒中": "sleep",
    "焦虑": "sleep", "肥胖": "sleep",
    "心脑血管": "cardiovascular",
    "胃脘痛": "tcm", "虚寒": "tcm",
}


# ============================================================
# 文档读取函数
# ============================================================

def read_docx(filepath):
    """读取 .docx 文件，返回段落列表"""
    if not HAS_DOCX:
        return []
    try:
        doc = python_docx.Document(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # 也读取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        return paragraphs
    except Exception as e:
        return [f"[读取错误: {e}]"]


def read_pdf(filepath):
    """读取 .pdf 文件，返回文本块列表"""
    if not HAS_PDF:
        return []
    paragraphs = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    # 按换行分割并合并短行
                    lines = text.split('\n')
                    current = ""
                    for line in lines:
                        line = line.strip()
                        if not line:
                            if current:
                                paragraphs.append(current)
                                current = ""
                            continue
                        if len(current) < 30:
                            current = (current + line).strip()
                        else:
                            if current:
                                paragraphs.append(current)
                            current = line
                    if current:
                        paragraphs.append(current)
        return paragraphs
    except Exception as e:
        return [f"[读取错误: {e}]"]


def read_document(filepath):
    """根据文件类型读取文档"""
    ext = Path(filepath).suffix.lower()
    if ext == '.docx':
        return read_docx(filepath)
    elif ext == '.pdf':
        return read_pdf(filepath)
    return []


# ============================================================
# 文件扫描与元数据
# ============================================================

def get_file_category(filepath, base_dir):
    """根据文件路径获取分类信息"""
    rel_path = os.path.relpath(filepath, base_dir)
    parts = Path(rel_path).parts

    if len(parts) == 1:
        return "根目录文件"
    elif parts[0] == "陈医生医疗内容":
        return "中医科普"
    elif parts[0] == "公众号医疗及产品内容":
        return "公众号推文"
    elif parts[0] == "企业资料汇总":
        return "企业资料"
    elif parts[0] == "医疗论文合集（原文）":
        return "学术论文"
    elif parts[0] == "热敏灸实用读本-陈日新":
        return "热敏灸专著"
    else:
        return parts[0]


def scan_directory(dirpath):
    """扫描目录下所有文档文件"""
    results = []
    for root, dirs, files in os.walk(dirpath):
        for f in files:
            if f.startswith('~$'):
                continue
            ext = Path(f).suffix.lower()
            if ext in ('.docx', '.pdf'):
                filepath = os.path.join(root, f)
                size = os.path.getsize(filepath)
                category = get_file_category(filepath, dirpath)
                results.append({
                    "file": f,
                    "path": filepath,
                    "relative_path": os.path.relpath(filepath, dirpath),
                    "category": category,
                    "type": ext.lstrip('.'),
                    "size_bytes": size,
                    "size_kb": round(size / 1024, 1)
                })
    return results


# ============================================================
# 检索函数
# ============================================================

def search_in_text(paragraphs, keywords):
    """在段落列表中搜索关键词，返回匹配结果"""
    matches = []
    for i, para in enumerate(paragraphs):
        matched = [kw for kw in keywords if kw in para]
        if matched:
            context_before = paragraphs[i-1] if i > 0 else ""
            context_after = paragraphs[i+1] if i < len(paragraphs)-1 else ""
            matches.append({
                "paragraph_index": i,
                "matched_keywords": matched,
                "content": para,
                "context_before": context_before,
                "context_after": context_after
            })
    return matches


def search_content(dirpath, keywords, max_results=50):
    """跨所有文档搜索关键词"""
    if isinstance(keywords, str):
        keywords = keywords.split()

    results = []
    files = scan_directory(dirpath)

    for file_info in files:
        paragraphs = read_document(file_info["path"])
        if not paragraphs:
            continue

        matches = search_in_text(paragraphs, keywords)
        if matches:
            total_chars = sum(len(p) for p in paragraphs)
            results.append({
                "file": file_info["file"],
                "path": file_info["path"],
                "relative_path": file_info["relative_path"],
                "category": file_info["category"],
                "type": file_info["type"],
                "total_paragraphs": len(paragraphs),
                "total_chars": total_chars,
                "match_count": len(matches),
                "matches": matches[:10]  # 每个文件最多返回10条匹配
            })

    # 按匹配数排序
    results.sort(key=lambda x: x["match_count"], reverse=True)
    return results[:max_results]


def search_by_topic(dirpath, topic, max_results=50):
    """按预定义主题检索"""
    if topic == "all":
        all_keywords = set()
        topic_map = {}
        for t, info in TOPIC_KEYWORDS.items():
            for kw in info["keywords"]:
                all_keywords.add(kw)
                topic_map[kw] = t
        # 搜索所有关键词
        results = search_content(dirpath, list(all_keywords), max_results)
        # 为每个匹配标注主题
        for r in results:
            for m in r["matches"]:
                m["topics"] = list(set(topic_map.get(kw, "unknown") for kw in m["matched_keywords"]))
        return {"topic": "all", "results": results}

    if topic not in TOPIC_KEYWORDS:
        return {"error": f"未知主题: {topic}。可用主题: {', '.join(TOPIC_KEYWORDS.keys())}"}

    keywords = TOPIC_KEYWORDS[topic]["keywords"]
    results = search_content(dirpath, keywords, max_results)
    return {
        "topic": topic,
        "label": TOPIC_KEYWORDS[topic]["label"],
        "keywords_used": keywords,
        "results": results
    }


def search_by_disease(dirpath, disease_name, max_results=50):
    """按疾病名称检索"""
    # 先查找疾病对应的主题
    topic = None
    for key, t in DISEASE_MAP.items():
        if key in disease_name or disease_name in key:
            topic = t
            break

    keywords = [disease_name]
    if topic and topic in TOPIC_KEYWORDS:
        keywords.extend(TOPIC_KEYWORDS[topic]["keywords"][:5])

    results = search_content(dirpath, keywords, max_results)
    return {
        "disease": disease_name,
        "matched_topic": topic,
        "keywords_used": keywords,
        "results": results
    }


def list_documents(dirpath):
    """列出所有文档及其元信息"""
    files = scan_directory(dirpath)
    # 按分类分组
    by_category = {}
    for f in files:
        cat = f["category"]
        if cat not in by_category:
            by_category[cat] = []
        by_category[cat].append(f)

    return {
        "total_files": len(files),
        "categories": by_category
    }


def get_topics():
    """返回所有可用主题"""
    return {t: info["label"] for t, info in TOPIC_KEYWORDS.items()}


# ============================================================
# 主入口
# ============================================================

def main():
    if len(sys.argv) < 2:
        print(json.dumps({
            "error": "缺少命令参数",
            "usage": "python search_medical.py <command> <dirpath> [args]",
            "commands": {
                "scan": "扫描全部文件，输出清单",
                "list": "列出所有文档及分类信息",
                "search": "按关键词搜索（空格分隔）",
                "topic": "按预定义主题检索",
                "disease": "按疾病名称检索",
                "topics": "列出所有可用主题"
            }
        }, ensure_ascii=False, indent=2))
        sys.exit(1)

    command = sys.argv[1]

    if command == "topics":
        print(json.dumps(get_topics(), ensure_ascii=False, indent=2))
        return

    if len(sys.argv) < 3:
        print(json.dumps({"error": "缺少目录路径参数"}, ensure_ascii=False))
        sys.exit(1)

    dirpath = sys.argv[2]

    if command == "scan":
        results = scan_directory(dirpath)
        print(json.dumps({
            "total_files": len(results),
            "files": results
        }, ensure_ascii=False, indent=2))

    elif command == "list":
        results = list_documents(dirpath)
        print(json.dumps(results, ensure_ascii=False, indent=2))

    elif command == "search":
        if len(sys.argv) < 4:
            print(json.dumps({"error": "缺少搜索关键词"}, ensure_ascii=False))
            sys.exit(1)
        keywords = sys.argv[3]
        results = search_content(dirpath, keywords)
        print(json.dumps({
            "keywords": keywords.split(),
            "total_matched_files": len(results),
            "results": results
        }, ensure_ascii=False, indent=2))

    elif command == "topic":
        if len(sys.argv) < 4:
            print(json.dumps({
                "error": "缺少主题参数",
                "available_topics": get_topics()
            }, ensure_ascii=False))
            sys.exit(1)
        topic = sys.argv[3]
        results = search_by_topic(dirpath, topic)
        print(json.dumps(results, ensure_ascii=False, indent=2))

    elif command == "disease":
        if len(sys.argv) < 4:
            print(json.dumps({"error": "缺少疾病名称"}, ensure_ascii=False))
            sys.exit(1)
        disease = sys.argv[3]
        results = search_by_disease(dirpath, disease)
        print(json.dumps(results, ensure_ascii=False, indent=2))

    else:
        print(json.dumps({"error": f"未知命令: {command}"}, ensure_ascii=False))
        sys.exit(1)


if __name__ == "__main__":
    main()
