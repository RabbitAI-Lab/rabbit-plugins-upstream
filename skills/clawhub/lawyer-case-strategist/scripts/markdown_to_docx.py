#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
markdown_to_docx.py - Markdown 批量转 DOCX
为建工实务·裁决视角系列技能提供文档导出能力
"""

import os
import sys
import zipfile
import argparse
from pathlib import Path
from typing import List, Optional

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("缺少依赖：python-docx")
    print("请运行：pip install python-docx")
    sys.exit(1)


def setup_default_styles(doc: Document) -> None:
    """设置文档默认样式：A4 / 宋体（正文）/ 黑体（标题）/ 1.5倍行距 / 标准页码"""
    # 设置页面为A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置标题样式
    for level, font_name in [(1, '黑体'), (2, '黑体'), (3, '黑体'), (4, '黑体')]:
        try:
            heading_style = doc.styles[f'Heading {level}']
            heading_style.font.name = font_name
            heading_style.font.size = Pt(16 - (level - 1) * 2)
            heading_style.font.bold = True
            heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        except KeyError:
            pass

    # 设置行距
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5


def add_page_numbers(doc: Document) -> None:
    """添加页码"""
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fldChar1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    instrText = run._element.makeelement(qn('w:instrText'), {})
    instrText.text = 'PAGE'
    fldChar2 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)


def parse_markdown_line(line: str) -> dict:
    """解析单行Markdown"""
    line = line.rstrip()

    # 标题
    if line.startswith('# '):
        return {'type': 'h1', 'text': line[2:].strip()}
    elif line.startswith('## '):
        return {'type': 'h2', 'text': line[3:].strip()}
    elif line.startswith('### '):
        return {'type': 'h3', 'text': line[4:].strip()}
    elif line.startswith('#### '):
        return {'type': 'h4', 'text': line[5:].strip()}
    elif line.startswith('##### '):
        return {'type': 'h5', 'text': line[6:].strip()}

    # 引用
    elif line.startswith('> '):
        return {'type': 'quote', 'text': line[2:].strip()}

    # 无序列表
    elif line.startswith('- ') or line.startswith('* '):
        return {'type': 'ul', 'text': line[2:].strip()}

    # 有序列表
    elif len(line) > 2 and line[0].isdigit() and (line[1] == '.' or (line[1].isdigit() and line[2] == '.')):
        # 找到第一个非数字字符
        idx = 0
        while idx < len(line) and (line[idx].isdigit() or line[idx] == '.'):
            idx += 1
        return {'type': 'ol', 'text': line[idx:].strip()}

    # 水平线
    elif line.strip() == '---':
        return {'type': 'hr'}

    # 表格行（简单处理）
    elif line.startswith('|') and line.endswith('|'):
        return {'type': 'table_row', 'text': line}

    # 空行
    elif line.strip() == '':
        return {'type': 'blank'}

    # 普通段落
    else:
        return {'type': 'p', 'text': line}


def remove_markdown_formatting(text: str) -> str:
    """移除行内Markdown格式"""
    # 粗体 **text**
    import re
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # 斜体 *text*
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # 代码 `text`
    text = re.sub(r'`(.+?)`', r'\1', text)
    # 链接 [text](url)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def add_paragraph_to_doc(doc: Document, element: dict) -> None:
    """将解析的Markdown元素添加到文档"""
    text = remove_markdown_formatting(element['text'])

    if element['type'] == 'h1':
        p = doc.add_heading(text, level=1)
    elif element['type'] == 'h2':
        p = doc.add_heading(text, level=2)
    elif element['type'] == 'h3':
        p = doc.add_heading(text, level=3)
    elif element['type'] == 'h4':
        p = doc.add_heading(text, level=4)
    elif element['type'] == 'h5':
        p = doc.add_heading(text, level=5)
    elif element['type'] == 'quote':
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Cm(0.74)
        p.runs[0].italic = True if p.runs else None
    elif element['type'] == 'ul':
        doc.add_paragraph(text, style='List Bullet')
    elif element['type'] == 'ol':
        doc.add_paragraph(text, style='List Number')
    elif element['type'] == 'hr':
        p = doc.add_paragraph('─' * 30)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif element['type'] == 'p':
        if text:
            doc.add_paragraph(text)
    elif element['type'] == 'blank':
        pass
    elif element['type'] == 'table_row':
        # 简单处理：作为段落
        if text and not all(c in '|-: ' for c in text):
            doc.add_paragraph(text)


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> bool:
    """将单个Markdown文件转为DOCX"""
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        doc = Document()
        setup_default_styles(doc)
        add_page_numbers(doc)

        # 跳过YAML frontmatter
        in_frontmatter = False
        start_idx = 0
        if lines and lines[0].strip() == '---':
            in_frontmatter = True
            for i in range(1, len(lines)):
                if lines[i].strip() == '---':
                    start_idx = i + 1
                    break

        for line in lines[start_idx:]:
            element = parse_markdown_line(line)
            add_paragraph_to_doc(doc, element)

        # 确保输出目录存在
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(docx_path))
        print(f"✓ 已转换：{md_path.name} → {docx_path.name}")
        return True
    except Exception as e:
        print(f"✗ 转换失败：{md_path.name} - {e}")
        return False


def batch_convert(input_paths: List[str], output_dir: Optional[str] = None, package: bool = False) -> List[Path]:
    """批量转换多个路径（文件或目录）"""
    md_files: List[Path] = []
    output_root = Path(output_dir) if output_dir else None

    for input_path in input_paths:
        path = Path(input_path)
        if not path.exists():
            print(f"⚠ 路径不存在：{input_path}")
            continue

        if path.is_file():
            if path.suffix.lower() == '.md':
                md_files.append(path)
        elif path.is_dir():
            # 递归扫描
            for md_file in path.rglob('*.md'):
                md_files.append(md_file)

    if not md_files:
        print("未找到Markdown文件")
        return []

    converted_files: List[Path] = []
    for md_file in md_files:
        # 计算输出路径
        if output_root:
            # 保持相对目录结构
            try:
                rel_path = md_file.resolve().relative_to(Path.cwd())
            except ValueError:
                rel_path = md_file.name
            docx_path = output_root / rel_path.with_suffix('.docx')
        else:
            docx_path = md_file.with_suffix('.docx')

        if convert_markdown_to_docx(md_file, docx_path):
            converted_files.append(docx_path)

    # 打包
    if package and converted_files:
        zip_path = (output_root or Path('.')) / 'markdown_docx_bundle.zip'
        zip_path.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for docx_file in converted_files:
                zf.write(docx_file, docx_file.name)
        print(f"\n📦 已打包：{zip_path}（共{len(converted_files)}个文件）")

    return converted_files


def main():
    parser = argparse.ArgumentParser(
        description='Markdown 批量转 DOCX - 建工实务·裁决视角系列技能',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例：
  # 转换单个文件
  python markdown_to_docx.py output.md

  # 转换多个文件
  python markdown_to_docx.py file1.md file2.md

  # 转换整个目录
  python markdown_to_docx.py assets/

  # 指定输出目录
  python markdown_to_docx.py assets/ -o output/

  # 批量转换并打包
  python markdown_to_docx.py assets/ -o output/ --package
        """
    )
    parser.add_argument('inputs', nargs='+', help='输入文件或目录')
    parser.add_argument('-o', '--output', help='输出目录')
    parser.add_argument('--package', action='store_true', help='打包为ZIP')

    args = parser.parse_args()
    batch_convert(args.inputs, args.output, args.package)


if __name__ == '__main__':
    main()
