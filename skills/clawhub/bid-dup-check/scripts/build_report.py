#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_report.py - 标书查重 Skill 的报告渲染脚本

读取大模型分析产出的 findings JSON，渲染为：
  1. Word (.docx) 检测报告（带风险等级、表格、定位信息）
  2. 可选：Markdown 版本（用于对话内展示）

findings JSON 结构（由 SKILL.md 的 Step 4 产出，权威；本文件仅作渲染说明）：
{
  "overview": {
    "file_count": 2,
    "risk_total": 5,
    "high_risk_count": 2,
    "avg_similarity": 0.42,
    "summary": "一句话概述"
  },
  "key_info_collisions": [
    {"field":"公司名称","value":"XX公司","files":["a.docx"],
     "locations":["a.docx 第3段"],"level":"高","reason":"两家投标文件单位全称相同"}
  ],
  "text_similarity": [
    {"file_a":"a.docx","file_b":"b.docx","score":0.90,"level":"高","type":"直接复制",
     "reason":"连续三段落结构完全一致且关键数据相同",
     "segments":[{"text":"...","location":"a.docx 第X段 / b.docx 第Y段"}]}
  ],
  "attribute_warnings": [
    {"file":"a.docx","level":"高","fields":{"author":"张三","company":"XX"},
     "detail":"作者/公司相同","reason":"不同投标方文档作者元数据相同"}
  ],
  "table_similarity": [
    {"file_a":"a.docx","file_b":"b.docx","level":"中",
     "detail":"表格结构相似","reason":"报价表行列结构一致仅单价尾数不同"}
  ],
  "diff": [
    {"type":"add|delete|change","file_a":"a.docx","file_b":"b.docx","content":"..."}
  ],
  "limitations": ["文档 c.txt 含扫描页未提取文字，覆盖率不完整"],
  "baseline_note": "已使用招标文件做基线剔除" or null,
  "conclusion": "综合建议文本"
}
注：reason / type / limitations 为可选新增字段，渲染脚本已向后兼容；level 仅取 高/中/低。

依赖：python-docx
用法：
    python build_report.py --findings findings.json --out report.docx [--markdown report.md]
"""

import argparse
import json
import os
import sys

LEVEL_ICON = {"高": "🔴", "中": "🟡", "低": "🟢", "high": "🔴", "mid": "🟡", "low": "🟢"}


def _require(module_name, pip_name):
    try:
        return __import__(module_name)
    except ImportError:
        sys.stderr.write(
            f"[错误] 缺少依赖 `{pip_name}`。请执行：pip install {pip_name}\n"
        )
        sys.exit(2)


def _lvl(level):
    return LEVEL_ICON.get(level, "") or ""


def _md_escape(s):
    if not isinstance(s, str):
        s = str(s)
    return s.replace("|", "\\|").replace("\n", " ")


# ----------------------------------------------------------------------------
# DOCX 渲染
# ----------------------------------------------------------------------------
def build_docx(findings, out_path):
    docx = _require("docx", "python-docx")
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading("标书查重检测报告", level=0)

    ov = findings.get("overview", {})
    p = doc.add_paragraph()
    p.add_run("检测概览\n").bold = True
    lines = [
        f"检测文件数：{ov.get('file_count', '-')}",
        f"风险项总数：{ov.get('risk_total', '-')}",
        f"高风险项数：{ov.get('high_risk_count', '-')}",
        f"平均相似度：{ov.get('avg_similarity', '-')}",
    ]
    if ov.get("summary"):
        lines.append(f"概述：{ov['summary']}")
    doc.add_paragraph("\n".join(lines))

    # 关键信息碰撞
    collisions = findings.get("key_info_collisions", [])
    doc.add_heading(f"一、关键信息碰撞检测（{len(collisions)} 项）", level=1)
    if collisions:
        t = doc.add_table(rows=1, cols=5)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(["风险", "字段", "出现值", "涉及文件", "位置"]):
            hdr[i].text = h
        for c in collisions:
            row = t.add_row().cells
            row[0].text = _lvl(c.get("level"))
            row[1].text = str(c.get("field", ""))
            row[2].text = str(c.get("value", ""))
            row[3].text = _md_escape("、".join(c.get("files", [])))
            loc = "；".join(c.get("locations", []))
            if c.get("reason"):
                loc += f" ｜ {c['reason']}"
            row[4].text = _md_escape(loc)
    else:
        doc.add_paragraph("未检测到关键信息雷同。")

    # 文本相似度
    sims = findings.get("text_similarity", [])
    doc.add_heading(f"二、文本语义相似度（{len(sims)} 项）", level=1)
    if sims:
        for s in sims:
            ph = doc.add_paragraph()
            type_tag = s.get("type")
            ph.add_run(
                f"{_lvl(s.get('level'))} {s.get('file_a','')} ↔ {s.get('file_b','')} "
                f"相似度 {s.get('score','-')}"
                + (f" 〔{_md_escape(type_tag)}〕" if type_tag else "")
            ).bold = True
            if s.get("reason"):
                doc.add_paragraph(f"定级依据：{_md_escape(s['reason'])}", style="List Bullet")
            for seg in s.get("segments", []):
                doc.add_paragraph(
                    f"[{_md_escape(seg.get('location',''))}] {_md_escape(seg.get('text',''))}",
                    style="List Bullet",
                )
    else:
        doc.add_paragraph("未检测到明显文本相似。")

    # 文档属性预警
    attrs = findings.get("attribute_warnings", [])
    doc.add_heading(f"三、文档属性预警（{len(attrs)} 项）", level=1)
    if attrs:
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(["风险", "文件", "异常字段 / 说明"]):
            hdr[i].text = h
        for a in attrs:
            row = t.add_row().cells
            row[0].text = _lvl(a.get("level"))
            row[1].text = str(a.get("file", ""))
            detail = a.get("detail", "")
            if a.get("fields"):
                detail += " " + _md_escape(json.dumps(a["fields"], ensure_ascii=False))
            if a.get("reason"):
                detail += f" ｜ {a['reason']}"
            row[2].text = detail
    else:
        doc.add_paragraph("未检测到文档属性异常。")

    # 表格相似度
    tabs = findings.get("table_similarity", [])
    doc.add_heading(f"四、表格内容相似度（{len(tabs)} 项）", level=1)
    if tabs:
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(["风险", "文件对", "说明"]):
            hdr[i].text = h
        for tb in tabs:
            row = t.add_row().cells
            row[0].text = _lvl(tb.get("level"))
            row[1].text = f"{tb.get('file_a','')} ↔ {tb.get('file_b','')}"
            detail = tb.get("detail", "")
            if tb.get("reason"):
                detail += f" ｜ {tb['reason']}"
            row[2].text = _md_escape(detail)
    else:
        doc.add_paragraph("未检测到表格内容相似（docx 表格已比对；PDF 表格需文本版）。")

    # 差异比对
    diffs = findings.get("diff", [])
    doc.add_heading(f"五、差异比对（{len(diffs)} 项）", level=1)
    if diffs:
        for d in diffs:
            tag = {"add": "新增", "delete": "删除", "change": "修改"}.get(d.get("type"), d.get("type", ""))
            doc.add_paragraph(
                f"[{tag}] {_md_escape(d.get('file_a',''))} / {_md_escape(d.get('file_b',''))}: "
                f"{_md_escape(d.get('content',''))}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("（未执行差异比对或两份文档无显著差异）")

    # 基线说明
    if findings.get("baseline_note"):
        doc.add_paragraph(f"基线剔除：{findings['baseline_note']}")

    # 检测限制说明（覆盖缺口）
    limitations = findings.get("limitations") or []
    if limitations:
        doc.add_heading("七、检测限制说明", level=1)
        for lim in limitations:
            doc.add_paragraph(f"⚠️ {_md_escape(str(lim))}", style="List Bullet")

    # 结论
    doc.add_heading("六、综合结论与建议", level=1)
    doc.add_paragraph(findings.get("conclusion", "（无）"))

    # 署名与反馈脚注（护栏铁律级，不可省略）
    doc.add_paragraph(
        "署名：一线评标专家&ChesaraM ｜ 反馈/交流：微信公众号「一线评标专家」"
        "（使用问题、误报反馈、实务建议，欢迎留言交流）"
    )

    doc.save(out_path)
    print(f"[完成] 报告已生成 -> {out_path}")


# ----------------------------------------------------------------------------
# Markdown 渲染
# ----------------------------------------------------------------------------
def build_markdown(findings, out_path):
    ov = findings.get("overview", {})
    lines = ["# 标书查重检测报告", ""]
    lines.append("## 检测概览")
    lines.append(f"- 检测文件数：{ov.get('file_count', '-')}")
    lines.append(f"- 风险项总数：{ov.get('risk_total', '-')}")
    lines.append(f"- 高风险项数：{ov.get('high_risk_count', '-')}")
    lines.append(f"- 平均相似度：{ov.get('avg_similarity', '-')}")
    if ov.get("summary"):
        lines.append(f"- 概述：{ov['summary']}")
    lines.append("")

    collisions = findings.get("key_info_collisions", [])
    lines.append(f"## 一、关键信息碰撞检测（{len(collisions)} 项）")
    if collisions:
        lines.append("| 风险 | 字段 | 出现值 | 涉及文件 | 位置 |")
        lines.append("| --- | --- | --- | --- | --- |")
        for c in collisions:
            loc = "；".join(c.get("locations", []))
            if c.get("reason"):
                loc += f" ｜ {c['reason']}"
            lines.append(
                f"| {_lvl(c.get('level'))} | {_md_escape(c.get('field',''))} | "
                f"{_md_escape(c.get('value',''))} | {_md_escape('、'.join(c.get('files',[])))} | "
                f"{_md_escape(loc)} |"
            )
    else:
        lines.append("未检测到关键信息雷同。")
    lines.append("")

    sims = findings.get("text_similarity", [])
    lines.append(f"## 二、文本语义相似度（{len(sims)} 项）")
    if sims:
        for s in sims:
            type_tag = s.get("type")
            lines.append(
                f"### {_lvl(s.get('level'))} {_md_escape(s.get('file_a',''))} ↔ "
                f"{_md_escape(s.get('file_b',''))} 相似度 {s.get('score','-')}"
                + (f" 〔{_md_escape(type_tag)}〕" if type_tag else "")
            )
            if s.get("reason"):
                lines.append(f"- 定级依据：{_md_escape(s['reason'])}")
            for seg in s.get("segments", []):
                lines.append(
                    f"- [{_md_escape(seg.get('location',''))}] {_md_escape(seg.get('text',''))}"
                )
    else:
        lines.append("未检测到明显文本相似。")
    lines.append("")

    attrs = findings.get("attribute_warnings", [])
    lines.append(f"## 三、文档属性预警（{len(attrs)} 项）")
    if attrs:
        lines.append("| 风险 | 文件 | 异常字段 / 说明 |")
        lines.append("| --- | --- | --- |")
        for a in attrs:
            detail = a.get("detail", "")
            if a.get("fields"):
                detail += " " + _md_escape(json.dumps(a["fields"], ensure_ascii=False))
            if a.get("reason"):
                detail += f" ｜ {a['reason']}"
            lines.append(f"| {_lvl(a.get('level'))} | {_md_escape(a.get('file',''))} | {detail} |")
    else:
        lines.append("未检测到文档属性异常。")
    lines.append("")

    tabs = findings.get("table_similarity", [])
    lines.append(f"## 四、表格内容相似度（{len(tabs)} 项）")
    if tabs:
        lines.append("| 风险 | 文件对 | 说明 |")
        lines.append("| --- | --- | --- |")
        for tb in tabs:
            detail = tb.get("detail", "")
            if tb.get("reason"):
                detail += f" ｜ {tb['reason']}"
            lines.append(
                f"| {_lvl(tb.get('level'))} | {_md_escape(tb.get('file_a',''))} ↔ "
                f"{_md_escape(tb.get('file_b',''))} | {_md_escape(detail)} |"
            )
    else:
        lines.append("未检测到表格内容相似。")
    lines.append("")

    diffs = findings.get("diff", [])
    lines.append(f"## 五、差异比对（{len(diffs)} 项）")
    if diffs:
        for d in diffs:
            tag = {"add": "新增", "delete": "删除", "change": "修改"}.get(d.get("type"), d.get("type", ""))
            lines.append(
                f"- [{tag}] {_md_escape(d.get('file_a',''))} / {_md_escape(d.get('file_b',''))}: "
                f"{_md_escape(d.get('content',''))}"
            )
    else:
        lines.append("（未执行差异比对）")
    lines.append("")

    if findings.get("baseline_note"):
        lines.append(f"**基线剔除**：{findings['baseline_note']}")
        lines.append("")

    limitations = findings.get("limitations") or []
    if limitations:
        lines.append("## 七、检测限制说明")
        for lim in limitations:
            lines.append(f"- ⚠️ {_md_escape(str(lim))}")
        lines.append("")

    lines.append("## 六、综合结论与建议")
    lines.append(findings.get("conclusion", "（无）"))

    lines.append("")
    lines.append("---")
    lines.append(
        "**署名**：一线评标专家&ChesaraM ｜ **反馈/交流**：微信公众号「一线评标专家」"
        "（使用问题、误报反馈、实务建议，欢迎留言交流）"
    )

    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"[完成] Markdown 报告 -> {out_path}")


def main():
    parser = argparse.ArgumentParser(description="渲染标书查重检测报告")
    parser.add_argument("--findings", required=True, help="大模型产出的 findings JSON")
    parser.add_argument("--out", required=True, help="输出 .docx 路径")
    parser.add_argument("--markdown", default=None, help="可选：同时输出 Markdown 路径")
    args = parser.parse_args()

    with open(args.findings, "r", encoding="utf-8") as f:
        findings = json.load(f)

    build_docx(findings, args.out)
    if args.markdown:
        build_markdown(findings, args.markdown)


if __name__ == "__main__":
    main()
