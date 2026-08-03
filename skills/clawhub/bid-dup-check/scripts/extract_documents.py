#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
extract_documents.py - 标书查重 Skill 的文档提取脚本

从 .docx / .pdf / .txt 文件中提取结构化文本、表格、元数据与内嵌图片路径，
输出统一的 JSON，供大模型做语义分析与碰撞检测。

依赖（在 WorkBuddy 托管 Python 环境中安装）：
    python-docx  (docx 读写)
    pypdf        (pdf 文本与元数据)

用法：
    python extract_documents.py --files a.docx b.pdf c.txt \
        [--bidding tender.docx] --out extracted.json

输出 JSON 结构（见 build_report.py 配套 schema）：
{
  "documents": [
    {
      "filename": "...",
      "type": "docx|pdf|txt",
      "page_count": int,
      "paragraph_count": int,
      "char_count": int,
      "metadata": { "author":..., "last_modified_by":..., "company":..., "title":..., ... },
      "paragraphs": [ { "idx": int, "text": "..." } ],
      "tables": [ { "idx": int, "rows": [ [ "cell", ... ], ... ] } ],
      "images": [ { "idx": int, "path": "绝对路径", "note": "需 OCR" } ],
      "warnings": [ "..." ]
    }
  ]
}
"""

import argparse
import json
import os
import sys
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime, timezone


# ----------------------------------------------------------------------------
# 依赖检查：缺失时给出明确指引，避免静默失败
# ----------------------------------------------------------------------------
def _require(module_name, pip_name):
    try:
        return __import__(module_name)
    except ImportError:
        sys.stderr.write(
            f"[错误] 缺少依赖 `{pip_name}`。请在托管 Python 环境中执行：\n"
            f"    pip install {pip_name}\n"
        )
        sys.exit(2)


# ----------------------------------------------------------------------------
# 元数据提取辅助
# ----------------------------------------------------------------------------
def _norm_meta(meta_dict):
    """将可能为 datetime / 其他类型的元数据转为可读字符串。"""
    out = {}
    for k, v in (meta_dict or {}).items():
        if v is None:
            continue
        if isinstance(v, (datetime,)):
            try:
                v = v.astimezone(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
            except Exception:
                v = str(v)
        out[k] = str(v)
    return out


def _read_docx_app_props(path):
    """从 docx 压缩包读取 docProps/app.xml，获取 company / manager 等扩展属性。"""
    props = {}
    try:
        with zipfile.ZipFile(path) as z:
            if "docProps/app.xml" in z.namelist():
                data = z.read("docProps/app.xml")
                root = ET.fromstring(data)
                ns = "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}"
                for tag in ("Company", "Manager"):
                    el = root.find(f"{ns}{tag}")
                    if el is not None and el.text:
                        props[tag.lower()] = el.text.strip()
    except Exception:
        pass
    return props


# ----------------------------------------------------------------------------
# .docx 提取
# ----------------------------------------------------------------------------
def extract_docx(path, image_dir):
    docx = _require("docx", "python-docx")
    from docx.oxml.ns import qn

    document = docx.Document(path)
    paragraphs = []
    for i, p in enumerate(document.paragraphs):
        txt = p.text.strip()
        if txt:
            paragraphs.append({"idx": i, "text": txt})

    tables = []
    for ti, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(cells)
        if any(any(c for c in r) for r in rows):
            tables.append({"idx": ti, "rows": rows})

    # 元数据
    cp = document.core_properties
    meta = _norm_meta(
        {
            "author": cp.author,
            "last_modified_by": cp.last_modified_by,
            "title": cp.title,
            "subject": cp.subject,
            "category": cp.category,
            "comments": cp.comments,
            "created": cp.created,
            "modified": cp.modified,
            "revision": cp.revision,
        }
    )
    meta.update(_read_docx_app_props(path))

    # 内嵌图片提取（供大模型视觉 OCR）
    images = []
    try:
        with zipfile.ZipFile(path) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
            if media:
                os.makedirs(image_dir, exist_ok=True)
            for idx, name in enumerate(media):
                ext = os.path.splitext(name)[1] or ".png"
                out_name = f"{os.path.splitext(os.path.basename(path))[0]}_img{idx}{ext}"
                out_path = os.path.join(image_dir, out_name)
                with open(out_path, "wb") as f:
                    f.write(z.read(name))
                images.append(
                    {
                        "idx": idx,
                        "path": os.path.abspath(out_path),
                        "note": "图片需 OCR 后纳入文本比对",
                    }
                )
    except Exception as e:
        images.append({"idx": -1, "path": "", "note": f"图片提取失败: {e}"})

    char_count = sum(len(p["text"]) for p in paragraphs)
    return {
        "type": "docx",
        "page_count": None,
        "paragraph_count": len(paragraphs),
        "char_count": char_count,
        "metadata": meta,
        "paragraphs": paragraphs,
        "tables": tables,
        "images": images,
    }


# ----------------------------------------------------------------------------
# .pdf 提取（文本层）
# ----------------------------------------------------------------------------
def extract_pdf(path):
    pypdf = _require("pypdf", "pypdf")
    from pypdf import PdfReader

    reader = PdfReader(path)
    paragraphs = []
    for i, page in enumerate(reader.pages):
        txt = (page.extract_text() or "").strip()
        # 按换行拆分为段落，保留页码标记
        for line in txt.split("\n"):
            line = line.strip()
            if line:
                paragraphs.append({"idx": i, "text": line})

    meta = {}
    try:
        if reader.metadata:
            meta = _norm_meta(
                {
                    "author": reader.metadata.author,
                    "title": reader.metadata.title,
                    "subject": reader.metadata.subject,
                    "creator": reader.metadata.creator,
                    "producer": reader.metadata.producer,
                    "created": reader.metadata.creation_date,
                    "modified": reader.metadata.modification_date,
                }
            )
    except Exception:
        pass

    warnings = []
    if not paragraphs:
        warnings.append(
            "PDF 无文本层（可能为扫描件）。当前环境未启用 OCR，"
            "请安装 poppler + pytesseract 或上传文本版 PDF；"
            "也可将 PDF 页面转为图片后用大模型视觉 OCR。"
        )

    char_count = sum(len(p["text"]) for p in paragraphs)
    return {
        "type": "pdf",
        "page_count": len(reader.pages),
        "paragraph_count": len(paragraphs),
        "char_count": char_count,
        "metadata": meta,
        "paragraphs": paragraphs,
        "tables": [],  # pypdf 不提供稳定表格结构，表格相似度以 docx 为主
        "images": [],
        "warnings": warnings,
    }


# ----------------------------------------------------------------------------
# .txt 提取
# ----------------------------------------------------------------------------
def extract_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    paragraphs = []
    for i, block in enumerate(text.split("\n")):
        block = block.strip()
        if block:
            paragraphs.append({"idx": i, "text": block})
    char_count = sum(len(p["text"]) for p in paragraphs)
    return {
        "type": "txt",
        "page_count": None,
        "paragraph_count": len(paragraphs),
        "char_count": char_count,
        "metadata": {},
        "paragraphs": paragraphs,
        "tables": [],
        "images": [],
        "warnings": [],
    }


# ----------------------------------------------------------------------------
# 主流程
# ----------------------------------------------------------------------------
def extract_one(path, image_root):
    ext = os.path.splitext(path)[1].lower()
    image_dir = os.path.join(image_root, os.path.splitext(os.path.basename(path))[0])
    if ext == ".docx":
        data = extract_docx(path, image_dir)
    elif ext == ".pdf":
        data = extract_pdf(path)
    elif ext == ".txt":
        data = extract_txt(path)
    else:
        return None, f"不支持的文件类型: {ext}"
    data["filename"] = os.path.basename(path)
    data["path"] = os.path.abspath(path)
    return data, None


def main():
    parser = argparse.ArgumentParser(description="标书文档结构化提取")
    parser.add_argument("--files", nargs="+", required=True, help="待检测文档路径")
    parser.add_argument("--bidding", default=None, help="可选：招标文件（用于基线剔除）")
    parser.add_argument("--out", required=True, help="输出 JSON 路径")
    args = parser.parse_args()

    out_dir = os.path.dirname(os.path.abspath(args.out))
    image_root = os.path.join(out_dir, "extracted_images")
    os.makedirs(out_dir, exist_ok=True)

    documents = []
    errors = []
    for f in args.files:
        if not os.path.exists(f):
            errors.append(f"文件不存在: {f}")
            continue
        data, err = extract_one(f, image_root)
        if err:
            errors.append(err)
        elif data:
            documents.append(data)

    bidding_doc = None
    if args.bidding:
        if not os.path.exists(args.bidding):
            errors.append(f"招标文件不存在: {args.bidding}")
        else:
            bidding_doc, err = extract_one(args.bidding, image_root)
            if err:
                errors.append(err)
            elif bidding_doc:
                bidding_doc["filename"] = os.path.basename(args.bidding)
                bidding_doc["path"] = os.path.abspath(args.bidding)

    result = {
        "generated_at": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC"),
        "documents": documents,
        "bidding_document": bidding_doc,
        "errors": errors,
    }

    with open(args.out, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"[完成] 提取 {len(documents)} 份文档 -> {args.out}")
    if errors:
        print("[警告]")
        for e in errors:
            print("  - " + e)


if __name__ == "__main__":
    main()
