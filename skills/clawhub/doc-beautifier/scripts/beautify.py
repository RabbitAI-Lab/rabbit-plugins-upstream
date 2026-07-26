#!/usr/bin/env python3
"""
doc-beautifier: beautify a .docx file with a professional general template.

Usage:
    python3 beautify.py <input.docx> <output.docx> [--template standard|compact|elegant]

The script reads the source document, detects its structure (title, headings, body),
and applies consistent formatting. All original TEXT content is preserved.
"""

import sys
import os
import re
from copy import deepcopy

# Ensure python-docx is available
try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


# ─── Template Definitions ────────────────────────────────────────────────────

TEMPLATES = {
    # Standard professional document (default)
    "standard": {
        "page": {
            "width": Cm(21.0),   # A4
            "height": Cm(29.7),
            "margin_top": Cm(2.54),
            "margin_bottom": Cm(2.54),
            "margin_left": Cm(3.18),
            "margin_right": Cm(3.18),
        },
        "title": {
            "font_name": "微软雅黑",
            "font_name_east_asia": "微软雅黑",
            "size": Pt(22),
            "bold": True,
            "color": RGBColor(0x1A, 0x1A, 0x1A),  # near-black
            "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            "space_before": Pt(0),
            "space_after": Pt(18),
            "line_spacing": 1.5,
        },
        "heading_1": {
            "font_name": "微软雅黑",
            "font_name_east_asia": "微软雅黑",
            "size": Pt(16),
            "bold": True,
            "color": RGBColor(0x33, 0x33, 0x33),
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(18),
            "space_after": Pt(8),
            "line_spacing": 1.5,
        },
        "heading_2": {
            "font_name": "微软雅黑",
            "font_name_east_asia": "微软雅黑",
            "size": Pt(14),
            "bold": True,
            "color": RGBColor(0x44, 0x44, 0x44),
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(12),
            "space_after": Pt(6),
            "line_spacing": 1.5,
        },
        "heading_3": {
            "font_name": "微软雅黑",
            "font_name_east_asia": "微软雅黑",
            "size": Pt(12),
            "bold": True,
            "color": RGBColor(0x55, 0x55, 0x55),
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(8),
            "space_after": Pt(4),
            "line_spacing": 1.5,
        },
        "body": {
            "font_name": "宋体",
            "font_name_east_asia": "宋体",
            "size": Pt(12),  # 小四
            "bold": False,
            "color": RGBColor(0x00, 0x00, 0x00),
            "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "first_line_indent": Cm(0.74),  # ~2 Chinese chars
            "space_before": Pt(0),
            "space_after": Pt(4),
            "line_spacing": 1.5,
        },
        "footer_page_number": True,
    },

    # Compact template (tighter spacing)
    "compact": {
        "page": {
            "width": Cm(21.0),
            "height": Cm(29.7),
            "margin_top": Cm(2.0),
            "margin_bottom": Cm(2.0),
            "margin_left": Cm(2.5),
            "margin_right": Cm(2.5),
        },
        "title": {
            "font_name": "微软雅黑",
            "font_name_east_asia": "微软雅黑",
            "size": Pt(18),
            "bold": True,
            "color": RGBColor(0x1A, 0x1A, 0x1A),
            "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            "space_before": Pt(0),
            "space_after": Pt(12),
            "line_spacing": 1.25,
        },
        "heading_1": {
            "font_name": "微软雅黑",
            "font_name_east_asia": "微软雅黑",
            "size": Pt(14),
            "bold": True,
            "color": RGBColor(0x33, 0x33, 0x33),
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(12),
            "space_after": Pt(4),
            "line_spacing": 1.25,
        },
        "heading_2": {
            "font_name": "微软雅黑",
            "font_name_east_asia": "微软雅黑",
            "size": Pt(12),
            "bold": True,
            "color": RGBColor(0x44, 0x44, 0x44),
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(8),
            "space_after": Pt(3),
            "line_spacing": 1.25,
        },
        "heading_3": {
            "font_name": "微软雅黑",
            "font_name_east_asia": "微软雅黑",
            "size": Pt(11),
            "bold": True,
            "color": RGBColor(0x55, 0x55, 0x55),
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(6),
            "space_after": Pt(2),
            "line_spacing": 1.25,
        },
        "body": {
            "font_name": "宋体",
            "font_name_east_asia": "宋体",
            "size": Pt(11),  # 五号
            "bold": False,
            "color": RGBColor(0x00, 0x00, 0x00),
            "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "first_line_indent": Cm(0.66),
            "space_before": Pt(0),
            "space_after": Pt(2),
            "line_spacing": 1.25,
        },
        "footer_page_number": True,
    },

    # Elegant template (larger margins, refined)
    "elegant": {
        "page": {
            "width": Cm(21.0),
            "height": Cm(29.7),
            "margin_top": Cm(3.0),
            "margin_bottom": Cm(3.0),
            "margin_left": Cm(3.5),
            "margin_right": Cm(3.5),
        },
        "title": {
            "font_name": "微软雅黑",
            "font_name_east_asia": "微软雅黑",
            "size": Pt(26),
            "bold": True,
            "color": RGBColor(0x2C, 0x3E, 0x50),
            "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            "space_before": Pt(0),
            "space_after": Pt(24),
            "line_spacing": 1.5,
        },
        "heading_1": {
            "font_name": "微软雅黑",
            "font_name_east_asia": "微软雅黑",
            "size": Pt(16),
            "bold": True,
            "color": RGBColor(0x2C, 0x3E, 0x50),
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(20),
            "space_after": Pt(8),
            "line_spacing": 1.5,
        },
        "heading_2": {
            "font_name": "微软雅黑",
            "font_name_east_asia": "微软雅黑",
            "size": Pt(14),
            "bold": True,
            "color": RGBColor(0x34, 0x49, 0x5E),
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(14),
            "space_after": Pt(6),
            "line_spacing": 1.5,
        },
        "heading_3": {
            "font_name": "微软雅黑",
            "font_name_east_asia": "微软雅黑",
            "size": Pt(12),
            "bold": True,
            "color": RGBColor(0x3C, 0x3C, 0x3C),
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(10),
            "space_after": Pt(4),
            "line_spacing": 1.5,
        },
        "body": {
            "font_name": "宋体",
            "font_name_east_asia": "宋体",
            "size": Pt(12),
            "bold": False,
            "color": RGBColor(0x33, 0x33, 0x33),
            "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "first_line_indent": Cm(0.74),
            "space_before": Pt(0),
            "space_after": Pt(6),
            "line_spacing": 1.5,
        },
        "footer_page_number": True,
    },
}


# ─── Helper Functions ─────────────────────────────────────────────────────────

def set_run_font(run, style_section, template):
    """Apply font formatting to a run from a style config."""
    cfg = template[style_section]
    run.font.name = cfg["font_name"]
    run.font.size = cfg["size"]
    run.font.bold = cfg["bold"]
    run.font.color.rgb = cfg["color"]

    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cfg["font_name_east_asia"])


def set_paragraph_format(para, style_section, template):
    """Apply paragraph-level formatting from a style config."""
    cfg = template[style_section]
    pf = para.paragraph_format
    pf.alignment = cfg["alignment"]
    pf.space_before = cfg["space_before"]
    pf.space_after = cfg["space_after"]
    pf.line_spacing = cfg["line_spacing"]
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    if "first_line_indent" in cfg and cfg["first_line_indent"]:
        pf.first_line_indent = cfg["first_line_indent"]
    else:
        pf.first_line_indent = None


def is_heading_paragraph(para, index, total, prev_was_heading):
    """
    Heuristic detection: is this paragraph likely a heading?
    Returns 0=body, 1=heading1, 2=heading2, 3=heading3
    """
    text = para.text.strip()
    if not text:
        return 0

    # Already styled as heading in source
    style_name = para.style.name if para.style else ""
    if "Heading" in style_name or "heading" in style_name:
        if "1" in style_name:
            return 1
        elif "2" in style_name:
            return 2
        elif "3" in style_name:
            return 3
        return 1

    # Check if bold with specific formatting
    is_bold = False
    has_large_font = False
    for run in para.runs:
        if run.bold:
            is_bold = True
        if run.font.size and run.font.size >= Pt(14):
            has_large_font = True

    # Heading heuristics:
    para_len = len(text)

    # Very short paragraph at start → likely title
    if index <= 1 and para_len < 40:
        return 1

    # Short bold paragraph not ending with punctuation → likely heading
    ends_with_punct = bool(re.search(r"[。！？，；：、.?!,;:]$", text))
    starts_with_num = bool(re.match(r"^[一二三四五六七八九十\d、\.\-\s]{1,4}", text))

    if is_bold and has_large_font:
        return 1 if para_len < 60 else 2

    if is_bold and para_len < 50 and not ends_with_punct:
        return 1

    # Patterns like "一、" "1." "(一)" "1.1"
    if starts_with_num and para_len < 60:
        if is_bold:
            return 1
        # Flowing headings: if short and not a full sentence
        if para_len < 40 and not ends_with_punct:
            return 2

    # Not a heading
    return 0


def add_page_number_footer(doc):
    """Add centered page number to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add page number field
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)


def beautify_document(input_path, output_path, template_name="standard"):
    """Main beautification function."""
    if template_name not in TEMPLATES:
        print(f"Unknown template '{template_name}'. Available: {', '.join(TEMPLATES.keys())}")
        print(f"Falling back to 'standard'.")
        template_name = "standard"

    tmpl = TEMPLATES[template_name]
    doc = Document(input_path)

    # ── Page Setup ──
    for section in doc.sections:
        section.page_width = tmpl["page"]["width"]
        section.page_height = tmpl["page"]["height"]
        section.top_margin = tmpl["page"]["margin_top"]
        section.bottom_margin = tmpl["page"]["margin_bottom"]
        section.left_margin = tmpl["page"]["margin_left"]
        section.right_margin = tmpl["page"]["margin_right"]

    # ── Add page numbers ──
    if tmpl.get("footer_page_number"):
        try:
            add_page_number_footer(doc)
        except Exception:
            pass  # Non-critical

    # ── Process Paragraphs ──
    paragraphs = list(doc.paragraphs)
    total = len(paragraphs)

    # First pass: detect structure
    # Treat first non-empty paragraph as document title
    title_idx = -1
    for i, p in enumerate(paragraphs):
        if p.text.strip():
            title_idx = i
            break

    # Determine paragraph roles
    prev_was_heading = False
    para_roles = []  # list of (role, level) where role = "title", "heading", "body", "empty"

    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            para_roles.append(("empty", 0))
            prev_was_heading = False
            continue

        if i == title_idx:
            para_roles.append(("title", 0))
            prev_was_heading = True
            continue

        h_level = is_heading_paragraph(p, i, total, prev_was_heading)
        if h_level > 0:
            para_roles.append(("heading", h_level))
            prev_was_heading = True
        else:
            para_roles.append(("body", 0))
            prev_was_heading = False

    # ── Second pass: apply formatting ──
    for i, p in enumerate(paragraphs):
        role, level = para_roles[i]
        text = p.text.strip()

        # Clear existing paragraph formatting (reset runs later)
        if role == "title":
            set_paragraph_format(p, "title", tmpl)
            for run in p.runs:
                set_run_font(run, "title", tmpl)

        elif role == "heading":
            style_key = f"heading_{level}"
            set_paragraph_format(p, style_key, tmpl)
            for run in p.runs:
                set_run_font(run, style_key, tmpl)

        elif role == "body":
            set_paragraph_format(p, "body", tmpl)
            for run in p.runs:
                set_run_font(run, "body", tmpl)

        # Empty paragraphs: remove formatting but keep them
        # (preserves document structure)

    doc.save(output_path)
    print(f"✅ Beautified: {input_path} → {output_path} (template: {template_name})")
    return True


# ─── CLI Entry Point ─────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python3 beautify.py <input.docx> <output.docx> [--template standard|compact|elegant]")
        print("\nTemplates:")
        for name, cfg in TEMPLATES.items():
            print(f"  {name}: {cfg.get('description', '')}")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]
    template_name = "standard"

    if "--template" in sys.argv:
        idx = sys.argv.index("--template")
        if idx + 1 < len(sys.argv):
            template_name = sys.argv[idx + 1]

    if not os.path.exists(input_path):
        print(f"ERROR: Input file not found: {input_path}")
        sys.exit(1)

    beautify_document(input_path, output_path, template_name)
