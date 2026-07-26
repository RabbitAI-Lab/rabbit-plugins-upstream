#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
质量周会策划方案Word文档生成脚本
接收结构化JSON数据，生成格式化的Word文档
"""

import argparse
import json
import sys
import os

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print(json.dumps({"status": "error", "message": "python-docx not installed. Run: pip install python-docx==0.8.11"}, ensure_ascii=False))
    sys.exit(1)


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)


def add_styled_paragraph(doc, text, style_name=None, bold=False, font_size=None, color=None, alignment=None, space_after=None, space_before=None):
    """添加带样式的段落"""
    para = doc.add_paragraph()
    if style_name:
        para.style = style_name
    run = para.add_run(text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        para.alignment = alignment
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        para.paragraph_format.space_before = Pt(space_before)
    return para


def validate_data(data):
    """验证JSON数据结构"""
    errors = []
    required_fields = ['meeting_title', 'meeting_date', 'meeting_time', 'meeting_location', 'organizer', 'attendees', 'meeting_objectives', 'agenda_items']
    for field in required_fields:
        if field not in data:
            errors.append(f"Missing required field: {field}")
        elif isinstance(data[field], str) and not data[field].strip():
            errors.append(f"Field '{field}' cannot be empty")
        elif isinstance(data[field], list) and len(data[field]) == 0:
            errors.append(f"Field '{field}' must have at least 1 element")

    if 'agenda_items' in data and isinstance(data['agenda_items'], list):
        for i, item in enumerate(data['agenda_items']):
            item_required = ['topic', 'presenter', 'time_allocation', 'background', 'discussion_points', 'expected_outcome']
            for field in item_required:
                if field not in item:
                    errors.append(f"agenda_items[{i}] missing field: {field}")
                elif isinstance(item[field], str) and not item[field].strip():
                    errors.append(f"agenda_items[{i}].{field} cannot be empty")
                elif isinstance(item[field], list) and len(item[field]) == 0:
                    errors.append(f"agenda_items[{i}].{field} must have at least 1 element")

    return errors


def generate_meeting_doc(data, output_path):
    """生成Word文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ===== 标题 =====
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data['meeting_title'])
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('策划方案')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    subtitle.paragraph_format.space_after = Pt(20)

    # ===== 会议基本信息(表格) =====
    doc.add_heading('一、会议基本信息', level=1)

    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置列宽
    for row in info_table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(5.5)
        row.cells[2].width = Cm(3)
        row.cells[3].width = Cm(5.5)

    info_data = [
        ('会议日期', data['meeting_date'], '会议时间', data['meeting_time']),
        ('会议地点', data['meeting_location'], '总时长', data.get('total_duration', '90分钟')),
        ('主持人', data['organizer'], '记录人', data.get('record_keeper', '待确认')),
        ('参会人员', '、'.join(data['attendees']), '', '')
    ]

    for row_idx, (label1, val1, label2, val2) in enumerate(info_data):
        row = info_table.rows[row_idx]
        # Label cells
        for cell_idx, text in [(0, label1), (2, label2)]:
            if text:
                cell = row.cells[cell_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(10)
                set_cell_shading(cell, 'E8EDF2')
        # Value cells
        for cell_idx, text in [(1, val1), (3, val2)]:
            if text or cell_idx == 1:
                cell = row.cells[cell_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(text if text else '')
                run.font.size = Pt(10)

    # 合并参会人员值单元格
    last_row = info_table.rows[3]
    last_row.cells[1].merge(last_row.cells[3])

    doc.add_paragraph()  # spacing

    # ===== 会议目标 =====
    doc.add_heading('二、会议目标', level=1)
    for idx, obj in enumerate(data['meeting_objectives'], 1):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1)
        run = para.add_run(f'{idx}. {obj}')
        run.font.size = Pt(10.5)
        para.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # ===== 会议议程 =====
    doc.add_heading('三、会议议程', level=1)

    for idx, item in enumerate(data['agenda_items'], 1):
        # 议题标题
        heading_para = doc.add_heading(level=2)
        run = heading_para.add_run(f'议题{idx}: {item["topic"]}')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        # 议题信息表
        item_table = doc.add_table(rows=2, cols=4)
        item_table.style = 'Table Grid'
        item_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        meta_row1 = [('负责人', item['presenter']), ('时间安排', item['time_allocation'])]
        meta_row2_labels = []

        for row_idx, meta_pairs in enumerate([meta_row1]):
            row = item_table.rows[row_idx]
            for col_idx, (label, value) in enumerate(meta_pairs):
                label_cell = row.cells[col_idx * 2]
                value_cell = row.cells[col_idx * 2 + 1]
                label_cell.text = ''
                p = label_cell.paragraphs[0]
                run = p.add_run(label)
                run.bold = True
                run.font.size = Pt(10)
                set_cell_shading(label_cell, 'E8EDF2')
                value_cell.text = ''
                p = value_cell.paragraphs[0]
                run = p.add_run(value)
                run.font.size = Pt(10)

        # 背景说明
        add_styled_paragraph(doc, '背景说明:', bold=True, font_size=10.5, space_before=8, space_after=2)
        bg_para = doc.add_paragraph()
        bg_para.paragraph_format.left_indent = Cm(0.5)
        run = bg_para.add_run(item['background'])
        run.font.size = Pt(10)
        bg_para.paragraph_format.space_after = Pt(6)

        # 关键数据
        if item.get('key_data') and len(item['key_data']) > 0:
            add_styled_paragraph(doc, '关键数据:', bold=True, font_size=10.5, space_before=4, space_after=2)
            for data_point in item['key_data']:
                data_para = doc.add_paragraph()
                data_para.paragraph_format.left_indent = Cm(1)
                run = data_para.add_run(f'  {data_point}')
                run.font.size = Pt(10)
                data_para.paragraph_format.space_after = Pt(2)

        # 讨论要点
        add_styled_paragraph(doc, '讨论要点:', bold=True, font_size=10.5, space_before=6, space_after=2)
        for dp_idx, dp in enumerate(item['discussion_points'], 1):
            dp_para = doc.add_paragraph()
            dp_para.paragraph_format.left_indent = Cm(1)
            run = dp_para.add_run(f'  {dp_idx}. {dp}')
            run.font.size = Pt(10)
            dp_para.paragraph_format.space_after = Pt(2)

        # 期望产出
        add_styled_paragraph(doc, '期望产出:', bold=True, font_size=10.5, space_before=6, space_after=2)
        outcome_para = doc.add_paragraph()
        outcome_para.paragraph_format.left_indent = Cm(0.5)
        run = outcome_para.add_run(item['expected_outcome'])
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x0B, 0x5D, 0x1E)
        outcome_para.paragraph_format.space_after = Pt(12)

    # ===== 会前准备要求 =====
    if data.get('preparation_requirements') and len(data['preparation_requirements']) > 0:
        doc.add_heading('四、会前准备要求', level=1)
        for idx, req in enumerate(data['preparation_requirements'], 1):
            req_para = doc.add_paragraph()
            req_para.paragraph_format.left_indent = Cm(1)
            run = req_para.add_run(f'{idx}. {req}')
            run.font.size = Pt(10.5)
            req_para.paragraph_format.space_after = Pt(4)
        doc.add_paragraph()

    # ===== 备注 =====
    if data.get('notes') and data['notes'].strip():
        doc.add_heading('五、备注', level=1)
        notes_para = doc.add_paragraph()
        run = notes_para.add_run(data['notes'])
        run.font.size = Pt(10)
        notes_para.paragraph_format.space_after = Pt(6)

    # ===== 页脚信息 =====
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_para.add_run(f'生成日期: {data["meeting_date"]}  |  组织者: {data["organizer"]}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # 保存文档
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description='质量周会策划方案Word文档生成')
    parser.add_argument('--input', required=True, help='JSON数据文件路径')
    parser.add_argument('--output', required=True, help='输出Word文档路径')
    args = parser.parse_args()

    # 读取JSON数据
    try:
        with open(args.input, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        result = {"status": "error", "message": f"Input file not found: {args.input}"}
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(1)
    except json.JSONDecodeError as e:
        result = {"status": "error", "message": f"Invalid JSON: {str(e)}"}
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(1)

    # 验证数据
    errors = validate_data(data)
    if errors:
        result = {"status": "error", "message": "Data validation failed", "errors": errors}
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(1)

    # 生成文档
    try:
        output_path = generate_meeting_doc(data, args.output)
        result = {
            "status": "success",
            "message": "Word document generated successfully",
            "output_path": output_path,
            "file_size": os.path.getsize(output_path)
        }
        print(json.dumps(result, ensure_ascii=False))
    except Exception as e:
        result = {"status": "error", "message": f"Document generation failed: {str(e)}"}
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(1)


if __name__ == "__main__":
    main()
