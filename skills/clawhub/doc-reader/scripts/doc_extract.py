#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""本地文档文本抽取：PDF / DOCX / TXT。

优先用本地库；缺失库时给出安装提示而非崩溃。

用法:
  python doc_extract.py <file> [--out out.txt] [--json] [--max N]
"""
import argparse
import json
import os
import sys


def extract_pdf(path, max_n=None):
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            return None, "⚠️ 需要 pypdf：pip install pypdf"
    r = PdfReader(path)
    paras = []
    for i, page in enumerate(r.pages):
        if max_n and i >= max_n:
            break
        t = page.extract_text() or ""
        if t.strip():
            paras.append({"page": i + 1, "text": t.strip()})
    return paras, None


def extract_docx(path, max_n=None):
    try:
        import docx
    except ImportError:
        return None, "⚠️ 需要 python-docx：pip install python-docx"
    d = docx.Document(path)
    items = []
    n = 0
    for p in d.paragraphs:
        if not p.text.strip():
            continue
        style = p.style.name if p.style else ""
        items.append({"type": "para", "style": style, "text": p.text.strip()})
        n += 1
        if max_n and n >= max_n:
            break
    # 表格（轻量：逐表逐行）
    for ti, tbl in enumerate(d.tables):
        rows = [[c.text.strip() for c in row.cells] for row in tbl.rows]
        items.append({"type": "table", "index": ti, "rows": rows})
    return items, None


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("file")
    ap.add_argument("--out", help="输出文件")
    ap.add_argument("--json", action="store_true", help="输出 JSON 结构")
    ap.add_argument("--max", type=int, default=None, help="最多前 N 页/段")
    args = ap.parse_args()

    if not os.path.exists(args.file):
        print("❌ 文件不存在:", args.file, file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(args.file)[1].lower()
    if ext == ".pdf":
        data, err = extract_pdf(args.file, args.max)
    elif ext in (".docx", ".doc"):
        if ext == ".doc":
            print("⚠️ .doc(旧格式) 需先转 docx（LibreOffice/离线工具）", file=sys.stderr)
            sys.exit(1)
        data, err = extract_docx(args.file, args.max)
    elif ext in (".txt", ".md"):
        with open(args.file, "r", encoding="utf-8", errors="replace") as f:
            data = [{"type": "para", "text": ln} for ln in f.read().splitlines() if ln.strip()]
        err = None
    else:
        print("❌ 不支持的格式:", ext, file=sys.stderr)
        sys.exit(1)

    if err:
        print(err, file=sys.stderr)
        sys.exit(1)

    if args.json:
        out = json.dumps(data, ensure_ascii=False, indent=2)
    else:
        if isinstance(data, list):
            out = "\n\n".join(
                (f"[{d.get('page') or d.get('style') or d.get('type')}] {d.get('text', '')}")
                if "text" in d else json.dumps(d, ensure_ascii=False)
                for d in data
            )
        else:
            out = str(data)

    print(f"✅ 抽取 {len(data)} 项")
    if args.out:
        with open(args.out, "w", encoding="utf-8") as f:
            f.write(out)
        print(f"💾 已写入 {args.out}")
    else:
        print(out[:3000])


if __name__ == "__main__":
    main()
