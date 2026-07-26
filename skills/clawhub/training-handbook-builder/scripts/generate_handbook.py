#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
培训手册Word文档生成器 (Training Handbook Generator)
====================================
生成专业级培训手册Word文档，含封面、自动目录、页码、彩色表格、重点提示框、实践案例标注。
适合打印印制分发给培训人员。

使用方法：
1. 修改下方 CONFIG 和 CONTENT 部分定义手册主题和内容
2. 运行: python generate_handbook.py
3. 打开生成的Word文档，右键目录页 → 更新域 → 更新整个目录

依赖: python-docx
安装: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# CONFIG - 修改此处配置手册基本信息
# ============================================================
CONFIG = {
    "title": "培训手册标题",           # 手册主标题
    "subtitle": "副标题",              # 手册副标题
    "version": "2026年版",             # 版本信息
    "audience": "目标受众描述",         # 面向对象
    "policy_basis": "政策依据描述",     # 编制依据
    "output_filename": "培训手册.docx", # 输出文件名
}

# ============================================================
# CONTENT - 修改此处定义手册各章节内容
# 格式: 每章为 (章标题, [(节标题, 节内容), ...])
# 节内容支持: "paragraph"字符串, ("table", headers, rows, col_widths),
#             ("tip", text), ("case", title, content), ("numbered", items)
# ============================================================
CONTENT = [
    # 示例：第一章
    ("第一章 时代背景", [
        ("一、背景概述", [
            "paragraph: 这里写正文段落内容。",
            "paragraph: 这里写第二段正文。",
        ]),
        ("二、政策框架", [
            ("table",
             ["政策名称", "发布时间", "核心要点"],
             [["示例政策", "2026年", "示例要点"]],
             [5, 3, 8]),
        ]),
    ]),
]

# ============================================================
# 排版引擎 - 以下代码一般不需要修改
# ============================================================

# 颜色常量
COLOR_TITLE = RGBColor(0x1A, 0x3C, 0x6E)      # 深蓝 - 主标题/H1
COLOR_H2 = RGBColor(0x2B, 0x57, 0x9A)          # 蓝色 - H2
COLOR_H3 = RGBColor(0x3A, 0x6B, 0xA5)          # 浅蓝 - H3
COLOR_TIP = RGBColor(0xC0, 0x39, 0x2B)         # 红色 - 重点提示
COLOR_CASE = RGBColor(0x27, 0xAE, 0x60)        # 绿色 - 案例标题
COLOR_TABLE_HEADER_BG = "1A3C6E"               # 表头背景深蓝
COLOR_TABLE_ROW_ODD = "EBF5FB"                 # 奇数行浅蓝
COLOR_TABLE_ROW_EVEN = "FFFFFF"                # 偶数行白色
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_GRAY = RGBColor(0x33, 0x33, 0x33)


def setup_document():
    """初始化文档和页面设置"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    return doc


def setup_styles(doc):
    """设置文档样式"""
    # 正文字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    font.color.rgb = COLOR_GRAY
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    _set_east_asia_font(style, '宋体')

    # 标题1
    h1 = doc.styles['Heading 1']
    h1.font.name = '黑体'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_TITLE
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing = 1.3
    _set_east_asia_font(h1, '黑体')

    # 标题2
    h2 = doc.styles['Heading 2']
    h2.font.name = '黑体'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_H2
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.line_spacing = 1.3
    _set_east_asia_font(h2, '黑体')

    # 标题3
    h3 = doc.styles['Heading 3']
    h3.font.name = '黑体'
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = COLOR_H3
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing = 1.3
    _set_east_asia_font(h3, '黑体')

    # 提示框样式
    tip = doc.styles.add_style('Tip Box', WD_STYLE_TYPE.PARAGRAPH)
    tip.font.name = '楷体'
    tip.font.size = Pt(11)
    tip.font.color.rgb = COLOR_TIP
    tip.font.bold = True
    tip.paragraph_format.space_before = Pt(8)
    tip.paragraph_format.space_after = Pt(8)
    tip.paragraph_format.line_spacing = 1.5
    tip.paragraph_format.left_indent = Cm(0.5)
    _set_east_asia_font(tip, '楷体')

    # 案例样式
    case = doc.styles.add_style('Case Text', WD_STYLE_TYPE.PARAGRAPH)
    case.font.name = '楷体'
    case.font.size = Pt(11)
    case.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    case.paragraph_format.space_before = Pt(4)
    case.paragraph_format.space_after = Pt(4)
    case.paragraph_format.line_spacing = 1.5
    case.paragraph_format.left_indent = Cm(1)
    _set_east_asia_font(case, '楷体')


def _set_east_asia_font(style, font_name):
    """设置样式的东亚字体"""
    rPr = style.element.rPr
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:eastAsia="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def _set_run_font(run, font_name):
    """设置run的中文字体"""
    r = run._element.rPr
    if r is None:
        r = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        run._element.insert(0, r)
    rFonts = r.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:eastAsia="{font_name}"/>')
        r.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def add_cover_page(doc, config):
    """添加封面页"""
    for _ in range(6):
        doc.add_paragraph()

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config["title"])
    run.font.name = '黑体'
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = COLOR_TITLE
    _set_run_font(run, '黑体')

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(config["subtitle"])
    run.font.name = '黑体'
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = COLOR_H2
    _set_run_font(run, '黑体')

    # 分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━')
    run.font.color.rgb = COLOR_H2
    run.font.size = Pt(14)

    # 版本信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config["version"])
    run.font.name = '楷体'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
    _set_run_font(run, '楷体')

    for _ in range(3):
        doc.add_paragraph()

    # 受众
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config["audience"])
    run.font.name = '楷体'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _set_run_font(run, '楷体')

    # 政策依据
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config["policy_basis"])
    run.font.name = '楷体'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _set_run_font(run, '楷体')

    doc.add_page_break()


def add_toc(doc):
    """添加自动目录页（Word域代码）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run('目  录')
    run.font.name = '黑体'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = COLOR_TITLE
    _set_run_font(run, '黑体')

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5

    for xml_str in [
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>',
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>',
        f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>',
    ]:
        run = p.add_run()
        run._element.append(parse_xml(xml_str))

    run = p.add_run('（请在Word中右键点击此处 → 更新域，即可自动生成完整目录）')
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    _set_run_font(run, '宋体')

    run = p.add_run()
    run._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

    doc.add_page_break()


def add_tip(doc, text):
    """添加重点提示"""
    p = doc.add_paragraph(style='Tip Box')
    p.add_run('★ 重点提示：' + text)


def add_case(doc, title, content):
    """添加实践案例"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('【实践案例】' + title)
    run.font.name = '黑体'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_CASE
    _set_run_font(run, '黑体')

    p2 = doc.add_paragraph(style='Case Text')
    p2.add_run(content)


def add_table(doc, headers, rows, col_widths=None):
    """添加彩色表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = '黑体'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_WHITE
        _set_run_font(run, '黑体')
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_TABLE_HEADER_BG}" w:val="clear"/>'))

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.name = '宋体'
            run.font.size = Pt(10)
            _set_run_font(run, '宋体')
            bg = COLOR_TABLE_ROW_ODD if r_idx % 2 == 0 else COLOR_TABLE_ROW_EVEN
            cell._tc.get_or_add_tcPr().append(
                parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>'))

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)


def add_numbered_item(doc, number, title, content):
    """添加编号条目"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(f'{number}. {title}')
    run.font.name = '黑体'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_H2
    _set_run_font(run, '黑体')

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.line_spacing = 1.5
    run2 = p2.add_run(content)
    run2.font.name = '宋体'
    run2.font.size = Pt(12)
    _set_run_font(run2, '宋体')


def add_page_number(doc):
    """添加页脚页码"""
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for xml_str in [
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>',
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>',
        f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>',
    ]:
        run = p.add_run()
        run._element.append(parse_xml(xml_str))
    p.add_run()
    run = p.add_run()
    run._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))


def render_content(doc, content):
    """渲染内容到文档"""
    for chapter_title, sections in content:
        doc.add_heading(chapter_title, level=1)
        for section_title, items in sections:
            doc.add_heading(section_title, level=2)
            for item in items:
                if isinstance(item, str) and item.startswith('paragraph:'):
                    doc.add_paragraph(item[len('paragraph:'):].strip())
                elif isinstance(item, tuple):
                    item_type = item[0]
                    if item_type == 'table':
                        _, headers, rows, col_widths = item
                        add_table(doc, headers, rows, col_widths)
                    elif item_type == 'tip':
                        add_tip(doc, item[1])
                    elif item_type == 'case':
                        add_case(doc, item[1], item[2])
                    elif item_type == 'numbered':
                        for idx, (title, content_text) in enumerate(item[1], 1):
                            add_numbered_item(doc, idx, title, content_text)
                    elif item_type == 'heading3':
                        doc.add_heading(item[1], level=3)
                    elif item_type == 'pagebreak':
                        doc.add_page_break()
        doc.add_page_break()


def generate(config, content, output_path=None):
    """
    生成培训手册Word文档

    Args:
        config: dict, 包含 title, subtitle, version, audience, policy_basis, output_filename
        content: list, 章节内容列表
        output_path: str, 输出路径（可选，默认当前目录）
    """
    doc = setup_document()
    setup_styles(doc)

    add_cover_page(doc, config)
    add_toc(doc)
    render_content(doc, content)
    add_page_number(doc)

    if output_path is None:
        output_path = os.path.join(os.getcwd(), config.get("output_filename", "培训手册.docx"))

    doc.save(output_path)
    print(f'文件已保存至: {output_path}')
    return output_path


if __name__ == '__main__':
    generate(CONFIG, CONTENT)
