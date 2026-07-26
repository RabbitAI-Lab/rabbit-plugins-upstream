"""
Jukan05 式分析报告生成器
输入：标的名称 + 分析结果（由 AI 根据 SKILL.md 指令完成分析后传入）
输出：Word 文档（Jukan 风格）
"""

import csv
import json
import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ===== 配置 =====
TEMPLATE = {
    'title_font': 'Arial',
    'title_size': 18,
    'heading_font': 'Arial',
    'heading_size': 14,
    'body_font': 'Arial',
    'body_size': 11,
    'accent_color': RGBColor(0x1D, 0xA1, 0xF2),  # Jukan 风格蓝色
    'warning_color': RGBColor(0xFF, 0x8C, 0x00),  # 橙色（风险提示）
}

def create_jukan_style_report(target, analysis_data, output_path):
    """
    生成 Jukan 风格的分析报告
    
    Args:
        target: 标的名称（如 "SK Hynix", "NVIDIA"）
        analysis_data: dict，包含以下 key（由 AI 分析后填充）：
            - trend: 趋势识别结果
            - bottleneck: 瓶颈定位结果
            - technical: 技术验证数据
            - source_rating: 信息源评级
            - quadrant: 四象限定位
            - execution: 执行建议
            - jukan_history: Jukan 历史观点（如有）
        output_path: 输出 Word 文档路径
    """
    doc = Document()
    
    # ===== 封面 =====
    title = doc.add_heading(f'{target} — Jukan-style Investment Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph('Based on @jukan05 (GF Securities) investment framework')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph()
    date_p = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d")}')
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(10)
    date_p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    doc.add_page_break()
    
    # ===== 免责声明（必须）=====
    disclaimer = doc.add_heading('Disclaimer (Mandatory)', 2)
    disclaimer.runs[0].font.color.rgb = TEMPLATE['warning_color']
    
    doc.add_paragraph(
        'This analysis is based on the investment framework of @jukan05 (GF Securities semiconductor analyst). '
        'It is for informational purposes only and does not constitute investment advice. '
        'All investment decisions should be made based on your own due diligence (DYODD).'
    )
    doc.add_paragraph(
        'Not investment advice. Do your own due diligence (DYODD).'
    )
    
    doc.add_page_break()
    
    # ===== 一、趋势识别（Jukan 出发点）=====
    h1 = doc.add_heading('1. Trend Identification (Starting Point)', 1)
    h1.runs[0].font.color.rgb = TEMPLATE['accent_color']
    
    if 'trend' in analysis_data:
        trend = analysis_data['trend']
        # 趋势描述
        if 'description' in trend:
            p = doc.add_paragraph('Trend: ')
            p.add_run(trend['description']).bold = True
        
        # 趋势够大吗
        if 'size' in trend:
            doc.add_paragraph(f"Is this trend big enough? {trend['size']}")
        
        # 还在早期吗
        if 'early' in trend:
            doc.add_paragraph(f"Is it still early? {trend['early']}")
        
        # Jukan 会关注吗
        if 'jukan_relevance' in trend:
            doc.add_paragraph(f"Would Jukan care? {trend['jukan_relevance']}")
    
    # ===== 二、瓶颈定位 =====
    h2 = doc.add_heading('2. Bottleneck Positioning (Screening Method)', 1)
    h2.runs[0].font.color.rgb = TEMPLATE['accent_color']
    
    if 'bottleneck' in analysis_data:
        b = analysis_data['bottleneck']
        if 'logic' in b:
            doc.add_paragraph(b['logic'])
        if 'why_bottleneck' in b:
            p = doc.add_paragraph()
            p.add_run('Why is this the most beneficiary segment? ').bold = True
            p.add_run(b['why_bottleneck'])
    
    # ===== 三、技术验证 =====
    h3 = doc.add_heading('3. Technical Verification (Core Metrics)', 1)
    h3.runs[0].font.color.rgb = TEMPLATE['accent_color']
    
    if 'technical' in analysis_data:
        tech = analysis_data['technical']
        # 用表格展示技术指标
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Metric'
        hdr[1].text = 'Value'
        hdr[2].text = 'Assessment'
        
        for metric, data in tech.items():
            row = table.add_row().cells
            row[0].text = metric
            row[1].text = str(data.get('value', 'N/A'))
            row[2].text = data.get('assessment', 'Unknown')
    
    # ===== 四、信息源评级 =====
    h4 = doc.add_heading('4. Source Reliability Rating (Jukan System)', 1)
    h4.runs[0].font.color.rgb = TEMPLATE['accent_color']
    
    if 'source_rating' in analysis_data:
        src = analysis_data['source_rating']
        for grade in ['S', 'A', 'B', 'C', 'D']:
            if grade in src:
                p = doc.add_paragraph()
                p.add_run(f'{grade}-grade: ').bold = True
                p.add_run(src[grade])
    
    # ===== 五、标的定位 =====
    h5 = doc.add_heading('5. Target Positioning (Jukan Quadrant)', 1)
    h5.runs[0].font.color.rgb = TEMPLATE['accent_color']
    
    if 'quadrant' in analysis_data:
        q = analysis_data['quadrant']
        p = doc.add_paragraph()
        p.add_run(f"Quadrant: {q['zone']}").bold = True
        p.add_run(f" — {q['description']}")
        
        if 'recommendation' in q:
            doc.add_paragraph(f"Jukan's likely view: {q['recommendation']}")
    
    # ===== 六、执行建议 =====
    h6 = doc.add_heading('6. Execution & Risk Management (Jukan Style)', 1)
    h6.runs[0].font.color.rgb = TEMPLATE['accent_color']
    
    if 'execution' in analysis_data:
        exe = analysis_data['execution']
        if 'position_sizing' in exe:
            doc.add_paragraph(f"Position sizing: {exe['position_sizing']}")
        if 'entry' in exe:
            doc.add_paragraph(f"Entry strategy: {exe['entry']}")
        if 'exit' in exe:
            doc.add_paragraph(f"Exit signals: {exe['exit']}")
        if 'risk' in exe:
            p = doc.add_paragraph()
            p.add_run('Risk factors: ').bold = True
            p.add_run(exe['risk']).font.color.rgb = TEMPLATE['warning_color']
    
    # ===== Jukan 历史观点（如有）=====
    if 'jukan_history' in analysis_data and analysis_data['jukan_history']:
        h7 = doc.add_heading('Appendix: Jukan Historical Views', 2)
        for entry in analysis_data['jukan_history']:
            p = doc.add_paragraph()
            p.add_run(f"[{entry['date']}] ").bold = True
            p.add_run(entry['content'])
    
    # ===== 结尾免责声明（必须）=====
    doc.add_page_break()
    final_disclaimer = doc.add_paragraph()
    final_disclaimer.add_run('Not investment advice. ').bold = True
    final_disclaimer.add_run('Do your own due diligence (DYODD).')
    
    # 保存
    doc.save(output_path)
    return output_path


def load_jukan_history(target):
    """
    从 references/jukan_views/ 加载 Jukan 历史观点
    """
    script_dir = os.path.dirname(os.path.abspath(__file__))
    skill_dir = os.path.dirname(script_dir)
    views_dir = os.path.join(skill_dir, 'references', 'jukan_views')
    
    # 匹配文件名
    target_lower = target.lower()
    history_files = {
        'sk hynix': 'sk_hynix.md',
        'samsung': 'samsung.md',
        'nvidia': 'nvidia.md',
        'tsmc': 'tsmc.md',
        'intel': 'intel.md',
        'memory': 'memory_hbm.md',
        'hbm': 'memory_hbm.md',
        'china': 'china_semiconductor.md',
        'foundry': 'foundry.md',
    }
    
    for key, filename in history_files.items():
        if key in target_lower:
            filepath = os.path.join(views_dir, filename)
            if os.path.exists(filepath):
                with open(filepath, encoding='utf-8') as f:
                    return f.read()
    return None


if __name__ == '__main__':
    # CLI 测试
    if len(sys.argv) < 3:
        print('Usage: python analyze.py <target> <output.docx>')
        sys.exit(1)
    
    target = sys.argv[1]
    output = sys.argv[2]
    
    # 加载 Jukan 历史观点
    history = load_jukan_history(target)
    
    # 示例分析数据（实际由 AI 根据 SKILL.md 分析后填充）
    sample_data = {
        'trend': {
            'description': 'AI compute demand driving HBM memory explosive growth',
            'size': 'CAGR 45%+, market size $150B by 2030',
            'early': 'Still early — Blackwell/Rubin not yet ramped',
            'jukan_relevance': 'High — Jukan\'s core coverage area',
        },
        'bottleneck': {
            'logic': 'In AI server supply chain, HBM is the bottleneck. SK Hynix has 50%+ share.',
            'why_bottleneck': 'HBM4 lead time 6+ weeks, LTA locked 30%+ capacity',
        },
        'technical': {
            'Yield rate': {'value': '70%+', 'assessment': 'Leading (Samsung ~50%)'},
            'Lead time': {'value': '6+ weeks', 'assessment': 'Shortage signal'},
            'LTA coverage': {'value': '30%+', 'assessment': 'Strong pricing power'},
        },
        'source_rating': {
            'A': 'The Elec report: SK Hynix HBM4 yield 70%+',
            'B': 'JPM upgrade price target (2026-06)',
            'C': 'SK Hynix Q2 2026 earnings call',
        },
        'quadrant': {
            'zone': 'Zone A',
            'description': 'Dual-driven by trend + technology',
            'recommendation': 'Strong buy and hold (Jukan style: 3-12 month horizon)',
        },
        'execution': {
            'position_sizing': 'Concentrated (Jukan style: "focused but not gambling")',
            'entry': 'On dip (if market overreacts to short-term noise)',
            'exit': 'When trend fully priced in (lead time normalizes, LTA renews at lower price)',
            'risk': 'HBM demand destruction (AI efficiency breakthrough), China export ban expands',
        },
        'jukan_history': [
            {'date': '2025-01', 'content': 'Started recommending SK Hynix based on HBM shortage + LTA pricing power'},
            {'date': '2025-05', 'content': 'Reiterated buy after HBM4 yield breakthrough (70%+)'},
        ] if history else [],
    }
    
    create_jukan_style_report(target, sample_data, output)
    print(f'✅ Report generated: {output}')
