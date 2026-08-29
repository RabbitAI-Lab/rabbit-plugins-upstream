#!/usr/bin/env python3
"""
创建默认模板文档（带占位符）
用于生成标准的标书模板文件
"""

import os
from pathlib import Path
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_font(run, font_name='宋体', font_size=12):
    """统一设置字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)


def set_heading_font(run, font_name='黑体', font_size=16):
    """设置标题字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = True


def create_cover_template(output_path):
    """创建封面模板"""
    doc = docx.Document()
    
    # 标题
    title = doc.add_heading("投 标 文 件", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_heading_font(run, '黑体', 24)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 信息表
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    
    rows_data = [
        ("项目名称", "{{project_name}}"),
        ("项目编号", "{{tender_number}}"),
        ("招标人", "{{tenderer}}"),
        ("投标截止日期", "{{deadline}}"),
        ("投标人名称", "{{company_name}}"),
        ("法定代表人", "{{legal_rep}}"),
        ("授权代表", "{{authorized_rep}}"),
        ("日  期", "{{deadline}}"),
    ]
    
    for i, (label, value) in enumerate(rows_data):
        cell0 = table.rows[i].cells[0]
        cell1 = table.rows[i].cells[1]
        cell0.text = label
        cell1.text = value
        for run in cell0.paragraphs[0].runs:
            set_font(run, '黑体', 12)
        for run in cell1.paragraphs[0].runs:
            set_font(run)
    
    doc.add_page_break()
    doc.save(str(output_path))
    print(f"✓ 创建封面模板: {output_path}")


def create_commitment_template(output_path):
    """创建承诺书模板"""
    doc = docx.Document()
    
    title = doc.add_heading("投 标 承 诺 书", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_heading_font(run, '黑体', 18)
    
    doc.add_paragraph()
    
    content = """致：{{tenderer}}

我方({{company_name}})作为参加{{project_name}}（项目编号：{{tender_number}}）投标的投标人，郑重承诺如下：

一、我方已仔细阅读并完全理解招标文件的全部内容，愿意按照招标文件的要求提供货物和服务。

二、我方承诺提供的货物和服务符合国家标准、行业标准及招标文件规定的技术要求。

三、我方承诺在投标有效期内不修改、撤销投标文件。

四、如我方中标，我方承诺在收到中标通知书后，在规定时间内与招标人签订合同，并严格履行合同约定的全部义务。

五、我方承诺投标文件中提供的所有资料都是真实、准确、完整的，不存在任何虚假记载、误导性陈述或重大遗漏。

六、我方承诺遵守相关法律法规，不参与任何形式的围标、串标等违法违规行为。

七、如有违反上述承诺，我方愿意承担相应的法律责任，并接受招标人提出的包括但不限于取消投标资格、中标资格等处理决定。

特此承诺。

投标人：{{company_name}}（盖章）
法定代表人或授权代表：{{authorized_rep}}（签字）
日期：    年    月    日
"""
    
    para = doc.add_paragraph(content)
    for run in para.runs:
        set_font(run)
    
    doc.save(str(output_path))
    print(f"✓ 创建承诺书模板: {output_path}")


def create_qualification_template(output_path):
    """创建资质证明文件模板"""
    doc = docx.Document()
    
    title = doc.add_heading("资 质 能 力 证 明", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_heading_font(run, '黑体', 18)
    
    doc.add_paragraph()
    
    content = """一、基本情况

公司名称：{{company_name}}
统一社会信用代码：{{credit_code}}
注册地址：{{company_address}}
法定代表人：{{legal_rep}}

二、银行信息

开户银行：{{bank_name}}
银行账号：{{bank_account}}
联行号：{{bank_code}}

三、授权代表

授权代表：{{authorized_rep}}
联系电话：{{authorized_rep_phone}}

四、声明

我公司声明以上信息真实有效，如有虚假，愿承担相应法律责任。

投标人：{{company_name}}（盖章）
日期：    年    月    日
"""
    
    para = doc.add_paragraph(content)
    for run in para.runs:
        set_font(run)
    
    doc.save(str(output_path))
    print(f"✓ 创建资质证明模板: {output_path}")


def create_quotation_template(output_path):
    """创建报价单模板"""
    doc = docx.Document()
    
    title = doc.add_heading("报 价 单", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_heading_font(run, '黑体', 18)
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    
    # 表头
    headers = ["序号", "项目", "金额（元）"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, '黑体', 11)
            run.bold = True
    
    # 内容
    rows_data = [
        ("1", "合计（含税）", "{{total_price}}"),
        ("2", "其中：硬件费用", ""),
        ("3", "其中：软件费用", ""),
        ("4", "其中：集成费用", ""),
        ("5", "其中：运维费用", ""),
    ]
    
    for i, row_data in enumerate(rows_data[1:], 1):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
    
    doc.add_paragraph()
    doc.add_paragraph("报价说明：以上报价为含税全包价，包含设备、软件、安装、调试、培训及售后服务等全部费用。")
    
    doc.add_paragraph()
    doc.add_paragraph(f"投标人：{{company_name}}（盖章）")
    doc.add_paragraph(f"法定代表人或授权代表：{{authorized_rep}}（签字）")
    doc.add_paragraph("日期：    年    月    日")
    
    doc.save(str(output_path))
    print(f"✓ 创建报价单模板: {output_path}")


def create_deviation_template(output_path):
    """创建偏离说明模板"""
    doc = docx.Document()
    
    title = doc.add_heading("技 术 / 商 务 偏 离 表", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_heading_font(run, '黑体', 16)
    
    doc.add_paragraph()
    
    content = """项目名称：{{project_name}}
项目编号：{{tender_number}}

我方对招标文件的响应情况如下：

经仔细阅读和研究招标文件，我方承诺：
1. 完全响应招标文件的所有技术要求。
2. 完全响应招标文件的所有商务条款。
3. 无任何偏离。

如有任何偏离，将在下表中详细说明：

| 序号 | 招标文件要求 | 投标文件响应 | 偏离说明 |
|------|-------------|-------------|----------|
| 1 | - | 完全响应 | 无偏离 |

投标人：{{company_name}}（盖章）
法定代表人或授权代表：{{authorized_rep}}（签字）
日期：    年    月    日
"""
    
    para = doc.add_paragraph(content)
    for run in para.runs:
        set_font(run)
    
    doc.save(str(output_path))
    print(f"✓ 创建偏离说明模板: {output_path}")


def create_technical_template(output_path):
    """创建技术方案模板"""
    doc = docx.Document()
    
    title = doc.add_heading("技 术 方 案", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_heading_font(run, '黑体', 18)
    
    doc.add_paragraph()
    
    content = """项目名称：{{project_name}}
项目编号：{{tender_number}}

一、项目概述

[请在此处描述对项目背景的理解，包括项目目标、建设内容、预期效果等]

二、技术方案

[请在此处详细描述技术解决方案，包括：]
- 技术架构设计
- 关键技术方案
- 创新点与优势

三、实施计划

[请在此处描述项目实施计划，包括：]
- 项目进度安排（建议使用甘特图）
- 人员配置方案
- 物资设备安排

四、质量保证

[请在此处描述质量保证措施，包括：]
- 质量管理体系
- 质量控制措施
- 验收标准

五、售后服务

[请在此处描述售后服务承诺，包括：]
- 服务范围
- 响应时间
- 培训计划
"""
    
    para = doc.add_paragraph(content)
    for run in para.runs:
        set_font(run)
    
    doc.save(str(output_path))
    print(f"✓ 创建技术方案模板: {output_path}")


def main():
    """创建所有默认模板"""
    base_dir = Path(__file__).parent.parent / "templates" / "default"
    base_dir.mkdir(parents=True, exist_ok=True)
    
    templates = [
        ("00-封面.docx", create_cover_template),
        ("01-承诺书.docx", create_commitment_template),
        ("02-资质文件.docx", create_qualification_template),
        ("03-报价明细.docx", create_quotation_template),
        ("04-偏离说明.docx", create_deviation_template),
        ("05-技术方案.docx", create_technical_template),
    ]
    
    print("\n=== 创建默认模板 ===\n")
    
    for filename, creator in templates:
        output_path = base_dir / filename
        if not output_path.exists():
            creator(output_path)
        else:
            print(f"~ 已存在，跳过: {filename}")
    
    print("\n✅ 默认模板创建完成!")
    print(f"   位置: {base_dir}")


if __name__ == "__main__":
    main()
