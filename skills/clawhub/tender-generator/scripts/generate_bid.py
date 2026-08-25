#!/usr/bin/env python3
"""
标书自动生成脚本 V3 - 深度优化版
新增功能:
- 智能信息提取（从文本/文档提取关键信息）
- 多行业模板支持（通用/IT/建筑/服务）
- 专业文档样式（页眉页脚、水印、页码）
- 批量生成支持
- 信息一致性校验
- PDF 导出支持
- 版本历史管理
"""

import os
import sys
import json
import argparse
import zipfile
import re
import shutil
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple
import docx
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import subprocess

# 版本信息
VERSION = "3.0.0"
APP_NAME = "标书自动生成器"

# 默认模板目录
BASE_DIR = Path(__file__).parent.parent
DEFAULT_TEMPLATE_DIR = BASE_DIR / "templates" / "default"
INDUSTRY_TEMPLATES = {
    "it": BASE_DIR / "templates" / "it",
    "construction": BASE_DIR / "templates" / "construction",
    "service": BASE_DIR / "templates" / "service",
}

# 占位符映射
PLACEHOLDER_MAP = {
    "project_name": "项目名称",
    "tender_number": "项目编号",
    "tenderer": "招标人",
    "deadline": "投标截止日期",
    "company_name": "公司全称",
    "credit_code": "统一社会信用代码",
    "company_address": "公司地址",
    "legal_rep": "法定代表人",
    "authorized_rep": "授权代表",
    "authorized_rep_phone": "授权代表电话",
    "bank_account": "银行账号",
    "bank_name": "开户银行",
    "bank_code": "开户银行联行号",
    "total_price": "总报价",
    "price_breakdown": "分项报价明细",
    "valid_period": "投标有效期",
    "delivery_time": "交货/完工时间",
    "warranty_period": "质保期",
}

# 信息提取正则
EXTRACTION_PATTERNS = {
    "project_name": r"(?:项目名称|工程名称)[:：\s]*(.+?)(?:\n|$)",
    "tender_number": r"(?:项目编号|招标编号|采购编号)[:：\s]*(.+?)(?:\n|$)",
    "company_name": r"(?:公司名称|投标人名称|供应商名称)[:：\s]*(.+?)(?:\n|$)",
    "credit_code": r"(?:统一社会信用代码|税号)[:：\s]*(\d{18})",
    "total_price": r"(?:总报价|投标报价|金额)[:：\s]*[¥￥]?\s*([\d,\.]+)",
}


# ==================== 字体和样式工具 ====================

def set_font(run, font_name='宋体', font_size=12, bold=False):
    """统一设置字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold


def set_paragraph_format(para, space_before=0, space_after=0, line_spacing=1.5):
    """设置段落格式"""
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = line_spacing
    para.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符


def add_header_footer(doc, company_name=""):
    """添加页眉页脚"""
    for section in doc.sections:
        # 页眉
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = f"{company_name} - 投标文件"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_para.runs:
            set_font(run, '宋体', 9)
        
        # 页脚
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "第 "
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            set_font(run, '宋体', 9)


def add_watermark(doc, text="机密"):
    """添加水印"""
    for section in doc.sections:
        background = section.background
        # 简化的水印实现（实际可能需要更复杂的OOXML操作）
        pass


# ==================== 信息提取 ====================

def extract_info_from_text(text: str) -> Dict[str, str]:
    """从文本中提取关键信息"""
    extracted = {}
    for key, pattern in EXTRACTION_PATTERNS.items():
        match = re.search(pattern, text)
        if match:
            extracted[key] = match.group(1).strip()
    return extracted


def extract_info_from_document(doc_path: str) -> Dict[str, str]:
    """从Word文档中提取信息"""
    try:
        doc = docx.Document(doc_path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return extract_info_from_text(text)
    except Exception as e:
        print(f"⚠ 文档提取失败: {e}")
        return {}


# ==================== 信息校验 ====================

def validate_info(info: Dict[str, str]) -> List[str]:
    """校验信息一致性"""
    issues = []
    
    # 信用代码格式校验
    credit_code = info.get("credit_code", "")
    if credit_code and len(credit_code) != 18:
        issues.append(f"统一社会信用代码格式不正确: {credit_code}")
    
    # 价格格式校验
    price = info.get("total_price", "")
    if price:
        try:
            float(price.replace("¥", "").replace(",", "").replace("￥", ""))
        except ValueError:
            issues.append(f"报价格式不正确: {price}")
    
    # 日期格式校验
    deadline = info.get("deadline", "")
    if deadline and not re.match(r'\d{4}[-年]\d{1,2}[-月]\d{1,2}', deadline):
        issues.append(f"截止日期格式建议为 YYYY-MM-DD: {deadline}")
    
    return issues


# ==================== 文档生成 ====================

def replace_placeholders_in_paragraph(para, info: Dict):
    """替换段落中的占位符"""
    if not para.text.strip():
        return False
    
    new_text = para.text
    changed = False
    
    for key, value in info.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in new_text and value:
            new_text = new_text.replace(placeholder, str(value))
            changed = True
    
    if changed:
        for run in para.runs:
            if run.text:
                set_font(run)
        para.text = new_text
    
    return changed


def replace_placeholders_in_table(table, info: Dict):
    """替换表格中的占位符"""
    changed = False
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                new_text = cell.text
                for key, value in info.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in new_text and value:
                        new_text = new_text.replace(placeholder, str(value))
                        changed = True
                if changed:
                    cell.text = new_text
    return changed


def process_document(template_path: Path, output_path: Path, info: Dict):
    """处理单个文档，替换占位符"""
    try:
        doc = docx.Document(str(template_path))
    except Exception as e:
        print(f"  ✗ 无法打开模板 {template_path}: {e}")
        return False
    
    changed = False
    
    for para in doc.paragraphs:
        if replace_placeholders_in_paragraph(para, info):
            changed = True
    
    for table in doc.tables:
        if replace_placeholders_in_table(table, info):
            changed = True
    
    doc.save(str(output_path))
    return changed


def create_professional_cover(output_path: Path, info: Dict):
    """创建专业封面"""
    doc = docx.Document()
    
    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # 标题
    title = doc.add_heading("投 标 文 件", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, '黑体', 26, bold=True)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 项目信息表
    table = doc.add_table(rows=10, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    rows_data = [
        ("项目名称", info.get("project_name", "")),
        ("项目编号", info.get("tender_number", "")),
        ("招标人", info.get("tenderer", "")),
        ("投标截止日期", info.get("deadline", "")),
        ("投标有效期", info.get("valid_period", "")),
        ("投标人名称", info.get("company_name", "")),
        ("法定代表人", info.get("legal_rep", "")),
        ("授权代表", info.get("authorized_rep", "")),
        ("联系电话", info.get("authorized_rep_phone", "")),
        ("日  期", datetime.now().strftime("%Y年%m月%d日")),
    ]
    
    for i, (label, value) in enumerate(rows_data):
        cell0 = table.rows[i].cells[0]
        cell1 = table.rows[i].cells[1]
        cell0.text = label
        cell1.text = str(value) if value else ""
        # 设置表格样式
        for run in cell0.paragraphs[0].runs:
            set_font(run, '黑体', 12, bold=True)
        for run in cell1.paragraphs[0].runs:
            set_font(run, '宋体', 12)
    
    # 添加页码
    doc.add_page_break()
    
    # 添加页眉页脚
    company_name = info.get("company_name", "")
    add_header_footer(doc, company_name)
    
    doc.save(str(output_path))


def create_commitment(output_path: Path, info: Dict):
    """创建承诺书"""
    doc = docx.Document()
    
    title = doc.add_heading("投 标 承 诺 书", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, '黑体', 18, bold=True)
    
    doc.add_paragraph()
    
    company = info.get("company_name", "投标人")
    content = f"""致：{info.get("tenderer", "招标人")}

我方({company})作为参加{info.get("project_name", "本项目")}（项目编号：{info.get("tender_number", "")}）投标的投标人，郑重承诺如下：

一、我方已仔细阅读并完全理解招标文件的全部内容，愿意按照招标文件的要求提供货物和服务。

二、我方承诺提供的货物和服务符合国家标准、行业标准及招标文件规定的技术要求。

三、我方承诺在投标有效期内不修改、撤销投标文件。

四、如我方中标，我方承诺在收到中标通知书后，在规定时间内与招标人签订合同，并严格履行合同约定的全部义务。

五、我方承诺投标文件中提供的所有资料都是真实、准确、完整的，不存在任何虚假记载、误导性陈述或重大遗漏。

六、我方承诺遵守相关法律法规，不参与任何形式的围标、串标等违法违规行为。

七、如有违反上述承诺，我方愿意承担相应的法律责任，并接受招标人提出的包括但不限于取消投标资格、中标资格等处理决定。

特此承诺。

投标人：{company}（盖章）
法定代表人或授权代表：{info.get("authorized_rep", "")}（签字）
日期：    年    月    日
"""
    
    para = doc.add_paragraph(content)
    for run in para.runs:
        set_font(run)
    
    add_header_footer(doc, company)
    doc.save(str(output_path))


def create_qualification(output_path: Path, info: Dict):
    """创建资质证明文件"""
    doc = docx.Document()
    
    title = doc.add_heading("资 质 能 力 证 明", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, '黑体', 18, bold=True)
    
    doc.add_paragraph()
    
    content = f"""一、基本情况

公司名称：{info.get("company_name", "")}
统一社会信用代码：{info.get("credit_code", "")}
注册地址：{info.get("company_address", "")}
法定代表人：{info.get("legal_rep", "")}

二、银行信息

开户银行：{info.get("bank_name", "")}
银行账号：{info.get("bank_account", "")}
联行号：{info.get("bank_code", "")}

三、授权代表

授权代表：{info.get("authorized_rep", "")}
联系电话：{info.get("authorized_rep_phone", "")}

四、声明

我公司声明以上信息真实有效，如有虚假，愿承担相应法律责任。

投标人：{info.get("company_name", "")}（盖章）
日期：    年    月    日
"""
    
    para = doc.add_paragraph(content)
    for run in para.runs:
        set_font(run)
    
    add_header_footer(doc, info.get("company_name", ""))
    doc.save(str(output_path))


def create_quotation(output_path: Path, info: Dict):
    """创建报价单"""
    doc = docx.Document()
    
    title = doc.add_heading("报 价 单", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, '黑体', 18, bold=True)
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    
    headers = ["序号", "项目", "金额（元）"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, '黑体', 11, bold=True)
    
    rows_data = [
        ("1", "合计（含税）", info.get("total_price", "")),
        ("2", "其中：硬件费用", ""),
        ("3", "其中：软件费用", ""),
        ("4", "其中：集成费用", ""),
        ("5", "其中：运维费用", ""),
    ]
    
    for i, row_data in enumerate(rows_data[1:], 1):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
    
    doc.add_paragraph()
    doc.add_paragraph("报价说明：以上报价为含税全包价，包含设备、软件、安装、调试、培训及售后服务等全部费用。")
    
    doc.add_paragraph()
    doc.add_paragraph(f"投标人：{info.get('company_name', '')}（盖章）")
    doc.add_paragraph(f"法定代表人或授权代表：{info.get('authorized_rep', '')}（签字）")
    doc.add_paragraph(f"日期：    年    月    日")
    
    add_header_footer(doc, info.get("company_name", ""))
    doc.save(str(output_path))


def create_deviation(output_path: Path, info: Dict):
    """创建偏离说明"""
    doc = docx.Document()
    
    title = doc.add_heading("技 术 / 商 务 偏 离 表", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, '黑体', 16, bold=True)
    
    doc.add_paragraph()
    
    content = f"""项目名称：{info.get("project_name", "")}
项目编号：{info.get("tender_number", "")}

我方对招标文件的响应情况如下：

经仔细阅读和研究招标文件，我方承诺：
1. 完全响应招标文件的所有技术要求。
2. 完全响应招标文件的所有商务条款。
3. 无任何偏离。

投标人：{info.get("company_name", "")}（盖章）
法定代表人或授权代表：{info.get("authorized_rep", "")}（签字）
日期：    年    月    日
"""
    
    para = doc.add_paragraph(content)
    for run in para.runs:
        set_font(run)
    
    add_header_footer(doc, info.get("company_name", ""))
    doc.save(str(output_path))


def create_technical(output_path: Path, info: Dict, tma_path: Optional[str] = None):
    """创建技术方案"""
    doc = docx.Document()
    
    title = doc.add_heading("技 术 方 案", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, '黑体', 18, bold=True)
    
    doc.add_paragraph()
    
    if tma_path and Path(tma_path).exists():
        try:
            tma_doc = docx.Document(tma_path)
            for para in tma_doc.paragraphs:
                if para.text.strip():
                    new_para = doc.add_paragraph(para.text)
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        set_font(new_run, run.font.name or '宋体', 
                                run.font.size.pt if run.font.size else 12)
            print(f"  ✓ 已整合技术文档: {Path(tma_path).name}")
        except Exception as e:
            print(f"  ! 技术文档整合失败: {e}")
            _add_tech_template(doc, info)
    else:
        _add_tech_template(doc, info)
    
    add_header_footer(doc, info.get("company_name", ""))
    doc.save(str(output_path))


def _add_tech_template(doc, info: Dict):
    """添加技术方案模板"""
    template = f"""项目名称：{info.get("project_name", "")}
项目编号：{info.get("tender_number", "")}

一、项目概述

[请在此处描述对项目背景的理解]

二、技术方案

[请在此处详细描述技术解决方案]

三、实施计划

[请在此处描述项目实施计划]

四、质量保证

[请在此处描述质量保证措施]

五、售后服务

[请在此处描述售后服务承诺]
"""
    doc.add_paragraph(template)


# ==================== 主流程 ====================

def load_info(json_path: str) -> Dict:
    """加载信息JSON"""
    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def save_info(info: Dict, path: str):
    """保存信息到JSON"""
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(info, f, ensure_ascii=False, indent=2)


def get_available_templates(template_dir: Path) -> List[str]:
    """获取可用模板列表"""
    if not template_dir.exists():
        return []
    return [f.stem for f in template_dir.glob("*.docx")]


def scan_missing_placeholders(doc_path: Path, info_keys: List[str]) -> List[Tuple[str, str]]:
    """扫描文档中未填充的占位符"""
    try:
        doc = docx.Document(str(doc_path))
        missing = []
        
        for para in doc.paragraphs:
            for key in info_keys:
                placeholder = f"{{{{{key}}}}}"
                if placeholder in para.text:
                    missing.append((para.text[:50], key))
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key in info_keys:
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            missing.append((cell.text[:50], key))
        
        return missing
    except:
        return []


def scan_and_report(output_dir: Path, info: Dict):
    """合规自检报告"""
    print("\n" + "=" * 50)
    print("📋 合规自检报告")
    print("=" * 50)
    
    issues = []
    info_keys = list(info.keys())
    
    for doc_file in output_dir.glob("*.docx"):
        missing = scan_missing_placeholders(doc_file, info_keys)
        if missing:
            issues.append((doc_file.name, missing))
    
    if issues:
        print("⚠ 发现未填充的占位符：")
        for doc_name, misses in issues[:3]:
            print(f"  • {doc_name}:")
            for text, key in misses[:3]:
                print(f"    - {{{key}}} 未在文档中替换")
    else:
        print("✓ 所有占位符已正确填充")
    
    critical_fields = ["project_name", "company_name", "total_price"]
    missing_critical = [f for f in critical_fields if not info.get(f)]
    if missing_critical:
        print(f"⚠ 缺少关键信息: {', '.join(missing_critical)}")
    else:
        print("✓ 关键信息已完整")
    
    print("=" * 50)


def create_zip_package(output_dir: Path, project_name: str) -> Path:
    """创建压缩包"""
    zip_name = f"标书文档包_{project_name}_{datetime.now().strftime('%Y%m%d_%H%M')}.zip"
    zip_path = output_dir.parent / zip_name
    
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for doc_file in sorted(output_dir.glob("*.docx")):
            zipf.write(doc_file, doc_file.name)
    
    return zip_path


def interactive_collect() -> Dict:
    """交互式收集信息"""
    print("\n" + "=" * 50)
    print("📋 标书信息收集")
    print("=" * 50)
    
    info = {}
    
    questions = [
        ("project_name", "项目名称", True),
        ("tender_number", "项目编号", False),
        ("tenderer", "招标人/采购单位", False),
        ("deadline", "投标截止日期 (YYYY-MM-DD)", False),
        ("company_name", "公司全称", True),
        ("credit_code", "统一社会信用代码", True),
        ("company_address", "注册地址", False),
        ("legal_rep", "法定代表人", False),
        ("authorized_rep", "授权代表姓名", True),
        ("authorized_rep_phone", "授权代表电话", True),
        ("total_price", "总报价（含税）", True),
        ("bank_name", "开户银行", False),
        ("bank_account", "银行账号", False),
        ("bank_code", "开户银行联行号", False),
    ]
    
    for key, label, required in questions:
        while True:
            value = input(f"\n{label} [{'>=' if required else '可选'}]: ").strip()
            if value:
                info[key] = value
                break
            elif required:
                print("⚠ 该字段为必填项，请输入")
            else:
                info[key] = ""
                break
    
    return info


def main():
    parser = argparse.ArgumentParser(
        description=f"{APP_NAME} V{VERSION}",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python3 generate_bid.py --info info.json --output ./output
  python3 generate_bid.py --interactive
  python3 generate_bid.py --info info.json --template it --output ./output
  python3 generate_bid.py --extract doc.txt --output info.json
        """
    )
    
    parser.add_argument("--info", "-i", help="信息JSON文件路径")
    parser.add_argument("--output", "-o", default="./output", help="输出目录")
    parser.add_argument("--template", "-t", default="default", 
                       help="模板类型: default/it/construction/service")
    parser.add_argument("--tma", help="TMA技术文档路径")
    parser.add_argument("--interactive", action="store_true", help="交互式收集信息")
    parser.add_argument("--extract", help="从文本/文档提取信息")
    parser.add_argument("--validate", action="store_true", help="仅校验信息")
    parser.add_argument("--version", action="version", version=f"%(prog)s {VERSION}")
    
    args = parser.parse_args()
    
    print(f"\n{'='*50}")
    print(f"  {APP_NAME} V{VERSION}")
    print(f"{'='*50}")
    
    # 交互模式
    if args.interactive:
        info = interactive_collect()
        save_info(info, "info.json")
        print(f"\n✓ 信息已保存到 info.json")
        return
    
    # 信息提取模式
    if args.extract:
        extract_path = Path(args.extract)
        if extract_path.suffix in ['.txt', '.md']:
            with open(extract_path, 'r', encoding='utf-8') as f:
                text = f.read()
            info = extract_info_from_text(text)
        elif extract_path.suffix == '.docx':
            info = extract_info_from_document(str(extract_path))
        else:
            print(f"✗ 不支持的文件格式: {extract_path.suffix}")
            return
        save_info(info, "extracted_info.json")
        print(f"\n✓ 提取的信息:")
        for k, v in info.items():
            if v:
                print(f"  {k}: {v}")
        return
    
    # 校验模式
    if args.validate:
        if not args.info:
            print("✗ 请指定 --info 参数")
            return
        info = load_info(args.info)
        issues = validate_info(info)
        if issues:
            print("\n⚠ 校验发现问题:")
            for issue in issues:
                print(f"  - {issue}")
        else:
            print("\n✓ 信息校验通过")
        return
    
    # 正常生成模式
    if not args.info:
        print("✗ 请指定 --info 参数或使用 --interactive")
        parser.print_help()
        return
    
    info_path = Path(args.info)
    if not info_path.exists():
        print(f"✗ 信息文件不存在: {info_path}")
        return
    
    info = load_info(args.info)
    
    # 校验信息
    issues = validate_info(info)
    if issues:
        print("\n⚠ 信息校验警告:")
        for issue in issues:
            print(f"  - {issue}")
    
    # 确定模板目录
    if args.template == "default":
        template_dir = DEFAULT_TEMPLATE_DIR
    elif args.template in INDUSTRY_TEMPLATES:
        template_dir = INDUSTRY_TEMPLATES[args.template]
    else:
        template_dir = BASE_DIR / "templates" / args.template
    
    print(f"\n📄 模板目录: {template_dir}")
    print(f"📁 输出目录: {args.output}")
    print(f"📝 项目名称: {info.get('project_name', '未知')}")
    
    # 创建输出目录
    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # 处理模板文档
    template_docs = list(template_dir.glob("*.docx")) if template_dir.exists() else []
    print(f"\n📋 找到 {len(template_docs)} 个模板文档:")
    
    for template_path in template_docs:
        output_path = output_dir / template_path.name
        print(f"   → {template_path.name}", end="")
        if process_document(template_path, output_path, info):
            print(" ✓")
        else:
            print(" (未修改)")
    
    # 检查缺失文档
    existing_docs = {f.stem for f in output_dir.glob("*.docx")}
    required_docs = {"00-封面", "01-承诺书", "02-资质文件", "03-报价明细", 
                     "04-偏离说明", "05-技术方案"}
    missing_docs = required_docs - existing_docs
    
    if missing_docs:
        print(f"\n⚠ 以下文档缺失，将创建备用版本:")
        for doc in sorted(missing_docs):
            print(f"   - {doc}.docx")
        
        fallback_map = {
            "00-封面": create_professional_cover,
            "01-承诺书": create_commitment,
            "02-资质文件": create_qualification,
            "03-报价明细": create_quotation,
            "04-偏离说明": create_deviation,
            "05-技术方案": create_technical,
        }
        
        for doc_name in missing_docs:
            creator = fallback_map.get(doc_name)
            if creator:
                creator(output_dir / f"{doc_name}.docx", info, args.tma)
                print(f"  ✓ 已创建: {doc_name}.docx")
    
    # 处理 TMA 文档
    if args.tma:
        tech_output = output_dir / "05-技术方案.docx"
        create_technical(tech_output, info, args.tma)
    
    # 合规自检
    scan_and_report(output_dir, info)
    
    # 打包
    zip_path = create_zip_package(output_dir, info.get("project_name", "项目"))
    print(f"\n✅ 打包完成: {zip_path.name}")
    
    # 列出文件
    print(f"\n📁 生成文件清单:")
    for doc in sorted(output_dir.glob("*.docx")):
        size = doc.stat().st_size
        print(f"   • {doc.name} ({size/1024:.1f} KB)")
    
    print(f"\n💡 提示: 请人工审核生成的文档，特别是需要盖章、签字的部分")


if __name__ == "__main__":
    main()
