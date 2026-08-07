# -*- coding: utf-8 -*-
"""
Stage 1: Document Parser
Parses .docx / .pdf / .txt / .md files into structured scenes.json

Usage:
    python parse_doc.py --input article.docx --output scenes.json
    python parse_doc.py --input report.pdf --output scenes.json --lang zh
"""

import argparse
import json
import os
import re
import sys
import hashlib
from pathlib import Path

# Ensure config is importable
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from config import PARSE_CONFIG, CONTENT_TYPE_KEYWORDS, AI_IMAGE_CONFIG


# ============================================================
# Document Format Parsers
# ============================================================

def parse_docx(file_path: str) -> list:
    """Parse .docx file using python-docx, preserving heading levels."""
    from docx import Document

    doc = Document(file_path)
    blocks = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Determine heading level from style name
        style_name = para.style.name.lower() if para.style else ""
        if "heading 1" in style_name or "title" in style_name:
            level = 1
        elif "heading 2" in style_name:
            level = 2
        elif "heading 3" in style_name:
            level = 3
        elif "heading" in style_name:
            # Extract level number from style name
            match = re.search(r'heading\s*(\d+)', style_name)
            level = int(match.group(1)) if match else 4
        else:
            level = 0  # Normal paragraph

        blocks.append({"type": "heading" if level > 0 else "paragraph",
                       "level": level, "text": text})

    # Also extract tables
    for table in doc.tables:
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        if rows_text:
            table_text = "\n".join(rows_text)
            blocks.append({"type": "table", "level": 0, "text": table_text})

    return blocks


def parse_pdf(file_path: str) -> list:
    """Parse .pdf file using pdfplumber, page by page."""
    import pdfplumber

    blocks = []
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            # Extract text
            text = page.extract_text() or ""
            if not text.strip():
                continue

            # Split text into paragraphs
            paragraphs = re.split(r'\n\s*\n', text)
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue

                # Heuristic: detect headings (short lines, often bold or larger font)
                lines = para.split('\n')
                first_line = lines[0].strip()
                is_heading = (
                    len(first_line) < 50 and
                    len(lines) <= 2 and
                    (first_line.isupper() or _looks_like_heading(first_line))
                )

                level = 1 if is_heading else 0
                blocks.append({
                    "type": "heading" if is_heading else "paragraph",
                    "level": level,
                    "text": para,
                    "page": page_num + 1
                })

            # Extract tables
            if PARSE_CONFIG.get("table_to_text", True):
                tables = page.extract_tables()
                for table in tables:
                    rows_text = []
                    for row in table:
                        cells = [c.strip() if c else "" for c in row]
                        rows_text.append(" | ".join(cells))
                    if rows_text:
                        table_text = "\n".join(rows_text)
                        blocks.append({"type": "table", "level": 0,
                                       "text": table_text, "page": page_num + 1})

    return blocks


def parse_txt(file_path: str) -> list:
    """Parse .txt file, splitting by blank lines into paragraphs."""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split by double newlines (blank line separators)
    paragraphs = re.split(r'\n\s*\n', content)
    blocks = []

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # Heuristic heading detection
        lines = para.split('\n')
        first_line = lines[0].strip()
        is_heading = (
            len(first_line) < 50 and
            len(lines) <= 2 and
            _looks_like_heading(first_line)
        )

        level = 1 if is_heading else 0
        blocks.append({
            "type": "heading" if is_heading else "paragraph",
            "level": level,
            "text": para
        })

    return blocks


def parse_md(file_path: str) -> list:
    """Parse .md file, preserving heading hierarchy from markdown syntax."""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    blocks = []
    lines = content.split('\n')
    current_para = []

    for line in lines:
        stripped = line.strip()

        # Detect markdown headings: # Title, ## Subtitle, etc.
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)

        if heading_match:
            # Flush current paragraph
            if current_para:
                text = ' '.join(current_para).strip()
                if text:
                    blocks.append({"type": "paragraph", "level": 0, "text": text})
                current_para = []

            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Remove markdown formatting
            text = _clean_markdown(text)
            blocks.append({"type": "heading", "level": level, "text": text})
        elif stripped == '':
            # Blank line = paragraph separator
            if current_para:
                text = ' '.join(current_para).strip()
                if text:
                    clean = _clean_markdown(text)
                    blocks.append({"type": "paragraph", "level": 0, "text": clean})
                current_para = []
        else:
            # Detect code blocks
            if stripped.startswith('```'):
                code_lines = [stripped]
                # This is a simplified approach; for full code block handling,
                # we'd need state tracking. For now, treat as paragraph.
                current_para.append(stripped)
            else:
                current_para.append(stripped)

    # Flush remaining paragraph
    if current_para:
        text = ' '.join(current_para).strip()
        if text:
            clean = _clean_markdown(text)
            blocks.append({"type": "paragraph", "level": 0, "text": clean})

    return blocks


# ============================================================
# Helper Functions
# ============================================================

def detect_content_type(blocks: list, manual_override: str = None) -> str:
    """Detect content type from parsed blocks using keyword matching.

    Args:
        blocks: List of parsed block dicts (each has 'text' field)
        manual_override: If provided, skip detection and return this value

    Returns:
        Content type string (e.g. "finance", "technology", "default")
    """
    if manual_override and manual_override in CONTENT_TYPE_KEYWORDS:
        return manual_override

    # Concatenate all text for keyword scanning
    all_text = " ".join(b.get("text", "") for b in blocks).lower()
    if not all_text.strip():
        return "default"

    # Score each content type by keyword hit count
    scores = {}
    for content_type, keywords in CONTENT_TYPE_KEYWORDS.items():
        score = 0
        for kw in keywords:
            kw_lower = kw.lower()
            # Count occurrences (case-insensitive)
            count = all_text.count(kw_lower)
            if count > 0:
                # Weight: longer keywords are more discriminative
                weight = max(1, len(kw) // 2)
                score += count * weight
        scores[content_type] = score

    # Pick the highest-scoring type
    best_type = max(scores, key=scores.get)
    best_score = scores[best_type]

    # If no keywords matched at all, return default
    if best_score == 0:
        return "default"

    # If the top score is less than 2x the second-best, it's ambiguous — use default
    sorted_scores = sorted(scores.values(), reverse=True)
    if len(sorted_scores) > 1 and sorted_scores[0] < sorted_scores[1] * 2:
        # Ambiguous: top two are close, check if one is clearly dominant
        if sorted_scores[0] < 3:
            return "default"

    print(f"  Content type detected: {best_type} (score={best_score})")
    return best_type


def _looks_like_heading(text: str) -> bool:
    """Heuristic: detect if a text line is likely a heading."""
    # Common heading patterns
    patterns = [
        r'^第[一二三四五六七八九十\d]+[章节部分]',    # Chinese chapter/section
        r'^[一二三四五六七八九十]+、',              # Chinese numbered list
        r'^\d+[\.、]',                            # Numbered heading
        r'^[A-Z][A-Z\s]{3,30}$',                 # All caps title
        r'^Chapter\s+\d+',                        # English chapter
        r'^Section\s+\d+',                        # English section
    ]
    for pattern in patterns:
        if re.match(pattern, text):
            return True
    return False


def _clean_markdown(text: str) -> str:
    """Remove markdown formatting markers from text."""
    # Remove bold/italic markers
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)   # **bold**
    text = re.sub(r'\*([^*]+)\*', r'\1', text)        # *italic*
    text = re.sub(r'__([^_]+)__', r'\1', text)        # __bold__
    text = re.sub(r'_([^_]+)_', r'\1', text)           # _italic_
    # Remove images: ![alt](url) → alt  (must run BEFORE link removal)
    text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', text)
    # Remove links: [text](url) → text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Remove inline code: `code` → code
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Remove code block markers
    text = re.sub(r'```[\w]*', '', text)
    # Remove list markers
    text = re.sub(r'^[\s]*[-*+]\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^[\s]*\d+\.\s+', '', text, flags=re.MULTILINE)
    # Remove blockquote markers
    text = re.sub(r'^>\s*', '', text, flags=re.MULTILINE)
    # Remove horizontal rules
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    # Clean up extra whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _generate_image_prompt(heading: str, text: str, content_type: str = "default") -> str:
    """Generate a text prompt for AI image generation based on scene content.

    The prompt is adapted based on content_type using AI_IMAGE_CONFIG style presets:
    - artistic_style: visual aesthetic (e.g. infographic, photography, illustration)
    - color_mood: color palette and emotional tone
    - composition: layout and focal point guidance

    Args:
        heading: Scene heading for context
        text: Scene narration text for content context
        content_type: Content type for style-aware prompt generation
    """
    # Take first 200 chars of content for context
    context = text[:200].replace('\n', ' ')

    # Get content-type-specific style from AI_IMAGE_CONFIG
    style_map = AI_IMAGE_CONFIG.get("type_style_map", {})
    style = style_map.get(content_type, style_map.get("default", {}))
    artistic_style = style.get("artistic_style", AI_IMAGE_CONFIG["default_style"])
    color_mood = style.get("color_mood", "balanced color palette")
    composition = style.get("composition", "centered composition")

    # Build the prompt with PURPOSE prefix for GenerateImage tool
    prompt = f"[PURPOSE]: Video scene illustration for '{heading}'. "
    prompt += f"Content context: {context}. "
    prompt += f"Artistic style: {artistic_style}. "
    prompt += f"Color mood: {color_mood}. "
    prompt += f"Composition: {composition}. "
    if AI_IMAGE_CONFIG.get("no_text", True):
        prompt += "No text in image. "
    prompt += "High quality, detailed, suitable for educational video."
    return prompt


# ============================================================
# Scene Building
# ============================================================

def build_scenes(blocks: list, lang: str = "zh", content_type: str = None) -> dict:
    """Convert parsed blocks into structured scenes for video generation.

    Args:
        blocks: Parsed document blocks
        lang: Language for duration estimation ("zh" or "en")
        content_type: Manual content type override (auto-detected if None)
    """

    max_chars = PARSE_CONFIG["max_chars_per_scene"]
    min_chars = PARSE_CONFIG["min_chars_per_scene"]
    chars_per_sec = PARSE_CONFIG["cn_chars_per_sec"] if lang == "zh" else PARSE_CONFIG["en_words_per_sec"]

    # Detect content type FIRST so image prompts can be style-aware
    detected_type = detect_content_type(blocks, content_type)

    scenes = []
    current_heading = ""
    current_level = 0
    current_texts = []
    scene_index = 0

    # Extract title (first H1 or first block)
    title = "Untitled Article"
    for b in blocks:
        if b["type"] == "heading" and b["level"] == 1:
            title = b["text"]
            break
    if title == "Untitled Article" and blocks:
        title = blocks[0]["text"][:100]

    def flush_scene():
        nonlocal scene_index, current_texts
        if not current_texts:
            return

        narration = " ".join(current_texts).strip()
        if len(narration) < min_chars and scenes:
            # Merge into previous scene
            prev = scenes[-1]
            prev["narration"] += " " + narration
            prev["char_count"] = len(prev["narration"])
            prev["estimated_duration"] = round(prev["char_count"] / chars_per_sec, 1)
            current_texts = []
            return

        # Split if too long
        if len(narration) > max_chars:
            sentences = re.split(r'(?<=[。！？.!?])\s*', narration)
            chunk = ""
            for sent in sentences:
                if len(chunk) + len(sent) > max_chars and chunk:
                    scenes.append(_make_scene(scene_index, current_heading,
                                             current_level, chunk, chars_per_sec,
                                             detected_type))
                    scene_index += 1
                    chunk = sent
                else:
                    chunk += sent
            if chunk:
                narration = chunk

        scenes.append(_make_scene(scene_index, current_heading, current_level,
                                  narration, chars_per_sec, detected_type))
        scene_index += 1
        current_texts = []

    for block in blocks:
        if block["type"] == "heading":
            # Flush previous scene
            flush_scene()
            current_heading = block["text"]
            current_level = block["level"]
        elif block["type"] == "table":
            # Convert table to narration text
            table_text = block["text"]
            # Summarize table: mention it's a table and read first few rows
            lines = table_text.split('\n')
            summary = "表格内容如下：" + "。".join(lines[:3])
            if len(lines) > 3:
                summary += f"。共{len(lines)}行数据。"
            current_texts.append(summary)
        else:
            current_texts.append(block["text"])

    # Flush last scene
    flush_scene()

    # Calculate totals
    total_chars = sum(s["char_count"] for s in scenes)
    total_duration = sum(s["estimated_duration"] for s in scenes)

    return {
        "title": title,
        "scenes": scenes,
        "total_chars": total_chars,
        "estimated_duration_sec": round(total_duration, 1),
        "language": lang,
        "content_type": detected_type,
    }


def _make_scene(index: int, heading: str, level: int,
                narration: str, chars_per_sec: float,
                content_type: str = "default") -> dict:
    """Create a single scene dict.

    Args:
        content_type: Content type for style-aware image prompt generation.
    """
    # Generate slide text (key points, max 200 chars)
    slide_text = narration[:200]
    if len(narration) > 200:
        # Find a natural break point
        for i in range(200, 150, -1):
            if narration[i] in '。！？.!?；;':
                slide_text = narration[:i+1]
                break
        else:
            slide_text = narration[:200] + "..."

    # Generate image prompt (content-type-aware)
    image_prompt = _generate_image_prompt(heading, narration, content_type)

    # Calculate text hash for caching
    text_hash = hashlib.sha256(narration.encode('utf-8')).hexdigest()[:16]

    return {
        "index": index,
        "heading": heading,
        "level": level,
        "narration": narration,
        "slide_text": slide_text,
        "image_prompt": image_prompt,
        "char_count": len(narration),
        "estimated_duration": round(len(narration) / chars_per_sec, 1),
        "text_hash": text_hash
    }


# ============================================================
# Main Entry Point
# ============================================================

def detect_format(file_path: str) -> str:
    """Detect file format from extension."""
    ext = Path(file_path).suffix.lower()
    format_map = {
        '.docx': 'docx',
        '.doc': 'docx',       # Treat .doc as .docx (may need conversion)
        '.pdf': 'pdf',
        '.txt': 'txt',
        '.md': 'md',
        '.markdown': 'md',
    }
    if ext not in format_map:
        raise ValueError(f"Unsupported file format: {ext}. "
                         f"Supported: .docx, .pdf, .txt, .md")
    return format_map[ext]


def main():
    parser = argparse.ArgumentParser(description="Parse document into scenes JSON")
    parser.add_argument("--input", "-i", required=True, help="Input file path")
    parser.add_argument("--output", "-o", required=True, help="Output JSON path")
    parser.add_argument("--lang", "-l", default="zh",
                       choices=["zh", "en"], help="Language for duration estimation")
    parser.add_argument("--content-type", "-t", default=None,
                       choices=list(CONTENT_TYPE_KEYWORDS.keys()) + ["default"],
                       help="Override content type detection (auto-detected if omitted)")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Error: Input file not found: {args.input}")
        sys.exit(1)

    try:
        print(f"Parsing {args.input} ...")
        fmt = detect_format(args.input)

        if fmt == "docx":
            blocks = parse_docx(args.input)
        elif fmt == "pdf":
            blocks = parse_pdf(args.input)
        elif fmt == "txt":
            blocks = parse_txt(args.input)
        elif fmt == "md":
            blocks = parse_md(args.input)
        else:
            print(f"Error: Unsupported format: {fmt}")
            sys.exit(1)

        print(f"  Extracted {len(blocks)} blocks")

        result = build_scenes(blocks, args.lang, args.content_type)

        print(f"  Built {len(result['scenes'])} scenes")
        print(f"  Content type: {result['content_type']}")
        print(f"  Total chars: {result['total_chars']}")
        print(f"  Estimated duration: {result['estimated_duration_sec']}s "
              f"({result['estimated_duration_sec']/60:.1f} min)")

        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        print(f"  Saved to {args.output}")
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
