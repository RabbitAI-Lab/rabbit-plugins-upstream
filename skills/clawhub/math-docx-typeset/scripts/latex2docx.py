#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
latex2docx.py —— 学术文稿 → Word 原生可编辑 OMML 公式 docx 生成器

技术链路（强制，不可绕过）：
    文稿文本 → 识别行内公式 $...$ / 块公式 $$...$$ → LaTeX→MathML→OMML → 写入 docx

铁律：
    - 主文档公式一律输出 Word 原生 OMML 对象（可双击编辑）
    - 单个公式转换失败时：原位保留完整 LaTeX 代码 + 标注提示，任务不中断
    - 不修改、不简化、不删减用户给出的数学符号与推导步骤
    - 图片模式只作为"对照版附件"（--image-variant 生成），由使用者自选，
      永远不替代主文档的 OMML 公式

中文排版预处理：
    - GB/T 15834 引号规范：中文内容直引号 " " 自动转弯引号 “ ”（保护代码与 URL）

依赖：pip install latex2mathml mathml2omml python-docx matplotlib
用法：python latex2docx.py input.md output.docx [--font 宋体] [--size 12]
              [--leading 1.5] [--image-variant]
"""

import argparse
import io
import re

import latex2mathml.converter
import mathml2omml
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt

M_NS = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'

# ---------- 公式识别 ----------

# 块公式：$$...$$（可跨行）或 \[...\]
BLOCK_RE = re.compile(r'\$\$(.+?)\$\$|\\\[([\s\S]+?)\\\]', re.DOTALL)
# 行内公式：$...$（单行、内容不含 $）或 \(...\)
INLINE_RE = re.compile(r'\$([^$\n]+?)\$|\\\((.+?)\\\)')
# \tag{...} 编号
TAG_RE = re.compile(r'\\tag\{([^}]*)\}')
# CJK 字符检测（公式内含中文时 matplotlib mathtext 渲染为 dummy 方块）
CJK_RE = re.compile(r'[\u4e00-\u9fff]')


def split_blocks(text):
    """先把文本切成 [(类型, 内容)] 序列；类型：'block' / 'rest'"""
    parts = []
    pos = 0
    for m in BLOCK_RE.finditer(text):
        parts.append(('rest', text[pos:m.start()]))
        latex = m.group(1) if m.group(1) is not None else m.group(2)
        parts.append(('block', latex.strip()))
        pos = m.end()
    parts.append(('rest', text[pos:]))
    return parts


def split_inline(text):
    """把纯文本切成 [(类型, 内容)] 序列；类型：'inline' / 'text'"""
    parts = []
    pos = 0
    for m in INLINE_RE.finditer(text):
        parts.append(('text', text[pos:m.start()]))
        latex = m.group(1) if m.group(1) is not None else m.group(2)
        parts.append(('inline', latex.strip()))
        pos = m.end()
    parts.append(('text', text[pos:]))
    return parts


# ---------- GB/T 15834 引号修正（中文正文预处理） ----------

def fix_punctuation(text):
    """中文内容的直引号转弯引号（GB/T 15834）。

    只转换"引号内含汉字"的情形；英文、代码、URL 保留直引号。
    公式（$...$ / $$...$$）与行内代码（`...`）先换占位符保护。
    """
    # 保护公式与行内代码
    stash = []
    def _stash(m):
        stash.append(m.group(0))
        return f'\x00{len(stash) - 1}\x00'
    text = BLOCK_RE.sub(_stash, text)
    text = INLINE_RE.sub(_stash, text)
    text = re.sub(r'`[^`]*`', _stash, text)
    text = re.sub(r'\]\(https?://[^)]+\)', _stash, text)

    def _fix_dq(m):
        inner = m.group(1)
        return '\u201c' + inner + '\u201d' if re.search(r'[\u4e00-\u9fff]', inner) else m.group(0)

    def _fix_sq(m):
        inner = m.group(1)
        return '\u2018' + inner + '\u2019' if re.search(r'[\u4e00-\u9fff]', inner) else m.group(0)

    text = re.sub(r'"([^"\n]*)"', _fix_dq, text)
    text = re.sub(r"'([^'\n]*)'", _fix_sq, text)

    # 还原
    for i, p in enumerate(stash):
        text = text.replace(f'\x00{i}\x00', p)
    return text


# ---------- LaTeX → OMML ----------

# OMML 文本节点中不允许出现反斜杠：LaTeX 命令若被静默当文本写入
# （如 \unknowncommand 未被识别），说明转换结果已"排版变乱"，必须判失败走降级
M_T_RE = re.compile(r'<m:t>([^<]*)</m:t>')
# align 系多行环境（含 align / align* / aligned）
ALIGN_ENV_RE = re.compile(r'^\\begin\{(align\*?|aligned)\}([\s\S]*)\\end\{\1\}$')


def validate_omml(omml):
    """校验 OMML 质量：文本节点残留 LaTeX 命令（\\ 开头）即判失败"""
    for t in M_T_RE.findall(omml):
        if '\\' in t:
            raise ValueError(f'OMML 文本残留未识别的 LaTeX 命令: {t[:40]}')
    return omml


def latex_to_omml(latex):
    """LaTeX → MathML → OMML（含质量校验），返回 OMML XML 字符串"""
    mathml = latex2mathml.converter.convert(latex)
    omml = mathml2omml.convert(mathml)
    return validate_omml(omml)


def split_align(latex):
    """align 多行环境拆分为多个单行公式（每行去 & 对齐符）。

    latex2mathml 不产生 OMML 方程数组（m:eqArr），多行会被压平丢行；
    拆成逐行居中是可靠降级：推导步骤逐行可读、公式原生可编辑。
    返回 None 表示不是 align 环境。
    """
    m = ALIGN_ENV_RE.match(latex.strip())
    if not m:
        return None
    body = m.group(2)
    lines = [ln.strip().rstrip(',').strip() for ln in re.split(r'\\\\', body)]
    lines = [ln.replace('&', '').strip() for ln in lines if ln.strip()]
    return lines or None


def make_omath_element(latex):
    """把 OMML 字符串解析为可插入 docx 的 lxml 元素"""
    wrapped = f'<m:oMath {M_NS}>{omml_inner(latex)}</m:oMath>'
    return parse_xml(wrapped)


def omml_inner(latex):
    omml = latex_to_omml(latex)
    # mathml2omml 输出 <m:oMath>...</m:oMath>，取内部内容重新包命名空间
    inner = re.sub(r'^<m:oMath[^>]*>|</m:oMath>$', '', omml)
    return inner


def make_block_omath_element(latex, tag_text=None):
    """块公式：m:oMathPara 居中容器包 m:oMath；\\tag 编号追加在公式尾部"""
    inner = omml_inner(latex)
    tag_run = ''
    if tag_text:
        tag_run = ('<m:r><m:rPr><m:sty m:val="p"/></m:rPr>'
                   f'<m:t>({tag_text})</m:t></m:r>')
    wrapped = (f'<m:oMathPara {M_NS}>'
               f'<m:oMathParaPr><m:jc m:val="center"/></m:oMathParaPr>'
               f'<m:oMath>{inner}{tag_run}</m:oMath>'
               f'</m:oMathPara>')
    return parse_xml(wrapped)


# ---------- 图片对照渲染（matplotlib mathtext） ----------

def render_formula_png(latex, fontsize=20, dpi=300):
    """把单条 LaTeX 渲染成 PNG，返回 (png_bytes, 自然宽度pt)。

    返回 None 表示不可渲染（含中文——mathtext 无 CJK 字形会出 dummy 方块；
    或 mathtext 不支持的命令）。
    300 DPI：打印级精度（出版标准要求 ≥300）；插入 docx 时按 px/dpi
    换算自然宽度，宽高比由 python-docx 单参数插入自动保持。
    """
    if CJK_RE.search(latex):
        return None
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        from PIL import Image
        fig = plt.figure(figsize=(10, 1.5))
        fig.text(0.5, 0.5, f'${latex}$', ha='center', va='center', fontsize=fontsize)
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight',
                    facecolor='white')
        plt.close(fig)
        # 按图片实际宽度计算插入宽度（px/dpi→inch→pt），上限 A4 正文宽 451pt
        buf.seek(0)
        w_px = Image.open(buf).size[0]
        width_pt = min(w_px / dpi * 72, 451)
        buf.seek(0)
        return buf.getvalue(), width_pt
    except Exception:
        return None


# ---------- docx 写入 ----------

def set_run_font(run, font='宋体', size=12):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font)


def write_text_paragraph(doc, text, font, size, leading):
    """把一行（可能含行内公式）写入 docx；失败公式降级为 LaTeX 原文 + 标注。

    返回行内公式降级数（供转换报告统计）。
    """
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = leading
    inline_fail = 0
    for kind, content in split_inline(text):
        if kind == 'text' and content:
            run = p.add_run(content)
            set_run_font(run, font, size)
        elif kind == 'inline':
            try:
                p._p.append(make_omath_element(content))
            except Exception:
                # 降级预案：保留原始 LaTeX + 标注，不中断
                run = p.add_run(f' {content} ')
                set_run_font(run, font, size)
                run = p.add_run('【公式转换失败，复制本段LaTeX粘贴进Word公式编辑器即可生成】')
                set_run_font(run, font, size)
                run.font.highlight_color = 7  # 黄色高亮
                inline_fail += 1
    return p, inline_fail


def add_image_variant(input_path, output_path, font, size, leading):
    """生成图片对照版附件：块公式渲染为 PNG（行内公式保留 LaTeX 源码形式）。

    定位：对照组附件，供使用者与主文档（OMML 原生公式版）自行比对选用；
    不替代主文档。含中文或 mathtext 不支持的公式，标注"图片对照不可用"。
    """
    with open(input_path, 'r', encoding='utf-8') as f:
        text = f.read()
    text = fix_punctuation(text)  # 与主文档一致的引号规范

    doc = Document()
    # 对照版说明页
    note = doc.add_paragraph()
    run = note.add_run('【图片对照版说明】本文件为对照附件：块公式以图片呈现，仅供视觉核对，'
                       '不可编辑。正式文稿请以主文档（Word 原生 OMML 公式版）为准，由使用者自选。'
                       '行内公式在本版中保留 LaTeX 源码形式。')
    set_run_font(run, font, size)
    run.font.highlight_color = 7

    stats = {'img': 0, 'unavail': 0}

    for kind, content in split_blocks(text):
        if kind == 'rest':
            for line in content.split('\n'):
                if not line.strip():
                    continue
                m = re.match(r'^(#{1,3})\s+(.*)$', line)
                if m:
                    h = doc.add_heading('', level=len(m.group(1)))
                    run = h.add_run(m.group(2))
                    set_run_font(run, font, size + 2)
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = leading
                    run = p.add_run(line)
                    set_run_font(run, font, size)
        elif kind == 'block':
            tag_text = None
            tagm = TAG_RE.search(content)
            if tagm:
                tag_text = tagm.group(1)
                content = TAG_RE.sub('', content).strip()
            # align 拆行后逐行渲染
            align_lines = split_align(content)
            lines = align_lines if align_lines else [content]
            doc.add_paragraph()
            for i, ln in enumerate(lines):
                last = (i == len(lines) - 1)
                result = render_formula_png(ln)
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = leading
                p.alignment = 1  # 居中
                if result:
                    png, width_pt = result
                    r = p.add_run()
                    r.add_picture(io.BytesIO(png), width=Pt(width_pt))
                    if last and tag_text:
                        run = p.add_run(f'  ({tag_text})')
                        set_run_font(run, font, size)
                    stats['img'] += 1
                else:
                    run = p.add_run(ln)
                    set_run_font(run, font, size)
                    run = p.add_run('\n【图片对照不可用（含中文或暂不支持的命令），请以主文档公式为准】')
                    set_run_font(run, font, size)
                    run.font.highlight_color = 7
                    stats['unavail'] += 1

    doc.save(output_path)
    return stats


def convert(input_path, output_path, font='宋体', size=12, leading=1.5,
            image_variant=False):
    with open(input_path, 'r', encoding='utf-8') as f:
        text = f.read()

    # GB/T 15834 引号修正（保护公式 / 代码 / URL）
    text = fix_punctuation(text)

    doc = Document()
    stats = {'ok': 0, 'fail': 0, 'fail_list': [], 'inline_ok': 0, 'inline_fail': 0}

    for kind, content in split_blocks(text):
        if kind == 'rest':
            for line in content.split('\n'):
                if not line.strip():
                    continue
                # markdown 标题轻量支持：# ## ###
                m = re.match(r'^(#{1,3})\s+(.*)$', line)
                if m:
                    level = len(m.group(1))
                    h = doc.add_heading('', level=level)
                    run = h.add_run(m.group(2))
                    set_run_font(run, font, size + 2)
                else:
                    _, ifail = write_text_paragraph(doc, line, font, size, leading)
                    stats['inline_fail'] += ifail
                    stats['inline_ok'] += len(INLINE_RE.findall(line)) - ifail
        elif kind == 'block':
            # 块公式：剥离 \tag{...}
            tag_text = None
            tagm = TAG_RE.search(content)
            if tagm:
                tag_text = tagm.group(1)
                content = TAG_RE.sub('', content).strip()
            try:
                doc.add_paragraph()  # 块公式前空行
                # align 多行环境：拆行逐个转换，逐行居中（编号挂在最后一行）
                align_lines = split_align(content)
                if align_lines:
                    for i, ln in enumerate(align_lines):
                        p = doc.add_paragraph()
                        p.paragraph_format.line_spacing = leading
                        last = (i == len(align_lines) - 1)
                        p._p.append(make_block_omath_element(ln, tag_text if last else None))
                        stats['ok'] += 1
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = leading
                    p._p.append(make_block_omath_element(content, tag_text))
                    stats['ok'] += 1
            except Exception:
                # 降级预案：原位保留完整 LaTeX + 标注
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = leading
                p.alignment = 1  # 居中
                run = p.add_run(content)
                set_run_font(run, font, size)
                run = p.add_run('\n【公式转换失败，复制本段LaTeX粘贴进Word公式编辑器即可生成】')
                set_run_font(run, font, size)
                run.font.highlight_color = 7
                stats['fail'] += 1
                stats['fail_list'].append(content[:50])

    doc.save(output_path)

    # 图片对照版附件（可选，由使用者自选比对）
    variant_stats = None
    if image_variant:
        base = re.sub(r'\.docx$', '', output_path)
        variant_path = f'{base}-图片对照版.docx'
        variant_stats = add_image_variant(input_path, variant_path, font, size, leading)
        variant_stats['path'] = variant_path

    return stats, variant_stats


def main():
    ap = argparse.ArgumentParser(description='学术文稿 → Word 原生 OMML 公式 docx')
    ap.add_argument('input', help='输入文稿（.md / .txt，含 $...$ 与 $$...$$ 公式）')
    ap.add_argument('output', help='输出 .docx 路径')
    ap.add_argument('--font', default='宋体', help='中文字体（默认宋体）')
    ap.add_argument('--size', type=float, default=12, help='正文字号 pt（默认 12=小四）')
    ap.add_argument('--leading', type=float, default=1.5, help='行距倍数（默认 1.5）')
    ap.add_argument('--image-variant', action='store_true',
                    help='额外生成图片对照版附件（块公式渲染为 PNG，仅供视觉核对，自选比对）')
    args = ap.parse_args()

    stats, variant_stats = convert(args.input, args.output, args.font, args.size,
                                   args.leading, args.image_variant)
    print(f'[完成] {args.output}')
    print(f"  块公式：成功 {stats['ok']} 个，降级 {stats['fail']} 个；"
          f"  行内公式：成功 {stats['inline_ok']} 个，降级 {stats['inline_fail']} 个")
    if stats['fail_list']:
        print('  降级公式（保留 LaTeX 原文 + 黄色高亮标注）：')
        for f in stats['fail_list']:
            print(f'    - {f}...')
    if variant_stats:
        print(f"[对照附件] {variant_stats['path']}")
        print(f"  图片渲染 {variant_stats['img']} 个，"
              f"不可用 {variant_stats['unavail']} 个（保留 LaTeX 原文标注）")
    print('  交付提示：请本地打开 Word/WPS 核对全部数学公式；'
          '确认无误后另存导出 PDF，即可得到矢量高清无模糊的论文 PDF。')


if __name__ == '__main__':
    main()
