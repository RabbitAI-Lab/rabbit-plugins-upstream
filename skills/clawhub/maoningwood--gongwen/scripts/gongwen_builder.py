# -*- coding: utf-8 -*-
"""公文 docx 构建库 —— 机关通行公文版式（非国标 GB/T 9704 红头版式，勿混）。

版式要点（规范核心，勿随意改动）：
- 大标题 方正小标宋简体 二号(22pt) 居中，行距固定 33 磅，长标题手工分行
- 副标题 楷体_GB2312 三号 居中
- 一级标题 黑体 三号；二级标题 楷体_GB2312 三号不加粗；三级标题 仿宋_GB2312 三号加粗
- 正文 仿宋_GB2312 三号(16pt)，首行缩进 32 磅（=2汉字），行距固定值 29 磅，两端对齐
- 数字/西文一律 Times New Roman（font.name=TNR + rFonts eastAsia 设中文字体）
- 页边距 上下 2.54cm、左右 3.17cm（Word 默认值，不是国标红头的 3.7/3.5/2.8/2.6）
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FANGSONG = '仿宋_GB2312'
KAITI = '楷体_GB2312'
HEITI = '黑体'
XIAOBIAOSONG = '方正小标宋简体'

SANHAO = 16  # 三号
ERHAO = 22   # 二号

BODY_LINE = 29   # 正文行距固定值（磅）
TITLE_LINE = 33  # 二号大标题行距（磅）


class GongwenDoc:
    def __init__(self):
        self.doc = Document()
        sec = self.doc.sections[0]
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
        sec.top_margin = sec.bottom_margin = Cm(2.54)
        sec.left_margin = sec.right_margin = Cm(3.17)

    def _run(self, p, text, cn, size, bold=False):
        r = p.add_run(text)
        r.font.name = 'Times New Roman'  # 先设 TNR 以创建 rPr/rFonts，再改 eastAsia
        r.font.size = Pt(size)
        r.font.bold = bold
        r._element.rPr.rFonts.set(qn('w:eastAsia'), cn)
        return r

    def _para(self, text='', cn=FANGSONG, size=SANHAO,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, bold=False,
              line=BODY_LINE, before=0, after=0):
        p = self.doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        if line:
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(line)
        pf.first_line_indent = Pt(32) if indent else Pt(0)
        pf.space_before, pf.space_after = Pt(before), Pt(after)
        if text:
            self._run(p, text, cn, size, bold)
        return p

    def title(self, lines):
        """大标题。lines 是手工分好行的列表——超过一行宽度的标题必须按词意
        单元拆成两三行传入，禁止整段塞一行依赖 Word 自动换行（会产生孤字行）。"""
        for ln in lines:
            self._para(ln, cn=XIAOBIAOSONG, size=ERHAO,
                       align=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
                       line=TITLE_LINE)
        self.spacer(10)

    def subtitle(self, text):
        self._para(text, cn=KAITI, size=SANHAO,
                   align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        self.spacer(10)

    def h1(self, text):
        """一级标题（一、二、…）：黑体 三号。"""
        return self._para(text, cn=HEITI, size=SANHAO)

    def h2(self, text):
        """二级标题（（一）（二）…）：楷体 三号，不加粗。"""
        return self._para(text, cn=KAITI, size=SANHAO)

    def h3(self, text):
        """三级标题（1. 2. …）：仿宋 三号，加粗。"""
        return self._para(text, cn=FANGSONG, size=SANHAO, bold=True)

    def para(self, text):
        """正文：仿宋 三号，首行缩进 32 磅，行距固定 29 磅，两端对齐。"""
        return self._para(text)

    def recipient(self, text):
        """主送机关抬头（"各处室、各有关单位："）：顶格不缩进，以全角冒号结尾。"""
        return self._para(text, indent=False)

    def signoff(self, org, date):
        """落款：单位名 + 成文日期，右对齐。需要右空二字等特殊缩进时再手调。"""
        self.spacer(20)
        self._para(org, align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
        self._para(date, align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)

    def spacer(self, pts=10):
        return self._para('', indent=False, line=pts)

    def table(self, headers, rows, header_size=12, cell_size=12, bold_rule=None):
        """通用表格：Table Grid 边框、整体居中、表头黑体。
        单元格文本里的 \\n 会渲染成换行（python-docx 自动转 <w:br/>）。
        bold_rule(values, idx) 返回 True 可整行加粗（如执行率低于红线的行）。"""
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            self._cell(t.rows[0].cells[i], h, HEITI, header_size)
        for idx, values in enumerate(rows):
            cells = t.add_row().cells
            bold = bold_rule(values, idx) if bold_rule else False
            for i, v in enumerate(values):
                self._cell(cells[i], v, FANGSONG, cell_size, bold)
        return t

    def _cell(self, cell, text, cn, size, bold=False):
        cell._tc.clear_content()  # cell.text='' 会残留无字体空 run，弃用
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(p, text, cn, size, bold)

    def save(self, path):
        self.doc.save(path)
        return path
