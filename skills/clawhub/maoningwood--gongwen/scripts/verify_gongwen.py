#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""公文 docx 数字通道验收：解 docx 核对版式，不靠视觉识读。

用法：python3 verify_gongwen.py <文件.docx>
全部通过输出 PASS 并退出码 0；否则逐条列出 FAIL/WARN，退出码 1。
"""
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn

ALLOWED_EA = {'方正小标宋简体', '黑体', '楷体_GB2312', '仿宋_GB2312'}
HALF_PUNCT = '"\''  # 公文用全角标点，半角引号报警

issues = []

def fail(msg):
    issues.append(('FAIL', msg))

def warn(msg):
    issues.append(('WARN', msg))


def eastAsia(run):
    rPr = run._element.rPr
    if rPr is None or rPr.rFonts is None:
        return None
    return rPr.rFonts.get(qn('w:eastAsia'))


def has_cjk(s):
    return any('\u4e00' <= c <= '\u9fff' for c in s)


def check_run(r, where):
    if not r.text.strip():
        return  # 空 run 无字体属性属正常，跳过
    if r.font.name != 'Times New Roman':
        fail(f'{where}: 西文字体是 {r.font.name!r}，应为 Times New Roman')
    ea = eastAsia(r)
    if ea not in ALLOWED_EA:
        fail(f'{where}: eastAsia 字体是 {ea!r}，应在 {sorted(ALLOWED_EA)} 中')
    if has_cjk(r.text) and ea is None:
        fail(f'{where}: 含中文但未设 eastAsia，会兜底成主题字体')


def check_para(p, idx):
    text = p.text
    runs = p.runs
    if not runs:
        return
    where = f'段{idx}({text[:12]!r})'
    for r in runs:
        check_run(r, where)
    pf = p.paragraph_format
    sizes = {r.font.size.pt for r in runs if r.font.size}
    eas = {eastAsia(r) for r in runs}
    # 大标题：二号 + 方正小标宋 + 居中 + 固定33磅 + 无首行缩进
    if sizes == {22}:
        if eas != {'方正小标宋简体'}:
            fail(f'{where}: 二号字标题字体应为方正小标宋简体，实为 {eas}')
        if pf.line_spacing_rule != WD_LINE_SPACING.EXACTLY or pf.line_spacing is None or abs(pf.line_spacing.pt - 33) > 0.6:
            fail(f'{where}: 大标题行距应固定 33 磅，实为 {pf.line_spacing}')
        if pf.first_line_indent and pf.first_line_indent.pt > 0.5:
            fail(f'{where}: 大标题不应首行缩进')
        align_name = p.alignment.name if p.alignment else ''
        if 'center' not in align_name.lower():
            fail(f'{where}: 大标题应居中')
        return
    if sizes == {16}:
        if pf.line_spacing_rule != WD_LINE_SPACING.EXACTLY or pf.line_spacing is None or abs(pf.line_spacing.pt - 29) > 0.6:
            got = f'{pf.line_spacing.pt:.0f}磅' if pf.line_spacing else '未设置固定值'
            fail(f'{where}: 三号段行距应固定 29 磅，实为 {got}')
        align = (p.alignment.name if p.alignment else '').lower()
        ind = pf.first_line_indent.pt if pf.first_line_indent is not None else 0
        # 顶格且以全角冒号结尾 → 主送机关抬头，允许不缩进
        is_recipient = ind <= 0.5 and text.rstrip().endswith('：')
        if (text.strip() and 'right' not in align and 'center' not in align
                and not is_recipient and abs(ind - 32) > 1.5):
            got = f'{pf.first_line_indent.pt:.0f}磅' if pf.first_line_indent is not None else '未设置(顶格)'
            fail(f'{where}: 首行缩进应 32 磅(=2字)，实为 {got}')
        if (text.strip() and 'right' not in align and 'center' not in align
                and not is_recipient and align != 'justify'):
            got = align if align else '未设置(默认左对齐)'
            fail(f'{where}: 正文应两端对齐(justify)，实为 {got}')
    for r in runs:
        if any(c in r.text for c in HALF_PUNCT):
            warn(f'{where}: 含半角引号，公文应为全角（\u201c\u201d）')


def main(path):
    d = Document(path)
    sec = d.sections[0]
    for name, got, want in [('上边距', sec.top_margin, 2.54), ('下边距', sec.bottom_margin, 2.54),
                            ('左边距', sec.left_margin, 3.17), ('右边距', sec.right_margin, 3.17)]:
        if abs(got.cm - want) > 0.05:
            fail(f'{name} {got.cm:.2f}cm ≠ {want}cm')
    for i, p in enumerate(d.paragraphs):
        check_para(p, i)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        check_run(r, '表格单元格')

    if not any(lv == 'FAIL' for lv, _ in issues):
        print('PASS：边距/字体/行距/缩进/对齐 全部符合公文规范')
        for lv, m in issues:
            print(f'{lv}: {m}')
        return 0
    for lv, m in issues:
        print(f'{lv}: {m}')
    print(f'\n共 {sum(1 for lv, _ in issues if lv == "FAIL")} 项 FAIL')
    return 1


if __name__ == '__main__':
    sys.exit(main(sys.argv[1]))
