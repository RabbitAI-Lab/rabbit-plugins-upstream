#!/usr/bin/env python3
"""
根据图片映射 JSON 生成图示说明 Word 文档。

将每道题的原图插入 Word，标注"第X题图示"，方便学生对照查看。
支持两种图片来源：
1. 独立图片文件（传统方式）
2. 从整页渲染图中裁剪子区域（扫描件 PDF 场景）

用法:
  # 方式1：使用独立图片文件
  python generate_image_doc.py \
    --mapping .temp/image_mapping.json \
    --images-dir .temp/images \
    --output 图示说明.docx \
    --title "圆的认识 图示说明"

  # 方式2：从整页渲染图中裁剪
  python generate_image_doc.py \
    --mapping .temp/image_mapping.json \
    --images-dir .temp/images \
    --pages-dir .temp/henan_exam_images_v2 \
    --output 图示说明.docx \
    --title "模拟试题八 图示说明"

输入映射 JSON 格式（由 extract_images.py 生成）:
  [
    {
      "image_file": "image_1.png",
      "question_number": 3,
      "page": 1,
      "context": "第3题题干预览...",
      "crop": {                    // 可选：裁剪坐标（百分比）
        "x": 15,                   // 左上角 X 百分比
        "y": 61,                   // 左上角 Y 百分比
        "width": 37,               // 宽度百分比
        "height": 16               // 高度百分比
      },
      "source_page": "page_2.png"  // 可选：来源整页图文件名（在 pages_dir 中）
    }
  ]
"""

import argparse
import json
import os
import sys
import tempfile


def crop_from_page(pages_dir: str, source_page: str, crop: dict, output_dir: str) -> str:
    """
    从整页渲染图中裁剪子区域，保存为独立图片文件。
    
    Args:
        pages_dir: 整页渲染图目录
        source_page: 来源页面文件名（如 page_2.png）
        crop: 裁剪坐标字典 {"x": 15, "y": 61, "width": 37, "height": 16}（百分比）
        output_dir: 裁剪后图片输出目录
    
    Returns:
        裁剪后图片的完整路径
    """
    from PIL import Image

    page_path = os.path.join(pages_dir, source_page)
    if not os.path.exists(page_path):
        raise FileNotFoundError(f"整页渲染图不存在: {page_path}")

    img = Image.open(page_path)
    page_w, page_h = img.size

    # 百分比转像素，加一点边距（±3%）
    margin = 0.03
    x_pct = max(0, crop["x"] / 100 - margin)
    y_pct = max(0, crop["y"] / 100 - margin)
    w_pct = min(100 - x_pct, crop["width"] / 100 + margin * 2)
    h_pct = min(100 - y_pct, crop["height"] / 100 + margin * 2)

    x1 = int(page_w * x_pct)
    y1 = int(page_h * y_pct)
    x2 = int(page_w * (x_pct + w_pct))
    y2 = int(page_h * (y_pct + h_pct))

    cropped = img.crop((x1, y1, x2, y2))

    # 生成输出文件名
    base_name = os.path.splitext(source_page)[0]
    out_name = f"{base_name}_crop_{crop['x']}_{crop['y']}_{crop['width']}_{crop['height']}.png"
    out_path = os.path.join(output_dir, out_name)
    cropped.save(out_path)

    return out_path


def generate_image_doc(mapping: list, images_dir: str, output_path: str, title: str,
                       pages_dir: str = None):
    """生成图示说明 Word 文档"""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT

    doc = Document()

    # 统一字体设置
    FONT_NAME = '宋体'
    FONT_NAME_ASCII = 'Times New Roman'
    TITLE_SIZE = Pt(18)   # 二号
    LABEL_SIZE = Pt(14)   # 四号
    BODY_SIZE = Pt(12)    # 小四
    NOTE_SIZE = Pt(10.5)  # 五号
    NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    def set_font(run, size=BODY_SIZE, bold=False, italic=False, color=None):
        """统一设置 run 的字体"""
        run.font.name = FONT_NAME_ASCII
        run.font.size = size
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
        run.element.rPr.rFonts.set(f'{NS}eastAsia', FONT_NAME)

    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = FONT_NAME_ASCII
    style.font.size = BODY_SIZE
    style.element.rPr.rFonts.set(f'{NS}eastAsia', FONT_NAME)

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    set_font(title_run, size=TITLE_SIZE, bold=True, color=(0x1A, 0x52, 0x76))

    # 分隔线
    sep_para = doc.add_paragraph()
    sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep_para.add_run("━" * 30)
    set_font(sep_run, size=NOTE_SIZE, color=(0x99, 0x99, 0x99))

    # 说明文字
    note_intro = doc.add_paragraph()
    note_bold = note_intro.add_run("说明：")
    set_font(note_bold, bold=True)
    note_run = note_intro.add_run(
        "本文件为 QuickClass 课堂作业 JSON 的配套图示说明文档。"
        "以下题目在 JSON 中以文字描述代替了原图，学生做题时需对照本文档中的图片查看。"
    )
    set_font(note_run)

    # 裁剪临时目录
    crop_temp_dir = None

    # 按题号排序
    sorted_mapping = sorted(
        [m for m in mapping if m.get("question_number", 0) > 0],
        key=lambda x: x["question_number"]
    )
    # 未确定题号的图片放最后
    unmapped = [m for m in mapping if m.get("question_number", 0) == 0]

    all_entries = sorted_mapping + unmapped

    for i, entry in enumerate(all_entries):
        image_file = entry.get("image_file", "")
        question_number = entry.get("question_number", 0)
        context = entry.get("context", "")
        crop = entry.get("crop")
        source_page = entry.get("source_page")

        # 确定最终图片路径
        image_path = None

        if crop and source_page and pages_dir:
            # 从整页渲染图中裁剪
            if crop_temp_dir is None:
                crop_temp_dir = tempfile.mkdtemp(prefix="img_crop_")
            try:
                image_path = crop_from_page(pages_dir, source_page, crop, crop_temp_dir)
            except Exception as e:
                error_para = doc.add_paragraph()
                error_run = error_para.add_run(f"【第{question_number}题图示】（裁剪失败: {e}）")
                set_font(error_run, color=(0xCC, 0x00, 0x00))
                continue
        elif image_file:
            # 使用独立图片文件
            image_path = os.path.join(images_dir, image_file)

        if not image_path or not os.path.exists(image_path):
            skip_para = doc.add_paragraph()
            skip_run = skip_para.add_run(
                f"【第{question_number}题图示】（图片文件缺失: {image_file or source_page}）"
            )
            set_font(skip_run, color=(0xCC, 0x00, 0x00))
            continue

        # 题号标注
        if question_number > 0:
            label = f"【第{question_number}题图示】"
        else:
            label = "【未确定题号的图示】"

        label_para = doc.add_paragraph()
        label_para.space_before = Pt(16)
        label_para.space_after = Pt(6)
        label_run = label_para.add_run(label)
        set_font(label_run, size=LABEL_SIZE, bold=True, color=(0x1A, 0x52, 0x76))

        # 上下文提示（如有）
        if context and context != "(未映射到具体题目，需手动确认)":
            ctx_para = doc.add_paragraph()
            ctx_para.space_after = Pt(4)
            ctx_run = ctx_para.add_run(f"题干预览：{context}")
            set_font(ctx_run, size=NOTE_SIZE, italic=True, color=(0x66, 0x66, 0x66))

        # 插入图片
        try:
            from PIL import Image as PILImage
            with PILImage.open(image_path) as img:
                img_width, img_height = img.size

            max_width_cm = 14
            aspect = img_height / img_width
            display_width = min(max_width_cm, img_width * 2.54 / 96)

            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_run = img_para.add_run()
            img_run.add_picture(image_path, width=Cm(display_width))

        except ImportError:
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_run = img_para.add_run()
            img_run.add_picture(image_path, width=Cm(12))

        except Exception as e:
            error_para = doc.add_paragraph()
            error_run = error_para.add_run(f"（图片插入失败: {e}）")
            set_font(error_run, size=NOTE_SIZE, color=(0xCC, 0x00, 0x00))

        # 题目之间加分隔
        if i < len(all_entries) - 1:
            dot_para = doc.add_paragraph()
            dot_para.space_before = Pt(4)
            dot_para.space_after = Pt(4)
            dot_run = dot_para.add_run("· · ·")
            set_font(dot_run, size=NOTE_SIZE, color=(0xCC, 0xCC, 0xCC))
            dot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 末尾分隔线
    sep_para2 = doc.add_paragraph()
    sep_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run2 = sep_para2.add_run("━" * 30)
    set_font(sep_run2, size=NOTE_SIZE, color=(0x99, 0x99, 0x99))

    # 说明文字
    note_para = doc.add_paragraph()
    note_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_run = note_para.add_run(
        "说明：本文件为 QuickClass 课堂作业配套图示，请对照题号查看对应图示。"
    )
    set_font(note_run, size=NOTE_SIZE, color=(0x66, 0x66, 0x66))

    doc.save(output_path)

    # 清理临时裁剪目录
    if crop_temp_dir:
        import shutil
        try:
            shutil.rmtree(crop_temp_dir)
        except Exception:
            pass


def main():
    parser = argparse.ArgumentParser(
        description="根据图片映射生成图示说明 Word 文档"
    )
    parser.add_argument("--mapping", required=True, help="映射 JSON 文件路径")
    parser.add_argument("--images-dir", required=True, help="图片目录")
    parser.add_argument("--pages-dir", default=None,
                        help="整页渲染图目录（扫描件 PDF 场景，用于裁剪子区域）")
    parser.add_argument("--output", required=True, help="输出 Word 文件路径")
    parser.add_argument("--title", default="题目图示说明", help="文档标题")

    args = parser.parse_args()

    with open(args.mapping, "r", encoding="utf-8") as f:
        mapping = json.load(f)

    if not mapping:
        print("映射为空，无需生成图示说明文档。")
        sys.exit(0)

    generate_image_doc(mapping, args.images_dir, args.output, args.title,
                       pages_dir=args.pages_dir)
    print(f"图示说明文档已生成: {args.output}")
    print(f"共包含 {len(mapping)} 张图片")


if __name__ == "__main__":
    main()
