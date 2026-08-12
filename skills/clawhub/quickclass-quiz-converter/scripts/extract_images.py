#!/usr/bin/env python3
"""
从 Word (.docx) 或 PDF 文件中提取图片，并映射到题号。

支持两种映射模式:
  1. 位置就近模式（默认）：根据图片在文档中的位置就近匹配题号
  2. 视觉识别模式（--vision-match）：用视觉模型识别图片内容，精准匹配题号

用法:
  # 位置就近模式
  python extract_images.py --input 试卷.docx --output-dir .temp/images --mapping .temp/mapping.json

  # 视觉识别模式（精准匹配，适合扫描件、答案集中放置等场景）
  python extract_images.py --input 试卷.pdf --output-dir .temp/images --mapping .temp/mapping.json --vision-match

  # 指定试卷题号范围（辅助视觉模型判断）
  python extract_images.py --input 试卷.pdf --output-dir .temp/images --mapping .temp/mapping.json --vision-match --question-range "1-20"

输出:
  1. --output-dir: 图片文件（image_1.png, image_2.png, ...）
  2. --mapping: 题号映射 JSON
"""

import argparse
import json
import os
import re
import sys


def extract_images_from_docx(input_path: str, output_dir: str) -> list:
    """从 Word 文件中提取内嵌图片并按位置映射题号"""
    from docx import Document
    from docx.oxml.ns import qn

    os.makedirs(output_dir, exist_ok=True)
    doc = Document(input_path)

    # Step 1: 提取所有图片
    image_map = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_ext = os.path.splitext(rel.target_ref)[1] or ".png"
            image_data = rel.target_part.blob
            idx = len(image_map) + 1
            image_file = f"image_{idx}{image_ext}"
            image_path = os.path.join(output_dir, image_file)
            with open(image_path, "wb") as f:
                f.write(image_data)
            image_map[rel.rId] = {
                "file": image_file,
                "path": image_path,
                "index": idx,
            }

    # Step 2: 遍历段落，按位置建立映射
    paragraphs_info = []
    current_question = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        q_match = re.match(r'[第]?[\s]*(\d+)[题\.、．\s]', text)
        if q_match:
            current_question = int(q_match.group(1))

        images_in_para = []
        for run in para.runs:
            drawings = run._element.findall(qn('w:drawing'))
            for drawing in drawings:
                blips = drawing.findall('.//' + qn('a:blip'))
                for blip in blips:
                    embed_id = blip.get(qn('r:embed'))
                    if embed_id and embed_id in image_map:
                        images_in_para.append(image_map[embed_id])
            picts = run._element.findall(qn('w:pict'))
            for pict in picts:
                for img_data in pict.findall('.//' + qn('v:imagedata')):
                    rid = img_data.get(qn('r:id'))
                    if rid and rid in image_map:
                        images_in_para.append(image_map[rid])

        if images_in_para:
            paragraphs_info.append({
                "text": text[:100],
                "question_number": current_question,
                "images": images_in_para,
            })

    # 也检查表格中的图片
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    q_match = re.match(r'[第]?[\s]*(\d+)[题\.、．\s]', text)
                    if q_match:
                        current_question = int(q_match.group(1))
                    for run in para.runs:
                        drawings = run._element.findall(qn('w:drawing'))
                        for drawing in drawings:
                            blips = drawing.findall('.//' + qn('a:blip'))
                            for blip in blips:
                                embed_id = blip.get(qn('r:embed'))
                                if embed_id and embed_id in image_map:
                                    paragraphs_info.append({
                                        "text": text[:100],
                                        "question_number": current_question,
                                        "images": [image_map[embed_id]],
                                    })

    # 生成映射（去重）
    seen_files = set()
    mapping = []
    for p in paragraphs_info:
        for img in p["images"]:
            if img["file"] not in seen_files:
                seen_files.add(img["file"])
                mapping.append({
                    "image_file": img["file"],
                    "question_number": p["question_number"],
                    "page": 1,
                    "context": p["text"],
                })

    # 未映射的图片
    for rId, img in image_map.items():
        if img["file"] not in seen_files:
            mapping.append({
                "image_file": img["file"],
                "question_number": 0,
                "page": 1,
                "context": "(未映射到具体题目，需手动确认)",
            })

    return mapping


def render_pdf_pages_as_images(input_path: str, output_dir: str, dpi: int = 200) -> list:
    """将 PDF 每页渲染为图片（用于扫描件 PDF）"""
    import pymupdf

    os.makedirs(output_dir, exist_ok=True)
    doc = pymupdf.open(input_path)

    mapping = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        # 将页面渲染为图片
        pix = page.get_pixmap(dpi=dpi)
        image_file = f"page_{page_num + 1}.png"
        image_path = os.path.join(output_dir, image_file)
        pix.save(image_path)

        mapping.append({
            "image_file": image_file,
            "question_number": 0,
            "page": page_num + 1,
            "context": f"PDF第{page_num + 1}页整页渲染",
            "source": "page_render",
        })

    return mapping


def extract_images_from_pdf(input_path: str, output_dir: str) -> list:
    """从 PDF 文件中提取内嵌图片并按位置映射题号"""
    import pymupdf

    os.makedirs(output_dir, exist_ok=True)
    doc = pymupdf.open(input_path)

    # 检测是否为扫描件：每页文字极少，且每页只有1张图
    total_text = ""
    page_image_counts = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text().strip()
        total_text += text
        img_count = len(page.get_images(full=True))
        page_image_counts.append(img_count)

    total_pages = len(doc)
    total_text_len = len(total_text.strip())
    is_scanned = total_text_len < total_pages * 20  # 平均每页少于20字 → 扫描件

    if is_scanned:
        print(f"检测到扫描件 PDF（{total_pages}页，平均每页文字{total_text_len // max(total_pages, 1)}字），"
              f"将每页渲染为图片。")
        return render_pdf_pages_as_images(input_path, output_dir)

    # 非扫描件：提取内嵌图片
    mapping = []
    image_idx = 0

    for page_num in range(len(doc)):
        page = doc[page_num]
        page_text = page.get_text()

        question_positions = []
        for match in re.finditer(r'[第]?[\s]*(\d+)[题\.、．\s]', page_text):
            q_num = int(match.group(1))
            question_positions.append(q_num)

        image_list = page.get_images(full=True)
        for img_info in image_list:
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                if base_image:
                    image_bytes = base_image["image"]
                    image_ext = base_image.get("ext", "png")
                    image_idx += 1
                    image_file = f"image_{image_idx}.{image_ext}"
                    image_path = os.path.join(output_dir, image_file)
                    with open(image_path, "wb") as f:
                        f.write(image_bytes)

                    question_number = 0
                    context = ""
                    try:
                        imgs_on_page = page.get_image_info(xrefs=True)
                        for img_detail in imgs_on_page:
                            if img_detail.get("xref") == xref:
                                img_y = img_detail.get("top", 0)
                                if question_positions:
                                    question_number = question_positions[0]
                                context = f"第{page_num+1}页，y坐标约{img_y:.0f}"
                                break
                    except Exception:
                        context = f"第{page_num+1}页"

                    mapping.append({
                        "image_file": image_file,
                        "question_number": question_number,
                        "page": page_num + 1,
                        "context": context,
                    })
            except Exception as e:
                print(f"警告: 无法提取图片 xref={xref}: {e}", file=sys.stderr)

    return mapping


def vision_match_images(images_dir: str, mapping: list, question_range: str = "") -> list:
    """
    用视觉模型识别每张图片内容，精准匹配题号。
    
    生成一个匹配指令文件，由 LLM 调用 image_understanding 工具逐张识别后，
    将结果写回映射 JSON。
    
    此函数输出待匹配的图片列表到 stdout，格式：
      图片文件 | 页码 | 当前映射题号 | 待识别
    
    然后从 stdin 读取 LLM 的识别结果更新映射。
    """
    # 输出匹配任务描述
    print("\n===== 视觉识别匹配任务 =====")
    print(f"共 {len(mapping)} 张图片需要识别，请逐张调用 image_understanding 工具")
    print(f"识别目标：判断图片属于第几题的配图，还是无关内容（如答案页、装饰图）")
    if question_range:
        print(f"题目范围：{question_range}")
    print("=" * 40 + "\n")

    # 打印待识别清单
    for i, m in enumerate(mapping):
        img_path = os.path.join(images_dir, m["image_file"])
        print(f"[{i+1}] {m['image_file']} | 第{m['page']}页 | 当前题号={m['question_number']} | 路径={img_path}")

    print("\n请对每张图片调用 image_understanding，prompt 如下：")
    print('  "这是一份试卷中的图片，请判断：1.这张图属于第几题的配图？（填写题号数字，如3或9）')
    print('   2.图片类型是什么？（流程图/几何图/函数图/数据表/答案/装饰/其他）')
    print('   3.如果不是任何题目的配图，回答0"')
    print("\n识别完成后，请将结果整理为 JSON 数组写入映射文件，格式：")
    print('[{"image_file": "xxx.png", "question_number": 9, "image_type": "流程图", "context": "第9题分支结构流程图"}]')

    return mapping


def main():
    parser = argparse.ArgumentParser(
        description="从 Word/PDF 中提取图片并映射到题号"
    )
    parser.add_argument("--input", required=True, help="输入文件路径（.docx 或 .pdf）")
    parser.add_argument("--output-dir", required=True, help="图片输出目录")
    parser.add_argument("--mapping", required=True, help="映射 JSON 输出路径")
    parser.add_argument("--vision-match", action="store_true",
                        help="启用视觉识别模式，用视觉模型精准匹配题号（适合扫描件、答案集中放置等场景）")
    parser.add_argument("--question-range", default="",
                        help="试卷题号范围（如'1-20'），辅助视觉模型判断")
    parser.add_argument("--dpi", type=int, default=200,
                        help="PDF 页面渲染 DPI（仅扫描件有效，默认200）")

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"错误: 文件不存在: {args.input}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(args.input)[1].lower()

    if ext == ".docx":
        mapping = extract_images_from_docx(args.input, args.output_dir)
    elif ext == ".pdf":
        mapping = extract_images_from_pdf(args.input, args.output_dir)
    else:
        print(f"错误: 不支持的文件格式: {ext}，仅支持 .docx 和 .pdf", file=sys.stderr)
        sys.exit(1)

    if args.vision_match:
        vision_match_images(args.output_dir, mapping, args.question_range)

    os.makedirs(os.path.dirname(os.path.abspath(args.mapping)), exist_ok=True)
    with open(args.mapping, "w", encoding="utf-8") as f:
        json.dump(mapping, f, ensure_ascii=False, indent=2)

    print(f"\n提取完成！共 {len(mapping)} 张图片")
    print(f"图片目录: {args.output_dir}")
    print(f"映射文件: {args.mapping}")

    unmapped = sum(1 for m in mapping if m["question_number"] == 0)
    if unmapped:
        print(f"注意: {unmapped} 张图片未映射到题号")
        if args.vision_match:
            print("请按照上述提示用 image_understanding 工具逐张识别，")
            print("然后将结果更新到映射 JSON 文件的 question_number 字段。")


if __name__ == "__main__":
    main()
