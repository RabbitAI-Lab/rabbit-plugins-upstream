#!/usr/bin/env python3
"""
docx-writer 模板脚本。
使用方法：
  1. 复制此文件到工作目录
  2. 修改 OUTPUT、FIG_DIR 路径
  3. 修改 FN_TEXT_MAP 为脚注内容
  4. 在 build_content() 中用 B('文本{F1}') 编写正文
  5. python build_docx.py
"""
import sys, os
from lxml import etree
import docx as dx
from docx.shared import Pt, Cm

sys.stdout.reconfigure(encoding='utf-8')

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
fb = lambda t: '{%s}%s' % (W_NS, t)

# ═══════════════════ 配置 ═══════════════════
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SKILL_DIR = os.path.dirname(SCRIPT_DIR)
TEMPLATE = os.path.join(SKILL_DIR, 'template.docx')
OUTPUT = r'输出路径.docx'
FIG_DIR = r'图片目录'

FN_TEXT_MAP = {
    'F1': '作者. 标题. 期刊, 年份, 卷(期): 页码.',
}

# ═══════════════════ 加载模板 ═══════════════════
doc = dx.Document(TEMPLATE)
body = doc.element.body
for child in list(body):
    local = etree.QName(child).localname if child.tag is not etree.Comment else ''
    if local != 'sectPr':
        body.remove(child)

for rel in doc.part.rels.values():
    if 'footnotes' in rel.reltype:
        fn_elem = rel.target_part.element
        for fn in list(fn_elem):
            if fn.get(fb('type'), '') not in ('separator', 'continuationSeparator'):
                fn_elem.remove(fn)
        for fn in list(fn_elem):
            if fn.get(fb('type')) == 'separator':
                fn.set(fb('id'), '-1')
            elif fn.get(fb('type')) == 'continuationSeparator':
                fn.set(fb('id'), '0')
        break

# ═══════════════════ 排版函数（不改） ═══════════════════
def set_font(r, cn='宋体', en='Times New Roman', sz=Pt(10.5)):
    r.font.name = en; r.font.size = sz
    rPr = r._element.find(fb('rPr'))
    if rPr is None: rPr = etree.SubElement(r._element, fb('rPr'))
    rf = rPr.find(fb('rFonts'))
    if rf is None: rf = etree.SubElement(rPr, fb('rFonts'))
    rf.set(fb('eastAsia'), cn); rf.set(fb('ascii'), en); rf.set(fb('hAnsi'), en)

def P(text, cn='宋体', en='Times New Roman', sz=Pt(10.5), align=None, fi=Cm(0.74), ls=1.5):
    p = doc.add_paragraph()
    if align == 'center': p.alignment = 1
    if fi: p.paragraph_format.first_line_indent = fi
    if ls: p.paragraph_format.line_spacing = ls
    segs = text.split('{')
    for seg in segs:
        if '}' in seg:
            fn_name, rest = seg.split('}', 1)
            fn = p.add_footnote(FN_TEXT_MAP[fn_name])
            fn_p = fn.find(fb('p'))
            if fn_p is not None:
                pp = fn_p.find(fb('pPr'))
                if pp is None: pp = etree.SubElement(fn_p, fb('pPr'))
                ps = pp.find(fb('pStyle'))
                if ps is None: ps = etree.SubElement(pp, fb('pStyle'))
                ps.set(fb('val'), '12')
            for r in fn.findall('.//' + fb('r')):
                t = r.find(fb('t'))
                if t is not None and t.text and t.text.strip():
                    rPr = r.find(fb('rPr'))
                    if rPr is None: rPr = etree.SubElement(r, fb('rPr'))
                    rf = rPr.find(fb('rFonts'))
                    if rf is None: rf = etree.SubElement(rPr, fb('rFonts'))
                    rf.set(fb('eastAsia'), '宋体')
                    rf.set(fb('ascii'), 'Times New Roman')
                    rf.set(fb('hAnsi'), 'Times New Roman')
                    for tag in ['sz', 'szCs']:
                        el = rPr.find(fb(tag))
                        if el is None: el = etree.SubElement(rPr, fb(tag))
                        el.set(fb('val'), '18')
            if rest:
                r2 = p.add_run(rest); set_font(r2, cn, en, sz)
        elif seg:
            r = p.add_run(seg); set_font(r, cn, en, sz)
    return p

def B(t): return P(t)
def Bs(t): return P(t, sz=Pt(9))
def H1(t): return P(t, cn='黑体', en='Arial', sz=Pt(12), align='center', fi=None)
def H2(t): return P(t, cn='黑体', en='Arial', sz=Pt(12))
def H3(t): return P(t, cn='黑体', en='Arial', sz=Pt(10.5))
def Tc(t): return P(t, sz=Pt(9), align='center', fi=None)

def add_img(path, cap, w=Cm(14)):
    if os.path.exists(path):
        doc.add_picture(path, width=w)
        p = doc.paragraphs[-1]; p.alignment = 1
        p.paragraph_format.first_line_indent = Pt(0)
    Tc(cap)

def add_tbl(h, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(h))
    t.style = 'Table Grid'
    for row in t.rows:
        for cell in row.cells:
            tc = cell._tc; tcPr = tc.find(fb('tcPr'))
            if tcPr is None: tcPr = etree.SubElement(tc, fb('tcPr'))
            b = etree.SubElement(tcPr, fb('tcBorders'))
            for s in ['top','left','bottom','right']:
                el = etree.SubElement(b, fb(s)); el.set(fb('val'), 'nil')
    for i, x in enumerate(h):
        cell = t.rows[0].cells[i]; cell.text = ''
        run = cell.paragraphs[0].add_run(x)
        run.font.size = Pt(9); run.font.bold = True
        cell.paragraphs[0].alignment = 1
        tc = cell._tc; tcPr = tc.find(fb('tcPr'))
        b = tcPr.find(fb('tcBorders'))
        for s in ['top','bottom']:
            el = etree.SubElement(b, fb(s))
            el.set(fb('val'), 'single'); el.set(fb('sz'), '4')
    for ri, r in enumerate(rows):
        for ci, v in enumerate(r):
            cell = t.rows[ri+1].cells[ci]; cell.text = ''
            run = cell.paragraphs[0].add_run(str(v))
            run.font.size = Pt(9)
            cell.paragraphs[0].alignment = 1
            if ri == len(rows) - 1:
                tc = cell._tc; tcPr = tc.find(fb('tcPr'))
                b = tcPr.find(fb('tcBorders'))
                el = etree.SubElement(b, fb('bottom'))
                el.set(fb('val'), 'single'); el.set(fb('sz'), '4')
    return t

# ═══════════════════ 正文 ═══════════════════
def build_content():
    # 在这里用 B/H1/H2/H3/Tc/add_img/add_tbl 编写内容
    # 示例：
    # B('这是正文{F1}')
    pass

build_content()

# ═══════════════════ 后处理：上标 ═══════════════════
for p in body.findall('.//' + fb('p')):
    for r in p.findall('.//' + fb('r')):
        if r.find(fb('footnoteReference')) is not None:
            rPr = r.find(fb('rPr'))
            if rPr is None: rPr = etree.SubElement(r, fb('rPr'))
            va = rPr.find(fb('vertAlign'))
            if va is None: va = etree.SubElement(rPr, fb('vertAlign'))
            va.set(fb('val'), 'superscript')

# ═══════════════════ 保存 ═══════════════════
doc.save(OUTPUT)
print(f'Saved: {OUTPUT}', file=sys.stderr)
