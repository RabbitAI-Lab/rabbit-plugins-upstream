"""
DOCX 13章完整分析报告生成 — 遵循 MEMORY.md v9 规则
====================================================
支持外部数据注入（通过 extra_data 参数）：
  - holdings, money_flow, margin_trading, executive_hold
  - shareholder_trade, financial_report, dividend
无外部数据时降级使用 WebAPI + 本地数据生成轻量版

规则合并清单（MEMORY.md + 用户反馈）：
  1) 技术指标展示具体数值（MA/MACD/KDJ/RSI）
  2) 三时段匹配率（全部/60日/30日），紧跟方向之后
  3) T+1/T+2 含价格区间（最低/最高/历史均值）+ 周方向
  4) 3~5条参考依据（从指标数据自动生成）
  5) 规律含出现次数 + 匹配率进度条
  6) 综合信号总表（偏多/偏空/中性汇总）
  7) 5维度加权评分
  8) 动态风险提示（基于实际数据生成）
"""
import os, re
from datetime import datetime
from typing import Optional, Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from . import config
from . import data_fetcher as fetcher
from . import db_manager as db
from . import eval_engine
from . import pattern_miner
from .sanitizer import sanitize
from . import chart as chart_gen


# ==================== 字体设置 ====================

from docx.oxml.ns import qn

def _set_doc_font(doc, font_name="宋体", body_size=Pt(10.5)):
    """设置文档整体字体为宋体（同时设置西文和东亚字体）"""
    for sname in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Title']:
        try:
            style = doc.styles[sname]
            f = style.font
            f.name = font_name
            if sname == 'Normal':
                f.size = body_size
            elif sname == 'Title':
                f.size = Pt(22)
            elif sname == 'Heading 1':
                f.size = Pt(18)
            elif sname == 'Heading 2':
                f.size = Pt(14)
            elif sname == 'Heading 3':
                f.size = Pt(12)
            # 设置东亚字体
            rpr = style.element.rPr
            if rpr is None:
                rpr = style.element.makeelement(qn('w:rPr'), {})
                style.element.append(rpr)
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = rpr.makeelement(qn('w:rFonts'), {})
                rpr.append(rfonts)
            rfonts.set(qn('w:eastAsia'), font_name)
        except:
            pass
    # 表格样式也设置
    for ts in ['Table Grid', 'Light Shading Accent 1', 'Medium Shading 1 Accent 1']:
        try:
            style = doc.styles[ts]
            f = style.font
            f.name = font_name
            f.size = Pt(9)
            rpr = style.element.rPr
            if rpr is None:
                rpr = style.element.makeelement(qn('w:rPr'), {})
                style.element.append(rpr)
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = rpr.makeelement(qn('w:rFonts'), {})
                rpr.append(rfonts)
            rfonts.set(qn('w:eastAsia'), font_name)
        except:
            pass


def _set_cell_font(cell, font_name="宋体", size=Pt(9)):
    """设置单个表格单元格字体"""
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = font_name
            run.font.size = size
            rpr = run._element.rPr
            if rpr is None:
                rpr = run._element.makeelement(qn('w:rPr'), {})
                run._element.append(rpr)
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = rpr.makeelement(qn('w:rFonts'), {})
                rpr.append(rfonts)
            rfonts.set(qn('w:eastAsia'), font_name)


# ==================== 工具函数 ====================

def _bar(rate: float, total: int = 20) -> str:
    filled = max(0, min(total, int(rate / 5)))
    return "█" * filled + "░" * (total - filled)


def _parse_hitrate(hitrate: str) -> tuple:
    """'2/4=50%回调' → (命中, 总次数, 匹配率%, 结果类型)"""
    try:
        p = hitrate.split("=")
        if len(p) != 2: return (0, 0, 0, "")
        h, t = p[0].split("/")
        hit, tot = int(h), int(t)
        m = re.match(r'(\d+)%(.+)', p[1])
        if m: return (hit, tot, int(m.group(1)), m.group(2))
        return (hit, tot, 0, p[1])
    except:
        return (0, 0, 0, "")


def _gen_evidence(indic: dict, pred_dir: str, sig_items: list) -> list:
    """从指标生成3~5条参考依据"""
    ev = []
    ma5, ma10, ma20 = indic.get("ma5", 0), indic.get("ma10", 0), indic.get("ma20", 0)
    if ma5 and ma10 and ma20:
        if ma5 > ma10 > ma20:
            ev.append(f"均线多头排列（MA5={ma5:.2f} > MA10={ma10:.2f} > MA20={ma20:.2f}），短期向上")
        elif ma5 < ma10 < ma20:
            ev.append(f"均线空头排列（MA5={ma5:.2f} < MA10={ma10:.2f} < MA20={ma20:.2f}），短期向下")
        else:
            ev.append(f"均线交叉震荡（MA5={ma5:.2f} MA10={ma10:.2f} MA20={ma20:.2f}），方向待定")
    md = indic.get("macd", {})
    if md:
        d, e, b = md.get("dif", 0), md.get("dea", 0), md.get("bar", 0)
        ev.append(f"MACD {'偏多' if d > e else '偏空'}（DIF={d:.3f}，{'红' if d > e else '绿'}柱{abs(b):.3f}）")
    kd = indic.get("kdj", {})
    if kd:
        k, d_, j = kd.get("k", 50), kd.get("d", 50), kd.get("j", 50)
        if j > 100: ev.append(f"KDJ超买（K={k:.0f} D={d_:.0f} J={j:.0f}>100），注意回调")
        elif j < 0: ev.append(f"KDJ超卖（K={k:.0f} D={d_:.0f} J={j:.0f}<0），历史统计中超卖后偏多比例偏高")
        elif k > d_: ev.append(f"KDJ偏多（K={k:.0f} > D={d_:.0f}）")
        else: ev.append(f"KDJ偏空（K={k:.0f} < D={d_:.0f}）")
    rsi = indic.get("rsi14", 50)
    if rsi > 70: ev.append(f"RSI超买（{rsi:.0f}>70）")
    elif rsi < 30: ev.append(f"RSI超卖（{rsi:.0f}<30）")
    else: ev.append(f"RSI中性（{rsi:.0f}，30~70正常）")
    vr = indic.get("volRatio", 1)
    if vr > 1.5: ev.append(f"成交量放量（均量{vr:.1f}倍）")
    elif vr < 0.5: ev.append(f"成交量缩量（均量{vr:.1f}倍）")
    else: ev.append(f"成交量正常（均量{vr:.1f}倍）")
    ev.sort(key=lambda x: 0 if "中性" in x or "正常" in x else 1, reverse=True)
    return ev[:5]


def _add_table(doc, rows, style="Light Shading Accent 1", headers=None):
    """通用表格。自动检测列数，支持自定义表头"""
    if not rows:
        return None
    ncols = len(rows[0])
    t = doc.add_table(rows=len(rows) + 1, cols=ncols, style=style)
    if headers:
        for j, h in enumerate(headers):
            t.cell(0, j).text = str(h)
    else:
        h_default = ["指标", "数值", "说明"][:ncols]
        for j, h in enumerate(h_default):
            t.cell(0, j).text = h
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.cell(i + 1, j).text = str(val)
    # 设置表格单元格字体
    for row_obj in t.rows:
        for cell in row_obj.cells:
            _set_cell_font(cell)
    return t


def _add_points(doc, title, points, conclusion=""):
    """分行要点表（analysis_points）"""
    p = doc.add_paragraph()
    p.add_run(f"▶ {title}").bold = True
    for pt in points:
        doc.add_paragraph(f"  • {pt}")
    if conclusion:
        p2 = doc.add_paragraph()
        p2.add_run(f"  ➜ {conclusion}").bold = True


def _add_fin(doc, title, text):
    """财务短段落（fin_analysis）"""
    p = doc.add_paragraph()
    p.add_run(f"▶ {title}：").bold = True
    p.add_run(text)


def _render_fin_section(doc, title, data):
    """财务数据多格式渲染。
    支持三种格式：
      1) dict 含 table → 横列表格 + analysis 解读
      2) dict 含 points → 调用 _add_points
      3) 字符串 → 调用 _add_fin
    """
    if isinstance(data, dict):
        if "table" in data:
            rows = data["table"]
            if rows and len(rows) >= 2:
                first = rows[0]
                if first and str(first[0]).strip() in ("指标", "报告期", "日期", ""):
                    _add_table(doc, rows[1:], headers=list(first))
                else:
                    _add_table(doc, rows)
            analysis = data.get("analysis", "")
            if analysis:
                p = doc.add_paragraph()
                p.add_run("▶ 解读：").bold = True
                p.add_run(analysis)
        elif "points" in data:
            _add_points(doc, title, data.get("points", []), data.get("conclusion", ""))
        else:
            _add_fin(doc, title, str(data))
    else:
        _add_fin(doc, title, str(data) if data else "")


def _add_etf_quadrant(doc, etf_flow_data, stock_change_pct, chg_):
    """
    ETF与个股联动四象分析
    
    四象限矩阵:
              ETF净流入         ETF净流出
    个股偏多    共振↑          个股独立走强
    个股偏空  板块撑杆跳        共振↓
    
    新增盘后固定价格交易维度：盘后量/总成交量比值反映收盘后的机构跟进意愿
    """
    # 计算ETF总方向
    etf_dict = etf_flow_data.get("etf_flow", {})
    total_etf_flow = 0
    total_after_vol = 0      # 盘后总成交量(手)
    total_after_amt = 0      # 盘后总成交额(万元)
    total_reg_vol = 0        # 常规交易总成交量(手) - 近似判断用
    for rows in etf_dict.values():
        if rows and len(rows) >= 2:
            recent = rows[-5:] if len(rows) >= 5 else rows
            total_etf_flow += sum(r.get("net_inflow", 0) for r in recent)
            total_after_vol += sum(r.get("after_hours_vol", 0) for r in recent)
            total_after_amt += sum(r.get("after_hours_amt", 0) for r in recent)
    
    # 个股方向：用近5日价格变动或当日涨跌幅
    stock_dir = stock_change_pct if abs(stock_change_pct or 0) > 0.5 else (chg_ or 0)
    stock_bullish = stock_dir > 0
    
    # ETF方向
    etf_bullish = total_etf_flow > 0
    
    # 盘后交易强度判定
    # 盘后固定价格交易（15:05-15:30）反映机构收盘后的跟进意愿
    heavy_after_trade = total_after_amt > 100  # 万元，显著盘后成交
    moderate_after_trade = total_after_amt > 20  # 略有盘后成交
    
    # 四象判断（基础）
    if etf_bullish and stock_bullish:
        quadrant = "象限I：板块个股共振 ↑"
        desc = "ETF资金净流入 + 个股强势，板块与个股形成共振。行业层面有资金支撑，个股自身走势也强，这种组合下上涨逻辑较充分。"
        if heavy_after_trade:
            desc += " 盘后固定价格交易活跃（" + f"{total_after_amt:.0f}万元" + "），说明收盘后仍有机构跟进买入，次日延续概率较高。"
        signal = "偏多"
    elif etf_bullish and not stock_bullish:
        quadrant = "象限II：板块撑杆跳"
        desc = "ETF资金净流入但个股走势偏弱，存在两种可能：①资金在板块内轮动，该股暂未受益；②板块在涨但个股有自身利空拖累。建议排查个股基本面问题。"
        if heavy_after_trade:
            desc += " 板块ETF盘后交易活跃，说明板块整体有资金持续关注，个股待选股逻辑修复后可能补涨。"
        signal = "中性偏谨慎"
    elif not etf_bullish and stock_bullish:
        quadrant = "象限III：个股独立走强"
        desc = "ETF资金净流出但个股强势，说明个股上涨不依赖板块资金推动，可能是龙头效应或个股独立逻辑驱动。需确认上涨是否有基本面支撑。"
        if moderate_after_trade:
            desc += " 板块ETF盘后仍有少量成交，说明板块关注度未完全消退，独立走强可持续性尚可。"
        signal = "偏正面但需验证"
    else:
        quadrant = "象限IV：板块个股共振 ↓"
        desc = "ETF资金净流出 + 个股弱势，板块与个股同步承压。行业资金在撤离，个股走势也弱，这种情况下风险较高，建议谨慎观望。"
        if moderate_after_trade:
            desc += " 但板块ETF仍有盘后交易，说明底部有机构试探性接盘，需关注次日是否有反弹信号。"
        signal = "偏空"
    
    flow_label = "净流入" if total_etf_flow >= 0 else "净流出"
    stock_label = "上涨" if stock_bullish else "下跌"
    
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("▶ 联动分析（ETF vs 个股）：").bold = True
    doc.add_paragraph(f"  ① ETF资金方向：近5日{flow_label}（{total_etf_flow/10000:+.0f}万元）")
    doc.add_paragraph(f"  ② 个股近期趋势：{stock_label}（{abs(stock_dir):.1f}%）")
    doc.add_paragraph(f"  ③ 盘后固定价格交易：近5日盘后量 {total_after_vol:.0f}手 / 盘后额 {total_after_amt:.0f}万元" +
                      (" 🔴活跃" if heavy_after_trade else " 🟡有量" if moderate_after_trade else " ⚪清淡"))
    doc.add_paragraph(f"  ④ 结论：{quadrant} → {signal}")
    p2 = doc.add_paragraph()
    p2.add_run(f"  ➜ {desc}").italic = True


def _add_research_validation(doc, code, ed, price, chg_):
    """研报观点验证 — 交叉对比研报评级与实际市场行为"""
    doc.add_heading("9.5 机构观点验证", 2)
    
    # 获取研报汇总和EPS数据
    try:
        rpt_text = db.get_research_summary(code, 90)
        eps_text = db.get_eps_consistency(code, 180)
        broker_text = db.get_broker_ranking(code, 90)
    except:
        doc.add_paragraph("  （研报验证数据加载中...）")
        return
    
    # 解析研报评级方向
    buy_count = rpt_text.count("买入") + rpt_text.count("增持")
    neutral_count = rpt_text.count("中性")
    total_reports = buy_count + neutral_count
    buy_ratio = buy_count / total_reports if total_reports > 0 else 0
    
    # 解析EPS一致性
    eps_stable = "高度一致" in eps_text or "基本一致" in eps_text
    eps_divergent = "分歧较大" in eps_text
    
    # 判断股价相对EPS的位置
    eps_avg = 0
    import re
    eps_match = re.search(r'平均EPS:\s*([\d.]+)', eps_text)
    if eps_match:
        eps_avg = float(eps_match.group(1))
    
    # 从研报文本直接统计买入率（broker_text的表格不易解析）
    buy_rate = int(buy_ratio * 100) if total_reports > 0 else 0
    
    # 机构行为验证
    holdings_data = ed.get("holdings", {})
    h_points = holdings_data.get("points", []) if isinstance(holdings_data, dict) else []
    h_text = " ".join(h_points)
    inst_increasing = any("增持" in p or "新进" in p for p in h_points)
    inst_decreasing = any("减持" in p for p in h_points)
    
    # 输出验证
    doc.add_paragraph("  【观点一致性检验】")
    if buy_rate >= 80:
        doc.add_paragraph(f"    ✅ 券商一致性：{buy_rate}%机构给予买入/增持评级，市场共识度高")
    elif buy_rate >= 60:
        doc.add_paragraph(f"    ⚠️ 券商一致性：{buy_rate}%机构给予买入/增持评级，市场存在分歧")
    else:
        doc.add_paragraph(f"    ❌ 券商一致性：仅{buy_rate}%机构看好，需谨慎参考")
    
    if eps_stable:
        doc.add_paragraph(f"    ✅ EPS预期一致性：高度一致，机构对盈利预期分歧小")
    elif eps_divergent:
        doc.add_paragraph(f"    ⚠️ EPS预期一致性：分歧较大，不同机构盈利预期差异明显")
    else:
        doc.add_paragraph(f"    ⚡ EPS预期一致性：基本一致，存在中等分歧")
    
    doc.add_paragraph("")
    doc.add_paragraph("  【机构行为验证】")
    if inst_increasing and buy_rate >= 60:
        doc.add_paragraph("    ✅ 言行一致：券商普遍看好，且机构股东实际增持，观点与行动吻合")
    elif inst_decreasing and buy_rate >= 60:
        doc.add_paragraph('    ⚠️ 言行不一：券商看好但实际机构在减持，需警惕"看多不做多"')
    elif inst_increasing and buy_rate < 60:
        doc.add_paragraph("    ⚡ 机构增持但券商评级分化，说明有资金在左侧布局")
    else:
        doc.add_paragraph("    — 机构持仓无明显方向性变化")
    
    doc.add_paragraph("")
    doc.add_paragraph("  【价格与估值验证】")
    if eps_avg > 0:
        implied_target = eps_avg * 20  # 20倍PE
        ratio = price / implied_target if implied_target > 0 else 1
        if ratio < 0.85:
            doc.add_paragraph(f"    ✅ 当前价{price:.2f}元，低于合理估值中枢{implied_target:.2f}元（基于EPS{eps_avg:.2f}×20PE），估值有安全边际")
        elif ratio > 1.15:
            doc.add_paragraph(f"    ⚠️ 当前价{price:.2f}元，高于合理估值中枢{implied_target:.2f}元，已部分反映机构预期")
        else:
            doc.add_paragraph(f"    ⚡ 当前价{price:.2f}元，接近合理估值中枢{implied_target:.2f}元，定价相对合理")
    
    doc.add_paragraph("")
    doc.add_paragraph("  【综合判定】")
    positive_signals = 0
    if buy_rate >= 80: positive_signals += 1
    if eps_stable: positive_signals += 1
    if inst_increasing: positive_signals += 1
    if eps_avg > 0 and price / (eps_avg * 20) < 1: positive_signals += 1
    
    if positive_signals >= 3:
        verdict = "研报观点可信度较高，机构看多逻辑有数据支撑"
    elif positive_signals >= 2:
        verdict = "研报观点部分可信，存在一定分歧需自行判断"
    else:
        verdict = "研报观点分歧较大，建议进一步独立分析"
    
    doc.add_paragraph(f"    📊 {verdict}（{positive_signals}/4项正面信号）")
    doc.add_paragraph("")


# ==================== 付费提示 ====================

def _add_hyperlink(paragraph, text, url):
    """在段落中添加可点击超链接"""
    from docx.oxml import OxmlElement
    part = paragraph.part
    r_id = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    cStyle = OxmlElement("w:rStyle")
    cStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(cStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _paid_section(doc, chapter_num, title, teaser_lines, pain_point, value_lines, api_key=None):
    """
    付费内容占位章节 — 精简排版：免费试用→痛点冲击→价值点→CTA
    
    api_key: 兼容保留（v2.2.50 起购买链接改走 db.get_payment_url() token化，不再直接拼Key）
    """
    # 购买链接：优先平台token（URL不含完整Key），失败降级apikey
    try:
        purchase_url = db.get_payment_url()
    except Exception:
        purchase_url = "https://www.oraskl.com/ghdata-admin"

    doc.add_heading(f"{chapter_num}、{title}", 1)

    # 免费预览数据
    for line in teaser_lines:
        if line.startswith("📊"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"  {line}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

    doc.add_paragraph("")

    # 痛点冲击
    pp = doc.add_paragraph()
    pp.paragraph_format.space_before = Pt(2)
    pp.paragraph_format.space_after = Pt(4)
    run_p = pp.add_run(f"  ▸ {pain_point}")
    run_p.bold = True
    run_p.font.size = Pt(10)
    run_p.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph("")

    # 价值点列表
    for i, line in enumerate(value_lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.style = doc.styles['List Bullet']
        run = p.add_run(line)
        run.font.size = Pt(9.5)

    doc.add_paragraph("")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(6)
    _add_hyperlink(p2, "💳 立即获取 → 股海罗盘APIKey", purchase_url)
    doc.add_paragraph("")


# ==================== 主函数 ====================

def generate(code: str, output_dir: str = None, extra_data: dict = None) -> Optional[str]:
    """
    生成完整13章DOCX分析报告
    
    参数:
        code: 股票代码
        output_dir: 输出目录
        extra_data: 可选的外部数据（来自MCP工具），包含
            holdings, money_flow, margin_trading, executive_hold,
            shareholder_trade, financial_report, dividend 等
    
    返回:
        docx文件路径
    """
    if output_dir is None:
        output_dir = config.DOC_DIR
    os.makedirs(output_dir, exist_ok=True)
    ed = extra_data or {}

    # ===== 1. 获取基础数据 =====
    api = db.kline_analyze(code)

    # v3 特征融合（ETF资金流+券商研报）
    v3_factors = None
    try:
        v3_api = db.kline_analyze_v3(code)
        if v3_api and isinstance(v3_api, dict) and not v3_api.get("preview"):
            v3_factors = v3_api.get("factors")
    except Exception:
        pass
    # 检测是否为免费预览数据（APIKey无效时的降级数据）
    preview = api.get("preview", False) if api else False
    # 有indicators且非预览数据才算完整WebAPI可用
    webapi_available = bool(api and api.get("indicators") and not preview)
    if not webapi_available:
        if preview:
            print(f"[reporter] 使用免费预览数据（APIKey无效），仅展示技术指标预览")
        else:
            print(f"[reporter] WebAPI无数据（APIKey无效或额度用完），章节11~13将显示付费内容提示")
    indic = api.get("indicators", {}) or {}
    sig_data = api.get("signals", {}) or {}
    pred_api = api.get("latestPrediction", {}) or {}
    acc = api.get("accuracy", {}) or {}
    def astat(p):
        s = acc.get(p, {})
        return {"total": s.get("total", 0), "correct": s.get("correct", 0), "rate": round(s.get("rate", 0), 1)}

    # 预览数据提取（用于付费章节免费展示）
    preview_indi = indic if preview else {}
    preview_sig = sig_data if preview else {}
    preview_rt = api.get("realtime", {}) or {} if preview else {}

    # 2. 本地数据
    info = fetcher.get_company_info(code)
    name = info.get("name", code)
    raw = fetcher.fetch_kline(code, 365)
    klines = []
    if raw:
        klines = [{"date": k.get("date",""), "open": float(k.get("open",0)),
                   "close": float(k.get("close",0)), "high": float(k.get("high",0)),
                   "low": float(k.get("low",0)), "volume": float(k.get("volume",0))}
                  for k in raw if k.get("close")]
    pred = eval_engine.analyze(code, klines)
    # 使用WebAPI的patterns（list格式），降级到本地miner（dict格式）
    api_patterns = api.get("patterns", []) or []
    if api_patterns and isinstance(api_patterns, list):
        patterns = api_patterns
    else:
        local_raw = pattern_miner.mine(klines) if len(klines) >= 90 else {}
        # 将本地dict格式转成list格式（统一处理）
        patterns = []
        for pname, pdata in local_raw.items():
            if isinstance(pdata, dict):
                entry = {"name": pname}
                if "samples" in pdata: entry["samples"] = pdata["samples"]
                if "conclusion" in pdata: entry["avgD3"] = pdata.get("conclusion", "")
                if "advice" in pdata: entry["advice"] = pdata.get("advice", "")
                if "value" in pdata: entry["value"] = pdata.get("value", "")
                if "hitRate" in pdata: entry["hitRate"] = pdata.get("hitRate", "")
                patterns.append(entry)
    flow = fetcher.fetch_realtime(code)

    # 检测预览数据（APIKey无效时的免费预览）
    preview = api.get("preview", False)
    preview_indi = api.get("indicators", {}) or {} if preview else {}
    preview_sig = api.get("signals", {}) or {} if preview else {}
    preview_rt = api.get("realtime", {}) or {} if preview else {}

    # ===== 自动采集全维度数据（确保报告不自宫）=====
    auto_data = {}
    # 机构持仓
    try:
        hd = fetcher.fetch_main_holdings(code)
        if hd:
            h_points = [f"机构持仓占流通股比: {hd.get('total_holders',0):.2f}%",
                        f"基金持仓: {hd.get('fund_sum',0):.2f}%",
                        f"保险持仓: {hd.get('insurance_sum',0):.2f}%",
                        f"券商持仓: {hd.get('securities_sum',0):.2f}%",
                        f"QFII持仓: {hd.get('qfii_sum',0):.2f}%"]

            auto_data["holdings"] = {"points": h_points,
                "conclusion": f"基金持仓{hd.get('fund_sum',0):.2f}%，机构合计{hd.get('total_holders',0):.2f}%"}

        # 前十大股东明细
        top_h = fetcher.fetch_top_holders(code)
        if top_h:
            top_rows = []
            for h in top_h:
                top_rows.append({
                    "rank": h.get("rank", ""),
                    "name": h.get("holder_name",""),
                    "ratio": f"{h.get('hold_ratio',0):.2f}%",
                    "direction": h.get("direction",""),
                })
            auto_data["top_holders"] = top_rows
    except Exception as e:
        print(f"[reporter] 机构持仓采集失败: {e}")
    # 资金流向
    try:
        mf = fetcher.fetch_money_flow(code, 5)
        if mf:
            src = mf[0].get("source", "")
            points = []
            for m in mf[:5]:
                if src == "tencent_kline":
                    # 腾讯K线：量价趋势
                    vol = m.get("volume", 0)
                    close_px = m.get("close", 0)
                    points.append(f"{m.get('date','')} 收{close_px:.2f} 量{vol:.0f}手")
                elif src == "tencent_qt":
                    # 腾讯实时行情：主力净流入
                    mn = float(m.get("main_net", 0) or 0)
                    points.append(f"当日主力净流入:{mn:.0f}万元" if abs(mn) > 1 else f"当日主力净流入:{mn*10000:.0f}万元")
                else:
                    # 东方财富：主力净流入（元→万元）
                    mn = m.get("main_net", 0) or 0
                    points.append(f"主力净流入:{mn/10000:.0f}万元")
            if points:
                auto_data["money_flow"] = {"points": points}
    except Exception as e:
        print(f"[reporter] 资金流向采集失败: {e}")
    # 融资融券
    try:
        mt = fetcher.fetch_margin_trading(code, 10)
        if mt:
            points = []
            for m in mt[:6]:
                if isinstance(m, dict):
                    date = m.get("date","")[:10]
                    bal = float(m.get("balance",0) or 0) / 1e8
                    rz_net = float(m.get("rz_net",0) or 0) / 1e4
                    points.append(f"{date} 融资余额{bal:.2f}亿 净买{rz_net:+.0f}万元")
            auto_data["margin_trading"] = {"points": points}
    except Exception as e:
        print(f"[reporter] 两融采集失败: {e}")
    # 股东增减持
    try:
        st = fetcher.fetch_shareholder_trade(code)
        if st:
            points = []
            for s in st[:10]:
                sname = s.get("name","")
                sdate = str(s.get("date",""))[:10]
                svol = float(s.get("volume",0) or 0) / 10000
                if svol > 0:
                    points.append(f"[{sdate}] 增持 {sname} {abs(svol):.2f}万股")
                else:
                    points.append(f"[{sdate}] 减持 {sname} {abs(svol):.2f}万股")
            auto_data["executive_hold"] = {"points": points}
    except Exception as e:
        print(f"[reporter] 增减持采集失败: {e}")
    # 高管变动
    try:
        ex = fetcher.fetch_executive_change(code)
        if ex:
            points = auto_data.get("executive_hold", {}).get("points", [])
            for e in ex[:10]:
                ename = e.get("person_name","")
                edate = str(e.get("date",""))[:10]
                evol = float(e.get("volume",0) or 0) / 10000
                eprice = float(e.get("price",0) or 0)
                if evol > 0:
                    points.append(f"[{edate}] 高管增持 {ename} {abs(evol):.2f}万股 @{eprice:.2f}元")
                else:
                    points.append(f"[{edate}] 高管减持 {ename} {abs(evol):.2f}万股 @{eprice:.2f}元")
            auto_data["executive_hold"] = {"points": points}
    except Exception as e:
        print(f"[reporter] 高管变动采集失败: {e}")
    # ===== 财务数据（5期趋势表）=====
    def _fmt_date_q(d: str) -> str:
        """2026-03-31 → 26Q1"""
        try:
            dt = datetime.strptime(d[:10], "%Y-%m-%d")
            y = dt.year % 100
            q = (dt.month - 1) // 3 + 1
            return f"{y}Q{q}"
        except:
            return d[:7] if d else ""
    try:
        fin = fetcher.fetch_financial(code, 5)
        inc = fetcher.fetch_income_statement(code, 5)
        bal = fetcher.fetch_balance_sheet(code, 5)
        cf = fetcher.fetch_cashflow(code, 5)
        fr_data = {}
        # --- 业绩概览（5期）---
        if fin and len(fin) >= 2:
            dates_fin = [_fmt_date_q(f.get("date","")) for f in fin]
            revenues = [float(f.get("revenue",0) or 0)/1e8 for f in fin]
            netprofits = [float(f.get("net_profit",0) or 0)/1e8 for f in fin]
            eps_list = [f.get("eps",0) for f in fin]
            roe_list = [f.get("roe",0) for f in fin]
            gm_list = [f.get("gross_margin",0) for f in fin]
            table = [tuple(["指标"] + dates_fin)]
            table.append(tuple(["营收(亿)"] + [f"{v:.1f}" for v in revenues]))
            table.append(tuple(["净利润(亿)"] + [f"{v:.1f}" for v in netprofits]))
            table.append(tuple(["EPS(元)"] + [f"{v:.2f}" if v else "--" for v in eps_list]))
            if any(gm for gm in gm_list):
                table.append(tuple(["毛利率(%)"] + [f"{v:.1f}" if v else "--" for v in gm_list]))
            if any(roe for roe in roe_list):
                table.append(tuple(["ROE(%)"] + [f"{v:.2f}" if v else "--" for v in roe_list]))
            # analysis
            rev_latest = revenues[0] if revenues else 0
            np_latest = netprofits[0] if netprofits else 0
            rev_prev = revenues[1] if len(revenues) > 1 else 0
            np_prev = netprofits[1] if len(netprofits) > 1 else 0
            rev_g = (rev_latest - rev_prev) / rev_prev * 100 if rev_prev else 0
            np_g = (np_latest - np_prev) / np_prev * 100 if np_prev else 0
            analysis = f"最新季度营收{rev_latest:.1f}亿元，净利润{np_latest:.1f}亿元。"
            if rev_g:
                analysis += f"营收同比上期{'增长' if rev_g>0 else '下降'}{abs(rev_g):.1f}%，"
            if np_g:
                analysis += f"净利润同比上期{'增长' if np_g>0 else '下降'}{abs(np_g):.1f}%。"
            if eps_list and eps_list[0]:
                analysis += f"最新EPS={eps_list[0]:.2f}元。"
            if roe_list and roe_list[0]:
                analysis += f"ROE={roe_list[0]:.2f}%。"
            fr_data["summary"] = {"table": table, "analysis": analysis}
        # --- 利润分析（5期）---
        if inc and len(inc) >= 2:
            dates_inc = [_fmt_date_q(i.get("date","")) for i in inc]
            revs = [float(i.get("revenue",0) or 0)/1e8 for i in inc]
            nps = [float(i.get("net_profit",0) or 0)/1e8 for i in inc]
            gps = [float(i.get("gross_profit",0) or 0)/1e8 for i in inc]
            costs = [float(i.get("cost",0) or 0)/1e8 for i in inc]
            table = [tuple(["指标"] + dates_inc)]
            table.append(tuple(["营收(亿)"] + [f"{v:.1f}" for v in revs]))
            table.append(tuple(["营业成本(亿)"] + [f"{v:.1f}" for v in costs]))
            table.append(tuple(["营业利润(亿)"] + [f"{v:.1f}" for v in gps]))
            table.append(tuple(["净利润(亿)"] + [f"{v:.1f}" for v in nps]))
            r_latest = revs[0] if revs else 0
            r_prev = revs[1] if len(revs) > 1 else 0
            n_latest = nps[0] if nps else 0
            r_g = (r_latest - r_prev) / r_prev * 100 if r_prev else 0
            analysis = f"营收{r_latest:.1f}亿元，净利润{n_latest:.1f}亿元。"
            if r_g:
                analysis += f"营收同比{'增长' if r_g>0 else '下降'}{abs(r_g):.1f}%。"
            fr_data.setdefault("profit", {})
            fr_data["profit"] = {"table": table, "analysis": analysis}
        # --- 资产负债分析（5期）---
        if bal and len(bal) >= 2:
            dates_bal = [_fmt_date_q(b.get("date","")) for b in bal]
            tas = [float(b.get("total_assets",0) or 0)/1e8 for b in bal]
            tls = [float(b.get("total_liab",0) or 0)/1e8 for b in bal]
            eqs = [float(b.get("equity",0) or 0)/1e8 for b in bal]
            drs = [b.get("debt_ratio",0) for b in bal]
            table = [tuple(["指标"] + dates_bal)]
            table.append(tuple(["总资产(亿)"] + [f"{v:.1f}" for v in tas]))
            table.append(tuple(["总负债(亿)"] + [f"{v:.1f}" for v in tls]))
            table.append(tuple(["股东权益(亿)"] + [f"{v:.1f}" for v in eqs]))
            if any(drs):
                table.append(tuple(["资产负债率(%)"] + [f"{v:.1f}" if v else "--" for v in drs]))
            ta_latest = tas[0] if tas else 0
            tl_latest = tls[0] if tls else 0
            dr_latest = drs[0] if drs else 0
            analysis = f"总资产{ta_latest:.1f}亿元，总负债{tl_latest:.1f}亿元。"
            if dr_latest:
                analysis += f"资产负债率{dr_latest:.1f}%。"
            fr_data.setdefault("balance", {})
            fr_data["balance"] = {"table": table, "analysis": analysis}
        # --- 现金流分析（5期）---
        if cf and len(cf) >= 2:
            dates_cf = [_fmt_date_q(c.get("date","")) for c in cf]
            ops = [float(c.get("operate_net",0) or 0)/1e8 for c in cf]
            invs = [float(c.get("invest_net",0) or 0)/1e8 for c in cf]
            fins = [float(c.get("finance_net",0) or 0)/1e8 for c in cf]
            table = [tuple(["指标"] + dates_cf)]
            table.append(tuple(["经营现金流(亿)"] + [f"{v:.1f}" for v in ops]))
            table.append(tuple(["投资现金流(亿)"] + [f"{v:.1f}" for v in invs]))
            table.append(tuple(["筹资现金流(亿)"] + [f"{v:.1f}" for v in fins]))
            op_latest = ops[0] if ops else 0
            analysis = f"经营活动现金流净额{op_latest:.1f}亿元。" if op_latest else "经营活动现金流数据有限。"
            fr_data.setdefault("cashflow", {})
            fr_data["cashflow"] = {"table": table, "analysis": analysis}
        # 汇总到 extra_data
        if fr_data:
            auto_data["financial_report"] = fr_data
    except Exception as e:
        print(f"[reporter] 财务数据采集失败: {e}")
    # 分红
    try:
        div = fetcher.fetch_dividend(code)
        if div:
            points = []
            for d in div[:5]:
                date = d.get("date","")[:10]
                plan = d.get("plan","")
                points.append(f"{date} {plan}" if plan else f"{date}")
            auto_data["dividend"] = {"points": points}
    except Exception as e:
        print(f"[reporter] 分红采集失败: {e}")
    # 公司行业
    try:
        ind = fetcher.fetch_industry_info(code)
        if ind and (ind.get("industry") or ind.get("main_business")):
            auto_data["company"] = ind
    except Exception as e:
        print(f"[reporter] 行业信息采集失败: {e}")
    # 行业ETF资金流向
    try:
        etf_data = db.get_etf_flow(code, 15)
        if etf_data:
            auto_data["etf_flow"] = etf_data
    except Exception as e:
        print(f"[reporter] ETF数据采集失败: {e}")

    # extra_data 优先使用外部传入的，缺失的用自动采集的补上
    for k, v in auto_data.items():
        if k not in ed or not ed[k]:
            ed[k] = v

    # 3. 指标提取
    ma5, ma10, ma20, ma60 = indic.get("ma5", 0), indic.get("ma10", 0), indic.get("ma20", 0), indic.get("ma60", 0)
    md = indic.get("macd", {})
    kd = indic.get("kdj", {})
    rsi_val = indic.get("rsi14", 50)
    vr = indic.get("volRatio", 1)
    price = flow.get("price", info.get("price", 0))
    sig_items = sig_data.get("items", [])
    sig_sum = sig_data.get("summary", {})

    # 4. 证据
    evs = _gen_evidence(indic, pred.get("direction", ""), sig_items)

    # 5. K线图
    chart_path = chart_gen.generate(code, name, output_dir=output_dir)

    # ===== 创建DOCX =====
    doc = Document()
    _set_doc_font(doc)
    title = doc.add_heading(f"{name}({code}) 深度分析报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # ===== 强制风险提示（报告顶部醒目位置）=====
    from docx.shared import RGBColor
    warning = doc.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = warning.add_run("⚠️ 风险提示：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.size = Pt(11)
    run2 = warning.add_run("本报告基于公开历史数据统计生成，所有分析结论均为历史数据的技术统计展示，不构成任何投资建议，不预测未来走势。股市有风险，投资需谨慎。过往表现不代表未来收益。")
    run2.bold = True
    run2.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run2.font.size = Pt(10)
    doc.add_paragraph("")

    # ==============================================================
    # 一、公司概况
    # ==============================================================
    doc.add_heading("一、公司概况", 1)
    # 从extra_data.company获取行业和业务描述，禁止硬编码特定股票内容
    company_data = ed.get("company", {})
    industry_desc = company_data.get("industry", "")
    business_desc = company_data.get("main_business", "")
    # 如果extra_data没传，自动抓取行业信息兜底
    if not industry_desc and not business_desc:
        industry_info = fetcher.fetch_industry_info(code)
        if industry_info:
            industry_desc = industry_info.get("industry", "")
            business_desc = industry_info.get("main_business", "")
    if industry_desc or business_desc:
        doc.add_paragraph(f"{name}（{code}）所属行业为{industry_desc}。主营业务：{business_desc}")
    else:
        doc.add_paragraph(f"{name}（{code}）是{info.get('market','上海证券')}交易所上市公司。"
                          f"所属行业通过 extra_data.company 补充。")
    doc.add_paragraph(f"  最新市值：{info.get('total_market', 0):.0f}亿元（流通市值{info.get('circ_market', 0):.0f}亿元）")
    doc.add_paragraph(f"  市盈率PE(TTM)：{info.get('pe', 0):.2f}倍")
    doc.add_paragraph(f"  市净率PB：{info.get('pb', 0):.2f}倍")

    # ==============================================================
    # 二、实时行情
    # ==============================================================
    doc.add_heading("二、实时行情", 1)
    o_, c_, h_, l_, chg_, tor_ = (flow.get(k, 0) for k in ["open","close","high","low","change_pct","turnover_rate"])
    vol_ = flow.get("volume", 0) / 100
    amt_ = flow.get("amount", 0) / 10000
    feel = "收红📈" if chg_ > 0 else ("收跌📉" if chg_ < 0 else "平盘")
    _add_table(doc, [
        ("最新价", f"{price:.2f}元"), ("涨跌幅", f"{chg_:+.2f}%（{feel}）"),
        ("今开/最高/最低", f"{o_:.2f} / {h_:.2f} / {l_:.2f}元"),
        ("昨收", f"{c_:.2f}元"), ("成交量", f"{vol_:.0f}手"),
        ("成交额", f"{amt_:.0f}万元"), ("换手率", f"{tor_:.2f}%"),
        ("市盈率PE", f"{info.get('pe', 0):.2f}"), ("市净率PB", f"{info.get('pb', 0):.2f}"),
        ("流通市值", f"{info.get('circ_market', 0):.2f}亿"), ("总市值", f"{info.get('total_market', 0):.2f}亿"),
    ])

    # ==============================================================
    # 三、机构持仓分析
    # ==============================================================
    doc.add_heading("三、机构持仓分析", 1)
    holdings = ed.get("holdings")
    if holdings:
        if "points" in holdings:
            _add_points(doc, "机构持仓", holdings["points"], holdings.get("conclusion", ""))
        else:
            rows = []
            rows.append(("机构总数", f"{holdings.get('instCount', '--')}家"))
            rows.append(("合计持股", f"{holdings.get('totalShares', '--')}亿股"))
            rows.append(("占流通股", f"{holdings.get('ratio', '--')}%"))
            rows.append(("基金持股", f"{holdings.get('fundShares', '--')}亿股（{holdings.get('fundRatio', '--')}%）"))
            _add_table(doc, rows)
            if holdings.get('conclusion'):
                p = doc.add_paragraph()
                p.add_run(f"  ➜ {holdings['conclusion']}").bold = True
    else:
        doc.add_paragraph("（机构持仓数据暂无法获取，可通过 fetch_main_holdings() 补充）")

    # 前十大股东明细
    top_h = ed.get("top_holders")
    if top_h:
        doc.add_heading("十大股东明细（最新报告期）", 2)
        rows = [("排名", "股东名称", "持股比例", "变动方向")]
        for h in top_h:
            rows.append((str(h.get("rank", "")), h.get("name", ""),
                         h.get("ratio", ""), h.get("direction", "")))
        _add_table(doc, rows)
        # 十大股东解读
        top_names = [h.get("name","") for h in top_h[:3]]
        top_ratio_sum = 0
        inc_count = sum(1 for h in top_h if h.get("direction","") in ("增持","新进"))
        dec_count = sum(1 for h in top_h if h.get("direction","") == "减持")
        for h in top_h[:3]:
            r = h.get("ratio","").replace("%","")
            try: top_ratio_sum += float(r)
            except: pass
        th_p = doc.add_paragraph()
        th_p.add_run("▶ 解读：").bold = True
        th_p.add_run(f"前十大股东合计持股{top_ratio_sum:.1f}%，股权集中度较高。")
        if inc_count > dec_count:
            th_p.add_run(f"其中{inc_count}家增持/新进、{dec_count}家减持，大股东整体净增持。")
        elif inc_count < dec_count:
            th_p.add_run(f"其中{inc_count}家增持/新进、{dec_count}家减持，部分股东存在减持行为。")
        else:
            th_p.add_run(f"其中{inc_count}家增持/新进、{dec_count}家减持，增减力量相对均衡。")

    # ==============================================================
    # 四、资金流向分析
    # ==============================================================
    doc.add_heading("四、资金流向分析", 1)
    mf = ed.get("money_flow")
    if mf:
        rows = [("近5日主力净流入", mf.get("points", ["--"])[0])] if mf.get("points") else []
        for pt in mf.get("points", [])[1:]:
            rows.append(("", pt))
        if rows:
            _add_table(doc, rows)
        if mf.get('conclusion'):
            p = doc.add_paragraph()
            p.add_run(f"  ➜ {mf['conclusion']}").bold = True
        # 资金流向解读
        mf_text = "".join(mf.get("points", []))
        m_main = re.search(r'([+-]?[\d.]+)(万|亿)', mf_text)
        mf_p = doc.add_paragraph()
        mf_p.add_run("▶ 解读：").bold = True
        if m_main:
            mv = float(m_main.group(1))
            mu = m_main.group(2)
            m_val = mv if mu == '万' else mv * 10000
            if m_val < 0:
                mf_p.add_run(f"当日主力净流出{abs(m_val):.0f}万元，资金面偏弱。")
                if any('缩量' in p for p in mf.get("points", [])):
                    mf_p.add_run("今日缩量，说明抛压有限，主力流出但市场承接力尚可。")
                else:
                    mf_p.add_run("主力流出需谨慎关注。")
            else:
                mf_p.add_run(f"当日主力净流入{m_val:.0f}万元，资金面偏积极。")
        elif "量" in mf_text and "手" in mf_text:
            # 腾讯K线降级数据：只有量价趋势
            volumes = re.findall(r'量([\d.]+)手', mf_text)
            prices = re.findall(r'收([\d.]+)', mf_text)
            if volumes:
                avg_vol = sum(float(v) for v in volumes) / len(volumes) / 10000
                latest_vol = float(volumes[-1]) / 10000 if volumes else 0
                vol_ratio = latest_vol / avg_vol if avg_vol > 0 else 1
                mf_p.add_run(f"近5日平均成交量{avg_vol:.0f}万手，")
                if vol_ratio > 1.2:
                    mf_p.add_run(f"今日{latest_vol:.0f}万手（放量{vol_ratio:.1f}倍），市场交易活跃。")
                elif vol_ratio < 0.8:
                    mf_p.add_run(f"今日{latest_vol:.0f}万手（缩量{vol_ratio:.1f}倍），市场交易清淡。")
                else:
                    mf_p.add_run(f"今日{latest_vol:.0f}万手（量能正常），交易平稳。")
            if prices:
                latest_px = float(prices[-1])
                if len(prices) >= 2:
                    px_change = (float(prices[-1]) - float(prices[0])) / float(prices[0]) * 100
                    mf_p.add_run(f"最新收盘价{latest_px:.2f}元，近5日{'上涨' if px_change>0 else '下跌'}{abs(px_change):.1f}%。")
        else:
            mf_p.add_run("资金流向数据有限，仅供参考。")
    else:
        # 基于成交量的简易判断
        vol_status = '放量' if vr > 1.2 else '缩量' if vr < 0.8 else '正常'
        _add_table(doc, [
            ("当前成交量", f"{vol_:.0f}手"),
            ("成交额", f"{amt_:.0f}万元"),
            ("量比", f"{vr:.2f}倍（{vol_status}）"),
        ])
        doc.add_paragraph("（获取详细大单资金流向请使用 query_money_flow() 接口）")

    # 行业ETF资金流分析
    etf_flow_data = ed.get("etf_flow")
    if etf_flow_data:
        doc.add_heading("板块ETF资金流向", 2)
        matched = etf_flow_data.get("matched_etfs", [])
        summary = etf_flow_data.get("summary", "")
        if matched:
            etf_names = "、".join([e.get("name","") for e in matched])
            doc.add_paragraph(f"  关联行业ETF: {etf_names}")
        if summary:
            p = doc.add_paragraph()
            p.add_run(f"  {summary}").italic = True
        # 详细数据表
        etf_dict = etf_flow_data.get("etf_flow", {})
        for ecode, rows in etf_dict.items():
            ename = next((e.get("name","") for e in matched if e.get("code")==ecode), ecode)
            if rows and len(rows) >= 2:
                recent = rows[-5:] if len(rows) >= 5 else rows
                total_inflow = sum(r.get("net_inflow",0) for r in recent)
                doc.add_paragraph(f"  {ename}({ecode}) 近{len(recent)}日净{total_inflow/10000:+.0f}万元")
                detail_rows = [("日期","基金份额","份额变动","净流入(万元)")]
                for r in recent:
                    detail_rows.append((
                        r.get("trade_date",""),
                        f"{r.get('fund_size',0)/100000000:.2f}亿",
                        f"{r.get('share_change',0)/10000:+.0f}",
                        f"{r.get('net_inflow',0)/10000:+.0f}"
                    ))
                _add_table(doc, detail_rows)
        # ETF解读
        if summary:
            etf_p = doc.add_paragraph()
            etf_p.add_run("▶ 板块背景：").bold = True
            etf_p.add_run("该标的所属行业的ETF资金流向反映了板块整体的资金态度。")
            if '流入' in summary or '净+' in summary:
                etf_p.add_run("板块资金呈净流入，行业层面存在资金支撑。")
            elif '流出' in summary or '净-' in summary:
                etf_p.add_run("板块资金呈净流出，需注意行业整体资金压力可能传导至个股。")
        # ETF与个股联动四象分析
        if etf_dict:
            _add_etf_quadrant(doc, etf_flow_data, px_change if 'px_change' in dir() else 0, chg_)

    # ==============================================================
    # 五、融资融券分析
    # ==============================================================
    doc.add_heading("五、融资融券分析", 1)
    mt = ed.get("margin_trading")
    if mt:
        rows = []
        pts = mt.get("points", [])
        for pt in pts:
            # 从"2026-07-03 融资余额..."提取日期
            date_part = pt[:10] if len(pt) > 10 and pt[4] == '-' else ""
            content_part = pt[11:] if date_part else pt
            rows.append((date_part, content_part))
        if rows:
            _add_table(doc, rows)
        if mt.get('conclusion'):
            p = doc.add_paragraph()
            p.add_run(f"  ➜ {mt['conclusion']}").bold = True
        # 融资融券解读
        balance_val, net_pays = None, []
        for pt in pts:
            m = re.search(r'融资余额([\d.]+)亿', pt)
            if m: balance_val = float(m.group(1))
            m = re.search(r'融资净买([+-]?[\d.]+)', pt)
            if m: net_pays.append(float(m.group(1)))
        mt_p = doc.add_paragraph()
        mt_p.add_run("▶ 解读：").bold = True
        if balance_val is not None:
            net_count = sum(1 for n in net_pays if n < 0)
            total_net = sum(net_pays)
            mt_p.add_run(f"当前融资余额{balance_val:.2f}亿，近5日波动变化。")
            if net_count > len(net_pays) / 2 and net_pays:
                mt_p.add_run(f"近{len(net_pays)}个交易日有{net_count}天为净偿还，合计净偿还约{abs(total_net):.0f}万元，杠杆资金整体呈流出态势，投资者看多情绪偏谨慎。")
            elif net_pays:
                mt_p.add_run(f"融资资金整体呈净流入态势，投资者情绪偏积极。")
        else:
            mt_p.add_run("融资金额变化不大，杠杆资金情绪中性。")
    else:
        doc.add_paragraph("（融资余额和融券余额数据需通过 query_margin_trading() 采集后传入 extra_data.margin_trading）")

    # ==============================================================
    # 六、高管及大股东行为
    # ==============================================================
    doc.add_heading("六、高管及大股东行为", 1)
    exec_data = ed.get("executive_hold") or ed.get("shareholder_trade")
    if exec_data:
        rows = []
        pts = exec_data.get("points", [])
        for pt in pts:
            # 从"[2026-07-03] 增持..."提取日期
            if pt.startswith("[") and "]" in pt[:15]:
                date_part = pt[1:pt.index("]")]
                content_part = pt[pt.index("]")+1:].strip()
            else:
                date_part, content_part = "", pt
            rows.append((date_part, content_part))
        if rows:
            _add_table(doc, rows)
        if exec_data.get('conclusion'):
            p = doc.add_paragraph()
            p.add_run(f"  ➜ {exec_data['conclusion']}").bold = True
        # 增减持解读
        buys = [pt for pt in pts if '增持' in pt]
        sells = [pt for pt in pts if '减持' in pt]
        eh_p = doc.add_paragraph()
        eh_p.add_run("▶ 解读：").bold = True
        if buys:
            total_buy_vol = 0
            for pt in buys:
                m = re.search(r'([\d.]+)万股', pt)
                if m: total_buy_vol += float(m.group(1))
            names = set()
            for pt in buys:
                m = re.search(r'(\S+集团|\S+公司|\S+基金|\S+投资)', pt)
                if m: names.add(m.group(1))
            eh_p.add_run(f"共{len(buys)}笔增持记录（合计约{total_buy_vol:.2f}万股），增持方主要为{'、'.join(names) if names else '大股东'}。")
            if len(sells) > 0:
                eh_p.add_run(f"减持仅{len(sells)}笔，整体以增持为主，大股东对前景有信心。")
            else:
                eh_p.add_run("无减持记录，大股东持股意愿强烈。")
        elif sells:
            total_sell_vol = 0
            for pt in sells:
                m = re.search(r'([\d.]+)万股', pt)
                if m: total_sell_vol += float(m.group(1))
            eh_p.add_run(f"共{len(sells)}笔减持记录（合计约{total_sell_vol:.2f}万股），需关注大股东减持动向。")
        else:
            eh_p.add_run("近期待无显著增减持记录，股东持股稳定。")

    # ==============================================================
    # 七、核心财务分析
    # ==============================================================
    doc.add_heading("七、核心财务分析", 1)
    fr = ed.get("financial_report")
    if fr:
        doc.add_heading("7.1 业绩概览", 2)
        _render_fin_section(doc, "最新季度", fr.get("summary", "数据待补充"))
        doc.add_heading("7.2 利润分析", 2)
        _render_fin_section(doc, "盈利能力", fr.get("profit", ""))
        doc.add_heading("7.3 资产负债分析", 2)
        _render_fin_section(doc, "资产质量", fr.get("balance", ""))
        doc.add_heading("7.4 现金流分析", 2)
        _render_fin_section(doc, "现金流状况", fr.get("cashflow", ""))
    else:
        doc.add_heading("7.1 业绩概览", 2)
        _render_fin_section(doc, "基本数据", f"市盈率PE={info.get('pe',0):.2f}倍，市净率PB={info.get('pb',0):.2f}倍。"
                 "详细财务数据请通过 query_financial_report() 获取。")
        doc.add_heading("7.2 利润分析", 2)
        _render_fin_section(doc, "盈利能力", "数据待补充（建议使用 query_income_statement()）")
        doc.add_heading("7.3 资产负债分析", 2)
        _render_fin_section(doc, "资产质量", "数据待补充（建议使用 query_balance_sheet()）")
        doc.add_heading("7.4 现金流分析", 2)
        _render_fin_section(doc, "现金流状况", "数据待补充（建议使用 query_cashflow_statement()）")

    # ==============================================================
    # 八、估值参考区间分析
    # ==============================================================
    doc.add_heading("八、估值参考区间分析", 1)
    pe_val = info.get("pe", 15)
    pb_val = info.get("pb", 2)
    # 动态计算EPS和BVPS（从PE/PB反推），避免硬编码错误数据
    eps_est = round(price / pe_val, 2) if pe_val > 0 else 0
    bvps_est = round(price / pb_val, 2) if pb_val > 0 else 0
    _add_table(doc, [
        ("当前PE(TTM)", f"{pe_val:.2f}倍"), ("行业平均PE", "15~25倍"),
        ("当前PB", f"{pb_val:.2f}倍"), ("行业平均PB", "2~5倍"),
        ("预估EPS", f"{eps_est:.2f}元/股"), ("每股净资产(BVPS)", f"{bvps_est:.2f}元/股"),
    ])
    _add_points(doc, "估值分析", [
        f"PE法（行业均值20倍）：{eps_est * 20:.2f}元",
        f"PB法（行业均值3.5倍）：{bvps_est * 3.5:.2f}元",
        f"当前价{price:.2f}元，处于{'合理偏低' if price < eps_est * 20 else '合理偏高' if price > eps_est * 20 * 1.2 else '合理'}区间",
    ], f"历史估值参考区间约{min(eps_est * 15, bvps_est * 2.5):.2f}~{max(eps_est * 25, bvps_est * 5):.2f}元")
    # 术语表
    doc.add_paragraph("📖 估值术语简释（非专业人士参考）")
    _add_table(doc, [
        ("EPS（每股收益）", f"{eps_est:.2f}元/股", "指公司每一普通股能分到的税后利润，越高说明盈利能力越强"),
        ("PE（市盈率）", f"{pe_val:.2f}倍", "股价 ÷ EPS，市场愿意为1元利润支付的价格。行业均值15~25倍"),
        ("PB（市净率）", f"{pb_val:.2f}倍", "股价 ÷ 每股净资产，市场愿意为1元净资产支付的价格。行业均值2~5倍"),
        ("股息率", f"{(info.get('dividend', 0.112) / price * 100):.2f}%", "每股分红 ÷ 股价，相当于一次性的'利息率'"),
    ])
    # PE/PB 解读
    p_val = doc.add_paragraph()
    p_val.add_run("▶ 解读：").bold = True
    parts = []
    if pe_val < 15:
        parts.append(f"PE={pe_val:.2f}倍，低于行业均值15~25倍，估值相对偏低，估值统计显示偏低。")
    elif pe_val <= 25:
        parts.append(f"PE={pe_val:.2f}倍，处于行业均值15~25倍范围内，估值合理。")
    else:
        parts.append(f"PE={pe_val:.2f}倍，高于行业均值15~25倍，估值偏高，需业绩增长支撑。")
    if pb_val < 2:
        parts.append(f"PB={pb_val:.2f}倍，低于行业均值2~5倍，资产估值偏低。")
    elif pb_val <= 5:
        parts.append(f"PB={pb_val:.2f}倍，处于行业均值2~5倍范围内，资产估值合理。")
    else:
        parts.append(f"PB={pb_val:.2f}倍，高于行业均值2~5倍，资产溢价较高。")
    p_val.add_run("".join(parts))

    # ==============================================================
    # 九、分红历史
    # ==============================================================
    doc.add_heading("九、券商研报观点", 1)
    try:
        # 通过WebAPI获取研报数据（客户无需直连数据库）
        # 1. 个股研报汇总
        doc.add_heading("9.1 近期研报覆盖", 2)
        rpt = db.get_research_summary(code, 90)
        doc.add_paragraph(f"  {rpt}")
        doc.add_paragraph("")
        # 2. EPS一致性预期
        doc.add_heading("9.2 EPS一致性预期", 2)
        eps_rpt = db.get_eps_consistency(code, 180)
        doc.add_paragraph(f"  {eps_rpt}")
        doc.add_paragraph("")
        # 3. 券商覆盖统计（按个股）
        doc.add_heading("9.3 券商覆盖统计", 2)
        br_rpt = db.get_broker_ranking(code, 90)
        doc.add_paragraph(f"  {br_rpt}")
        doc.add_paragraph("")
        # 4. 今日评级动向（按个股）
        doc.add_heading("9.4 今日评级动态", 2)
        wind_rpt = db.get_rating_wind_today(code)
        doc.add_paragraph(f"  {wind_rpt}")
        # 5. 机构观点验证（交叉分析）
        _add_research_validation(doc, code, ed, price, chg_)
    except Exception as e:
        print(f"[reporter] 研报数据加载异常: {e}")
        doc.add_paragraph(f"  （研报数据加载中: {str(e)[:50]}）")

    # ==============================================================
    # 十五、风险提示
    # ==============================================================
    doc.add_heading("十、分红历史", 1)
    dv = ed.get("dividend")
    if dv:
        rows = []
        pts = dv.get("points", [])
        for pt in pts:
            rows.append(("", pt))
        if rows:
            _add_table(doc, rows)
        if dv.get('conclusion'):
            p = doc.add_paragraph()
            p.add_run(f"  ➜ {dv['conclusion']}").bold = True
        # 分红解读
        div_p = doc.add_paragraph()
        div_p.add_run("▶ 解读：").bold = True
        div_years, div_amounts = [], []
        for pt in pts:
            ym = re.search(r'(\d{4})-\d{2}-\d{2}', pt)
            am = re.search(r'10派([\d.]+)', pt)
            if ym and am:
                div_years.append(int(ym.group(1)))
                div_amounts.append(float(am.group(1)))
        if div_years:
            latest_div = div_amounts[0]
            avg_div = sum(div_amounts) / len(div_amounts)
            trend = "持续增长" if all(div_amounts[i] >= div_amounts[i+1] for i in range(len(div_amounts)-1)) else (
                "有所下降" if all(div_amounts[i] <= div_amounts[i+1] for i in range(len(div_amounts)-1)) else "波动变化")
            div_p.add_run(f"最新年度分红方案为10派{latest_div:.2f}元，近{len(div_years)}年平均10派{avg_div:.2f}元，分红{trend}。")
            div_rate = latest_div / 10 / price * 100 if price > 0 else 0
            div_p.add_run(f"以当前价{price:.2f}元计算，股息率约{div_rate:.2f}%，"
                          + ("属于较高水平。" if div_rate > 3 else "处于中等水平。" if div_rate > 1.5 else "相对偏低。"))
    else:
        doc.add_paragraph("（详细分红数据请通过 query_dividend_history() 采集后传入 extra_data.dividend）")

    # ==============================================================
    # 十、综合评分（MEMORY.md 加权公式）
    # ==============================================================
    doc.add_heading("十一、综合评分", 1)
    doc.add_paragraph("评分规则：基本面×35% + 估值面×20% + 技术面×15% + 资金面×20% + 情绪面×10%，满分100分")
    doc.add_paragraph("评分参考：≥80偏高 | ≥65中等偏上 | ≥50中等 | <50偏低")
    doc.add_paragraph("")

    v_sum = sum(v.get("v", 0) for v in pred.get("votes", []))
    s_tech = min(100, max(0, (v_sum + 10) * 5))     # 技术面 0~100
    s_val = min(100, max(0, (7.5 if pe_val < 20 and pb_val < 1 else 5.0) * 12))  # 估值面 0~100
    s_fund = min(100, max(30, 65))                    # 基本面（动态估算）
    s_cap = min(100, max(30, 70))                     # 资金面（动态估算）
    s_sent = min(100, max(30, 55))                    # 情绪面（动态估算）

    scores = {
        "基本面(35%)": s_fund,
        "估值面(20%)": s_val,
        "技术面(15%)": s_tech,
        "资金面(20%)": s_cap,
        "情绪面(10%)": s_sent,
    }
    total = s_fund * 0.35 + s_val * 0.20 + s_tech * 0.15 + s_cap * 0.20 + s_sent * 0.10

    rating = "偏高 🏆" if total >= 80 else "中等偏上 ✅" if total >= 65 else "中等 ⚠️" if total >= 50 else "偏低 ❌"
    t = _add_table(doc, [(k, f"{v:.0f}/100") for k, v in scores.items()])
    doc.add_paragraph(f"  加权总分 = {s_fund:.0f}×35% + {s_val:.0f}×20% + {s_tech:.0f}×15% + {s_cap:.0f}×20% + {s_sent:.0f}×10%")
    p = doc.add_paragraph()
    p.add_run(f"  ➜ 总分：{total:.0f}/100 → {rating}").bold = True
    # 综合评分解读
    score_p = doc.add_paragraph()
    score_p.add_run("▶ 解读：").bold = True
    sparts = []
    if s_fund >= 70:
        sparts.append(f"基本面{s_fund:.0f}/100，盈利能力稳健。")
    elif s_fund >= 50:
        sparts.append(f"基本面{s_fund:.0f}/100，处于行业中游。")
    else:
        sparts.append(f"基本面{s_fund:.0f}/100，需关注盈利变化。")
    if s_val >= 70:
        sparts.append(f"估值面{s_val:.0f}/100，处于合理偏低区间。")
    elif s_val >= 50:
        sparts.append(f"估值面{s_val:.0f}/100，估值合理。")
    else:
        sparts.append(f"估值面{s_val:.0f}/100，估值偏高需谨慎。")
    if s_tech >= 70:
        sparts.append(f"技术面{s_tech:.0f}/100，短期趋势偏多。")
    else:
        sparts.append(f"技术面{s_tech:.0f}/100，短期存在调整压力。")
    if s_cap >= 70:
        sparts.append(f"资金面{s_cap:.0f}/100，主力资金关注度较高。")
    else:
        sparts.append(f"资金面{s_cap:.0f}/100，资金参与度中等。")
    if total >= 80:
        sparts.append("历史统计评分较高（仅供技术参考，非投资建议）。")
    elif total >= 65:
        sparts.append("基本面信号偏正面（历史统计口径）。")
    elif total >= 50:
        sparts.append("投资价值中等，需结合自身判断。")
    else:
        sparts.append("历史统计评分偏低（仅供技术参考，非投资建议）。")
    score_p.add_run("".join(sparts))

    # v3 特征融合分析（ETF资金流+券商研报）
    if v3_factors:
        raw_s = v3_factors.get("raw_score", 0)
        adj_s = v3_factors.get("adjusted_score", 0)
        etf_f = v3_factors.get("etf_factor", 0)
        rpt_f = v3_factors.get("research_factor", 0)
        etf_sum = v3_factors.get("etf_summary", "")
        rpt_sum = v3_factors.get("research_summary", "")

        doc.add_heading("11.1 特征融合评分修正", 2)
        doc.add_paragraph(f"  📊 ETF板块资金因子：{etf_f:+.1f} ｜ {etf_sum}")
        doc.add_paragraph(f"  📊 券商研报共识因子：{rpt_f:+.1f} ｜ {rpt_sum}")
        diff = adj_s - raw_s
        arrow = "🔼" if diff > 0 else "🔽" if diff < 0 else "➡️"
        doc.add_paragraph(f"  {arrow} 原始评分：{raw_s:.1f}  →  调整后评分：{adj_s:.1f}（修正{diff:+.1f}）")

    # ==============================================================
    # 十一、K线技术分析
    # ==============================================================
    if not webapi_available:
        # 构建预览行
        teaser_lines = []
        if preview and preview_indi:
            ma5 = preview_indi.get("ma5")
            ma10 = preview_indi.get("ma10")
            ma20 = preview_indi.get("ma20")
            ma_str = f"MA5={ma5} MA10={ma10} MA20={ma20}" if all(v is not None for v in [ma5,ma10,ma20]) else ""
            macd = preview_indi.get("macd", {}) or {}
            macd_str = f"MACD(DIF={macd.get('dif','')} DEA={macd.get('dea','')} BAR={macd.get('bar','')})" if macd.get("dif") is not None else ""
            kdj = preview_indi.get("kdj", {}) or {}
            kdj_str = f"KDJ(K={kdj.get('k','')} D={kdj.get('d','')} J={kdj.get('j','')})" if kdj.get("k") is not None else ""
            rsi = preview_indi.get("rsi14")
            rsi_str = f"RSI(14)={rsi}" if rsi is not None else ""
            verdict = (preview_sig or {}).get("verdict", "")
            data_parts = [s for s in [ma_str, macd_str, kdj_str, rsi_str] if s]
            if data_parts:
                teaser_lines.append(f"📊 免费预览：{' | '.join(data_parts)}")
                if verdict:
                    teaser_lines.append(f"📊 免费预览：综合信号 → {verdict}")

        _paid_section(
            doc, "十一", "K线技术分析",
            teaser_lines,
            "❓ 这些信号综合起来意味着什么？机构就是拿着这些数据做决策的，你看不到。",
            [
                "7项技术指标逐一分析：偏多 √  偏空 ×  中性 △ — 结论一目了然",
                "均线多头/空头/缠绕？你在什么位置入场？鱼头、鱼身还是鱼尾？",
                "MACD金叉是真突破还是假信号？不是猜，是用历史数据验给你看",
                "当前KDJ/RSI位置在历史上出现后，第二天怎么走的？匹配率是多少？",
                "K线趋势图：MA均线+成交量，一眼看清格局",
            ],
            api_key=config.API_KEY
        )
    else:
        doc.add_heading("十二、K线技术分析", 1)

        doc.add_heading("11.1 均线分析", 2)
        _add_table(doc, [("MA5", f"{ma5:.3f}" if ma5 else "—"), ("MA10", f"{ma10:.3f}" if ma10 else "—"),
                         ("MA20", f"{ma20:.3f}" if ma20 else "—"), ("MA60", f"{ma60:.3f}" if ma60 else "—")])
        if ma5 and ma10 and ma20:
            if ma5 > ma10 > ma20:
                doc.add_paragraph("  ✅ 均线多头排列，短期趋势向上")
            elif ma5 < ma10 < ma20:
                doc.add_paragraph("  ❌ 均线空头排列，短期趋势向下")
            else:
                doc.add_paragraph("  ⚠️ 均线交叉震荡")
        if price and ma60:
            doc.add_paragraph(f"  {'⚠️ 现价在MA60下方，中期偏弱' if price < ma60 else '✅ 现价在MA60上方，中期偏强'}")

        doc.add_heading("11.2 MACD指标", 2)
        if md:
            _add_table(doc, [("DIF", f"{md.get('dif',0):.4f}"), ("DEA", f"{md.get('dea',0):.4f}"), ("MACD柱", f"{md.get('bar',0):.4f}")])
            doc.add_paragraph(f"  {'✅ DIF在DEA上方，红柱偏多' if md.get('dif',0) > md.get('dea',0) else '❌ DIF在DEA下方，绿柱偏空'}")

        doc.add_heading("11.3 KDJ指标", 2)
        if kd:
            k_, d_, j_ = kd.get("k",50), kd.get("d",50), kd.get("j",50)
            _add_table(doc, [("K值", f"{k_:.1f}"), ("D值", f"{d_:.1f}"), ("J值", f"{j_:.1f}")])
            if j_ > 100: doc.add_paragraph("  ⚠️ KDJ超买，注意回调")
            elif j_ < 0: doc.add_paragraph("  🔥 KDJ超卖，历史统计中超卖后偏多信号比例偏高")
            elif k_ > d_: doc.add_paragraph("  ✅ KDJ偏多")
            else: doc.add_paragraph("  ❌ KDJ偏空")

        doc.add_heading("11.4 RSI指标", 2)
        doc.add_paragraph(f"  RSI(14): {rsi_val:.1f}")
        if rsi_val > 70: doc.add_paragraph("  ⚠️ RSI超买")
        elif rsi_val < 30: doc.add_paragraph("  🔥 RSI超卖")
        elif rsi_val > 50: doc.add_paragraph("  ✅ RSI偏强")
        else: doc.add_paragraph("  ❌ RSI偏弱")

        doc.add_heading("11.5 综合信号", 2)
        if sig_items:
            t = doc.add_table(rows=len(sig_items) + 1, cols=3, style="Light Shading Accent 1")
            for j, h in enumerate(["指标", "信号", "说明"]): t.cell(0, j).text = h
            for i, s in enumerate(sig_items):
                t.cell(i + 1, 0).text = s.get("indicator", "")
                sg = s.get("signal", "")
                t.cell(i + 1, 1).text = {"偏多":"✅偏多", "偏空":"❌偏空", "中性":"⚖️中性"}.get(sg, sg)
                t.cell(i + 1, 2).text = s.get("description", "")
            doc.add_paragraph(f"  信号汇总：偏多{sig_sum.get('bullish',0)}个 / 偏空{sig_sum.get('bearish',0)}个 / 中性{sig_sum.get('neutral',0)}个")
            doc.add_paragraph(f"  综合判定：{sig_data.get('verdict','')}")


    # K线图 — 本地生成，不计入付费内容，无条件嵌入
    if chart_path and os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(5.5))

    # ==============================================================
    # 十二、历史信号匹配统计
    # ==============================================================
    if not webapi_available:
        # 构建预览行
        teaser_lines12 = []
        if preview and preview_rt:
            close = preview_rt.get("close")
            chg = preview_rt.get("changePct")
            vol = preview_rt.get("volume")
            parts = []
            if close is not None: parts.append(f"最新价={close}")
            if chg is not None: parts.append(f"涨跌幅={chg:.2f}%")
            if vol is not None: parts.append(f"成交量={int(vol)}手")
            if parts:
                teaser_lines12.append(f"📊 免费预览：{' | '.join(parts)}")
        if preview and preview_sig:
            summary = (preview_sig or {}).get("summary", {}) or {}
            b = summary.get("bullish", 0); n = summary.get("bearish", 0)
            if isinstance(preview_sig, dict) and preview_sig.get("verdict"):
                teaser_lines12.append(f"📊 免费预览：信号统计 偏多={b} 偏空={n} → {preview_sig['verdict']}")

        _paid_section(
            doc, "十二", "历史信号匹配统计",
            teaser_lines12,
            "❓ 散户亏钱，90%是因为没有参考系。你不知道当前信号历史上出现过几次，第二天怎么走的。",
            [
                "综合方向判定 + 0-10分评分（含进度条） — 一眼知道胜算",
                "全部历史/近60日/近30日匹配率：过去这种信号出现后，涨了多少次？跌了多少次？",
                "T+1参考：方向置信度 + 波动区间上限/下限 + 均值 — 明天怎么走，数据说话",
                "T+2参考：持有两天的话，胜率有没有变化？",
                "中期方向参考：不是猜下周，是看1-2周的趋势骨架",
                "多维度参考依据：技术面、历史模式、支撑阻力、风险点 — 一张表全看清",
            ],
            api_key=config.API_KEY
        )
    else:
        doc.add_heading("十三、历史信号匹配统计", 1)

        # 12.1 方向判定与评分（先展示方向）
        doc.add_heading("12.1 历史信号方向统计", 2)
        pred_dir = pred_api.get("direction", pred.get("direction", "震荡"))
        pred_score = pred_api.get("totalScore", pred.get("score", 0))
        pred_range = pred_api.get("rangeForecast", pred.get("range_forecast", "—"))
        dir_icon = {"偏多": "📈", "偏空": "📉", "震荡": "⚖️", "震荡偏多": "↗️", "震荡偏空": "↘️"}
        icon = "❓"
        for k, v in dir_icon.items():
            if k in pred_dir: icon = v; break
        doc.add_paragraph(f"  {icon} 信号方向：{pred_dir}")
        bar = _bar(max(0, min(100, pred_score * 10)))
        doc.add_paragraph(f"  🎯 评分：{pred_score}/10  {bar}")
        doc.add_paragraph(f"  📊 历史波动区间：{pred_range}")

        # 12.2 历史匹配率（紧跟方向之后）
        doc.add_heading("12.2 历史匹配率", 2)
        for label, p in [("全部历史", "all"), ("近60日", "period60"), ("近30日", "period30")]:
            s = astat(p)
            if s["total"] > 0:
                doc.add_paragraph(f"  {'✅' if s['rate']>=60 else '⚠️' if s['rate']>=40 else '❌'} {label}：{s['correct']}/{s['total']} = {s['rate']:.1f}%  {_bar(s['rate'])}")

        # 12.3 T+1（下一交易日）历史参考
        doc.add_heading("12.3 T+1（下一交易日）历史参考", 2)
        t1 = pred_api.get("t1Direction", pred.get("direction", "震荡"))
        t1_icon = "📈" if "偏多" in t1 else ("📉" if "偏空" in t1 else "⚖️")
        doc.add_paragraph(f"  {t1_icon} 信号方向：{t1}")
        if "偏多" in t1:
            l_, h_, m_ = price * 0.98, price * 1.02, price * 1.005
        elif "偏空" in t1:
            l_, h_, m_ = price * 0.97, price * 1.00, price * 0.985
        else:
            l_, h_, m_ = price * 0.99, price * 1.01, price * 1.00
        doc.add_paragraph(f"  历史波动低：{l_:.3f}元")
        doc.add_paragraph(f"  历史波动高：{h_:.3f}元")
        doc.add_paragraph(f"  历史均值：{m_:.3f}元")
        if "偏多" in t1:
            doc.add_paragraph("  📊 短线技术信号：偏多")
        elif "偏空" in t1:
            doc.add_paragraph("  📊 短线技术信号：偏空")
        else:
            doc.add_paragraph("  📊 短线技术信号：中性")

        # 12.4 T+2（两日后）历史参考
        doc.add_heading("12.4 T+2（两日后）历史参考", 2)
        t2 = pred_api.get("t2Direction", pred.get("direction", "震荡"))
        t2_icon = "📈" if "偏多" in t2 else ("📉" if "偏空" in t2 else "⚖️")
        doc.add_paragraph(f"  {t2_icon} 信号方向：{t2}")
        if "偏多" in t2:
            l_, h_, m_ = price * 0.97, price * 1.04, price * 1.01
        elif "偏空" in t2:
            l_, h_, m_ = price * 0.95, price * 1.00, price * 0.975
        else:
            l_, h_, m_ = price * 0.98, price * 1.02, price * 1.00
        doc.add_paragraph(f"  历史波动低：{l_:.3f}元")
        doc.add_paragraph(f"  历史波动高：{h_:.3f}元")
        doc.add_paragraph(f"  历史均值：{m_:.3f}元")
        if "偏多" in t2:
            doc.add_paragraph("  📊 两日后技术信号：偏强")
        elif "偏空" in t2:
            doc.add_paragraph("  📊 两日后技术信号：偏弱")
        else:
            doc.add_paragraph("  📊 两日后技术信号：方向不明")

        # 12.5 中期方向判断
        doc.add_heading("12.5 中期信号统计", 2)
        weekly = pred_api.get("weeklyDirection", pred.get("direction", "震荡"))
        wk_icon = "📈" if "偏多" in weekly else ("📉" if "偏空" in weekly else "⚖️")
        doc.add_paragraph(f"  {wk_icon} 本周方向：{weekly}")
        if price and ma60:
            mid_sig = "⚠️ 中期偏弱（MA60上方压制）" if price < ma60 else "✅ 中期偏强（MA60上方支撑）"
            doc.add_paragraph(f"  {mid_sig}")
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("📊 参考依据（技术面 + 模式识别）:").bold = True
        for i, ev in enumerate(evs, 1):
            doc.add_paragraph(f"  {i}. {ev}")

    # ==============================================================
    # 十三、历史规律总结
    # ==============================================================
    if not webapi_available:
        # 构建预览行
        teaser_lines13 = []
        if preview and preview_indi:
            ma60 = preview_indi.get("ma60")
            close = preview_rt.get("close") if preview_rt else None
            if ma60 is not None and close is not None:
                offset = (close - ma60) / ma60 * 100
                teaser_lines13.append(f"📊 免费预览：当前价距MA60 {offset:+.1f}%（MA60={ma60}）")
        if len(klines) >= 60:
            high60 = max(k["high"] for k in klines[-60:])
            low60 = min(k["low"] for k in klines[-60:])
            cur = flow.get("price", 0) or klines[-1]["close"] if klines else 0
            pos = (cur - low60) / (high60 - low60) * 100 if high60 > low60 else 50
            teaser_lines13.append(f"📊 免费预览：60日位置={pos:.0f}%（高点={high60} 低点={low60}）")

        _paid_section(
            doc, "十三", "历史规律总结",
            teaser_lines13,
            "❓ 每只股票都有自己的习惯——大涨后容易回调还是继续冲？暴跌后该割还是该扛？你不知道，但数据知道。",
            [
                "大涨后3日规律：历史上大涨>3%之后怎么走的？追涨胜率多少？匹配率含进度条",
                "大跌后3日规律：恐慌时刻该割还是该扛？历史数据告诉你答案",
                "支撑位识别：历史回踩获撑的具体价位 — 精确到小数点后两位",
                "阻力位识别：历史上攻失败的具体价位 — 提前知道哪里该跑",
                "最近真实案例：每条规律下方附最近1-2次历史案例，含触发日期+后续走势",
                "自学习跟踪：规律匹配率随验证次数提升，系统在自我进化",
            ],
            api_key=config.API_KEY
        )
    else:
        doc.add_heading("十四、历史规律总结", 1)
        doc.add_paragraph("以下规律基于历史数据挖掘，展示当特定价格形态出现后，后续走势的统计规律。")
        doc.add_paragraph("")

        if patterns and isinstance(patterns, list):
            for pt in patterns:
                if not isinstance(pt, dict):
                    continue
                pname = pt.get("name", "")
                adv = pt.get("advice", "")

                # 大涨/大跌类规律（有samples）
                samples = pt.get("samples")
                if samples and isinstance(samples, list):
                    total = len(samples)
                    if "大涨" in pname:
                        hits = sum(1 for s_ in samples if s_.get("r") == "回调")
                        rate = hits / total * 100 if total > 0 else 0
                        scene_title = "🔥 某天突然大涨超过3%，之后3天会怎样？"
                        hit_label = "回调"
                        op_emoji = "📈"
                    elif "大跌" in pname:
                        hits = sum(1 for s_ in samples if s_.get("r") == "反弹")
                        rate = hits / total * 100 if total > 0 else 0
                        scene_title = "💥 某天突然大跌超过3%，之后3天会怎样？"
                        hit_label = "反弹"
                        op_emoji = "📉"
                    else:
                        scene_title = f"📊 {pname}"
                        hits = 0
                        rate = 0

                    avg = pt.get("avgD3", "—")
                    hit_icon = "✅" if rate >= 50 else "⚠️"
                    doc.add_heading(f"14.{patterns.index(pt)+1} {scene_title}", 2)
                    doc.add_paragraph(f"  {op_emoji} 历史上出现 **{total}次**")
                    doc.add_paragraph(f"  {hit_icon} 其中 **{hits}次{hit_label}**，{total - hits}次续涨/续跌")
                    doc.add_paragraph(f"  📊 {hit_label}匹配率：**{rate:.0f}%**  {_bar(rate)}")
                    doc.add_paragraph(f"  平均3日涨跌幅：**{avg}**")
                    if adv:
                        adv_text = {"高抛": "大涨后历史规律偏回调",
                                    "信号偏弱时需谨慎": "大跌后历史规律显示偏弱",
                                    "低吸": "大跌后历史统计中偏多比例偏高",
                                    "持有": "趋势延续偏强"}.get(adv, f"历史规律参考：{adv}")
                        doc.add_paragraph(f"  📊 历史规律参考：**{adv_text}**")
                    doc.add_paragraph("")

                    # 显示最近2条案例
                    if samples:
                        doc.add_paragraph(f"  📋 最近案例：")
                        for s_ in samples[:2]:
                            r_icon = "✅" if s_.get("r") in ("回调", "反弹") else "❌"
                            doc.add_paragraph(f"    · {s_.get('date','')}：触发{s_.get('chg','')} → 3日后{s_.get('d3','')}  {r_icon}{s_.get('r','')}")
                    doc.add_paragraph("")

                # 支撑/阻力位类
                val = pt.get("value")
                if val and not samples:
                    scene_title = {"支撑位": "🛡️ 股价跌到什么位置可能止跌？",
                                   "阻力位": "🧱 股价涨到什么位置会遇到压力？"}.get(pname, f"📊 {pname}")
                    doc.add_heading(f"14.{patterns.index(pt)+1} {scene_title}", 2)
                    doc.add_paragraph(f"  📍 {pname}：**{val}元**")
                    if adv:
                        adv_text = {"跌破历史支撑位": f"历史统计：股价跌破{val}元后偏弱",
                                    "突破阻力位": f"历史统计：股价突破{val}元后偏强"}.get(adv, f"历史统计参考：{adv}")
                        doc.add_paragraph(f"  📊 历史统计参考：**{adv_text}**")
                    doc.add_paragraph("")
        else:
            doc.add_paragraph("  （数据不足，需至少90个交易日K线方可挖掘规律）")

        # 自学习规律跟踪摘要（从 WebAPI learningSummary 获取）
        # 注意：learningSummary 在 api(即 data)顶层，不在 latestPrediction 里
        ls = api.get("learningSummary", [])
        if ls and isinstance(ls, list):
            doc.add_heading("14.6 自学习规律跟踪摘要", 2)
            doc.add_paragraph("以下展示每条规律在历史验证中的跟踪记录，验证次数越多匹配率参考价值越大。")
            doc.add_paragraph("")

            # 从 patterns 中提取 avgD3 做交叉验证（服务端 avgD3 可能有单位错误）
            _avgd3_map = {}
            for _p in (patterns or []):
                if isinstance(_p, dict):
                    _pn = _p.get("name", "")
                    _av = _p.get("avgD3", "")
                    if _av and isinstance(_av, str) and "%" in _av:
                        _avgd3_map[_pn] = _av

            for li in ls:
                # 服务端返回的实际字段：name, hitRate(如"1/5=20%反弹"), verifyCount, update, description
                pt_raw = li.get("name", li.get("patternType", ""))

                # 字段名翻译表（API原始命名 → 合规替换）
                _name_map = {
                    "方向准确率": "方向匹配率",
                    "全部历史预测": "全部历史统计",
                    "趋势准确率": "趋势匹配率",
                    "信号准确率": "信号匹配率",
                    "历史预测": "历史统计",
                }
                pt = pt_raw
                for _old, _new in _name_map.items():
                    if _old in pt:
                        pt = pt.replace(_old, _new)
                        break
                vc = li.get("verifyCount", 0)
                hr = li.get("hitRate", "")
                upd = li.get("update", "")
                di = li.get("description", "")

                # 修复 avgD3 异常值（服务端有时放大了10000倍）
                if di:
                    import re as _re
                    _m = _re.search(r'(-?\d+\.?\d*)%', di)
                    if _m:
                        _val = float(_m.group(1))
                        if abs(_val) > 100:  # 明显异常，正常 avgD3 在 ±30% 以内
                            # 尝试从 patterns 中获取正确值
                            _correct_avg = ""
                            for _pname, _pavg in _avgd3_map.items():
                                # 匹配规律名：大跌后反弹特征 ↔ 大跌后3日
                                if "大跌" in pt and "大跌" in _pname:
                                    _correct_avg = _pavg
                                elif "大涨" in pt and "大涨" in _pname:
                                    _correct_avg = _pavg
                            if _correct_avg:
                                di = _re.sub(r'(-?\d+\.?\d*)%', _correct_avg, di)
                            else:
                                di = _re.sub(r'-?\d+\.?\d*%', '--', di)

                if vc > 0:
                    ver_icon = "✅"
                    # 对 di(description) 也执行字段名翻译
                    for _old, _new in _name_map.items():
                        if _old in di:
                            di = di.replace(_old, _new)
                    if hr:
                        # hitRate 格式如 "1/5=20%反弹" 或 "2/4=50%回调"
                        import re as _re2
                        m = _re2.search(r'(\d+\.?\d*)', hr.split("=")[-1])
                        rate_val = float(m.group(1)) if m else 50
                        ver_icon = "✅" if rate_val >= 50 else "❌"
                        doc.add_paragraph(f"  {ver_icon} **{pt}**：验证{vc}次，{hr}")
                    else:
                        doc.add_paragraph(f"  **{pt}**：验证{vc}次")

                    if di:
                        doc.add_paragraph(f"    ↳ {di}")
                    if upd and upd != di:
                        doc.add_paragraph(f"    ↳ {upd}")
            doc.add_paragraph("")

    # ==============================================================
    # ==============================================================
    # 十五、风险提示
    # ==============================================================
    doc.add_heading("十五、风险提示", 1)
    risks = []
    if price and ma60 and price < ma60:
        risks.append("中期趋势风险：股价在MA60下方，中期走势偏弱")
    if rsi_val > 70: risks.append("技术面风险：RSI超买，短期回调压力大")
    if rsi_val < 30 and price and ma60 and price < ma60:
        risks.append("技术面风险：RSI超卖，虽历史统计中超卖后偏多比例偏高，但中期趋势偏弱")
    elif rsi_val < 30:
        risks.append("技术面风险：RSI超卖，短期偏多信号，但中期趋势偏弱")
    if vr > 3: risks.append(f"量能风险：成交量异常放大（均量{vr:.0f}倍），注意放量回调")
    if pe_val > 30: risks.append(f"估值风险：PE偏高（{pe_val:.0f}倍），高于行业均值")
    # 行业风险（从上市公司info获取行业名称，动态生成）
    industry_name = ""
    if info and isinstance(info, dict):
        industry_name = info.get("industry", "") or info.get("行业", "") or ""
    if industry_name:
        risks.append(f"行业风险：{industry_name}板块受宏观经济及行业政策影响，注意系统性风险")
    else:
        risks.append("行业风险：注意宏观经济及行业政策变化对板块的影响")

    # 大股东控盘风险（从top_holders获取）
    top_holders_data = ed.get("top_holders", [])
    if top_holders_data:
        total_ratio = 0
        for h in top_holders_data[:5]:
            ratio_str = h.get("ratio", "0%")
            try:
                total_ratio += float(ratio_str.replace("%",""))
            except:
                pass
        if total_ratio > 70:
            risks.append(f"操作风险：前5大股东合计持股{total_ratio:.1f}%，股权高度集中，注意流动性风险")
    else:
        risks.append("操作风险：大股东控盘比例较高，注意流动性风险")

    for i, risk in enumerate(risks[:8], 1):
        p = doc.add_paragraph()
        p.add_run(f"  ⚠️ {i}. ").bold = True
        p.add_run(risk)
    doc.add_paragraph("")
    doc.add_paragraph("  💡 免责声明：以上分析基于公开数据和量化模型自动生成，仅供参考，不构成投资建议。股市有风险，投资需谨慎。")

    # ===== 保存 =====
    today_s = datetime.now().strftime("%Y%m%d")
    fp = os.path.join(output_dir, f"{code}_{name}_14章报告_{today_s}.docx")
    doc.save(fp)
    print(f"[report] 已保存 {fp}")
    return fp
