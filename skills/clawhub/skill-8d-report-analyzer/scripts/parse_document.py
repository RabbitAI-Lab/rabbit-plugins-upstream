#!/usr/bin/env python3
"""
8D报告文档解析脚本
支持Word(.docx)、PDF(.pdf)、Excel(.xlsx)格式的8D报告解析
输出JSON格式的结构化数据
"""

import argparse
import json
import sys
import re
from pathlib import Path

def clean_text(text):
    """清洗文本内容"""
    # 去除多余空白字符
    text = re.sub(r'\s+', ' ', text)
    # 去除特殊控制字符
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f]', '', text)
    return text.strip()

def parse_docx(file_path):
    """解析Word文档"""
    try:
        from docx import Document
    except ImportError:
        print(json.dumps({"error": "python-docx未安装，请运行: pip install python-docx==1.1.0"}, ensure_ascii=False))
        sys.exit(1)
    
    doc = Document(file_path)
    paragraphs = []
    
    for para in doc.paragraphs:
        text = clean_text(para.text)
        if text:
            paragraphs.append({
                "text": text,
                "style": para.style.name if para.style else "Normal"
            })
    
    # 提取表格内容
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [clean_text(cell.text) for cell in row.cells]
            table_data.append(row_data)
        if table_data:
            tables.append(table_data)
    
    return {
        "format": "docx",
        "paragraphs": paragraphs,
        "tables": tables,
        "full_text": "\n".join([p["text"] for p in paragraphs])
    }

def parse_pdf(file_path):
    """解析PDF文档"""
    try:
        import PyPDF2
    except ImportError:
        print(json.dumps({"error": "PyPDF2未安装，请运行: pip install PyPDF2==3.0.1"}, ensure_ascii=False))
        sys.exit(1)
    
    paragraphs = []
    
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                # 分行处理，保持段落结构
                lines = text.split('\n')
                for line in lines:
                    line = clean_text(line)
                    if line:
                        paragraphs.append({
                            "text": line,
                            "page": page_num + 1
                        })
    
    return {
        "format": "pdf",
        "paragraphs": paragraphs,
        "tables": [],
        "full_text": "\n".join([p["text"] for p in paragraphs])
    }

def parse_xlsx(file_path):
    """解析Excel文档"""
    try:
        import openpyxl
    except ImportError:
        print(json.dumps({"error": "openpyxl未安装，请运行: pip install openpyxl==3.1.2"}, ensure_ascii=False))
        sys.exit(1)
    
    wb = openpyxl.load_workbook(file_path)
    all_content = []
    tables = []
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        sheet_data = []
        
        for row in sheet.iter_rows(values_only=True):
            row_data = [str(cell) if cell is not None else "" for cell in row]
            # 清洗每个单元格
            row_data = [clean_text(cell) for cell in row_data]
            # 跳过空行
            if any(cell for cell in row_data):
                sheet_data.append(row_data)
                all_content.append(" | ".join(row_data))
        
        if sheet_data:
            tables.append({
                "sheet": sheet_name,
                "data": sheet_data
            })
    
    return {
        "format": "xlsx",
        "paragraphs": [{"text": line} for line in all_content],
        "tables": tables,
        "full_text": "\n".join(all_content)
    }

def detect_format(file_path):
    """根据文件扩展名检测格式"""
    ext = Path(file_path).suffix.lower()
    format_map = {
        '.docx': 'docx',
        '.doc': 'docx',  # 尝试兼容，失败再说
        '.pdf': 'pdf',
        '.xlsx': 'xlsx',
        '.xls': 'xlsx'
    }
    return format_map.get(ext)

def main():
    parser = argparse.ArgumentParser(description='8D报告文档解析工具')
    parser.add_argument('--file', required=True, help='文件路径')
    parser.add_argument('--format', choices=['docx', 'pdf', 'xlsx'], 
                       help='文件格式，不指定则根据扩展名自动检测')
    
    args = parser.parse_args()
    
    file_path = Path(args.file)
    if not file_path.exists():
        print(json.dumps({"error": f"文件不存在: {file_path}"}))
        sys.exit(1)
    
    # 检测格式
    fmt = args.format or detect_format(str(file_path))
    if not fmt:
        print(json.dumps({"error": "不支持的文件格式，仅支持 .docx, .pdf, .xlsx"}))
        sys.exit(1)
    
    # 解析文档
    try:
        if fmt == 'docx':
            result = parse_docx(str(file_path))
        elif fmt == 'pdf':
            result = parse_pdf(str(file_path))
        elif fmt == 'xlsx':
            result = parse_xlsx(str(file_path))
        
        # 输出JSON结果
        print(json.dumps(result, ensure_ascii=False, indent=2))
        
    except Exception as e:
        print(json.dumps({"error": f"解析失败: {str(e)}"}))
        sys.exit(1)

if __name__ == "__main__":
    main()
