#!/usr/bin/env python3
"""Generate an evidence-backed digital-construction Word report.

Usage:
    python generate_word_report.py --input reviewed_report.json --output report.docx
    python generate_word_report.py --write-schema report-schema.json

The generator deliberately contains no patent, company, policy, project, ranking,
date or performance conclusions. All report facts must arrive through a reviewed
UTF-8 JSON file and retain evidence identifiers.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any, Iterable


SCHEMA_EXAMPLE = {
    "metadata": {
        "title": "Digital Construction Technology Intelligence",
        "subtitle": "Bridges, tunnels and highways",
        "decision": "Describe the decision supported by this report",
        "geographies": ["Define geography"],
        "cutoff_date": "YYYY-MM-DD",
        "prepared_date": "YYYY-MM-DD",
        "counting_unit": "simple patent families",
        "confidentiality": "Internal",
        "status": "Draft — requires technical and IP review",
    },
    "executive_findings": [
        {
            "finding": "Evidence-backed finding",
            "confidence": "medium",
            "evidence_ids": ["S1"],
            "implication": "Decision implication",
        }
    ],
    "methodology": {
        "scope": "Included and excluded technology meanings",
        "sources": ["Source categories used"],
        "search_routes": ["Semantic", "Keyword/classification", "Entity/date"],
        "coverage": "Coverage, pagination and known limitations",
        "counting_rules": "Family/publication and entity normalization rules",
    },
    "technology_architecture": [
        {
            "layer": "Sensing",
            "function": "Function",
            "technologies": "Technologies",
            "maturity": "Assessed maturity",
            "dependencies": "Interfaces/dependencies",
            "failure_modes": "Material failure modes",
            "evidence_ids": ["S1"],
        }
    ],
    "patents": [
        {
            "publication_number": "Publication number",
            "title": "Patent title",
            "applicant": "Normalized applicant",
            "earliest_priority": "YYYY-MM-DD or unknown",
            "jurisdiction": "Jurisdiction",
            "family_id": "Provider family identifier",
            "status_as_of": "Status and observation date",
            "claim_relevance": "Why a claim is relevant",
            "url": "Exact returned global PatSnap URL",
            "evidence_ids": ["S1"],
        }
    ],
    "scientific_evidence": [
        {
            "study": "Study title/identifier",
            "question": "Research question",
            "method": "Method and test environment",
            "result": "Result with metric and uncertainty",
            "translation_stage": "Simulation/lab/pilot/operation",
            "limitations": "Limitations",
            "evidence_ids": ["S1"],
        }
    ],
    "competitors": [
        {
            "actor": "Normalized actor",
            "role": "Platform/OEM/contractor/etc.",
            "capability": "Evidence-backed capability",
            "deployment": "Deployment evidence",
            "patent_position": "Scoped patent evidence",
            "limitations": "Missing data or caveats",
            "confidence": "low/medium/high",
            "evidence_ids": ["S1"],
        }
    ],
    "project_cases": [
        {
            "project": "Project name",
            "location": "Location",
            "period": "Period",
            "intervention": "Technology intervention",
            "baseline": "Baseline and comparator",
            "measured_outcome": "Outcome, unit and method",
            "transfer_limits": "Conditions limiting transfer",
            "evidence_ids": ["S1"],
        }
    ],
    "trends": [
        {
            "trend": "Trend or scenario",
            "supporting_signals": "Signals",
            "counter_signals": "Counter-signals",
            "time_horizon": "Horizon",
            "confidence": "low/medium/high",
            "evidence_ids": ["S1"],
        }
    ],
    "actions": [
        {
            "action": "Recommended action",
            "owner": "Owner",
            "timing": "Timing",
            "success_measure": "Measurable success criterion",
            "dependency": "Dependency",
        }
    ],
    "limitations": ["Material limitation"],
    "evidence_register": [
        {
            "id": "S1",
            "title": "Source title",
            "publisher": "Publisher",
            "date": "Publication date",
            "accessed": "Access date",
            "url_or_identifier": "Stable URL or identifier",
            "source_type": "Patent/standard/paper/project/company/news",
            "notes": "Relevant fields, passage or limitations",
        }
    ],
}


REQUIRED_TOP_LEVEL = (
    "metadata",
    "executive_findings",
    "methodology",
    "technology_architecture",
    "patents",
    "scientific_evidence",
    "competitors",
    "project_cases",
    "trends",
    "actions",
    "limitations",
    "evidence_register",
)

TABLES = {
    "technology_architecture": (
        "Technology architecture",
        ("layer", "function", "technologies", "maturity", "dependencies", "failure_modes", "evidence_ids"),
    ),
    "patents": (
        "Patent landscape",
        ("publication_number", "title", "applicant", "earliest_priority", "jurisdiction", "status_as_of", "claim_relevance", "evidence_ids"),
    ),
    "scientific_evidence": (
        "Scientific evidence and translation",
        ("study", "question", "method", "result", "translation_stage", "limitations", "evidence_ids"),
    ),
    "competitors": (
        "Competitive landscape",
        ("actor", "role", "capability", "deployment", "patent_position", "limitations", "confidence", "evidence_ids"),
    ),
    "project_cases": (
        "Deployment and project cases",
        ("project", "location", "period", "intervention", "baseline", "measured_outcome", "transfer_limits", "evidence_ids"),
    ),
    "trends": (
        "Trends, constraints and scenarios",
        ("trend", "supporting_signals", "counter_signals", "time_horizon", "confidence", "evidence_ids"),
    ),
    "actions": (
        "Recommended actions",
        ("action", "owner", "timing", "success_measure", "dependency"),
    ),
}


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    mode = parser.add_mutually_exclusive_group(required=True)
    mode.add_argument("--input", type=Path, help="Reviewed UTF-8 JSON evidence package")
    mode.add_argument("--write-schema", type=Path, help="Write an example JSON input and exit")
    parser.add_argument("--output", type=Path, help="Output DOCX path; required with --input")
    parser.add_argument("--force", action="store_true", help="Replace an existing output file")
    return parser.parse_args(argv)


def read_json(path: Path) -> dict[str, Any]:
    if not path.is_file():
        raise ValueError(f"Input file does not exist: {path}")
    try:
        data = json.loads(path.read_text(encoding="utf-8-sig"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(f"Cannot read valid JSON from {path}: {exc}") from exc
    if not isinstance(data, dict):
        raise ValueError("Top-level JSON value must be an object")
    return data


def text_value(value: Any) -> str:
    if value is None:
        return "Not reported"
    if isinstance(value, list):
        return "; ".join(text_value(item) for item in value) or "Not reported"
    if isinstance(value, (str, int, float, bool)):
        rendered = str(value).strip()
        return rendered or "Not reported"
    return json.dumps(value, ensure_ascii=False, sort_keys=True)


def evidence_ids(value: Any) -> list[str]:
    if isinstance(value, str):
        values = [value]
    elif isinstance(value, list):
        values = [str(item) for item in value]
    else:
        values = []
    return [item.strip() for item in values if item.strip()]


def validate(data: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    for key in REQUIRED_TOP_LEVEL:
        if key not in data:
            errors.append(f"Missing top-level field: {key}")

    metadata = data.get("metadata")
    if not isinstance(metadata, dict):
        errors.append("metadata must be an object")
    else:
        for key in ("title", "decision", "geographies", "cutoff_date", "counting_unit", "status"):
            if not text_value(metadata.get(key)) or text_value(metadata.get(key)) == "Not reported":
                errors.append(f"metadata.{key} is required")

    register = data.get("evidence_register")
    if not isinstance(register, list):
        errors.append("evidence_register must be an array")
        register = []
    registered: set[str] = set()
    for index, record in enumerate(register):
        if not isinstance(record, dict):
            errors.append(f"evidence_register[{index}] must be an object")
            continue
        identifier = text_value(record.get("id"))
        if identifier == "Not reported":
            errors.append(f"evidence_register[{index}].id is required")
        elif identifier in registered:
            errors.append(f"Duplicate evidence ID: {identifier}")
        else:
            registered.add(identifier)
        for key in ("title", "publisher", "url_or_identifier", "source_type"):
            if text_value(record.get(key)) == "Not reported":
                errors.append(f"evidence_register[{index}].{key} is required")

    cited: set[str] = set()
    for section in ("executive_findings",) + tuple(TABLES):
        records = data.get(section)
        if not isinstance(records, list):
            errors.append(f"{section} must be an array")
            continue
        for index, record in enumerate(records):
            if not isinstance(record, dict):
                errors.append(f"{section}[{index}] must be an object")
                continue
            if section != "actions":
                ids = evidence_ids(record.get("evidence_ids"))
                if not ids:
                    errors.append(f"{section}[{index}] needs evidence_ids")
                cited.update(ids)

    unknown = sorted(cited - registered)
    if unknown:
        errors.append("Unregistered evidence IDs: " + ", ".join(unknown))
    return errors


def set_cell_shading(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_table(document: Any, columns: Iterable[str], rows: list[dict[str, Any]]) -> None:
    from docx.shared import Pt, RGBColor

    fields = tuple(columns)
    table = document.add_table(rows=1, cols=len(fields))
    table.style = "Table Grid"
    table.autofit = True
    for index, field in enumerate(fields):
        cell = table.rows[0].cells[index]
        cell.text = field.replace("_", " ").title()
        set_cell_shading(cell, "16324F")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for index, field in enumerate(fields):
            cells[index].text = text_value(row.get(field))
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    document.add_paragraph()


def add_source_link(paragraph: Any, label: str, url: str) -> None:
    """Add a true external hyperlink to a paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F5A85")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    text = OxmlElement("w:t")
    text.text = label
    run.extend((properties, text))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_document(document: Any) -> None:
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    for section in document.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.color.rgb = RGBColor(22, 50, 79)
    for style_name, size in (("Heading 1", 16), ("Heading 2", 12)):
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = RGBColor(22, 50, 79)

    footer = document.sections[0].footer.paragraphs[0]
    footer.text = "Evidence-backed digital construction analysis"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_report(data: dict[str, Any]) -> Any:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    document = Document()
    configure_document(document)
    metadata = data["metadata"]

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run(text_value(metadata.get("title")))
    subtitle = document.add_paragraph()
    subtitle.add_run(text_value(metadata.get("subtitle"))).bold = True
    document.add_paragraph(
        f"Decision: {text_value(metadata.get('decision'))}\n"
        f"Geographies: {text_value(metadata.get('geographies'))}\n"
        f"Evidence cutoff: {text_value(metadata.get('cutoff_date'))}\n"
        f"Counting unit: {text_value(metadata.get('counting_unit'))}\n"
        f"Status: {text_value(metadata.get('status'))}"
    )
    document.add_paragraph("Scientific and patent evidence must be interpreted within the documented scope and limitations.")

    document.add_heading("Executive findings", level=1)
    for record in data["executive_findings"]:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(text_value(record.get("finding"))).bold = True
        paragraph.add_run(
            f" — Confidence: {text_value(record.get('confidence'))}; "
            f"evidence: {text_value(record.get('evidence_ids'))}. "
            f"Implication: {text_value(record.get('implication'))}"
        )

    document.add_heading("Method, scope and coverage", level=1)
    methodology = data["methodology"]
    for field in ("scope", "sources", "search_routes", "coverage", "counting_rules"):
        paragraph = document.add_paragraph()
        paragraph.add_run(field.replace("_", " ").title() + ": ").bold = True
        paragraph.add_run(text_value(methodology.get(field)))

    for section, (heading, columns) in TABLES.items():
        document.add_heading(heading, level=1)
        rows = data[section]
        if rows:
            add_table(document, columns, rows)
        else:
            document.add_paragraph("No evidence-backed records were supplied for this section.")

    document.add_heading("Limitations", level=1)
    for limitation in data["limitations"]:
        document.add_paragraph(text_value(limitation), style="List Bullet")

    document.add_heading("Evidence register", level=1)
    for record in data["evidence_register"]:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"[{text_value(record.get('id'))}] {text_value(record.get('title'))}").bold = True
        paragraph.add_run(
            f". {text_value(record.get('publisher'))}; {text_value(record.get('date'))}. "
            f"Accessed {text_value(record.get('accessed'))}. Type: {text_value(record.get('source_type'))}. "
        )
        url = text_value(record.get("url_or_identifier"))
        if re.match(r"^https?://", url, flags=re.IGNORECASE):
            add_source_link(paragraph, url, url)
        else:
            paragraph.add_run(url)
        paragraph.add_run(f". Notes: {text_value(record.get('notes'))}")

    return document


def write_schema(path: Path, force: bool) -> None:
    if path.exists() and not force:
        raise ValueError(f"Refusing to replace existing file without --force: {path}")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(SCHEMA_EXAMPLE, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    try:
        if args.write_schema:
            write_schema(args.write_schema, args.force)
            print(f"Wrote example schema: {args.write_schema}")
            return 0

        if args.output is None:
            raise ValueError("--output is required with --input")
        if args.output.exists() and not args.force:
            raise ValueError(f"Refusing to replace existing output without --force: {args.output}")

        data = read_json(args.input)
        errors = validate(data)
        if errors:
            raise ValueError("Input validation failed:\n- " + "\n- ".join(errors))

        try:
            import docx  # noqa: F401
        except ImportError as exc:
            raise ValueError("python-docx is required to generate Word output") from exc

        args.output.parent.mkdir(parents=True, exist_ok=True)
        document = build_report(data)
        document.save(args.output)
        print(f"Generated reviewed report: {args.output}")
        return 0
    except (OSError, ValueError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
