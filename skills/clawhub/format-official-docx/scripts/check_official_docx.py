#!/usr/bin/env python3
"""Check common formatting signals in a generated official DOCX."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def east_asia_font(run) -> str | None:
    r_pr = run._element.rPr
    if r_pr is None or r_pr.rFonts is None:
        return None
    return r_pr.rFonts.get(qn("w:eastAsia"))


def run_color(run) -> str | None:
    color = run.font.color.rgb if run.font.color else None
    return str(color) if color else None


def nonempty_paragraphs(doc: Document):
    return [p for p in doc.paragraphs if p.text.strip()]


def check_docx(path: Path, doc_type: str) -> dict:
    doc = Document(str(path))
    section = doc.sections[0]
    paragraphs = nonempty_paragraphs(doc)
    text = "\n".join(p.text.strip() for p in paragraphs)
    result = {
        "file": str(path),
        "doc_type": doc_type,
        "paragraph_count": len(paragraphs),
        "section_count": len(doc.sections),
        "table_count": len(doc.tables),
        "checks": [],
        "warnings": [],
    }

    def add(name: str, ok: bool, detail: str):
        result["checks"].append({"name": name, "ok": ok, "detail": detail})

    add(
        "a4_margins",
        round(section.page_width.cm, 1) == 21.0 and round(section.page_height.cm, 1) == 29.7,
        f"page={section.page_width.cm:.2f}x{section.page_height.cm:.2f}cm",
    )
    expected_top = 3.7
    expected_bottom = 3.5
    add(
        "margins",
        abs(section.top_margin.cm - expected_top) < 0.05
        and abs(section.bottom_margin.cm - expected_bottom) < 0.05
        and abs(section.left_margin.cm - 2.8) < 0.05
        and abs(section.right_margin.cm - 2.6) < 0.05,
        f"top={section.top_margin.cm:.2f}, bottom={section.bottom_margin.cm:.2f}, left={section.left_margin.cm:.2f}, right={section.right_margin.cm:.2f}",
    )

    if paragraphs:
        # 标题/红头通常为居中段落；密级、主送机关等可能位于其前，故按居中段落定位
        title_para = next((p for p in paragraphs if p.alignment == WD_ALIGN_PARAGRAPH.CENTER), paragraphs[0])
        first_run = title_para.runs[0] if title_para.runs else None
        add(
            "first_title_font",
            bool(first_run and east_asia_font(first_run) == "方正小标宋简体"),
            east_asia_font(first_run) if first_run else "no run",
        )
        add(
            "first_title_alignment",
            title_para.alignment == WD_ALIGN_PARAGRAPH.CENTER,
            str(title_para.alignment),
        )
        if doc_type in {"request", "letter"}:
            add(
                "red_header_or_title",
                bool(first_run and run_color(first_run) == "FF0000") if doc_type == "request" else True,
                run_color(first_run) if first_run else "no run",
            )

    add("has_recipient_colon", bool(re.search(r"：", text)), "found Chinese colon" if "：" in text else "not found")

    # 页码字体（宋体四号，规范第六条）
    footer_fonts = set()
    for fp in section.footer.paragraphs:
        for r in fp.runs:
            ea = east_asia_font(r)
            if ea:
                footer_fonts.add(ea)
    add(
        "page_number_font_song",
        ("宋体" in " ".join(footer_fonts)) or not footer_fonts,
        "footer fonts=" + ",".join(sorted(footer_fonts)),
    )
    if doc_type == "request":
        request_closing_count = text.count("妥否，请批示。")
        add("has_request_closing", request_closing_count >= 1, "妥否，请批示。")
        add("no_duplicate_request_closing", request_closing_count <= 1, f"count={request_closing_count}")
        add("has_signer_label", "签发人：" in text, "签发人：")
        add("has_outgoing_number", bool(re.search(r"〔[^〕]+〕\S+号", text)), "发文字号 pattern")
    if doc_type == "letter":
        add("has_letter_reply_closing", any(closing in text for closing in ("此复。", "函复", "特此函复。")), "此复/函复")
        if len(paragraphs) >= 4:
            add(
                "letter_title_has_issuer_line",
                paragraphs[0].text.strip() == paragraphs[2].text.strip(),
                "red header agency should repeat as first centered title line",
            )
        tail = paragraphs[-4:] if len(paragraphs) >= 4 else paragraphs
        right_aligned_tail = [p for p in tail if p.alignment == WD_ALIGN_PARAGRAPH.RIGHT]
        add("signature_right_aligned", len(right_aligned_tail) >= 2, f"right-aligned tail paragraphs={len(right_aligned_tail)}")
    if doc_type == "report":
        if "特此报告。" in text:
            result["warnings"].append("报告正文含『特此报告。』结尾，符合常见结尾习惯。")
        else:
            result["warnings"].append("报告未检测到『特此报告。』结尾（如正文确无结尾或仅为引用，可忽略）。")

    if doc.tables:
        result["warnings"].append("Tables are present; verify table formatting manually.")
    if "印发" in text:
        has_even_page_section = any(section.start_type == WD_SECTION.EVEN_PAGE for section in doc.sections)
        add("copy_block_even_page_section", has_even_page_section, "even-page section break before copy/printing block")
    if not paragraphs:
        result["warnings"].append("No non-empty paragraphs found.")

    result["ok"] = all(item["ok"] for item in result["checks"])
    return result


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx")
    parser.add_argument("--doc-type", default="generic", choices=["generic", "request", "report", "letter", "notice", "approval", "opinion", "minutes"])
    parser.add_argument("--json", action="store_true", help="Emit JSON instead of plain text")
    args = parser.parse_args()

    result = check_docx(Path(args.docx), args.doc_type)
    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return

    print(f"File: {result['file']}")
    print(f"Overall: {'OK' if result['ok'] else 'NEEDS REVIEW'}")
    for item in result["checks"]:
        mark = "OK" if item["ok"] else "FAIL"
        print(f"[{mark}] {item['name']}: {item['detail']}")
    for warning in result["warnings"]:
        print(f"[WARN] {warning}")


if __name__ == "__main__":
    main()
