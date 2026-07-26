#!/usr/bin/env python3
"""Generate or reformat a Chinese Party/government official DOCX."""

from __future__ import annotations

import argparse
import copy
import datetime
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.shared import RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


BODY_FONT = "仿宋_GB2312"
TITLE_FONT = "方正小标宋简体"
HEITI = "黑体"
KAITI = "楷体_GB2312"
TIMES = "Times New Roman"
SONGTI = "宋体"
TEXT_ENCODINGS = ("utf-8-sig", "utf-8", "gb18030", "gbk", "gb2312")
USABLE_PAGE_WIDTH = Cm(15.6)


def set_run_font(
    run,
    east_asia: str = BODY_FONT,
    size_pt: float = 16,
    bold: bool | None = None,
    color: str | None = None,
):
    run.font.name = TIMES
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), TIMES)
    r_fonts.set(qn("w:hAnsi"), TIMES)


def set_paragraph_common(paragraph, first_line: bool = True, align=None):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(28.9)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Cm(1.13) if first_line else None
    if align is not None:
        paragraph.alignment = align


def add_text_run(
    paragraph,
    text: str,
    east_asia: str = BODY_FONT,
    size_pt: float = 16,
    bold: bool | None = None,
    color: str | None = None,
):
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=east_asia, size_pt=size_pt, bold=bold, color=color)
    return run


def add_mixed_text(
    paragraph,
    text: str,
    east_asia: str = BODY_FONT,
    size_pt: float = 16,
    bold: bool | None = None,
    color: str | None = None,
):
    if not text:
        return
    parts = re.findall(r"[A-Za-z0-9.\-—]+|[^A-Za-z0-9.\-—]+", text)
    for part in parts:
        add_text_run(paragraph, part, east_asia=east_asia, size_pt=size_pt, bold=bold, color=color)


def read_text_with_fallback(path: Path) -> str:
    data = path.read_bytes()
    errors = []
    for encoding in TEXT_ENCODINGS:
        try:
            return data.decode(encoding)
        except UnicodeDecodeError as exc:
            errors.append(f"{encoding}: {exc}")
    raise UnicodeDecodeError(
        "official-docx",
        data,
        0,
        min(len(data), 1),
        "unable to decode text file; tried " + ", ".join(TEXT_ENCODINGS),
    )


def probe_fonts() -> dict:
    """Best-effort check for official-document fonts on Windows.

    Returns {canonical_name: bool_found}. Non-Windows platforms return {} (skipped)
    so generation is never blocked; the canonical GB2312 names are kept on purpose
    to match the unit's formatting standard.
    """
    windir = os.environ.get("WINDIR")
    if not windir:
        return {}
    fdir = os.path.join(windir, "Fonts")
    if not os.path.isdir(fdir):
        return {}
    present = [p.lower() for p in os.listdir(fdir)]
    hints = {
        "方正小标宋简体": ("fzxbs", "方正小标宋", "fzxiaobiaosong"),
        "仿宋_GB2312": ("仿宋", "fangsong", "fangsong_gb2312"),
        "楷体_GB2312": ("楷体", "kaiti", "kai"),
    }
    return {font: any(h.lower() in p for p in present for h in hs) for font, hs in hints.items()}


def warn_missing_fonts():
    info = probe_fonts()
    if not info:
        return
    missing = [name for name, ok in info.items() if not ok]
    if missing:
        sys.stderr.write(
            "警告：本机未安装公文规范字体 " + "、".join(missing)
            + "，Word 打开后将自动替换字体，可能导致字号/字间距偏移；请安装对应字体以获得规范排版。\n"
        )


def add_tabbed_line(
    paragraph,
    left: str,
    right: str,
    east_asia: str = BODY_FONT,
    size_pt: float = 16,
    leader=WD_TAB_LEADER.SPACES,
):
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        USABLE_PAGE_WIDTH,
        WD_TAB_ALIGNMENT.RIGHT,
        leader,
    )
    add_mixed_text(paragraph, left, east_asia=east_asia, size_pt=size_pt)
    paragraph.add_run("\t")
    add_mixed_text(paragraph, right, east_asia=east_asia, size_pt=size_pt)


def setup_section(section, doc_type: str, letter_special_margins: bool = False):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    if doc_type == "letter" and letter_special_margins:
        section.top_margin = Cm(5.1)
        section.bottom_margin = Cm(2.1)
    else:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(2.75)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_run(p, "— ", east_asia=SONGTI, size_pt=14)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    set_run_font(run, east_asia=SONGTI, size_pt=14)
    add_text_run(p, " —", east_asia=SONGTI, size_pt=14)


def set_paragraph_border(paragraph, edge: str = "bottom", color: str = "FF0000", size: str = "18"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), color)


def add_red_footer_line(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_border(p, edge="top", color="FF0000", size="18")


def add_letterhead(doc, agency: str, outgoing_no: str | None = None):
    p = doc.add_paragraph()
    set_paragraph_common(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_mixed_text(p, agency, east_asia=TITLE_FONT, size_pt=28, color="FF0000")
    set_paragraph_border(p, edge="bottom", color="FF0000", size="18")
    if outgoing_no:
        no = doc.add_paragraph()
        set_paragraph_common(no, first_line=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
        add_mixed_text(no, outgoing_no, east_asia=BODY_FONT, size_pt=16)
    doc.add_paragraph()


def add_red_document_header(doc, header: str, outgoing_no: str | None = None, signer: str | None = None):
    p = doc.add_paragraph()
    set_paragraph_common(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_mixed_text(p, header, east_asia=TITLE_FONT, size_pt=26, color="FF0000")
    info = doc.add_paragraph()
    set_paragraph_common(info, first_line=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    left = outgoing_no or "XXXX〔XXXX〕XX号"
    right = f"签发人：{signer}" if signer else "签发人：XXX"
    add_tabbed_line(info, left, right, east_asia=BODY_FONT, size_pt=16, leader=WD_TAB_LEADER.DOTS)
    set_paragraph_border(info, edge="bottom", color="FF0000", size="18")
    doc.add_paragraph()


def classify_paragraph(text: str) -> str:
    stripped = text.strip()
    if re.match(r"^[一二三四五六七八九十]+、", stripped):
        return "h1"
    if re.match(r"^（[一二三四五六七八九十]+）", stripped):
        return "h2"
    if re.match(r"^\d+[．.]", stripped):
        return "h3"
    if re.match(r"^（\d+）", stripped):
        return "h4"
    return "body"


def add_body_paragraph(doc, text: str, headings_bold: bool = False):
    p = doc.add_paragraph()
    kind = classify_paragraph(text)
    set_paragraph_common(p, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    if kind == "h1":
        add_mixed_text(p, text, east_asia=HEITI, size_pt=16, bold=headings_bold or None)
    elif kind == "h2":
        add_mixed_text(p, text, east_asia=KAITI, size_pt=16)
    else:
        add_mixed_text(p, text, east_asia=BODY_FONT, size_pt=16)
    return p


def add_centered(doc, text: str, east_asia: str, size_pt: float):
    p = doc.add_paragraph()
    set_paragraph_common(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_mixed_text(p, text, east_asia=east_asia, size_pt=size_pt)
    return p


def read_input(path: Path) -> list[str]:
    if path.suffix.lower() == ".docx":
        source = Document(str(path))
        return [p.text.strip().lstrip("\ufeff") for p in source.paragraphs if p.text.strip()]
    return [line.strip().lstrip("\ufeff") for line in read_text_with_fallback(path).splitlines() if line.strip()]


def _restyle_block_runs(element, doc, headings_bold: bool):
    """Apply the official body style (font/size/line spacing/first-line indent) to a
    cloned block element (a w:p paragraph or w:tbl table). Headings keep their
    canonical font (黑体/楷体); level-1 headings are bolded only when requested.
    """
    text_blocks = []
    if element.tag.endswith("}tbl"):
        tbl = Table(element, doc)
        for row in tbl.rows:
            for cell in row.cells:
                text_blocks.extend(cell.paragraphs)
    else:
        text_blocks.append(Paragraph(element, doc))
    for p in text_blocks:
        text = p.text.strip()
        kind = classify_paragraph(text) if text else "body"
        if kind == "h1":
            ea, bold = HEITI, (True if headings_bold else None)
        elif kind == "h2":
            ea, bold = KAITI, None
        else:
            ea, bold = BODY_FONT, None
        for run in p.runs:
            set_run_font(run, east_asia=ea, size_pt=16, bold=bold)
        set_paragraph_common(p, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)


def format_only_clone(src_doc: Document, out_doc: Document, headings_bold: bool = False):
    """Re-layout mode: clone every paragraph and table from the source .docx into the
    output document, restyling them to the official body format. Tables are preserved
    (previously dropped), keeping their structure and cell text.
    """
    body = out_doc.element.body
    for child in src_doc.element.body:
        tag = child.tag
        if tag.endswith("}p") or tag.endswith("}tbl"):
            clone = copy.deepcopy(child)
            body.append(clone)
            _restyle_block_runs(clone, out_doc, headings_bold)


def add_attachment_block(doc, attachments: list[str]):
    if not attachments:
        return
    doc.add_paragraph()
    if len(attachments) == 1:
        p = doc.add_paragraph()
        set_paragraph_common(p, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        add_mixed_text(p, f"附件：{attachments[0]}", east_asia=BODY_FONT, size_pt=16)
        return
    for idx, name in enumerate(attachments, start=1):
        p = doc.add_paragraph()
        set_paragraph_common(p, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        prefix = "附件：" if idx == 1 else "      "
        add_mixed_text(p, f"{prefix}{idx}．{name}", east_asia=BODY_FONT, size_pt=16)


def add_attachment_markers(doc, attachments: list[str]):
    """Insert per-attachment '附件N' placeholder lines (top-aligned, 仿宋_GB2312, 三号).

    Per the unit's standard, each attachment's first page carries a top-left '附件N'
    marker. Because attachment bodies are not generated here, a placeholder line is
    added after the attachment list for the user to replace with real content.
    """
    if not attachments:
        return
    for idx in range(1, len(attachments) + 1):
        doc.add_paragraph()
        p = doc.add_paragraph()
        set_paragraph_common(p, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        add_mixed_text(p, f"附件{idx}", east_asia=BODY_FONT, size_pt=16)


def add_signature(doc, issuer: str | None, date: str | None, contact: str | None, doc_type: str = "generic"):
    if not issuer and not date and not contact:
        return
    for _ in range(3):
        doc.add_paragraph()
    # 发文机关署名：信函落款靠右（右空四字）；其余文种按规范第九条居中排版
    if issuer:
        if doc_type == "letter":
            p = doc.add_paragraph()
            set_paragraph_common(p, first_line=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
            p.paragraph_format.right_indent = Cm(2.26)
            add_mixed_text(p, issuer, east_asia=BODY_FONT, size_pt=16)
        else:
            add_centered(doc, issuer, BODY_FONT, 16)
    # 成文日期：右空四字（规范第九条）
    if date:
        p = doc.add_paragraph()
        set_paragraph_common(p, first_line=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
        p.paragraph_format.right_indent = Cm(2.26)
        add_mixed_text(p, date, east_asia=BODY_FONT, size_pt=16)
    if contact:
        p = doc.add_paragraph()
        set_paragraph_common(p, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        add_mixed_text(p, contact, east_asia=BODY_FONT, size_pt=16)


def add_disclosure_note(doc, disclosure_note: str | None):
    if not disclosure_note:
        return
    p = doc.add_paragraph()
    set_paragraph_common(p, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_mixed_text(p, disclosure_note, east_asia=BODY_FONT, size_pt=16)


def add_secret_level(doc, secret: str | None):
    """左上角标注密级与保密期限（三号黑体，密级★期限）。"""
    if not secret:
        return
    p = doc.add_paragraph()
    set_paragraph_common(p, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after = Pt(0)
    add_mixed_text(p, secret, east_asia=HEITI, size_pt=16, bold=True)


def add_copy_block(doc, copy_to: str | None, print_org: str | None, print_date: str | None):
    if not any([copy_to, print_org, print_date]):
        return
    doc.add_section(WD_SECTION.EVEN_PAGE)
    if copy_to:
        p = doc.add_paragraph()
        set_paragraph_common(p, first_line=False)
        set_paragraph_border(p, edge="top", color="000000", size="6")
        set_paragraph_border(p, edge="bottom", color="000000", size="6")
        p.paragraph_format.left_indent = Cm(1.48)
        p.paragraph_format.first_line_indent = Cm(-0.99)
        add_mixed_text(p, f"抄送：{copy_to.rstrip('。')}。", east_asia=BODY_FONT, size_pt=14)
    p = doc.add_paragraph()
    set_paragraph_common(p, first_line=False)
    set_paragraph_border(p, edge="bottom", color="000000", size="6")
    p.paragraph_format.first_line_indent = Cm(0.49)
    left = print_org or "XXXXXXXXX"
    right = f"{print_date}印发" if print_date else "XXXX年XX月XX日印发"
    add_tabbed_line(p, left, right, east_asia=BODY_FONT, size_pt=14)


def build_doc(args):
    if args.date and args.date.lower() == "today":
        today = datetime.date.today()
        args.date = f"{today.year}年{today.month}月{today.day}日"
    doc = Document()
    setup_section(doc.sections[0], args.doc_type, args.letter_special_margins)
    warn_missing_fonts()
    if args.secret_level:
        add_secret_level(doc, args.secret_level)
        doc.add_paragraph()
    if args.red_header:
        add_red_document_header(doc, args.red_header, args.outgoing_no, args.signer)
    if args.doc_type == "letter" and args.letterhead:
        add_red_footer_line(doc.sections[0])
        add_letterhead(doc, args.letterhead, args.outgoing_no)
    if args.doc_type != "letter" and not args.no_page_number:
        add_page_number(doc.sections[0])

    clone_mode = bool(
        args.format_only and args.input and Path(args.input).suffix.lower() == ".docx"
    )
    if clone_mode:
        # 重排模式：克隆源文档的段落与表格，套用公文正文字体/行距/缩进，保留表格结构
        src_doc = Document(str(args.input))
        format_only_clone(src_doc, doc, headings_bold=args.bold_headings)
        paragraphs = []
    else:
        paragraphs = read_input(Path(args.input)) if args.input else []

    if not clone_mode and args.outgoing_no and not (args.doc_type == "letter" and args.letterhead) and not args.red_header:
        p = doc.add_paragraph()
        set_paragraph_common(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        line = args.outgoing_no
        if args.signer:
            add_tabbed_line(p, args.outgoing_no, f"签发人：{args.signer}", east_asia=BODY_FONT, size_pt=16, leader=WD_TAB_LEADER.DOTS)
        else:
            add_mixed_text(p, line, east_asia=BODY_FONT, size_pt=16)
        doc.add_paragraph()

    title = args.title
    if not clone_mode and title:
        if args.doc_type == "letter" and args.letterhead and args.issuer and args.issuer not in title:
            add_centered(doc, args.issuer, TITLE_FONT, 22)
        for line in title.split("\\n"):
            add_centered(doc, line.strip(), TITLE_FONT, 22)
        doc.add_paragraph()

    if args.recipient and not args.format_only:
        recipient = args.recipient
        if not recipient.endswith("：") and not recipient.endswith(":"):
            recipient += "："
        p = doc.add_paragraph()
        set_paragraph_common(p, first_line=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        add_mixed_text(p, recipient, east_asia=BODY_FONT, size_pt=16)

    if not clone_mode:
        for idx, text in enumerate(paragraphs):
            if title and text == title:
                continue
            if args.promote_first_paragraph_title and not title and idx == 0:
                add_centered(doc, text, TITLE_FONT, 22)
                doc.add_paragraph()
                continue
            add_body_paragraph(doc, text, headings_bold=args.bold_headings)

    if not args.format_only and args.add_closing:
        closing = {"request": "妥否，请批示。", "report": "特此报告。"}.get(args.doc_type)
        body_text = "\n".join(paragraphs)
        if closing and closing not in body_text:
            add_body_paragraph(doc, closing)

    if not args.format_only:
        add_attachment_block(doc, args.attachment)
        add_attachment_markers(doc, args.attachment)
        add_signature(doc, args.issuer, args.date, args.contact, args.doc_type)
        add_disclosure_note(doc, args.disclosure_note)
        add_copy_block(doc, args.copy_to, args.print_org, args.print_date)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", help="UTF-8 .txt or .docx input file")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--doc-type", default="generic", choices=["generic", "request", "report", "letter", "notice"])
    parser.add_argument("--format-only", action="store_true", help="Do not add metadata, endings, attachments, or signature blocks")
    parser.add_argument("--title")
    parser.add_argument("--recipient")
    parser.add_argument("--issuer")
    parser.add_argument("--date")
    parser.add_argument("--outgoing-no")
    parser.add_argument("--red-header", help="Red official-document header text, e.g. 某单位文件")
    parser.add_argument("--letterhead", help="Agency letterhead text for formal letters, rendered in red at the top")
    parser.add_argument("--signer")
    parser.add_argument("--contact")
    parser.add_argument("--disclosure-note", help="Disclosure note such as （此件不公开）")
    parser.add_argument("--secret-level", help="Secret level line, e.g. 秘密★1年 (top-left, 三号黑体)")
    parser.add_argument("--attachment", action="append", default=[])
    parser.add_argument("--copy-to")
    parser.add_argument("--print-org")
    parser.add_argument("--print-date")
    parser.add_argument("--add-closing", action="store_true", help="Add standard 请示/报告 closing sentence")
    parser.add_argument("--no-page-number", action="store_true")
    parser.add_argument(
        "--letter-special-margins",
        action="store_true",
        help="Use legacy/special letter first-page margins: top 5.1 cm, bottom 2.1 cm",
    )
    parser.add_argument(
        "--promote-first-paragraph-title",
        action="store_true",
        help="Format the first input paragraph as the centered official title without changing its text",
    )
    parser.add_argument(
        "--bold-headings",
        action="store_true",
        help="Bold level-1 (黑体) headings; off by default to match the standard's 三号黑体 (no bold) convention",
    )
    args = parser.parse_args()
    build_doc(args)


if __name__ == "__main__":
    main()
