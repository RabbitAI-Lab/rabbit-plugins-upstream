"""
招标文件格式解析器

从招标文件 (.docx) 中提取：
1. 章节结构（大纲）
2. 页面设置
3. 字体格式信息
4. 匹配最接近的编号体系
5. 输出 JSON 配置，可直接喂给 generate_bid_template.js

用法:
    python parse_bidding_docx.py <input.docx> [--output config.json]
"""

import sys
import os
import json
import re
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.style import WD_STYLE_TYPE


def extract_headings(doc):
    """提取文档大纲：按出现顺序提取 Heading 段落"""
    headings = []
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            try:
                level = int(para.style.name.replace('Heading', '').strip())
            except ValueError:
                level = 1
            text = para.text.strip()
            if text:
                headings.append({
                    'level': level,
                    'title': text
                })
    return headings


def detect_numbering_scheme(headings):
    """
    从标题文本中检测最匹配的编号体系。
    返回: scheme1 | scheme2 | scheme3 | scheme4 | unknown
    """
    if not headings:
        return 'unknown'

    # 取前几个 H1 标题分析编号模式
    h1_samples = [h['title'] for h in headings if h['level'] == 1][:3]
    if not h1_samples:
        return 'unknown'

    first = h1_samples[0]

    # 体系四：第一章、第二章……
    if re.match(r'^第[一二三四五六七八九十百零]+[章节]', first):
        return 'scheme4'

    # 体系三：一、二、三……
    if re.match(r'^[一二三四五六七八九十百零]+[、，,.]', first):
        return 'scheme3'

    # 体系一或二：1. 2. 3.……
    if re.match(r'^\d+[\.\、\s]', first):
        # 需要看四级标题判断是体系一还是体系二
        h4_samples = [h['title'] for h in headings if h['level'] == 4][:2]
        if h4_samples:
            if re.match(r'^\(\d+\)', h4_samples[0]):
                return 'scheme2'
            if re.match(r'^\d+\.\d+\.\d+\.\d+', h4_samples[0]):
                return 'scheme1'
        # 默认推荐体系一
        return 'scheme1'

    return 'unknown'


def detect_font_from_run(run):
    """提取一个 run 的字体信息"""
    info = {}
    if run.font.name:
        info['name'] = run.font.name
    if run.font.size:
        info['size'] = run.font.size.pt
    if run.font.bold:
        info['bold'] = True
    return info


def extract_format_info(doc):
    """提取文档中的格式信息"""
    fmt = {
        'page': {},
        'styles': {}
    }

    # 页面设置（从第一个 section）
    if doc.sections:
        sec = doc.sections[0]
        fmt['page'] = {
            'width': sec.page_width.emu if sec.page_width else None,
            'height': sec.page_height.emu if sec.page_height else None,
            'margin_top': sec.top_margin.emu if sec.top_margin else None,
            'margin_bottom': sec.bottom_margin.emu if sec.bottom_margin else None,
            'margin_left': sec.left_margin.emu if sec.left_margin else None,
            'margin_right': sec.right_margin.emu if sec.right_margin else None,
        }

    # 样式信息
    style_targets = ['Normal', 'Heading 1', 'Heading 2', 'Heading 3',
                     'Heading 4', 'Heading 5']
    for style_name in style_targets:
        try:
            style = doc.styles[style_name]
            info = {}
            if style.font:
                if style.font.name:
                    info['font'] = style.font.name
                if style.font.size:
                    info['size'] = style.font.size.pt
                info['bold'] = style.font.bold or False
            fmt['styles'][style_name] = info
        except KeyError:
            pass

    # 从正文段落中采样字体
    font_samples = {}
    for para in doc.paragraphs[:50]:
        for run in para.runs:
            if run.font.name:
                name = run.font.name
                font_samples[name] = font_samples.get(name, 0) + 1

    if font_samples:
        fmt['font_samples'] = dict(sorted(
            font_samples.items(), key=lambda x: -x[1]
        )[:5])

    return fmt


def normalize_headings_for_config(headings, scheme):
    """
    将提取的标题转换为配置格式。
    去掉原有编号，保留纯标题文字，后续由生成脚本按所选体系重新编号。

    关键：按「最长匹配优先」顺序剥离各级编号前缀，避免先剥单级 "1."
    把 "1.1 标题" 误处理成 "1 标题" 残留悬空数字（P0-2 修复）。
    """
    cleaned = []
    for h in headings:
        text = h['title']
        # 五级数字编号：1.1.1.1.1
        text = re.sub(r'^\d+(?:\.\d+){4}[\.\、\s]\s*', '', text)
        # 四级数字编号：1.1.1.1
        text = re.sub(r'^\d+(?:\.\d+){3}[\.\、\s]\s*', '', text)
        # 三级数字编号：1.1.1
        text = re.sub(r'^\d+(?:\.\d+){2}[\.\、\s]\s*', '', text)
        # 二级数字编号：1.1
        text = re.sub(r'^\d+(?:\.\d+){1}[\.\、\s]\s*', '', text)
        # 体系四：第一章 / 第一节 XXX → XXX
        text = re.sub(r'^第[一二三四五六七八九十百零]+[章节]\s*', '', text)
        # 体系四 H4 / 体系三 H2 全角括号：（一）（二）XXX → XXX
        # （修复：原仅剥离 ASCII (1)，导致 scheme4 H4 全角（一）残留，生成器重加后出现双重前缀）
        text = re.sub(r'^（[一二三四五六七八九十百零]+）\s*', '', text)
        # 体系三：一、 二、 XXX → XXX
        text = re.sub(r'^[一二三四五六七八九十百零]+[、，,.\s]\s*', '', text)
        # 一级数字编号：1. 1、 1 XXX → XXX
        text = re.sub(r'^\d+[\.\、\s]\s*', '', text)
        # 体系二四级：(1) XXX → XXX
        text = re.sub(r'^\(\d+\)\s*', '', text)
        # 体系二五级：① XXX → XXX
        text = re.sub(r'^①\s*', '', text)
        text = text.strip()
        if text:
            cleaned.append({
                'level': h['level'],
                'title': text
            })
    return cleaned


def main():
    if len(sys.argv) < 2:
        print("用法: python parse_bidding_docx.py <input.docx> [--output config.json]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = None
    if '--output' in sys.argv:
        idx = sys.argv.index('--output')
        if idx + 1 < len(sys.argv):
            output_path = sys.argv[idx + 1]

    if not os.path.exists(input_path):
        print(f"[ERROR] 文件不存在: {input_path}")
        sys.exit(1)

    print(f"[INFO] 解析招标文件: {input_path}")
    doc = Document(input_path)

    # 1. 提取章节结构
    headings = extract_headings(doc)
    print(f"[INFO] 提取到 {len(headings)} 个标题段落")

    if not headings:
        print("[WARN] 未检测到标题段落（可能文件未使用样式标记标题）")
        cleaned = []
    else:
        # 2. 检测编号体系
        scheme = detect_numbering_scheme(headings)
        print(f"[INFO] 检测到编号体系: {scheme}")

        # 3. 清理标题
        cleaned = normalize_headings_for_config(headings, scheme)

    # 4. 提取格式信息
    fmt = extract_format_info(doc)

    # 5. 生成配置
    config = {
        'source_file': os.path.basename(input_path),
        'numbering': scheme,            # 生成器读取此键（P0-1 修复：此前两端键名不一致导致体系丢失）
        'numbering_scheme': scheme,     # 保留别名，便于人工阅读与向下兼容
        'chapters': cleaned,
        'format': fmt,
    }

    # 输出
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        print(f"[OK] 配置已保存: {output_path}")
    else:
        print("\n[OUTPUT] 生成的配置:\n")
        print(json.dumps(config, ensure_ascii=False, indent=2))

    print(f"\n[INFO] 章节总数: {len(cleaned)}")
    if cleaned:
        print(f"[INFO] 首章节: [{cleaned[0]['level']}] {cleaned[0]['title']}")
        print(f"[INFO] 末章节: [{cleaned[-1]['level']}] {cleaned[-1]['title']}")


if __name__ == '__main__':
    main()
