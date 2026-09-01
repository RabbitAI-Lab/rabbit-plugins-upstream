#!/usr/bin/env python3
"""
doc_parser.py - Extract plain text from bid documents (BidHunter v2.0, B1).

Supports PDF and DOCX. Uses the best available backend, gracefully degrading:
  1. PyPDF2 (if installed) for PDF
  2. `pdftotext` (poppler CLI, if available) for PDF
  3. python-docx (if installed) for DOCX
  4. last-resort regex stream extraction (PDF) — partial text only

This keeps the skill dependency-free by default; users who want best-in-class
parsing just `pip install PyPDF2 python-docx` (optional).

Usage:
  python3 doc_parser.py <file.pdf|file.docx>   # prints extracted text to stdout
"""
import os
import sys
import re


def extract_pdf(path):
    # 1. PyPDF2
    try:
        import PyPDF2
        r = PyPDF2.PdfReader(path)
        parts = [p.extract_text() or "" for p in r.pages]
        txt = "\n".join(parts)
        if txt.strip():
            return txt, "PyPDF2"
    except Exception:
        pass
    # 2. pdftotext CLI
    try:
        import subprocess
        r = subprocess.run(["pdftotext", "-layout", path, "-"],
                           capture_output=True, text=True, timeout=60)
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout, "pdftotext"
    except Exception:
        pass
    # 3. regex fallback (between stream/endstream, decode printable runs)
    try:
        with open(path, "rb") as f:
            data = f.read()
        chunks = re.findall(rb"stream\r?\n(.*?)\r?\nendstream", data, re.S)
        texts = []
        for c in chunks:
            try:
                s = c.decode("utf-8", "ignore")
            except Exception:
                continue
            s = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", s)
            if len(s.strip()) > 20:
                texts.append(s)
        if texts:
            return "\n".join(texts), "regex-fallback"
    except Exception:
        pass
    return "", "none"


def extract_docx(path):
    try:
        import docx
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        # tables
        for t in d.tables:
            for row in t.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        txt = "\n".join(parts)
        if txt.strip():
            return txt, "python-docx"
    except Exception:
        pass
    return "", "none"


def parse(path):
    if not os.path.exists(path):
        return "", "missing"
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_pdf(path)
    if ext in (".docx", ".doc"):
        if ext == ".doc":
            return "", "unsupported-doc"
        return extract_docx(path)
    # plain text
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read(), "text"
    except Exception:
        return "", "none"


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 doc_parser.py <file>", file=sys.stderr)
        sys.exit(1)
    txt, backend = parse(sys.argv[1])
    if not txt:
        print(f"[doc_parser] 未能提取文本（backend={backend}）。"
              f"如需更好效果：pip install PyPDF2 python-docx，或安装 poppler 的 pdftotext。",
              file=sys.stderr)
        sys.exit(1)
    print(txt)


if __name__ == "__main__":
    main()
