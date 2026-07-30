#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
diagrams.py — word-formatter 图件处理模块

能力：
  1. render   把"结构化规格(structured)"或"mermaid 文本"渲染为 PNG 图片
  2. insert   把图件按占位符/标题插入到 docx，自动加分章编号图题（图1-1）
  3. scan     识别 docx 中已有的图片（图件识别）
  4. renumber 重排图题编号，并修正正文中的交叉引用（见图1-1 -> 图2-3）

图表类型覆盖：流程图(flowchart)、股权架构图(equity/hierarchy)、
数据流向图(data-flow)、功能说明图(feature) —— 本质上都是 节点+连线 的有向图，
统一用 structured 引擎渲染；mermaid 引擎作为其更精美的可选后端。

依赖：python-docx, matplotlib, pillow；可选 mermaid-cli(npx @mermaid-js/mermaid-cli)。
"""

import argparse
import copy
import glob
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from collections import defaultdict

# ---------- docx ----------
from docx import Document
from docx.shared import Cm, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- matplotlib (结构化渲染) ----------
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, Rectangle
from matplotlib import font_manager

# ----------------------------------------------------------------------------
# 字体：优先使用系统中的中文黑体（图表标签用黑体，正式、清晰）
# ----------------------------------------------------------------------------
_CANDIDATE_FONTS = [
    "/System/Library/Fonts/STHeiti Medium.ttc",
    "/System/Library/Fonts/STHeiti Light.ttc",
    "/Library/Fonts/STHeiti Medium.ttc",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
]
_FONT_PATH = None
for _f in _CANDIDATE_FONTS:
    if os.path.exists(_f):
        _FONT_PATH = _f
        break

if _FONT_PATH:
    _FONT_PROP = font_manager.FontProperties(fname=_FONT_PATH)
    matplotlib.rcParams["font.family"] = _FONT_PROP.get_name()
else:
    _FONT_PROP = font_manager.FontProperties()
    print("[warn] 未找到系统中文字体，图表中的中文可能显示为方框。", file=sys.stderr)

matplotlib.rcParams["axes.unicode_minus"] = False


# ----------------------------------------------------------------------------
# 中文换行
# ----------------------------------------------------------------------------
def _wrap_cn(text, width=8):
    """按字符数把中文标签折行（中文无空格，按 width 切分）。"""
    text = str(text)
    if "\n" in text:
        return [line for seg in text.split("\n") for line in _wrap_cn(seg, width)]
    lines = []
    cur = ""
    for ch in text:
        cur += ch
        if len(cur) >= width:
            lines.append(cur)
            cur = ""
    if cur:
        lines.append(cur)
    return lines or [""]


# ----------------------------------------------------------------------------
# 有向图分层布局（Sugiyama-lite）
# ----------------------------------------------------------------------------
def _layout(nodes, edges, direction="TB"):
    ids = [n["id"] for n in nodes]
    idset = set(ids)
    adj = defaultdict(list)
    radj = defaultdict(list)
    for e in edges:
        if e["from"] in idset and e["to"] in idset:
            adj[e["from"]].append(e["to"])
            radj[e["to"]].append(e["from"])

    # 入度 -> 根节点
    indeg = {i: len(radj[i]) for i in ids}
    roots = [i for i in ids if indeg[i] == 0]
    if not roots:  # 存在环，任选一个起点
        roots = [ids[0]]

    # 最长路径分层（Bellman-Ford 风格松弛）
    layer = {i: 0 for i in ids}
    for _ in range(len(ids)):
        changed = False
        for e in edges:
            if e["from"] in idset and e["to"] in idset:
                if layer[e["to"]] < layer[e["from"]] + 1:
                    layer[e["to"]] = layer[e["from"]] + 1
                    changed = True
        if not changed:
            break

    by_layer = defaultdict(list)
    for i in ids:
        by_layer[layer[i]].append(i)
    max_layer = max(by_layer) if by_layer else 0

    # 层内排序（重心法减少交叉，单遍）
    order = {}
    for l in range(max_layer + 1):
        ns = by_layer[l]
        if l == 0:
            ns.sort()
        else:
            def bary(n):
                preds = radj[n]
                if not preds:
                    return 0.0
                return sum(order.get(p, 0) for p in preds) / len(preds)
            ns.sort(key=bary)
        for idx, n in enumerate(ns):
            order[n] = idx

    # 坐标（TB：x=层内序号，y=-层）
    pos = {}
    for l in range(max_layer + 1):
        ns = by_layer[l]
        n = len(ns)
        for idx, i in enumerate(ns):
            x = idx - (n - 1) / 2.0
            y = -l
            pos[i] = (x, y)
    if direction == "LR":
        # 层号 → X 轴（从左到右递增）；层内序号 → Y 轴
        pos = {i: (-y, x) for i, (x, y) in pos.items()}
    return pos, by_layer


def _box_size(label, direction="TB"):
    lines = _wrap_cn(label, width=8)
    w_chars = max(len(l) for l in lines)
    h = max(1.0, 0.55 * len(lines) + 0.45)
    w = max(2.0, 0.95 * w_chars + 0.8)
    if direction == "LR":
        return h, w  # 交换
    return w, h


def _trim_to_box(center, other, half):
    """把 center->other 这条线的端点裁剪到 other 所在矩形边框上。"""
    cx, cy = center
    ox, oy = other
    dx, dy = ox - cx, oy - cy
    if dx == 0 and dy == 0:
        return other
    hw, hh = half
    tx = hw / abs(dx) if dx != 0 else float("inf")
    ty = hh / abs(dy) if dy != 0 else float("inf")
    t = min(tx, ty)
    return (ox - dx * t, oy - dy * t)


def render_structured(spec, out_path, style=None):
    """把 structured 规格（nodes/edges/direction）渲染为 PNG。"""
    style = style or {}
    direction = spec.get("direction", "TB")
    nodes = spec.get("nodes", [])
    edges = spec.get("edges", [])
    titles = spec.get("title", "")

    pos, _ = _layout(nodes, edges, direction)
    fig, ax = plt.subplots(figsize=(10, 7), dpi=150)
    ax.set_axis_off()
    ax.set_aspect("equal")

    # 极度克制的黑灰白配色：浅灰填充 + 深灰描边，无任何彩色
    node_color = style.get("node_fill", "#F2F2F2")
    node_edge = style.get("node_edge", "#404040")
    edge_color = style.get("edge_color", "#595959")
    font_color = style.get("font_color", "#1A1A1A")

    sizes = {}
    for n in nodes:
        w, h = _box_size(n.get("label", n["id"]), direction)
        sizes[n["id"]] = (w, h)

    # 先画线，后画框，保证框在上层
    for e in edges:
        a, b = e["from"], e["to"]
        if a not in pos or b not in pos:
            continue
        wa, ha = sizes[a]
        wb, hb = sizes[b]
        sa = _trim_to_box(pos[a], pos[b], (wa / 2, ha / 2))
        eb = _trim_to_box(pos[b], pos[a], (wb / 2, hb / 2))
        arrow = FancyArrowPatch(
            sa, eb,
            arrowstyle="-|>", mutation_scale=14,
            color=edge_color, lw=1.6, zorder=1,
        )
        ax.add_patch(arrow)
        lbl = e.get("label", "")
        if lbl:
            mx = (sa[0] + eb[0]) / 2
            my = (sa[1] + eb[1]) / 2
            ax.text(mx, my, lbl, fontsize=8, color=edge_color,
                    ha="center", va="center",
                    bbox=dict(boxstyle="round,pad=0.15", fc="white", ec="none", alpha=0.85),
                    zorder=3, fontproperties=_FONT_PROP)

    for n in nodes:
        x, y = pos[n["id"]]
        w, h = sizes[n["id"]]
        rect = Rectangle((x - w / 2, y - h / 2), w, h,
                         facecolor=node_color, edgecolor=node_edge,
                         linewidth=1.4, zorder=2)
        ax.add_patch(rect)
        ax.text(x, y, "\n".join(_wrap_cn(n.get("label", n["id"]), width=8)),
                fontsize=10.5, color=font_color, ha="center", va="center",
                zorder=3, fontproperties=_FONT_PROP)

    # 坐标轴范围
    xs = [p[0] for p in pos.values()]
    ys = [p[1] for p in pos.values()]
    if xs and ys:
        padx = max(sizes[i][0] for i in sizes) / 2 + 0.6
        pady = max(sizes[i][1] for i in sizes) / 2 + 0.6
        ax.set_xlim(min(xs) - padx, max(xs) + padx)
        ax.set_ylim(min(ys) - pady, max(ys) + pady)

    if titles:
        ax.set_title(titles, fontsize=12, fontproperties=_FONT_PROP,
                     color="#1A1A1A", pad=10)

    fig.savefig(out_path, bbox_inches="tight", transparent=True,
                dpi=150, pad_inches=0.15)
    plt.close(fig)
    return out_path


def render_mermaid(spec_text, out_path):
    """用 mermaid-cli 渲染。需要 npx @mermaid-js/mermaid-cli。"""
    node_bin = "/Users/jackwang/.workbuddy/binaries/node/versions/22.22.2/bin/node"
    npx = "/Users/jackwang/.workbuddy/binaries/node/versions/22.22.2/bin/npx"
    mmd_js = "/Users/jackwang/.workbuddy/binaries/node/workspace/node_modules/.bin/mmdc"
    if not os.path.exists(mmd_js):
        raise RuntimeError(
            "未找到 mermaid-cli。请先在隔离 node 工作区安装：\n"
            "  cd /Users/jackwang/.workbuddy/binaries/node/workspace && "
            "npm install @mermaid-js/mermaid-cli\n"
            "或把该图改用 engine=structured（纯 Python，离线可用）。"
        )
    with tempfile.NamedTemporaryFile("w", suffix=".mmd", delete=False, encoding="utf-8") as f:
        f.write(spec_text)
        mmd_file = f.name
    try:
        out_dir = os.path.dirname(os.path.abspath(out_path))
        out_name = os.path.basename(out_path)
        env = dict(os.environ)
        env["PUPPETEER_SKIP_DOWNLOAD"] = "1"
        cmd = [node_bin, mmd_js, "-i", mmd_file, "-o", out_name,
               "-b", "transparent", "-t", "default"]
        # mmdc 在工作目录输出
        subprocess.run(cmd, cwd=out_dir, env=env, check=True,
                       stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    finally:
        os.unlink(mmd_file)
    return out_path


def render_diagram(spec, out_path, style=None):
    """按 spec['engine'] 选择渲染后端。返回图片路径。"""
    engine = spec.get("engine", "structured")
    if engine == "mermaid":
        return render_mermaid(spec.get("mermaid", ""), out_path)
    return render_structured(spec, out_path, style)


# ----------------------------------------------------------------------------
# docx 辅助：在指定段落之后插入新段落
# ----------------------------------------------------------------------------
def _insert_para_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def _is_heading1(p):
    if p.style is None:
        return False
    name = p.style.name or ""
    if name == "Heading 1":
        return True
    # 兼容直接用 outlineLvl 的情况
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None:
        ol = pPr.find(qn("w:outlineLvl"))
        if ol is not None and ol.get(qn("w:val")) == "0":
            return True
    return False


def _chapter_of(anchor_element, doc):
    """统计 anchor 之前出现的 Heading 1 数量，得到章节号（从1开始）。"""
    chap = 0
    for p in doc.paragraphs:
        if _is_heading1(p):
            chap += 1
        if p._p is anchor_element:
            break
    return max(chap, 1)


def _find_anchor(doc, target):
    """target 形如：
       'heading:三、数据处理流程'  -> 该标题段落之后
       'placeholder:{{DIAGRAM:flow}}' -> 替换为该占位段落
       'append' -> 文档末尾
    返回 (anchor_paragraph, mode)。mode='after' 或 'replace'。
    """
    if target == "append":
        return doc.paragraphs[-1] if doc.paragraphs else None, "after"
    if target.startswith("heading:"):
        key = target[len("heading:"):].strip()
        for p in doc.paragraphs:
            if key in (p.text or "").strip():
                return p, "after"
        raise ValueError(f"未找到标题包含「{key}」的段落，无法插入图件。")
    if target.startswith("placeholder:"):
        token = target[len("placeholder:"):].strip()
        for p in doc.paragraphs:
            if token in (p.text or ""):
                return p, "replace"
        raise ValueError(f"未找到占位符「{token}」，无法插入图件。")
    # 默认当 heading 处理
    for p in doc.paragraphs:
        if target in (p.text or "").strip():
            return p, "after"
    raise ValueError(f"未找到插入锚点：「{target}」")


# ----------------------------------------------------------------------------
# 图件插入 + 分章编号
# ----------------------------------------------------------------------------
def insert_diagrams(docx_path, specs_path, config_path, out_path=None):
    """读取 specs.json（数组），把每个图件渲染并插入 docx。"""
    doc = Document(docx_path)
    with open(specs_path, encoding="utf-8") as f:
        specs = json.load(f)
    with open(config_path, encoding="utf-8") as f:
        cfg = json.load(f)
    fig_cfg = cfg.get("figures", {})
    width_cm = float(fig_cfg.get("width_cm", 15))
    caption_font = fig_cfg.get("caption_font_cn", "宋体")
    caption_size = float(fig_cfg.get("caption_size_pt", 9))

    out_path = out_path or _default_out(docx_path)
    tmp_doc = Document(docx_path)  # 用于定位（避免边插边遍历混乱）

    # 记录每个图件渲染出的临时 PNG
    rendered = []
    for spec in specs:
        diagram_spec = spec.get("spec", spec)
        ext = ".png"
        fd, png = tempfile.mkstemp(suffix=ext)
        os.close(fd)
        render_diagram(diagram_spec, png, style=fig_cfg.get("style"))
        rendered.append((spec, png))

    # 逐个插入（基于原始 doc 重新定位，避免索引漂移）
    doc2 = Document(docx_path)
    fig_counter = 0
    for spec, png in rendered:
        target = spec.get("target", "append")
        anchor, mode = _find_anchor(doc2, target)
        if anchor is None:
            anchor = doc2.paragraphs[-1]
            mode = "after"
        # 用临时占位编号（不含章节），renumber 会统一重排为正确的分章编号
        fig_counter += 1
        caption_text = f"图-{fig_counter} {spec.get('caption', '')}"

        if mode == "replace":
            # 清空占位段落，作为图件段落
            for r in list(anchor.runs):
                r._r.getparent().remove(r._r)
            img_p = anchor
        else:
            img_p = _insert_para_after(anchor)
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_p.add_run()
        w = float(spec.get("width_cm", width_cm))
        run.add_picture(png, width=Cm(w))
        if spec.get("page_break"):
            _add_page_break_before(img_p)

        cap_p = _insert_para_after(img_p)
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap_p.add_run(caption_text)
        cr.font.name = caption_font
        cr.font.size = Pt(caption_size)
        rpr = cr._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), caption_font)
        cr.font.bold = False

    doc2.save(out_path)
    for _, png in rendered:
        try:
            os.unlink(png)
        except OSError:
            pass

    # 统一重排编号 + 修正交叉引用
    renumber_figures(out_path, config_path)
    print(f"[ok] 已插入 {len(rendered)} 个图件 -> {out_path}")
    return out_path


def _add_page_break_before(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    pPr.append(br)


def _default_out(path):
    base, ext = os.path.splitext(path)
    return f"{base}_figs{ext}"


# ----------------------------------------------------------------------------
# 图件识别（scan）
# ----------------------------------------------------------------------------
def scan_docx(docx_path):
    """识别 docx 中已有图片，输出 Markdown 清单。"""
    doc = Document(docx_path)
    rows = []
    for idx, p in enumerate(doc.paragraphs):
        drawings = p._p.findall(".//" + qn("w:drawing"))
        if not drawings:
            continue
        # 读取图片尺寸
        extents = p._p.findall(".//" + qn("wp:extent"))
        w = h = None
        if extents:
            w = extents[0].get("cx")
            h = extents[0].get("cy")
        # 后继段落当作图题
        caption = ""
        if idx + 1 < len(doc.paragraphs):
            nxt = doc.paragraphs[idx + 1].text.strip()
            if nxt.startswith("图"):
                caption = nxt
        cm_w = (int(w) / 360000.0) if w else None
        rows.append((idx + 1, cm_w, caption))
    print(f"# 图件识别结果：{os.path.basename(docx_path)}\n")
    if not rows:
        print("（未发现图片）")
        return
    print("| 序号 | 段落位置 | 宽度(cm) | 图题 |")
    print("|---|---|---|---|")
    for i, (pos, w, cap) in enumerate(rows, 1):
        wtxt = f"{w:.1f}" if w is not None else "N/A"
        print(f"| {i} | 第{pos}段 | {wtxt} | {cap or '（无图题）'} |")


# ----------------------------------------------------------------------------
# 图题重编号 + 交叉引用一致性
# ----------------------------------------------------------------------------
_CAP_RE = re.compile(r"^图\s*(\d+)[-－]\s*(\d+)\s*(.*)$")
_CAP_RE_INLINE = re.compile(r"图\s*(\d+)[-－]\s*(\d+)")
_CAP_RE_PLAIN = re.compile(r"^图\s*(\d+)\s*(.*)$")  # 无分章：图1
_CAP_RE_PROVISIONAL = re.compile(r"^图\s*-\s*(\d+)\s+(.*)$")  # 临时占位：图-1 标题


def renumber_figures(docx_path, config_path):
    """重排图题编号（分章 图X-Y），并修正正文交叉引用。

    约定（源文档写作契约）：图题与正文交叉引用统一使用「顺序占位编号」
    图-N（N 为全文档第 N 个图件，由插入流程自动生成）。本函数将其统一
    转换为分章编号「图{章}-{章内序号}」，并保证图题与正文引用一一对应。
    """
    with open(config_path, encoding="utf-8") as f:
        json.load(f)  # 预留：后续可按 config 调整编号/样式

    doc = Document(docx_path)
    cur_chap = 0
    # 第一遍：按文档顺序收集所有图题（临时占位格式 图-N）
    captions = []  # (para, src_idx, chap, title)
    for p in doc.paragraphs:
        if _is_heading1(p):
            cur_chap += 1
        m = _CAP_RE_PROVISIONAL.match(p.text.strip())
        if m:
            oi = int(m.group(1))
            chap = max(cur_chap, 1)
            captions.append((p, oi, chap, m.group(2).strip()))

    # 第二遍：按章节顺序分配最终编号，建立 顺序序号 -> (章, 章内序号)
    per_chapter = defaultdict(int)
    src_to_final = {}
    for p, oi, chap, title in captions:
        per_chapter[chap] += 1
        src_to_final[oi] = (chap, per_chapter[chap])

    # 重写图题
    for p, oi, chap, title in captions:
        fchap, fidx = src_to_final[oi]
        new_text = f"图{fchap}-{fidx} {title}".rstrip()
        if p.runs:
            first = p.runs[0]
            first.text = new_text
            for r in p.runs[1:]:
                r._r.getparent().remove(r._r)
        else:
            p.add_run(new_text)

    # 修正交叉引用（正文中出现的 图-N，N 为顺序序号）
    ref_re = re.compile(r"图\s*-\s*(\d+)")
    cap_set = set(cp for cp, _, _, _ in captions)
    for p in doc.paragraphs:
        if p in cap_set:
            continue  # 图题本身不替换
        if "图" not in p.text:
            continue

        def repl(m, _map=src_to_final):
            k = int(m.group(1))
            if k in _map:
                c, n = _map[k]
                return f"图{c}-{n}"
            return m.group(0)

        new_text = ref_re.sub(repl, p.text)
        if new_text != p.text:
            if p.runs:
                p.runs[0].text = new_text
                for r in p.runs[1:]:
                    r._r.getparent().remove(r._r)
            else:
                p.add_run(new_text)

    doc.save(docx_path)
    return src_to_final


# ----------------------------------------------------------------------------
# 命令行
# ----------------------------------------------------------------------------
def main():
    ap = argparse.ArgumentParser(description="word-formatter 图件处理")
    sub = ap.add_subparsers(dest="cmd", required=True)

    p_rend = sub.add_parser("render", help="渲染单个图件为 PNG")
    p_rend.add_argument("spec", help="规格 JSON 文件（含 engine/spec 或 mermaid）")
    p_rend.add_argument("out", help="输出 PNG 路径")
    p_rend.add_argument("--style", help="样式 JSON（可选）")

    p_ins = sub.add_parser("insert", help="把图件插入 docx（自动分章编号）")
    p_ins.add_argument("docx")
    p_ins.add_argument("specs", help="图件规格 JSON 数组文件")
    p_ins.add_argument("config", help="文档类型配置 JSON")
    p_ins.add_argument("-o", "--out", default=None)

    p_scan = sub.add_parser("scan", help="识别 docx 中已有图片")
    p_scan.add_argument("docx")

    p_ren = sub.add_parser("renumber", help="重排图题编号并修正交叉引用")
    p_ren.add_argument("docx")
    p_ren.add_argument("config")

    args = ap.parse_args()
    if args.cmd == "render":
        with open(args.spec, encoding="utf-8") as f:
            spec = json.load(f)
        style = None
        if args.style:
            with open(args.style, encoding="utf-8") as f:
                style = json.load(f)
        out = render_diagram(spec, args.out, style)
        print(f"[ok] 已渲染 -> {out}")
    elif args.cmd == "insert":
        insert_diagrams(args.docx, args.specs, args.config, args.out)
    elif args.cmd == "scan":
        scan_docx(args.docx)
    elif args.cmd == "renumber":
        renumber_figures(args.docx, args.config)
        print("[ok] 已重排图题编号并修正交叉引用。")


if __name__ == "__main__":
    main()
