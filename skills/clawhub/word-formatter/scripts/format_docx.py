#!/usr/bin/env python3
"""按 JSON 配置对 .docx 文档进行后期排版。

用法:
    python format_docx.py <input.docx> <config.json> [-o output.docx]

处理内容:
    1. 页面设置: 纸张 A4、页边距
    2. 样式重定义: Normal / Heading 1-3 / Note (中西文字体、字号、行距、缩进、对齐、段前后、颜色)
    3. 段落级兜底: 对使用直接格式的正文段落, 覆盖字体/颜色与缩进
    4. 表格: 统一单元格字体字号、表头加粗、三线表/全线表边框、去除斑马纹与彩色填充
    5. 清理: 全角空格、连续空段
输出文件永不覆盖输入文件。

设计语言（正式报告 / 全类型统一）：
    - 配色极度克制，仅黑/白/灰三色，不使用任何彩色（含深蓝、深红等均禁用）
    - 强调处一律用加粗或灰度层次区分，绝不用彩色
    - 表格采用三线表或全线表，表头加粗、行距均匀、无填色、无彩色
    - 统一生成页眉（报告标题）与页脚页码（第 N 页 共 M 页；公文按 GB/T 9704 用 —N— 单双页对齐）
"""
import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.dml import MSO_COLOR_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

HEADING_STYLE_ALIASES = {
    "Heading 1": ["Heading 1", "heading 1", "1", "标题 1"],
    "Heading 2": ["Heading 2", "heading 2", "2", "标题 2"],
    "Heading 3": ["Heading 3", "heading 3", "3", "标题 3"],
}


def _rgb(hex_str):
    """把 '1A1A1A' 转为 RGBColor；空值返回 None。"""
    if not hex_str:
        return None
    try:
        return RGBColor.from_string(hex_str.replace("#", ""))
    except ValueError:
        return None


def set_run_font(run, font_cn, font_en, size_pt, bold, color_rgb=None, enforce_color=False):
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)
    rfonts.set(qn("w:eastAsia"), font_cn)
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    color = _rgb(color_rgb)
    if color is not None:
        # 尊重作者已有的强调色（深蓝/深红等），除非配置要求强制覆盖
        has_explicit = False
        try:
            if run.font.color is not None and run.font.color.type == MSO_COLOR_TYPE.RGB:
                has_explicit = True
        except Exception:
            has_explicit = False
        if enforce_color or not has_explicit:
            run.font.color.rgb = color


def apply_style_def(style, cfg):
    f = style.font
    f.name = cfg["font_en"]
    f.size = Pt(cfg["size_pt"])
    f.bold = cfg.get("bold", False)
    f.italic = cfg.get("italic", False)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), cfg["font_en"])
    rfonts.set(qn("w:hAnsi"), cfg["font_en"])
    rfonts.set(qn("w:eastAsia"), cfg["font_cn"])
    color = _rgb(cfg.get("color_rgb"))
    if color is not None:
        f.color.rgb = color
    pf = style.paragraph_format
    rule = cfg.get("line_spacing_rule", "multiple")
    if rule == "exact":
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(cfg["line_spacing_pt"])
    else:
        pf.line_spacing = cfg.get("line_spacing", 1.5)
    pf.space_before = Pt(cfg.get("space_before_pt", 0))
    pf.space_after = Pt(cfg.get("space_after_pt", 0))
    al = cfg.get("alignment", "justify")
    pf.alignment = ALIGN_MAP.get(al)
    indent_chars = cfg.get("first_line_indent_chars", 0)
    ppr = style.element.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = ppr.makeelement(qn("w:ind"), {})
        ppr.append(ind)
    if indent_chars:
        ind.set(qn("w:firstLineChars"), str(int(indent_chars * 100)))
        ind.set(qn("w:firstLine"), str(int(indent_chars * cfg["size_pt"] * 20)))
    else:
        for attr in ("w:firstLineChars", "w:firstLine"):
            if ind.get(qn(attr)) is not None:
                del ind.attrib[qn(attr)]


def setup_page(doc, page_cfg):
    m = page_cfg["margins_cm"]
    # 自定义纸张（申请文件准则要求 209×295mm，相当于 A4；默认 21.0×29.7cm）
    pw = page_cfg.get("page_width_cm", 21.0)
    ph = page_cfg.get("page_height_cm", 29.7)
    for sec in doc.sections:
        sec.page_width = Cm(pw)
        sec.page_height = Cm(ph)
        sec.top_margin = Cm(m["top"])
        sec.bottom_margin = Cm(m["bottom"])
        sec.left_margin = Cm(m["left"])
        sec.right_margin = Cm(m["right"])


def style_key_of(paragraph, custom_names):
    name = paragraph.style.name if paragraph.style is not None else "Normal"
    if name in custom_names:
        return name
    for canonical, aliases in HEADING_STYLE_ALIASES.items():
        if name in aliases:
            return canonical
    return "Normal"


def fix_paragraph_direct(paragraph, cfg):
    """对段落内所有 run 强制统一字体/颜色（兜底直接格式）。

    极度克制配色要求：默认 enforce_color=true，即一律按配置色（黑/深灰）覆盖，
    任何 stray 彩色（含深蓝/深红）都不会保留。如需保留个别手设颜色可设 enforce_color=false。
    """
    bold = cfg.get("bold", False) or None
    enforce = cfg.get("enforce_color", True)
    for run in paragraph.runs:
        set_run_font(run, cfg["font_cn"], cfg["font_en"], cfg["size_pt"], bold, cfg.get("color_rgb"), enforce)


# ----------------------------------------------------------------------------
# 表格边框（三线表 / 全线表）
# ----------------------------------------------------------------------------
def _set_border(el, edge, val, sz, color):
    """在 el 的 w:tcBorders / w:tblBorders 下设置某条边。"""
    borders = el.find(qn("w:tcBorders"))
    if borders is None:
        borders = el.makeelement(qn("w:tcBorders"), {})
        el.append(borders)
    tag = qn("w:" + edge)
    e = borders.find(tag)
    if e is None:
        e = borders.makeelement(tag, {})
        borders.append(e)
    e.set(qn("w:val"), val)
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "0")
    e.set(qn("w:color"), color)


def _clear_borders(el):
    borders = el.find(qn("w:tcBorders"))
    if borders is None:
        borders = el.makeelement(qn("w:tcBorders"), {})
        el.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = borders.find(qn("w:" + edge))
        if e is None:
            e = borders.makeelement(qn("w:" + edge), {})
            borders.append(e)
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")


def apply_table_borders(table, style):
    """按 style 设置全表外边框；三线表额外在表头下方加一条细线。"""
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    else:
        tblPr.remove(tblBorders)
        tblBorders = OxmlElement("w:tblBorders")

    blk = "#000000"
    if style == "three_line":
        edges = {
            "top": ("single", 12, blk),
            "bottom": ("single", 12, blk),
            "left": ("none", 0, "auto"),
            "right": ("none", 0, "auto"),
            "insideH": ("none", 0, "auto"),
            "insideV": ("none", 0, "auto"),
        }
    else:  # full_line
        edges = {
            "top": ("single", 6, blk),
            "bottom": ("single", 6, blk),
            "left": ("single", 6, blk),
            "right": ("single", 6, blk),
            "insideH": ("single", 4, blk),
            "insideV": ("single", 4, blk),
        }
    for edge, (val, sz, col) in edges.items():
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), val)
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), col)
        tblBorders.append(e)
    tblPr.append(tblBorders)

    # 三线表：表头（首行）下边加细线 0.75pt
    if style == "three_line" and len(table.rows) > 0:
        for cell in table.rows[0].cells:
            _set_border(cell._tc.get_or_add_tcPr(), "bottom", "single", 6, blk)


def clear_cell_shading(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is not None:
        tcPr.remove(shd)


def format_table(table, tbl_cfg):
    style = tbl_cfg.get("style")  # "three_line" / "full_line" / None
    if style:
        apply_table_borders(table, style)
    body_color = tbl_cfg.get("color_rgb")
    header_color = tbl_cfg.get("header_color_rgb", body_color)
    header_bold = tbl_cfg.get("header_bold", True)
    cell_align = tbl_cfg.get("alignment", "center")  # 表格内文字对齐
    align_enum = ALIGN_MAP.get(cell_align, WD_ALIGN_PARAGRAPH.CENTER)
    rows = table.rows
    for ri, row in enumerate(rows):
        is_header = (ri == 0)
        for cell in row.cells:
            clear_cell_shading(cell)
            for p in cell.paragraphs:
                # 关键修复：清除 Normal 样式继承的「首行缩进 2 字符」等段格式
                # 表格单元格应无缩进、无段前段后间距、独立对齐
                ppr = p._p.get_or_add_pPr()
                for tag in ("w:ind", "w:contextualSpacing"):
                    for el in list(ppr.findall(qn(tag))):
                        ppr.remove(el)
                # 显式清零首行缩进
                ind = ppr.find(qn("w:ind"))
                if ind is None:
                    ind = OxmlElement("w:ind")
                    ppr.append(ind)
                ind.set(qn("w:firstLineChars"), "0")
                ind.set(qn("w:firstLine"), "0")
                # 显式对齐（防止继承 Normal 的 justify）
                for el in list(ppr.findall(qn("w:jc"))):
                    ppr.remove(el)
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), cell_align if cell_align in ("left", "center", "right", "justify") else "center")
                ppr.append(jc)
                # 清除段前/段后间距
                sp = ppr.find(qn("w:spacing"))
                if sp is not None:
                    for attr in ("w:before", "w:after"):
                        if sp.get(qn(attr)):
                            sp.set(qn(attr), "0")
                # 给段落 run 设置字体
                for run in p.runs:
                    set_run_font(
                        run,
                        tbl_cfg["font_cn"],
                        tbl_cfg["font_en"],
                        tbl_cfg["size_pt"],
                        header_bold if is_header else None,
                        header_color if is_header else body_color,
                        enforce_color=True,
                    )


def add_section_dividers(doc, cfg):
    """在一级标题（节）前插入分页符，满足申请文件准则「节与节之间应有明显的分隔标识」。

    通过在 Heading 1 段落首部追加一个 w:br type=page 实现；仅当配置
    page_break_before_h1=true 时生效，默认为关闭（不影响其他配置）。
    """
    if not cfg.get("page_break_before_h1"):
        return
    n = 0
    for p in doc.paragraphs:
        if p.style.name != "Heading 1":
            continue
        # 避免重复插入
        for r in p.runs:
            if r._element.find(qn("w:br")) is not None:
                break
        else:
            run = p.add_run()
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            run._r.append(br)
            n += 1
    return n


def cleanup_doc(doc, cleanup_cfg):
    removed = 0
    if cleanup_cfg.get("remove_fullwidth_spaces"):
        for p in doc.paragraphs:
            for run in p.runs:
                if "\u3000" in run.text and p.style.name.startswith(("Heading", "标题")) is False:
                    if run.text.strip("\u3000 ") == "" and len(p.runs) == 1:
                        continue
                    run.text = re.sub(r"^\u3000+", "", run.text)
    if cleanup_cfg.get("normalize_empty_paragraphs"):
        prev_empty = False
        for p in list(doc.paragraphs):
            is_empty = not p.text.strip() and not p._p.findall(qn("w:r") + "/" + qn("w:drawing"))
            if is_empty and prev_empty:
                p._p.getparent().remove(p._p)
                removed += 1
            else:
                prev_empty = is_empty
    return removed


# ----------------------------------------------------------------------------
# 页眉 / 页脚 / 页码
# ----------------------------------------------------------------------------
def _add_field(paragraph, field_code):
    """在段落末尾追加一个 Word 域（如 PAGE / NUMPAGES）。"""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    return run


def _style_header_run(run, font_cn, font_en, size_pt, color_hex):
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)
    rfonts.set(qn("w:eastAsia"), font_cn)
    if size_pt:
        run.font.size = Pt(size_pt)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex.replace("#", ""))


def _set_para_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pBdr.append(bottom)


def add_header_footer(doc, hf):
    """按 header_footer 配置块统一加入页眉（左上客户名 + 右上报告名）、页脚页码。

    header_footer 结构（均为可选，缺失则用合理默认）：
        header_client     客户全称（非空则显示在页眉左上角）
        header_text       页眉右上角文字
        header_alignment  仅对纯文字生效；两栏模式自动取消
        header_font_cn/en/size_pt/color_rgb  页眉字体
        header_border     bool，页眉下方细线（默认 True）
        footer_mode       page_of / page_only / gb9704
        footer_alignment  默认 center
        footer_font_cn/en/size_pt/color_rgb
        show_page_number  bool（默认 True）
        different_odd_even bool
    """
    if not hf:
        return
    show_pn = hf.get("show_page_number", True)
    htext = (hf.get("header_text") or "").strip()
    hclient = (hf.get("header_client") or "").strip()

    # ---- 页眉（左上客户名 + 右上报告名） ----
    if htext or hclient:
        fcn = hf.get("header_font_cn", "宋体")
        fen = hf.get("header_font_en", "Times New Roman")
        fsz = hf.get("header_size_pt", 9)
        fco = hf.get("header_color_rgb", "#595959")
        for sec in doc.sections:
            hp = sec.header.paragraphs[0]
            hp.text = ""
            # 彻底重建 pPr：清空所有子元素后重新构建
            pPr = hp._p.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                hp._p.insert(0, pPr)
            # 清空 pPr 全部子元素（避免样式、制表位、对齐等残留干扰）
            for child in list(pPr):
                pPr.remove(child)
            if hclient and htext:
                # 两栏布局：显式左对齐 + 右制表位（从左边距计算）
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), "left")
                pPr.append(jc)
                tabs = OxmlElement("w:tabs")
                tab = OxmlElement("w:tab")
                tab.set(qn("w:val"), "right")
                pw = sec.page_width
                lm = sec.left_margin
                rm = sec.right_margin
                if pw is not None and lm is not None and rm is not None:
                    tab.set(qn("w:pos"), str(int((pw - lm - rm) / 635)))
                else:
                    tab.set(qn("w:pos"), "8845")
                tabs.append(tab)
                pPr.append(tabs)
                # 文字：客户名 → 制表符 → 报告名（同一 run 内）
                r1 = hp.add_run(hclient)
                _style_header_run(r1, fcn, fen, fsz, fco)
                tab_el = OxmlElement("w:tab")
                r1._r.append(tab_el)
                r1.add_text(htext)
            elif htext:
                hp.add_run(htext)
                _style_header_run(hp.runs[0], fcn, fen, fsz, fco)
            # 页眉下方细线
            if hf.get("header_border", True):
                _set_para_bottom_border(hp)

    # ---- 页脚页码 ----
    if show_pn:
        mode = hf.get("footer_mode", "page_of")
        fcn = hf.get("footer_font_cn", "宋体")
        fen = hf.get("footer_font_en", "Times New Roman")
        fsize = hf.get("footer_size_pt", 9)
        fcolor = hf.get("footer_color_rgb", "#595959")
        falign = ALIGN_MAP.get(hf.get("footer_alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)

        if mode == "gb9704":
            # 公文 GB/T 9704：页码 — N —，单页居右、双页居左
            for sec in doc.sections:
                sec.odd_and_even_pages_header_footer = True
                odd = sec.footer.paragraphs[0]          # 奇数页（默认 footer）
                even = sec.even_page_footer.paragraphs[0]
                for p, align in ((odd, WD_ALIGN_PARAGRAPH.RIGHT), (even, WD_ALIGN_PARAGRAPH.LEFT)):
                    p.text = ""
                    p.add_run("— ")
                    _add_field(p, "PAGE")
                    p.add_run(" —")
                    for run in p.runs:
                        _style_header_run(run, fcn, fen, fsize, fcolor)
                    p.alignment = align
        else:
            for sec in doc.sections:
                fp = sec.footer.paragraphs[0]
                fp.text = ""
                if mode == "page_only":
                    _add_field(fp, "PAGE")
                else:  # page_of
                    fp.add_run("第 ")
                    _add_field(fp, "PAGE")
                    fp.add_run(" 页 共 ")
                    _add_field(fp, "NUMPAGES")
                    fp.add_run(" 页")
                for run in fp.runs:
                    _style_header_run(run, fcn, fen, fsize, fcolor)
                fp.alignment = falign


def _sect_break_cover(page_w_emu, page_h_emu, margins):
    """封面末分节符（下一页），继承原文档页面尺寸。

    参数从文档原 sectPr 复制。封面标题块的垂直定位由 add_cover_page 中的
    段落间距控制，而非 vAlign=center——以实现「标题偏上、落款贴底」的专业布局。
    """
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    sectPr = OxmlElement("w:sectPr")
    # 分节类型：下一页
    br_type = OxmlElement("w:type")
    br_type.set(qn("w:val"), "nextPage")
    sectPr.append(br_type)
    # 页面尺寸（EMU→Twips，因 python-docx 用 ST_TwipsMeasure 解析）
    if page_w_emu and page_h_emu:
        pgSz = OxmlElement("w:pgSz")
        pgSz.set(qn("w:w"), str(int(page_w_emu / 635)))
        pgSz.set(qn("w:h"), str(int(page_h_emu / 635)))
        sectPr.append(pgSz)
    if margins:
        pgMar = OxmlElement("w:pgMar")
        for k in ("top", "bottom", "left", "right", "header", "footer"):
            v = margins.get(k)
            if v is not None:
                pgMar.set(qn("w:" + k), str(int(v / 635)))
        sectPr.append(pgMar)
    # 不含 vAlign=center —— 标题块与落款间距由段落 spacing 控制
    pPr.append(sectPr)
    p.append(pPr)
    return p


def add_cover_page(doc, cfg):
    """按配置生成封面页，独占一页，内容页面内垂直居中。

    封面布局（黑/白/灰，无任何彩色）：
      - 客户名称（16pt 居中）
      - 报告名称（22pt 居中加粗）
      - 副标题（可选，14pt 居中）
      - 出具机构 + 日期（可选，12pt 居中）
      - 分节符（nextPage）—— 封面页垂直居中，正文页正常顶部对齐

    仅在 cfg['cover']['title'] 非空时生成；封面插入文档最前。
    不使用 Word 内置 Title/Subtitle 样式（避免继承主题色边框）。
    """
    cover = cfg.get("cover")
    if not cover:
        return
    title = (cover.get("title") or "").strip()
    if not title:
        return

    cn = cover.get("title_font_cn", "黑体")
    en = cover.get("title_font_en", "Arial")
    tsize = cover.get("title_size_pt", 22)
    sub = (cover.get("subtitle") or "").strip()
    ssize = cover.get("subtitle_size_pt", 16)
    client = (cover.get("client") or "").strip()
    csize = cover.get("client_size_pt", 14)
    issuer = (cover.get("issuer") or "").strip()
    date_text = (cover.get("date") or "").strip()
    isize = cover.get("info_size_pt", 12)
    color = cover.get("color_rgb", "#000000").replace("#", "")

    # 客户名是否并入主标题
    client_in_title = cover.get("client_in_title", True)

    def _line(text, fcn, fen, size_pt, bold, align, before=0, after=0):
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc"); jc.set(qn("w:val"), align); pPr.append(jc)
        sp = OxmlElement("w:spacing")
        if before: sp.set(qn("w:before"), str(int(before * 20)))
        if after: sp.set(qn("w:after"), str(int(after * 20)))
        pPr.append(sp)
        p.append(pPr)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), fen); rf.set(qn("w:hAnsi"), fen); rf.set(qn("w:eastAsia"), fcn)
        rPr.append(rf)
        for tag in ("sz", "szCs"):
            s = OxmlElement("w:" + tag); s.set(qn("w:val"), str(int(size_pt * 2))); rPr.append(s)
        if bold:
            b = OxmlElement("w:b"); rPr.append(b)
        rc = OxmlElement("w:color"); rc.set(qn("w:val"), color); rPr.append(rc)
        r.append(rPr)
        t = OxmlElement("w:t"); t.text = text; r.append(t)
        p.append(r)
        return p

    # ── 封面排版 ──
    # 专业布局（对标四大/中伦/金杜等事务所报告封面）：
    #   「标题块」位于页面约 25% 处（客户名→报告名→副标题，自上而下紧凑排列）
    #   「出具信息」位于页面约 80% 处（机构名→日期，靠下但不贴底）
    #   中间留有充分"呼吸"空间，视觉重心上移、信息完整。
    #
    # 参数说明：before/after 单位=pt，在 _line 内部 ×20 转为 twips 写入 w:before/w:after

    body = doc.element.body
    els = []

    # 顶部呼吸空间（将标题块推到约页面 25% 处）
    els.append(_line("", cn, en, tsize, False, "center", before=160))

    # 客户名/报告名标题行（始终拆分两行）
    if client_in_title and client:
        els.append(_line(client, cn, en, csize, False, "center", after=30))
        els.append(_line(title, cn, en, tsize, True, "center", after=40))
    else:
        els.append(_line(title, cn, en, tsize, True, "center", after=40))

    # 副标题（紧凑跟随后）
    if sub:
        els.append(_line(sub, cn, en, ssize, False, "center", after=40))

    # 委托方行（仅当 client 未并入标题）
    if client and not client_in_title:
        els.append(_line("委托方：" + client, "宋体", "Times New Roman", csize, False, "center", after=40))

    # 出具信息——大间距推到页面约 80% 处
    if issuer or date_text:
        if issuer:
            els.append(_line(issuer, "宋体", "Times New Roman", isize, False, "center", before=220, after=10))
        if date_text:
            els.append(_line(date_text, "宋体", "Times New Roman", isize, False, "center"))
    # 从原文档 section 读取页尺寸和边距，传递给封面分节符
    # 确保封面节与正文节 A4 / 边距完全一致
    try:
        src_sec = doc.sections[0]
        pg_w = src_sec.page_width
        pg_h = src_sec.page_height
        margins = {
            "top": src_sec.top_margin,
            "bottom": src_sec.bottom_margin,
            "left": src_sec.left_margin,
            "right": src_sec.right_margin,
            "header": src_sec.header_distance,
            "footer": src_sec.footer_distance,
        }
    except (IndexError, AttributeError):
        pg_w = pg_h = None
        margins = None

    # 封面末：分节符（下一页 + 垂直居中）
    els.append(_sect_break_cover(pg_w, pg_h, margins))

    for el in reversed(els):
        body.insert(0, el)

    # 封面后正文不再重复插入致送对象行：客户名已并入主标题（封面加粗主标题即含），
    # 正文绪言中亦保留"致：【委托方名称】"，故封面只保留一条加粗主标题，避免重复行。


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input")
    ap.add_argument("config")
    ap.add_argument("-o", "--output", default=None)
    args = ap.parse_args()

    inp = Path(args.input)
    if not inp.exists():
        sys.exit(f"输入文件不存在: {inp}")
    out = Path(args.output) if args.output else inp.with_name(inp.stem + "_formatted.docx")
    if out.resolve() == inp.resolve():
        sys.exit("输出文件不能与输入文件相同")

    cfg = json.loads(Path(args.config).read_text(encoding="utf-8"))
    doc = Document(str(inp))

    setup_page(doc, cfg["page"])

    styles_cfg = cfg["styles"]
    applied = []
    for style_name, style_cfg in styles_cfg.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            candidates = HEADING_STYLE_ALIASES.get(style_name, [])
            style = None
            for cand in candidates:
                try:
                    style = doc.styles[cand]
                    break
                except KeyError:
                    continue
            if style is None:
                continue
        apply_style_def(style, style_cfg)
        applied.append(style_name)

    custom_names = set(styles_cfg.keys())
    normal_cfg = styles_cfg.get("Normal")
    n_para = 0
    for p in doc.paragraphs:
        key = style_key_of(p, custom_names)
        target = styles_cfg.get(key, normal_cfg)
        if target:
            fix_paragraph_direct(p, target)
            n_para += 1

    tbl_cfg = cfg.get("table")
    n_cells = 0
    if tbl_cfg:
        for table in doc.tables:
            format_table(table, tbl_cfg)
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        n_cells += 1

    removed = cleanup_doc(doc, cfg.get("cleanup", {}))

    n_break = add_section_dividers(doc, cfg)
    add_cover_page(doc, cfg)                              # 先插入封面（含分节符），后续 headers 自动应用于新 section
    add_header_footer(doc, cfg.get("header_footer"))

    doc.save(str(out))
    print(f"排版完成: {out}")
    print(f"  规范: {cfg['doc_type']} ({cfg['standard']})")
    print(f"  样式已重定义: {', '.join(applied)}")
    print(f"  段落处理: {n_para} 个 | 表格单元格: {n_cells} 个 | 清理空段: {removed} 个 | 插入分页: {n_break or 0} 处")


if __name__ == "__main__":
    main()
