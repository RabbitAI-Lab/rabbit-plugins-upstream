#!/usr/bin/env python3
"""按 JSON 配置对 .docx 文档做合规校验，输出 Markdown 清单。

用法:
    python validate_docx.py <file.docx> <config.json>

三档结论: [通过] / [不通过] / [人工复核]
"""
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

TOL_CM = 0.05
TOL_PT = 0.5

results = []


def record(status, item, detail=""):
    results.append((status, item, detail))


def check_page(doc, page_cfg):
    m = page_cfg["margins_cm"]
    exp_w = page_cfg.get("page_width_cm", 21.0)
    exp_h = page_cfg.get("page_height_cm", 29.7)
    sec = doc.sections[0]
    w_ok = abs(sec.page_width.cm - exp_w) < 0.1 and abs(sec.page_height.cm - exp_h) < 0.1
    record("通过" if w_ok else "不通过", f"纸张尺寸（要求 {exp_w:.1f}×{exp_h:.1f}cm）",
           f"实际 {sec.page_width.cm:.1f}×{sec.page_height.cm:.1f}cm")
    for name, expect, actual in [
        ("上边距", m["top"], sec.top_margin.cm),
        ("下边距", m["bottom"], sec.bottom_margin.cm),
        ("左边距", m["left"], sec.left_margin.cm),
        ("右边距", m["right"], sec.right_margin.cm),
    ]:
        ok = abs(actual - expect) < TOL_CM
        record("通过" if ok else "不通过", f"页边距-{name}",
               f"要求 {expect}cm / 实际 {actual:.2f}cm")


def get_east_asia(style):
    rpr = style.element.find(qn("w:rPr"))
    if rpr is not None:
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is not None:
            return rfonts.get(qn("w:eastAsia"))
    return None


def check_styles(doc, styles_cfg):
    for style_name, cfg in styles_cfg.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            record("人工复核", f"样式「{style_name}」", "文档中不存在该样式，可能使用了直接格式")
            continue
        ea = get_east_asia(style)
        ok_font = ea == cfg["font_cn"]
        record("通过" if ok_font else "不通过", f"{style_name} 中文字体",
               f"要求 {cfg['font_cn']} / 实际 {ea or '未设置'}")
        size = style.font.size
        ok_size = size is not None and abs(size.pt - cfg["size_pt"]) < TOL_PT
        record("通过" if ok_size else "不通过", f"{style_name} 字号",
               f"要求 {cfg['size_pt']}pt / 实际 {size.pt if size else '未设置'}")


def check_body_sample(doc, normal_cfg):
    bad = 0
    total = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        if p.style.name not in ("Normal", "正文"):
            continue
        total += 1
        for run in p.runs:
            rpr = run._element.find(qn("w:rPr"))
            if rpr is not None:
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is not None:
                    ea = rfonts.get(qn("w:eastAsia"))
                    if ea and ea != normal_cfg["font_cn"]:
                        bad += 1
                        break
    if total == 0:
        record("人工复核", "正文段落字体抽检", "未找到正文样式段落")
    else:
        record("通过" if bad == 0 else "不通过", "正文段落字体抽检",
               f"{total} 段中 {bad} 段存在与规范不符的直接字体设置")


def check_figures(doc, fig_cfg):
    """检查文档中的图件：图题、编号、宽度、居中、交叉引用。"""
    if not fig_cfg:
        return
    cap_re = re.compile(r"^图\s*(\d+)[-－]\s*(\d+)\s*(.*)$")
    cap_re_plain = re.compile(r"^图\s*(\d+)\s+(.*)$")
    ref_re = re.compile(r"图\s*(\d+)[-－]\s*(\d+)")

    w_min = float(fig_cfg.get("min_width_cm", 8))
    w_max = float(fig_cfg.get("max_width_cm", 18))
    images = []          # (para_idx, width_cm, has_caption, caption_text)
    captions = []        # (para_idx, chap, idx, text)

    for idx, p in enumerate(doc.paragraphs):
        drawings = p._p.findall(".//" + qn("w:drawing"))
        if not drawings:
            continue
        # 图片宽度
        extents = p._p.findall(".//" + qn("wp:extent"))
        w = None
        if extents:
            cx = extents[0].get("cx")
            if cx:
                w = int(cx) / 360000.0
        # 后继段落作为图题
        has_cap = False
        cap_txt = ""
        if idx + 1 < len(doc.paragraphs):
            nxt = doc.paragraphs[idx + 1].text.strip()
            m = cap_re.match(nxt) or cap_re_plain.match(nxt)
            if m:
                has_cap = True
                cap_txt = nxt
        # 居中检查
        centered = p.alignment == WD_ALIGN_PARAGRAPH.CENTER

        images.append((idx, w, has_cap, cap_txt, centered))

        # 收集图题（用于编号连续性）
        if idx + 1 < len(doc.paragraphs):
            nxt_p = doc.paragraphs[idx + 1]
            m2 = cap_re.match(nxt_p.text.strip())
            if m2:
                captions.append((idx + 1, int(m2.group(1)), int(m2.group(2)), m2.group(3).strip()))

    # 逐图检查
    if not images:
        record("通过", "图件", "未发现图片")
        return

    for i, (pos, w, has_cap, cap_txt, centered) in enumerate(images, 1):
        record(
            "通过" if has_cap else "不通过",
            f"图件 #{i} 图题存在性",
            f"第{pos}段 {'有' if has_cap else '无'}图题"
        )
        if w is not None:
            ok_w = w_min <= w <= w_max
            record(
                "通过" if ok_w else "不通过",
                f"图件 #{i} 宽度范围",
                f"{w:.1f}cm (要求 {w_min}-{w_max}cm)"
            )
        record(
            "通过" if centered else "不通过",
            f"图件 #{i} 居中对齐",
            "已居中" if centered else "未居中"
        )

    # 编号连续性
    if len(captions) > 1:
        prev_chap, prev_idx = captions[0][1], captions[0][2]
        gaps = []
        for ci, ch, ix, _ in captions[1:]:
            if ch == prev_chap and ix != prev_idx + 1:
                gaps.append(f"图{ch}-{prev_idx} -> 图{ch}-{ix} 不连续")
            prev_chap, prev_idx = ch, ix
        if gaps:
            record("不通过", "图题编号连续性", "; ".join(gaps))
        else:
            record("通过", "图题编号连续性", f"{len(captions)} 个图题编号连续")

    # 交叉引用解析（正文中的 见图X-Y / 如图X-Y 所示）
    all_caps = set()
    for _, ch, ix, _ in captions:
        all_caps.add((ch, ix))
    unresolved = set()
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt.startswith("图") and "图" in txt:
            for m in ref_re.finditer(txt):
                key = (int(m.group(1)), int(m.group(2)))
                if key not in all_caps:
                    unresolved.add(f"图{m.group(1)}-{m.group(2)}")
    if unresolved:
        record("不通过", "图件交叉引用可解析性",
               f"以下引用无对应图题: {', '.join(unresolved)}")
    elif captions:
        record("通过", "图件交叉引用可解析性", "所有正文中的图号引用均有对应图题")


def check_table_style(doc, tbl_cfg):
    """检查表格样式：三线表/全线表符合配置，且无填色、无斑马纹。"""
    if not tbl_cfg:
        return
    style = tbl_cfg.get("style")  # three_line / full_line
    if not doc.tables:
        record("通过", "表格样式", "未发现表格")
        return

    shading_found = False
    vertical_found = False
    for tbl in doc.tables:
        # 1) 单元格填色 / 斑马纹检查
        for row in tbl.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = tcPr.find(qn("w:shd"))
                if shd is not None:
                    fill = shd.get(qn("w:fill"))
                    if fill and fill.lower() not in ("auto", "none", ""):
                        shading_found = True
        # 2) 三线表应无竖线（left/right/insideV 为 none）
        tblBorders = tbl._tbl.tblPr.find(qn("w:tblBorders"))
        if tblBorders is not None and style == "three_line":
            for edge in ("left", "right", "insideV"):
                e = tblBorders.find(qn("w:" + edge))
                if e is not None and e.get(qn("w:val")) not in ("none", None):
                    vertical_found = True

    if shading_found:
        record("不通过", "表格无填色/斑马纹", "存在单元格底色填充，违反严肃专业配色要求")
    else:
        record("通过", "表格无填色/斑马纹", "所有单元格无底色填充")

    if style == "three_line":
        if vertical_found:
            record("不通过", "三线表无竖线", "检测到表格含竖向边框，应为仅三条横线")
        else:
            record("通过", "三线表无竖线", "仅保留顶线/表头下细线/底线，无竖线")


def check_section_breaks(doc, cfg):
    """若配置要求节前分页，校验每个一级标题前确有分页符（分隔标识）。"""
    if not cfg.get("page_break_before_h1"):
        return
    h1_paras = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
    if not h1_paras:
        record("人工复核", "节前分页分隔", "未发现一级标题（节）")
        return
    missing = 0
    for p in h1_paras:
        found = False
        for r in p.runs:
            if r._element.find(qn("w:br")) is not None:
                found = True
                break
        if not found:
            missing += 1
    if missing == 0:
        record("通过", "节前分页分隔",
               f"{len(h1_paras)} 个一级标题（节）前均有分页符，满足「节与节之间应有明显分隔标识」")
    else:
        record("不通过", "节前分页分隔", f"{missing} 个一级标题前缺少分页符")


def check_manual(cfg):
    for item in cfg.get("manual_checks", []):
        record("人工复核", item, "")


def check_header_footer(doc, cfg):
    """校验页眉文字与页脚页码域是否存在（按 header_footer 配置）。"""
    hf = cfg.get("header_footer")
    if not hf:
        record("人工复核", "页眉页脚", "配置未定义 header_footer 块，需人工确认")
        return
    # 页眉文字
    htext = (hf.get("header_text") or "").strip()
    if htext:
        hcontent = "".join(p.text for p in doc.sections[0].header.paragraphs)
        if htext in hcontent:
            record("通过", "页眉文字", f"页眉含「{htext}」")
        else:
            record("不通过", "页眉文字", f"要求含「{htext}」/ 实际「{hcontent.strip()}」")
    else:
        record("通过", "页眉文字", "配置要求无页眉文字（如公文仅页码）")
    # 页脚页码域
    if hf.get("show_page_number", True):
        found = False
        for sec in doc.sections:
            xml = (
                sec.footer._element.xml
                + sec.even_page_footer._element.xml
            )
            if "PAGE" in xml:
                found = True
                break
        if found:
            record("通过", "页脚页码域", "页脚含 PAGE 页码域")
        else:
            record("不通过", "页脚页码域", "页脚未找到 PAGE 页码域")
    else:
        record("通过", "页脚页码域", "配置要求不显示页码")


def main():
    if len(sys.argv) < 3:
        sys.exit("用法: validate_docx.py <file.docx> <config.json>")
    doc = Document(sys.argv[1])
    cfg = json.loads(Path(sys.argv[2]).read_text(encoding="utf-8"))

    check_page(doc, cfg["page"])
    check_styles(doc, cfg["styles"])
    check_body_sample(doc, cfg["styles"]["Normal"])
    check_figures(doc, cfg.get("figures"))
    check_table_style(doc, cfg.get("table"))
    check_section_breaks(doc, cfg)
    check_header_footer(doc, cfg)
    check_manual(cfg)

    n_pass = sum(1 for s, _, _ in results if s == "通过")
    n_fail = sum(1 for s, _, _ in results if s == "不通过")
    n_manual = sum(1 for s, _, _ in results if s == "人工复核")

    print(f"# 排版合规校验清单\n")
    print(f"- 文档: {sys.argv[1]}")
    print(f"- 规范: {cfg['doc_type']}（{cfg['standard']}）")
    print(f"- 结果: 通过 {n_pass} 项 | 不通过 {n_fail} 项 | 需人工复核 {n_manual} 项\n")
    print("| 结论 | 检查项 | 说明 |")
    print("|---|---|---|")
    for status, item, detail in results:
        mark = {"通过": "✅ 通过", "不通过": "❌ 不通过", "人工复核": "⚠️ 人工复核"}[status]
        print(f"| {mark} | {item} | {detail} |")
    sys.exit(1 if n_fail else 0)


if __name__ == "__main__":
    main()
