"""外部交叉验证：产出的文件能不能被别的程序正常打开。

自己写的测试只能证明自己的假设成立，证明不了 Word 会认。
这里用两个独立实现来验：LibreOffice 和 python-docx，都不是本模块的代码。

更强的做法是传上原始文件做对照，比 PDF 页数：

    python3 verify_external.py 改写后.docx --compare 原始.docx

实测一篇 30 页论文，改写前后都是 30 页、体积差 34 字节，
这比「我的测试全绿」有说服力得多。

注意 LibreOffice 装了不代表能跑。本机 7.3.7.2 转任何文件都报
source file could not be loaded，换一台就正常。所以它报失败时，
先拿一个没经过本模块的文件做对照，确认是环境问题还是产出有问题。

跑法: python3 verify_external.py <.docx> [--compare <原始.docx>]
"""
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path


def _to_pdf(path):
    """转成 PDF，返回 (成功?, 字节数, 页数, 说明)。页数取不到时为 None。"""
    soffice = shutil.which('soffice') or shutil.which('libreoffice')
    if not soffice:
        return None, 0, None, '本机没有 LibreOffice，跳过'
    tmp = tempfile.mkdtemp()
    try:
        r = subprocess.run(
            [soffice, '-env:UserInstallation=file://%s/profile' % tmp,
             '--headless', '--convert-to', 'pdf', '--outdir', tmp, str(path)],
            capture_output=True, timeout=300)
        pdfs = list(Path(tmp).glob('*.pdf'))
        if not pdfs:
            return False, 0, None, ('转换失败（退出码 %d）：%s'
                                    % (r.returncode, r.stderr.decode()[:160]))
        size = pdfs[0].stat().st_size
        pages = None
        if shutil.which('pdfinfo'):
            out = subprocess.run(['pdfinfo', str(pdfs[0])],
                                 capture_output=True, timeout=60).stdout.decode()
            for line in out.splitlines():
                if line.lower().startswith('pages:'):
                    pages = int(line.split(':')[1].strip())
        return True, size, pages, '%d 字节%s' % (size, '，%d 页' % pages if pages else '')
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def by_libreoffice(path):
    ok, _size, _pages, msg = _to_pdf(path)
    return ok, ('LibreOffice 能完整解析（%s）' % msg) if ok else msg


def by_python_docx(path):
    """用另一个 OOXML 实现读一遍，看段落数和文本长度。"""
    try:
        import docx
    except ImportError:
        return None, '没装 python-docx，跳过'
    d = docx.Document(str(path))
    n_p = len(d.paragraphs)
    n_c = sum(len(p.text) for p in d.paragraphs)
    if n_p == 0:
        return False, 'python-docx 读出 0 个段落，文件可能坏了'
    return True, 'python-docx 读到 %d 段、%d 字符' % (n_p, n_c)


def main(path, compare=None):
    path = Path(path)
    if not path.exists():
        print('找不到文件：%s' % path)
        return 2
    print('外部交叉验证：%s' % path.name)
    bad = False
    lo_ok = None
    other_failed = False
    for name, fn in (('LibreOffice', by_libreoffice), ('python-docx', by_python_docx)):
        ok, msg = fn(path)
        if name == 'LibreOffice':
            lo_ok = ok
        elif ok is False:
            other_failed = True
        print('  %s %-12s %s' % ({True: '通过', False: '失败', None: '跳过'}[ok], name, msg))
        if ok is False:
            bad = True

    if compare:
        ok_a, size_a, pages_a, msg_a = _to_pdf(Path(compare))
        ok_b, size_b, pages_b, _ = _to_pdf(path)
        if ok_a is False and ok_b is False:
            print('  提示 原始文件同样转不出来，是 LibreOffice 环境问题，不是产出有问题')
            # 只撤销 LibreOffice 那一项的判定，别把其它检查的失败一起抹掉
            if lo_ok is False and not other_failed:
                bad = False
        elif ok_a and ok_b:
            same_pages = (pages_a == pages_b)
            drift = abs(size_b - size_a)
            print('  %s 页数对照 原始 %s 页 / 改写后 %s 页' %
                  ('通过' if same_pages else '失败', pages_a, pages_b))
            print('  参考 体积 %d -> %d 字节（差 %d）' % (size_a, size_b, drift))
            if not same_pages:
                bad = True
        else:
            print('  对照 一侧转换失败：原始 %s' % msg_a)

    return 1 if bad else 0


if __name__ == '__main__':
    args = sys.argv[1:]
    if not args or '--help' in args:
        print(__doc__)
        sys.exit(2)
    cmp_path = None
    if '--compare' in args:
        i = args.index('--compare')
        cmp_path = args[i + 1]
        args = args[:i] + args[i + 2:]
    sys.exit(main(args[0], cmp_path))
