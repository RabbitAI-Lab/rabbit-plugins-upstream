"""report_spec → .docx 组装器（设计感版本）。

设计语言（与 PPTX/HTML 报告统一）：
  - 配色：primary #0E6BA8 / accent #E07A3B / text #26292E / sub #5C6066 / rule #DCE3EA
  - 字体：标题 = 宋体（Source Han Serif SC），正文 = 微软雅黑，数字 = Calibri
  - 节奏：H1 大字号 + accent 装饰条 + 标题前留白；正文 1.5 倍行距；图表居中带边框

页面结构：
  第 1 页    封面（顶部蓝色色块 + 大标题 + 副标题 + 元信息）
  第 2 页    目录（章节列表）
  第 3 页    执行摘要（标题 + 项目符号列表）
  第 4..页   各章节（H1 + 正文 + 居中图 + 图脚）
  末页       附录·关键数据溯源表

依赖：python-docx（可选）。模块级 try-import；缺包时抛 ExportDependencyError。
"""
from __future__ import annotations

from pathlib import Path


class ExportDependencyError(Exception):
    def __init__(self, message: str, pkg: str):
        super().__init__(message)
        self.pkg = pkg


def _ensure_docx():
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        return Document, Pt, Inches, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, Cm
    except ImportError as e:
        raise ExportDependencyError(
            f"python-docx 不可用: {e}",
            pkg="python-docx",
        ) from e


# 模块级 try-import（缺包时 _ensure_docx 重检；build_docx 顶部还会再确认）
try:
    Document, Pt, Inches, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, Cm = _ensure_docx()
    _DX_OK = True
except Exception:
    _DX_OK = False
    Document = Pt = Inches = RGBColor = qn = WD_ALIGN_PARAGRAPH = WD_TABLE_ALIGNMENT = Cm = None


# 调色板（与 pptx_export.py 一致）
C_PRIMARY = (0x0E, 0x6B, 0xA8)
C_ACCENT  = (0xE0, 0x7A, 0x3B)
C_TEXT    = (0x26, 0x29, 0x2E)
C_SUB     = (0x5C, 0x60, 0x66)
C_RULE    = (0xDC, 0xE3, 0xEA)
C_BG_CARD = (0xF4, 0xF6, 0xF9)
C_INK_REV = (0xFF, 0xFF, 0xFF)
C_INK_SUB = (0xDC, 0xE6, 0xF1)

F_HAN_HEAD = "Source Han Serif SC"
F_HAN_BODY = "Microsoft YaHei"
F_NUM      = "Calibri"


# ---------- 字体辅助 ----------

def _set_run(run, ascii_name: str, east_asia: str, size_pt: float,
             bold: bool = False, color_rgb=None):
    """设字体 + 字号 + bold + 颜色，含 eastAsia 修正。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from lxml import etree
    from docx.shared import RGBColor
    run.font.name = ascii_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = etree.SubElement(rpr, qn('w:rFonts'))
    rfonts.set(qn('w:ascii'), ascii_name)
    rfonts.set(qn('w:hAnsi'), ascii_name)
    rfonts.set(qn('w:eastAsia'), east_asia)
    if color_rgb is not None:
        if isinstance(color_rgb, tuple):
            color_rgb = RGBColor(*color_rgb)
        run.font.color.rgb = color_rgb


def _add_paragraph(doc, *, style=None):
    return doc.add_paragraph(style=style) if style else doc.add_paragraph()


def _para(doc, *, alignment=None, before=0, after=0, line_space=1.5, style=None):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = _add_paragraph(doc, style=style)
    if alignment is not None:
        p.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }.get(alignment, alignment)
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_space
    return p


def _run(p, text, ascii_name, east_asia, size_pt, bold=False, color_rgb=None):
    r = p.add_run(text)
    _set_run(r, ascii_name, east_asia, size_pt, bold=bold, color_rgb=color_rgb)
    return r


# ---------- 视觉块 ----------

def _horizontal_rule(doc, color=C_RULE, height_pt=1.5, width_inches=6.0):
    """用底部边框模拟横线。"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr')
    if pBdr is None:
        from lxml import etree
        pBdr = etree.SubElement(pPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr')
    from lxml import etree
    bottom = etree.SubElement(pBdr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom')
    bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
    bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', str(int(height_pt * 8)))
    bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', '%02X%02X%02X' % color)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)


def _accent_block(doc, width_inches=0.7, height_pt=4, color=C_ACCENT):
    """用一短段下边线模拟色块。"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr')
    if pBdr is None:
        from lxml import etree
        pBdr = etree.SubElement(pPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr')
    from lxml import etree
    bottom = etree.SubElement(pBdr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom')
    bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
    bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', str(height_pt * 8))
    bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', '%02X%02X%02X' % color)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)


def _shade_paragraph(p, color):
    """给段落加底色（用 pPr/shd）。"""
    from lxml import etree
    pPr = p._element.get_or_add_pPr()
    shd = etree.SubElement(pPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'clear')
    shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', 'auto')
    shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '%02X%02X%02X' % color)


def _shade_cell(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    from lxml import etree
    if isinstance(color, tuple):
        from docx.shared import RGBColor
        color = RGBColor(*color)
    shd = etree.SubElement(tcPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'clear')
    shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', 'auto')
    shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '%02X%02X%02X' % (color[0], color[1], color[2]))


# ---------- 1. 封面页 ----------

def _add_cover(doc, spec):
    from docx.shared import Inches, Pt
    # 上方蓝色色块（用 1x1 表格）
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.0)
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.0)
    _shade_cell(cell, C_PRIMARY)
    cell.paragraphs[0].text = ""
    # 顶部装饰：橙条
    p_top = cell.paragraphs[0]
    p_top.paragraph_format.space_before = Pt(8)
    _accent_block_in_cell(cell, color=C_ACCENT, height_pt=2)
    # 角标
    p_tag = cell.add_paragraph()
    _run(p_tag, "SMART REPORT", F_NUM, F_HAN_BODY, 10, bold=True, color_rgb=C_INK_REV)
    _run(p_tag, "  · 数据报告", F_HAN_BODY, F_HAN_BODY, 10, color_rgb=C_INK_SUB)
    p_tag.paragraph_format.space_before = Pt(4)
    p_tag.paragraph_format.space_after = Pt(0)
    # 主标题
    p_title = cell.add_paragraph()
    _run(p_title, spec.get("title") or "数据报告",
         F_HAN_HEAD, F_HAN_HEAD, 26, bold=True, color_rgb=C_INK_REV)
    p_title.paragraph_format.space_before = Pt(14)
    p_title.paragraph_format.space_after = Pt(8)
    p_title.paragraph_format.line_spacing = 1.2
    # 副标题
    if spec.get("subtitle"):
        p_sub = cell.add_paragraph()
        _run(p_sub, spec["subtitle"], F_HAN_BODY, F_HAN_BODY, 12,
             color_rgb=C_INK_SUB)
        p_sub.paragraph_format.space_after = Pt(8)
        p_sub.paragraph_format.line_spacing = 1.3
    # 数据口径
    if spec.get("appendix", {}).get("methodology"):
        p_meta = cell.add_paragraph()
        _run(p_meta, "数据口径 · " + (spec["appendix"]["methodology"][:120] + ("…" if len(spec["appendix"]["methodology"]) > 120 else "")),
             F_HAN_BODY, F_HAN_BODY, 9, color_rgb=C_INK_SUB)
        p_meta.paragraph_format.space_after = Pt(8)
    # 底部日期
    from datetime import datetime
    p_date = cell.add_paragraph()
    _run(p_date, datetime.now().strftime("%Y-%m-%d"),
         F_NUM, F_HAN_BODY, 10, color_rgb=C_INK_SUB)
    p_date.paragraph_format.space_before = Pt(8)
    p_date.alignment = 2  # right

    # 色块下留白
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_after = Pt(0)


def _accent_block_in_cell(cell, color, height_pt=2):
    """在单元格段落底部加橙色细线（替代段落色块）。"""
    p = cell.paragraphs[0]
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr')
    if pBdr is None:
        from lxml import etree
        pBdr = etree.SubElement(pPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr')
    from lxml import etree
    bottom = etree.SubElement(pBdr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom')
    bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
    bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', str(height_pt * 8))
    bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', '%02X%02X%02X' % color)


# ---------- 2. 目录页 ----------

def _add_toc(doc, sections):
    # H1 标题
    p = _para(doc, before=0, after=2, line_space=1.2)
    _run(p, "目录", F_HAN_HEAD, F_HAN_HEAD, 22, bold=True, color_rgb=C_TEXT)
    _accent_block(doc, width_inches=0.6, height_pt=3, color=C_ACCENT)
    p2 = _para(doc, before=0, after=14, line_space=1.0)
    _run(p2, "CONTENTS", F_NUM, F_HAN_BODY, 9, color_rgb=C_SUB)

    # 章节列表（用表格：编号 / 标题 / 页码占位）
    from docx.shared import Inches
    from docx.oxml.ns import qn
    table = doc.add_table(rows=len(sections), cols=3)
    table.autofit = False
    widths = [Inches(0.6), Inches(4.6), Inches(0.8)]
    for i, w in enumerate(widths):
        table.columns[i].width = w
    for i, sec in enumerate(sections):
        cells = table.rows[i].cells
        for j, w in enumerate(widths):
            cells[j].width = w
        # 编号
        cells[0].text = ""
        p = cells[0].paragraphs[0]
        _run(p, f"{i + 1:02d}", F_NUM, F_HAN_BODY, 14, bold=True, color_rgb=C_PRIMARY)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # 标题
        cells[1].text = ""
        p = cells[1].paragraphs[0]
        _run(p, sec.get("title") or "", F_HAN_HEAD, F_HAN_HEAD,
             13, bold=True, color_rgb=C_TEXT)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # 页码占位
        cells[2].text = ""
        p = cells[2].paragraphs[0]
        _run(p, f"P.{i + 3:02d}", F_NUM, F_HAN_BODY, 10, color_rgb=C_SUB)
        p.alignment = 2
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)


# ---------- 3. 执行摘要 ----------

def _add_summary(doc, summary: str):
    p = _para(doc, before=0, after=2, line_space=1.2)
    _run(p, "执行摘要", F_HAN_HEAD, F_HAN_HEAD, 22, bold=True, color_rgb=C_TEXT)
    _accent_block(doc, width_inches=0.6, height_pt=3, color=C_ACCENT)
    p2 = _para(doc, before=0, after=12, line_space=1.0)
    _run(p2, "EXECUTIVE SUMMARY", F_NUM, F_HAN_BODY, 9, color_rgb=C_SUB)

    # 按句号切分，每句一行（带左 accent 模拟）
    sentences = [s.strip() for s in
                 summary.replace("\n\n", "\n").split("\n") if s.strip()]
    if not sentences:
        sentences = [s.strip() for s in
                     summary.replace("\n\n", "\u3002\n").split("\u3002\n") if s.strip()]
    sentences = [s.rstrip("。.") for s in sentences if s.strip()]

    for idx, s in enumerate(sentences, 1):
        p = _para(doc, before=4, after=4, line_space=1.55)
        # 左边条
        pPr = p._element.get_or_add_pPr()
        from lxml import etree
        pBdr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr')
        if pBdr is None:
            pBdr = etree.SubElement(pPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr')
        left = etree.SubElement(pBdr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left')
        left.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
        left.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', '18')
        left.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', '%02X%02X%02X' % C_ACCENT)
        p.paragraph_format.left_indent = Pt(12)
        _run(p, f"{idx:02d}  ", F_NUM, F_HAN_BODY, 11, bold=True, color_rgb=C_PRIMARY)
        _run(p, s + "。", F_HAN_BODY, F_HAN_BODY, 12, color_rgb=C_TEXT)


# ---------- 4. 章节页 ----------

def _add_section(doc, idx, sec, image_path):
    # H1：编号 + 标题
    p = _para(doc, before=8, after=4, line_space=1.2)
    _run(p, f"{idx:02d}  ", F_NUM, F_HAN_BODY, 16, bold=True, color_rgb=C_PRIMARY)
    _run(p, sec.get("title") or "", F_HAN_HEAD, F_HAN_HEAD,
         18, bold=True, color_rgb=C_TEXT)
    # H1 下方装饰线
    _horizontal_rule(doc, color=C_RULE, height_pt=1.0)

    # 正文（按 \n\n 切段）
    narrative = sec.get("narrative") or sec.get("annotation") or ""
    for chunk in [c.strip() for c in narrative.split("\n\n") if c.strip()]:
        p = _para(doc, before=2, after=4, line_space=1.6)
        p.paragraph_format.first_line_indent = Pt(22)
        _run(p, chunk, F_HAN_BODY, F_HAN_BODY, 11, color_rgb=C_TEXT)

    # 图表：居中 + 浅色边框（用 1x1 表格实现）
    if image_path and Path(image_path).is_file():
        from docx.shared import Inches
        table = doc.add_table(rows=1, cols=1)
        table.alignment = 1  # center
        cell = table.rows[0].cells[0]
        cell.width = Inches(6.0)
        # 单元格边框 + 底色
        tcPr = cell._tc.get_or_add_tcPr()
        from lxml import etree
        tcBorders = etree.SubElement(tcPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
        for side in ('top', 'left', 'bottom', 'right'):
            b = etree.SubElement(tcBorders, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}' + side)
            b.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
            b.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', '4')
            b.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', '%02X%02X%02X' % C_RULE)
        _shade_cell(cell, (0xFF, 0xFF, 0xFF))
        cell.paragraphs[0].text = ""
        # 图内文
        cell.paragraphs[0].paragraph_format.space_before = Pt(4)
        run = cell.paragraphs[0].add_run()
        try:
            run.add_picture(str(image_path), width=Inches(5.6))
        except Exception as e:
            _run(cell.paragraphs[0], f"[图片插入失败: {e}]",
                 F_HAN_BODY, F_HAN_BODY, 9, color_rgb=C_SUB)

        # 图脚 caption
        ann = sec.get("annotation") or ""
        p_cap = _para(doc, before=2, after=10, line_space=1.2)
        p_cap.alignment = 1
        _run(p_cap, f"图 {idx}", F_HAN_BODY, F_HAN_BODY, 10,
             bold=True, color_rgb=C_PRIMARY)
        if ann:
            _run(p_cap, "  ·  ", F_HAN_BODY, F_HAN_BODY, 10, color_rgb=C_SUB)
            _run(p_cap, ann, F_HAN_BODY, F_HAN_BODY, 10, color_rgb=C_SUB)


# ---------- 5. 附录·关键数据溯源表 ----------

def _add_ledger_appendix(doc, resolver):
    if not resolver or not resolver.entries:
        return
    p = _para(doc, before=12, after=2, line_space=1.2)
    _run(p, "附录·关键数据溯源", F_HAN_HEAD, F_HAN_HEAD,
         18, bold=True, color_rgb=C_TEXT)
    _accent_block(doc, width_inches=0.6, height_pt=3, color=C_ACCENT)
    p2 = _para(doc, before=0, after=8, line_space=1.4)
    _run(p2, "本表列出正文引用的关键数字及其在数据台账中的出处，便于读者复核口径。",
         F_HAN_BODY, F_HAN_BODY, 10, color_rgb=C_SUB)

    from docx.shared import Inches
    rows = resolver.entries
    headers = ["指标", "数值", "单位", "出处"]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.autofit = False
    widths = [Inches(2.0), Inches(1.2), Inches(0.6), Inches(2.2)]
    for i, w in enumerate(widths):
        table.columns[i].width = w
    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.width = widths[j]
        cell.text = ""
        _shade_cell(cell, C_PRIMARY)
        _run(cell.paragraphs[0], h, F_HAN_BODY, F_HAN_BODY, 10,
             bold=True, color_rgb=C_INK_REV)
        cell.paragraphs[0].paragraph_format.space_before = Pt(4)
        cell.paragraphs[0].paragraph_format.space_after = Pt(4)
    # 数据行
    for i, e in enumerate(rows, 1):
        row_bg = (0xFF, 0xFF, 0xFF) if i % 2 == 1 else (0xF7, 0xF9, 0xFC)
        for j, txt in enumerate([e.metric, str(e.value), e.unit or "", e.source]):
            cell = table.rows[i].cells[j]
            cell.width = widths[j]
            cell.text = ""
            _shade_cell(cell, row_bg)
            fn = F_NUM if j == 1 else F_HAN_BODY
            _run(cell.paragraphs[0], str(txt), fn, F_HAN_BODY, 10,
                 bold=(j == 1), color_rgb=C_TEXT)
            cell.paragraphs[0].paragraph_format.space_before = Pt(3)
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)


# ---------- 6. 附录：数据与方法 / 口径与局限 ----------

def _add_appendix(doc, appendix):
    items = [(k, label) for k, label in [
        ("methodology", "数据与方法"),
        ("caveats", "口径与局限"),
    ] if appendix.get(k)]
    if not items:
        return
    p = _para(doc, before=12, after=2, line_space=1.2)
    _run(p, "附录", F_HAN_HEAD, F_HAN_HEAD, 18, bold=True, color_rgb=C_TEXT)
    _accent_block(doc, width_inches=0.6, height_pt=3, color=C_ACCENT)
    for k, label in items:
        p = _para(doc, before=8, after=2, line_space=1.3)
        _run(p, label, F_HAN_HEAD, F_HAN_HEAD, 13, bold=True, color_rgb=C_PRIMARY)
        for chunk in [c.strip() for c in (appendix.get(k) or "").split("\n\n") if c.strip()]:
            p2 = _para(doc, before=2, after=4, line_space=1.6)
            _run(p2, chunk, F_HAN_BODY, F_HAN_BODY, 10, color_rgb=C_TEXT)


# ---------- 入口 ----------

def build_docx(spec: dict, sections: list,
               chart_images: dict | None,
               resolver, output_path) -> Path:
    Document, _, _, _, _, _, _, _ = _ensure_docx()
    doc = Document()

    # 页面：A4，1 inch 边距
    for section in doc.sections:
        section.left_margin = section.right_margin = 914400
        section.top_margin = section.bottom_margin = 914400

    _add_cover(doc, spec)
    doc.add_page_break()

    _add_toc(doc, sections)
    doc.add_page_break()

    if (spec.get("executive_summary") or "").strip():
        _add_summary(doc, spec["executive_summary"])
        doc.add_page_break()

    for i, sec in enumerate(sections):
        _add_section(doc, i + 1, sec,
                     (chart_images or {}).get(sec.get("id")) if chart_images else None)
        if i < len(sections) - 1:
            # 留白而非强制分页（让 Word 自然分页）
            pass

    appendix = spec.get("appendix") or {}
    _add_appendix(doc, appendix)

    _add_ledger_appendix(doc, resolver)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path