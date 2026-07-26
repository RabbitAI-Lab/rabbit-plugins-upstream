# -*- coding: utf-8 -*-
"""
Score Analysis - Template Generator
Generate Word report template with placeholders
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Default color scheme
COLORS = {
    'primary': RGBColor(0x00, 0x6B, 0x6B),
    'gold': RGBColor(0xB8, 0x86, 0x0B),
    'text': RGBColor(0x33, 0x33, 0x33),
    'light_gray': RGBColor(0x99, 0x99, 0x99),
}


def setup_styles(doc):
    """Set document styles"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = COLORS['text']
    style.paragraph_format.line_spacing = 1.5


def add_header_footer(doc, logo_path=None):
    """Add header and footer"""
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
        
        # Header
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if logo_path and os.path.exists(logo_path):
            run = p.add_run()
            run.add_picture(logo_path, width=Inches(0.4))
            p.add_run('  ')
        
        run = p.add_run('{{SCHOOL_NAME}}')
        run.font.size = Pt(10)
        run.font.color.rgb = COLORS['primary']
        
        # Header underline
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="006B6B"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)
        
        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="6" w:space="1" w:color="006B6B"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)
        
        run = p.add_run('— ')
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS['light_gray']
        
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run1 = p.add_run()
        run1._r.append(fldChar1)
        
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2 = p.add_run()
        run2._r.append(instrText)
        
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3 = p.add_run()
        run3._r.append(fldChar2)
        
        run4 = p.add_run(' —')
        run4.font.size = Pt(9)
        run4.font.color.rgb = COLORS['light_gray']


def add_cover(doc, logo_path=None):
    """Add cover page"""
    if logo_path and os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(1.2))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('{{SCHOOL_NAME}}')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = COLORS['primary']
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run('━' * 25)
    run.font.color.rgb = COLORS['gold']
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('{{CLASS_NAME}}')
    run.font.size = Pt(20)
    run.font.color.rgb = COLORS['primary']
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(15)
    run = p.add_run('{{EXAM_NAME}} Analysis Report')
    run.font.size = Pt(20)
    run.font.color.rgb = COLORS['primary']
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(15)
    run = p.add_run('━' * 25)
    run.font.color.rgb = COLORS['gold']
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run('Date: {{DATE}}')
    run.font.size = Pt(13)
    run.font.color.rgb = COLORS['text']
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run('Prepared by: {{CLASS_NAME}}')
    run.font.size = Pt(13)
    run.font.color.rgb = COLORS['text']
    
    doc.add_page_break()


def add_placeholder_sections(doc):
    """Add placeholder sections"""
    doc.add_heading('Table of Contents', level=1)
    for item in ['1. Basic Statistics', '2. Class Comparison', '3. Score Distribution', 
                 '4. Critical Students', '5. Subject Imbalance', '6. Conclusions']:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)
    doc.add_page_break()
    
    sections = [
        ('1. Basic Statistics', '{{SECTION_1_CONTENT}}'),
        ('2. Class Comparison', '{{SECTION_2_CONTENT}}'),
        ('3. Score Distribution', '{{SECTION_3_CONTENT}}'),
        ('4. Critical Students', '{{SECTION_4_CONTENT}}'),
        ('5. Subject Imbalance', '{{SECTION_5_CONTENT}}'),
        ('6. Conclusions & Recommendations', '{{SECTION_6_CONTENT}}'),
    ]
    
    for title, placeholder in sections:
        doc.add_heading(title, level=1)
        p = doc.add_paragraph(placeholder)
        p.paragraph_format.space_before = Pt(20)
        p.runs[0].font.color.rgb = COLORS['light_gray']
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 25)
    run.font.color.rgb = COLORS['gold']
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('{{SCHOOL_NAME}}')
    run.font.size = Pt(10)
    run.font.color.rgb = COLORS['light_gray']


def create_template(output_path, logo_path=None):
    """Create template file"""
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc, logo_path)
    add_cover(doc, logo_path)
    add_placeholder_sections(doc)
    doc.save(output_path)
    print(f'Template generated: {output_path}')


if __name__ == '__main__':
    create_template('./report_template.docx')
