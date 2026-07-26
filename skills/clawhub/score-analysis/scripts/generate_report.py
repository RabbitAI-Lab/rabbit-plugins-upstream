# -*- coding: utf-8 -*-
"""
Score Analysis - Report Generator
Generates professional Word reports with charts
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Default color scheme (customizable)
DEFAULT_COLORS = {
    'primary': RGBColor(0x00, 0x6B, 0x6B),
    'secondary': RGBColor(0x2E, 0x86, 0x86),
    'accent': RGBColor(0xC0, 0x39, 0x2B),
    'gold': RGBColor(0xB8, 0x86, 0x0B),
    'text': RGBColor(0x33, 0x33, 0x33),
    'light_gray': RGBColor(0x99, 0x99, 0x99),
}


def create_three_line_table(doc, headers, data, col_widths=None, colors=None):
    """Create research-style three-line table"""
    if colors is None:
        colors = DEFAULT_COLORS
    
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="12" w:space="0" w:color="006B6B"/>'
        '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="006B6B"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = colors['primary']
        
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="006B6B"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F4F4" w:val="clear"/>')
        tcPr.append(shading_elm)
    
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_data))
            run.font.size = Pt(10.5)
            run.font.name = 'Times New Roman'
            
            if i % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
                tcPr.append(shading_elm)
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    doc.add_paragraph()


def add_chart(doc, chart_path, caption):
    """Add chart image to document"""
    if chart_path and os.path.exists(chart_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(chart_path, width=Inches(5))
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.font.size = Pt(10.5)
        run.font.italic = True
    doc.add_paragraph()


def add_highlight_box(doc, text, colors=None):
    """Add highlight box for emphasis"""
    if colors is None:
        colors = DEFAULT_COLORS
    
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    cell = table.rows[0].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = colors['primary']
    
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F4F4" w:val="clear"/>')
    tcPr.append(shading_elm)
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="6" w:space="0" w:color="006B6B"/>'
        '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="006B6B"/>'
        '  <w:left w:val="single" w:sz="18" w:space="0" w:color="006B6B"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    doc.add_paragraph()


def generate_report(class_name, exam_name, date, analysis_data, charts_dir, output_path, logo_path=None):
    """
    Generate complete analysis report
    
    Args:
        class_name: str, class name (e.g., "Class 1, Grade 10")
        exam_name: str, exam name (e.g., "Mid-term Exam")
        date: str, analysis date
        analysis_data: dict, analysis results
        charts_dir: str, charts directory
        output_path: str, output file path
        logo_path: str, optional school logo path
    """
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Add cover page
    if logo_path and os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(1.2))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(class_name)
    run.font.size = Pt(24)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{exam_name} Analysis Report')
    run.font.size = Pt(20)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Date: {date}')
    run.font.size = Pt(14)
    
    doc.add_page_break()
    
    # Add content sections
    # Section 1: Basic Statistics
    doc.add_heading('1. Basic Statistics', level=1)
    if analysis_data.get('basic_stats'):
        headers = ['Subject', 'Average', 'Max', 'Min', 'Pass Rate']
        data = analysis_data['basic_stats']
        create_three_line_table(doc, headers, data)
    
    # Section 2: Class Comparison
    doc.add_heading('2. Class Comparison', level=1)
    if analysis_data.get('comparison'):
        headers = ['Class', 'Average', 'Top Score', 'Special Line', 'Undergraduate Line']
        data = analysis_data['comparison']
        create_three_line_table(doc, headers, data)
    
    # Section 3: Score Distribution
    doc.add_heading('3. Score Distribution', level=1)
    chart_path = os.path.join(charts_dir, 'score_distribution.png') if charts_dir else None
    add_chart(doc, chart_path, 'Figure 1: Score Distribution')
    
    # Section 4: Critical Students
    doc.add_heading('4. Critical Students', level=1)
    if analysis_data.get('critical_special'):
        doc.add_heading('4.1 Special Control Line', level=2)
        headers = ['Name', 'Score', 'Gap', 'Weak Subject', 'Suggestion']
        data = [[s['name'], s['score'], s['gap'], s['weak'], s['suggestion']] 
                for s in analysis_data['critical_special']]
        create_three_line_table(doc, headers, data)
        
        chart_path = os.path.join(charts_dir, 'radar_critical_special.png') if charts_dir else None
        add_chart(doc, chart_path, 'Figure 2: Critical Students Radar')
    
    # Section 5: Subject Imbalance
    doc.add_heading('5. Subject Imbalance Analysis', level=1)
    if analysis_data.get('poor_balance'):
        headers = ['Name', 'Total', 'Chinese', 'Math', 'English', 'Physics', 'Chemistry', 'Biology', 'Imbalance Index']
        data = [[s['name'], s['total'], s['chinese'], s['math'], s['english'], 
                 s['physics'], s['chemistry'], s['biology'], s['imbalance']] 
                for s in analysis_data['poor_balance']]
        create_three_line_table(doc, headers, data)
        
        chart_path = os.path.join(charts_dir, 'radar_poor_balance.png') if charts_dir else None
        add_chart(doc, chart_path, 'Figure 3: Imbalanced Students Radar')
    
    # Section 6: Conclusions
    doc.add_heading('6. Conclusions & Recommendations', level=1)
    if analysis_data.get('conclusions'):
        for item in analysis_data['conclusions']:
            p = doc.add_paragraph(item, style='List Bullet')
    
    doc.save(output_path)
    print(f'Report generated: {output_path}')


# Example usage
if __name__ == '__main__':
    example_data = {
        'basic_stats': [
            ['Chinese', '104.2', '118', '88', '85%'],
            ['Math', '87.9', '135', '37', '62%'],
            ['English', '87.2', '117.5', '54', '65%'],
        ],
        'comparison': [
            ['Class A', '480.6', '609', '18', '51'],
            ['Class B', '478.5', '585', '13', '50'],
            ['Class C', '478.4', '590', '18', '52'],
        ],
        'critical_special': [
            {'name': 'Student A', 'score': 506, 'gap': -2, 'weak': 'Physics', 'suggestion': 'Focus on physics'},
        ],
        'poor_balance': [
            {'name': 'Student B', 'total': 563, 'chinese': 112, 'math': 45, 'english': 108.5, 
             'physics': 65, 'chemistry': 88, 'biology': 73, 'imbalance': 28.6},
        ],
        'conclusions': [
            'Strong performance in top tier',
            'Subject imbalance needs attention',
            'Critical students require focused support',
        ]
    }
    
    generate_report(
        class_name='Class 1, Grade 10',
        exam_name='Mid-term Exam',
        date='2026-01',
        analysis_data=example_data,
        charts_dir='./charts',
        output_path='./example_report.docx'
    )
