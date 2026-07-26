#!/usr/bin/env python3
"""
产品检验规程 Word 文档生成脚本。
根据 JSON 结构化数据，按照标准模板格式生成 .docx 检验规程文档。
"""

import argparse
import json
import os
import sys
import copy
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print(json.dumps({"status": "error", "message": "python-docx 未安装，请执行 pip install python-docx"}))
    sys.exit(1)


# ── 格式常量 ──────────────────────────────────────────────
FONT_TITLE = "黑体"
FONT_HEADING = "黑体"
FONT_BODY = "宋体"
FONT_TABLE = "宋体"
SIZE_TITLE = Pt(16)
SIZE_HEADING = Pt(14)
SIZE_BODY = Pt(10.5)
SIZE_TABLE = Pt(10.5)
SIZE_SMALL = Pt(9)

PAGE_WIDTH = Cm(21)
PAGE_HEIGHT = Cm(29.7)
MARGIN_LEFT = Cm(1.8)
MARGIN_RIGHT = Cm(1.8)
MARGIN_TOP = Cm(2.54)
MARGIN_BOTTOM = Cm(2.54)

BORDER_SZ = "4"  # half-point, ~0.5pt


def set_cell_border(cell, **kwargs):
    """设置单元格边框。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{val.get("sz", BORDER_SZ)}" '
            f'w:space="0" w:color="000000"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_table_borders(table):
    """设置表格整体边框。"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        f'  <w:insideH w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        f'  <w:insideV w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def set_run_font(run, font_name=FONT_BODY, font_size=SIZE_BODY, bold=False):
    """设置 run 的字体属性。"""
    run.font.size = font_size
    run.font.name = font_name
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_paragraph(doc, text, font_name=FONT_BODY, font_size=SIZE_BODY, bold=False,
                  alignment=None, space_before=None, space_after=None, first_line_indent=None):
    """添加段落并设置格式。"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    run = p.add_run(text)
    set_run_font(run, font_name, font_size, bold)
    return p


def add_heading_paragraph(doc, text, level=1):
    """添加章节标题段落。"""
    if level == 1:
        return add_paragraph(doc, text, FONT_HEADING, SIZE_HEADING, bold=True,
                             space_before=Pt(12), space_after=Pt(6))
    else:
        return add_paragraph(doc, text, FONT_HEADING, SIZE_BODY, bold=True,
                             space_before=Pt(6), space_after=Pt(3))


def add_body_paragraph(doc, text, indent=True):
    """添加正文段落。"""
    return add_paragraph(doc, text, FONT_BODY, SIZE_BODY,
                         first_line_indent=Cm(0.74) if indent else None,
                         space_after=Pt(3))


def add_list_item(doc, text, bold_prefix=None, indent=True):
    """添加列表项，可选加粗前缀。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if indent:
        pf.left_indent = Cm(0.74)
    pf.space_after = Pt(2)

    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        set_run_font(run_b, FONT_BODY, SIZE_BODY, bold=True)
        run_t = p.add_run(text)
        set_run_font(run_t, FONT_BODY, SIZE_BODY, False)
    else:
        run = p.add_run(text)
        set_run_font(run, FONT_BODY, SIZE_BODY, False)
    return p


def set_cell_text(cell, text, font_name=FONT_TABLE, font_size=SIZE_TABLE, bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """设置单元格文本和格式。"""
    for paragraph in cell.paragraphs:
        p_elem = paragraph._element
        p_elem.getparent().remove(p_elem)

    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.alignment = alignment
    run = p.add_run(str(text))
    set_run_font(run, font_name, font_size, bold)


def set_cell_shading(cell, color):
    """设置单元格背景色。"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def merge_cells_in_row(table, row_idx, start_col, end_col):
    """合并同一行中的单元格。"""
    cell_start = table.cell(row_idx, start_col)
    cell_end = table.cell(row_idx, end_col)
    cell_start.merge(cell_end)


def create_sampling_table(doc, sampling_data):
    """创建抽样计划表。"""
    plan = sampling_data.get("plan", [])
    if not plan:
        return

    num_rows = len(plan) + 2  # header row 1 + header row 2 + data rows
    table = doc.add_table(rows=num_rows, cols=9)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # 设置列宽
    col_widths = [Cm(3.0), Cm(1.8), Cm(2.0), Cm(1.5), Cm(1.8), Cm(1.7),
                  Cm(1.5), Cm(1.8), Cm(1.7)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            if idx < len(row.cells):
                row.cells[idx].width = width

    # 第一行表头（含合并）
    headers_r1 = [
        (0, 0, "批次数量\n（最小包装数）"),
        (1, 1, "样本量字码"),
        (2, 2, "抽样数量（n）"),
        (3, 5, "合格判定数（Ac）"),
        (6, 8, "不合格判定数（Re）"),
    ]
    for start, end, text in headers_r1:
        cell = table.cell(0, start)
        if start != end:
            cell = cell.merge(table.cell(0, end))
        set_cell_text(cell, text, bold=True)
        set_cell_shading(cell, "D9E2F3")

    # 第二行表头（Ac/Re 子列）
    aql = sampling_data.get("aql", {})
    cri_val = aql.get("CRI", 0)
    maj_val = aql.get("MAJ", 0.65)
    min_val = aql.get("MIN", 1.5)

    sub_headers = ["CRI: " + str(cri_val), "MAJ: " + str(maj_val), "MIN: " + str(min_val),
                   "CRI: 1", "MAJ: " + str(maj_val), "MIN: " + str(min_val)]
    # 合并前3列到上一行
    for col_idx in range(3):
        table.cell(0, col_idx).merge(table.cell(1, col_idx))

    for i, text in enumerate(sub_headers):
        set_cell_text(table.cell(1, 3 + i), text, font_size=SIZE_SMALL)
        set_cell_shading(table.cell(1, 3 + i), "E2EFDA")

    # 数据行
    for row_idx, row_data in enumerate(plan):
        r = row_idx + 2
        values = [
            row_data.get("batch_range", ""),
            row_data.get("sample_code", ""),
            str(row_data.get("sample_size", "")),
            str(row_data.get("ac_cri", "")),
            str(row_data.get("ac_maj", "")),
            str(row_data.get("ac_min", "")),
            str(row_data.get("re_cri", "")),
            str(row_data.get("re_maj", "")),
            str(row_data.get("re_min", "")),
        ]
        for c_idx, val in enumerate(values):
            set_cell_text(table.cell(r, c_idx), val)

    return table


def create_inspection_table(doc, inspection_items):
    """创建检验项目表。"""
    if not inspection_items:
        return

    num_rows = len(inspection_items) + 1  # header + data
    table = doc.add_table(rows=num_rows, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # 列宽
    col_widths = [Cm(1.2), Cm(3.0), Cm(2.0), Cm(3.5), Cm(7.0)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            if idx < len(row.cells):
                row.cells[idx].width = width

    # 表头
    headers = ["序号", "检验项目", "缺陷等级", "检验方法 / 工具", "合格标准 / 要求"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        set_cell_shading(table.cell(0, i), "D9E2F3")

    # 数据行
    for row_idx, item in enumerate(inspection_items):
        r = row_idx + 1
        values = [
            item.get("seq", ""),
            item.get("name", ""),
            item.get("defect_level", ""),
            item.get("method", ""),
            item.get("standard", ""),
        ]
        # 序号和缺陷等级居中，合格标准左对齐
        alignments = [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
        ]
        for c_idx, (val, align) in enumerate(zip(values, alignments)):
            set_cell_text(table.cell(r, c_idx), val, alignment=align)

    return table


def generate_document(data, output_path):
    """根据结构化数据生成检验规程 Word 文档。"""
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM

    # ── 提取数据 ──
    product = data.get("product_info", {})
    doc_info = data.get("document_info", {})
    purpose = data.get("purpose", "")
    scope = data.get("scope", {})
    criteria = data.get("inspection_criteria", {})
    sampling = data.get("sampling", {})
    env = data.get("environment", {})
    inspection_items = data.get("inspection_items", [])
    acceptance = data.get("acceptance_criteria", "")
    nonconf = data.get("nonconformance_handling", {})
    records = data.get("records", [])
    signatures = data.get("signatures", {})

    product_name = product.get("product_name", "待确认")
    product_model = product.get("product_model", "待确认")
    product_category = product.get("product_category", "")
    company_name = doc_info.get("company_name", "待确认")
    company_address = doc_info.get("company_address", "待确认")
    doc_number = doc_info.get("doc_number", "待确认")
    version = doc_info.get("version", "A/0")
    effective_date = doc_info.get("effective_date", "待确认")

    # ── 标题 ──
    title_text = f"{product_name}（{product_model}）入厂检验规程" if product_model else f"{product_name}入厂检验规程"
    if product_category:
        title_text = f"{product_category}—{title_text}"
    add_paragraph(doc, title_text, FONT_TITLE, SIZE_TITLE, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # ── 文件信息 ──
    add_paragraph(doc, f"{company_name}", FONT_BODY, SIZE_BODY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_paragraph(doc, f"地址：{company_address}", FONT_BODY, SIZE_SMALL,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_paragraph(doc, f"文件编号: [{doc_number}]        版本号: {version}        生效日期: [{effective_date}]",
                  FONT_BODY, SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # ── 分隔线 ──
    add_paragraph(doc, "—" * 50, FONT_BODY, SIZE_SMALL,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

    # ── 1. 目的 ──
    add_heading_paragraph(doc, "1  目的")
    if purpose:
        add_body_paragraph(doc, purpose)
    else:
        add_body_paragraph(doc, f"为确保{product_name}（{product_model}）在尺寸、材质、性能等方面完全符合"
                          f"我公司提供的技术图纸与标准要求，防止不合格品流入产线，保障产品质量与可靠性，特制定本检验规范。")

    # ── 2. 范围 ──
    add_heading_paragraph(doc, "2  范围")
    scope_desc = scope.get("description", "")
    if scope_desc:
        add_body_paragraph(doc, scope_desc)
    else:
        add_body_paragraph(doc, f"本规范适用于{product_name}（型号：{product_model}）的入厂检验。")

    product_types = scope.get("product_types", [])
    if product_types:
        for pt in product_types:
            add_list_item(doc, pt)

    # ── 3. 检验准则 ──
    add_heading_paragraph(doc, "3  检验准则")
    add_body_paragraph(doc, "所有检验活动必须依据以下文件，优先级从高到低：")

    criteria_docs = criteria.get("documents", [])
    if criteria_docs:
        for i, d in enumerate(criteria_docs, 1):
            add_list_item(doc, f"{i}. {d}")
    else:
        add_list_item(doc, "1. 本检验规范")
        add_list_item(doc, "2. 经双方签字确认的技术协议")
        add_list_item(doc, "3. 我公司提供的最终版产品图纸（含尺寸、公差、形状等）")
        add_list_item(doc, "4. 我公司提供的材质标准要求")

    standards = criteria.get("standards", [])
    if standards:
        std_text = "引用标准：" + "、".join(standards) + "。"
        add_list_item(doc, std_text, bold_prefix="")

    # ── 4. 抽样方法 ──
    add_heading_paragraph(doc, "4  抽样方法")

    sampling_std = sampling.get("standard", "GB/T 2828.1-2012《计数抽样检验程序》")
    add_list_item(doc, f"采用 {sampling_std}。", bold_prefix="抽样标准：")

    sampling_level = sampling.get("level", "一般检验水平 II")
    add_list_item(doc, f"{sampling_level}。", bold_prefix="检验水平：")

    aql = sampling.get("aql", {})
    add_list_item(doc, "", bold_prefix="接受质量限（AQL）：")
    if aql:
        add_list_item(doc, f"关键项（CRI）：AQL={aql.get('CRI', 0)}（致命缺陷，不允许出现）", indent=True)
        add_list_item(doc, f"重要项（MAJ）：AQL={aql.get('MAJ', 0.65)}（主要缺陷）", indent=True)
        add_list_item(doc, f"次要项（MIN）：AQL={aql.get('MIN', 1.5)}（轻微缺陷）", indent=True)
    else:
        add_list_item(doc, "关键项（CRI）：AQL=0（致命缺陷，不允许出现）", indent=True)
        add_list_item(doc, "重要项（MAJ）：AQL=0.65（主要缺陷）", indent=True)
        add_list_item(doc, "次要项（MIN）：AQL=1.5（轻微缺陷）", indent=True)

    plan_desc = sampling.get("plan_description", "一次正常抽样方案。根据每批交货数量查询抽样计划表确定样本量。")
    add_list_item(doc, plan_desc, bold_prefix="抽样方案：")

    # 抽样计划表
    if sampling.get("plan"):
        add_paragraph(doc, "")
        create_sampling_table(doc, sampling)
        add_paragraph(doc, "（注：需从完整抽样表中查询具体数值，此表为常用范围示例）",
                      FONT_BODY, SIZE_SMALL, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

    # ── 5. 检验流程与项目 ──
    add_heading_paragraph(doc, "5  检验流程与项目")

    # 5.1 来料确认
    add_heading_paragraph(doc, "5.1  来料确认", level=2)
    incoming = data.get("incoming_check", {})
    if incoming:
        items = incoming.get("items", [])
        if items:
            for item in items:
                add_list_item(doc, item)
        else:
            add_body_paragraph(doc, incoming.get("description", "核对送货单、采购订单信息是否一致。检查包装是否完好。"))
    else:
        add_list_item(doc, "核对送货单、采购订单信息（供应商、物料代码、品名、规格、数量）是否一致。")
        add_list_item(doc, "检查包装是否完好，有无破损、受潮、挤压变形。外包装标识是否清晰。")

    # 5.2 检验环境与设备
    add_heading_paragraph(doc, "5.2  检验环境与设备", level=2)
    if env:
        temp = env.get("temperature", "待确认")
        humidity = env.get("humidity", "待确认")
        add_list_item(doc, f"温度：{temp}，湿度：{humidity}。", bold_prefix="环境：")
        equipment = env.get("equipment", [])
        if equipment:
            add_list_item(doc, "、".join(equipment) + "。", bold_prefix="设备：")
        else:
            add_list_item(doc, "待确认。", bold_prefix="设备：")
    else:
        add_list_item(doc, "温度（23±2）℃，湿度（50±5）%RH（优先）。若无条件，需在常温常湿、光照充足的环境下进行。",
                      bold_prefix="环境：")
        add_list_item(doc, "待确认。", bold_prefix="设备：")

    # 5.3 检验项目与方法
    add_heading_paragraph(doc, "5.3  检验项目与方法", level=2)
    if inspection_items:
        create_inspection_table(doc, inspection_items)
    else:
        add_body_paragraph(doc, "检验项目待确认。")

    # ── 6. 合格标准 ──
    add_heading_paragraph(doc, "6  合格标准")
    if acceptance:
        if isinstance(acceptance, list):
            for item in acceptance:
                add_list_item(doc, item)
        else:
            add_body_paragraph(doc, str(acceptance))
    else:
        add_body_paragraph(doc,
            "批合格判定：抽样样本中，所有关键项（CRI）的不合格数为0，"
            "且重要项和次要项的不合格数均小于或等于规定的合格判定数（Ac），则判定该批次为合格。")
        add_body_paragraph(doc,
            "单项判定：任何关键项（CRI）出现一个不合格，即判定该批次不合格。")

    # ── 7. 不合格品处置 ──
    add_heading_paragraph(doc, "7  不合格品处置")
    if nonconf:
        if isinstance(nonconf, dict):
            for key, val in nonconf.items():
                if isinstance(val, list):
                    add_list_item(doc, "", bold_prefix=f"{key}：")
                    for v in val:
                        add_list_item(doc, v, indent=True)
                else:
                    add_list_item(doc, str(val), bold_prefix=f"{key}：")
        elif isinstance(nonconf, str):
            add_body_paragraph(doc, nonconf)
    else:
        add_list_item(doc, '立即对不合格品及其同批产品粘贴红色"不合格"标签，并移至不合格品区隔离，防止误用。',
                      bold_prefix="标识与隔离：")
        add_list_item(doc, "详细填写《进货检验报告》，记录不合格项、比例，并附照片证据。通知采购部和供应商。",
                      bold_prefix="记录与报告：")
        add_list_item(doc, "", bold_prefix="处置方式：")
        add_list_item(doc, "退货/拒收：对于关键项不合格或主要缺陷严重的批次，作整批退货/拒收处理。", indent=True)
        add_list_item(doc, "挑选/返工：仅适用于轻微缺陷且可全数挑选的情况。需供应商派人处理，费用由供应商承担。", indent=True)
        add_list_item(doc, "特采/让步接收：仅在极其特殊的情况下，需经技术、质量部门最高级别评审批准，并记录追溯。", indent=True)
        add_list_item(doc, "向供应商发出《供应商纠正措施报告（SCAR）》，要求分析根本原因并制定改进措施。",
                      bold_prefix="纠正措施：")

    # ── 8. 相关记录 ──
    add_heading_paragraph(doc, "8  相关记录")
    if records:
        for r in records:
            add_list_item(doc, f"《{r}》")
    else:
        add_list_item(doc, "《进货检验报告》")
        add_list_item(doc, "《不合格品评审单》")
        add_list_item(doc, "《供应商纠正措施报告（SCAR）》")

    # ── 签署栏 ──
    add_paragraph(doc, "", space_before=Pt(24))
    compiled = signatures.get("compiled_by", "______________")
    reviewed = signatures.get("reviewed_by", "______________")
    approved = signatures.get("approved_by", "______________")
    add_paragraph(doc, f"编制： {compiled}        审核： {reviewed}        批准： {approved}",
                  FONT_BODY, SIZE_BODY, space_before=Pt(12))

    # ── 保存 ──
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="产品检验规程 Word 文档生成")
    parser.add_argument("--data", required=True, help="JSON 数据文件路径")
    parser.add_argument("--output", required=True, help="输出 .docx 文件路径")
    args = parser.parse_args()

    # 读取 JSON 数据
    try:
        with open(args.data, "r", encoding="utf-8") as f:
            data = json.load(f)
    except FileNotFoundError:
        print(json.dumps({"status": "error", "message": f"数据文件不存在: {args.data}"}, ensure_ascii=False))
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(json.dumps({"status": "error", "message": f"JSON 解析失败: {str(e)}"}, ensure_ascii=False))
        sys.exit(1)

    # 校验必填字段
    required_fields = {
        "product_info": ["product_name"],
        "inspection_items": None,  # 列表不能为空
    }
    warnings = []
    product_info = data.get("product_info", {})
    if not product_info.get("product_name"):
        warnings.append("缺少产品名称(product_info.product_name)")
    if not data.get("inspection_items"):
        warnings.append("缺少检验项目(inspection_items)")

    # 生成文档
    try:
        output_path = generate_document(data, args.output)
        result = {
            "status": "success",
            "output_path": output_path,
            "file_size": os.path.getsize(output_path),
        }
        if warnings:
            result["warnings"] = warnings
        print(json.dumps(result, ensure_ascii=False))
    except Exception as e:
        print(json.dumps({"status": "error", "message": f"文档生成失败: {str(e)}"}, ensure_ascii=False))
        sys.exit(1)


if __name__ == "__main__":
    main()
