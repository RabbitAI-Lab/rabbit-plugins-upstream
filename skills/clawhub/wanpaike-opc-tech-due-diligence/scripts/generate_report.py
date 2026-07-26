#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
技术尽调报告生成脚本 - 通用版
支持参数化调用，也保留原有硬编码示例行为。

用法:
    # 无参数：使用默认硬编码示例（河北青山鼎信）
    python generate_report.py

    # 带参数：生成指定公司的报告
    python generate_report.py --company "某科技公司" --output "报告路径" [--json "findings.json"]

参数:
    --company  公司名称（必填，使用参数时）
    --output   输出文件路径（可选，默认自动生成）
    --json     JSON格式的尽调数据文件路径（可选，默认使用内置示例数据）
"""

import json
import sys
import os
import argparse
from datetime import datetime
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    tcPr.append(shading_elm)


def add_table_with_style(doc, rows, cols, data, header_color="4472C4"):
    """创建表格并应用样式"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 表头样式
            if i == 0:
                set_cell_shading(cell, header_color)
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

    return table


def set_table_width(table, width=8800):
    """设置表格宽度"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{width}" w:type="dxa"/>')
        tblPr.append(tblW)
    else:
        tblW.set(qn('w:w'), str(width))
        tblW.set(qn('w:type'), 'dxa')


def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.size = Pt(16)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    elif level == 2:
        heading.runs[0].font.size = Pt(14)
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(4)
    elif level == 3:
        heading.runs[0].font.size = Pt(12)
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(3)
    return heading


def add_paragraph(doc, text):
    """添加段落"""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    for run in para.runs:
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return para


def add_bullet_list(doc, items):
    """添加项目符号列表"""
    for item in items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        for run in para.runs:
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


# ====================================================================
# 示例数据：河北青山鼎信新能源科技有限公司
# 当无参数调用时，使用此示例数据生成报告
# ====================================================================
SAMPLE_COMPANY_NAME = "河北青山鼎信新能源科技有限公司"

SAMPLE_DATA = {
    "basic_info": {
        "公司名称": {"claimed": "河北青山鼎信新能源科技有限公司", "verified": "✅ 已确认", "note": "企查查可查"},
        "成立日期": {"claimed": "2019年8月21日", "verified": "✅ 已确认", "note": "国家企业信用信息公示系统"},
        "注册资本": {"claimed": "2001万元人民币", "verified": "✅ 已确认", "note": "实缴暂无公开信息"},
        "法定代表人": {"claimed": "王一蔚（2024年11月变更）", "verified": "✅ 已确认", "note": "原为李肥子"},
        "参保人数": {"claimed": "10人（2024年）", "verified": "✅ 已确认", "note": "企查查"},
        "企业类型": {"claimed": "其他有限责任公司", "verified": "✅ 已确认", "note": "2024年9月变更"},
    },
    "team_info": {
        "核心团队": [
            {"name": "王一蔚", "role": "法定代表人、执行董事", "verified": "🟢可信"},
            {"name": "陈炯", "role": "总经理（海澜电力）", "verified": "🟢可信"},
            {"name": "李肥子", "role": "原法定代表人（已退出）", "verified": "🟡待核实"},
        ],
        "研发人员": "0人（无专利）"
    },
    "patent_info": {
        "专利总数": {"claimed": "0项", "verified": "✅ 已确认", "note": "无任何专利申请记录"},
        "商标": {"claimed": "1项", "verified": "🟢已确认", "note": "青山鼎信（2024年注册）"},
        "著作权": {"claimed": "0项", "verified": "✅ 已确认", "note": "无软件著作权登记"},
    },
    "tech_claims": [
        {"声明": "新能源科技公司", "核实结果": "❌ 名不副实", "说明": "0专利0软著，科创等级入门", "质疑点": "科技属性薄弱"},
        {"声明": "服务1400家客户", "核实结果": "⚠️ 待核实", "说明": "10人支撑1400家客户存疑", "质疑点": "人效比极度异常"},
        {"声明": "年交易量70亿千瓦时", "核实结果": "🟢 可信", "说明": "冀南60亿+冀北10亿", "质疑点": "依赖海澜电力"},
    ],
    "risk_assessment": [
        {"风险类型": "团队能力不足", "等级": "高", "描述": "10人支撑1400家客户，人效比极度异常"},
        {"风险类型": "技术能力薄弱", "等级": "高", "描述": "0专利、0软著，科创等级入门"},
        {"风险类型": "股权结构复杂", "等级": "高", "描述": "保定如亿壳公司嫌疑，持股33%"},
        {"风险类型": "市场竞争激烈", "等级": "高", "描述": "河北223家售电公司竞争"},
        {"风险类型": "盈利不确定性", "等级": "中", "描述": "暂无公开财务数据"},
    ],
    "overall_score": 5.5,
    "investment_advice": "⚠️ 有条件投资（需完成深入尽调）",
    "data_sources": ["企查查", "启信宝", "水滴信用", "国家企业信用信息公示系统", "河北电力交易中心"],
    "follow_up": [
        {"事项": "核实10人团队劳动关系", "优先级": "极高", "说明": "灵活用工/外包情况"},
        {"事项": "核实与海澜电力的整合深度", "优先级": "极高", "说明": "独立运营还是完全依赖"},
        {"事项": "核实保定如亿实际控制关系", "优先级": "极高", "说明": "是否存在代持安排"},
        {"事项": "获取近三年财务报表", "优先级": "高", "说明": "盈利能力、现金流核实"},
        {"事项": "核查核心资质", "优先级": "高", "说明": "电力业务许可证等"},
    ],
    "pros": ["有河北售电市场基础业务", "海澜电力控股后获品牌背书", "客户基础明确"],
    "cons": ["人员规模与业务不匹配", "技术能力薄弱", "股东结构复杂"],
}


def create_report(company_name: str = SAMPLE_COMPANY_NAME, data: Optional[Dict] = None, output_path: Optional[str] = None):
    """
    生成技术尽调报告docx

    Args:
        company_name: 公司名称
        data: 尽调数据字典（可选，不传则使用内置示例数据）
        output_path: 输出文件路径（可选，默认自动生成）
    """
    if data is None:
        data = SAMPLE_DATA

    if output_path is None:
        output_path = f"{company_name}技术尽调报告.docx"

    doc = Document()

    # 设置文档默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(10.5)

    # ========== 封面 ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    title_run = title.add_run('技术尽调报告')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(24)
    sub_run = subtitle.add_run(company_name)
    sub_run.font.size = Pt(20)
    sub_run.font.bold = True
    sub_run.font.name = '黑体'
    sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_before = Pt(400)
    date_run = date_para.add_run(f'报告日期：{datetime.now().strftime("%Y年%m月%d日")}')
    date_run.font.size = Pt(14)
    date_run.font.name = '宋体'
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 保密声明
    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_para.paragraph_format.space_before = Pt(200)
    conf_run = conf_para.add_run('【保密文件】')
    conf_run.font.size = Pt(12)
    conf_run.font.color.rgb = RGBColor(192, 0, 0)
    conf_run.font.bold = True

    doc.add_page_break()

    # ========== 目录占位 ==========
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_run = toc_title.add_run('目  录')
    toc_run.font.size = Pt(18)
    toc_run.font.bold = True
    toc_run.font.name = '黑体'
    toc_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    toc_items = [
        '第一章 执行摘要',
        '第二章 公司概况',
        '第三章 第一层：材料完整性检查',
        '第四章 第二层：技术可行性分析',
        '第五章 第三层：商业逻辑验证',
        '第六章 第四层：风险识别与评级',
        '第七章 红旗信号清单',
        '第八章 综合风险评级',
        '第九章 投资建议',
        '附录'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ========== 第一章 执行摘要 ==========
    add_heading(doc, '第一章 执行摘要', 1)
    score = data.get("overall_score", 5.0)
    if score >= 7:
        summary_text = f"本报告对{company_name}进行了全面的技术尽调分析。公司整体技术实力较强，团队背景可靠，商业模式清晰，建议重点关注。"
    elif score >= 5:
        summary_text = f"本报告对{company_name}进行了全面的技术尽调分析。公司有一定技术或业务基础，但存在若干风险点，需进一步核实后决策。"
    else:
        summary_text = f"本报告对{company_name}进行了全面的技术尽调分析。公司存在较多不确定性或重大风险，建议谨慎评估。"

    add_paragraph(doc, summary_text)

    # 核心发现表格
    add_heading(doc, '1.1 核心发现概览', 2)
    risk_items = data.get("risk_assessment", [])
    findings_rows = [['维度', '关键发现', '风险等级']]
    for r in risk_items[:6]:
        findings_rows.append([r.get("风险类型", ""), r.get("描述", ""), f'🔴{r.get("等级", "")}' if r.get("等级") == "高" else f'🟠{r.get("等级", "")}'])

    table = add_table_with_style(doc, len(findings_rows), 3, findings_rows)
    set_table_width(table)

    score = data.get("overall_score", 5.0)
    if score >= 7:
        add_paragraph(doc, f"综合评估结论：{company_name}技术实力较强，综合评分{score}/10，建议进入深度尽调。")
    elif score >= 5:
        add_paragraph(doc, f"综合评估结论：{company_name}存在一定风险，综合评分{score}/10，建议补充核实后决策。")
    else:
        add_paragraph(doc, f"综合评估结论：{company_name}存在重大风险，综合评分{score}/10，不建议投资。")

    # ========== 第二章 公司概况 ==========
    add_heading(doc, '第二章 公司概况', 1)
    add_paragraph(doc, f'本章节从公司基本信息、股权结构、融资历程和发展里程碑四个维度全面呈现{company_name}的企业画像。')

    add_heading(doc, '2.1 基本信息', 2)
    basic = data.get("basic_info", {})
    basic_rows = [['项目', '内容']]
    for key, value in basic.items():
        basic_rows.append([key, value.get("claimed", str(value))])
    table = add_table_with_style(doc, len(basic_rows), 2, basic_rows)
    set_table_width(table)

    # ========== 第三章 第一层：材料完整性检查 ==========
    add_heading(doc, '第三章 第一层：材料完整性检查', 1)
    add_paragraph(doc, f'本章节从团队背景、专利/资质、备案信息、负面信息四个维度核查{company_name}的材料完整性和真实性。')

    # 团队核查
    add_heading(doc, '3.1 团队背景核查', 2)
    team = data.get("team_info", {})
    core_team = team.get("核心团队", [])
    if core_team:
        team_rows = [['姓名', '职位', '核查结论']]
        for member in core_team:
            team_rows.append([member.get("name", ""), member.get("role", ""), member.get("verified", "待核实")])
        table = add_table_with_style(doc, len(team_rows), 3, team_rows)
        set_table_width(table)

    # 专利核查
    add_heading(doc, '3.2 专利/资质核查', 2)
    patent = data.get("patent_info", {})
    patent_rows = [['类型', '声明内容', '核实结果', '说明']]
    for key, value in patent.items():
        patent_rows.append([key, value.get("claimed", ""), value.get("verified", "待核实"), value.get("note", "")])
    table = add_table_with_style(doc, len(patent_rows), 4, patent_rows)
    set_table_width(table)

    # ========== 第四章 第二层：技术可行性分析 ==========
    add_heading(doc, '第四章 第二层：技术可行性分析', 1)
    add_paragraph(doc, '本章节基于物理定律核查、技术成熟度评估、学术发表验证三个维度评估技术可行性。')

    # 技术声明核查
    add_heading(doc, '4.1 技术声明核查', 2)
    claims = data.get("tech_claims", [])
    if claims:
        claim_rows = [['技术声明', '核实结果', '说明', '质疑点']]
        for c in claims:
            claim_rows.append([c.get("声明", ""), c.get("核实结果", ""), c.get("说明", ""), c.get("质疑点", "")])
        table = add_table_with_style(doc, len(claim_rows), 4, claim_rows)
        set_table_width(table)

    # ========== 第五章 第三层：商业逻辑验证 ==========
    add_heading(doc, '第五章 第三层：商业逻辑验证', 1)
    add_paragraph(doc, '本章节从商业模式画布、竞争格局、市场规模、估值分析四个维度验证商业逻辑。')

    add_heading(doc, '5.1 商业模式画布', 2)
    bm_rows = [['模块', '内容']]
    bm_data = [
        ['价值主张', '待评估'],
        ['客户细分', '待评估'],
        ['渠道通路', '待评估'],
        ['收入来源', '待评估'],
        ['核心资源', '待评估'],
        ['成本结构', '待评估'],
    ]
    for row in bm_data:
        bm_rows.append(row)
    table = add_table_with_style(doc, len(bm_rows), 2, bm_rows)
    set_table_width(table)

    # ========== 第六章 第四层：风险识别与评级 ==========
    add_heading(doc, '第六章 第四层：风险识别与评级', 1)
    add_paragraph(doc, '本章节从技术风险、市场风险、团队风险、财务风险、法律风险五个维度识别潜在风险。')

    risks = data.get("risk_assessment", [])
    # 分组
    risk_types = set(r.get("风险类型", "") for r in risks)
    for rtype in ["技术风险", "市场风险", "团队风险", "财务风险", "法律风险"]:
        type_risks = [r for r in risks if rtype in r.get("风险类型", "")]
        if type_risks:
            add_heading(doc, f'6.{list(risk_types).index(rtype)+1 if rtype in risk_types else 0} {rtype}', 2)
            r_rows = [['风险点', '具体描述', '等级']]
            for r in type_risks:
                r_rows.append([r.get("风险类型", ""), r.get("描述", ""), f'🔴{r.get("等级", "")}' if r.get("等级") == "高" else f'🟠{r.get("等级", "")}'])
            table = add_table_with_style(doc, len(r_rows), 3, r_rows)
            set_table_width(table)

    # ========== 第七章 红旗信号清单 ==========
    add_heading(doc, '第七章 红旗信号清单', 1)
    add_paragraph(doc, '以下为本次尽调中发现的重大风险信号，需要重点关注或进一步核实：')

    # 严重红旗
    add_heading(doc, '7.1 严重红旗（🔴需立即关注）', 2)
    high_risks = [r for r in risks if r.get("等级") == "高"]
    if high_risks:
        items = [f'【{r.get("风险类型", "")}】{r.get("描述", "")}' for r in high_risks]
        add_bullet_list(doc, items)
    else:
        add_paragraph(doc, '未发现严重红旗信号。')

    # ========== 第八章 综合风险评级 ==========
    add_heading(doc, '第八章 综合风险评级', 1)
    add_paragraph(doc, f'本章节基于四层尽调发现，对{company_name}进行综合风险评级。')

    add_heading(doc, '8.1 最终评级结论', 2)
    final_rows = [
        ['评级维度', '结论', '依据'],
        ['综合评分', f'{score}/10', '基于四层尽调框架综合评估'],
        ['投资建议', data.get("investment_advice", "待评估"), '见第九章'],
    ]
    table = add_table_with_style(doc, len(final_rows), 3, final_rows)
    set_table_width(table)

    # ========== 第九章 投资建议 ==========
    add_heading(doc, '第九章 投资建议', 1)
    add_paragraph(doc, f'基于本次技术尽调的四层评估框架，对{company_name}提出如下投资建议：')

    pros = data.get("pros", [])
    cons = data.get("cons", [])
    if pros:
        add_heading(doc, '9.1 推荐理由', 2)
        add_bullet_list(doc, [f'🟢 {p}' for p in pros])
    if cons:
        add_heading(doc, '9.2 谨慎理由', 2)
        add_bullet_list(doc, [f'🔴 {c}' for c in cons])

    # ========== 附录 ==========
    add_heading(doc, '附录', 1)
    add_paragraph(doc, '本附录包含信息核查清单、数据来源和免责声明等内容。')

    add_heading(doc, '附录A：信息核查清单', 2)
    check_rows = [['核查项', '状态', '说明']]
    for key, value in basic.items():
        check_rows.append([key, value.get("verified", "待核实"), value.get("note", "")])
    table = add_table_with_style(doc, len(check_rows), 3, check_rows)
    set_table_width(table)

    add_heading(doc, '附录B：数据来源', 2)
    sources = data.get("data_sources", [])
    if sources:
        add_bullet_list(doc, sources)

    add_heading(doc, '附录C：免责声明', 2)
    add_paragraph(doc, '本报告仅基于公开可获取的信息编制，不构成任何投资建议。报告中的分析、评级和建议仅供参考，不对报告使用人或第三方因使用本报告而产生的任何损失承担责任。投资决策应由投资方在充分尽调和独立判断的基础上做出。本报告禁止未经授权的传播和引用。')

    # 保存文档
    doc.save(output_path)
    print(f"✅ 报告已生成：{output_path}")
    return output_path


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(description='技术尽调报告生成器')
    parser.add_argument('--company', type=str, help='公司名称')
    parser.add_argument('--output', type=str, help='输出文件路径')
    parser.add_argument('--json', type=str, help='JSON格式的尽调数据文件路径')

    args = parser.parse_args()

    if args.company:
        # 参数化模式
        data = None
        if args.json:
            with open(args.json, 'r', encoding='utf-8') as f:
                data = json.load(f)
        output_path = args.output or f"{args.company}技术尽调报告.docx"
        create_report(args.company, data, output_path)
    else:
        # 无参数：使用默认示例
        print("ℹ️ 无参数传入，使用默认示例数据（河北青山鼎信）")
        print("   用法: python generate_report.py --company \"公司名称\" [--json data.json]")
        create_report()


if __name__ == "__main__":
    main()