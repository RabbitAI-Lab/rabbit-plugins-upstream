#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成docx格式技术尽调报告 - 通用版
支持参数化调用，也保留原有硬编码示例行为。

用法:
    # 无参数：使用默认硬编码示例（菲瑞药业）
    python generate_docx.py

    # 带参数：生成指定公司的报告
    python generate_docx.py --company "某科技公司" --output "报告路径" [--json "findings.json"]

参数:
    --company  公司名称（必填，使用参数时）
    --output   输出文件路径（可选，默认自动生成）
    --json     JSON格式的尽调数据文件路径（可选，默认使用内置示例数据）
"""

import json
import sys
import os
import argparse
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraph(doc, text):
    """添加段落"""
    p = doc.add_paragraph(text)
    return p


def add_table_from_list(doc, headers, rows):
    """从列表添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 添加表头
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        hdr_cells[idx].paragraphs[0].runs[0].bold = True

    # 添加数据行
    for row_idx, row in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row):
            row_cells[col_idx].text = str(cell_text)

    return table


# ====================================================================
# 示例数据：菲瑞药业
# 当无参数调用时，使用此示例数据生成报告
# ====================================================================
SAMPLE_DOCX_DATA = {
    "company_name": "湖北菲瑞生物药业有限公司",
    "overall_score": 4.5,
    "investment_advice": "谨慎推荐（下调观察）",
    "scores": {
        "技术壁垒": 5.0,
        "团队实力": 4.0,
        "市场前景": 6.0,
        "财务健康": 4.0,
        "商业化能力": 5.0,
        "风险可控性": 3.0,
    },
    "layer1": [
        ["公司成立时间", "2019年11月", "⚠️ 存疑", "企查查显示主体公司成立于2022年12月28日"],
        ["注册资本", "2596万元", "✅ 基本属实", "企查查显示2596.9232万元，实缴一致"],
        ["员工人数", "近100人", "❌ 严重不符", "企查查显示2024年参保人数仅4人"],
        ["办公地址", "鄂州市", "✅ 属实", "鄂州市鄂城区樊口街道旭光大道18号"],
        ["瞪羚企业", "2025年7月获评", "✅ 属实", "湖北省2025年入库瞪羚企业"],
    ],
    "layer1_patent": [
        ["专利总数", "18项", "⚠️ 13项", "企查查显示专利13项（部分实质审查中）"],
        ["发明专利", "8项", "⚠️ 存疑", "部分专利如布洛芬状态为「实质审查」"],
        ["实用新型", "7项", "✅ 吻合", "-"],
        ["软件著作权", "3项", "✅ 吻合", "-"],
    ],
    "layer2_claims": [
        "❌ 「世界唯二、中国唯一」说法严重失实：陕西量子高科明确声称自己是该领域的开创者和引领者",
        "❌ 市场规模夸大数千倍：万亿 vs 实际12亿，相差过于悬殊",
        "❌ 员工人数严重不符：声称100人 vs 实际参保4人，差异无法合理解释",
        "❌ 专利数量存在夸大：18项 vs 实际13项",
    ],
    "risk_assessment": [
        ["技术壁垒夸大", "高", "「5-8年领先优势」基于「世界唯二、中国唯一」，但实际存在强劲竞争对手"],
        ["竞争对手突破", "高", "齐鲁、石药等上市公司具备更强研发和资金实力"],
        ["估值偏高", "高", "PS达14倍（5000万营收/7亿估值），远高于行业可比公司"],
        ["营收规模小", "高", "年营收5000万，抵御风险能力有限"],
        ["资金缺口大", "高", "计划2029年前投入7.65亿元，资金来源不明"],
        ["申报周期长", "高", "2类药5-6年，3类药2-3年，时间窗口存在不确定性"],
    ],
    "pros": [
        "冻干闪释技术确有一定市场应用价值，符合老龄化社会的用药需求趋势",
        "瞪羚企业认证、ISO认证等资质具备一定可信度",
        "湖北地方政府产业基金背书",
        "已有营养品OEM业务实现营收",
    ],
    "cons": [
        "「世界唯二、中国唯一」说法严重失实",
        "市场规模夸大数千倍",
        "员工人数严重不符：声称100人 vs 实际参保4人",
        "专利数量存在夸大：18项 vs 实际13项",
        "竞争对手众多：齐鲁、石药、扬子江等多家上市公司均有布局",
        "估值偏高：PS达14倍",
        "IPO目标激进：2029年上市不确定性大",
    ],
    "follow_up": [
        ["实际员工人数和团队构成", "极高", "要求提供社保缴纳明细、员工名册"],
        ["药品备案号真实性", "极高", "要求提供备案证原件或药监局官网截图"],
        ["技术评估报告原件", "高", "要求提供鄂技交评字【2023】第S130号原件"],
        ["财务数据核实", "高", "要求提供经审计的财务报表"],
        ["陕西量子高科对比分析", "高", "详细了解竞争对手的技术、产品、商业模式"],
    ],
    "data_sources": ["企查查", "启信宝", "官网", "新闻报道", "药监局网站"],
}


def create_report(company_name: str = SAMPLE_DOCX_DATA["company_name"],
                  data: Optional[Dict] = None,
                  output_path: Optional[str] = None):
    """
    生成技术尽调报告docx

    Args:
        company_name: 公司名称
        data: 尽调数据字典（可选，不传则使用内置示例数据）
        output_path: 输出文件路径（可选，默认自动生成）
    """
    if data is None:
        data = SAMPLE_DOCX_DATA

    if output_path is None:
        output_path = f"{company_name}技术尽调报告.docx"

    doc = Document()

    # 设置文档标题
    title = doc.add_heading(f'{company_name}技术尽调报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加文档信息
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.add_run('报告日期：').bold = True
    info_para.add_run(f'{datetime.now().strftime("%Y年%m月%d日")}\n')
    info_para.add_run('尽调对象：').bold = True
    info_para.add_run(f'{company_name}\n')
    info_para.add_run('尽调级别：').bold = True
    info_para.add_run('全面技术尽调（四层分析法）\n')
    info_para.add_run('报告密级：').bold = True
    info_para.add_run('内部资料')

    doc.add_paragraph()
    doc.add_paragraph('_' * 60)

    # 摘要与核心结论
    add_heading(doc, '摘要与核心结论', 1)
    add_heading(doc, '一句话结论', 2)

    score = data.get("overall_score", 5.0)
    if score >= 7:
        doc.add_paragraph(f'{company_name}技术实力强，团队背景可靠，建议重点关注。')
    elif score >= 5:
        doc.add_paragraph(f'{company_name}有一定技术或业务基础，但存在若干风险点，需进一步核实后决策。')
    else:
        doc.add_paragraph(f'{company_name}存在较多不确定性或重大风险，建议谨慎评估。')

    add_heading(doc, '投资建议等级', 2)
    doc.add_paragraph()

    # 评分表格
    scores = data.get("scores", {})
    score_rows = []
    for dim, val in scores.items():
        score_rows.append([dim, str(val), ''])
    if score_rows:
        headers = ['评估维度', '评分（1-10）', '核心发现']
        add_table_from_list(doc, headers, score_rows)

    doc.add_paragraph()

    # 第一层
    add_heading(doc, '第一层：信息提取与核实', 1)

    add_heading(doc, '1.1 公司基本信息核实', 2)
    layer1 = data.get("layer1", [])
    if layer1:
        headers = ['声明内容', '融资计划书', '核实结果', '说明']
        add_table_from_list(doc, headers, layer1)

    add_heading(doc, '1.2 专利情况核实', 2)
    layer1_patent = data.get("layer1_patent", [])
    if layer1_patent:
        headers = ['声明内容', '融资计划书', '核实结果', '说明']
        add_table_from_list(doc, headers, layer1_patent)

    # 第二层
    add_heading(doc, '第二层：交叉验证与质疑', 1)
    layer2_claims = data.get("layer2_claims", [])
    for claim in layer2_claims:
        doc.add_paragraph(claim)

    # 第三层
    add_heading(doc, '第三层：深度风险分析', 1)
    risks = data.get("risk_assessment", [])
    risk_categories = [
        ("技术风险", [r for r in risks if "技术" in r[0] or "专利" in r[0]]),
        ("商业风险", [r for r in risks if "估值" in r[0] or "营收" in r[0] or "资金" in r[0]]),
        ("财务风险", [r for r in risks if "营收" in r[0] or "资金" in r[0] or "盈利" in r[0]]),
        ("药品申报风险", [r for r in risks if "申报" in r[0] or "周期" in r[0]]),
    ]
    for cat_name, cat_risks in risk_categories:
        if cat_risks:
            add_heading(doc, f'3.{risk_categories.index((cat_name, cat_risks))+1} {cat_name}', 2)
            headers = ['风险类型', '风险等级', '具体描述']
            add_table_from_list(doc, headers, cat_risks)

    # 第四层
    add_heading(doc, '第四层：投资建议与评分', 1)

    add_heading(doc, '4.1 综合评分', 2)
    if scores:
        score_rows_2 = []
        for dim, val in scores.items():
            score_rows_2.append([dim, str(val), ''])
        headers = ['评估维度', '评分（1-10）', '说明']
        add_table_from_list(doc, headers, score_rows_2)

    add_heading(doc, '4.2 投资建议', 2)
    p = doc.add_paragraph()
    p.add_run(data.get("investment_advice", "待评估")).bold = True

    pros = data.get("pros", [])
    cons = data.get("cons", [])
    if pros:
        doc.add_paragraph()
        doc.add_paragraph('推荐理由：')
        for r in pros:
            doc.add_paragraph(f'• 🟢 {r}')
    if cons:
        doc.add_paragraph()
        doc.add_paragraph('谨慎理由：')
        for c in cons:
            doc.add_paragraph(f'• 🔴 {c}')

    # 附录
    add_heading(doc, '附录', 1)

    add_heading(doc, '附录A：需要进一步核实的事项清单', 2)
    follow_up = data.get("follow_up", [])
    if follow_up:
        headers = ['序号', '核实事项', '优先级', '说明']
        numbered_rows = [[str(i+1)] + row for i, row in enumerate(follow_up)]
        add_table_from_list(doc, headers, numbered_rows)

    add_heading(doc, '附录B：数据来源', 2)
    sources = data.get("data_sources", [])
    for s in sources:
        doc.add_paragraph(f'• {s}')

    # 页脚
    doc.add_paragraph()
    doc.add_paragraph('_' * 60)
    footer = doc.add_paragraph()
    footer.add_run('报告完成时间：').bold = True
    footer.add_run(f'{datetime.now().strftime("%Y年%m月%d日")}\n')
    footer.add_run('报告编制人：').bold = True
    footer.add_run('技术尽调团队\n')
    footer.add_run('免责声明：').bold = True
    footer.add_run('本报告基于公开信息和有限核查编制，仅供参考，不构成投资建议。建议投资前进行更深入的尽职调查。')

    # 保存文档
    doc.save(output_path)
    print(f"✅ docx报告已生成：{output_path}")
    return output_path


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(description='技术尽调报告docx生成器')
    parser.add_argument('--company', type=str, help='公司名称')
    parser.add_argument('--output', type=str, help='输出文件路径')
    parser.add_argument('--json', type=str, help='JSON格式的尽调数据文件路径')

    args = parser.parse_args()

    if args.company:
        data = None
        if args.json:
            with open(args.json, 'r', encoding='utf-8') as f:
                data = json.load(f)
        output_path = args.output or f"{args.company}技术尽调报告.docx"
        create_report(args.company, data, output_path)
    else:
        # 无参数：使用默认示例
        print("ℹ️ 无参数传入，使用默认示例数据（菲瑞药业）")
        print("   用法: python generate_docx.py --company \"公司名称\" [--json data.json]")
        create_report()


if __name__ == '__main__':
    main()