#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
gongwen_check.py — 公文格式合规质检 (GB/T 9704-2012)

对已有 .docx 公文做机械式格式扫描，输出分级检查报告。
只检查「可机器判定」的客观格式项，不对内容作价值判断。

用法:
    python gongwen_check.py --input 待检公文.docx
    python gongwen_check.py --input 待检公文.docx --report 检查报告.md
    python gongwen_check.py --input 待检公文.docx --json          # 机器可读输出

风险等级:
    [高] 明显偏离国标强制性格式规定，正式行文前应当修改
    [中] 与国标推荐做法不一致，建议核对本单位行文细则
    [低] 提示性事项，需人工确认
"""

import argparse
import json
import os
import re
import shutil
import sys
from datetime import date, datetime

try:
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn
except ImportError:
    sys.stderr.write("缺少依赖 python-docx。请先安装：pip install python-docx\n")
    sys.exit(2)


TOL = 0.03  # cm 容差

SPEC = {
    "top": 3.7, "bottom": 3.5, "left": 2.8, "right": 2.6,
    "page_w": 21.0, "page_h": 29.7,
}

FANGSONG_OK = {"仿宋_GB2312", "仿宋", "FangSong", "FangSong_GB2312", "SimSun", "宋体"}
HEITI_OK = {"黑体", "SimHei"}
KAITI_OK = {"楷体_GB2312", "楷体", "KaiTi", "SimKai"}
TITLE_OK = {"方正小标宋简体", "FZXiaoBiaoSong-B05S", "小标宋体", "宋体", "SimSun"}

L1 = re.compile(r"^[一二三四五六七八九十百]+、")
L2 = re.compile(r"^[（(][一二三四五六七八九十百]+[）)]")
L3 = re.compile(r"^\d+[\.．、](?!\d)")
L4 = re.compile(r"^[（(]\d+[）)]")
DOCNUM = re.compile(r"[〔\[【(（]\s*(\d{4})\s*[〕\]】)）]\s*(\d+)\s*号")
DATE_CN = re.compile(r"(\d{4})年(\d{1,2})月(\d{1,2})日")
DATE_CN_ZH = re.compile(r"[二〇一二三四五六七八九十]{4}年")


class Issue:
    __slots__ = ("level", "item", "found", "expect", "advice", "loc")

    def __init__(self, level, item, found, expect, advice, loc=""):
        self.level, self.item = level, item
        self.found, self.expect = found, expect
        self.advice, self.loc = advice, loc

    def as_dict(self):
        return {"level": self.level, "item": self.item, "found": self.found,
                "expect": self.expect, "advice": self.advice, "loc": self.loc}


def cm(v):
    return round(v.cm, 2) if v is not None else None


def _backup(path):
    """写报告前，先把已存在的原报告复制到同目录 .bak 备份，防覆盖丢失。返回备份路径或 None。"""
    if not path or not os.path.exists(path):
        return None
    d = os.path.dirname(os.path.abspath(path)) or "."
    base = os.path.basename(path)
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    bak = os.path.join(d, f".{base}.bak-{ts}")
    n = 1
    while os.path.exists(bak):
        bak = os.path.join(d, f".{base}.bak-{ts}-{n}")
        n += 1
    try:
        shutil.copy2(path, bak)
        return bak
    except OSError:
        return None


def _safe_write_text(path, content, label="文件"):
    """写文本报告前自动备份；捕获占用/权限/磁盘错误并给可操作提示。"""
    bak = _backup(path)
    directory = os.path.dirname(os.path.abspath(path)) or "."
    os.makedirs(directory, exist_ok=True)
    try:
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
    except PermissionError:
        sys.stderr.write(
            f"[err] 无法写入 {label} {path}：文件可能被其他程序占用或无写入权限。\n"
            f"      请关闭该文件后重试，或更换输出路径。\n")
        sys.exit(3)
    except OSError as e:
        sys.stderr.write(f"[err] 写入 {label} {path} 失败：{e}\n")
        sys.exit(3)
    if bak:
        sys.stderr.write(f"[i ] 原文件已备份：{bak}\n")


def run_font(run):
    """取 run 的东亚字体名。"""
    rPr = run._element.find(qn("w:rPr"))
    if rPr is not None:
        rF = rPr.find(qn("w:rFonts"))
        if rF is not None:
            ea = rF.get(qn("w:eastAsia"))
            if ea:
                return ea
    return run.font.name or ""


def run_size_pt(run, default=None):
    return run.font.size.pt if run.font.size else default


def run_is_red(run):
    """红色文字为发文机关标志、印章占位等版头/印章要素，不参与正文字体字号判定。"""
    try:
        c = run.font.color
        if c is not None and c.rgb is not None:
            return str(c.rgb).upper() in ("FF0000", "C00000", "CC0000")
    except Exception:
        pass
    return False


def para_info(p):
    txt = p.text.strip()
    fonts, sizes, reds = [], [], []
    for r in p.runs:
        if not r.text.strip():
            continue
        fonts.append(run_font(r))
        reds.append(run_is_red(r))
        s = run_size_pt(r)
        if s:
            sizes.append(s)
    pf = p.paragraph_format
    return {
        "text": txt,
        "font": fonts[0] if fonts else "",
        "fonts": fonts,
        "size": sizes[0] if sizes else None,
        "red": bool(reds) and all(reds),
        "first_indent": cm(pf.first_line_indent),
        "left_indent": cm(pf.left_indent) or 0.0,
        "line_rule": str(pf.line_spacing_rule),
        "line_val": pf.line_spacing.pt if hasattr(pf.line_spacing, "pt") else pf.line_spacing,
        "align": str(p.alignment),
    }


# 不适用「首行缩进 2 字」的要素：这些要素国标另有编排规则
NO_INDENT_OK = re.compile(
    r"^(附\s*件|抄\s*送|印\s*发|签发人|附\s*注|[（(]此件|[〔【].{0,4}印章|"
    r"\d{6}$|绝密|机密|秘密|特急|加急|平急)")


def is_body_like(d):
    """判断该段是否应当适用正文排版规则（3 号仿宋 + 首行缩进 2 字）。"""
    t = d["text"]
    if not t or d["red"]:
        return False
    if "RIGHT" in d["align"] or "CENTER" in d["align"]:
        return False            # 署名、成文日期、标题、发文字号（居中）
    if NO_INDENT_OK.match(t):
        return False
    if DOCNUM.search(t) and len(t) < 30:
        return False            # 发文字号
    # 命中层次序号的一律计入正文体系（即使很短，如「一、总体要求」），
    # 必须排在短行排除规则之前，否则会漏检层级标题的字体错误
    if L1.match(t) or L2.match(t) or L3.match(t) or L4.match(t):
        return True
    if t.endswith(("：", ":")) and len(t) < 80:
        return False            # 主送机关顶格，不缩进
    if DATE_CN.fullmatch(t):
        return False            # 成文日期
    if len(t) <= 14 and not re.search(r"[。；：，！？]$", t):
        return False            # 疑似机关署名等短行
    return True


# ---------------- 各检查项 ----------------

def check_page(doc, issues):
    sec = doc.sections[0]
    got = {"top": cm(sec.top_margin), "bottom": cm(sec.bottom_margin),
           "left": cm(sec.left_margin), "right": cm(sec.right_margin)}
    label = {"top": "上边距 37mm", "bottom": "下边距 35mm",
             "left": "左边距 28mm", "right": "右边距 26mm"}
    for k, v in got.items():
        if v is None or abs(v - SPEC[k]) > TOL:
            issues.append(Issue("高", f"页边距·{label[k]}", f"{v}cm",
                                f"{SPEC[k]}cm",
                                "在页面设置中改回国标值，否则版心尺寸不符合 156mm×225mm"))
    pw, ph = cm(sec.page_width), cm(sec.page_height)
    if pw is None or abs(pw - 21.0) > TOL or ph is None or abs(ph - 29.7) > TOL:
        issues.append(Issue("高", "纸张规格", f"{pw}×{ph}cm", "21.0×29.7cm（A4）",
                            "公文用纸采用 GB/T 148 规定的 A4 型"))
    grid = sec._sectPr.find(qn("w:docGrid"))
    if grid is None or grid.get(qn("w:type")) not in ("linesAndChars", "snapToChars"):
        issues.append(Issue("中", "文档网格", "未设置指定行网格",
                            "每页 22 行、每行 28 字",
                            "在页面设置-文档网格中指定行和字符网格"))


def classify(paras):
    """定位公文标题：第一个字号 ≥20pt 且非红色的段落。

    红色大字为发文机关标志（版头要素），不是标题，必须先排除，
    否则会把红头字号误判为标题字号。
    """
    for i, d in enumerate(paras):
        if not d["text"] or d["red"]:
            continue
        if d["size"] and d["size"] >= 20:
            return i
    return None


def check_title(paras, issues):
    idx = classify(paras)
    if idx is None:
        issues.append(Issue("高", "标题", "未发现 2 号字标题", "2 号小标宋体居中",
                            "标题应使用 2 号（22pt）小标宋体，居中编排"))
        return None
    d = paras[idx]
    if d["size"] and abs(d["size"] - 22) > 0.6:
        issues.append(Issue("高", "标题字号", f"{d['size']}pt", "22pt（2 号）",
                            "标题用 2 号小标宋体", d["text"][:20]))
    if d["font"] and d["font"] not in TITLE_OK:
        issues.append(Issue("中", "标题字体", d["font"], "方正小标宋简体",
                            "缺字体时可暂用宋体，正式行文应装公文字体", d["text"][:20]))
    if "CENTER" not in d["align"]:
        issues.append(Issue("高", "标题对齐", d["align"], "居中",
                            "标题居中编排，回行时词意完整、排列对称", d["text"][:20]))
    return idx


def check_body(paras, start, issues):
    body_cnt = 0
    bad_indent, bad_size, bad_font, bad_line = [], [], [], []
    for d in paras[start:]:
        t = d["text"]
        if not t or d["size"] is None or d["red"]:
            continue
        if d["size"] >= 20:
            continue
        if abs(d["size"] - 14) < 0.6:                 # 4 号：页码 / 版记，另行检查
            continue
        if not is_body_like(d):
            continue
        body_cnt += 1
        lv = 1 if L1.match(t) else 2 if L2.match(t) else 3 if L3.match(t) else \
             4 if L4.match(t) else 0
        if abs(d["size"] - 16) > 0.6:
            bad_size.append((t[:18], d["size"]))
        # 无层次序号但形如小标题的行（不以句读收尾、长度较短），
        # 实务中常用黑体/楷体提行，不按普通正文的仿宋要求判错
        unnumbered_head = (lv == 0 and len(t) <= 30
                           and not re.search(r"[。；：，！？、]$", t))
        if lv == 0 and not unnumbered_head and d["font"] \
                and d["font"] not in FANGSONG_OK:
            bad_font.append((t[:18], d["font"], "仿宋_GB2312"))
        if lv == 1 and d["font"] and d["font"] not in HEITI_OK:
            bad_font.append((t[:18], d["font"], "黑体"))
        if lv == 2 and d["font"] and d["font"] not in KAITI_OK:
            bad_font.append((t[:18], d["font"], "楷体_GB2312"))
        fi = d["first_indent"]
        # 整段左缩进的行属于附件说明续行、抄送回行等对齐排法，
        # 国标另有编排规则，不适用「首行缩进 2 字」
        if d["left_indent"] < 0.5 and (fi is None or abs(fi - 1.13) > 0.15):
            bad_indent.append((t[:18], fi))
        lval = d["line_val"]
        if lval and "EXACTLY" in d["line_rule"] and abs(float(lval) - 28) > 2.0:
            bad_line.append((t[:18], lval))

    if bad_line:
        s = "；".join(f"「{t}…」{v}磅" for t, v in bad_line[:3])
        issues.append(Issue("中", f"正文行距（{len(bad_line)} 处）", s, "固定值 28 磅",
                            "行距 28 磅可使每页约 22 行，与版心 225mm 匹配"))

    if body_cnt == 0:
        issues.append(Issue("高", "正文", "未识别到正文段落", "3 号仿宋正文",
                            "确认文件是否为公文正文，或正文字号异常"))
        return
    if bad_size:
        s = "；".join(f"「{t}…」{v}pt" for t, v in bad_size[:5])
        issues.append(Issue("高", f"正文字号（{len(bad_size)} 处）", s, "16pt（3 号）",
                            "正文及各级标题一律 3 号字"))
    if bad_font:
        s = "；".join(f"「{t}…」{f}→应为{e}" for t, f, e in bad_font[:5])
        issues.append(Issue("中", f"层级字体（{len(bad_font)} 处）", s,
                            "一级黑体·二级楷体·三四级仿宋",
                            "按层次序号切换字体，便于阅读层级"))
    if bad_indent:
        s = "；".join(f"「{t}…」缩进{v}cm" for t, v in bad_indent[:5])
        issues.append(Issue("中", f"首行缩进（{len(bad_indent)} 处）", s, "1.13cm（2 字）",
                            "正文每自然段左空二字"))


def check_elements(paras, issues, full_text):
    # 发文字号
    m = DOCNUM.search(full_text)
    if m:
        raw = m.group(0)
        if "〔" not in raw:
            issues.append(Issue("高", "发文字号括号", raw, "〔〕六角括号",
                                "年份应用六角括号〔〕，不得用方括号或圆括号"))
        if re.search(r"〕\s*0\d", raw):
            issues.append(Issue("高", "发文字号虚位", raw, "序号不编虚位",
                                "如〔2026〕12号，不写〔2026〕012号"))
    # 成文日期
    dates = DATE_CN.findall(full_text)
    if DATE_CN_ZH.search(full_text):
        issues.append(Issue("高", "成文日期数字", "使用了汉字数字年份",
                            "阿拉伯数字", "成文日期用阿拉伯数字将年、月、日标全"))
    for y, mo, d in dates:
        if (mo.startswith("0") and len(mo) == 2) or (d.startswith("0") and len(d) == 2):
            issues.append(Issue("中", "日期虚位", f"{y}年{mo}月{d}日",
                                f"{y}年{int(mo)}月{int(d)}日", "月、日不编虚位"))
            break
    # 主送机关冒号
    for d in paras:
        t = d["text"]
        if t and re.search(r"(各|全体|省|市|县|区|委|厅|局|办)", t) and len(t) < 80 \
                and (d["first_indent"] in (None, 0.0)) and t.endswith(("：", ":")):
            if t.endswith(":"):
                issues.append(Issue("中", "主送机关冒号", "半角冒号", "全角冒号「：」",
                                    "主送机关名称后标全角冒号", t[:20]))
            break
    # 附件说明标点
    for d in paras:
        t = d["text"]
        if t.startswith("附件") and re.search(r"[。；;，,]$", t):
            issues.append(Issue("中", "附件说明标点", t[:24], "名称后不加标点",
                                "附件名称后不加标点符号", t[:20]))
            break
    # 版记
    if "抄送" in full_text or "印发" in full_text:
        for d in paras:
            if d["text"].startswith(("抄送", "印发")) or "印发" in d["text"][-6:]:
                if d["size"] and abs(d["size"] - 14) > 0.6:
                    issues.append(Issue("中", "版记字号", f"{d['size']}pt", "14pt（4 号）",
                                        "抄送机关、印发机关和印发日期用 4 号仿宋",
                                        d["text"][:20]))
                break


def check_footer(doc, issues):
    sec = doc.sections[0]
    txt = ""
    try:
        for p in sec.footer.paragraphs:
            txt += p.text
            for r in p.runs:
                xml = r._element.xml
                if "PAGE" in xml:
                    txt += "«PAGE»"
    except Exception:
        pass
    if "«PAGE»" not in txt:
        issues.append(Issue("中", "页码", "页脚未发现页码域", "—N— 形式页码",
                            "页码用 4 号半角宋体阿拉伯数字，两侧各加一条一字线"))
        return
    if "—" not in txt and "-" in txt:
        issues.append(Issue("中", "页码一字线", "使用了半角连字符 -", "一字线 —",
                            "数字左右各放一条一字线（—），不是短横线"))
    fd = cm(sec.footer_distance)
    if fd is None or abs(fd - 2.8) > 0.3:
        issues.append(Issue("低", "页码位置", f"页脚距边界 {fd}cm", "约 2.8cm",
                            "使页码落在版心下边缘之下 7mm"))


def check_dates_logic(paras, issues, full_text):
    """可机械验证的内容逻辑校验（低级别，不替代人工校核）。"""
    # 1. 无效日期：月份 / 日越界
    for y, mo, d in DATE_CN.findall(full_text):
        yi, moi, di = int(y), int(mo), int(d)
        if moi < 1 or moi > 12:
            issues.append(Issue("低", "日期合法性", f"{y}年{mo}月{d}日", "月份应在 1–12",
                                "请核对月份是否录入错误", f"{y}年{mo}月{d}日"))
            continue
        if di < 1 or di > 31:
            issues.append(Issue("低", "日期合法性", f"{y}年{mo}月{d}日", "日应在 1–31",
                                "请核对日期是否录入错误", f"{y}年{mo}月{d}日"))
            continue
        if moi == 2 and di > 29:
            issues.append(Issue("低", "日期合法性", f"{y}年{mo}月{d}日", "2 月最多 29 日",
                                "请核对日期是否录入错误", f"{y}年{mo}月{d}日"))
            continue
        if moi in (4, 6, 9, 11) and di > 30:
            issues.append(Issue("低", "日期合法性", f"{y}年{mo}月{d}日", "该月最多 30 日",
                                "请核对日期是否录入错误", f"{y}年{mo}月{d}日"))

    # 2. 成文日期年份 早于 发文字号年份（如文号〔2026〕但写成 2024 年）
    m = DOCNUM.search(full_text)
    if m:
        doc_year = int(m.group(1))
        for y, mo, d in DATE_CN.findall(full_text):
            yi = int(y)
            if yi < doc_year:
                issues.append(Issue(
                    "低", "日期逻辑", f"成文日期 {y}年 早于发文字号年份 {doc_year}年",
                    "成文日期不应早于发文年份",
                    "公文成文日期一般晚于或等于发文字号年份，请核对", f"{y}年{mo}月{d}日"))
            break


def check_structure(paras, issues, full_text):
    seq = []
    for d in paras:
        t = d["text"]
        if L1.match(t):
            seq.append(("1", t[:12]))
        elif L2.match(t):
            seq.append(("2", t[:12]))
        elif L3.match(t):
            seq.append(("3", t[:12]))
        elif L4.match(t):
            seq.append(("4", t[:12]))
    # 层级跳跃：出现「（一）」但没有「一、」
    lv_set = {s for s, _ in seq}
    if "2" in lv_set and "1" not in lv_set:
        issues.append(Issue("低", "层次序号", "出现（一）但缺一级「一、」",
                            "一、→（一）→1.→（1）",
                            "确认是否有意跳级，公文层次序号一般逐级使用"))
    # 一级序号连续性
    nums = "一二三四五六七八九十"
    ones = [t for s, t in seq if s == "1"]
    expect = 0
    for t in ones:
        ch = t[0]
        if ch in nums:
            i = nums.index(ch)
            if i != expect:
                issues.append(Issue("中", "一级序号连续性", f"「{t}」",
                                    f"应为「{nums[expect]}、」",
                                    "一级层次序号应连续，检查是否漏项或重号"))
                break
            expect += 1
    if len(full_text) > 0 and "请示" in full_text[:120]:
        if full_text.count("请示") and "妥否" not in full_text and "当否" not in full_text:
            issues.append(Issue("低", "请示结语", "未发现「妥否，请批示」类结语",
                                "请示应有请示语", "请示一般以「妥否，请批示。」收束"))


def build_report(path, issues, meta):
    hi = [i for i in issues if i.level == "高"]
    mid = [i for i in issues if i.level == "中"]
    lo = [i for i in issues if i.level == "低"]
    verdict = "不通过（存在明显国标偏离）" if hi else \
              ("基本合规（有待优化项）" if mid else "合规")
    L = []
    L.append("# 公文格式合规检查报告")
    L.append("")
    L.append(f"- 受检文件：`{os.path.basename(path)}`")
    L.append(f"- 检查依据：《党政机关公文格式》GB/T 9704-2012")
    L.append(f"- 检查时间：{date.today().isoformat()}")
    L.append(f"- 段落总数：{meta.get('paras', 0)}｜节数：{meta.get('sections', 0)}")
    L.append(f"- **结论：{verdict}**（高 {len(hi)} · 中 {len(mid)} · 低 {len(lo)}）")
    L.append("")
    if not issues:
        L.append("未发现可机器判定的格式偏离项。建议再人工核对印章、密级与内容表述。")
        return "\n".join(L)
    L.append("## 一、问题清单")
    L.append("")
    L.append("| 等级 | 检查项 | 实际 | 国标要求 | 处理建议 |")
    L.append("|---|---|---|---|---|")
    for i in hi + mid + lo:
        f = str(i.found).replace("|", "/")[:60]
        e = str(i.expect).replace("|", "/")[:40]
        a = str(i.advice).replace("|", "/")[:60]
        L.append(f"| {i.level} | {i.item} | {f} | {e} | {a} |")
    L.append("")
    L.append("## 二、需人工确认事项")
    L.append("")
    L.append("以下事项无法由脚本判定，请人工核对：")
    L.append("")
    L.append("- 印章是否加盖、是否压成文日期（骑年盖月）")
    L.append("- 密级与保密期限标注是否与定密结论一致")
    L.append("- 文种选用是否恰当，行文方向与主送机关是否匹配")
    L.append("- 发文机关标志用字是否为规范化简称或全称")
    L.append("- 内容表述的政策依据、数据、人名地名是否准确")
    L.append("")
    L.append("> 本报告仅覆盖可机械判定的格式项，不替代人工校核与合法性审查。")
    return "\n".join(L)


def main():
    ap = argparse.ArgumentParser(description="公文格式合规质检 (GB/T 9704-2012)")
    ap.add_argument("--input", required=True, help="待检 .docx 文件")
    ap.add_argument("--report", default="", help="报告输出路径（.md）")
    ap.add_argument("--json", action="store_true", help="输出 JSON 结果")
    a = ap.parse_args()

    if not os.path.exists(a.input):
        sys.stderr.write(f"[err] 文件不存在：{a.input}\n")
        sys.exit(1)

    try:
        doc = Document(a.input)
    except PermissionError:
        sys.stderr.write(f"[err] 无法读取 {a.input}：文件可能被占用（如已在 Word 中打开）或无读取权限。\n"
                         f"      请关闭该文件后重试。\n")
        sys.exit(3)
    except Exception as e:
        sys.stderr.write(f"[err] 读取 {a.input} 失败：{e}\n")
        sys.exit(3)

    paras = [para_info(p) for p in doc.paragraphs]
    full_text = "\n".join(d["text"] for d in paras)
    issues = []

    check_page(doc, issues)
    idx = check_title(paras, issues)
    check_body(paras, (idx + 1) if idx is not None else 0, issues)
    check_elements(paras, issues, full_text)
    check_footer(doc, issues)
    check_structure(paras, issues, full_text)
    check_dates_logic(paras, issues, full_text)

    meta = {"paras": len([d for d in paras if d["text"]]),
            "sections": len(doc.sections)}

    if a.json:
        print(json.dumps({"file": a.input, "meta": meta,
                          "issues": [i.as_dict() for i in issues]},
                         ensure_ascii=False, indent=2))
        return

    report = build_report(a.input, issues, meta)
    if a.report:
        _safe_write_text(a.report, report, "检查报告")
        print(f"[ok] 检查报告已生成：{a.report}")
    print(report if not a.report else
          f"[i ] 高 {sum(1 for i in issues if i.level=='高')} · "
          f"中 {sum(1 for i in issues if i.level=='中')} · "
          f"低 {sum(1 for i in issues if i.level=='低')}")


if __name__ == "__main__":
    main()
