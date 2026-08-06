#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
gongwen_format.py — 党政机关公文标准排版引擎 (GB/T 9704-2012)

将 Markdown / TXT / DOCX / JSON 转换为符合国标的 .docx 公文。

依赖: python-docx  (pip install python-docx)

两种驱动方式：
  1) 命令行参数（快速）
     python gongwen_format.py --title "标题" --input body.md --output out.docx \
            --redhead "××市人民政府" --doc-number "×政发〔2026〕12号"
  2) JSON 配置（要素多时更清晰）
     python gongwen_format.py --config gongwen.json --output out.docx

自检：
     python gongwen_format.py --demo        # 生成示例公文与示例配置
"""

import argparse
import json
import os
import re
import shutil
import sys
import unicodedata
from datetime import date, datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_SECTION_START
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    sys.stderr.write(
        "缺少依赖 python-docx。请先安装：\n"
        "    pip install python-docx\n"
    )
    sys.exit(2)


# ===================== 一、国标常量 =====================
# GB/T 9704-2012 §6 版面 / §7 要素

# 版面（A4，页边距 上37 / 下35 / 左28 / 右26 mm，版心 156mm × 225mm）
PAGE_W, PAGE_H = Cm(21.0), Cm(29.7)
MARGIN_TOP, MARGIN_BOTTOM = Cm(3.7), Cm(3.5)
MARGIN_LEFT, MARGIN_RIGHT = Cm(2.8), Cm(2.6)
FOOTER_DISTANCE = Cm(2.8)          # 页码距纸张下边缘，使其落在版心下边缘之下 7mm

# 字号（号 → 磅）
SIZE_XIAOCHU = Pt(36)   # 小初（长机关名压缩用）
SIZE_1 = Pt(26)         # 一号
SIZE_2 = Pt(22)         # 二号：标题、发文机关标志
SIZE_3 = Pt(16)         # 三号：正文及各级标题、版头要素
SIZE_4 = Pt(14)         # 四号：页码、版记

# 行距与缩进（3 号字：1 字 = 16pt = 0.564cm）
CHAR_W_CM = 0.564
LINE_EXACT = Pt(28)     # 正文固定行距 28 磅（每页约 22 行）
LINE_TITLE = Pt(32)     # 标题固定行距
INDENT_2 = Cm(CHAR_W_CM * 2)   # 首行缩进 2 字 ≈ 1.13cm

RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

# 发文机关标志上边缘至版心上边缘：一般 35mm；上报的公文 80mm
REDHEAD_OFFSET_MM = 35.0
REDHEAD_OFFSET_UPWARD_MM = 80.0

SECRET_LEVELS = ("绝密", "机密", "秘密")
URGENCY_LEVELS = ("特急", "加急", "平急")


# ===================== 二、字体解析（带回退链） =====================
# 国标规定的字体在部分系统缺失，按候选链探测，避免 Word 回退到错误字体。

FONT_CANDIDATES = {
    "xiaobiaosong": ["方正小标宋简体", "FZXiaoBiaoSong-B05S", "小标宋体", "宋体", "SimSun"],
    "fangsong":     ["仿宋_GB2312", "仿宋", "FangSong_GB2312", "FangSong", "SimSun"],
    "heiti":        ["黑体", "SimHei"],
    "kaiti":        ["楷体_GB2312", "楷体", "KaiTi", "SimKai"],
    "songti":       ["宋体", "SimSun"],
}

# 字体名 → 可能的字库文件名（用于探测系统是否安装）
FONT_FILES = {
    "方正小标宋简体": ["FZXBSJW.TTF", "FZXBSJW.ttf", "方正小标宋简体.ttf"],
    "仿宋_GB2312": ["FS_GB2312.ttf", "仿宋_GB2312.ttf", "FZFSJW.TTF"],
    "仿宋": ["simfang.ttf"],
    "黑体": ["simhei.ttf"],
    "楷体_GB2312": ["KAI_GB2312.ttf", "楷体_GB2312.ttf"],
    "楷体": ["simkai.ttf"],
    "宋体": ["simsun.ttc", "simsun.ttf"],
}

FONT_LATIN = "Times New Roman"      # 正文中西文、阿拉伯数字统一用 Times New Roman
_FONT_CACHE = {}


def _font_dirs():
    dirs = []
    if os.name == "nt":
        dirs.append(os.path.join(os.environ.get("WINDIR", r"C:\Windows"), "Fonts"))
        local = os.environ.get("LOCALAPPDATA")
        if local:
            dirs.append(os.path.join(local, "Microsoft", "Windows", "Fonts"))
    else:
        dirs += ["/usr/share/fonts", "/usr/local/share/fonts",
                 os.path.expanduser("~/.fonts"), "/System/Library/Fonts",
                 "/Library/Fonts"]
    return [d for d in dirs if os.path.isdir(d)]


def _installed_files():
    if "_files" in _FONT_CACHE:
        return _FONT_CACHE["_files"]
    names = set()
    for d in _font_dirs():
        try:
            for root, _dirs, files in os.walk(d):
                for f in files:
                    names.add(f.lower())
        except OSError:
            continue
    _FONT_CACHE["_files"] = names
    return names


def resolve_font(key, strict=False):
    """返回该字体角色在当前系统上最合适的字体名。

    strict=True 时不做探测，直接返回国标首选名（交由 Word 自行回退），
    适合最终要在装齐公文字体的机关电脑上打印的场景。
    """
    cands = FONT_CANDIDATES[key]
    if strict:
        return cands[0]
    cache_key = f"r:{key}"
    if cache_key in _FONT_CACHE:
        return _FONT_CACHE[cache_key]
    files = _installed_files()
    chosen = cands[-1]
    for name in cands:
        pats = FONT_FILES.get(name, [])
        if any(p.lower() in files for p in pats):
            chosen = name
            break
    else:
        chosen = cands[0]        # 探测不到就用国标首选名，让 Word 决定
    _FONT_CACHE[cache_key] = chosen
    return chosen


class Fonts:
    """本次运行使用的字体集合。"""

    def __init__(self, strict=False):
        self.title = resolve_font("xiaobiaosong", strict)
        self.body = resolve_font("fangsong", strict)
        self.h1 = resolve_font("heiti", strict)
        self.h2 = resolve_font("kaiti", strict)
        self.song = resolve_font("songti", strict)

    def report(self):
        return (f"字体映射：标题/机关标志={self.title}｜正文={self.body}｜"
                f"一级标题={self.h1}｜二级标题={self.h2}｜页码/版记={self.song}")


# ===================== 三、底层排版工具 =====================

def set_run_font(run, cn_font, size, bold=False, color=None, latin=None):
    """设置 run 的中文字体（eastAsia）与西文字体，避免数字被中文字体拉宽。"""
    run.font.size = size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    latin_font = latin or FONT_LATIN
    run.font.name = latin_font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), latin_font)
    rFonts.set(qn("w:hAnsi"), latin_font)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:hint"), "eastAsia")


def new_para(doc, align=None, line=LINE_EXACT, first_indent=None,
             left_indent=None, right_indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if line is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if right_indent is not None:
        pf.right_indent = right_indent
    return p


def blank_lines(doc, n=1, line=LINE_EXACT):
    for _ in range(n):
        new_para(doc, line=line)


def set_border(par, edge="bottom", color="000000", sz=8, val="single"):
    """给段落加边框线（用于红色分隔线、版记反线）。sz 单位为 1/8 磅。"""
    pPr = par._element.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}/>')
        pPr.append(pBdr)
    el = parse_xml(
        f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" '
        f'w:space="0" w:color="{color}"/>'
    )
    pBdr.append(el)


def set_page(section, doc_grid=True):
    section.page_width, section.page_height = PAGE_W, PAGE_H
    section.top_margin, section.bottom_margin = MARGIN_TOP, MARGIN_BOTTOM
    section.left_margin, section.right_margin = MARGIN_LEFT, MARGIN_RIGHT
    section.footer_distance = FOOTER_DISTANCE
    if doc_grid:
        # 文档网格：每页 22 行、每行 28 字（GB/T 9704 §6.3）
        sectPr = section._sectPr
        for old in list(sectPr.findall(qn("w:docGrid"))):
            sectPr.remove(old)
        sectPr.append(parse_xml(
            f'<w:docGrid {nsdecls("w")} w:type="linesAndChars" '
            f'w:linePitch="580" w:charSpace="0"/>'
        ))


def disp_width(text):
    """按公文排版口径计算显示宽度（全角=1，半角=0.5）。"""
    w = 0.0
    for ch in text:
        w += 1.0 if unicodedata.east_asian_width(ch) in ("F", "W", "A") else 0.5
    return w


def chars(n):
    return Cm(CHAR_W_CM * n)


# ===================== 四、要素编排 =====================

def add_header_marks(doc, F, copies="", secret="", urgency=""):
    """版头三要素：份号、密级和保密期限、紧急程度（版心左上角，3 号黑体，顶格）。

    份号为阿拉伯数字，国标要求 6 位、顶格编排。
    返回已占用的行数，供发文机关标志定位使用。
    """
    used = 0
    if copies:
        copies = str(copies).strip()
        if copies.isdigit():
            copies = copies.zfill(6)
        p = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_run_font(p.add_run(copies), F.h1, SIZE_3)
        used += 1
    if secret:
        p = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_run_font(p.add_run(secret.strip()), F.h1, SIZE_3)
        used += 1
    if urgency:
        p = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_run_font(p.add_run(urgency.strip()), F.h1, SIZE_3)
        used += 1
    return used


def _redhead_font_size(name):
    """发文机关标志字号自适应：不大于上级机关文件字号，且不超版心宽度。"""
    n = disp_width(name)
    if n <= 6:
        return Pt(30)
    if n <= 8:
        return Pt(26)
    if n <= 12:
        return SIZE_2          # 22pt
    if n <= 16:
        return Pt(18)
    return Pt(16)


def add_redhead(doc, F, org_names, doc_number="", signer="", used_lines=0,
                is_upward=False, minutes=False):
    """发文机关标志 + 发文字号（+签发人）+ 红色分隔线。

    org_names: 字符串或列表（联合行文时多个机关，主办机关排首位）。
    """
    if isinstance(org_names, str):
        org_names = [org_names] if org_names else []
    if not org_names:
        return

    # 定位：标志上边缘至版心上边缘 35mm；上报的公文为 80mm
    target_mm = REDHEAD_OFFSET_UPWARD_MM if is_upward else REDHEAD_OFFSET_MM
    target_pt = target_mm * 72.0 / 25.4
    used_pt = used_lines * 28.0
    lead_pt = max(0.0, target_pt - used_pt)

    for i, name in enumerate(org_names):
        p = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, line=Pt(48))
        if i == 0 and lead_pt > 0:
            p.paragraph_format.space_before = Pt(lead_pt)
        run = p.add_run(name)
        size = _redhead_font_size(name)
        # 会议纪要标志同为红色小标宋体，但不加分隔线
        set_run_font(run, F.title, size, bold=False, color=RED)

    if minutes:
        blank_lines(doc, 1)
        return

    # 发文字号：一般居中；上报的公文左空一字，签发人右空一字同行
    if doc_number or signer:
        blank_lines(doc, 1)
        if signer:
            p = new_para(doc, left_indent=chars(1), right_indent=chars(1))
            if doc_number:
                set_run_font(p.add_run(doc_number), F.body, SIZE_3)
            tab = p.paragraph_format.tab_stops
            from docx.enum.text import WD_TAB_ALIGNMENT
            usable_cm = 21.0 - 2.8 - 2.6 - CHAR_W_CM * 2
            tab.add_tab_stop(Cm(usable_cm), WD_TAB_ALIGNMENT.RIGHT)
            set_run_font(p.add_run("\t"), F.body, SIZE_3)
            set_run_font(p.add_run("签发人："), F.body, SIZE_3)
            for idx, s in enumerate([x.strip() for x in re.split(r"[、,，;；]", signer) if x.strip()]):
                if idx:
                    set_run_font(p.add_run("  "), F.body, SIZE_3)
                set_run_font(p.add_run(s), F.h2, SIZE_3)
        else:
            p = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
            if doc_number:
                set_run_font(p.add_run(doc_number), F.body, SIZE_3)

    # 红色分隔线：与版心等宽，武文线在发文字号下 4mm
    line_p = new_para(doc, line=Pt(6))
    set_border(line_p, "bottom", color="FF0000", sz=16)


def add_title(doc, F, title, minutes=False):
    """标题：红色分隔线下空二行，2 号小标宋体，居中，回行时词意完整。"""
    blank_lines(doc, 2 if not minutes else 1, line=Pt(24))
    for seg in [s for s in title.split("\n") if s.strip()]:
        p = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, line=LINE_TITLE)
        set_run_font(p.add_run(seg.strip()), F.title, SIZE_2, bold=False)


def add_recipient(doc, F, recipient):
    """主送机关：标题下空一行，3 号仿宋，居左顶格，末尾全角冒号。"""
    if not recipient:
        return
    blank_lines(doc, 1)
    text = recipient.strip()
    if not text.endswith("："):
        text = text.rstrip(":：") + "："
    p = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.first_line_indent = Cm(0)
    set_run_font(p.add_run(text), F.body, SIZE_3)


# --- 正文层级识别 ---
LEVEL_PATTERNS = [
    (1, [re.compile(r"^[一二三四五六七八九十百]+、"),
         re.compile(r"^第[一二三四五六七八九十百]+[部分章节篇]")]),
    (2, [re.compile(r"^[（(][一二三四五六七八九十百]+[）)]")]),
    (3, [re.compile(r"^\d+[\.．、](?!\d)")]),
    (4, [re.compile(r"^[（(]\d+[）)]")]),
]
ATTACH_RE = re.compile(r"^附\s*件\s*[:：]\s*(.*)$")
MD_STRIP = re.compile(r"\*\*(.+?)\*\*")


def detect_level(raw):
    """识别段落层级。返回 (level, text)；level=0 为普通正文，None 为空行。"""
    s = raw.strip()
    if not s:
        return None, None
    # Markdown 标题。若标题文字自带中文层次序号（一、／（一）／1.／（1）），
    # 以序号为准——序号是公文的正式层级标记，# 的深度只是作者的排版习惯，
    # 二者冲突时（如「## 一、提高思想认识」）按 # 深度会把一级标题错排成楷体。
    m = re.match(r"^(#{1,6})\s+(.*)$", s)
    if m:
        txt = MD_STRIP.sub(r"\1", m.group(2)).strip()
        for lv, pats in LEVEL_PATTERNS:
            if any(p.match(txt) for p in pats):
                return lv, txt
        return min(len(m.group(1)), 4), txt
    if s.startswith(("- ", "* ", "+ ")):
        s = s[2:].strip()
    s = MD_STRIP.sub(r"\1", s).strip()
    if not s:
        return None, None
    for lv, pats in LEVEL_PATTERNS:
        if any(p.match(s) for p in pats):
            return lv, s
    return 0, s


def split_inline_heading(text, level):
    """二级标题常写成「（一）压实责任。具体内容……」，拆成标题行 + 正文段。"""
    if level not in (2, 3, 4):
        return [(text, level)]
    pos = text.find("。")
    if pos == -1 or pos >= len(text) - 1:
        return [(text, level)]
    head, body = text[: pos + 1], text[pos + 1:].strip()
    if not body or disp_width(head) > 30:
        return [(text, level)]
    return [(head, level), (body, 0)]


def add_body_para(doc, F, text, level=0):
    """正文：3 号仿宋，首行缩进 2 字，固定行距 28 磅。
    层级字体：一级黑体、二级楷体、三/四级与正文同为仿宋。"""
    font = {0: F.body, 1: F.h1, 2: F.h2, 3: F.body, 4: F.body}.get(level, F.body)
    p = new_para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=INDENT_2)
    set_run_font(p.add_run(text), font, SIZE_3)
    return p


def add_attachment_note(doc, F, items):
    """附件说明：正文下空一行，左空二字编排；名称后不加标点。"""
    if not items:
        return
    blank_lines(doc, 1)
    for i, name in enumerate(items):
        name = re.sub(r"[。；;，,\.！!？?]+$", "", str(name).strip())
        if not name:
            continue
        p = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=INDENT_2)
        if i == 0:
            prefix = "附件："
            body = name if len(items) == 1 and not re.match(r"^\d+[\.．]", name) else name
            set_run_font(p.add_run(prefix + body), F.body, SIZE_3)
        else:
            # 续行与首行「附件：」之后的文字左对齐
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = chars(5)
            set_run_font(p.add_run(name), F.body, SIZE_3)


def add_signature(doc, F, org_sign, date_text, sealed=False, seal_image=None):
    """发文机关署名、成文日期与印章（GB/T 9704 §7.3.5）。

    加盖印章时：成文日期右空四字，一般不单独署机关名称（以印章代）。
    不加盖印章时：署名右空二字；成文日期首字比署名首字右移二字；
                  若日期长于署名，则日期右空二字，相应增加署名右空字数。
    """
    if not (org_sign or date_text):
        return
    orgs = [org_sign] if isinstance(org_sign, str) else list(org_sign or [])
    orgs = [o for o in orgs if o and o.strip()]
    blank_lines(doc, 1)

    if sealed:
        for o in orgs:
            p = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, right_indent=chars(4))
            set_run_font(p.add_run(o), F.body, SIZE_3)
        if date_text:
            p = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, right_indent=chars(4))
            set_run_font(p.add_run(date_text), F.body, SIZE_3)
        _seal_placeholder(doc, F, seal_image)
        return

    name = orgs[0] if orgs else ""
    w_n, w_d = disp_width(name), disp_width(date_text or "")
    if name and date_text:
        r_n, r_d = 2.0, w_n - w_d
        if r_d < 2.0:                    # 日期长于署名，改为日期右空二字
            r_d = 2.0
            r_n = 4.0 + w_d - w_n
    else:
        r_n = r_d = 2.0

    for i, o in enumerate(orgs):
        p = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT,
                     right_indent=chars(r_n if i == 0 else r_n))
        set_run_font(p.add_run(o), F.body, SIZE_3)
    if date_text:
        p = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, right_indent=chars(r_d))
        set_run_font(p.add_run(date_text), F.body, SIZE_3)


def _seal_placeholder(doc, F, seal_image=None):
    """印章位：印章中心压成文日期（骑年盖月）。脚本仅留占位或插入图片。"""
    p = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, right_indent=chars(2))
    if seal_image and os.path.exists(seal_image):
        try:
            p.add_run().add_picture(seal_image, width=Cm(4.2))
            return
        except Exception:
            pass
    run = p.add_run("〔此处加盖机关印章，印章中心压成文日期〕")
    set_run_font(run, F.body, Pt(10.5), color=RED)


def add_notes(doc, F, notes):
    """附注：成文日期下一行，左空二字，加圆括号。"""
    if not notes:
        return
    if isinstance(notes, str):
        notes = [notes]
    for note in notes:
        note = str(note).strip()
        if not note:
            continue
        if not note.startswith(("（", "(")):
            note = f"（{note}）"
        p = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=INDENT_2)
        set_run_font(p.add_run(note), F.body, SIZE_3)


def add_attachment_pages(doc, F, attach_docs):
    """附件正文另面编排：「附件1」顶格于版心第一行，标题下空一行接正文。"""
    for idx, item in enumerate(attach_docs, 1):
        path, sep, title = str(item).partition("::")
        path = path.strip()
        if not os.path.exists(path):
            sys.stderr.write(f"[warn] 附件文件不存在，已跳过：{path}\n")
            continue
        sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
        set_page(sec)
        p = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.first_line_indent = Cm(0)
        set_run_font(p.add_run(f"附件{idx}"), F.h1, SIZE_3)
        if title.strip():
            blank_lines(doc, 1)
            tp = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, line=LINE_TITLE)
            set_run_font(tp.add_run(title.strip()), F.title, SIZE_2)
        blank_lines(doc, 1)
        content, _ = read_input(path)
        render_body(doc, F, content)


def add_banji(doc, F, cc="", print_org="", print_date="", force_page=True):
    """版记：抄送机关、印发机关和印发日期，4 号仿宋，排在公文最后一面。

    版记上下各一条分隔线（首末为粗线），版记最后一行为公文最后一行。
    """
    if not (cc or print_org or print_date):
        return
    if force_page:
        sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
        set_page(sec)
        sec._sectPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="bottom"/>'))

    def _line(sz=12):
        p = new_para(doc, line=Pt(6))
        set_border(p, "bottom", color="000000", sz=sz)

    _line(12)                                     # 版记首条：粗线
    if cc:
        text = cc.strip()
        if not text.endswith("。"):
            text += "。"
        p = new_para(doc, left_indent=chars(1), right_indent=chars(1), line=Pt(22))
        p.paragraph_format.first_line_indent = Cm(0)
        set_run_font(p.add_run(f"抄送：{text}"), F.body, SIZE_4)
        _line(4)                                  # 抄送与印发之间：细线
    if print_org or print_date:
        p = new_para(doc, left_indent=chars(1), right_indent=chars(1), line=Pt(22))
        p.paragraph_format.first_line_indent = Cm(0)
        from docx.enum.text import WD_TAB_ALIGNMENT
        usable_cm = 21.0 - 2.8 - 2.6 - CHAR_W_CM * 2
        p.paragraph_format.tab_stops.add_tab_stop(Cm(usable_cm), WD_TAB_ALIGNMENT.RIGHT)
        set_run_font(p.add_run(print_org or ""), F.body, SIZE_4)
        set_run_font(p.add_run("\t"), F.body, SIZE_4)
        tail = f"{print_date}印发" if print_date else "印发"
        set_run_font(p.add_run(tail), F.body, SIZE_4)
    _line(12)                                     # 版记末条：粗线


def add_page_numbers(doc, F, first_page_no_number=False):
    """页码：4 号半角宋体阿拉伯数字，数字左右各一条一字线；
    单页码居右空一字，双页码居左空一字。

    注意：国标未规定首页不编页码，默认全篇连续编排；
    确有单位惯例需首页不显示时，用 --no-first-page-num 打开。
    """
    settings = doc.settings.element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(parse_xml(f'<w:evenAndOddHeaders {nsdecls("w")}/>'))

    def _fill(footer, align, indent_side):
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = align
        pf = p.paragraph_format
        if indent_side == "right":
            pf.right_indent = chars(1)
        else:
            pf.left_indent = chars(1)
        for r in list(p.runs):
            p._element.remove(r._element)
        r1 = p.add_run("—")
        set_run_font(r1, F.song, SIZE_4)
        r2 = p.add_run()
        set_run_font(r2, F.song, SIZE_4)
        r2._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        r2._element.append(parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        r2._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
        r3 = p.add_run("—")
        set_run_font(r3, F.song, SIZE_4)

    sec = doc.sections[0]
    _fill(sec.footer, WD_ALIGN_PARAGRAPH.RIGHT, "right")
    _fill(sec.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT, "left")

    if first_page_no_number:
        sectPr = sec._sectPr
        if sectPr.find(qn("w:titlePg")) is None:
            sectPr.append(parse_xml(f'<w:titlePg {nsdecls("w")}/>'))

    for s in doc.sections[1:]:
        s.footer.is_linked_to_previous = True
        s.even_page_footer.is_linked_to_previous = True


# ===================== 五、输入解析 =====================

def normalize_text(content):
    """轻量文本规范：直引号转弯引号，半角括号中的中文标点保持原样。"""
    out, i, n = [], 0, len(content)
    while i < n:
        ch = content[i]
        if ch == '"':
            close = content.find('"', i + 1)
            if close != -1:
                out.append("“" + content[i + 1:close] + "”")
                i = close + 1
                continue
        out.append(ch)
        i += 1
    return "".join(out)


def _backup(path):
    """写目标文件前，先把已存在的原稿复制到同目录 .bak 备份，防覆盖丢失。返回备份路径或 None。"""
    if not path or not os.path.exists(path):
        return None
    d = os.path.dirname(os.path.abspath(path)) or "."
    base = os.path.basename(path)
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    bak = os.path.join(d, f".{base}.bak-{ts}")
    n = 1
    while os.path.exists(bak):
        bak = os.path.join(d, f".{base}.bak-{ts}-{n}")
        n += 1
    try:
        shutil.copy2(path, bak)
        return bak
    except OSError:
        return None


def read_pdf_text(path):
    """从 PDF 文本层抽取文字（不做 OCR，扫描件无文本层会提示）。优先 pypdf，回退 pdfminer。"""
    try:
        from pypdf import PdfReader
        try:
            r = PdfReader(path)
        except Exception as e:
            sys.stderr.write(f"[err] 读取 PDF 失败：{e}\n"
                             f"      若该文件为扫描件图片，本技能只提取文本层、不做 OCR，"
                             f"请先转为可编辑文本后再处理。\n")
            sys.exit(3)
        parts = []
        for pg in r.pages:
            try:
                parts.append(pg.extract_text() or "")
            except Exception:
                parts.append("")
        text = "\n".join(parts)
    except ImportError:
        try:
            from pdfminer.high_level import extract_text
            text = extract_text(path)
        except ImportError:
            sys.stderr.write("[err] 读取 PDF 需要文本抽取库：请执行 pip install pypdf\n")
            sys.exit(2)
    if not text.strip():
        sys.stderr.write(
            "[warn] 该 PDF 未包含可提取的文本层（疑似扫描件）。\n"
            "       本技能只做文本层提取，不做 OCR；请先转为可编辑文本后再处理。\n")
    return text


def _safe_save_docx(doc, path):
    """保存 .docx，写前自动备份；捕获 Word 占用/权限/磁盘错误并给可操作提示。"""
    bak = _backup(path)
    directory = os.path.dirname(os.path.abspath(path)) or "."
    os.makedirs(directory, exist_ok=True)
    try:
        doc.save(path)
    except PermissionError:
        sys.stderr.write(
            f"[err] 无法保存 {path}：文件可能正被 Word 等程序占用，或无写入权限。\n"
            f"      请关闭该文件后重试，或用 --output 另存为新文件名（如 新文件名.docx）。\n")
        sys.exit(3)
    except OSError as e:
        sys.stderr.write(f"[err] 保存 {path} 失败：{e}\n"
                         f"      请检查磁盘空间与目标目录权限。\n")
        sys.exit(3)
    if bak:
        sys.stderr.write(f"[i ] 原文件已备份：{bak}\n")


def read_input(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        try:
            d = Document(path)
        except PermissionError:
            sys.stderr.write(f"[err] 无法读取 {path}：文件可能被占用或无读取权限。\n")
            sys.exit(3)
        except Exception as e:
            sys.stderr.write(f"[err] 读取 {path} 失败：{e}\n")
            sys.exit(3)
        lines = [p.text.rstrip() for p in d.paragraphs]
        return normalize_text("\n".join(lines)), False
    if ext == ".pdf":
        return read_pdf_text(path), False
    try:
        with open(path, "r", encoding="utf-8") as f:
            return normalize_text(f.read()), (ext == ".md")
    except FileNotFoundError:
        sys.stderr.write(f"[err] 文件不存在：{path}\n")
        sys.exit(1)
    except UnicodeDecodeError:
        # 部分 Windows 文本以 GBK 编码保存，退一步再试
        try:
            with open(path, "r", encoding="gbk") as f:
                return normalize_text(f.read()), (ext == ".md")
        except Exception:
            sys.stderr.write(
                f"[err] 无法以 UTF-8/GBK 解码 {path}。\n"
                f"      请将其转为 UTF-8 编码后重试（记事本'另存为'可选 UTF-8）。\n")
            sys.exit(3)
    except PermissionError:
        sys.stderr.write(f"[err] 无法读取 {path}：文件可能被占用或无读取权限。\n")
        sys.exit(3)


def _same_title(a, b):
    """判断正文首行是否与公文标题重复。

    草稿常把标题写成省略发文机关的短式（「关于××的通知」），而 --title 传的是
    含机关名的全称，因此不能只做全等比较，需允许一方是另一方的尾部子串。
    """
    if not a or not b:
        return False
    norm = lambda x: re.sub(r"[\s　·—\-_、，,。.：:；;“”\"'（）()《》〈〉]", "", x)
    x, y = norm(a), norm(b)
    if not x or not y:
        return False
    if x == y:
        return True
    short, long_ = (x, y) if len(x) <= len(y) else (y, x)
    # 短式须足够长且构成长式的结尾，避免把正常正文首句误删
    return len(short) >= 8 and len(short) / len(long_) >= 0.5 and long_.endswith(short)


def render_body(doc, F, content, title=None):
    """逐行渲染正文，处理附件说明块与层级。

    title 用于剔除草稿正文里与公文标题重复的首行——用户常在 Markdown
    草稿开头写「# 标题」，同时又通过 --title 传入，不去重会出现两个标题。
    """
    lines = content.split("\n")
    i, pending_attach = 0, []
    title_dropped = title is None
    while i < len(lines):
        raw = lines[i]
        s = raw.strip()
        if not s:
            i += 1
            continue
        if not title_dropped:
            _lv, _t = detect_level(raw)
            if _t and _same_title(_t, title):
                title_dropped = True
                i += 1
                continue
            title_dropped = True     # 只检查正文第一个非空行
        m = ATTACH_RE.match(s)
        if m:
            items, i = _collect_attachments(lines, i, m.group(1))
            pending_attach.extend(items)
            continue
        lv, text = detect_level(raw)
        if lv is None or not text:
            i += 1
            continue
        for t, l in split_inline_heading(text, lv):
            add_body_para(doc, F, t, level=l)
        i += 1
    if pending_attach:
        add_attachment_note(doc, F, pending_attach)


def _collect_attachments(lines, idx, first):
    """收集附件说明块，支持「附件：1.×× 2.××」单行多项与多行续排。"""
    items = []
    first = first.strip()
    if first:
        parts = re.findall(r"\d+[\.．、][^\d]*(?:(?!\s\d+[\.．、]).)*", first)
        items.extend([p.strip() for p in parts] if len(parts) > 1 else [first])
    i = idx + 1
    while i < len(lines):
        s = lines[i].strip()
        if not s:
            i += 1
            if items:
                break
            continue
        if re.match(r"^\d+[\.．、]", s) or re.match(r"^附件\s*\d+\s*[:：]", s):
            items.append(re.sub(r"^附件\s*\d+\s*[:：]\s*", "", s))
            i += 1
        else:
            break
    return ([x for x in items if x] or ["（见附件）"]), i


def format_cn_date(s):
    """成文日期用阿拉伯数字，年月日标全，月、日不编虚位。"""
    s = str(s or "").strip()
    if not s:
        return ""
    if s in ("today", "今天", "auto"):
        t = date.today()
        return f"{t.year}年{t.month}月{t.day}日"
    m = re.match(r"^(\d{4})\s*[年\-/\.]\s*(\d{1,2})\s*[月\-/\.]\s*(\d{1,2})\s*日?$", s)
    if m:
        return f"{int(m.group(1))}年{int(m.group(2))}月{int(m.group(3))}日"
    return s


def normalize_doc_number(s):
    """发文字号：年份用六角括号〔〕，序号不编虚位。"""
    s = str(s or "").strip()
    if not s:
        return ""
    s = re.sub(r"[\[【(（]\s*(\d{4})\s*[\]】)）]", r"〔\1〕", s)
    s = re.sub(r"〔(\d{4})〕\s*0*(\d+)\s*号", r"〔\1〕\2号", s)
    return s


# ===================== 六、装配 =====================

def build(cfg, out_path):
    F = Fonts(strict=bool(cfg.get("strict_font")))
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.size = SIZE_3
    _rPr = normal.element.get_or_add_rPr()
    _rf = _rPr.find(qn("w:rFonts"))
    if _rf is None:
        _rf = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        _rPr.insert(0, _rf)
    _rf.set(qn("w:eastAsia"), F.body)
    _rf.set(qn("w:ascii"), FONT_LATIN)
    _rf.set(qn("w:hAnsi"), FONT_LATIN)

    set_page(doc.sections[0])

    minutes = bool(cfg.get("minutes"))
    signer = cfg.get("signer", "")
    is_upward = bool(signer)

    used = add_header_marks(doc, F, cfg.get("copies", ""),
                            cfg.get("secret_level", ""), cfg.get("urgency", ""))

    redhead = cfg.get("redhead") or cfg.get("org_mark") or ""
    if redhead:
        add_redhead(doc, F, redhead,
                    normalize_doc_number(cfg.get("doc_number", "")),
                    signer, used_lines=used, is_upward=is_upward, minutes=minutes)

    if cfg.get("title"):
        add_title(doc, F, cfg["title"], minutes=minutes)

    add_recipient(doc, F, cfg.get("recipient", ""))
    blank_lines(doc, 1)

    body = cfg.get("body_text", "")
    if not body and cfg.get("input"):
        body, _ = read_input(cfg["input"])
    if body:
        render_body(doc, F, body, title=cfg.get("title", ""))

    if cfg.get("attachments"):
        add_attachment_note(doc, F, cfg["attachments"])

    if not minutes:
        add_signature(doc, F, cfg.get("author") or cfg.get("sign_office", ""),
                      format_cn_date(cfg.get("date", "")),
                      sealed=bool(cfg.get("seal")),
                      seal_image=cfg.get("seal_image"))
        add_notes(doc, F, cfg.get("notes", []))

    if cfg.get("attachment_docs"):
        add_attachment_pages(doc, F, cfg["attachment_docs"])

    if not minutes:
        add_banji(doc, F, cfg.get("cc", ""), cfg.get("print_author", ""),
                  format_cn_date(cfg.get("print_date", "")))

    if not cfg.get("no_page_num"):
        add_page_numbers(doc, F, first_page_no_number=bool(cfg.get("no_first_page_num")))

    out_dir = os.path.dirname(os.path.abspath(out_path))
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    _safe_save_docx(doc, out_path)
    return out_path, F


def export_pdf(docx_path):
    import subprocess
    out_dir = os.path.dirname(os.path.abspath(docx_path)) or "."
    pdf = os.path.splitext(docx_path)[0] + ".pdf"
    for exe in ("soffice", "libreoffice"):
        try:
            r = subprocess.run([exe, "--headless", "--convert-to", "pdf",
                                "--outdir", out_dir, os.path.abspath(docx_path)],
                               capture_output=True, text=True, timeout=120)
            if r.returncode == 0 and os.path.exists(pdf):
                return pdf
        except FileNotFoundError:
            continue
        except subprocess.TimeoutExpired:
            sys.stderr.write("[warn] PDF 转换超时\n")
            return None
    sys.stderr.write("[warn] 未找到 LibreOffice，跳过 PDF 导出\n")
    return None


# ===================== 七、CLI =====================

DEMO_BODY = """一、提高思想认识

近年来，极端天气频发，城市排水防涝形势严峻。各区县要坚持人民至上、生命至上，压实责任、补齐短板，切实保障人民群众生命财产安全。

（一）压实属地责任。各区县人民政府对本行政区域排水防涝工作负总责，主要负责同志亲自抓。

1. 建立台账。对辖区内易涝点位逐一登记，动态更新，实行销号管理。

2. 限期整改。2026年主汛期前完成整治，逾期未完成的应当书面说明原因。

（二）强化部门协同。住建、水务、气象、应急等部门应当健全会商研判和联合调度机制。

二、健全工作机制

各区县、各部门应当于本通知印发之日起三十日内，制定本地区本部门实施方案并报市政府办公室备案。

附件：1. 城市易涝点位整治清单
2. 排水防涝责任分工表
"""


def parse_args():
    ap = argparse.ArgumentParser(
        description="党政机关公文标准排版 (GB/T 9704-2012)",
        formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--config", help="JSON 配置文件（与命令行参数可混用，命令行优先）")
    ap.add_argument("--title", help="公文标题（2 号小标宋体居中）")
    ap.add_argument("--input", help="正文来源文件 .md / .txt / .docx")
    ap.add_argument("--output", "--out", dest="output", help="输出 .docx 路径")

    g1 = ap.add_argument_group("版头要素")
    g1.add_argument("--redhead", default="", help="发文机关标志，联合行文用「;」分隔")
    g1.add_argument("--doc-number", default="", help="发文字号，如 ×政发〔2026〕12号")
    g1.add_argument("--copies", default="", help="份号，6 位阿拉伯数字")
    g1.add_argument("--secret-level", default="", help="密级和保密期限，如 机密★20年")
    g1.add_argument("--urgency", default="", choices=["", *URGENCY_LEVELS], help="紧急程度")
    g1.add_argument("--signer", default="", help="签发人（上报的公文必备），多人用「、」分隔")

    g2 = ap.add_argument_group("主体要素")
    g2.add_argument("--recipient", default="", help="主送机关")
    g2.add_argument("--attachment", action="append", default=[], help="附件说明条目，可重复")
    g2.add_argument("--attachment-doc", action="append", default=[],
                    help="附件正文文件，另面编排。格式 path 或 path::附件标题")
    g2.add_argument("--author", default="", help="发文机关署名，联合行文用「;」分隔")
    g2.add_argument("--date", default="", help="成文日期，支持 2026-08-06 / today")
    g2.add_argument("--seal", action="store_true", help="加盖印章版式（日期右空四字）")
    g2.add_argument("--seal-image", default="", help="印章图片路径（可选）")
    g2.add_argument("--notes", nargs="*", default=[], help="附注，如 此件公开发布")

    g3 = ap.add_argument_group("版记与页面")
    g3.add_argument("--cc", default="", help="抄送机关")
    g3.add_argument("--print-author", default="", help="印发机关")
    g3.add_argument("--print-date", default="", help="印发日期")
    g3.add_argument("--no-page-num", action="store_true", help="不编页码")
    g3.add_argument("--no-first-page-num", action="store_true",
                    help="首页不显示页码（国标未作此要求，仅供单位惯例使用）")
    g3.add_argument("--minutes", action="store_true", help="会议纪要版式（无版记、无印章）")
    g3.add_argument("--strict-font", action="store_true",
                    help="强制使用国标字体名，不做本机可用性回退")
    g3.add_argument("--pdf", action="store_true", help="同时导出 PDF（需 LibreOffice）")
    g3.add_argument("--demo", action="store_true", help="生成示例公文与示例配置")
    return ap.parse_args()


def main():
    a = parse_args()

    if a.demo:
        cfg = {
            "redhead": "××市人民政府",
            "doc_number": "×政发〔2026〕12号",
            "title": "××市人民政府关于进一步加强城市排水防涝工作的通知",
            "recipient": "各区县人民政府，市政府各部门、各直属机构",
            "body_text": DEMO_BODY,
            "author": "××市人民政府",
            "date": "2026-08-06",
            "seal": True,
            "notes": ["此件公开发布"],
            "cc": "市委办公室，市人大常委会办公室，市政协办公室",
            "print_author": "××市人民政府办公室",
            "print_date": "2026-08-07",
        }
        out = os.path.join("output", "示例公文_GB9704.docx")
        with open("gongwen_config.demo.json", "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
        path, F = build(cfg, out)
        print(f"[ok] 示例配置：gongwen_config.demo.json")
        print(f"[ok] 示例公文：{path}")
        print(f"[i ] {F.report()}")
        return

    cfg = {}
    if a.config:
        with open(a.config, "r", encoding="utf-8") as f:
            cfg = json.load(f)

    def pick(key, val, default=None):
        if val not in (None, "", [], False):
            cfg[key] = val
        elif key not in cfg and default is not None:
            cfg[key] = default

    pick("title", a.title)
    pick("input", a.input)
    pick("redhead", [x.strip() for x in a.redhead.split(";") if x.strip()] if a.redhead else None)
    pick("doc_number", a.doc_number)
    pick("copies", a.copies)
    pick("secret_level", a.secret_level)
    pick("urgency", a.urgency)
    pick("signer", a.signer)
    pick("recipient", a.recipient)
    pick("attachments", a.attachment)
    pick("attachment_docs", a.attachment_doc)
    pick("author", [x.strip() for x in a.author.split(";") if x.strip()] if a.author else None)
    pick("date", a.date)
    pick("seal", a.seal)
    pick("seal_image", a.seal_image)
    pick("notes", a.notes)
    pick("cc", a.cc)
    pick("print_author", a.print_author)
    pick("print_date", a.print_date)
    pick("no_page_num", a.no_page_num)
    pick("no_first_page_num", a.no_first_page_num)
    pick("minutes", a.minutes)
    pick("strict_font", a.strict_font)

    out = a.output or cfg.get("output") or "output/公文成品.docx"
    want_pdf = a.pdf or out.lower().endswith(".pdf")
    if out.lower().endswith(".pdf"):
        out = os.path.splitext(out)[0] + ".docx"

    if not cfg.get("title"):
        sys.stderr.write("[err] 缺少 --title（公文标题）\n")
        sys.exit(1)
    if not (cfg.get("body_text") or cfg.get("input")):
        sys.stderr.write("[err] 缺少正文来源：--input 文件 或 配置中的 body_text\n")
        sys.exit(1)

    path, F = build(cfg, out)
    print(f"[ok] 公文已生成：{path}")
    print(f"[i ] {F.report()}")
    if F.title in ("宋体", "SimSun"):
        print("[!] 本机未安装方正小标宋简体，标题已回退为宋体；"
              "正式行文建议在装有公文字体的机器上重排。")
    if want_pdf:
        pdf = export_pdf(path)
        if pdf:
            print(f"[ok] PDF 已生成：{pdf}")


if __name__ == "__main__":
    main()
