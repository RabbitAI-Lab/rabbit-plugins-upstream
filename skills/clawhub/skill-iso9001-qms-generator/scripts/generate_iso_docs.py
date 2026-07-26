#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ISO9001 质量管理体系文件生成器
将 Markdown 格式的内容转换为 Word 文档
支持：质量手册、程序文件、作业指导书、表单模板
"""

import argparse
import json
import os
import re
from typing import Dict, Optional, List

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.table import Table
except ImportError:
    print("错误: 需要安装 python-docx 库")
    print("请运行: pip install python-docx==1.1.2")
    exit(1)


def parse_markdown_to_docx(
    content: str,
    output_path: str,
    doc_type: str = "程序文件",
    title: str = "",
    company_info: Optional[Dict] = None,
    include_cover: bool = True
) -> None:
    """
    将 Markdown 内容转换为 Word 文档
    
    参数:
        content: Markdown 格式的文档内容
        output_path: 输出 Word 文档路径
        doc_type: 文档类型（质量手册/程序文件/作业指导书/表单）
        title: 文档标题
        company_info: 企业信息字典
        include_cover: 是否包含封面
    """
    if company_info is None:
        company_info = {}
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置中文字体
    set_chinese_font(doc)
    
    # 添加封面
    if include_cover:
        add_cover(doc, title, doc_type, company_info)
    
    # 添加目录占位符
    doc.add_heading('目录', level=1)
    p = doc.add_paragraph('（请使用 Word 自动生成目录功能：引用 -> 目录 -> 自动目录）')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()
    
    # 解析 Markdown 内容
    parse_content(doc, content)
    
    # 添加页眉页脚
    add_header_footer(doc, title, doc_type, company_info)
    
    # 保存文档
    doc.save(output_path)
    print(f"✓ 文档已生成: {output_path}")


def set_chinese_font(doc: Document) -> None:
    """
    设置文档默认中文字体
    """
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_cover(doc: Document, title: str, doc_type: str, company_info: Dict) -> None:
    """
    添加封面
    """
    # 标题
    title_heading = doc.add_heading('', level=0)
    title_run = title_heading.add_run(title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 文档类型
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    type_para = doc.add_paragraph(f'{doc_type}')
    type_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    type_para.runs[0].font.size = Pt(18)
    type_para.runs[0].font.bold = True
    
    # 企业信息
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    company_name = company_info.get('name', '')
    if company_name:
        para = doc.add_paragraph(company_name)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.runs[0].font.size = Pt(16)
        para.runs[0].font.bold = True
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # 版本信息
    version_para = doc.add_paragraph('版本：A/0')
    version_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    version_para.runs[0].font.size = Pt(14)
    
    date_para = doc.add_paragraph('日期：2024年')
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.runs[0].font.size = Pt(14)
    
    # 分页
    doc.add_page_break()


def generate_form_template(
    content: str,
    output_path: str,
    doc_type: str = "表单",
    title: str = "",
    company_info: Optional[Dict] = None,
    include_cover: bool = False
) -> None:
    """
    生成表单模板
    
    参数:
        content: 表单内容的 Markdown 格式（可能包含表格定义）
        output_path: 输出 Word 文档路径
        doc_type: 文档类型（表单）
        title: 文档标题
        company_info: 企业信息字典
        include_cover: 是否包含封面
    """
    if company_info is None:
        company_info = {}
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置中文字体
    set_chinese_font(doc)
    
    # 添加封面（如果需要）
    if include_cover:
        add_cover(doc, title, doc_type, company_info)
    
    # 添加标题
    if title:
        doc.add_heading(title, level=1)
    
    # 解析内容并转换为表单格式
    parse_form_content(doc, content, title, company_info)
    
    # 添加页眉页脚
    add_header_footer(doc, title, doc_type, company_info)
    
    # 保存文档
    doc.save(output_path)
    print(f"✓ 表单模板已生成: {output_path}")


def parse_form_content(doc: Document, content: str, title: str, company_info: Dict) -> None:
    """
    解析表单内容并转换为表单格式
    """
    lines = content.split('\n')
    i = 0
    
    # 添加表单基本信息
    add_form_header(doc, title, company_info)
    
    # 解析表格内容
    current_table_data = []
    in_table = False
    
    while i < len(lines):
        line = lines[i]
        
        # 检测表格开始
        if line.strip().startswith('|') and line.strip().endswith('|'):
            in_table = True
            current_table_data = []
            
            # 收集所有表格行
            while i < len(lines) and lines[i].strip().startswith('|'):
                current_table_data.append(lines[i])
                i += 1
            
            # 跳过分隔行
            if len(current_table_data) > 1 and re.match(r'^[\s|\-:]+$', current_table_data[1]):
                table_content = current_table_data[2:]
            else:
                table_content = current_table_data
            
            # 生成表格
            if table_content:
                generate_form_table(doc, table_content)
            
            in_table = False
            continue
        
        # 普通段落
        if line.strip():
            doc.add_paragraph(line.strip())
        
        i += 1
    
    # 添加表单底部签名区
    add_form_footer(doc)


def add_form_header(doc: Document, title: str, company_info: Dict) -> None:
    """
    添加表单头部信息
    """
    # 表单编号
    p = doc.add_paragraph()
    p.add_run('表单编号：').font.bold = True
    p.add_run('【待补充】')
    
    # 表单版本
    p = doc.add_paragraph()
    p.add_run('表单版本：').font.bold = True
    p.add_run('A/0')
    
    # 表单名称
    if title:
        p = doc.add_paragraph()
        p.add_run('表单名称：').font.bold = True
        p.add_run(title)
    
    # 分隔线
    doc.add_paragraph('_' * 50)
    
    # 表单日期
    p = doc.add_paragraph()
    p.add_run('填写日期：').font.bold = True
    p.add_run('20  年  月  日')


def generate_form_table(doc: Document, table_content: List[str]) -> None:
    """
    生成表单表格
    """
    if not table_content:
        return
    
    # 计算列数（以第一行为准）
    cols = len([cell for cell in table_content[0].split('|') if cell.strip()])
    rows = len(table_content)
    
    # 创建表格
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    
    # 填充表格内容
    for r_idx, t_line in enumerate(table_content):
        # 解析单元格内容
        cells = []
        current_cell = ''
        for char in t_line:
            if char == '|':
                cells.append(current_cell.strip())
                current_cell = ''
            else:
                current_cell += char
        
        # 填充单元格
        for c_idx, cell_text in enumerate(cells):
            if c_idx < cols:
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = cell_text
                
                # 第一行加粗
                if r_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True


def add_form_footer(doc: Document) -> None:
    """
    添加表单底部签名区
    """
    doc.add_paragraph('_' * 50)
    doc.add_paragraph('')
    
    # 创建签名表格
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    
    # 第一行
    table.rows[0].cells[0].text = '编制人：'
    table.rows[0].cells[1].text = '审核人：'
    table.rows[0].cells[2].text = '批准人：'
    
    # 第二行
    table.rows[1].cells[0].text = '日期：'
    table.rows[1].cells[1].text = '日期：'
    table.rows[1].cells[2].text = '日期：'


def parse_content(doc: Document, content: str) -> None:
    """
    解析 Markdown 内容并转换为 Word 格式
    """
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # 空行
        if not line.strip():
            doc.add_paragraph('')
            i += 1
            continue
        
        # 标题处理 (# ## ### ####)
        if line.startswith('#'):
            level = 0
            while level < len(line) and line[level] == '#':
                level += 1
            title_text = line[level:].strip()
            if level <= 4:
                doc.add_heading(title_text, level=level)
            else:
                doc.add_heading(title_text, level=4)
            i += 1
            continue
        
        # 流程图标记
        if line.strip() == '```flowchart':
            # 添加流程图占位符
            doc.add_paragraph('')
            p = doc.add_paragraph('【流程图占位符】')
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            p = doc.add_paragraph('此处建议插入流程图。请使用 Word 的 SmartArt 功能：')
            p = doc.add_paragraph('插入 -> SmartArt -> 流程 -> 选择合适的布局')
            doc.add_paragraph('')
            
            # 跳过流程图内容直到 ```
            i += 1
            while i < len(lines) and lines[i].strip() != '```':
                i += 1
            i += 1
            continue
        
        # 代码块
        if line.strip().startswith('```'):
            code_type = line.strip()[3:]
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                # 添加代码块
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            
            i += 1
            continue
        
        # 列表处理
        if line.strip().startswith(('-', '*', '+')):
            text = line.strip()[1:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        # 有序列表处理
        if re.match(r'^\s*\d+\.', line):
            text = re.sub(r'^\s*\d+\.', '', line).strip()
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # 表格处理
        if line.strip().startswith('|') and line.strip().endswith('|'):
            # 解析表格
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            # 创建表格
            rows = len(table_lines)
            cols = len([cell for cell in table_lines[0].split('|') if cell.strip()])
            
            # 跳过分隔行
            if rows > 1 and re.match(r'^[\s|\-:]+$', table_lines[1]):
                table_content = table_lines[2:]
                table_rows = len(table_content)
            else:
                table_content = table_lines
                table_rows = rows
            
            if table_rows > 0:
                table = doc.add_table(rows=table_rows, cols=cols)
                table.style = 'Light Grid Accent 1'
                
                for r_idx, t_line in enumerate(table_content):
                    cells = [cell.strip() for cell in t_line.split('|') if cell.strip() or cell == '']
                    for c_idx, cell_text in enumerate(cells):
                        if c_idx < cols:
                            cell = table.rows[r_idx].cells[c_idx]
                            cell.text = cell_text
                            # 表头加粗
                            if r_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
            
            continue
        
        # 普通段落
        doc.add_paragraph(line.strip())
        i += 1


def add_header_footer(doc: Document, title: str, doc_type: str, company_info: Dict) -> None:
    """
    添加页眉页脚
    """
    section = doc.sections[0]
    
    # 页眉
    header = section.header
    header_para = header.paragraphs[0]
    company_name = company_info.get('name', '')
    header_text = f"{company_name} - {title}" if company_name else title
    header_para.text = header_text
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 页脚
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"第 {''} 页"
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def main():
    parser = argparse.ArgumentParser(description='ISO9001 质量管理体系文件生成器')
    parser.add_argument('--markdown', type=str, help='Markdown 格式的文档内容')
    parser.add_argument('--output', type=str, required=True, help='输出 Word 文档路径')
    parser.add_argument('--doc-type', type=str, default='程序文件', 
                        choices=['质量手册', '程序文件', '作业指导书', '表单', '表单模板'],
                        help='文档类型')
    parser.add_argument('--title', type=str, default='', help='文档标题')
    parser.add_argument('--company-name', type=str, default='', help='企业名称')
    parser.add_argument('--company-info', type=str, default='{}', 
                        help='企业信息 JSON 字符串')
    parser.add_argument('--no-cover', action='store_true', help='不包含封面')
    
    args = parser.parse_args()
    
    # 获取内容
    if args.markdown:
        content = args.markdown
    elif not sys.stdin.isatty():
        content = sys.stdin.read()
    else:
        print("错误: 需要通过 --markdown 参数或标准输入提供内容")
        parser.print_help()
        exit(1)
    
    # 解析企业信息
    try:
        company_info = json.loads(args.company_info)
        if args.company_name:
            company_info['name'] = args.company_name
    except json.JSONDecodeError:
        company_info = {'name': args.company_name}
    
    # 如果没有标题，使用文档类型
    if not args.title:
        args.title = args.doc_type
    
    # 根据文档类型选择生成方法
    if args.doc_type in ['表单', '表单模板']:
        # 使用表单模板生成方法
        generate_form_template(
            content=content,
            output_path=args.output,
            doc_type=args.doc_type,
            title=args.title,
            company_info=company_info,
            include_cover=not args.no_cover
        )
    else:
        # 使用标准文档生成方法
        parse_markdown_to_docx(
            content=content,
            output_path=args.output,
            doc_type=args.doc_type,
            title=args.title,
            company_info=company_info,
            include_cover=not args.no_cover
        )


if __name__ == '__main__':
    import sys
    main()
    main()
