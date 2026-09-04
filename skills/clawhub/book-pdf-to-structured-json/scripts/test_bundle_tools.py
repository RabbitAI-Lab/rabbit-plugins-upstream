#!/usr/bin/env python3
"""Regression tests for the deterministic builder and fail-closed validator."""

from __future__ import annotations

import json
import hashlib
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path


SCRIPT_DIR = Path(__file__).resolve().parent
BUILDER = SCRIPT_DIR / "build_derivatives.py"
VALIDATOR = SCRIPT_DIR / "validate_bundle.py"
CONVERTER = SCRIPT_DIR / "convert_book.py"


def valid_tree() -> dict[str, object]:
    content = "这是经过来源复核的正文。"
    return {
        "book_id": "demo-book",
        "title": "示例书",
        "authority": "printed_toc",
        "node_count": 1,
        "nodes": [
            {
                "key": "toc-001",
                "title": "第一章 示例",
                "level": 1,
                "parent_key": None,
                "sort": 1,
                "logical_page": 1,
                "source_page": 3,
                "heading_start": 10,
                "content_start": 16,
                "content_end": 40,
                "heading_score": 1.0,
                "structural_only": False,
                "content": content,
                "content_chars": len(content),
                "source_file": "demo-source.txt",
            }
        ],
    }


class BundleToolsTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary.name)
        self.canonical = self.root / "canonical"
        self.bundle = self.root / "bundle"
        self.canonical.mkdir()
        self.tree_path = self.canonical / "01_demo_tree.json"
        self.tree_path.write_text(
            json.dumps(valid_tree(), ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        self.review = self.root / "ocr-review.tsv"
        self.review.write_text(
            "candidate_id\tbook_id\tnode_key\toriginal\tsuggestion\tcontext\tdecision\treviewer\n",
            encoding="utf-8",
        )

    def tearDown(self) -> None:
        self.temporary.cleanup()

    def run_tool(self, *arguments: object) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [sys.executable, *(str(item) for item in arguments)],
            text=True,
            capture_output=True,
            check=False,
        )

    def build(self) -> None:
        result = self.run_tool(BUILDER, self.tree_path, "--output-dir", self.bundle)
        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)

    def validate(self) -> subprocess.CompletedProcess[str]:
        return self.run_tool(
            VALIDATOR,
            self.tree_path,
            "--artifact-dir",
            self.bundle,
            "--ocr-review",
            self.review,
            "--require-ocr-review",
            "--fail-on-warnings",
        )

    def test_valid_bundle_passes(self) -> None:
        self.build()
        result = self.validate()
        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        self.assertIn("PASS\t1 books\t1 nodes", result.stdout)

    def test_rebuild_is_byte_deterministic(self) -> None:
        self.build()
        first = {
            path.relative_to(self.bundle): hashlib.sha256(path.read_bytes()).hexdigest()
            for path in self.bundle.rglob("*")
            if path.is_file()
        }
        result = self.run_tool(BUILDER, self.tree_path, "--output-dir", self.bundle, "--replace")
        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        second = {
            path.relative_to(self.bundle): hashlib.sha256(path.read_bytes()).hexdigest()
            for path in self.bundle.rglob("*")
            if path.is_file()
        }
        self.assertEqual(first, second)

    def test_stale_derivative_fails(self) -> None:
        self.build()
        target = self.bundle / "01_demo_all-content.txt"
        target.write_text(target.read_text(encoding="utf-8") + "旧内容\n", encoding="utf-8")
        result = self.validate()
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("stale or non-deterministic derivative", result.stderr)

    def test_missing_required_field_fails(self) -> None:
        tree = valid_tree()
        del tree["nodes"][0]["structural_only"]  # type: ignore[index]
        self.tree_path.write_text(json.dumps(tree, ensure_ascii=False), encoding="utf-8")
        self.build()
        result = self.validate()
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("missing fields structural_only", result.stderr)

    def conversion_source(self) -> tuple[Path, Path]:
        source = self.root / "source.txt"
        source.write_text(
            "===== 第 1 页 =====\n"
            "目录\n第一章 开始 …… 1\n第二章 继续 …… 2\n"
            "===== 第 2 页 =====\n"
            "第一章 开始\n这是第一章经过复核的正文。\n"
            "===== 第 3 页 =====\n"
            "第二章 继续\n这是第二章经过复核的正文。\n",
            encoding="utf-8",
        )
        toc = self.root / "toc.json"
        toc.write_text(
            json.dumps(
                [
                    {"title": "第一章 开始", "level": 1, "logical_page": 1},
                    {"title": "第二章 继续", "level": 1, "logical_page": 2},
                ],
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        return source, toc

    def test_txt_to_final_json_end_to_end(self) -> None:
        source, toc = self.conversion_source()
        output = self.root / "converted"
        result = self.run_tool(
            CONVERTER,
            source,
            "--output-dir",
            output,
            "--book-id",
            "converted-book",
            "--title",
            "转换示例",
            "--toc-json",
            toc,
            "--toc-pages",
            "1",
            "--reviewer",
            "test-reviewer",
        )
        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        tree = json.loads((output / "converted-book_tree.json").read_text(encoding="utf-8"))
        self.assertEqual(tree["node_count"], 2)
        self.assertIn("第一章经过复核的正文", tree["nodes"][0]["content"])
        self.assertIn("第二章经过复核的正文", tree["nodes"][1]["content"])

    def test_auto_toc_writes_json_but_stays_draft_without_review(self) -> None:
        source, _ = self.conversion_source()
        output = self.root / "auto-draft"
        result = self.run_tool(
            CONVERTER,
            source,
            "--output-dir",
            output,
            "--book-id",
            "auto-book",
            "--title",
            "自动目录示例",
        )
        self.assertEqual(result.returncode, 3, result.stdout + result.stderr)
        self.assertTrue((output / "auto-book_tree.json").is_file())
        tree = json.loads((output / "auto-book_tree.json").read_text(encoding="utf-8"))
        self.assertEqual(tree["review_status"], "pending_toc_review")

    def test_docx_to_final_json_end_to_end(self) -> None:
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is unavailable")
        source = self.root / "source.docx"
        document = Document()
        document.add_paragraph("第一章 开始")
        document.add_paragraph("这是第一章经过复核的正文。")
        document.add_paragraph("第二章 继续")
        document.add_paragraph("这是第二章经过复核的正文。")
        document.save(source)
        _, toc = self.conversion_source()
        output = self.root / "docx-converted"
        result = self.run_tool(
            CONVERTER,
            source,
            "--output-dir",
            output,
            "--book-id",
            "docx-book",
            "--title",
            "DOCX 转换示例",
            "--toc-json",
            toc,
            "--reviewer",
            "test-reviewer",
        )
        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        tree = json.loads((output / "docx-book_tree.json").read_text(encoding="utf-8"))
        self.assertEqual(tree["source_format"], "docx")
        self.assertEqual(tree["node_count"], 2)

    def test_reviewed_correction_is_applied_before_validation(self) -> None:
        source, toc = self.conversion_source()
        source.write_text(
            source.read_text(encoding="utf-8").replace("第一章经过复核", "第一章错宇"),
            encoding="utf-8",
        )
        ledger = self.root / "reviewed-corrections.tsv"
        ledger.write_text(
            "candidate_id\tbook_id\tnode_key\toriginal\tsuggestion\tcontext\tdecision\treviewer\n"
            "manual-001\tcorrected-book\ttoc-001\t错宇\t错字\t第一章错宇正文\tcorrected\ttest-reviewer\n",
            encoding="utf-8",
        )
        output = self.root / "corrected"
        result = self.run_tool(
            CONVERTER,
            source,
            "--output-dir",
            output,
            "--book-id",
            "corrected-book",
            "--title",
            "自动校正示例",
            "--toc-json",
            toc,
            "--toc-pages",
            "1",
            "--review-ledger",
            ledger,
            "--reviewer",
            "test-reviewer",
        )
        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        tree = json.loads((output / "corrected-book_tree.json").read_text(encoding="utf-8"))
        self.assertIn("第一章错字的正文", tree["nodes"][0]["content"])
        self.assertNotIn("错宇", tree["nodes"][0]["content"])

    def test_text_pdf_to_final_json_end_to_end(self) -> None:
        try:
            from reportlab.pdfgen import canvas
        except ImportError:
            self.skipTest("reportlab is unavailable")
        source = self.root / "source.pdf"
        document = canvas.Canvas(str(source))
        document.drawString(72, 760, "Contents")
        document.drawString(72, 730, "Chapter One ........ 1")
        document.drawString(72, 710, "Chapter Two ........ 2")
        document.showPage()
        document.drawString(72, 760, "Chapter One")
        document.drawString(72, 730, "Reviewed body for the first chapter.")
        document.showPage()
        document.drawString(72, 760, "Chapter Two")
        document.drawString(72, 730, "Reviewed body for the second chapter.")
        document.save()
        toc = self.root / "pdf-toc.json"
        toc.write_text(
            json.dumps(
                [
                    {"title": "Chapter One", "level": 1, "logical_page": 1},
                    {"title": "Chapter Two", "level": 1, "logical_page": 2},
                ]
            ),
            encoding="utf-8",
        )
        output = self.root / "pdf-converted"
        result = self.run_tool(
            CONVERTER,
            source,
            "--output-dir",
            output,
            "--book-id",
            "pdf-book",
            "--title",
            "PDF conversion example",
            "--toc-json",
            toc,
            "--toc-pages",
            "1",
            "--reviewer",
            "test-reviewer",
        )
        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        tree = json.loads((output / "pdf-book_tree.json").read_text(encoding="utf-8"))
        self.assertEqual(tree["extraction_method"], "text_layer")
        self.assertEqual(tree["node_count"], 2)


if __name__ == "__main__":
    unittest.main()
