"""提取模板原子结构：段落、表格（含每个 cell 行列索引与文本）、标题层级、文档顺序。

支持嵌套表格与合并单元格识别。"""

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


def _iter_block_items(parent):
    """按文档顺序产出段落与表格，覆盖嵌套表格。"""
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, Table):
        # 表格内单元格内的 block items 通过 cell.paragraphs 间接处理
        return
    else:
        parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield ("paragraph", Paragraph(child, parent))
        elif isinstance(child, CT_Tbl):
            yield ("table", Table(child, parent))


def _extract_merged_cells(table):
    """识别合并单元格，返回 [{row, col, row_span, col_span}]。"""
    merged = []
    tbl = table._tbl
    grid = tbl.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblGrid"
    )
    seen = set()
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            tc = cell._tc
            tc_id = id(tc)
            if tc_id in seen:
                continue
            # 同一 tc 跨多列/多行时，row.cells 会出现重复引用
            same_tc_positions = []
            for r_idx, r in enumerate(table.rows):
                for c_idx, c in enumerate(r.cells):
                    if id(c._tc) == tc_id:
                        same_tc_positions.append((r_idx, c_idx))
            if not same_tc_positions:
                continue
            rows = [p[0] for p in same_tc_positions]
            cols = [p[1] for p in same_tc_positions]
            r_min, r_max = min(rows), max(rows)
            c_min, c_max = min(cols), max(cols)
            seen.add(tc_id)
            if r_max > r_min or c_max > c_min:
                merged.append({
                    "row": r_min,
                    "col": c_min,
                    "row_span": r_max - r_min + 1,
                    "col_span": c_max - c_min + 1,
                })
    return merged


def _extract_table(table, table_index):
    """提取单个表格的完整结构。"""
    rows = []
    for r_idx, row in enumerate(table.rows):
        cells_data = []
        seen_tc = set()
        for c_idx, cell in enumerate(row.cells):
            tc_id = id(cell._tc)
            if tc_id in seen_tc:
                cells_data.append({"merged_from": True, "row": r_idx, "col": c_idx})
                continue
            seen_tc.add(tc_id)
            cell_paragraphs = []
            for p in cell.paragraphs:
                cell_paragraphs.append({
                    "text": p.text,
                    "style": p.style.name if p.style else None,
                })
            nested_tables = []
            for nt in cell.tables:
                nested_tables.append(_extract_table(nt, f"{table_index}.{r_idx}.{c_idx}"))
            cells_data.append({
                "row": r_idx,
                "col": c_idx,
                "text": cell.text,
                "paragraphs": cell_paragraphs,
                "nested_tables": nested_tables,
            })
        rows.append({"row_index": r_idx, "cells": cells_data})

    return {
        "table_index": table_index,
        "rows": len(table.rows),
        "cols": len(table.columns),
        "grid": rows,
        "merged_cells": _extract_merged_cells(table),
    }


def extract_structure(template_path: str) -> dict:
    """提取模板的原子结构。

    输出:
        paragraphs: [{index, text, style}]
        tables: [{table_index, rows, cols, grid, merged_cells}]
        headings: [{level, text, para_index}]
        body_order: [{type: "paragraph"|"table", index, table_index?}]
    """
    doc = Document(template_path)
    paragraphs = []
    tables = []
    headings = []
    body_order = []

    para_idx = 0
    table_idx = 0

    for block_type, block in _iter_block_items(doc):
        if block_type == "paragraph":
            text = block.text
            style_name = block.style.name if block.style else ""
            paragraphs.append({
                "index": para_idx,
                "text": text,
                "style": style_name,
            })
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading", "").strip() or "1")
                except ValueError:
                    level = 1
                headings.append({
                    "level": level,
                    "text": text,
                    "para_index": para_idx,
                })
            body_order.append({"type": "paragraph", "index": para_idx})
            para_idx += 1
        elif block_type == "table":
            table_data = _extract_table(block, table_idx)
            tables.append(table_data)
            body_order.append({"type": "table", "table_index": table_idx})
            table_idx += 1

    return {
        "template_path": str(template_path),
        "paragraphs": paragraphs,
        "tables": tables,
        "headings": headings,
        "body_order": body_order,
    }


def main():
    parser = argparse.ArgumentParser(description="提取 docx 模板原子结构")
    parser.add_argument("--template", required=True, help="模板文件路径")
    parser.add_argument("--output", required=True, help="输出 JSON 路径")
    args = parser.parse_args()

    if not Path(args.template).exists():
        print(json.dumps({"error": f"模板不存在: {args.template}"}, ensure_ascii=False))
        sys.exit(1)

    try:
        result = extract_structure(args.template)
        Path(args.output).write_text(
            json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        print(json.dumps({
            "success": True,
            "output": args.output,
            "paragraphs": len(result["paragraphs"]),
            "tables": len(result["tables"]),
            "headings": len(result["headings"]),
        }, ensure_ascii=False, indent=2))
    except Exception as e:
        print(json.dumps({"success": False, "error": str(e)}, ensure_ascii=False))
        sys.exit(1)


if __name__ == "__main__":
    main()
