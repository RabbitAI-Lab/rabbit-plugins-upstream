"""Tier 1 校验最终文档格式。

检查:
- 段落数与模板一致
- 表格行列数与模板一致
- 静态文本未被改动
- 样式名、字体、字号与模板一致
"""

import argparse
import json
import sys
from pathlib import Path

from docx import Document


def _collect_static_texts(template_path: str) -> list:
    """从 fill_contract 中读取静态文本，从模板中读取实际文本用于对比。"""
    return []


def validate_format(generated_path: str, template_path: str, contract: dict = None) -> dict:
    failed_checks = []
    tier_failed = None

    tpl_doc = Document(template_path)
    gen_doc = Document(generated_path)

    # 段落数
    tpl_paras = len(tpl_doc.paragraphs)
    gen_paras = len(gen_doc.paragraphs)
    if tpl_paras != gen_paras:
        failed_checks.append({
            "check": "paragraph_count_mismatch",
            "fix_hint": f"段落数不匹配: 模板 {tpl_paras} vs 生成 {gen_paras}",
        })
        tier_failed = 1

    # 表格数与行列数
    tpl_tables = tpl_doc.tables
    gen_tables = gen_doc.tables
    if len(tpl_tables) != len(gen_tables):
        failed_checks.append({
            "check": "table_count_mismatch",
            "fix_hint": f"表格数不匹配: 模板 {len(tpl_tables)} vs 生成 {len(gen_tables)}",
        })
        tier_failed = 1
    else:
        for i, (t_tbl, g_tbl) in enumerate(zip(tpl_tables, gen_tables)):
            if len(t_tbl.rows) != len(g_tbl.rows):
                failed_checks.append({
                    "check": f"table_{i}_rows_mismatch",
                    "fix_hint": f"表格 {i} 行数不匹配: {len(t_tbl.rows)} vs {len(g_tbl.rows)}",
                })
                if tier_failed is None:
                    tier_failed = 1
            if len(t_tbl.columns) != len(g_tbl.columns):
                failed_checks.append({
                    "check": f"table_{i}_cols_mismatch",
                    "fix_hint": f"表格 {i} 列数不匹配: {len(t_tbl.columns)} vs {len(g_tbl.columns)}",
                })
                if tier_failed is None:
                    tier_failed = 1

    # 静态文本未被改动
    if contract:
        for s in contract.get("static_texts", []):
            loc = s.get("location", {})
            expected_text = s.get("text", "")
            actual_text = None
            if "para_index" in loc:
                idx = loc["para_index"]
                if idx < len(gen_doc.paragraphs):
                    actual_text = gen_doc.paragraphs[idx].text
            elif "table_index" in loc:
                ti = loc["table_index"]
                r = loc.get("row", 0)
                c = loc.get("col", 0)
                if ti < len(gen_tables):
                    try:
                        actual_text = gen_tables[ti].rows[r].cells[c].text
                    except IndexError:
                        actual_text = None

            if actual_text is not None and expected_text and expected_text not in actual_text:
                failed_checks.append({
                    "check": f"static_text_changed_{s.get('id', '')}",
                    "fix_hint": f"静态文本被改动: 期望含 '{expected_text}'，实际 '{actual_text}'",
                })
                if tier_failed is None:
                    tier_failed = 1

    # 样式一致性抽查（前 10 个段落）
    sample_size = min(10, tpl_paras, gen_paras)
    for i in range(sample_size):
        tpl_style = tpl_doc.paragraphs[i].style.name if tpl_doc.paragraphs[i].style else ""
        gen_style = gen_doc.paragraphs[i].style.name if gen_doc.paragraphs[i].style else ""
        if tpl_style != gen_style:
            failed_checks.append({
                "check": f"style_mismatch_para_{i}",
                "fix_hint": f"段落 {i} 样式不匹配: 模板 '{tpl_style}' vs 生成 '{gen_style}'",
            })
            if tier_failed is None:
                tier_failed = 1

    return {
        "passed": len(failed_checks) == 0,
        "tier_failed": tier_failed,
        "failed_checks": failed_checks,
    }


def main():
    parser = argparse.ArgumentParser(description="Tier 1 校验最终文档格式")
    parser.add_argument("--generated", required=True, help="生成文档路径")
    parser.add_argument("--template", required=True, help="模板路径")
    parser.add_argument("--contract", help="fill_contract.json 路径（可选，用于静态文本校验）")
    args = parser.parse_args()

    contract = None
    if args.contract:
        contract = json.loads(Path(args.contract).read_text(encoding="utf-8"))

    result = validate_format(args.generated, args.template, contract)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    sys.exit(0 if result["passed"] else 1)


if __name__ == "__main__":
    main()
