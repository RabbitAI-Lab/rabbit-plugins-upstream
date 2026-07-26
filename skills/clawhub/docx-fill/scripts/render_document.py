"""将内容按 Fill Contract 的 location 填回模板，输出 .docx。

实现:
- 用 python-docx 打开原始模板（保留所有样式）
- 遍历 contract.placeholders，按 location 定位
- 替换占位符文本为生成内容
- static_texts 原样保留
- 渲染前清空既未声明为 placeholder 也未声明为 static 的非空 cell（防模板示例数据残留）
- 保存到 output_path

支持:
- 段落占位符 (para_index)
- 表格单元格占位符 (table_index + row + col)
- 多段内容（按换行符拆分为多个 Paragraph）
"""

import argparse
import json
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document


def _set_cell_text(cell, text: str):
    """替换单元格文本，保留首段样式，多行内容拆分为多段。"""
    if not cell.paragraphs:
        return
    lines = text.split("\n")
    # 第一行写入首个 paragraph
    first_para = cell.paragraphs[0]
    # 清空首段已有 runs
    for run in list(first_para.runs):
        run.text = ""
    if first_para.runs:
        first_para.runs[0].text = lines[0]
    else:
        first_para.add_run(lines[0])
    # 移除多余段落
    parent = first_para._element.getparent()
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    # 追加剩余行
    for line in lines[1:]:
        new_para = cell.add_paragraph()
        new_para.style = first_para.style
        new_para.add_run(line)


def _set_paragraph_text(para, text: str):
    """替换段落文本，保留首 run 样式。"""
    lines = text.split("\n")
    # 清空已有 runs
    for run in list(para.runs):
        run.text = ""
    if para.runs:
        para.runs[0].text = lines[0]
    else:
        para.add_run(lines[0])
    # 多行时追加新段落（紧随当前段后）— 简化处理：仅写第一行，多行内容警告
    if len(lines) > 1:
        # 把后续行追加为同段内换行 run
        for line in lines[1:]:
            para.add_run().add_break()
            para.runs[-1].text = line


def _collect_table_locs(items: list) -> set:
    """从 placeholder/static_text 列表中收集 (table_index, row, col) 集合。

    仅收集 table_index 为 int 的项（顶层表格），点分路径的嵌套表格不在
    本函数处理范围。
    """
    locs = set()
    for item in items:
        loc = item.get("location", {})
        if "table_index" not in loc:
            continue
        ti = loc["table_index"]
        # 仅处理顶层表格（int 类型）；嵌套表格的点分路径不在本清空逻辑范围
        if not isinstance(ti, int):
            continue
        locs.add((ti, loc.get("row", 0), loc.get("col", 0)))
    return locs


def _clear_unfilled_cells(doc, filled_locs: set, static_locs: set) -> list:
    """清空既未填充也未声明为静态的非空表格 cell。

    防止模板示例数据（如 "张三" / "示例：XXX"）残留到输出文档。
    合并单元格按 _tc id 去重，避免对同一单元格多次清空。
    仅处理顶层 doc.tables；嵌套表格不在本函数处理范围。

    返回被清空的 cell location 列表，便于审计。
    """
    cleared = []
    for tbl_idx, table in enumerate(doc.tables):
        seen_tc = set()
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                tc_id = id(cell._tc)
                if tc_id in seen_tc:
                    # 合并单元格的后续位置，跳过避免重复
                    continue
                seen_tc.add(tc_id)

                loc_key = (tbl_idx, r_idx, c_idx)
                if loc_key in filled_locs or loc_key in static_locs:
                    continue

                # 仅当 cell 有非空文本时才清空，避免无意义写入
                if cell.text and cell.text.strip():
                    _set_cell_text(cell, "")
                    cleared.append({"table_index": tbl_idx, "row": r_idx, "col": c_idx})
    return cleared


def render_document(template_path: str, contract: dict, content: dict, output_path: str) -> dict:
    doc = Document(template_path)

    content_map = {c["placeholder_id"]: c.get("text", "") for c in content.get("contents", [])}
    placeholders = contract.get("placeholders", [])
    static_texts = contract.get("static_texts", [])

    # === 方案 D: 清空未声明的非空 cell（防模板示例数据残留） ===
    # 仅对 is_placeholder=true 的占位符 location 视为"已填充"位置
    filled_locs = _collect_table_locs(
        [p for p in placeholders if p.get("is_placeholder", False)]
    )
    static_locs = _collect_table_locs(static_texts)
    cleared_cells = _clear_unfilled_cells(doc, filled_locs, static_locs)

    # === 填充占位符 ===
    filled_count = 0
    skipped = []

    for p in placeholders:
        if not p.get("is_placeholder", False):
            continue
        pid = p.get("id")
        if pid not in content_map:
            skipped.append({"placeholder_id": pid, "reason": "no content"})
            continue

        text = content_map[pid]
        loc = p.get("location", {})
        try:
            if "para_index" in loc:
                idx = loc["para_index"]
                if idx < len(doc.paragraphs):
                    _set_paragraph_text(doc.paragraphs[idx], text)
                    filled_count += 1
                else:
                    skipped.append({"placeholder_id": pid, "reason": f"para_index {idx} out of range"})
            elif "table_index" in loc:
                ti = loc["table_index"]
                r = loc.get("row", 0)
                c = loc.get("col", 0)
                if isinstance(ti, int) and ti < len(doc.tables):
                    try:
                        cell = doc.tables[ti].rows[r].cells[c]
                        _set_cell_text(cell, text)
                        filled_count += 1
                    except IndexError:
                        skipped.append({
                            "placeholder_id": pid,
                            "reason": f"table {ti} cell ({r},{c}) out of range",
                        })
                else:
                    skipped.append({"placeholder_id": pid, "reason": f"table_index {ti} out of range or unsupported nested path"})
        except Exception as e:
            skipped.append({"placeholder_id": pid, "reason": str(e)})

    # 确保输出目录存在
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    return {
        "output_path": str(output_path),
        "filled_count": filled_count,
        "cleared_cells_count": len(cleared_cells),
        "cleared_cells": cleared_cells,
        "skipped": skipped,
    }


def main():
    parser = argparse.ArgumentParser(description="渲染：将内容填回模板生成 .docx")
    parser.add_argument("--template", required=True, help="模板文件路径")
    parser.add_argument("--contract", required=True, help="fill_contract.json 路径")
    parser.add_argument("--content", required=True, help="generated_content.json 路径")
    parser.add_argument("--output", required=True, help="输出 .docx 路径")
    args = parser.parse_args()

    contract = json.loads(Path(args.contract).read_text(encoding="utf-8"))
    content = json.loads(Path(args.content).read_text(encoding="utf-8"))

    result = render_document(args.template, contract, content, args.output)
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
