# -*- coding: utf-8 -*-
"""
从 Markdown 源文件生成排版精良的 DOCX。
支持：多级标题、正文、无序/有序列表、任务清单、表格、引用块、提示块、
行内粗体，以及防止标题/列表/表格被分页断开的排版控制。
"""

import argparse
import re
import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── 默认样式常量 ──
FONT_HEI = "黑体"
FONT_BODY = "黑体"

COLOR_TITLE = RGBColor(0x11, 0x11, 0x11)
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)
COLOR_HINT = RGBColor(0x66, 0x66, 0x66)
COLOR_TABLE_HEADER_BG = "E8E8E8"
COLOR_BORDER = "CCCCCC"
COLOR_HINT_BAR = "999999"

SZ_DOC_TITLE = 22
SZ_CHAPTER = 16
SZ_SECTION = 14
SZ_METHOD = 12
SZ_SUBHEAD = 11
SZ_BODY = 10.5
SZ_TABLE = 10
SZ_HINT = 9.5
SZ_QUOTE = 10.5


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, color="CCCCCC", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def set_paragraph_spacing(para, before=0, after=60, line=280):
    pPr = para._p.get_or_add_pPr()
    spacing = parse_xml(
        f'<w:spacing {nsdecls("w")} '
        f'w:before="{before}" w:after="{after}" '
        f'w:line="{line}" w:lineRule="exact"/>'
    )
    pPr.append(spacing)


def set_keep_with_next(para, keep=True):
    para.paragraph_format.keep_with_next = keep


def set_keep_together(para, keep=True):
    para.paragraph_format.keep_together = keep


def set_table_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cant_split = parse_xml(f'<w:cantSplit {nsdecls("w")}/>')
    trPr.append(cant_split)


def add_page_number(section):
    """在页脚添加居中的「第 X 页」页码（使用 PAGE 字段，自动更新）"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line=240)

    run_pre = p.add_run("第 ")
    set_run_font(run_pre, font_name=FONT_BODY, size=SZ_HINT, bold=False, color=COLOR_HINT)

    run_fld = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run_fld._r.append(fld_begin)
    run_fld._r.append(instr)
    run_fld._r.append(fld_end)
    set_run_font(run_fld, font_name=FONT_BODY, size=SZ_HINT, bold=False, color=COLOR_HINT)

    run_post = p.add_run(" 页")
    set_run_font(run_post, font_name=FONT_BODY, size=SZ_HINT, bold=False, color=COLOR_HINT)


def set_run_font(run, font_name=FONT_BODY, size=SZ_BODY, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_left_border(para, color=COLOR_HINT_BAR, sz="12"):
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="8" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bottom_border(para, color="CCCCCC", sz="6"):
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="4" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def parse_inline_text(para, text, font_name=FONT_BODY, size=SZ_BODY, bold=False, color=None):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            run = para.add_run(inner)
            set_run_font(run, font_name=font_name, size=size, bold=True, color=color)
        elif part:
            run = para.add_run(part)
            set_run_font(run, font_name=font_name, size=size, bold=bold, color=color)


def parse_table_row(line):
    cells = [c.strip() for c in line.split('|')]
    if cells and cells[0] == '':
        cells = cells[1:]
    if cells and cells[-1] == '':
        cells = cells[:-1:]
    return cells


def is_table_separator(line):
    return bool(re.match(r'^\s*\|[\s\-:|]+\|\s*$', line))


def classify_hint(text):
    text = text.strip()
    if text.startswith('[提示]'):
        return 'hint', text
    elif text.startswith('[示例]'):
        return 'example', text
    elif text.startswith('[练习]'):
        return 'exercise', text
    elif text.startswith('[规则]'):
        return 'rule', text
    elif text.startswith('[核心观点]'):
        return 'key', text
    elif text.startswith('[关键]'):
        return 'key', text
    else:
        return 'quote', text


def build_doc(src_path, options=None):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)
        section.footer_distance = Cm(1.2)
        add_page_number(section)

    with open(src_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_table = False
    table_rows = []
    prev_was_heading = False  # 标记前一段是否为标题，用于下一段正文粘连

    while i < len(lines):
        line = lines[i].rstrip('\n')

        if not line.strip():
            if in_table and table_rows:
                _build_table(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        if line.strip() == '---':
            if in_table and table_rows:
                _build_table(doc, table_rows)
                table_rows = []
                in_table = False
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and lines[j].strip().startswith('## '):
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=120, after=0)
                add_bottom_border(p, color="DDDDDD", sz="4")
            i += 1
            continue

        if line.strip().startswith('|'):
            if is_table_separator(line):
                i += 1
                continue
            in_table = True
            table_rows.append(parse_table_row(line))
            i += 1
            continue

        if in_table and table_rows:
            _build_table(doc, table_rows)
            table_rows = []
            in_table = False

        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:].strip()
            main_title, sub_title = title, ''
            if '：' in title:
                main_title, sub_title = title.split('：', 1)
            elif ':' in title:
                main_title, sub_title = title.split(':', 1)
            main_title = main_title.strip()
            sub_title = sub_title.strip()

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=400, after=120, line=360)
            set_keep_with_next(p, True)
            set_keep_together(p, True)
            run = p.add_run(main_title)
            set_run_font(run, font_name=FONT_HEI, size=SZ_DOC_TITLE, bold=True, color=COLOR_TITLE)

            if sub_title:
                p2 = doc.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(p2, before=0, after=360, line=320)
                set_keep_with_next(p2, True)
                set_keep_together(p2, True)
                run2 = p2.add_run(sub_title)
                set_run_font(run2, font_name=FONT_HEI, size=SZ_SECTION, bold=False, color=COLOR_HINT)
            prev_was_heading = True
            i += 1
            continue

        if line.startswith('## '):
            title = line[3:].strip()
            p = doc.add_paragraph()
            # 章标题强制从新页开始，避免被挤在页面底部
            p.paragraph_format.page_break_before = True
            # 减小段前间距，让标题出现在新页顶部附近
            set_paragraph_spacing(p, before=120, after=200, line=320)
            set_keep_with_next(p, True)
            set_keep_together(p, True)
            run = p.add_run(title)
            set_run_font(run, font_name=FONT_HEI, size=SZ_CHAPTER, bold=True, color=COLOR_TITLE)
            prev_was_heading = True
            i += 1
            continue

        if line.startswith('### '):
            title = line[4:].strip()
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=280, after=140, line=300)
            set_keep_with_next(p, True)
            set_keep_together(p, True)
            run = p.add_run(title)
            set_run_font(run, font_name=FONT_HEI, size=SZ_SECTION, bold=True, color=COLOR_TITLE)
            prev_was_heading = True
            i += 1
            continue

        if line.strip().startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                qline = lines[i].strip()
                content = re.sub(r'^>\s*', '', qline)
                quote_lines.append(content)
                i += 1

            quote_text = '\n'.join(quote_lines).strip()
            if not quote_text:
                continue

            hint_type, hint_text = classify_hint(quote_text)

            if hint_type in ('hint', 'example', 'exercise', 'rule', 'key'):
                _add_hint_block(doc, hint_text, hint_type)
            else:
                _add_quote_block(doc, quote_text)
            continue

        if re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=40, after=40, line=280)
            set_keep_together(p, True)
            if i + 1 < len(lines) and re.match(r'^\d+\.\s', lines[i + 1].strip()):
                set_keep_with_next(p, True)
            content = re.sub(r'^\d+\.\s*', '', line)
            num_match = re.match(r'^(\d+)\.\s', line)
            num = num_match.group(1) if num_match else ''
            run_num = p.add_run(f"{num}. ")
            set_run_font(run_num, font_name=FONT_BODY, size=SZ_BODY, bold=True, color=COLOR_BODY)
            parse_inline_text(p, content, font_name=FONT_BODY, size=SZ_BODY, bold=False, color=COLOR_BODY)
            prev_was_heading = False
            i += 1
            continue

        if line.strip().startswith('- ') or line.strip().startswith('- [ ]'):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=30, after=30, line=270)
            set_keep_together(p, True)
            if i + 1 < len(lines) and re.match(r'^(\-\s|\-\s\[\s\])', lines[i + 1].strip()):
                set_keep_with_next(p, True)

            if '- [ ]' in line:
                content = line.replace('- [ ]', '').strip()
                run_box = p.add_run("☐ ")
                set_run_font(run_box, font_name=FONT_BODY, size=SZ_BODY, bold=False, color=COLOR_BODY)
                parse_inline_text(p, content, font_name=FONT_BODY, size=SZ_BODY, bold=False, color=COLOR_BODY)
            else:
                content = line.strip()[2:]
                run_bullet = p.add_run("· ")
                set_run_font(run_bullet, font_name=FONT_BODY, size=SZ_BODY, bold=False, color=COLOR_HINT)
                parse_inline_text(p, content, font_name=FONT_BODY, size=SZ_BODY, bold=False, color=COLOR_BODY)
            prev_was_heading = False
            i += 1
            continue

        if line.strip().startswith('**方法') or re.match(r'^\*\*方法', line.strip()):
            title_text = line.strip().strip('*').strip()
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=200, after=80, line=290)
            set_keep_with_next(p, True)
            set_keep_together(p, True)
            run = p.add_run(title_text)
            set_run_font(run, font_name=FONT_HEI, size=SZ_METHOD, bold=True, color=COLOR_TITLE)
            prev_was_heading = True
            i += 1
            continue

        if line.strip().startswith('**') and line.strip().endswith('**') and line.strip().count('**') == 2:
            title_text = line.strip().strip('*').strip()
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=160, after=60, line=280)
            set_keep_with_next(p, True)
            set_keep_together(p, True)
            run = p.add_run(title_text)
            set_run_font(run, font_name=FONT_HEI, size=SZ_SUBHEAD, bold=True, color=COLOR_TITLE)
            prev_was_heading = True
            i += 1
            continue

        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=30, after=80, line=290)
        set_keep_together(p, True)

        # 如果前一段是标题，当前正文段落需要与标题保持同页
        if prev_was_heading:
            set_keep_with_next(p, True)
            prev_was_heading = False

        next_is_list = False
        j = i + 1
        while j < len(lines) and not lines[j].strip():
            j += 1
        if j < len(lines):
            next_line = lines[j].strip()
            if re.match(r'^(\d+\.\s|-\s|-\s\[)', next_line):
                next_is_list = True
        if next_is_list:
            set_keep_with_next(p, True)

        parse_inline_text(p, line.strip(), font_name=FONT_BODY, size=SZ_BODY, bold=False, color=COLOR_BODY)
        i += 1

    if in_table and table_rows:
        _build_table(doc, table_rows)

    return doc


def _add_hint_block(doc, text, hint_type):
    label_map = {
        'hint': '提示',
        'example': '示例',
        'exercise': '练习',
        'rule': '规则',
        'key': '关键',
    }

    label_match = re.match(r'\[(.+?)\]', text)
    if label_match:
        label = label_match.group(1)
        content = text[label_match.end():].strip()
    else:
        label = ''
        content = text

    content_lines = content.split('\n')

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=80, after=20, line=260)
    set_keep_with_next(p, True)
    set_keep_together(p, True)
    add_left_border(p, color=COLOR_HINT_BAR, sz="12")

    if label:
        run_label = p.add_run(f"【{label}】")
        set_run_font(run_label, font_name=FONT_HEI, size=SZ_HINT, bold=True, color=COLOR_HINT)

    first_content = content_lines[0].strip() if content_lines else ''
    if first_content:
        if first_content.startswith('>'):
            first_content = first_content.lstrip('>').strip()
        run = p.add_run(first_content)
        set_run_font(run, font_name=FONT_BODY, size=SZ_HINT, bold=False, color=COLOR_HINT)

    for cl in content_lines[1:]:
        cl = cl.strip()
        if not cl:
            continue
        if cl.startswith('>'):
            cl = cl.lstrip('>').strip()
        if not cl:
            continue
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=10, after=20, line=260)
        set_keep_together(p, True)
        add_left_border(p, color=COLOR_HINT_BAR, sz="12")
        parse_inline_text(p, cl, font_name=FONT_BODY, size=SZ_HINT, bold=False, color=COLOR_HINT)

    p_spacer = doc.add_paragraph()
    set_paragraph_spacing(p_spacer, before=0, after=40, line=100)


def _add_quote_block(doc, text):
    lines = text.split('\n')
    for cl in lines:
        cl = cl.strip()
        if cl.startswith('>'):
            cl = cl.lstrip('>').strip()
        if not cl:
            continue
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=40, after=60, line=280)
        add_left_border(p, color="BBBBBB", sz="8")
        parse_inline_text(p, cl, font_name=FONT_BODY, size=SZ_QUOTE, bold=False, color=COLOR_BODY)


def _build_table(doc, rows):
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row in table.rows:
        set_table_row_cant_split(row)

    if table.rows:
        first_cell_para = table.rows[0].cells[0].paragraphs[0]
        set_keep_with_next(first_cell_para, True)

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= num_cols:
                break
            cell = table.rows[r_idx].cells[c_idx]
            cell.paragraphs[0].text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=40, after=40, line=260)

            is_header = (r_idx == 0)
            parse_inline_text(
                p, cell_text.strip(),
                font_name=FONT_BODY,
                size=SZ_TABLE,
                bold=is_header,
                color=COLOR_TITLE if is_header else COLOR_BODY
            )

            if is_header:
                set_cell_shading(cell, COLOR_TABLE_HEADER_BG)
            set_cell_border(cell, color=COLOR_BORDER, sz="4")

    p_spacer = doc.add_paragraph()
    set_paragraph_spacing(p_spacer, before=0, after=60, line=100)


def main():
    parser = argparse.ArgumentParser(description="从 Markdown 生成精排版 DOCX")
    parser.add_argument("input", help="输入 Markdown 文件路径")
    parser.add_argument("-o", "--output", required=True, help="输出 DOCX 文件路径")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"错误：输入文件不存在 {args.input}")
        sys.exit(1)

    print(f"开始生成：{args.input} → {args.output}")
    doc = build_doc(args.input)
    doc.save(args.output)
    print(f"已保存：{args.output} ({os.path.getsize(args.output)} bytes)")


if __name__ == '__main__':
    main()
