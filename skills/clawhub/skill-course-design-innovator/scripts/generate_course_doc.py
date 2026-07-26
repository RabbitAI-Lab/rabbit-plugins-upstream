#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
课程方案Word文档生成器
功能：接收结构化JSON数据，生成格式化的课程方案Word文档
"""

import json
import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


def add_heading_with_number(doc, text, level, number=None):
    """
    添加带编号的标题
    """
    if number is not None:
        full_text = f"{number}. {text}"
    else:
        full_text = text
    return doc.add_heading(full_text, level=level)


def add_list_item(doc, text, level=0):
    """
    添加列表项
    """
    para = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        para.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    return para


def set_cell_background(cell, color):
    """
    设置单元格背景色
    """
    from docx.oxml import parse_xml
    shading_elm = parse_xml(f'<w:shd {{{color}}} w:fill="D9D9D9"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)


def format_file_size(size_bytes):
    """
    格式化文件大小
    """
    if size_bytes < 1024:
        return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.2f} KB"
    else:
        return f"{size_bytes / (1024 * 1024):.2f} MB"


def generate_course_doc(input_json=None, output_file=None):
    """
    生成课程方案Word文档

    Args:
        input_json: JSON格式的字符串或字典，包含课程方案数据
        output_file: 输出文件路径，默认为 course-proposal.docx

    Returns:
        dict: 包含文件信息的字典，成功返回{'success': True, 'file_path': str, 'file_size': str}
              失败返回{'success': False, 'error': str}
    """
    # 解析输入数据
    if input_json is None:
        error_msg = "错误：未提供输入数据"
        print(error_msg)
        return {'success': False, 'error': error_msg}

    if isinstance(input_json, str):
        try:
            course_data = json.loads(input_json)
        except json.JSONDecodeError as e:
            error_msg = f"错误：JSON解析失败 - {e}"
            print(error_msg)
            return {'success': False, 'error': error_msg}
    else:
        course_data = input_json

    # 设置输出文件路径（使用绝对路径）
    if output_file is None:
        # 默认输出到工作区根目录
        output_file = "/workspace/projects/course-proposal.docx"
    else:
        # 转换为绝对路径
        output_file = os.path.abspath(output_file)

    # 确保输出目录存在
    output_dir = os.path.dirname(output_file)
    if output_dir and not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir, exist_ok=True)
            print(f"✓ 创建输出目录: {output_dir}")
        except Exception as e:
            error_msg = f"错误：无法创建输出目录 - {e}"
            print(error_msg)
            return {'success': False, 'error': error_msg}

    try:
        # 创建Word文档
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)

        # 添加文档标题
        title = doc.add_heading('课程设计方案', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加副标题信息
        if 'course_subtitle' in course_data:
            subtitle = doc.add_paragraph(course_data.get('course_subtitle', ''))
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].bold = True
            subtitle.runs[0].font.size = Pt(14)

        # 添加元信息表格
        meta_table = doc.add_table(rows=2, cols=4)
        meta_table.style = 'Table Grid'

        meta_data = [
            ('课程名称', course_data.get('course_name', '')),
            ('版本', course_data.get('version', 'v1.0')),
            ('创建日期', course_data.get('created_date', datetime.now().strftime('%Y-%m-%d'))),
            ('课程时长', course_data.get('section_1_background', {}).get('duration', ''))
        ]

        for i, (key, value) in enumerate(meta_data):
            row = 0 if i < 2 else 1
            col = 0 if i % 2 == 0 else 2
            meta_table.cell(row, col).text = key
            meta_table.cell(row, col + 1).text = str(value)

            # 设置表头单元格样式
            meta_table.cell(row, col).paragraphs[0].runs[0].bold = True

        doc.add_paragraph()  # 空行

        # 第一部分：项目背景
        add_heading_with_number(doc, '项目背景', 1, 1)
        background = course_data.get('section_1_background', {})

        doc.add_heading('课程背景', 2)
        doc.add_paragraph(background.get('course_background', ''))

        doc.add_heading('目标受众', 2)
        doc.add_paragraph(background.get('target_audience', ''))

        doc.add_heading('课程主题', 2)
        doc.add_paragraph(background.get('course_theme', ''))

        doc.add_heading('期望目标', 2)
        expected_goals = background.get('expected_goals', {})

        # 业务目标
        doc.add_heading('业务目标', 3)
        doc.add_paragraph(expected_goals.get('business_objectives', ''))

        # 学习目标
        doc.add_heading('学习目标', 3)
        learning_obj = expected_goals.get('learning_objectives', [])
        if learning_obj:
            for obj in learning_obj:
                add_list_item(doc, obj)
        else:
            doc.add_paragraph('无')

        # 能力提升目标
        doc.add_heading('能力提升目标', 3)
        competency_targets = expected_goals.get('competency_targets', [])
        if competency_targets:
            for target in competency_targets:
                add_list_item(doc, target)
        else:
            doc.add_paragraph('无')

        # 第二部分：核心创新点
        add_heading_with_number(doc, '核心创新点', 1, 2)
        innovation = course_data.get('section_2_innovation', {})

        innovation_points = innovation.get('innovation_points', [])
        if innovation_points:
            for i, point in enumerate(innovation_points, 1):
                doc.add_heading(f'创新点{i}', 2)
                doc.add_paragraph(point)
        else:
            doc.add_paragraph('无')

        # 第三部分：课程内容（表格形式）
        add_heading_with_number(doc, '课程内容', 1, 3)
        course_content = course_data.get('section_3_course_content', [])

        if not course_content:
            doc.add_paragraph('未提供课程内容信息')
        else:
            # 创建课程内容表格
            # 列：模块名称、章节名称、章节时长、大纲内容、教学方法、活动设计、输出成果
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'

            # 设置表头
            header_cells = table.rows[0].cells
            headers = ['模块名称', '章节名称', '章节时长', '大纲内容（细化到2级）', '教学方法', '活动设计（编注序号）', '输出成果（编注序号）']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True

            # 填充每个模块的每个章节内容
            for module in course_content:
                module_title = module.get('module_title', '')
                chapters = module.get('chapters', [])

                for chapter in chapters:
                    row_cells = table.add_row().cells

                    # 模块名称
                    row_cells[0].text = module_title

                    # 章节名称
                    row_cells[1].text = chapter.get('chapter_title', '')

                    # 章节时长
                    row_cells[2].text = chapter.get('duration', '')

                    # 大纲内容（细化到2级）
                    content_outline = chapter.get('content_outline', {})
                    level1_list = content_outline.get('level1', [])
                    outline_text = ''
                    for item in level1_list:
                        title = item.get('title', '')
                        level2_list = item.get('level2', [])
                        outline_text += f'{title}\n'
                        for level2 in level2_list:
                            outline_text += f'  {level2}\n'
                    row_cells[3].text = outline_text.strip()

                    # 教学方法
                    teaching_methods = chapter.get('teaching_methods', [])
                    row_cells[4].text = '\n'.join(teaching_methods)

                    # 活动设计（编注序号）
                    activity_design = chapter.get('activity_design', [])
                    row_cells[5].text = '\n'.join(activity_design)

                    # 输出成果（编注序号）
                    output_deliverables = chapter.get('output_deliverables', [])
                    row_cells[6].text = '\n'.join(output_deliverables)

            # 设置表格列宽
            for row in table.rows:
                row.cells[0].width = Cm(3.0)   # 模块名称
                row.cells[1].width = Cm(3.5)   # 章节名称
                row.cells[2].width = Cm(2.5)   # 章节时长
                row.cells[3].width = Cm(4.5)   # 大纲内容
                row.cells[4].width = Cm(3.0)   # 教学方法
                row.cells[5].width = Cm(3.5)   # 活动设计
                row.cells[6].width = Cm(3.5)   # 输出成果

        doc.add_paragraph()  # 空行

        # 第四部分：课程评估方式建议
        add_heading_with_number(doc, '课程评估方式建议', 1, 4)
        evaluation = course_data.get('section_4_evaluation', {})

        doc.add_heading('评估方式建议', 2)
        doc.add_paragraph(evaluation.get('evaluation_methods', ''))

        doc.add_heading('评估维度', 2)
        evaluation_dimensions = evaluation.get('evaluation_dimensions', [])
        if evaluation_dimensions:
            for dimension in evaluation_dimensions:
                add_list_item(doc, dimension)
        else:
            doc.add_paragraph('无')

        # 添加页脚
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"课程设计方案 - {course_data.get('course_name', '')} - 生成日期：{datetime.now().strftime('%Y-%m-%d')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 保存文档
        doc.save(output_file)

        # 验证文件是否成功创建
        if not os.path.exists(output_file):
            error_msg = f"错误：文档保存后未找到文件 - {output_file}"
            print(error_msg)
            return {'success': False, 'error': error_msg}

        # 获取文件信息
        file_size = os.path.getsize(output_file)
        file_size_formatted = format_file_size(file_size)
        creation_time = datetime.fromtimestamp(os.path.getctime(output_file)).strftime('%Y-%m-%d %H:%M:%S')

        # 输出成功信息
        print("=" * 60)
        print("✓ 课程方案文档已成功生成")
        print("=" * 60)
        print(f"文件路径: {output_file}")
        print(f"文件大小: {file_size_formatted}")
        print(f"创建时间: {creation_time}")
        print("=" * 60)
        print("提示: 请使用文件路径下载文档")
        print("=" * 60)

        # 返回文件信息
        return {
            'success': True,
            'file_path': output_file,
            'file_size': file_size,
            'file_size_formatted': file_size_formatted,
            'creation_time': creation_time
        }

    except Exception as e:
        error_msg = f"错误：文档生成失败 - {e}"
        print(error_msg)
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': error_msg}


def main():
    """
    主函数：从命令行参数或标准输入读取JSON数据
    """
    import argparse

    parser = argparse.ArgumentParser(description='生成课程方案Word文档')
    parser.add_argument('--input', '-i', type=str,
                       help='输入JSON文件路径')
    parser.add_argument('--output', '-o', type=str,
                       help='输出Word文件路径，默认为/workspace/projects/course-proposal.docx')

    args = parser.parse_args()

    # 读取输入数据
    if args.input:
        try:
            with open(args.input, 'r', encoding='utf-8') as f:
                input_data = f.read()
        except Exception as e:
            print(f"✓ 无法读取输入文件：{e}")
            return 1
    else:
        # 从标准输入读取
        print("请输入课程方案的JSON数据（按Ctrl+D结束输入）：")
        input_data = sys.stdin.read()

    # 生成文档
    result = generate_course_doc(input_data, args.output)

    if result['success']:
        print(f"\n文档生成成功！文件位于: {result['file_path']}")
        return 0
    else:
        print(f"\n文档生成失败: {result.get('error', '未知错误')}")
        return 1


if __name__ == '__main__':
    sys.exit(main())
