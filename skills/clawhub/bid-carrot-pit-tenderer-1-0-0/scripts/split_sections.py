#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
split_sections.py — 招标文件分章节切分（招采发标合规自检专家·可选加速器）

把一份长招标文件按中文层级标题切成结构化章节，帮助 LLM 精准定位
合规风险条款所在章节（第X章第Y条），避免超长上下文丢失条款。

本脚本与「招采萝卜坑识别专家」完全通用——盲区标记（★/偏离/付款/附件）
对招标人自检同样关键（这些区域最易被投标人异议）。

用法：
  python split_sections.py <招标文件.docx|.txt> [--out sections.json]

输出 sections.json:
  [{"id": "S01", "heading": "第一章 投标人须知", "level": 1, "text": "...",
    "flags": {"hard_marker": false, "appendix": false, "payment": false}}, ...]
  # flags 标记该章节是否含 ★/偏离/废标、附件/附录/图纸、付款/保证金 —— 配合 §2.2 防漏扫回溯

只做机械切分，不判内容；内容判断交给 SKILL.md §3 信号标尺 + LLM。
脚本不可用/解析失败时，启用 SKILL.md §2.2 无脚本降级模式（手动按标题切）。
"""

import sys
import os
import re
import json
import argparse


# ---- 标题识别：覆盖中文招标文件的常见层级写法 ----
# 关键规则：标题必须后接「≥2 个汉字」的标题短语，避免把「1.1 / 4.2」这类
# 条款号误判成标题（它们在招标文档里是子条款，不是章节断点）。
# 真正的结构锚点 = 第X章 / 一、 / N. 中文标题。
HEADING_PATTERNS = [
    (1, re.compile(r'^\s*第[一二三四五六七八九十百千0-9]+[章节目编部分]')),        # 第一章 / 第二节 / 第三部分
    (1, re.compile(r'^\s*[0-9]+[\.、]\s*[一-龥]{2,}')),                            # 1. 投标人须知 / 2、资格条件
    (2, re.compile(r'^\s*[一二三四五六七八九十百千]+、[一-龥]{2,}')),               # 一、项目概况
]

# ---- 盲区标记（防漏扫核心，对应 SKILL.md §2.2 三大盲区回溯）----
# 这些区域是合规风险高发且 LLM 易漏扫的死角，切分时直接打标，
# 让报告概览能机械驱动"已读/未读"红字警告，而非依赖 LLM 记忆。
RISK_MARKERS = re.compile(r'[★▲]|实质性响应|不可偏离|废标|否决|无效投标|关键条款')  # 盲区1：★/偏离/否决
APPENDIX_MARKERS = re.compile(r'附件|附录|图纸|附[图表明单]')                        # 盲区3：附件/附录/图纸
PAYMENT_MARKERS = re.compile(r'付款|保证金|履约担保|预付款|违约金|垫资')            # 盲区2：付款/保证金


def detect_heading(line):
    """返回 (level, heading_text) 或 None。"""
    s = line.strip()
    if not s:
        return None
    # 标题不应过长（招标标题一般 < 40 字），过长视为正文
    if len(s) > 40:
        return None
    for level, pat in HEADING_PATTERNS:
        if pat.match(s):
            return level, s
    return None


def read_text(path):
    """读 docx / txt 为纯文本（按行）。"""
    ext = os.path.splitext(path)[1].lower()
    if ext == '.txt':
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read().splitlines()
    if ext in ('.docx',):
        try:
            from docx import Document
        except ImportError:
            print("⚠️ 未安装 python-docx，无法解析 .docx；请转 .txt 后重试，或启用降级模式。", file=sys.stderr)
            sys.exit(2)
        doc = Document(path)
        lines = []
        for p in doc.paragraphs:
            lines.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                lines.append(' | '.join(cells))
        return lines
    raise ValueError(f"不支持的文件类型：{ext}（仅支持 .txt / .docx）")


def split_sections(path):
    lines = read_text(path)
    sections = []
    cur = None
    buf = []

    def flush():
        nonlocal cur, buf
        if cur is not None:
            text = '\n'.join(buf).strip()
            cur['text'] = text
            if text:
                # 盲区标记：让 LLM 与报告概览能机械确认"这些死角已扫/未扫"
                cur['flags'] = {
                    'hard_marker': bool(RISK_MARKERS.search(text)),  # ★/偏离/否决
                    'appendix': bool(APPENDIX_MARKERS.search(text)),  # 附件/附录/图纸
                    'payment': bool(PAYMENT_MARKERS.search(text)),    # 付款/保证金
                }
                sections.append(cur)
        cur = None
        buf = []

    for raw in lines:
        h = detect_heading(raw)
        if h is not None:
            level, heading = h
            flush()
            cur = {'id': '', 'heading': heading, 'level': level, 'text': ''}
            buf = []
        else:
            if cur is None:
                # 标题前的导语/正文，归入一个未命名节
                if not sections or sections[-1]['heading'] != '__preamble__':
                    cur = {'id': '', 'heading': '__preamble__', 'level': 0, 'text': ''}
                    buf = []
            buf.append(raw)

    flush()
    # 编号 + 过滤空节
    idx = 0
    out = []
    for s in sections:
        if not s['text']:
            continue
        idx += 1
        s['id'] = f'S{idx:02d}'
        out.append(s)
    return out


def main():
    ap = argparse.ArgumentParser(description='招标文件分章节切分')
    ap.add_argument('input', help='招标文件 .txt / .docx')
    ap.add_argument('--out', default='sections.json', help='输出 JSON 路径（默认 sections.json）')
    args = ap.parse_args()

    secs = split_sections(args.input)
    with open(args.out, 'w', encoding='utf-8') as f:
        json.dump(secs, f, ensure_ascii=False, indent=2)

    total = sum(len(s['text']) for s in secs)
    hm = sum(1 for s in secs if s['flags'].get('hard_marker'))
    ap_ = sum(1 for s in secs if s['flags'].get('appendix'))
    pm = sum(1 for s in secs if s['flags'].get('payment'))
    print(f"分章节完成 → {args.out}：{len(secs)} 个章节，正文约 {total} 字")
    print(f"盲区标记 → 高危条款(★/偏离/废标):{hm}  附件/附录/图纸:{ap_}  付款/保证金:{pm}")
    for s in secs:
        tag = ''
        f = s['flags']
        if f.get('hard_marker'):
            tag += ' ★'
        if f.get('appendix'):
            tag += ' 附'
        if f.get('payment'):
            tag += ' 款'
        print(f"  {s['id']} [L{s['level']}] {s['heading'][:28]}{tag}  ({len(s['text'])}字)")


if __name__ == '__main__':
    main()
