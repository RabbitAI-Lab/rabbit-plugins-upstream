#!/usr/bin/env python3
"""Extract structured clauses/tables from tender documents for version diff.

Outputs JSON: { "documents": [ { "meta": {...}, "clauses": [...], "tables": [...] } ] }
Clause segmentation is heuristic (numbered/heading lines start a new clause).
Each clause carries:
  - id: hierarchical clause id (e.g. "一.1", "表2-R3")
  - heading: short heading text (or the id if heading too long)
  - text: clause body (heading line included as first line)
  - section: current top-level section context (chapter/一./二.)
  - page: 1-based page number if known (PDF only), else None
  - context: human-readable location string (page + section) for report anchoring

Tables (DOCX) are flattened into pseudo-clauses so their content enters the
diff pipeline instead of being silently dropped.
"""
import argparse
import json
import os
import re
import sys

try:
    from docx import Document
except Exception:
    Document = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    from docx.oxml.ns import qn
except Exception:
    qn = None

# Markdown heading markers (e.g. "## ") that IMA / exports may prefix lines with.
MD_HEAD = re.compile(r'^#{1,6}\s*')
# 第一章 / 第二节 / 第三条
CHAPTER_RE = re.compile(r'^第[一二三四五六七八九十百零0-9]+[章节目]')
# 一. 二、 三、 (major divisions)
MAJOR_RE = re.compile(r'^[一二三四五六七八九十百]+[.、]')
# Dotted clause numbers (5.1, 5.1.10, 5.2.1) — the dominant structure in
# tender specs. Followed by any non-digit (incl. Chinese text glued with no
# space, e.g. "5.1.10文件系统") or end of line.
NUM_DOTTED = re.compile(r'^\s*(\d+\.\d+(?:\.\d+)*)(?=[^\d]|$)')
# Plain single-integer numbering (1. / 2、 / 3）) — must be followed by a
# separator, whitespace, or the ★/# mandatory marker, so bare quantities like
# "2台" / "500米" are NOT mistaken for clause numbers.
NUM_PLAIN = re.compile(r'^\s*(\d+)(?=[.、)）\s★#])')
# （1） (1)
PAREN_RE = re.compile(r'^[（(]\d+[)）]')
# 1） 2）  (half-width open, full-width close)
SUB_RE = re.compile(r'^\d+[)）]')
# Word/property style names that denote a heading (e.g. "Heading 1", "标题 2").
# Prose-style bids / 参选文件 use Heading styles instead of numeric clause
# numbering; treating them as clause boundaries lets the diff pipeline run on
# documents that have no "5.2.1"-style numbering.
HEADING_RE = re.compile(r'(?:heading|标题)\s*\d+\s*$', re.I)


def _is_heading(style):
    return bool(style) and bool(HEADING_RE.search(str(style).strip()))


def _is_top_heading(style):
    return bool(style) and bool(re.search(r'(?:heading|标题)\s*1\s*$', str(style).strip(), re.I))
SHORT_MAX = 40

# Unicode subscript / superscript translation for run-level vertAlign normalization.
# Word stores many subscripts (e.g. CO₂) as a plain "2" with w:vertAlign=subscript,
# so python-docx's run.text yields "CO2" and the formatting change is invisible to a
# text diff. Mapping subscript/superscript runs to their Unicode forms surfaces these
# format-only changes (e.g. CO2 -> CO₂) to the diff pipeline.
_SUBSCRIPT = str.maketrans("0123456789+-=()", "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎")
_SUPERSCRIPT = str.maketrans("0123456789+-=()", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾")


def _run_is_script(run, kind):
    """True if a run is subscript/superscript. Checks both the high-level font flag
    and the raw w:vertAlign element (font.subscript can be None when inherited)."""
    flag = run.font.subscript if kind == 'sub' else run.font.superscript
    if flag:
        return True
    if qn is None:
        return False
    rPr = run._element.rPr
    if rPr is None:
        return False
    va = rPr.find(qn('w:vertAlign'))
    if va is None:
        return False
    val = va.get(qn('w:val'))
    return val in ('subscript', 'sub') if kind == 'sub' else val in ('superscript', 'super')


def _runs_text(runs):
    """Join run texts, mapping subscript/superscript runs to Unicode forms so that
    format-only script changes become visible to the diff."""
    parts = []
    for r in runs:
        t = r.text or ""
        if not t:
            continue
        if _run_is_script(r, 'sub'):
            t = t.translate(_SUBSCRIPT)
        elif _run_is_script(r, 'sup'):
            t = t.translate(_SUPERSCRIPT)
        parts.append(t)
    return "".join(parts)


def _clause_parts(text):
    """Return (cleaned_text, id_fragment). id_fragment is not None when the
    line starts a new clause."""
    t = MD_HEAD.sub('', text).strip()
    if not t:
        return t, None
    if CHAPTER_RE.match(t):
        return t, t
    m = MAJOR_RE.match(t)
    if m:
        return t, m.group(0).rstrip('.、')
    m = NUM_DOTTED.match(t)
    if m:
        return t, m.group(1)
    m = PAREN_RE.match(t)
    if m:
        return t, m.group(0)
    m = SUB_RE.match(t)
    if m:
        return t, m.group(0)
    m = NUM_PLAIN.match(t)
    if m:
        return t, m.group(1)
    return t, None


def _loc(clause):
    """Build a human-readable location string for report anchoring."""
    parts = []
    if clause.get('page') is not None:
        parts.append(f"第{clause['page']}页")
    if clause.get('section'):
        parts.append(clause['section'])
    return ' · '.join(parts)


def segment(items):
    """Split items into clause dicts. A clause starts at a numbered/section
    line OR a Word Heading-style paragraph (prose bids / 参选文件 that use
    Heading 1-4 instead of "5.2.1" numbering). Following lines are its body
    until the next start.

    `items` are (text, page) or (text, page, style) tuples; `style` is the
    optional paragraph style name (e.g. "Heading 2"). Section-level headings
    (chapter / 一. / 二、 / Heading 1) reset the id context so that
    numbering which restarts per chapter yields unique hierarchical ids such as
    '一.1', '二.1'. The heading line is kept as the first body line so its
    content is never lost during diffing. `section` tracks the current
    top-level context for location anchors."""
    clauses = []
    cur = None
    buf = []
    section = None
    for item in items:
        if len(item) == 3:
            raw, page, style = item
        else:
            raw, page = item
            style = None
        text = raw.strip()
        if not text:
            continue
        clean, frag = _clause_parts(text)
        is_head = _is_heading(style)
        if frag is not None or is_head:
            # close current clause
            if cur is not None:
                cur['text'] = '\n'.join(buf).strip()
                clauses.append(cur)
            if frag is not None:
                # numbered / section line takes priority
                if CHAPTER_RE.match(clean) or MAJOR_RE.match(clean):
                    section = frag
                    cid = frag
                else:
                    cid = f'{section}.{frag}' if section else frag
            else:
                # Heading-style paragraph: use its text as id; a top-level
                # heading (Heading 1) resets the section context.
                hid = clean[:SHORT_MAX]
                if _is_top_heading(style):
                    section = hid
                cid = hid
            cur = {'id': cid,
                   'heading': clean if len(clean) <= SHORT_MAX else cid,
                   'text': '', 'section': section, 'page': page}
            buf = [clean]
        else:
            if cur is None:
                cid = f'p{len(clauses)+1}'
                cur = {'id': cid,
                       'heading': clean if len(clean) <= SHORT_MAX else cid,
                       'text': '', 'section': section, 'page': page}
                buf = [clean]
            else:
                buf.append(clean)
    if cur is not None:
        cur['text'] = '\n'.join(buf).strip()
        clauses.append(cur)
    return [c for c in clauses if c['text']]


def flatten_table(ti, rows, page=None):
    """Flatten a table (DOCX or PDF) into pseudo-clauses so its cells enter the
    diff pipeline. Header row becomes its own clause; each data row becomes a
    clause with column=value pairs for easy cell-level diffing. `page` carries
    the 1-based page number (PDF) for report anchoring; None for DOCX."""
    if not rows:
        return []
    header = rows[0]
    out = [{'id': f'表{ti+1}-表头', 'heading': f'表{ti+1} 表头',
            'text': ' | '.join(header), 'section': f'表{ti+1}', 'page': page}]
    for ri, row in enumerate(rows[1:], 1):
        cells = []
        for ci, cell in enumerate(row):
            col = header[ci] if ci < len(header) else f'列{ci+1}'
            cells.append(f'{col}={cell}')
        out.append({'id': f'表{ti+1}-R{ri}', 'heading': f'表{ti+1} 行{ri}',
                    'text': ' | '.join(cells), 'section': f'表{ti+1}', 'page': page})
    return out


def _word_in_tables(w, tbboxes):
    """True if a pdfplumber word's center falls inside any detected table bbox."""
    cx = (w['x0'] + w['x1']) / 2
    cy = (w['top'] + w['bottom']) / 2
    for (tx0, ttop, tx1, tbottom) in tbboxes:
        if tx0 <= cx <= tx1 and ttop <= cy <= tbottom:
            return True
    return False


def _words_to_lines(words, y_tol=3):
    """Group pdfplumber words into reading-order lines. Words sharing a row
    (top within y_tol) are joined left-to-right; a horizontal gap > 2pt adds a
    space so English words stay separated while CJK runs stay tight."""
    if not words:
        return []
    words = sorted(words, key=lambda w: (round(w['top'] / y_tol), w['x0']))
    lines, cur, cur_top = [], [], None
    for w in words:
        if cur_top is None or abs(w['top'] - cur_top) <= y_tol:
            cur.append(w)
            cur_top = w['top'] if cur_top is None else (cur_top + w['top']) / 2
        else:
            lines.append(cur)
            cur, cur_top = [w], w['top']
    if cur:
        lines.append(cur)
    out = []
    for ln in lines:
        ln = sorted(ln, key=lambda w: w['x0'])
        txt, prev = '', None
        for w in ln:
            if prev is None:
                txt = w['text']
            else:
                gap = w['x0'] - prev['x1']
                txt += (' ' + w['text']) if gap > 2 else w['text']
            prev = w
        if txt.strip():
            out.append(txt.strip())
    return out


def _finalize(clauses):
    for c in clauses:
        c['context'] = _loc(c) or c['heading']
    return clauses


def _empty_doc(path, doc_type, reason):
    """降级产物：文件不可解析时返回零条款、带 data_gap 标记的友好结构，
    不抛出原始异常。下游阶段④/⑤ 据此在报告中显式标注「数据缺口」。"""
    return {
        'meta': {'file': os.path.basename(path), 'type': doc_type,
                 'paras': 0, 'tables': 0, 'parse_status': 'failed',
                 'parse_error': reason},
        'clauses': [],
        'tables': [],
        'data_gap': f"文件无法解析（{reason}）。请确认文件未损坏、未加密、非纯图片扫描件；"
                    f"若为扫描件请改用可编辑 docx 或手动粘贴关键条款。",
    }


def extract_docx(path):
    if Document is None:
        raise RuntimeError('python-docx not installed')
    try:
        doc = Document(path)
    except Exception as e:  # 损坏/加密/非 docx 文件
        return _empty_doc(path, 'docx', type(e).__name__)
    paras = [(_runs_text(p.runs), None, p.style.name)
              for p in doc.paragraphs if p.runs]
    clauses = segment(paras)
    tables = []
    for ti, t in enumerate(doc.tables):
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        tables.append({'index': ti, 'rows': rows})
        clauses.extend(flatten_table(ti, rows))
    return {
        'meta': {'file': os.path.basename(path), 'type': 'docx',
                 'paras': len(paras), 'tables': len(tables),
                 'parse_status': 'ok'},
        'clauses': _finalize(clauses),
        'tables': tables,
    }


def extract_pdf(path):
    """Extract a PDF. Uses pdfplumber when available: text lines are rebuilt
    from words with table-region words excluded (so table cells are NOT
    double-counted in the free text), and detected tables are flattened into
    pseudo-clauses carrying their page number for report anchoring. Falls back
    to pypdf text-only (tables dropped) when pdfplumber is unavailable.

    任何解析异常（损坏/加密/非 PDF/纯图片扫描件）均降级为带 data_gap 的友好
    产物，不抛出原始异常。"""
    def _from_pypdf():
        if PdfReader is None:
            return None
        try:
            reader = PdfReader(path)
            paras = []
            for pi, page in enumerate(reader.pages, 1):
                txt = page.extract_text() or ''
                paras.extend((line.strip(), pi) for line in txt.split('\n') if line.strip())
            if not paras:
                return None  # 无文本层（疑似扫描件）
            return segment(paras), 0
        except Exception:
            return None

    if pdfplumber is not None:
        try:
            paras = []
            table_clauses = []
            table_idx = 0
            table_count = 0
            with pdfplumber.open(path) as pdf:
                for pi, page in enumerate(pdf.pages, 1):
                    tbboxes = [t.bbox for t in page.find_tables()]
                    words = page.extract_words()
                    kept = [w for w in words if not _word_in_tables(w, tbboxes)]
                    for ln in _words_to_lines(kept):
                        paras.append((ln, pi))
                    for trows in page.extract_tables():
                        rows = [[(c or '').strip() for c in r] for r in trows]
                        table_clauses.extend(flatten_table(table_idx, rows, page=pi))
                        table_idx += 1
                        table_count += 1
            if not paras and not table_clauses:
                # 无可提取文本/表格，疑似纯图片扫描件
                return _empty_doc(path, 'pdf', 'no extractable text layer (likely scanned image)')
            text_clauses = segment(paras)
            clauses = text_clauses + table_clauses
            return {
                'meta': {'file': os.path.basename(path), 'type': 'pdf',
                         'paras': len(paras), 'tables': table_count,
                         'parse_status': 'ok'},
                'clauses': _finalize(clauses),
                'tables': [{'index': i, 'rows': []} for i in range(table_count)],
            }
        except Exception as e:
            # pdfplumber 解析失败（损坏/加密/非 PDF），退回 pypdf
            res = _from_pypdf()
            if res is None:
                return _empty_doc(path, 'pdf', type(e).__name__)
            clauses, table_count = res
            return {
                'meta': {'file': os.path.basename(path), 'type': 'pdf',
                         'paras': len(clauses), 'tables': table_count,
                         'parse_status': 'ok-fallback-pypdf'},
                'clauses': _finalize(clauses),
                'tables': [],
            }
    # Fallback: pypdf text only (tables silently dropped — see Limitations).
    res = _from_pypdf()
    if res is None:
        return _empty_doc(path, 'pdf',
                          'pypdf produced no text (corrupted or scanned image)')
    clauses, table_count = res
    return {
        'meta': {'file': os.path.basename(path), 'type': 'pdf',
                 'paras': len(clauses), 'tables': table_count,
                 'parse_status': 'ok-fallback-pypdf'},
        'clauses': _finalize(clauses),
        'tables': [],
    }


def extract_txt(path):
    with open(path, encoding='utf-8', errors='ignore') as fh:
        paras = [(line.strip(), None) for line in fh if line.strip()]
    clauses = segment(paras)
    return {
        'meta': {'file': os.path.basename(path), 'type': 'txt',
                 'paras': len(paras), 'tables': 0},
        'clauses': _finalize(clauses),
        'tables': [],
    }


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--files', nargs='+', required=True)
    ap.add_argument('--out', required=True)
    args = ap.parse_args()
    docs = []
    for f in args.files:
        low = f.lower()
        if low.endswith('.docx'):
            docs.append(extract_docx(f))
        elif low.endswith('.pdf'):
            docs.append(extract_pdf(f))
        elif low.endswith(('.txt', '.md', '.markdown')):
            docs.append(extract_txt(f))
        else:
            # Unknown extension: many exports (e.g. IMA markdown) are plain
            # text. Fall back to txt extraction instead of dropping the doc.
            docs.append(extract_txt(f))
    # 数据质量检查：文件不可解析则终止并给出友好提示，不静默产出空结果
    failed = [d for d in docs if d.get('meta', {}).get('parse_status') == 'failed']
    if failed:
        for d in failed:
            sys.stderr.write(
                f"[ERROR] 文件解析失败: {d['meta'].get('file')} — {d['meta'].get('parse_error', '')}\n"
                f"[HINT] {d.get('data_gap', '请检查文件后重试')}\n"
            )
        sys.exit(2)
    with open(args.out, 'w', encoding='utf-8') as fh:
        json.dump({'documents': docs}, fh, ensure_ascii=False, indent=2)
    total = sum(len(d['clauses']) for d in docs)
    print(f'Extracted {len(docs)} doc(s), {total} clauses -> {args.out}')


if __name__ == '__main__':
    main()
