#!/usr/bin/env python3
"""Render findings into Markdown + DOCX report with role-based templates.

Role-based rendering:
  --role bidder   → 投标人版：红线→投标策略汇总→报价影响→时限权利→逐条明细
  --role tenderer → 招标人版（待完善）：合规自检→称谓一致性→质疑风险预判→逐条明细

Each finding item may carry: clause_id, change_type, old_text, new_text,
numeric_delta, dimension, sentiment, severity, impact, basis, basis_source,
action, confidence, context, compliance_check, timeliness_warning,
final_severity, bid_impact (bidder), implicit_barrier (bidder),
objection_eligible (bidder), pricing_note/pricing_risk/pricing_action (bidder).
"""
import argparse
import json
from collections import Counter

try:
    from docx import Document
except Exception:
    Document = None

SENT = {'风险': '[风险]', '中性': '[中性]', '利好': '[利好]'}
BID_IMPACT_LABELS = {
    '报价需调整': '💰',
    '技术方案需改': '🔧',
    '资格/业绩需补': '📋',
    '时限权利受影响': '⏰',
    '仅关注不需动作': '👁',
    '无影响': '✅',
}


def load(path):
    with open(path, encoding='utf-8') as fh:
        data = json.load(fh)
    if isinstance(data, list):
        return data, {}
    return data.get('items', []), data.get('meta', {})


def _loc(it):
    ctx = it.get('context')
    return f"位置：{ctx}" if ctx else ''


def _severity(it):
    return it.get('final_severity') or it.get('severity', '')


def _flagged(items):
    return [i for i in items
            if i.get('is_redline')
            or i.get('sentiment') == '风险'
            or i.get('severity') == '高'
            or i.get('final_severity') == '高']


# ============================================================
# BIDDER REPORT TEMPLATES
# ============================================================

def _bidder_summary_section(items):
    """投标策略汇总 — 报告首屏核心决策信息。"""
    lines = []
    c = Counter(i.get('sentiment', '中性') for i in items)
    redline = sum(1 for i in items if i.get('is_redline'))
    ibarrier = sum(1 for i in items if i.get('implicit_barrier'))
    tw = any(i.get('timeliness_warning') for i in items)
    obj_ok = sum(1 for i in items if i.get('objection_eligible'))
    pricing_items = [i for i in items if i.get('bid_impact') == '报价需调整']
    tech_items = [i for i in items if i.get('bid_impact') == '技术方案需改']
    qual_items = [i for i in items if i.get('bid_impact') == '资格/业绩需补']

    lines.append('## 影响概览')
    for k in ['风险', '中性', '利好']:
        if c.get(k):
            lines.append(f"- {SENT.get(k, k)}：{c[k]} 条")
    if redline:
        lines.append(f"- 🔴 **[红线]**：{redline} 条（强制合规风险，建议优先处理）")
    if ibarrier:
        lines.append(f"- ⚠️ **[隐性门槛]**：{ibarrier} 条（准入范围可能收窄）")
    lines.append('')

    # 投标行动清单
    actions = []
    if redline or obj_ok:
        actions.append("存在合规红线或可质疑项 —— 评估是否提交书面异议/质疑")
    if pricing_items:
        actions.append(f"**报价需重新测算**（{len(pricing_items)} 条变更影响成本/收入项）")
    if tech_items:
        actions.append(f"**技术方案需调整**（{len(tech_items)} 条变更涉及技术响应内容）")
    if qual_items:
        actions.append(f"**资格/业绩需核查**（{len(qual_items)} 条涉及准入条件变化）")
    if tw:
        actions.append("**时限权利** —— 变更距截止<15日，有权要求顺延")

    if actions:
        lines.append('## 📋 投标行动清单')
        for idx, a in enumerate(actions, 1):
            lines.append(f"{idx}. {a}")
        lines.append('')

    # 报价影响专段
    if pricing_items:
        lines.append('## 💰 报价影响评估')
        for it in pricing_items:
            note = it.get('pricing_note') or it.get('impact', '')
            risk = it.get('pricing_risk', '')
            paction = it.get('pricing_action') or it.get('action', '')
            risk_tag = f"（风险等级：{risk}）" if risk else ''
            lines.append(f"- **{it.get('clause_id','')}**：{note} {risk_tag}")
            if paction:
                lines.append(f"  → {paction}")
        lines.append('')

    # 时限权利专段
    if tw:
        lines.append('## ⏰ 投标人时限权利')
        lines.append('- **顺延权**：本次变更可能影响投标文件编制且距截止不足15日，依法你有权要求顺延投标截止时间。')
        lines.append('- **依据**：87号令第二十七条 / 招标投标法实施条例第二十一条')
        lines.append('- **行动**：立即致函招标人确认新截止时间；保留要求顺延的证据。')
        lines.append('')

    return '\n'.join(lines)


def _bidder_detail(idx, it):
    """投标人版单条详情模板。"""
    out = []
    loc = _loc(it)
    tag = f"{it.get('dimension','')}/{it.get('sentiment','')}"
    if it.get('is_redline'):
        tag += " 🔴"
    if it.get('implicit_barrier'):
        tag += " ⚠️"

    bi = it.get('bid_impact', '')
    bi_icon = BID_IMPACT_LABELS.get(bi, '')
    head = f"### {idx}. [{tag}] {it.get('clause_id','')}"
    if bi_icon:
        head += f" {bi_icon}"
    if loc:
        head += f"\n> {loc}"
    out.append(head)

    # Meta line
    meta_parts = [
        f"变更类型：{it.get('change_type','')}",
        f"严重度：{_severity(it)}",
        f"置信度：{it.get('confidence','')}",
    ]
    if bi:
        meta_parts.insert(2, f"投标影响：{bi}")
    out.append(f"- {'　'.join(meta_parts)}")

    if it.get('numeric_delta'):
        out.append(f"- 数值变更：{it['numeric_delta']}")
    if it.get('old_text'):
        out.append(f"- 原版：{it['old_text'][:500]}")
    if it.get('new_text'):
        out.append(f"- 新版：{it['new_text'][:500]}")
    out.append(f"- 影响：{it.get('impact','')}")

    if it.get('pricing_note'):
        out.append(f"- 📊 **报价影响**：{it['pricing_note']}（风险：{it.get('pricing_risk','N/A')}）")
        if it.get('pricing_action'):
            out.append(f"  → {it['pricing_action']}")

    out.append(f"- 依据：{it.get('basis','')}（{it.get('basis_source','')}）")
    if it.get('compliance_check'):
        out.append(f"- 合规核查：{it['compliance_check']}")
    if it.get('timeliness_warning'):
        out.append('- ⏰ **时限预警**：你有权要求顺延投标截止时间（<15日 + 影响编制）')
    if it.get('objection_eligible') and not it.get('is_redline'):
        out.append('- ✉️ 该条存在隐性风险，若认为不合理可在质疑期内提出书面质疑')

    # Action as final actionable item
    action = it.get('action', '')
    if action:
        out.append(f"- **建议操作**：{action}")
    out.append('')
    return '\n'.join(out)


def markdown_bidder(items, meta):
    """投标人版完整 Markdown 报告。"""
    lines = ['# 招标文件版本比对报告 · 投标人版', '']
    lines.append(f"- 原版：{meta.get('old_file','')}（{meta.get('old_date','')}）")
    lines.append(f"- 新版：{meta.get('new_file','')}（{meta.get('new_date','')}）")
    lines.append(f"- 投标截止：{meta.get('bid_deadline','未提供')}")
    lines.append(f"- 差异条数：{len(items)}")
    lines.append('')

    # Section 1: Summary + Action Items
    lines.append(_bidder_summary_section(items))

    # Section 2: Redlines / High Risk first
    lines.append('## 🔴 红线与高风险明细（优先处理）')
    flagged = _flagged(items)
    if not flagged:
        lines.append('（无）\n')
    for idx, it in enumerate(flagged, 1):
        lines.append(_bidder_detail(idx, it))

    # Section 3: All details
    lines.append('## 全部差异明细')
    for idx, it in enumerate(items, 1):
        lines.append(_bidder_detail(idx, it))
    return '\n'.join(lines)


# ============================================================
# TENDERER REPORT TEMPLATE (full)
# ============================================================

SAFETY_ICONS = {'合规安全': '✅', '需关注': '⚠️', '仅格式': '📝'}
COMPETITION_ICONS = {
    '无影响': '✅', '轻微收窄': '🔻', '明显收窄': '⚠️', '可能涉嫌排斥': '🔴'
}
PRIORITY_LABELS = {
    'P0': '🔴 P0 — 立即处理',
    'P1': '🟠 P1 — 本批次处理',
    'P2': '🟡 P2 — 尽快处理',
    'P3': '🟢 P3 — 记录备查',
    'P4': '⬜ P4 — 无需处理',
}
RISK_LEVEL_ICONS = {'🔴 高危': 0, '🟡 中危': 1, '🟢 低危': 2, '✅ 安全': 3}


def _tenderer_build_summary(items):
    """Build global summary dict from findings."""
    sc = Counter(i.get('safety_level', '合规安全') for i in items)
    pc = Counter(i.get('priority', 'P4') for i in items)
    complaint_count = sum(1 for i in items if i.get('is_complaint_risk'))
    competition_narrow = sum(1 for i in items
                             if i.get('competition_impact') in ('明显收窄', '可能涉嫌排斥'))
    term_issues = sum(1 for i in items if i.get('terminology_consistency') in ('有不一致', '需全局核查'))

    # Determine global risk level
    has_p0 = pc.get('P0', 0) > 0
    has_high_complaint_and_competition = any(
        i.get('is_complaint_risk') and i.get('competition_impact') == '可能涉嫌排斥'
        for i in items
    )
    has_medium = any(i.get('is_complaint_risk') for i in items) and not has_p0
    all_safe = all(i.get('safety_level') == '合规安全' or i.get('safety_level') == '仅格式'
                   for i in items)

    if has_p0 or has_high_complaint_and_competition:
        risk = '🔴 高危'
        release_rec = '暂停发布 → 修正问题 → 重走公告流程（必要时顺延截止）'
    elif has_medium:
        risk = '🟡 中危'
        release_rec = '可发布但建议附带说明，准备好答复口径，关注异议期内反馈'
    elif competition_narrow > 0 or term_issues > 0:
        risk = '🟢 低危'
        release_rec = '正常发布，内部记录备查，下次避免同类问题'
    else:
        risk = '✅ 安全'
        release_rec = '直接发布'

    top_concerns = []
    # Collect P0/P1 concerns
    for it in sorted(items, key=lambda x: ['P0','P1','P2','P3','P4'].index(x.get('priority','P4'))):
        p = it.get('priority', 'P4')
        if p in ('P0', 'P1'):
            concern = f"{it.get('clause_id','')}：{it.get('impact','')[:80]}"
            if it.get('complaint_trigger'):
                trigger_brief = it['complaint_trigger'][:60]
                concern += f"（{trigger_brief}）"
            top_concerns.append(concern)
        if len(top_concerns) >= 5:
            break

    # Pre-release action checklist
    pre_release_actions = []
    if pc.get('P0', 0) > 0:
        p0_items = [i for i in items if i.get('priority') == 'P0']
        for it in p0_items[:3]:
            lc = it.get('legal_check', {})
            breaches = lc.get('breached_items', [])
            breach_str = '; '.join(breaches) if breaches else ''
            pre_release_actions.append(
                f"[P0] {it.get('clause_id','')}：{breach_str or it.get('impact','')[:60]}——必须修正后再发"
            )
    if pc.get('P1', 0) > 0:
        p1_items = [i for i in items if i.get('priority') == 'P1']
        for it in p1_items[:3]:
            pre_release_actions.append(
                f"[P1] {it.get('clause_id','')}：{it.get('complaint_trigger','')[:60]}——建议补充说明或修正"
            )
    if term_issues > 0:
        pre_release_actions.append(f"称谓一致性：共 {term_issues} 处称谓不一致或需全局核查——建议全文搜索统一")
    if any(i.get('timeliness_check', {}).get('extension_required') for i in items):
        pre_release_actions.append("时限顺延：变更距截止不足15日且影响编制——建议发出顺延公告")

    return {
        'total': len(items),
        'by_safety': dict(sc),
        'by_priority': dict(pc),
        'complaint_count': complaint_count,
        'competition_narrow': competition_narrow,
        'term_issues': term_issues,
        'risk_level': risk,
        'release_recommendation': release_rec,
        'top_concerns': top_concerns,
        'pre_release_actions': pre_release_actions,
    }


def _tenderer_summary_section(items):
    """招标人版报告首屏——全局风险判定 + 发布决策。"""
    s = _tenderer_build_summary(items)
    lines = []

    lines.append('## 发布决策概览')
    lines.append('')
    lines.append(f'| 指标 | 数值 |')
    lines.append(f'|------|------|')
    lines.append(f'| 差异总数 | **{s['total']}** 条 |')
    lines.append(f"| 全局风险等级 | **{s['risk_level']}** |")
    lines.append(f'| 合规安全 | {s["by_safety"].get("合规安全", 0)} 条 |')
    lines.append(f"| 需关注 | {s['by_safety'].get('需关注', 0)} 条 |")
    lines.append(f"| 仅格式 | {s['by_safety'].get('仅格式', 0)} 条 |")
    lines.append(f"| 质疑风险项 | {s['complaint_count']} 条 |")
    lines.append(f"| 竞争收窄项 | {s['competition_narrow']} 条 |")
    lines.append(f"| 称谓一致性问题 | {s['term_issues']} 处 |")
    lines.append('')

    # Release recommendation box
    lines.append('> ### 📋 发布建议')
    lines.append(f'> ')
    lines.append(f'> **{s['release_recommendation']}**')
    lines.append('')

    # Top concerns
    if s['top_concerns']:
        lines.append('### ⚠️ 重点关注的变更')
        for idx, c in enumerate(s['top_concerns'], 1):
            lines.append(f'{idx}. {c}')
        lines.append('')

    # Priority breakdown
    lines.append('### 📊 优先级分布')
    for p_key in ['P0', 'P1', 'P2', 'P3', 'P4']:
        count = s['by_priority'].get(p_key, 0)
        label = PRIORITY_LABELS.get(p_key, p_key)
        if count:
            lines.append(f"- **{label}**：{count} 条")
    lines.append('')

    # Pre-release checklist
    if s['pre_release_actions']:
        lines.append('## ✅ 发布前处置清单')
        for idx, act in enumerate(s['pre_release_actions'], 1):
            lines.append(f'{idx}. {act}')
        lines.append('')

    # Timeliness section
    timeliness_items = [i for i in items if i.get('timeliness_check')]
    if timeliness_items:
        lines.append('## ⏰ 时限合规检查')
        any_ext_req = any(t.get('extension_required') for t in timeliness_items if isinstance(t, dict))
        if any_ext_req:
            lines.append('- ⚠️ **需要顺延**：存在距截止不足15日 + 影响投标文件编制的变更')
            for ti in timeliness_items:
                if isinstance(ti, dict) and ti.get('extension_required'):
                    note = ti.get('extension_note', '')
                    clause = [i.get('clause_id','?') for i in items
                              if i.get('timeliness_check') is ti][0] if False else ''
                    if note:
                        lines.append(f'  - {note}')
        else:
            lines.append('- ✅ 时限合规：本次变更预计不影响投标人编制周期')
        lines.append('')

    # Consistency issues summary
    consistency_issues = []
    for it in items:
        cc = it.get('consistency_check')
        if isinstance(cc, dict) and cc.get('issues_found'):
            for iss in cc['issues_found']:
                consistency_issues.append(f"**{it.get('clause_id','')}**：{iss}")
    if consistency_issues:
        lines.append('## 🔗 一致性扫描结果')
        for idx, iss in enumerate(consistency_issues, 1):
            lines.append(f'{idx}. {iss}')
        lines.append('')

    return '\n'.join(lines)


def _tenderer_detail(idx, it):
    """招标人版单条详情模板。"""
    out = []
    ctx = _loc(it)

    # Header: safety level icon + clause ID + priority
    sl = it.get('safety_level', '合规安全')
    sl_icon = SAFETY_ICONS.get(sl, '')
    pri = it.get('priority', 'P4')
    pri_label = PRIORITY_LABELS.get(pri, pri)

    cr_icon = '🚨' if it.get('is_complaint_risk') else ''
    comp_icon = COMPETITION_ICONS.get(it.get('competition_impact', '无影响'), '')
    term_status = it.get('terminology_consistency', '一致')
    term_badge = '' if term_status == '一致' else f' 🔀{term_status}'

    head = f"### {idx}. {sl_icon} [{sl}] {it.get('clause_id','')}"
    head += f" {cr_icon}{comp_icon}{term_badge}"
    head += f"\n> {pri_label} | 置信度：{it.get('confidence','')}"
    if ctx:
        head += f" | {ctx}"
    out.append(head)

    # Impact line
    out.append(f"**影响**：{it.get('impact','')}")

    # Complaint trigger (if any)
    ct = it.get('complaint_trigger', '')
    if ct:
        out.append(f"")
        out.append(f"> 💬 **质疑触发点**：{ct}")

    # Self-check items
    sci = it.get('selfcheck_items', [])
    if sci:
        out.append(f"")
        out.append(f"**自检清单**：")
        for item in sci:
            out.append(f"- [ ] {item}")
        out.append("")

    # Text diff
    if it.get('old_text'):
        out.append(f"**原版**：{it['old_text'][:400]}")
    if it.get('new_text'):
        out.append(f"**新版**：{it['new_text'][:400]}")
    if it.get('numeric_delta'):
        out.append(f"**数值变更**：{it['numeric_delta']}")

    # Legal check
    lc = it.get('legal_check')
    if isinstance(lc, dict) and lc.get('threshold_breached'):
        out.append("")
        out.append(f"🔴 **法定阈值超限**：{'；'.join(lc.get('breached_items', []))}")
        basis = lc.get('legal_basis', '')
        if basis:
            out.append(f"依据：{basis}")

    # Consistency check
    cc = it.get('consistency_check')
    if isinstance(cc, dict) and cc.get('issues_found'):
        out.append(f"")
        out.append(f"🔗 **一致性问题**：{'；'.join(cc['issues_found'])}")

    # Timeliness check
    tc = it.get('timeliness_check')
    if isinstance(tc, dict) and tc.get('extension_required'):
        out.append(f"")
        out.append(f"⏰ **时限注意**：{tc.get('extension_note', '')}")

    # Basis
    out.append(f"")
    out.append(f"**依据**：{it.get('basis','')}（{it.get('basis_source','')}）")

    # Release decision / action as final actionable item
    rd = it.get('release_decision', '')
    action = it.get('action', '')
    if rd or action:
        out.append(f"")
        if rd:
            out.append(f"> **发布决策**：{rd}")
        if action:
            out.append(f"> **处置建议**：{action}")
    out.append('')
    return '\n'.join(out)


def markdown_tenderer(items, meta):
    """招标人版完整 Markdown 报告。"""
    lines = ['# 招标文件版本自检报告 · 招标人版', '']
    lines.append(f"- 原版：{meta.get('old_file','')}（{meta.get('old_date','')}）")
    lines.append(f"- 新版：{meta.get('new_file','')}（{meta.get('new_date','')}）")
    lines.append(f"- 差异条数：{len(items)}")
    gen_time = meta.get('generated_at', '')
    if gen_time:
        lines.append(f"- 生成时间：{gen_time}")
    lines.append('')

    # Section 1: Executive Summary + Decision
    lines.append(_tenderer_summary_section(items))

    # Section 2: All detailed findings (sorted by priority)
    # Sort: P0 first, then P1, P2, P3, P4; within same priority, by severity desc
    pri_order = {'P0': 0, 'P1': 1, 'P2': 2, 'P3': 3, 'P4': 4}
    sev_order = {'高': 0, '中': 1, '低': 2}
    sorted_items = sorted(items, key=lambda x: (
        pri_order.get(x.get('priority', 'P4'), 9),
        sev_order.get(x.get('final_severity') or x.get('severity', '低'), 9)
    ))

    lines.append('## 全部差异明细（按优先级排序）')
    for idx, it in enumerate(sorted_items, 1):
        lines.append(_tenderer_detail(idx, it))

    # Disclaimer
    lines.append('')
    lines.append('---')
    lines.append('')
    lines.append('> ⚠️ **免责声明**：本报告为辅助自筛工具，不替代法律顾问正式意见。高利害场景（P0项、涉嫌排斥竞争项）建议在发布前由法务或合规部门复核。')
    return '\n'.join(lines)


# ============================================================
# DOCX RENDERER (shared structure, role-aware heading)
# ============================================================

def docx_render(items, meta, out, role='bidder'):
    if Document is None:
        raise RuntimeError('python-docx not installed')
    doc = Document()
    title = '招标文件版本比对报告'
    if role == 'tenderer':
        title += ' · 招标人版'
    else:
        title += ' · 投标人版'
    doc.add_heading(title, 0)

    doc.add_paragraph(f"原版：{meta.get('old_file','')}（{meta.get('old_date','')}）")
    doc.add_paragraph(f"新版：{meta.get('new_file','')}（{meta.get('new_date','')}）")
    doc.add_paragraph(f"投标截止：{meta.get('bid_deadline','未提供')}")
    doc.add_paragraph(f"差异条数：{len(items)}")

    # Use markdown content to populate body (simple approach: parse sections)
    if role == 'tenderer':
        md = markdown_tenderer(items, meta)
    else:
        md = markdown_bidder(items, meta)

    # Render markdown sections into docx paragraphs
    for line in md.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph('')
        elif line.startswith('# '):
            doc.add_heading(line[2:], 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
        elif line.startswith('- '):
            # Bold markers within list items
            text = line[2:]
            p = doc.add_paragraph(text, style='List Bullet')
        else:
            doc.add_paragraph(line)

    doc.save(out)
    print(f'Report ({role}) -> {out}')


def main():
    ap = argparse.ArgumentParser(description='Build version-compare report (role-aware)')
    ap.add_argument('--findings', required=True, help='findings.json path')
    ap.add_argument('--out', required=True, help='output .docx path')
    ap.add_argument('--markdown', default='', help='optional output .md path')
    ap.add_argument('--role', default='bidder', choices=['bidder', 'tenderer'],
                    help='report role: bidder (default) or tenderer')
    args = ap.parse_args()

    items, meta = load(args.findings)

    if args.markdown:
        md_fn = markdown_bidder if args.role == 'bidder' else markdown_tenderer
        with open(args.markdown, 'w', encoding='utf-8') as fh:
            fh.write(md_fn(items, meta))
        print(f'Markdown -> {args.markdown}')

    docx_render(items, meta, args.out, role=args.role)


if __name__ == '__main__':
    main()
