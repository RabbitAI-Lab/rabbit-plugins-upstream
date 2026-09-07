from __future__ import annotations

from pathlib import Path
from contextlib import contextmanager
import hashlib
import io
import json
import subprocess
import sys
import tempfile
import unittest
from unittest import mock

from zipfile import ZipFile

try:
    import tomllib
except ModuleNotFoundError:  # Python 3.10 test extra
    import tomli as tomllib

from docx import Document
from lxml import etree

from scripts import docx_export
from scripts.adapters.pathway_bridge import (
    bridge_pathway_policies,
    bridge_pathway_policy_evidence,
)
from scripts.contracts import EvidenceStatus, RecommendationResult, SourceTier
from scripts.decision_policy import DecisionPolicySnapshot
from scripts.evidence import EvidenceStore
from scripts.generate_report import build_pathway_atlas_model
from scripts.path_recommend import (
    PATHWAY_DISPLAY_EVIDENCE_FIELDS,
    PathwayFieldEvidenceOrigin,
)
from scripts.research_snapshot import build_research_snapshot
from scripts.report_model import ReportModel, build_report_model, render_markdown
from scripts.school_recommend import recommend_schools
from tests.test_generate_report_evidence import (
    capability,
    evidence_snapshot,
    formal_pathway_result,
    pathway_rank_model,
    pathway_rank_scenario,
    pathway_result,
    partial_task3_recommendations,
    rank_estimate,
    recommendations,
    student,
)
from tests.test_pathway_atlas_blackbox import school_anchor_bridge
from tests.test_pathway_evidence_bridge import candidate as pathway_candidate
from tests.test_pathway_evidence_bridge import project as pathway_projection
from tests.test_rank_evidence_bridge import candidate as rank_candidate
from tests.test_research_snapshot import bridges
from tests.test_scenario_recommendations import policy, profile, rows, scenario


ROOT = Path(__file__).resolve().parents[1]
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def canonical_digest(value) -> str:
    encoded = json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        allow_nan=False,
    ).encode("utf-8")
    return "sha256:" + hashlib.sha256(encoded).hexdigest()


def expected_origin_binding(origin: PathwayFieldEvidenceOrigin, payload) -> str:
    return canonical_digest(
        {
            "contract": "pathway-field-origin-v1",
            "origin": origin.value,
            "payload": payload,
        }
    )


def expected_record_digest(record) -> str:
    payload = record.to_dict()
    payload.pop("digest")
    return canonical_digest(
        {"contract": "pathway-field-evidence-v2", "record": payload}
    )


def model(**overrides) -> ReportModel:
    values = {
        "profile": student(secondary_subjects=("化学", "生物")),
        "recommendations": recommendations(),
        "rank": rank_estimate(),
        "pathways": formal_pathway_result(),
        "evidence": evidence_snapshot(),
    }
    values.update(overrides)
    return build_report_model(**values)


def document_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def xml_part(path: Path, name: str):
    with ZipFile(path) as package:
        return etree.fromstring(package.read(name))


@contextmanager
def typed_atlas_artifacts(*, admission_bridge=None, admission_candidates=()):
    """Persist a full v3 bundle, optionally with an adapter-produced admission bridge."""
    planning, query_plan, rank_bridge, default_admission_bridge = bridges()
    active_admission_bridge = (
        default_admission_bridge if admission_bridge is None else admission_bridge
    )
    active_admission_candidates = (
        (
            rank_candidate(
                "official-admission",
                publisher="湖北省普通批发布机关",
                host="admission.hubei.gov.cn",
            ),
        )
        if admission_bridge is None
        else tuple(admission_candidates)
    )
    school_2025 = school_anchor_bridge(planning, query_plan, 2025, 20000)
    school_2026 = school_anchor_bridge(planning, query_plan, 2026, 20000)
    pathway_source = pathway_candidate()
    pathway_bridge = bridge_pathway_policy_evidence(
        pathway_projection(
            student=planning,
            query_plan=query_plan,
            candidates=(pathway_source,),
        )
    )
    registered_sources = (
        rank_candidate(),
        *active_admission_candidates,
        *school_2025.candidates,
        *school_2026.candidates,
        pathway_source,
    )
    with tempfile.TemporaryDirectory() as temporary:
        root = Path(temporary)
        store = EvidenceStore.create(root, capability())
        for source in registered_sources:
            store.add_candidate(source)
        rank_bridge.persist(store)
        school_2025.persist(store)
        school_2026.persist(store)
        active_admission_bridge.persist(store)
        pathway_bridge.persist(store)
        store.finalize()
        profile_path = root / "profile.json"
        profile_payload = planning.to_dict()
        profile_payload.pop("mode")
        profile_payload.pop("digest")
        profile_path.write_text(
            json.dumps(profile_payload, ensure_ascii=False), encoding="utf-8"
        )
        yield planning, query_plan, store.session_path, profile_path


class DocxSemanticParityTest(unittest.TestCase):
    def test_authenticated_field_trails_bind_query_task_year_plan_and_rank_scenario(self):
        with typed_atlas_artifacts() as (planning, query_plan, bundle, _profile):
            reviewed = DecisionPolicySnapshot.load_default()
            research = build_research_snapshot(planning, query_plan, bundle, reviewed)
            report = build_pathway_atlas_model(
                planning, research, bundle, query_plan, decision_policy=reviewed
            )
            plan_digest = canonical_digest(query_plan.to_dict())
            policies = bridge_pathway_policies(
                bundle,
                province=planning.province,
                subject_mode=planning.subject_mode,
                target_year=query_plan.research_year,
                expected_profile_digest=planning.digest,
                expected_query_plan_digest=plan_digest,
            )

        self.assertEqual(len(policies), 1)
        policy = policies[0]
        evaluated = tuple(
            item for item in report.pathways if item.policy_id == policy.policy_id
        )
        observations = tuple(
            item for item in report.pathways if item.policy_id != policy.policy_id
        )
        expected_titles = {
            task.target_name
            for task in query_plan.tasks
            if task.target_name is not None
            and task.kind
            in {
                "strong_foundation",
                "comprehensive_evaluation",
                "hk_macao_admission",
                "special_pathway",
            }
        }
        self.assertEqual(len(evaluated), 1)
        self.assertEqual({item.title for item in report.pathways}, expected_titles)
        self.assertEqual(len(observations), len(expected_titles) - 1)
        self.assertTrue(
            all(
                item.status == "pending_verification"
                and item.investment_decision == "观察"
                and item.qualification_status == "待核验"
                and item.evidence_status is EvidenceStatus.MISSING
                and not item.source_ids
                and item.target_rank is None
                and item.target_year is None
                and item.data_year is None
                and not item.timeline
                and not item.professional_options
                and item.institution == "待核验"
                and item.training_arrangements is None
                and item.transition_rules is None
                and item.outcomes is None
                for item in observations
            )
        )
        for observation in observations:
            self.assertEqual(
                tuple(record.field for record in observation.field_evidence),
                PATHWAY_DISPLAY_EVIDENCE_FIELDS,
            )
            self.assertTrue(
                all(not record.source_ids for record in observation.field_evidence)
            )
        item = evaluated[0]
        projection = policy._authenticated_projection
        self.assertIsNotNone(projection)
        task = projection.input_projection["task"]
        self.assertEqual(
            task["task_digest"],
            canonical_digest(
                {name: value for name, value in task.items() if name != "task_digest"}
            ),
        )
        self.assertEqual(policy.profile_digest, planning.digest)
        self.assertEqual(policy.query_plan_digest, plan_digest)
        self.assertEqual(projection.query_plan_digest, plan_digest)

        records = {record.field: record for record in item.field_evidence}
        self.assertEqual(
            tuple(record.field for record in item.field_evidence),
            PATHWAY_DISPLAY_EVIDENCE_FIELDS,
        )
        self.assertEqual(len(records), 22)
        self.assertTrue(
            all(record.digest == expected_record_digest(record) for record in records.values())
        )

        title = records["title"]
        institution = records["institution"]
        self.assertIs(title.origin, PathwayFieldEvidenceOrigin.QUERY_CONTEXT)
        self.assertEqual(item.title, task["target_name"])
        self.assertEqual(title.upstream_fields, ("query_task.target_name",))
        self.assertEqual(
            title.locators,
            (f"query-task:{task['task_id']}/target_name",),
        )
        self.assertFalse(set(title.locators).intersection(institution.locators))
        self.assertEqual(
            title.origin_binding,
            expected_origin_binding(
                PathwayFieldEvidenceOrigin.QUERY_CONTEXT,
                {
                    "projection_digest": projection.digest,
                    "task": task,
                    "field": "target_name",
                },
            ),
        )

        policy_records = {record.field: record for record in policy.field_evidence}
        data_year = policy_records["data_year"]
        year = records["year_basis"]
        self.assertIs(year.origin, PathwayFieldEvidenceOrigin.DERIVED_DECISION)
        self.assertEqual(
            year.upstream_fields,
            ("data_year", "query_plan.research_year", "query_task.target_year"),
        )
        self.assertEqual(year.profile_fields, ("query_plan.research_year",))
        self.assertEqual(year.upstream_evidence_digests, (data_year.digest,))
        self.assertIn(f"query-task:{task['task_id']}/target_year", year.locators)
        self.assertEqual(task["target_year"], item.target_year)
        self.assertEqual(query_plan.research_year, item.target_year)
        self.assertEqual(
            year.origin_binding,
            expected_origin_binding(
                PathwayFieldEvidenceOrigin.DERIVED_DECISION,
                {
                    "policy_id": policy.policy_id,
                    "field": "year_basis",
                    "decision_reasons": [
                        reason.to_dict() for reason in item.decision_reasons
                    ],
                    "year_basis": item.year_basis,
                    "target_year": item.target_year,
                    "data_year": item.data_year,
                    "fallback_distance": item.fallback_distance,
                    "research_year": query_plan.research_year,
                    "query_task_digest": task["task_digest"],
                },
            ),
        )

        self.assertIsNotNone(report.rank)
        rank_scenario = report.rank
        calculation = records["calculation_basis"]
        self.assertIs(
            calculation.origin, PathwayFieldEvidenceOrigin.DERIVED_DECISION
        )
        self.assertEqual(
            calculation.upstream_fields,
            (
                "data_year",
                "rank_scenario.basis",
                "rank_scenario.central_rank",
                "rank_scenario.source_ids",
                "rank_scenario.status",
            ),
        )
        self.assertEqual(calculation.upstream_evidence_digests, (data_year.digest,))
        self.assertEqual(
            calculation.source_ids,
            tuple(sorted(set(policy.policy_source_ids) | set(rank_scenario.source_ids))),
        )
        self.assertEqual(
            calculation.origin_binding,
            expected_origin_binding(
                PathwayFieldEvidenceOrigin.DERIVED_DECISION,
                {
                    "policy_id": policy.policy_id,
                    "field": "calculation_basis",
                    "decision_reasons": [
                        reason.to_dict() for reason in item.decision_reasons
                    ],
                    "target_rank": None,
                    "transformation": None,
                    "rank_scenario": rank_scenario.to_dict(),
                    "rank_model": None,
                },
            ),
        )

    def test_model_backed_calculation_trail_binds_rank_scenario_and_model(self):
        rank_scenario = pathway_rank_scenario()
        rank_model = pathway_rank_model()
        pathways = formal_pathway_result()
        empty_recommendations = recommendations(
            items=(),
            coverage_status=EvidenceStatus.REFERENCE,
            empty_reason="no_match_within_verified_coverage",
            warnings=(),
        )
        report = model(
            rank=rank_scenario,
            pathways=pathways,
            recommendations=empty_recommendations,
        )
        item = report.pathways[0]
        calculation = next(
            record
            for record in item.field_evidence
            if record.field == "calculation_basis"
        )

        self.assertIs(
            calculation.origin, PathwayFieldEvidenceOrigin.DERIVED_DECISION
        )
        self.assertEqual(
            calculation.upstream_fields,
            (
                "data_year",
                "rank_model.cohort_years",
                "rank_model.evidence_status",
                "rank_model.method",
                "rank_model.model_id",
                "rank_model.source_ids",
                "rank_scenario.basis",
                "rank_scenario.central_rank",
                "rank_scenario.source_ids",
                "rank_scenario.status",
            ),
        )
        self.assertEqual(
            calculation.source_ids,
            tuple(
                sorted(
                    {"s5"}
                    | set(rank_scenario.source_ids)
                    | set(rank_model.source_ids)
                )
            ),
        )
        self.assertTrue(calculation.upstream_evidence_digests)
        self.assertEqual(calculation.digest, expected_record_digest(calculation))
        self.assertEqual(
            calculation.origin_binding,
            expected_origin_binding(
                PathwayFieldEvidenceOrigin.DERIVED_DECISION,
                {
                    "policy_id": item.policy_id,
                    "field": "calculation_basis",
                    "decision_reasons": [],
                    "target_rank": pathways.target_rank,
                    "transformation": pathways.transformation,
                    "rank_scenario": rank_scenario.to_dict(),
                    "rank_model": rank_model.to_dict(),
                },
            ),
        )
        markdown = render_markdown(report)
        with tempfile.TemporaryDirectory() as temporary:
            output = docx_export.export_docx(
                report, Path(temporary) / "model-backed-field-evidence.docx"
            )
            text = document_text(output)
        for literal in (
            "rank_model.source_ids",
            "rank_model.evidence_status",
            "rank_model.method",
            "rank_scenario.source_ids",
            "rank_scenario.status",
            "s3、s4、s5、s6",
        ):
            self.assertIn(literal, markdown)
            self.assertIn(literal, text)

    def test_markdown_and_docx_reject_mutated_field_evidence_record(self):
        with typed_atlas_artifacts() as (planning, query_plan, bundle, _profile):
            reviewed = DecisionPolicySnapshot.load_default()
            research = build_research_snapshot(planning, query_plan, bundle, reviewed)
            report = build_pathway_atlas_model(
                planning, research, bundle, query_plan, decision_policy=reviewed
            )

        trail = report.pathways[0].field_evidence[0]
        object.__setattr__(trail, "source_ids", ("forged-source",))
        with self.assertRaisesRegex(ValueError, "field evidence digest"):
            render_markdown(report)
        with tempfile.TemporaryDirectory() as temporary:
            output = Path(temporary) / "mutated-trail.docx"
            with self.assertRaisesRegex(ValueError, "field evidence digest"):
                docx_export.export_docx(report, output)
            self.assertFalse(output.exists())

    def test_markdown_and_docx_reject_mutated_displayed_pathway_value(self):
        with typed_atlas_artifacts() as (planning, query_plan, bundle, _profile):
            reviewed = DecisionPolicySnapshot.load_default()
            research = build_research_snapshot(planning, query_plan, bundle, reviewed)
            report = build_pathway_atlas_model(
                planning, research, bundle, query_plan, decision_policy=reviewed
            )

        object.__setattr__(
            report.pathways[0], "professional_options", ("FAKE-MAJOR",)
        )
        with self.assertRaisesRegex(
            ValueError, "professional_options.*value digest"
        ):
            render_markdown(report)
        with tempfile.TemporaryDirectory() as temporary:
            output = Path(temporary) / "mutated-value.docx"
            with self.assertRaisesRegex(
                ValueError, "professional_options.*value digest"
            ):
                docx_export.export_docx(report, output)
            self.assertFalse(output.exists())

    def test_rank_scenarios_and_decisive_pathways_have_markdown_docx_parity(self):
        with typed_atlas_artifacts() as (planning, query_plan, bundle, _profile):
            reviewed = DecisionPolicySnapshot.load_default()
            research = build_research_snapshot(planning, query_plan, bundle, reviewed)
            report = build_pathway_atlas_model(
                planning, research, bundle, query_plan, decision_policy=reviewed
            )
            markdown = render_markdown(report)
            with tempfile.TemporaryDirectory() as temporary:
                output = docx_export.export_docx(report, Path(temporary) / "parity.docx")
                text = document_text(output)

        rank_values = (
            report.rank.optimistic_rank,
            report.rank.central_rank,
            report.rank.conservative_rank,
        )
        self.assertGreater(len(set(rank_values)), 1)
        for literal in (
            f"乐观位次：{rank_values[0]}",
            f"中性位次：{rank_values[1]}",
            f"保守位次：{rank_values[2]}",
            "official-pathway",
            "2025",
        ):
            self.assertIn(literal, markdown)
            self.assertIn(literal, text)

    def test_direct_and_derived_pathway_field_audits_have_markdown_docx_parity(self):
        report = model(pathways=formal_pathway_result())
        markdown = render_markdown(report)
        with tempfile.TemporaryDirectory() as temporary:
            output = docx_export.export_docx(
                report, Path(temporary) / "field-evidence-parity.docx"
            )
            text = document_text(output)

        expected = (
            "字段：professional_options（专业选项）；证据状态：官方；覆盖：完整；"
            "来源编号：s5；证据定位：policy-record:policy-formal:professional_options；"
            "抽取方式：legacy-policy-record；证据方法："
            "legacy-policy-field-v1；上游字段：professional_options；"
            "画像字段：无；提示：无",
            "字段：investment_decision（投入结论）；证据状态：推断；覆盖：完整；"
            "来源编号：s5；证据定位：policy-record:policy-formal:activity_requirements",
            "证据方法：pathway-investment-decision-v1",
        )
        self.assertIn("逐字段证据审计", markdown)
        self.assertIn("逐字段证据审计", text)
        for literal in expected:
            self.assertIn(literal, markdown)
            self.assertIn(literal, text)

    def test_docx_revalidates_pathway_field_trails_after_frozen_object_mutation(self):
        report = model(pathways=formal_pathway_result())
        pathway = report.pathways[0]
        object.__setattr__(pathway, "field_evidence", pathway.field_evidence[:-1])

        with tempfile.TemporaryDirectory() as temporary:
            output = Path(temporary) / "invalid-field-evidence.docx"
            with self.assertRaisesRegex(ValueError, "field evidence is incomplete"):
                docx_export.export_docx(report, output)
            self.assertFalse(output.exists())

    def test_docx_gate_uses_shared_semantics_and_exception_never_echoes_amount(self):
        safe = "武汉大学学费 30000元；国家助学金 6000元"
        rejected_amount = "30600元"
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            with mock.patch.object(docx_export, "_document_text", return_value=safe):
                output = docx_export.export_docx(model(), root / "safe.docx")
            self.assertTrue(output.is_file())

            with mock.patch.object(
                docx_export,
                "_document_text",
                return_value=f"升学规划产品售价 {rejected_amount}",
            ):
                with self.assertRaises(docx_export.DocumentComplianceError) as captured:
                    docx_export.export_docx(model(), root / "rejected.docx")
            self.assertFalse((root / "rejected.docx").exists())

        self.assertNotIn(rejected_amount, str(captured.exception))

    def test_task3_inside_coverage_partial_result_has_markdown_docx_parity(self):
        report = model(recommendations=partial_task3_recommendations())
        markdown = render_markdown(report)
        with tempfile.TemporaryDirectory() as temporary:
            output = docx_export.export_docx(report, Path(temporary) / "partial.docx")
            text = document_text(output)
        self.assertEqual(report.recommendations, ())
        for literal in ("部分覆盖大学", "部分覆盖", "s2", "仅作方向性观察", "不进入冲稳保"):
            self.assertIn(literal, markdown)
            self.assertIn(literal, text)
        self.assertNotIn("| 620 |", markdown)
        self.assertNotIn("| 4300 |", markdown)
        for observation in report.school_observations:
            self.assertFalse(hasattr(observation, "min_score"))
            self.assertFalse(hasattr(observation, "min_rank"))
        self.assertNotIn("620", text)
        self.assertNotIn("4300", text)

    def test_task3_outside_coverage_partial_result_has_markdown_docx_parity(self):
        result = partial_task3_recommendations(rank=6000)
        report = model(
            profile=student(
                rank=6000,
                secondary_subjects=("化学", "生物"),
            ),
            recommendations=result,
            rank=None,
        )
        markdown = render_markdown(report)
        with tempfile.TemporaryDirectory() as temporary:
            output = docx_export.export_docx(
                report,
                Path(temporary) / "partial-outside.docx",
            )
            text = document_text(output)

        for literal in (
            "部分覆盖大学",
            "部分覆盖",
            "仅作方向性观察",
            "不进入冲稳保",
            "s2",
        ):
            self.assertIn(literal, markdown)
            self.assertIn(literal, text)
        for forbidden in ("620", "4300"):
            self.assertNotIn(forbidden, markdown)
            self.assertNotIn(forbidden, text)

    def test_docx_accepts_strict_machine_ids_with_phone_shaped_digits(self):
        from scripts.contracts import EvidenceManifest
        from scripts.validate_evidence import ValidatedEvidenceSnapshot

        snapshot = evidence_snapshot()
        session_id = "a13800138000bcdef123456789abcdef"
        manifest_hash = "sha256:" + session_id * 2
        machine_snapshot = ValidatedEvidenceSnapshot._create(
            EvidenceManifest(
                schema_version=snapshot.manifest.schema_version,
                session_id=session_id,
                capability_tier=snapshot.manifest.capability_tier,
                candidates_filename=snapshot.manifest.candidates_filename,
                facts_filename=snapshot.manifest.facts_filename,
                rejected_count=snapshot.manifest.rejected_count,
                manifest_hash=manifest_hash,
            ),
            snapshot.capability,
            snapshot.retrieval_dates,
            snapshot.facts,
            snapshot.rejections,
        )
        try:
            report = model(evidence=machine_snapshot)
        except ValueError as error:
            self.fail(f"strict machine identifiers were treated as human text: {error}")

        with tempfile.TemporaryDirectory() as temporary:
            output = docx_export.export_docx(
                report, Path(temporary) / "machine-identifiers.docx"
            )
            text = document_text(output)

        self.assertIn(session_id, text)
        self.assertIn(manifest_hash, text)

    def test_export_projects_the_complete_report_model_semantics(self):
        """Catches a renderer that drops evidence, rank, pathway, or action fields."""
        report = model()
        with tempfile.TemporaryDirectory() as temporary:
            output = Path(temporary) / "report.docx"
            result = docx_export.export_docx(report, output)
            text = document_text(result)

        for literal in (
            "匿名升学规划报告（演示甲省）",
            "化学、生物",
            "证据状态",
            "部分覆盖",
            "检索日期",
            "2026-08-23",
            "仅覆盖 2026",
            "synthetic-ordinary-batch-v1",
            "synthetic-policy-basis-v1",
            "冲=3、稳=4、保=5",
            "虚构甲大学",
            "4300",
            "school_rank_offset_median_observed_spread",
            "虚构正式专项",
            "合成培养安排",
            "合成转段规则",
            "合成出口说明",
            "合成服务就业说明",
            "合成退出规则",
            "合成费用说明",
            "model-report",
            "documented_rank_delta",
            "当前最需要做的事",
            "战略价值：high",
            "依赖行动：补齐或复核关键证据缺口（evidence-gap-review）",
            "关联院校：虚构甲大学",
            "关联路径：虚构正式专项",
            "基于公开数据由 AI 整理，仅供参考",
        ):
            self.assertIn(literal, text)
        self.assertNotIn("http://", text)
        self.assertNotIn("https://", text)
        self.assertNotIn(str(ROOT), text)
        self.assertGreaterEqual(
            text.count("基于公开数据由 AI 整理，仅供参考"),
            2,
        )
        self.assertIn("八、证据披露", render_markdown(report))

    def test_optional_and_unusable_sections_degrade_without_proxy_values(self):
        """Catches DOCX-only fallbacks that invent rank or pathway values."""
        empty = RecommendationResult(
            ordinary_batch_policy=recommendations().ordinary_batch_policy,
            items=(),
            excluded_by_subject_count=2,
            zero_score_excluded_count=1,
            input_years=(2026,),
            usable_years=(),
            verified_rank_coverage=None,
            coverage_status=EvidenceStatus.MASKED,
            empty_reason="unusable_evidence",
            warnings=("屏蔽值未进入计算",),
        )
        report = model(recommendations=empty, rank=None, pathways=None)
        with tempfile.TemporaryDirectory() as temporary:
            output = docx_export.export_docx(report, Path(temporary))
            text = document_text(output)

        self.assertEqual(output.name, "匿名升学规划报告-演示甲省-2026.docx")
        self.assertIn("屏蔽、冲突或不可精确使用", text)
        self.assertIn("不执行校排名折算", text)
        self.assertIn("多元升学数据不足", text)
        self.assertNotIn("− 4000", text)
        self.assertNotIn("-4000", text)

    def test_export_rejects_detached_or_missing_document_capability(self):
        """Catches constructor bypass and silent optional-dependency skips."""
        with tempfile.TemporaryDirectory() as temporary:
            output = Path(temporary) / "report.docx"
            with self.assertRaisesRegex(TypeError, "ReportModel"):
                docx_export.export_docx({"profile": "detached"}, output)
            with mock.patch.object(docx_export, "Document", None):
                with self.assertRaisesRegex(
                    docx_export.DocumentDependencyError, "documents"
                ):
                    docx_export.export_docx(model(), output)
                stderr = io.StringIO()
                stdout = io.StringIO()
                with mock.patch("sys.stderr", stderr), mock.patch("sys.stdout", stdout):
                    self.assertEqual(docx_export.main([]), 3)
                self.assertIn("python-docx", stderr.getvalue())
                self.assertIn("缺少能力", stderr.getvalue())
                self.assertEqual(stdout.getvalue(), "")
                self.assertFalse(output.exists())

    def test_cli_rejects_secondary_subject_override_when_bundle_digest_mismatches(self):
        """Catches a DOCX CLI that replays a v3 evidence bundle for another profile."""
        with typed_atlas_artifacts() as (_planning, _query_plan, bundle, profile_path):
            with tempfile.TemporaryDirectory() as temporary:
                output = Path(temporary) / "anonymous-admission-report.docx"
                command = [
                    sys.executable,
                    str(ROOT / "scripts" / "docx_export.py"),
                    "--dataset",
                    str(ROOT / "tests" / "fixtures" / "provinces" / "demo-312"),
                    "--profile",
                    str(profile_path),
                    "--evidence",
                    str(bundle),
                    "--secondary-subject",
                    "化学",
                    "--secondary-subject",
                    "生物",
                    "--output",
                    str(output),
                ]
                completed = subprocess.run(
                    command, capture_output=True, text=True, encoding="utf-8"
                )
                self.assertEqual(completed.returncode, 2)
                self.assertFalse(output.exists())
                self.assertEqual(completed.stdout, "")
                self.assertEqual(completed.stderr, "错误[DOCX_002]：DOCX 生成或发布失败\n")

    def test_public_cli_defaults_to_exclusive_canonical_output_in_cwd(self):
        with typed_atlas_artifacts() as (_planning, _query_plan, bundle, profile_path):
            command = [
                sys.executable,
                str(ROOT / "scripts" / "docx_export.py"),
                "--dataset", str(ROOT / "tests" / "fixtures" / "provinces" / "demo-312"),
                "--profile", str(profile_path),
                "--evidence", str(bundle),
            ]
            with tempfile.TemporaryDirectory() as temporary:
                sandbox = Path(temporary)
                success_dir = sandbox / "success"
                competing_dir = sandbox / "competing"
                success_dir.mkdir()
                competing_dir.mkdir()

                completed = subprocess.run(
                    command,
                    cwd=success_dir,
                    capture_output=True,
                    text=True,
                    encoding="utf-8",
                )
                output = success_dir / "anonymous-admission-report.docx"
                self.assertEqual(completed.returncode, 0, completed.stderr)
                self.assertEqual(
                    json.loads(completed.stdout)["filename"],
                    "anonymous-admission-report.docx",
                )
                self.assertTrue(output.is_file())

                competitor = competing_dir / "anonymous-admission-report.docx"
                competitor.write_bytes(b"competitor-owned")
                refused = subprocess.run(
                    command,
                    cwd=competing_dir,
                    capture_output=True,
                    text=True,
                    encoding="utf-8",
                )
                self.assertEqual(refused.returncode, 2)
                self.assertEqual(refused.stdout, "")
                self.assertEqual(
                    refused.stderr, "错误[DOCX_002]：DOCX 生成或发布失败\n"
                )
                self.assertEqual(competitor.read_bytes(), b"competitor-owned")

    def test_public_cli_rejects_pii_output_name_with_path_neutral_error(self):
        with tempfile.TemporaryDirectory() as temporary:
            output = Path(temporary) / "张三-13800138000-secret.docx"
            completed = subprocess.run(
                [
                    sys.executable,
                    str(ROOT / "scripts" / "docx_export.py"),
                    "--dataset", str(ROOT / "tests" / "fixtures" / "provinces" / "demo-312"),
                    "--profile", str(ROOT / "tests" / "fixtures" / "profiles" / "demo.json"),
                    "--evidence", str(ROOT / "tests" / "fixtures" / "evidence" / "three-source-consensus"),
                    "--output", str(output),
                ],
                capture_output=True,
                text=True,
                encoding="utf-8",
            )
        self.assertEqual(completed.returncode, 2)
        self.assertEqual(completed.stdout, "")
        self.assertEqual(completed.stderr, "错误[DOCX_002]：DOCX 生成或发布失败\n")
        for forbidden in (str(output), output.name, "张三", "13800138000", "secret"):
            self.assertNotIn(forbidden, completed.stderr)

    def test_pending_pathway_preserves_missing_constraints_and_real_details(self):
        """Catches a DOCX renderer that hides why a pathway is not formal."""
        report = model(pathways=pathway_result())
        with tempfile.TemporaryDirectory() as temporary:
            path = docx_export.export_docx(report, Path(temporary))
            text = document_text(path)
        for literal in (
            "待核实",
            "服务期未核实",
            "合成培养安排",
            "合成转段规则",
            "合成出口说明",
            "当前证据未提供",
            "未提供有依据的位次模型",
        ):
            self.assertIn(literal, text)

    def test_documents_extra_is_a_real_runtime_dependency(self):
        """Catches an optional feature whose tests can only skip at import time."""
        payload = tomllib.loads((ROOT / "pyproject.toml").read_text(encoding="utf-8"))
        self.assertIn("python-docx>=1.1,<2", payload["project"]["optional-dependencies"]["documents"])
        self.assertGreaterEqual(tuple(map(int, __import__("docx").__version__.split(".")[:2])), (1, 1))


class DocxStructureTest(unittest.TestCase):
    def export(self, temporary: str) -> Path:
        return docx_export.export_docx(model(), Path(temporary) / "report.docx")

    def test_compact_reference_guide_geometry_styles_and_metadata(self):
        """Catches reliance on Word defaults or personal core properties."""
        with tempfile.TemporaryDirectory() as temporary:
            path = self.export(temporary)
            document = Document(path)
            section = document.sections[0]
            self.assertEqual(section.page_width.twips, 12240)
            self.assertEqual(section.page_height.twips, 15840)
            self.assertEqual(section.top_margin.twips, 1440)
            self.assertEqual(section.right_margin.twips, 1440)
            self.assertEqual(section.bottom_margin.twips, 1440)
            self.assertEqual(section.left_margin.twips, 1440)
            self.assertEqual(section.header_distance.twips, 708)
            self.assertEqual(section.footer_distance.twips, 708)
            self.assertEqual(document.styles["Normal"].font.name, "Calibri")
            self.assertEqual(document.styles["Normal"].font.size.pt, 11)
            self.assertEqual(document.styles["Normal"].paragraph_format.space_after.pt, 6)
            self.assertEqual(document.styles["Normal"].paragraph_format.line_spacing, 1.25)
            for style_name, size, color, before, after in (
                ("Heading 1", 16, "2E74B5", 18, 10),
                ("Heading 2", 13, "2E74B5", 14, 7),
                ("Heading 3", 12, "1F4D78", 10, 5),
            ):
                style = document.styles[style_name]
                self.assertEqual(style.font.size.pt, size)
                self.assertEqual(str(style.font.color.rgb), color)
                self.assertEqual(style.paragraph_format.space_before.pt, before)
                self.assertEqual(style.paragraph_format.space_after.pt, after)
            self.assertEqual(document.core_properties.author, "")
            self.assertEqual(document.core_properties.last_modified_by, "")
            self.assertNotIn("张三", str(document.core_properties.__dict__))

    def test_tables_use_fixed_matching_dxa_geometry_without_fixed_rows(self):
        """Catches percentage/autofit tables, drifting cell widths, and clipped rows."""
        with tempfile.TemporaryDirectory() as temporary:
            root = xml_part(self.export(temporary), "word/document.xml")
        tables = root.xpath(".//w:tbl", namespaces=NS)
        self.assertGreaterEqual(len(tables), 2)
        for table in tables:
            width = table.xpath("./w:tblPr/w:tblW", namespaces=NS)[0]
            indent = table.xpath("./w:tblPr/w:tblInd", namespaces=NS)[0]
            self.assertEqual(width.get(f"{{{NS['w']}}}type"), "dxa")
            self.assertEqual(width.get(f"{{{NS['w']}}}w"), "9360")
            self.assertEqual(indent.get(f"{{{NS['w']}}}type"), "dxa")
            self.assertEqual(indent.get(f"{{{NS['w']}}}w"), "120")
            grid = [
                int(node.get(f"{{{NS['w']}}}w"))
                for node in table.xpath("./w:tblGrid/w:gridCol", namespaces=NS)
            ]
            self.assertEqual(sum(grid), 9360)
            for row in table.xpath("./w:tr", namespaces=NS):
                widths = [
                    int(node.get(f"{{{NS['w']}}}w"))
                    for node in row.xpath("./w:tc/w:tcPr/w:tcW", namespaces=NS)
                ]
                self.assertEqual(widths, grid)
        self.assertFalse(root.xpath(".//w:tblW[@w:type='pct']", namespaces=NS))
        self.assertFalse(root.xpath(".//w:tblLayout[@w:type='autofit']", namespaces=NS))
        self.assertFalse(root.xpath(".//w:trHeight", namespaces=NS))

    def test_lists_use_real_numbering_and_never_fake_markers(self):
        """Catches manual bullets/numbers and wrapped-line misalignment."""
        with tempfile.TemporaryDirectory() as temporary:
            path = self.export(temporary)
            document_xml = xml_part(path, "word/document.xml")
            numbering_xml = xml_part(path, "word/numbering.xml")
        numbered = document_xml.xpath(".//w:p[w:pPr/w:numPr]", namespaces=NS)
        self.assertGreaterEqual(len(numbered), 3)
        for paragraph in document_xml.xpath(".//w:body/w:p", namespaces=NS):
            text = "".join(paragraph.xpath(".//w:t/text()", namespaces=NS)).lstrip()
            self.assertFalse(text.startswith(("•", "·", "- ")))
            self.assertIsNone(__import__("re").match(r"^\d+[.、]\s", text))
        levels = numbering_xml.xpath(".//w:abstractNum/w:lvl[@w:ilvl='0']", namespaces=NS)
        self.assertTrue(
            any(
                level.xpath("./w:numFmt[@w:val='bullet']", namespaces=NS)
                and level.xpath("./w:pPr/w:ind[@w:left='540'][@w:hanging='271']", namespaces=NS)
                for level in levels
            )
        )
        self.assertTrue(
            any(
                level.xpath("./w:numFmt[@w:val='decimal']", namespaces=NS)
                and level.xpath("./w:pPr/w:ind[@w:left='540'][@w:hanging='271']", namespaces=NS)
                for level in levels
            )
        )

    def test_package_is_byte_deterministic_and_contains_no_local_identity(self):
        """Catches timestamps, temp paths, or author identity leaking into the ZIP."""
        with tempfile.TemporaryDirectory() as temporary:
            first = docx_export.export_docx(model(), Path(temporary) / "a.docx")
            second = docx_export.export_docx(model(), Path(temporary) / "b.docx")
            first_bytes = first.read_bytes()
            second_bytes = second.read_bytes()
            self.assertEqual(hashlib.sha256(first_bytes).digest(), hashlib.sha256(second_bytes).digest())
            with ZipFile(first) as package:
                xml = b"\n".join(
                    package.read(name)
                    for name in package.namelist()
                    if name.endswith((".xml", ".rels"))
                ).decode("utf-8", errors="replace")
        # OOXML namespace declarations are themselves HTTP URIs; raw-source
        # URLs are therefore asserted against visible text in the semantic test.
        for forbidden in (str(ROOT), "C:\\Users\\hp", "张三", "13800138000"):
            self.assertNotIn(forbidden, xml)

    def test_publish_race_never_deletes_a_competing_owner_file(self):
        """Catches cleanup that unlinks a path this process never created."""
        with tempfile.TemporaryDirectory() as temporary:
            output = Path(temporary) / "race.docx"
            original_resolver = docx_export._output_path

            def create_competitor(report, requested):
                destination = original_resolver(report, requested)
                destination.write_bytes(b"RIVAL-OWNER")
                return destination

            with mock.patch.object(
                docx_export, "_output_path", side_effect=create_competitor
            ):
                with self.assertRaises(FileExistsError):
                    docx_export.export_docx(model(), output)

            self.assertEqual(output.read_bytes(), b"RIVAL-OWNER")

    def test_final_path_appears_only_after_complete_archive_is_closed(self):
        """Catches streaming a partial ZIP directly into the public final path."""
        real_zip_file = docx_export.ZipFile
        visibility_during_writes = []

        class ObservingTargetZip:
            def __init__(self, archive):
                self.archive = archive

            def __enter__(self):
                return self

            def __exit__(self, *exc_info):
                self.archive.close()

            def writestr(self, info, data):
                visibility_during_writes.append(output.exists())
                return self.archive.writestr(info, data)

        def zip_factory(file, mode="r", *args, **kwargs):
            archive = real_zip_file(file, mode, *args, **kwargs)
            return ObservingTargetZip(archive) if mode == "w" else archive

        with tempfile.TemporaryDirectory() as temporary:
            output = Path(temporary).resolve() / "atomic.docx"
            with mock.patch.object(docx_export, "ZipFile", side_effect=zip_factory):
                result = docx_export.export_docx(model(), output)

            self.assertTrue(visibility_during_writes)
            self.assertFalse(any(visibility_during_writes))
            self.assertEqual(result, output)
            self.assertTrue(output.is_file())
            self.assertIn("匿名升学规划报告", document_text(output))

    def test_owned_partial_publish_is_removed_after_write_failure(self):
        """Catches an exclusive publisher that leaves its own corrupt artifact."""
        real_zip_file = docx_export.ZipFile
        visibility_at_failure = []

        class FailingTargetZip:
            def __init__(self, archive):
                self.archive = archive

            def __enter__(self):
                return self

            def __exit__(self, *exc_info):
                self.archive.close()

            def writestr(self, _info, _data):
                visibility_at_failure.append(output.exists())
                raise OSError("synthetic write failure")

        def zip_factory(file, mode="r", *args, **kwargs):
            archive = real_zip_file(file, mode, *args, **kwargs)
            return FailingTargetZip(archive) if mode in {"w", "x"} else archive

        with tempfile.TemporaryDirectory() as temporary:
            output = Path(temporary) / "partial.docx"
            with mock.patch.object(docx_export, "ZipFile", side_effect=zip_factory):
                with self.assertRaisesRegex(OSError, "synthetic write failure"):
                    docx_export.export_docx(model(), output)
            self.assertEqual(visibility_at_failure, [False])
            self.assertFalse(output.exists())
            self.assertEqual(list(Path(temporary).iterdir()), [])

    def test_first_private_temp_is_removed_when_second_creation_fails(self):
        """Catches temp ownership beginning only after both files exist."""
        real_named_temporary_file = docx_export.tempfile.NamedTemporaryFile
        created_paths = []

        def fail_second_creation(*args, **kwargs):
            if created_paths:
                raise OSError("synthetic second temp creation failure")
            handle = real_named_temporary_file(*args, **kwargs)
            created_paths.append(Path(handle.name))
            return handle

        with tempfile.TemporaryDirectory() as temporary:
            output = Path(temporary) / "never-published.docx"
            with mock.patch.object(
                docx_export.tempfile,
                "NamedTemporaryFile",
                side_effect=fail_second_creation,
            ):
                with self.assertRaisesRegex(
                    OSError, "synthetic second temp creation failure"
                ):
                    docx_export.export_docx(model(), output)

            self.assertEqual(len(created_paths), 1)
            self.assertFalse(output.exists())
            self.assertEqual(list(Path(temporary).glob("*.source.docx")), [])
            self.assertEqual(list(Path(temporary).glob("*.ready.docx")), [])

    def test_cleanup_continues_without_overriding_the_publish_error(self):
        """Catches one unlink failure masking the cause and skipping its peer."""
        real_unlink = docx_export.Path.unlink
        cleanup_attempts = []

        def fail_source_cleanup(path, *args, **kwargs):
            cleanup_attempts.append(path)
            if path.name.endswith(".source.docx"):
                raise OSError("synthetic source cleanup failure")
            return real_unlink(path, *args, **kwargs)

        with tempfile.TemporaryDirectory() as temporary:
            output = Path(temporary) / "never-published.docx"
            try:
                with mock.patch.object(
                    docx_export.os,
                    "link",
                    side_effect=OSError("synthetic publish failure"),
                ), mock.patch.object(
                    docx_export.Path,
                    "unlink",
                    autospec=True,
                    side_effect=fail_source_cleanup,
                ):
                    with self.assertRaisesRegex(OSError, "synthetic publish failure"):
                        docx_export.export_docx(model(), output)

                self.assertEqual(len(cleanup_attempts), 2)
                self.assertTrue(cleanup_attempts[0].name.endswith(".source.docx"))
                self.assertTrue(cleanup_attempts[1].name.endswith(".ready.docx"))
                self.assertFalse(cleanup_attempts[1].exists())
                self.assertFalse(output.exists())
            finally:
                for path in cleanup_attempts:
                    real_unlink(path, missing_ok=True)


if __name__ == "__main__":
    unittest.main()
