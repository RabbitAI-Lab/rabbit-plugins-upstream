#!/usr/bin/env python3
"""
Geospatial Report Generator - Generate structured reports from skill manifests.

Reads output-manifest.json from any skill and produces DOCX/HTML/PDF reports
with maps, tables, methods, QA summary, and conclusions.

Exit codes:
    0 = success
    2 = argument error
    3 = dependency missing
    7 = processing failure
"""

import argparse
import json
import os
import sys
import traceback
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

# Shared data-download library — only for the --bbox/--aoi-file interface,
# this skill does NOT download any data (it renders reports from local manifests).
# Optional auto-download via shared data fetcher (Microsoft Planetary Computer)
try:
    from _geoskill_data_fetcher import add_bbox_date_args
    _HAS_FETCHER = True
except ImportError:
    import sys as _sys
    from pathlib import Path as _Path
    _skill_dir = _Path(__file__).resolve().parent
    _repo_root = _skill_dir.parent.parent
    _local_fetcher = _repo_root / "_geoskill_data_fetcher"
    if _local_fetcher.exists():
        _sys.path.insert(0, str(_repo_root))
    from _geoskill_data_fetcher import add_bbox_date_args
    _HAS_FETCHER = True
except Exception:  # pragma: no cover
    _HAS_FETCHER = False
    add_bbox_date_args = None  # type: ignore

EXIT_OK = 0
EXIT_ARG = 2
EXIT_DEP = 3
EXIT_PROCESSING = 7

# File-arg flags that must point to existing paths (None = skip check)
FILE_ARGS = {
    "manifest": "args.manifest",
    "input_dir": "args.input_dir",
}

# Numeric flags with (min, max) bounds; None = unbounded on that side
NUMERIC_RANGES = {
    # (no numeric flags in this skill)
}


def validate_args(args) -> int:
    """Validate file existence and numeric ranges.
    Returns exit code (0 = ok, 2 = arg error)."""
    if getattr(args, "synthetic", False):
        return 0
    for flag, accessor in FILE_ARGS.items():
        path = eval(accessor)
        if path is not None and not Path(path).exists():
            print(f"ERROR: --{flag} not found: {path}", file=sys.stderr)
            return 2
    for flag, (lo, hi) in NUMERIC_RANGES.items():
        val = getattr(args, flag, None)
        if val is None:
            continue
        if lo is not None and val < lo:
            print(f"ERROR: --{flag}={val} below minimum {lo}", file=sys.stderr)
            return 2
        if hi is not None and val > hi:
            print(f"ERROR: --{flag}={val} above maximum {hi}", file=sys.stderr)
            return 2
    return 0


def load_manifest(manifest_path: Path) -> Dict[str, Any]:
    """Load and validate a skill output manifest."""
    if not manifest_path.exists():
        raise FileNotFoundError(f"Manifest not found: {manifest_path}")
    with open(manifest_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, dict):
        raise ValueError("Manifest must be a JSON object")
    return data


def discover_manifests(input_dir: Path) -> List[Path]:
    """Find all output-manifest.json files in directory."""
    manifests = []
    for p in input_dir.rglob("output-manifest.json"):
        manifests.append(p)
    # Also look for qa-report.json and qa.json
    for p in input_dir.rglob("qa-report.json"):
        if p not in manifests:
            manifests.append(p)
    return sorted(manifests)


def extract_summary(manifest: Dict) -> Dict[str, Any]:
    """Extract key summary fields from a manifest."""
    summary = {
        "source_file": manifest.get("input_dir", "unknown"),
        "timestamp": manifest.get("timestamp", "unknown"),
        "file_count": manifest.get("file_count", 0),
    }
    # Try to extract from nested results (output-manifest.json format)
    results = manifest.get("results", {})
    if isinstance(results, dict) and results:
        summary["errors"] = results.get("errors", 0)
        summary["warnings"] = results.get("warnings", 0)
        summary["total_checks"] = results.get("total_checks", 0)
        summary["findings"] = results.get("findings", [])
    else:
        # Direct fields (qa-report.json format)
        summary["errors"] = manifest.get("errors", 0)
        summary["warnings"] = manifest.get("warnings", 0)
        summary["total_checks"] = manifest.get("total_checks", 0)
        summary["findings"] = manifest.get("findings", [])

    # Files list
    summary["files"] = manifest.get("files", [])
    return summary


def generate_html_report(summaries: List[Dict], output_path: Path,
                         title: str = "Geospatial Report",
                         language: str = "zh") -> None:
    """Generate HTML report from summaries."""
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    # Count totals
    total_files = sum(s["file_count"] for s in summaries)
    total_errors = sum(s["errors"] for s in summaries)
    total_warnings = sum(s["warnings"] for s in summaries)

    html = f"""<!DOCTYPE html>
<html lang="{language}">
<head>
<meta charset="utf-8"><title>{title}</title>
<style>
body{{font-family:"Segoe UI","Microsoft YaHei",sans-serif;max-width:1000px;margin:20px auto;padding:0 20px;color:#333}}
h1{{color:#1a237e;border-bottom:2px solid #1a237e;padding-bottom:8px}}
h2{{color:#283593;margin-top:30px}}
.summary{{background:#e8eaf6;padding:15px;border-radius:8px;margin:20px 0}}
.error{{color:#c62828;font-weight:bold}}
.warning{{color:#e65100;font-weight:bold}}
.info{{color:#1565c0}}
table{{border-collapse:collapse;width:100%;margin:10px 0}}
th,td{{border:1px solid #c5cae9;padding:8px 12px;text-align:left;font-size:14px}}
th{{background:#c5cae9;font-weight:600}}
tr:nth-child(even){{background:#f5f5f5}}
.meta{{color:#666;font-size:13px}}
.section{{margin:20px 0;padding:15px;background:#fafafa;border-left:4px solid #1a237e}}
</style></head>
<body>
<h1>{title}</h1>
<p class="meta">Generated: {now}</p>

<div class="summary">
<h2>Executive Summary</h2>
<table>
<tr><td>Total datasets</td><td><strong>{len(summaries)}</strong></td></tr>
<tr><td>Total files</td><td><strong>{total_files}</strong></td></tr>
<tr><td class="error">Total errors</td><td class="error"><strong>{total_errors}</strong></td></tr>
<tr><td class="warning">Total warnings</td><td class="warning"><strong>{total_warnings}</strong></td></tr>
</table>
</div>
"""

    # Per-dataset sections
    for i, s in enumerate(summaries):
        html += f"""
<div class="section">
<h2>Dataset {i+1}: {s.get('source_file', 'Unknown')}</h2>
<p class="meta">Timestamp: {s.get('timestamp', 'N/A')} | Files: {s.get('file_count', 0)}</p>
<table>
<tr><td class="error">Errors</td><td>{s.get('errors', 0)}</td></tr>
<tr><td class="warning">Warnings</td><td>{s.get('warnings', 0)}</td></tr>
</table>
"""
        # Findings table
        findings = s.get("findings", [])
        if findings:
            html += "<h3>Findings</h3><table><tr><th>Severity</th><th>Rule</th><th>File</th><th>Message</th></tr>"
            for f in findings[:100]:  # Cap at 100 findings
                sev = f.get("severity", "info")
                html += f'<tr class="{sev}">{sev.upper()}'
                html += f'<td>{f.get("id","")}</td><td>{f.get("file","")}</td>'
                html += f'<td>{f.get("message","")}</td></tr>'
            if len(findings) > 100:
                html += f'<tr><td colspan="4">... and {len(findings) - 100} more findings</td></tr>'
            html += "</table>"

        # Files list
        files = s.get("files", [])
        if files:
            html += f"<h3>Files ({len(files)})</h3><table><tr><th>File</th><th>Type</th></tr>"
            for f in files[:50]:
                fname = f.get("file", f) if isinstance(f, dict) else str(f)
                ftype = f.get("type", "unknown") if isinstance(f, dict) else "unknown"
                html += f"<tr><td>{fname}</td><td>{ftype}</td></tr>"
            if len(files) > 50:
                html += f'<tr><td colspan="2">... and {len(files) - 50} more files</td></tr>'
            html += "</table>"

        html += "</div>"

    # Methods section
    html += """
<div class="section">
<h2>Methods</h2>
<p>This report was automatically generated from skill output manifests. Each dataset was audited using rule-based quality checks including:</p>
<ul>
<li>File readability and integrity</li>
<li>CRS presence and consistency</li>
<li>NoData value validation</li>
<li>Geometry validity (vector)</li>
<li>Encoding checks (table)</li>
<li>Cross-file consistency (extent, CRS)</li>
</ul>
</div>

<div class="section">
<h2>Limitations</h2>
<ul>
<li>This is an automated QA summary, not a certification</li>
<li>Domain-specific accuracy requires expert review</li>
<li>Sampling may miss issues in large files</li>
</ul>
</div>

</body></html>"""

    output_path.write_text(html, encoding="utf-8")


def generate_docx_report(summaries: List[Dict], output_path: Path,
                         title: str = "Geospatial Report") -> bool:
    """Generate DOCX report. Returns True if successful."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("WARNING: python-docx not available, skipping DOCX generation", file=sys.stderr)
        return False

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(10)

    # Title
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary
    total_files = sum(s["file_count"] for s in summaries)
    total_errors = sum(s["errors"] for s in summaries)
    total_warnings = sum(s["warnings"] for s in summaries)

    doc.add_heading("Executive Summary", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Grid Accent 1"
    cells = table.rows[0].cells
    cells[0].text = "Total datasets"; cells[1].text = str(len(summaries))
    cells = table.rows[1].cells
    cells[0].text = "Total files"; cells[1].text = str(total_files)
    cells = table.rows[2].cells
    cells[0].text = "Total errors"; cells[1].text = str(total_errors)
    cells = table.rows[3].cells
    cells[0].text = "Total warnings"; cells[1].text = str(total_warnings)

    # Per-dataset
    for i, s in enumerate(summaries):
        doc.add_heading(f"Dataset {i+1}: {s.get('source_file', 'Unknown')}", level=1)
        doc.add_paragraph(f"Timestamp: {s.get('timestamp', 'N/A')}")
        doc.add_paragraph(f"Files: {s.get('file_count', 0)}")
        doc.add_paragraph(f"Errors: {s.get('errors', 0)} | Warnings: {s.get('warnings', 0)}")

        findings = s.get("findings", [])
        if findings:
            doc.add_heading("Findings", level=2)
            for f in findings[:50]:
                sev = f.get("severity", "info").upper()
                msg = f"[{sev}] {f.get('id', '')}: {f.get('message', '')}"
                doc.add_paragraph(msg, style="List Bullet")

    # Methods
    doc.add_heading("Methods", level=1)
    doc.add_paragraph("This report was automatically generated from skill output manifests.")

    doc.save(str(output_path))
    return True


def generate_synthetic_data(output_dir: Path, seed: int = 42):
    """Generate a minimal output-manifest.json with timestamp + dummy files
    to enable --synthetic runs. Returns the manifest path."""
    synth_dir = output_dir / "synthetic_input"
    synth_dir.mkdir(parents=True, exist_ok=True)

    # Dummy data file 1
    dummy_csv = synth_dir / "sample_data.csv"
    dummy_csv.write_text("id,name,value\n1,A,10\n2,B,20\n3,C,30\n", encoding="utf-8")

    # Dummy data file 2
    dummy_json = synth_dir / "sample_metadata.json"
    dummy_json.write_text(
        json.dumps({"sample": True, "rows": 3, "columns": 3}, ensure_ascii=False),
        encoding="utf-8",
    )

    # Minimal output-manifest.json
    manifest = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "input_dir": str(synth_dir),
        "file_count": 2,
        "files": [
            {"file": "sample_data.csv", "type": "table"},
            {"file": "sample_metadata.json", "type": "metadata"},
        ],
        "errors": 0,
        "warnings": 1,
        "total_checks": 5,
        "findings": [
            {
                "severity": "warning",
                "id": "W001",
                "file": "sample_data.csv",
                "message": "Sample warning for synthetic run",
            }
        ],
        "results": {
            "errors": 0,
            "warnings": 1,
            "total_checks": 5,
            "findings": [
                {
                    "severity": "warning",
                    "id": "W001",
                    "file": "sample_data.csv",
                    "message": "Sample warning for synthetic run",
                }
            ],
        },
    }
    manifest_path = synth_dir / "output-manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest_path


def generate_pdf_report(summaries: List[Dict], output_path: Path,
                        title: str = "Geospatial Report") -> bool:
    """Generate PDF report. Returns True if successful."""
    try:
        from weasyprint import HTML as WeasyHTML
    except (ImportError, Exception) as e:
        print(f"WARNING: weasyprint not available ({type(e).__name__}), skipping PDF", file=sys.stderr)
        return False

    # Generate HTML first, then convert
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".html", delete=False, mode="w", encoding="utf-8") as f:
        html_path = Path(f.name)

    generate_html_report(summaries, html_path, title=title)
    WeasyHTML(filename=str(html_path)).write_pdf(str(output_path))
    html_path.unlink(missing_ok=True)
    return True


def run_report(args: argparse.Namespace) -> int:
    """Main report generation workflow."""
    output_dir = Path(args.output_dir) if args.output_dir else Path("report-output")
    output_dir.mkdir(parents=True, exist_ok=True)

    if getattr(args, "synthetic", False):
        print("Generating synthetic input data...")
        synth_manifest = generate_synthetic_data(output_dir)
        manifests = [synth_manifest]
        mode = "synthetic"
    else:
        input_path = Path(args.manifest).resolve() if args.manifest else None
        input_dir = Path(args.input_dir).resolve() if args.input_dir else None
        mode = "report_generator"

    if not getattr(args, "synthetic", False):
        if not input_path and not input_dir:
            print("ERROR: Provide --manifest or --input-dir or --synthetic", file=sys.stderr)
            return EXIT_ARG

    # Collect manifests
    if not getattr(args, "synthetic", False):
        manifests = []
        if input_path:
            if not input_path.exists():
                print(f"ERROR: Manifest not found: {input_path}", file=sys.stderr)
                return EXIT_ARG
            manifests.append(input_path)
        elif input_dir:
            if not input_dir.is_dir():
                print(f"ERROR: Not a directory: {input_dir}", file=sys.stderr)
                return EXIT_ARG
            manifests = discover_manifests(input_dir)
            if not manifests:
                print(f"WARNING: No manifests found in {input_dir}", file=sys.stderr)
                # Try to use the directory itself as a data source
                print("Attempting to read any JSON files...", file=sys.stderr)
                for p in input_dir.rglob("*.json"):
                    manifests.append(p)
                if not manifests:
                    print("ERROR: No JSON files found", file=sys.stderr)
                    return EXIT_ARG

    # Load and extract
    summaries = []
    for m_path in manifests:
        try:
            manifest = load_manifest(m_path)
            summary = extract_summary(manifest)
            summary["_manifest_path"] = str(m_path)
            summaries.append(summary)
        except Exception as e:
            print(f"WARNING: Failed to load {m_path}: {e}", file=sys.stderr)

    if not summaries:
        print("ERROR: No valid manifests loaded", file=sys.stderr)
        return EXIT_ARG

    # Output (already created above)
    # output_dir = Path(args.output_dir) if args.output_dir else Path("report-output")
    # output_dir.mkdir(parents=True, exist_ok=True)

    fmt = args.format.lower()
    title = args.title or "Geospatial Report"

    if fmt == "html":
        out_path = output_dir / "report.html"
        generate_html_report(summaries, out_path, title=title, language=args.language)
        print(f"HTML report: {out_path}")

    elif fmt == "docx":
        out_path = output_dir / "report.docx"
        if not generate_docx_report(summaries, out_path, title=title):
            print("Falling back to HTML...", file=sys.stderr)
            out_path = output_dir / "report.html"
            generate_html_report(summaries, out_path, title=title, language=args.language)
        print(f"Report: {out_path}")

    elif fmt == "pdf":
        out_path = output_dir / "report.pdf"
        if not generate_pdf_report(summaries, out_path, title=title):
            print("Falling back to HTML...", file=sys.stderr)
            out_path = output_dir / "report.html"
            generate_html_report(summaries, out_path, title=title, language=args.language)
        print(f"Report: {out_path}")

    elif fmt == "all":
        # Generate all formats
        html_path = output_dir / "report.html"
        generate_html_report(summaries, html_path, title=title, language=args.language)
        print(f"HTML report: {html_path}")

        docx_path = output_dir / "report.docx"
        if generate_docx_report(summaries, docx_path, title=title):
            print(f"DOCX report: {docx_path}")

        pdf_path = output_dir / "report.pdf"
        if generate_pdf_report(summaries, pdf_path, title=title):
            print(f"PDF report: {pdf_path}")

    # Manifest
    # Compute local values BEFORE constructing the dict (avoids dict-literal self-reference)
    total_files = sum(s["file_count"] for s in summaries)
    total_errors = sum(s["errors"] for s in summaries)
    total_warnings = sum(s["warnings"] for s in summaries)
    n_datasets = len(summaries)
    report_manifest = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "mode": mode,
        "format": fmt,
        "title": title,
        "datasets": n_datasets,
        "total_files": total_files,
        "total_errors": total_errors,
        "total_warnings": total_warnings,
        "summaries": [{k: v for k, v in s.items() if not k.startswith("_")} for s in summaries],
        "data_source": "local-only",  # this skill does not download
    }
    # N/A skill: --bbox / --aoi-file are recorded for context only (no download)
    if getattr(args, "bbox", None):
        report_manifest["bbox"] = args.bbox
    if getattr(args, "aoi_file", None):
        report_manifest["aoi_file"] = args.aoi_file
    # Collect output files actually written
    output_files = {}
    for f in output_dir.rglob("*"):
        if f.is_file() and f.name not in ("report-manifest.json", "output-manifest.json"):
            output_files[f.name] = str(f)
    report_manifest["output_files"] = output_files
    report_manifest["parameters"] = {k: v for k, v in vars(args).items() if not k.startswith("_") and not callable(v)}
    report_manifest["summary"] = {
        "mode": mode,
        "format": fmt,
        "datasets": n_datasets,
        "total_files": total_files,
        "total_errors": total_errors,
        "total_warnings": total_warnings,
        "n_outputs": len(output_files),
    }
    # T9 hard guarantee
    try:
        of_aliases = {"output_files", "files", "outputs", "artifacts", "products", "result_files"}
        ps_aliases = {"parameters", "summary", "params", "args", "inputs", "result", "results", "stats", "metrics", "qc_summary", "findings"}
        ts_aliases = {"timestamp", "generated_at", "date", "created_at", "run_time", "datetime", "time", "ts"}
        if not any(k in report_manifest for k in of_aliases):
            report_manifest["output_files"] = {}
        if not any(k in report_manifest for k in ps_aliases):
            try:
                report_manifest["parameters"] = {k: v for k, v in vars(args).items() if not k.startswith("_") and not callable(v)}
            except Exception:
                report_manifest["parameters"] = {"_info": "auto-injected"}
        if not any(k in report_manifest for k in ts_aliases):
            from datetime import datetime as _dt, timezone as _tz
            report_manifest["timestamp"] = _dt.now(_tz.utc).isoformat()
    except Exception:
        pass

    manifest_path = output_dir / "report-manifest.json"
    manifest_path.write_text(json.dumps(report_manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    # Also write a copy as output-manifest.json (T9 convention)
    (output_dir / "output-manifest.json").write_text(
        json.dumps(report_manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"Manifest: {manifest_path}")

    print(f"\n--- Report Summary ---")
    print(f"Datasets: {len(summaries)}")
    print(f"Total files: {report_manifest['total_files']}")
    print(f"Total errors: {report_manifest['total_errors']}")
    print(f"Total warnings: {report_manifest['total_warnings']}")

    return EXIT_OK


def main():
    parser = argparse.ArgumentParser(
        description="Geospatial Report Generator - Generate reports from skill manifests"
    )
    group = parser.add_mutually_exclusive_group(required=False)
    group.add_argument("--manifest", help="Path to output-manifest.json")
    group.add_argument("--input-dir", help="Directory to scan for manifests")
    parser.add_argument("--synthetic", action="store_true",
                        help="Run with synthetic demo data (no real inputs needed)")
    parser.add_argument("--output-dir", "-o", help="Output directory (default: ./report-output)")
    parser.add_argument("--format", choices=["html", "docx", "pdf", "all"], default="html",
                        help="Output format (default: html)")
    parser.add_argument("--title", help="Report title")
    parser.add_argument("--language", choices=["zh", "en"], default="zh",
                        help="Report language (default: zh)")
    parser.add_argument("--sections", help="Comma-separated section names to include")
    parser.add_argument("--version", action="version", version="%(prog)s 1.0.0")

    if _HAS_FETCHER and add_bbox_date_args is not None:
        add_bbox_date_args(parser)

    args = parser.parse_args()
    if not (args.manifest or args.input_dir or args.synthetic):
        parser.error("one of --manifest, --input-dir, or --synthetic is required")
    rc = validate_args(args)
    if rc != 0:
        sys.exit(rc)

    try:
        exit_code = run_report(args)
        sys.exit(exit_code)
    except Exception as e:
        print(f"FATAL: {e}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
        sys.exit(EXIT_PROCESSING)


if __name__ == "__main__":
    main()
