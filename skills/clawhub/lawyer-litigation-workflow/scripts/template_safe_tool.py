# -*- coding: utf-8 -*-
"""
模板安全替换工具 v1.0
SkillHub 适配版 — 从 template_safe_tool.py v3.3 精简适配。
强制规则:
  1. safe_set_para: 清空所有run → 写入全文到run[0] → 校验
  2. scan_and_clean: 遍历全部段落清除旧关键词残留
  3. verify_doc: 生成后完整性验证
  4. precheck_and_warn: 生成前骨架预检
"""

import os
import re
import shutil
import sys
from pathlib import Path
from docx import Document
from skeleton_checker import precheck_and_warn, sniff_doc_type

# 旧模板关键词黑名单
TEMPLATE_RESIDUE_KEYWORDS = [
    "中电科思仪", "广州辰创", "辰创科技", "方特远", "方特元",
    "和润达", "中辉尚德", "赵明明", "孙静", "丁晓涵", "杜丕志",
    "西岸兰海", "庚盛建设", "日照荣德", "珠光路", "兰东路",
    "青岛中辉", "青岛方特远", "秦延东", "乐信韵达",
    "孙威龙", "李润民", "薛总", "宋经理",
    "【】", "【  】", "【论点序号】", "【论点标题】", "【论点正文】",
    "【原告姓名】", "【被告姓名】", "【案由】", "【案号】",
    "【律所全称】", "【代理方角色】", "【当事人姓名】", "【律师姓名】",
    "【日期】", "【法院名称】",
]


def clear_all_runs(para):
    """清空段落所有run的文本"""
    for r in para.runs:
        r.text = ""


def safe_set_para(para, text, expected=None):
    """安全设置段落文本：
    1. 清空所有run
    2. 写入全文到run[0]
    3. 校验 para.text 包含 expected
    4. 校验不含【】占位符
    """
    if expected is None:
        expected = text[:80] if len(text) > 80 else text

    run_count = len(para.runs)
    clear_all_runs(para)

    if run_count == 0:
        para.add_run(text)
    else:
        para.runs[0].text = text

    # 完整性校验
    if expected not in para.text:
        raise RuntimeError(
            f"safe_set_para 校验失败: expected '{expected}' 不在 para.text '{para.text[:100]}'"
        )

    # 占位符残留校验
    placeholder_match = re.search(r"【[^】]{0,20}】", para.text)
    if placeholder_match:
        raise RuntimeError(
            f"safe_set_para 占位符残留: '{placeholder_match.group()}' in '{para.text[:80]}'"
        )

    return para


def scan_and_clean(doc, doc_type=""):
    """遍历全部段落，清除模板旧案件关键词残留"""
    issues = []
    for i, p in enumerate(doc.paragraphs):
        for kw in TEMPLATE_RESIDUE_KEYWORDS:
            if kw in p.text:
                old_text = p.text[:60]
                new_text = p.text.replace(kw, "")
                safe_set_para(p, new_text)
                issues.append(f"P{i}: 清除旧关键词 '{kw}' [{old_text}]")
                break
    if issues:
        print(f"[scan_and_clean] {doc_type}: 清除 {len(issues)} 处旧关键词残留")
        for issue in issues[:5]:
            print(f"  {issue}")
    return issues


def verify_doc(doc, doc_type=""):
    """生成后完整性验证"""
    full_text = "\n".join([p.text for p in doc.paragraphs])

    # 空文档检查
    if len(full_text.strip()) < 50:
        return {"passed": False, "reason": "文档内容过短（疑似生成失败）"}

    # 旧关键词残留检查
    residues = []
    for kw in TEMPLATE_RESIDUE_KEYWORDS[:15]:  # 仅检查前15个高频关键词
        if kw in full_text:
            residues.append(kw)

    if residues:
        return {"passed": False, "reason": f"旧关键词残留: {residues}"}

    return {"passed": True, "reason": "文档完整性通过"}


def verify_after_write(file_path, expected_min_kb=1):
    """Write后立即Read验证文件完整性"""
    if not os.path.exists(file_path):
        return {"passed": False, "reason": f"文件不存在: {file_path}"}

    size = os.path.getsize(file_path)
    if size < expected_min_kb * 1024:
        return {"passed": False, "reason": f"文件异常小: {size} bytes"}

    try:
        doc = Document(file_path)
        if len(doc.paragraphs) < 3:
            return {"passed": False, "reason": "段落数过少"}
    except Exception as e:
        return {"passed": False, "reason": f"无法打开文件: {e}"}

    return {"passed": True, "reason": "文件完整性通过"}


def open_template(template_path):
    """安全打开模板（带备份和校验）"""
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"模板文件不存在: {template_path}")

    doc = Document(template_path)
    par_count = len(doc.paragraphs)
    if par_count < 3:
        raise RuntimeError(f"模板段落异常: {par_count}段")

    print(f"[模板] 已加载: {os.path.basename(template_path)} ({par_count}段)")
    return doc


def guarded_generate(template_path, output_path, replacements, doc_type="", case_data_path=None):
    """安全生成主函数：
    1. 打开模板
    2. 逐段替换占位符
    3. scan_and_clean 清除残留
    4. precheck_and_warn 骨架预检
    5. 保存
    6. verify_after_write 文件完整性验证
    """
    doc = open_template(template_path)

    # 逐段逐run替换
    for p in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in p.text:
                old_text = p.text
                new_text = old_text.replace(placeholder, str(value))
                safe_set_para(p, new_text)
                break  # 每个段落只处理第一个匹配的占位符

    # 清除旧关键词残留
    scan_and_clean(doc, doc_type)

    # 生成前骨架预检
    precheck_and_warn(doc, doc_type, case_data_path)

    # 保存
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    # 写入后验证
    verify_result = verify_after_write(output_path)
    if not verify_result["passed"]:
        raise RuntimeError(f"文件写入验证失败: {verify_result['reason']}")

    print(f"[生成] {doc_type}: {output_path}")
    return output_path
