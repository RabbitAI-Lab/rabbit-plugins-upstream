#!/usr/bin/env python3
"""
文书质量强制检查器 v1.0
SkillHub 适配版 — 从 doc_quality_checker.py v2.4 精简适配。
任何一项不通过 → exit 1，禁止降级。

检查项:
  1. 模板来源（结构标记 + 段落数 + 旧关键词 + 占位符）
  2. 字体一致性
  3. 占位符残留
  4. 文本漂移
  5. 多Run完整性
  6. 内容骨架（skeleton_checker）
  7. 当事人覆盖（skeleton_checker）
"""

import sys
import re
import os
import json
from datetime import datetime
from skeleton_checker import check_skeleton, check_parties, sniff_doc_type

GLOBAL_FORBIDDEN_KEYWORDS = [
    "中电科思仪", "广州辰创", "辰创科技", "方特远", "方特元",
    "和润达", "中辉尚德", "赵明明", "孙静", "丁晓涵", "杜丕志",
    "西岸兰海", "庚盛建设", "日照荣德", "珠光路", "兰东路",
    "青岛中辉", "青岛方特远", "秦延东", "乐信韵达",
    "孙威龙", "李润民", "薛总", "宋经理",
]


def load_template_spec():
    return {
        "答辩状": {"struct_markers": ["答辩状", "答辩人", "此致"], "min_paragraphs": 8, "forbidden_keywords": GLOBAL_FORBIDDEN_KEYWORDS},
        "授权委托书": {"struct_markers": ["授", "权", "委", "托", "委托人", "受委托人", "委托权限"], "min_paragraphs": 10, "forbidden_keywords": GLOBAL_FORBIDDEN_KEYWORDS},
        "出庭函": {"struct_markers": ["律师参加诉讼", "贵院", "律师事务所"], "min_paragraphs": 5, "forbidden_keywords": GLOBAL_FORBIDDEN_KEYWORDS},
        "律师事务所函": {"struct_markers": ["律师参加诉讼", "贵院", "律师事务所"], "min_paragraphs": 5, "forbidden_keywords": GLOBAL_FORBIDDEN_KEYWORDS},
        "委托代理协议": {"struct_markers": ["委托代理协议", "甲方", "乙方", "代理事项", "代理权限"], "min_paragraphs": 50, "intentional_placeholders": 15, "allowed_font_mix": True, "forbidden_keywords": GLOBAL_FORBIDDEN_KEYWORDS},
        "起诉状": {"struct_markers": ["民事起诉状", "原告", "被告", "诉讼请求", "事实与理由", "此致"], "min_paragraphs": 10, "forbidden_keywords": GLOBAL_FORBIDDEN_KEYWORDS},
        "谈话笔录": {"struct_markers": ["律师接待当事人谈话笔录", "律师", "服务风险", "签字"], "min_paragraphs": 20, "forbidden_keywords": GLOBAL_FORBIDDEN_KEYWORDS},
        "证据目录": {"struct_markers": ["证据目录", "证据名称", "提交人"], "min_paragraphs": 5, "forbidden_keywords": GLOBAL_FORBIDDEN_KEYWORDS},
        "代理词": {"struct_markers": ["代理意见", "尊敬的审判长", "以上意见"], "min_paragraphs": 8, "forbidden_keywords": GLOBAL_FORBIDDEN_KEYWORDS},
    }


def check_template_origin(doc, doc_type, template_path=None):
    spec = load_template_spec()
    matched_spec = None
    for key, val in spec.items():
        if key in doc_type:
            matched_spec = val
            break
    if matched_spec is None:
        return {"passed": True, "reason": "无对应模板"}

    full_text = "\n".join([p.text for p in doc.paragraphs])
    par_count = len(doc.paragraphs)
    reasons = []

    markers_missing = [m for m in matched_spec["struct_markers"] if m not in full_text]
    par_ok = par_count >= matched_spec["min_paragraphs"]

    forbidden_found = [kw for kw in matched_spec.get("forbidden_keywords", []) if kw in full_text]

    placeholder_patterns = [
        r"【\s*】", r"【论点[^】]*】", r"【原告[^】]*】", r"【被告[^】]*】",
        r"【案由】", r"【案号】", r"【律师[^】]*】", r"【日期】",
        r"【法院[^】]*】", r"【律所[^】]*】",
    ]
    placeholder_found = []
    for pat in placeholder_patterns:
        placeholder_found.extend(re.findall(pat, full_text))

    all_ok = len(markers_missing) == 0 and par_ok and len(forbidden_found) == 0 and len(placeholder_found) == 0

    if markers_missing:
        reasons.append(f"缺少结构标记: {markers_missing}")
    if not par_ok:
        reasons.append(f"段落数不足: {par_count} < {matched_spec['min_paragraphs']}")
    if forbidden_found:
        reasons.append(f"旧关键词残留({len(forbidden_found)}): {forbidden_found[:3]}")
    if placeholder_found:
        reasons.append(f"占位符残留({len(placeholder_found)}): {placeholder_found[:5]}")

    return {"passed": all_ok, "reason": "; ".join(reasons) if reasons else "模板结构完整",
            "markers_missing": markers_missing, "paragraph_count": par_count,
            "forbidden_found": forbidden_found, "placeholder_found": placeholder_found}


def check_font_consistency(doc, doc_type=""):
    spec = load_template_spec().get(doc_type, {})
    if spec.get("allowed_font_mix", False):
        return {"passed": True, "drift_count": 0, "drifts": []}
    drifts = []
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue
        fonts_in_para = set()
        for run in p.runs:
            if run.text.strip():
                fname = run.font.name or "(inherit)"
                fonts_in_para.add(fname)
        if len(fonts_in_para) > 1:
            drifts.append({"paragraph": i, "text": p.text[:60], "fonts": list(fonts_in_para)})
    return {"passed": len(drifts) == 0, "drift_count": len(drifts), "drifts": drifts[:10]}


def check_placeholder_residue(doc, doc_type=""):
    spec = load_template_spec().get(doc_type, {})
    max_intentional = spec.get("intentional_placeholders", 0)
    residues = []
    patterns = [r"【[\s]*】", r"【[^】\n]{1,30}】"]
    for i, p in enumerate(doc.paragraphs):
        for pat in patterns:
            found = re.findall(pat, p.text)
            if found:
                residues.append({"paragraph": i, "text": p.text[:80], "count": len(found), "matches": found[:3]})
    total = sum(r["count"] for r in residues)
    return {"passed": total <= max_intentional, "residue_count": len(residues),
            "total_residues": total, "allowed": max_intentional, "residues": residues[:10]}


def check_text_drift(doc, doc_type=""):
    full_text = "\n".join([p.text for p in doc.paragraphs])
    indicators = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if len(t) > 100 and full_text.count(t) > 1:
            indicators.append(f"重复信息: {t[:50]}")
    for i, p in enumerate(doc.paragraphs):
        if len(p.text.strip()) > 800:
            indicators.append(f"异常长段落P{i}: {len(p.text)}字")
    return {"passed": len(indicators) == 0, "drift_count": len(indicators), "indicators": indicators[:10]}


def check_run_integrity(doc):
    issues = []
    truncation_keywords = ["本着友好", "甲乙双方本", "综上所述"]
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        if len(t) > 20 and not t.rstrip().endswith(("。", "）", "!", "：", "\n")):
            for kw in truncation_keywords:
                if t.endswith(kw):
                    issues.append(f"P{i}: 文本疑似截断 [{t[-40:]}]")
        non_empty_runs = [r for r in p.runs if r.text.strip()]
        if len(non_empty_runs) >= 2:
            for j in range(len(non_empty_runs) - 1):
                r1 = non_empty_runs[j].text.strip()
                r2 = non_empty_runs[j + 1].text.strip()
                if r1.endswith("本") and not r1.endswith("文本"):
                    issues.append(f"P{i} run{j}: 文本截断标记")
    return {"passed": len(issues) == 0, "issue_count": len(issues), "issues": issues[:10]}


def run_full_check(docx_path, doc_type="", template_path=None, case_data_path=None):
    from docx import Document

    if not os.path.exists(docx_path):
        return {"error": f"文件不存在: {docx_path}", "exit_code": 1}

    doc = Document(docx_path)

    if not doc_type:
        fname = os.path.basename(docx_path)
        if "答辩状" in fname:
            doc_type = "答辩状"
        elif "授权" in fname:
            doc_type = "授权委托书"
        elif "出庭" in fname or "律师事务" in fname:
            doc_type = "出庭函"
        elif "代理协议" in fname or "委托协议" in fname:
            doc_type = "委托代理协议"
        elif "起诉状" in fname:
            doc_type = "起诉状"
        elif "笔录" in fname or "谈话" in fname:
            doc_type = "谈话笔录"
        elif "证据" in fname:
            doc_type = "证据目录"
        elif "代理词" in fname or "代理意见" in fname:
            doc_type = "代理词"

        if not doc_type:
            full_text = "\n".join([p.text for p in doc.paragraphs])
            doc_type = sniff_doc_type(full_text)

    results = {
        "file": docx_path, "doc_type": doc_type,
        "check_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables),
    }

    results["template_origin"] = check_template_origin(doc, doc_type, template_path)
    results["font_consistency"] = check_font_consistency(doc, doc_type)
    results["placeholder_residue"] = check_placeholder_residue(doc, doc_type)
    results["text_drift"] = check_text_drift(doc, doc_type)
    results["run_integrity"] = check_run_integrity(doc)

    full_text = "\n".join([p.text for p in doc.paragraphs])
    results["content_skeleton"] = check_skeleton(full_text, doc_type)
    results["party_coverage"] = check_parties(full_text, doc_type, case_data_path)

    all_passed = all([
        results["template_origin"]["passed"],
        results["font_consistency"]["passed"],
        results["placeholder_residue"]["passed"],
        results["text_drift"]["passed"],
        results["run_integrity"]["passed"],
        results["content_skeleton"]["passed"],
        results["party_coverage"]["passed"],
    ])
    results["all_passed"] = all_passed

    print(format_report(results))
    return results


def format_report(results):
    lines = [
        "=" * 60,
        f"[QC] 文书质量检查 v1.0",
        f"   文件: {results['file']}",
        f"   类型: {results['doc_type']}  段落: {results['paragraph_count']}",
        "=" * 60,
    ]

    checks = [
        ("模板来源", results.get("template_origin", {})),
        ("字体一致", results.get("font_consistency", {})),
        ("占位符", results.get("placeholder_residue", {})),
        ("文本漂移", results.get("text_drift", {})),
        ("Run完整性", results.get("run_integrity", {})),
        ("内容骨架", results.get("content_skeleton", {})),
        ("当事人覆盖", results.get("party_coverage", {})),
    ]

    for name, check in checks:
        passed = check.get("passed", False)
        skipped = check.get("skipped", False)
        icon = "[SKIP]" if skipped else ("[OK]" if passed else "[FAIL]")
        detail = check.get("reason", "")
        if not detail and "issues" in check and check["issues"]:
            detail = check["issues"][0][:60]
        lines.append(f"   {icon} {name}: {detail}")

    lines.append("=" * 60)
    lines.append("[OK] 全部检查通过" if results.get("all_passed", False) else "[FAIL] 质量检查未通过！")
    lines.append("=" * 60)
    return "\n".join(lines)


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="文书质量强制检查器 v1.0")
    parser.add_argument("docx_path")
    parser.add_argument("--type", "-t", default="")
    parser.add_argument("--json", action="store_true")
    parser.add_argument("--case-data", default="")
    args = parser.parse_args()

    case_data_path = args.case_data if args.case_data else None
    if not case_data_path:
        docx_dir = os.path.dirname(os.path.abspath(args.docx_path))
        for parent_dir in [docx_dir, os.path.dirname(docx_dir), os.path.dirname(os.path.dirname(docx_dir))]:
            candidate = os.path.join(parent_dir, "案件数据.json")
            if os.path.exists(candidate):
                case_data_path = candidate
                break

    results = run_full_check(args.docx_path, args.type, None, case_data_path)
    if args.json:
        print(json.dumps(results, ensure_ascii=False, indent=2))
    sys.exit(0 if results.get("all_passed", False) else 1)
