# -*- coding: utf-8 -*-
"""
规格驱动文书生成管道 v1.0
SkillHub 适配版 — 从原版 specification_pipeline 精简适配。

管道五阶段:
  1. read_spec   — 读取规格 JSON
  2. enforce_spec — 强制执行不可变异元素
  3. generate    — 模板替换 + narrative_fill
  4. quality_check — QC 检查
  5. output      — 保存并验证
"""

import json
import os
import re
from pathlib import Path
from docx import Document
from template_safe_tool import safe_set_para, scan_and_clean, verify_after_write, guarded_generate


def load_spec(spec_name, specs_dir=None):
    """加载规格文件"""
    if specs_dir is None:
        specs_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "assets", "workflow_specs")

    spec_path = os.path.join(specs_dir, f"{spec_name}_spec.json")
    if not os.path.exists(spec_path):
        # 尝试模糊匹配
        for fname in os.listdir(specs_dir):
            if spec_name in fname and fname.endswith("_spec.json"):
                spec_path = os.path.join(specs_dir, fname)
                break

    if not os.path.exists(spec_path):
        raise FileNotFoundError(f"规格文件不存在: {spec_name}_spec.json (搜索路径: {specs_dir})")

    with open(spec_path, "r", encoding="utf-8") as f:
        return json.load(f)


def enforce_immutables(doc, spec):
    """强制执行不可变异元素（标题/法院/结尾等不可修改内容）"""
    immutables = spec.get("immutables", [])
    for item in immutables:
        marker = item.get("marker", "")
        text = item.get("text", "")
        for p in doc.paragraphs:
            if marker in p.text:
                safe_set_para(p, text)
                break
    return doc


def narrative_fill(doc, spec, case_data, output_dir):
    """根据案件数据填充叙事内容（事实描述/证据列表/论点区等）"""
    fill_config = spec.get("narrative_fill", {})

    for section_name, section_config in fill_config.items():
        content_source = section_config.get("source", "")
        template_placeholder = section_config.get("placeholder", "")

        if content_source.startswith("case_data."):
            # 从案件数据提取
            key_path = content_source.split(".", 1)[1]
            content = _get_nested(case_data, key_path, "")
        elif content_source == "generated":
            # 由 AI 生成（调用方通过 narrative_fill 参数传入）
            content = section_config.get("_generated_content", "")
        else:
            content = content_source

        if content and template_placeholder:
            for p in doc.paragraphs:
                if template_placeholder in p.text:
                    new_text = p.text.replace(template_placeholder, str(content))
                    safe_set_para(p, new_text)
                    break

    return doc


def _get_nested(data, key_path, default=""):
    """从嵌套字典中按路径取值，如 'parties.plt.name'"""
    keys = key_path.split(".")
    current = data
    for k in keys:
        if isinstance(current, dict):
            current = current.get(k, default)
        elif isinstance(current, list) and k.isdigit():
            idx = int(k)
            current = current[idx] if idx < len(current) else default
        else:
            return default
    return current if current else default


def run_spec_pipeline(template_path, spec_name, case_data, output_path, specs_dir=None, case_data_path=None):
    """
    规格驱动生成主函数

    Args:
        template_path: 模板 .docx 路径
        spec_name: 规格名称（如 "起诉状"）
        case_data: 案件数据 dict
        output_path: 输出 .docx 路径
        specs_dir: 规格文件目录
        case_data_path: 案件数据 JSON 路径（用于 QC 检查）

    Returns:
        output_path
    """
    # 1. Read spec
    spec = load_spec(spec_name, specs_dir)

    # 2. Read template
    doc = Document(template_path)

    # 3. Enforce immutables
    doc = enforce_immutables(doc, spec)

    # 4. Apply replacements from spec
    replacements = spec.get("replacements", {})
    resolved = {}
    for placeholder, source in replacements.items():
        if source.startswith("case_data."):
            key_path = source.split(".", 1)[1]
            resolved[placeholder] = _get_nested(case_data, key_path, placeholder)
        else:
            resolved[placeholder] = source

    # 5. Generate with safe replacements
    from doc_quality_checker import run_full_check

    output_path = guarded_generate(
        template_path, output_path, resolved, spec_name, case_data_path
    )

    # 6. Narrative fill
    doc = Document(output_path)
    doc = narrative_fill(doc, spec, case_data, os.path.dirname(output_path))
    scan_and_clean(doc, spec_name)
    doc.save(output_path)

    # 7. Quality check
    qc_result = run_full_check(output_path, spec_name, template_path, case_data_path)
    if not qc_result.get("all_passed", False):
        raise RuntimeError(f"规格驱动管道 [{spec_name}] QC 检查未通过")

    print(f"[规格管道] {spec_name}: {output_path}")
    return output_path


def generate_from_spec_only(spec_name, case_data, output_path, specs_dir=None):
    """
    无模板降级模式：仅从规格文件生成文书（当模板缺失时使用）
    生成的是纯文本格式，建议补齐模板以获得最佳格式。
    """
    spec = load_spec(spec_name, specs_dir)

    sections = []
    sections.append(spec.get("title", f"【{spec_name}】"))
    sections.append("")

    for section_name, section_info in spec.get("sections", {}).items():
        sections.append(section_info.get("label", section_name))
        content = ""
        source = section_info.get("source", "")
        if source.startswith("case_data."):
            content = _get_nested(case_data, source.split(".", 1)[1], "")
        elif source:
            content = source
        sections.append(content)
        sections.append("")

    if "footer" in spec:
        sections.append(spec["footer"])

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(sections))

    print(f"[规格管道-降级] {spec_name}: {output_path} (纯文本，建议补齐模板)")
    return output_path
