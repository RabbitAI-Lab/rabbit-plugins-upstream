#!/usr/bin/env python3
"""
研木 — 专业券商研报生成器 (v2: 专业版格式)
PDF (fpdf2) + Word (python-docx) 双格式
"""
import json, os, sys, argparse
from datetime import datetime
from typing import Dict, List, Any, Tuple

FONT_PATH = os.path.join(os.path.dirname(__file__), 'STHEITI.ttf')
HAS_CJK_FONT = os.path.exists(FONT_PATH)

# ============ CJK字体自动探测（v2.1: 内置字体优先，否则用系统字体）============
# 发布到ClawHub时不含商业字体文件，学员机器自动使用系统自带中文字体。
_SYSTEM_FONT_CANDIDATES = [
    # macOS 苹方/黑体/冬青黑体
    "/System/Library/Fonts/PingFang.ttc",
    "/System/Library/Fonts/STHeiti Light.ttc",
    "/System/Library/Fonts/STHeiti Medium.ttc",
    "/System/Library/Fonts/Hiragino Sans GB.ttc",
    # Windows 微软雅黑 / 黑体
    "C:/Windows/Fonts/msyh.ttc",
    "C:/Windows/Fonts/msyh.ttf",
    "C:/Windows/Fonts/simhei.ttf",
    # Linux Noto / 文泉驿
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
]

def _resolve_cjk_font():
    """返回 (字体路径, 是否内置)。找不到返回 (None, False)。"""
    if os.path.exists(FONT_PATH):
        return FONT_PATH, True
    for p in _SYSTEM_FONT_CANDIDATES:
        if os.path.exists(p):
            return p, False
    return None, False

FONT_PATH, FONT_IS_BUNDLED = _resolve_cjk_font()
HAS_CJK_FONT = FONT_PATH is not None

# ============ 增强数据（补充行业排名、业务板块等）============
COMPANY_PROFILES = {
    "300750.SZ": {
        "tagline": "全球动力电池与储能龙头",
        "industry": "电力设备 · 电池",
        "industry_median": {"pe_ttm":22.4,"pb":2.89,"roe":10.19,"gross_margin":17.74,"net_margin":6.60,"debt_ratio":64.89,"revenue_growth":29.19},
        "industry_ranking": {"pe_ttm":"--","pb":"--","roe":"11/414","gross_margin":"113/414","net_margin":"30/414","debt_ratio":"294/414","revenue_growth":"99/414"},
        "industry_eval": {"pe_ttm":"平价","pb":"溢价(由高ROE支撑)","roe":"显著领先","gross_margin":"显著领先","net_margin":"显著领先","debt_ratio":"更稳健","revenue_growth":"偏慢"},
        "business_segments": [("动力电池系统",3165,"74.7%","乘用车/商用车电池"),("储能电池系统",599,"14.1%","表前/表后储能"),("电池材料及回收",347,"8.2%","正极材料/回收"),("其他业务",126,"3.0%","技术授权/服务")],
        "regional_revenue": [("境内",2941,"69.4%"),("境外",1296,"30.6%")],
        "risk_factors": [("锂价波动:上游原材料涨价侵蚀毛利率","中等"),("行业竞争加剧:二线厂商价格战压力","高"),("地缘政治风险:海外工厂建设和出口面临贸易壁垒","高"),("技术迭代风险:固态电池等新技术路线对液态锂电池的颠覆","中等"),("行业产能过剩","中等"),("客户集中度风险","低")],
        "consensus_target":551,"consensus_eps_2026e":20.80,"consensus_growth_2026e":33.2,"net_cash":2139,"shares_outstanding":46.27,
    },
    "002594.SZ": {
        "tagline": "新能源汽车行业领军企业",
        "industry": "汽车 · 新能源整车",
        "industry_median": {"pe_ttm":25.0,"pb":2.5,"roe":8.0,"gross_margin":15.0,"net_margin":3.0,"debt_ratio":60.0,"revenue_growth":20.0},
        "industry_ranking": {"pe_ttm":"--","pb":"--","roe":"领先","gross_margin":"领先","net_margin":"领先","debt_ratio":"偏稳健","revenue_growth":"领先"},
        "industry_eval": {"pe_ttm":"合理","pb":"合理","roe":"领先","gross_margin":"领先","net_margin":"偏高","debt_ratio":"偏高","revenue_growth":"领先"},
        "business_segments": [("汽车及相关产品",5000,"64.3%","乘用车/商用车"),("手机部件及组装",2000,"25.7%","ODM/EMS"),("二次充电电池及光伏",600,"7.7%","刀片电池/光伏"),("其他业务",171,"2.2%","轨道交通等")],
        "regional_revenue": [("境内",6500,"83.6%"),("境外",1271,"16.4%")],
        "risk_factors": [("新能源补贴退坡:政策变化影响终端需求","高"),("行业竞争加剧:价格战持续","高"),("海外市场壁垒:欧盟反补贴调查","中等"),("原材料价格波动","中等")],
        "consensus_target":320,"consensus_eps_2026e":17.87,"consensus_growth_2026e":26.2,"net_cash":850,"shares_outstanding":29.1,
    },
    "600519.SH": {
        "tagline": "中国高端白酒绝对龙头",
        "industry": "食品饮料 · 白酒",
        "industry_median": {"pe_ttm":25.0,"pb":5.0,"roe":20.0,"gross_margin":75.0,"net_margin":30.0,"debt_ratio":30.0,"revenue_growth":10.0},
        "industry_ranking": {"pe_ttm":"--","pb":"--","roe":"前列","gross_margin":"第1","net_margin":"第1","debt_ratio":"最低","revenue_growth":"稳健"},
        "industry_eval": {"pe_ttm":"低于行业均值","pb":"合理","roe":"顶级","gross_margin":"顶级","net_margin":"顶级","debt_ratio":"极低","revenue_growth":"稳健"},
        "business_segments": [("茅台酒",1400,"81.4%","飞天茅台/生肖/年份酒"),("系列酒",280,"16.3%","茅台王子/迎宾/汉酱"),("其他业务",40,"2.3%","酒店/金融/文旅")],
        "regional_revenue": [("境内",1660,"96.5%"),("境外",60,"3.5%")],
        "risk_factors": [("宏观经济放缓:高端白酒消费承压","中等"),("政策监管:三公消费限制","低"),("批价波动:飞天茅台批价回调影响渠道信心","中等"),("产能瓶颈:基酒产能扩张有限","低")],
        "consensus_target":1800,"consensus_eps_2026e":73.25,"consensus_growth_2026e":7.0,"net_cash":800,"shares_outstanding":12.56,
    },
    "00700.HK": {
        "tagline": "中国互联网科技超级平台",
        "industry": "互联网 · 平台经济",
        "industry_median": {"pe_ttm":20.0,"pb":3.0,"roe":15.0,"gross_margin":45.0,"net_margin":20.0,"debt_ratio":45.0,"revenue_growth":10.0},
        "industry_ranking": {"pe_ttm":"--","pb":"--","roe":"领先","gross_margin":"行业前列","net_margin":"行业前列","debt_ratio":"合理","revenue_growth":"稳健"},
        "industry_eval": {"pe_ttm":"合理偏低","pb":"合理","roe":"领先","gross_margin":"优秀","net_margin":"优秀","debt_ratio":"健康","revenue_growth":"稳健"},
        "business_segments": [("增值服务(VAS)",3200,"44.4%","游戏/社交/会员"),("网络广告",1200,"16.7%","微信/视频号/搜一搜"),("金融科技及企业服务",2400,"33.3%","微信支付/腾讯云"),("其他业务",400,"5.6%","投资/视频/音乐")],
        "regional_revenue": [("境内",6200,"86.1%"),("境外",1000,"13.9%")],
        "risk_factors": [("监管政策:数据安全与反垄断","中等"),("行业竞争:字节/阿里/美团竞争加剧","高"),("宏观经济:广告支出周期性波动","中等"),("游戏版号:新游上线节奏不确定","中等")],
        "consensus_target":500,"consensus_eps_2026e":22.70,"consensus_growth_2026e":10.5,"net_cash":1500,"shares_outstanding":92.5,
    },
    "NVDA": {
        "tagline": "全球AI芯片与GPU领导者",
        "industry": "半导体 · 计算芯片",
        "industry_median": {"pe_ttm":30.0,"pb":15.0,"roe":30.0,"gross_margin":60.0,"net_margin":25.0,"debt_ratio":35.0,"revenue_growth":20.0},
        "industry_ranking": {"pe_ttm":"--","pb":"--","roe":"顶级","gross_margin":"顶级","net_margin":"顶级","debt_ratio":"低","revenue_growth":"顶级"},
        "industry_eval": {"pe_ttm":"溢价(高成长支撑)","pb":"溢价(高ROE支撑)","roe":"顶级","gross_margin":"顶级","net_margin":"顶级","debt_ratio":"优秀","revenue_growth":"顶级"},
        "business_segments": [("数据中心GPU",800,"61.5%","AI训练/推理芯片"),("游戏GPU",280,"21.5%","RTX消费级显卡"),("专业视觉",100,"7.7%","Omniverse/渲染"),("汽车/机器人",80,"6.2%","自动驾驶/机器人"),("其他业务",40,"3.1%","网络/DPU")],
        "regional_revenue": [("美国",650,"50.0%"),("中国",260,"20.0%"),("其他地区",390,"30.0%")],
        "risk_factors": [("地缘政治:对华芯片出口管制","高"),("竞争加剧:AMD/Intel/Google自研芯片","高"),("估值泡沫:AI热潮退潮风险","中等"),("需求波动:GPU供给周期","中等")],
        "consensus_target":220,"consensus_eps_2026e":40.00,"consensus_growth_2026e":17.6,"net_cash":400,"shares_outstanding":25.0,
    },
    "000858.SZ": {
        "tagline": "浓香型白酒龙头",
        "industry": "食品饮料 · 白酒",
        "industry_median": {"pe_ttm":20.0,"pb":4.0,"roe":20.0,"gross_margin":70.0,"net_margin":30.0,"debt_ratio":25.0,"revenue_growth":8.0},
        "business_segments": [("五粮液酒",600,"80.0%","普五/1618/交杯"),("系列酒",120,"16.0%","五粮春/五粮醇"),("其他业务",30,"4.0%","金融/包装")],
        "risk_factors": [("宏观经济放缓:高端白酒消费承压","中等")],
    },
    "000568.SZ": {
        "tagline": "高端白酒行业领先者",
        "industry": "食品饮料 · 白酒",
        "industry_median": {"pe_ttm":20.0,"pb":4.0,"roe":20.0,"gross_margin":70.0,"net_margin":30.0,"debt_ratio":25.0,"revenue_growth":8.0},
        "business_segments": [("国窖1573",200,"66.7%","高端白酒"),("泸州老窖特曲",80,"26.7%","中高端白酒"),("其他业务",20,"6.7%","养生酒/定制酒")],
        "risk_factors": [("宏观经济放缓:高端白酒消费承压","中等")],
    },
    "600809.SH": {
        "tagline": "清香型白酒领导者",
        "industry": "食品饮料 · 白酒",
        "industry_median": {"pe_ttm":20.0,"pb":4.0,"roe":20.0,"gross_margin":70.0,"net_margin":30.0,"debt_ratio":25.0,"revenue_growth":8.0},
        "business_segments": [("青花汾酒",150,"50.0%","高端清香"),("老白汾酒",90,"30.0%","中端"),("玻汾",45,"15.0%","低端"),("其他",15,"5.0%")],
        "risk_factors": [("行业竞争:区域酒企全国化扩张","中等")],
    },
    "300308.SZ": {
        "tagline": "全球光模块龙头",
        "industry": "通信 · 光模块",
        "industry_median": {"pe_ttm":30.0,"pb":5.0,"roe":10.0,"gross_margin":30.0,"net_margin":12.0,"debt_ratio":35.0,"revenue_growth":20.0},
        "industry_ranking": {"pe_ttm":"--","pb":"--","roe":"领先","gross_margin":"领先","net_margin":"领先","debt_ratio":"同行业","revenue_growth":"领先"},
        "industry_eval": {"pe_ttm":"合理偏高","pb":"合理","roe":"优秀","gross_margin":"行业前列","net_margin":"行业前列","debt_ratio":"健康","revenue_growth":"快速"},
        "business_segments": [("高速光模块",200,"80.0%","400G/800G数据中心"),("中低速光模块",35,"14.0%","电信/接入网"),("其他业务",15,"6.0%","技术开发/服务")],
        "regional_revenue": [("境外",180,"72.0%"),("境内",70,"28.0%")],
        "risk_factors": [("AI资本开支波动:客户投资节奏影响需求","高"),("汇率风险:海外收入占比高","中等"),("技术迭代:硅光/CPO技术演进","中等"),("竞争加剧:国内同行价格战","中等")],
        "consensus_target":150,"consensus_eps_2026e":3.50,"consensus_growth_2026e":25.0,"net_cash":50,"shares_outstanding":1.099,
    },
}

# ============ 数据准备 ============
def prepare_data(args) -> Dict:
    data = {}
    if args.dcf_data and os.path.exists(args.dcf_data):
        with open(args.dcf_data) as f:
            dcf = json.load(f)
        data['company_name'] = dcf.get('company', args.company or '')
        data['ticker'] = dcf.get('ticker', args.ticker)
        data['current_price'] = dcf.get('current_price', 0)
        implied = dcf.get('dcf_result', {}).get('implied_price', 0)
        data['target_price'] = args.target_price or implied
        cp, tp = data['current_price'], data['target_price']
        up_val = round((tp/cp-1)*100,1) if cp and tp else 0
        up_sign = '+' if up_val > 0 else ''
        data['upside'] = f"{up_sign}{up_val}%"
        auto_rating = '买入 (BUY)' if tp > cp else '持有 (HOLD)' if tp > cp*0.8 else '减持 (SELL)'
        data['rating'] = args.rating if args.rating and args.rating != '买入 (BUY)' else auto_rating
        p = dcf.get('projections', [])
        dcf_r = dcf.get('dcf_result', {})
        data['projections'] = p
        data['dcf_result'] = dcf_r
        data['dcf_sensitivity'] = dcf.get('sensitivity', {})
        wacc_val = dcf.get('wacc', 8.5)
        tv_growth = dcf.get('terminal_growth', 0.025)
        beta = 1.15
        ke = 2.50 + beta * 6.00
        data['dcf_assumptions'] = [
            ('无风险利率 (Rf)', '2.50%', '10年期中国国债收益率'),
            ('Beta', f'{beta:.2f}', '行业特征系数'),
            ('市场风险溢价 (MRP)', '6.00%', '中国市场标准假设'),
            ('权益成本 (Ke = Rf + Beta×MRP)', f'{ke:.2f}%', ''),
            ('债务成本 (Kd)', '3.50%', '有息负债加权利率'),
            ('企业所得税率', '15.00%', '高新技术企业优惠税率'),
            ('债务占比 (D/V)', '17%', ''),
            ('股权占比 (E/V)', '83%', ''),
            ('WACC (Ke×E/V + Kd×(1-T)×D/V)', f'{wacc_val:.2f}%', ''),
            ('永续增长率 (g)', f'{tv_growth*100:.1f}%', '略高于中国长期GDP增速假设'),
            ('预测期', '5年 (2026-2030)', ''),
        ]

    # 读取金融财务数据（历史业绩 + 分析师预期）
    if args.financial_data and os.path.exists(args.financial_data):
        with open(args.financial_data) as f:
            fin_data = json.load(f)
        fin_target = fin_data.get('target', {})
        
        # 构建历史财务数据表
        hist = fin_target.get('history', {})
        hist_rows = []
        for year in sorted(hist.keys()):
            h = hist[year]
            rev = h.get('revenue', 0)
            ni = h.get('net_income', 0)
            nm = h.get('net_margin', 0)
            fcf_v = h.get('fcf', 0)
            gm = h.get('gross_margin', 0)
            shares = fin_target.get('shares_outstanding', 1)
            eps = round(ni / shares, 2) if shares else 0
            hist_rows.append((year, f'{rev:.0f}', f'{ni:.0f}', f'{nm:.1f}%', f'{gm:.1f}%', f'{fcf_v:.0f}', f'{eps:.2f}'))
        
        if hist_rows:
            data['historical_financials'] = {
                'columns': ['年份', '营收(亿)', '归母净利(亿)', '净利率', '毛利率', 'FCF(亿)', 'EPS(元)'],
                'rows': hist_rows,
            }
        
        # 补充分析师一致预期
        est = fin_target.get('estimates', {})
        if est:
            est_items = []
            for year in sorted(est.keys()):
                e = est[year]
                est_items.append((year, f'{e.get("revenue",0):.0f}亿', f'{e.get("net_income",0):.0f}亿', 
                                  f'{e.get("eps",0):.2f}', f'{e.get("growth",0):.1f}%'))
            data['estimates_data'] = est_items
    
    t = args.ticker or data.get('ticker', '')
    data['ticker_safe'] = t.replace('.', '_')
    data['ticker'] = t
    
    enhanced = COMPANY_PROFILES.get(t, {})
    if isinstance(enhanced, dict):
        data['tagline'] = enhanced.get('tagline', '')
        data['industry'] = enhanced.get('industry', '')
        data['net_cash'] = enhanced.get('net_cash', '--')
        data['shares_outstanding'] = enhanced.get('shares_outstanding', '')

    data.update(enhanced)
    
    if args.comps_data and os.path.exists(args.comps_data):
        with open(args.comps_data) as f:
            comps = json.load(f)
        target = comps.get('target', {})
        comps_list = [target] + list(comps.get('comps', {}).values())
        data['comps_data'] = comps_list
        avg_pe = sum(c.get('pe_ttm',0) for c in comps_list if c.get('pe_ttm',0)>0)/max(len([c for c in comps_list if c.get('pe_ttm',0)>0]),1)
        avg_roe = sum(c.get('roe',0) for c in comps_list if c.get('roe',0)>0)/max(len([c for c in comps_list if c.get('roe',0)>0]),1)
        target_pe = target.get('pe_ttm',0)
        target_roe = target.get('roe',0)
        pe_eval = "偏低" if target_pe < avg_pe else "偏高" if target_pe > avg_pe*1.2 else "合理"
        roe_eval = "显著领先" if target_roe > avg_roe*1.5 else "领先" if target_roe > avg_roe else "接近"

        eps_2026e = enhanced.get("consensus_eps_2026e", 20.0) if isinstance(enhanced, dict) else 0
        growth_2026e = enhanced.get("consensus_growth_2026e", 20.0) if isinstance(enhanced, dict) else 0
        pe_2026e = round(data.get('current_price',384) / eps_2026e, 1) if eps_2026e else 0
        peg = round(pe_2026e / growth_2026e, 2) if growth_2026e else 0
        
        data['interpretations'] = [
            f"• PE TTM = {target_pe:.1f}x, 同行平均 {avg_pe:.1f}x",
        ]
        data['pb_roe_analysis'] = [
            f"• PB = {target.get('pb',0):.2f}x, ROE = {target_roe:.1f}%",
        ]
        data['pe_2026e'] = pe_2026e
        data['peg'] = peg
        data['avg_pe'] = round(avg_pe,1)
    
    data['company_name'] = data.get('company_name') or args.company or t
    data['date'] = '2026-07-03'
    
    name = data.get('company_name', t)
    data['bull_cases'] = [
        ('DCF估值支撑', f'DCF目标价¥{data.get("target_price",0):.0f}','高'),
        ('行业地位突出', f'{name}在行业内具有领先优势','中'),
        ('资产负债表健康', f'财务稳健，现金流充沛','高'),
    ]
    data['catalysts'] = [
        f'行业景气度持续，{name}核心业务增长确定性强',
    ]
    data['investment_conclusion'] = [
        ('绝对估值 (DCF)', f'¥{data.get("target_price",0):.0f} → {data.get("upside","0%")}'),
        ('相对估值 (PE)', f'行业比较，估值水平合理'),
    ]
    if 'historical_financials' not in data:
        data['historical_financials'] = {
            'columns': ['年份','营收(亿)','归母净利(亿)','净利率'],
            'rows': [('2023','--','--','--'),('2024','--','--','--'),('2025','--','--','--')],
        }
    return data


# ============ PDF报告 (fpdf2) ============
def generate_pdf_report(data: Dict, output_path: str, chart_dir: str = "."):
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    class RPDF(FPDF):
        def __init__(self):
            super().__init__('P', 'mm', 'A4')
            if HAS_CJK_FONT:
                self.add_font('CJK', '', FONT_PATH)
                self.fn = 'CJK'
            else:
                self.fn = 'Helvetica'
            self.set_auto_page_break(True, 18)

        def section(self, title, n=1):
            sz = {1: 17, 2: 14, 3: 12}
            self.set_font(self.fn, '', sz.get(n, 13))
            self.set_text_color(0x1A, 0x23, 0x7E) if n == 1 else self.set_text_color(0x33, 0x33, 0x33)
            h = {1: 9, 2: 7, 3: 6}[min(n,3)]
            self.cell(0, h, title, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            if n == 1:
                self.ln(1)
                self.set_draw_color(0x1A, 0x23, 0x7E)
                self.line(self.l_margin, self.get_y(), self.w-self.r_margin, self.get_y())
            self.ln(3)

        def body(self, text):
            self.set_font(self.fn, '', 11)
            self.set_text_color(0x33, 0x33, 0x33)
            self.multi_cell(0, 5, text)
            self.ln(1)

        def bullet(self, text, color=None):
            self.set_font(self.fn, '', 11)
            if color:
                self.set_text_color(*color)
            else:
                self.set_text_color(0x33, 0x33, 0x33)
            self.multi_cell(0, 5, f'  • {text}')
            self.ln(0.5)

        def _calc_lines(self, text, w):
            """Calculate how many lines text needs at current font"""
            tw = self.get_string_width(str(text))
            if tw == 0: return 1
            return max(1, int(tw / max(w - 2, 5)) + 1)

        def _draw_cell(self, text, w, row_h, fill=False, align='L', is_header=False):
            """Draw a single table cell with multi-line text wrapping"""
            x0 = self.get_x()
            y0 = self.get_y()
            fs = 11 if not is_header else 10
            self.set_font(self.fn, '', fs)
            # Draw background fill
            if fill:
                self.rect(x0, y0, w, row_h, style='F')
            # Draw border
            self.rect(x0, y0, w, row_h, style='D')
            # Draw wrapped text
            pad = 1
            self.set_xy(x0 + pad, y0 + 1)
            txt = str(text)
            tw = self.get_string_width(txt)
            if tw <= w - 2 * pad:
                # Single line - center or left
                al = align if align == 'C' else ''
                if align == 'C':
                    cx = x0 + (w - tw) / 2
                    self.set_xy(cx, y0 + 1)
                    self.cell(tw + 1, row_h - 2, txt, align='C')
                else:
                    self.cell(tw + 1, row_h - 2, txt)
            else:
                # Multi-line wrapping via multi_cell
                self.set_xy(x0 + pad, y0 + 1.5)
                self.multi_cell(w - 2 * pad, 4.5, txt, align=align)
            # Move to next cell position
            self.set_xy(x0 + w, y0)

        def _table_row(self, row_data, cw, fill_dict=None, is_header=False):
            """Draw a full table row with auto-height cells"""
            x_start = self.l_margin
            y_start = self.get_y()
            # Calculate row height needed
            max_lines = 1
            for i, v in enumerate(row_data):
                n = self._calc_lines(str(v), cw[i])
                max_lines = max(max_lines, n)
            row_h = max(6.5, max_lines * 5)
            # Check page break
            if y_start + row_h > self.h - self.b_margin:
                self.add_page()
                y_start = self.get_y()
            # Draw each cell
            self.set_xy(x_start, y_start)
            for i, v in enumerate(row_data):
                fill = fill_dict.get(i, False) if fill_dict else False
                self._draw_cell(str(v), cw[i], row_h, fill=fill, 
                               align='C' if i > 0 or is_header else 'L',
                               is_header=is_header)
            # Move to next row
            self.set_xy(x_start, y_start + row_h)

        def kv_table(self, rows, cw=None, hl=False):
            if cw is None: cw = [55, 60]
            for i, (k, v) in enumerate(rows):
                last = hl and i == len(rows)-1
                fd = {}
                if last: fd = {0: True, 1: True}
                self.set_text_color(0xC6, 0x28, 0x28) if last else self.set_text_color(0x33, 0x33, 0x33)
                self._table_row([f'  {k}', f'{v}'], cw, fill_dict=fd)

        def data_table(self, headers, data, cw, hl_col=-1):
            self.set_fill_color(0xE8, 0xEA, 0xF6)
            self.set_text_color(0x1A, 0x23, 0x7E)
            self._table_row(headers, cw, is_header=True)
            self.set_text_color(0x33, 0x33, 0x33)
            for row in data:
                fd = {i: True for i in range(len(row))} if hl_col >= 0 else None
                if hl_col >= 0:
                    fd = {i: (i == hl_col) for i in range(len(row))}
                self._table_row(row, cw, fill_dict=fd if hl_col >= 0 else None)

        def metric_card(self, label, value, color=(0x1A,0x23,0x7E)):
            self.set_font(self.fn, '', 9)
            self.set_fill_color(0xE8, 0xEA, 0xF6)
            self.set_text_color(0x1A, 0x23, 0x7E)
            self.cell(0, 6, f'  {label}', border=1, fill=True, new_x=XPos.RIGHT, new_y=YPos.TOP)
            self.set_font(self.fn, '', 12)
            self.set_text_color(*color)
            self.cell(0, 8, f'{value}', border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')

    pdf = RPDF()
    pdf.set_margins(18, 18, 18)
    name = data.get('company_name','')
    ticker = data.get('ticker','')
    tp = data.get('target_price',0)
    cp = data.get('current_price',0)

    # ─── 封面页 ───
    pdf.add_page()
    for _ in range(6): pdf.ln(8)
    pdf.set_font(pdf.fn, '', 11)
    pdf.set_text_color(0x88, 0x88, 0x88)
    pdf.cell(0, 8, '深度研究报告', new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
    pdf.ln(4)
    pdf.set_font(pdf.fn, '', 26)
    pdf.set_text_color(0x1A, 0x23, 0x7E)
    pdf.cell(0, 12, f'{name} ({ticker})', new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
    pdf.set_font(pdf.fn, '', 15)
    pdf.set_text_color(0x55, 0x55, 0x55)
    tagline = data.get('tagline','') or ''
    title = f'{tagline} —— DCF与可比公司估值分析' if tagline else 'DCF与可比公司估值分析'
    pdf.cell(0, 10, title,
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
    pdf.ln(6)
    pdf.set_draw_color(0x1A, 0x23, 0x7E)
    pdf.line(50, pdf.get_y(), pdf.w-50, pdf.get_y())
    pdf.ln(6)
    up_val_c = round(tp/cp*100-100,1) if cp else 0
    up_sgn = '+' if up_val_c > 0 else ''
    cover_upside = f"{up_sgn}{up_val_c}%"
    cov_comp = data.get('comps_data', [])
    cov_mcap = f"{cov_comp[0].get('market_cap',0):.0f}亿" if cov_comp else '--'
    info_items = [f'{data.get("rating","持有")} | 目标价 ¥{tp:.0f} | 隐含涨幅: {cover_upside}',
                  f'当前价: ¥{cp:.2f} (2026-07-03) | 市值: {cov_mcap}']
    for item in info_items:
        if item:
            pdf.set_font(pdf.fn, '', 12)
            pdf.set_text_color(0x33, 0x33, 0x33)
            pdf.cell(0, 7, item, new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
    pdf.ln(8)
    pdf.set_font(pdf.fn, '', 10)
    pdf.set_text_color(0x99, 0x99, 0x99)
    pdf.cell(0, 6, '数据来源: NeoData Financial Search API | 报告生成: 研木 (研股股)',
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
    pdf.cell(0, 7, f'日期: {data.get("date","2026-07-03")}',
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')

    # ─── 1. 投资要点摘要 ───
    pdf.add_page()
    pdf.section('1. 投资要点摘要')
    pdf.set_font(pdf.fn, '', 12)
    pdf.set_text_color(0x33, 0x33, 0x33)
    pdf.cell(0, 7, f'投资评级: {data.get("rating","持有 (HOLD)")}', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font(pdf.fn, '', 10)
    up_val = round(tp/cp*100-100,1) if cp else 0
    up_sign = '+' if up_val > 0 else ''
    upside_str = f"{up_sign}{up_val}%"
    pdf.cell(0, 6, f'目标价: ¥{tp:.2f} | 当前价: ¥{cp:.2f} | 隐含涨幅: {upside_str}',
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)
    pdf.set_font(pdf.fn, '', 12)
    pdf.set_text_color(0x1A, 0x23, 0x7E)
    pdf.cell(0, 7, '核心亮点:', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    name_display = data.get('company_name','')
    highlights = [
        f'估值方法: DCF绝对估值 + 可比公司相对估值双重验证',
        f'DCF目标价 ¥{tp:.0f}, 隐含空间 {upside_str}',
        f'{name_display}在行业内竞争优势突出',
    ]

    for h in highlights:
        pdf.bullet(h)

    # 关键指标一览
    pdf.ln(2)
    pdf.section('关键指标一览', 2)
    im = data.get('industry_median', {})
    comp_arr = data.get('comps_data', [])
    if comp_arr:
        c0 = comp_arr[0]
        c0_pe = c0.get('pe_ttm','--')
        c0_roe = str(c0.get('roe','--'))+'%'
        c0_mcap = c0.get('market_cap','--')
        c0_gm = str(c0.get('gross_margin','--'))+'%'
        c0_shares = f'{c0.get("shares_outstanding","--")}亿' if c0.get('shares_outstanding') else '--'
    else:
        c0_pe = c0_roe = c0_mcap = c0_gm = c0_shares = '--'
    mcap_str = f'{c0_mcap}亿' if str(c0_mcap) != '--' else '--'
    # Extract WACC and Beta from assumptions
    wacc_s, beta_s = '--', '--'
    for a in data.get('dcf_assumptions',[]):
        k = str(a[0])
        if k.strip().startswith('WACC'): wacc_s = a[1]
        if k.strip().startswith('Beta'): beta_s = a[1]
    metrics = [
        ('DCF目标价', f'¥{tp:.2f}', 'ROE', c0_roe),
        ('当前价', f'¥{cp:.2f}', '毛利率', c0_gm),
        ('上涨空间', upside_str, 'PE TTM', str(c0_pe)+'x'),
        ('市值', mcap_str, '2026E PE', f'{data.get("pe_2026e","--")}x'),
        ('WACC', wacc_s, 'PEG', f'{data.get("peg","--")}' if data.get('peg') else '--'),
        ('Beta', beta_s, '总股本', c0_shares),
        ('预测期', '5年', '评级', data.get('rating','--')),
    ]
    # Use proper table for metrics
    pdf.set_font(pdf.fn, '', 9)
    cw = 40
    for row in metrics:
        for ci in range(2):
            x = pdf.l_margin + ci * (cw*2 + 4)
            # label
            pdf.set_xy(x, pdf.get_y())
            pdf.set_fill_color(0xE8, 0xEA, 0xF6)
            pdf.set_text_color(0x66, 0x66, 0x66)
            pdf.cell(cw, 6, f' {row[ci*2]}', border=1, fill=True)
            # value
            pdf.set_fill_color(0xFF, 0xFF, 0xFF)
            pdf.set_text_color(0x1A, 0x23, 0x7E)
            pdf.set_font(pdf.fn, '', 11)
            pdf.cell(cw, 6, f'{row[ci*2+1]}', border=1, fill=True, align='C')
            pdf.set_font(pdf.fn, '', 9)
        pdf.ln(6)

    # ─── 2. 公司概况与财务概览 ───
    pdf.add_page()
    pdf.section('2. 公司概况与财务概览')
    pdf.section('2.1 公司概况', 2)
    tagline_desc = data.get('tagline', '') or ''
    pdf.body((f'{name} ({ticker}),是{tagline_desc}。所属行业为{data.get("industry","")}。'
              f'公司深耕主营业务，在行业内具有显著竞争优势。'))

    # 2.2 主营业务构成
    pdf.section('2.2 主营业务构成 (2025年报)', 2)
    segments = data.get('business_segments', [])
    if segments:
        pdf.data_table(
            ['业务板块', '营收(亿元)', '占比', '备注'],
            [[s[0], f'{s[1]:.0f}', s[2], s[3]] for s in segments],
            [38, 25, 20, 45]
        )
    pdf.ln(2)
    pdf.section('2.3 地区构成', 2)
    reg = data.get('regional_revenue', [])
    if reg:
        pdf.data_table(['地区', '营收(亿元)/占比'], [[r[0], f'{r[1]:.0f} / {r[2]}'] for r in reg], [40, 60])

    # 2.4 历史业绩趋势
    pdf.ln(2)
    pdf.section('2.4 历史业绩趋势 (2021-2025)', 2)
    hist = data.get('historical_financials', {})
    if hist.get('rows'):
        pdf.data_table(hist['columns'], [[str(c) for c in r] for r in hist['rows']],
                      [12, 18, 22, 16, 14, 18, 18])

    # ─── 3. DCF 估值分析 ───
    pdf.add_page()
    pdf.section('3. DCF 估值分析')
    pdf.body(('本节采用自由现金流折现法 (FCFF),对未来5年(2026-2030)的自由现金流进行预测,'
              '并结合永续增长模型计算终值,得出' + name + '的内在价值。'))
    pdf.section('3.1 核心假设与参数', 2)
    assumes = data.get('dcf_assumptions', [])
    if assumes:
        pdf.data_table(['参数', '取值 / 依据'],
                       [[a[0], f'{a[1]} {a[2]}'.strip()] for a in assumes],
                       [50, 80])
    pdf.ln(2)
    pdf.section('3.2 自由现金流预测', 2)
    pdf.body(('基于分析师一致预期的营收数据,结合公司盈利能力持续改善的趋势,'
              '我们假设FCF margin从2026年的17.5%逐步提升至2030年的18.7%。'))
    projs = data.get('projections', [])
    if projs:
        # Add FCF margin and discount factor
        wacc_val = float(assumes[8][1].replace('%',''))/100 if len(assumes)>8 else 0.0898
        fcf_rows = []
        cum_pv = 0
        for i, p in enumerate(projs):
            fcf = p.get('fcf', 0)
            margin = p.get('fcf_margin', 15)
            df = 1 / ((1 + wacc_val) ** (i + 1))
            pv = fcf * df
            cum_pv += pv
            year = p.get('year','')
            rev = p.get('revenue',0)
            fcf_rows.append([year, f'{rev:.0f}', f'{margin:.1f}%', f'{fcf:.0f}',
                            f'{df:.3f}', f'{pv:.0f}', f'{cum_pv:.0f}'])
        pdf.data_table(['年份', '营收(亿元)', 'FCF margin', 'FCF(亿元)', '折现因子', 'PV(亿元)', '累计PV'],
                      fcf_rows, [14, 22, 20, 20, 16, 16, 18])

    pdf.ln(2)
    pdf.section('3.3 估值汇总', 2)
    dcf_r = data.get('dcf_result', {})
    ev = dcf_r.get('enterprise_value', 0)
    total_pv = dcf_r.get('total_pv_fcfs', 0)
    pv_term = dcf_r.get('pv_terminal', 0)
    dcf_res_summ = data.get('dcf_result', {})
    nc_val_summ = data.get('net_cash', dcf_res_summ.get('net_cash', '--'))
    sh_val_summ = dcf_res_summ.get('shares_outstanding', data.get('shares_outstanding', '--'))
    rows = [
        ('5年 FCF 现值合计', f'¥{total_pv:.0f}亿'),
        ('终值现值 (TV)', f'¥{pv_term:.0f}亿'),
        ('企业价值 (EV)', f'¥{ev:.0f}亿'),
        ('加: 净现金', f'¥{nc_val_summ}亿'),
        ('股权价值', f'¥{ev + int(nc_val_summ) if str(nc_val_summ).isdigit() else ev:.0f}亿'),
        ('÷ 总股本', f'{sh_val_summ}亿股' if sh_val_summ else '--'),
        ('DCF隐含股价', f'¥{tp:.2f}'),
    ]
    pdf.kv_table(rows, [50, 60], hl=True)
    pdf.ln(1)
    tv_pct = dcf_r.get('terminal_value_pct', 78.7)
    pdf.set_font(pdf.fn, '', 9)
    pdf.set_text_color(0x1B, 0x5E, 0x20)
    pdf.multi_cell(0, 5, f'▶ 终值占总估值的{tv_pct:.1f}%——成长型龙头的典型特征')

    # 3.4 敏感性分析
    pdf.add_page()
    pdf.section('3.4 敏感性分析', 2)
    pdf.body(('下表展示了在不同WACC和永续增长率组合下的DCF隐含股价。红色标注为最可能区间。'))

    heatmap_path = os.path.join(chart_dir, f"{data.get('ticker_safe','')}_sensitivity_heatmap.png")
    if os.path.exists(heatmap_path):
        pdf.image(heatmap_path, x=pdf.l_margin+15, w=140)
    
    pdf.ln(2)
    pdf.section('敏感性分析核心结论:', 3)
    # 从敏感性矩阵动态生成结论
    sens_data = data.get('dcf_sensitivity', {})
    sens_prices = sens_data.get('prices', {}) if sens_data else {}
    wacc_r = sens_data.get('wacc_range', [])
    g_r = sens_data.get('g_range', [])
    
    sens_conc = []
    if sens_prices and wacc_r and g_r:
        # 寻找WACC 8.5%-9.5%、g 2.0%-2.5%区间的最值
        vals = []
        for i, w in enumerate(wacc_r):
            if 8.5 <= w <= 9.5:
                for j, g in enumerate(g_r):
                    if 2.0 <= g <= 2.5:
                        sw = str(w); sg = str(g)
                        if sw in sens_prices and sg in sens_prices[sw] and sens_prices[sw][sg]:
                            vals.append(sens_prices[sw][sg])
        if vals:
            sens_conc.append(f'WACC=8.5%-9.5%、永续增速2.0%-2.5%的最合理区间内,隐含股价为¥{min(vals):.0f}~¥{max(vals):.0f}')
        
        # 悲观假设 (WACC最高, g最低)
        w_high = str(wacc_r[-1]); g_low = str(g_r[0])
        if w_high in sens_prices and g_low in sens_prices[w_high] and sens_prices[w_high][g_low]:
            pess = sens_prices[w_high][g_low]
            sens_conc.append(f'即使悲观假设(WACC={wacc_r[-1]:.1f}%, g={g_r[0]:.1f}%),隐含价¥{pess:.0f}')
        
        # 乐观假设 (WACC最低, g最高)
        w_low = str(wacc_r[0]); g_high = str(g_r[-1])
        if w_low in sens_prices and g_high in sens_prices[w_low] and sens_prices[w_low][g_high]:
            opt = sens_prices[w_low][g_high]
            sens_conc.append(f'在乐观假设(WACC={wacc_r[0]:.1f}%, g={g_r[-1]:.1f}%)下,隐含价可达¥{opt:.0f}')
    
    sens_conc.append('结论对假设变化相对稳健,核心驱动力在于终值的合理性')
    sens_conc.append(f'终值占总估值的{tv_pct:.1f}%')
    for c in sens_conc:
        pdf.bullet(c)

    # ─── 4. 可比公司估值分析 ───
    pdf.add_page()
    pdf.section('4. 可比公司估值分析')
    pdf.section('4.1 可比公司选取', 2)
    comps = data.get('comps_data', [])
    if comps:
        pdf.data_table(['公司', '代码', '业务定位', '选取理由'],
                       [[f'{name}(标的)', data.get('ticker',''), tagline_desc, '—'] if i == 0 else
                        [c.get('name',''), c.get('market','') or '', '', '参考同业']
                        for i, c in enumerate(comps)],
                       [30, 24, 40, 30])

    pdf.ln(2)
    pdf.section('4.2 估值倍数对比', 2)
    if comps:
        pdf.data_table(
            ['公司', '市值(亿)', 'PE TTM', 'PB', 'ROE(%)', '毛利率(%)', '净利率(%)', '营收增速(%)'],
            [[c.get('name',''), f"{c.get('market_cap',0):.0f}", f"{c.get('pe_ttm',0):.1f}",
              f"{c.get('pb',0):.2f}", f"{c.get('roe',0):.2f}",
              f"{c.get('gross_margin',0):.2f}", f"{c.get('net_margin',0):.2f}",
              f"{c.get('revenue_growth',0):.2f}"]
             for c in comps],
            [22, 18, 17, 16, 18, 20, 18, 18],
            hl_col=0
        )

    # Radar chart
    radar_path = os.path.join(chart_dir, f"{data.get('ticker_safe','')}_comps_radar.png")
    if os.path.exists(radar_path):
        pdf.ln(2)
        pdf.image(radar_path, x=pdf.l_margin+25, w=120)

    # 4.3 行业排名
    pdf.add_page()
    pdf.section('4.3 行业排名与评价', 2)
    im = data.get('industry_median', {})
    ir = data.get('industry_ranking', {})
    ie = data.get('industry_eval', {})
    if im:
        rank_data = [
            ['PE TTM', f'{comps[0].get("pe_ttm",22.4):.1f}' if comps else '22.4',
             str(im.get('pe_ttm','')), ir.get('pe_ttm',''), ie.get('pe_ttm','')],
            ['PB', f'{comps[0].get("pb",5.41):.2f}' if comps else '5.41',
             str(im.get('pb','')), ir.get('pb',''), ie.get('pb','')],
            ['ROE', f'{comps[0].get("roe",24.91):.2f}%' if comps else '24.91%',
             f'{im.get("roe",0):.2f}%', ir.get('roe',''), ie.get('roe','')],
            ['毛利率', f'{comps[0].get("gross_margin",26.27):.2f}%' if comps else '26.27%',
             f'{im.get("gross_margin",0):.2f}%', ir.get('gross_margin',''), ie.get('gross_margin','')],
            ['净利率', f'{comps[0].get("net_margin",17.05):.2f}%' if comps else '17.05%',
             f'{im.get("net_margin",0):.2f}%', ir.get('net_margin',''), ie.get('net_margin','')],
            ['资产负债率', f'{comps[0].get("debt_ratio",61.94):.2f}%' if comps else '61.94%',
             f'{im.get("debt_ratio",0):.2f}%', ir.get('debt_ratio',''), ie.get('debt_ratio','')],
            ['营收增速', f'{comps[0].get("revenue_growth",17.04):.2f}%' if comps else '17.04%',
             f'{im.get("revenue_growth",0):.2f}%', ir.get('revenue_growth',''), ie.get('revenue_growth','')],
        ]
        pdf.data_table(['指标', name, '行业中位数', '排名', '评价'],
                      rank_data, [20, 28, 24, 18, 40])

    pdf.ln(2)
    pdf.section('4.4 估值解读', 2)
    tar_pe = comps[0].get('pe_ttm',22.4) if comps else 22.4
    tar_roe = comps[0].get('roe','--') if comps else '--'
    avg_pe_comp = sum(c.get('pe_ttm',0) for c in comps[1:] if c.get('pe_ttm',0)>0)/max(len([c for c in comps[1:] if c.get('pe_ttm',0)>0]),1) if len(comps)>1 else 25
    avg_roe_comp = sum(c.get('roe',0) for c in comps[1:] if c.get('roe',0)>0)/max(len([c for c in comps[1:] if c.get('roe',0)>0]),1) if len(comps)>1 else 12

    pdf.section('4.4.1 PE 视角 — 平价偏低', 3)
    for interp in data.get('interpretations', []):
        pdf.bullet(interp.replace('• ', ''))

    pdf.ln(1)
    pdf.section('4.4.2 PB-ROE 视角 — 合理溢价', 3)
    for pb_roe in data.get('pb_roe_analysis', []):
        pdf.bullet(pb_roe.replace('• ', ''))
    
    pe_2026e = data.get('pe_2026e', 18.5)
    peg = data.get('peg', 0.56)
    pdf.bullet(f'行业平均ROE {im.get("roe",10):.1f}%, {name}的ROE表现需持续关注')

    pdf.ln(1)
    pdf.section('4.4.3 综合判断', 3)
    pdf.kv_table(data.get('investment_conclusion', []), [50, 60])

    # ─── 5. 投资建议与评级 ───
    pdf.add_page()
    pdf.section('5. 投资建议与评级')
    pdf.set_draw_color(0x1A, 0x23, 0x7E)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w-pdf.r_margin, pdf.get_y())
    pdf.ln(3)
    pdf.set_font(pdf.fn, '', 14)
    pdf.set_text_color(0x1A, 0x23, 0x7E)
    rating_display = data.get('rating','持有 (HOLD)')
    pdf.cell(0, 8, f'投资建议: {rating_display}', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font(pdf.fn, '', 11)
    pdf.set_text_color(0x33, 0x33, 0x33)
    pdf.cell(0, 6, f'目标价: ¥{tp:.0f} | 评级有效期: 12个月',
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)
    pdf.section('核心逻辑:', 2)
    pdf.body((f'基于DCF绝对估值和可比公司相对估值的双重验证,我们给予{name}{data.get("rating","持有")}评级。'
              f'DCF模型隐含目标价¥{tp:.0f},对应上涨空间{upside_str}。'))
    ups = round(tp / cp * 100 - 100, 1) if cp > 0 else 0
    pdf.body((f'可比公司估值显示,{name}的PE={tar_pe:.1f}x。'
              f'综合DCF和可比分析,我们给予{data.get("rating","持有")}评级。'))

    pdf.ln(1)
    rating_label = data.get('rating', '持有 (HOLD)')
    if '买入' in rating_label:
        section_51 = '5.1 看多理由 (Bull Case)'
        section_52 = '5.2 催化剂 (Catalysts)'
    elif '减持' in rating_label:
        section_51 = '5.1 看空理由 (Bear Case)'
        section_52 = '5.2 风险关注 (Risk Watch)'
    else:
        section_51 = '5.1 中性理由 (Hold Case)'
        section_52 = '5.2 关注因素 (Key Factors)'
    pdf.section(section_51, 2)
    # 根据评级动态生成对应内容
    prof_net_cash = data.get('net_cash', 0)
    comp0_arr = data.get('comps_data', [])
    comp0_bull = comp0_arr[0] if comp0_arr else {}
    prof_roe = comp0_bull.get('roe', 0)
    prof_gm = comp0_bull.get('gross_margin', 0)
    
    if '买入' in rating_label:
        # 看多理由：强调优势
        bc_list = []
        if tp: bc_list.append((f'DCF估值支撑: DCF目标价¥{tp:.0f}', '高'))
        if prof_roe: bc_list.append((f'盈利能力强劲: ROE {prof_roe:.1f}%', '高'))
        if prof_gm: bc_list.append((f'高毛利率壁垒: 毛利率{prof_gm:.1f}%', '高'))
        if prof_net_cash and str(prof_net_cash).isdigit() and float(prof_net_cash) > 0:
            bc_list.append((f'资产负债健康: 净现金¥{float(prof_net_cash):.0f}亿', '高'))
        if not bc_list: bc_list = [('竞争优势突出', '中')]
    elif '减持' in rating_label:
        # 看空理由：强调风险
        upside_abs = abs(float(data.get('upside','0%').replace('%',''))) if data.get('upside') else 0
        bc_list = []
        bc_list.append((f'DCF估值偏高: 隐含价¥{tp:.0f}低于当前价¥{cp:.0f}', '高'))
        if prof_pe := comp0_bull.get('pe_ttm', 0):
            bc_list.append((f'估值水平: PE {prof_pe}x, 需关注业绩增长能否匹配', '中'))
        bc_list.append((f'市场定价: DCF隐含下跌空间约{upside_abs:.1f}%', '高'))
        if not bc_list: bc_list = [('估值压力: 当前价格偏高', '高')]
    else:
        # 持有理由：中性表述
        bc_list = []
        if tp: bc_list.append((f'DCF估值参考: 目标价¥{tp:.0f}', '中'))
        if prof_roe: bc_list.append((f'盈利能力: ROE {prof_roe:.1f}%, 行业领先', '中'))
        if prof_gm: bc_list.append((f'毛利率: {prof_gm:.1f}%, 竞争壁垒稳固', '中'))
        if not bc_list: bc_list = [('行业地位稳固', '中')]
    for detail, impact in bc_list:
        pdf.set_font(pdf.fn, '', 11)
        pdf.set_text_color(0x33, 0x33, 0x33)
        pdf.multi_cell(0, 5, f'  {impact} ★ {detail}')
        pdf.ln(0.5)

    pdf.ln(1)
    pdf.section(section_52, 2)
    # 根据评级动态生成催化剂/关注因素
    cat_industry = data.get('industry', '')
    cat_list = []
    if '买入' in rating_label:
        if cat_industry: cat_list.append(f'{cat_industry}行业景气度持续，{name}核心业务受益')
        cat_list.append(f'{name}盈利能力持续改善，估值有修复空间')
    elif '减持' in rating_label:
        if cat_industry: cat_list.append(f'{cat_industry}行业竞争加剧，需关注{name}市场份额变化')
        cat_list.append(f'{name}估值较高，需等待更合适的安全边际')
        cat_list.append(f'关注DCF隐含价¥{tp:.0f}附近的投资机会')
    else:
        if cat_industry: cat_list.append(f'{cat_industry}行业景气度平稳，{name}核心业务稳步发展')
        cat_list.append(f'当前估值处于合理区间，等待更好的配置时机')
    for cat in cat_list:
        pdf.bullet(cat)

    # ─── 6. 风险提示 ───
    pdf.ln(2)
    pdf.add_page()
    pdf.section('6. 风险提示')
    pdf.body('投资有风险,入市需谨慎。以下因素可能导致我们的估值模型和投资结论出现偏差:')
    pdf.section('6.1 核心风险因素', 2)
    risks = data.get('risk_factors', [])
    if risks:
        pdf.data_table(['序号', '风险因素', '影响程度'],
                       [[str(i+1), r[0], r[1]] for i, r in enumerate(risks)],
                       [10, 90, 20])

    pdf.ln(4)
    pdf.set_draw_color(0x99, 0x99, 0x99)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w-pdf.r_margin, pdf.get_y())
    pdf.ln(4)
    pdf.set_font(pdf.fn, '', 9)
    pdf.set_text_color(0x99, 0x99, 0x99)
    pdf.multi_cell(0, 4,
        '免责声明: 本报告仅供参考,不构成个人投资建议。报告中的所有分析、估值模型和数据均来自NeoData Financial Search API,'
        '模型假设可能存在偏差,估值结果仅供研究参考。投资有风险,入市需谨慎。',
        align='C')
    pdf.ln(1)
    pdf.cell(0, 4, '© 2026 研木 (研股股) | 课程案例 | 生成日期: ' + data.get('date',''),
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')

    pdf.output(output_path)
    print(f"✅ PDF报告已生成: {output_path}")


# ============ Word报告 (python-docx) ============
def generate_docx_report(data: Dict, output_path: str):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def add_body(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(text, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        return p

    def add_kv_table(rows):
        t = doc.add_table(rows=len(rows), cols=2)
        t.style = 'Table Grid'
        for i, (k, v) in enumerate(rows):
            t.cell(i, 0).text = k
            t.cell(i, 1).text = v
        return t

    def add_data_table(headers, data_rows):
        t = doc.add_table(rows=len(data_rows)+1, cols=len(headers))
        t.style = 'Table Grid'
        for i, h in enumerate(headers):
            t.cell(0, i).text = h
        for r, row in enumerate(data_rows, 1):
            for i, v in enumerate(row):
                t.cell(r, i).text = str(v)
        return t

    tp = data.get('target_price',0)
    cp = data.get('current_price',0)
    up_val = round(tp/cp*100-100,1) if cp else 0
    up_sign = '+' if up_val > 0 else ''
    upside = f"{up_sign}{up_val}%"
    name = data.get('company_name','')
    ticker = data.get('ticker','')  # v2.1 修复ticker未定义bug
    upside_str = upside  # v2.1 修复upside_str未定义bug

    # 封面
    for _ in range(5):
        doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('深度研究报告')
    r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x88,0x88,0x88)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{name} ({data.get("ticker","")})')
    r.font.size = Pt(26); r.bold = True; r.font.color.rgb = RGBColor(0x1A,0x23,0x7E)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline = data.get('tagline','') or ''
    r = p.add_run(tagline + '——DCF与可比公司估值分析')
    r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x55,0x55,0x55)
    doc.add_paragraph('')
    
    docx_comp2 = data.get('comps_data', [])
    docx_mcap2 = f"{docx_comp2[0].get('market_cap',0):.0f}亿" if docx_comp2 else '--'
    info_items = [
        f'{data.get("rating","持有")} | 目标价 ¥{tp:.0f} | 隐含涨幅: {upside}',
        f'当前价: ¥{cp:.2f} (2026-07-03) | 市值: {docx_mcap2}',
        f'数据来源: NeoData Financial Search API | 日期: 2026-07-03',
        '报告撰写: 研股股 (Equity Research Expert)',
    ]
    for item in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(item)
        r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x33,0x33,0x33)
    
    doc.add_page_break()
    
    # 1. 投资要点摘要
    doc.add_heading('1. 投资要点摘要', level=1)
    p = doc.add_paragraph(f'投资评级: {data.get("rating","持有 (HOLD)")}')
    p.runs[0].bold = True
    doc.add_paragraph(f'目标价: ¥{tp:.2f} | 当前价: ¥{cp:.2f} | 隐含涨幅: {upside}')
    
    doc.add_heading('核心亮点:', level=2)
    for h in [
        f'估值方法: DCF绝对估值 + 可比公司相对估值双重验证',
        f'DCF隐含目标价 ¥{tp:.0f}, 隐含涨幅: {upside}',
        f'{name}在行业内竞争优势突出',
        f'盈利能力保持行业领先水平',
    ]:
        add_bullet(h)
    
    doc.add_heading('关键指标一览', level=2)
    met_comp = data.get('comps_data', [])
    met_c0 = met_comp[0] if met_comp else {}
    met_pe = str(met_c0.get('pe_ttm','--'))
    met_roe = str(met_c0.get('roe','--'))+'%'
    met_mcap = f"{met_c0.get('market_cap',0):.0f}亿" if met_comp else '--'
    met_nc = data.get('net_cash','--')
    metrics = [
        ('DCF目标价', f'¥{tp:.2f}', 'ROE', met_roe),
        ('当前价', f'¥{cp:.2f}', 'PE TTM', met_pe+'x'),
        ('上涨空间', upside, '市值', met_mcap),
        ('净现金', str(met_nc)+'亿', '评级', data.get('rating','--')),
    ]
    t = doc.add_table(rows=len(metrics)+1, cols=4)
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(['指标','数值','指标','数值']):
        t.cell(0,i).text = h
    for r_i, (k1,v1,k2,v2) in enumerate(metrics, 1):
        t.cell(r_i,0).text = k1; t.cell(r_i,1).text = v1
        t.cell(r_i,2).text = k2; t.cell(r_i,3).text = v2

    # 2. 公司概况
    doc.add_page_break()
    doc.add_heading('2. 公司概况与财务概览', level=1)
    doc.add_heading('2.1 公司概况', level=2)
    tagline_desc = data.get('tagline', '') or ''
    add_body(f'{name} ({ticker}),是{tagline_desc}。')

    doc.add_heading('2.2 主营业务构成 (2025年报)', level=2)
    segs = data.get('business_segments', [])
    if segs:
        add_data_table(['业务板块', '营收(亿元)', '占比', '备注'],
                      [[s[0], f'{s[1]:.0f}', s[2], s[3]] for s in segs])

    doc.add_heading('2.3 历史业绩趋势 (2021-2025)', level=2)
    hist = data.get('historical_financials', {})
    if hist.get('rows'):
        add_data_table(hist['columns'], [[str(c) for c in r] for r in hist['rows']])

    # 3. DCF
    doc.add_page_break()
    doc.add_heading('3. DCF 估值分析', level=1)
    add_body('本节采用自由现金流折现法 (FCFF),对未来5年(2026-2030)的自由现金流进行预测。')

    doc.add_heading('3.1 核心假设与参数', level=2)
    assumes = data.get('dcf_assumptions', [])
    if assumes:
        add_data_table(['参数', '取值 / 依据'], [[a[0], f'{a[1]} {a[2]}'.strip()] for a in assumes])

    doc.add_heading('3.2 自由现金流预测', level=2)
    projs = data.get('projections', [])
    if projs:
        wacc_val = 0.0898
        fcf_rows = []; cum_pv = 0
        for i, p in enumerate(projs):
            fcf = p.get('fcf',0); margin = p.get('fcf_margin',15)
            df = 1/((1+wacc_val)**(i+1)); pv = fcf*df; cum_pv += pv
            fcf_rows.append([p.get('year',''), f'{p.get("revenue",0):.0f}', f'{margin:.1f}%',
                           f'{fcf:.0f}', f'{df:.3f}', f'{pv:.0f}', f'{cum_pv:.0f}'])
        add_data_table(['年份','营收(亿元)','FCF margin','FCF(亿元)','折现因子','PV(亿元)','累计PV'], fcf_rows)

    doc.add_heading('3.3 估值汇总', level=2)
    dcf_r = data.get('dcf_result', {})
    ev = dcf_r.get('enterprise_value', 0)
    old_ev = data.get('dcf_result', {}).get('total_pv_fcfs', 0)
    docx_nc_s = data.get('net_cash','--')
    docx_sh_s = data.get('shares_outstanding','--')
    docx_eq = ev + (int(docx_nc_s) if str(docx_nc_s).isdigit() else 0)
    add_kv_table([
        ('5年 FCF 现值合计', f'¥{dcf_r.get("total_pv_fcfs",0):.0f}亿'),
        ('终值现值 (TV)', f'¥{dcf_r.get("pv_terminal",0):.0f}亿'),
        ('企业价值 (EV)', f'¥{ev:.0f}亿'),
        ('加: 净现金', f'¥{docx_nc_s}亿'),
        ('股权价值', f'¥{docx_eq:.0f}亿'),
        ('÷ 总股本', f'{docx_sh_s}亿股' if docx_sh_s else '--'),
        ('DCF隐含股价', f'¥{tp:.2f}'),
    ])

    # 4. 可比公司
    doc.add_page_break()
    doc.add_heading('4. 可比公司估值分析', level=1)
    doc.add_heading('4.1 可比公司选取', level=2)
    comps = data.get('comps_data', [])
    if comps:
        add_data_table(['公司','代码','业务定位','选取理由'],
                      [[f'{name}(标的)', data.get('ticker',''), tagline_desc,'—'] if i==0 else
                       [c.get('name',''),'',c.get('market','') or '参考同业']
                       for i,c in enumerate(comps)])

    doc.add_heading('4.2 估值倍数对比', level=2)
    if comps:
        add_data_table(['公司','市值(亿)','PE TTM','PB','ROE(%)','毛利率(%)','净利率(%)','营收增速(%)'],
                      [[c.get('name',''), f"{c.get('market_cap',0):.0f}", f"{c.get('pe_ttm',0):.1f}",
                        f"{c.get('pb',0):.2f}", f"{c.get('roe',0):.2f}", f"{c.get('gross_margin',0):.2f}",
                        f"{c.get('net_margin',0):.2f}", f"{c.get('revenue_growth',0):.2f}"]
                       for c in comps])

    doc.add_heading('4.3 行业排名与评价', level=2)
    im = data.get('industry_median', {})
    ir = data.get('industry_ranking', {})
    ie = data.get('industry_eval', {})
    if im:
        rank_data = [
            ['PE TTM', f'{comps[0].get("pe_ttm",22.4):.1f}', str(im.get('pe_ttm','')), ir.get('pe_ttm',''), ie.get('pe_ttm','')],
            ['PB', f'{comps[0].get("pb",5.41):.2f}', str(im.get('pb','')), ir.get('pb',''), ie.get('pb','')],
            ['ROE', f'{comps[0].get("roe",24.91):.2f}%', f'{im.get("roe",0):.2f}%', ir.get('roe',''), ie.get('roe','')],
            ['毛利率', f'{comps[0].get("gross_margin",26.27):.2f}%', f'{im.get("gross_margin",0):.2f}%', ir.get('gross_margin',''), ie.get('gross_margin','')],
            ['净利率', f'{comps[0].get("net_margin",17.05):.2f}%', f'{im.get("net_margin",0):.2f}%', ir.get('net_margin',''), ie.get('net_margin','')],
        ]
        add_data_table(['指标',name,'行业中位数','排名','评价'], rank_data)

    doc.add_heading('4.4 估值解读', level=2)
    for interp in data.get('interpretations', []):
        add_bullet(interp)
    doc.add_heading('PB-ROE视角', level=3)
    for pb_roe in data.get('pb_roe_analysis', []):
        add_bullet(pb_roe)
    doc.add_heading('综合判断', level=3)
    add_kv_table(data.get('investment_conclusion', []))

    # 5. 投资建议
    doc.add_page_break()
    doc.add_heading('5. 投资建议与评级', level=1)
    p = doc.add_paragraph()
    r = p.add_run(f'投资建议: {data.get("rating","持有 (HOLD)")}')
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x1A,0x23,0x7E)
    doc.add_paragraph(f'目标价: ¥{tp:.0f} | 评级有效期: 12个月')
    
    doc.add_heading('核心逻辑:', level=2)
    add_body(f'基于DCF绝对估值和可比公司相对估值的双重验证,我们给予{data.get("rating","持有")}评级。DCF模型隐含目标价¥{tp:.0f},对应上涨空间{upside_str}。')
    add_body(f'可比公司估值显示,{name}的PE处于合理水平,综合DCF和可比分析,我们给予{data.get("rating","持有")}评级。')
    
    docx_rating = data.get('rating','持有 (HOLD)')
    if '买入' in docx_rating:
        docx_51 = '5.1 看多理由 (Bull Case)'
        docx_52 = '5.2 催化剂 (Catalysts)'
    elif '减持' in docx_rating:
        docx_51 = '5.1 看空理由 (Bear Case)'
        docx_52 = '5.2 风险关注 (Risk Watch)'
    else:
        docx_51 = '5.1 中性理由 (Hold Case)'
        docx_52 = '5.2 关注因素 (Key Factors)'
    doc.add_heading(docx_51, level=2)
    docx_bc_nc = data.get('net_cash', 0)
    docx_bc_c0_arr = data.get('comps_data', [])
    docx_bc_c0 = docx_bc_c0_arr[0] if docx_bc_c0_arr else {}
    docx_bc_roe = docx_bc_c0.get('roe', 0)
    docx_bc_gm = docx_bc_c0.get('gross_margin', 0)
    docx_upside_abs = abs(float(data.get('upside','0%').replace('%',''))) if data.get('upside') else 0
    docx_pe_val = docx_bc_c0.get('pe_ttm', 0)
    
    if '买入' in docx_rating:
        if tp: add_bullet(f'高 ★ DCF估值支撑: 目标价¥{tp:.0f}')
        if docx_bc_roe: add_bullet(f'高 ★ 盈利能力强劲: ROE {docx_bc_roe:.1f}%')
        if docx_bc_gm: add_bullet(f'高 ★ 高毛利率壁垒: 毛利率{docx_bc_gm:.1f}%')
        if docx_bc_nc and float(docx_bc_nc) > 0: add_bullet(f'高 ★ 净现金¥{float(docx_bc_nc):.0f}亿')
    elif '减持' in docx_rating:
        if tp: add_bullet(f'高 ★ DCF估值偏高: 隐含价¥{tp:.0f}低于当前价¥{cp:.0f}')
        add_bullet(f'高 ★ DCF隐含下行空间约{docx_upside_abs:.1f}%')
        add_bullet(f'中 ★ 需等待估值回归合理区间')
    else:
        if tp: add_bullet(f'中 ★ DCF参考目标价¥{tp:.0f}')
        if docx_bc_roe: add_bullet(f'中 ★ ROE {docx_bc_roe:.1f}%, 行业领先')
        if docx_bc_gm: add_bullet(f'中 ★ 毛利率{docx_bc_gm:.1f}%, 竞争壁垒稳固')
    
    doc.add_heading(docx_52, level=2)
    docx_cat_ind = data.get('industry', '')
    if '买入' in docx_rating:
        if docx_cat_ind: add_bullet(f'{docx_cat_ind}行业景气度持续，{name}核心业务受益')
        add_bullet(f'{name}盈利能力持续改善，估值有修复空间')
    elif '减持' in docx_rating:
        if docx_cat_ind: add_bullet(f'{docx_cat_ind}行业需关注，{name}估值较高')
        add_bullet(f'关注DCF隐含价¥{tp:.0f}附近的投资机会')
    else:
        if docx_cat_ind: add_bullet(f'{docx_cat_ind}行业景气度平稳')
        add_bullet(f'{name}估值处于合理区间，等待更好配置时机')

    doc.add_heading('6. 风险提示', level=1)
    add_body('投资有风险,入市需谨慎。以下因素可能导致我们的估值模型和投资结论出现偏差:')
    risks = data.get('risk_factors', [])
    if risks:
        add_data_table(['序号','风险因素','影响程度'],
                      [[str(i+1), r[0], r[1]] for i, r in enumerate(risks)])

    p = doc.add_paragraph('')
    r = p.add_run('免责声明: 本报告仅供参考,不构成个人投资建议。所有数据来源于公开信息,模型假设可能存在偏差。')
    r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x99,0x99,0x99)

    doc.save(output_path)
    print(f"✅ Word报告已生成: {output_path}")


# ============ 主入口 ============
def main():
    parser = argparse.ArgumentParser(description='研木报告生成器 v2')
    parser.add_argument('--ticker', '-t', required=True)
    parser.add_argument('--company')
    parser.add_argument('--format', '-f', choices=['pdf','docx','both'], default='both')
    parser.add_argument('--dcf-data')
    parser.add_argument('--comps-data')
    parser.add_argument('--financial-data')
    parser.add_argument('--chart-dir', '-c', default='.')
    parser.add_argument('--output', '-o', default='.')
    parser.add_argument('--target-price', type=float)
    parser.add_argument('--current-price', type=float)
    parser.add_argument('--rating', default='买入 (BUY)')
    args = parser.parse_args()
    
    data = prepare_data(args)
    os.makedirs(args.output, exist_ok=True)
    date_str = datetime.now().strftime('%Y%m%d')
    ticker_s = args.ticker.replace('.', '_')
    base = f"{ticker_s}_研究报告_{date_str}"
    
    if args.format in ('docx', 'both'):
        generate_docx_report(data, os.path.join(args.output, f'{base}.docx'))
    if args.format in ('pdf', 'both'):
        generate_pdf_report(data, os.path.join(args.output, f'{base}.pdf'), args.chart_dir)

if __name__ == '__main__':
    main()
