#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
国家电网规范格式 Word 文档格式化脚本
按辉总确认的 12 条规范调整 .docx 格式。

用法:
    python3 format_sgcc.py 输入.docx [输出.docx]
    --no-auto-toc  不自动插入目录域（默认自动）
"""

import re
import sys
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

BLACK = RGBColor(0, 0, 0)

# 字体常量
F_TITLE = "方正小标宋_GBK"   # 主标题
F_SUB = "方正楷体_GBK"       # 副标题
F_TOC_TITLE = "方正黑体_GBK" # 目录标题
F_H1 = "方正黑体_GBK"
F_H2 = "方正楷体_GBK"
F_H3 = "方正仿宋_GBK"
F_H4 = "方正仿宋_GBK"
F_BODY = "方正仿宋_GBK"
F_PAGE = "方正仿宋_GBK"
F_TBL_HEAD = "方正黑体_GBK"
F_TBL_BODY = "方正仿宋_GBK"

SZ_2 = 22   # 二号
SZ_3 = 16   # 三号
SZ_4 = 14   # 四号

RE_H1 = re.compile(r"^[一二三四五六七八九十百]+、")
RE_H2 = re.compile(r"^（[一二三四五六七八九十百]+）")
RE_H3 = re.compile(r"^\d+[.．]")
RE_H4 = re.compile(r"^（\d+）")

HEADING_STYLE_IDS = {"Heading1": 1, "Heading2": 2, "Heading3": 3, "Heading4": 4}


# ---------------- 基础工具 ----------------

def set_run_font(run, font_cn, size_pt, bold=None, color=None):
    """设置 run 中英文字体 + 中文字体 + 字号 + 加粗 + 颜色"""
    run.font.name = font_cn
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_cn)
    run.font.size = Pt(size_pt)
    # 同步 szCs（复杂文种字号）
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), str(size_pt * 2))
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def fmt_para(p, align=None, first_line_chars=0, line_pt=28, before=0, after=0):
    """段落格式：对齐 / 首行缩进(字符) / 固定行距 / 段前段后"""
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(line_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if first_line_chars:
        ind.set(qn("w:firstLineChars"), str(first_line_chars * 100))
        ind.set(qn("w:firstLine"), str(first_line_chars * 320))  # 兜底磅值(twips)
    else:
        for attr in ("w:firstLineChars", "w:firstLine"):
            if ind.get(qn(attr)) is not None:
                del ind.attrib[qn(attr)]


def format_paragraph_runs(p, font_cn, size_pt, bold=None):
    for run in p.runs:
        set_run_font(run, font_cn, size_pt, bold=bold, color=BLACK)


def new_paragraph_element():
    return OxmlElement("w:p")


def add_field_to_element(p_el, instr, placeholder_text):
    """向段落元素追加一个 Word 域：begin / instrText / separate / 占位 / end"""
    def _run_with(child):
        r = OxmlElement("w:r")
        r.append(child)
        return r

    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin")
    p_el.append(_run_with(fld))
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve")
    instrText.text = instr
    p_el.append(_run_with(instrText))
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    p_el.append(_run_with(sep))
    t = OxmlElement("w:t"); t.text = placeholder_text
    r = OxmlElement("w:r"); r.append(t)
    p_el.append(r)
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    p_el.append(_run_with(end))


def ensure_style(doc, style_id, name, kind=WD_STYLE_TYPE.PARAGRAPH, base=None):
    """按 styleId 查找样式，不存在则创建并设置 styleId"""
    for s in doc.styles:
        if s.style_id == style_id:
            return s
    try:
        s = doc.styles[name]
        return s
    except KeyError:
        pass
    s = doc.styles.add_style(name, kind)
    s.element.set(qn("w:styleId"), style_id)
    if base:
        try:
            s.base_style = doc.styles[base]
        except KeyError:
            pass
    return s


def set_style_font(style, font_cn, size_pt, bold=None):
    style.font.name = font_cn
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_cn)
    style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    style.font.color.rgb = BLACK


def set_style_outline(style, level):
    """设置样式大纲级别（0 起），供目录识别"""
    pPr = style.element.get_or_add_pPr()
    ol = pPr.find(qn("w:outlineLvl"))
    if ol is None:
        ol = OxmlElement("w:outlineLvl")
        pPr.append(ol)
    ol.set(qn("w:val"), str(level - 1))


def set_style_para(style, line_pt=28, before=0, after=0):
    pPr = style.element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(int(before * 20)))
    spacing.set(qn("w:after"), str(int(after * 20)))
    spacing.set(qn("w:line"), str(int(line_pt * 20)))
    spacing.set(qn("w:lineRule"), "exact")


def heading_level_of(p):
    """识别段落标题级别：按样式优先，其次按编号规则"""
    try:
        sid = p.style.style_id if p.style is not None else None
    except Exception:
        sid = None
    if sid in HEADING_STYLE_IDS:
        return HEADING_STYLE_IDS[sid]
    text = p.text.strip()
    if not text:
        return None
    if RE_H1.match(text):
        return 1
    if RE_H2.match(text):
        return 2
    if RE_H3.match(text):
        return 3
    if RE_H4.match(text):
        return 4
    return None


def replace_quotes_in_paragraph(p):
    """英文双引号成对替换为中文引号"""
    state = {"open": True}

    def _repl(m):
        if state["open"]:
            state["open"] = False
            return "\u201c"
        state["open"] = True
        return "\u201d"

    for run in p.runs:
        if run.text and '"' in run.text:
            run.text = re.sub(r'"', _repl, run.text)


# ---------------- 主流程 ----------------

def main():
    if len(sys.argv) < 2:
        print("用法: python3 format_sgcc.py 输入.docx [输出.docx] [--no-auto-toc]")
        sys.exit(1)
    in_path = sys.argv[1]
    auto_toc = "--no-auto-toc" not in sys.argv
    out_path = None
    for a in sys.argv[2:]:
        if a.startswith("--"):
            continue
        out_path = a
    if out_path is None:
        out_path = re.sub(r"\.docx$", "", in_path, flags=re.I) + "_规范格式.docx"

    doc = Document(in_path)
    paras = list(doc.paragraphs)

    # 1) 引号替换（正文 + 表格）
    for p in paras:
        replace_quotes_in_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_quotes_in_paragraph(p)

    # 2) 结构识别
    toc_para = None
    for p in paras:
        if p.text.strip() == "目录":
            toc_para = p
            break

    # 正文起始 = 目录标题之后第一个标题段（无目录则从文档开头找）
    start_idx = 0
    if toc_para is not None:
        start_idx = paras.index(toc_para) + 1
    body_start_para = None
    for p in paras[start_idx:]:
        if heading_level_of(p) is not None:
            body_start_para = p
            break

    # 3) 目录处理：删除旧目录条目，插入目录域
    if toc_para is not None:
        # 删除目录标题与正文之间所有旧目录条目段落
        i = paras.index(toc_para)
        j = paras.index(body_start_para) if body_start_para is not None else len(paras)
        removed = 0
        for p in paras[i + 1:j]:
            p._p.getparent().remove(p._p)
            removed += 1
        if auto_toc and not _has_toc_field(toc_para):
            _insert_toc_field_after(doc, toc_para)
        print(f"[目录] 已有目录标题，清理旧条目 {removed} 段，插入可点击目录域")
    elif auto_toc and body_start_para is not None:
        # 自动创建：目录标题 + 目录域，插在正文起始段之前
        new_toc_p_el = new_paragraph_element()
        body_start_para._p.addprevious(new_toc_p_el)
        toc_para = Paragraph(new_toc_p_el, doc._body)
        toc_para.add_run("目录")
        _insert_toc_field_after(doc, toc_para)
        print("[目录] 自动插入：目录标题 + 可点击目录域")

    # 4) 主标题 / 副标题识别（封面区域，正文之前）
    main_title_para = None
    subtitle_para = None
    cover_paras = []
    for p in paras:
        if p is toc_para or p is body_start_para:
            break
        cover_paras.append(p)
    # 主标题：第一个居中段落（正文前）
    for p in cover_paras:
        if p.text.strip() and p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            main_title_para = p
            break
    if main_title_para is None and cover_paras:
        # 兜底：正文前第一个非空段
        for p in cover_paras:
            if p.text.strip() and p.text.strip() != "目录":
                main_title_para = p
                break
    if main_title_para is not None:
        idx = cover_paras.index(main_title_para)
        for p in cover_paras[idx + 1:]:
            if p.text.strip() and p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                subtitle_para = p
                break

    # 5) 正文标题统计
    h_counts = {1: 0, 2: 0, 3: 0, 4: 0}
    for p in paras:
        lv = heading_level_of(p)
        if lv:
            h_counts[lv] += 1

    # 6) 格式化段落
    body_started = body_start_para is None
    for p in paras:
        if p is main_title_para:
            _fmt_main_title(p)
        elif p is subtitle_para:
            _fmt_subtitle(p)
        elif p is toc_para:
            _fmt_toc_title(p)
        elif _is_toc_field_para(p):
            _fmt_toc_field(p)
        elif p is body_start_para:
            body_started = True
            lv = heading_level_of(p)
            _fmt_heading(p, lv)
        elif body_started:
            lv = heading_level_of(p)
            if lv:
                _fmt_heading(p, lv)
            elif p.text.strip():
                _fmt_body(p)
        # 封面其他段落保持原样

    # 7) 表格
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(table.rows):
            font_cn = F_TBL_HEAD if ri == 0 else F_TBL_BODY
            for cell in row.cells:
                for p in cell.paragraphs:
                    fmt_para(p, line_pt=20, before=0, after=0)
                    format_paragraph_runs(p, font_cn, SZ_4, bold=None)

    # 8) 标题样式 & 目录样式
    _setup_styles(doc)

    # 9) 分节 + 页码（正文从 1 开始，-1- 样式）
    if body_start_para is not None and paras.index(body_start_para) > 0:
        _split_section_before(doc, body_start_para)
        _setup_page_number(doc, body_section_index=1)
    else:
        _setup_page_number(doc, body_section_index=None)

    # 10) 打开时自动更新域（目录刷新）
    settings = doc.settings.element
    uf = settings.find(qn("w:updateFields"))
    if uf is None:
        uf = OxmlElement("w:updateFields")
        settings.append(uf)
    uf.set(qn("w:val"), "true")

    doc.save(out_path)

    print(f"\n✅ 已生成: {out_path}")
    print(f"[主标题] {main_title_para.text.strip() if main_title_para else '未识别'}")
    print(f"[副标题] {subtitle_para.text.strip() if subtitle_para else '未识别'}")
    print(f"[标题统计] 一级:{h_counts[1]} 二级:{h_counts[2]} 三级:{h_counts[3]} 四级:{h_counts[4]}")
    print(f"[表格] {len(doc.tables)} 个")
    print("提示: 用 Word 打开后目录会自动更新（或全选按 F9），即可点击跳转")


# ---------------- 格式化函数 ----------------

def _set_normal_style(p):
    try:
        p.style = p.part.document.styles["Normal"]
    except Exception:
        pass


def _fmt_main_title(p):
    _set_normal_style(p)
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_chars=0, line_pt=28, before=0, after=0)
    format_paragraph_runs(p, F_TITLE, SZ_2, bold=False)


def _fmt_subtitle(p):
    _set_normal_style(p)
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_chars=0, line_pt=28, before=0, after=0)
    format_paragraph_runs(p, F_SUB, SZ_3, bold=True)


def _fmt_toc_title(p):
    _set_normal_style(p)
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_chars=0, line_pt=28, before=0, after=0)
    format_paragraph_runs(p, F_TOC_TITLE, SZ_2, bold=False)


def _is_toc_field_para(p):
    try:
        pPr = p._p.pPr
        if pPr is None:
            return False
        return pPr.find(qn("w:rPr")) is not None and False
    except Exception:
        pass
    # 检查段内是否有 TOC 域指令
    return any(
        t.text and "TOC" in t.text
        for t in p._p.iter(qn("w:instrText"))
    )


def _fmt_toc_field(p):
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_chars=0, line_pt=28, before=0, after=0)
    for run in p.runs:
        set_run_font(run, F_BODY, SZ_4, bold=False, color=BLACK)


def _fmt_heading(p, lv):
    style = _heading_style_of(p, lv)
    try:
        p.style = style
    except Exception:
        pass
    spec = {
        1: (F_H1, SZ_3, False),
        2: (F_H2, SZ_3, True),
        3: (F_H3, SZ_3, True),
        4: (F_H4, SZ_3, False),
    }
    font_cn, size, bold = spec[lv]
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_chars=2, line_pt=28, before=0, after=0)
    format_paragraph_runs(p, font_cn, size, bold=bold)


def _fmt_body(p):
    _set_normal_style(p)
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_chars=2, line_pt=28)
    format_paragraph_runs(p, F_BODY, SZ_3, bold=False)


def _heading_style_of(p, lv):
    doc = p.part.document
    sid = f"Heading{lv}"
    return ensure_style(doc, sid, f"Heading {lv}", base="Normal")


# ---------------- 目录域 ----------------

def _has_toc_field(p):
    return any(
        t.text and "TOC" in t.text
        for t in p._p.iter(qn("w:instrText"))
    )


def _insert_toc_field_after(doc, ref_para):
    p_el = new_paragraph_element()
    ref_para._p.addnext(p_el)
    add_field_to_element(p_el, ' TOC \\o "1-4" \\h \\z \\u ', "（目录：打开文档后自动更新，或全选按 F9）")
    new_p = Paragraph(p_el, doc._body)
    fmt_para(new_p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_chars=0, line_pt=28, before=0, after=0)
    for run in new_p.runs:
        set_run_font(run, F_BODY, SZ_4, bold=False, color=BLACK)


# ---------------- 样式准备 ----------------

def _setup_styles(doc):
    # 标题 1-4 样式
    for lv, sid in enumerate(["Heading1", "Heading2", "Heading3", "Heading4"], start=1):
        style = ensure_style(doc, sid, f"Heading {lv}", base="Normal")
        spec = {
            1: (F_H1, SZ_3, False),
            2: (F_H2, SZ_3, True),
            3: (F_H3, SZ_3, True),
            4: (F_H4, SZ_3, False),
        }[lv]
        set_style_font(style, spec[0], spec[1], bold=spec[2])
        set_style_outline(style, lv)
        set_style_para(style, line_pt=28, before=0, after=0)
    # 目录样式 TOC1-4：字体与各级标题一致，四号，行距 28
    toc_spec = {
        "TOC1": (F_H1, False),
        "TOC2": (F_H2, True),
        "TOC3": (F_H3, True),
        "TOC4": (F_H4, False),
    }
    for sid, (font_cn, bold) in toc_spec.items():
        style = ensure_style(doc, sid, sid.replace("TOC", "TOC "), base="Normal")
        set_style_font(style, font_cn, SZ_4, bold=bold)
        set_style_para(style, line_pt=28, before=0, after=0)


# ---------------- 分节与页码 ----------------

def _split_section_before(doc, body_start_para):
    """在正文起始段前插入分节符（封面/目录为一节，正文为另一节）"""
    body = doc.element.body
    body_sectPr = body.find(qn("w:sectPr"))
    if body_sectPr is None:
        return
    sect = copy.deepcopy(body_sectPr)
    for tag in ("w:footerReference", "w:headerReference", "w:pgNumType"):
        el = sect.find(qn(tag))
        if el is not None:
            sect.remove(el)
    new_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pPr.append(sect)
    new_p.append(pPr)
    body_start_para._p.addprevious(new_p)


def _setup_page_number(doc, body_section_index):
    """正文节页脚：-1- 样式，从 1 开始编号；前置节无页码"""
    sections = doc.sections
    # 正文节 = 最后一节（body 级 sectPr）；前置节 = 之前各节
    body_sec = sections[-1]
    # 前置节页脚置空（删除 footerReference）
    for sec in sections[:-1]:
        sectPr = sec._sectPr
        for tag in ("w:footerReference", "w:pgNumType"):
            el = sectPr.find(qn(tag))
            if el is not None:
                sectPr.remove(el)
    # 正文节：footer 独立，写入 -1- 页码域
    body_sec.footer.is_linked_to_previous = False
    footer = body_sec.footer
    fp = footer.paragraphs[0]
    for r in list(fp.runs):
        r._element.getparent().remove(r._element)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = fp.add_run("-")
    set_run_font(run1, F_PAGE, SZ_4, bold=False, color=BLACK)
    # PAGE 域
    r = OxmlElement("w:r")
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin")
    r.append(fld); fp._p.append(r)
    r = OxmlElement("w:r")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " PAGE "
    r.append(it); fp._p.append(r)
    r = OxmlElement("w:r")
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "separate")
    r.append(fld); fp._p.append(r)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t"); t.text = "1"
    r.append(t); fp._p.append(r)
    r = OxmlElement("w:r")
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "end")
    r.append(fld); fp._p.append(r)
    run2 = fp.add_run("-")
    set_run_font(run2, F_PAGE, SZ_4, bold=False, color=BLACK)
    # 域 run 字体统一
    for r in fp._p.findall(qn("w:r")):
        _ensure_run_font(r, F_PAGE, SZ_4)
    # 正文节页码从 1 开始
    sectPr = body_sec._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        cols = sectPr.find(qn("w:cols"))
        if cols is not None:
            cols.addprevious(pg)
        else:
            sectPr.append(pg)
    pg.set(qn("w:start"), "1")


def _ensure_run_font(r_el, font_cn, size_pt):
    rPr = r_el.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r_el.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_cn)
    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rPr.append(sz)
    sz.set(qn("w:val"), str(size_pt * 2))


if __name__ == "__main__":
    main()
