#!/usr/bin/env python3
"""
Auto-generate a financial analysis Word report from financial data.
No Word template needed - generates a complete report from scratch.

Usage (CLI):
  python generate_report.py \
    --company "XX公司" \
    --period "2025年12月" \
    --bs-data bs.json \
    --is-data is.json \
    [--cf-data cf.json] \
    --output 贷后分析报告.docx \
    [--prev-bs-data prev_bs.json] \
    [--prev-is-data prev_is.json] \
    [--prev-cf-data prev_cf.json] \
    [--loan-amount 50000000] \
    [--loan-type "流动资金贷款"] \
    [--report-type "贷后分析报告"]

Usage (module):
  from generate_report import generate_report
  generate_report(
      company="XX公司",
      period="2025年12月",
      bs_data={"货币资金": "35677380.72", ...},
      is_data={"营业收入": "645841626.73", ...},
      cf_data={"经营活动产生的现金流量净额": "...", ...},
      output_path="报告.docx",
  )

JSON data format (key = account name in Chinese, value = amount string in yuan):
  {"货币资金": "35677380.72", "应收账款": "219143915.18", ...}
"""

import argparse
import json
import os
import sys
from datetime import datetime
from pathlib import Path

# Import utility functions from merge_reports
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from merge_reports import (
    to_wan_rounded, parse_numeric, find_in_data,
    calculate_financial_ratios, normalize_quotes,
)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

BLUE_RGB = RGBColor(0, 0, 255)

# Key items to include in each statement table
BS_ASSET_ITEMS = [
    '货币资金', '应收票据', '应收账款', '预付款项', '其他应收款', '存货',
    '合同资产', '其他流动资产',
    '流动资产合计',
    '固定资产', '在建工程', '使用权资产', '无形资产', '商誉',
    '长期待摊费用', '递延所得税资产', '其他非流动资产',
    '非流动资产合计',
    '资产总计',
]

BS_LIAB_ITEMS = [
    '短期借款', '应付票据', '应付账款', '预收款项', '应付职工薪酬',
    '应交税费', '其他应付款', '一年内到期的非流动负债',
    '流动负债合计',
    '长期借款', '长期应付款', '递延收益', '其他非流动负债',
    '非流动负债合计',
    '负债合计',
]

BS_EQUITY_ITEMS = [
    '实收资本', '资本公积', '其他综合收益', '盈余公积', '未分配利润',
    '少数股东权益',
    '所有者权益合计',
]

IS_KEY_ITEMS = [
    '营业收入', '营业成本', '税金及附加', '销售费用', '管理费用',
    '研发费用', '财务费用', '其他收益', '投资收益',
    '营业利润', '营业外收入', '营业外支出',
    '利润总额', '所得税费用', '净利润',
]

CF_KEY_ITEMS = [
    '销售商品、提供劳务收到的现金',
    '经营活动现金流入小计',
    '购买商品、接受劳务支付的现金',
    '经营活动现金流出小计',
    '经营活动产生的现金流量净额',
    '投资活动现金流入小计',
    '投资活动现金流出小计',
    '投资活动产生的现金流量净额',
    '筹资活动现金流入小计',
    '筹资活动现金流出小计',
    '筹资活动产生的现金流量净额',
    '现金及现金等价物净增加额',
    '期末现金及现金等价物余额',
]

# Financial ratio reference values for risk assessment
# unit: 'percent' -> values shown with '%' ; 'ratio' -> plain number
# 周转率/增长率行业差异大，不给参考值（参考值与评估列显示 '-'）
RATIO_REFERENCES = {
    '流动比率': {'good': 2.0, 'warning': 1.0, 'direction': 'higher_better', 'unit': 'ratio'},
    '速动比率': {'good': 1.0, 'warning': 0.5, 'direction': 'higher_better', 'unit': 'ratio'},
    '资产负债率': {'good': 60, 'warning': 70, 'direction': 'lower_better', 'unit': 'percent'},
    '权益乘数': {'good': 2.0, 'warning': 3.0, 'direction': 'lower_better', 'unit': 'ratio'},
    '利息保障倍数': {'good': 3.0, 'warning': 1.0, 'direction': 'higher_better', 'unit': 'ratio'},
    '现金流量负债率': {'good': 20, 'warning': 10, 'direction': 'higher_better', 'unit': 'percent'},
    '负债权益比': {'good': 100, 'warning': 150, 'direction': 'lower_better', 'unit': 'percent'},
    '净资产收益率': {'good': 10, 'warning': 3, 'direction': 'higher_better', 'unit': 'percent'},
    '毛利率': {'good': 20, 'warning': 10, 'direction': 'higher_better', 'unit': 'percent'},
    '营业净利率': {'good': 10, 'warning': 3, 'direction': 'higher_better', 'unit': 'percent'},
    '总资产收益率': {'good': 8, 'warning': 3, 'direction': 'higher_better', 'unit': 'percent'},
    '总资产报酬率': {'good': 8, 'warning': 3, 'direction': 'higher_better', 'unit': 'percent'},
}

# 4-category layout of the 18 financial ratios (v1.1)
RATIO_CATEGORIES = [
    {
        'name': '偿债能力',
        'subgroups': [
            ('短期', ['流动比率', '速动比率']),
            ('长期', ['资产负债率', '权益乘数', '利息保障倍数', '现金流量负债率', '负债权益比']),
        ],
    },
    {'name': '营运效率', 'items': ['存货周转率', '总资产周转率', '应收账款周转率']},
    {'name': '盈利能力', 'items': ['净资产收益率', '毛利率', '营业净利率', '总资产收益率', '总资产报酬率']},
    {'name': '发展能力', 'items': ['销售增长率', '净利润增长率', '每股收益增长率']},
]


# ---------------------------------------------------------------------------
# Document Formatting Helpers
# ---------------------------------------------------------------------------

def _set_cell_borders(cell):
    """Add single black borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    # Remove existing borders if any
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)


def _set_table_borders(table):
    """Add borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def _set_cell_shading(cell, color_hex):
    """Set background shading color for a cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_run_font(run, name='宋体', size=10.5, bold=False, color=None):
    """Set font properties for a run, including East Asian font."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_title(doc, text, size=16):
    """Add a centered title paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_run_font(run, name='黑体', size=size, bold=True)
    return p


def add_heading1(doc, text):
    """Add a level-1 heading (一、二、三...)."""
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, name='黑体', size=14, bold=True)
    return p


def add_heading2(doc, text):
    """Add a level-2 heading (（一）（二）...)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, name='宋体', size=12, bold=True)
    return p


def add_body(doc, text, blue=True):
    """Add a body paragraph with optional blue font."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)  # 2 char indent
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    _set_run_font(run, name='宋体', size=10.5,
                  color=BLUE_RGB if blue else None)
    return p


def add_data_table(doc, headers, rows, col_widths=None):
    """Add a formatted data table.

    Args:
        headers: List of header strings.
        rows: List of row lists (each row is a list of cell value strings).
        col_widths: Optional list of column widths in cm.
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    # Header row
    for ci, header in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        _set_run_font(run, name='宋体', size=9, bold=True)
        _set_cell_shading(cell, 'D9E2F3')

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            # Right-align numeric columns (skip first column)
            if ci > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            _set_run_font(run, name='宋体', size=9,
                          color=BLUE_RGB)
            # Highlight subtotal/total rows
            item_name = str(row_data[0]) if row_data else ''
            if any(kw in item_name for kw in ['合计', '总计', '净额']):
                run.font.bold = True
                _set_cell_shading(cell, 'F2F2F2')

    # Set column widths
    if col_widths:
        for ci, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(width)

    return table


# ---------------------------------------------------------------------------
# Analysis Text Generators
# ---------------------------------------------------------------------------

def _fmt_pct(part, total):
    """Format percentage: part/total * 100, rounded to 2 decimals."""
    p = parse_numeric(part)
    t = parse_numeric(total)
    if t is None or t == 0 or p is None:
        return '-'
    return f'{p / t * 100:.2f}%'


def _fmt_change(curr, prev):
    """Format change between current and previous period.
    Returns (change_wan_str, change_pct_str, direction_str).
    """
    c = parse_numeric(curr)
    p = parse_numeric(prev)
    if c is None or p is None:
        return None, None, None
    change = c - p
    change_wan = to_wan_rounded(str(abs(change)))
    if p == 0:
        change_pct = '-'
    else:
        change_pct = f'{abs(change / p * 100):.1f}'
    direction = '增加' if change > 0 else ('减少' if change < 0 else '持平')
    return change_wan, change_pct, direction


def generate_asset_analysis(bs_data, prev_bs_data=None):
    """Generate asset analysis text."""
    total_assets = find_in_data(bs_data, '资产总计')
    curr_assets = find_in_data(bs_data, '流动资产合计')
    noncurr_assets = find_in_data(bs_data, '非流动资产合计')

    parts = []
    ta_wan = to_wan_rounded(total_assets)
    ca_wan = to_wan_rounded(curr_assets)
    nca_wan = to_wan_rounded(noncurr_assets)

    parts.append(
        f'截至报告期末，公司资产总额为{ta_wan}万元，'
        f'其中流动资产{ca_wan}万元，占资产总额的{_fmt_pct(curr_assets, total_assets)}；'
        f'非流动资产{nca_wan}万元，占资产总额的{_fmt_pct(noncurr_assets, total_assets)}。'
    )

    # Key asset breakdown
    monetary = find_in_data(bs_data, '货币资金')
    ar = find_in_data(bs_data, '应收账款')
    inventory = find_in_data(bs_data, '存货')
    fa = find_in_data(bs_data, '固定资产')

    detail_parts = []
    if monetary and monetary != '0':
        detail_parts.append(f'货币资金{to_wan_rounded(monetary)}万元')
    if ar and ar != '0':
        detail_parts.append(f'应收账款{to_wan_rounded(ar)}万元')
    if inventory and inventory != '0':
        detail_parts.append(f'存货{to_wan_rounded(inventory)}万元')
    if detail_parts:
        parts.append(
            f'流动资产中，{"，".join(detail_parts)}，是流动资产的主要构成部分。'
        )

    if fa and fa != '0':
        parts.append(f'非流动资产中，固定资产{to_wan_rounded(fa)}万元。')

    # Period-over-period comparison
    if prev_bs_data:
        prev_ta = find_in_data(prev_bs_data, '资产总计')
        chg_wan, chg_pct, direction = _fmt_change(total_assets, prev_ta)
        if chg_wan:
            parts.append(
                f'与上期末相比，资产总额{direction}了{chg_wan}万元，'
                f'{"增幅" if direction == "增加" else "降幅"}为{chg_pct}%。'
            )

    return ''.join(parts)


def generate_liability_analysis(bs_data, prev_bs_data=None):
    """Generate liability analysis text."""
    total_liab = find_in_data(bs_data, '负债合计')
    curr_liab = find_in_data(bs_data, '流动负债合计')
    noncurr_liab = find_in_data(bs_data, '非流动负债合计')
    total_assets = find_in_data(bs_data, '资产总计')

    parts = []
    tl_wan = to_wan_rounded(total_liab)
    cl_wan = to_wan_rounded(curr_liab)
    ncl_wan = to_wan_rounded(noncurr_liab)

    parts.append(
        f'截至报告期末，公司负债总额为{tl_wan}万元，'
        f'其中流动负债{cl_wan}万元，占负债总额的{_fmt_pct(curr_liab, total_liab)}；'
        f'非流动负债{ncl_wan}万元，占负债总额的{_fmt_pct(noncurr_liab, total_liab)}。'
    )

    # Key liability breakdown
    st_loan = find_in_data(bs_data, '短期借款')
    ap = find_in_data(bs_data, '应付账款')
    lt_loan = find_in_data(bs_data, '长期借款')

    detail_parts = []
    if st_loan and st_loan != '0':
        detail_parts.append(f'短期借款{to_wan_rounded(st_loan)}万元')
    if ap and ap != '0':
        detail_parts.append(f'应付账款{to_wan_rounded(ap)}万元')
    if detail_parts:
        parts.append(f'流动负债中，{"，".join(detail_parts)}。')

    if lt_loan and lt_loan != '0':
        parts.append(f'非流动负债中，长期借款{to_wan_rounded(lt_loan)}万元。')

    # Debt ratio
    debt_ratio = _fmt_pct(total_liab, total_assets)
    parts.append(f'资产负债率为{debt_ratio}。')

    # Period-over-period comparison
    if prev_bs_data:
        prev_tl = find_in_data(prev_bs_data, '负债合计')
        chg_wan, chg_pct, direction = _fmt_change(total_liab, prev_tl)
        if chg_wan:
            parts.append(
                f'与上期末相比，负债总额{direction}了{chg_wan}万元，'
                f'{"增幅" if direction == "增加" else "降幅"}为{chg_pct}%。'
            )

    return ''.join(parts)


def generate_equity_analysis(bs_data, prev_bs_data=None):
    """Generate equity analysis text."""
    equity = find_in_data(bs_data, '所有者权益合计', '股东权益合计')
    paid_in = find_in_data(bs_data, '实收资本', '股本')
    capital_reserve = find_in_data(bs_data, '资本公积')
    surplus_reserve = find_in_data(bs_data, '盈余公积')
    retained = find_in_data(bs_data, '未分配利润')

    parts = []
    eq_wan = to_wan_rounded(equity)

    parts.append(f'截至报告期末，公司所有者权益为{eq_wan}万元。')

    detail_parts = []
    if paid_in and paid_in != '0':
        detail_parts.append(f'实收资本{to_wan_rounded(paid_in)}万元')
    if capital_reserve and capital_reserve != '0':
        detail_parts.append(f'资本公积{to_wan_rounded(capital_reserve)}万元')
    if surplus_reserve and surplus_reserve != '0':
        detail_parts.append(f'盈余公积{to_wan_rounded(surplus_reserve)}万元')
    if retained and retained != '0':
        detail_parts.append(f'未分配利润{to_wan_rounded(retained)}万元')
    if detail_parts:
        parts.append(f'其中，{"，".join(detail_parts)}。')

    # Period-over-period comparison
    if prev_bs_data:
        prev_eq = find_in_data(prev_bs_data, '所有者权益合计', '股东权益合计')
        chg_wan, chg_pct, direction = _fmt_change(equity, prev_eq)
        if chg_wan:
            parts.append(
                f'与上期末相比，所有者权益{direction}了{chg_wan}万元，'
                f'{"增幅" if direction == "增加" else "降幅"}为{chg_pct}%。'
            )

    return ''.join(parts)


def generate_profit_analysis(is_data, prev_is_data=None):
    """Generate profit analysis text."""
    revenue = find_in_data(is_data, '营业收入')
    cost = find_in_data(is_data, '营业成本')
    op_profit = find_in_data(is_data, '营业利润')
    total_profit = find_in_data(is_data, '利润总额')
    income_tax = find_in_data(is_data, '所得税费用')
    net_profit = find_in_data(is_data, '净利润')

    parts = []
    parts.append(
        f'报告期内，公司实现营业收入{to_wan_rounded(revenue)}万元，'
        f'营业利润{to_wan_rounded(op_profit)}万元，'
        f'利润总额{to_wan_rounded(total_profit)}万元，'
        f'净利润{to_wan_rounded(net_profit)}万元。'
    )

    # Profitability ratios
    np_val = parse_numeric(net_profit)
    rev_val = parse_numeric(revenue)
    if rev_val and rev_val != 0 and np_val is not None:
        parts.append(f'营业净利率为{np_val / rev_val * 100:.2f}%。')

    cost_val = parse_numeric(cost)
    if rev_val and rev_val != 0 and cost_val is not None:
        parts.append(f'毛利率为{(rev_val - cost_val) / rev_val * 100:.2f}%。')

    # Period-over-period comparison
    if prev_is_data:
        prev_rev = find_in_data(prev_is_data, '营业收入')
        prev_np = find_in_data(prev_is_data, '净利润')
        chg_rev_wan, chg_rev_pct, dir_rev = _fmt_change(revenue, prev_rev)
        chg_np_wan, chg_np_pct, dir_np = _fmt_change(net_profit, prev_np)
        if chg_rev_wan:
            parts.append(
                f'与上年同期相比，营业收入{dir_rev}了{chg_rev_wan}万元，'
                f'{"增幅" if dir_rev == "增加" else "降幅"}为{chg_rev_pct}%；'
            )
        if chg_np_wan:
            parts.append(
                f'净利润{dir_np}了{chg_np_wan}万元，'
                f'{"增幅" if dir_np == "增加" else "降幅"}为{chg_np_pct}%。'
            )

    return ''.join(parts)


def generate_cashflow_analysis(cf_data, prev_cf_data=None):
    """Generate cash flow analysis text."""
    if not cf_data:
        return '报告期内现金流量数据暂缺。'

    op_net = find_in_data(cf_data, '经营活动产生的现金流量净额', '经营活动现金流量净额')
    inv_net = find_in_data(cf_data, '投资活动产生的现金流量净额', '投资活动现金流量净额')
    fin_net = find_in_data(cf_data, '筹资活动产生的现金流量净额', '筹资活动现金流量净额')
    cash_change = find_in_data(cf_data, '现金及现金等价物净增加额', '现金净增加额')
    cash_end = find_in_data(cf_data, '期末现金及现金等价物余额', '期末现金余额')

    parts = []
    parts.append(
        f'报告期内，公司经营活动产生的现金流量净额为{to_wan_rounded(op_net)}万元，'
        f'投资活动产生的现金流量净额为{to_wan_rounded(inv_net)}万元，'
        f'筹资活动产生的现金流量净额为{to_wan_rounded(fin_net)}万元。'
    )

    if cash_change and cash_change != '0':
        parts.append(f'现金及现金等价物净增加额为{to_wan_rounded(cash_change)}万元。')

    if cash_end and cash_end != '0':
        parts.append(f'期末现金及现金等价物余额为{to_wan_rounded(cash_end)}万元。')

    # Quality assessment
    op_val = parse_numeric(op_net)
    if op_val is not None:
        if op_val > 0:
            parts.append('经营活动现金流量净额为正，表明公司主营业务现金创造能力较好。')
        else:
            parts.append('经营活动现金流量净额为负，需关注公司经营性现金流状况。')

    return ''.join(parts)


def generate_ratio_analysis(ratios):
    """Generate 4-paragraph financial ratio analysis text (by category)."""
    parts = []
    for cat in RATIO_CATEGORIES:
        items = []
        if 'subgroups' in cat:
            for _, names in cat['subgroups']:
                items.extend(names)
        else:
            items = cat['items']
        seg_parts = []
        for name in items:
            val = ratios.get(name, '')
            if val and val != '-':
                seg_parts.append(f'{name}{val}')
        if seg_parts:
            parts.append(f'{cat["name"]}方面，{"，".join(seg_parts)}。')
    return ''.join(parts)


def assess_risk_level(ratios, bs_data=None, is_data=None, cf_data=None):
    """Assess risk level and return structured rating data.
    
    Returns:
        dict with 'overall_level', 'factors' (list of dicts), 'text'
    """
    risk_factors = []
    risk_level = '可控'
    level_priority = {'可控': 0, '需关注': 1, '较高': 2}

    # --- Short-term solvency (速动比率) ---
    quick_str = ratios.get('速动比率', '')
    if quick_str:
        val = parse_numeric(quick_str)
        if val is not None:
            if val < 0.5:
                rating, desc = '不足', f'速动比率{quick_str}偏低，短期偿债能力不足'
                risk_level = '较高'
            elif val < 1.0:
                rating, desc = '一般', f'速动比率{quick_str}低于1，短期偿债能力一般'
                if level_priority.get(risk_level, 0) < 1:
                    risk_level = '需关注'
            else:
                rating, desc = '正常', f'速动比率{quick_str}，短期偿债能力良好'
            risk_factors.append({
                'dimension': '短期偿债能力', 'indicator': '速动比率',
                'value': quick_str, 'rating': rating, 'desc': desc,
            })

    # --- Long-term solvency (资产负债率) ---
    debt_str = ratios.get('资产负债率', '')
    if debt_str:
        val = parse_numeric(debt_str.replace('%', ''))
        if val is not None:
            if val > 70:
                rating, desc = '偏高', f'资产负债率{debt_str}偏高，偿债压力较大'
                risk_level = '较高'
            elif val > 60:
                rating, desc = '偏高', f'资产负债率{debt_str}处于中等偏高水平'
                if level_priority.get(risk_level, 0) < 1:
                    risk_level = '需关注'
            else:
                rating, desc = '正常', f'资产负债率{debt_str}处于合理水平'
            risk_factors.append({
                'dimension': '长期偿债能力', 'indicator': '资产负债率',
                'value': debt_str, 'rating': rating, 'desc': desc,
            })

    # --- Profitability (净资产收益率) ---
    roe_str = ratios.get('净资产收益率', '')
    if roe_str:
        val = parse_numeric(roe_str.replace('%', ''))
        if val is not None:
            if val < 0:
                rating, desc = '堪忧', f'净资产收益率{roe_str}为负，盈利能力堪忧'
                risk_level = '较高'
            elif val < 3:
                rating, desc = '较弱', f'净资产收益率{roe_str}偏低，盈利能力较弱'
                if level_priority.get(risk_level, 0) < 1:
                    risk_level = '需关注'
            else:
                rating, desc = '正常', f'净资产收益率{roe_str}，盈利能力尚可'
            risk_factors.append({
                'dimension': '盈利能力', 'indicator': '净资产收益率',
                'value': roe_str, 'rating': rating, 'desc': desc,
            })

    # --- Cash flow quality ---
    if cf_data:
        op_net = find_in_data(cf_data, '经营活动产生的现金流量净额', '经营活动现金流量净额')
        op_val = parse_numeric(op_net) if op_net else None
        if op_val is not None:
            if op_val > 0:
                rating, desc = '正常', '经营活动现金流为正，现金创造能力良好'
            else:
                rating, desc = '需关注', '经营活动现金流为负，需关注现金流状况'
                if level_priority.get(risk_level, 0) < 1:
                    risk_level = '需关注'
            risk_factors.append({
                'dimension': '现金流状况', 'indicator': '经营净现金流',
                'value': to_wan_rounded(op_net) + '万元', 'rating': rating, 'desc': desc,
            })

    # --- Operating efficiency (soft dimension, added in v1.1) ---
    if bs_data and is_data:
        ar = find_in_data(bs_data, '应收账款')
        revenue = find_in_data(is_data, '营业收入')
        if ar and revenue:
            ar_val = parse_numeric(ar)
            rev_val = parse_numeric(revenue)
            if ar_val is not None and rev_val and rev_val != 0:
                ar_ratio = ar_val / rev_val * 100
                if ar_ratio > 30:
                    risk_factors.append({
                        'dimension': '营运效率', 'indicator': '应收账款占营收',
                        'value': f'{ar_ratio:.1f}%', 'rating': '需关注',
                        'desc': f'应收账款占营收{ar_ratio:.1f}%，营运资金占用偏高',
                    })
                    if level_priority.get(risk_level, 0) < 1:
                        risk_level = '需关注'
        inventory = find_in_data(bs_data, '存货')
        curr_assets = find_in_data(bs_data, '流动资产合计')
        if inventory and curr_assets:
            inv_val = parse_numeric(inventory)
            ca_val = parse_numeric(curr_assets)
            if inv_val is not None and ca_val and ca_val != 0:
                inv_pct = inv_val / ca_val * 100
                if inv_pct > 40:
                    risk_factors.append({
                        'dimension': '营运效率', 'indicator': '存货占流动资产',
                        'value': f'{inv_pct:.1f}%', 'rating': '需关注',
                        'desc': f'存货占流动资产{inv_pct:.1f}%，存在积压风险',
                    })
                    if level_priority.get(risk_level, 0) < 1:
                        risk_level = '需关注'

    # --- Growth capability (soft dimension, added in v1.1) ---
    np_growth = ratios.get('净利润增长率', '')
    rev_growth = ratios.get('销售增长率', '')
    if np_growth and np_growth != '-':
        npg_val = parse_numeric(np_growth.replace('%', ''))
        if npg_val is not None and npg_val < 0:
            rev_neg = False
            if rev_growth and rev_growth != '-':
                rvg_val = parse_numeric(rev_growth.replace('%', ''))
                rev_neg = rvg_val is not None and rvg_val < 0
            rating = '较高' if rev_neg else '需关注'
            risk_factors.append({
                'dimension': '发展能力', 'indicator': '净利润增长率',
                'value': np_growth, 'rating': rating,
                'desc': f'净利润增长率{np_growth}，成长性不足',
            })
            if rating == '较高':
                risk_level = '较高'
            elif level_priority.get(risk_level, 0) < 1:
                risk_level = '需关注'

    # Build summary text
    desc_list = [f['desc'] for f in risk_factors]
    text = '综合各项财务指标分析，' + '；'.join(desc_list) + '。' if desc_list else ''
    text += f'总体来看，公司财务风险{risk_level}。'

    return {
        'overall_level': risk_level,
        'factors': risk_factors,
        'text': text,
    }


def generate_risk_assessment(ratios, bs_data=None, is_data=None, cf_data=None):
    """Generate risk assessment text based on financial ratios."""
    result = assess_risk_level(ratios, bs_data, is_data, cf_data)
    return result['text']


def build_risk_rating_table(risk_result):
    """Build rows for risk rating table.
    
    Returns (headers, rows, overall_level).
    """
    headers = ['评估维度', '指标', '数值', '评级']
    rows = []
    for f in risk_result['factors']:
        rows.append([f['dimension'], f['indicator'], f['value'], f['rating']])
    # Add overall row
    rows.append(['综合评级', '', '', risk_result['overall_level']])
    return headers, rows, risk_result['overall_level']


def generate_conclusion(bs_data, is_data, ratios, company='', loan_amount=None):
    """Generate conclusion and recommendations."""
    parts = []

    total_assets = find_in_data(bs_data, '资产总计')
    net_profit = find_in_data(is_data, '净利润')
    debt_ratio = ratios.get('资产负债率', '')

    parts.append(
        f'综合以上分析，{company}截至报告期末资产总额{to_wan_rounded(total_assets)}万元，'
        f'资产负债率{debt_ratio}，报告期实现净利润{to_wan_rounded(net_profit)}万元。'
    )

    if loan_amount:
        loan_wan = to_wan_rounded(str(loan_amount))
        parts.append(
            f'公司现有贷款余额{loan_wan}万元，'
            f'结合其财务状况和经营情况，建议继续加强贷后管理，'
            f'定期跟踪财务指标变化，确保贷款资金安全。'
        )
    else:
        parts.append('建议持续关注公司财务状况变化，定期进行财务分析。')

    return ''.join(parts)


def generate_recommendations(ratios, bs_data, is_data, cf_data=None):
    """Generate numbered post-loan management recommendations based on data."""
    recs = []

    # 1. Debt ratio monitoring
    debt_str = ratios.get('资产负债率', '')
    if debt_str:
        debt_val = parse_numeric(debt_str.replace('%', ''))
        if debt_val is not None and debt_val > 60:
            recs.append(f'持续关注公司资产负债率变化趋势（当前{debt_str}），防范偿债风险。')
        else:
            recs.append(f'持续关注公司资产负债率变化趋势（当前{debt_str}）。')

    # 2. Receivables monitoring
    ar = find_in_data(bs_data, '应收账款')
    revenue = find_in_data(is_data, '营业收入')
    if ar and revenue:
        ar_val = parse_numeric(ar)
        rev_val = parse_numeric(revenue)
        if ar_val and rev_val and rev_val != 0:
            ar_ratio = ar_val / rev_val * 100
            if ar_ratio > 30:
                recs.append(f'跟踪应收账款回收情况（占营收{ar_ratio:.1f}%），关注资产质量与坏账风险。')
            else:
                recs.append('跟踪应收账款回收情况，关注资产质量。')

    # 3. Inventory monitoring
    inventory = find_in_data(bs_data, '存货')
    if inventory:
        inv_val = parse_numeric(inventory)
        if inv_val and inv_val > 0:
            curr_assets = find_in_data(bs_data, '流动资产合计')
            ca_val = parse_numeric(curr_assets) if curr_assets else None
            if ca_val and ca_val != 0:
                inv_pct = inv_val / ca_val * 100
                if inv_pct > 40:
                    recs.append(f'关注存货周转效率（存货占流动资产{inv_pct:.1f}%），防范积压风险。')

    # 4. Cash flow monitoring
    if cf_data:
        op_net = find_in_data(cf_data, '经营活动产生的现金流量净额', '经营活动现金流量净额')
        op_val = parse_numeric(op_net) if op_net else None
        if op_val is not None:
            if op_val > 0:
                recs.append('关注公司经营性现金流状况，确保还款来源稳定。')
            else:
                recs.append('经营活动现金流为负，需重点监控现金流改善情况及还款来源。')

    # 5. Profitability monitoring
    npm_str = ratios.get('营业净利率', '')
    if npm_str:
        npm_val = parse_numeric(npm_str.replace('%', ''))
        if npm_val is not None and npm_val < 3:
            recs.append(f'关注公司盈利能力改善情况（营业净利率{npm_str}偏低）。')

    # 6. Growth monitoring (v1.1)
    np_growth = ratios.get('净利润增长率', '')
    if np_growth and np_growth != '-':
        npg_val = parse_numeric(np_growth.replace('%', ''))
        if npg_val is not None and npg_val < 0:
            recs.append(f'关注公司成长性（净利润增长率{np_growth}为负），跟踪经营改善情况。')

    # Ensure at least 3 recommendations
    if len(recs) < 3:
        recs.append('定期复核财务报表，及时掌握经营变化。')

    return recs


# ---------------------------------------------------------------------------
# Table Data Builders
# ---------------------------------------------------------------------------

def build_bs_table_data(bs_data, prev_bs_data=None):
    """Build rows for balance sheet table."""
    rows = []

    # Asset section
    for item in BS_ASSET_ITEMS:
        val = find_in_data(bs_data, item)
        if val and val != '0':
            wan = to_wan_rounded(val)
            pct = _fmt_pct(val, find_in_data(bs_data, '资产总计'))
            row = [item, wan, pct]
            if prev_bs_data:
                prev_val = find_in_data(prev_bs_data, item)
                row.append(to_wan_rounded(prev_val) if prev_val and prev_val != '0' else '-')
            rows.append(row)

    # Liability section
    for item in BS_LIAB_ITEMS:
        val = find_in_data(bs_data, item)
        if val and val != '0':
            wan = to_wan_rounded(val)
            pct = _fmt_pct(val, find_in_data(bs_data, '负债合计'))
            row = [item, wan, pct]
            if prev_bs_data:
                prev_val = find_in_data(prev_bs_data, item)
                row.append(to_wan_rounded(prev_val) if prev_val and prev_val != '0' else '-')
            rows.append(row)

    # Equity section
    for item in BS_EQUITY_ITEMS:
        val = find_in_data(bs_data, item)
        if val and val != '0':
            wan = to_wan_rounded(val)
            pct = _fmt_pct(val, find_in_data(bs_data, '所有者权益合计', '股东权益合计'))
            row = [item, wan, pct]
            if prev_bs_data:
                prev_val = find_in_data(prev_bs_data, item)
                row.append(to_wan_rounded(prev_val) if prev_val and prev_val != '0' else '-')
            rows.append(row)

    return rows


def build_is_table_data(is_data, prev_is_data=None):
    """Build rows for income statement table."""
    rows = []
    for item in IS_KEY_ITEMS:
        val = find_in_data(is_data, item)
        if val and val != '0':
            wan = to_wan_rounded(val)
            row = [item, wan]
            if prev_is_data:
                prev_val = find_in_data(prev_is_data, item)
                row.append(to_wan_rounded(prev_val) if prev_val and prev_val != '0' else '-')
            rows.append(row)
    return rows


def build_cf_table_data(cf_data, prev_cf_data=None):
    """Build rows for cash flow statement table."""
    if not cf_data:
        return []
    rows = []
    for item in CF_KEY_ITEMS:
        val = find_in_data(cf_data, item)
        if val and val != '0':
            wan = to_wan_rounded(val)
            row = [item, wan]
            if prev_cf_data:
                prev_val = find_in_data(prev_cf_data, item)
                row.append(to_wan_rounded(prev_val) if prev_val and prev_val != '0' else '-')
            rows.append(row)
    return rows


def rate_ratio_value(name, value_str):
    """Rate a ratio value against its reference thresholds.

    Returns: 优秀/正常/一般/偏高/较弱, or '-' when no reference / value missing.
    lower_better: ≤good→优秀, ≤warning→正常, else→偏高
    higher_better: ≥good→优秀, ≥warning→正常, >0→一般, else→较弱
    """
    ref = RATIO_REFERENCES.get(name)
    if not ref or value_str is None or str(value_str) in ('-', ''):
        return '-'
    val = parse_numeric(str(value_str).replace('%', '').replace(',', ''))
    if val is None:
        return '-'
    if ref['direction'] == 'lower_better':
        if val <= ref['good']:
            return '优秀'
        if val <= ref['warning']:
            return '正常'
        return '偏高'
    else:  # higher_better
        if val >= ref['good']:
            return '优秀'
        if val >= ref['warning']:
            return '正常'
        if val > 0:
            return '一般'
        return '较弱'


def _ratio_row(name, ratios, prev_ratios=None):
    """Build a single ratio row: [指标, 本期值, 上期值, 参考值, 评估]."""
    value = ratios.get(name, '-')
    prev = prev_ratios.get(name, '-') if prev_ratios else '-'
    ref = RATIO_REFERENCES.get(name)
    if ref:
        op = '<=' if ref['direction'] == 'lower_better' else '>='
        if ref['unit'] == 'percent':
            ref_str = f'{op} {ref["warning"]}%'
        else:
            ref_str = f'{op} {ref["warning"]}'
    else:
        ref_str = '-'
    return [name, value, prev, ref_str, rate_ratio_value(name, value)]


def build_ratio_tables(ratios, prev_ratios=None):
    """Build 4-category ratio tables for the report.

    Returns a list of dicts:
      [{'title': '偿债能力（短期）', 'headers': ['指标','本期值','上期值','参考值','评估'],
        'rows': [[...], ...]}, ...]

    Every indicator row is always kept; missing values show '-'.
    """
    tables = []
    for cat in RATIO_CATEGORIES:
        if 'subgroups' in cat:
            for sub_name, items in cat['subgroups']:
                tables.append({
                    'title': f'{cat["name"]}（{sub_name}）',
                    'headers': ['指标', '本期值', '上期值', '参考值', '评估'],
                    'rows': [_ratio_row(name, ratios, prev_ratios) for name in items],
                })
        else:
            tables.append({
                'title': cat['name'],
                'headers': ['指标', '本期值', '上期值', '参考值', '评估'],
                'rows': [_ratio_row(name, ratios, prev_ratios) for name in cat['items']],
            })
    return tables


# ---------------------------------------------------------------------------
# Main Report Generator
# ---------------------------------------------------------------------------

def generate_report(
    company: str,
    period: str,
    bs_data: dict,
    is_data: dict,
    cf_data: dict | None = None,
    output_path: str = '贷后分析报告.docx',
    prev_bs_data: dict | None = None,
    prev_is_data: dict | None = None,
    prev_cf_data: dict | None = None,
    loan_amount: float | None = None,
    loan_type: str = '',
    report_type: str = '贷后分析报告',
    guarantor: str = '',
    blue_all: bool = True,
    enterprise_info: dict | None = None,
):
    """Generate a complete financial analysis Word report.

    Args:
        company: Enterprise name.
        period: Report period (e.g., "2025年12月", "2025年一季度").
        bs_data: Balance sheet data dict {account_name: amount_string}.
        is_data: Income statement data dict.
        cf_data: Cash flow statement data dict (optional).
        output_path: Output .docx file path.
        prev_bs_data: Previous period balance sheet (for comparison).
        prev_is_data: Previous period income statement.
        prev_cf_data: Previous period cash flow statement.
        loan_amount: Loan amount in yuan (optional).
        loan_type: Loan type description (optional).
        report_type: Report type label (default: "贷后分析报告").
        guarantor: Guarantor/collateral info (optional).
        blue_all: If True, all generated text is blue (for review).
        enterprise_info: Dict with keys like 'establish_date', 'registered_capital',
                         'main_business', 'industry', 'loan_period', 'collateral'.
    """
    doc = Document()

    # Page setup: A4, standard margins
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ---- Title ----
    add_title(doc, f'{company}{report_type}', size=18)

    # Report info
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_p.add_run(f'报告期间：{period}')
    _set_run_font(run, name='宋体', size=11, color=BLUE_RGB if blue_all else None)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = datetime.now().strftime('%Y年%m月%d日')
    run = date_p.add_run(f'编制日期：{today}')
    _set_run_font(run, name='宋体', size=11, color=BLUE_RGB if blue_all else None)

    doc.add_paragraph()  # spacer

    # ---- Section 1: Enterprise Overview ----
    add_heading1(doc, '一、企业基本情况')

    # (1) Enterprise profile
    add_heading2(doc, '（一）企业概况')
    if enterprise_info:
        info_parts = [f'企业名称：{company}']
        if enterprise_info.get('establish_date'):
            info_parts.append(f'成立日期：{enterprise_info["establish_date"]}')
        if enterprise_info.get('registered_capital'):
            info_parts.append(f'注册资本：{enterprise_info["registered_capital"]}万元')
        if enterprise_info.get('main_business'):
            info_parts.append(f'主营业务：{enterprise_info["main_business"]}')
        if enterprise_info.get('industry'):
            info_parts.append(f'所属行业：{enterprise_info["industry"]}')
        add_body(doc, '。'.join(info_parts) + '。', blue=blue_all)
    else:
        add_body(doc, f'企业名称：{company}。', blue=blue_all)

    # (2) Loan info
    has_loan_info = loan_amount or loan_type or guarantor or (
        enterprise_info and (enterprise_info.get('loan_period') or enterprise_info.get('collateral'))
    )
    if has_loan_info:
        add_heading2(doc, '（二）贷款情况')
        loan_parts = []
        if loan_type:
            loan_parts.append(f'贷款类型：{loan_type}')
        if loan_amount:
            loan_parts.append(f'贷款金额：{to_wan_rounded(str(loan_amount))}万元')
        if enterprise_info:
            if enterprise_info.get('loan_period'):
                loan_parts.append(f'贷款期限：{enterprise_info["loan_period"]}')
            if enterprise_info.get('collateral'):
                loan_parts.append(f'抵押物：{enterprise_info["collateral"]}')
        if guarantor:
            loan_parts.append(f'担保方式：{guarantor}')
        add_body(doc, '。'.join(loan_parts) + '。', blue=blue_all)

    section_num = 2

    # ---- Section 2: Financial Analysis ----
    add_heading1(doc, f'{"一二三四五六"[section_num-1]}、财务状况分析')

    # (1) Balance Sheet Analysis
    add_heading2(doc, '（一）资产负债分析')

    bs_headers = ['项目', '期末数(万元)', '占比']
    if prev_bs_data:
        bs_headers.append('上期末(万元)')
    bs_rows = build_bs_table_data(bs_data, prev_bs_data)
    add_data_table(doc, bs_headers, bs_rows)
    doc.add_paragraph()  # spacer

    add_body(doc, generate_asset_analysis(bs_data, prev_bs_data), blue=blue_all)
    add_body(doc, generate_liability_analysis(bs_data, prev_bs_data), blue=blue_all)
    add_body(doc, generate_equity_analysis(bs_data, prev_bs_data), blue=blue_all)

    # (2) Income Statement Analysis
    add_heading2(doc, '（二）利润分析')

    is_headers = ['项目', '本期金额(万元)']
    if prev_is_data:
        is_headers.append('上年同期(万元)')
    is_rows = build_is_table_data(is_data, prev_is_data)
    add_data_table(doc, is_headers, is_rows)
    doc.add_paragraph()

    add_body(doc, generate_profit_analysis(is_data, prev_is_data), blue=blue_all)

    # (3) Cash Flow Analysis
    if cf_data:
        add_heading2(doc, '（三）现金流量分析')

        cf_headers = ['项目', '本期金额(万元)']
        if prev_cf_data:
            cf_headers.append('上期(万元)')
        cf_rows = build_cf_table_data(cf_data, prev_cf_data)
        if cf_rows:
            add_data_table(doc, cf_headers, cf_rows)
            doc.add_paragraph()

        add_body(doc, generate_cashflow_analysis(cf_data, prev_cf_data), blue=blue_all)

    # ---- Section 3: Financial Ratios ----
    section_num += 1
    add_heading1(doc, f'{"一二三四五六"[section_num-1]}、主要财务指标分析')

    ratios = calculate_financial_ratios(bs_data, is_data, cf_data,
                                        prev_bs_data, prev_is_data, prev_cf_data)
    prev_ratios = None
    if prev_bs_data and prev_is_data:
        # 不传上上期数据 → 平均余额/增长率类指标为 '-'
        prev_ratios = calculate_financial_ratios(prev_bs_data, prev_is_data)

    for i, t in enumerate(build_ratio_tables(ratios, prev_ratios), 1):
        add_heading2(doc, f'（{"一二三四五六七八九十"[i-1]}）{t["title"]}')
        add_data_table(doc, t['headers'], t['rows'])
        doc.add_paragraph()

    add_body(doc, generate_ratio_analysis(ratios), blue=blue_all)

    # ---- Section 4: Risk Assessment ----
    section_num += 1
    add_heading1(doc, f'{"一二三四五六"[section_num-1]}、贷款风险评估')

    risk_result = assess_risk_level(ratios, bs_data, is_data, cf_data)
    add_body(doc, risk_result['text'], blue=blue_all)
    doc.add_paragraph()

    # Risk rating table
    risk_headers, risk_rows, overall = build_risk_rating_table(risk_result)
    add_data_table(doc, risk_headers, risk_rows)
    doc.add_paragraph()

    # ---- Section 5: Conclusion & Recommendations ----
    section_num += 1
    add_heading1(doc, f'{"一二三四五六"[section_num-1]}、结论与建议')
    add_body(doc, generate_conclusion(bs_data, is_data, ratios, company, loan_amount), blue=blue_all)

    # Numbered recommendations
    recs = generate_recommendations(ratios, bs_data, is_data, cf_data)
    rec_intro = doc.add_paragraph()
    rec_intro.paragraph_format.first_line_indent = Cm(0.74)
    rec_intro.paragraph_format.line_spacing = 1.5
    run = rec_intro.add_run('贷后管理建议：')
    _set_run_font(run, name='宋体', size=10.5, bold=True,
                  color=BLUE_RGB if blue_all else None)

    for i, rec in enumerate(recs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f'{i}. {rec}')
        _set_run_font(run, name='宋体', size=10.5,
                      color=BLUE_RGB if blue_all else None)

    # Save
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _load_json(path):
    """Load a JSON file into a dict."""
    if not path:
        return None
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def main():
    parser = argparse.ArgumentParser(
        description='Auto-generate a financial analysis Word report from financial data.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument('--company', required=True, help='Enterprise name')
    parser.add_argument('--period', required=True, help='Report period (e.g., "2025年12月")')
    parser.add_argument('--bs-data', required=True, help='JSON file with balance sheet data')
    parser.add_argument('--is-data', required=True, help='JSON file with income statement data')
    parser.add_argument('--cf-data', help='JSON file with cash flow data (optional)')
    parser.add_argument('--output', required=True, help='Output .docx file path')
    parser.add_argument('--prev-bs-data', help='JSON file with previous period balance sheet')
    parser.add_argument('--prev-is-data', help='JSON file with previous period income statement')
    parser.add_argument('--prev-cf-data', help='JSON file with previous period cash flow')
    parser.add_argument('--loan-amount', type=float, help='Loan amount in yuan')
    parser.add_argument('--loan-type', help='Loan type description')
    parser.add_argument('--guarantor', help='Guarantor/collateral info')
    parser.add_argument('--report-type', default='贷后分析报告', help='Report type label')
    parser.add_argument('--no-blue', action='store_true', help='Do not color text blue')
    parser.add_argument('--establish-date', help='Enterprise establishment date')
    parser.add_argument('--registered-capital', help='Registered capital (万元)')
    parser.add_argument('--main-business', help='Main business description')
    parser.add_argument('--industry', help='Industry classification')
    parser.add_argument('--loan-period', help='Loan period')
    parser.add_argument('--collateral', help='Collateral description')

    args = parser.parse_args()

    enterprise_info = {}
    if args.establish_date:
        enterprise_info['establish_date'] = args.establish_date
    if args.registered_capital:
        enterprise_info['registered_capital'] = args.registered_capital
    if args.main_business:
        enterprise_info['main_business'] = args.main_business
    if args.industry:
        enterprise_info['industry'] = args.industry
    if args.loan_period:
        enterprise_info['loan_period'] = args.loan_period
    if args.collateral:
        enterprise_info['collateral'] = args.collateral

    output = generate_report(
        company=args.company,
        period=args.period,
        bs_data=_load_json(args.bs_data),
        is_data=_load_json(args.is_data),
        cf_data=_load_json(args.cf_data),
        output_path=args.output,
        prev_bs_data=_load_json(args.prev_bs_data),
        prev_is_data=_load_json(args.prev_is_data),
        prev_cf_data=_load_json(args.prev_cf_data),
        loan_amount=args.loan_amount,
        loan_type=args.loan_type or '',
        report_type=args.report_type,
        guarantor=args.guarantor or '',
        blue_all=not args.no_blue,
        enterprise_info=enterprise_info if enterprise_info else None,
    )

    print(f'Report generated: {output}')


if __name__ == '__main__':
    main()
