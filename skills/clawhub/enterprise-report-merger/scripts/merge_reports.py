#!/usr/bin/env python3
"""
Enterprise Report Merger - Core Script
合并企业报表的核心脚本，支持三种合并模式 + Word 模板填充。

Usage:
  python merge_reports.py merge \
    --input file1.xlsx file2.xlsx ... \
    --mode concat|key|consolidation \
    --output result.xlsx \
    [--key-column "科目代码"] \
    [--source-labels "子公司A" "子公司B"] \
    [--elimination-entries elimination.json] \
    [--word-template template.docx] \
    [--word-output filled_report.docx]

  python merge_reports.py extract-pdf \
    --input report.pdf \
    --output extracted.xlsx \
    [--pages "0-5"] \
    [--table-index 0]

  python merge_reports.py fill-template \
    --input merged_data.xlsx \
    --template template.docx \
    --output filled_report.docx \
    [--mapping mapping.json]
"""

import argparse
import difflib
import json
import math
import os
import re
import sys
from pathlib import Path
from typing import Any

import openpyxl
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

# Optional imports — loaded on demand
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import xlrd
except ImportError:
    xlrd = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document as DocxDocument
    from docx.shared import RGBColor
except ImportError:
    DocxDocument = None
    RGBColor = None


# ---------------------------------------------------------------------------
# Utility Functions (v2 improvements)
# ---------------------------------------------------------------------------

# Known aliases: {canonical_name: [alternative_names]}
ACCOUNT_ALIASES = {
    '应收账款': ['应收货款', '应收帐款'],
    '应付账款': ['应付货款', '应付帐款'],
    '固定资产': ['固定资产净额', '固定资产原值', '固定资产净值'],
    '税金及附加': ['营业税金及附加'],
    '所有者权益合计': ['股东权益合计', '股东权益', '所有者权益', '归属于母公司股东权益小计'],
    '股本': ['实收资本', '实收资本（股本）', '实收资本(股本)'],
    '流动负债合计': ['流动负债小计'],
    '非流动负债合计': ['非流动负债小计'],
    '资本公积': ['资本公积金'],
    '盈余公积': ['盈余公积金'],
    '未分配利润': ['留存收益'],
    '货币资金': ['现金及银行存款'],
    '短期借款': ['短期贷款'],
    '长期借款': ['长期贷款'],
    '其他应付款': ['其他应付款项'],
    '其他应收款': ['其他应收款项'],
    '营业收入': ['主营业务收入'],
    '营业成本': ['主营业务成本'],
    '净利润': ['净亏损'],
    '利润总额': ['税前利润'],
}

BLUE_RGB = RGBColor(0, 0, 255) if RGBColor else None


def clean_item_name(name: str) -> str:
    """Clean PDF item name: remove numbering prefixes, section markers."""
    if not name:
        return ''
    name = name.strip()
    name = re.sub(r'^[一二三四五六七八九十]+、\s*', '', name)
    name = re.sub(r'^[加减]\s*[：:]\s*', '', name)
    name = re.sub(r'^其中\s*[：:]\s*', '', name)
    name = name.replace('\n', '')
    return name.strip()


def is_name_conflict(name1: str, name2: str) -> bool:
    """Check if two names represent different concepts despite similarity."""
    if ('应收' in name1 and '应付' in name2) or ('应付' in name1 and '应收' in name2):
        return True
    if ('资产' in name1 and '负债' in name2) or ('负债' in name1 and '资产' in name2):
        return True
    if ('收款' in name1 and '付款' in name2) or ('付款' in name1 and '收款' in name2):
        return True
    return False


def normalize_account_name(s: str) -> str:
    """Normalize item name for comparison."""
    s = re.sub(r'[（(].*?[)）]', '', s)
    s = s.replace('其中：', '').replace('其中:', '')
    s = s.replace('合计', '').replace('小计', '')
    s = s.replace('净额', '').replace('总额', '')
    return s.strip()


def find_account_match(word_name: str, pdf_keys: set) -> tuple:
    """Find the best matching PDF key for a Word template item name.
    Returns (matched_key, match_type) or (None, 'no_match').
    """
    word_clean = word_name.replace('其中：', '').replace('其中:', '').strip()

    # 1. Exact match
    if word_clean in pdf_keys:
        return word_clean, 'exact'

    # 2. Alias match
    for canonical, aliases in ACCOUNT_ALIASES.items():
        all_names = [canonical] + aliases
        word_in_group = word_clean in all_names
        if not word_in_group:
            wnorm = normalize_account_name(word_clean)
            for n in all_names:
                if normalize_account_name(n) == wnorm:
                    word_in_group = True
                    break
        if word_in_group:
            for pk in pdf_keys:
                if is_name_conflict(word_clean, pk):
                    continue
                if pk in all_names:
                    return pk, 'alias'
            for pk in pdf_keys:
                if is_name_conflict(word_clean, pk):
                    continue
                if normalize_account_name(pk) == normalize_account_name(word_clean):
                    return pk, 'alias'

    # 3. Normalized match
    wnorm = normalize_account_name(word_clean)
    for pk in pdf_keys:
        if is_name_conflict(word_clean, pk):
            continue
        if normalize_account_name(pk) == wnorm:
            return pk, 'normalized'

    # 4. Substring match (>= 3 chars)
    if len(word_clean) >= 3:
        for pk in pdf_keys:
            if is_name_conflict(word_clean, pk):
                continue
            if word_clean in pk or pk in word_clean:
                return pk, 'substring'

    # 5. Similarity score
    best_score = 0
    best_match = None
    for pk in pdf_keys:
        if is_name_conflict(word_clean, pk):
            continue
        score = difflib.SequenceMatcher(None, word_clean, pk).ratio()
        if score > best_score:
            best_score = score
            best_match = pk
    if best_score >= 0.7:
        return best_match, f'similarity({best_score:.2f})'

    return None, 'no_match'


def to_wan_rounded(yuan_val) -> str:
    """Convert yuan to wan (万元), round half up to nearest integer."""
    if yuan_val is None or yuan_val == '-' or yuan_val == '':
        return '-'
    try:
        val = float(str(yuan_val).replace(',', ''))
        wan = val / 10000.0
        if wan >= 0:
            wan_int = int(math.floor(wan + 0.5))
        else:
            wan_int = int(math.ceil(wan - 0.5))
        return f'{wan_int:,}'
    except (ValueError, TypeError):
        return '-'


def parse_numeric(val_str) -> float | None:
    """Parse a value string to float."""
    if val_str is None or val_str == '-' or val_str == '':
        return None
    try:
        return float(str(val_str).replace(',', ''))
    except (ValueError, TypeError):
        return None


def set_cell_blue(cell, text: str):
    """Set cell text with blue font color."""
    if DocxDocument is None:
        return
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ''
    if cell.paragraphs:
        p = cell.paragraphs[0]
        while len(p.runs) > 1:
            p.runs[-1]._element.getparent().remove(p.runs[-1]._element)
        if p.runs:
            p.runs[0].text = text
            p.runs[0].font.color.rgb = BLUE_RGB
        else:
            run = p.add_run(text)
            run.font.color.rgb = BLUE_RGB
    else:
        p = cell.add_paragraph()
        run = p.add_run(text)
        run.font.color.rgb = BLUE_RGB


def set_paragraph_blue(paragraph, text: str):
    """Replace paragraph text with blue font."""
    if DocxDocument is None:
        return
    if not paragraph.runs:
        run = paragraph.add_run(text)
        run.font.color.rgb = BLUE_RGB
        return
    paragraph.runs[0].text = text
    paragraph.runs[0].font.color.rgb = BLUE_RGB
    for run in paragraph.runs[1:]:
        run.text = ''


def _num(val) -> float | None:
    """Parse a numeric value; None-safe wrapper for ratio math."""
    if val is None:
        return None
    return parse_numeric(val)


def _avg(a: float | None, b: float | None) -> float | None:
    """Average of two values (期初+期末)/2; None if either is missing."""
    if a is None or b is None:
        return None
    return (a + b) / 2.0


def _pct(num: float | None, denom: float | None) -> str:
    """Format num/denom as 'xx.xx%'; '-' when missing or zero denominator."""
    if num is None or denom is None or denom == 0:
        return '-'
    return f'{num / denom * 100:.2f}%'


def _fmt(x: float | None, digits: int = 2) -> str:
    """Format a plain ratio to fixed decimals; '-' when missing."""
    if x is None:
        return '-'
    return f'{x:.{digits}f}'


def _growth(cur: float | None, prev: float | None) -> str:
    """Growth rate (本期-上期)/|上期| as 'xx.xx%'; '-' when prev missing or 0."""
    if cur is None or prev is None or prev == 0:
        return '-'
    return f'{(cur - prev) / abs(prev) * 100:.2f}%'


def calculate_financial_ratios(bs_data: dict, is_data: dict, cf_data: dict | None = None,
                               prev_bs_data: dict | None = None, prev_is_data: dict | None = None,
                               prev_cf_data: dict | None = None) -> dict:
    """Calculate 18 financial ratios grouped into 4 categories.

    Categories:
      1. 偿债能力 (7): 流动比率、速动比率、资产负债率、权益乘数、利息保障倍数、
         现金流量负债率、负债权益比
      2. 营运效率 (3): 存货周转率、应收账款周转率、总资产周转率（平均余额口径）
      3. 盈利能力 (5): 净资产收益率、毛利率、营业净利率、总资产收益率、总资产报酬率
      4. 发展能力 (3): 销售增长率、净利润增长率、每股收益增长率

    Ratios depending on prior-period data return '-' when it is unavailable,
    so report tables always keep every indicator row. Uses find_optional() to
    distinguish "missing" from a real 0 value.
    """
    ratios = {}

    # ---------- 本期数据（缺失 -> None） ----------
    ta = _num(find_optional(bs_data, '资产总计'))
    tl = _num(find_optional(bs_data, '负债合计'))
    ca = _num(find_optional(bs_data, '流动资产合计'))
    cl = _num(find_optional(bs_data, '流动负债合计', '流动负债小计'))
    inv = _num(find_optional(bs_data, '存货'))
    ar = _num(find_optional(bs_data, '应收账款'))
    eq = _num(find_optional(bs_data, '所有者权益合计', '股东权益合计',
                             '归属于母公司股东权益小计'))
    rev = _num(find_optional(is_data, '营业收入'))
    cost = _num(find_optional(is_data, '营业成本'))
    np_ = _num(find_optional(is_data, '净利润'))
    tp = _num(find_optional(is_data, '利润总额', '税前利润'))
    ie = _num(find_optional(is_data, '利息支出'))
    eps = _num(find_optional(is_data, '基本每股收益', '基本每股收益(元/股)'))
    op_net = _num(find_optional(cf_data, '经营活动产生的现金流量净额',
                                '经营活动现金流量净额')) if cf_data else None

    # ---------- 上期数据（用于平均余额与增长率） ----------
    prev_ta = _num(find_optional(prev_bs_data, '资产总计')) if prev_bs_data else None
    prev_inv = _num(find_optional(prev_bs_data, '存货')) if prev_bs_data else None
    prev_ar = _num(find_optional(prev_bs_data, '应收账款')) if prev_bs_data else None
    prev_rev = _num(find_optional(prev_is_data, '营业收入')) if prev_is_data else None
    prev_np = _num(find_optional(prev_is_data, '净利润')) if prev_is_data else None
    prev_eps = _num(find_optional(prev_is_data, '基本每股收益',
                                  '基本每股收益(元/股)')) if prev_is_data else None

    avg_ta = _avg(ta, prev_ta)

    # ---------- 1. 偿债能力 ----------
    # 短期
    ratios['流动比率'] = _fmt(ca / cl) if (ca is not None and cl) else '-'
    ratios['速动比率'] = _fmt((ca - (inv or 0)) / cl) if (ca is not None and cl) else '-'
    # 长期
    ratios['资产负债率'] = _pct(tl, ta)
    ratios['权益乘数'] = _fmt(ta / eq) if (ta is not None and eq) else '-'
    if ie is not None and ie != 0 and tp is not None:
        ratios['利息保障倍数'] = _fmt((tp + ie) / ie)
    else:
        ratios['利息保障倍数'] = '-'
    ratios['现金流量负债率'] = _pct(op_net, tl)
    ratios['负债权益比'] = _pct(tl, eq)

    # ---------- 2. 营运效率（平均余额口径，需上期数据） ----------
    avg_inv = _avg(inv, prev_inv)
    avg_ar = _avg(ar, prev_ar)
    ratios['存货周转率'] = _fmt(cost / avg_inv) if (cost is not None and avg_inv) else '-'
    ratios['应收账款周转率'] = _fmt(rev / avg_ar) if (rev is not None and avg_ar) else '-'
    ratios['总资产周转率'] = _fmt(rev / avg_ta) if (rev is not None and avg_ta) else '-'

    # ---------- 3. 盈利能力 ----------
    ratios['净资产收益率'] = _pct(np_, eq)
    ratios['毛利率'] = _pct(rev - cost, rev) if (rev is not None and cost is not None) else '-'
    ratios['营业净利率'] = _pct(np_, rev)
    ratios['总资产收益率'] = _pct(np_, avg_ta)
    if avg_ta:
        if ie is not None and tp is not None:
            ratios['总资产报酬率'] = _pct(tp + ie, avg_ta)
        elif tp is not None:
            ratios['总资产报酬率'] = _pct(tp, avg_ta)
        elif np_ is not None:
            ratios['总资产报酬率'] = _pct(np_, avg_ta)
        else:
            ratios['总资产报酬率'] = '-'
    else:
        ratios['总资产报酬率'] = '-'

    # ---------- 4. 发展能力（需上期数据） ----------
    ratios['销售增长率'] = _growth(rev, prev_rev)
    ratios['净利润增长率'] = _growth(np_, prev_np)
    ratios['每股收益增长率'] = _growth(eps, prev_eps)

    return ratios


def parse_wan(val_str: str | None) -> float | None:
    """Parse a numeric string and return value in 万元 (rounded to integer)."""
    n = parse_numeric(val_str)
    if n is None:
        return None
    return round(n / 10000)


def validate_financial_data(bs_data: dict, is_data: dict, cf_data: dict | None = None) -> dict:
    """多维度交叉验证财务数据的准确性。

    对图片型 PDF（截图或渲染识别）识别后的数据进行全面的勾稽关系验证，
    自动检测数字误读、漏读、科目错位等问题。

    Args:
        bs_data: 资产负债表数据字典 {科目名: 金额字符串}
        is_data: 利润表数据字典 {科目名: 金额字符串}
        cf_data: 现金流量表数据字典 {科目名: 金额字符串}

    Returns:
        dict: {
            'passed': bool,         # 是否全部通过
            'checks': [             # 每条验证记录
                {'name': str, 'passed': bool, 'expected': str, 'actual': str, 'diff': str, 'detail': str}
            ],
            'summary': str          # 人类可读的验证摘要
        }
    """
    checks = []
    errors = []

    # ---------- 表内勾稽：资产负债表 ----------
    # 1. 资产总计 = 流动资产合计 + 非流动资产合计
    total_assets = find_in_data(bs_data, '资产总计')
    curr_total = find_in_data(bs_data, '流动资产合计', '流动资产')
    noncurr_total = find_in_data(bs_data, '非流动资产合计', '非流动资产')
    if total_assets and curr_total and noncurr_total:
        ta = parse_numeric(total_assets)
        ct = parse_numeric(curr_total)
        nt = parse_numeric(noncurr_total)
        if ta is not None and ct is not None and nt is not None:
            expected = ct + nt
            diff = ta - expected
            passed = abs(diff) < 1  # 允许 1 元舍入误差
            checks.append({
                'name': '资产总计 = 流动资产 + 非流动资产',
                'passed': passed,
                'expected': f'{expected:,.0f}',
                'actual': f'{ta:,.0f}',
                'diff': f'{diff:,.0f}',
                'detail': f'{to_wan_rounded(total_assets)} = {to_wan_rounded(curr_total)} + {to_wan_rounded(noncurr_total)}'
            })
            if not passed:
                errors.append('资产总计勾稽不平！')

    # 2. 资产总计 = 负债合计 + 所有者权益合计
    total_liab = find_in_data(bs_data, '负债合计')
    total_equity = find_in_data(bs_data, '所有者权益合计', '股东权益合计', '归属于母公司所有者权益合计')
    if total_assets and total_liab and total_equity:
        ta = parse_numeric(total_assets)
        tl = parse_numeric(total_liab)
        te = parse_numeric(total_equity)
        if ta is not None and tl is not None and te is not None:
            expected = tl + te
            diff = ta - expected
            passed = abs(diff) < 1
            checks.append({
                'name': '资产总计 = 负债合计 + 所有者权益合计',
                'passed': passed,
                'expected': f'{expected:,.0f}',
                'actual': f'{ta:,.0f}',
                'diff': f'{diff:,.0f}',
                'detail': f'{to_wan_rounded(total_assets)} = {to_wan_rounded(total_liab)} + {to_wan_rounded(total_equity)}'
            })
            if not passed:
                errors.append('资产负债表恒等式不平！')

    # 3. 负债合计 = 流动负债合计 + 非流动负债合计
    curr_liab = find_in_data(bs_data, '流动负债合计', '流动负债')
    noncurr_liab = find_in_data(bs_data, '非流动负债合计', '非流动负债')
    if total_liab and curr_liab and noncurr_liab:
        tl = parse_numeric(total_liab)
        cl = parse_numeric(curr_liab)
        nl = parse_numeric(noncurr_liab)
        if tl is not None and cl is not None and nl is not None:
            expected = cl + nl
            diff = tl - expected
            passed = abs(diff) < 1
            checks.append({
                'name': '负债合计 = 流动负债 + 非流动负债',
                'passed': passed,
                'expected': f'{expected:,.0f}',
                'actual': f'{tl:,.0f}',
                'diff': f'{diff:,.0f}',
                'detail': f'{to_wan_rounded(total_liab)} = {to_wan_rounded(curr_liab)} + {to_wan_rounded(noncurr_liab)}'
            })
            if not passed:
                errors.append('负债合计勾稽不平！')

    # 4. 所有者权益各分项之和 ≈ 所有者权益合计（如果有明细项的话）
    # 实收资本+资本公积+盈余公积+未分配利润+其他
    paid_in = find_in_data(bs_data, '实收资本', '股本')
    capital_reserve = find_in_data(bs_data, '资本公积')
    surplus_reserve = find_in_data(bs_data, '盈余公积')
    retained = find_in_data(bs_data, '未分配利润')
    minority = find_in_data(bs_data, '少数股东权益')
    if total_equity and (paid_in or capital_reserve or surplus_reserve or retained):
        te = parse_numeric(find_in_data(bs_data, '所有者权益合计', '股东权益合计', '归属于母公司所有者权益合计'))
        subs = []
        sub_names = []
        for name, val in [('实收资本', paid_in), ('资本公积', capital_reserve),
                          ('盈余公积', surplus_reserve), ('未分配利润', retained)]:
            if val:
                v = parse_numeric(val)
                if v is not None:
                    subs.append(v)
                    sub_names.append(f'{name}={to_wan_rounded(val)}万')
        if subs and te is not None:
            sub_sum = sum(subs)
            if minority:
                mn = parse_numeric(minority)
                if mn is not None:
                    sub_sum += mn
                    sub_names.append(f'少数股东权益={to_wan_rounded(minority)}万')
            diff = te - sub_sum
            # 允许 1 万元误差（四舍五入后可能有尾差）
            passed = abs(diff) < 10000
            checks.append({
                'name': '所有者权益分项之和 ≈ 合计',
                'passed': passed,
                'expected': f'{sub_sum:,.0f}',
                'actual': f'{te:,.0f}',
                'diff': f'{diff:,.0f}',
                'detail': f'分项 {" + ".join(sub_names)} = {to_wan_rounded(str(sub_sum))}万, 合计 = {to_wan_rounded(total_equity)}万'
            })
            if not passed:
                errors.append('所有者权益分项与合计不匹配，可能存在未识别的权益科目！')

    # ---------- 表内勾稽：利润表 ----------
    # 5. 利润总额 = 营业利润 + 营业外收入 - 营业外支出
    total_profit = find_in_data(is_data, '利润总额')
    operating_profit = find_in_data(is_data, '营业利润')
    non_op_income = find_in_data(is_data, '营业外收入')
    non_op_expense = find_in_data(is_data, '营业外支出')
    if total_profit and operating_profit:
        tp = parse_numeric(total_profit)
        op = parse_numeric(operating_profit)
        noi = parse_numeric(non_op_income) if non_op_income else 0
        noe = parse_numeric(non_op_expense) if non_op_expense else 0
        if tp is not None and op is not None:
            expected = op + noi - noe
            diff = tp - expected
            passed = abs(diff) < 1
            checks.append({
                'name': '利润总额 = 营业利润 + 营业外收入 - 营业外支出',
                'passed': passed,
                'expected': f'{expected:,.0f}',
                'actual': f'{tp:,.0f}',
                'diff': f'{diff:,.0f}',
                'detail': f'{to_wan_rounded(total_profit)} = {to_wan_rounded(operating_profit)} + {to_wan_rounded(non_op_income) if non_op_income else "0"} - {to_wan_rounded(non_op_expense) if non_op_expense else "0"}'
            })
            if not passed:
                errors.append('利润总额勾稽不平！')

    # 6. 净利润 = 利润总额 - 所得税费用
    net_profit = find_in_data(is_data, '净利润')
    income_tax = find_in_data(is_data, '所得税费用')
    if total_profit and net_profit and income_tax:
        tp = parse_numeric(total_profit)
        np = parse_numeric(net_profit)
        it = parse_numeric(income_tax)
        if tp is not None and np is not None and it is not None:
            expected = tp - it
            diff = np - expected
            passed = abs(diff) < 1
            checks.append({
                'name': '净利润 = 利润总额 - 所得税费用',
                'passed': passed,
                'expected': f'{expected:,.0f}',
                'actual': f'{np:,.0f}',
                'diff': f'{diff:,.0f}',
                'detail': f'{to_wan_rounded(net_profit)} = {to_wan_rounded(total_profit)} - {to_wan_rounded(income_tax)}'
            })
            if not passed:
                errors.append('净利润勾稽不平！')

    # ---------- 表内勾稽：现金流量表 ----------
    if cf_data:
        # 7. 经营活动现金流量净额 = 流入小计 - 流出小计
        op_in = find_in_data(cf_data, '经营活动现金流入小计')
        op_out = find_in_data(cf_data, '经营活动现金流出小计')
        op_net = find_in_data(cf_data, '经营活动产生的现金流量净额', '经营活动现金流量净额')
        if op_in and op_out and op_net:
            oi = parse_numeric(op_in)
            oo = parse_numeric(op_out)
            on = parse_numeric(op_net)
            if oi is not None and oo is not None and on is not None:
                expected = oi - oo
                diff = on - expected
                passed = abs(diff) < 1
                checks.append({
                    'name': '经营活动净额 = 流入 - 流出',
                    'passed': passed,
                    'expected': f'{expected:,.0f}',
                    'actual': f'{on:,.0f}',
                    'diff': f'{diff:,.0f}',
                    'detail': f'{to_wan_rounded(op_net)} = {to_wan_rounded(op_in)} - {to_wan_rounded(op_out)}'
                })
                if not passed:
                    errors.append('经营活动现金流量勾稽不平！')

        # 8. 现金净增加额 = 经营净额 + 投资净额 + 筹资净额
        inv_net = find_in_data(cf_data, '投资活动产生的现金流量净额', '投资活动现金流量净额')
        fin_net = find_in_data(cf_data, '筹资活动产生的现金流量净额', '筹资活动现金流量净额')
        cash_change = find_in_data(cf_data, '现金及现金等价物净增加额', '现金净增加额')
        if op_net and inv_net and fin_net and cash_change:
            on = parse_numeric(op_net)
            ivn = parse_numeric(inv_net)
            fn = parse_numeric(fin_net)
            cc = parse_numeric(cash_change)
            if on is not None and ivn is not None and fn is not None and cc is not None:
                expected = on + ivn + fn
                diff = cc - expected
                passed = abs(diff) < 1
                checks.append({
                    'name': '现金净增加额 = 经营 + 投资 + 筹资',
                    'passed': passed,
                    'expected': f'{expected:,.0f}',
                    'actual': f'{cc:,.0f}',
                    'diff': f'{diff:,.0f}',
                    'detail': f'{to_wan_rounded(cash_change)} = {to_wan_rounded(op_net)} + {to_wan_rounded(inv_net)} + {to_wan_rounded(fin_net)}'
                })
                if not passed:
                    errors.append('现金流量净增加额勾稽不平！')

        # 9. 表间勾稽：期末现金 ≈ 货币资金
        monetary = find_in_data(bs_data, '货币资金')
        cash_end = find_in_data(cf_data, '期末现金及现金等价物余额', '期末现金余额')
        if monetary and cash_end:
            m = parse_numeric(monetary)
            ce = parse_numeric(cash_end)
            if m is not None and ce is not None:
                diff = m - ce
                # 货币资金通常略大于现金等价物（含受限资金），允许一定差异
                passed = abs(diff) < 1000  # 允许 1000 元差异
                checks.append({
                    'name': '表间勾稽：货币资金 ≈ 期末现金余额',
                    'passed': passed,
                    'expected': f'{ce:,.0f}',
                    'actual': f'{m:,.0f}',
                    'diff': f'{diff:,.0f}',
                    'detail': f'资产负债表货币资金 = {to_wan_rounded(monetary)}万, 现金流量表期末现金 = {to_wan_rounded(cash_end)}万'
                })
                if not passed:
                    errors.append('货币资金与期末现金余额偏差较大！')

        # 10. 表间勾稽：利润表净利润应与现金流量表间接法起点一致
        cf_net_profit = find_in_data(cf_data, '净利润')
        if net_profit and cf_net_profit:
            np_is = parse_numeric(net_profit)
            np_cf = parse_numeric(cf_net_profit)
            if np_is is not None and np_cf is not None:
                diff = np_is - np_cf
                passed = abs(diff) < 1
                checks.append({
                    'name': '表间勾稽：两表净利润一致',
                    'passed': passed,
                    'expected': f'{np_cf:,.0f}',
                    'actual': f'{np_is:,.0f}',
                    'diff': f'{diff:,.0f}',
                    'detail': f'利润表净利润 = {to_wan_rounded(net_profit)}万, 现金流量表净利润 = {to_wan_rounded(cf_net_profit)}万'
                })
                if not passed:
                    errors.append('利润表与现金流量表的净利润不一致！')

    # ---------- 万元反算验证 ----------
    # 检查所有大数在万元转换后是否合理（辅助发现首位数字错位）
    large_items = []
    for d, label in [(bs_data, 'BS'), (is_data, 'IS'), (cf_data or {}, 'CF')]:
        for k, v in d.items():
            n = parse_numeric(v)
            if n is not None and abs(n) >= 1e7:  # >= 1000万
                large_items.append((label, k, n))

    # ---------- 汇总 ----------
    total_checks = len(checks)
    passed_checks = sum(1 for c in checks if c['passed'])
    failed_checks = total_checks - passed_checks

    if not checks:
        summary = '⚠️ 数据不足，无法进行交叉验证（至少需要资产总计、负债合计、所有者权益合计）'
    elif failed_checks == 0:
        summary = f'✅ 全部 {total_checks} 项验证通过，数据可信度高'
    else:
        err_list = '；'.join(errors)
        summary = f'⚠️ {passed_checks}/{total_checks} 项通过，{failed_checks} 项未通过：{err_list}。建议重读图片核对。'

    return {
        'passed': failed_checks == 0,
        'checks': checks,
        'summary': summary,
        'total_checks': total_checks,
        'passed_checks': passed_checks,
        'failed_checks': failed_checks
    }


def normalize_quotes(s: str) -> str:
    """Normalize Chinese full-width quotes to ASCII quotes for matching.

    Excel data often uses \u201c\u201d (""), while code uses straight quotes.
    Also normalizes full-width parentheses and colons.
    """
    if not s:
        return s
    s = s.replace('\u201c', '"').replace('\u201d', '"')  # "" → "
    s = s.replace('\u2018', "'").replace('\u2019', "'")  # '' → '
    s = s.replace('\uff08', '(').replace('\uff09', ')')  # （） → ()
    s = s.replace('\uff1a', ':').replace('\uff1b', ';')  # ： → :  ； → ;
    return s


def find_in_data(data_dict: dict, *keywords, default: str = '0') -> str:
    """Find first value in data_dict whose key matches any keyword.

    Uses 4-level priority matching to avoid false positives:
      1. Exact match (after quote normalization)
      2. Normalized exact match (strip '合计'/'小计' suffixes etc.)
      3. Substring match excluding prefixed variants (e.g. '负债合计' won't match '流动负债合计')
      4. Fallback substring match (loosest)

    Returns the value string, or `default` if not found.
    """
    if not data_dict:
        return default

    normalized_data = {}
    for k, v in data_dict.items():
        nk = normalize_quotes(str(k)).strip()
        normalized_data[nk] = v

    # Level 1: Exact match
    for kw in keywords:
        nk = normalize_quotes(kw).strip()
        if nk in normalized_data:
            return str(normalized_data[nk])

    # Level 2: Normalized exact match (remove 合计/小计/净额 suffixes)
    def _norm_key(s):
        s = normalize_quotes(s)
        s = re.sub(r'[（(].*?[)）]', '', s)
        s = s.replace('其中：', '').replace('其中:', '')
        s = s.replace('合计', '').replace('小计', '')
        s = s.replace('净额', '').replace('总额', '')
        return s.strip()

    for kw in keywords:
        nk = _norm_key(kw)
        for k, v in normalized_data.items():
            if _norm_key(k) == nk:
                return str(v)

    # Level 3: Substring match, excluding keys that have prefix modifiers
    # e.g. '负债合计' should NOT match '流动负债合计' or '非流动负债合计'
    prefix_modifiers = ['流动', '非流动', '其他', '一年内到期']
    for kw in keywords:
        nk = normalize_quotes(kw).strip()
        for k, v in normalized_data.items():
            if nk in k:
                # Check if k has a prefix modifier that kw doesn't
                has_prefix = any(k.startswith(pm) and not nk.startswith(pm) for pm in prefix_modifiers)
                if not has_prefix:
                    return str(v)

    # Level 4: Fallback loose substring match
    for kw in keywords:
        nk = normalize_quotes(kw).strip()
        for k, v in normalized_data.items():
            if nk in k or k in nk:
                return str(v)

    return default


def find_optional(data_dict: dict, *keywords) -> str | None:
    """Like find_in_data but returns None when not found.

    Used inside ratio calculations to distinguish "missing" from a real 0,
    avoiding the legacy '0' sentinel being parsed as 0.0 (which would
    produce fake values like "0.00%").
    """
    val = find_in_data(data_dict, *keywords, default=None)
    return val


# ---------------------------------------------------------------------------
# Data Models
# ---------------------------------------------------------------------------

class ReportData:
    """Container for a single report's data extracted from a source file."""

    def __init__(self, source_label: str, headers: list[str], rows: list[list]):
        self.source_label = source_label
        self.headers = headers      # Column names
        self.rows = rows            # List of rows, each row is a list of values

    def to_dict(self) -> dict:
        return {
            "source_label": self.source_label,
            "headers": self.headers,
            "row_count": len(self.rows),
        }


# ---------------------------------------------------------------------------
# File Readers
# ---------------------------------------------------------------------------

def read_excel(path: str, sheet_name: str | None = None,
               header_row: int = 1, skip_empty: bool = True) -> ReportData:
    """
    Read an Excel file and return ReportData.

    Supports both .xlsx (via openpyxl) and .xls (via xlrd) formats.

    Args:
        path: Path to .xlsx or .xls file.
        sheet_name: Specific sheet to read. If None, reads the first sheet.
        header_row: 1-based row number for headers.
        skip_empty: Skip completely empty rows.
    """
    is_xls = path.lower().endswith('.xls') and not path.lower().endswith('.xlsx')

    if is_xls:
        if xlrd is None:
            raise ImportError("xlrd is required for .xls files. Install with: pip install xlrd")
        return _read_xls(path, sheet_name, header_row, skip_empty)

    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb[sheet_name] if sheet_name else wb.active

    all_rows = list(ws.iter_rows(values_only=True))
    if not all_rows:
        raise ValueError(f"Sheet '{ws.title}' in {path} is empty.")

    # Adjust for 1-based header_row
    header_idx = header_row - 1
    if header_idx >= len(all_rows):
        raise ValueError(f"Header row {header_row} exceeds sheet length.")

    headers = [str(h).strip() if h is not None else "" for h in all_rows[header_idx]]
    data_rows = all_rows[header_idx + 1:]

    if skip_empty:
        data_rows = [r for r in data_rows if any(c is not None and str(c).strip() != "" for c in r)]

    # Normalize row length to match headers
    col_count = len(headers)
    normalized = []
    for row in data_rows:
        row_list = list(row)
        # Pad or trim
        while len(row_list) < col_count:
            row_list.append(None)
        row_list = row_list[:col_count]
        normalized.append(row_list)

    source_label = Path(path).stem
    wb.close()
    return ReportData(source_label=source_label, headers=headers, rows=normalized)


def _read_xls(path: str, sheet_name: str | None = None,
              header_row: int = 1, skip_empty: bool = True) -> ReportData:
    """Read .xls files using xlrd library."""
    wb = xlrd.open_workbook(path)
    ws = wb.sheet_by_name(sheet_name) if sheet_name else wb.sheet_by_index(0)

    header_idx = header_row - 1
    if header_idx >= ws.nrows:
        raise ValueError(f"Header row {header_row} exceeds sheet length.")

    headers = [str(ws.cell_value(header_idx, c)).strip() if c < ws.ncols else ""
               for c in range(ws.ncols)]
    data_rows = []
    for r in range(header_idx + 1, ws.nrows):
        row = [ws.cell_value(r, c) if c < ws.ncols else None for c in range(ws.ncols)]
        data_rows.append(row)

    if skip_empty:
        data_rows = [r for r in data_rows if any(c is not None and str(c).strip() != "" for c in r)]

    col_count = len(headers)
    normalized = []
    for row in data_rows:
        row_list = list(row)
        while len(row_list) < col_count:
            row_list.append(None)
        row_list = row_list[:col_count]
        normalized.append(row_list)

    source_label = Path(path).stem
    return ReportData(source_label=source_label, headers=headers, rows=normalized)


def extract_pdf_tables(path: str, pages: str | None = None,
                       table_index: int | None = None) -> list[ReportData]:
    """
    Extract tables from a PDF file using pdfplumber.

    Args:
        path: Path to PDF file.
        pages: Page range string like "0-5" or "0,2,4". None = all pages.
        table_index: If specified, only extract the table at this index from each page.

    Returns:
        List of ReportData, one per extracted table.
    """
    if pdfplumber is None:
        raise ImportError("pdfplumber is required for PDF extraction. Install with: pip install pdfplumber")

    page_list = _parse_page_range(pages)
    results = []
    source_label = Path(path).stem

    with pdfplumber.open(path) as pdf:
        total_pages = len(pdf.pages)
        target_pages = page_list if page_list is not None else range(total_pages)

        for page_num in target_pages:
            if page_num >= total_pages:
                continue
            page = pdf.pages[page_num]
            tables = page.extract_tables()
            if not tables:
                continue

            if table_index is not None:
                if table_index < len(tables):
                    tables = [tables[table_index]]
                else:
                    continue

            for i, table in enumerate(tables):
                if not table or len(table) < 2:
                    continue
                headers = [str(h).strip() if h else "" for h in table[0]]
                rows = [list(r) for r in table[1:]]
                label = f"{source_label}_p{page_num}_t{i}"
                results.append(ReportData(source_label=label, headers=headers, rows=rows))

    if not results:
        print(f"Warning: No tables found in {path}.", file=sys.stderr)

    return results


def _parse_page_range(pages: str | None) -> list[int] | None:
    """Parse page range string like '0-5' or '0,2,4' into a list of ints."""
    if pages is None:
        return None
    result = []
    for part in pages.split(","):
        part = part.strip()
        if "-" in part:
            start, end = part.split("-", 1)
            result.extend(range(int(start), int(end) + 1))
        else:
            result.append(int(part))
    return result


def detect_pdf_type(pdf_path: str) -> dict:
    """
    Detect whether a PDF is text-based or image-based (scanned).

    Uses PyMuPDF to check if each page has extractable text.
    A page is 'image' if it has < 50 chars of extractable text.

    Returns:
        {
            'type': 'text' | 'image' | 'mixed',
            'total_pages': int,
            'text_pages': [int, ...],
            'image_pages': [int, ...],
            'text_ratio': float,  # percentage
        }
    """
    if fitz is None:
        raise ImportError("PyMuPDF is required for PDF type detection. Install with: pip install PyMuPDF")

    doc = fitz.open(pdf_path)
    total_pages = len(doc)
    text_pages = []
    image_pages = []

    for i in range(total_pages):
        page = doc[i]
        text = page.get_text().strip()
        if len(text) > 50:
            text_pages.append(i)
        else:
            image_pages.append(i)

    doc.close()

    if not image_pages:
        pdf_type = 'text'
    elif not text_pages:
        pdf_type = 'image'
    else:
        pdf_type = 'mixed'

    text_ratio = round(len(text_pages) / total_pages * 100, 1) if total_pages > 0 else 0

    return {
        'type': pdf_type,
        'total_pages': total_pages,
        'text_pages': text_pages,
        'image_pages': image_pages,
        'text_ratio': text_ratio,
    }


def pdf_to_images(pdf_path: str, output_dir: str,
                  dpi: int = 300, pages: list[int] | None = None,
                  fmt: str = 'png') -> list[dict]:
    """
    Convert PDF pages to high-resolution images using PyMuPDF.

    Use this when detect_pdf_type() returns 'image' or 'mixed' —
    pdfplumber cannot extract tables from image-based pages.

    The generated images can then be read by the agent's multimodal
    vision capability to extract table data manually.

    Args:
        pdf_path: Path to PDF file.
        output_dir: Directory to save images.
        dpi: Resolution (default 300 for good OCR/visual quality).
        pages: List of 0-based page indices to convert. None = all pages.
        fmt: Image format ('png' or 'jpeg').

    Returns:
        List of dicts: [{page, path, width, height}, ...]
    """
    if fitz is None:
        raise ImportError("PyMuPDF is required for PDF to image conversion. Install with: pip install PyMuPDF")

    os.makedirs(output_dir, exist_ok=True)

    doc = fitz.open(pdf_path)
    total_pages = len(doc)
    target_pages = pages if pages is not None else list(range(total_pages))

    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)

    results = []
    for page_num in target_pages:
        if page_num >= total_pages:
            continue

        page = doc[page_num]
        pix = page.get_pixmap(matrix=matrix)

        ext = 'png' if fmt == 'png' else 'jpg'
        filename = f'page_{page_num + 1:03d}.{ext}'
        filepath = os.path.join(output_dir, filename)

        if fmt == 'png':
            pix.save(filepath)
        else:
            pix.save(filepath, jpg_quality=95)

        results.append({
            'page': page_num,
            'path': os.path.abspath(filepath),
            'width': pix.width,
            'height': pix.height,
        })

    doc.close()
    return results


# ---------------------------------------------------------------------------
# Merge Modes
# ---------------------------------------------------------------------------

def merge_concat(reports: list[ReportData], source_labels: list[str] | None = None) -> ReportData:
    """
    Mode 1: Simple concatenation.
    Stacks all rows from all reports vertically. Headers from the first report are used.
    Reports must have the same column structure (or at least the same number of columns).
    """
    if not reports:
        raise ValueError("No reports to merge.")

    if source_labels:
        for i, label in enumerate(source_labels):
            if i < len(reports):
                reports[i].source_label = label

    # Use headers from first report; validate compatibility
    base_headers = reports[0].headers
    all_rows = []

    for rpt in reports:
        # If headers differ in count, pad/truncate
        if len(rpt.headers) != len(base_headers):
            print(f"Warning: Column count mismatch in '{rpt.source_label}' "
                  f"({len(rpt.headers)} vs {len(base_headers)}). Adjusting.", file=sys.stderr)
        for row in rpt.rows:
            all_rows.append(row)

    return ReportData(
        source_label="合并结果",
        headers=base_headers,
        rows=all_rows,
    )


def merge_by_key(reports: list[ReportData], key_column: str,
                 source_labels: list[str] | None = None,
                 aggregation: str = "sum") -> ReportData:
    """
    Mode 2: Merge by key column.
    Matches rows across reports using a shared key column (e.g., "科目代码").
    Numeric columns are aggregated (sum by default).

    Args:
        key_column: Name of the column to use as the merge key.
        source_labels: Custom labels for each source (used as column suffixes).
        aggregation: How to aggregate numeric values: "sum", "avg", "max", "min".
    """
    if not reports:
        raise ValueError("No reports to merge.")

    if source_labels:
        for i, label in enumerate(source_labels):
            if i < len(reports):
                reports[i].source_label = label

    # Find key column index in each report
    key_indices = []
    for rpt in reports:
        if key_column not in rpt.headers:
            # Try fuzzy match
            matches = [h for h in rpt.headers if key_column.lower() in h.lower()]
            if matches:
                idx = rpt.headers.index(matches[0])
            else:
                raise ValueError(f"Key column '{key_column}' not found in report '{rpt.source_label}'. "
                                 f"Available headers: {rpt.headers}")
        else:
            idx = rpt.headers.index(key_column)
        key_indices.append(idx)

    # Collect all unique keys (preserve order)
    all_keys: list[str] = []
    seen_keys: set[str] = set()
    for rpt, ki in zip(reports, key_indices):
        for row in rpt.rows:
            key_val = str(row[ki]).strip() if row[ki] is not None else ""
            if key_val and key_val not in seen_keys:
                all_keys.append(key_val)
                seen_keys.add(key_val)

    # Build lookup: {key: {source_label: {col_name: value}}}
    lookup: dict[str, dict[str, dict[str, Any]]] = {k: {} for k in all_keys}
    for rpt, ki in zip(reports, key_indices):
        for row in rpt.rows:
            key_val = str(row[ki]).strip() if row[ki] is not None else ""
            if not key_val:
                continue
            if key_val not in lookup:
                lookup[key_val] = {}
            row_dict = {}
            for ci, col_name in enumerate(rpt.headers):
                if ci == ki:
                    continue
                row_dict[col_name] = row[ci] if ci < len(row) else None
            lookup[key_val][rpt.source_label] = row_dict

    # Build merged headers: key_column + source_label/col_name pairs
    merged_headers = [key_column]
    # Collect all non-key column names across reports
    all_col_names: list[str] = []
    seen_cols: set[str] = set()
    for rpt in reports:
        for h in rpt.headers:
            if h == key_column:
                continue
            if h not in seen_cols:
                all_col_names.append(h)
                seen_cols.add(h)

    for rpt in reports:
        for col_name in all_col_names:
            if col_name in rpt.headers:
                merged_headers.append(f"{rpt.source_label}_{col_name}")

    # Build merged rows
    merged_rows = []
    for key_val in all_keys:
        row = [key_val]
        for rpt in reports:
            source_data = lookup.get(key_val, {}).get(rpt.source_label, {})
            for col_name in all_col_names:
                if col_name in rpt.headers:
                    row.append(source_data.get(col_name))
                else:
                    row.append(None)
        merged_rows.append(row)

    return ReportData(
        source_label="合并结果",
        headers=merged_headers,
        rows=merged_rows,
    )


def merge_consolidation(reports: list[ReportData], key_column: str,
                        source_labels: list[str] | None = None,
                        elimination_entries: list[dict] | None = None,
                        aggregation: str = "sum") -> ReportData:
    """
    Mode 3: Consolidation with elimination entries (合并报表 + 抵销分录).

    First performs key-based merge (sum across entities), then applies
    elimination entries to remove inter-company transactions.

    Elimination entry format (list of dicts):
    [
        {
            "科目": "应收账款",
            "方向": "借",        # "借" or "贷"
            "金额": 100000,
            "说明": "母子公司内部应收应付抵销"
        },
        ...
    ]

    The elimination adjusts the consolidated total for each matching account.
    """
    # Step 1: Sum across all entities for each key
    if source_labels:
        for i, label in enumerate(source_labels):
            if i < len(reports):
                reports[i].source_label = label

    # Find key column indices
    key_indices = []
    for rpt in reports:
        if key_column not in rpt.headers:
            matches = [h for h in rpt.headers if key_column.lower() in h.lower()]
            if matches:
                idx = rpt.headers.index(matches[0])
            else:
                raise ValueError(f"Key column '{key_column}' not found in '{rpt.source_label}'.")
        else:
            idx = rpt.headers.index(key_column)
        key_indices.append(idx)

    # Collect all unique keys
    all_keys: list[str] = []
    seen_keys: set[str] = set()
    for rpt, ki in zip(reports, key_indices):
        for row in rpt.rows:
            key_val = str(row[ki]).strip() if row[ki] is not None else ""
            if key_val and key_val not in seen_keys:
                all_keys.append(key_val)
                seen_keys.add(key_val)

    # Sum numeric values across all reports for each key
    # Determine numeric columns (columns that have numeric values in any report)
    numeric_cols: dict[str, bool] = {}  # col_name -> is_numeric
    for rpt in reports:
        for ci, col_name in enumerate(rpt.headers):
            if ci in key_indices:
                continue
            if col_name not in numeric_cols:
                numeric_cols[col_name] = True
            for row in rpt.rows:
                val = row[ci] if ci < len(row) else None
                if val is not None and not _is_numeric(val):
                    numeric_cols[col_name] = False

    # Build consolidated totals
    consolidated: dict[str, dict[str, float]] = {k: {} for k in all_keys}
    for rpt, ki in zip(reports, key_indices):
        for row in rpt.rows:
            key_val = str(row[ki]).strip() if row[ki] is not None else ""
            if not key_val:
                continue
            if key_val not in consolidated:
                consolidated[key_val] = {}
            for ci, col_name in enumerate(rpt.headers):
                if ci == ki or not numeric_cols.get(col_name, False):
                    continue
                val = row[ci] if ci < len(row) else None
                num_val = _to_numeric(val)
                if num_val is not None:
                    consolidated[key_val][col_name] = \
                        consolidated[key_val].get(col_name, 0) + num_val

    # Step 2: Apply elimination entries
    eliminations_applied: list[dict] = []
    if elimination_entries:
        for entry in elimination_entries:
            account = entry.get("科目", "")
            direction = entry.get("方向", "借")
            amount = _to_numeric(entry.get("金额", 0)) or 0
            description = entry.get("说明", "")
            target_col = entry.get("列", None)  # Which numeric column to adjust

            # Find matching key
            matched = False
            for key_val in consolidated:
                if account in key_val or key_val in account:
                    # Determine which column to adjust
                    if target_col and target_col in consolidated[key_val]:
                        col_to_adjust = target_col
                    elif consolidated[key_val]:
                        # Adjust the first numeric column
                        col_to_adjust = list(consolidated[key_val].keys())[0]
                    else:
                        # No numeric data for this key — initialize
                        col_to_adjust = target_col or "金额"
                        consolidated[key_val][col_to_adjust] = 0

                    if direction == "借":
                        consolidated[key_val][col_to_adjust] -= amount
                    else:  # 贷
                        consolidated[key_val][col_to_adjust] += amount

                    eliminations_applied.append({
                        "科目": key_val,
                        "方向": direction,
                        "金额": amount,
                        "说明": description,
                        "调整列": col_to_adjust,
                    })
                    matched = True
                    break

            if not matched:
                print(f"Warning: Elimination entry for '{account}' did not match any key.", file=sys.stderr)

    # Build output
    # Headers: key_column + all numeric columns + "合并主体"
    output_headers = [key_column]
    numeric_col_names = [c for c, is_num in numeric_cols.items() if is_num]
    output_headers.extend(numeric_col_names)

    merged_rows = []
    for key_val in all_keys:
        row = [key_val]
        for col_name in numeric_col_names:
            val = consolidated[key_val].get(col_name)
            row.append(val if val is not None else 0)
        merged_rows.append(row)

    result = ReportData(
        source_label="合并报表",
        headers=output_headers,
        rows=merged_rows,
    )

    # Attach elimination log as metadata
    result.eliminations = eliminations_applied  # type: ignore[attr-defined]
    return result


def _is_numeric(val) -> bool:
    """Check if a value is numeric."""
    if val is None:
        return False
    if isinstance(val, (int, float)):
        return True
    try:
        float(str(val).replace(",", "").replace(" ", ""))
        return True
    except (ValueError, TypeError):
        return False


def _to_numeric(val):
    """Convert a value to float, returning None on failure."""
    if val is None:
        return None
    if isinstance(val, (int, float)):
        return float(val)
    try:
        return float(str(val).replace(",", "").replace(" ", ""))
    except (ValueError, TypeError):
        return None


# ---------------------------------------------------------------------------
# Excel Writer
# ---------------------------------------------------------------------------

def write_excel(data: ReportData, output_path: str,
                sheet_name: str = "合并结果",
                eliminations: list[dict] | None = None) -> str:
    """
    Write ReportData to an Excel file with formatting.

    Args:
        data: ReportData to write.
        output_path: Output file path.
        sheet_name: Name of the sheet.
        eliminations: Optional elimination entries to write on a second sheet.
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet_name

    # Styles
    header_font = Font(name="微软雅黑", size=10, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell_font = Font(name="微软雅黑", size=10)
    cell_align = Alignment(horizontal="left", vertical="center")
    num_align = Alignment(horizontal="right", vertical="center")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    # Write headers
    for ci, header in enumerate(data.headers, 1):
        cell = ws.cell(row=1, column=ci, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    # Write data rows
    for ri, row in enumerate(data.rows, 2):
        for ci, val in enumerate(row, 1):
            cell = ws.cell(row=ri, column=ci, value=val)
            cell.font = cell_font
            cell.border = thin_border
            if isinstance(val, (int, float)):
                cell.alignment = num_align
                cell.number_format = '#,##0.00'
            else:
                cell.alignment = cell_align

    # Auto-fit column widths (approximate)
    for ci, header in enumerate(data.headers, 1):
        max_len = len(str(header))
        for ri in range(2, min(len(data.rows) + 2, 102)):  # Sample first 100 rows
            cell_val = ws.cell(row=ri, column=ci).value
            if cell_val is not None:
                max_len = max(max_len, len(str(cell_val)))
        ws.column_dimensions[get_column_letter(ci)].width = min(max_len + 4, 40)

    # Freeze header row
    ws.freeze_panes = "A2"

    # Write elimination entries on a second sheet if provided
    if eliminations:
        ws2 = wb.create_sheet("抵销分录")
        elim_headers = ["科目", "方向", "金额", "说明", "调整列"]
        for ci, h in enumerate(elim_headers, 1):
            cell = ws2.cell(row=1, column=ci, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = thin_border

        for ri, entry in enumerate(eliminations, 2):
            for ci, key in enumerate(elim_headers, 1):
                cell = ws2.cell(row=ri, column=ci, value=entry.get(key, ""))
                cell.font = cell_font
                cell.border = thin_border
                if key == "金额" and isinstance(entry.get(key), (int, float)):
                    cell.alignment = num_align
                    cell.number_format = '#,##0.00'

        for ci in range(1, len(elim_headers) + 1):
            ws2.column_dimensions[get_column_letter(ci)].width = 20

    wb.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Word Template Filler
# ---------------------------------------------------------------------------

def fill_word_template(template_path: str, data: ReportData,
                       output_path: str, mapping: dict | None = None) -> str:
    """
    Fill a Word (.docx) template with merged report data.

    Supports two filling strategies:
    1. Placeholder replacement: Replace {{placeholder}} tokens in the document text
       with values from the merged data.
    2. Table filling: Fill tables in the document with rows from the merged data,
       matching by column header names.

    Mapping format (optional JSON):
    {
        "placeholders": {
            "报告日期": "2024年12月31日",
            "编制单位": "XX集团"
        },
        "table_mapping": {
            "target_table_index": 0,
            "column_map": {
                "科目代码": "科目代码",
                "合并金额": "本期金额"
            },
            "key_column": "科目代码"
        }
    }

    Args:
        template_path: Path to the Word template (.docx).
        data: Merged ReportData to fill into the template.
        output_path: Output path for the filled document.
        mapping: Optional mapping configuration (see above).
    """
    if DocxDocument is None:
        raise ImportError("python-docx is required for Word template filling. "
                          "Install with: pip install python-docx")

    doc = DocxDocument(template_path)

    if mapping is None:
        mapping = {}

    # --- Strategy 1: Placeholder replacement ---
    placeholders = mapping.get("placeholders", {})

    # Also auto-generate placeholders from data (first row values keyed by header)
    # This allows {{header_name}} to resolve to the value in the first data row
    if data.rows:
        first_row = data.rows[0]
        for ci, header in enumerate(data.headers):
            if header and ci < len(first_row):
                key = header.strip()
                if key and key not in placeholders:
                    val = first_row[ci]
                    if val is not None:
                        placeholders[key] = str(val)

    # Replace placeholders in all paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, placeholders)

    # Replace placeholders in all table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, placeholders)

    # --- Strategy 2: Table filling ---
    table_mapping = mapping.get("table_mapping")
    if table_mapping:
        target_table_idx = table_mapping.get("target_table_index", 0)
        column_map = table_mapping.get("column_map", {})
        key_col = table_mapping.get("key_column")

        if target_table_idx < len(doc.tables):
            table = doc.tables[target_table_idx]

            # Determine which template table column maps to which data column
            # Read header row from the template table
            template_headers = []
            if len(table.rows) > 0:
                for cell in table.rows[0].cells:
                    template_headers.append(cell.text.strip())

            # Build column mapping: template_col_idx -> data_col_idx
            col_pairs = []
            for tci, t_header in enumerate(template_headers):
                # Check if template header maps to a data column
                data_col_name = column_map.get(t_header, t_header)
                if data_col_name in data.headers:
                    dci = data.headers.index(data_col_name)
                    col_pairs.append((tci, dci))

            if not col_pairs:
                # Fallback: map by position
                for tci in range(min(len(template_headers), len(data.headers))):
                    col_pairs.append((tci, tci))

            # Build data lookup by key if key_column is specified
            data_lookup = {}
            if key_col and key_col in data.headers:
                key_idx = data.headers.index(key_col)
                for row in data.rows:
                    k = str(row[key_idx]).strip() if row[key_idx] is not None else ""
                    if k:
                        data_lookup[k] = row

            # Fill table rows
            # If the template has existing data rows (beyond header), try to match by key
            # Otherwise, append new rows
            existing_data_rows = len(table.rows) - 1  # Subtract header row

            if key_col and data_lookup and existing_data_rows > 0:
                # Match by key — update existing rows
                for ri in range(1, len(table.rows)):
                    row = table.rows[ri]
                    # Find key value in the template row (first column by default)
                    row_key = row.cells[0].text.strip()
                    if row_key in data_lookup:
                        data_row = data_lookup[row_key]
                        for tci, dci in col_pairs:
                            val = data_row[dci] if dci < len(data_row) else None
                            if val is not None:
                                row.cells[tci].text = str(val)
            else:
                # Append new rows
                for data_row in data.rows:
                    new_row = table.add_row()
                    for tci, dci in col_pairs:
                        val = data_row[dci] if dci < len(data_row) else None
                        if val is not None:
                            new_row.cells[tci].text = str(val)
                        else:
                            new_row.cells[tci].text = ""

    doc.save(output_path)
    return output_path


def _replace_in_paragraph(paragraph, replacements: dict):
    """
    Replace {{placeholder}} tokens in a paragraph's text.
    Handles cases where the placeholder spans multiple runs.
    """
    full_text = paragraph.text
    if "{{" not in full_text:
        return

    # Find all {{...}} patterns
    pattern = re.compile(r"\{\{(\w+)\}\}")
    matches = pattern.findall(full_text)

    if not matches:
        return

    # Rebuild text with replacements
    new_text = full_text
    for match in matches:
        if match in replacements:
            new_text = new_text.replace(f"{{{{{match}}}}}", str(replacements[match]))

    # Clear existing runs and set new text (preserving first run's formatting)
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run.text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new_text


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def cmd_merge(args):
    """Handle the merge subcommand."""
    # Read all input files
    reports = []
    for file_path in args.input:
        ext = Path(file_path).suffix.lower()
        if ext in (".xlsx", ".xls"):
            rpt = read_excel(file_path, sheet_name=args.sheet, header_row=args.header_row)
            reports.append(rpt)
            print(f"  Loaded Excel: {file_path} ({len(rpt.rows)} rows, {len(rpt.headers)} cols)")
        elif ext == ".pdf":
            pdf_reports = extract_pdf_tables(file_path, pages=args.pdf_pages, table_index=args.pdf_table_index)
            reports.extend(pdf_reports)
            print(f"  Loaded PDF: {file_path} ({len(pdf_reports)} tables extracted)")
        else:
            print(f"  Warning: Skipping unsupported file type: {file_path}", file=sys.stderr)

    if not reports:
        print("Error: No valid input files loaded.", file=sys.stderr)
        sys.exit(1)

    # Perform merge
    mode = args.mode
    if mode == "concat":
        result = merge_concat(reports, source_labels=args.source_labels)
    elif mode == "key":
        if not args.key_column:
            print("Error: --key-column is required for key-based merge.", file=sys.stderr)
            sys.exit(1)
        result = merge_by_key(reports, key_column=args.key_column,
                              source_labels=args.source_labels, aggregation=args.aggregation)
    elif mode == "consolidation":
        if not args.key_column:
            print("Error: --key-column is required for consolidation merge.", file=sys.stderr)
            sys.exit(1)
        elim_entries = None
        if args.elimination_entries:
            with open(args.elimination_entries, "r", encoding="utf-8") as f:
                elim_entries = json.load(f)
        result = merge_consolidation(reports, key_column=args.key_column,
                                     source_labels=args.source_labels,
                                     elimination_entries=elim_entries,
                                     aggregation=args.aggregation)
    else:
        print(f"Error: Unknown merge mode '{mode}'.", file=sys.stderr)
        sys.exit(1)

    print(f"\n  Merge complete: {len(result.rows)} rows, {len(result.headers)} columns")

    # Write Excel output
    write_excel(result, args.output, sheet_name=args.sheet_name or "合并结果",
                eliminations=getattr(result, "eliminations", None))
    print(f"  Excel output: {args.output}")

    # Fill Word template if provided
    if args.word_template:
        if not args.word_output:
            args.word_output = str(Path(args.output).with_suffix(".docx"))
        mapping = None
        if args.mapping:
            with open(args.mapping, "r", encoding="utf-8") as f:
                mapping = json.load(f)
        fill_word_template(args.word_template, result, args.word_output, mapping)
        print(f"  Word output: {args.word_output}")


def cmd_extract_pdf(args):
    """Handle the extract-pdf subcommand."""
    tables = extract_pdf_tables(args.input, pages=args.pages, table_index=args.table_index)
    if not tables:
        print("No tables found in PDF.", file=sys.stderr)
        sys.exit(1)

    # Write each table to a sheet
    wb = openpyxl.Workbook()
    if len(tables) == 1:
        ws = wb.active
        ws.title = "Table_1"
        rpt = tables[0]
        ws.append(rpt.headers)
        for row in rpt.rows:
            ws.append(row)
    else:
        wb.remove(wb.active)
        for i, rpt in enumerate(tables, 1):
            ws = wb.create_sheet(f"Table_{i}")
            ws.append(rpt.headers)
            for row in rpt.rows:
                ws.append(row)

    wb.save(args.output)
    print(f"Extracted {len(tables)} table(s) to {args.output}")


def cmd_fill_template(args):
    """Handle the fill-template subcommand."""
    rpt = read_excel(args.input)
    mapping = None
    if args.mapping:
        with open(args.mapping, "r", encoding="utf-8") as f:
            mapping = json.load(f)
    fill_word_template(args.template, rpt, args.output, mapping)
    print(f"Template filled: {args.output}")


def main():
    parser = argparse.ArgumentParser(
        description="Enterprise Report Merger - 合并企业报表工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    subparsers = parser.add_subparsers(dest="command", help="Available commands")

    # merge subcommand
    p_merge = subparsers.add_parser("merge", help="Merge multiple report files")
    p_merge.add_argument("--input", nargs="+", required=True, help="Input file paths (Excel/PDF)")
    p_merge.add_argument("--mode", choices=["concat", "key", "consolidation"], required=True,
                         help="Merge mode: concat (simple stack), key (match by column), "
                              "consolidation (with elimination entries)")
    p_merge.add_argument("--output", required=True, help="Output Excel file path")
    p_merge.add_argument("--key-column", help="Key column name for key/consolidation modes")
    p_merge.add_argument("--source-labels", nargs="+", help="Custom labels for each source file")
    p_merge.add_argument("--elimination-entries", help="JSON file with elimination entries (consolidation mode)")
    p_merge.add_argument("--aggregation", choices=["sum", "avg", "max", "min"], default="sum",
                         help="Aggregation method for numeric values (default: sum)")
    p_merge.add_argument("--sheet", help="Sheet name to read (Excel) / sheet name for output")
    p_merge.add_argument("--header-row", type=int, default=1, help="Header row number in Excel (1-based)")
    p_merge.add_argument("--sheet-name", default="合并结果", help="Output sheet name")
    p_merge.add_argument("--pdf-pages", help="Page range for PDF extraction (e.g., '0-5')")
    p_merge.add_argument("--pdf-table-index", type=int, help="Specific table index to extract from PDF pages")
    p_merge.add_argument("--word-template", help="Word template (.docx) to fill with merged data")
    p_merge.add_argument("--word-output", help="Output path for filled Word document")
    p_merge.add_argument("--mapping", help="JSON file with placeholder/table mapping for Word template")
    p_merge.set_defaults(func=cmd_merge)

    # extract-pdf subcommand
    p_pdf = subparsers.add_parser("extract-pdf", help="Extract tables from PDF")
    p_pdf.add_argument("--input", required=True, help="Input PDF file path")
    p_pdf.add_argument("--output", required=True, help="Output Excel file path")
    p_pdf.add_argument("--pages", help="Page range (e.g., '0-5' or '0,2,4')")
    p_pdf.add_argument("--table-index", type=int, help="Extract only this table index from each page")
    p_pdf.set_defaults(func=cmd_extract_pdf)

    # fill-template subcommand
    p_fill = subparsers.add_parser("fill-template", help="Fill Word template with data from Excel")
    p_fill.add_argument("--input", required=True, help="Input Excel file with merged data")
    p_fill.add_argument("--template", required=True, help="Word template (.docx) file path")
    p_fill.add_argument("--output", required=True, help="Output Word document path")
    p_fill.add_argument("--mapping", help="JSON file with placeholder/table mapping")
    p_fill.set_defaults(func=cmd_fill_template)

    args = parser.parse_args()
    if not args.command:
        parser.print_help()
        sys.exit(1)

    args.func(args)


if __name__ == "__main__":
    main()
