#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
scan_docx.py —— 暗标格式"防盲"合规扫描仪 · Word 确定性提取引擎

功能：
  1. 提取 docx 的版面/排版/字体/页码等机械属性，输出结构化 JSON（供 SKILL 比对评标盲审条款）。
  2. 提取文档元数据（作者/公司/最后修改人），用于"身份显形"扫描。
  3. 全文扫描投标人身份信息残留（精确标识符 + 启发式泄漏词），返回命中片段与位置。

依赖：python-docx（pip install python-docx）。缺失时脚本会给出明确报错，由 WorkBuddy 在受管 venv 中安装。
输出：单行 JSON 到 stdout。大文档的全文不输出，仅输出泄漏命中，避免上下文爆炸。

注意：本脚本只做"机械格式提取 + 文本扫描"，不评价技术方案内容优劣，符合技能护栏铁律。
"""

import sys
import json
import zipfile
import re
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
except ImportError:
    sys.stderr.write("缺少依赖 python-docx。请在受管 venv 执行：pip install python-docx\n")
    sys.exit(2)

EMU_PER_CM = 360000.0
EMU_PER_PT = 12700.0


def emu_to_cm(emu):
    if emu is None:
        return None
    return round(emu / EMU_PER_CM, 2)


def emu_to_pt(emu):
    if emu is None:
        return None
    return round(emu / EMU_PER_PT, 2)


# 启发式泄漏词：出现在技术暗标正文中，往往暗示投标人身份。
# 脚本做基础扫描；更细的语义判定由 SKILL 结合 references/identity_patterns.md 完成。
HEURISTIC_LEAK_PATTERNS = [
    r"我\s*公司", r"我\s*司", r"本\s*公司", r"我\s*们\s*公司", r"我\s*单位",
    r"荣\s*获", r"中\s*标", r"承\s*建", r"独\s*家", r"独\s*占", r"唯\s*一",
    r"行\s*业\s*领\s*先", r"一\s*流", r"领\s*先\s*地\s*位", r"龙\s*头\s*企\s*业",
    r"工\s*法", r"专\s*利", r"知\s*名\s*品\s*牌", r"著\s*名", r"首\s*家",
    r"自\s*主\s*研\s*发", r"核\s*心\s*技\s*术", r"荣\s*获\s*.*奖", r"ISO\s*9\d{3,}",
    r"三\s*标\s*一\s*体", r"AAA\s*级", r"省\s*级\s*.*奖", r"国\s*家\s*级\s*.*奖",
]


def read_doc_metadata(path: Path):
    """直接从 docx 包内 core.xml / app.xml 读取元数据（比 python-docx 的 core_properties 更全）。"""
    meta = {"author": None, "last_modified_by": None, "company": None,
            "title": None, "creator": None, "pages": None}
    try:
        with zipfile.ZipFile(path) as z:
            if "docProps/core.xml" in z.namelist():
                xml = z.read("docProps/core.xml").decode("utf-8", "ignore")
                m = re.search(r"<dc:creator>(.*?)</dc:creator>", xml)
                if m:
                    meta["author"] = m.group(1)
                m = re.search(r"<cp:lastModifiedBy>(.*?)</cp:lastModifiedBy>", xml)
                if m:
                    meta["last_modified_by"] = m.group(1)
                m = re.search(r"<dc:title>(.*?)</dc:title>", xml)
                if m:
                    meta["title"] = m.group(1)
            if "docProps/app.xml" in z.namelist():
                xml = z.read("docProps/app.xml").decode("utf-8", "ignore")
                m = re.search(r"<Company>(.*?)</Company>", xml)
                if m:
                    meta["company"] = m.group(1)
                m = re.search(r"<Pages>(\d+)</Pages>", xml)
                if m:
                    meta["pages"] = int(m.group(1))
            if "docProps/custom.xml" in z.namelist():
                xml = z.read("docProps/custom.xml").decode("utf-8", "ignore")
                m = re.search(r'<property[^>]*name="[^"]*Creator"[^>]*>(.*?)</property>', xml)
                if m and not meta["creator"]:
                    inner = re.search(r"<vt:lpwstr>(.*?)</vt:lpwstr>", m.group(1))
                    if inner:
                        meta["creator"] = inner.group(1)
    except Exception as e:
        meta["_read_error"] = str(e)
    return meta


def get_run_fonts(run):
    """返回 run 的字体（中/英）、字号、加粗/倾斜/下划线/颜色标记。"""
    info = {"eastasia": None, "ascii": None, "size": None,
            "bold": False, "italic": False, "underline": False, "colored": False}
    f = run.font
    # 中文（eastAsia）字体需从 XML 读取
    rpr = run._element.rPr
    if rpr is not None:
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is not None:
            info["eastasia"] = rfonts.get(qn('w:eastAsia'))
            info["ascii"] = rfonts.get(qn('w:ascii'))
    info["ascii"] = info["ascii"] or f.name
    if f.size is not None:
        info["size"] = round(f.size.pt, 2)
    info["bold"] = bool(f.bold)
    info["italic"] = bool(f.italic)
    info["underline"] = bool(f.underline)
    if f.color is not None and f.color.rgb is not None:
        try:
            rgb = str(f.color.rgb)
            if rgb.upper() not in ("000000", "AUTO", "FF000000"):
                info["colored"] = True
        except Exception:
            pass
    return info


def detect_page_number_in_part(part_xml: str) -> bool:
    """检测页眉/页脚 XML 中是否含 PAGE 域（页码）。"""
    if part_xml is None:
        return False
    return "PAGE" in part_xml


def scan_identity(text: str, identifiers: list, run_heuristics: bool, allow: set = None):
    """返回身份泄漏命中列表。

    allow: 白名单集合（小写）。命中的 match 若落在白名单内（如项目允许使用的"我公司"自称），
           则不计入结果，避免误报。精确到字符串匹配。
    """
    allow = set(a.strip().lower() for a in (allow or []) if a.strip())
    hits = []
    # 精确标识符扫描
    for ident in identifiers:
        ident = ident.strip()
        if not ident:
            continue
        for m in re.finditer(re.escape(ident), text):
            start = max(0, m.start() - 50)
            end = min(len(text), m.end() + 50)
            hits.append({"type": "identifier", "match": ident,
                         "pos": m.start(), "context": text[start:end]})
    # 启发式扫描
    if run_heuristics:
        for pat in HEURISTIC_LEAK_PATTERNS:
            for m in re.finditer(pat, text):
                start = max(0, m.start() - 50)
                end = min(len(text), m.end() + 50)
                hits.append({"type": "heuristic", "match": m.group(0),
                             "pos": m.start(), "context": text[start:end]})
    # 白名单过滤（精确匹配 match 字符串，大小写不敏感）
    if allow:
        hits = [h for h in hits if h["match"].strip().lower() not in allow]
    # 去重（按 pos + match）
    seen = set()
    uniq = []
    for h in sorted(hits, key=lambda x: x["pos"]):
        key = (h["pos"], h["match"])
        if key in seen:
            continue
        seen.add(key)
        uniq.append(h)
    return uniq


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx_path", help="待扫描的 .docx 文件路径")
    ap.add_argument("--identifiers", default="", help="投标人可识别身份信息，逗号分隔（公司全称/简称/Logo文字/项目负责人姓名/既往项目名称等）")
    ap.add_argument("--allow", default="", help="白名单（逗号分隔）。命中的措辞若在本名单内则不报，用于项目允许的自称/术语，如'我公司,我单位'。仅精确匹配")
    ap.add_argument("--no-heuristics", action="store_true", help="关闭启发式泄漏词扫描")
    ap.add_argument("--max-text", type=int, default=200000, help="提取全文的最大字符数（仅用于身份扫描，不输出全文）")
    args = ap.parse_args()

    path = Path(args.docx_path)
    if not path.exists():
        sys.stderr.write(f"文件不存在: {path}\n")
        sys.exit(1)

    identifiers = [x for x in args.identifiers.split(",") if x.strip()]

    result = {"file": str(path), "format": "docx", "sections": [], "font_stats": {},
              "alignment_dist": {}, "line_spacing_dist": {}, "indent_dist": {},
              "doc_properties": {}, "leakage_hits": [], "errors": []}

    try:
        doc = Document(str(path))
    except Exception as e:
        sys.stderr.write(f"无法打开 docx（可能损坏或非 Word 格式）: {e}\n")
        sys.exit(3)

    # —— 元数据（身份显形）——
    result["doc_properties"] = read_doc_metadata(path)

    # —— 版面/页边距/页眉页脚/页码 ——
    for i, sec in enumerate(doc.sections):
        try:
            orientation = "portrait" if sec.orientation == WD_ORIENT.PORTRAIT else "landscape"
            has_header = False
            has_footer = False
            header_page_no = False
            footer_page_no = False
            try:
                if sec.header is not None:
                    has_header = any(p.text.strip() for p in sec.header.paragraphs)
                    try:
                        hxml = sec.header.part.blob.decode("utf-8", "ignore")
                        header_page_no = detect_page_number_in_part(hxml)
                    except Exception:
                        pass
                if sec.footer is not None:
                    has_footer = any(p.text.strip() for p in sec.footer.paragraphs)
                    try:
                        fxml = sec.footer.part.blob.decode("utf-8", "ignore")
                        footer_page_no = detect_page_number_in_part(fxml)
                    except Exception:
                        pass
            except Exception as e:
                result["errors"].append(f"section {i} header/footer 读取失败: {e}")
            # 含页码域即视为存在页眉/页脚（即便无可见文字）
            has_header = has_header or header_page_no
            has_footer = has_footer or footer_page_no
            result["sections"].append({
                "index": i,
                "page_width_cm": emu_to_cm(sec.page_width),
                "page_height_cm": emu_to_cm(sec.page_height),
                "orientation": orientation,
                "margin_top_cm": emu_to_cm(sec.top_margin),
                "margin_bottom_cm": emu_to_cm(sec.bottom_margin),
                "margin_left_cm": emu_to_cm(sec.left_margin),
                "margin_right_cm": emu_to_cm(sec.right_margin),
                "has_header": has_header,
                "has_footer": has_footer,
                "header_has_page_number": header_page_no,
                "footer_has_page_number": footer_page_no,
            })
        except Exception as e:
            result["errors"].append(f"section {i} 读取失败: {e}")

    # —— 字体/排版统计 ——
    fonts_eastasia = set()
    fonts_ascii = set()
    sizes = set()
    any_bold = any_italic = any_underline = any_colored = False
    title_number_samples = []

    for p in doc.paragraphs:
        # 对齐分布
        al = p.alignment
        al_name = str(al).split("(")[-1].rstrip(")") if al is not None else "None"
        result["alignment_dist"][al_name] = result["alignment_dist"].get(al_name, 0) + 1

        # 行距分布
        pf = p.paragraph_format
        ls_rule = pf.line_spacing_rule
        ls_val = pf.line_spacing
        if ls_rule is not None:
            rule_name = str(ls_rule).split("(")[-1].rstrip(")")
            if ls_val is not None and rule_name in ("EXACTLY", "AT_LEAST"):
                key = f"{rule_name}:{round(float(ls_val), 1)}pt"
            else:
                key = f"{rule_name}:{ls_val}"
            result["line_spacing_dist"][key] = result["line_spacing_dist"].get(key, 0) + 1

        # 首行缩进分布（cm）
        if pf.first_line_indent is not None:
            ind = emu_to_cm(pf.first_line_indent)
            result["indent_dist"][f"{ind}cm"] = result["indent_dist"].get(f"{ind}cm", 0) + 1

        # 标题编号采样（取段落开头）
        txt = p.text.strip()
        if txt and len(title_number_samples) < 40:
            m = re.match(r"^([0-9一二三四五六七八九十百千]+[、.．\)）]|\(\s*[一二三四五六七八九十]+\s*\)|[（(][0-9]+[）)])", txt)
            if m:
                title_number_samples.append(txt[:30])

        # 字体统计
        for run in p.runs:
            fi = get_run_fonts(run)
            if fi["eastasia"]:
                fonts_eastasia.add(fi["eastasia"])
            if fi["ascii"]:
                fonts_ascii.add(fi["ascii"])
            if fi["size"]:
                sizes.add(fi["size"])
            if fi["bold"]:
                any_bold = True
            if fi["italic"]:
                any_italic = True
            if fi["underline"]:
                any_underline = True
            if fi["colored"]:
                any_colored = True

    result["font_stats"] = {
        "eastasia_fonts": sorted(fonts_eastasia),
        "ascii_fonts": sorted(fonts_ascii),
        "sizes_pt": sorted(sizes),
        "any_bold": any_bold,
        "any_italic": any_italic,
        "any_underline": any_underline,
        "any_colored_text": any_colored,
    }
    result["title_number_samples"] = title_number_samples

    # —— 身份扫描（仅计算命中，不输出全文）——
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if len(full_text) > args.max_text:
        full_text = full_text[:args.max_text]
    allow_set = [x for x in args.allow.split(",") if x.strip()]
    result["leakage_hits"] = scan_identity(full_text, identifiers, not args.no_heuristics, allow_set)
    result["text_length"] = len(full_text)

    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
