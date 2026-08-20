#!/usr/bin/env python3
"""
md_to_docx.py — Convert Markdown to a beautifully formatted DOCX (Word) document.

Part of the hormozi-business-kickstart skill.
Generates a polished, professional Word document from any Markdown file.

Design focus:
- A4 landscape default, centered headings + table cells, body justified
- Strong typography hierarchy: large bold H1, accent H3, readable body
- Color theme: primary navy (#1A3A52) for headings, accent orange (#D35400) for H3 + quote
- Tables: navy header row with white text, alternating row backgrounds, centered cells
- Header/footer with branding

Usage:
    python3 md_to_docx.py <input.md> <output.docx> [options]

Options:
    --title "TITLE"         Document title (defaults to first H1)
    --author "NAME"         Author name
    --business "NAME"       Business name (shown in header)
    --tagline "TEXT"        Tagline (shown in header center)
    --location "PLACE"      Location (shown in header right)
    --primary "#1A3A52"     Primary color (hex, default navy)
    --accent "#D35400"      Accent color (hex, default orange)
    --size a4|a3|a2         Page size (default a4 landscape)
    --portrait              Portrait orientation (default landscape)
    --no-header             Skip header
    --no-footer             Skip footer

Examples:
    python3 md_to_docx.py manual.md manual.docx
    python3 md_to_docx.py manual.md manual.docx --business "Ponovo Novo" --location "Belgrade"
"""
import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


PAGE_SIZES = {
    "a4": (Cm(29.7), Cm(21.0)),    # landscape
    "a3": (Cm(42.0), Cm(29.7)),    # landscape
    "a2": (Cm(59.4), Cm(42.0)),    # landscape
}

# Typography hierarchy
FONT_SIZES = {
    'H1': 22, 'H2': 16, 'H3': 13, 'H4': 11,
    'BODY': 10, 'SMALL': 8, 'TABLE_HEAD': 9, 'TABLE_CELL': 9,
}


def hex_to_rgb(hex_str):
    hex_str = hex_str.lstrip('#')
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex.lstrip('#').upper())
    tc_pr.append(shd)


def set_cell_vertical_alignment(cell, alignment='center'):
    tc_pr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), alignment)
    tc_pr.append(vAlign)


def add_table_borders(table, color_hex='BDC3C7', size='4'):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color_hex)
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def set_paragraph_border(paragraph, color_hex, size='6', position='bottom'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    border = OxmlElement(f'w:{position}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), size)
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), color_hex.lstrip('#').upper())
    pBdr.append(border)
    pPr.append(pBdr)


def add_page_number_field(paragraph, font_size=8, color=None):
    if color is None:
        color = RGBColor(0x7f, 0x8c, 0x8d)
    run = paragraph.add_run()
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_inline(p, text, size, color, bold=False, center=False, italic=False):
    """Add text with markdown bold (**text**) and inline code (`text`) support."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(size)
            run.font.color.rgb = color
            if italic:
                run.italic = True
        else:
            subparts = re.split(r'(`[^`]+`)', part)
            for sp in subparts:
                if sp.startswith('`') and sp.endswith('`'):
                    run = p.add_run(sp[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(size - 1)
                    run.font.color.rgb = color
                    if italic:
                        run.italic = True
                else:
                    if sp:
                        run = p.add_run(sp)
                        run.font.size = Pt(size)
                        run.font.color.rgb = color
                        if bold:
                            run.bold = True
                        if italic:
                            run.italic = True
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def md_to_docx(input_path, output_path, title=None, author=None, business=None,
               tagline=None, location=None, primary="#1A3A52", accent="#D35400",
               size="a4", portrait=False, header=True, footer=True):
    md_path = Path(input_path)
    if not md_path.exists():
        print(f"ERROR: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    md = md_path.read_text(encoding='utf-8')
    lines = md.split('\n')

    PRIMARY = hex_to_rgb(primary)
    ACCENT = hex_to_rgb(accent)
    TEXT = RGBColor(0x2c, 0x3e, 0x50)
    MUTED = RGBColor(0x7f, 0x8c, 0x8d)
    WHITE = RGBColor(0xff, 0xff, 0xff)
    PRIMARY_HEX = primary.lstrip('#').upper()
    BG = 'F4F6F7'

    page_w, page_h = PAGE_SIZES.get(size.lower(), PAGE_SIZES["a4"])
    if portrait:
        page_w, page_h = page_h, page_w

    doc = Document()
    section = doc.sections[0]
    section.page_width = page_w
    section.page_height = page_h
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.4)

    # Default Normal style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(FONT_SIZES['BODY'])
    style.font.color.rgb = TEXT
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15

    # Document metadata
    if title is None:
        for line in lines:
            s = line.rstrip()
            if s.startswith('# '):
                title = s[2:].strip()
                break
    if title is None:
        title = md_path.stem.replace('_', ' ').replace('-', ' ').title()
    doc.core_properties.title = title
    if author:
        doc.core_properties.author = author

    # ---- Header (3-column) ----
    if header and (business or location):
        h_para = section.header.paragraphs[0]
        tab_stops = h_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(13.5), WD_TAB_ALIGNMENT.CENTER)
        tab_stops.add_tab_stop(Cm(page_w.cm - 3), WD_TAB_ALIGNMENT.RIGHT)
        if business:
            r = h_para.add_run(business.upper())
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = PRIMARY
        h_para.add_run("\t").font.size = Pt(9)
        if tagline:
            r2 = h_para.add_run(tagline)
            r2.italic = True
            r2.font.size = Pt(8)
            r2.font.color.rgb = MUTED
        h_para.add_run("\t").font.size = Pt(9)
        if location:
            r3 = h_para.add_run(location)
            r3.font.size = Pt(9)
            r3.font.color.rgb = PRIMARY

    # ---- Footer ----
    if footer:
        f_para = section.footer.paragraphs[0]
        tab_stops = f_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(13.5), WD_TAB_ALIGNMENT.CENTER)
        tab_stops.add_tab_stop(Cm(page_w.cm - 3), WD_TAB_ALIGNMENT.RIGHT)
        fr1 = f_para.add_run(title)
        fr1.font.size = Pt(FONT_SIZES['SMALL'])
        fr1.font.color.rgb = MUTED
        f_para.add_run("\t").font.size = Pt(FONT_SIZES['SMALL'])
        fr2 = f_para.add_run("Stranica ")
        fr2.font.size = Pt(FONT_SIZES['SMALL'])
        fr2.font.color.rgb = MUTED
        add_page_number_field(f_para, font_size=FONT_SIZES['SMALL'], color=MUTED)
        f_para.add_run("\t").font.size = Pt(FONT_SIZES['SMALL'])
        fr3 = f_para.add_run("Built on Hormozi's $100M Trilogy · 2026")
        fr3.font.size = Pt(FONT_SIZES['SMALL'])
        fr3.font.color.rgb = MUTED

    # ---- Main content ----
    i = 0
    table_rows = []
    list_items = []
    in_code_block = False
    code_buffer = []

    def flush_list():
        nonlocal list_items
        if list_items:
            for it in list_items:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Cm(0.6)
                # Bullet character
                bullet_run = p.add_run("•  ")
                bullet_run.font.size = Pt(FONT_SIZES['BODY'])
                bullet_run.font.color.rgb = ACCENT
                bullet_run.bold = True
                add_inline(p, it, size=FONT_SIZES['BODY'], color=TEXT)
            list_items = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            n_cols = max(len(r) for r in table_rows)
            for r in table_rows:
                while len(r) < n_cols:
                    r.append('')
            t = doc.add_table(rows=len(table_rows), cols=n_cols)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            # Disable autofit so we can set explicit widths
            t.autofit = False
            t.allow_autofit = False
            add_table_borders(t)

            # Calculate the available content width on the page, then size
            # the table to 80% of that. This leaves 10% white space on each
            # side, making the table's CENTER alignment visually obvious.
            page_w_cm = section.page_width.cm
            available_cm = page_w_cm - section.left_margin.cm - section.right_margin.cm
            table_w_cm = available_cm * 0.80
            col_w_cm = table_w_cm / n_cols
            col_w_dxa = int(col_w_cm * 567)   # 1 cm = 567 dxa/twips
            table_w_dxa = int(table_w_cm * 567)
            # Indent needed to push the table to the center: half of the
            # white space that will exist on each side of the table.
            # (available - table) / 2 = the offset from the content area edge
            indent_dxa = int((available_cm - table_w_cm) / 2 * 567)

            # Set the table's TOTAL width DIRECTLY in the XML.
            # python-docx's Table.width setter is broken in this version
            # (it does not modify <w:tblW>), so we have to do it ourselves.
            tblW = t._tbl.tblPr.find(qn('w:tblW'))
            if tblW is None:
                from docx.oxml import OxmlElement
                tblW = OxmlElement('w:tblW')
                t._tbl.tblPr.insert(0, tblW)
            tblW.set(qn('w:type'), 'dxa')
            tblW.set(qn('w:w'), str(table_w_dxa))

            # Set the table layout to FIXED so the explicit widths stick
            tblLayout = t._tbl.tblPr.find(qn('w:tblLayout'))
            if tblLayout is None:
                from docx.oxml import OxmlElement
                tblLayout = OxmlElement('w:tblLayout')
                t._tbl.tblPr.append(tblLayout)
            tblLayout.set(qn('w:type'), 'fixed')

            # Force the table to the center using <w:tblInd>. This is the
            # method that works across Word desktop, Word mobile, and
            # LibreOffice — <w:jc> alone is sometimes ignored on mobile.
            tblInd = t._tbl.tblPr.find(qn('w:tblInd'))
            if tblInd is None:
                from docx.oxml import OxmlElement
                tblInd = OxmlElement('w:tblInd')
                t._tbl.tblPr.append(tblInd)
            tblInd.set(qn('w:type'), 'dxa')
            tblInd.set(qn('w:w'), str(indent_dxa))

            # Now set each column's width and the gridCol widths
            for col in t.columns:
                col.width = Cm(col_w_cm)
            for gridCol in t._tbl.findall(qn('w:tblGrid') + '/' + qn('w:gridCol')):
                gridCol.set(qn('w:w'), str(col_w_dxa))

            for ri, row in enumerate(table_rows):
                for ci, cell_text in enumerate(row):
                    cell = t.cell(ri, ci)
                    # Remove ALL existing paragraphs from the cell (otherwise
                    # the default empty paragraph keeps the cell's default
                    # LEFT alignment even when we set a new paragraph's
                    # alignment to CENTER).
                    for old_p in list(cell.paragraphs):
                        old_p._element.getparent().remove(old_p._element)
                    # Force the cell to its column width in the XML directly
                    tcW = cell._tc.tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        from docx.oxml import OxmlElement
                        tcW = OxmlElement('w:tcW')
                        cell._tc.tcPr.append(tcW)
                    tcW.set(qn('w:type'), 'dxa')
                    tcW.set(qn('w:w'), str(col_w_dxa))
                    # Now add a fresh, properly-centered paragraph
                    p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    if ri == 0:
                        add_inline(p, cell_text, size=FONT_SIZES['TABLE_HEAD'],
                                   color=WHITE, bold=True, center=True)
                        set_cell_shading(cell, PRIMARY_HEX)
                    else:
                        add_inline(p, cell_text, size=FONT_SIZES['TABLE_CELL'],
                                   color=TEXT, center=True)
                        if ri % 2 == 0:
                            set_cell_shading(cell, BG)
                    set_cell_vertical_alignment(cell, 'center')
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(2)
            table_rows = []

    while i < len(lines):
        line = lines[i]
        s = line.rstrip()

        # Code fences
        if s.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_buffer = []
            else:
                in_code_block = False
                code_text = '\n'.join(code_buffer).replace('<', '&lt;').replace('>', '&gt;')
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.4)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(FONT_SIZES['SMALL'])
                run.font.color.rgb = MUTED
                code_buffer = []
            i += 1
            continue
        if in_code_block:
            code_buffer.append(s)
            i += 1
            continue

        # H1 — Big, centered, primary color, with thick bottom border
        if s.startswith('# '):
            flush_list(); flush_table()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(s[2:].strip())
            run.bold = True
            run.font.size = Pt(FONT_SIZES['H1'])
            run.font.color.rgb = PRIMARY
            set_paragraph_border(p, PRIMARY_HEX, size='12')

        # H2 — Centered, primary color, with thin bottom border
        elif s.startswith('## '):
            flush_list(); flush_table()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(s[3:].strip())
            run.bold = True
            run.font.size = Pt(FONT_SIZES['H2'])
            run.font.color.rgb = PRIMARY
            set_paragraph_border(p, PRIMARY_HEX, size='6')

        # H3 — Centered, accent (orange) color
        elif s.startswith('### '):
            flush_list(); flush_table()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(s[4:].strip())
            run.bold = True
            run.font.size = Pt(FONT_SIZES['H3'])
            run.font.color.rgb = ACCENT

        # H4 — Centered, primary color
        elif s.startswith('#### '):
            flush_list(); flush_table()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(s[5:].strip())
            run.bold = True
            run.font.size = Pt(FONT_SIZES['H4'])
            run.font.color.rgb = PRIMARY

        # Tables
        elif '|' in s and s.strip().startswith('|'):
            flush_list()
            cells = [c.strip() for c in s.strip('|').split('|')]
            if not all(re.match(r'^[-:\s]+$', c) for c in cells):
                table_rows.append(cells)

        # Empty line
        elif s == '':
            flush_list(); flush_table()

        # Blockquote — centered, accent, italic+bold
        elif s.startswith('> '):
            flush_list(); flush_table()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.right_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(s[2:].strip())
            run.italic = True
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = ACCENT

        # List item
        elif s.startswith('- ') or s.startswith('* ') or re.match(r'^\d+\.\s', s):
            flush_table()
            cleaned = re.sub(r'^[-*\d.]+\s+', '', s)
            list_items.append(cleaned)

        # Horizontal rule
        elif s == '---':
            flush_list(); flush_table()

        # Indented code
        elif s.startswith('    '):
            flush_list(); flush_table()
            p = doc.add_paragraph()
            run = p.add_run(s.strip())
            run.font.name = 'Courier New'
            run.font.size = Pt(FONT_SIZES['SMALL'])
            run.font.color.rgb = MUTED

        # Body
        else:
            flush_list(); flush_table()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(3)
            add_inline(p, s, size=FONT_SIZES['BODY'], color=TEXT)

        i += 1

    flush_list()
    flush_table()

    doc.save(output_path)
    print(f"✓ DOCX generated: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown to a beautifully formatted DOCX document.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__
    )
    parser.add_argument("input", help="Input Markdown file (.md)")
    parser.add_argument("output", help="Output DOCX file (.docx)")
    parser.add_argument("--title", help="Document title (defaults to first H1)")
    parser.add_argument("--author", help="Author name")
    parser.add_argument("--business", help="Business name (shown in header)")
    parser.add_argument("--tagline", help="Tagline (shown in header center)")
    parser.add_argument("--location", help="Location (shown in header right)")
    parser.add_argument("--primary", default="#1A3A52", help="Primary color hex (default: #1A3A52 navy)")
    parser.add_argument("--accent", default="#D35400", help="Accent color hex (default: #D35400 orange)")
    parser.add_argument("--size", default="a4", choices=["a2", "a3", "a4"], help="Page size (default: a4)")
    parser.add_argument("--portrait", action="store_true", help="Portrait orientation (default: landscape)")
    parser.add_argument("--no-header", action="store_true", help="Skip the page header")
    parser.add_argument("--no-footer", action="store_true", help="Skip the page footer")

    args = parser.parse_args()
    md_to_docx(
        input_path=args.input,
        output_path=args.output,
        title=args.title,
        author=args.author,
        business=args.business,
        tagline=args.tagline,
        location=args.location,
        primary=args.primary,
        accent=args.accent,
        size=args.size,
        portrait=args.portrait,
        header=not args.no_header,
        footer=not args.no_footer,
    )


if __name__ == "__main__":
    main()
