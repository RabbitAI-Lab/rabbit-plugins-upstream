#!/usr/bin/env python3

from __future__ import annotations

import json
import sys
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


# Design preset: compact_reference_guide.
# Named overrides:
# - cross_platform_font: Hiragino Sans GB for Latin and CJK glyph stability in Word/LibreOffice.
# - editorial_cover: 28 pt centered navy title and a compact three-metric strip.
# - case_penalty_table: 8.5 pt body with fixed five-column geometry for dense records.

ASCII_FONT = "Hiragino Sans GB"
CJK_FONT = "Hiragino Sans GB"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
INK = "243746"
MUTED = "5F6B76"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GOLD = "8C6A18"
RED = "9B1C1C"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    ascii_font: str = ASCII_FONT,
    cjk_font: str = CJK_FONT,
):
    run.font.name = ascii_font
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cjk_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: float, color: str, bold: bool = False):
    style.font.name = ASCII_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)


def set_style_spacing(style, *, before: float, after: float, line: float):
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.widow_control = True


def add_paragraph_border(paragraph, *, color: str, size: int = 8, space: int = 4):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def set_cell_shading(cell, fill: str):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "C8D0D8", size: int = 6):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int], *, indent: int = TABLE_INDENT_DXA):
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {TABLE_WIDTH_DXA}, got {sum(widths)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblpr = table._tbl.tblPr

    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent))
    tblind.set(qn("w:type"), "dxa")

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    set_table_borders(table)


def set_repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def prevent_row_split(row):
    trpr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trpr.append(cant_split)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def configure_cell_text(
    cell,
    text: str,
    *,
    size: float = 9.0,
    bold: bool = False,
    color: str = INK,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    for idx, part in enumerate(clean_text(text).split("\n")):
        if idx:
            p.add_run().add_break()
        run = p.add_run(part)
        set_run_font(run, size=size, color=color, bold=bold)


def clean_text(value) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    return "\n".join(line.strip() for line in text.split("\n") if line.strip()).strip()


def unique_strings(values) -> list[str]:
    out = []
    seen = set()
    for value in values:
        text = clean_text(value)
        if text and text not in seen:
            seen.add(text)
            out.append(text)
    return out


def add_page_field(paragraph):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    set_run_font(run, size=8.5, color=MUTED)
    run._r.extend([begin, instr, separate, value, end])


def case_bookmark_name(index: int) -> str:
    return f"case_{index:03d}"


def add_bookmark(paragraph, name: str, bookmark_id: int):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    insert_at = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(insert_at, start)
    paragraph._p.append(end)


def add_internal_hyperlink(
    paragraph,
    text: str,
    anchor: str,
    *,
    size: float = 8.5,
):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=BLUE, bold=True)
    run.font.underline = True
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_real_bullet_definition(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet")
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.extend([tabs, ind, spacing])
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    rpr.append(rfonts)
    lvl.extend([start, fmt, text, suff, jc, ppr, rpr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num_id(paragraph, num_id: int):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, num])


def configure_styles(doc: Document) -> int:
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 11, INK)
    set_style_spacing(normal, before=0, after=6, line=1.25)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        set_style_font(style, size, color, bold=True)
        set_style_spacing(style, before=before, after=after, line=1.0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    custom_specs = {
        "Report Subtitle": (13, MUTED, False, 0, 8, 1.1),
        "Report Meta": (10, MUTED, False, 0, 3, 1.15),
        "Source Extract": (10.5, INK, False, 0, 7, 1.25),
        "Small Note": (9, MUTED, False, 3, 5, 1.15),
        "Quality Note": (9.5, MUTED, False, 4, 7, 1.15),
    }
    for name, (size, color, bold, before, after, line) in custom_specs.items():
        style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(style, size, color, bold=bold)
        set_style_spacing(style, before=before, after=after, line=line)

    bullet_num_id = add_real_bullet_definition(doc)
    bullet_style = styles["Report Bullet"] if "Report Bullet" in styles else styles.add_style(
        "Report Bullet", WD_STYLE_TYPE.PARAGRAPH
    )
    set_style_font(bullet_style, 10.5, INK)
    set_style_spacing(bullet_style, before=0, after=4, line=1.25)
    return bullet_num_id


def add_bullet(doc: Document, text: str, num_id: int):
    p = doc.add_paragraph(style="Report Bullet")
    apply_num_id(p, num_id)
    run = p.add_run(clean_text(text) or "原记录未取得")
    set_run_font(run, size=10.5, color=INK)
    return p


def add_label_value(doc: Document, label: str, value: str):
    p = doc.add_paragraph(style="Report Meta")
    p.paragraph_format.keep_together = True
    label_run = p.add_run(f"{label}：")
    set_run_font(label_run, size=10, color=DARK_BLUE, bold=True)
    value_run = p.add_run(clean_text(value) or "原记录未取得")
    set_run_font(value_run, size=10, color=INK)
    return p


def add_source_text(doc: Document, value: str, *, empty_label: str):
    p = doc.add_paragraph(style="Source Extract")
    text = clean_text(value) or empty_label
    for idx, line in enumerate(text.split("\n")):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, size=10.5, color=INK)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    for run in p.runs:
        set_run_font(
            run,
            size={1: 16, 2: 13, 3: 12}[level],
            color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
            bold=True,
        )
    return p


def configure_page(doc: Document, start: int = 0):
    for section in doc.sections[start:]:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = False


def configure_headers_footers(doc: Document, window_label: str):
    section = doc.sections[0]
    section.different_first_page_header_footer = False

    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run("上市公司违规案例周报")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    p.add_run("\t")
    right = p.add_run(window_label)
    set_run_font(right, size=8.5, color=MUTED)
    add_paragraph_border(p, color="D7DBE2", size=6, space=4)

    for footer in (section.footer,):
        fp = footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(4)
        run = fp.add_run("董小屿违规案例库  |  第 ")
        set_run_font(run, size=8.5, color=MUTED)
        add_page_field(fp)
        run2 = fp.add_run(" 页")
        set_run_font(run2, size=8.5, color=MUTED)


def add_metric_strip(doc: Document, metrics: list[tuple[str, str]]):
    table = doc.add_table(rows=1, cols=len(metrics))
    widths = [TABLE_WIDTH_DXA // len(metrics)] * len(metrics)
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    set_table_geometry(table, widths)
    for idx, (value, label) in enumerate(metrics):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        value_run = p.add_run(value)
        set_run_font(value_run, size=20, color=NAVY, bold=True)
        value_run.add_break()
        label_run = p.add_run(label)
        set_run_font(label_run, size=9, color=MUTED, bold=True)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    set_repeat_header(table.rows[0])
    return table


def add_cover(doc: Document, data: dict):
    meta = data["report_meta"]
    stats = data["statistics"]
    start = meta["window_start"]
    end = meta["window_end"]

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(74)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("合规风险观察  ·  WEEKLY REVIEW")
    set_run_font(run, size=10, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("上市公司违规案例周报")
    set_run_font(run, size=28, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run(f"{start} — {end}")
    set_run_font(run, size=14, color=DARK_BLUE)

    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope.paragraph_format.space_after = Pt(26)
    run = scope.add_run(
        "统计口径：董小屿违规案例库 https://www.dxy-aiagent.com/website/weigui"
    )
    set_run_font(run, size=9.5, color=MUTED)

    add_metric_strip(
        doc,
        [
            (str(data["total"]), "违规案例"),
            (str(stats["company_count"]), "涉及公司"),
            (str(stats["party_count"]), "处罚对象"),
        ],
    )

    doc.add_page_break()


def add_distribution_table(doc: Document, data: dict):
    rows = []
    groups = (
        ("板块", data["statistics"]["by_bankuai"]),
        ("触发机构", data["statistics"]["by_trigger_institution"]),
        ("高频违规类型", data["statistics"]["by_violation_type"][:6]),
        ("主要处罚类型", data["statistics"]["by_penalty_type"][:6]),
    )
    for category, items in groups:
        for item in items:
            rows.append((category, item["name"], str(item["count"])))

    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1800, 6060, 1500])
    headers = ("维度", "项目", "案例数")
    for idx, value in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
        configure_cell_text(
            table.rows[0].cells[idx],
            value,
            size=9.5,
            bold=True,
            color=DARK_BLUE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    set_repeat_header(table.rows[0])
    previous = None
    for category, item, count in rows:
        cells = table.add_row().cells
        show_category = category if category != previous else ""
        configure_cell_text(cells[0], show_category, size=9, bold=bool(show_category), color=DARK_BLUE)
        configure_cell_text(cells[1], item, size=9)
        configure_cell_text(cells[2], count, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        prevent_row_split(table.rows[-1])
        previous = category
    set_table_geometry(table, [1800, 6060, 1500])
    return table


def add_case_index_table(doc: Document, cases: list[dict]):
    table = doc.add_table(rows=1, cols=5)
    widths = [600, 1500, 2200, 1100, 3960]
    set_table_geometry(table, widths)
    headers = ("序号", "发布时间", "公司", "板块", "触发机构")
    for idx, value in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
        configure_cell_text(
            table.rows[0].cells[idx],
            value,
            size=8.8,
            bold=True,
            color=DARK_BLUE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    set_repeat_header(table.rows[0])
    for idx, case in enumerate(cases, 1):
        company = case["company"]
        trigger = "、".join(unique_strings(x.get("name") for x in case.get("trigger_institutions", [])))
        values = (
            str(idx),
            case.get("publish_date") or "",
            company.get("bankuai") or "未映射",
            trigger or "原记录未取得",
        )
        row = table.add_row()
        configure_cell_text(row.cells[0], values[0], size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        configure_cell_text(row.cells[1], values[1], size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        configure_cell_text(row.cells[2], "", size=8.5)
        company_paragraph = row.cells[2].paragraphs[0]
        add_internal_hyperlink(
            company_paragraph,
            company.get("short_name") or "未映射",
            case_bookmark_name(idx),
        )
        stock_code = clean_text(company.get("stock_code"))
        if stock_code:
            company_paragraph.add_run().add_break()
            code_run = company_paragraph.add_run(stock_code)
            set_run_font(code_run, size=8.5, color=MUTED)
        configure_cell_text(row.cells[3], values[2], size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        configure_cell_text(row.cells[4], values[3], size=8.5)
        prevent_row_split(row)
    set_table_geometry(table, widths)
    return table


def format_identity(penalty: dict, company: dict) -> str:
    object_name = clean_text(penalty.get("object_name"))
    listed_company_names = {
        clean_text(company.get("short_name")),
        clean_text(company.get("full_name")),
    }
    listed_company_names.discard("")
    if object_name in listed_company_names:
        return "上市公司"
    identity = penalty.get("object_identity") or {}
    role = clean_text(identity.get("role_raw"))
    return role or "——"


def format_violation_types(penalty: dict) -> str:
    values = unique_strings(
        item.get("normalized_type") or item.get("violation_statement")
        for item in penalty.get("violation_types", [])
    )
    return "、".join(values) or "原记录未取得"


def format_amount(value, currency) -> str:
    if value is None:
        return ""
    amount = float(value)
    if currency == "CNY":
        if abs(amount) >= 10000:
            return f"{amount / 10000:g}万元"
        return f"{amount:g}元"
    return f"{amount:g} {currency or ''}".strip()


def format_penalty_types(penalty: dict) -> str:
    values = []
    for item in penalty.get("penalty_types", []):
        name = clean_text(item.get("name")) or "未解码处罚类型"
        amount = format_amount(item.get("amount_value"), item.get("currency"))
        values.append(f"{name}（{amount}）" if amount else name)
    return "、".join(unique_strings(values)) or "原记录未取得"


def format_penalty_term(penalty: dict) -> str:
    term = penalty.get("penalty_term")
    if not term:
        return "原记录未载明处罚期限"
    if term.get("permanent"):
        return "终身"
    if term.get("duration_value") is not None and term.get("duration_unit"):
        return f'{term["duration_value"]}{term["duration_unit"]}'
    start = clean_text(term.get("start_date"))
    end = clean_text(term.get("end_date"))
    raw = clean_text(term.get("raw_text"))
    if start and end:
        return f"{start} 至 {end}"
    if start:
        return f"开始日期：{start}；处罚期限原记录未载明"
    if end:
        return f"截至 {end}；开始日期原记录未载明"
    return raw or "原记录未载明处罚期限"


def add_penalty_table(doc: Document, penalties: list[dict], company: dict):
    show_term = any(penalty.get("penalty_term") for penalty in penalties)
    column_count = 5 if show_term else 4
    table = doc.add_table(rows=1, cols=column_count)
    widths = [1500, 1600, 2240, 2200, 1820] if show_term else [1700, 1800, 2860, 3000]
    set_table_geometry(table, widths)
    headers = ("处罚对象名称", "对象身份", "违规类型", "处罚类型")
    if show_term:
        headers += ("处罚期限",)
    for idx, value in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
        configure_cell_text(
            table.rows[0].cells[idx],
            value,
            size=8.8,
            bold=True,
            color=DARK_BLUE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    set_repeat_header(table.rows[0])
    if not penalties:
        row = table.add_row()
        configure_cell_text(row.cells[0], "原记录未取得处罚对象明细", size=8.5, color=RED)
        row.cells[0].merge(row.cells[-1])
    else:
        for penalty in penalties:
            row = table.add_row()
            values = (
                clean_text(penalty.get("object_name")) or "原记录未取得",
                format_identity(penalty, company),
                format_violation_types(penalty),
                format_penalty_types(penalty),
            )
            if show_term:
                values += (format_penalty_term(penalty),)
            for idx, value in enumerate(values):
                configure_cell_text(row.cells[idx], value, size=8.5)
            prevent_row_split(row)
    set_table_geometry(table, widths)
    return table


def legal_citations(case: dict) -> list[str]:
    return unique_strings(item.get("citation_text") for item in case.get("legal_basis", []))


def add_quality_callout(doc: Document, text: str, *, caution: bool):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_cell_shading(table.cell(0, 0), "FFF8E8" if caution else CALLOUT)
    configure_cell_text(
        table.cell(0, 0),
        text,
        size=9.2,
        color=GOLD if caution else MUTED,
    )
    prevent_row_split(table.rows[0])
    set_repeat_header(table.rows[0])
    return table


def add_case(doc: Document, case: dict, index: int, num_id: int, *, first_case: bool):
    company = case["company"]
    short_name = company.get("short_name") or "未映射公司"
    code = company.get("stock_code") or "代码未取得"
    case_heading = add_heading(doc, f"案例 {index:02d} | {short_name}（{code}）", 2)
    add_bookmark(case_heading, case_bookmark_name(index), 1000 + index)
    if not first_case:
        case_heading.paragraph_format.page_break_before = True

    title = doc.add_paragraph(style="Report Subtitle")
    run = title.add_run(clean_text(case.get("title")) or "案例标题原记录未取得")
    set_run_font(run, size=11, color=MUTED, italic=True)

    triggers = "、".join(unique_strings(x.get("name") for x in case.get("trigger_institutions", [])))
    numbers = "、".join(case.get("penalty_situation", {}).get("admin_numbers") or [])
    add_label_value(doc, "公司全称", company.get("full_name") or "")
    add_label_value(
        doc,
        "基本信息",
        f'{case.get("publish_date") or "原记录未取得"}  |  '
        f'{company.get("bankuai") or "板块未映射"}  |  案例ID {case.get("case_id")}',
    )
    add_label_value(doc, "触发机构", triggers)
    add_label_value(doc, "处罚文号", numbers)

    add_heading(doc, "违规事项", 3)
    add_source_text(
        doc,
        (case.get("violation_matters") or {}).get("text"),
        empty_label="原记录未取得违规事项内容",
    )

    add_heading(doc, "处罚情况", 3)
    add_source_text(
        doc,
        (case.get("penalty_situation") or {}).get("text"),
        empty_label="原记录未取得处罚情况内容",
    )

    add_heading(doc, "案例处罚情况", 3)
    add_penalty_table(doc, case.get("case_penalties") or [], company)

    add_heading(doc, "法规依据", 3)
    citations = legal_citations(case)
    if citations:
        for citation in citations:
            law_paragraph = add_bullet(doc, citation, num_id)
            law_paragraph.paragraph_format.space_after = Pt(2)
            law_paragraph.paragraph_format.line_spacing = 1.15
            for run in law_paragraph.runs:
                set_run_font(run, size=9.8, color=INK)
    else:
        law_paragraph = add_bullet(doc, "原记录未取得法规依据", num_id)
        law_paragraph.paragraph_format.space_after = Pt(2)
        law_paragraph.paragraph_format.line_spacing = 1.15
        for run in law_paragraph.runs:
            set_run_font(run, size=9.8, color=INK)


def build_report(data: dict, output_path: Path):
    doc = Document()
    configure_page(doc)
    bullet_num_id = configure_styles(doc)
    meta = data["report_meta"]
    window_label = f'{meta["window_start"]} 至 {meta["window_end"]}'
    configure_headers_footers(doc, window_label)

    props = doc.core_properties
    props.title = "上市公司违规案例周报"
    props.subject = f"董小屿违规案例库，{window_label}"
    props.author = "Codex"
    props.keywords = "董小屿,违规案例,处罚,周报"

    add_cover(doc, data)

    add_heading(doc, "一、周报概览", 1)
    lead = (
        f'本期共收录 {data["total"]} 起违规案例，涉及 '
        f'{data["statistics"]["company_count"]} 家上市公司、'
        f'{data["statistics"]["party_count"]} 个去重处罚对象。'
        "案例均按董小屿违规案例库发布时间纳入，处罚对象多条记录已合并展示。"
    )
    add_quality_callout(doc, lead, caution=False)

    add_heading(doc, "分布概览", 2)
    add_distribution_table(doc, data)

    add_heading(doc, "案例索引", 2)
    add_case_index_table(doc, data["cases"])

    doc.add_page_break()
    add_heading(doc, "二、案例明细", 1)
    for index, case in enumerate(data["cases"], 1):
        add_case(doc, case, index, bullet_num_id, first_case=index == 1)

    # Keep section geometry explicit after all content has been added.
    configure_page(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build-weekly-word.py DATA_JSON OUTPUT_DOCX")
    data_path = Path(sys.argv[1]).resolve()
    output_path = Path(sys.argv[2]).resolve()
    data = json.loads(data_path.read_text(encoding="utf-8"))
    build_report(data, output_path)
    print(json.dumps({
        "output": str(output_path),
        "cases": data.get("total"),
        "window_start": data.get("report_meta", {}).get("window_start"),
        "window_end": data.get("report_meta", {}).get("window_end"),
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
