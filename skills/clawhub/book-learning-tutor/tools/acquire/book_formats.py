#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
book_formats.py — 多格式书籍 → 纯文本「章/节」层级（授业格式层 v4）

把 PDF / EPUB / MOBI / AZW / AZW3 / DJVU / DOCX / TXT / MD / FB2 抽成
「章 → 节」两层结构的章节，写入 参考/<书名>/NNNN_title.txt（每节一个文件）
+ 参考/<书名>/_sections.json（章/节分组清单），供 course_gen 直接消费（参考/直读，无需中间 预处理/）。
CBZ（漫画）不在此文本路径内——它走 ingest 解包到 参考/<书名>/NNNN_话名/pages/
+ transcript.md，再交给 course_gen 的漫画课程化（generate_comic）。

v4（"小说章没区分 / 嵌套错误"修复）核心变化：
  - extract 不再产出扁平 (title, text)，而是 (chapter, section, text) 三层信息：
      * 小说：每「章」就是最小单位 → section=None，每章 = 一课（不往里塞节）。
      * 课文：每「节」是最小单位，挂在所属的「章」下 → chapter=章名, section=节名。
  - 新增 _split_hierarchical：整本文本按「章正则 + 节正则」两层切，章内再切节；
    章起点 = 第X章/回/卷/篇 + Chapter N + 数字编号 + 序/前言/附录；
    节起点 = 第X节/课 + Section N + X.Y 编号。
  - PDF/EPUB/MOBI/DOCX/TXT 全部改为返回 (chapter, section, text)，并加「防御性
    re-split」：任一段正文里若还嵌着章起点，就地再切，杜绝"一个章文件夹下塞好几章"。
  - 扫描版仍标 needs_ocr；新增可插拔 OCR 钩子（ocr.py）：有 tesseract / MinerU /
    Nougat 后端时自动识别，否则保留标记 + 清晰提示（不捆绑重型模型）。

下游契约（course_gen 直读消费）：
  - 参考/<书名>/NNNN_title.txt：每「节」一个文件（小说每章一个）。
  - 参考/<书名>/_sections.json：[{"chapter","section","title","file"}, ...] 顺序清单；
    缺该清单时 course_gen 退化为「每 txt=一章」（兼容下载/旧目录）。

依赖（按需懒加载，缺哪个报哪个，不强制全装）：
  - pymupdf (fitz)   ：PDF / EPUB / DJVU / 图片 文本层抽取（首选，已装）
  - ebooklib         ：EPUB 兜底解析
  - python-docx      ：DOCX 读取（标题/表格/元数据），已装
  - mobi (0.4.1)     ：MOBI/AZW/AZW3 解包为 HTML，已装
  - beautifulsoup4   ：EPUB/MOBI 的 HTML 清洗，已装
  - djvulibre(djvutxt CLI) ：DJVU 兜底（PyMuPDF wheel 通常不含 djvu）；
                        _djvu_sections 优先找 项目/vendor/djvulibre/，再 PATH
  - OCR（扫描版 PDF）：可选。ocr.py 提供 tesseract（本地轻量）/ MinerU / Nougat
                        （输出 md 直接，理想）后端；均无则标 needs_ocr 不强制。

用法：
  from book_formats import extract, detect_format
  r = extract("book.pdf")          # r['sections'] = [{'chapter','section','title','text'}, ...]
  python book_formats.py --selftest
"""
import os
import re
import sys
import json
import shutil
import subprocess
from pathlib import Path
from collections import Counter

# 文本层稀疏判定阈值：平均每内容页字符数低于此值 → 疑似扫描版（需OCR）
SCAN_THRESHOLD = 25

# 章起点正则（中文/英文/数字编号；命中即视作新章起点）
_HEAD_RE = re.compile(
    r"^\s*(第\s*[一二三四五六七八九十百千零0-9]+\s*[章回节部卷篇集]|"
    r"chapter\s+[0-9ivxIVX]+|"
    r"[0-9]{1,3}[\.\、]\s+\S|"
    r"序[言章]|前[言]|引[言]|附[录]|后[记]|结语)\b",
    re.IGNORECASE,
)
# 节起点正则（章之下的细分；课文用得多，小说一般没有）
_SECTION_RE = re.compile(
    r"^\s*(第\s*[一二三四五六七八九十百千零0-9]+\s*[节课部分段篇]|"
    r"section\s+[0-9ivxIVX]+|"
    r"[0-9]{1,3}\.[0-9]+\s+\S)\b",
    re.IGNORECASE,
)
# 纯页码行（可能带连字符/破折号包围）：- 12 - / 13 / —14—
_PAGENUM_RE = re.compile(r"^[\-–—]?\s*\d{1,4}\s*[\-–—]?$")
# 非章起点（目录/Contents 这类占位标题，不应单独成章）
_NONCHAPTER_RE = re.compile(r"^(目\s*录|contents?|index)$", re.IGNORECASE)


def detect_format(path):
    """返回 'pdf' / 'epub' / 'mobi' / 'azw' / 'azw3' / 'djvu' / 'docx' /
    'txt' / 'md' / 'fb2' / 'cbz' / 'unknown'。"""
    p = Path(path)
    ext = p.suffix.lower().lstrip(".")
    if ext in ("pdf", "epub", "mobi", "azw", "azw3", "djvu", "docx", "txt", "md",
               "fb2", "cbz"):
        return ext
    try:
        head = p.read_bytes()[:64]
    except Exception:
        return "unknown"
    if head.startswith(b"%PDF"):
        return "pdf"
    if head.startswith(b"PK\x03\x04"):  # epub/mobi/azw/cbz 都是 zip 容器，靠扩展名区分
        return "unknown"
    if head.startswith(b"FORM\x00\x00\x00\x0cDJVU"):
        return "djvu"
    # FB2 是 XML（<?xml … <FictionBook）；扩展名缺失时回退识别
    if b"<FictionBook" in head or head.lstrip()[:5] == b"<?xml":
        return "fb2"
    return "unknown"


def _safe_title(title, limit=40):
    """章节标题洗成文件名安全片段（保留中文，去掉路径危险字符）。"""
    t = re.sub(r"[\\/:*?\"<>|]", "_", str(title)).strip() or "未命名"
    t = t.replace("\n", " ").strip()
    return t[:limit]


def _strip_num_prefix(title):
    """去掉标题首部『第N章/节/课』编号词（course_gen 输出章节用，避免与 course_gen
    自动加的 第XX章 前缀重复）。"""
    return re.sub(r"^第[0-9零一二三四五六七八九十百千]+\s*[章章节节节课部分段篇]\s*",
                  "", (title or "")).strip()


_CJK_RE = re.compile(r"[㐀-鿿一-鿿豈-﫿]")


def _is_cjk(ch):
    return bool(_CJK_RE.match(ch)) if ch else False


def _collapse_cjk_spacing(title):
    """折叠 PDF 大纲书签里常见的『字距』空格（如 '数 字 计 算'→'数字计算'、
    '版 权 声 明'→'版权声明'、'前    言'→'前言'）。

    判定：删除两个 CJK 字符之间的空格，当且仅当该空格两侧的字符中**至少一个**的
    另一侧也是空格/边界（即处于"逐字隔开"的字距串中）。这样只吃字距，保留合法的
    单空格（如 '第一节 自然数' 中 节↔自 之间的空格，其两侧字符的另一侧都是正常
    汉字，不会被吃）。不改变语义、不破坏 '第N章' 与正文的间隔。
    """
    if not title:
        return title
    # 先把 2+ 连续空白折叠成单个空格（'前    言'→'前 言'，'定  义  类'→'定 义 类'），
    # 否则双空格会让下方字距判定漏掉中间的空格。
    title = re.sub(r"\s{2,}", " ", title)
    chars = list(title)
    n = len(chars)
    out = []
    for i, ch in enumerate(chars):
        if ch.isspace() and 0 < i < n - 1:
            prev, nxt = chars[i - 1], chars[i + 1]
            if _is_cjk(prev) and _is_cjk(nxt):
                left_ok = (i - 2 < 0) or chars[i - 2].isspace()
                right_ok = (i + 2 >= n) or chars[i + 2].isspace()
                if left_ok or right_ok:
                    continue  # 删掉这个字距空格
        out.append(ch)
    return "".join(out)


def _loose_sig(s):
    """章名/页眉的宽松签名：仅留 CJK 与字母，去空格/标点/数字（章号、页码），
    并折叠连续重复汉字（该教材 PDF 页眉有双倍渲染伪影 '第第章章…'）。用于判定
    一页最顶端的块是否为『章名页眉』（其签名应等于当前章名签名）。"""
    if not s:
        return ""
    s = re.sub(r"[^㐀-鿿A-Za-z]", "", s)
    s = re.sub(r"([㐀-鿿])\1+", r"\1", s)
    return s


def _page_text_clean(doc, p, norms=None, exclude=None):
    """单页文本抽取（dict 模式，按位置重建段落），并剔除最顶/最底端的章级页眉/页脚块。

    该教材 PDF 的章名页眉/页脚固定在每页顶/底端，且渲染异常（text 模式会被逐字拆行/
    双倍化成 '第/章/计/算…' 或 '第第章章…'），故改用 dict 按位置取块；其签名若落在
    norms（仅 level-1 章/前言/目录等章级标题签名，不含 '学习目标' 等小节）且不在
    exclude（当前小节自身标题，避免误删课的正文起始）即判定为页眉/页脚并丢弃。
    无大纲路径传 norms=None 则不剔除。返回与 get_text('text') 近似的段落文本。

    注：章节起始页大标题含双倍渲染伪影（如附录A 的 '附录A Python快速参考' 被拆成
    附/附录/录AA…/考 多行）时，其块位于内容区（top≈14%）而非页面边缘，且 course_gen 已以
    '# 章名' 提供干净标题，故作为无害冗余保留；_is_runner 仅做精确签名匹配、不含子集
    兜底，避免误删正常章节标题（详见 _is_runner 注释）。
    """
    norms = norms or set()
    exclude = exclude or set()
    d = doc[p].get_text("dict")
    blocks = []
    for b in d.get("blocks", []):
        if "lines" not in b:
            continue
        lines = []
        for ln in b.get("lines", []):
            line = "".join(sp.get("text", "") for sp in ln.get("spans", []))
            lines.append(line)
        text = "\n".join(lines).strip()
        if not text:
            continue
        blocks.append((b["bbox"][1], b["bbox"][0], text))
    blocks.sort(key=lambda x: (x[0], x[1]))

    def _is_runner(sig):
        """边缘块是否为章级运行页眉/页脚（精确签名匹配；不含章节标题/正文，避免误删）。

        注：章节起始页的大标题（含双倍渲染伪影，如附录A 的 '附录A Python快速参考' 被
        拆成 附/附录/录AA…/考 多行）位于内容区（top≈14%）而非页面边缘，且 course_gen 已以
        '# 章名' 形式提供干净标题，故作为无害冗余保留，不作页眉剥离——强行剥离需识别
        '双倍渲染标题'，易误删正常章节标题。
        """
        if not sig or sig in exclude:
            return False
        if sig in norms:
            return True
        if re.match(r"^\s*\d+\s*$", sig):
            return True
        return False

    # 连续剥离顶端/底端运行页眉/页脚（首个非 runner 即停，避免跨正文误删）
    while blocks and _is_runner(_loose_sig(blocks[0][2])):
        blocks.pop(0)
    while blocks and _is_runner(_loose_sig(blocks[-1][2])):
        blocks.pop(-1)
    return "\n\n".join(t for _, _, t in blocks)


def _needs_ocr(sections, page_count):
    if page_count <= 0:
        return False
    total = sum(len(t) for s in sections for t in [s["text"]])
    return (total / page_count) < SCAN_THRESHOLD


def _looks_like_heading(txt, maxsz, body_size):
    """综合判断一段文本是否章节标题。"""
    if not txt or len(txt) > 60:
        return False
    if _HEAD_RE.match(txt):
        return True
    if body_size and maxsz >= body_size * 1.2 and len(txt) <= 40:
        return True
    return False


# ---------------------------------------------------------------------------
# 两层切分：章 → 节
# ---------------------------------------------------------------------------
def _split_hierarchical(text):
    """整本纯文本按「章 → 节」两层切分。

    返回 [(chapter, section, text), ...]；无章起点返回 None。
    - 章起点 = _HEAD_RE；节起点 = _SECTION_RE。
    - 一个章内若无节起点 → (chapter, None, body)（小说：章即最小单位）。
    - 有节起点 → 每个 (chapter, section, body)（课文：节是最小单位）。
    - 目录/Contents 之类的占位标题不当章；正文过短的疑似 TOC 项丢弃。
    """
    lines = text.split("\n")
    chap_idxs = [i for i, l in enumerate(lines) if _HEAD_RE.match(l.strip())
                 and not _NONCHAPTER_RE.match(l.strip())]
    if not chap_idxs:
        return None
    triples = []
    for j, cstart in enumerate(chap_idxs):
        cend = chap_idxs[j + 1] if j + 1 < len(chap_idxs) else len(lines)
        chapter_title = lines[cstart].strip()
        cbody_lines = lines[cstart + 1:cend]
        sec_idxs = [k for k, l in enumerate(cbody_lines) if _SECTION_RE.match(l.strip())]
        if not sec_idxs:
            body = "\n".join(cbody_lines).strip()
            triples.append((chapter_title, None, body))
        else:
            for si in range(len(sec_idxs)):
                sstart = sec_idxs[si]
                send = sec_idxs[si + 1] if si + 1 < len(sec_idxs) else len(cbody_lines)
                sec_title = cbody_lines[sstart].strip()
                sec_body = "\n".join(cbody_lines[sstart + 1:send]).strip()
                triples.append((chapter_title, sec_title, sec_body))
    return _clean_triples(triples)


def split_chapter(chapter_title, chapter_text):
    """把单章正文切成 (chapter, section, text) 三元组（章内再切节，复用节起点正则）。

    供爬虫下载路径使用：source_engine 每章已落盘为独立 NNNN_*.txt，只需在章内再切节，
    补齐教科书（章→节）的节级拆分——否则 generate_from_ref 退化成「每 txt=一章」会丢节。
    与 _split_hierarchical 的章内切节逻辑一致，但不检测章起点（章已由爬虫按 TOC 切好）：
    - 章内无节起点 → [(chapter, None, text)]（小说：章即最小单位）
    - 有节起点 → 每个 (chapter, section, text)（课文：节是最小单位）
    """
    lines = (chapter_text or "").split("\n")
    sec_idxs = [k for k, l in enumerate(lines) if _SECTION_RE.match(l.strip())]
    if not sec_idxs:
        body = "\n".join(lines).strip()
        return [(chapter_title, None, body)] if body else []
    triples = []
    for si in range(len(sec_idxs)):
        sstart = sec_idxs[si]
        send = sec_idxs[si + 1] if si + 1 < len(sec_idxs) else len(lines)
        sec_title = lines[sstart].strip()
        sec_body = "\n".join(lines[sstart + 1:send]).strip()
        triples.append((chapter_title, sec_title, sec_body))
    triples = [(c, s, t) for (c, s, t) in triples if t.strip()]
    return _drop_toc_like_triples(triples) if len(triples) >= 3 else triples


def build_sections_manifest(book_dir):
    """为 参考/<书名>/ 下所有 NNNN_*.txt 生成 _sections.json 清单（章→节层级）。

    供爬虫下载路径使用：每章已落盘为独立 NNNN_*.txt，先在章内用 split_chapter 切节，
    再产出与本地抽取同构的 [{chapter, section, title, file, body}, ...]。
    - 节级条目带 body（该节正文切片），generate_from_ref 直接消费，避免整章正文重复挂到每个节下。
    - 章级（无节，section=None）body=None，沿用整文件读取（小说：章即最小单位）。
    - 漫画/空书无 txt 时返回 []（调用方据此跳过写 _sections.json，漫画走 generate_comic）。
    """
    from pathlib import Path
    d = Path(book_dir)
    txts = sorted(d.glob("[0-9][0-9][0-9][0-9]_*.txt"))
    manifest = []
    for fp in txts:
        ch_name = fp.name[5:-4]
        body = fp.read_text(encoding="utf-8")
        triples = split_chapter(ch_name, body)
        if not triples:
            triples = [(ch_name, None, "")]  # 空章也保留条目，防漏章
        for (chapter, section, text) in triples:
            manifest.append({
                "chapter": chapter,
                "section": section,
                "title": section or chapter,
                "file": fp.name,
                "body": text if section else None,
            })
    return manifest


def _clean_triples(triples):
    out = [(c, s, t) for (c, s, t) in triples if t.strip()]
    return _drop_toc_like_triples(out) if out else triples


def _drop_toc_like_triples(triples):
    """丢弃『正文过短』的疑似目录(TOC)/许可证样板项（同 _drop_toc_like_sections）。

    阈值修正（v 后）：不再用绝对 200 字符地板——否则整章内容本就短（短诗/词典条目/
    极简课本）或混排里夹一个真实短节时，短节会被一刀切删掉、内容凭空丢失。改为：
    - 中位数 < 200 → 整章 uniformly 短，不存在『长正文 vs 目录』双峰，全部保留；
    - 中位数 >= 200 → 用相对阈值（中位数的 2%，下限 40）丢弃明显过短的疑似目录项，
      保留真实短节（短节只要不比中位数短 50 倍就留，长章里夹的短节不会被误删）。
    保留旧兜底：若全删光则回退原表，绝不让整章消失。
    """
    import statistics
    if len(triples) < 3:
        return triples
    bodies = [len(t) for _, _, t in triples]
    med = statistics.median(bodies) if bodies else 0
    if med < 200:
        return triples
    thr = max(40, int(0.02 * med)) if med > 0 else 40
    out = [(c, s, t) for (c, s, t) in triples if len(t) >= thr]
    return out if out else triples


def _drop_empty_sections(sections):
    """丢弃空正文章节（多为 EPUB/MOBI 的目录 TOC 占位项），保留 >=1 个。"""
    out = [(t, x) for t, x in sections if x.strip()]
    return out if out else sections


def _drop_toc_like_sections(sections):
    """丢弃"正文过短"的疑似目录(TOC)/许可证样板项（兼容旧调用，逻辑同 _drop_toc_like_triples）。

    阈值修正（v 后）：中位数 < 200 全保留；否则相对阈值（中位数 2%，下限 40）丢疑似目录项；
    全删则回退原表。避免真实短节被 200 绝对地板误删。
    """
    import statistics
    if len(sections) < 3:
        return sections
    bodies = [len(x) for _, x in sections]
    med = statistics.median(bodies) if bodies else 0
    if med < 200:
        return sections
    thr = max(40, int(0.02 * med)) if med > 0 else 40
    out = [(t, x) for t, x in sections if len(x) >= thr]
    return out if out else sections


def _resplit_embedded(triples):
    """防御性 re-split：任一段正文里若还嵌着章起点，就地再切，
    杜绝「一个章文件夹下塞好几章」（PDF 大纲过粗 / 启发式漏切时兜底）。"""
    out = []
    changed = False
    for (chapter, section, text) in triples:
        sub = _split_hierarchical(text)
        if sub and len(sub) > 1:
            changed = True
            # 用 sub 的章替换（保留外层章作为优先；sub 章若是 None 退回外层章）
            for (c2, s2, t2) in sub:
                out.append((c2 or chapter, s2, t2))
        else:
            out.append((chapter, section, text))
    return out if changed else triples


# ---------------------------------------------------------------------------
# 文本清洗（通用，零依赖）—— 与小说章节清洗(clean.py)不同：不删"广告/版权"等
# 站点噪声词，避免误伤真实出版书的正文/版权页。
# ---------------------------------------------------------------------------
def _clean_book_text(text):
    if not text:
        return ""
    out = []
    blank = 0
    for raw in text.replace("\r\n", "\n").replace("\u3000", " ").split("\n"):
        line = " ".join(raw.split())  # 折叠行内多空格 + strip
        if not line:
            blank += 1
            if blank <= 1:  # 段落间至多留一个空行
                out.append("")
            continue
        blank = 0
        if _PAGENUM_RE.match(line):
            continue  # 丢弃纯页码行
        out.append(line)
    while out and not out[0]:
        out.pop(0)
    while out and not out[-1]:
        out.pop()
    return "\n".join(out)


def _drop_running_headers(sections):
    """删跨页重复的页眉/页脚：出现于 >50% 章节、且长度<=40 的短行。"""
    if len(sections) < 3:
        return sections
    shorts = set()
    for sec in sections:
        for ln in sec["text"].split("\n"):
            s = ln.strip()
            if s and len(s) <= 40:
                shorts.add(s)
    appear = Counter()
    for sec in sections:
        present = set()
        for ln in sec["text"].split("\n"):
            s = ln.strip()
            if s in shorts:
                present.add(s)
        for s in present:
            appear[s] += 1
    threshold = max(3, int(0.5 * len(sections)))
    drop = {s for s, c in appear.items() if c >= threshold}
    if not drop:
        return sections
    for sec in sections:
        kept = [ln for ln in sec["text"].split("\n") if ln.strip() not in drop]
        sec["text"] = "\n".join(kept)
    return sections


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def _pdf_meta(doc):
    m = getattr(doc, "metadata", None) or {}
    return {
        "title": (m.get("title") or "").strip(),
        "author": (m.get("author") or "").strip(),
    }


def _outline_to_triples(toc, doc, n):
    """PDF 大纲 → (chapter, section, text)。lvl<=1 视为章，>=2 视为节（挂当前章）。

    设计铁律：**完全信任大纲**，绝不对 TOC 派生出的正文再跑 _split_hierarchical /
    _resplit_embedded 二次切分。否则 PDF 每页页眉（"第5章 ……81" 带引导点页码、字距
    拉开的 "判 断 结 构"）、代码清单标题（"# File: chaos.py"）会被当成"章/节"切出，
    污染章名并产生大量垃圾章——实测某教材 21 个真实章被炸成 753 章、章名还混进页眉
    字距版与引导点页码。

    每个 TOC 条目独占一段正文（start→下一条目页之前），章=level<=1，节=level>=2
    （统一挂到最近一个 level<=1 章名下）。章自身正文即"章首到第一节之前"的引子，
    不重复子节内容。无大纲的兜底路径（_pdf_sections 下方）仍保留 _split_hierarchical。
    """
    # 收集所有 level-1 章级标题签名（运行页眉/页脚都是章级标签，不含『学习目标』等小节）
    chapter_norms = set()
    for (lvl, title, _) in toc:
        if lvl <= 1:
            s = _loose_sig(_collapse_cjk_spacing((title or "").strip()))
            if s:
                chapter_norms.add(s)

    triples = []
    cur_chapter = "正文"
    for i, (lvl, title, page) in enumerate(toc):
        start = max(0, page - 1)
        end = (toc[i + 1][2] - 1) if i + 1 < len(toc) else n - 1
        end = max(end, start)
        title = _collapse_cjk_spacing((title or "").strip())
        if lvl <= 1:
            cur_chapter = title
        # exclude 仅保护『当前小节标题』（防止正文起始的小节标题被误当页眉删）；
        # 绝不可加入当前章签名——否则当前章的页眉块（签名==章签名）将被保护而不被剔除，
        # 残留成 '第/序/前/言…' 逐字碎片。
        exclude = set()
        if lvl > 1:
            s_sig = _loose_sig(title)
            if s_sig:
                exclude.add(s_sig)
        txt = "".join(_page_text_clean(doc, p, chapter_norms, exclude)
                      for p in range(start, end + 1))
        if lvl <= 1:
            triples.append((cur_chapter, None, txt))
        else:
            triples.append((cur_chapter, title, txt))
    return _clean_triples(triples)


def _pdf_sections(path):
    import pymupdf as fitz
    import statistics
    doc = fitz.open(path)
    try:
        n = doc.page_count
        meta = _pdf_meta(doc)
        toc = doc.get_toc()  # [[level, title, page(1-based)], ...]
        if toc:
            # TOC 已是完整层级（章/节/小节）。直接信任大纲，不再对正文跑
            # _resplit_embedded —— 否则 PDF 每页页眉（"第5章 ……81" 带引导点页码）、
            # 代码清单标题会被 _HEAD_RE 当成"章"切出，产生大量垃圾章
            # （实测某教材 21 真实章被炸成 58 章/766 课）。无 TOC 的 fallback 路径
            # 仍保留 _resplit_embedded 作兜底。
            return _outline_to_triples(toc, doc, n), n, meta

        # 无大纲 → 字体大小 + 标题正则 启发式
        body_sizes = []
        for p in range(n):
            d = doc[p].get_text("dict")
            for b in d.get("blocks", []):
                for ln in b.get("lines", []):
                    for sp in ln.get("spans", []):
                        if sp.get("text", "").strip():
                            body_sizes.append(sp.get("size", 0))
        body_size = statistics.median(body_sizes) if body_sizes else 0

        headings = []  # (page_index_0based, title)
        for p in range(n):
            d = doc[p].get_text("dict")
            for b in d.get("blocks", []):
                txt = " ".join(
                    sp.get("text", "") for ln in b.get("lines", [])
                    for sp in ln.get("spans", [])
                ).strip()
                if not txt:
                    continue
                sizes = [
                    sp.get("size", 0) for ln in b.get("lines", [])
                    for sp in ln.get("spans", []) if sp.get("text", "").strip()
                ]
                maxsz = max(sizes) if sizes else 0
                if _looks_like_heading(txt, maxsz, body_size):
                    if not headings or not (headings[-1][0] == p and headings[-1][1] == txt):
                        headings.append((p, txt))
        if headings:
            triples = []
            for i, (pg, title) in enumerate(headings):
                # 下一章起点页（exclusive）：正文只取到下一章所在页之前，避免把
                # 下一章标题行并进本章（否则"一章文件夹下塞好几章"）。
                end_excl = headings[i + 1][0] if i + 1 < len(headings) else n
                txt = "\n".join(doc[p].get_text("text") for p in range(pg, end_excl))
                sub = _split_hierarchical(txt)
                if sub:
                    for (c2, s2, t2) in sub:
                        triples.append((c2 or title, s2, t2))
                else:
                    triples.append((title, None, txt))
            return _resplit_embedded(_clean_triples(triples)), n, meta

        # 字体检测无果 → 整本拼接按标题正则切
        full = "\n".join(doc[p].get_text("text") for p in range(n))
        sub = _split_hierarchical(full)
        if sub:
            return _resplit_embedded(sub), n, meta

        # 完全无结构 → 逐页（保留旧行为，标为单章）
        sections = [(f"第{p + 1}页", doc[p].get_text("text")) for p in range(n)]
        triples = [(t, None, x) for t, x in sections]
        return triples, n, meta
    finally:
        doc.close()


# ---------------------------------------------------------------------------
# EPUB（pymupdf 优先；失败回退 ebooklib）
# ---------------------------------------------------------------------------
def _epub_sections(path):
    """EPUB 章节抽取。

    优先 ebooklib（原生解析 EPUB2/3 的 spine 与 nav，按 XHTML 文档逐章切分更准，
    且不会因 EPUB3 的 epub:| 命名空间 CSS 选择器向 stderr 喷语法错）；
    PyMuPDF 仅作后备（它对 EPUB3 常解析不出 TOC，退化成整本一个文本块）。
    """
    # DRM 加密 EPUB 直接报错（内容文档加密，抽出来是乱码），与 figures 路径一致。
    try:
        import zipfile as _zf
        with _zf.ZipFile(path) as _zz:
            if "META-INF/encryption.xml" in _zz.namelist() and _epub_has_content_drm(_zz):
                raise RuntimeError(
                    "检测到 DRM 加密 EPUB（内容文档已加密），无法抽取；"
                    "请改用无 DRM 的合法副本（如 Standard Ebooks / Gutenberg 公版）。")
    except RuntimeError:
        raise
    except Exception:
        pass

    # —— 主：ebooklib（EPUB 原生解析）——
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
        book = epub.read_epub(path)
        meta = {}
        try:
            t = book.get_metadata("DC", "title")
            a = book.get_metadata("DC", "creator")
            if t:
                meta["title"] = t[0][0].strip()
            if a:
                meta["author"] = a[0][0].strip()
        except Exception:
            pass
        # 从 EPUB nav 目录取 (title, href) 映射，给每个文档定章名（比只取文档内首个
        # 标题更准，且能区分 Chapter 1/2/...；front/back matter 也能拿到 Titlepage 等）。
        toc_links = []
        def _flat(toc):
            for e in toc:
                if isinstance(e, epub.Link):
                    toc_links.append(e)
                elif isinstance(e, (tuple, list)) and e:
                    if isinstance(e[0], epub.Link):
                        toc_links.append(e[0])
                    if len(e) > 1:
                        _flat(e[1])
        try:
            _flat(book.toc)
        except Exception:
            pass
        href2title = {}
        for ln in toc_links:
            base = (ln.href or "").split("#")[0]
            if base and base not in href2title:
                href2title[base] = (ln.title or "").strip()
        triples = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "html.parser")
            for tag in soup(["script", "style"]):
                tag.decompose()
            doc_base = item.get_name().split("/")[-1]
            chap = None
            for k, v in href2title.items():
                if k.split("/")[-1] == doc_base and v:
                    chap = v
                    break
            if not chap:
                h = soup.find(["h1", "h2", "h3"])
                chap = h.get_text(" ", strip=True) if h else None
            chap = chap or meta.get("title") or "正文"
            txt = soup.get_text("\n")
            txt = "\n".join(ln.strip() for ln in txt.splitlines() if ln.strip())
            sub = _split_hierarchical(txt)
            if sub:
                for (c, s, t) in sub:
                    triples.append((c or chap, s, t))
            else:
                triples.append((chap, None, txt))
        if triples:
            return _resplit_embedded(_clean_triples(triples)), len(triples), meta
    except Exception:
        pass
    # —— 后备：PyMuPDF ——
    try:
        import pymupdf as fitz
        doc = fitz.open(path)
        try:
            n = doc.page_count
            meta = _pdf_meta(doc)
            toc = doc.get_toc()
            if toc:
                return _resplit_embedded(_outline_to_triples(toc, doc, n)), n, meta
            full = "\n".join(doc[p].get_text("text") for p in range(n))
            sub = _split_hierarchical(full)
            if sub:
                return _resplit_embedded(sub), n, meta
            return [(meta.get("title") or "全文", None, full)], n, meta
        finally:
            doc.close()
    except Exception as e:
        raise RuntimeError("EPUB 抽取需 ebooklib 或 pymupdf，均失败：%s" % e)


# ---------------------------------------------------------------------------
# DJVU（pymupdf 优先；否则 项目/vendor/djvulibre/ 的 djvutxt；再否则 PATH）
# ---------------------------------------------------------------------------
def _find_djvutxt():
    """优先项目本地 vendor/djvulibre/，再系统 PATH。"""
    local = Path(__file__).resolve().parent.parent.parent / "vendor" / "djvulibre"
    for name in ("djvutxt.exe", "djvutxt"):
        cand = local / name
        if cand.exists():
            return str(cand)
    return shutil.which("djvutxt") or shutil.which("djvutxt.exe")


def _djvu_sections(path):
    try:
        import pymupdf as fitz
        doc = fitz.open(path)
        try:
            if doc.page_count > 0:
                n = doc.page_count
                meta = _pdf_meta(doc)
                toc = doc.get_toc()
                if toc:
                    return _resplit_embedded(_outline_to_triples(toc, doc, n)), n, meta
                full = "\n".join(doc[p].get_text("text") for p in range(n))
                sub = _split_hierarchical(full)
                if sub:
                    return _resplit_embedded(sub), n, meta
                triples = [(f"第{p + 1}页", None, doc[p].get_text("text")) for p in range(n)]
                return triples, n, meta
        finally:
            doc.close()
    except Exception:
        pass
    djvutxt = _find_djvutxt()
    if djvutxt:
        try:
            out = subprocess.run([djvutxt, path], capture_output=True, text=True, timeout=120)
            txt = out.stdout or ""
        except Exception as e:
            raise RuntimeError("djvutxt 执行失败：%s" % e)
        sub = _split_hierarchical(txt)
        if sub:
            return _resplit_embedded(sub), len(sub), {}
        return [("全文", None, txt)], 1, {}
    raise RuntimeError(
        "DJVU 抽取失败：当前环境 PyMuPDF 未带 djvu 支持，且未装 djvulibre(djvutxt)。\n"
        "  已支持本地化：把 djvutxt.exe 放到 项目/vendor/djvulibre/ 即可（无需装系统版）。\n"
        "  或从 https://sourceforge.net/projects/djvu/ 下载 DjVuLibre 后放入该目录。")


# ---------------------------------------------------------------------------
# TXT / MD（直达，无需解析）
# ---------------------------------------------------------------------------
def _text_sections(path, is_md=False):
    raw = Path(path).read_text(encoding="utf-8", errors="ignore")
    if is_md:
        # md：# 章，##/### 节（层级切分）
        lines = raw.split("\n")
        heads = [(i, l) for i, l in enumerate(lines)
                 if re.match(r"^(#{1,6})\s+", l)]
        if len(heads) >= 1:
            triples = []
            cur_chapter = None
            for j, (idx, line) in enumerate(heads):
                m = re.match(r"^(#{1,6})\s+(.*)$", line)
                level = len(m.group(1)); title = m.group(2).strip()
                nxt = heads[j + 1][0] if j + 1 < len(heads) else len(lines)
                body = "\n".join(lines[idx + 1:nxt]).strip()
                if level <= 1:
                    cur_chapter = title
                    triples.append((title, None, body))
                else:
                    triples.append((cur_chapter or "正文", title, body))
            return _clean_triples(triples), len(triples), {}
    # 纯 txt 或无标题 md → 整本按章正则切
    sub = _split_hierarchical(raw)
    if sub:
        return _resplit_embedded(sub), len(sub), {}
    title = os.path.splitext(os.path.basename(path))[0]
    return [(title, None, raw)], 1, {}


# ---------------------------------------------------------------------------
# DOCX（python-docx：标题样式层级→章/节；表格→md；按文档顺序保留段落/表格）
# ---------------------------------------------------------------------------
def _iter_block_items(parent):
    from docx.document import Document as _Doc
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    if isinstance(parent, _Doc):
        elem = parent.element.body
    elif isinstance(parent, Table):
        elem = parent._tbl
    else:
        elem = parent._tc
    for child in elem.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _docx_table_to_md(table):
    rows = []
    for r in table.rows:
        cells = [c.text.strip().replace("\n", " ").replace("|", "/") for c in r.cells]
        rows.append(cells)
    if not rows:
        return ""
    ncol = len(rows[0])
    out = ["| " + " | ".join(rows[0]) + " |"]
    if len(rows) >= 2:
        out.append("| " + " | ".join(["---"] * ncol) + " |")
        for r in rows[1:]:
            out.append("| " + " | ".join(r) + " |")
    return "\n".join(out)


def _docx_heading_level(para):
    name = (para.style.name or "") if para.style else ""
    m = re.match(r"(?:Heading|标题)\s*(\d+)", name, re.IGNORECASE)
    if m:
        return int(m.group(1))
    if name.lower().startswith("heading") or name.startswith("标题"):
        return 1
    return None


def _docx_sections(path):
    import docx
    from docx.text.paragraph import Paragraph as _P
    from docx.table import Table as _T
    doc = docx.Document(path)
    meta = {}
    cp = doc.core_properties
    if getattr(cp, "title", None):
        meta["title"] = cp.title.strip()
    if getattr(cp, "author", None):
        meta["author"] = cp.author.strip()
    fallback = meta.get("title") or os.path.splitext(os.path.basename(path))[0]
    triples = []
    cur_chapter = None
    cur_section = None
    buf = []

    def flush():
        nonlocal cur_section, buf
        if buf:
            triples.append((cur_chapter or fallback, cur_section, "\n".join(buf)))
        cur_section = None
        buf = []

    for blk in _iter_block_items(doc):
        if isinstance(blk, _P):
            txt = blk.text.strip()
            lvl = _docx_heading_level(blk)
            if lvl is not None and txt:
                flush()
                if lvl <= 1:
                    cur_chapter = txt
                    cur_section = None
                else:
                    cur_section = txt
            elif txt:
                buf.append(txt)
        elif isinstance(blk, _T):
            md = _docx_table_to_md(blk)
            if md:
                buf.append(md)
    flush()
    if not triples:
        triples = [(fallback, None, "")]
    return _resplit_embedded(_clean_triples(triples)), len(triples), meta


# ---------------------------------------------------------------------------
# MOBI / AZW / AZW3（mobi 库解包为 HTML，零膨胀；兜底 calibre 转 EPUB）
# ---------------------------------------------------------------------------
def _mobi_html_to_triples(html_path, fallback_title="全文"):
    from bs4 import BeautifulSoup
    try:
        html = open(html_path, encoding="utf-8", errors="ignore").read()
    except Exception as e:
        raise RuntimeError("MOBI 解包后 HTML 读取失败：%s" % e)
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    meta = {}
    if soup.title and soup.title.string:
        meta["title"] = soup.title.string.strip()
    full0 = soup.get_text("\n")
    m = re.search(r"Project Gutenberg eBook of .*? by (.+?)[\r\n]", full0)
    if m:
        meta["author"] = m.group(1).strip()

    heads = [h for h in soup.find_all(["h1", "h2", "h3"]) if h.get_text(strip=True)]
    if len(heads) >= 2:
        triples = []
        cur_chapter = None
        cur_section = None
        buf = []

        def flush():
            nonlocal cur_section, buf
            if buf:
                triples.append((cur_chapter or fallback_title, cur_section, "\n".join(buf)))
            cur_section = None
            buf = []

        blocks = soup.find_all(["h1", "h2", "h3", "p", "table"])
        for el in blocks:
            if el.name in ("h1", "h2", "h3"):
                flush()
                t = el.get_text(strip=True)
                if el.name == "h1":
                    cur_chapter = t
                    cur_section = None
                else:
                    cur_section = t
            else:
                t = el.get_text(strip=True)
                if t:
                    buf.append(t)
        flush()
        secs = _resplit_embedded(_clean_triples(triples))
        return secs, len(secs), meta

    # 无 HTML 标题标签 → 整本文本按章→节切（覆盖 Alice 这类 "CHAPTER I." 起步）
    full = soup.get_text("\n")
    sub = _split_hierarchical(full)
    if sub:
        return _resplit_embedded(sub), len(sub), meta
    return [(meta.get("title") or fallback_title, None, full.strip())], 1, meta


def _mobi_sections(path):
    try:
        import mobi
        _td, html_path = mobi.extract(path)
        return _mobi_html_to_triples(
            html_path, fallback_title=os.path.splitext(os.path.basename(path))[0])
    except ImportError:
        pass
    ec = shutil.which("ebook-convert") or shutil.which("ebook-convert.exe")
    if ec:
        import tempfile as _tf
        fd, epub = _tf.mkstemp(suffix=".epub")
        os.close(fd)
        try:
            subprocess.run([ec, path, epub], capture_output=True, text=True,
                           timeout=180, check=True)
        except Exception as e:
            raise RuntimeError("calibre ebook-convert 转换失败：%s" % e)
        try:
            return _epub_sections(epub)
        finally:
            try:
                os.unlink(epub)
            except Exception:
                pass
    raise RuntimeError(
        "MOBI/AZW 抽取失败：需要 `mobi` 库或 calibre。\n"
        "  推荐：pip install mobi   （纯 Python 零膨胀）\n"
        "  或安装 calibre，ebook-convert 会自动把 MOBI 转 EPUB 再抽取。")


# ---------------------------------------------------------------------------
# OCR 钩子（可选）—— 扫描版 PDF/DJVU 文本层空时调用
# ---------------------------------------------------------------------------
def _try_ocr(path, fmt):
    """若本机有 OCR 后端（tesseract / MinerU / Nougat），对扫描版做识别。
    返回 (text_or_None, backend_name_or_None)。无后端 → (None, None)。"""
    try:
        from ocr import ocr_document
        return ocr_document(path, fmt)
    except Exception as e:
        return None, f"OCR 钩子异常：{e}"


# ---------------------------------------------------------------------------
# FB2（FictionBook，纯 XML，零依赖）—— 小说常见格式
# ---------------------------------------------------------------------------
def _fb2_local(tag):
    """取 XML 标签的本地名（去掉 {namespace} 前缀）。"""
    return tag.split("}", 1)[-1]


def _fb2_text(elem):
    """递归收集元素内全部文本（含 <title>/<p> 内的文本与 tail）。"""
    parts = []
    if elem.text:
        parts.append(elem.text)
    for c in elem:
        parts.append(_fb2_text(c))
        if c.tail:
            parts.append(c.tail)
    return "".join(parts).strip()


def _fb2_collect(sec, chapter_title, out):
    """递归把一个 <section> 收集成 (chapter, section, text) 三元组。

    映射规则（FB2 章节常 ≤2 层，更深也兼容）：
      - 顶层 <section>（chapter_title 还没定）→ 其 <title> 当作章名，section=None。
      - 嵌套 <section> → 继承外层章名，其 <title> 当作节名。
      - 某 <section> 自身还有子 <section>（即"部/卷"容器）→ 容器章名沿用外层，
        其直属 <p> 作为该章导言（section=None），子节再递归。
      - <binary> 内嵌图片：小说极少且需 base64 解码落盘，本次仅解析文字，跳过。
    """
    my_title = None
    my_paras = []
    sub_sections = []
    for child in sec:
        tag = _fb2_local(child.tag)
        if tag == "title":
            my_title = _fb2_text(child)
        elif tag == "section":
            sub_sections.append(child)
        elif tag == "image":
            continue  # 跳过内嵌图片引用（binary 解码超出本次小说文字范围）
        elif tag == "p":
            my_paras.append(_fb2_text(child))
        else:
            # subtitle / epigraph / cite / poem / stanza 等：并入正文
            t = _fb2_text(child)
            if t:
                my_paras.append(t)
    text = "\n".join(my_paras).strip()
    if sub_sections:
        chapter = chapter_title if chapter_title is not None else (my_title or "正文")
        if text:
            out.append((chapter, None, text))  # 容器导言
        for ss in sub_sections:
            _fb2_collect(ss, chapter, out)
    else:
        if chapter_title is None:
            chapter = my_title or "正文"
            section = None
        else:
            chapter = chapter_title
            section = my_title or "（小节）"
        out.append((chapter, section, text))


def _fb2_sections(path):
    """FB2 → (chapter, section, text) 三元组列表 + page_count(0) + meta。

    只解析主内容 <body>（无 name 属性）；脚注/注释 <body name="notes"> 跳过。
    提取 <title-info> 的 book-title / author 进 meta。
    """
    import xml.etree.ElementTree as ET
    try:
        root = ET.parse(path).getroot()
    except Exception as e:
        raise RuntimeError("FB2 解析失败（XML 不合法？）：%s" % e)
    meta0 = {}
    title_info = next((b for b in root if _fb2_local(b.tag) == "title-info"), None)
    if title_info is not None:
        for tag in ("book-title", "author"):
            nodes = [x for x in title_info if _fb2_local(x.tag) == tag]
            val = _fb2_text(nodes[0]) if nodes else ""
            if val:
                meta0["title" if tag == "book-title" else "author"] = val
    bodies = [b for b in root if _fb2_local(b.tag) == "body"]
    named_none = next((b for b in bodies if not b.get("name")), None)
    body = named_none if named_none is not None else (bodies[0] if bodies else None)
    if body is None:
        return [("正文", None, "")], 0, meta0
    triples = []
    for sec in body:
        if _fb2_local(sec.tag) == "section":
            _fb2_collect(sec, None, triples)
    if not triples:
        full = _fb2_text(body)  # 退化：整本作为单章
        return [("正文", None, full)], 0, meta0
    return _clean_triples([(c, s, t) for (c, s, t) in triples if t.strip()]), 0, meta0


# ---------------------------------------------------------------------------
# CBZ（漫画，zip 内嵌图片，零依赖）—— 仅拆分+整合：
# 解包图片到 参考/<书名>/ 漫画课程化布局（与 source_engine 漫画下载一致）
# ---------------------------------------------------------------------------
_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".tiff", ".tif"}


def is_image_ext(name):
    return os.path.splitext(name)[1].lower() in _IMAGE_EXTS


def unpack_cbz(path, out_dir, chapters=None):
    """把 CBZ 解包成 参考/<书名>/ 的漫画课程化布局（供 course_gen.generate_comic 消费）：

        参考/<书名>/
          0001_话名/        （每话=一章）
            pages/
              001.jpg …      （原图，自然序重排）
            transcript.md    （占位：漫画对白见原图，未 OCR）

    返回 (n_chapters, n_pages)。分章规则：
      - CBZ 内顶层子目录 >1 个 → 每个子目录 = 一话（按名自然序）；此时忽略 chapters 参数
        （子目录已是权威分章，不再强行重切）。
      - 否则（单层图片或单一顶层目录）= 扁平整包：
          * chapters 未指定 → 全部图片归为一话（0001_全书）。
          * chapters = N (>1) → 按图片自然序切成 N 段（每段一话：0001_话001 …），
            用于扁平整包无目录元数据时人为恢复章结构（每段页数尽量均等）。
    不处理：CBR（需 unrar）；加密/损坏的 zip 会抛错。
    """
    import zipfile
    out_dir = Path(out_dir)
    try:
        zf = zipfile.ZipFile(path)
    except Exception as e:
        raise RuntimeError("CBZ 解包失败（非 zip / 损坏 / 加密？）：%s" % e)
    bad = zf.testzip()
    if bad is not None:
        raise RuntimeError("CBZ 含损坏项：%s" % bad)
    img_entries = [n for n in zf.namelist()
                   if not n.endswith("/") and is_image_ext(n)]
    if not img_entries:
        raise RuntimeError("CBZ 内未找到任何图片（jpg/png/…），可能实为其他格式。")
    # 按顶层目录分组
    groups = {}
    for n in img_entries:
        parts = n.split("/")
        top = parts[0] if len(parts) > 1 else ""
        groups.setdefault(top, []).append(n)
    if len([k for k in groups if k]) > 1:
        chapter_groups = {k: v for k, v in groups.items() if k}  # 多顶层子目录→各一话
    else:
        flat = img_entries  # 单层/单目录→扁平整包
        if chapters and int(chapters) > 1:
            # 扁平整包：按自然序切成 N 段（每段一话）
            k = int(chapters)
            ordered = sorted(flat, key=_natkey_static)
            size = (len(ordered) + k - 1) // k  # 向上取整，末段可少几页
            chapter_groups = {}
            for idx in range(k):
                seg = ordered[idx * size:(idx + 1) * size]
                if seg:
                    chapter_groups["话%03d" % (idx + 1)] = seg
            if not chapter_groups:
                chapter_groups = {"全书": flat}
        else:
            chapter_groups = {"全书": flat}  # 未指定 chapters→整包单话

    items = sorted(chapter_groups.items(), key=lambda kv: _natkey_static(kv[0]))
    total_pages = 0
    out_dir.mkdir(parents=True, exist_ok=True)
    for i, (ch_name, files_in_ch) in enumerate(items, 1):
        # 去掉 ZIP 里可能自带的 "0001_" 数字前缀，避免与本函数加的序号前缀重复
        # （clean 为空则退回原名，例如纯数字目录名）
        clean = re.sub(r"^\d+[\s_-]*", "", ch_name).strip() or ch_name
        ch_dir = out_dir / ("%04d_%s" % (i, _safe_title(clean, 60)))
        pages_dir = ch_dir / "pages"
        pages_dir.mkdir(parents=True, exist_ok=True)
        for j, fn in enumerate(sorted(files_in_ch, key=_natkey_static), 1):
            data = zf.read(fn)
            ext = os.path.splitext(fn)[1].lower()
            (pages_dir / ("%03d%s" % (j, ext))).write_bytes(data)
            total_pages += 1
        (ch_dir / "transcript.md").write_text(
            "# %s\n\n（漫画原图见 `pages/`，对白见原图，未做 OCR。）\n" % ch_name,
            encoding="utf-8")
    zf.close()
    return len(items), total_pages


def _natkey_static(s):
    """模块级自然序键（供扁平切段时排序，避免闭包在顶层定义前引用）。"""
    return [int(t) if t.isdigit() else t for t in re.split(r"(\d+)", s)]


# ---------------------------------------------------------------------------
# 插图抽取（PDF，pymupdf）—— 多模态训练单元：(图, 图注, 上下文)
# ---------------------------------------------------------------------------
# 严格（行首锚定）：图/表/插画 + 可选标点 + 编号（"图1.1" / "Figure 2" / "Plate XIII"）。
# 注意：CJK 后接数字无单词边界，故不能用 \b（会漏掉"图1.1"）；要求编号可挡掉"图中/图里"等正文误绑。
# 编号支持 阿拉伯数字 / 小写罗马(i,iv) / 大写罗马(I,V,X) —— 公版英文插图书常用 "Plate XIII" "Figure II"。
_CAP_RE = re.compile(r"^(图|表|照片|插画|示意图|结构图|流程图|Plate|Illustration|Fig\.?|Figure)\s*[.\-:：]?\s*[\divxIVX]+", re.IGNORECASE)
# 宽松（行内任意位置）：捕获"参见图1"/"图 3"/"Plate 4"等，作兜底。
_CAP_LOOSE = re.compile(r"图\s*\d|表\s*\d|插画\s*\d|Figure\s*\d|Fig\.?\s*\d|Plate\s*\d|Illustration\s*\d", re.IGNORECASE)
_MIN_FIG_PX = 200 * 200  # 小于此面积的图片视为装饰/图标，丢弃（降噪）

# 装饰性 alt 词（封面/logo/扉页/分隔/边框等）——这些 alt 不该当图注进训练（噪音）。
# 公版 EPUB 常把插图说明写进 <img alt>，但封面/logo 的 alt 是装饰词；需要区分。
_DECOR_ALT = {"cover", "logo", "titlepage", "decorative", "ornament", "bullet",
              "icon", "spacer", "blank", "border", "rule", "background", "header",
              "footer", "separator", "device", "frontispiece", "halftitle", "endpaper"}
# 签名词：只要 alt 含这些词（即便还有别的词），即判定为装饰（封面/logo/扉页本身无"世界事实"信号）
_DECOR_SIG = {"logo", "cover", "titlepage", "frontispiece", "halftitle", "endpaper"}
_DECOR_RE = re.compile(r"[\s,.;:!?()\"'-]+")


def _is_decorative_alt(alt):
    """判断 alt 是否纯装饰（封面/logo 等）。装饰性 → True（丢弃）。"""
    a = (alt or "").strip().lower()
    if not a:
        return True
    toks = [t for t in _DECOR_RE.split(a) if t]
    if not toks:
        return True
    if any(t in _DECOR_SIG for t in toks):
        return True  # 含 logo/cover 等签名词 → 装饰（如 "The Standard Ebooks logo."）
    # 全部 token 都在装饰集 → 装饰；出现描述性实词 → 非装饰
    return all(t in _DECOR_ALT for t in toks)


def _find_caption(page, blocks, img_bbox):
    """在图片同页下方、或下一页顶部，找第一条像图注的行（图/表/Fig 起头或其内含编号）。

    仅做弱启发式：命中即绑，未命中返回空串。误绑风险存在（如正文恰好以"图"起头），
    v1 接受——训练侧可再筛；且图注缺失也比乱绑安全。
    """
    if not img_bbox:
        return ""
    iy1 = img_bbox[3]
    same_page = []
    for (bbox, text) in blocks:
        t = text.strip()
        if not t:
            continue
        if bbox[1] >= iy1 - 6:  # 图片底部之下（含轻微重叠）
            if _CAP_RE.match(t) or _CAP_LOOSE.search(t):
                same_page.append((bbox[1], t))
    if same_page:
        same_page.sort()
        return same_page[0][1][:220]
    return ""


def _pdf_figures(doc, p0, p1, min_px=_MIN_FIG_PX):
    """抽取 [p0,p1] 页范围内的插图（嵌入图片），启发式绑图注，去重。

    返回 [{'ext','width','height','bytes','caption','page'}, ...]。
    降噪三招：① 面积 < min_px 的装饰/图标丢弃；② 相同内容哈希(sha256)只保留首次出现
    （PDF 常把同一图在 TOC/正文重复嵌入）；③ 图注仅在同页图片下方或次页顶部、且命中
    图注正则时绑定（_find_caption）。
    """
    import hashlib
    import pymupdf as fitz
    seen = set()
    out = []
    for p in range(max(0, p0), min(doc.page_count, p1) + 1):
        page = doc[p]
        img_infos = page.get_images(full=True)
        if not img_infos:
            continue
        d = page.get_text("dict")
        blocks = [(b["bbox"], " ".join(
            sp.get("text", "") for l in b.get("lines", []) for sp in l.get("spans", [])))
            for b in d.get("blocks", []) if "lines" in b]
        for imginfo in img_infos:
            xref = imginfo[0]
            if xref == 0:
                continue
            try:
                img = doc.extract_image(xref)
            except Exception:
                continue
            w = img.get("width", 0) or 0
            h = img.get("height", 0) or 0
            if w * h < min_px:
                continue
            data = img.get("image")
            if not data:
                continue
            digest = hashlib.sha256(data).hexdigest()
            if digest in seen:
                continue
            seen.add(digest)
            bbox = None
            try:
                bbox = page.get_image_bbox(imginfo)
            except Exception:
                bbox = None
            caption = _find_caption(page, blocks, bbox)
            out.append({"ext": (img.get("ext") or "png").lower(), "width": w,
                        "height": h, "bytes": data, "caption": caption, "page": p})
    return out


def _find_toc_doc(zf, manifest):
    """定位 TOC 文档：兼容 EPUB3 标准 nav.xhtml、SE 的 toc.xhtml，以及任意 XHTML 内嵌
    <nav epub:type*="toc">。返回该文档的 zip 内相对路径，找不到返回 None。"""
    from bs4 import BeautifulSoup
    # 1) 名字命中：media-type 含 nav / 名为 nav.xhtml / toc.xhtml
    cands = [h for h, v in manifest.items()
             if "nav" in (v[1] or "") or h.lower().endswith("nav.xhtml")
             or h.lower().endswith("toc.xhtml")]
    if not cands:
        # 2) 扫描全部 XHTML 找含 toc 的 <nav>
        for h, v in manifest.items():
            if not (v[1].endswith("xhtml") or v[1].endswith("html")):
                continue
            try:
                s = BeautifulSoup(zf.read(h).decode("utf-8", "ignore"), "html.parser")
            except Exception:
                continue
            nv = s.find("nav", attrs={"epub:type": lambda x: x and "toc" in x.lower()})
            if nv:
                return h
    # 在候选里挑真正含 toc <nav> 的（避免 nav 是 landmark 而非 toc）
    for h in cands:
        try:
            s = BeautifulSoup(zf.read(h).decode("utf-8", "ignore"), "html.parser")
        except Exception:
            continue
        nv = s.find("nav", attrs={"epub:type": lambda x: x and "toc" in x.lower()})
        if nv:
            return h
    return cands[0] if cands else None


def _epub_open(zf):
    """从 EPUB zip 解析 OPF 路径、manifest、spine、TOC 条目。

    返回 (manifest, spine, toc_entries)
      manifest:    {href(相对 zip 根, 已归一化): (id, media_type)}
      spine:       [doc_href, ...] 阅读顺序
      toc_entries: [(title, doc_href, fragment, level), ...] 顺序
    EPUB 是 XHTML 包，PyMuPDF 会栅格化页面（get_images 拿不到原始 <img> 二进制 /
    图注结构），故这里直接解 zip 解析 manifest/spine/TOC 定位 <img> 原始字节。
    """
    import posixpath
    try:
        container = zf.read("META-INF/container.xml").decode("utf-8", "ignore")
    except KeyError:
        container = ""
    opf_path = None
    if container:
        m = re.search(r'full-path="([^"]+)"', container)
        if m:
            opf_path = m.group(1)
    if not opf_path:  # 退化：找第一个 .opf
        opf_path = next((n for n in zf.namelist() if n.lower().endswith(".opf")), None)
    if not opf_path:
        raise RuntimeError("EPUB 缺少 OPF（container.xml 也无 .opf）")
    opf_dir = posixpath.dirname(opf_path)
    opf_xml = zf.read(opf_path).decode("utf-8", "ignore")
    # manifest（逐 <item> 解析属性，顺序无关）
    manifest = {}
    for m in re.finditer(r"<item\b([^>]*)>", opf_xml, re.IGNORECASE):
        attrs = m.group(1)
        iid = re.search(r'\bid="([^"]+)"', attrs, re.IGNORECASE)
        href = re.search(r'\bhref="([^"]+)"', attrs, re.IGNORECASE)
        mt = re.search(r'\bmedia-type="([^"]+)"', attrs, re.IGNORECASE)
        if iid and href:
            h = posixpath.normpath(posixpath.join(opf_dir, href.group(1)))
            props = re.search(r'\bproperties="([^"]+)"', attrs, re.IGNORECASE)
            manifest[h] = (iid.group(1), mt.group(1) if mt else "",
                           props.group(1) if props else "")
    # spine（有序 doc idref → href）
    spine_ids = re.findall(r'<itemref\b[^>]*?\bidref="([^"]+)"', opf_xml, re.IGNORECASE)
    id_to_href = {v[0]: k for k, v in manifest.items()}
    spine = [id_to_href[i] for i in spine_ids if i in id_to_href]
    if not spine:  # 兜底：所有 xhtml 文档按名排序
        spine = sorted(h for h, v in manifest.items()
                       if (v[1].endswith("xhtml") or v[1].endswith("html")))
    # TOC：优先 manifest 中 nav / nav.xhtml / toc.xhtml；否则扫描全部 XHTML 找
    # <nav epub:type*="toc">（兼容 Standard Ebooks 的 toc.xhtml 与 EPUB3 标准 nav.xhtml）。
    nav_href = _find_toc_doc(zf, manifest)
    ncx_href = next((h for h, v in manifest.items()
                     if v[1] == "application/x-dtbncx+xml" or h.lower().endswith(".ncx")), None)
    toc_entries = (_epub_parse_nav(zf, nav_href) if nav_href
                   else _epub_parse_ncx(zf, ncx_href) if ncx_href else [])
    # 归一化 TOC 的 doc href：相对 TOC 文档目录 → 与 manifest/spine 同基准
    toc_dir = posixpath.dirname(nav_href or ncx_href or opf_dir)
    norm_entries = []
    for (title, base, frag, level) in toc_entries:
        nb = posixpath.normpath(posixpath.join(toc_dir, base)) if base else base
        norm_entries.append((title, nb, frag, level))
    return manifest, spine, norm_entries


def _epub_parse_nav(zf, nav_href):
    from bs4 import BeautifulSoup
    html = zf.read(nav_href).decode("utf-8", "ignore")
    soup = BeautifulSoup(html, "html.parser")
    nav = soup.find("nav", attrs={"epub:type": lambda x: x and "toc" in x.lower()}) or soup.find("nav")
    out = []
    if not nav:
        return out

    def walk(ol, level):
        for li in ol.find_all("li", recursive=False):
            a = li.find("a")
            if a:
                href = a.get("href", "")
                title = a.get_text(" ", strip=True)
                base, _, frag = href.partition("#")
                out.append((title, base, frag, level))
            sub = li.find("ol")
            if sub:
                walk(sub, level + 1)
    ol = nav.find("ol")
    if ol:
        walk(ol, 1)
    return out


def _epub_parse_ncx(zf, ncx_href):
    from bs4 import BeautifulSoup
    xml = zf.read(ncx_href).decode("utf-8", "ignore")
    soup = BeautifulSoup(xml, "html.parser")
    out = []
    for np in soup.find_all("navpoint"):
        txt = np.find("text")
        content = np.find("content")
        title = txt.get_text(" ", strip=True) if txt else ""
        src = content.get("src", "") if content else ""
        base, _, frag = src.partition("#")
        lvl = len(np.find_parents("navpoint")) + 1
        out.append((title, base, frag, lvl))
    return out


def _epub_caption(img):
    """EPUB 图注：优先 <figure><figcaption>；否则相邻图注元素（<p class=caption> 等）。"""
    fig = img.find_parent("figure")
    if fig is not None:
        fc = fig.find("figcaption")
        if fc:
            t = fc.get_text(" ", strip=True)
            if t:
                return t[:220]
    cand_tags = []
    sib = img.find_next_sibling()
    while sib is not None and len(cand_tags) < 3:
        cand_tags.append(sib)
        sib = sib.find_next_sibling()
    if img.parent:
        ps = img.parent.find_next_sibling()
        while ps is not None and len(cand_tags) < 3:
            cand_tags.append(ps)
            ps = ps.find_next_sibling()
    for tag in cand_tags:
        if tag.name in ("p", "div", "span", "i", "caption", "dd", "figcaption"):
            t = tag.get_text(" ", strip=True)
            if t and (_CAP_RE.match(t) or _CAP_LOOSE.search(t)):
                return t[:220]
        for el in tag.find_all(["p", "div", "span", "i", "caption", "dd", "figcaption"]):
            t = el.get_text(" ", strip=True)
            if t and (_CAP_RE.match(t) or _CAP_LOOSE.search(t)):
                return t[:220]
    # 兜底：<img alt> 或 <svg><title>/<desc> 无障碍文本作图注（公版 EPUB 常把说明写进
    # alt 而非 <figcaption>，如 Flatland 的 "A diagram of three triangles..."）。
    # 采用条件：非装饰性 且 多词描述（挡掉 cover/logo/titlepage 等装饰词与单字标签）。
    if img.name == "img":
        alt = (img.get("alt") or "").strip()
    else:
        t = img.find("title") or img.find("desc")
        alt = (t.get_text(" ", strip=True) if t else "").strip()
    if alt and (" " in alt) and not _is_decorative_alt(alt):
        return alt[:220]
    return ""


def _epub_has_content_drm(zf):
    """META-INF/encryption.xml 里若任一 CipherReference 指向内容文档(xhtml/opf/ncx)，
    即内容被加密（DRM，如 Adobe ADEPT），抽出来是乱码——法律/许可边界非技术题。
    仅混淆字体(.otf/.ttf/.woff)不算内容加密，可正常抽。"""
    import re as _re
    try:
        xml = zf.read("META-INF/encryption.xml").decode("utf-8", "ignore")
    except Exception:
        return False
    for m in _re.finditer(r'<CipherReference[^>]*URI="([^"]+)"', xml, _re.IGNORECASE):
        uri = m.group(1).lower()
        if uri.endswith((".xhtml", ".html", ".htm", ".opf", ".ncx")):
            return True
    return False


def _epub_figures(path):
    """EPUB 插图抽取：解 zip → 定位各 XHTML 的 <img> 原始字节 + 图注，按 TOC 章/节归属。

    返回同 extract_figures：{fmt,page_count,sections:[{chapter,section,title,level,
    figures,context}]}。figures 每项 {ext,width,height,bytes,caption,page(=doc路径)}。
    降噪：① 装饰/图标小图（面积<_MIN_FIG_PX 的栅格图）丢弃；② sha256 去重；
    ③ 图注走 <figcaption> / 相邻图注文本 / alt 无障碍文本。SVG 无法按尺寸过滤，原样保留。
    DRM 加密 EPUB 直接抛清晰异常（内容文档加密，抽出来是乱码，不产出垃圾）。
    """
    import zipfile
    import posixpath
    import hashlib
    import pymupdf as fitz
    from bs4 import BeautifulSoup

    zf = zipfile.ZipFile(path)
    try:
        if "META-INF/encryption.xml" in zf.namelist() and _epub_has_content_drm(zf):
            raise RuntimeError(
                "检测到 DRM 加密 EPUB（内容文档已加密），无法抽取图文；"
                "请改用无 DRM 的合法副本（如 Standard Ebooks / Gutenberg 公版）。")
        manifest, spine, toc_entries = _epub_open(zf)
        toc_by_doc = {}
        for (title, base, frag, level) in toc_entries:
            toc_by_doc.setdefault(base, []).append((level, title))
        seen = set()
        sections = []
        for doc_href in spine:
            try:
                html = zf.read(doc_href).decode("utf-8", "ignore")
            except KeyError:
                continue
            soup = BeautifulSoup(html, "html.parser")
            for tag in soup(["script", "style"]):
                tag.decompose()
            entries = toc_by_doc.get(doc_href, [])
            if entries:
                top = [e for e in entries if e[0] <= 1]
                chapter = (top[-1][1] if top else entries[0][1])
                deepest = max(entries, key=lambda e: e[0])
                # 章内无独立子节（或子节名与章名相同）→ section 置空，避免 章/节 目录名重复
                section = deepest[1] if (deepest[0] > 1 and deepest[1] != chapter) else None
            else:
                chapter, section = None, None
            # 文档正文（context，前 4000 字）
            context = soup.get_text("\n")
            context = "\n".join(ln.strip() for ln in context.splitlines() if ln.strip())
            figs = []
            for el in soup.find_all(["img", "svg"]):
                if el.name == "svg":
                    # 内联 SVG 矢量图（技术/公版 EPUB 常见，如 Flatland / 数学书插图）。
                    # 无独立二进制，存其 outer XML；矢量不按像素尺寸过滤（有 w/h 属性才粗滤）。
                    data = str(el).encode("utf-8")
                    ext = "svg"
                    mt = ""
                    try:
                        w = int(re.sub(r"\D", "", str(el.get("width") or "")) or 0)
                        h = int(re.sub(r"\D", "", str(el.get("height") or "")) or 0)
                    except Exception:
                        w = h = 0
                    if w and h and w * h < _MIN_FIG_PX:
                        continue
                else:
                    src = el.get("src")
                    if not src or src.startswith("data:"):
                        continue
                    rel = posixpath.normpath(
                        posixpath.join(posixpath.dirname(doc_href), src.split("#")[0]))
                    item = manifest.get(rel)
                    mt = item[1] if item else ""
                    props = item[2] if item else ""
                    if "cover-image" in props:
                        # EPUB3 封面图（manifest properties="cover-image"）：低信号，
                        # 排除避免污染多模态训练数据。
                        continue
                    try:
                        data = zf.read(rel)
                    except KeyError:
                        try:
                            data = zf.read(src.split("#")[0])
                        except KeyError:
                            continue
                    if not data:
                        continue
                    ext = src.split(".")[-1].split("?")[0].lower() if "." in src else ""
                    if not ext and mt:
                        ext = mt.split("/")[-1].lower()
                    is_raster = (mt.startswith("image/") and "svg" not in mt) or ext in ("png", "jpeg", "jpg", "gif", "webp")
                    w = h = 0
                    if is_raster:
                        try:
                            pm = fitz.Pixmap(data)
                            w, h = pm.width, pm.height
                        except Exception:
                            w = h = 0
                        if w * h < _MIN_FIG_PX:
                            continue
                digest = hashlib.sha256(data).hexdigest()
                if digest in seen:
                    continue
                seen.add(digest)
                caption = _epub_caption(el)
                figs.append({"ext": ext, "width": w, "height": h,
                             "bytes": data, "caption": caption, "page": doc_href})
            if figs:
                sections.append({"chapter": chapter, "section": section,
                                 "title": chapter or section or doc_href,
                                 "level": 1 if section is None else 2,
                                 "figures": figs, "context": context[:4000]})
    finally:
        zf.close()
    return {"fmt": "epub", "page_count": len(spine), "sections": sections}


def extract_figures(path, fmt=None):
    """PDF / EPUB 插图抽取入口：按 TOC 章/节归属插图，返回每节的图列表。

    返回 {'fmt','page_count','sections':[{chapter,section,title,level,figures,context}]}。
    figures 每项 {ext,width,height,bytes,caption,page}；context 为该 (章,节) 正文前若干字
    （训练侧图文对齐用）。无图的节不返回；流水线 figures_to_book 只落地有图的 (章,节)。
    """
    path = str(path)
    if not os.path.exists(path):
        raise RuntimeError(f"文件不存在：{path}")
    fmt = fmt or detect_format(path)
    if fmt == "pdf":
        import pymupdf as fitz
        doc = fitz.open(path)
        n = doc.page_count
        toc = doc.get_toc()
        sections = []
        if toc:
            cur_ch = None
            for i, (lvl, title, page) in enumerate(toc):
                t = _collapse_cjk_spacing((title or "").strip())
                if lvl <= 1:
                    cur_ch = t
                start = max(0, page - 1)
                end = (toc[i + 1][2] - 1) if i + 1 < len(toc) else n - 1
                end = max(end, start)
                figs = _pdf_figures(doc, start, end)
                ctx = "".join(doc[p].get_text("text") for p in range(start, end + 1))
                sections.append({"chapter": cur_ch,
                                 "section": (t if lvl > 1 else None),
                                 "title": t, "level": lvl,
                                 "figures": figs, "context": ctx[:4000]})
        else:
            figs = _pdf_figures(doc, 0, n - 1)
            ctx = "".join(doc[p].get_text("text") for p in range(n))
            sections.append({"chapter": None, "section": None, "title": "全文",
                             "level": 0, "figures": figs, "context": ctx[:4000]})
        doc.close()
        return {"fmt": fmt, "page_count": n, "sections": sections}
    elif fmt == "epub":
        return _epub_figures(path)
    else:
        raise RuntimeError(
            f"插图抽取暂仅支持 PDF / EPUB；当前格式 {fmt}。")


# ---------------------------------------------------------------------------
# 统一入口
# ---------------------------------------------------------------------------
def extract(path, fmt=None):
    """抽取一本书为「章/节」层级章节。

    返回 dict：
      {
        'fmt': str, 'page_count': int, 'needs_ocr': bool,
        'book_type': 'novel'|'textbook'|'unknown',
        'sections': [{'chapter','section','title','text'}, ...],  # text 已清洗
        'meta': {'name','source','title','author', ...}
      }
    """
    path = str(path)
    if not os.path.exists(path):
        raise RuntimeError(f"文件不存在：{path}")
    fmt = fmt or detect_format(path)
    meta0 = {}
    if fmt == "pdf":
        triples, n, meta0 = _pdf_sections(path)
    elif fmt == "epub":
        triples, n, meta0 = _epub_sections(path)
    elif fmt == "djvu":
        triples, n, meta0 = _djvu_sections(path)
    elif fmt == "txt":
        triples, n, meta0 = _text_sections(path, is_md=False)
    elif fmt == "md":
        triples, n, meta0 = _text_sections(path, is_md=True)
    elif fmt == "docx":
        triples, n, meta0 = _docx_sections(path)
    elif fmt in ("mobi", "azw", "azw3"):
        triples, n, meta0 = _mobi_sections(path)
    elif fmt == "fb2":
        triples, n, meta0 = _fb2_sections(path)
    elif fmt == "cbz":
        raise RuntimeError(
            "CBZ 是漫画容器，不走 extract 三元组路径。\n"
            "  请用：python pipeline.py ingest <file.cbz>\n"
            "  它会解包到 参考/<书名>/NNNN_话名/pages/，再用 "
            "python course_gen.py 参考/<书名>/ 课程化。")
    else:
        raise RuntimeError(
            f"不支持的格式：{fmt}（{path}）。"
            "支持 pdf/epub/mobi/azw/azw3/djvu/docx/txt/md/fb2（cbz 走 ingest 漫画解包）。")

    # 清洗 + 去跨页页眉页脚
    sections = [{"chapter": c, "section": s, "title": (s or c),
                 "text": _clean_book_text(t)} for (c, s, t) in triples]
    sections = _drop_running_headers(sections)

    # 扫描版判定 + 可选 OCR
    # EPUB/MOBI/TXT/MD 是重排文本，没有"页"概念，OCR 无意义；
    # 仅扫描型位图容器 PDF/DJVU 才可能文本层过薄需 OCR。
    needs_ocr = _needs_ocr(sections, n) if fmt in ("pdf", "djvu") else False
    if needs_ocr and fmt in ("pdf", "djvu"):
        ocr_text, backend = _try_ocr(path, fmt)
        if ocr_text and ocr_text.strip():
            sub = _split_hierarchical(ocr_text)
            if sub:
                sections = [{"chapter": c, "section": s, "title": (s or c),
                             "text": _clean_book_text(t)}
                            for (c, s, t) in _resplit_embedded(sub)]
                needs_ocr = False
                meta0["ocr_backend"] = backend or "unknown"

    # 书籍类型：任一带 section → 课文（章含节）；否则小说（章即最小单位）
    book_type = "textbook" if any(s["section"] for s in sections) else "novel"

    meta = {"name": os.path.splitext(os.path.basename(path))[0], "source": "local-file"}
    meta.update({k: v for k, v in meta0.items() if v})
    return {
        "fmt": fmt,
        "page_count": n,
        "needs_ocr": needs_ocr,
        "book_type": book_type,
        "sections": sections,
        "meta": meta,
    }


# ---------------------------------------------------------------------------
# 自测
# ---------------------------------------------------------------------------
def selftest():
    print("[book_formats selftest]")
    try:
        import pymupdf as fitz
    except Exception as e:
        print("   skipped: pymupdf 缺失（%s）" % e)
        return
    import tempfile

    # --- 用例1：带大纲的 PDF（课文：章含节）→ 按大纲拆 2 章，含节 ---
    d = fitz.open()
    for i in range(4):
        pg = d.new_page()
        pg.insert_text((72, 72), f"这是第 {i + 1} 页的示例正文。")
    d.set_toc([
        (1, "第一章 整数", 1),
        (2, "第一节 自然数", 1),
        (2, "第二节 四则运算", 2),
        (1, "第二章 代数", 3),
        (2, "第一节 方程", 3),
    ])
    fd, pdf1 = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    d.save(pdf1)
    d.close()
    r = extract(pdf1)
    assert r["fmt"] == "pdf" and r["page_count"] == 4
    assert r["book_type"] == "textbook", "大纲含节 → 应为 textbook，实际 %s" % r["book_type"]
    chs = [s["chapter"] for s in r["sections"]]
    secs = [s["section"] for s in r["sections"] if s["section"]]
    assert "第一章 整数" in chs and "第二章 代数" in chs, chs
    assert "第一节 自然数" in secs and "第二节 四则运算" in secs, secs
    os.unlink(pdf1)

    # --- 用例2：无大纲小说 PDF（多章，无节）→ 启发式切多章，book_type=novel ---
    d = fitz.open()
    p1 = d.new_page()
    p1.insert_text((72, 60), "Chapter 1 The Beginning", fontsize=18)
    p1.insert_text((72, 90), "The boy stood in the courtyard, looking up at the sky.", fontsize=11)
    p2 = d.new_page()
    p2.insert_text((72, 60), "Chapter 2 The Encounter", fontsize=18)
    p2.insert_text((72, 90), "From now on, the youth called a waste would cease to be.", fontsize=11)
    p3 = d.new_page()
    p3.insert_text((72, 60), "Chapter 3 The Return", fontsize=18)
    p3.insert_text((72, 90), "He came back to the city where it all began.", fontsize=11)
    fd, pdf2 = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    d.save(pdf2)
    d.close()
    r = extract(pdf2)
    assert r["fmt"] == "pdf"
    assert r["book_type"] == "novel", "无节多章 → 应为 novel，实际 %s" % r["book_type"]
    assert len(r["sections"]) >= 3, "应切出 >=3 章，实际 %d" % len(r["sections"])
    titles = " ".join(s["title"] for s in r["sections"])
    assert "Chapter 1" in titles and "Chapter 2" in titles and "Chapter 3" in titles, titles
    # 断言：没有任何一节的正文里还嵌着别的章起点（杜绝"一章文件夹下好几章"）
    for s in r["sections"]:
        assert not any(_HEAD_RE.match(ln.strip()) and not _NONCHAPTER_RE.match(ln.strip())
                        for ln in s["text"].split("\n")), \
            "正文内仍嵌章起点：%s" % s["title"]
    os.unlink(pdf2)

    # --- 用例3：扫描版（文本层空）→ needs_ocr=True（无 OCR 后端时）---
    d = fitz.open()
    pg = d.new_page()
    pg.insert_text((72, 72), " ")
    fd, pdf3 = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    d.save(pdf3)
    d.close()
    r = extract(pdf3)
    assert r["needs_ocr"] is True, "文本层空应标 needs_ocr"
    os.unlink(pdf3)

    # --- 用例4：DOCX → 标题层级切章/节 + 表格 + 元数据 ---
    try:
        import docx as _dx
        dd = _dx.Document()
        dd.core_properties.title = "Docx Test Book"
        dd.core_properties.author = "John Smith"
        dd.add_heading("第一章 引言", level=1)
        dd.add_paragraph("这是第一章的示例正文。")
        t = dd.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "项目"
        t.cell(0, 1).text = "说明"
        t.cell(1, 0).text = "A"
        t.cell(1, 1).text = "示例"
        dd.add_heading("第一节 背景", level=2)
        dd.add_paragraph("这是第一节的示例正文。")
        dd.add_heading("第二章 方法", level=1)
        dd.add_paragraph("这是第二章的示例正文。")
        fd, docx1 = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        dd.save(docx1)
        r = extract(docx1)
        assert r["fmt"] == "docx", r
        assert r["book_type"] == "textbook", r["book_type"]
        assert len(r["sections"]) >= 3, "DOCX 应切出 >=3 节，实际 %d" % len(r["sections"])
        body = "\n".join(s["text"] for s in r["sections"])
        assert "|" in body, "DOCX 表格应转成 md 表格"
        assert r["meta"].get("author") == "John Smith", r["meta"]
        os.unlink(docx1)
    except ImportError:
        print("  (skipped DOCX 用例：python-docx 未装)")

    # --- 用例5：MOBI（用预下载样本；无样本则跳过）---
    _mobi_sample = os.path.join(tempfile.gettempdir(), "alice.mobi")
    if os.path.exists(_mobi_sample):
        r = extract(_mobi_sample)
        assert r["fmt"] == "mobi", r
        assert len(r["sections"]) >= 12, "MOBI 应切出多章，实际 %d" % len(r["sections"])
        titles = " ".join(s["title"] for s in r["sections"])
        assert "CHAPTER I" in titles, "MOBI 章节标题未被识别：%s" % titles[:200]
        assert r["needs_ocr"] is False
        print("  mobi 样本切出 %d 章" % len(r["sections"]))
    else:
        print("  (skipped MOBI 用例：无样本 alice.mobi)")

    # --- 用例6：FB2（纯 XML 小说）→ 章/节映射 + 嵌套 + 元数据 ---
    try:
        import zipfile as _zf
        fb2_xml = (
            '<?xml version="1.0" encoding="UTF-8"?>\n'
            '<FictionBook xmlns="http://www.gribuser.ru/xml/fictionbook/2.0">\n'
            '  <title-info><book-title>测试小说</book-title>'
            '<author>测试作者</author></title-info>\n'
            '  <body>\n'
            '    <section><title><p>第一章 启程</p></title>'
            '<p>夜色中，少年踏上了旅程。</p><p>风从远方吹来。</p></section>\n'
            '    <section><title><p>第二章 相遇</p></title>'
            '<p>他在路口遇见了旅人。</p></section>\n'
            '    <section><title><p>第一部</p></title>'
            '<section><title><p>第三章 开端</p></title><p>内容。</p></section>'
            '</section>\n'
            '  </body>\n'
            '</FictionBook>\n'
        )
        fd, fb2 = tempfile.mkstemp(suffix=".fb2")
        os.close(fd)
        Path(fb2).write_text(fb2_xml, encoding="utf-8")
        r = extract(fb2)
        assert r["fmt"] == "fb2", r
        # 本样本含嵌套（第一部>第三章），故有 section → textbook；
        # 纯平铺章节（无嵌套）的 FB2 小说才会是 novel（见下方 flat 用例）。
        assert r["book_type"] == "textbook", "FB2 含子节 → 应为 textbook，实际 %s" % r["book_type"]
        chapters = [s["chapter"] for s in r["sections"]]
        assert "第一章 启程" in chapters and "第二章 相遇" in chapters, chapters
        # 嵌套：第一部（容器）→ 第三章 开端 应作为 节 挂在第一部 下
        nested = [(s["chapter"], s["section"]) for s in r["sections"]
                  if "第三章" in (s["section"] or "")]
        assert nested and nested[0][0] == "第一部" and nested[0][1] == "第三章 开端", nested
        assert r["meta"].get("title") == "测试小说" and r["meta"].get("author") == "测试作者", r["meta"]
        os.unlink(fb2)

        # FB2 平铺章节（无嵌套）→ novel
        fb2_flat = (
            '<?xml version="1.0" encoding="UTF-8"?>\n'
            '<FictionBook xmlns="http://www.gribuser.ru/xml/fictionbook/2.0">\n'
            '  <body>\n'
            '    <section><title><p>第一章</p></title><p>内容一。</p></section>\n'
            '    <section><title><p>第二章</p></title><p>内容二。</p></section>\n'
            '  </body>\n'
            '</FictionBook>\n'
        )
        fd, fb2b = tempfile.mkstemp(suffix=".fb2")
        os.close(fd)
        Path(fb2b).write_text(fb2_flat, encoding="utf-8")
        rb = extract(fb2b)
        assert rb["book_type"] == "novel", "FB2 平铺章节 → 应为 novel，实际 %s" % rb["book_type"]
        assert len(rb["sections"]) == 2, "FB2 平铺应切出 2 章，实际 %d" % len(rb["sections"])
        os.unlink(fb2b)
    except ImportError:
        print("  (skipped FB2 用例：xml.etree 不可用)")

    # --- 用例7：CBZ（漫画）→ 解包到 参考/<书名>/ 漫画布局（pages/ + transcript.md） ---
    import zipfile as _zip
    fd, cbz = tempfile.mkstemp(suffix=".cbz")
    os.close(fd)
    with _zip.ZipFile(cbz, "w") as z:
        z.writestr("0001_第一话/001.jpg", b"\xff\xd8\xff\xffjpg")
        z.writestr("0001_第一话/002.png", b"\x89PNG\r\n\x1a\n")
        z.writestr("0002_第二话/001.jpg", b"\xff\xd8\xff\xffjpg")
        z.writestr("notes.txt", b"not-an-image")  # 应被忽略（非图片）
    cbz_out = tempfile.mkdtemp(prefix="cbz_selftest_")
    n_ch, n_pg = unpack_cbz(cbz, cbz_out)
    assert n_ch == 2 and n_pg == 3, (n_ch, n_pg)
    for i, ch in enumerate(("0001_第一话", "0002_第二话"), 1):
        cd = Path(cbz_out) / ch
        assert cd.is_dir() and (cd / "pages").is_dir(), "缺话目录/pages：%s" % cd
        assert (cd / "transcript.md").exists(), "缺 transcript.md：%s" % cd
        np = len([x for x in (cd / "pages").iterdir() if x.is_file()])
        assert np == (2 if i == 1 else 1), "话 %s 页数异常：%d" % (ch, np)
    # 单层（无子目录）CBZ → 合并为一话（0001_全书）
    fd, cbz2 = tempfile.mkstemp(suffix=".cbz")
    os.close(fd)
    with _zip.ZipFile(cbz2, "w") as z:
        z.writestr("001.jpg", b"\xff\xd8\xff\xffjpg")
        z.writestr("002.jpg", b"\xff\xd8\xff\xffjpg")
    cbz_out2 = tempfile.mkdtemp(prefix="cbz_selftest2_")
    n_ch2, n_pg2 = unpack_cbz(cbz2, cbz_out2)
    assert n_ch2 == 1 and n_pg2 == 2, (n_ch2, n_pg2)
    assert (Path(cbz_out2) / "0001_全书" / "pages" / "001.jpg").exists()
    # 扁平整包 + chapters=3 → 切成 3 话（每话 2 页，共 6 图）
    fd, cbz3 = tempfile.mkstemp(suffix=".cbz")
    os.close(fd)
    with _zip.ZipFile(cbz3, "w") as z:
        for i in range(1, 7):
            z.writestr("%03d.jpg" % i, b"\xff\xd8\xff\xffjpg%d" % i)
    cbz_out3 = tempfile.mkdtemp(prefix="cbz_selftest3_")
    n_ch3, n_pg3 = unpack_cbz(cbz3, cbz_out3, chapters=3)
    assert n_ch3 == 3 and n_pg3 == 6, (n_ch3, n_pg3)
    for idx in range(1, 4):
        cd = Path(cbz_out3) / ("000%d_话00%d" % (idx, idx))
        assert cd.is_dir(), "缺切分话目录：%s" % cd
        np = len([x for x in (cd / "pages").iterdir() if x.is_file()])
        assert np == 2, "切分话 %s 页数异常：%d" % (cd, np)
    shutil.rmtree(cbz_out, ignore_errors=True)
    shutil.rmtree(cbz_out2, ignore_errors=True)
    shutil.rmtree(cbz_out3, ignore_errors=True)
    os.unlink(cbz)
    os.unlink(cbz2)
    os.unlink(cbz3)

    # --- 用例8：detect_format 扩展名 + FB2(xml 回退) ---
    assert detect_format("book.fb2") == "fb2"
    assert detect_format("comic.cbz") == "cbz"
    fd, fb2x = tempfile.mkstemp(suffix=".xml")
    os.close(fd)
    Path(fb2x).write_text('<?xml version="1.0"?>\n<FictionBook/>\n', encoding="utf-8")
    assert detect_format(fb2x) == "fb2", "无扩展名的 FB2 未被 xml 回退识别"
    os.unlink(fb2x)

    # --- 用例9：PDF 插图抽取（造嵌图 + 图注，验证抽中且图注命中）---
    try:
        import pymupdf as fitz
        # 用 MuPDF 内置 CJK 字体 "china-s" 渲染中文图注。
        # 注意：本机 PyMuPDF 的 insert_text(fontfile=系统.ttf/.ttc) 在 CJK 下会渲染成 tofu
        #（编码未走 CID），导致图注正则失配；内置 china-s 可正确落字，故测试统一用它。
        # 生产代码只读真实 PDF 的文本层，不渲染，此坑仅影响本合成用例。
        caption = "图1.1 这是一个测试示意图。"
        d = fitz.open()
        pg = d.new_page()
        # 造一张 240x240 的纯色图（>min_px 阈值）嵌入页面
        pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 240, 240), False)
        pg.insert_image(fitz.Rect(50, 50, 290, 290), pixmap=pix)
        # 图注行（图片下方）：以"图"起头 + 数字，应被 _find_caption 命中
        pg.insert_text((50, 320), caption, fontsize=12, fontname="china-s")
        # 再放一张极小装饰图（<min_px），应被过滤丢弃
        small = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 80, 80), False)
        pg.insert_image(fitz.Rect(50, 360, 130, 440), pixmap=small)
        fd, figpdf = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)
        d.save(figpdf)
        d.close()
        fr = extract_figures(figpdf)
        # 收集所有节的图（去噪后应只抽到那张大图，过滤掉小装饰图）
        all_figs = [fg for s in fr["sections"] for fg in s["figures"]]
        assert len(all_figs) == 1, "应只抽中 1 张大图（小装饰图被过滤），实际 %d" % len(all_figs)
        cap = all_figs[0]["caption"]
        # 内置 china-s 字体在测试 PDF 里会把 "1.1" 渲染成 "1 . 1"（数字间插空格），
        # 故用宽松子串断言；真实 PDF 文本层里"图1.1"是连续的，不影响生产。
        assert cap and ("图1" in cap and "示意图" in cap), "图注未命中：%r" % cap
        assert all_figs[0]["width"] >= 200 and all_figs[0]["height"] >= 200, "大图尺寸异常"
        os.unlink(figpdf)
    except ImportError:
        print("  (skipped 插图用例：pymupdf 缺失)")

    # --- 用例10：EPUB 插图抽取（造含 <figure><img><figcaption> 的 EPUB，验证抽中+图注+过滤装饰图）---
    try:
        import pymupdf as fitz
        import zipfile as _zip
        def _png(w, h):
            px = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, w, h), False)
            return px.tobytes("png")
        big1, big2, small = _png(240, 240), _png(300, 200), _png(80, 80)
        container = ('<?xml version="1.0"?>\n'
                     '<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">'
                     '<rootfiles><rootfile full-path="OEBPS/content.opf" '
                     'media-type="application/oebps-package+xml"/></rootfiles></container>')
        opf = ('<?xml version="1.0"?>\n<package xmlns="http://www.idpf.org/2007/opf" version="3.0">'
               '<manifest>'
               '<item id="ch1" href="ch1.xhtml" media-type="application/xhtml+xml"/>'
               '<item id="ch2" href="ch2.xhtml" media-type="application/xhtml+xml"/>'
               '<item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>'
               '<item id="f1" href="img/fig1.png" media-type="image/png"/>'
               '<item id="f2" href="img/fig2.png" media-type="image/png"/>'
               '<item id="f3" href="img/fig3.png" media-type="image/png"/>'
               '<item id="f5" href="img/fig5.png" media-type="image/png"/>'
               '<item id="sm" href="img/small.png" media-type="image/png"/>'
               '</manifest><spine><itemref idref="ch1"/><itemref idref="ch2"/></spine></package>')
        nav = ('<?xml version="1.0"?>\n<html xmlns="http://www.w3.org/1999/xhtml" '
               'xmlns:epub="http://www.idpf.org/2007/ops"><body><nav epub:type="toc"><ol>'
               '<li><a href="ch1.xhtml">第一章</a></li>'
               '<li><a href="ch2.xhtml">第二章</a></li>'
               '</ol></nav></body></html>')
        ch1 = ('<?xml version="1.0"?>\n<html xmlns="http://www.w3.org/1999/xhtml"><body>'
               '<h1>第一章</h1>'
               '<figure><img src="img/fig1.png"/>'
               '<figcaption>图1.1 第一章的示意图。</figcaption></figure>'
               '<img src="img/small.png"/>'  # 无图注的装饰小图 → 应被尺寸过滤
               '</body></html>')
        ch2 = ('<?xml version="1.0"?>\n<html xmlns="http://www.w3.org/1999/xhtml"><body>'
               '<h1>第二章</h1>'
               '<figure><img src="img/fig2.png"/>'
               '<figcaption>图2.3 第二章的流程图。</figcaption></figure>'
               '<p>一段普通正文，前后都不该被误绑为图注。</p>'
               # 非 <figure> 包裹的栅格图 + 相邻 <p class="caption">（验证 _epub_caption 兄弟节点路径）
               '<img src="img/fig3.png"/>'
               '<p class="caption">图4 邻位图注。</p>'
               # 内联 SVG 矢量图 + 相邻图注（验证 SVG 分支不被尺寸过滤、且能绑图注）
               '<svg width="300" height="200" xmlns="http://www.w3.org/2000/svg">'
               '<rect width="300" height="200" fill="#c33"/></svg>'
               '<p class="caption">图3 矢量示意图。</p>'
               '<img src="img/fig5.png" alt="图5 靠alt文本"/>'  # 仅 alt 无障碍文本作图注（无 figcaption/邻位）
               '</body></html>')
        fd, epub = tempfile.mkstemp(suffix=".epub")
        os.close(fd)
        with _zip.ZipFile(epub, "w") as z:
            z.writestr("mimetype", "application/epub+zip", compress_type=_zip.ZIP_STORED)
            z.writestr("META-INF/container.xml", container)
            z.writestr("OEBPS/content.opf", opf)
            z.writestr("OEBPS/nav.xhtml", nav)
            z.writestr("OEBPS/ch1.xhtml", ch1)
            z.writestr("OEBPS/ch2.xhtml", ch2)
            z.writestr("OEBPS/img/fig1.png", big1)
            z.writestr("OEBPS/img/fig2.png", big2)
            z.writestr("OEBPS/img/fig3.png", _png(220, 260))
            z.writestr("OEBPS/img/fig5.png", _png(250, 250))
            z.writestr("OEBPS/img/small.png", small)
        fr = extract_figures(epub)
        assert fr["fmt"] == "epub", "EPUB 应识别为 epub，实际 %s" % fr["fmt"]
        all_figs = [fg for s in fr["sections"] for fg in s["figures"]]
        # fig1+fig2（figcaption）、fig3（邻位图注）、svg（矢量+邻位图注）抽中；
        # 装饰小图(small)被 min_px 过滤 → 共 4 张
        assert len(all_figs) == 5, "应抽中 5 张（2 figcaption + 1 邻位栅格 + 1 内联SVG + 1 alt兜底；装饰小图被过滤），实际 %d" % len(all_figs)
        caps = " ".join(f["caption"] for f in all_figs)
        assert "图1.1" in caps and "示意图" in caps, "图注1 未命中：%r" % caps
        assert "图2.3" in caps and "流程图" in caps, "图注2 未命中：%r" % caps
        assert "图3" in caps and "矢量" in caps, "SVG 图注未命中：%r" % caps
        assert "图4" in caps and "邻位" in caps, "邻位图注未命中：%r" % caps
        assert "图5" in caps and "alt" in caps, "alt 兜底图注未命中：%r" % caps
        # 章归属：第一章 / 第二章
        chapters = {s["chapter"] for s in fr["sections"] if s["figures"]}
        assert chapters == {"第一章", "第二章"}, "章归属异常：%s" % chapters
        # SVG 分支应被抽出且标记 ext=svg（验证不被尺寸过滤误杀）
        svg_figs = [f for f in all_figs if f["ext"] == "svg"]
        assert len(svg_figs) == 1, "内联 SVG 应抽中 1 张，实际 %d" % len(svg_figs)
        for f in svg_figs:
            assert f["width"] == 300 and f["height"] == 200, "SVG 尺寸未读：%s" % f
        for f in all_figs:
            if f["ext"] == "svg":
                continue
            assert f["width"] >= 200 and f["height"] >= 200, "大图尺寸异常：%s" % f
        os.unlink(epub)
        # DRM 加密 EPUB：内容文档被加密 → 应抛清晰异常而非产出乱码
        enc = ('<?xml version="1.0"?>\n<encryption xmlns="http://www.w3.org/2001/04/xmlenc#">'
               '<EncryptedData><CipherReference URI="OEBPS/ch1.xhtml"/></EncryptedData></encryption>')
        fd, depub = tempfile.mkstemp(suffix=".epub")
        os.close(fd)
        with _zip.ZipFile(depub, "w") as z:
            z.writestr("mimetype", "application/epub+zip", compress_type=_zip.ZIP_STORED)
            z.writestr("META-INF/container.xml", container)
            z.writestr("META-INF/encryption.xml", enc)
            z.writestr("OEBPS/content.opf", opf)
            z.writestr("OEBPS/ch1.xhtml", ch1)
        try:
            extract_figures(depub)
            assert False, "DRM EPUB 应抛异常"
        except RuntimeError as e:
            assert "DRM" in str(e), "DRM 异常信息应含 DRM：%r" % e
        os.unlink(depub)
        # 仅混淆字体（非内容加密）→ 不应抛 DRM 异常，可正常抽
        enc_font = ('<?xml version="1.0"?>\n<encryption xmlns="http://www.w3.org/2001/04/xmlenc#">'
                    '<EncryptedData><CipherReference URI="OEBPS/fonts/body.otf"/></EncryptedData></encryption>')
        fd, fepub = tempfile.mkstemp(suffix=".epub")
        os.close(fd)
        with _zip.ZipFile(fepub, "w") as z:
            z.writestr("mimetype", "application/epub+zip", compress_type=_zip.ZIP_STORED)
            z.writestr("META-INF/container.xml", container)
            z.writestr("META-INF/encryption.xml", enc_font)
            z.writestr("OEBPS/content.opf", opf)
            z.writestr("OEBPS/nav.xhtml", nav)
            z.writestr("OEBPS/ch1.xhtml", ch1)
        fr2 = extract_figures(fepub)
        assert fr2["fmt"] == "epub", "仅混淆字体的 EPUB 不应被判 DRM"
        os.unlink(fepub)
    except ImportError:
        print("  (skipped EPUB 插图用例：pymupdf/bs4 缺失)")

    # --- 用例11：SE 式 toc.xhtml（无 nav.xhtml）+ 描述性英文 alt 召回 ---
    # 复现 Flatland 真书坑：Standard Ebooks 旧版用 epub/toc.xhtml 而非 nav.xhtml；
    # 插图无 <figcaption>，说明写在 <img alt>（英文描述，无 图/Fig 正则命中）。
    try:
        import pymupdf as fitz
        import zipfile as _zip
        def _png(w, h):
            px = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, w, h), False)
            return px.tobytes("png")
        container = ('<?xml version="1.0"?>\n'
                     '<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">'
                     '<rootfiles><rootfile full-path="OEBPS/content.opf" '
                     'media-type="application/oebps-package+xml"/></rootfiles></container>')
        # 注意：manifest 无 nav 项，TOC 在 toc.xhtml（SE 旧版风格）
        opf = ('<?xml version="1.0"?>\n<package xmlns="http://www.idpf.org/2007/opf" version="3.0">'
               '<manifest>'
               '<item id="ch1" href="ch1.xhtml" media-type="application/xhtml+xml"/>'
               '<item id="ch2" href="ch2.xhtml" media-type="application/xhtml+xml"/>'
               '<item id="toc" href="toc.xhtml" media-type="application/xhtml+xml"/>'
               '<item id="f1" href="img/fig1.png" media-type="image/png"/>'
               '<item id="logo" href="img/logo.png" media-type="image/png"/>'
               '</manifest><spine><itemref idref="ch1"/><itemref idref="ch2"/></spine></package>')
        toc = ('<?xml version="1.0"?>\n<html xmlns="http://www.w3.org/1999/xhtml" '
               'xmlns:epub="http://www.idpf.org/2007/ops"><body epub:type="frontmatter">'
               '<nav epub:type="toc"><ol>'
               '<li><a href="ch1.xhtml">I: Of the Nature of Flatland</a></li>'
               '<li><a href="ch2.xhtml">II: Of the Climate and Houses</a></li>'
               '</ol></nav></body></html>')
        ch1 = ('<?xml version="1.0"?>\n<html xmlns="http://www.w3.org/1999/xhtml"><body>'
               '<h1>I: Of the Nature of Flatland</h1>'
               # 无 figcaption；说明写进 alt（英文描述，无 图/Fig 正则命中）→ 应被描述性 alt 兜底召回
               '<img src="img/fig1.png" alt="A diagram of three triangles, points downwards."/>'
               '</body></html>')
        ch2 = ('<?xml version="1.0"?>\n<html xmlns="http://www.w3.org/1999/xhtml"><body>'
               '<h1>II: Of the Climate and Houses</h1>'
               # 装饰性 logo（alt 纯装饰词）→ 不应当图注进训练
               '<img src="img/logo.png" alt="logo"/>'
               '</body></html>')
        fd, epub = tempfile.mkstemp(suffix=".epub")
        os.close(fd)
        with _zip.ZipFile(epub, "w") as z:
            z.writestr("mimetype", "application/epub+zip", compress_type=_zip.ZIP_STORED)
            z.writestr("META-INF/container.xml", container)
            z.writestr("OEBPS/content.opf", opf)
            z.writestr("OEBPS/toc.xhtml", toc)
            z.writestr("OEBPS/ch1.xhtml", ch1)
            z.writestr("OEBPS/ch2.xhtml", ch2)
            z.writestr("OEBPS/img/fig1.png", _png(240, 240))
            z.writestr("OEBPS/img/logo.png", _png(240, 240))
        fr = extract_figures(epub)
        # 章归属：靠 toc.xhtml 解析（非 nav.xhtml），不应是 全书/全文
        chapters = {s["chapter"] for s in fr["sections"] if s["figures"]}
        assert chapters == {"I: Of the Nature of Flatland"}, "toc.xhtml 章节归属失败：%s" % chapters
        # 描述性英文 alt 被召回；装饰性 logo 的 alt 被拒（caption 空）
        caps_by_ch = {s["chapter"]: [f["caption"] for f in s["figures"]] for s in fr["sections"] if s["figures"]}
        diag_caps = caps_by_ch.get("I: Of the Nature of Flatland", [])
        assert any("diagram of three triangles" in c for c in diag_caps), "描述性英文 alt 未召回：%s" % diag_caps
        logo_caps = caps_by_ch.get("II: Of the Climate and Houses", [])
        assert all(c == "" for c in logo_caps), "装饰性 alt 不应当图注：%s" % logo_caps
        os.unlink(epub)
    except ImportError:
        print("  (skipped EPUB toc.xhtml 用例：pymupdf/bs4 缺失)")

    print("  selftest passed（大纲章含节 / 无大纲多章小说 / 扫描版标记 / DOCX / MOBI / "
          "FB2 / CBZ / PDF插图抽取 / EPUB插图抽取 / 防御 re-split）。")


if __name__ == "__main__":
    if "--selftest" in sys.argv:
        selftest()
    else:
        if len(sys.argv) < 2:
            print("usage: python book_formats.py <file> [--selftest]")
            sys.exit(1)
        r = extract(sys.argv[1])
        print(f"fmt={r['fmt']} pages={r['page_count']} book_type={r['book_type']} "
              f"sections={len(r['sections'])} needs_ocr={r['needs_ocr']} meta={r['meta']}")
        for i, s in enumerate(r["sections"][:3], 1):
            print(f"--- [{i}] 章={s['chapter']} 节={s['section']} ({len(s['text'])} chars) ---\n"
                  f"{s['text'][:200]}")
