#!/usr/bin/env python3
"""
Fill Word template with provided content.
Replaces {{placeholder}} patterns with actual values.
"""

import sys
import json
import re
from pathlib import Path
from datetime import datetime

def replace_in_text(text: str, values: dict) -> str:
    """Replace {{placeholder}} patterns in text."""
    def replacer(match):
        key = match.group(1).strip()
        return str(values.get(key, match.group(0)))  # Keep original if not found
    
    return re.sub(r'\{\{([^}]+)\}\}', replacer, text)

def fill_template(docx_path: str, output_path: str, values: dict) -> dict:
    """Fill template with values and save to output path."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}
    
    try:
        doc = Document(docx_path)
        
        # Replace in paragraphs
        for para in doc.paragraphs:
            if '{{' in para.text:
                # Need to preserve formatting, so we work with runs
                full_text = para.text
                new_text = replace_in_text(full_text, values)
                if full_text != new_text:
                    # Clear and rebuild (simple approach, loses some formatting)
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = new_text
                    else:
                        para.add_run(new_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if '{{' in para.text:
                            new_text = replace_in_text(para.text, values)
                            for run in para.runs:
                                run.text = ""
                            if para.runs:
                                para.runs[0].text = new_text
                            else:
                                para.add_run(new_text)
        
        # Replace in headers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        if '{{' in para.text:
                            new_text = replace_in_text(para.text, values)
                            for run in para.runs:
                                run.text = ""
                            if para.runs:
                                para.runs[0].text = new_text
                            else:
                                para.add_run(new_text)
        
        # Replace in footers
        for section in doc.sections:
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        if '{{' in para.text:
                            new_text = replace_in_text(para.text, values)
                            for run in para.runs:
                                run.text = ""
                            if para.runs:
                                para.runs[0].text = new_text
                            else:
                                para.add_run(new_text)
        
        # Save
        doc.save(output_path)
        
        return {
            "success": True,
            "output_path": output_path,
            "placeholders_filled": len([k for k in values if values[k]])
        }
    
    except Exception as e:
        return {"error": str(e)}

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print(json.dumps({
            "error": "Usage: fill_template.py <input.docx> <output.docx> <values_json>"
        }))
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    try:
        values = json.loads(sys.argv[3])
    except json.JSONDecodeError as e:
        print(json.dumps({"error": f"Invalid JSON for values: {e}"}))
        sys.exit(1)
    
    result = fill_template(input_path, output_path, values)
    print(json.dumps(result, ensure_ascii=False, indent=2))
