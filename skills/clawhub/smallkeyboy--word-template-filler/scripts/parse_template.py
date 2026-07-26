#!/usr/bin/env python3
"""
Parse Word template and extract {{placeholder}} patterns.
Returns JSON with found placeholders.
"""

import sys
import json
import re
from pathlib import Path

def extract_placeholders(docx_path: str) -> dict:
    """Extract all {{placeholder}} patterns from a Word document."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}
    
    placeholders = set()
    
    try:
        doc = Document(docx_path)
        
        # Search in paragraphs
        for para in doc.paragraphs:
            matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
            placeholders.update(m.strip() for m in matches)
        
        # Search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = re.findall(r'\{\{([^}]+)\}\}', cell.text)
                    placeholders.update(m.strip() for m in matches)
        
        # Search in headers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
                        placeholders.update(m.strip() for m in matches)
        
        # Search in footers
        for section in doc.sections:
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
                        placeholders.update(m.strip() for m in matches)
        
        return {
            "success": True,
            "placeholders": sorted(list(placeholders)),
            "count": len(placeholders)
        }
    
    except Exception as e:
        return {"error": str(e)}

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: parse_template.py <docx_file>"}))
        sys.exit(1)
    
    result = extract_placeholders(sys.argv[1])
    print(json.dumps(result, ensure_ascii=False, indent=2))
