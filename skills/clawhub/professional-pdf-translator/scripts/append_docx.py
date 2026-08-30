"""
PDF Batch Translator — Markdown 译文追加写入 DOCX
用法: python append_docx.py --docx <path> --textfile <path> --page <num> [--highlights "regex:#RRGGBB" ...]
"""

import argparse
import os
import re
import json
import sys

try:
    from docx import Document
    from docx.shared import RGBColor, Pt
except ImportError:
    print(json.dumps({"error": "缺失 python-docx 依赖，请执行: python -m pip install python-docx"}, ensure_ascii=False))
    sys.exit(1)


def parse_highlights(highlight_args):
    """
    解析 --highlights 参数。
    格式: "regex:#RRGGBB"
    返回: [(compiled_regex, RGBColor), ...]
    """
    result = []
    if not highlight_args:
        return result

    for arg in highlight_args:
        if ':' not in arg:
            continue
        # 分割最后出现的冒号（因为正则本身可能包含冒号）
        parts = arg.rsplit(':', 1)
        if len(parts) != 2:
            continue
        pattern_str, color_str = parts
        try:
            color_str = color_str.strip().lstrip('#')
            r = int(color_str[0:2], 16)
            g = int(color_str[2:4], 16)
            b = int(color_str[4:6], 16)
            compiled = re.compile(pattern_str.strip())
            result.append((compiled, RGBColor(r, g, b)))
        except (ValueError, re.error):
            continue

    return result


def build_pattern(highlights):
    """
    构建用于 re.split 的复合正则表达式。
    优先级: **粗体** > *斜体* > 自定义高亮规则
    返回预编译的 pattern 对象。
    """
    patterns = [r'\*\*.*?\*\*', r'\*.*?\*']
    patterns.extend([h[0].pattern for h in highlights])
    combined = '|'.join(f'({p})' for p in patterns)
    return re.compile(combined)


def add_formatted_runs(paragraph, text, highlights):
    """将包含 Markdown 标记的文本解析并添加为多个 run（粗体/斜体/自定义高亮/普通）"""
    if not highlights:
        # 无自定义高亮时只需处理粗体和斜体
        pattern_str = r'(\*\*.*?\*\*)|(\*.*?\*)'
    else:
        hl_patterns = [h[0].pattern for h in highlights]
        parts = [r'\*\*.*?\*\*', r'\*.*?\*'] + hl_patterns
        pattern_str = '|'.join(f'({p})' for p in parts)

    compiled = re.compile(pattern_str)
    tokens = compiled.split(text)

    # split 结果: [non_match, group1, group2, ..., non_match, group1, group2, ...]
    # group 数量 = 2 (bold, italic) + len(highlights)
    num_groups = 2 + len(highlights)

    for i in range(0, len(tokens), num_groups + 1):
        # 非匹配文本
        plain = tokens[i]
        if plain:
            paragraph.add_run(plain)

        # 各捕获组
        for g in range(num_groups):
            idx = i + 1 + g
            if idx < len(tokens) and tokens[idx] is not None:
                token = tokens[idx]
                if not token:
                    continue

                if g == 0:
                    # **粗体**
                    run = paragraph.add_run(token[2:-2])
                    run.bold = True
                elif g == 1:
                    # *斜体*
                    run = paragraph.add_run(token[1:-1])
                    run.italic = True
                else:
                    # 自定义高亮 (g - 2 为 highlights 索引)
                    hl_idx = g - 2
                    if hl_idx < len(highlights):
                        _, color = highlights[hl_idx]
                        run = paragraph.add_run(token)
                        run.bold = True
                        run.font.color.rgb = color


def append_markdown_to_docx(docx_path, text, page_num, highlights):
    try:
        if os.path.exists(docx_path):
            doc = Document(docx_path)
        else:
            doc = Document()

        # 页码分隔线
        p_sep = doc.add_paragraph()
        r_sep = p_sep.add_run(f"\n--- 原文第 {page_num} 页 ---")
        r_sep.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        r_sep.font.italic = True

        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue

            # 处理标题
            m_heading = re.match(r'^(#{1,6})\s+(.*)', line)
            if m_heading:
                level = len(m_heading.group(1))
                heading_text = m_heading.group(2)
                heading = doc.add_heading(level=level)
                # 标题内也解析粗体/斜体格式
                add_formatted_runs(heading, heading_text, [])
                continue

            # 普通段落
            p = doc.add_paragraph()
            add_formatted_runs(p, line, highlights)

        doc.save(docx_path)
        return {"status": "success", "file": docx_path, "page": page_num}

    except Exception as e:
        return {"error": str(e)}


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="将 Markdown 译文追加写入 DOCX 文档")
    parser.add_argument("--docx", required=True, help="目标 DOCX 文档路径")
    parser.add_argument("--textfile", required=True, help="包含 Markdown 译文的 TXT 临时文件")
    parser.add_argument("--page", type=int, required=True, help="对应的原文页码")
    parser.add_argument("--highlights", nargs="*", default=[],
                        help='自定义高亮规则，格式: "regex:#RRGGBB"')
    args = parser.parse_args()

    if sys.stdout.encoding and sys.stdout.encoding.lower() != 'utf-8':
        sys.stdout.reconfigure(encoding='utf-8')

    with open(args.textfile, 'r', encoding='utf-8') as f:
        md_text = f.read()

    highlights = parse_highlights(args.highlights)
    result = append_markdown_to_docx(args.docx, md_text, page_num=args.page, highlights=highlights)

    # 写入成功后清理临时文件（异常时保留供调试）
    try:
        os.remove(args.textfile)
    except OSError:
        pass

    print(json.dumps(result, ensure_ascii=False, indent=2))
