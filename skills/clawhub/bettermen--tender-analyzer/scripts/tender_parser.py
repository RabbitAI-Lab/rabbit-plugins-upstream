#!/usr/bin/env python3
"""
Tender Document Parser — 投标文件智能解析引擎
支持 PDF/DOCX/XLSX 格式的结构化文本提取
"""

import sys
import json
import re
from pathlib import Path
from typing import Optional


def parse_pdf(filepath: str) -> dict:
    """解析PDF文件，提取文本和表格"""
    try:
        import fitz  # pymupdf
    except ImportError:
        return {"error": "pymupdf not installed. Run: pip install pymupdf"}

    doc = fitz.open(filepath)
    result = {
        "metadata": {
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "pages": len(doc),
            "format": doc.metadata.get("format", ""),
        },
        "pages": [],
        "tables": [],
    }

    for page_num, page in enumerate(doc, 1):
        text = page.get_text("text")
        result["pages"].append({
            "page": page_num,
            "text": text,
            "char_count": len(text),
        })

        # Extract tables
        tables = page.find_tables()
        for table in tables:
            rows = table.extract()
            if rows:
                result["tables"].append({
                    "page": page_num,
                    "rows": rows,
                    "row_count": len(rows),
                })

    doc.close()
    result["total_chars"] = sum(p["char_count"] for p in result["pages"])
    result["total_tables"] = len(result["tables"])
    return result


def parse_docx(filepath: str) -> dict:
    """解析DOCX文件，提取文本和表格"""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    doc = Document(filepath)
    result = {
        "metadata": {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "paragraphs": len(doc.paragraphs),
        },
        "paragraphs": [],
        "tables": [],
    }

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            result["paragraphs"].append({
                "index": i,
                "text": text,
                "style": para.style.name if para.style else "",
            })

    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            result["tables"].append({
                "table_index": table_idx,
                "rows": rows,
                "row_count": len(rows),
            })

    result["total_paragraphs"] = len(result["paragraphs"])
    result["total_tables"] = len(result["tables"])
    return result


def parse_xlsx(filepath: str) -> dict:
    """解析XLSX文件，提取工作表数据"""
    try:
        import openpyxl
    except ImportError:
        return {"error": "openpyxl not installed. Run: pip install openpyxl"}

    wb = openpyxl.load_workbook(filepath, data_only=True)
    result = {
        "metadata": {
            "sheets": wb.sheetnames,
        },
        "worksheets": [],
    }

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append([str(cell) if cell is not None else "" for cell in row])

        result["worksheets"].append({
            "name": sheet_name,
            "rows": rows[:500],  # Limit to 500 rows per sheet
            "row_count": min(ws.max_row, 500),
            "col_count": ws.max_column,
        })

    wb.close()
    return result


def extract_key_info(parsed_data: dict, file_format: str) -> dict:
    """从解析结果中提取关键招标信息"""
    full_text = ""

    if file_format == "pdf":
        full_text = "\n".join(p["text"] for p in parsed_data.get("pages", []))
    elif file_format == "docx":
        full_text = "\n".join(p["text"] for p in parsed_data.get("paragraphs", []))
    elif file_format == "xlsx":
        for ws in parsed_data.get("worksheets", []):
            full_text += f"\n--- {ws['name']} ---\n"
            for row in ws["rows"]:
                full_text += "\t".join(row) + "\n"

    key_info = {
        "project_name": _extract_pattern(full_text, r"项目名称[：:]\s*(.+)"),
        "project_number": _extract_pattern(full_text, r"(?:招标编号|项目编号)[：:]\s*(\S+)"),
        "budget": _extract_pattern(full_text, r"(?:预算|最高限价)[：:]\s*([\d.,]+)\s*万?"),
        "bidder": _extract_pattern(full_text, r"(?:招标人|采购人)[：:]\s*(.+)"),
        "bid_deadline": _extract_pattern(full_text, r"投标截止[：:]\s*(.+)"),
        "has_star_clause": "★" in full_text or "*" in full_text,
        "qualification_count": len(re.findall(r"资质|资格", full_text)),
        "score_keywords": len(re.findall(r"评分|分值|权重", full_text)),
    }

    return key_info


def _extract_pattern(text: str, pattern: str) -> Optional[str]:
    """正则提取单行匹配"""
    match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
    if match:
        return match.group(1).strip()
    return None


def main():
    if len(sys.argv) < 2:
        print("Usage: python tender_parser.py <filepath> [--json]")
        print("Supported formats: PDF, DOCX, XLSX")
        sys.exit(1)

    filepath = sys.argv[1]
    output_json = "--json" in sys.argv

    path = Path(filepath)
    suffix = path.suffix.lower()

    parsers = {
        ".pdf": (parse_pdf, "pdf"),
        ".docx": (parse_docx, "docx"),
        ".doc": (parse_docx, "docx"),
        ".xlsx": (parse_xlsx, "xlsx"),
        ".xls": (parse_xlsx, "xlsx"),
    }

    if suffix not in parsers:
        print(f"Unsupported format: {suffix}")
        sys.exit(1)

    parser, fmt = parsers[suffix]
    parsed = parser(filepath)

    if "error" in parsed:
        print(parsed["error"])
        sys.exit(1)

    # Extract key info
    key_info = extract_key_info(parsed, fmt)
    parsed["key_info"] = key_info

    if output_json:
        # Print summary JSON (without full text to avoid huge output)
        summary = {
            "metadata": parsed.get("metadata", {}),
            "key_info": key_info,
            "total_tables": parsed.get("total_tables", 0),
        }
        if fmt == "pdf":
            summary["pages"] = parsed["metadata"]["pages"]
            summary["total_chars"] = parsed["total_chars"]
        elif fmt == "docx":
            summary["paragraphs"] = parsed["total_paragraphs"]
        print(json.dumps(summary, ensure_ascii=False, indent=2))
    else:
        # Print full parsed data
        print(json.dumps(parsed, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
