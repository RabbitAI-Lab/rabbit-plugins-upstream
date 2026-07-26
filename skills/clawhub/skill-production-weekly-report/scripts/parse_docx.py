#!/usr/bin/env python3
"""Word文档解析脚本 - 提取生产周报数据"""
import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print(json.dumps({"status": "error", "message": "python-docx未安装"}))
    sys.exit(1)


def parse_docx(file_path: str) -> dict:
    """解析Word文档，提取生产周报相关数据"""
    try:
        doc = Document(file_path)
    except Exception as e:
        return {"status": "error", "message": f"无法打开文档: {str(e)}"}

    result = {
        "status": "success",
        "source": "word",
        "file": file_path,
        "data": {
            "overview": "",
            "production_metrics": [],
            "completed_items": [],
            "issues": [],
            "next_week_plan": [],
            "pending_matters": []
        }
    }

    current_section = None
    section_keywords = {
        "overview": ["概况", "概要", "总结"],
        "metrics": ["数据", "指标", "产量", "良率"],
        "completed": ["完成", "已完", "进展"],
        "issues": ["异常", "问题", "故障"],
        "plan": ["计划", "下周", "安排"],
        "pending": ["协调", "需支", "支持"]
    }

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 检测章节标题
        for section, keywords in section_keywords.items():
            if any(kw in text for kw in keywords) and len(text) < 20:
                current_section = section
                result["data"][section] = [] if isinstance(result["data"][section], list) else result["data"][section]
                break

        # 提取表格数据
        if para.style.name.startswith("Table"):
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        # 尝试解析为指标
                        if len(cells) >= 2 and any(c in cells[0] for c in ["产量", "良率", "OEE", "交付"]):
                            result["data"]["production_metrics"].append({
                                "name": cells[0],
                                "value": cells[1] if len(cells) > 1 else "",
                                "unit": cells[2] if len(cells) > 2 else ""
                            })

        # 内容行处理
        if current_section and text and not any(kw in text for kw in ["概况", "数据", "完成", "异常", "计划", "协调"]):
            if isinstance(result["data"][current_section], list):
                result["data"][current_section].append(text)
            elif current_section == "overview":
                result["data"]["overview"] += text + "\n"

    # 处理表格数据（收集所有表格）
    for table in doc.tables:
        for row in table.rows[1:]:  # 跳过表头
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2 and cells[0]:
                result["data"]["production_metrics"].append({
                    "name": cells[0],
                    "value": cells[1],
                    "unit": cells[2] if len(cells) > 2 else ""
                })

    return result


def main():
    parser = argparse.ArgumentParser(description="解析Word文档提取生产周报数据")
    parser.add_argument("--file", required=True, help="Word文档路径")
    args = parser.parse_args()

    if not Path(args.file).exists():
        print(json.dumps({"status": "error", "message": f"文件不存在: {args.file}"}))
        sys.exit(1)

    result = parse_docx(args.file)
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
