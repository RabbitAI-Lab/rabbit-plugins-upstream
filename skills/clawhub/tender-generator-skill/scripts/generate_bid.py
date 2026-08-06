#!/usr/bin/env python3
"""
标书自动生成脚本 - 完善版
根据用户提供的信息和模板文件，自动生成全套标书文档。
支持：TMA.docx 集成、更多文档类型、模板占位符、打包功能、字体统一、内容丰富
"""

import os
import sys
import json
import argparse
import zipfile
from pathlib import Path
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import shutil

# 统一字体设置
def set_font(run, font_name='宋体', font_size=12):
    """设置字体和大小"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)

def set_heading_font(run, font_name='黑体', font_size=16):
    """设置标题字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = True

def apply_document_style(doc):
    """应用统一的文档样式"""
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    font.size = Pt(12)
    
    # 设置段落间距为 0，行间距为 1.5 倍
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.5

def load_info_from_json(json_path):
    """从 JSON 文件加载信息"""
    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def extract_content_from_tma(tma_path):
    """从 TMA.docx 提取技术内容"""
    try:
        doc = docx.Document(tma_path)
        content = []
        
        # 提取所有段落
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text.strip())
        
        # 提取表格内容
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_content.append(" | ".join(row_text))
            if table_content:
                content.append("\n".join(table_content))
        
        return "\n\n".join(content)
    except Exception as e:
        print(f"警告：无法读取 TMA.docx: {e}")
        return None

def replace_placeholders_in_text(text, info):
    """替换文本中的占位符"""
    if not text:
        return text
    
    for key, value in info.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in text:
            text = text.replace(placeholder, str(value) if value else "")
    
    return text

def replace_placeholders_in_doc(doc, info):
    """替换文档中的所有占位符"""
    # 替换段落
    for para in doc.paragraphs:
        if para.text:
            new_text = replace_placeholders_in_text(para.text, info)
            if new_text != para.text:
                para.text = new_text
    
    # 替换表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    new_text = replace_placeholders_in_text(cell.text, info)
                    if new_text != cell.text:
                        cell.text = new_text

def create_cover_document(output_path, info):
    """创建封面文档"""
    doc = docx.Document()
    apply_document_style(doc)
    
    # 添加标题（黑体，小二）
    title = doc.add_heading(info.get("项目名称", "项目名称"), 0)
    for run in title.runs:
        set_heading_font(run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 添加项目信息表格
    info_table = doc.add_table(rows=9, cols=2)
    info_table.style = 'Table Grid'
    
    # 设置表格样式
    for row in info_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_font(run, '宋体', 12)
    
    rows_data = [
        ("项目名称", info.get("项目名称", "")),
        ("项目编号", info.get("项目编号", "")),
        ("招标人", info.get("招标人", "")),
        ("投标截止日期", info.get("投标截止日期", "")),
        ("投标人名称", info.get("公司全称", "")),
        ("法定代表人", info.get("法定代表人", "")),
        ("授权代表", info.get("授权代表", "")),
        ("联系电话", info.get("授权代表电话", "")),
        ("日期", ""),
    ]
    
    for i, (label, value) in enumerate(rows_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value) if value else ""
    
    # 添加底部说明
    doc.add_paragraph()
    bottom_para = doc.add_paragraph()
    bottom_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = bottom_para.add_run("（此页 intentionally 空白）")
    set_font(run, '宋体', 10)
    
    doc.save(output_path)
    print(f"✓ 已生成：{output_path}")
    return True

def create_commitment_document(output_path, info):
    """创建承诺书文档"""
    doc = docx.Document()
    apply_document_style(doc)
    
    # 标题
    title = doc.add_heading("承 诺 书", 0)
    for run in title.runs:
        set_heading_font(run, '黑体', 20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 正文内容（丰富版）
    content = f"""{info.get("公司全称", "")}（以下简称"我方"）作为本次投标的投标人，在此郑重承诺：

一、关于投标文件的真实性承诺

1. 我方已仔细阅读并完全理解招标文件（项目编号：{info.get("项目编号", "")}）的全部内容，包括所有附件、补充文件及澄清文件。

2. 我方投标文件中提供的所有资料、文件、证明、业绩材料等都是真实、准确、完整、有效的，不存在任何虚假记载、误导性陈述或重大遗漏。

3. 如我方提供虚假材料，愿意承担由此产生的一切法律责任，包括但不限于取消投标资格、没收投标保证金、列入不良行为记录名单等。

二、关于履约能力的承诺

4. 我方具备履行本合同所需的全部资质、许可、人员、设备和技术能力。

5. 我方承诺按照招标文件和合同约定的时间、质量、数量要求提供货物和服务。

6. 如我方中标，将在收到中标通知书后 30 日内与招标人签订合同，并严格按照合同约定履行义务。

三、关于服务质量的承诺

7. 我方承诺提供的货物和服务符合国家现行标准、行业标准和招标文件要求。

8. 我方承诺提供完善的售后服务，包括技术支持、维修保养、备件供应等。

9. 我方承诺建立专门的项目服务团队，确保项目顺利实施。

四、关于廉洁从业的承诺

10. 我方承诺在投标和履约过程中，严格遵守国家法律法规，坚持诚信经营，不进行任何形式的不正当竞争。

11. 我方承诺不向招标人工作人员、评标委员会成员行贿或提供其他不正当利益。

12. 我方承诺自觉接受有关部门的监督和检查。

五、其他承诺

13. 我方承诺在投标有效期内不修改、撤销投标文件。

14. 如我方中标，承诺不转包、不违法分包。

15. 我方愿意承担因违反本承诺书内容而产生的一切法律责任和经济责任。

特此承诺！

投标人名称（盖章）：{info.get("公司全称", "")}
法定代表人或授权代表（签字）：{info.get("授权代表", "")}
联系电话：{info.get("授权代表电话", "")}
日    期：2026 年    月    日
"""
    
    # 添加段落并设置字体
    para = doc.add_paragraph(content)
    for run in para.runs:
        set_font(run, '宋体', 12)
    
    # 调整段落格式
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    
    doc.save(output_path)
    print(f"✓ 已生成：{output_path}")
    return True

def create_business_license_doc(output_path, info):
    """创建营业执照信息文档"""
    doc = docx.Document()
    apply_document_style(doc)
    
    # 标题
    title = doc.add_heading("营业执照信息", 0)
    for run in title.runs:
        set_heading_font(run, '黑体', 20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 添加说明文字
    intro_para = doc.add_paragraph("以下是我方营业执照登记的主要信息：")
    for run in intro_para.runs:
        set_font(run, '宋体', 12)
    
    doc.add_paragraph()
    
    # 信息表格
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    
    rows_data = [
        ("公司名称", info.get("公司全称", "")),
        ("统一社会信用代码", info.get("统一社会信用代码", "")),
        ("注册地址", info.get("公司地址", "")),
        ("法定代表人", info.get("法定代表人", "")),
        ("成立日期", "（详见营业执照原件）"),
    ]
    
    for i, (label, value) in enumerate(rows_data):
        cell1 = info_table.rows[i].cells[0]
        cell2 = info_table.rows[i].cells[1]
        cell1.text = label
        cell2.text = str(value) if value else ""
        
        # 设置单元格字体
        for para in cell1.paragraphs:
            for run in para.runs:
                set_font(run, '黑体', 12)
        for para in cell2.paragraphs:
            for run in para.runs:
                set_font(run, '宋体', 12)
    
    doc.add_paragraph()
    
    # 添加说明
    note_para = doc.add_paragraph("注：本表信息以营业执照原件为准，复印件附后。")
    for run in note_para.runs:
        set_font(run, '宋体', 10)
    note_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.save(output_path)
    print(f"✓ 已生成：{output_path}")
    return True

def create_quote_document(output_path, info):
    """创建报价单文档"""
    doc = docx.Document()
    apply_document_style(doc)
    
    # 标题
    title = doc.add_heading("报 价 单", 0)
    for run in title.runs:
        set_heading_font(run, '黑体', 20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 项目信息
    info_para = doc.add_paragraph(f"""
项目名称：{info.get("项目名称", "")}
项目编号：{info.get("项目编号", "")}
投标人名称：{info.get("公司全称", "")}
报价日期：2026 年    月    日
""")
    for run in info_para.runs:
        set_font(run, '宋体', 12)
    
    doc.add_paragraph()
    
    # 报价表格
    quote_table = doc.add_table(rows=8, cols=4)
    quote_table.style = 'Table Grid'
    
    # 表头
    headers = ["序号", "项目名称", "规格型号/说明", "金额（元）"]
    for i, header in enumerate(headers):
        quote_table.rows[0].cells[i].text = header
        for run in quote_table.rows[0].cells[i].paragraphs[0].runs:
            set_font(run, '黑体', 12)
    
    # 报价内容
    quote_items = [
        ("1", "技术服务费", "系统设计、开发、实施、培训等", ""),
        ("2", "设备费", "硬件设备、软件许可等", ""),
        ("3", "材料费", "项目实施所需材料", ""),
        ("4", "运输保险费", "设备运输及保险", ""),
        ("5", "安装调试费", "设备安装、调试、验收", ""),
        ("6", "培训费", "人员培训", ""),
        ("7", "其他费用", "", ""),
    ]
    
    for i, item in enumerate(quote_items):
        for j, value in enumerate(item):
            quote_table.rows[i+1].cells[j].text = value
            for run in quote_table.rows[i+1].cells[j].paragraphs[0].runs:
                set_font(run, '宋体', 12)
    
    # 总计行
    total_row = quote_table.add_row()
    total_row.cells[0].text = ""
    total_row.cells[1].text = ""
    total_row.cells[2].text = "总计（含税）"
    total_row.cells[3].text = info.get("总报价", "待定")
    
    for run in total_row.cells[2].paragraphs[0].runs:
        set_font(run, '黑体', 12)
    for run in total_row.cells[3].paragraphs[0].runs:
        set_font(run, '宋体', 12)
        run.font.bold = True
    
    doc.add_paragraph()
    
    # 报价说明
    note_content = """
报价说明：
1. 本报价为含税全包价，包含完成本项目所需的一切费用。
2. 报价有效期：自投标之日起 90 天。
3. 付款方式：按合同约定执行。
4. 交货/服务期限：按合同约定执行。
"""
    note_para = doc.add_paragraph(note_content)
    for run in note_para.runs:
        set_font(run, '宋体', 11)
    
    doc.save(output_path)
    print(f"✓ 已生成：{output_path}")
    return True

def create_deviation_document(output_path, info):
    """创建偏离表文档"""
    doc = docx.Document()
    apply_document_style(doc)
    
    # 标题
    title = doc.add_heading("技术/商务偏离表", 0)
    for run in title.runs:
        set_heading_font(run, '黑体', 20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 项目信息
    info_para = doc.add_paragraph(f"""
项目名称：{info.get("项目名称", "")}
项目编号：{info.get("项目编号", "")}
投标人名称：{info.get("公司全称", "")}
""")
    for run in info_para.runs:
        set_font(run, '宋体', 12)
    
    doc.add_paragraph()
    
    # 偏离表内容
    content = """
我方对招标文件（项目编号：{tender_number}）的响应情况如下：

一、技术部分

经仔细研究和比对，我方承诺：

1. 完全响应招标文件的所有技术要求，无任何技术偏离。

2. 所提供的产品/服务均符合国家现行标准、行业标准和招标文件规定的技术参数。

3. 如招标文件中有不明确的技术要求，我方将按照国家标准和行业最高标准执行。

二、商务部分

经仔细研究和比对，我方承诺：

1. 完全响应招标文件的所有商务条款，无任何商务偏离。

2. 接受招标文件规定的付款方式、交货期限、质保期等商务条件。

3. 承诺按照招标文件要求提供相应的售后服务和保障措施。

三、特别说明

1. 如我方中标，将严格按照招标文件和合同约定履行义务。

2. 如有任何偏离，将在合同签订前以书面形式详细说明，并取得招标人的书面同意。

3. 本偏离表作为投标文件的组成部分，与投标文件具有同等法律效力。

投标人名称（盖章）：{company_name}
法定代表人或授权代表（签字）：{authorized_rep}
日    期：2026 年    月    日
""".format(
        tender_number=info.get("项目编号", ""),
        company_name=info.get("公司全称", ""),
        authorized_rep=info.get("授权代表", "")
    )
    
    para = doc.add_paragraph(content)
    for run in para.runs:
        set_font(run, '宋体', 12)
    para.paragraph_format.line_spacing = 1.5
    
    doc.save(output_path)
    print(f"✓ 已生成：{output_path}")
    return True

def create_qualification_document(output_path, info):
    """创建资格能力证明文件"""
    doc = docx.Document()
    apply_document_style(doc)
    
    # 标题
    title = doc.add_heading("资格能力证明文件", 0)
    for run in title.runs:
        set_heading_font(run, '黑体', 20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 目录
    toc_content = """
目  录

一、基本情况…………………………………………………1
二、银行信息…………………………………………………2
三、授权代表信息……………………………………………3
四、资质证明…………………………………………………4
五、业绩材料…………………………………………………5
"""
    toc_para = doc.add_paragraph(toc_content)
    for run in toc_para.runs:
        set_font(run, '宋体', 12)
    
    doc.add_page_break()
    
    # 一、基本情况
    section1 = doc.add_heading("一、基本情况", 1)
    for run in section1.runs:
        set_heading_font(run, '黑体', 16)
    
    content1 = f"""
1.1 公司名称：{info.get("公司全称", "")}

1.2 统一社会信用代码：{info.get("统一社会信用代码", "")}

1.3 注册地址：{info.get("公司地址", "")}

1.4 法定代表人：{info.get("法定代表人", "")}

1.5 公司类型：（详见营业执照）

1.6 注册资本：（详见营业执照）

1.7 经营范围：（详见营业执照）

我方是一家依法注册成立的合法企业，具备独立承担民事责任的能力。
"""
    para1 = doc.add_paragraph(content1)
    for run in para1.runs:
        set_font(run, '宋体', 12)
    para1.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()
    
    # 二、银行信息
    section2 = doc.add_heading("二、银行信息", 1)
    for run in section2.runs:
        set_heading_font(run, '黑体', 16)
    
    content2 = f"""
2.1 开户银行：{info.get("开户银行", "")}

2.2 银行账号：{info.get("银行账号", "")}

2.3 联行号：{info.get("开户银行联行号", "")}

我方账户信息真实有效，具备履行合同所需的资金结算能力。
"""
    para2 = doc.add_paragraph(content2)
    for run in para2.runs:
        set_font(run, '宋体', 12)
    para2.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()
    
    # 三、授权代表信息
    section3 = doc.add_heading("三、授权代表信息", 1)
    for run in section3.runs:
        set_heading_font(run, '黑体', 16)
    
    content3 = f"""
3.1 授权代表姓名：{info.get("授权代表", "")}

3.2 联系电话：{info.get("授权代表电话", "")}

3.3 职务：（详见授权委托书）

我方授权代表有权代表我方处理本次投标及后续合同履行的相关事宜。
"""
    para3 = doc.add_paragraph(content3)
    for run in para3.runs:
        set_font(run, '宋体', 12)
    para3.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()
    
    # 四、资质证明
    section4 = doc.add_heading("四、资质证明", 1)
    for run in section4.runs:
        set_heading_font(run, '黑体', 16)
    
    content4 = """
我方具备履行本合同所需的全部资质，包括但不限于：

4.1 营业执照（复印件附后）

4.2 相关行业资质证书（如有）

4.3 质量管理体系认证证书（如有）

4.4 其他相关资质证明（如有）

以上资质证书均在有效期内，真实有效。
"""
    para4 = doc.add_paragraph(content4)
    for run in para4.runs:
        set_font(run, '宋体', 12)
    para4.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()
    
    # 五、业绩材料
    section5 = doc.add_heading("五、业绩材料", 1)
    for run in section5.runs:
        set_heading_font(run, '黑体', 16)
    
    content5 = """
我方具有类似项目的实施经验，主要业绩包括：

（此处列出相关业绩，附合同复印件或证明材料）

以上业绩真实有效，可查证。
"""
    para5 = doc.add_paragraph(content5)
    for run in para5.runs:
        set_font(run, '宋体', 12)
    para5.paragraph_format.line_spacing = 1.5
    
    doc.save(output_path)
    print(f"✓ 已生成：{output_path}")
    return True

def create_technical_document(output_path, info, tma_content):
    """创建技术文件文档"""
    doc = docx.Document()
    apply_document_style(doc)
    
    # 封面
    title = doc.add_heading("技术实施方案", 0)
    for run in title.runs:
        set_heading_font(run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    rows_data = [
        ("项目名称", info.get("项目名称", "")),
        ("项目编号", info.get("项目编号", "")),
        ("投标人名称", info.get("公司全称", "")),
        ("日期", "2026 年    月    日"),
    ]
    
    for i, (label, value) in enumerate(rows_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value) if value else ""
        
        for para in info_table.rows[i].cells[0].paragraphs:
            for run in para.runs:
                set_font(run, '黑体', 12)
        for para in info_table.rows[i].cells[1].paragraphs:
            for run in para.runs:
                set_font(run, '宋体', 12)
    
    doc.add_page_break()
    
    # 目录
    toc_content = """
目  录

第一章 项目概述………………………………………………1
第二章 技术方案………………………………………………3
第三章 实施计划………………………………………………8
第四章 质量保证………………………………………………12
第五章 售后服务………………………………………………15
第六章 培训方案………………………………………………18
第七章 风险管理………………………………………………20
"""
    toc_para = doc.add_paragraph(toc_content)
    for run in toc_para.runs:
        set_font(run, '宋体', 12)
    
    doc.add_page_break()
    
    # 第一章 项目概述
    chapter1 = doc.add_heading("第一章 项目概述", 1)
    for run in chapter1.runs:
        set_heading_font(run, '黑体', 18)
    
    content1 = f"""
1.1 项目背景

{info.get("项目名称", "")}（项目编号：{info.get("项目编号", "")}）是由{info.get("招标人", "")}发起的重要项目。

1.2 项目目标

本项目旨在通过先进的技术方案，实现以下目标：

（1）提高效率和性能
（2）降低成本和能耗
（3）提升用户体验
（4）确保系统稳定性和安全性

1.3 项目范围

本项目的实施范围包括：

（1）系统设计
（2）设备供应
（3）安装调试
（4）人员培训
（5）售后服务

1.4 编制依据

本技术方案的编制依据包括：

（1）招标文件及其附件
（2）国家现行标准和规范
（3）行业最佳实践
（4）我方技术积累和经验
"""
    para1 = doc.add_paragraph(content1)
    for run in para1.runs:
        set_font(run, '宋体', 12)
    para1.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()
    
    # 第二章 技术方案
    chapter2 = doc.add_heading("第二章 技术方案", 1)
    for run in chapter2.runs:
        set_heading_font(run, '黑体', 18)
    
    # 如果有 TMA 内容，插入其中
    if tma_content:
        content2 = f"""
2.1 技术架构

{tma_content}

2.2 核心技术

我方将采用以下核心技术：

（1）先进的设计理念和架构
（2）成熟的技术路线和方案
（3）创新的技术亮点和优势
（4）完善的技术保障措施

2.3 技术优势

与同类方案相比，我方技术方案具有以下优势：

（1）技术先进性
（2）系统稳定性
（3）可扩展性
（4）经济性
"""
    else:
        content2 = """
2.1 技术架构

我方将采用先进、成熟、可靠的技术架构，确保系统的稳定性和可扩展性。

2.2 核心技术

我方将采用以下核心技术：

（1）先进的设计理念和架构
（2）成熟的技术路线和方案
（3）创新的技术亮点和优势
（4）完善的技术保障措施

2.3 技术优势

与同类方案相比，我方技术方案具有以下优势：

（1）技术先进性
（2）系统稳定性
（3）可扩展性
（4）经济性
"""
    
    para2 = doc.add_paragraph(content2)
    for run in para2.runs:
        set_font(run, '宋体', 12)
    para2.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()
    
    # 第三章 实施计划
    chapter3 = doc.add_heading("第三章 实施计划", 1)
    for run in chapter3.runs:
        set_heading_font(run, '黑体', 18)
    
    content3 = """
3.1 项目实施阶段

本项目实施分为以下阶段：

（1）准备阶段：项目启动、需求调研、方案设计
（2）实施阶段：系统开发、设备采购、安装调试
（3）验收阶段：系统测试、用户验收、正式交付
（4）运维阶段：售后服务、技术支持、持续优化

3.2 项目进度安排

| 阶段 | 时间节点 | 主要工作内容 |
|------|----------|--------------|
| 准备阶段 | 第 1-2 周 | 项目启动、需求调研 |
| 设计阶段 | 第 3-4 周 | 方案设计、评审 |
| 实施阶段 | 第 5-10 周 | 开发、采购、安装 |
| 验收阶段 | 第 11-12 周 | 测试、验收、交付 |
| 运维阶段 | 第 13 周起 | 售后服务 |

3.3 项目团队

我方将组建专业的项目团队，包括：

（1）项目经理：负责项目整体管理
（2）技术负责人：负责技术方案和实施
（3）工程师：负责具体实施工作
（4）质检员：负责质量检查
（5）安全员：负责安全管理
"""
    para3 = doc.add_paragraph(content3)
    for run in para3.runs:
        set_font(run, '宋体', 12)
    para3.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()
    
    # 第四章 质量保证
    chapter4 = doc.add_heading("第四章 质量保证", 1)
    for run in chapter4.runs:
        set_heading_font(run, '黑体', 18)
    
    content4 = """
4.1 质量目标

我方承诺：

（1）工程质量达到国家现行标准和要求
（2）一次验收合格率 100%
（3）客户满意度 95% 以上

4.2 质量管理体系

我方将严格执行 ISO9001 质量管理体系，确保项目质量。

4.3 质量控制措施

（1）事前控制：方案评审、材料检验
（2）事中控制：过程检查、测试验证
（3）事后控制：竣工验收、质量评估

4.4 质量保证承诺

我方承诺：

（1）提供质量保证期
（2）承担质量责任
（3）及时处理质量问题
"""
    para4 = doc.add_paragraph(content4)
    for run in para4.runs:
        set_font(run, '宋体', 12)
    para4.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()
    
    # 第五章 售后服务
    chapter5 = doc.add_heading("第五章 售后服务", 1)
    for run in chapter5.runs:
        set_heading_font(run, '黑体', 18)
    
    content5 = """
5.1 售后服务承诺

我方承诺提供以下售后服务：

（1）质保期：按合同约定
（2）响应时间：24 小时内响应
（3）解决时间：48 小时内解决一般问题
（4）定期巡检：每季度一次

5.2 售后服务内容

（1）技术支持：电话、邮件、现场支持
（2）维修保养：定期保养、故障维修
（3）备件供应：保证备件供应
（4）技术培训：免费培训操作人员

5.3 售后服务团队

我方将组建专门的售后服务团队，提供全天候服务。
"""
    para5 = doc.add_paragraph(content5)
    for run in para5.runs:
        set_font(run, '宋体', 12)
    para5.paragraph_format.line_spacing = 1.5
    
    doc.save(output_path)
    print(f"✓ 已生成：{output_path}")
    return True

def create_power_of_attorney(output_path, info):
    """创建授权委托书"""
    doc = docx.Document()
    apply_document_style(doc)
    
    # 标题
    title = doc.add_heading("授 权 委 托 书", 0)
    for run in title.runs:
        set_heading_font(run, '黑体', 20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    content = f"""
{info.get("招标人", "")}：

我方 {info.get("公司全称", "")} 授权 {info.get("授权代表", "")}（身份证号：__________）为我方合法代理人，参加贵方组织的 {info.get("项目名称", "")}（项目编号：{info.get("项目编号", "")}）的投标活动。

该代理人在投标过程中所签署的所有文件和处理与之相关的一切事务，我方均予以承认并承担全部法律责任。

代理人无转委托权。

本授权委托书自签发之日起生效，有效期至本次投标活动结束或合同签订之日止。

特此授权。

投标人（盖章）：{info.get("公司全称", "")}
法定代表人（签字）：{info.get("法定代表人", "")}
授权代表（签字）：{info.get("授权代表", "")}
日    期：2026 年    月    日
"""
    
    para = doc.add_paragraph(content)
    for run in para.runs:
        set_font(run, '宋体', 12)
    para.paragraph_format.line_spacing = 1.5
    
    doc.save(output_path)
    print(f"✓ 已生成：{output_path}")
    return True

def create_invoice_info_document(output_path, info):
    """创建开票信息文档"""
    doc = docx.Document()
    apply_document_style(doc)
    
    # 标题
    title = doc.add_heading("开票信息", 0)
    for run in title.runs:
        set_heading_font(run, '黑体', 20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    content = f"""
我方开票信息如下：

一、基本信息

公司名称：{info.get("公司全称", "")}
统一社会信用代码：{info.get("统一社会信用代码", "")}
地址：{info.get("公司地址", "")}

二、银行账户信息

开户银行：{info.get("开户银行", "")}
银行账号：{info.get("银行账号", "")}
联行号：{info.get("开户银行联行号", "")}

三、联系方式

联系人：{info.get("授权代表", "")}
联系电话：{info.get("授权代表电话", "")}

以上信息真实有效，如有变更将及时书面通知。
"""
    
    para = doc.add_paragraph(content)
    for run in para.runs:
        set_font(run, '宋体', 12)
    para.paragraph_format.line_spacing = 1.5
    
    doc.save(output_path)
    print(f"✓ 已生成：{output_path}")
    return True

def generate_all_documents(info, output_dir, template_dir, use_tma):
    """生成所有标书文档"""
    # 创建输出目录
    os.makedirs(output_dir, exist_ok=True)
    
    # 提取 TMA 内容
    tma_content = None
    if use_tma:
        # TMA.docx 应该在脚本所在目录下，而不是 templates 目录的父目录
        script_dir = os.path.dirname(os.path.abspath(__file__))
        project_root = os.path.dirname(script_dir)
        tma_path = os.path.join(project_root, "TMA.docx")
        if os.path.exists(tma_path):
            print("正在提取 TMA.docx 内容...")
            tma_content = extract_content_from_tma(tma_path)
            if tma_content:
                print(f"✓ 成功提取 TMA.docx 内容（{len(tma_content)} 字符）")
            else:
                print("⚠ TMA.docx 内容为空")
        else:
            print("⚠ TMA.docx 不存在，将使用默认技术内容")
    
    # 定义要生成的文档
    documents = [
        ("00-封面.docx", "cover"),
        ("01-承诺书.docx", "commitment"),
        ("02-营业执照信息.docx", "business_license"),
        ("03-报价单.docx", "quote"),
        ("04-偏离表.docx", "deviation"),
        ("05-资格证明.docx", "qualification"),
        ("06-技术文件.docx", "technical"),
        ("07-授权委托书.docx", "power_of_attorney"),
        ("08-开票信息.docx", "invoice_info"),
    ]
    
    print("\n开始生成标书文档...")
    print("=" * 60)
    
    for filename, doc_type in documents:
        output_path = os.path.join(output_dir, filename)
        
        if doc_type == "cover":
            create_cover_document(output_path, info)
        elif doc_type == "commitment":
            create_commitment_document(output_path, info)
        elif doc_type == "business_license":
            create_business_license_doc(output_path, info)
        elif doc_type == "quote":
            create_quote_document(output_path, info)
        elif doc_type == "deviation":
            create_deviation_document(output_path, info)
        elif doc_type == "qualification":
            create_qualification_document(output_path, info)
        elif doc_type == "technical":
            create_technical_document(output_path, info, tma_content)
        elif doc_type == "power_of_attorney":
            create_power_of_attorney(output_path, info)
        elif doc_type == "invoice_info":
            create_invoice_info_document(output_path, info)
    
    print("=" * 60)
    print(f"\n✅ 所有文档已生成到：{output_dir}")
    
    # 打包文件
    zip_path = os.path.join(os.path.dirname(output_dir), f"标书文档包_{info.get('项目名称', '项目')}.zip")
    print(f"\n正在打包文档到：{zip_path}")
    
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for filename, _ in documents:
            file_path = os.path.join(output_dir, filename)
            if os.path.exists(file_path):
                zipf.write(file_path, filename)
                print(f"  ✓ 已添加：{filename}")
    
    print(f"\n📦 打包完成：{zip_path}")
    return True

def main():
    parser = argparse.ArgumentParser(description="标书自动生成脚本 - 完善版")
    parser.add_argument("--info", "-i", required=True, help="包含项目信息的 JSON 文件路径")
    parser.add_argument("--output", "-o", default="./output", help="输出目录")
    parser.add_argument("--template", "-t", default="./templates/默认", help="模板目录")
    parser.add_argument("--use-tma", action="store_true", help="使用 TMA.docx 作为技术文件内容源")
    
    args = parser.parse_args()
    
    # 加载信息
    if not os.path.exists(args.info):
        print(f"错误：信息文件 {args.info} 不存在")
        sys.exit(1)
    
    info = load_info_from_json(args.info)
    
    # 生成文档
    generate_all_documents(info, args.output, args.template, args.use_tma)

if __name__ == "__main__":
    main()