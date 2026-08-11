#!/usr/bin/env python3
"""
标书自动生成脚本 - 最终版
彻底修复空格问题：每个段落单独添加，不使用包含 \n\n 的长字符串
"""

import os
import sys
import json
import argparse
import zipfile
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name='宋体', font_size=12):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)

def set_heading_font(run, font_name='黑体', font_size=16):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = True

def set_para_spacing(para, space_before=Pt(0), space_after=Pt(0), line_spacing=1.5):
    para.paragraph_format.space_before = space_before
    para.paragraph_format.space_after = space_after
    para.paragraph_format.line_spacing = line_spacing

def apply_document_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5

def add_para(doc, text, font_size=12, font_name='宋体', bold=False, alignment=None, space_before=Pt(0), space_after=Pt(0)):
    """添加段落，每个段落单独添加，无空行"""
    para = doc.add_paragraph(text)
    for run in para.runs:
        set_font(run, font_name, font_size)
        if bold:
            run.font.bold = True
    set_para_spacing(para, space_before, space_after)
    if alignment:
        para.alignment = alignment
    return para

def add_heading(doc, text, font_size=18):
    """添加标题"""
    para = doc.add_paragraph(text)
    for run in para.runs:
        set_heading_font(run, '黑体', font_size)
    set_para_spacing(para, Pt(12), Pt(6))
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para

def create_commitment_document(output_path, info):
    """创建承诺书文档 - 每个段落单独添加"""
    doc = docx.Document()
    apply_document_style(doc)
    
    # 标题
    add_heading(doc, "承  诺  书", 20)
    
    # 第一段
    add_para(doc, f'{info.get("公司全称", "")}（以下简称"我方"）作为本次投标的投标人，在此郑重承诺：', space_before=Pt(12))
    
    # 一、关于投标文件的真实性承诺
    add_para(doc, "一、关于投标文件的真实性承诺", font_size=14, bold=True, space_before=Pt(12))
    add_para(doc, "1. 我方已仔细阅读并完全理解招标文件（项目编号：{}）的全部内容，包括所有附件、补充文件及澄清文件。".format(info.get("项目编号", "")))
    add_para(doc, "2. 我方投标文件中提供的所有资料、文件、证明、业绩材料等都是真实、准确、完整、有效的，不存在任何虚假记载、误导性陈述或重大遗漏。")
    add_para(doc, "3. 如我方提供虚假材料，愿意承担由此产生的一切法律责任，包括但不限于取消投标资格、没收投标保证金、列入不良行为记录名单等。")
    
    # 二、关于履约能力的承诺
    add_para(doc, "二、关于履约能力的承诺", font_size=14, bold=True, space_before=Pt(12))
    add_para(doc, "4. 我方具备履行本合同所需的全部资质、许可、人员、设备和技术能力。")
    add_para(doc, "5. 我方承诺按照招标文件和合同约定的时间、质量、数量要求提供货物和服务。")
    add_para(doc, "6. 如我方中标，将在收到中标通知书后 30 日内与招标人签订合同，并严格按照合同约定履行义务。")
    
    # 三、关于服务质量的承诺
    add_para(doc, "三、关于服务质量的承诺", font_size=14, bold=True, space_before=Pt(12))
    add_para(doc, "7. 我方承诺提供的货物和服务符合国家现行标准、行业标准和招标文件要求。")
    add_para(doc, "8. 我方承诺提供完善的售后服务，包括技术支持、维修保养、备件供应等。")
    add_para(doc, "9. 我方承诺建立专门的项目服务团队，确保项目顺利实施。")
    
    # 四、关于廉洁从业的承诺
    add_para(doc, "四、关于廉洁从业的承诺", font_size=14, bold=True, space_before=Pt(12))
    add_para(doc, "10. 我方承诺在投标和履约过程中，严格遵守国家法律法规，坚持诚信经营，不进行任何形式的不正当竞争。")
    add_para(doc, "11. 我方承诺不向招标人工作人员、评标委员会成员行贿或提供其他不正当利益。")
    add_para(doc, "12. 我方承诺自觉接受有关部门的监督和检查。")
    
    # 五、其他承诺
    add_para(doc, "五、其他承诺", font_size=14, bold=True, space_before=Pt(12))
    add_para(doc, "13. 我方承诺在投标有效期内不修改、撤销投标文件。")
    add_para(doc, "14. 如我方中标，承诺不转包、不违法分包。")
    add_para(doc, "15. 我方愿意承担因违反本承诺书内容而产生的一切法律责任和经济责任。")
    
    # 结尾
    add_para(doc, "特此承诺！", space_before=Pt(12))
    add_para(doc, "", space_before=Pt(12))  # 空行用于分隔
    add_para(doc, "投标人名称（盖章）：{}".format(info.get("公司全称", "")))
    add_para(doc, "法定代表人或授权代表（签字）：{}".format(info.get("授权代表", "")))
    add_para(doc, "联系电话：{}".format(info.get("授权代表电话", "")))
    add_para(doc, "日    期：2026 年    月    日")
    
    doc.save(output_path)
    print(f"✓ 已生成：{output_path}")
    return True

def main():
    info_path = "/home/ym/.openclaw/workspace/tender-generator-skill/info.json"
    output_path = "/home/ym/.openclaw/workspace/tender-generator-skill/output_final/01-承诺书.docx"
    
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    with open(info_path, 'r', encoding='utf-8') as f:
        info = json.load(f)
    
    create_commitment_document(output_path, info)
    print("\n✅ 完成！请检查 output_final/01-承诺书.docx 是否有空行问题。")

if __name__ == "__main__":
    main()
