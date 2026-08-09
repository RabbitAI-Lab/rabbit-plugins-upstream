#!/usr/bin/env python3
"""
批量修复所有生成文档的段落间距问题
"""

import os
import docx
from docx.shared import Pt
from docx.oxml.ns import qn

def fix_document_paragraphs(file_path):
    """修复文档中的段落间距"""
    try:
        doc = docx.Document(file_path)
        
        # 设置默认段落样式
        style = doc.styles['Normal']
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.5
        
        # 修复所有段落
        for para in doc.paragraphs:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.5
        
        # 修复表格中的段落
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        para.paragraph_format.line_spacing = 1.5
        
        # 保存
        doc.save(file_path)
        print(f"✓ 已修复：{file_path}")
        return True
    except Exception as e:
        print(f"✗ 修复失败 {file_path}: {e}")
        return False

def main():
    output_dir = "/home/ym/.openclaw/workspace/tender-generator-skill/output"
    
    print("开始修复文档段落间距...")
    print("=" * 50)
    
    files = [f for f in os.listdir(output_dir) if f.endswith('.docx')]
    
    for filename in files:
        file_path = os.path.join(output_dir, filename)
        fix_document_paragraphs(file_path)
    
    print("=" * 50)
    print("✅ 所有文档修复完成！")

if __name__ == "__main__":
    main()
