"""
pipeline/formatter.py — 输出格式化模块
从 pipeline.py 拆分而来 (2026-07-18)
职责：JudgmentDraft → Text/Markdown/HTML 输出
"""

import re


def _esc(s: str) -> str:
    """HTML 转义"""
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _nl2br(s: str) -> str:
    """换行转 <br>"""
    return _esc(s).replace("\n", "<br>\n")


# ─── 纯文本格式 ────────────────────────────────────────
def format_judgment_text(draft) -> str:
    """将草稿格式化为纯文本判决书"""
    lines = []

    lines.append("××××××人民法院")
    lines.append("民 事 判 决 书")
    lines.append("")
    lines.append("(××××)××××民初字第××××号")
    lines.append("")
    lines.append(draft.parties_section)
    lines.append("")
    lines.append(f"案由:{draft.cause_of_action}")
    lines.append("")
    lines.append(draft.claims_section)
    lines.append("")
    lines.append(draft.facts_section)
    lines.append("")
    lines.append(draft.evidence_section)
    lines.append("")
    lines.append(draft.reasoning_section)
    lines.append("")
    lines.append(draft.verdict_section)
    lines.append("")
    lines.append(draft.footer_section)

    return "\n".join(lines)


# ─── Markdown 格式 ─────────────────────────────────────
def format_judgment_markdown(draft) -> str:
    """将草稿格式化为 Markdown 判决书"""
    import random
    import datetime
    
    md = []
    
    # 生成案号
    year = datetime.datetime.now().year
    case_num = random.randint(1000, 9999)
    case_no = f"（{year}）××××民初字第{case_num}号"

    md.append("# ××××××人民法院民事判决书")
    md.append(f"**{case_no}**\n")

    md.append("## 当事人")
    md.append(draft.parties_section + "\n")

    md.append(f"**案由**：{draft.cause_of_action}\n")

    md.append("## 诉讼请求")
    md.append(draft.claims_section + "\n")

    md.append("## 事实认定")
    # 确保有经审理查明
    facts = draft.facts_section
    if not facts.startswith("经审理查明"):
        facts = "经审理查明：\n" + facts
    md.append(facts + "\n")

    md.append("## 证据")
    md.append(draft.evidence_section + "\n")

    md.append("## 本院认为")
    md.append(draft.reasoning_section + "\n")

    md.append("## 判决主文")
    # 确保有判决如下
    verdict = draft.verdict_section
    if not verdict.startswith("判决如下"):
        verdict = "判决如下：\n" + verdict
    md.append(verdict + "\n")

    md.append("## 审判人员")
    md.append(draft.footer_section + "\n")

    md.append("---")

    if draft.retrieved_cases:
        md.append("\n## 📚 参考入库案例")
        for c in draft.retrieved_cases[:3]:
            md.append(f"- **{c['title'][:60]}**")
            if c['content']:
                md.append(f"  > {c['content'][:200]}...")

    if draft.retrieved_laws:
        md.append("\n## 📜 参考法律法规")
        for l in draft.retrieved_laws[:3]:
            md.append(f"- {l['title']}")
            if l['content']:
                md.append(f"  > {l['content'][:150]}...")

    if draft.retrieved_patterns:
        md.append("\n## ✍️ 参考优秀文书范式")
        for p in draft.retrieved_patterns[:3]:
            md.append(f"- **{p['title'][:50]}**")
            if p['content']:
                md.append(f"  > {p['content'][:200]}...")

    if draft.warnings:
        md.append("\n## ⚠️ 注意事项")
        for w in draft.warnings:
            md.append(f"- {w}")

    return "\n".join(md)


# ─── HTML 格式 ─────────────────────────────────────────
def format_judgment_html(draft) -> str:
    """将草稿格式化为 HTML 判决书(符合最高法排版规范 + 精致视觉设计)"""
    html = []

    html.append("""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>民事判决书草稿</title>
<style>
@page { size: A4; margin: 3.7cm 2.6cm 3.5cm 2.8cm; }
* { margin: 0; padding: 0; box-sizing: border-box; }

body {
  font-family: "FangSong", "仿宋", "STFangsong", "仿宋_GB2312", serif;
  font-size: 16pt; line-height: 28pt; color: #1a1a1a;
  background: #e8e4df; padding: 32px 20px;
  -webkit-font-smoothing: antialiased;
}

.judgment {
  max-width: 210mm; margin: 0 auto; background: #fff;
  padding: 3.7cm 2.6cm 3.5cm 2.8cm;
  box-shadow: 0 1px 3px rgba(0,0,0,0.06), 0 8px 32px rgba(0,0,0,0.08);
  border-radius: 2px; position: relative;
}
.judgment::before {
  content: ""; position: absolute; top: 0; left: 0; right: 0; bottom: 0;
  background: repeating-linear-gradient(0deg, transparent, transparent 27pt, rgba(0,0,0,0.015) 27pt, rgba(0,0,0,0.015) 28pt);
  pointer-events: none; z-index: 0;
}
.judgment > * { position: relative; z-index: 1; }

.court-name {
  font-family: "FZXiaoBiaoSong-B05S", "方正小标宋简体", "SimSun", serif;
  font-size: 22pt; text-align: center; line-height: 28pt;
  letter-spacing: 4pt; font-weight: normal; color: #000; margin-bottom: 2pt;
}
.doc-title {
  font-family: "FZXiaoBiaoSong-B05S", "方正小标宋简体", "SimSun", serif;
  font-size: 26pt; text-align: center; line-height: 36pt;
  letter-spacing: 6pt; font-weight: normal; color: #000;
  margin-bottom: 12pt; padding-bottom: 10pt;
  border-bottom: 1.5pt solid #8b0000;
}
.case-no {
  font-size: 14pt; text-align: center; line-height: 28pt;
  margin-bottom: 20pt; color: #333; letter-spacing: 1pt;
}

.parties { margin-bottom: 16pt; padding: 10pt 0; }
.parties p { line-height: 28pt; text-indent: 0; }

.cause {
  font-weight: bold; margin-bottom: 16pt; line-height: 28pt;
  padding: 6pt 0; border-bottom: 0.5pt solid #ccc;
}

.section-title {
  font-size: 16pt; font-weight: bold; line-height: 28pt;
  margin: 16pt 0 6pt; color: #000;
  position: relative; padding-left: 12pt;
}
.section-title::before {
  content: ""; position: absolute;
  left: 0; top: 6pt; bottom: 6pt;
  width: 3pt; background: #8b0000; border-radius: 1pt;
}

section { margin-bottom: 8pt; }
section p { line-height: 28pt; text-indent: 2em; }

.verdict {
  margin: 16pt 0; padding: 12pt 16pt;
  background: #fafaf8; border: 1pt solid #d4d0c8;
  border-left: 3pt solid #8b0000; border-radius: 2pt;
}
.verdict p { line-height: 28pt; text-indent: 2em; }

.footer { margin-top: 32pt; text-align: right; line-height: 32pt; }
.footer p { text-indent: 0; }

.ref-section {
  margin-top: 40pt; border-top: 1pt solid #d4d0c8;
  padding-top: 20pt; page-break-before: always;
  font-size: 13pt; line-height: 22pt;
}
.ref-section h2 {
  font-size: 15pt; margin-bottom: 12pt; color: #333;
  padding-bottom: 6pt; border-bottom: 0.5pt solid #e0e0e0;
}
.ref-section h3 { font-size: 13pt; margin: 12pt 0 6pt; color: #555; }
.ref-item {
  margin: 6pt 0; padding: 8pt 12pt;
  background: #f8f7f5; border-radius: 3pt;
  border-left: 2pt solid #d4d0c8; font-size: 12pt; line-height: 20pt;
}
.ref-item strong { color: #333; }
.ref-item blockquote {
  color: #666; font-size: 11pt; margin-top: 4pt;
  padding-left: 10pt; border-left: 1.5pt solid #ccc; line-height: 18pt;
}

.warnings {
  background: #fffdf0; border: 1pt solid #e6d88a;
  border-left: 3pt solid #c9a800; padding: 10pt 14pt;
  margin-top: 20pt; font-size: 13pt; line-height: 22pt; border-radius: 2pt;
}
.warnings p { text-indent: 0; }

@media print {
  body { background: #fff; padding: 0; }
  .judgment { box-shadow: none; padding: 0; max-width: none; border-radius: 0; }
  .judgment::before { display: none; }
  .ref-section { page-break-before: always; }
  .warnings { page-break-inside: avoid; }
  .verdict { page-break-inside: avoid; }
  .footer { page-break-before: avoid; }
}
@media (max-width: 800px) {
  body { padding: 8px; background: #f5f5f5; }
  .judgment { padding: 24px 16px; border-radius: 0; }
}
</style>
</head>
<body>
<div class="judgment">
""")

    html.append('<p class="court-name">\u00d7\u00d7\u00d7\u00d7\u00d7\u00d7\u4eba\u6c11\u6cd5\u9662</p>')
    html.append('<p class="doc-title">\u6c11 \u4e8b \u5224 \u51b3 \u4e66</p>')
    html.append('<p class="case-no">\uff08\u00d7\u00d7\u00d7\u00d7\uff09\u00d7\u00d7\u00d7\u00d7\u6c11\u521d\u5b57\u7b2c\u00d7\u00d7\u00d7\u00d7\u53f7</p>')

    html.append('<div class="parties">')
    for line in draft.parties_section.split('\n'):
        if line.strip():
            html.append(f'<p>{_esc(line)}</p>')
    html.append('</div>')

    html.append(f'<p class="cause">\u6848\u7531\uff1a{_esc(draft.cause_of_action)}</p>')

    sections = [
        ('\u8bc9\u8bbc\u8bf7\u6c42', draft.claims_section),
        ('\u4e8b\u5b9e\u8ba4\u5b9a', draft.facts_section),
        ('\u8bc1\u636e', draft.evidence_section),
        ('\u672c\u9662\u8ba4\u4e3a', draft.reasoning_section),
    ]
    for title, content in sections:
        html.append(f'<p class="section-title">{title}</p>')
        html.append(f'<section>{_nl2br(content)}</section>')

    html.append('<p class="section-title">\u5224\u51b3\u5982\u4e0b</p>')
    html.append(f'<div class="verdict">{_nl2br(draft.verdict_section)}</div>')

    html.append('<div class="footer">')
    for line in draft.footer_section.split('\n'):
        if line.strip():
            html.append(f'<p>{_esc(line)}</p>')
    html.append('</div>')

    if draft.retrieved_cases or draft.retrieved_laws or draft.retrieved_patterns:
        html.append('<div class="ref-section">')
        html.append('<h2>\u53c2\u8003\u8d44\u6599</h2>')
        if draft.retrieved_cases:
            html.append('<h3>\u5165\u5e93\u6848\u4f8b</h3>')
            for c in draft.retrieved_cases[:3]:
                html.append(f'<div class="ref-item"><strong>{_esc(c["title"][:60])}</strong>')
                if c['content']:
                    html.append(f'<blockquote>{_esc(c["content"][:200])}</blockquote>')
                html.append('</div>')
        if draft.retrieved_laws:
            html.append('<h3>\u6cd5\u5f8b\u6cd5\u89c4</h3>')
            for l in draft.retrieved_laws[:3]:
                html.append(f'<div class="ref-item">{_esc(l["title"])}')
                if l['content']:
                    html.append(f'<blockquote>{_esc(l["content"][:150])}</blockquote>')
                html.append('</div>')
        if draft.retrieved_patterns:
            html.append('<h3>\u4f18\u79c0\u6587\u4e66\u8303\u5f0f</h3>')
            for p in draft.retrieved_patterns[:3]:
                html.append(f'<div class="ref-item"><strong>{_esc(p["title"][:50])}</strong>')
                if p['content']:
                    html.append(f'<blockquote>{_esc(p["content"][:200])}</blockquote>')
                html.append('</div>')
        html.append('</div>')

    if draft.warnings:
        html.append('<div class="warnings">')
        for w in draft.warnings:
            html.append(f'<p>{_esc(w)}</p>')
        html.append('</div>')

    html.append('</div></body></html>')
    return "\n".join(html)


# ─── DOCX 格式 ─────────────────────────────────────────
def format_judgment_docx(draft, output_path: str) -> str:
    """
    将草稿格式化为 DOCX 文件（符合法院排版规范）。
    返回文件路径。
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # ── 默认字体 ──
    style = doc.styles['Normal']
    font = style.font
    font.name = '仿宋'
    font.size = Pt(16)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    def add_para(text, bold=False, size=None, align=None, space_after=None, font_name=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.bold = True
        if size:
            run.font.size = Pt(size)
        if font_name:
            run.font.name = font_name
            run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if align:
            p.alignment = align
        if space_after is not None:
            p.paragraph_format.space_after = Pt(space_after)
        return p

    # ── 法院名称 ──
    add_para('××××××人民法院', bold=True, size=22,
             align=WD_ALIGN_PARAGRAPH.CENTER, font_name='方正小标宋简体')

    # ── 文书标题 ──
    add_para('民 事 判 决 书', bold=True, size=26,
             align=WD_ALIGN_PARAGRAPH.CENTER, font_name='方正小标宋简体',
             space_after=6)

    # ── 案号 ──
    add_para('（××××）××××民初字第××××号', size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # ── 当事人 ──
    for line in draft.parties_section.split('\n'):
        if line.strip():
            add_para(line.strip())

    # ── 案由 ──
    add_para(f'案由：{draft.cause_of_action}', bold=True, space_after=6)

    # ── 诉讼请求 ──
    add_para('诉讼请求', bold=True, size=16)
    for line in draft.claims_section.split('\n'):
        if line.strip():
            add_para(line.strip())

    # ── 事实认定 ──
    add_para('事实认定', bold=True, size=16)
    for line in draft.facts_section.split('\n'):
        if line.strip():
            add_para(line.strip())

    # ── 证据 ──
    add_para('证据', bold=True, size=16)
    for line in draft.evidence_section.split('\n'):
        if line.strip():
            add_para(line.strip())

    # ── 本院认为 ──
    add_para('本院认为', bold=True, size=16)
    for line in draft.reasoning_section.split('\n'):
        if line.strip():
            add_para(line.strip())

    # ── 判决如下 ──
    add_para('判决如下', bold=True, size=16)
    for line in draft.verdict_section.split('\n'):
        if line.strip():
            add_para(line.strip())

    # ── 审判人员 ──
    doc.add_page_break()
    for line in draft.footer_section.split('\n'):
        if line.strip():
            add_para(line.strip(), align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── 保存 ──
    doc.save(output_path)
    return output_path
