#!/usr/bin/env python3
"""Generate civil complaint DOCX files from baozheng skill templates."""

from __future__ import annotations

import argparse
from html import escape
import json
import re
from zipfile import ZIP_DEFLATED, ZipFile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable


CASE_NAME_RE = re.compile(r"^template-(?P<code>\d{2})-(?P<slug>.+)\.md$")
CRIMINAL_NAME_RE = re.compile(r"^(?P<slug>[a-z-]+)-template\.md$")
FIELD_RE = re.compile(r"^\s*-\s+\*\*(?P<body>.+?)\*\*\s*$")
NORMALIZE_RE = re.compile(r"[\s:：()（）/、，,。；;|]+")
PARAGRAPH_STYLE_TITLES = {"民事起诉状", "刑事附带民事起诉状"}


GENERAL_CIVIL_FIELDS = [
    "原告",
    "法定代理人/指定代理人",
    "委托诉讼代理人",
    "被告",
    "诉讼请求",
    "事实和理由",
    "证据和证据来源，证人姓名和住所",
    "受诉法院",
    "副本份数",
    "具状人",
    "日期",
]


@dataclass(frozen=True)
class CaseTemplate:
    code: str
    slug: str
    case_path: Path
    template_path: Path


@dataclass(frozen=True)
class TemplateField:
    label: str
    placeholder: str


def resolve_skill_root(skill_root: str | None) -> Path:
    if skill_root:
        return Path(skill_root).resolve()
    return Path(__file__).resolve().parents[1]


def discover_cases(skill_root: Path) -> list[CaseTemplate]:
    assets_dir = skill_root / "assets"
    references_dir = skill_root / "references"
    cases: list[CaseTemplate] = []

    for template_path in sorted(assets_dir.glob("template-*.md")):
        match = CASE_NAME_RE.match(template_path.name)
        if match is None:
            continue

        code = match.group("code")
        slug = match.group("slug")
        case_path = references_dir / f"case-{code}-{slug}.md"
        if not case_path.exists():
            raise FileNotFoundError(f"Missing matching case file: {case_path}")

        cases.append(
            CaseTemplate(
                code=code,
                slug=slug,
                case_path=case_path,
                template_path=template_path,
            )
        )

    for template_path in sorted(assets_dir.glob("*-template.md")):
        match = CRIMINAL_NAME_RE.match(template_path.name)
        if match is None:
            continue

        slug = match.group("slug")
        code = slug  # 刑事模板使用 slug 作为 code
        case_path = references_dir / "module-d-criminal.md"
        cases.append(
            CaseTemplate(
                code=code,
                slug=slug,
                case_path=case_path,
                template_path=template_path,
            )
        )

    if not cases:
        raise FileNotFoundError(f"No template-*.md or criminal template files found in {assets_dir}")

    return cases


def select_case(cases: Iterable[CaseTemplate], case_id: str) -> CaseTemplate:
    normalized = case_id.strip().lower()
    for item in cases:
        if normalized in {item.code, item.slug.lower(), f"{item.code}-{item.slug}".lower()}:
            return item
    # 显示时，如果 code 和 slug 相同，只显示一个
    available = ", ".join(
        item.code if item.code == item.slug else f"{item.code}-{item.slug}" 
        for item in cases
    )
    raise ValueError(f"Unknown case '{case_id}'. Available: {available}")


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


CRIMINAL_TITLE_MAP = {
    "刑事控告材料模板": "刑事控告书",
    "取保候审申请模板": "取保候审申请书",
    "刑事辩护意见要点模板": "辩护意见",
    "刑事附带民事起诉状模板": "刑事附带民事起诉状",
    "羁押阶段家属沟通提纲模板": "家属沟通提纲",
}


def parse_title(template_text: str) -> tuple[str, str]:
    title = "民事起诉状"
    subtitle = ""
    for line in template_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            heading = stripped[2:].strip()
            if heading in CRIMINAL_TITLE_MAP:
                title = CRIMINAL_TITLE_MAP[heading]
            subtitle = heading
        elif stripped in {"民事起诉状", "行政起诉状"}:
            title = stripped
        elif stripped.startswith("(") and stripped.endswith(")"):
            subtitle = stripped
            break
    return title, subtitle


def parse_fields(template_text: str) -> list[TemplateField]:
    fields: list[TemplateField] = []
    for line in template_text.splitlines():
        match = FIELD_RE.match(line)
        if match is None:
            continue

        body = match.group("body").strip()
        if " | " in body:
            label, placeholder = body.split(" | ", 1)
        else:
            label, placeholder = body, ""
        fields.append(TemplateField(label=label.strip(), placeholder=placeholder.strip()))

    if not fields and "通用民事起诉状" in template_text:
        return [TemplateField(label=label, placeholder="") for label in GENERAL_CIVIL_FIELDS]

    if not fields:
        raise ValueError("No template fields found. Expected Markdown lines like '- **label | value**'.")
    return fields


def load_data(data_path: str | None) -> dict[str, str]:
    if not data_path:
        return {}

    raw_data = json.loads(Path(data_path).read_text(encoding="utf-8"))
    if not isinstance(raw_data, dict):
        raise ValueError("--data JSON must be an object")

    source = raw_data.get("fields", raw_data)
    if not isinstance(source, dict):
        raise ValueError("--data JSON 'fields' must be an object when provided")

    return {str(key): stringify_value(value) for key, value in source.items() if value is not None}


def stringify_value(value: object) -> str:
    if isinstance(value, str):
        return value
    if isinstance(value, (list, tuple)):
        return "；".join(stringify_value(item) for item in value)
    if isinstance(value, dict):
        return "；".join(f"{key}：{stringify_value(item)}" for key, item in value.items())
    return str(value)


def normalize_label(value: str) -> str:
    return NORMALIZE_RE.sub("", value).lower()


def find_field_value(field: TemplateField, data: dict[str, str]) -> str | None:
    if field.label in data:
        return data[field.label]

    normalized_label = normalize_label(field.label)
    normalized_data = {normalize_label(key): value for key, value in data.items()}
    if normalized_label in normalized_data:
        return normalized_data[normalized_label]

    for key, value in normalized_data.items():
        if key and (key in normalized_label or normalized_label in key):
            return value
    return None


def apply_data(fields: list[TemplateField], data: dict[str, str]) -> tuple[list[TemplateField], list[str]]:
    if not data:
        return fields, []

    filled_fields: list[TemplateField] = []
    matched_labels: list[str] = []
    for field in fields:
        value = find_field_value(field, data)
        if value is None:
            filled_fields.append(field)
            continue

        filled_fields.append(TemplateField(label=field.label, placeholder=value))
        matched_labels.append(field.label)
    return filled_fields, matched_labels


def create_docx(title: str, subtitle: str, fields: list[TemplateField], output_path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
        from docx.oxml.ns import qn
    except ImportError:
        create_docx_with_stdlib(title, subtitle, fields, output_path)
        return

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "FangSong"
    normal_style.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '仿宋')
    normal_style.font.size = Pt(12)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)

    if subtitle:
        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        subtitle_run.font.size = Pt(14)

    if title in PARAGRAPH_STYLE_TITLES:
        for field in fields:
            label_run = document.add_paragraph().add_run(field.label + "：")
            label_run.bold = True
            label_run.font.name = "FangSong"
            label_run.font.size = Pt(12)
            value_paragraph = document.add_paragraph(field.placeholder)
            for run in value_paragraph.runs:
                run.font.name = "FangSong"
                run.font.size = Pt(12)
    else:
        table = document.add_table(rows=len(fields), cols=2, style="Table Grid")
        for index, field in enumerate(fields):
            table.cell(index, 0).text = field.label
            table.cell(index, 1).text = field.placeholder

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def paragraph_xml(text: str, bold: bool = False, center: bool = False, size: str | None = None) -> str:
    align_xml = '<w:pPr><w:jc w:val="center"/></w:pPr>' if center else ""
    bold_xml = "<w:b/>" if bold else ""
    size_xml = f'<w:sz w:val="{size}"/>' if size else ""
    run_pr = f"<w:rPr>{bold_xml}{size_xml}</w:rPr>" if bold_xml or size_xml else ""
    return f"<w:p>{align_xml}<w:r>{run_pr}<w:t>{escape(text)}</w:t></w:r></w:p>"


def cell_xml(text: str) -> str:
    return f"<w:tc><w:tcPr><w:tcW w:w=\"4500\" w:type=\"dxa\"/></w:tcPr>{paragraph_xml(text)}</w:tc>"


def create_docx_with_stdlib(title: str, subtitle: str, fields: list[TemplateField], output_path: Path) -> None:
    if title in PARAGRAPH_STYLE_TITLES:
        body_parts = [paragraph_xml(title, bold=True, center=True, size="32")]
        if subtitle:
            body_parts.append(paragraph_xml(subtitle, center=True, size="28"))
        for field in fields:
            body_parts.append(paragraph_xml(field.label + "\uff1a", bold=True))
            body_parts.append(paragraph_xml(field.placeholder))
        body_content = "\n    ".join(body_parts)
    else:
        rows = ["<w:tr>" + cell_xml(field.label) + cell_xml(field.placeholder) + "</w:tr>" for field in fields]
        table_xml = f"""<w:tbl>
      <w:tblPr>
        <w:tblW w:w="0" w:type="auto"/>
        <w:tblBorders>
          <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
          <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
          <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
          <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
          <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
          <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        </w:tblBorders>
      </w:tblPr>
      {''.join(rows)}
    </w:tbl>"""
        body_content = f"""{paragraph_xml(title, bold=True, center=True, size="32")}
    {paragraph_xml(subtitle, center=True, size="28") if subtitle else ""}
    {table_xml}"""

    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {body_content}
    <w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>
  </w:body>
</w:document>
"""
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
"""
    package_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with ZipFile(output_path, "w", ZIP_DEFLATED) as docx:
        docx.writestr("[Content_Types].xml", content_types)
        docx.writestr("_rels/.rels", package_rels)
        docx.writestr("word/document.xml", document_xml)


def build_summary(
    case: CaseTemplate,
    fields: list[TemplateField],
    title: str,
    subtitle: str,
    matched_labels: list[str] | None = None,
) -> dict[str, object]:
    matched_labels = matched_labels or []
    # 显示时，如果 code 和 slug 相同，只显示一个
    case_display = case.code if case.code == case.slug else f"{case.code}-{case.slug}"
    return {
        "case": case_display,
        "case_path": str(case.case_path),
        "template_path": str(case.template_path),
        "title": title,
        "subtitle": subtitle,
        "field_count": len(fields),
        "filled_count": len(matched_labels),
        "filled_labels": matched_labels,
        "first_fields": [field.label for field in fields[:5]],
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate complaint DOCX from baozheng skill templates.")
    parser.add_argument("--skill-root", help="Path to baozheng-skills root. Defaults to script parent.")
    parser.add_argument("--case", default="00", help="Case code or slug, for example 01 or private-lending.")
    parser.add_argument("--data", help="Optional UTF-8 JSON file used to fill right-column values.")
    parser.add_argument("--output", help="Output .docx path. Required unless --dry-run or --list is used.")
    parser.add_argument("--dry-run", action="store_true", help="Parse template and print JSON summary only.")
    parser.add_argument("--list", action="store_true", help="List available case templates.")
    args = parser.parse_args()

    skill_root = resolve_skill_root(args.skill_root)
    cases = discover_cases(skill_root)

    if args.list:
        # 显示时，如果 code 和 slug 相同，只显示一个
        case_list = [
            item.code if item.code == item.slug else f"{item.code}-{item.slug}"
            for item in cases
        ]
        print(json.dumps(case_list, ensure_ascii=False, indent=2))
        return 0

    selected = select_case(cases, args.case)
    template_text = read_text(selected.template_path)
    title, subtitle = parse_title(template_text)
    fields = parse_fields(template_text)
    data = load_data(args.data)
    fields, matched_labels = apply_data(fields, data)

    if args.dry_run:
        print(json.dumps(build_summary(selected, fields, title, subtitle, matched_labels), ensure_ascii=False, indent=2))
        return 0

    if not args.output:
        parser.error("--output is required unless --dry-run or --list is used")

    create_docx(title, subtitle, fields, Path(args.output).resolve())
    print(json.dumps({"output": str(Path(args.output).resolve()), **build_summary(selected, fields, title, subtitle, matched_labels)}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
