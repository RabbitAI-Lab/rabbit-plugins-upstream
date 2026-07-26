#!/usr/bin/env python3
"""DramaLex · export_hub.py  (orchestrator)
Reads words.json + listening.json + annotated.json + tasks.json, generates audio
for every referenced line, and emits:
  - practice.html : four-skill dashboard (audio embedded, flip cards, reflection,
                   exit check, author footer). Works in any browser, no app.
  - per --mode A|B|C:
      A (layered, default): recall.tsv + recall.apkg = vocab+dictation+listening+cloze
      B (all-to-Anki):       above + speaking/writing production cards
      C (all-HTML):          no Anki file; everything lives in practice.html
Pure stdlib + optional genanki. Agent-neutral.
"""
import argparse, json, os, sys, html, base64, shutil, subprocess, time

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
import gen_audio as ga
import exam_map as exm

AUTHOR = "yinjianheng（殷健恒）"
CONTACT = "📧 yinjianheng@foxmail.com · 💬 WeChat: YJH-yinjianheng"
LEGAL = ("⚖️ 法律声明：本工具协助你从互联网公开渠道检索字幕，仅供个人非商业语言学习使用。"
         "生成物含第三方台词，请勿再分发或用于商业用途；如权利人主张权利，请联系作者下架。")

# ---------- audio ----------
def b64(path):
    if not path or not os.path.exists(path):
        return ""
    ext = os.path.splitext(path)[1].lower()
    mime = 'audio/mpeg' if ext == '.mp3' else 'audio/wav'
    with open(path, 'rb') as f:
        return f"data:{mime};base64," + base64.b64encode(f.read()).decode()

def collect_line_audio(text, media_dir, backend, voice, cache):
    key = (text or "").strip().lower()
    if not key:
        return ""
    if key in cache:
        return cache[key]
    idx = len(cache) + 1
    out = os.path.join(media_dir, f"line_{idx:03d}.wav")
    p = ga.gen(text, out, backend, voice)
    if p:
        cache[key] = p
        return p
    return ""

def collect_all_audio(words, listening, annotated, tasks, media_dir, backend, voice):
    """Generate TTS for every non-word line referenced by listening/annotated/speaking.
    Returns (cache, missing). cache = {lowercased_text: path}, idempotent.
    missing = list of source texts whose TTS failed (for the build-time warning summary)."""
    cache = {}
    missing = []
    texts = []
    for q in listening.get('comprehension', []):
        if q.get('audio_line'): texts.append(q['audio_line'])
    for x in listening.get('dictation', []):
        texts.append(x['line'])
    for a in annotated.get('annotations', []):
        texts.append(a['line'])
    for s in tasks.get('speaking', []):
        if s.get('model_line'): texts.append(s['model_line'])
    for t in texts:
        p = collect_line_audio(t, media_dir, backend, voice, cache)
        if not p:
            missing.append(t)
        # 让出语音合成服务，避免连续调用抢占失败（macOS say 服务易在高频调用时丢帧）
        time.sleep(0.2)
    return cache, missing

# ---------- helpers ----------
def esc(s):
    return html.escape(str(s or ""))

def vocab_back(w):
    p = [f"<b>{esc(w['term'])}</b>"]
    if w.get('ipa'): p.append(f" {esc(w['ipa'])}")
    if w.get('pos'): p.append(f" <i>({esc(w['pos'])})</i>")
    if w.get('cefr'): p.append(f" [{esc(w['cefr'])}]")
    exam = w.get('exam') or (exm.exam_label(w['cefr']) if w.get('cefr') else '')
    if exam:
        p.append(f" <span class='exam'>{esc(exam)}</span>")
    p.append(f"<br>{esc(w.get('gloss',''))}")
    if w.get('collocation'): p.append(f"<br><i>搭配:</i> {esc(w['collocation'])}")
    if w.get('line'):
        sp = f" ({esc(w['line_speaker'])})" if w.get('line_speaker') else ""
        p.append(f"<br>🎬 {esc(w['line'])}{sp}")
    if w.get('example'): p.append(f"<br>✏️ {esc(w['example'])}")
    if w.get('why'): p.append(f"<br><i>🎯 为什么学：</i> {esc(w['why'])}")
    if w.get('l1_note'): p.append(f"<br>🇨🇳 <i>中文易错：</i> {esc(w['l1_note'])}")
    if w.get('tags'): p.append(f"<br><small>{esc(', '.join(w['tags']))}</small>")
    return ''.join(p)

# ---------- Excel / Word UI localization ----------
XL_UI = {
    'zh': {
        'vocab': ('词汇卡', ["词/语块","类型","CEFR","考试对照","音标","词性","释义","搭配","原句","例句","为什么学","中文易错","标签"]),
        'watch': ('观看', ["遍数","名称","说明"]),
        'listen': ('听力理解', ["ID","类型","问题","选项","答案","解析"]),
        'dict': ('听写', ["ID","挖空句","答案","目标词"]),
        'mp': ('最小对立体', ["ID","词A","IPA-A","词B","IPA-B","本集实际","提示"]),
        'cs': ('连读拆解', ["ID","原句","拆解(自然读法/注解)","意思"]),
        'anno': ('精读标注', ["ID","焦点","原句","规则","注解","提示","更多例句"]),
        'cloze': ('完形', ["ID","挖空","答案"]),
        'speak': ('口语', ["ID","类型","指令","需用词","重点发音","检核表","可评分目标句"]),
        'write': ('写作', ["ID","类型","语域","指令","需用词","量规","范文","自动量规"]),
        'review': '复习',
        'review_note': "说明：音频播放请配合 practice.html 或 deck.apkg（Excel 内无法内嵌音频）。",
        'pre': "学前目标：这集我想搞懂/记住什么？",
        'post': "学后反思：我注意到了什么？还有什么模糊？",
    },
    'en': {
        'vocab': ('Vocab', ["Term/Chunk","Type","CEFR","Exam","IPA","POS","Gloss","Colloc","Line","Example","Why","L1-note","Tags"]),
        'watch': ('Watch', ["Pass","Name","Instruction"]),
        'listen': ('Listening', ["ID","Type","Question","Options","Answer","Rationale"]),
        'dict': ('Dictation', ["ID","Cloze","Answer","Target"]),
        'mp': ('Minimal pairs', ["ID","Word A","IPA-A","Word B","IPA-B","In episode","Hint"]),
        'cs': ('Connected speech', ["ID","Line","Breakdown (natural read/notes)","Gloss"]),
        'anno': ('Annotation', ["ID","Focus","Line","Rule","Note","Tip","More"]),
        'cloze': ('Cloze', ["ID","Cloze","Answer"]),
        'speak': ('Speaking', ["ID","Type","Instruction","Use","Focus sounds","Checklist","ASR target"]),
        'write': ('Writing', ["ID","Type","Register","Instruction","Use","Rubric","Model","Auto-checks"]),
        'review': 'Review',
        'review_note': "Note: play audio via practice.html or deck.apkg (Excel can't embed audio).",
        'pre': "Pre-goal: what do I want to get / remember from this ep?",
        'post': "Post-reflection: what did I notice? What is still fuzzy?",
    },
}
WD_UI = {
    'zh': {
        'title_sub': "看剧学英语终极闭环 · 听/说/读/写 四技能报告",
        'sec': ["0 · 词汇卡（目标词链）","1 · 看（观看协议）","2 · 听","3 · 读 · 台词精读","4 · 说","5 · 写","6 · 复习"],
        'listen_sub': "听力理解", 'dict_sub': "听写", 'cloze_sub': "完形",
        'mp_sub': "最小对立体", 'cs_sub': "连读拆解",
        'pre': "学前目标：这集我想搞懂/记住什么？",
        'post': "学后反思：我注意到了什么？还有什么模糊？",
        'note': "说明：音频播放请配合 practice.html 或 deck.apkg（Word 内无法内嵌音频）。",
        'vocab_hdr': ["词/语块","CEFR","音标","释义","搭配","原句","为什么学","中文易错"],
        'listen_hdr': ["问题","选项","答案","解析"],
        'dict_hdr': ["挖空句","答案","目标词"],
        'mp_hdr': ["词A","IPA-A","词B","IPA-B","本集实际","提示"],
        'cs_hdr': ["原句","拆解","意思"],
        'anno_hdr': ["焦点","原句","规则","注解","提示","更多例句"],
        'cloze_hdr': ["挖空","答案"],
        'speak_hdr': ["类型","指令","需用词","重点发音","检核表","可评分目标句"],
        'write_hdr': ["类型","语域","指令","需用词","量规","范文","自动量规"],
    },
    'en': {
        'title_sub': "Learn English from TV · four-skill report",
        'sec': ["0 · Vocab (target lexicon)","1 · Watch (protocol)","2 · Listen","3 · Read (annotation)","4 · Speak","5 · Write","6 · Review"],
        'listen_sub': "Listening comprehension", 'dict_sub': "Dictation", 'cloze_sub': "Cloze",
        'mp_sub': "Minimal pairs", 'cs_sub': "Connected speech",
        'pre': "Pre-goal: what do I want to get / remember?",
        'post': "Post-reflection: what did I notice? What is still fuzzy?",
        'note': "Note: play audio via practice.html or deck.apkg (Word can't embed audio).",
        'vocab_hdr': ["Term/Chunk","CEFR","IPA","Gloss","Colloc","Line","Why","L1-note"],
        'listen_hdr': ["Question","Options","Answer","Rationale"],
        'dict_hdr': ["Cloze","Answer","Target"],
        'mp_hdr': ["Word A","IPA-A","Word B","IPA-B","In episode","Hint"],
        'cs_hdr': ["Line","Breakdown","Gloss"],
        'anno_hdr': ["Focus","Line","Rule","Note","Tip","More"],
        'cloze_hdr': ["Cloze","Answer"],
        'speak_hdr': ["Type","Instruction","Use","Focus","Checklist","ASR target"],
        'write_hdr': ["Type","Register","Instruction","Use","Rubric","Model","Auto-checks"],
    },
}

# ---------- HTML builder ----------
UI = {
    'zh': {
        'title': '🎬 DramaLex · 看剧学英语完整闭环',
        'sub': '词汇预习 · 观看 · 听/说/读/写 · 复习 · 浏览器本地（无需安装 App）',
        'tabs': ['📚 词汇预习', '🎬 看', '🎧 听', '📖 读', '🗣️ 说', '✍️ 写', '🔁 复习'],
        'prime_h': '词汇预习（目标词链 · 看前先过一遍）',
        'watch_h': '🎬 观看协议（三步字幕法）',
        'catch_label': '📝 记下让你印象最深的一句／一个表达：',
        'listen_warn': '⚠️ 本环节音频为 TTS 合成参照音，真实原速的连读/弱读/语调请观看正片。',
        'comp_h': '听力理解（先只听，不看字幕）', 'dict_h': '听写',
        'mp_h': '最小对立体（音素级辨音）', 'cs_h': '连读拆解（自然语速读法）',
        'ann_h': '台词精读标注', 'cloze_h': 'Cloze 完形',
        'speak_h': '口语产出', 'write_h': '写作产出',
        'review_h': '跨技能复习（词汇·听写·完形 本地 SRS）', 'meta_h': '元认知',
        'exit_h': '出口小测', 'pre': '🎯 学前目标', 'pre_ph': '这集我想搞懂/记住什么？',
        'refl': '🪞 学后反思', 'refl_ph': '我注意到了什么？还有什么模糊？',
        'tts_note': 'TTS 为发音参照，真实跟读请对照正片。',
        'tts_warn2': '把你的作文贴给 agent，获取批改与“目标语块该在哪”的反馈。',
        'foot': "个人非商业学习用途 · 字幕版权归制片方所有，请勿再分发 · TTS 为合成音，原速听力请观看正片",
        'week_h': '本周打卡', 'week_lbl': ['一','二','三','四','五','六','日'],
        'prog_export': '⤓ 导出进度', 'prog_import': '⤒ 导入进度',
        'prod_h': '口语/写作产出（纳入复习）',
    },
    'en': {
        'title': '🎬 DramaLex · Learn English from TV',
        'sub': 'Vocab priming · Watch · Listen/Speak/Read/Write · Review · runs locally in your browser',
        'tabs': ['📚 Vocab', '🎬 Watch', '🎧 Listen', '📖 Read', '🗣️ Speak', '✍️ Write', '🔁 Review'],
        'prime_h': 'Vocab priming (target lexicon · scan before watching)',
        'watch_h': '🎬 Viewing protocol (3-pass subtitling)',
        'catch_label': '📝 Note the line / expression that stuck with you:',
        'listen_warn': '⚠️ Audio here is TTS reference only; real connected speech (liaison/reduction/intonation) comes from the actual episode.',
        'comp_h': 'Listening comprehension (listen first, no subtitles)', 'dict_h': 'Dictation',
        'mp_h': 'Minimal pairs (phoneme ear-training)', 'cs_h': 'Connected speech breakdown',
        'ann_h': 'Transcript annotation', 'cloze_h': 'Cloze',
        'speak_h': 'Speaking', 'write_h': 'Writing',
        'review_h': 'Cross-skill review (vocab · dictation · cloze, local SRS)',
        'meta_h': 'Metacognition',
        'exit_h': 'Exit check', 'pre': '🎯 Pre-goal', 'pre_ph': 'What do I want to get / remember from this ep?',
        'refl': '🪞 Post-reflection', 'refl_ph': 'What did I notice? What is still fuzzy?',
        'tts_note': 'TTS is a pronunciation reference; shadow the real audio for authentic delivery.',
        'tts_warn2': 'Paste your writing to the agent for feedback on where target chunks belong.',
        'foot': "Personal non-commercial study use only · subtitle copyright belongs to the rightsholder, do not redistribute · TTS is synthetic; real listening needs the episode",
        'week_h': 'This week', 'week_lbl': ['M','T','W','T','F','S','S'],
        'prog_export': '⤓ Export', 'prog_import': '⤒ Import',
        'prod_h': 'Speaking/Writing (in review)',
    },
}

def build_html(words, listening, annotated, tasks, audio_cache, mode, deck, watch=None, media_dir='media', ui_lang='zh', recall=None):
    import random
    U = UI.get(ui_lang, UI['zh'])
    def b64_local(p):
        if not p:
            return ""
        if os.path.exists(p):
            return b64(p)
        alt = os.path.join(media_dir, os.path.basename(p))
        if os.path.exists(alt):
            return b64(alt)
        return ""
    # target lexicon chips
    chips = "".join(f"<span class='chip'>{esc(w['term'])} <i>{esc(w.get('cefr',''))}</i></span>" for w in words)

    # ---- Listen tab ----
    comp = ""
    for q in listening.get('comprehension', []):
        opts = "".join(f"<label class='opt'><input type='radio' name='c{q['id']}'> {esc(o)}</label>" for o in q['options'])
        btn = f"<button class='ap' onclick='playTxt(\"t{q['id']}\")'>🔊 听</button>" if q.get('audio_line') else ""
        comp += f"""<div class='q'><div class='qh'>Q{q['id']} <span class='tag'>{esc(q.get('type',''))}</span> {btn}</div>
        <div class='qt'>{esc(q['question'])}</div><div class='opts'>{opts}</div>
        <div class='ans hide' id='ca{q['id']}'>✅ {esc(q['answer'])} — {esc(q.get('rationale',''))}</div>
        <button class='mini' onclick='toggle(\"ca{q['id']}\")'>对答案</button></div>"""
    dic = ""
    for x in listening.get('dictation', []):
        sp = f" ({esc(x.get('speaker'))})" if x.get('speaker') else ""
        dic += f"""<div class='q dcard' id='dc{x['id']}'><div class='qh'>听写 {x['id']} {esc(sp)} <button class='ap' onclick='playTxt(\"d{x['id']}\")'>🔊 听</button></div>
        <div class='qt'>{esc(x['blanked'])}</div>
        <input class='fill' id='df{x['id']}' placeholder='听写填空…'>
        <div class='ans hide' id='da{x['id']}'>✅ {esc(' / '.join(x['answers']))}</div>
        <button class='mini' onclick='toggle(\"da{x['id']}\")'>对答案</button>
        <div class='srs'><button class='again' onclick='grade(\"dc{x['id']}\",0)'>没记住</button><button class='good' onclick='grade(\"dc{x['id']}\",1)'>记住了</button></div></div>"""

    # ---- Listen tab: minimal pairs (phoneme-level) ----
    mp = ""
    for p in listening.get('minimal_pairs', []):
        actual = p.get('in_episode', '')
        actual_txt = ("A" if actual == 'a' else "B" if actual == 'b' else "?")
        mp += f"""<div class='q'><div class='qh'>最小对立体 {p['id']}</div>
        <div class='note'>A: <b>{esc(p['word_a'])}</b> <span class='ip'>{esc(p.get('ipa_a',''))}</span> &nbsp;|&nbsp; B: <b>{esc(p['word_b'])}</b> <span class='ip'>{esc(p.get('ipa_b',''))}</span></div>
        <div class='qt'>原句：“{esc(p.get('line',''))}”</div>
        {f"<div class='note warn'>🔤 辨音提示：{esc(p['hint'])}</div>" if p.get('hint') else ""}
        <div class='ans hide' id='mp{p['id']}'>✅ 本集实际说的是 <b>{actual_txt}</b>（{esc(p.get('word_a' if actual=='a' else 'word_b',''))}）</div>
        <button class='mini' onclick='toggle(\"mp{p['id']}\")'>听完后对答案</button></div>"""

    # ---- Listen tab: connected speech breakdown ----
    cs = ""
    for c in listening.get('connected_speech', []):
        brk = "".join(f"<div class='note'>▶ {esc(b.get('text',''))} <span class='note ok'>— {esc(b.get('note',''))}</span></div>" for b in c.get('breakdown', []))
        cs += f"""<div class='q'><div class='qh'>连读拆解 {c['id']}</div>
        <div class='qt'>“{esc(c.get('line',''))}”</div>
        {f"<div class='note'>💡 {esc(c['gloss'])}</div>" if c.get('gloss') else ""}
        <div class='note ok'>自然语速读法：</div>{brk}</div>"""

    # audio data for lines (comprehension + dictation)
    line_audio_js = {}
    for q in listening.get('comprehension', []):
        if q.get('audio_line'):
            p = audio_cache.get(q['audio_line'].strip().lower())
            if p: line_audio_js[f"t{q['id']}"] = b64(p)
    for x in listening.get('dictation', []):
        p = audio_cache.get(x['line'].strip().lower())
        if p: line_audio_js[f"d{x['id']}"] = b64(p)
    for a in annotated.get('annotations', []):
        p = audio_cache.get(a['line'].strip().lower())
        if p: line_audio_js[f"a{a['id']}"] = b64(p)
    for s in tasks.get('speaking', []):
        if s.get('model_line'):
            p = audio_cache.get(s['model_line'].strip().lower())
            if p: line_audio_js[f"s{s['id']}"] = b64(p)
    la_js = json.dumps(line_audio_js)

    # ---- Read tab ----
    ann = ""
    for a in annotated.get('annotations', []):
        sp = f" ({esc(a.get('speaker'))})" if a.get('speaker') else ""
        ann += f"""<div class='q'><div class='qh'><span class='tag'>{esc(a['focus'])}</span> <button class='ap' onclick='playTxt(\"a{a['id']}\")'>🔊</button></div>
        <div class='qt'>“{esc(a['line'])}”{sp}</div>
        {f"<div class='note'>📐 <b>规则：</b>{esc(a['rule'])}</div>" if a.get('rule') else ""}
        <div class='note'>💡 {esc(a['note'])}</div><div class='note ok'>✅ {esc(a['tip'])}</div>
        {f"<div class='note'>➕ <b>更多例句：</b>{esc(a['more'])}</div>" if a.get('more') else ""}</div>"""
    clo = ""
    for c in annotated.get('cloze', []):
        clo += f"""<div class='q ccard' id='cc{c['id']}'><div class='qh'>Cloze {c['id']}</div>
        <div class='qt'>{esc(c['blanked'])}</div>
        <div class='ans hide' id='clo{c['id']}'>✅ {esc(' / '.join(c['answers']))}</div>
        <button class='mini' onclick='toggle(\"clo{c['id']}\")'>对答案</button>
        <div class='srs'><button class='again' onclick='grade(\"cc{c['id']}\",0)'>没记住</button><button class='good' onclick='grade(\"cc{c['id']}\",1)'>记住了</button></div></div>"""

    # ---- Speak tab ----
    spk = ""
    for s in tasks.get('speaking', []):
        ch = f" as {esc(s['character'])}" if s.get('character') else ""
        ml = f"<div class='note'>🎧 Model: “{esc(s['model_line'])}” <button class='ap' onclick='playTxt(\"s{s['id']}\")'>🔊</button></div>" if s.get('model_line') else ""
        spk += f"""<div class='q'><div class='qh'>{esc(s['type'])}{ch}</div>
        <div class='qt'>{esc(s['instruction'])}</div>{ml}
        <div class='note'>Use: <b>{esc(', '.join(s.get('use_words', [])))}</b></div>
        <div class='note ok'>Checklist: {esc(', '.join(s.get('checklist', [])))}</div>
        {f"<div class='note'>🔤 重点发音：{esc(', '.join(s.get('focus_sounds', [])))}</div>" if s.get('focus_sounds') else ""}
        {f"<div class='note ok'>🎤 可评分：录下你说的话，运行 <code>score_speaking.py --audio 你的录音 --target \"{esc(s['asr_target'])}\"</code> 获取 Whisper 转写比对（标记丢失/错误词，非发音评分）。</div>" if s.get('asr_target') else ""}
        <div class='note warn'>{esc(U['tts_note'])}</div></div>"""

    # ---- Write tab ----
    wrt = ""
    for w in tasks.get('writing', []):
        mdl = f"<div class='note'>📝 Model: {esc(w['model'])}</div>" if w.get('model') else ""
        wrt += f"""<div class='q'><div class='qh'>{esc(w.get('type',''))} / {esc(w.get('register',''))}</div>
        <div class='qt'>{esc(w['instruction'])}</div>
        <div class='note'>Must use: <b>{esc(', '.join(w.get('require_words', [])))}</b></div>
        <div class='note ok'>Rubric: {esc('; '.join(w.get('rubric', [])))}</div>{mdl}
        {f"<div class='note ok'>🤖 自动量规：把作文存为 essay.txt，运行 <code>score_writing.py --task {w.get('id','')} --text essay.txt --tasks tasks.json</code> 自动校验（目标词命中/字数/时态）。</div>" if w.get('checks') else ""}
        <div class='note warn'>{esc(U['tts_warn2'])}</div></div>"""

    # ---- Review tab: vocab flip + reflection + exit ----
    vcards = ""
    for i, w in enumerate(words):
        ta = b64_local(w.get('term_audio')) if w.get('term_audio') else ""
        la = b64_local(w.get('line_audio')) if w.get('line_audio') else ""
        tab = f"<button class='ap' onclick='playb64(\"{ta}\")'>🔊 词</button>" if ta else ""
        lab = f"<button class='ap' onclick='playb64(\"{la}\")'>🔊 原句</button>" if la else ""
        vcards += f"""<div class='vcard' id='vc{i}'><div class='vf'>{esc(w['term'])} <span class='ip'>{esc(w.get('ipa',''))}</span></div>
        <div class='vb hide'>{vocab_back(w)}</div><div class='aud'>{tab} {lab}</div>
        <div class='srs'><button class='again' onclick='grade(\"vc{i}\",0)'>没记住</button><button class='good' onclick='grade(\"vc{i}\",1)'>记住了</button></div></div>"""

    # 口语/写作产出卡（纳入本地 SRS，确保 HTML 用户也有说/写复习入口，与 Anki 模式互补）
    pcards = ""
    for i, s in enumerate(tasks.get('speaking', [])):
        ch = f" as {esc(s['character'])}" if s.get('character') else ""
        ml = f"<div class='note'>🎧 Model: “{esc(s['model_line'])}” <button class='ap' onclick='playTxt(\"s{s['id']}\")'>🔊</button></div>" if s.get('model_line') else ""
        pcards += f"""<div class='q pcard' id='ps{i}'><div class='qh'>{esc(s['type'])}{ch}</div>
        <div class='qt'>{esc(s['instruction'])}</div>{ml}
        <div class='note'>Use: <b>{esc(', '.join(s.get('use_words', [])))}</b></div>
        <div class='note ok'>Checklist: {esc(', '.join(s.get('checklist', [])))}</div>
        {f"<div class='note'>🔤 重点发音：{esc(', '.join(s.get('focus_sounds', [])))}</div>" if s.get('focus_sounds') else ""}
        <div class='srs'><button class='again' onclick='grade(\"ps{i}\",0)'>未完成</button><button class='good' onclick='grade(\"ps{i}\",1)'>已完成</button></div></div>"""
    for j, w in enumerate(tasks.get('writing', [])):
        mdl = f"<div class='note'>📝 Model: {esc(w['model'])}</div>" if w.get('model') else ""
        pcards += f"""<div class='q wcard' id='pw{j}'><div class='qh'>{esc(w['type'])} / {esc(w['register'])}</div>
        <div class='qt'>{esc(w['instruction'])}</div>
        <div class='note'>Must use: <b>{esc(', '.join(w.get('require_words', [])))}</b></div>
        <div class='note ok'>Rubric: {esc('; '.join(w.get('rubric', [])))}</div>{mdl}
        <div class='srs'><button class='again' onclick='grade(\"pw{j}\",0)'>未完成</button><button class='good' onclick='grade(\"pw{j}\",1)'>已完成</button></div></div>"""

    # ---- Cross-episode recall hints ----
    rec = ""
    h3_recall = ""
    if recall and recall.get('hints'):
        rec = "<div class='note'>旧语境 → 新语境：把已学词再撞见一次，复用记忆更牢（Webb & Rodgers）。</div>"
        for h in recall['hints']:
            ctx = "".join(f"<div class='note'>▶ “{esc(c)}”</div>" for c in h.get('new_contexts', []))
            rec += f"<div class='q'><div class='qh'>{esc(h['term'])} <i>{esc(h.get('cefr',''))}</i></div>" \
                   f"<div class='note'>旧语境：{esc(h.get('old_context',''))}</div>{ctx}</div>"
        h3_recall = f"<h3>🔁 跨集复现（{recall.get('recalled', len(recall['hints']))} 个已学词的新语境）</h3>"

    refl = f"""<div class='q'><div class='qh'>{U['pre']}</div><textarea class='ta' id='pre' placeholder='{esc(U['pre_ph'])}'></textarea></div>
    <div class='q'><div class='qh'>{U['refl']}</div><textarea class='ta' id='refl' placeholder='{esc(U['refl_ph'])}'></textarea></div>
    <div class='note ok'>✓ 目标与反思自动保存在本机浏览器，下次打开仍在。</div></div>"""

    exitq = []
    if listening.get('comprehension'):
        q = random.choice(listening['comprehension'])
        exitq.append(("听", q['question'], q['answer']))
    if annotated.get('cloze'):
        c = random.choice(annotated['cloze'])
        exitq.append(("读", c['blanked'], " / ".join(c['answers'])))
    if tasks.get('speaking'):
        s = random.choice(tasks['speaking'])
        exitq.append(("说", s['instruction'], "用: " + ", ".join(s.get('use_words', []))))
    exith = ""
    for j, (sk, q, a) in enumerate(exitq):
        exith += f"""<div class='q'><div class='qh'>出口小测 · {sk}</div><div class='qt'>{esc(q)}</div>
        <div class='ans hide' id='ex{j}'>✅ {esc(a)}</div><button class='mini' onclick='toggle(\"ex{j}\")'>对答案</button></div>"""

    # ---- Prime tab: full target-lexicon reference (i+1 priming first) ----
    prime_cards = ""
    for w in words:
        ta = b64_local(w.get('term_audio')) if w.get('term_audio') else ""
        la = b64_local(w.get('line_audio')) if w.get('line_audio') else ""
        tab = f"<button class='ap' onclick='playb64(\"{ta}\")'>🔊 词</button>" if ta else ""
        lab = f"<button class='ap' onclick='playb64(\"{la}\")'>🔊 原句</button>" if la else ""
        _exam = w.get('exam') or (exm.exam_label(w['cefr']) if w.get('cefr') else '')
        prime_cards += f"""<div class='q'><div class='qh'>{esc(w['term'])} <span class='ip'>{esc(w.get('ipa',''))}</span> <i>{esc(w.get('cefr',''))}</i></div>
        <div class='qt'>{esc(w.get('gloss',''))}</div>
        {('<div class="note">🎓 ' + esc(_exam) + '</div>') if _exam else ''}
        <div class='note'>搭配: {esc(w.get('collocation',''))}</div>
        <div class='note'>例句: {esc(w.get('example',''))}</div>
        <div class='note'>原句: {esc(w.get('line',''))}</div>
        <div class='aud'>{tab} {lab}</div></div>"""
    # ---- Watch tab: viewing protocol + catch-expression slot ----
    wp = watch or {}
    watch_html = f"<div class='q'><div class='qh'>{U['watch_h']}</div>"
    for p in wp.get('protocol', []):
        watch_html += f"<div class='note'><b>第{p.get('pass','?')}遍 · {esc(p.get('name',''))}</b>：{esc(p.get('instruction',''))}</div>"
    for n in wp.get('notice', []):
        watch_html += f"<div class='note warn'>⚠️ {esc(n)}</div>"
    watch_html += f"<div class='note'>{U['catch_label']}</div><textarea class='ta' id='catch' placeholder='catch-expression…'></textarea></div>"

    html_doc = f"""<!DOCTYPE html><html lang='zh'><head><meta charset='utf-8'>
<meta name='viewport' content='width=device-width,initial-scale=1'>
<title>DramaLex · {esc(deck)}</title>
<style>
*{{box-sizing:border-box}} body{{font-family:-apple-system,'PingFang SC',Segoe UI,sans-serif;background:#f5f6f8;margin:0;color:#1c1c1e;padding:14px}}
h1{{font-size:18px;margin:0}} .sub{{color:#666;font-size:12px;margin:2px 0 10px}}
.chips{{margin:6px 0 6px}} .chip{{display:inline-block;background:#eef3ff;color:#0a55d6;border-radius:20px;padding:3px 10px;font-size:12px;margin:2px}}
.habit{{display:inline-block;background:#fff3e0;color:#b06b00;border-radius:20px;padding:5px 12px;font-size:12px;margin:0 0 8px;font-weight:600}}
.week{{margin:0 0 10px;font-size:12px;color:#666}} .week .wc{{display:inline-block;min-width:20px;text-align:center;margin:0 2px;padding:3px 0;border-radius:6px;background:#eee;color:#999}}
.week .wc.on{{background:#1a7f37;color:#fff;font-weight:700}}
.tabs{{display:flex;flex-wrap:wrap;gap:4px;position:sticky;top:0;background:#f5f6f8;padding:6px 0;z-index:5}}
.tab{{border:none;background:#fff;border-radius:8px;padding:7px 12px;font-size:13px;cursor:pointer}}
.tab.on{{background:#0a84ff;color:#fff}}
.panel{{display:none}} .panel.on{{display:block}}
.q{{background:#fff;border-radius:12px;padding:14px;margin:10px 0;box-shadow:0 1px 3px rgba(0,0,0,.07)}}
.qh{{font-weight:700;font-size:14px;margin-bottom:6px}} .qt{{font-size:15px;margin:4px 0;line-height:1.5}}
.tag{{background:#ffe9c7;color:#9a5b00;border-radius:6px;padding:1px 7px;font-size:11px;margin-right:4px}}
.note{{font-size:13px;margin:4px 0;line-height:1.5}} .note.ok{{color:#1a7f37}} .note.warn{{color:#b06b00}}
.exam{{display:inline-block;background:#f0ecff;color:#5b2bd6;border-radius:6px;padding:1px 7px;font-size:11px;margin-left:4px}}
.opts .opt{{display:block;font-size:14px;margin:3px 0}} .fill{{width:100%;padding:8px;border:1px solid #ccc;border-radius:8px;font-size:14px;margin-top:4px}}
.ans{{color:#1a7f37;font-size:14px;margin-top:4px}} .mini{{margin-top:6px;border:1px solid #0a84ff;background:#fff;color:#0a84ff;border-radius:8px;padding:5px 10px;font-size:12px;cursor:pointer}}
.ap{{border:1px solid #0a84ff;background:#fff;color:#0a84ff;border-radius:8px;padding:5px 9px;font-size:12px;margin-right:5px;cursor:pointer}}
.ta{{width:100%;min-height:60px;border:1px solid #ccc;border-radius:8px;padding:8px;font-size:14px}}
.vcard{{background:#fff;border-radius:12px;padding:14px;margin:10px 0;box-shadow:0 1px 3px rgba(0,0,0,.07);cursor:pointer}}
.vf{{font-size:20px;font-weight:700}} .ip{{color:#888;font-size:14px;font-weight:400}}
.aud{{margin:8px 0}} .srs{{display:flex;gap:8px;margin-top:6px}}
.again,.good{{flex:1;border:none;border-radius:8px;padding:9px;font-weight:600;cursor:pointer}}
.again{{background:#ffe5e5;color:#c0392b}} .good{{background:#e3f9e5;color:#1a7f37}}
.foot{{text-align:center;color:#888;font-size:12px;margin-top:18px;line-height:1.6}}
</style></head><body>
<h1>{esc(U['title'])} · {esc(deck)}</h1>
<div class='sub'>{esc(U['sub'])}</div>
<div class='chips'>{chips}</div>
<div id='habit' class='habit'>🔥 连续 1 天 · 📚 今日待复习 0 · 📅 明日 0</div>
<div id='week' class='week'></div>
<div class='tabs'>
  <button class='tab on' onclick='show("prime",this)'>{U['tabs'][0]}</button>
  <button class='tab' onclick='show("watch",this)'>{U['tabs'][1]}</button>
  <button class='tab' onclick='show("listen",this)'>{U['tabs'][2]}</button>
  <button class='tab' onclick='show("read",this)'>{U['tabs'][3]}</button>
  <button class='tab' onclick='show("speak",this)'>{U['tabs'][4]}</button>
  <button class='tab' onclick='show("write",this)'>{U['tabs'][5]}</button>
  <button class='tab' onclick='show("review",this)'>{U['tabs'][6]}</button>
</div>
<div class='panel on' id='prime'><h3>{U['prime_h']}</h3>{prime_cards}</div>
<div class='panel' id='watch'>{watch_html}</div>
<div class='panel' id='listen'><div class='note warn'>{U['listen_warn']}</div>
<h3>{U['comp_h']}</h3>{comp}<h3>{U['dict_h']}</h3>{dic}
<h3>{U['mp_h']}</h3>{mp}<h3>{U['cs_h']}</h3>{cs}</div>
<div class='panel' id='read'>{ann}<h3>{U['cloze_h']}</h3>{clo}</div>
<div class='panel' id='speak'>{spk}</div>
<div class='panel' id='write'>{wrt}</div>
<div class='panel' id='review'><h3>{U['review_h']}</h3>{vcards}
<h3>{esc(U['prod_h'])}</h3>{pcards}
{h3_recall}{rec}
<div class='note'>📦 {esc(U['prog_export'])} / {esc(U['prog_import'])}（换设备/清缓存不丢复习进度，单文件离线可用）：
<button class='mini' onclick='exportProgress()'>{esc(U['prog_export'])}</button>
<input type='file' id='imp' accept='application/json' onchange='importProgress(this)' style='display:inline-block;font-size:12px'></div>
<h3>{U['meta_h']}</h3>{refl}<h3>{U['exit_h']}</h3>{exith}</div>
<div class='foot'>{esc(AUTHOR)}<br>{esc(CONTACT)}<br><sub>{esc(U['foot'])}</sub></div>
<script>
const LA={la_js};
const WK={json.dumps(U['week_lbl'])}; const WKPRE={json.dumps(U['week_h'])};
function playTxt(k){{ if(LA[k]) new Audio(LA[k]).play(); }}
function playb64(u){{ if(u) new Audio(u).play(); }}
function toggle(id){{ document.getElementById(id).classList.toggle('hide'); }}
function show(p,el){{ document.querySelectorAll('.panel').forEach(p=>p.classList.remove('on'));
  document.getElementById(p).classList.add('on');
  document.querySelectorAll('.tab').forEach(t=>t.classList.remove('on')); el.classList.add('on'); }}
const DK={json.dumps(deck)}; const KEY='dlex_'+JSON.stringify(DK);
let DB=JSON.parse(localStorage.getItem(KEY)||'null')||{{cards:{{}},streak:0,last:'',days:[]}};
function saveDB(){{ localStorage.setItem(KEY, JSON.stringify(DB)); }}
function ymd(d){{ return d.getFullYear()+'-'+(d.getMonth()+1)+'-'+d.getDate(); }}
function updateStreak(){{ const t=ymd(new Date()); if(DB.last!==t){{ const y=new Date(); y.setDate(y.getDate()-1);
  if(DB.last===ymd(y)) DB.streak=(DB.streak||0)+1; else DB.streak=1; DB.last=t;
  const ds=DB.days||[]; if(!ds.includes(t)){{ ds.push(t); DB.days=ds; }} saveDB(); }} }}
function updateHabit(){{ const now=Date.now(), DAY=86400000; let dT=0,dM=0;
  document.querySelectorAll('.vcard,.dcard,.ccard,.pcard,.wcard').forEach(function(el){{ const c=DB.cards[el.id];
    if(!c){{ dT++; return; }} if(c.due<=now) dT++; else if(c.due<=now+DAY) dM++; }});
  const b=document.getElementById('habit'); if(b){{ let m='🔥 连续 '+(DB.streak||0)+' 天 · 📚 今日待复习 '+dT+' · 📅 明日 '+dM;
    if(dT===0) m+=' · ✅ 今日打卡完成'; b.textContent=m; }} }}
function updateWeek(){{ const el=document.getElementById('week'); if(!el) return;
  const days=DB.days||[]; const today=new Date(); let h=WKPRE+' ';
  for(let i=6;i>=0;i--){{ const d=new Date(today); d.setDate(d.getDate()-i); const key=ymd(d);
    const on=days.includes(key); h+="<span class='wc "+(on?'on':'')+"'>"+(on?'✓':WK[(d.getDay()+6)%7])+"</span>"; }}
  el.innerHTML=h; }}
function exportProgress(){{ const blob=new Blob([JSON.stringify(DB)],{{type:'application/json'}});
  const a=document.createElement('a'); a.href=URL.createObjectURL(blob); a.download=KEY+'.json'; a.click(); }}
function importProgress(inp){{ const f=inp.files[0]; if(!f) return; const r=new FileReader();
  r.onload=function(){{ try{{ const d=JSON.parse(r.result); if(d&&d.cards){{ DB=d; DB.days=DB.days||[]; saveDB(); location.reload(); }} }}
    catch(e){{ alert('导入失败：'+(e.message||e)); }} }}; r.readAsText(f); }}
function grade(el,g){{ let c=DB.cards[el]||{{iv:0,due:0}}; c.iv = g? (c.iv>=1? c.iv*2:1):0;
  c.due = Date.now()+(g? c.iv*86400000:600000); DB.cards[el]=c; saveDB();
  const node=document.getElementById(el); if(node) node.classList.add('hide'); updateHabit(); }}
document.querySelectorAll('.vcard').forEach(c=>c.onclick=e=>{{ if(e.target.tagName==='BUTTON')return;
        const f=c.querySelector('.vf'),b=c.querySelector('.vb'); f.classList.toggle('hide'); b.classList.toggle('hide'); }});
['pre','refl','catch'].forEach(function(id){{var el=document.getElementById(id); if(el){{var k='dlex_'+id; el.value=localStorage.getItem(k)||''; el.addEventListener('input',function(){{localStorage.setItem(k,el.value);}});}}}});
updateStreak(); updateHabit(); updateWeek();
</script></body></html>"""
    return html_doc

# ---------- Anki recall builder ----------
def build_recall(words, listening, annotated, tasks, audio_cache, mode):
    rows = []  # (front, back_html, audio_sound, tags)
    for w in words:
        front = f"[sound:{os.path.basename(w['term_audio'])}]" if w.get('term_audio') else esc(w['term'])
        back = vocab_back(w)
        audio = f"[sound:{os.path.basename(w['line_audio'])}]" if w.get('line_audio') else ""
        rows.append((front, back, audio, f"dramalex {w.get('cefr','').lower()} {w.get('type','')}"))
    for x in listening.get('dictation', []):
        p = audio_cache.get(x['line'].strip().lower())
        sound = f"[sound:{os.path.basename(p)}]" if p else ""
        back = f"<b>Dictation</b><br>{esc(x['line'])}<br>✅ {esc(' / '.join(x['answers']))}"
        rows.append((f"✍️ {esc(x['blanked'])}", back, sound, "dramalex dictation"))
    for q in listening.get('comprehension', []):
        p = audio_cache.get(q['audio_line'].strip().lower()) if q.get('audio_line') else None
        sound = f"[sound:{os.path.basename(p)}]" if p else ""
        back = f"<b>{esc(q['question'])}</b><br>✅ {esc(q['answer'])}<br>{esc(q['rationale'])}"
        rows.append((f"🎧 {esc(q['question'])}", back, sound, "dramalex listening"))
    for c in annotated.get('cloze', []):
        back = f"{esc(c['line'])}<br>✅ {esc(' / '.join(c['answers']))}"
        rows.append((f"📖 {esc(c['blanked'])}", back, "", "dramalex cloze"))
    if mode == 'B':
        for s in tasks.get('speaking', []):
            back = f"<b>{esc(s['instruction'])}</b><br>Use: {esc(', '.join(s.get('use_words',[])))}<br>Checklist: {esc(', '.join(s.get('checklist',[])))}"
            rows.append((f"🗣️ {esc(s['type'])}: {esc(s['instruction'])}", back, "", "dramalex speaking"))
        for w in tasks.get('writing', []):
            back = f"<b>{esc(w['instruction'])}</b><br>Must use: {esc(', '.join(w.get('require_words',[])))}<br>{esc(w.get('model',''))}"
            rows.append((f"✍️ {esc(w['type'])}: {esc(w['instruction'])}", back, "", "dramalex writing"))
    return rows

def build_anki(words, listening, annotated, tasks, cache, mode, deck, out_dir, media_dir):
    """Build a single <deck>.apkg with recall items. Returns path or None on failure."""
    try:
        import genanki
    except Exception as e:
        print("genanki 缺失:", e, file=sys.stderr); return None
    rows = build_recall(words, listening, annotated, tasks, cache, mode)
    notes = [genanki.Note(model=genanki.BASIC_MODEL, fields=[fr, bk]) for fr, bk, au, tg in rows]
    deck_obj = genanki.Deck(1234567890, deck + " · DramaLex")
    for n in notes: deck_obj.add_note(n)
    def resolve_audio(p):
        if not p:
            return None
        if os.path.exists(p):
            return p
        alt = os.path.join(media_dir, os.path.basename(p))
        return alt if os.path.exists(alt) else None
    media_files = []
    for w in words:
        a = resolve_audio(w.get('term_audio'))
        if a: media_files.append(a)
        a = resolve_audio(w.get('line_audio'))
        if a: media_files.append(a)
    for p in cache.values():
        if os.path.exists(p): media_files.append(p)
    apkg = os.path.join(out_dir, deck + '.apkg')
    genanki.Package(deck_obj, media_files=media_files).write_to_file(apkg)
    return apkg

def build_excel(words, listening, annotated, tasks, deck, out_dir, watch=None, ui_lang='zh', recall=None):
    """One file: <deck>.xlsx with one sheet per skill. Audio not embeddable in xlsx."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill
    except Exception as e:
        print("openpyxl 缺失:", e, file=sys.stderr); return None
    X = XL_UI.get(ui_lang, XL_UI['zh'])
    wb = Workbook()
    head_font = Font(bold=True, color="FFFFFF")
    head_fill = PatternFill("solid", fgColor="0A84FF")
    def sheet(title, headers, rows):
        ws = wb.create_sheet(title)
        ws.append(headers)
        for c in ws[1]:
            c.font = head_font; c.fill = head_fill
            c.alignment = Alignment(vertical="center")
        for r in rows:
            ws.append([str(x) for x in r])
        for col in ws.columns:
            width = max((len(str(c.value)) for c in col if c.value is not None), default=10)
            ws.column_dimensions[col[0].column_letter].width = min(max(width + 2, 12), 60)
        ws.freeze_panes = "A2"
    # 词汇卡
    sheet(*X['vocab'],
          [[w['term'],w.get('type',''),w.get('cefr',''),
            w.get('exam') or (exm.exam_label(w['cefr']) if w.get('cefr') else ''),
            w.get('ipa',''),w.get('pos',''),
            w.get('gloss',''),w.get('collocation',''),w.get('line',''),w.get('example',''),
            w.get('why',''),w.get('l1_note',''),
            ", ".join(w.get('tags',[]))] for w in words])
    # 观看（三步字幕法）
    sheet(*X['watch'],
          [[p.get('pass',''), p.get('name',''), p.get('instruction','')] for p in (watch or {}).get('protocol', [])])
    # 听力理解
    sheet(*X['listen'],
          [[q['id'],q.get('type',''),q['question']," | ".join(q.get('options',[])),
            q.get('answer',''),q.get('rationale','')] for q in listening.get('comprehension',[])])
    # 听写
    sheet(*X['dict'],
          [[x['id'],x.get('blanked','')," / ".join(x.get('answers',[])),
            ", ".join(x.get('target_words',[]))] for x in listening.get('dictation',[])])
    # 最小对立体
    sheet(*X['mp'],
          [[p['id'],p.get('word_a',''),p.get('ipa_a',''),p.get('word_b',''),p.get('ipa_b',''),
            ("A" if p.get('in_episode')=='a' else "B" if p.get('in_episode')=='b' else ""),
            p.get('hint','')] for p in listening.get('minimal_pairs',[])])
    # 连读拆解
    sheet(*X['cs'],
          [[c['id'],c.get('line',''),
            " | ".join(f"{b.get('text','')}（{b.get('note','')}）" for b in c.get('breakdown',[])),
            c.get('gloss','')] for c in listening.get('connected_speech',[])])
    # 精读标注
    sheet(*X['anno'],
          [[a['id'],a.get('focus',''),a.get('line',''),a.get('rule',''),a.get('note',''),a.get('tip',''),a.get('more','')]
           for a in annotated.get('annotations',[])])
    # 完形
    sheet(*X['cloze'],
          [[c['id'],c.get('blanked','')," / ".join(c.get('answers',[]))] for c in annotated.get('cloze',[])])
    # 口语
    sheet(*X['speak'],
          [[s['id'],s.get('type',''),s.get('instruction',''),", ".join(s.get('use_words',[])),
            ", ".join(s.get('focus_sounds',[])),", ".join(s.get('checklist',[])),
            s.get('asr_target','')] for s in tasks.get('speaking',[])])
    # 写作
    sheet(*X['write'],
          [[w['id'],w.get('type',''),w.get('register',''),w.get('instruction',''),
            ", ".join(w.get('require_words',[])),"; ".join(w.get('rubric',[])),w.get('model',''),
            "; ".join(f"{c.get('type')}:{c.get('value')}" for c in (w.get('checks') or []))]
           for w in tasks.get('writing',[])])
    # 复习
    ws = wb.create_sheet(X['review'])
    ws.append(["目标词链"]); ws["A1"].font = Font(bold=True)
    for w in words:
        ws.append([f"{w['term']} ({w.get('cefr','')})"])
    ws.append([]); ws.append(["元认知"]); ws["A" + str(ws.max_row)].font = Font(bold=True)
    ws.append([X['pre']])
    ws.append([X['post']])
    if recall and recall.get('hints'):
        ws.append([]); ws.append(["🔁 跨集复现"]); ws["A" + str(ws.max_row)].font = Font(bold=True)
        ws.append(["词", "CEFR", "旧语境", "新语境(本集)"])
        for h in recall['hints']:
            ws.append([h['term'], h.get('cefr', ''), h.get('old_context', ''),
                       " | ".join(h.get('new_contexts', []))])
    ws.append([]); ws.append([X['review_note']])
    ws.append([])
    ws.append([f"👨‍💻 {AUTHOR} · {CONTACT}"])
    ws.append([LEGAL])
    ws.column_dimensions["A"].width = 60
    # remove default empty sheet
    if wb.sheetnames[0] == "Sheet":
        del wb["Sheet"]
    path = os.path.join(out_dir, deck + ".xlsx")
    wb.save(path)
    return path

def build_word(words, listening, annotated, tasks, deck, out_dir, watch=None, ui_lang='zh', recall=None):
    """One file: <deck>.docx — a readable four-skill report."""
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as e:
        print("python-docx 缺失:", e, file=sys.stderr); return None
    D = WD_UI.get(ui_lang, WD_UI['zh'])
    doc = Document()
    doc.add_heading(f"DramaLex · {deck}", 0)
    doc.add_paragraph(D['title_sub']).italic = True
    doc.add_paragraph("👨‍💻 yinjianheng（殷健恒） · yinjianheng@foxmail.com · WeChat: YJH-yinjianheng").runs[0].font.size = Pt(9)
    doc.add_paragraph(LEGAL).runs[0].font.size = Pt(8)

    def table(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = str(h)
        for r in rows:
            cells = t.add_row().cells
            for i, v in enumerate(r):
                cells[i].text = str(v)
    # 词汇卡
    doc.add_heading(D['sec'][0], 1)
    table(D['vocab_hdr'],
          [[w['term'],w.get('cefr',''),w.get('ipa',''),w.get('gloss',''),w.get('collocation',''),w.get('line',''),w.get('why',''),w.get('l1_note','')] for w in words])
    doc.add_heading(D['sec'][1], 1)
    for p in (watch or {}).get('protocol', []):
        doc.add_paragraph(f"第{p.get('pass','?')}遍 · {p.get('name','')}：{p.get('instruction','')}")
    for n in (watch or {}).get('notice', []):
        doc.add_paragraph(f"⚠️ {n}")
    # 听
    doc.add_heading(D['sec'][2], 1)
    doc.add_heading(D['listen_sub'], 2)
    table(D['listen_hdr'],
          [[q['question']," | ".join(q.get('options',[])),q.get('answer',''),q.get('rationale','')] for q in listening.get('comprehension',[])])
    doc.add_heading(D['dict_sub'], 2)
    table(D['dict_hdr'],
          [[x.get('blanked','')," / ".join(x.get('answers',[])),", ".join(x.get('target_words',[]))] for x in listening.get('dictation',[])])
    doc.add_heading(D['mp_sub'], 2)
    table(D['mp_hdr'],
          [[p.get('word_a',''),p.get('ipa_a',''),p.get('word_b',''),p.get('ipa_b',''),
            ("A" if p.get('in_episode')=='a' else "B" if p.get('in_episode')=='b' else ""),p.get('hint','')] for p in listening.get('minimal_pairs',[])])
    doc.add_heading(D['cs_sub'], 2)
    table(D['cs_hdr'],
          [[c.get('line','')," | ".join(f"{b.get('text','')}（{b.get('note','')}）" for b in c.get('breakdown',[])),c.get('gloss','')] for c in listening.get('connected_speech',[])])
    # 读
    doc.add_heading(D['sec'][3], 1)
    table(D['anno_hdr'],
          [[a.get('focus',''),a.get('line',''),a.get('rule',''),a.get('note',''),a.get('tip',''),a.get('more','')] for a in annotated.get('annotations',[])])
    doc.add_heading(D['cloze_sub'], 2)
    table(D['cloze_hdr'], [[c.get('blanked','')," / ".join(c.get('answers',[]))] for c in annotated.get('cloze',[])])
    # 说
    doc.add_heading(D['sec'][4], 1)
    table(D['speak_hdr'],
          [[s.get('type',''),s.get('instruction',''),", ".join(s.get('use_words',[])),", ".join(s.get('focus_sounds',[])),", ".join(s.get('checklist',[])),s.get('asr_target','')] for s in tasks.get('speaking',[])])
    # 写
    doc.add_heading(D['sec'][5], 1)
    table(D['write_hdr'],
          [[w.get('type',''),w.get('register',''),w.get('instruction',''),", ".join(w.get('require_words',[])),
            "; ".join(w.get('rubric',[])),w.get('model',''),"; ".join(f"{c.get('type')}:{c.get('value')}" for c in (w.get('checks') or []))] for w in tasks.get('writing',[])])
    # 复习
    doc.add_heading(D['sec'][6], 1)
    doc.add_paragraph("目标词链：" + "、".join(w['term'] for w in words))
    doc.add_paragraph(D['pre'])
    doc.add_paragraph(D['post'])
    if recall and recall.get('hints'):
        doc.add_heading("🔁 跨集复现", 2)
        for h in recall['hints']:
            doc.add_paragraph(f"{h['term']}（{h.get('cefr','')}）：旧「{h.get('old_context','')}」")
            for c in h.get('new_contexts', []):
                doc.add_paragraph(f"  ▶ 新：「{c}」", style="List Bullet")
    doc.add_paragraph(D['note']).italic = True
    path = os.path.join(out_dir, deck + ".docx")
    doc.save(path)
    return path

def build_markdown(words, listening, annotated, tasks, deck, out_dir, watch=None, ui_lang='zh', recall=None):
    """One file: <deck>.md — Obsidian / 双链笔记友好，纯文本、可检索、可互相链接。"""
    X = XL_UI.get(ui_lang, XL_UI['zh'])
    L = {
        'zh': {
            'h_vocab':'## 0 · 词汇卡（目标词链）', 'h_watch':'## 1 · 看（观看协议）',
            'h_listen':'## 2 · 听', 'h_listen_c':'### 听力理解', 'h_dict':'### 听写',
            'h_mp':'### 最小对立体（音素级辨音）', 'h_cs':'### 连读拆解（自然语速读法）',
            'h_read':'## 3 · 读 · 台词精读', 'h_cloze':'### 完形',
            'h_speak':'## 4 · 说', 'h_write':'## 5 · 写', 'h_review':'## 6 · 复习',
            'exam':'考试对照', 'gloss':'释义', 'colloc':'搭配', 'line':'原句', 'example':'例句',
            'pre':'学前目标：', 'post':'学后反思：', 'note':'说明：音频请配合 practice.html / deck.apkg。',
        },
        'en': {
            'h_vocab':'## 0 · Vocab (target lexicon)', 'h_watch':'## 1 · Watch (protocol)',
            'h_listen':'## 2 · Listen', 'h_listen_c':'### Listening comprehension', 'h_dict':'### Dictation',
            'h_mp':'### Minimal pairs (phoneme ear-training)', 'h_cs':'### Connected speech breakdown',
            'h_read':'## 3 · Read (annotation)', 'h_cloze':'### Cloze',
            'h_speak':'## 4 · Speak', 'h_write':'## 5 · Write', 'h_review':'## 6 · Review',
            'exam':'Exam', 'gloss':'Gloss', 'colloc':'Colloc', 'line':'Line', 'example':'Example',
            'pre':'Pre-goal: ', 'post':'Post-reflection: ', 'note':'Note: play audio via practice.html / deck.apkg.',
        },
    }[ui_lang]
    lines = [f"# DramaLex · {deck}", "", f"> {L['note']}", ""]
    # 词汇
    lines += [L['h_vocab'], ""]
    for w in words:
        exam = w.get('exam') or (exm.exam_label(w['cefr']) if w.get('cefr') else '')
        tag = f" `{w.get('cefr','')}`" if w.get('cefr') else ""
        ex = f" · {L['exam']}: {exam}" if exam else ""
        lines.append(f"- **{w['term']}**{tag}{ex} — {w.get('gloss','')}")
        if w.get('ipa'): lines.append(f"  - IPA: {w['ipa']}")
        if w.get('collocation'): lines.append(f"  - {L['colloc']}: {w['collocation']}")
        if w.get('why'): lines.append(f"  - 🎯 为什么学：{w['why']}")
        if w.get('l1_note'): lines.append(f"  - 🇨🇳 中文易错：{w['l1_note']}")
        if w.get('line'): lines.append(f"  - {L['line']}: {w['line']}")
        if w.get('example'): lines.append(f"  - {L['example']}: {w['example']}")
    # 看
    lines += ["", L['h_watch'], ""]
    for p in (watch or {}).get('protocol', []):
        lines.append(f"- 第{p.get('pass','?')}遍 · {p.get('name','')}：{p.get('instruction','')}")
    for n in (watch or {}).get('notice', []):
        lines.append(f"- ⚠️ {n}")
    # 听
    lines += ["", L['h_listen'], "", L['h_listen_c'], ""]
    for q in listening.get('comprehension', []):
        opts = " / ".join(q.get('options', []))
        lines.append(f"- Q{q.get('id')}（{q.get('type','')}）：{q.get('question','')}")
        if opts: lines.append(f"  - 选项：{opts}")
        lines.append(f"  - ✅ {q.get('answer','')} — {q.get('rationale','')}")
    lines += ["", L['h_dict'], ""]
    for x in listening.get('dictation', []):
        lines.append(f"- 听写 {x.get('id')}：{x.get('blanked','')} → ✅ {' / '.join(x.get('answers',[]))}")
    lines += ["", L['h_mp'], ""]
    for p in listening.get('minimal_pairs', []):
        actual = ("A" if p.get('in_episode')=='a' else "B" if p.get('in_episode')=='b' else "?")
        lines.append(f"- 最小对立体 {p.get('id')}：A **{p.get('word_a','')}** {p.get('ipa_a','')} | B **{p.get('word_b','')}** {p.get('ipa_b','')}")
        if p.get('line'): lines.append(f"  - 原句：{p['line']}")
        if p.get('hint'): lines.append(f"  - 🔤 辨音提示：{p['hint']}")
        lines.append(f"  - ✅ 本集实际：{actual}（{p.get('word_a' if p.get('in_episode')=='a' else 'word_b','')}）")
    lines += ["", L['h_cs'], ""]
    for c in listening.get('connected_speech', []):
        lines.append(f"- 连读拆解 {c.get('id')}：“{c.get('line','')}”")
        if c.get('gloss'): lines.append(f"  - 💡 {c['gloss']}")
        for b in c.get('breakdown', []):
            lines.append(f"  - ▶ {b.get('text','')} — {b.get('note','')}")
    # 读
    lines += ["", L['h_read'], ""]
    for a in annotated.get('annotations', []):
        lines.append(f"- （{a.get('focus','')}）{a.get('line','')}")
        if a.get('rule'): lines.append(f"  - 📐 规则：{a['rule']}")
        lines.append(f"  - 💡 {a.get('note','')}")
        lines.append(f"  - ✅ {a.get('tip','')}")
        if a.get('more'): lines.append(f"  - ➕ 更多例句：{a['more']}")
    lines += ["", L['h_cloze'], ""]
    for c in annotated.get('cloze', []):
        lines.append(f"- 完形 {c.get('id')}：{c.get('blanked','')} → ✅ {' / '.join(c.get('answers',[]))}")
    # 说
    lines += ["", L['h_speak'], ""]
    for s in tasks.get('speaking', []):
        ch = f" as {s['character']}" if s.get('character') else ""
        lines.append(f"- {s.get('type','')}{ch}：{s.get('instruction','')}")
        lines.append(f"  - Use: {', '.join(s.get('use_words', []))}")
        if s.get('focus_sounds'): lines.append(f"  - 🔤 重点发音：{', '.join(s['focus_sounds'])}")
        lines.append(f"  - Checklist: {', '.join(s.get('checklist', []))}")
        if s.get('asr_target'): lines.append(f"  - 🎤 可评分：录下你说的话，运行 `score_speaking.py --audio 录音 --target \"{s['asr_target']}\"`")
    # 写
    lines += ["", L['h_write'], ""]
    for w in tasks.get('writing', []):
        lines.append(f"- {w.get('type','')} / {w.get('register','')}：{w.get('instruction','')}")
        lines.append(f"  - Must use: {', '.join(w.get('require_words', []))}")
        if w.get('checks'):
            chk = '; '.join(f"{c.get('type')}:{c.get('value')}" for c in w['checks'])
            lines.append(f"  - 🤖 自动量规：{chk}（运行 `score_writing.py --task {w.get('id')} --text essay.txt --tasks tasks.json`）")
        if w.get('model'): lines.append(f"  - Model: {w['model']}")
    # 复习
    lines += ["", L['h_review'], "", f"- 目标词链：{'、'.join(w['term'] for w in words)}", "",
              f"- {L['pre']}", f"- {L['post']}", ""]
    if recall and recall.get('hints'):
        lines += ["### 🔁 跨集复现（已学词的新语境）", ""]
        for h in recall['hints']:
            lines.append(f"- **{h['term']}**（{h.get('cefr','')}）：旧「{h.get('old_context','')}」")
            for c in h.get('new_contexts', []):
                lines.append(f"  - ▶ 新：「{c}」")
        lines.append("")
    lines += [f"> 👨‍💻 {AUTHOR} · {CONTACT}", f"> {LEGAL}"]
    path = os.path.join(out_dir, deck + ".md")
    open(path, 'w', encoding='utf-8').write("\n".join(lines) + "\n")
    return path

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--words', required=True)
    ap.add_argument('--listening', required=True)
    ap.add_argument('--annotated', required=True)
    ap.add_argument('--tasks', required=True)
    ap.add_argument('--media-dir', default='media')
    ap.add_argument('--deck', default='DramaLex')
    ap.add_argument('--mode', default='A', choices=['A','B','C'])
    ap.add_argument('--format', default='html', choices=['html','anki','excel','word','md'])
    ap.add_argument('--out-dir', default='out')
    ap.add_argument('--backend', default='auto')
    ap.add_argument('--voice', default='Samantha')
    ap.add_argument('--watch', default=None, help='可选：观看协议 JSON（缺省用内置三步字幕法）')
    ap.add_argument('--ui-lang', default='zh', choices=['zh', 'en'], help='界面语言（影响 Excel/Word/Markdown 表头与声明）')
    args = ap.parse_args()

    try:
        words = json.load(open(args.words, encoding='utf-8'))
        listening = json.load(open(args.listening, encoding='utf-8'))
        annotated = json.load(open(args.annotated, encoding='utf-8'))
        tasks = json.load(open(args.tasks, encoding='utf-8'))
    except FileNotFoundError as e:
        print("缺少输入文件:", e, file=sys.stderr); sys.exit(2)
    except json.JSONDecodeError as e:
        print("JSON 解析失败:", e, file=sys.stderr); sys.exit(2)

    DEFAULT_WATCH = {
        "protocol": [
            {"pass": 1, "name": "无字幕抓大意", "instruction": "先关字幕看一遍，只抓谁、在哪、发生什么，允许听不懂。"},
            {"pass": 2, "name": "开字幕抓细节", "instruction": "开着字幕看，对照不懂的词句，注意连读/弱读/语调。"},
            {"pass": 3, "name": "无字幕再听", "instruction": "再关字幕看一遍，检验刚学的词是否 now 听懂。"},
        ],
        "notice": [
            "合成音(TTS)只是发音参照，真实原速听力请观看正片。",
            "本工具不替代看剧，观看正片才是听力提升的主场。",
        ],
    }
    watch = DEFAULT_WATCH
    if args.watch:
        try:
            watch = json.load(open(args.watch, encoding='utf-8'))
        except Exception as e:
            print("watch 读取失败，改用内置协议:", e, file=sys.stderr)
    os.makedirs(args.media_dir, exist_ok=True)
    os.makedirs(args.out_dir, exist_ok=True)
    backend = ga.pick_backend(args.backend)
    print("TTS 后端:", backend)

    # audio only needed for html / anki
    cache = {}
    if args.format in ('html', 'anki'):
        cache, _ = collect_all_audio(words, listening, annotated, tasks, args.media_dir, backend, args.voice)
        print(f"线音频已生成: {len(cache)} 条 -> {args.media_dir}")

    if args.format == 'html':
        html_doc = build_html(words, listening, annotated, tasks, cache, args.mode, args.deck, watch=watch, media_dir=args.media_dir, ui_lang=args.ui_lang)
        html_path = os.path.join(args.out_dir, 'practice.html')
        open(html_path, 'w', encoding='utf-8').write(html_doc)
        print("已写出（单文件）:", html_path)

    elif args.format == 'anki':
        apkg = build_anki(words, listening, annotated, tasks, cache, args.mode, args.deck, args.out_dir, args.media_dir)
        if not apkg:
            print(".apkg 生成失败（请 pip install genanki）", file=sys.stderr); sys.exit(1)
        print("已写出（单文件）:", apkg)

    elif args.format == 'excel':
        path = build_excel(words, listening, annotated, tasks, args.deck, args.out_dir, watch=watch, ui_lang=args.ui_lang)
        if not path:
            print("Excel 生成失败（请 pip install openpyxl）", file=sys.stderr); sys.exit(1)
        print("已写出（单文件）:", path)

    elif args.format == 'word':
        path = build_word(words, listening, annotated, tasks, args.deck, args.out_dir, watch=watch, ui_lang=args.ui_lang)
        if not path:
            print("Word 生成失败（请 pip install python-docx）", file=sys.stderr); sys.exit(1)
        print("已写出（单文件）:", path)

    elif args.format == 'md':
        path = build_markdown(words, listening, annotated, tasks, args.deck, args.out_dir, watch=watch, ui_lang=args.ui_lang)
        if not path:
            print("Markdown 生成失败", file=sys.stderr); sys.exit(1)
        print("已写出（单文件）:", path)

    print("done. format =", args.format)

if __name__ == '__main__':
    main()
