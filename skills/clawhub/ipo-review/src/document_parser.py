from __future__ import annotations

import csv
import re
from pathlib import Path

from .models import DocumentInfo, Evidence


def parse_documents(docs: list[DocumentInfo]) -> tuple[list[Evidence], list[dict]]:
    evidence: list[Evidence] = []
    unparsed: list[dict] = []
    for doc in docs:
        try:
            before = len(evidence)
            if doc.suffix == ".txt":
                evidence.extend(_parse_txt(doc))
            elif doc.suffix == ".pdf":
                evidence.extend(_parse_pdf(doc))
            elif doc.suffix == ".docx":
                evidence.extend(_parse_docx(doc))
            elif doc.suffix == ".xlsx":
                evidence.extend(_parse_xlsx(doc))
            doc.pages = max([e.page for e in evidence if e.doc_id == doc.doc_id] or [0])
            doc.parse_status = "parsed" if len(evidence) > before else "empty"
        except Exception as exc:  # noqa: BLE001
            doc.parse_status = "failed"
            doc.error = str(exc)
            unparsed.append({"file": doc.filename, "page": "", "item": "document", "reason": str(exc)})
    return evidence, unparsed


def _eid(doc: DocumentInfo, page: int, seq: int, kind: str = "P") -> str:
    return f"{doc.doc_id}-P{page:04d}-{kind}{seq:04d}"


def default_unit_for_role(doc: DocumentInfo) -> tuple[str, str]:
    if doc.role in {"financial_statements", "notes"}:
        return "元", "document_default"
    if doc.role in {"prospectus", "inquiry_round_1", "inquiry_round_2"}:
        return "万元", "document_default"
    return "unknown", "unknown"


def _section(text: str) -> str:
    match = re.match(r"^[一二三四五六七八九十]+[、.．]\s*(.+)|^第[一二三四五六七八九十\d]+[章节]\s*(.+)", text)
    return next((g for g in match.groups() if g), "") if match else ""


def _question(text: str) -> str:
    match = re.search(r"(问题\s*[一二三四五六七八九十\d]+|第\s*[一二三四五六七八九十\d]+\s*题)", text)
    return match.group(1).replace(" ", "") if match else ""


def _make_text_evidence(doc: DocumentInfo, page: int, lines: list[str], docx: bool = False) -> list[Evidence]:
    rows = []
    current_section = ""
    nearby_unit, nearby_ttl = "unknown", 0
    for idx, line in enumerate(lines, 1):
        line = re.sub(r"\s+", " ", line).strip()
        if not line:
            nearby_unit, nearby_ttl = "unknown", 0
            continue
        current_section = _section(line) or current_section
        if _resets_nearby(line):
            nearby_unit, nearby_ttl = "unknown", 0
        explicit = _detect_unit_with_source(line)
        has_number = bool(re.search(r"\d", line))
        if explicit[0] != "unknown" and has_number:
            unit, source = explicit
            nearby_unit, nearby_ttl = explicit[0], 1
        elif explicit[0] != "unknown":
            unit, source = "unknown", "unknown"
            nearby_unit, nearby_ttl = explicit[0], 1
        elif has_number and nearby_ttl > 0:
            unit, source = nearby_unit, "nearby_text"
            nearby_ttl -= 1
        else:
            unit, source = "unknown", "unknown"
            if not has_number:
                nearby_ttl = 0
        evidence_id = f"{doc.doc_id}-PARA{idx:05d}" if docx else _eid(doc, page, idx)
        position = f"PARA{idx:05d}" if docx else f"P{page}"
        rows.append(Evidence(
            evidence_id=evidence_id,
            doc_id=doc.doc_id,
            filename=doc.filename,
            page=0 if docx else page,
            kind="text",
            position=position,
            section=current_section,
            question_no=_question(line),
            paragraph_no=idx if docx else 0,
            text=line,
            unit=unit,
            unit_source=source,
        ))
    return rows


def _resets_nearby(text: str) -> bool:
    if _section(text) or _question(text):
        return True
    if re.fullmatch(r"目\s*录|董事长[:：]?|年\s*月\s*日|（?本页无正文.*", text):
        return True
    return False


def _detect_unit(text: str) -> str:
    return _detect_unit_with_source(text)[0]


def _detect_unit_with_source(text: str, source: str = "explicit_value") -> tuple[str, str]:
    match = re.search(r"单位[:：]\s*([人民币美元港币]*[万亿千]*元|%|百分比|百分点|倍)", text)
    if match:
        return match.group(1), "nearby_text" if not re.search(r"\d", text) else source
    for unit in ["亿元", "万元", "千元", "元", "百分点", "倍", "%"]:
        if re.search(rf"\d\s*{re.escape(unit)}|{re.escape(unit)}\s*\d", text):
            return unit, source
    return "unknown", "unknown"


def _parse_txt(doc: DocumentInfo) -> list[Evidence]:
    text = Path(doc.path).read_text(encoding="utf-8", errors="ignore")
    pages = re.split(r"\f|第\s*\d+\s*页", text)
    out: list[Evidence] = []
    for page, content in enumerate(pages, 1):
        out.extend(_make_text_evidence(doc, page, content.splitlines()))
    return out


def _parse_pdf(doc: DocumentInfo) -> list[Evidence]:
    try:
        import fitz  # type: ignore
    except ImportError as exc:
        raise RuntimeError("缺少 PyMuPDF，无法解析PDF；请安装 requirements.txt") from exc
    out: list[Evidence] = []
    pdf = fitz.open(doc.path)
    for page_index, page in enumerate(pdf, 1):
        text = page.get_text("text")
        if not text.strip():
            raise RuntimeError(f"第{page_index}页未提取到文本，可能为扫描件，需OCR")
        out.extend(_make_text_evidence(doc, page_index, text.splitlines()))
    return out


def _parse_docx(doc: DocumentInfo) -> list[Evidence]:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，无法解析DOCX；请安装 requirements.txt") from exc
    parsed = Document(doc.path)
    lines = [p.text for p in parsed.paragraphs if p.text.strip()]
    out = _make_text_evidence(doc, 0, lines, docx=True)
    seq = len(out) + 1
    table_sections = _docx_table_sections(parsed)
    for t_index, table in enumerate(parsed.tables, 1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        section = table_sections[t_index - 1] if t_index - 1 < len(table_sections) else ""
        table_unit, table_source = _table_unit(rows, doc)
        header_idx = _header_index(rows)
        headers = rows[header_idx] if header_idx < len(rows) else rows[0]
        for r_idx, row in enumerate(rows[header_idx + 1:], header_idx + 2):
            row_name = row[0] if row else ""
            for c_idx, value in enumerate(row[1:], 2):
                if not value:
                    continue
                col_name = headers[c_idx - 1] if c_idx - 1 < len(headers) else f"C{c_idx}"
                unit, source = _cell_unit(value, col_name, table_unit, table_source, doc)
                position = f"T{t_index:02d}-R{r_idx:03d}-C{c_idx:03d}"
                out.append(Evidence(f"{doc.doc_id}-{position}", doc.doc_id, doc.filename, 0, "table",
                                    position=position, section=section, table_no=f"T{t_index:02d}", row_index=r_idx,
                                    col_index=c_idx, row_name=row_name, col_name=col_name,
                                    text=value, unit=unit, unit_source=source,
                                    table_structure_confidence=_table_confidence(headers, table_unit)))
                seq += 1
    return out


def _iter_block_items(parent):
    """按文档本体顺序产出段落和表格，用于把区分见出（母公司/合并报表）下推到其后的表格。"""
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table as _Table
    from docx.text.paragraph import Paragraph as _Paragraph
    body = parent.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield _Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield _Table(child, parent)


def _financial_section_hint(text: str) -> str:
    if re.search(r"母公司.{0,8}(财务报表|资产负债表|利润表|现金流量表|所有者权益|报表项目|报表主要项目)", text):
        return "母公司财务报表"
    if re.search(r"合并.{0,8}(财务报表|资产负债表|利润表|现金流量表|所有者权益|报表项目|报表主要项目)", text):
        return "合并财务报表"
    return ""


def _docx_table_sections(parsed) -> list[str]:
    """返回与 parsed.tables 等长的列表，每项是该表格所属的财务报表区分见出。"""
    from docx.table import Table as _Table
    sections: list[str] = []
    current = ""
    for block in _iter_block_items(parsed):
        if isinstance(block, _Table):
            sections.append(current)
        else:
            text = (block.text or "").strip()
            if text:
                hint = _financial_section_hint(text)
                if hint:
                    current = hint
    return sections


def _parse_xlsx(doc: DocumentInfo) -> list[Evidence]:
    try:
        import openpyxl  # type: ignore
    except ImportError as exc:
        raise RuntimeError("缺少 openpyxl，无法解析XLSX；请安装 requirements.txt") from exc
    wb = openpyxl.load_workbook(doc.path, data_only=True, read_only=True)
    out: list[Evidence] = []
    seq = 1
    for s_idx, sheet in enumerate(wb.worksheets, 1):
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        table_unit, table_source = _table_unit([[str(c).strip() if c is not None else "" for c in row] for row in rows[:5]], doc)
        header_idx = _header_index([[str(c).strip() if c is not None else "" for c in row] for row in rows[:10]])
        headers = [str(v).strip() if v is not None else "" for v in rows[header_idx]]
        for r_idx, row in enumerate(rows[header_idx + 1:], header_idx + 2):
            row_name = str(row[0]).strip() if row and row[0] is not None else ""
            for c_idx, cell in enumerate(row[1:], 2):
                if cell is None or str(cell).strip() == "":
                    continue
                text = str(cell).strip()
                col_name = headers[c_idx - 1] if c_idx - 1 < len(headers) else f"C{c_idx}"
                unit, source = _cell_unit(text, col_name, table_unit, table_source, doc)
                position = f"{sheet.title}!R{r_idx}C{c_idx}"
                out.append(Evidence(f"{doc.doc_id}-{sheet.title}-R{r_idx:03d}-C{c_idx:03d}", doc.doc_id, doc.filename, s_idx, "table",
                                    position=position,
                                    table_no=sheet.title, table_name=sheet.title,
                                    row_index=r_idx, col_index=c_idx,
                                    row_name=row_name, col_name=col_name,
                                    text=text, unit=unit, unit_source=source,
                                    table_structure_confidence=_table_confidence(headers, table_unit)))
                seq += 1
    return out


def _header_index(rows: list[list[str]]) -> int:
    for idx, row in enumerate(rows[:8]):
        joined = " ".join(str(x) for x in row)
        if any(k in joined for k in ["项目", "年度", "202", "金额", "占比", "比例"]):
            return idx
    return 0


def _table_unit(rows: list[list[str]], doc: DocumentInfo) -> tuple[str, str]:
    for row in rows[:5]:
        text = " ".join(str(x) for x in row if x)
        match = re.search(r"单位[:：]\s*([人民币美元港币]*[万亿千]*元|%|百分比|百分点|倍)", text)
        if match:
            return match.group(1), "table_header"
    return default_unit_for_role(doc)


def _cell_unit(value: str, col_name: str, table_unit: str, table_source: str, doc: DocumentInfo) -> tuple[str, str]:
    unit, source = _detect_unit_with_source(value, "explicit_value")
    if unit != "unknown":
        return unit, source
    unit, source = _detect_unit_with_source(col_name, "column_header")
    if unit != "unknown":
        return unit, source
    if table_unit != "unknown":
        return table_unit, table_source
    return default_unit_for_role(doc)


def _table_confidence(headers: list[str], unit: str) -> float:
    score = 0.0
    if unit != "unknown":
        score += 0.35
    if any(h for h in headers):
        score += 0.25
    if any(re.search(r"20\d{2}|金额|占比|比例|项目", h) for h in headers):
        score += 0.3
    return min(score, 1.0)


def write_unparsed_csv(path: Path, rows: list[dict]) -> None:
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=["file", "page", "item", "reason"])
        writer.writeheader()
        writer.writerows(rows)
