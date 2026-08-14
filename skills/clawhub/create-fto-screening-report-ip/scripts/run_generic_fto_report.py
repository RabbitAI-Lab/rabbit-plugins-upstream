#!/usr/bin/env python3
"""Run a source-faithful, evidence-preserving PatSnap FTO screening workflow.

This command creates research artifacts for human review. It does not issue a
legal opinion, determine infringement, or represent that a search is complete.
REST mode uses the global PatSnap Connect API through ``PatSnapClient``. MCP
mode is intentionally orchestration-only: an MCP-capable host must supply its
results as JSON because this local script cannot honestly claim an MCP call.
"""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import os
import pathlib
import re
import sys
from collections import OrderedDict
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from render_report import render_from_structured_data
from zhihuiya_api import PatSnapApiError, PatSnapClient


SCRIPT_DIR = pathlib.Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
DEFAULT_API_CONFIG_PATH = SKILL_DIR / "references" / "zhihuiya_config.json"
DEFAULT_BUSINESS_CONFIG_PATH = SKILL_DIR / "references" / "config.json"
DEFAULT_TEMPLATE_PATH = SKILL_DIR / "assets" / "FTO\u62a5\u544a\u6a21\u677f.docx"
SCHEMA_VERSION = "2.0"
DISCLAIMER = (
    "Research screening only—not a legal opinion, freedom-to-operate clearance, "
    "or determination of infringement. Qualified counsel must review material "
    "claims, prosecution history, ownership, legal status, and jurisdiction-specific law."
)
SAFE_NAME = re.compile(r"[^A-Za-z0-9._-]+")
PATENT_NUMBER_KEYS = (
    "publication_number", "patent_number", "publication_no", "pn", "apno", "application_number"
)


def utc_now() -> str:
    return dt.datetime.now(dt.timezone.utc).replace(microsecond=0).isoformat()


def load_json(path: pathlib.Path, *, expected: type | None = None) -> Any:
    value = json.loads(path.read_text(encoding="utf-8-sig"))
    if expected is not None and not isinstance(value, expected):
        raise ValueError(f"{path} must contain a {expected.__name__}.")
    return value


def write_json(path: pathlib.Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    temporary.replace(path)


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def as_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, tuple):
        return list(value)
    return [value]


def first(record: dict[str, Any], *keys: str, default: Any = "") -> Any:
    for key in keys:
        value = record.get(key)
        if value is not None and value != "":
            return value
    return default


def safe_stem(value: str, fallback: str = "fto-screening") -> str:
    stem = SAFE_NAME.sub("-", clean_text(value)).strip("-._")
    return stem[:80] or fallback


def stable_id(prefix: str, *values: Any) -> str:
    payload = "|".join(clean_text(value).lower() for value in values)
    return f"{prefix}-{hashlib.sha256(payload.encode('utf-8')).hexdigest()[:12]}"


def parse_delimited(value: str | None) -> list[str]:
    if not value:
        return []
    return [item.strip() for item in re.split(r"[,;\n]", value) if item.strip()]


METADATA_ALIASES = {
    "product": "product_name",
    "product name": "product_name",
    "technology": "product_name",
    "version": "product_version",
    "design baseline": "product_version",
    "decision context": "decision_context",
    "purpose": "decision_context",
    "jurisdiction": "target_jurisdictions",
    "jurisdictions": "target_jurisdictions",
    "target market": "target_jurisdictions",
    "relevant acts": "relevant_acts",
    "commercial acts": "relevant_acts",
    "search cutoff": "search_cutoff",
    "data cutoff": "search_cutoff",
    "status cutoff": "status_cutoff",
    "legal status cutoff": "status_cutoff",
    "family rule": "family_counting_convention",
    "family counting convention": "family_counting_convention",
    "competitors": "competitors",
    "assignees": "competitors",
    "ipc": "classifications",
    "cpc": "classifications",
    "classifications": "classifications",
    "manual queries": "manual_queries",
    "reviewed queries": "manual_queries",
}


def normalize_metadata_key(value: Any) -> str:
    key = clean_text(value).lower().rstrip(":")
    key = re.sub(r"[_-]+", " ", key)
    key = re.sub(r"\s+", " ", key)
    return METADATA_ALIASES.get(key, key.replace(" ", "_"))


def parse_metadata_tables(tables: list[list[list[str]]]) -> dict[str, Any]:
    """Parse explicit two-column metadata only; do not classify tables by title."""
    metadata: dict[str, Any] = {}
    for table in tables:
        for row in table:
            if len(row) != 2:
                continue
            raw_key, raw_value = clean_text(row[0]), clean_text(row[1])
            if not raw_key or not raw_value:
                continue
            key = normalize_metadata_key(raw_key)
            if key not in set(METADATA_ALIASES.values()):
                continue
            if key in {"target_jurisdictions", "relevant_acts", "competitors", "classifications", "manual_queries"}:
                metadata[key] = parse_delimited(raw_value)
            else:
                metadata[key] = raw_value
    return metadata


def parse_feature_tables(tables: list[list[list[str]]]) -> list[dict[str, Any]]:
    """Find tables with explicit English feature headers and preserve locators."""
    output: list[dict[str, Any]] = []
    accepted = {"feature", "technical feature", "feature description", "element", "technical element"}
    id_headers = {"id", "feature id", "feature number", "no", "number"}
    source_headers = {"source", "source reference", "specification reference", "locator"}
    for table_index, table in enumerate(tables, start=1):
        if len(table) < 2:
            continue
        headers = [clean_text(cell).lower().rstrip(":") for cell in table[0]]
        feature_col = next((index for index, name in enumerate(headers) if name in accepted), None)
        if feature_col is None:
            continue
        id_col = next((index for index, name in enumerate(headers) if name in id_headers), None)
        source_col = next((index for index, name in enumerate(headers) if name in source_headers), None)
        for row_index, row in enumerate(table[1:], start=2):
            feature = clean_text(row[feature_col]) if feature_col < len(row) else ""
            if not feature:
                continue
            source = clean_text(row[source_col]) if source_col is not None and source_col < len(row) else f"Table {table_index}, row {row_index}"
            feature_id = clean_text(row[id_col]) if id_col is not None and id_col < len(row) else f"F-{len(output) + 1:03d}"
            output.append({
                "feature_id": feature_id,
                "feature": feature,
                "source_reference": source,
                "review_status": "needs_human_confirmation",
            })
    return output


def quote_query_term(value: str) -> str:
    """Quote a literal term for a review draft; this is not database validation."""
    return '"' + clean_text(value).replace('"', '\\"') + '"'


def draft_queries(features: list[dict[str, Any]], *, max_terms: int = 8) -> list[dict[str, Any]]:
    """Create visibly unapproved drafts from feature language.

    The original source sent generated queries directly. The localized workflow
    keeps the useful drafting step but prevents execution until a human has
    checked field syntax, jurisdiction scope, synonyms, translations, and
    classification strategy.
    """
    output: list[dict[str, Any]] = []
    stop = {
        "a", "an", "and", "as", "at", "by", "for", "from", "in", "into",
        "is", "of", "on", "or", "that", "the", "to", "using", "with",
    }
    for index, feature in enumerate(features, start=1):
        tokens = re.findall(r"[A-Za-z][A-Za-z0-9-]{2,}", clean_text(feature.get("feature")))
        unique: list[str] = []
        for token in tokens:
            lowered = token.lower()
            if lowered in stop or lowered in {item.lower() for item in unique}:
                continue
            unique.append(token)
            if len(unique) >= max_terms:
                break
        if not unique:
            continue
        output.append({
            "query_id": f"QD-{index:03d}",
            "label": f"Draft from {feature.get('feature_id') or f'feature {index}'}",
            "query": " AND ".join(quote_query_term(token) for token in unique),
            "origin": "generated_draft",
            "human_approved": False,
            "review_note": "Validate PatSnap field syntax, synonyms, classifications, translations, exclusions, and jurisdiction scope before approval.",
        })
    return output


def validate_project(project: dict[str, Any]) -> list[dict[str, Any]]:
    errors = []
    required = {
        "product_name": "Define the product or technology under review.",
        "product_version": "Freeze a design version or technical baseline.",
        "decision_context": "State the business decision this screening informs.",
        "target_jurisdictions": "Identify each jurisdiction to be screened.",
        "relevant_acts": "Identify legally relevant commercial acts.",
        "search_cutoff": "Record the search cutoff date.",
        "status_cutoff": "Record the legal-status verification cutoff date.",
        "family_counting_convention": "Define the patent-family counting rule.",
    }
    for field, message in required.items():
        if not project.get(field):
            errors.append({"stage": "input_validation", "field": field, "state": "missing", "message": message})
    return errors


def extract_docx_input(path: pathlib.Path) -> dict[str, Any]:
    """Read paragraphs and tables without inventing facts from formatting."""
    document = Document(path)
    paragraphs = [clean_text(p.text) for p in document.paragraphs if clean_text(p.text)]
    tables: list[list[list[str]]] = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([clean_text(cell.text) for cell in row.cells])
        if rows:
            tables.append(rows)
    return {
        "source_file": path.name,
        "paragraphs": paragraphs,
        "tables": tables,
        "metadata": parse_metadata_tables(tables),
        "features": parse_feature_tables(tables),
        "extraction_note": "Text and table cells extracted; formatting was not treated as evidence.",
    }


def normalize_input(raw: dict[str, Any], source_path: pathlib.Path) -> dict[str, Any]:
    project = raw.get("project") if isinstance(raw.get("project"), dict) else {}
    scope = raw.get("scope") if isinstance(raw.get("scope"), dict) else {}
    features = raw.get("features") or raw.get("technical_features") or []
    queries = raw.get("queries") or raw.get("search_queries") or []
    return {
        "project": {
            "product_name": clean_text(first(project, "product_name", "name", default=first(raw, "product_name", "title"))),
            "product_version": clean_text(first(project, "product_version", "version", default=first(raw, "product_version"))),
            "decision_context": clean_text(first(project, "decision_context", "purpose", default=first(raw, "decision_context"))),
            "target_jurisdictions": as_list(first(project, "target_jurisdictions", "jurisdictions", default=first(scope, "target_jurisdictions", "target_market", default=[]))),
            "relevant_acts": as_list(first(project, "relevant_acts", default=first(scope, "relevant_acts", default=[]))),
            "search_cutoff": clean_text(first(project, "search_cutoff", default=first(scope, "search_cutoff", "data_cutoff"))),
            "status_cutoff": clean_text(first(project, "status_cutoff", default=first(scope, "status_cutoff", "status_checked_as_of"))),
            "family_counting_convention": clean_text(first(project, "family_counting_convention", default=first(scope, "family_counting_convention", "family_method"))),
        },
        "features": normalize_features(features),
        "queries": normalize_queries(queries),
        "input_candidates": [item for item in as_list(raw.get("candidates") or raw.get("patent_list")) if isinstance(item, dict)],
        "source": {"path": str(source_path), "kind": source_path.suffix.lower().lstrip(".") or "json"},
    }


def normalize_features(values: Any) -> list[dict[str, Any]]:
    result = []
    for index, value in enumerate(as_list(values), start=1):
        if isinstance(value, str):
            record = {"feature": clean_text(value)}
        elif isinstance(value, dict):
            record = dict(value)
        else:
            continue
        text = clean_text(first(record, "feature", "text", "description", "technical_feature"))
        if not text:
            continue
        result.append({
            "feature_id": clean_text(first(record, "feature_id", "id", default=f"F-{index:03d}")),
            "feature": text,
            "source_reference": clean_text(first(record, "source_reference", "source", "citation")),
            "review_status": clean_text(first(record, "review_status", default="needs_human_confirmation")),
        })
    return result


def normalize_queries(values: Any) -> list[dict[str, Any]]:
    if isinstance(values, dict):
        values = [{"label": key, "query": value} for key, value in values.items()]
    result = []
    for index, value in enumerate(as_list(values), start=1):
        record = {"query": value} if isinstance(value, str) else dict(value) if isinstance(value, dict) else {}
        query = clean_text(first(record, "query", "q", "value"))
        if not query:
            continue
        origin = clean_text(first(record, "origin", "provenance", default="user_supplied"))
        approved = bool(first(record, "approved", "human_approved", default=origin == "user_supplied"))
        result.append({
            "query_id": clean_text(first(record, "query_id", "id", default=f"Q-{index:03d}")),
            "label": clean_text(first(record, "label", "name", default=f"Query {index}")),
            "query": query,
            "origin": origin,
            "human_approved": approved,
        })
    return result


def normalize_patent(record: dict[str, Any], *, query_ids: list[str] | None = None) -> dict[str, Any]:
    number = clean_text(first(record, *PATENT_NUMBER_KEYS))
    family_id = clean_text(first(record, "family_id", "simple_family_id", "inpadoc_family_id"))
    return {
        "candidate_id": clean_text(first(record, "candidate_id", default=stable_id("C", number, family_id, first(record, "title")))),
        "publication_number": number,
        "application_number": clean_text(first(record, "application_number", "apno")),
        "title": clean_text(first(record, "title", "patent_title")),
        "assignee": clean_text(first(record, "assignee", "current_assignee", "applicant")),
        "jurisdiction": clean_text(first(record, "jurisdiction", "country", "authority")),
        "family_id": family_id,
        "filing_date": clean_text(first(record, "filing_date", "application_date")),
        "publication_date": clean_text(first(record, "publication_date")),
        "legal_status_raw": first(record, "legal_status", "legal_status_raw", "simple_legal_status"),
        "status_checked_as_of": clean_text(first(record, "status_checked_as_of", "status_cutoff")),
        "source_url": clean_text(first(record, "source_url", "url")),
        "matched_query_ids": list(query_ids or as_list(record.get("matched_query_ids"))),
        "screening_state": "not_assessed",
        "evidence_gaps": ["Material claims and jurisdiction-specific legal status require review."],
    }


def _extract_search_rows(payload: Any) -> list[dict[str, Any]]:
    if isinstance(payload, list):
        return [item for item in payload if isinstance(item, dict)]
    if not isinstance(payload, dict):
        return []
    for key in ("results", "patents", "data", "records", "items", "list"):
        value = payload.get(key)
        if isinstance(value, list):
            return [item for item in value if isinstance(item, dict)]
        if isinstance(value, dict):
            nested = _extract_search_rows(value)
            if nested:
                return nested
    return []


def deduplicate_candidates(records: Iterable[dict[str, Any]]) -> list[dict[str, Any]]:
    merged: OrderedDict[str, dict[str, Any]] = OrderedDict()
    for record in records:
        key = clean_text(record.get("publication_number") or record.get("application_number") or record.get("candidate_id")).upper()
        if not key:
            key = record["candidate_id"]
        if key not in merged:
            merged[key] = record
            continue
        existing = merged[key]
        existing["matched_query_ids"] = list(dict.fromkeys(as_list(existing.get("matched_query_ids")) + as_list(record.get("matched_query_ids"))))
        for field, value in record.items():
            if not existing.get(field) and value:
                existing[field] = value
    return list(merged.values())


def collect_rest_candidates(client: PatSnapClient, queries: list[dict[str, Any]], *, limit: int) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    candidates: list[dict[str, Any]] = []
    errors: list[dict[str, Any]] = []
    for query in queries:
        if not query.get("human_approved"):
            errors.append({"stage": "search", "query_id": query["query_id"], "state": "skipped", "message": "Generated query was not human-approved."})
            continue
        try:
            payload = client.search_all_patents(query["query"], max_total=limit)
            for row in _extract_search_rows(payload):
                candidates.append(normalize_patent(row, query_ids=[query["query_id"]]))
        except (PatSnapApiError, OSError, ValueError) as exc:
            errors.append({"stage": "search", "query_id": query["query_id"], "state": "error", "message": clean_text(exc)})
    return deduplicate_candidates(candidates), errors


def attach_claim_evidence(client: PatSnapClient, candidates: list[dict[str, Any]], *, max_candidates: int) -> list[dict[str, Any]]:
    errors = []
    for candidate in candidates[:max_candidates]:
        number = candidate.get("publication_number") or candidate.get("application_number")
        if not number:
            errors.append({"stage": "claims", "candidate_id": candidate["candidate_id"], "state": "skipped", "message": "No patent identifier."})
            continue
        try:
            claims = client.get_claims(str(number))
            claim_rows = _extract_search_rows(claims)
            candidate["claims"] = claim_rows
            candidate["claim_review_scope"] = {
                "retrieved_count": len(claim_rows),
                "screening_shortcut": "Claim 1 may be triaged first, but every material independent and dependent claim requires counsel review.",
            }
        except (PatSnapApiError, OSError, ValueError) as exc:
            candidate["claims"] = []
            candidate["evidence_gaps"].append("Claim retrieval failed; no mapping conclusion may be inferred.")
            errors.append({"stage": "claims", "candidate_id": candidate["candidate_id"], "state": "error", "message": clean_text(exc)})
    return errors


def build_comparisons(features: list[dict[str, Any]], candidates: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Create blank review rows; keyword overlap is never an infringement conclusion."""
    rows = []
    for candidate in candidates:
        claim_records = as_list(candidate.get("claims"))
        for feature in features:
            rows.append({
                "comparison_id": stable_id("M", candidate["candidate_id"], feature["feature_id"]),
                "candidate_id": candidate["candidate_id"],
                "publication_number": candidate.get("publication_number", ""),
                "feature_id": feature["feature_id"],
                "feature": feature["feature"],
                "claim_reference": "",
                "claim_text": "",
                "mapping_state": "not_assessed" if claim_records else "evidence_missing",
                "mapping_rationale": "Human claim construction and element-by-element review required.",
                "evidence_locator": "",
                "reviewer": "",
                "reviewed_at": "",
            })
    return rows


def build_structured_data(normalized: dict[str, Any], candidates: list[dict[str, Any]], errors: list[dict[str, Any]], *, mode: str, dry_run: bool) -> dict[str, Any]:
    features = normalized["features"]
    queries = normalized["queries"]
    comparisons = build_comparisons(features, candidates)
    status = "dry_run" if dry_run else "partial" if errors else "complete_for_screening_scope"
    return {
        "schema_version": SCHEMA_VERSION,
        "report_title": f"{normalized['project'].get('product_name') or 'Product'} FTO Screening Report",
        "generated_at": utc_now(),
        "disclaimer": DISCLAIMER,
        "project": normalized["project"],
        "run_provenance": {
            "mode": mode,
            "status": status,
            "source": normalized["source"],
            "rest_service": "PatSnap Connect" if mode == "rest" else "Not called",
            "mcp_note": "MCP evidence was imported from an MCP-capable host; this script did not call MCP." if mode == "mcp-import" else "Not applicable",
        },
        "features": features,
        "queries": queries,
        "patent_list": candidates,
        "comparisons": comparisons,
        "pending_application_watchlist": [],
        "conclusion": {
            "screening_state": "not_assessed",
            "statement": "No legal conclusion is generated automatically. Candidate records require evidence-backed human review.",
        },
        "recommendations": [
            {"priority": "Required", "action": "Have qualified counsel review all material claims, family members, prosecution history, ownership, and current legal status in each target jurisdiction."},
            {"priority": "Required", "action": "Resolve every evidence gap and document claim-construction assumptions before a go/no-go decision."},
        ],
        "limitations": [
            DISCLAIMER,
            "Search results depend on approved query scope, database coverage, family rules, and cutoff dates.",
            "A simplified legal-status filter is discovery metadata, not proof that a right is enforceable.",
            "Pending and unpublished applications may change the risk picture.",
        ],
        "errors": errors,
        "sources": [],
    }


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor(45, 55, 65)
    normal.paragraph_format.space_after = Pt(5)
    for name, size, color in (("Title", 28, "17324D"), ("Subtitle", 11, "526575"), ("Heading 1", 17, "17324D"), ("Heading 2", 12, "246B84"), ("Heading 3", 10, "526575")):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.text = "PATENT SCREENING  /  CONFIDENTIAL WORK PRODUCT"
    header.style = styles["Caption"]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("PatSnap-assisted research screening  •  ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_table(document: Document, headers: list[str], rows: list[list[Any]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, "17324D")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(8)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = clean_text(value) or "—"
            set_cell_margins(cells[index])
            if row_index % 2:
                set_cell_shading(cells[index], "EEF3F5")
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(7.8)
        if widths:
            for index, width in enumerate(widths):
                cells[index].width = Inches(width)


def add_bullets(document: Document, values: Iterable[Any], fallback: str = "Not supplied") -> None:
    rendered = False
    for value in values:
        text = clean_text(value.get("action") if isinstance(value, dict) else value)
        if text:
            document.add_paragraph(text, style="List Bullet")
            rendered = True
    if not rendered:
        document.add_paragraph(fallback)


def render_docx(data: dict[str, Any], path: pathlib.Path) -> None:
    document = Document()
    configure_document(document)
    project = data.get("project", {})
    document.core_properties.title = clean_text(data.get("report_title"))
    document.core_properties.subject = "Evidence-preserving freedom-to-operate screening"
    document.core_properties.author = "PatSnap-assisted research workflow"
    document.core_properties.keywords = "FTO, patents, claim mapping, screening"
    document.add_paragraph("FREEDOM-TO-OPERATE SCREENING", style="Title")
    document.add_paragraph(clean_text(project.get("product_name")) or "Product / technology under review", style="Subtitle")
    document.add_paragraph(DISCLAIMER)
    add_table(document, ["Decision context", "Jurisdictions", "Relevant acts", "Cutoffs"], [[
        project.get("decision_context"), ", ".join(as_list(project.get("target_jurisdictions"))),
        ", ".join(as_list(project.get("relevant_acts"))),
        f"Search: {project.get('search_cutoff') or 'not supplied'}; status: {project.get('status_cutoff') or 'not supplied'}",
    ]], [2.0, 1.4, 1.4, 1.8])
    sections = [
        ("1. Executive screening statement", [data.get("conclusion", {}).get("statement"), DISCLAIMER]),
        ("2. Scope and assumptions", [f"Product version: {project.get('product_version') or 'not supplied'}", f"Family convention: {project.get('family_counting_convention') or 'not supplied'}"]),
        ("3. Technical feature set", []),
    ]
    for heading, body in sections:
        document.add_paragraph(heading, style="Heading 1")
        add_bullets(document, body)
    add_table(document, ["Feature ID", "Feature", "Source", "Review status"], [[item.get("feature_id"), item.get("feature"), item.get("source_reference"), item.get("review_status")] for item in data.get("features", [])], [0.7, 3.2, 1.3, 1.4])
    document.add_paragraph("4. Search strategy and provenance", style="Heading 1")
    add_table(document, ["Query ID", "Query", "Origin", "Approved"], [[item.get("query_id"), item.get("query"), item.get("origin"), item.get("human_approved")] for item in data.get("queries", [])], [0.7, 3.8, 1.2, 0.9])
    document.add_paragraph("5. Candidate patent set", style="Heading 1")
    add_table(document, ["Publication", "Title", "Assignee", "Jurisdiction", "Status metadata"], [[item.get("publication_number"), item.get("title"), item.get("assignee"), item.get("jurisdiction"), item.get("legal_status_raw")] for item in data.get("patent_list", [])], [1.1, 2.5, 1.3, 0.8, 1.0])
    document.add_paragraph("6. Claim-to-feature review matrix", style="Heading 1")
    add_table(document, ["Publication", "Feature", "Claim", "State", "Rationale"], [[item.get("publication_number"), item.get("feature"), item.get("claim_reference"), item.get("mapping_state"), item.get("mapping_rationale")] for item in data.get("comparisons", [])], [1.0, 2.0, 0.8, 0.9, 2.0])
    for heading, key in (("7. Pending-application watchlist", "pending_application_watchlist"), ("8. Evidence gaps and errors", "errors"), ("9. Required actions", "recommendations"), ("10. Limitations", "limitations")):
        document.add_paragraph(heading, style="Heading 1")
        add_bullets(document, data.get(key, []), "None recorded; reviewer confirmation still required.")
    document.add_paragraph("11. Run provenance", style="Heading 1")
    provenance = data.get("run_provenance", {})
    add_table(document, ["Mode", "Run status", "Generated", "Schema"], [[provenance.get("mode"), provenance.get("status"), data.get("generated_at"), data.get("schema_version")]], [1.4, 1.6, 2.3, 1.3])
    document.add_paragraph("12. Reviewer sign-off", style="Heading 1")
    add_table(document, ["Role", "Name", "Date", "Decision / conditions"], [["Technical reviewer", "", "", ""], ["Patent counsel", "", "", ""]], [1.4, 1.4, 1.0, 2.8])
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def resolve_input(path: pathlib.Path) -> dict[str, Any]:
    if path.suffix.lower() == ".json":
        raw = load_json(path, expected=dict)
    elif path.suffix.lower() == ".docx":
        extracted = extract_docx_input(path)
        metadata = extracted["metadata"]
        features = extracted["features"]
        raw = {
            "project": metadata,
            "features": features,
            "queries": [{"query": query, "origin": "document_supplied", "human_approved": True} for query in as_list(metadata.get("manual_queries"))],
            "source_extraction": extracted,
        }
    else:
        raise ValueError("Input must be .json or .docx.")
    return normalize_input(raw, path)


def build_client(config_path: pathlib.Path, api_key: str | None) -> PatSnapClient:
    config = load_json(config_path, expected=dict) if config_path.exists() else {}
    key = api_key or os.getenv("PATSNAP_API_KEY") or clean_text(config.get("api_key"))
    if not key:
        raise ValueError("PatSnap API key missing. Set PATSNAP_API_KEY or pass --api-key; never place secrets in reports.")
    return PatSnapClient(
        api_key=key,
        base_url=clean_text(config.get("base_url")) or "https://connect.patsnap.com",
        connect_timeout=float(config.get("connect_timeout_seconds", 10)),
        read_timeout=float(config.get("read_timeout_seconds", config.get("timeout_seconds", 60))),
        max_retries=int(config.get("max_retries", 3)),
    )


def run(args: argparse.Namespace) -> dict[str, pathlib.Path]:
    normalized = resolve_input(args.input)
    if not normalized["queries"] and args.draft_queries:
        normalized["queries"] = draft_queries(normalized["features"])
    candidates = [normalize_patent(item) for item in normalized.pop("input_candidates")]
    errors: list[dict[str, Any]] = validate_project(normalized["project"])
    mode = args.mode
    if args.mcp_results:
        imported = load_json(args.mcp_results)
        imported_rows = imported if isinstance(imported, list) else _extract_search_rows(imported)
        candidates.extend(normalize_patent(item) for item in imported_rows)
        mode = "mcp-import"
    if args.dry_run:
        mode = "dry-run"
    elif mode == "rest":
        client = build_client(args.api_config, args.api_key)
        searched, search_errors = collect_rest_candidates(client, normalized["queries"], limit=args.max_records)
        candidates.extend(searched)
        errors.extend(search_errors)
        candidates = deduplicate_candidates(candidates)
        errors.extend(attach_claim_evidence(client, candidates, max_candidates=args.max_claim_candidates))
    candidates = deduplicate_candidates(candidates)
    data = build_structured_data(normalized, candidates, errors, mode=mode, dry_run=args.dry_run)
    output = args.output_dir
    output.mkdir(parents=True, exist_ok=True)
    stem = safe_stem(args.report_name or normalized["project"].get("product_name") or "fto-screening")
    paths = {
        "queries": output / "queries.json",
        "patent_list": output / "patent_list.json",
        "claim_chart": output / "claim_chart.json",
        "structured": output / "fto_structured_data.json",
        "html": output / f"{stem}.html",
        "docx": output / f"{stem}.docx",
    }
    write_json(paths["queries"], data["queries"])
    write_json(paths["patent_list"], data["patent_list"])
    write_json(paths["claim_chart"], data["comparisons"])
    write_json(paths["structured"], data)
    render_from_structured_data(data, paths["html"])
    render_docx(data, paths["docx"])
    return paths


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Create an evidence-preserving English FTO screening package from JSON or DOCX input.")
    parser.add_argument("input", type=pathlib.Path, help="Structured JSON input or source DOCX")
    parser.add_argument("output_dir", type=pathlib.Path, help="Directory for JSON, HTML, and DOCX artifacts")
    parser.add_argument("--mode", choices=("rest", "offline"), default="offline", help="REST calls PatSnap Connect; offline uses supplied candidates only")
    parser.add_argument("--mcp-results", type=pathlib.Path, help="JSON exported by a verified PatSnap MCP host; switches provenance to mcp-import")
    parser.add_argument("--api-key", help="PatSnap API key; prefer PATSNAP_API_KEY to avoid shell-history exposure")
    parser.add_argument("--api-config", type=pathlib.Path, default=DEFAULT_API_CONFIG_PATH)
    parser.add_argument("--business-config", type=pathlib.Path, default=DEFAULT_BUSINESS_CONFIG_PATH, help="Reserved source-compatible path; report facts still come from input")
    parser.add_argument("--max-records", type=int, default=100)
    parser.add_argument("--max-claim-candidates", type=int, default=25)
    parser.add_argument("--report-name", help="Safe output filename stem")
    parser.add_argument("--draft-queries", action="store_true", help="Create unapproved query drafts; generated drafts are never executed")
    parser.add_argument("--dry-run", action="store_true", help="Validate and render without any network call")
    return parser


def main() -> int:
    try:
        args = build_parser().parse_args()
        if args.max_records < 1 or args.max_claim_candidates < 0:
            raise ValueError("Limits must be positive (claim limit may be zero).")
        paths = run(args)
        for label, path in paths.items():
            print(f"{label}: {path}")
        return 0
    except (OSError, UnicodeError, json.JSONDecodeError, ValueError, PatSnapApiError) as exc:
        print(f"FTO screening failed: {clean_text(exc)}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
