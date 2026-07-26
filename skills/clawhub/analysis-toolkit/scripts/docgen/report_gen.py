"""
报告生成模块

将分析结果组装为结构化Word报告，支持图表、数据表、结论。
"""
import os
import tempfile
from datetime import datetime

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START


def generate_report(sections, output_path, title="分析报告"):
    """
    组合报告生成。
    
    Parameters
    ----------
    sections : list of dict
        报告章节列表。每个章节为:
        {
            "type": "heading" | "chart" | "table" | "text" | "page_break",
            "level": 1/2 (heading only),
            "title": str,
            "figure": matplotlib Figure (chart only),
            "dataframe": pd.DataFrame (table only),
            "content": str (text only),
            "analysis": str (可选结论)
        }
    output_path : str
        输出的.docx文件路径
    title : str
        报告标题
    
    Returns
    -------
    str
        输出文件路径
    """
    doc = Document()
    
    # 设置默认中文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10)
    
    # 封面
    _add_cover(doc, title)
    
    # 临时目录存放图片
    temp_dir = tempfile.mkdtemp()
    
    for section in sections:
        _add_section(doc, section, temp_dir)
    
    doc.save(output_path)
    
    # 清理临时图片
    for f in os.listdir(temp_dir):
        try:
            os.remove(os.path.join(temp_dir, f))
        except (PermissionError, OSError):
            pass
    try:
        os.rmdir(temp_dir)
    except (PermissionError, OSError):
        pass
    
    return output_path


def _add_cover(doc, title):
    """添加封面页"""
    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    
    # 空行
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f"\n{title}\n")
    run.font.size = Pt(26)
    run.bold = True
    
    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    info.add_run(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run("―" * 39).bold = True
    
    doc.add_paragraph()
    _add_section_break(doc)


def _add_section(doc, section, temp_dir):
    """添加一个章节"""
    stype = section.get("type", "text")
    
    if stype == "heading":
        level = section.get("level", 1)
        title = section.get("title", "")
        h = doc.add_heading(title, level=level)
        for run in h.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    elif stype == "chart":
        title = section.get("title", "")
        figure = section.get("figure")
        analysis = section.get("analysis", "")
        
        if title:
            doc.add_heading(title, level=2)
        
        if figure:
            img_path = os.path.join(temp_dir, f"chart_{id(figure)}.png")
            figure.savefig(img_path, dpi=150, bbox_inches='tight')
            doc.add_picture(img_path, width=Inches(5.5))
            plt.close(figure)
        
        if analysis:
            doc.add_paragraph(analysis)
    
    elif stype == "table":
        title = section.get("title", "")
        df = section.get("dataframe")
        
        if title:
            doc.add_paragraph(title).bold = True
        
        if df is not None and not df.empty:
            rows = len(df) + 1
            cols = len(df.columns)
            table = doc.add_table(rows=rows, cols=cols, style='Light Shading Accent 1')
            
            # 表头
            for j, col in enumerate(df.columns):
                table.rows[0].cells[j].text = str(col)
            
            # 数据
            for i, (_, row) in enumerate(df.iterrows()):
                for j, val in enumerate(row):
                    text = f"{val:.2%}" if isinstance(val, float) and 0 < abs(val) < 1 else str(val)
                    table.rows[i + 1].cells[j].text = text
    
    elif stype == "text":
        content = section.get("content", "")
        if content:
            doc.add_paragraph(content)
    
    elif stype == "page_break":
        _add_section_break(doc)


def _add_section_break(doc):
    """添加分节符（新页）"""
    new_section = doc.add_section()
    new_section.start_type = WD_SECTION_START.NEW_PAGE


def export_to_pdf(docx_path):
    """
    Word转PDF。
    
    Parameters
    ----------
    docx_path : str
        .docx 文件路径
    
    Returns
    -------
    str
        .pdf 文件路径
    """
    pdf_path = docx_path.replace('.docx', '.pdf')
    
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        return pdf_path
    except ImportError:
        raise ImportError("需要安装 docx2pdf: pip install docx2pdf")
    except Exception as e:
        raise RuntimeError(f"PDF转换失败: {e}")
