#!/usr/bin/env python3
"""DOCX后处理脚本：为pandoc生成的DOCX添加表格边框、调整格式"""

import argparse
import json
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, Cm


def add_table_borders(doc):
    """为文档中所有表格添加边框"""
    for t in doc.tables:
        tblPr = t._element.find(qn("w:tblPr"))
        if tblPr is None:
            from docx.oxml import OxmlElement
            tblPr = OxmlElement("w:tblPr")
            t._element.insert(0, tblPr)

        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)

        borders_xml = (
            '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            "</w:tblBorders>"
        )
        borders_element = parse_xml(borders_xml)
        tblPr.append(borders_element)


def set_table_header_bold(doc):
    """为所有表格的第一行（表头）设置加粗+居中"""
    for t in doc.tables:
        if len(t.rows) == 0:
            continue
        for cell in t.rows[0].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1  # CENTER
                for run in paragraph.runs:
                    run.bold = True
                    if run.font.size is None:
                        run.font.size = Pt(10)


def set_table_cell_font_size(doc, size_pt=10):
    """设置表格单元格字体大小"""
    for t in doc.tables:
        for row in t.rows[1:]:  # 跳过表头
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.size is None or run.font.size > Pt(size_pt + 1):
                            run.font.size = Pt(size_pt)


def set_page_margins(doc, top=2.54, bottom=2.54, left=3.17, right=3.17):
    """设置页边距（单位：厘米）"""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def main():
    parser = argparse.ArgumentParser(description="DOCX后处理：添加表格边框与格式调整")
    parser.add_argument("--input", required=True, help="输入DOCX文件路径")
    parser.add_argument("--output", required=True, help="输出DOCX文件路径")
    parser.add_argument("--no-borders", action="store_true", help="跳过添加表格边框")
    parser.add_argument("--no-header-bold", action="store_true", help="跳过表头加粗")
    parser.add_argument("--font-size", type=int, default=10, help="表格正文字号(磅)，默认10")
    parser.add_argument(
        "--margins",
        type=str,
        default=None,
        help="页边距(厘米)，格式: top,bottom,left,right",
    )
    args = parser.parse_args()

    doc = Document(args.input)
    table_count = len(doc.tables)
    image_count = len(doc.inline_shapes)

    if not args.no_borders:
        add_table_borders(doc)

    if not args.no_header_bold:
        set_table_header_bold(doc)

    if args.font_size:
        set_table_cell_font_size(doc, args.font_size)

    if args.margins:
        parts = [float(x) for x in args.margins.split(",")]
        if len(parts) == 4:
            set_page_margins(doc, *parts)

    doc.save(args.output)

    result = {
        "status": "success",
        "input": args.input,
        "output": args.output,
        "tables_processed": table_count,
        "images_found": image_count,
        "borders_added": not args.no_borders,
        "header_bold_set": not args.no_header_bold,
    }
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
