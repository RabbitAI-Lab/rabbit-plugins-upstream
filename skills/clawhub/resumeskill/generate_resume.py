#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将简历 JSON 渲染为统一排版的 Word 文档（.docx）。

用法：
    python generate_resume.py <resume.json> [output.docx]

输入 JSON 格式见 sample_resume.json。
"""

import json
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx is required.")
    print("Install with: pip install python-docx")
    sys.exit(1)


def set_run_font(run, font_name='Microsoft YaHei', size=10.5, bold=False, color=None):
    """统一设置 run 字体。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    """添加标题。"""
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(text)
        set_run_font(run, font_name='Microsoft YaHei', size=14, bold=True, color=(0, 0, 0))
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
    else:
        run = p.add_run(text)
        set_run_font(run, font_name='Microsoft YaHei', size=11, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet_line(doc, text, bold_prefix=None):
    """添加带项目符号的段落。"""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix and text.startswith(bold_prefix):
        run1 = p.add_run(bold_prefix)
        set_run_font(run1, bold=True)
        run2 = p.add_run(text[len(bold_prefix):])
        set_run_font(run2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_normal_text(doc, text, bold=False):
    """添加普通段落。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_two_column_line(doc, left_text, right_text, left_bold=True, right_bold=False):
    """添加左右对齐的一行（左侧公司/项目名，右侧时间）。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    # 添加右对齐制表位，让时间显示在页面右侧
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT)

    run_left = p.add_run(left_text)
    set_run_font(run_left, bold=left_bold)

    p.add_run('\t')

    run_right = p.add_run(right_text)
    set_run_font(run_right, bold=right_bold)
    run_right.italic = True
    return p


def render_resume(data: dict, output_path: str):
    doc = Document()

    # 页面边距
    sections = doc.sections[0]
    sections.top_margin = Inches(0.6)
    sections.bottom_margin = Inches(0.6)
    sections.left_margin = Inches(0.7)
    sections.right_margin = Inches(0.7)

    # 默认正文字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(10.5)

    # 1. 个人信息
    name = data.get('name', '姓名')
    phone = data.get('phone', '')
    email = data.get('email', '')
    wechat = data.get('wechat', '')
    location = data.get('location', '')
    links = data.get('links', [])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    set_run_font(run, font_name='Microsoft YaHei', size=18, bold=True)
    p.paragraph_format.space_after = Pt(4)

    contact_parts = [p for p in [phone, email, wechat, location] if p]
    if links:
        contact_parts.extend(links)
    if contact_parts:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(' | '.join(contact_parts))
        set_run_font(run2, size=9.5)
        p2.paragraph_format.space_after = Pt(10)

    # 2. 教育背景
    if data.get('education'):
        add_heading(doc, '教育背景')
        for edu in data['education']:
            school = edu.get('school', '')
            major = edu.get('major', '')
            degree = edu.get('degree', '')
            time = edu.get('time', '')
            gpa = edu.get('gpa', '')
            highlights = edu.get('highlights', [])

            title = f"{school} | {major} | {degree}"
            if gpa:
                title += f" | GPA: {gpa}"
            add_two_column_line(doc, title, time)

            for h in highlights:
                add_bullet_line(doc, h)

    # 3. 实习/工作经历
    if data.get('experience'):
        add_heading(doc, '实习/工作经历')
        for exp in data['experience']:
            company = exp.get('company', '')
            title = exp.get('title', '')
            time = exp.get('time', '')
            location = exp.get('location', '')
            header = f"{company} | {title}"
            if location:
                header += f" | {location}"
            add_two_column_line(doc, header, time)

            for desc in exp.get('description', []):
                add_bullet_line(doc, desc)

    # 4. 项目经历
    if data.get('projects'):
        add_heading(doc, '项目经历')
        for proj in data['projects']:
            name = proj.get('name', '')
            role = proj.get('role', '')
            time = proj.get('time', '')
            header = name
            if role:
                header += f" | {role}"
            add_two_column_line(doc, header, time)

            for desc in proj.get('description', []):
                add_bullet_line(doc, desc)

    # 5. 校园/组织经历
    if data.get('campus'):
        add_heading(doc, '校园/组织经历')
        for c in data['campus']:
            org = c.get('organization', '')
            role = c.get('role', '')
            time = c.get('time', '')
            header = org
            if role:
                header += f" | {role}"
            add_two_column_line(doc, header, time)

            for desc in c.get('description', []):
                add_bullet_line(doc, desc)

    # 6. 技能与证书
    if data.get('skills'):
        add_heading(doc, '技能与证书')
        skills = data['skills']
        if isinstance(skills, list):
            for s in skills:
                add_bullet_line(doc, s)
        elif isinstance(skills, dict):
            for category, items in skills.items():
                line = f"{category}：{items}" if isinstance(items, str) else f"{category}：{', '.join(items)}"
                add_bullet_line(doc, line)

    # 7. 其他
    if data.get('others'):
        add_heading(doc, '其他')
        for o in data['others']:
            add_bullet_line(doc, o)

    doc.save(output_path)
    print(f"Resume saved to: {output_path}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_resume.py <resume.json> [output.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else '改后简历.docx'

    with open(input_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    render_resume(data, output_path)


if __name__ == '__main__':
    main()
