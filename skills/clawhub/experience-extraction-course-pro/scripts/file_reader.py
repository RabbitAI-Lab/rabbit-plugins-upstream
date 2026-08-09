#!/usr/bin/env python3
"""完整、只读地提取课程资料中的文本，不修改或截断原文件。"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path


SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}
MAX_FILE_SIZE_MB = 50


def read_text_file(path: Path) -> str:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少依赖：python-docx==0.8.11") from exc

    document = Document(str(path))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    table_lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_lines.append(" | ".join(cells))

    return "\n".join(paragraphs + table_lines)


def read_pdf(path: Path) -> str:
    try:
        from PyPDF2 import PdfReader
    except ImportError as exc:
        raise RuntimeError("缺少依赖：PyPDF2==3.0.1") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"--- 第 {index} 页 ---\n{text.strip()}")
    return "\n\n".join(pages)


def read_file(path: Path) -> str:
    resolved = path.expanduser().resolve()
    if not resolved.exists() or not resolved.is_file():
        raise FileNotFoundError(f"文件不存在：{resolved}")

    size_mb = resolved.stat().st_size / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise ValueError(
            f"文件过大：{size_mb:.1f} MB；当前上限为 {MAX_FILE_SIZE_MB} MB"
        )

    extension = resolved.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        allowed = "、".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"不支持的文件类型：{extension}；支持类型：{allowed}")

    if extension in {".txt", ".md"}:
        return read_text_file(resolved)
    if extension == ".docx":
        return read_docx(resolved)
    if extension == ".pdf":
        return read_pdf(resolved)

    raise ValueError(f"不支持的文件类型：{extension}")


def configure_stdout() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")


def main() -> int:
    configure_stdout()
    parser = argparse.ArgumentParser(
        description="提取课程资料中的可读文本。",
        add_help=False,
    )
    parser.add_argument("-h", "--help", action="help", help="显示帮助并退出")
    parser.add_argument(
        "--file_path",
        required=True,
        help="资料文件路径，支持 .txt、.md、.docx 和 .pdf",
    )
    args = parser.parse_args()

    try:
        print(read_file(Path(args.file_path)))
    except (FileNotFoundError, OSError, RuntimeError, ValueError) as exc:
        parser.exit(1, f"错误：{exc}\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
