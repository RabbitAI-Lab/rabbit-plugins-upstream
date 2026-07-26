#!/usr/bin/env python3
"""Extract text from a .docx file.

Usage:
    python extract_docx_text.py <docx_path> [out_txt]

Tries (in order): python-docx, docx2txt. If neither installed, prints hint.
"""
from __future__ import annotations

import sys
from pathlib import Path


def _try_python_docx(docx_path: str) -> str:
    from docx import Document

    doc = Document(docx_path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _try_docx2txt(docx_path: str) -> str:
    import docx2txt

    return docx2txt.process(docx_path) or ""


def main() -> int:
    if len(sys.argv) < 2:
        print("Usage: extract_docx_text.py <docx_path> [out_txt]", file=sys.stderr)
        return 1

    docx_path = sys.argv[1]
    out_path = sys.argv[2] if len(sys.argv) >= 3 else None

    if not Path(docx_path).exists():
        print(f"ERROR: file not found: {docx_path}", file=sys.stderr)
        return 1

    text = None
    last_err: Exception | None = None
    for fn in (_try_python_docx, _try_docx2txt):
        try:
            text = fn(docx_path)
            break
        except Exception as e:  # noqa: BLE001
            last_err = e
            continue

    if text is None:
        print(
            "ERROR: no DOCX library available. Install one of:\n"
            "  pip install python-docx\n"
            "  pip install docx2txt\n"
            f"Last error: {last_err}",
            file=sys.stderr,
        )
        return 2

    if out_path:
        Path(out_path).write_text(text, encoding="utf-8")
        print(f"Wrote {len(text):,} chars to {out_path}")
    else:
        sys.stdout.write(text)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
