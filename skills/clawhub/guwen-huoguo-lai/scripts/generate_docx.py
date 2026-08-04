#!/usr/bin/env python3
"""
古文活过来 - Word 剧本生成器（提示词参考卡放最前方版）
用法: python3 generate_docx.py <输出路径>
"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def border(table):
    tbl = table._tbl
    pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    b = OxmlElement('w:tblBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{e}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'6')
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),'999999')
        b.append(el)
    pr.append(b)

def create_cover(doc, title, account, duration, style, version="v2.4"):
    for _ in range(6): doc.add_paragraph("")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('\u300a'+title+'\u300b'); r.font.size=Pt(28); r.font.bold=True
    r.font.color.rgb = RGBColor(0x8B,0x45,0x13)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('AI \u53e4\u6587\u52a8\u753b\u77ed\u89c6\u9891 \u00b7 \u62cd\u6444\u811a\u672c'); r.font.size=Pt(16)
    r.font.color.rgb = RGBColor(0x66,0x66,0x66)
    doc.add_paragraph("")
    data = [('\u8d26\u53f7\u540d',account),('\u603b\u65f6\u957f',duration),('\u753b\u98ce',style),('\u7248\u672c',version),('\u751f\u6210\u65e5\u671f','2026-07-29')]
    t = doc.add_table(rows=len(data),cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER; border(t)
    for i,(k,v) in enumerate(data):
        c0=t.rows[i].cells[0]; c0.text=k; c0.width=Cm(4); shade(c0,'F5F0E8')
        for pp in c0.paragraphs: pp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        for rr in c0.paragraphs[0].runs: rr.font.size=Pt(12); rr.font.bold=True
        c1=t.rows[i].cells[1]; c1.text=v; c1.width=Cm(10)
        for rr in c1.paragraphs[0].runs: rr.font.size=Pt(12)

def add_ref_card(doc, characters, scene_base, style_desc):
    """在封面之后、分镜之前添加统一的提示词参考卡"""
    p = doc.add_paragraph()
    r = p.add_run('\U0001f4cb \u7edf\u4e00\u63d0\u793a\u8bcd\u53c2\u8003\u5361'); r.font.size=Pt(18); r.font.bold=True
    r.font.color.rgb = RGBColor(0x8B,0x45,0x13)

    p = doc.add_paragraph()
    r = p.add_run('\u2501'*40); r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC); r.font.size=Pt(8)

    # 角色提示词
    p = doc.add_paragraph()
    r = p.add_run('\U0001f9d1 \u89d2\u8272\u63d0\u793a\u8bcd'); r.font.size=Pt(14); r.font.bold=True
    r.font.color.rgb = RGBColor(0x2C,0x3E,0x50)
    for char in characters:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(f'\u2022 {char["name"]}'); r.font.bold=True; r.font.size=Pt(11)
        r.font.color.rgb = RGBColor(0x2C,0x3E,0x50)
        r = p.add_run(f'\uff1a{char["prompt"]}'); r.font.size=Pt(9)
        r.font.color.rgb = RGBColor(0x66,0x66,0x66)

    doc.add_paragraph("")

    # 场景基调提示词
    p = doc.add_paragraph()
    r = p.add_run('\U0001f3d7\ufe0f \u573a\u666f\u57fa\u8c03\u63d0\u793a\u8bcd'); r.font.size=Pt(14); r.font.bold=True
    r.font.color.rgb = RGBColor(0x8B,0x45,0x13)
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(scene_base); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x66,0x66,0x66)

    doc.add_paragraph("")

    # 画面风格提示词
    p = doc.add_paragraph()
    r = p.add_run('\U0001f3a8 \u753b\u9762\u98ce\u683c\u63d0\u793a\u8bcd'); r.font.size=Pt(14); r.font.bold=True
    r.font.color.rgb = RGBColor(0x8B,0x45,0x13)
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(style_desc); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x66,0x66,0x66)

    doc.add_page_break()

def add_scene(doc, num, time_range, scene_desc, camera, chars, subtitles, sound):
    p = doc.add_paragraph()
    r = p.add_run(f'\u573a\u666f{num} \uff5c {time_range}'); r.font.size=Pt(16); r.font.bold=True
    r.font.color.rgb = RGBColor(0x8B,0x45,0x13)
    p = doc.add_paragraph()
    r = p.add_run('\u2501'*40); r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC); r.font.size=Pt(8)

    # 画面
    if scene_desc:
        p=doc.add_paragraph()
        r=p.add_run('\U0001f3a8 \u753b\u9762\uff1a'); r.font.bold=True; r.font.size=Pt(11)
        r.font.color.rgb = RGBColor(0x8B,0x45,0x13)
        r=p.add_run(scene_desc); r.font.size=Pt(11)

    # 运镜
    if camera:
        p=doc.add_paragraph()
        r=p.add_run('\U0001f3a5 \u8fd0\u955c\uff1a'); r.font.bold=True; r.font.size=Pt(11)
        r.font.color.rgb = RGBColor(0x00,0x55,0x99)
        r=p.add_run(camera); r.font.size=Pt(10); r.font.color.rgb=RGBColor(0x00,0x55,0x99)

    # 角色台词
    for c in chars:
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(1)
        r=p.add_run(f'\u2502 {c["name"]}\uff08{c["tone"]}\uff09')
        r.font.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x2C,0x3E,0x50)
        r=p.add_run(f'\uff1a\u201c{c["line"]}\u201d'); r.font.size=Pt(11); r.font.italic=True

    # 字幕
    if subtitles:
        p=doc.add_paragraph()
        r=p.add_run('\U0001f4dd \u5b57\u5e55\uff1a'); r.font.bold=True; r.font.size=Pt(11)
        r=p.add_run(subtitles); r.font.size=Pt(10); r.font.color.rgb=RGBColor(0x88,0x44,0x00)

    # 音效
    if sound:
        p=doc.add_paragraph()
        r=p.add_run('\U0001f50a \u97f3\u6548\uff1a'); r.font.bold=True; r.font.size=Pt(11)
        r=p.add_run(sound); r.font.size=Pt(10); r.font.color.rgb=RGBColor(0x55,0x55,0x55)


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "output.docx"
    doc = Document()
    for s in doc.sections:
        s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Cm(2.5)

    # 封面
    create_cover(doc, '陈太丘与友期行', '古文活过来', '3分00秒', '写实古风CG（暖金/琥珀色调）')

    # 统一提示词参考卡（在封面后、分镜前）
    characters = [
        {'name': '陈太丘', 'prompt': '45岁左右中年男子，相貌威严端庄，深蓝色直裾汉服，束发冠。表情平静略带忧虑，站姿挺拔。电影级CG写实人像。'},
        {'name': '友人', 'prompt': '40岁左右，体面富贵相，棕色镶金边广袖汉服，束发高冠。骑马，愤怒时眉头紧皱、动作幅度大。电影级CG写实人像。'},
        {'name': '元方', 'prompt': '7岁男孩，浅蓝色小汉服，束小发髻（总角），白净面容，目光平静坚定，身型不高但站姿笔直。电影级CG写实人像。'},
    ]
    scene_base = '时代：东汉。场景元素：青瓦屋顶庭院、木门铜环、垂柳、石板路、仆人牵马。光线：下午至黄昏暖金色阳光，琥珀色调。构图：电影感写实CG，浅景深。调色：暖金色/琥珀色主调。统一9:16竖屏。'
    style_desc = '写实古风CG动画风格，电影质感，暖金色调。人物写实，场景渲染结合水墨渲染元素（转场时用淡墨过渡）。自然光影，温暖怀旧氛围。'
    add_ref_card(doc, characters, scene_base, style_desc)

    # 分镜
    scenes = [
        ('1','00:00-00:12',
         '黄昏奶茶店门口，一个穿蓝白校服的13岁初中男生独自等待，看手机、看表、脸上有些焦急。街灯刚刚亮起，湿润的柏油路面折射着霓虹灯的暖光。幽蓝色调，忧郁氛围。',
         '【固定→缓慢拉远】开场特写男孩手指在手机屏幕上滑动→拉远到中景，显示他站在奶茶店门口的全身。6秒后切拉远到路口远景，孤独背影。',
         [], '朋友迟到该不该等？', '街灯亮起"咔"声，环境音渐弱'),
        ('2','00:12-00:22',
         '手机屏幕融化为墨液——墨汁四散流动，渐渐形成一座东汉庭院的水墨画，木门铜环、青瓦屋顶、垂柳如丝。水墨淡去，写实场景越来越清晰。暖金色夕阳透出。',
         '【融镜变形】手机屏幕占满画面→黑色墨汁从中心扩散→流动形成水墨庭院→缓缓融入写实场景。4秒平滑过渡，水墨与写实并存。',
         [], '', '水墨流动声 + 古琴低声起'),
        ('3','00:22-00:50',
         '东汉庭院门口，陈太丘穿深蓝汉服站在门前，望向远方的空旷道路。下午阳光在地上拉出长影。青瓦屋顶、木门、仆人牵马。暖金色光线，琥珀色调。',
         '【横摇→推】庭院全景横摇到大门→缓缓推向陈太丘背影→切侧脸特写，目光望向远方。推到眼神时停顿2秒。横移跟随上马。',
         [{'name':'陈太丘','tone':'平静自语','line':'已经正午了。'}],
         '陈太丘与友期行，期日中。过中不至，太丘舍去。',
         '马蹄声，木门关闭"吱——咔"'),
        ('4','00:50-01:20',
         '友人穿棕色金边汉服骑马赶到，见大门紧闭，怒气冲冲重重敲门。马拴在身后。下午阳光打出强烈的阴影分界线，紧张氛围。',
         '【手持微晃 + 跟】友人骑马入画，手持微晃跟马跑动→翻身下马切固定中景→大步到门前，镜头横移→敲门时急剧推到拳头特写。',
         [{'name':'友人','tone':'怒喝','line':'约好中午同行，你竟丢下我自己走了？非人哉！'}],
         '友人便怒曰："非人哉！与人期行，相委而去！"',
         '急促马蹄声，重重敲门声'),
        ('5','01:20-02:10',
         '大门缓缓打开，7岁元方穿浅蓝汉服站在门内，目光平静坚定。夕阳在他身上打出金色轮廓光。门外友人怒目而视。',
         '【固定→推】固定全景，大门缓缓打开→元方身影出现在门内→缓缓推到他面部特写。他说话时镜头定格在他脸上。对方声音从背景传来，不切镜。',
         [{'name':'元方','tone':'平静直视','line':'您与我父亲约好正午。您正午未到——是您不守信。现在对着儿子骂父亲——是您没有礼貌。'}],
         '元方曰："君与家君期日中。日中不至，则是无信；对子骂父，则是无礼。"',
         '开门声，安静中只有风声'),
        ('6','02:10-02:40',
         '友人呆站在关上的大门外，表情从怒气到惊愕到羞愧，低下头。元方在院内走开不回头。落日余晖洒在空荡的门口，悲伤氛围。',
         '【缓慢拉远 + 升降下降】从元方转身特写缓缓拉远→镜头随他升高俯拍门外友人背影。友人定格3秒，然后缓缓拉远到全景。',
         [{'name':'友人','tone':'低声羞愧','line':'元方……'}],
         '友人惭，下车引之。元方入门不顾。',
         '门关上"吱"声，轻叹'),
        ('7','02:40-03:00',
         '元方站在黄昏庭院门内，缓缓转身面对镜头，微微一笑。微风吹动衣角。暖琥珀色夕阳照在脸上，气氛温暖坚定。',
         '【缓缓推近 + 环绕】元方转身面对镜头→推到面部特写→镜头微微环绕左右。说到最后一句时缓缓拉远到中景，画面渐暗到黑屏。',
         [{'name':'元方','tone':'温和坚定，面向观众','line':'守信和礼貌——这是1800年前我坚持的事。换到今天，你的朋友迟到了——你会等，还是转身离开？'}],
         '如果是你，会等吗？评论区聊聊', '古琴BGM渐弱，轻柔结束。灯光渐暗到黑屏'),
    ]

    for i, s in enumerate(scenes):
        add_scene(doc, s[0], s[1], s[2], s[3], s[4], s[5], s[6])
        if i < len(scenes) - 1:
            doc.add_page_break()

    # 结尾
    for _ in range(2): doc.add_paragraph("")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('— END —'); r.font.size=Pt(14); r.font.color.rgb=RGBColor(0x99,0x99,0x99)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('关注「古文活过来」· 让课文活过来 📖')
    r.font.size=Pt(12); r.font.color.rgb=RGBColor(0x8B,0x45,0x13)

    doc.save(out)
    print(f'✅ Done: {out}')
