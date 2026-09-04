# -*- coding: utf-8 -*-
"""Assemble the double-sided Word file from the all_front/all_back PNGs.
Page order: front1, back1, front2, back2, ... (ALTERNATING — required so a
long-edge flipped double-sided printer pairs each sheet's front/back correctly).
Back page = that group's front with the two columns MIRRORED.
"""
import os, sys
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
from words100 import WORDS

N = len(WORDS); GRP = 4
assert N % GRP == 0, "word count must be a multiple of 4"
PAGES = N // GRP
FR_DIR = os.path.join(HERE, "all_front"); BK_DIR = os.path.join(HERE, "all_back")
OUT = os.environ.get("OUT_DOCX", os.path.join(HERE, "English_Words_DoubleSided.docx"))

MARGIN = Mm(10); COL_W = Mm(95); COL_H = Mm(135)

def ffile(i): return f"{FR_DIR}/c{i:03d}_{WORDS[i][0]}.png"
def bfile(i): return f"{BK_DIR}/c{i:03d}_{WORDS[i][0]}.png"

def no_borders(t):
    tblPr = t._tbl.tblPr
    for o in tblPr.findall(qn("w:tblBorders")): tblPr.remove(o)
    b = OxmlElement("w:tblBorders")
    for e in ["top","left","bottom","right","insideH","insideV"]:
        el = OxmlElement(f"w:{e}"); el.set(qn("w:val"),"none"); el.set(qn("w:sz"),"0")
        el.set(qn("w:space"),"0"); el.set(qn("w:color"),"auto"); b.append(el)
    tblPr.append(b)
    w = OxmlElement("w:tblW"); w.set(qn("w:w"),"5000"); w.set(qn("w:type"),"pct"); tblPr.append(w)

def fill(cell, path):
    cell.width = COL_W; cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(p.runs): r._element.getparent().remove(r._element)
    tcPr = cell._tc.get_or_add_tcPr(); tcMar = OxmlElement("w:tcMar")
    for e in ["start","end","top","bottom"]:
        el = OxmlElement(f"w:{e}"); el.set(qn("w:w"),"0"); el.set(qn("w:type"),"dxa"); tcMar.append(el)
    tcPr.append(tcMar)
    p.add_run().add_picture(path, width=COL_W, height=COL_H)

doc = Document()
def section_margins(sec):
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = MARGIN

def add_page(files, mirror, first):
    if first: sec = doc.sections[0]; section_margins(sec)
    else: sec = doc.add_section(WD_SECTION.NEW_PAGE); section_margins(sec)
    t = doc.add_table(rows=2, cols=2); t.autofit = False; no_borders(t)
    layout = ([(0,1,files[0]),(0,0,files[1]),(1,1,files[2]),(1,0,files[3])] if mirror
              else [(0,0,files[0]),(0,1,files[1]),(1,0,files[2]),(1,1,files[3])])
    for r,c,img in layout:
        row = t.rows[r]
        fill(row.cells[c], img)

first = True
for k in range(PAGES):
    base = k*GRP
    add_page([ffile(base+i) for i in range(GRP)], mirror=False, first=first); first = False
    add_page([bfile(base+i) for i in range(GRP)], mirror=True,  first=False)

doc.core_properties.title   = f"English Words · {N} · double-sided"
doc.core_properties.subject = f"Kid flashcards · {N} English words · double-sided (long-edge flip)"
doc.core_properties.author  = "english-flashcards-word"
doc.save(OUT)

from docx import Document as D
print("Wrote", OUT, os.path.getsize(OUT)//1024, "KB |",
      len(D(OUT).tables), "tables |", len(D(OUT).sections), "sections | pages:", PAGES)
