#!/usr/bin/env python3
"""提取简历文本(PDF / DOCX / TXT / MD),供测评流程使用。

用法:
  uv run --with pymupdf --with python-docx python extract_resume.py 简历.pdf
  uv run --with pymupdf --with python-docx python extract_resume.py 简历.pdf -o resume.txt

说明:
  - PDF 取文本层;扫描件/图片 PDF 请先 OCR(如 ocrmypdf)再跑本脚本
  - 输出到 stdout,或 -o 指定文件
"""
import argparse
import sys
from pathlib import Path


def extract(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        import pymupdf  # pymupdf 新接口(旧名 fitz 已弃用)
        doc = pymupdf.open(path)
        return "\n".join(page.get_text() for page in doc)
    if ext == ".docx":
        import docx
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    if ext in (".txt", ".md", ".json", ".csv"):
        return path.read_text(encoding="utf-8", errors="replace")
    sys.exit(f"不支持的文件类型: {ext}(支持 pdf/docx/txt/md)")


def main() -> None:
    ap = argparse.ArgumentParser(description="提取简历文本")
    ap.add_argument("resume", help="简历文件路径")
    ap.add_argument("-o", "--output", default=None, help="输出文本文件(默认 stdout)")
    args = ap.parse_args()
    p = Path(args.resume)
    if not p.exists():
        sys.exit(f"文件不存在: {p}")
    text = extract(p)
    if args.output:
        Path(args.output).write_text(text, encoding="utf-8")
        print(f"已写入 {args.output} ({len(text)} 字符)")
    else:
        print(text)


if __name__ == "__main__":
    main()
