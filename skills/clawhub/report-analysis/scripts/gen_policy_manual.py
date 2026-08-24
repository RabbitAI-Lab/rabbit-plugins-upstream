# -*- coding: utf-8 -*-
# 政策速查手册生成器：把三份销售合作伙伴政策文档整理为带目录的 Word 速查手册
# 结构：概览表 → 按文档分章（章节树+核心条款带出处）→ 附录A 术语映射 → 附录B 政策<->数据联动速查
import os
import json
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from analyzer import load_analyzed, summarize, load_county_summary, fmt_wan

HERE = os.path.dirname(os.path.abspath(__file__))
IDX = os.path.join(HERE, 'policy_docs.json')
TM = os.path.join(HERE, 'term_map.json')
OUT = os.environ.get('YJ_POLICY_OUT', os.path.join(HERE, '中国移动云南公司销售合作伙伴政策速查手册.docx'))

DARK_BLUE = RGBColor(0x1F, 0x3B, 0x73)
GRAY = RGBColor(0x88, 0x88, 0x88)

RULES = re.compile(r'(必须|应当|应|不得|禁止|严禁|需要|可以|扣|罚|奖|激励|返还|折算|系数|门槛|否决|档|投诉|合约|融合|签约|退出|清退|解除|星级|评定|报备|审批|备案|公示|考核|评估|标准|公式|上限|占比|比例|%|折|/|＋|\+)')

with open(IDX, encoding='utf-8') as f:
    DOCS = json.load(f)
with open(TM, encoding='utf-8') as f:
    TERMS = json.load(f)

chans = load_analyzed()
S = summarize(chans)
CS = load_county_summary()


def set_run(r, size=10.5, bold=False, color=None):
    r.font.name = '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = color


def add_para(doc, text, size=10.5, bold=False, color=None, align=None, space_after=4, indent=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size, bold, color)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def style_table(t, widths):
    for r in t.rows:
        for i, c in enumerate(r.cells):
            if i < len(widths):
                c.width = Cm(widths[i])
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_run(run, 9)


def add_table(doc, headers, rows, widths, header_fill='D9E2F3'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        shade_cell(cell, header_fill)
        for p in cell.paragraphs:
            for run in p.runs:
                set_run(run, 9.5, True)
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = str(v)
    style_table(t, widths)
    return t


def sec_text(item):
    return ' / '.join(item['sec']) if item['sec'] else ''


def core_items(doc, limit=30):
    """抽取规则性条款，按章节聚合"""
    out = []
    seen = set()
    for it in doc['items']:
        t = it['text']
        if len(t) < 18 or t in seen:
            continue
        if RULES.search(t):
            seen.add(t)
            out.append(it)
        if len(out) >= limit:
            break
    return out


def build():
    doc = Document()
    # 封面标题
    add_para(doc, '中国移动云南公司销售合作伙伴', 20, True, DARK_BLUE, WD_ALIGN_PARAGRAPH.CENTER, 0)
    add_para(doc, '政策速查手册', 20, True, DARK_BLUE, WD_ALIGN_PARAGRAPH.CENTER, 6)
    add_para(doc, '整理自三份政策文件 · 含版本标注与术语对照', 11, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER, 2)
    add_para(doc, '配套「评优激励通报分析」技能使用 · 整理日期 2026-08-20', 9, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER, 12)

    # 一、概览
    add_para(doc, '一、政策文件概览', 14, True, DARK_BLUE, space_after=6)
    rows = []
    for d in DOCS:
        rows.append([d['short'], d['year'] + '年', d['layer'], d['desc'],
                     str(d['n_paras']) + '段/' + str(d['n_tables']) + '表/' + '{:,}'.format(d['chars']) + '字'])
    add_table(doc, ['文档', '版本', '层级', '内容定位', '篇幅'], rows, [3.6, 1.2, 1.4, 7.2, 2.6])
    add_para(doc, '', 4, space_after=2)

    # 二、按文档分章
    add_para(doc, '二、核心政策规则速查（带原文出处）', 14, True, DARK_BLUE, space_after=6)
    for d in DOCS:
        add_para(doc, '【' + d['short'] + '】（' + d['year'] + '年版 · ' + d['layer'] + '）', 12, True, DARK_BLUE, space_after=4)
        add_para(doc, '定位：' + d['desc'], 9.5, False, GRAY, space_after=4)
        # 章节树
        seen = set()
        tree = []
        for it in d['items']:
            if it['sec'] and it['sec'][0] not in seen:
                seen.add(it['sec'][0])
                tree.append(it['sec'][0])
        if tree:
            add_para(doc, '章节结构：' + ' → '.join(tree[:12]) + (' …' if len(tree) > 12 else ''), 9, False, GRAY, space_after=4)
        # 核心条款
        items = core_items(d)
        if not items:
            add_para(doc, '（未提取到明显规则条款）', 9.5, space_after=4)
        for i, it in enumerate(items, 1):
            sec = sec_text(it)
            head = str(i) + '. ' + (sec if sec else '（无章节）')
            add_para(doc, head, 9.5, True, space_after=1, indent=0.3)
            add_para(doc, it['text'], 9.5, space_after=3, indent=0.6)
        add_para(doc, '', 4, space_after=2)

    # 附录A：术语映射
    doc.add_page_break()
    add_para(doc, '附录A：业务术语对照表（Excel 核算术语 ↔ 政策文档表述）', 13, True, DARK_BLUE, space_after=6)
    rows = []
    for k, v in TERMS.items():
        rows.append([k, '、'.join(v)])
    add_table(doc, ['Excel / 业务术语', '政策文档中的相关表述（检索关键词）'], rows, [4.5, 10.5])

    # 附录B：政策<->数据联动
    add_para(doc, '', 4, space_after=2)
    add_para(doc, '附录B：政策 ⇄ 数据联动速查（2026年3季度评优激励实测）', 13, True, DARK_BLUE, space_after=6)
    qz = CS.get('全州', {})
    gate_zero = sum(1 for c in chans if c['gate_level'] == '门槛未完成' and c['tiger'] == 0 and c['ai5'] == 0 and c['rights_up'] == 0 and c['member88'] == 0)
    coef08 = sum(1 for c in chans if abs(c['app_coef'] - 0.8) < 0.001)
    rows = [
        ['门槛/否决项', '未达档 ' + str(S['gate_notdone']) + ' 家（' + '{:.1f}'.format(S['gate_notdone_rate']) + '%），完全无业务 ' + str(gate_zero) + ' 家，单道损收 ' + fmt_wan(abs(qz.get('gate_loss', 0)) * 10000)],
        ['激励总额', '原始 ' + fmt_wan(S['raw_total']) + ' → 最终 ' + fmt_wan(S['final_total']) + '，损失 ' + fmt_wan(S['loss_total']) + '（' + '{:.1f}'.format(S['loss_rate']) + '%），0 元渠道 ' + str(S['zero_final']) + ' 家'],
        ['APP融合率', '新入网 ' + '{:.1f}'.format(S['avg_newnet_fuse']) + '% / 新终端 ' + '{:.1f}'.format(S['avg_newterm_fuse']) + '% / 宽带 ' + '{:.1f}'.format(S['avg_bb_fuse']) + '%，系数 0.8 渠道 ' + str(coef08) + ' 家'],
        ['终端合约率', '均值 ' + '{:.1f}'.format(S['avg_term']) + '%，为 0 渠道 ' + str(S['zero_term']) + ' 家，单道损收 ' + fmt_wan(abs(qz.get('term_loss', 0)) * 10000)],
        ['重点业务牵引', '系数 0.81 共 ' + str(S['coef081']) + ' 家 / 0.9 共 ' + str(S['coef09']) + ' 家 / 1.0 共 ' + str(S['coef1']) + ' 家，单道损收 ' + fmt_wan(abs(qz.get('focus_loss', 0)) * 10000)],
        ['弱势网格', '单道损收 ' + fmt_wan(abs(qz.get('weakgrid_loss', 0)) * 10000)],
        ['有责投诉', '当期未出数按不扣罚计，存在追溯扣罚风险'],
    ]
    add_table(doc, ['政策主题', '2026年3季度实测数据'], rows, [3.2, 11.8])

    doc.save(OUT)
    print('速查手册已生成:', OUT)


if __name__ == '__main__':
    build()
