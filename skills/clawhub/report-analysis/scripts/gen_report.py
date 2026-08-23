# -*- coding: utf-8 -*-
# 功能C：总结分析 Word 报告生成器
# 结构：标题+数据来源 → 一、总体评价(县市损失一览) → 二、主要欠缺问题
#      → 三、改进建议 → 附1 重点网格整改清单 → 附2 网格汇总表(横向)
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from analyzer import load_analyzed, summarize, by_county, by_grid, load_county_summary, fmt_wan

SRC_NAME = os.environ.get('YJ_SRC_NAME', '2026年3季度合作伙伴评优激励通报20260811.xlsx')
OUT = os.environ.get('YJ_OUT',
                     r'D:/AI/WorkBuddy/2026-08-20-16-02-35/analysis/2026年3季度合作伙伴评优激励通报-数据欠缺分析报告.docx')

DARK_BLUE = RGBColor(0x1F, 0x3B, 0x73)
RED = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x70, 0x00)
GRAY = RGBColor(0x88, 0x88, 0x88)

chans = load_analyzed()
S = summarize(chans)
CS = load_county_summary()
counties = by_county(chans)
grids = by_grid(chans)


def set_cn_font(style, size=None, bold=None):
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if size:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def add_para(doc, text, size=10.5, bold=False, color=None, align=None, space_after=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def make_table(doc, header, rows, widths=None, header_fill='1F3B73', font_size=8.5):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(font_size)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        shade_cell(hdr[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
                    r.font.name = '微软雅黑'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    return t


# ============ 文档主体 ============
doc = Document()
style = doc.styles['Normal']
set_cn_font(style, size=10.5)
for s in doc.sections:
    s.left_margin = Cm(1.8)
    s.right_margin = Cm(1.8)

add_para(doc, '2026年3季度合作伙伴评优激励 · 数据欠缺分析报告', 16, True, DARK_BLUE,
         WD_ALIGN_PARAGRAPH.CENTER, 2)
add_para(doc, '（基于 ' + SRC_NAME + '）', 9, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER, 8)

# 数据来源
add_para(doc, '数据来源与说明', 12, True, DARK_BLUE)
add_para(doc, '本报告基于《' + SRC_NAME + '》核算，主表「渠道完成情况通报0811」共 397 家渠道、90 项指标；'
         '激励金额经否决门槛→终端合约搭载→重点业务牵引系数→有责投诉考核→弱势网格扣罚→APP融合率六道核算逐级递减；'
         '有责投诉当期未出数，主表按先按不扣罚计处理。金额口径：原始兑换金额合计 '
         + fmt_wan(S['raw_total']) + '，最终可发金额 ' + fmt_wan(S['final_total']) + '，'
         '全州损失 ' + fmt_wan(S['loss_total']) + '（损失率 ' + '{:.1f}'.format(S['loss_rate']) + '%）。', 9.5)

# 一、总体评价
add_para(doc, '一、总体评价', 13, True, DARK_BLUE)
got5 = sorted([x for x in chans if x['final'] > 0], key=lambda c: -c['final'])[:5]
add_para(doc, '本季激励核算结果整体偏差：397 家渠道中 369 家（92.9%）未通过否决门槛，344 家（86.6%）'
         '最终可发金额为 0 元，另有 31 家因扣罚超过可发额出现负值（按 0 计）。全州原始兑换金额 '
         + fmt_wan(S['raw_total']) + '，经六道核算后仅剩 ' + fmt_wan(S['final_total']) + '，'
         '流失 ' + fmt_wan(S['loss_total']) + '，损失率 ' + '{:.1f}'.format(S['loss_rate']) + '%。'
         '激励向头部少数渠道集中：有激励的渠道仅 ' + str(sum(1 for c in chans if c['final'] > 0)) + ' 家，'
         '头部前 5 名合计 ' + fmt_wan(sum(c['final'] for c in got5)) + '。'
         '问题主要集中在门槛业务（尤其 88 会员年包、AI 五件套）大面积空白、重点业务牵引系数普遍处于最低档、'
         '弱势网格攻坚参与不足三个方面，属渠道整体业务能力与参与意愿的双重短板，而非个别数据异常。', 10)

add_para(doc, '（一）各县（市）损失一览', 11, True, DARK_BLUE)
rows = []
for c in counties:
    cs = CS[c['county']]
    rows.append([
        c['county'], str(c['n']) + '家', fmt_wan(c['raw']), fmt_wan(c['final']),
        fmt_wan(c['loss']), '{:.0f}%'.format(c['loss_rate']),
        '{:.2f}万'.format(abs(cs['gate_loss'])), '{:.2f}万'.format(abs(cs['term_loss'])),
        '{:.2f}万'.format(abs(cs['focus_loss'])), '{:.2f}万'.format(abs(cs['weakgrid_loss'])),
        '{:.2f}万'.format(abs(cs['app_loss'])),
    ])
tot = CS['全州']
rows.append(['全州', str(S['n']) + '家', fmt_wan(S['raw_total']), fmt_wan(S['final_total']),
             fmt_wan(S['loss_total']), '{:.0f}%'.format(S['loss_rate']),
             '{:.2f}万'.format(abs(tot['gate_loss'])), '{:.2f}万'.format(abs(tot['term_loss'])),
             '{:.2f}万'.format(abs(tot['focus_loss'])), '{:.2f}万'.format(abs(tot['weakgrid_loss'])),
             '{:.2f}万'.format(abs(tot['app_loss']))])
t = make_table(doc, ['县（市）', '渠道数', '原始金额', '最终金额', '损失', '损失率',
                     '门槛损收', '终端损收', '重点业务损收', '弱势网格损收', 'APP融合损收'],
               rows, widths=[1.7, 1.2, 2.0, 2.0, 2.0, 1.3, 1.8, 1.8, 2.2, 2.2, 2.2])
for r in t.rows[1:]:
    if r.cells[0].text == '全州':
        for cell in r.cells:
            shade_cell(cell, 'DCE6F1')
add_para(doc, '注：景洪市终端搭载环节为正（+0.10万），因其合约率高的渠道获得系数加成；投诉考核未出数，损收暂为 0。', 8.5,
         False, GRAY, space_after=8)

# 二、主要欠缺问题
add_para(doc, '二、主要欠缺问题（按严重程度排序）', 13, True, DARK_BLUE)
n = S['n']
got_cnt = sum(1 for c in chans if c['final'] > 0)
problems = [
    ('1. 否决门槛大面积未达档（最严重）',
     '369 家（92.9%）门槛未完成，仅 28 家达档（2档4家/3档7家/4档8家/5档9家，无1档）。'
     '门槛单道损收 ' + fmt_wan(abs(tot['gate_loss']) * 10000) + '，占原始金额 49.0%，为六道核算中最大出血点。'
     '其中 215 家（58%）四项门槛业务完成量全为 0——完全无业务；其余 154 家有量但未达档位线。',
     '门槛业务未前置、渠道无抓手：88会员/99会员年包完成量为 0 的渠道占 92.9%（完成率最低单品），'
     'AI五件套 71.0%、金虎 76.8%、权益+升档 63.2% 的渠道为零。',
     '按档位逐户下发四项差距（如1档需 AI五件套≥20、金虎≥100、权益升档≥1000、88会员≥500），季初签约前置宣贯；'
     '对四项全 0 渠道建帮扶台账，季度中每周通报进度。'),
    ('2. 关键会员/套餐业务乏力',
     '88会员/99会员年包：369 家（92.9%）为 0；金虎：305 家（76.8%）为 0；AI五件套：282 家（71.0%）为 0。',
     '高价值产品依赖柜台开口推荐，渠道缺乏话术与激励刺激，普遍停留在等客上门、只办卡的粗放经营。',
     '把年包/AI五件套列为月度必做动作并按周排名；对头部渠道试点 AI 体验区；会员年包主推老客关怀回访。'),
    ('3. 重点业务发展不均衡，牵引系数普遍最低档',
     '牵引系数 0.81（最低档）渠道 ' + str(S['coef081']) + ' 家（' + '{:.0f}%'.format(S['coef081'] / n * 100) + '），'
     '0.9 档 ' + str(S['coef09']) + ' 家，1.0 满档仅 ' + str(S['coef1']) + ' 家。该道损收 '
     + fmt_wan(abs(tot['focus_loss']) * 10000) + '。',
     '存量/产品运营积分占比不达标，渠道重新增、轻存量运营，运营积分结构失衡。',
     '对 0.81 档渠道逐户核对存量运营积分、产品运营积分缺口；将牵引系数纳入季度中预警指标而非季末结果。'),
    ('4. 弱势网格攻坚参与度低、履约差',
     '签约渠道仅 ' + str(S['signed']) + ' 家（' + '{:.1f}%'.format(S['signed'] / n * 100) + '），'
     '其中 ' + str(S['signed_zero_done']) + ' 家签约后 4 项业务为 0（占签约数 '
     + '{:.0f}%'.format(S['signed_zero_done'] / S['signed'] * 100) + '）。扣罚损收 '
     + fmt_wan(abs(tot['weakgrid_loss']) * 10000) + '。',
     '签约门槛低、缺乏过程跟踪，部分渠道签而不战。',
     '签约前评估承接能力；按月跟踪 4 项业务进度；零完成渠道按协议扣罚并取消下季签约资格。'),
    ('5. APP 融合率偏低，新终端/宽带两条线最弱',
     '新入网融合率均值 {:.1f}%、新终端 {:.1f}%、宽带 {:.1f}%；APP 系数 0.8（扣减档）渠道占 98.0%。损收 '.format(
         S['avg_newnet_fuse'], S['avg_newterm_fuse'], S['avg_bb_fuse'])
     + fmt_wan(abs(tot['app_loss']) * 10000) + '。',
     '办终端、装宽带时未把 APP 开通做成默认动作，融合靠渠道自觉。',
     '办机装宽时默认开通 APP（弹窗+话术）；按周输出三类融合率红黑榜。'),
    ('6. 终端合约搭载不足',
     '合约率均值 {:.1f}%，273 家（68.8%）合约率为 0。损收 '.format(S['avg_term'])
     + fmt_wan(abs(tot['term_loss']) * 10000) + '（勐海 2.10万、勐腊 0.50万为主）。',
     '卖了终端没搭合约，合约机政策传导不到位。',
     '清点办终端未搭合约的渠道逐户给目标；合约率 0 渠道主推合约机政策。'),
    ('7. 激励极度集中，尾部大量渠道颗粒无收',
     '最终 0 元渠道 ' + str(S['zero_final']) + ' 家（' + '{:.1f}%'.format(S['zero_final_rate']) + '），'
     '加 31 家负值（按0计）近九成无激励。有激励渠道仅 ' + str(got_cnt) + ' 家。',
     '门槛与系数双重过滤下中小渠道投入产出失衡，存在下季参与意愿下降风险。',
     '对零激励渠道分层：纯业务差→给帮扶；门槛差→季初宣贯；连续两季零激励→退出评估。'),
    ('8. 数据质量隐患',
     '有责投诉当期未出数（主表按不扣罚计），存在追溯扣罚风险；31 家渠道最终金额为负（扣罚超可发额），需人工复核；'
     '金豆拍照值与原始金豆口径需再核对。',
     '投诉出数滞后、扣罚规则未封底（未限制最低为 0）。',
     '与投诉主管部门确认出数时间；对负值渠道复核扣罚规则并增设最低0元兜底。'),
]
for title, desc, cause, advice in problems:
    add_para(doc, title, 11, True, DARK_BLUE)
    add_para(doc, '问题表现：' + desc, 9.5)
    add_para(doc, '可能原因：' + cause, 9.5, False, GRAY)
    add_para(doc, '改进建议：' + advice, 9.5, False, GREEN)

# 三、改进建议
add_para(doc, '三、改进建议', 13, True, DARK_BLUE)
add_para(doc, '（一）业务短板方面', 11, True, DARK_BLUE)
for x in [
    '门槛业务季初前置：与渠道签约时同步下发四项门槛量化目标（88会员年包、AI五件套、金虎、权益升档），按档位给出差距清单。',
    '抓三条最弱线：88会员年包（92.9%空白）、新终端/宽带 APP 融合（15%上下）、终端合约搭载（68.8%为0），逐户定目标、按周通报。',
    '存量运营积分补课：对 0.81 牵引系数渠道核对积分缺口，把重新增、轻运营的经营结构拉回来。',
]:
    add_para(doc, '· ' + x, 9.5)
add_para(doc, '（二）参与与管理方面', 11, True, DARK_BLUE)
for x in [
    '渠道分层管理：按门槛全0/有量未达档/达档分三档建台账，分别给帮扶、催办、维持策略，避免一刀切。',
    '弱势网格攻坚严进严出：签约前评估承接能力，月度跟踪 4 项业务，零完成按协议扣罚并取消下季资格。',
    '过程预警替代结果考核：牵引系数、APP融合率、门槛差距改为季度中（第5/9周）双预警节点，季末只做确认。',
    '数据治理：投诉出数时间表、扣罚金额最低0元封底、负值渠道复核，纳入季度复盘清单。',
]:
    add_para(doc, '· ' + x, 9.5)

# ============ 附1：重点网格整改清单（横向节） ============
sec = doc.add_section(WD_SECTION.NEW_PAGE)
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.left_margin = sec.right_margin = Cm(1.5)

add_para(doc, '附1  重点网格整改清单（损失最大 8 个网格）', 13, True, DARK_BLUE)
top8 = grids[:8]
add_para(doc, '以下网格按损失额降序排列，合计损失 ' + fmt_wan(sum(g['loss'] for g in top8))
         + '，占全州损失 ' + '{:.0f}%'.format(sum(g['loss'] for g in top8) / S['loss_total'] * 100) + '。', 9.5)

for g in top8:
    zero_term = sum(1 for c in g['chans'] if c['term_ratio'] == 0)
    low_fuse = sum(1 for c in g['chans'] if c['newterm_fuse'] < 15 or c['bb_fuse'] < 15)
    zero_all = sum(1 for c in g['chans'] if c['gate_level'] == '门槛未完成'
                   and c['tiger'] == 0 and c['ai5'] == 0 and c['rights_up'] == 0 and c['member88'] == 0)
    grid_title = g['grid'] if g['grid'].endswith('网格') else g['grid'] + '网格'
    add_para(doc, g['county'] + ' · ' + grid_title + '（' + str(g['n']) + '家渠道）——损失 '
             + fmt_wan(g['loss']) + '（原始 ' + fmt_wan(g['raw']) + ' → 最终 ' + fmt_wan(g['final']) + '）',
             10.5, True, DARK_BLUE)
    problems_txt = []
    if g['gate_notdone'] / g['n'] >= 0.8:
        problems_txt.append('门槛未完成 ' + str(g['gate_notdone']) + '/' + str(g['n'])
                            + ' 家（其中 ' + str(zero_all) + ' 家四项业务全 0）')
    if g['zero_final'] / g['n'] >= 0.7:
        problems_txt.append('0 元渠道 ' + str(g['zero_final']) + ' 家（' + '{:.0f}%'.format(g['zero_final'] / g['n'] * 100) + '）')
    if zero_term:
        problems_txt.append('合约率为 0 的 ' + str(zero_term) + ' 家')
    if low_fuse:
        problems_txt.append('新终端/宽带融合率低于 15% 的 ' + str(low_fuse) + ' 家')
    if not problems_txt:
        problems_txt.append('各指标相对均衡，重点维持门槛档位')
    add_para(doc, '问题：' + '；'.join(problems_txt) + '。', 9.5, color=RED)
    advice = []
    if g['gate_notdone'] / g['n'] >= 0.8:
        advice.append('先逐户补四项门槛（优先 88 会员年包、AI五件套），这是最大止血点')
    if g['zero_final'] / g['n'] >= 0.7:
        advice.append('对 0 元渠道逐户归类原因，业务缺口给帮扶、门槛缺口季初宣贯')
    if zero_term:
        advice.append('主推合约机搭载')
    if low_fuse:
        advice.append('办机装宽时默认开通 APP')
    add_para(doc, '整改：' + '；'.join(advice) + '。', 9.5, color=GREEN)
    add_para(doc, '', 2)

add_para(doc, '共性提醒：以上网格均以门槛未完成作为第一症结，若季度初不动员渠道补门槛，'
         '季末激励仍将大面积归零。', 9.5, True, RED)

# ============ 附2：网格汇总表（横向节） ============
sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
sec2.orientation = WD_ORIENT.LANDSCAPE
sec2.page_width, sec2.page_height = sec2.page_height, sec2.page_width
sec2.left_margin = sec2.right_margin = Cm(1.2)

add_para(doc, '附2  网格汇总表（按县分组、县内按损失降序）', 13, True, DARK_BLUE)


def grid_shortboard(g):
    gz = g['chans']
    gate_ok = sum(1 for c in gz if c['gate_level'] != '门槛未完成')
    tiger_sum = sum(c['tiger'] for c in gz)
    m88_sum = sum(c['member88'] for c in gz)
    avg_term = sum(c['term_ratio'] for c in gz) / len(gz) * 100
    avg_nt = sum(c['newterm_fuse'] for c in gz) / len(gz)
    avg_bb = sum(c['bb_fuse'] for c in gz) / len(gz)
    tags = []
    if g['raw'] <= 0:
        tags.append('无业务量')
    else:
        if g['gate_notdone'] / g['n'] >= 0.8:
            tags.append('门槛未达')
        if gate_ok == 0:
            tags.append('零达档')
        if tiger_sum == 0:
            tags.append('金虎空白')
        if m88_sum == 0:
            tags.append('会员空白')
        if avg_term < 20:
            tags.append('合约低')
        if avg_nt < 15 or avg_bb < 15:
            tags.append('APP融合低')
    if not tags:
        tags.append('相对均衡')
    return '、'.join(tags[:3]), gate_ok, tiger_sum, m88_sum, avg_term, avg_nt, avg_bb


header = ['网格', '县', '渠道数', '原始(万)', '损失(万)', '达档', '金虎', '88会员',
          '合约率', 'APP融合(终/宽)', '主要短板']
rows = []
for g in grids:
    sb, gate_ok, tiger_sum, m88_sum, avg_term, avg_nt, avg_bb = grid_shortboard(g)
    fuse = '{:.0f}/{:.0f}'.format(avg_nt, avg_bb)
    rows.append([g['grid'], g['county'], str(g['n']), '{:.1f}'.format(g['raw'] / 10000),
                 '{:.1f}'.format(g['loss'] / 10000), str(gate_ok), '{:.0f}'.format(tiger_sum),
                 '{:.0f}'.format(m88_sum), '{:.0f}%'.format(avg_term), fuse, sb])
rows.append(['全州合计', '', str(S['n']), '{:.1f}'.format(S['raw_total'] / 10000),
             '{:.1f}'.format(S['loss_total'] / 10000),
             str(sum(1 for c in chans if c['gate_level'] != '门槛未完成')),
             '{:.0f}'.format(sum(c['tiger'] for c in chans)),
             '{:.0f}'.format(sum(c['member88'] for c in chans)),
             '{:.0f}%'.format(S['avg_term']),
             '{:.0f}/{:.0f}'.format(S['avg_newterm_fuse'], S['avg_bb_fuse']),
             ''])
t2 = make_table(doc, header, rows,
                widths=[3.0, 1.8, 1.2, 1.6, 1.6, 1.1, 1.1, 1.3, 1.5, 3.2, 4.0], font_size=8)

for ri, g in enumerate(grids, 1):
    row = t2.rows[ri].cells
    loss_wan = g['loss'] / 10000
    if loss_wan >= 5:
        shade_cell(row[4], 'FFC7CE')
    elif loss_wan >= 1:
        shade_cell(row[4], 'FFEB9C')
    else:
        shade_cell(row[4], 'C6EFCE')
    _, gate_ok, tiger_sum, m88_sum, avg_term, avg_nt, avg_bb = grid_shortboard(g)
    if gate_ok == 0 and g['raw'] > 0:
        shade_cell(row[5], 'FFC7CE')
    elif gate_ok > 0:
        shade_cell(row[5], 'C6EFCE')
    if tiger_sum == 0 and g['raw'] > 0:
        shade_cell(row[6], 'FFC7CE')
    elif tiger_sum > 0:
        shade_cell(row[6], 'C6EFCE')
    if m88_sum == 0 and g['raw'] > 0:
        shade_cell(row[7], 'FFC7CE')
    elif m88_sum > 0:
        shade_cell(row[7], 'C6EFCE')
    if avg_term < 20:
        shade_cell(row[8], 'FFC7CE')
    elif avg_term < 50:
        shade_cell(row[8], 'FFEB9C')
    else:
        shade_cell(row[8], 'C6EFCE')
    if avg_nt < 15 or avg_bb < 15:
        shade_cell(row[9], 'FFC7CE')
    elif avg_nt < 30 or avg_bb < 30:
        shade_cell(row[9], 'FFEB9C')
    else:
        shade_cell(row[9], 'C6EFCE')
for cell in t2.rows[-1].cells:
    shade_cell(cell, 'DCE6F1')

add_para(doc, '图例：红=严重，黄=一般，绿=达标（损失≥5万红/1~5万黄；达档、金虎、88会员为0红；合约率<20%红、20~50%黄；'
         'APP融合率<15%红、15~30%黄）。', 8, False, GRAY)

doc.save(OUT)
print('已生成:', OUT)
