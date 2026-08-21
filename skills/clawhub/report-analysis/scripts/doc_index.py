# -*- coding: utf-8 -*-
"""政策文档索引构建：提取三份 Word 文档全文（段落+表格），识别章节结构，输出 JSON 索引。

用法:
    python doc_index.py            # 构建索引到 policy_docs.json
环境变量:
    YJ_DOC_DIR  文档目录（必填，通过环境变量指定）
    YJ_DOC_OUT  索引输出路径（默认 analysis/policy_docs.json）
"""
import os
import json
import re
from docx import Document

DOC_DIR = os.environ.get('YJ_DOC_DIR', '')
if not DOC_DIR:
    raise SystemExit('未指定政策文档目录：请通过环境变量 YJ_DOC_DIR 传入存放三份 Word 文档的文件夹路径')
OUT = os.environ.get('YJ_DOC_OUT', os.path.join(os.path.dirname(os.path.abspath(__file__)), 'policy_docs.json'))

DOCS = [
    {
        'file': '管理：2024年中国移动云南公司销售合作伙伴日常运营执行要求.docx',
        'short': '日常运营执行要求',
        'year': '2024',
        'layer': '操作层',
        'desc': '日常运营执行要求：合作引入流程、资质要求、日常运营规范',
    },
    {
        'file': '西双版纳分公司2023年销售合作伙伴分级施策指导意见V4.0(1).docx',
        'short': '西双版纳分级施策指导意见',
        'year': '2023',
        'layer': '激励层',
        'desc': '分级施策指导意见：月度激励公式、季度星级激励、否决与扣罚项',
    },
    {
        'file': '中国移动云南公司销售合作伙伴管理办法（2023年版）.docx',
        'short': '云南公司销售合作伙伴管理办法',
        'year': '2023',
        'layer': '制度层',
        'desc': '管理办法：分类布局、准入退出、费用管理、星级评定、运营管理',
    },
]

HEAD_RE = [
    (r'^第[一二三四五六七八九十百\d]+[章章节部分]', 1),
    (r'^[一二三四五六七八九十]+、', 1),
    (r'^（[一二三四五六七八九十\d]+）', 2),
    (r'^\d+[、.]\s*', 2),
]

def is_heading(text):
    for pat, lv in HEAD_RE:
        if re.match(pat, text):
            return lv
    if len(text) <= 30 and not re.search(r'[。；，]', text):
        return 3
    return 0

def extract_doc(path):
    doc = Document(path)
    paras = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            paras.append(t)
    tables = []
    for t in doc.tables:
        rows = []
        for r in t.rows:
            cells = [c.text.strip() for c in r.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)
    return paras, tables

def build_index():
    out = []
    for d in DOCS:
        path = os.path.join(DOC_DIR, d['file'])
        if not os.path.exists(path):
            print('WARN 文件不存在，跳过:', path)
            continue
        paras, tables = extract_doc(path)
        # 章节栈：用标题行维护当前章节路径
        sections = []
        items = []  # {text, sec}
        for p in paras:
            lv = is_heading(p)
            if lv == 1:
                sections = [p]
            elif lv == 2 and len(sections) < 2:
                sections = sections[:1] + [p]
            elif lv == 3 and len(sections) < 3:
                sections = sections[:2] + [p]
            items.append({'text': p, 'sec': list(sections)})
        entry = {
            'file': d['file'],
            'short': d['short'],
            'year': d['year'],
            'layer': d['layer'],
            'desc': d['desc'],
            'n_paras': len(items),
            'n_tables': len(tables),
            'chars': sum(len(x['text']) for x in items) + sum(len(c) for t in tables for r in t for c in r),
            'sections': sections,  # 顶层章节
            'items': items,
            'tables': tables,
        }
        out.append(entry)
        print('OK', d['short'], '| 段落', len(items), '| 表格', len(tables), '| 章节', ' / '.join(sections))
    with open(OUT, 'w', encoding='utf-8') as f:
        json.dump(out, f, ensure_ascii=False, indent=1)
    print('索引已保存:', OUT)

# 术语映射：Excel/业务术语 -> 文档检索关键词（用于跨表联动检索）
TERM_MAP = {
    '门槛': ['否决', '门槛'],
    '否决': ['否决', '门槛'],
    '星级': ['星级'],
    '金虎': ['金虎'],
    'AI五件套': ['AI五件套', 'AI 五件套', '五件套'],
    '会员': ['会员', '年包'],
    '88会员': ['88会员', '88会员年包'],
    '终端合约': ['合约'],
    '合约率': ['合约'],
    '终端合约率': ['合约率', '合约'],
    '不达标': ['达标', '完成率'],
    '未达标': ['达标', '完成率'],
    '达标': ['达标', '完成率', '门槛'],
    '标准': ['标准', '要求', '规定'],
    '月度激励': ['月度激励', '月度贡献'],
    '季度激励': ['季度星级', '星级激励', '星级奖励'],
    '牵引系数': ['协同系数', '协同考核'],
    '重点业务': ['重点业务', '协同'],
    '弱势网格': ['弱势网格'],
    'APP融合': ['融合率', '融合'],
    '融合率': ['融合率', '融合'],
    '投诉': ['投诉'],
    '激励': ['激励'],
    '月度激励': ['月度激励', '月度贡献'],
    '贡献奖励': ['贡献奖励'],
    '贡献完成率': ['贡献完成率'],
    '处罚': ['违规', '处罚', '扣罚'],
    '扣罚': ['扣罚', '处罚', '违规'],
    '退出': ['退出', '清退', '解除'],
    '准入': ['准入', '引入', '资质'],
    '资质': ['资质'],
    '星级评定': ['星级评定', '评定'],
    '评级': ['星级', '评定'],
    '费用': ['费用'],
    '结算': ['结算', '费用支付'],
    '返利': ['返利', '激励'],
    '佣金': ['佣金', '激励'],
}

def save_term_map(path):
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(TERM_MAP, f, ensure_ascii=False, indent=1)
    print('术语映射已保存:', path)

if __name__ == '__main__':
    build_index()
    save_term_map(os.path.join(os.path.dirname(OUT), 'term_map.json'))
