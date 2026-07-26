#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
著作权登记辅助 - Word文档生成工具
用于生成著作权登记相关的各类Word文档
"""

import os
import json
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn


# ==================== 样式工具 ====================

def set_cell_shading(cell, color):
    """设置单元格底色"""
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_paragraph(doc, text, style=None, bold=False, size=None,
                            alignment=None, color=None, space_before=None,
                            space_after=None, font_name=None):
    """添加带格式的段落"""
    p = doc.add_paragraph()
    if style:
        p.style = style

    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
        # 设置中文字体
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)

    if alignment:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)

    return p


def add_table_row(table, cells_data, bold=False, header=False):
    """添加表格行"""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        if i < len(row.cells):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.bold = bold
            run.font.size = Pt(10)
            run.font.name = '宋体'
            if header:
                set_cell_shading(cell, 'E8F0FE')


# ==================== 文档生成函数 ====================

def generate_registration_form(data, output_path):
    """
    生成著作权登记申请表
    data: dict 包含作品信息、作者信息、著作权人信息等

    【参考】中国版权保护中心官方登记申请表格式：
    - 官网：https://www.ccopyright.com.cn
    - 在线登记平台：https://register.ccopyright.com.cn
    - 官方字段包括：作品名称、作品类别、创作完成日期、创作完成地点、
      创作性质（原创/改编/翻译/汇编等）、作品篇幅/字数、首次发表信息、
      作者姓名/类别、著作权人姓名/类别、权利取得方式、权利归属方式等
    """
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10)

    # 标题
    add_formatted_paragraph(doc, '著作权登记申请表',
                            bold=True, size=18, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_after=6)
    add_formatted_paragraph(doc, 'Copyright Registration Application Form',
                            bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            color=(128, 128, 128), space_after=12)

    # 说明
    add_formatted_paragraph(doc, '填写说明：带 * 项为必填项，请如实填写。本表仅供参考，正式表格请以中国版权保护中心官方平台为准。',
                            size=9, color=(180, 80, 80), space_after=12)

    # 一、作品信息
    add_formatted_paragraph(doc, '一、作品信息', bold=True, size=12,
                            space_before=12, space_after=6)

    sections = [
        ('作品名称 *', data.get('work_name', '')),
        ('作品类别 *', data.get('work_category', '')),
        ('创作完成日期 *', data.get('completion_date', '')),
        ('创作完成地点 *', data.get('completion_place', '')),
        ('创作性质 *', data.get('creation_nature', '原创/改编/翻译/汇编/其他')),
        ('作品篇幅/字数', data.get('work_length', '')),
        ('是否发表 *', '已发表' if data.get('is_published') else '未发表'),
        ('首次发表国家/地区', data.get('first_publication_place', '')),
        ('首次发表日期', data.get('first_publication_date', '')),
        ('首次发表方式', data.get('first_publication_method', '')),
        ('作品内容简述', data.get('work_description', '')),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '字段名称'
    header_cells[1].text = '填写内容'
    for cell in header_cells:
        set_cell_shading(cell, 'D6E4F0')
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for label, value in sections:
        add_table_row(table, [label, value])

    # 二、作者信息
    add_formatted_paragraph(doc, '二、作者信息', bold=True, size=12,
                            space_before=18, space_after=6)

    author = data.get('author', {})
    author_sections = [
        ('姓名/名称 *', author.get('name', '')),
        ('作者类别 *', author.get('author_category', '自然人/法人/其他组织')),
        ('证件类型 *', author.get('id_type', '身份证/护照/营业执照等')),
        ('证件号码 *', author.get('id_number', '')),
        ('国籍/地区', author.get('nationality', '')),
        ('住所地', author.get('address', '')),
        ('联系电话', author.get('phone', '')),
        ('电子邮箱', author.get('email', '')),
    ]

    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells2 = table2.rows[0].cells
    header_cells2[0].text = '字段名称'
    header_cells2[1].text = '填写内容'
    for cell in header_cells2:
        set_cell_shading(cell, 'D6E4F0')
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for label, value in author_sections:
        add_table_row(table2, [label, value])

    # 三、著作权人信息
    add_formatted_paragraph(doc, '三、著作权人信息', bold=True, size=12,
                            space_before=18, space_after=6)

    owner = data.get('owner', author)
    owner_sections = [
        ('姓名/名称 *', owner.get('name', '')),
        ('著作权人类别 *', owner.get('owner_category', '自然人/法人/其他组织')),
        ('证件类型 *', owner.get('id_type', '身份证/护照/营业执照等')),
        ('证件号码 *', owner.get('id_number', '')),
        ('国籍/地区', owner.get('nationality', '')),
        ('住所地', owner.get('address', '')),
        ('联系电话', owner.get('phone', '')),
        ('电子邮箱', owner.get('email', '')),
        ('权利取得方式 *', data.get('right_acquisition', '原始取得/继承/受让')),
        ('权利归属方式 *', data.get('right_ownership', '个人独自创作/合作创作/职务创作/委托创作')),
    ]

    table3 = doc.add_table(rows=1, cols=2)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells3 = table3.rows[0].cells
    header_cells3[0].text = '字段名称'
    header_cells3[1].text = '填写内容'
    for cell in header_cells3:
        set_cell_shading(cell, 'D6E4F0')
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for label, value in owner_sections:
        add_table_row(table3, [label, value])

    # 提交声明
    add_formatted_paragraph(doc, '', space_before=12)
    add_formatted_paragraph(doc, '本人/本单位声明：以上填写信息真实、准确、完整。如有不实，愿承担由此产生的一切法律责任。',
                            size=10, color=(80, 80, 80), space_before=12)

    add_formatted_paragraph(doc, '申请人签名/盖章：_______________        日期：_______年_______月_______日',
                            size=10, space_before=18)

    add_formatted_paragraph(doc, '', space_before=6)
    add_formatted_paragraph(doc, f'文档生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}  |  本文件为辅助参考材料，正式提交请以官方平台为准',
                            size=8, color=(160, 160, 160), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(output_path)
    return output_path


def generate_work_description(data, output_path):
    """
    生成作品说明书
    """
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10)

    # 标题
    add_formatted_paragraph(doc, '作品说明书',
                            bold=True, size=18, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_after=6)
    add_formatted_paragraph(doc, 'Work Description Statement',
                            bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            color=(128, 128, 128), space_after=12)

    # 作品基本信息
    add_formatted_paragraph(doc, '一、作品基本信息', bold=True, size=12,
                            space_before=12, space_after=6)

    info_items = [
        ('作品名称', data.get('work_name', '')),
        ('作品类别', data.get('work_category', '')),
        ('作者', data.get('author', {}).get('name', '')),
        ('著作权人', data.get('owner', {}).get('name', data.get('author', {}).get('name', ''))),
        ('创作完成日期', data.get('completion_date', '')),
        ('创作完成地点', data.get('completion_place', '')),
        ('发表状态', '已发表' if data.get('is_published') else '未发表'),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    header_cells[0].text = '项目'
    header_cells[1].text = '内容'
    for cell in header_cells:
        set_cell_shading(cell, 'D6E4F0')
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for label, value in info_items:
        add_table_row(table, [label, value])

    # 二、创作背景
    add_formatted_paragraph(doc, '二、创作背景', bold=True, size=12,
                            space_before=18, space_after=6)
    add_formatted_paragraph(doc, data.get('creation_background',
        '（请在此描述作品的创作动机、灵感来源、创作目的等背景信息）'),
        size=10, space_after=6)

    # 三、创作过程
    add_formatted_paragraph(doc, '三、创作过程', bold=True, size=12,
                            space_before=12, space_after=6)
    add_formatted_paragraph(doc, data.get('creation_process',
        '（请在此描述作品从构思到完成的完整创作过程，包括创作时间线、创作方法、创作工具等）'),
        size=10, space_after=6)

    # 四、作品内容说明
    add_formatted_paragraph(doc, '四、作品内容说明', bold=True, size=12,
                            space_before=12, space_after=6)
    add_formatted_paragraph(doc, data.get('work_description',
        '（请在此详细描述作品的内容、结构、主要特征等）'),
        size=10, space_after=6)

    # 五、独创性声明
    add_formatted_paragraph(doc, '五、独创性声明', bold=True, size=12,
                            space_before=12, space_after=6)
    add_formatted_paragraph(doc, data.get('originality_statement',
        '本人/本单位声明：本作品由本人/本单位独立创作完成，具有独创性，'
        '未侵犯他人著作权及其他知识产权。对作品中所引用他人作品的部分，'
        '已注明出处并取得了合法授权。'),
        size=10, space_after=6)

    # 六、创作证明材料说明
    add_formatted_paragraph(doc, '六、创作证明材料说明', bold=True, size=12,
                            space_before=12, space_after=6)
    add_formatted_paragraph(doc, data.get('creation_evidence',
        '（请在此列明可以证明创作过程和创作时间的材料，如创作手稿、创作日志、'
        '时间戳记录、电子邮件记录等）'),
        size=10, space_after=12)

    # 签名区
    add_formatted_paragraph(doc, '申请人签名/盖章：_______________        日期：_______年_______月_______日',
                            size=10, space_before=18)
    add_formatted_paragraph(doc, '', space_before=6)
    add_formatted_paragraph(doc, f'文档生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}  |  本文件为辅助参考材料',
                            size=8, color=(160, 160, 160), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(output_path)
    return output_path


def generate_rights_guarantee(data, output_path):
    """
    生成权利保证书
    """
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10)

    # 标题
    add_formatted_paragraph(doc, '权利保证书',
                            bold=True, size=18, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_after=6)
    add_formatted_paragraph(doc, 'Rights Guarantee Statement',
                            bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            color=(128, 128, 128), space_after=12)

    author_name = data.get('author', {}).get('name', '【作者姓名】')
    owner_name = data.get('owner', {}).get('name', author_name)
    work_name = data.get('work_name', '【作品名称】')
    today = datetime.now().strftime('%Y年%m月%d日')

    guarantee_text = f"""
保证人：{owner_name}
被保证作品：{work_name}
作品类别：{data.get('work_category', '【作品类别】')}

鉴于保证人拟就上述作品向著作权登记机构申请著作权登记，为明确权利归属、保障权利真实性，保证人特此作出如下保证：

一、权利来源保证
1. 保证人系上述作品的合法著作权人，对该作品享有完整的著作权。
2. {data.get('right_ownership_desc', '该作品由保证人独立创作完成（合作作品的，保证人已与其他合作作者就权利归属作出明确约定）。')}
3. 该作品的权利取得方式为{data.get('right_acquisition', '原始取得')}，不存在权利瑕疵。

二、权利真实性保证
1. 保证人保证该作品为原创作品，具有独创性，未被任何第三方主张权利。
2. 该作品中使用的他人作品（如有）已取得合法授权，并已按约定方式使用。
3. 保证人保证所提交的登记申请材料真实、准确、完整。

三、无侵权承诺
保证人承诺该作品不侵犯任何第三方的著作权、商标权、专利权、肖像权、名誉权等合法权益。如因该作品的著作权登记引发任何第三方索赔或法律纠纷，保证人愿意承担全部法律责任，并赔偿由此给著作权登记机构和相关方造成的全部损失。

四、权利维持承诺
1. 保证人承诺在著作权有效期内，不主动放弃该作品的著作权。
2. 如该作品的权利状况发生任何变化（如权利转让、许可等），保证人将及时通知著作权登记机构。

五、法律效力
本保证书自保证人签署之日起生效。保证人确认，本保证书是著作权登记申请材料的组成部分，保证人对其内容的真实性承担法律责任。

保证人签名/盖章：_______________

签署日期：{today}
"""

    add_formatted_paragraph(doc, guarantee_text.strip(), size=10, space_after=12)

    add_formatted_paragraph(doc, f'文档生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}  |  本文件为辅助参考材料',
                            size=8, color=(160, 160, 160), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(output_path)
    return output_path


def generate_template(template_type, output_path):
    """
    生成模板文档
    template_type: work_description / rights_guarantee / creation_evidence / power_of_attorney

    【参考】中国版权保护中心官方格式要求：
    - 申请表格式：https://register.ccopyright.com.cn
    - 作品说明书应包含：创作背景、创作过程、作品内容说明、独创性声明
    - 权利保证书应包含：权利来源声明、权利归属确认、无侵权承诺
    - 创作证明材料应包括：创作时间线、创作过程记录、发表记录等
    """
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10)

    now = datetime.now()

    if template_type == 'work_description':
        add_formatted_paragraph(doc, '作品说明书（模板）',
                                bold=True, size=18, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                space_after=12)

        add_formatted_paragraph(doc, '一、作品基本信息', bold=True, size=12,
                                space_before=6, space_after=6)
        add_formatted_paragraph(doc, '作品名称：【　　　　　　】')
        add_formatted_paragraph(doc, '作品类别：【　　　　　　】')
        add_formatted_paragraph(doc, '作者姓名：【　　　　　　】')
        add_formatted_paragraph(doc, '创作完成日期：【　年　月　日　】')
        add_formatted_paragraph(doc, '创作完成地点：【　　　　　　】')

        add_formatted_paragraph(doc, '二、创作背景', bold=True, size=12,
                                space_before=12, space_after=6)
        add_formatted_paragraph(doc, '（请在此描述作品的创作动机、灵感来源、创作目的等背景信息）\n【　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　】')

        add_formatted_paragraph(doc, '三、创作过程', bold=True, size=12,
                                space_before=12, space_after=6)
        add_formatted_paragraph(doc, '（请在此描述作品从构思到完成的创作历程，包括创作时间、创作地点、创作方式等）\n【　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　】')

        add_formatted_paragraph(doc, '四、作品内容说明', bold=True, size=12,
                                space_before=12, space_after=6)
        add_formatted_paragraph(doc, '（请在此详细描述作品的内容、结构、风格、主要特征等）\n【　　　　　　　　　　　　　　　　　　　　　　　　　　　　　　】')

        add_formatted_paragraph(doc, '五、独创性声明', bold=True, size=12,
                                space_before=12, space_after=6)
        add_formatted_paragraph(doc, '本人声明：本作品由本人独立创作完成，具有独创性，未侵犯他人著作权。')

        add_formatted_paragraph(doc, '申请人签名：_______________        日期：_______年_______月_______日',
                                space_before=18)

    elif template_type == 'rights_guarantee':
        add_formatted_paragraph(doc, '权利保证书（模板）',
                                bold=True, size=18, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                space_after=12)

        add_formatted_paragraph(doc, '保证人：【　　　　　　】')
        add_formatted_paragraph(doc, '被保证作品名称：【　　　　　　】')
        add_formatted_paragraph(doc, '作品类别：【　　　　　　】')
        add_formatted_paragraph(doc, '')

        guarantee_template = """鉴于保证人拟就上述作品申请著作权登记，保证人特此作出如下保证：

一、权利来源保证
保证人系上述作品的合法著作权人，对该作品享有完整的著作权。该作品权利来源合法，不存在权利瑕疵。

二、权利真实性保证
保证人保证该作品为原创作品，具有独创性。所提交的登记申请材料真实、准确、完整。

三、无侵权承诺
保证人承诺该作品不侵犯任何第三方的合法权益。如引发纠纷，保证人承担全部法律责任。

四、法律效力
本保证书自签署之日起生效，保证人对其内容的真实性承担法律责任。"""
        add_formatted_paragraph(doc, guarantee_template.strip(), size=10, space_after=12)
        add_formatted_paragraph(doc, '保证人签名/盖章：_______________')
        add_formatted_paragraph(doc, f'签署日期：______年______月______日')

    elif template_type == 'creation_evidence':
        add_formatted_paragraph(doc, '创作证明材料清单（模板）',
                                bold=True, size=18, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                space_after=12)

        add_formatted_paragraph(doc, '作品名称：【　　　　　　】')
        add_formatted_paragraph(doc, '作者姓名：【　　　　　　】')
        add_formatted_paragraph(doc, '创作完成日期：【　年　月　日　】')
        add_formatted_paragraph(doc, '创作完成地点：【　　　　　　】')
        add_formatted_paragraph(doc, '')

        add_formatted_paragraph(doc, '以下为可证明作品创作过程和创作时间的材料清单：', bold=True,
                                space_before=6, space_after=6)

        evidence_table = doc.add_table(rows=1, cols=4)
        evidence_table.style = 'Table Grid'
        evidence_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_cells = evidence_table.rows[0].cells
        for i, h in enumerate(['序号', '材料名称', '材料说明', '备注']):
            header_cells[i].text = h
            set_cell_shading(header_cells[i], 'D6E4F0')
            for p in header_cells[i].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        for i in range(5):
            add_table_row(evidence_table, [str(i + 1), '', '', ''])

        add_formatted_paragraph(doc, '', space_before=6)
        add_formatted_paragraph(doc, '说明：可提交的创作证明材料包括但不限于——')
        add_formatted_paragraph(doc, '1. 创作手稿/草稿/素描原稿（含日期标注）')
        add_formatted_paragraph(doc, '2. 创作日志/创作笔记（记录创作时间线）')
        add_formatted_paragraph(doc, '3. 创作过程中的电子邮件/聊天记录（含时间戳）')
        add_formatted_paragraph(doc, '4. 创作软件/工具的版本记录和保存记录')
        add_formatted_paragraph(doc, '5. 第三方平台上传记录（如网盘、社交平台等）')
        add_formatted_paragraph(doc, '6. 可信时间戳认证证书（如联合信任时间戳）')
        add_formatted_paragraph(doc, '7. 发表记录（如刊登的报刊、发布的网页等）')
        add_formatted_paragraph(doc, '8. 证人证言（可附证人身份证明和联系方式）')
        add_formatted_paragraph(doc, '9. 创作过程中的照片/视频记录')
        add_formatted_paragraph(doc, '10. 作品迭代版本对比说明')

        add_formatted_paragraph(doc, '', space_before=6)
        add_formatted_paragraph(doc, '创作证明材料梳理建议：', bold=True, size=10)
        add_formatted_paragraph(doc, '• 建议按时间顺序排列材料，形成完整的创作时间线')
        add_formatted_paragraph(doc, '• 每个材料应标注时间、地点、参与人员')
        add_formatted_paragraph(doc, '• 电子材料建议提供哈希值或时间戳验证')
        add_formatted_paragraph(doc, '• 材料应能证明作品从构思到完成的完整创作过程')

    elif template_type == 'power_of_attorney':
        add_formatted_paragraph(doc, '授权委托书（模板）',
                                bold=True, size=18, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                space_after=12)

        add_formatted_paragraph(doc, '委托人：【　　　　　　】')
        add_formatted_paragraph(doc, '证件类型：【　　　　　　】')
        add_formatted_paragraph(doc, '证件号码：【　　　　　　】')
        add_formatted_paragraph(doc, '受托人：【　　　　　　】')
        add_formatted_paragraph(doc, '证件类型：【　　　　　　】')
        add_formatted_paragraph(doc, '证件号码：【　　　　　　】')
        add_formatted_paragraph(doc, '')

        attorney_text = f"""委托人因办理著作权登记事宜，现委托受托人作为代理人，代理权限包括：

一、代为提交著作权登记申请及相关材料；
二、代为修改、补正申请文件；
三、代为接收与著作权登记相关的通知书、文件及证书；
四、代为缴纳著作权登记相关费用；
五、其他与著作权登记相关的事项。

委托期限：自本委托书签署之日起至著作权登记事宜办结之日止。

受托人在上述委托权限范围内所实施的行为和签署的文件，委托人均予以承认，并承担由此产生的法律后果。

委托人（签名/盖章）：_______________

受托人（签名/盖章）：_______________

签署日期：______年______月______日"""
        add_formatted_paragraph(doc, attorney_text.strip(), size=10, space_after=12)

    add_formatted_paragraph(doc, '', space_before=6)
    add_formatted_paragraph(doc, f'文档生成时间：{now.strftime("%Y-%m-%d %H:%M")}  |  本文件为辅助参考材料',
                            size=8, color=(160, 160, 160), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(output_path)
    return output_path


# ==================== 批量登记功能 ====================

def batch_generate(works: list, output_dir: str) -> list:
    """
    批量生成多个作品的著作权登记材料

    Args:
        works: 作品信息列表，每个元素为包含完整登记信息的字典
        output_dir: 输出目录

    Returns:
        生成的文档路径列表
    """
    os.makedirs(output_dir, exist_ok=True)
    generated_files = []

    for i, work_data in enumerate(works):
        work_name = work_data.get('work_name', f'作品_{i+1}')
        safe_name = work_name.replace('/', '_').replace('\\', '_')

        # 生成申请表
        form_path = os.path.join(output_dir, f'著作权登记申请表_{safe_name}.docx')
        generate_registration_form(work_data, form_path)
        generated_files.append(form_path)

        # 生成作品说明书
        desc_path = os.path.join(output_dir, f'作品说明书_{safe_name}.docx')
        generate_work_description(work_data, desc_path)
        generated_files.append(desc_path)

        # 生成权利保证书
        guarantee_path = os.path.join(output_dir, f'权利保证书_{safe_name}.docx')
        generate_rights_guarantee(work_data, guarantee_path)
        generated_files.append(guarantee_path)

    return generated_files


# ==================== 主入口 ====================

def main():
    """
    命令行入口，接受JSON参数
    使用方式：
        python main.py '{"action":"registration_form","data":{...}}'
        python main.py '{"action":"work_description","data":{...}}'
        python main.py '{"action":"rights_guarantee","data":{...}}'
        python main.py '{"action":"template","template_type":"work_description"}'
        python main.py '{"action":"template","template_type":"rights_guarantee"}'
        python main.py '{"action":"template","template_type":"creation_evidence"}'
        python main.py '{"action":"template","template_type":"power_of_attorney"}'
        python main.py '{"action":"batch_generate","works":[{...},{...}]}'
    """
    if len(sys.argv) < 2:
        print(json.dumps({"error": "请传入JSON参数"}))
        sys.exit(1)

    try:
        params = json.loads(sys.argv[1])
    except json.JSONDecodeError:
        print(json.dumps({"error": "参数格式错误，请传入有效的JSON"}))
        sys.exit(1)

    action = params.get('action', '')
    data = params.get('data', {})

    # 确保输出目录存在
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_dir = os.path.join(script_dir, 'output')
    os.makedirs(output_dir, exist_ok=True)

    work_name = data.get('work_name', '作品').replace('/', '_').replace('\\', '_')

    if action == 'registration_form':
        filename = f'著作权登记申请表_{work_name}.docx'
        output_path = os.path.join(output_dir, filename)
        generate_registration_form(data, output_path)
        print(json.dumps({"success": True, "file_path": output_path, "file_name": filename}))

    elif action == 'work_description':
        filename = f'作品说明书_{work_name}.docx'
        output_path = os.path.join(output_dir, filename)
        generate_work_description(data, output_path)
        print(json.dumps({"success": True, "file_path": output_path, "file_name": filename}))

    elif action == 'rights_guarantee':
        filename = f'权利保证书_{work_name}.docx'
        output_path = os.path.join(output_dir, filename)
        generate_rights_guarantee(data, output_path)
        print(json.dumps({"success": True, "file_path": output_path, "file_name": filename}))

    elif action == 'template':
        template_type = params.get('template_type', 'work_description')
        type_names = {
            'work_description': '作品说明书',
            'rights_guarantee': '权利保证书',
            'creation_evidence': '创作证明材料清单',
            'power_of_attorney': '授权委托书'
        }
        type_name = type_names.get(template_type, template_type)
        filename = f'{type_name}模板.docx'
        output_path = os.path.join(output_dir, filename)
        generate_template(template_type, output_path)
        print(json.dumps({"success": True, "file_path": output_path, "file_name": filename}))

    elif action == 'batch_generate':
        works = params.get('works', [])
        if not works:
            print(json.dumps({"error": "请提供works列表"}))
            sys.exit(1)
        result_files = batch_generate(works, output_dir)
        print(json.dumps({
            "success": True,
            "file_count": len(result_files),
            "files": result_files
        }))

    else:
        print(json.dumps({"error": f"未知的操作类型: {action}"}))
        sys.exit(1)


if __name__ == '__main__':
    main()