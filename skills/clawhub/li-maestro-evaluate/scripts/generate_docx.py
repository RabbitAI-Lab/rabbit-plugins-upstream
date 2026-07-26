#!/usr/bin/env python3
"""Generate standardized .docx threat model output from MAESTRO evaluation.

Usage:
    python generate_docx.py <project-dir> [--output OUTPUT]

Reads state.json and phase files from <project-dir> to produce a styled
Word document with title page, executive summary, threat register,
mitigation plan, AI risk classification (三级九子类), and disclaimer.

The standard structure follows Chinese risk report conventions:
封面 → 摘要 → 风险分类 → 威胁登记表 → 缓解措施 → 业务背景 → 免责声明
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


# ── constants ────────────────────────────────────────────────────────

AI_RISK_SUBCATEGORIES = {
    "E1": "数据安全风险", "E2": "算法安全风险", "E3": "模型安全风险",
    "A1": "网络系统安全风险", "A2": "供应链安全风险", "A3": "隐私保护风险",
    "D1": "滥用风险", "D2": "信任风险", "D3": "合规风险",
}

THEME_BLUE = RGBColor(47, 84, 150)
THEME_GRAY = RGBColor(128, 128, 128)


# ── helpers ──────────────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _make_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_shading(cell, "2F5496")
    ncols = len(headers)
    for row_data in rows:
        row = table.add_row()
        for i in range(min(len(row_data), ncols)):
            cell = row.cells[i]
            cell.text = str(row_data[i]) if row_data[i] else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


# ── readers ──────────────────────────────────────────────────────────

def _read_json(path):
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return None


def _read_md(path):
    try:
        with open(path, encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        return None


def _parse_horizontal_table(md_text, header_keywords, min_cols=5):
    """Parse horizontal markdown tables (| a | b | c |)."""
    rows = []
    if not md_text:
        return rows
    in_table = False
    for line in md_text.splitlines():
        stripped = line.strip()
        if not stripped.startswith("|"):
            if in_table:
                break
            continue
        if any(kw in stripped for kw in header_keywords):
            in_table = True
            continue
        if in_table and stripped.startswith("|---"):
            continue
        if in_table:
            cells = [c.strip() for c in stripped.strip("| ").split("|")]
            if len(cells) >= min_cols:
                rows.append(cells)
    return rows


def _parse_threat_cards(md_text):
    """Parse vertical threat card format to extract threat register rows.
    
    Accepts both bold (**Field**) and non-bold (Field) markers.
    Recognises 'Name' and 'Threat Name' as the same key.
    """
    threats = []
    if not md_text:
        return threats
    current = {}
    for line in md_text.splitlines():
        stripped = line.strip()
        m = re.match(r"\|\s*(?:\*\*)?(ID|Local ID|Threat Name|Name|Layer|Severity|Risk Level|ASI ID|ASI Mapping|AI Risk Code)(?:\*\*)?\s*\|\s*(.+)", stripped)
        if m:
            key = m.group(1).lower().replace("local id", "id").replace("threat name", "name").replace("risk level", "risk_level").replace("asi id", "asi_id").replace("asi mapping", "asi_id").replace("ai risk code", "ai_risk_code")
            val = m.group(2).strip().strip("`").strip()
            if key == "id":
                if current.get("id"):
                    threats.append(current)
                    current = {}
                current["id"] = val
            elif key in ("name", "layer", "severity", "risk_level", "asi_id", "ai_risk_code"):
                current[key] = val
    if current.get("id"):
        threats.append(current)
    return threats


def _extract_threats(md_text):
    """Try horizontal table first, then vertical card format."""
    threats = _parse_horizontal_table(md_text, ["Local ID", "Threat Name", "Name"], 5)
    if threats:
        return threats
    cards = _parse_threat_cards(md_text)
    if cards:
        return cards
    return []


# ── builders ─────────────────────────────────────────────────────────

def build_docx(project_dir, output_path=None):
    project_dir = Path(project_dir)
    state = _read_json(project_dir / "state.json")
    analysis_mode = (state or {}).get("analysis_mode", "full_assessment")

    if analysis_mode == "mvtm_checklist":
        # MVTM mode: single checklist file contains all data
        mvtm_checklist = _read_md(project_dir / "01-mvtm-checklist.md") or ""
        phase1 = mvtm_checklist
        phase6 = mvtm_checklist
        phase7 = mvtm_checklist
        summary = mvtm_checklist
    else:
        # Full Assessment mode: individual phase files
        phase1 = _read_md(project_dir / "01-business-context.md")
        phase6 = _read_md(project_dir / "06-threat-register.md")
        phase7 = _read_md(project_dir / "07-mitigations.md")
        summary = _read_md(project_dir / "10-output-summary.md")

    risk_cls = _read_md(project_dir / "11-ai-risk-classification.md")

    doc = Document()

    # ── styles ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── 封面 ──
    for _ in range(6):
        doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CSA MAESTRO 风险评估报告")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = THEME_BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    proj_name = (state or {}).get("project", "未知项目")
    run = subtitle.add_run(f"项目名称：{proj_name}")
    run.font.size = Pt(16)
    run.font.color.rgb = THEME_GRAY

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    created = (state or {}).get("created", "N/A")
    depth = (state or {}).get("analysis_depth", "N/A")
    meta.add_run(f"评估日期：{created}  |  分析深度：{depth}").font.size = Pt(11)

    doc.add_paragraph("")
    refs = doc.add_paragraph()
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    refs.add_run("框架依据：CSA MAESTRO (OWASP GenAI Security Project)").font.size = Pt(10)
    refs.add_run("\nAI风险分类：按《人工智能安全治理框架》2.0版 三级九子类").font.size = Pt(10)

    doc.add_page_break()

    # ── 1. 摘要 ──
    doc.add_heading("一、摘要", level=1)
    if summary:
        m = re.search(r"##\s+1\.\s*Executive Summary(.+?)(?=##|\Z)", summary, re.DOTALL)
        if m:
            text = m.group(1).strip()[:3000]
            if len(m.group(1).strip()) > 3000:
                text += "\n\n【注：原文超长，此处仅展示前3000字符】"
            doc.add_paragraph(text)
    if state:
        _make_table(doc,
                    ["指标", "数值"],
                    [
                        ["威胁总数", str(state.get("threat_count", 0))],
                        ["缓解措施数", str(state.get("mitigation_count", 0))],
                        ["分析深度", depth],
                        ["系统类型", state.get("system_type", "unknown")],
                    ])

    # ── 2. AI风险分类 ──
    doc.add_heading("二、AI风险分类（三级九子类）", level=1)
    doc.add_paragraph(
        "依据《人工智能安全治理框架》2.0版，AI风险分为三大类九子类："
    )
    doc.add_paragraph(
        "技术内生安全风险：E1(数据安全风险) / E2(算法安全风险) / E3(模型安全风险)"
    )
    doc.add_paragraph(
        "技术应用安全风险：A1(网络系统安全风险) / A2(供应链安全风险) / A3(隐私保护风险)"
    )
    doc.add_paragraph(
        "应用衍生安全风险：D1(滥用风险) / D2(信任风险) / D3(合规风险)"
    )
    if risk_cls:
        threats = _extract_threats(risk_cls)
        if threats:
            _make_table(doc,
                        ["ID", "威胁名称", "层级", "严重性", "ASI映射"],
                        [[t.get("id",""), t.get("name",""), t.get("layer",""),
                          t.get("severity",""), t.get("asi_id","")]
                         for t in threats[:50]])

    # ── 3. 威胁登记表 ──
    doc.add_heading("三、威胁登记表", level=1)
    if phase6:
        threats = _extract_threats(phase6)
        if threats:
            _make_table(doc,
                        ["ID", "威胁名称", "层级", "严重性", "ASI映射", "风险等级"],
                        [[t.get("id",""), t.get("name",""), t.get("layer",""),
                          t.get("severity",""), t.get("asi_id",""), t.get("risk_level","")]
                         for t in threats[:100]])
    else:
        doc.add_paragraph("（威胁登记表数据不可用）")

    # ── 4. 缓解措施汇总 ──
    doc.add_heading("四、缓解措施汇总", level=1)
    if phase7:
        mit_rows = _parse_horizontal_table(phase7, ["Mitigation ID", "Catalog ID"], 5)
        if mit_rows:
            _make_table(doc,
                        ["措施ID", "目录ID", "类型", "成本", "效果"],
                        [r[:5] for r in mit_rows])
        else:
            doc.add_paragraph("（未找到表格形式的缓解措施数据）")
    else:
        doc.add_paragraph("（缓解措施数据不可用）")

    # ── 5. 业务背景 ──
    if phase1:
        doc.add_heading("五、业务背景", level=1)
        doc.add_paragraph(phase1[:5000] if phase1 else "")

    # ── 免责声明 ──
    doc.add_page_break()
    doc.add_heading("免责声明", level=1)
    p = doc.add_paragraph(
        "本威胁模型系借助AI辅助，基于OWASP MAESTRO Playbook框架生成，"
        "AI风险分类映射至《人工智能安全治理框架》2.0版。"
        "在用于生产环境风险决策之前，必须经合格安全专业人员审核。"
    )
    p.runs[0].font.italic = True
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = THEME_GRAY

    # ── save ──
    if output_path is None:
        output_path = project_dir / "11-ai-risk-classification.docx"
    doc.save(str(output_path))
    print(f"OK: {output_path}")
    return output_path


# ── cli ──────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Generate .docx threat model from MAESTRO evaluation output"
    )
    parser.add_argument("project_dir", help="Path to <project> directory containing state.json and phase files")
    parser.add_argument("--output", "-o", help="Output .docx path (default: <project_dir>/11-ai-risk-classification.docx)")
    args = parser.parse_args()

    if not Path(args.project_dir).is_dir():
        print(f"ERROR: directory not found: {args.project_dir}")
        sys.exit(1)

    state = _read_json(Path(args.project_dir) / "state.json")
    analysis_mode = (state or {}).get("analysis_mode", "full_assessment")
    if analysis_mode == "mvtm_checklist":
        prereqs = ["state.json", "01-mvtm-checklist.md"]
    else:
        prereqs = ["state.json", "06-threat-register.md", "07-mitigations.md"]
    missing = [f for f in prereqs if not (Path(args.project_dir) / f).is_file()]
    if missing:
        print(f"WARNING: missing prerequisite files: {', '.join(missing)} (output may be incomplete)")

    out = build_docx(args.project_dir, args.output)
    print(f"Done: {out}")


if __name__ == "__main__":
    main()
