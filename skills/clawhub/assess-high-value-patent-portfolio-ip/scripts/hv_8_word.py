#!/usr/bin/env python3
"""Stage 8: render an optional scientific US-Letter DOCX from final_records.json."""

from __future__ import annotations

import argparse
import io
import pathlib
from collections import Counter
from typing import Any
from urllib.parse import urlparse

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from hv_common import jload, require_checkpoint

try:
    import requests
except ImportError:
    requests = None  # type: ignore[assignment]


NAVY = RGBColor(23, 50, 77)
TEAL = RGBColor(36, 107, 132)
INK = RGBColor(39, 55, 70)
MUTED = RGBColor(93, 108, 120)
WHITE = RGBColor(255, 255, 255)
MAX_IMAGE_BYTES = 8 * 1024 * 1024


def text(value: Any, fallback: str = "Not available") -> str:
    if value is None or value == "":
        return fallback
    return " ".join(str(value).split())


def shade(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    node = properties.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        properties.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell: Any, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    properties = cell._tc.get_or_add_tcPr()
    container = properties.first_child_found_in("w:tcMar")
    if container is None:
        container = OxmlElement("w:tcMar")
        properties.append(container)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = container.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            container.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def keep_row(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    properties.append(node)


def add_page_number(paragraph: Any) -> None:
    paragraph.add_run("PatSnap-assisted research screening  ·  ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_hyperlink(paragraph: Any, label: str, url: str) -> bool:
    parsed = urlparse(str(url or ""))
    if parsed.scheme.lower() not in {"http", "https"} or not parsed.netloc:
        paragraph.add_run(label)
        return False
    relationship = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "185D78")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = label
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return True


def configure(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4)
    for name, size, color in (("Title", 25, NAVY), ("Subtitle", 10, MUTED), ("Heading 1", 15, NAVY), ("Heading 2", 11, TEAL), ("Heading 3", 9, MUTED)):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.text = "PATENT PORTFOLIO SCREENING  /  EVIDENCE TRACE"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(7)
        run.font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(7)
        run.font.color.rgb = MUTED


def table(document: Document, headers: list[str], rows: list[list[Any]], widths: list[float] | None = None) -> Any:
    output = document.add_table(rows=1, cols=len(headers))
    output.style = "Table Grid"
    output.alignment = WD_TABLE_ALIGNMENT.CENTER
    output.autofit = False
    header = output.rows[0]
    repeat_header(header)
    for index, label in enumerate(headers):
        cell = header.cells[index]
        cell.text = label
        shade(cell, "17324D")
        margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.name = "Arial"
            run.font.size = Pt(7)
            run.font.bold = True
            run.font.color.rgb = WHITE
    for row_index, values in enumerate(rows):
        cells = output.add_row().cells
        keep_row(output.rows[-1])
        for index, value in enumerate(values):
            cells[index].text = text(value, "—")
            margins(cells[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_index % 2:
                shade(cells[index], "EEF3F5")
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(6.8)
        if widths:
            for index, width in enumerate(widths):
                cells[index].width = Inches(width)
    return output


def bullet_list(document: Document, values: list[Any], fallback: str) -> None:
    if not values:
        document.add_paragraph(fallback)
        return
    for value in values:
        document.add_paragraph(text(value), style="List Bullet")


def fetch_image(url: Any) -> bytes | None:
    if requests is None:
        return None
    parsed = urlparse(str(url or ""))
    if parsed.scheme.lower() not in {"http", "https"} or not parsed.netloc:
        return None
    try:
        response = requests.get(str(url), timeout=(8, 20), allow_redirects=False, stream=True, headers={"User-Agent": "patsnap-high-value-screening/2.0"})
        if response.status_code != 200 or "image/" not in response.headers.get("Content-Type", "").lower():
            return None
        declared = response.headers.get("Content-Length")
        if declared and int(declared) > MAX_IMAGE_BYTES:
            return None
        content = bytearray()
        for block in response.iter_content(64 * 1024):
            content.extend(block)
            if len(content) > MAX_IMAGE_BYTES:
                return None
        return bytes(content)
    except (requests.RequestException, ValueError):
        return None


def event_summary(record: dict[str, Any]) -> str:
    categories = record.get("legal_event_categories") or []
    if not categories:
        evidence = record.get("legal_event_evidence") or {}
        complete = evidence and all((detail or {}).get("state") in {"available", "empty"} for detail in evidence.values())
        return "No event records returned in all checked categories" if complete else "Incomplete event evidence"
    evidence = record.get("legal_event_evidence") or {}
    return "; ".join(f"{category} ({int((evidence.get(category) or {}).get('count') or 0)})" for category in categories)


def screening_rows(selected: list[dict[str, Any]]) -> list[list[Any]]:
    output = []
    for record in selected:
        components = record.get("score_components") or {}
        component_text = f"C {components.get('forward_citations')}/30; F {components.get('family_size')}/30; I {components.get('core_inventor')}/20; E {components.get('legal_event_activity')}/20"
        inventor = "Yes — " + ", ".join(record.get("matched_inventors") or []) if record.get("core_inventor") else "No"
        output.append([
            record.get("rank"),
            f"{record.get('score')}\n{component_text}",
            record.get("pn") or record.get("patent_id"),
            record.get("title"),
            record.get("current_assignee"),
            f"{text(record.get('legal_status'))} [{record.get('legal_status_state')}]",
            f"{text(record.get('cited_by_simple_family'))} [P{record.get('citation_percentile')}; {record.get('citation_state')}]",
            f"{text(record.get('simple_family_count'))} [P{record.get('family_percentile')}; {record.get('simple_family_state')}]",
            inventor,
            event_summary(record),
            "; ".join(record.get("gaps") or []) or "None recorded",
        ])
    return output


def render(data: dict[str, Any], output: pathlib.Path, *, include_images: bool) -> None:
    meta = data.get("meta") if isinstance(data.get("meta"), dict) else {}
    selected = [item for item in data.get("selected", []) if isinstance(item, dict)]
    errors = data.get("errors") or []
    document = Document()
    configure(document)
    document.core_properties.title = "High-Value Patent Portfolio Screening"
    document.core_properties.subject = "Evidence-traceable candidate-universe ranking"
    document.core_properties.author = "PatSnap-assisted research workflow"
    document.core_properties.keywords = "patent portfolio, screening, citations, families, legal events"
    document.add_paragraph("HIGH-VALUE PATENT PORTFOLIO SCREENING", style="Title")
    document.add_paragraph("Evidence-traceable ranking within one documented PatSnap candidate universe", style="Subtitle")
    notice = document.add_paragraph()
    notice.add_run("SCREENING BOUNDARY  ").bold = True
    notice.add_run("Scores are not monetary valuations, legal opinions, or conclusions on validity or enforceability.")
    document.add_paragraph("1. Screening summary", style="Heading 1")
    table(document, ["P002 reported", "Retrieved", "Deduplicated", "Selected", "Selection ratio"], [[meta.get("p002_reported_total"), meta.get("retrieved_count"), meta.get("deduplicated_count"), meta.get("selected_count"), f"{meta.get('ratio')}%"]], [1.5, 1.4, 1.5, 1.4, 1.6])
    document.add_paragraph("2. Query and methodology", style="Heading 1")
    document.add_paragraph(text(meta.get("query_text"), "Query was not retained in this checkpoint."))
    table(document, ["Indicator", "Weight", "Interpretation"], [
        ["Simple-family forward-citation position", "30", "Candidate-set relative; age, field, authority and coverage dependent"],
        ["Simple-family size position", "30", "Family breadth; not market coverage or enforceability"],
        ["Core-inventor membership", "20", "Exact-name concentration within the scoped candidate set"],
        ["Verified legal-event activity", "20", "Activity presence; not positive value"],
    ], [3.0, 0.7, 5.0])
    document.add_paragraph("3. Core-inventor calculation", style="Heading 1")
    inventor_rows = [[index, item.get("name"), item.get("candidate_patent_count")] for index, item in enumerate(meta.get("top5_inventors") or [], start=1) if isinstance(item, dict)]
    table(document, ["Rank", "Exact-returned inventor name", "Candidate patents"], inventor_rows, [0.7, 4.8, 1.5])
    document.add_paragraph("Commas remain inside Western names. Transliteration variants and homonyms are not merged automatically.")
    document.add_paragraph(f"4. Selected patent portfolio ({len(selected)} records)", style="Heading 1")
    document.add_paragraph("Publication identifiers remain plain text unless a verified stable global record URL is supplied. The trace JSON retains all evidence fields.")
    selected_table = table(document, ["Rank", "Score", "Publication", "Title", "Assignee", "Status", "Citations", "Family", "Core inventor", "Event activity", "Data gaps"], screening_rows(selected), [0.4, 0.7, 0.9, 1.45, 1.1, 0.9, 0.7, 0.65, 0.85, 0.9, 1.15])
    for row_index, record in enumerate(selected, start=1):
        publication_cell = selected_table.rows[row_index].cells[2]
        publication_cell.text = ""
        add_hyperlink(publication_cell.paragraphs[0], text(record.get("pn") or record.get("patent_id")), text(record.get("record_url"), ""))
    document.add_paragraph("5. Technical summaries and selection rationales", style="Heading 1")
    for record in selected:
        document.add_paragraph(f"{record.get('rank')}. {text(record.get('pn') or record.get('patent_id'))} — {text(record.get('title'))}", style="Heading 2")
        document.add_paragraph(text(record.get("rationale")))
        table(document, ["PatSnap title", "Technical problem", "Technical approach", "Benefit / effect"], [[record.get("patsnap_title"), record.get("tech_problem"), record.get("tech_approach"), record.get("benefit")]], [2.0, 2.45, 2.7, 2.45])
        if include_images:
            image = fetch_image(record.get("drawing"))
            paragraph = document.add_paragraph()
            if image:
                paragraph.add_run().add_picture(io.BytesIO(image), width=Inches(2.1))
                paragraph.add_run(f"  Abstract drawing for {text(record.get('pn'))}")
            else:
                paragraph.add_run(f"Abstract drawing: {text(record.get('drawing_state'))}; no image embedded.")
    document.add_paragraph("6. Evidence quality and limitations", style="Heading 1")
    gaps = Counter(str(gap) for record in selected for gap in record.get("gaps") or [])
    bullet_list(document, [f"{label} ({count} selected records)" for label, count in gaps.most_common()], "No selected-record gap was recorded; reviewer confirmation remains required.")
    document.add_paragraph("Pipeline errors", style="Heading 2")
    bullet_list(document, [text(item.get("message") or item) if isinstance(item, dict) else text(item) for item in errors], "No pipeline error was recorded.")
    document.add_paragraph("Limitations", style="Heading 2")
    bullet_list(document, list(meta.get("limitations") or []), "No limitation text supplied.")
    document.add_paragraph("7. Provenance and reviewer sign-off", style="Heading 1")
    table(document, ["Run ID", "Source mode", "Query SHA-256", "Schema", "Generated"], [[meta.get("run_id"), meta.get("source_mode"), meta.get("query_sha256"), data.get("schema_version"), data.get("generated_at")]], [1.5, 1.1, 3.3, 0.8, 1.7])
    table(document, ["Review role", "Name", "Date", "Decision / conditions"], [["Patent analyst", "", "", ""], ["IP counsel / portfolio owner", "", "", ""]], [1.8, 1.8, 1.3, 4.2])
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"Wrote optional DOCX report to {output}.")


def main() -> int:
    parser = argparse.ArgumentParser(description="Render the optional high-value patent screening DOCX report.")
    parser.add_argument("--input", default="final_records.json")
    parser.add_argument("--output", default="high_value_patent_portfolio_screening.docx")
    parser.add_argument("--images", action="store_true", help="Download safe P021 HTTP(S) images with size/type limits")
    parser.add_argument("--noimg", action="store_true", help="Source-compatible alias; images are already off by default")
    args = parser.parse_args()
    data = require_checkpoint(jload(args.input), keys=("meta", "selected"), filename=args.input)
    render(data, pathlib.Path(args.output), include_images=args.images and not args.noimg)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
