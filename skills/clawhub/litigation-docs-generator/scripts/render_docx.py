#!/usr/bin/env python3
"""
诉讼文书模板渲染引擎 v2
读取 YAML 格式模板 → 解析格式规则 → 替换变量 → 生成 .docx
"""

import re, os, yaml, sys
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ========================= 格式解析 =========================

def parse_template(md_path):
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    # 分离 frontmatter 和 body
    m = re.match(r'^---\s*\n(.*?)\n---\s*\n(.*)', content, re.DOTALL)
    if not m:
        raise ValueError(f"模板缺少 YAML frontmatter: {md_path}")
    
    config = yaml.safe_load(m.group(1))
    body = m.group(2)
    
    return config, body


def parse_units(config):
    """将 pt、cm 等字符串转为 Pt/Cm 对象"""
    for section in ["styles"]:
        if section not in config:
            continue
        for key, style in config[section].items():
            if isinstance(style, dict):
                for prop in ["size", "space_before", "space_after", "gap_before"]:
                    if prop in style and isinstance(style[prop], str):
                        val = style[prop]
                        if "cm" in val:
                            style[prop] = Cm(float(val.replace("cm", "")))
                        elif "pt" in val:
                            style[prop] = Pt(float(val.replace("pt", "")))
                if "indent" in style and isinstance(style["indent"], str):
                    style["indent"] = Pt(float(style["indent"].replace("pt", "")))
    return config


# ========================= 渲染引擎 =========================

def render_docx(config, body, variables, output_path):
    """根据模板渲染生成 Word 文档"""
    page = config.get("page", {})
    styles = config.get("styles", {})
    
    doc = Document()
    section = doc.sections[0]
    _apply_page(section, page)
    
    # 获取样式定义
    style_h1 = styles.get("#", {})
    style_h2 = styles.get("##", {})
    style_body = styles.get("body", {})
    style_sign = styles.get("sign", {})
    style_date = styles.get("date", {})
    style_note = styles.get("note", {})
    
    # "此致" 样式 = body + gap_before
    style_cizhi = style_body.copy()
    if "cizhi" in styles:
        style_cizhi.update({k: v for k, v in styles["cizhi"].items() if k != "from"})
    
    # 内联样式标签
    INLINE_TAGS = {
        "meta": "meta",
        "qa_gap": "qa_gap",
        "sign": "sign",
        "date": "date",
        "note": "note",
        "court": "court",
    }
    
    lines = body.strip().split("\n")
    
    for line in lines:
        stripped = line.strip()
        
        # 空行 → 跳过
        if not stripped:
            continue
        # 分隔线 --- → 跳过
        if stripped == "---":
            continue
        
        # 检测内联样式标签 [tag]
        inline_style = None
        tag_content = stripped
        for tag in INLINE_TAGS:
            if stripped.startswith(f"[{tag}]"):
                inline_style = tag
                tag_content = stripped[len(f"[{tag}]"):].strip()
                break
        
        # 变量替换
        text = _sub_vars(tag_content, variables)
        
        # 按内联标签分发
        if inline_style == "meta":
            _add_para(doc, text, styles.get("meta", style_body), indent=False)
        elif inline_style == "qa_gap":
            _add_para(doc, "", styles.get("qa_gap", {}), indent=False)
        elif inline_style == "sign":
            _add_para(doc, text, style_sign)
        elif inline_style == "date":
            year = variables.get("year", datetime.now().year)
            tmpl = style_date.get("template", "{year}年  月  日")
            _add_para(doc, tmpl.format(year=year), style_date)
        elif inline_style == "note":
            _add_para(doc, text, style_note)
        elif inline_style == "court":
            _add_para(doc, text, styles.get("court", style_body), indent=False)
        # 标题 #
        elif stripped.startswith("# ") and not stripped.startswith("## "):
            _add_para(doc, text[2:], style_h1)
        # 段落标题 ##
        elif stripped.startswith("## "):
            _add_para(doc, text[3:], style_h2, indent=False)
        # "此致"
        elif stripped == "此致":
            _add_para(doc, text, style_cizhi)
        # 普通 body
        else:
            _add_para(doc, text, style_body)
    
    doc.save(output_path)
    return output_path


def _apply_page(section, page_config):
    if "width" in page_config:
        section.page_width = _parse_size(page_config["width"])
    if "height" in page_config:
        section.page_height = _parse_size(page_config["height"])
    if "margin" in page_config:
        m = page_config["margin"]
        if isinstance(m, list) and len(m) == 4:
            section.top_margin, section.bottom_margin = _parse_size(m[0]), _parse_size(m[1])
            section.left_margin, section.right_margin = _parse_size(m[2]), _parse_size(m[3])


def _parse_size(val):
    if isinstance(val, (Cm, Pt)):
        return val
    s = str(val)
    if "cm" in s:
        return Cm(float(s.replace("cm", "")))
    if "pt" in s:
        return Pt(float(s.replace("pt", "")))
    return Cm(float(s))


def _sub_vars(text, variables):
    """替换 [变量] 为实际值"""
    def replacer(m):
        key = m.group(1)
        val = variables.get(key, m.group(0))
        if isinstance(val, (int, float)):
            return f"{val:,}"
        return str(val)
    return re.sub(r'\[([^\[\]]+)\]', replacer, text)


def _add_para(doc, text, style_dict, indent=True):
    p = doc.add_paragraph()
    
    # space
    p.paragraph_format.space_before = style_dict.get("space_before", Pt(0))
    p.paragraph_format.space_after = style_dict.get("space_after", Pt(0))
    
    # gap_before (used for "此致")
    if "gap_before" in style_dict:
        p.paragraph_format.space_before = style_dict["gap_before"]
    
    # line spacing
    ls = style_dict.get("line_spacing", 1.5)
    if ls == 1.0:
        p.paragraph_format.line_spacing = 1.0
    else:
        p.paragraph_format.line_spacing = ls
    
    # alignment
    align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                 "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
    if "align" in style_dict:
        p.alignment = align_map.get(style_dict["align"], WD_ALIGN_PARAGRAPH.LEFT)
    
    # indent: false=无缩进, true=使用默认 28pt, 具体值=使用该值
    if "indent" in style_dict:
        val = style_dict["indent"]
        if val is False:
            p.paragraph_format.first_line_indent = 0
        elif val is True:
            p.paragraph_format.first_line_indent = Pt(28)   # 默认缩进
        elif isinstance(val, (int, float)) and val > 0:
            p.paragraph_format.first_line_indent = Pt(val) if not isinstance(val, (Cm, Pt)) else val
    
    # font
    run = p.add_run(text)
    font_name = style_dict.get("font", "仿宋")
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if "size" in style_dict:
        run.font.size = style_dict["size"]
    if style_dict.get("bold"):
        run.bold = True
    
    return p


def _add_empty(doc, size):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run("")
    run.font.size = size


# ========================= CLI =========================

if __name__ == "__main__":
    import argparse, json
    p = argparse.ArgumentParser()
    p.add_argument("template", help="模板 .md 文件路径")
    p.add_argument("output", help="输出 .docx 文件路径")
    p.add_argument("--vars", default="{}", help="变量 JSON")
    args = p.parse_args()
    
    config, body = parse_template(args.template)
    config = parse_units(config)
    variables = json.loads(args.vars)
    
    render_docx(config, body, variables, args.output)
    print(f"✅ {args.output}")
