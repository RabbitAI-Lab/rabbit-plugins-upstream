#!/usr/bin/env python3
"""
export_resume.py — 将优化后的简历导出为 .docx 和 .pdf

用法:
  python export_resume.py <input.md> [--output-dir <dir>] [--name <filename>] [--format <docx|pdf|both>]

输入: 符合中间格式约定的 Markdown 简历（中文或英文标题均可）
输出: <name>.docx 和/或 <name>.pdf

字体策略:
  - .docx 指定宋体(正文)/黑体(标题)，Windows 原生渲染，macOS 自动替换为 STSong/STHeiti
  - .pdf 主路线用 LibreOffice 转换（字体自动嵌入），备用路线用 reportlab 注册系统 TTF 字体并嵌入
  - 确保在 Windows / macOS / Linux / iOS / Android 上中文均能正确显示
"""

import argparse
import os
import re
import subprocess
import sys
from pathlib import Path

# ── 字体配置 ──────────────────────────────────────────────
BODY_FONT_CN = "宋体"          # SimSun — ATS 标准，Windows 原生
BODY_FONT_EN = "Calibri"       # 英文正文
HEADER_FONT_CN = "黑体"        # SimHei — 标题用
HEADER_FONT_EN = "Arial"       # 英文标题
NAME_FONT_CN = "黑体"
NAME_FONT_EN = "Arial"

# 字号 (pt)
NAME_SIZE = 18
HEADER_SIZE = 13
SUBHEADER_SIZE = 11
BODY_SIZE = 10.5
BULLET_SIZE = 10.5

# 页面边距 (cm)
MARGIN_TOP = 2.0
MARGIN_BOTTOM = 2.0
MARGIN_LEFT = 2.5
MARGIN_RIGHT = 2.5


# ── Markdown 解析 ─────────────────────────────────────────
def parse_markdown(md_text):
    """将简历 Markdown 解析为结构化段落列表。

    返回: [{"type": "name"|"header"|"subheader"|"bullet"|"text", "text": ...}]
    """
    lines = md_text.strip().split("\n")
    blocks = []
    for line in lines:
        line = line.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            blocks.append({"type": "name", "text": line[2:].strip()})
        elif line.startswith("## "):
            blocks.append({"type": "header", "text": line[3:].strip()})
        elif line.startswith("### "):
            blocks.append({"type": "subheader", "text": line[4:].strip()})
        elif line.startswith("- "):
            blocks.append({"type": "bullet", "text": line[2:].strip()})
        else:
            blocks.append({"type": "text", "text": line.strip()})
    return blocks


# ── .docx 生成 ───────────────────────────────────────────
def generate_docx(blocks, output_path):
    """生成 .docx 文件。"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def set_run_font(run, cn_font, en_font, size, bold=False):
        """设置 run 的字体（同时处理中文 eastAsia 和西文字体）。"""
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = en_font
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:eastAsia"), cn_font)
        rfonts.set(qn("w:ascii"), en_font)
        rfonts.set(qn("w:hAnsi"), en_font)

    def add_styled_paragraph(doc, text, cn_font, en_font, size, bold=False,
                             space_before=0, space_after=4, alignment=None):
        """添加一个带样式的段落，支持 **加粗** 标记。"""
        p = doc.add_paragraph()
        if alignment is not None:
            p.alignment = alignment
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15

        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                set_run_font(run, cn_font, en_font, size, bold=True)
            else:
                run = p.add_run(part)
                set_run_font(run, cn_font, en_font, size, bold=bold)
        return p

    doc = Document()

    # 页面设置 — A4 尺寸（210mm × 297mm）
    from docx.shared import Mm
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Cm(MARGIN_TOP)
        section.bottom_margin = Cm(MARGIN_BOTTOM)
        section.left_margin = Cm(MARGIN_LEFT)
        section.right_margin = Cm(MARGIN_RIGHT)

    # 默认 Normal 样式
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT_EN
    style.font.size = Pt(BODY_SIZE)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_CN)

    for block in blocks:
        btype = block["type"]
        text = block["text"]

        if btype == "name":
            add_styled_paragraph(
                doc, text, NAME_FONT_CN, NAME_FONT_EN, NAME_SIZE,
                bold=True, space_before=0, space_after=6,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )

        elif btype == "header":
            p = add_styled_paragraph(
                doc, text, HEADER_FONT_CN, HEADER_FONT_EN, HEADER_SIZE,
                bold=True, space_before=10, space_after=4,
            )
            # 底部边框线
            ppr = p._element.get_or_add_pPr()
            pborder = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "000000")
            pborder.append(bottom)
            ppr.append(pborder)

        elif btype == "subheader":
            add_styled_paragraph(
                doc, text, HEADER_FONT_CN, HEADER_FONT_EN, SUBHEADER_SIZE,
                bold=True, space_before=6, space_after=2,
            )

        elif btype == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            parts = re.split(r"(\*\*.*?\*\*)", text)
            for part in parts:
                if not part:
                    continue
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    set_run_font(run, BODY_FONT_CN, BODY_FONT_EN, BULLET_SIZE, bold=True)
                else:
                    run = p.add_run(part)
                    set_run_font(run, BODY_FONT_CN, BODY_FONT_EN, BULLET_SIZE)

        elif btype == "text":
            add_styled_paragraph(
                doc, text, BODY_FONT_CN, BODY_FONT_EN, BODY_SIZE,
                space_before=0, space_after=3,
            )

    doc.save(output_path)
    return output_path


# ── PDF 生成（LibreOffice 主路线）─────────────────────────
def find_soffice():
    """查找 LibreOffice 可执行文件。"""
    candidates = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "soffice",
        "libreoffice",
    ]
    for c in candidates:
        try:
            subprocess.run([c, "--version"], capture_output=True, timeout=10)
            return c
        except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
            continue
    return None


def convert_to_pdf_libreoffice(docx_path, output_dir):
    """用 LibreOffice headless 将 .docx 转为 .pdf。

    LibreOffice 自动处理字体替换和嵌入，确保跨平台显示一致。
    """
    soffice = find_soffice()
    if not soffice:
        return None, "LibreOffice not found"

    profile_dir = os.path.join(output_dir, ".lo_profile")
    cmd = [
        soffice,
        "-env:UserInstallation=file://" + profile_dir,
        "--headless",
        "--norestore",
        "--convert-to", "pdf",
        "--outdir", output_dir,
        docx_path,
    ]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode == 0:
            pdf_path = str(Path(docx_path).with_suffix(".pdf"))
            if Path(pdf_path).exists():
                return pdf_path, None
        return None, f"soffice failed: {result.stderr.strip() or result.stdout.strip()}"
    except subprocess.TimeoutExpired:
        return None, "soffice timeout"
    except Exception as e:
        return None, str(e)


# ── PDF 生成（reportlab 备用路线）─────────────────────────
def find_chinese_fonts():
    """查找系统中可用的中文 TTF/TTC 字体文件。

    返回: {"body": path, "header": path}
    """
    body_candidates = [
        # macOS 宋体
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        # Linux Noto Serif CJK
        "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSerifCJK-Regular.ttc",
        # Windows 宋体
        "C:/Windows/Fonts/simsun.ttc",
        # 通用备选
        "/System/Library/Fonts/PingFang.ttc",
    ]
    header_candidates = [
        # macOS 黑体
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        # Linux Noto Sans CJK
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        # Windows 黑体/雅黑
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/msyh.ttc",
    ]
    result = {}
    for p in body_candidates:
        if os.path.exists(p):
            result["body"] = p
            break
    for p in header_candidates:
        if os.path.exists(p):
            result["header"] = p
            break
    return result


def generate_pdf_reportlab(blocks, output_path):
    """备用路线：用 reportlab 直接生成 PDF，注册并嵌入中文 TTF 字体。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.colors import black
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    fonts = find_chinese_fonts()

    # 注册正文字体
    body_font_name = "Helvetica"
    if "body" in fonts:
        try:
            if fonts["body"].endswith(".ttc"):
                pdfmetrics.registerFont(TTFont("ResumeBody", fonts["body"], subfontIndex=1))
            else:
                pdfmetrics.registerFont(TTFont("ResumeBody", fonts["body"]))
            body_font_name = "ResumeBody"
        except Exception:
            pass

    # 注册标题字体
    header_font_name = body_font_name
    if "header" in fonts:
        try:
            if fonts["header"].endswith(".ttc"):
                pdfmetrics.registerFont(TTFont("ResumeHeader", fonts["header"], subfontIndex=0))
            else:
                pdfmetrics.registerFont(TTFont("ResumeHeader", fonts["header"]))
            header_font_name = "ResumeHeader"
        except Exception:
            header_font_name = body_font_name

    # 样式定义
    styles = {
        "name": ParagraphStyle(
            "Name", fontName=header_font_name, fontSize=NAME_SIZE,
            alignment=TA_CENTER, spaceAfter=8, leading=NAME_SIZE * 1.3,
        ),
        "header": ParagraphStyle(
            "Header", fontName=header_font_name, fontSize=HEADER_SIZE,
            spaceBefore=12, spaceAfter=2, leading=HEADER_SIZE * 1.3,
        ),
        "subheader": ParagraphStyle(
            "Subheader", fontName=header_font_name, fontSize=SUBHEADER_SIZE,
            spaceBefore=6, spaceAfter=2, leading=SUBHEADER_SIZE * 1.3,
        ),
        "body": ParagraphStyle(
            "Body", fontName=body_font_name, fontSize=BODY_SIZE,
            spaceAfter=3, leading=BODY_SIZE * 1.5,
        ),
        "bullet": ParagraphStyle(
            "Bullet", fontName=body_font_name, fontSize=BULLET_SIZE,
            spaceAfter=2, leading=BULLET_SIZE * 1.5,
            leftIndent=18, bulletIndent=6,
        ),
    }

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        topMargin=MARGIN_TOP * cm, bottomMargin=MARGIN_BOTTOM * cm,
        leftMargin=MARGIN_LEFT * cm, rightMargin=MARGIN_RIGHT * cm,
    )

    story = []
    for block in blocks:
        btype = block["type"]
        text = block["text"]
        # XML 转义
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        # **bold** → <b>bold</b>
        text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)

        if btype == "name":
            story.append(Paragraph(text, styles["name"]))
        elif btype == "header":
            story.append(Paragraph(text, styles["header"]))
            story.append(HRFlowable(width="100%", thickness=0.5, color=black, spaceAfter=4))
        elif btype == "subheader":
            story.append(Paragraph(text, styles["subheader"]))
        elif btype == "bullet":
            story.append(Paragraph(f"\u2022 {text}", styles["bullet"]))
        elif btype == "text":
            story.append(Paragraph(text, styles["body"]))

    doc.build(story)
    return output_path


# ── 主入口 ───────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description="将优化后的简历导出为 .docx 和 .pdf",
    )
    parser.add_argument("input", help="输入 Markdown 文件路径")
    parser.add_argument("-o", "--output-dir", default=".", help="输出目录（默认当前目录）")
    parser.add_argument("-n", "--name", default="resume_optimized",
                        help="输出文件名（不含扩展名，默认 resume_optimized）")
    parser.add_argument("-f", "--format", choices=["docx", "pdf", "both"],
                        default="both", help="输出格式（默认 both）")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"错误: 输入文件不存在: {args.input}", file=sys.stderr)
        sys.exit(1)

    md_text = input_path.read_text(encoding="utf-8")
    blocks = parse_markdown(md_text)

    if not blocks:
        print("错误: 未能从输入文件解析出任何内容", file=sys.stderr)
        sys.exit(1)

    os.makedirs(args.output_dir, exist_ok=True)
    results = {}

    # 生成 .docx
    if args.format in ("docx", "both"):
        docx_path = os.path.join(args.output_dir, f"{args.name}.docx")
        try:
            generate_docx(blocks, docx_path)
            results["docx"] = docx_path
            print(f"OK DOCX: {docx_path}")
        except Exception as e:
            print(f"FAIL DOCX: {e}", file=sys.stderr)

    # 生成 .pdf
    if args.format in ("pdf", "both"):
        pdf_path = os.path.join(args.output_dir, f"{args.name}.pdf")

        # 主路线：LibreOffice 转换
        if "docx" in results:
            lo_path, lo_err = convert_to_pdf_libreoffice(results["docx"], args.output_dir)
            if lo_path:
                if lo_path != pdf_path:
                    os.rename(lo_path, pdf_path)
                results["pdf"] = pdf_path
                print(f"OK PDF (LibreOffice): {pdf_path}")
            else:
                print(f"WARN LibreOffice 失败: {lo_err}，尝试 reportlab", file=sys.stderr)
                try:
                    generate_pdf_reportlab(blocks, pdf_path)
                    results["pdf"] = pdf_path
                    print(f"OK PDF (reportlab): {pdf_path}")
                except Exception as e:
                    print(f"FAIL PDF: {e}", file=sys.stderr)
        else:
            try:
                generate_pdf_reportlab(blocks, pdf_path)
                results["pdf"] = pdf_path
                print(f"OK PDF (reportlab): {pdf_path}")
            except Exception as e:
                print(f"FAIL PDF: {e}", file=sys.stderr)

    if not results:
        print("错误: 未能生成任何输出文件", file=sys.stderr)
        sys.exit(1)

    return results


if __name__ == "__main__":
    main()
