#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
parse_score_table.py — 评标表 / 响应文件 解析脚本（投标模拟评标 · 评标表逆向工程 配套）

功能：
  1) 默认模式(table)：从《评标表》docx（抽取评分表）或 txt（启发式）提取评分项 → criteria.json
  2) --mode response：从《响应文件》docx / txt 提取正文 → response.txt（供 LLM 对照打分/审计）
  3) --mode audit：读 criteria.json + 响应正文，做关键词覆盖初判 → coverage_hints.json（含 format_sensitive 形式要件标记，供 LLM 精修）

依赖：
  - docx 路径：python-docx（已置于托管 venv）
  - txt 路径 / audit 模式：零依赖（标准库）

输出约定：
  criteria.json: [{id, factor, score, standard, type, evidence_required, is_knockout, needs_review}]
  coverage_hints.json: [{id, factor, matched_keywords, coverage_hint, format_sensitive}]
  type ∈ {objective, subjective, unknown}
  is_knockout: 含否决/废标/★/一票等红线 → true
  needs_review: 抽取不确定时置 true，提示 LLM 核定
"""
import argparse
import json
import os
import re
import sys

# ---------- 评分表识别关键词 ----------
TABLE_HEADER_HINTS = ["评审因素", "评分项", "评分因素", "评审项目", "评分标准", "分值", "得分", "评分细则"]
SUBJECTIVE_HINTS = ["优", "良", "中", "差", "合理", "较合理", "完善", "0-", "0－", "酌情", "酌情打分", "横向", "横向比较"]
OBJECTIVE_HINTS = ["提供", "具备", "具有", "证书", "复印件", "承诺", "须", "不得分", "每", "满分", "得", "分，", "以上", "以下", "齐全"]
# 否决/红线关键词（用于 is_knockout 标记，支撑模拟评标与合规审查）
KNOCKOUT_HINTS = ["否决", "废标", "无效", "一票", "★", "星号", "不得参与", "投标无效", "未提供.*不得分", "否则不得分", "不符合.*否决"]

# 覆盖审计用的停用词（过于通用的词不作为"命中"依据）
STOP = set("的 了 与 或 及 和 在 为 由 其 该 各 每 项 分 得 提供 具备 具有 证书 复印件 承诺 证明 盖章 签字 须 等 至 内 后 前 上 下 中 不 无 有 可 应 须 所 此 这 那 一 一个 一项 以上 以下".split())


def read_docx_tables(path):
    """返回 [(表序号, [[cell_text, ...], ...]), ...]"""
    import docx
    doc = docx.Document(path)
    out = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(cells)
        out.append((ti, rows))
    return out


def looks_like_score_table(rows):
    """根据表头/内容判定是否为评分表"""
    if not rows:
        return False
    flat = " ".join(sum(rows, []))
    head_hits = sum(1 for h in TABLE_HEADER_HINTS if h in flat)
    has_score_num = bool(re.search(r"\d+\s*分", flat))
    return head_hits >= 2 and has_score_num


def classify_type(standard):
    s = standard or ""
    obj = sum(1 for h in OBJECTIVE_HINTS if h in s)
    sub = sum(1 for h in SUBJECTIVE_HINTS if h in s)
    if obj > sub:
        return "objective"
    if sub > obj:
        return "subjective"
    return "unknown"


def is_knockout(standard):
    s = standard or ""
    return any(re.search(p, s) for p in KNOCKOUT_HINTS)


def extract_score(text):
    """从一段文本里尽量拿到分值数字"""
    m = re.search(r"(\d+(?:\.\d+)?)\s*分", text)
    if m:
        return float(m.group(1))
    return None


def parse_docx_criteria(path):
    tables = read_docx_tables(path)
    criteria = []
    cid = 0
    used_any = False
    for ti, rows in tables:
        if not looks_like_score_table(rows):
            continue
        used_any = True
        header_row = max(rows, key=lambda r: sum(1 for c in r for h in TABLE_HEADER_HINTS if h in c))
        col_factor = col_score = col_standard = None
        for ci, cell in enumerate(header_row):
            for h in ["评审因素", "评分项", "评审项目", "评分因素"]:
                if h in cell:
                    col_factor = ci
            if "分值" in cell or "得分" in cell:
                col_score = ci
            for h in ["评分标准", "评分细则", "标准", "评审标准"]:
                if h in cell:
                    col_standard = ci
        if col_factor is None:
            col_factor = 0
        if col_standard is None:
            col_standard = len(header_row) - 1
        if col_score is None and len(header_row) >= 2:
            col_score = len(header_row) - 2 if len(header_row) >= 3 else None
        for row in rows:
            if row == header_row:
                continue
            if not any(row):
                continue
            factor = row[col_factor] if col_factor < len(row) else ""
            standard_cells = [row[col_standard]] if col_standard < len(row) else []
            standard = " ".join(c for c in standard_cells if c).strip()
            score = None
            if col_score is not None and col_score < len(row):
                score = extract_score(row[col_score])
            if score is None:
                score = extract_score(factor + " " + standard)
            if not factor.strip() and not standard.strip():
                continue
            t = classify_type(standard) if standard else "unknown"
            criteria.append({
                "id": f"C{cid:02d}",
                "factor": factor.strip(),
                "score": score,
                "standard": standard.strip(),
                "type": t,
                "evidence_required": bool(re.search(r"提供|证书|复印件|承诺|证明|盖章|签字", standard or "")),
                "is_knockout": is_knockout(standard),
                "needs_review": (t == "unknown" or score is None),
            })
            cid += 1
    return criteria, used_any


# ---------- Markdown 表格解析（IMA/PDF 抽取的常见形态） ----------
def _split_md_blocks(text):
    """把全文切成若干 markdown 表块（连续 |...| 行）"""
    blocks, cur = [], []
    for line in text.splitlines():
        s = line.strip()
        if len(s) > 1 and s.startswith("|") and s.endswith("|"):
            cur.append(s)
        else:
            if cur:
                blocks.append(cur)
                cur = []
    if cur:
        blocks.append(cur)
    return blocks


def _md_cells(line):
    s = line.strip().strip("|")
    return [c.strip() for c in s.split("|")]


def _is_sep(cells):
    return len(cells) > 0 and all(re.fullmatch(r":?-+:?", c.strip()) for c in cells if c.strip() != "")


def _looks_header(cells):
    return any("序号" in c for c in cells) and any(k in " ".join(cells) for k in ["评分", "分值", "得分"])


def _detect_role(header_cells):
    role = dict(seq=None, factor=None, score=None, standard=None, type=None)
    for i, c in enumerate(header_cells):
        if "序号" in c:
            role["seq"] = i
        if c == "评分项" or any(k in c for k in ["评审因素", "评标因素", "评分因素", "评审项目"]):
            role["factor"] = i
        if "评分类型" in c or "评审类型" in c:
            role["type"] = i
        if "分值" in c or "得分" in c:
            role["score"] = i
        if any(k in c for k in ["评标内容", "评分标准", "评分细则", "评分项目内容",
                                 "评审标准", "评分办法", "标准", "评审内容"]):
            role["standard"] = i
    # 窄表兜底：序号|评分项|分值|评标内容 → 标准列取末列
    if role["standard"] is None and role["factor"] is not None and role["score"] is not None:
        n = len(header_cells)
        if role["score"] == n - 2 and role["factor"] == n - 3:
            role["standard"] = n - 1
    return role


def _parse_score_value(text):
    if not text:
        return None
    m = re.search(r"(\d+(?:\.\d+)?)\s*分", text)
    if m:
        return float(m.group(1))
    m = re.search(r"(\d+(?:\.\d+)?)\s*[-~]\s*(\d+(?:\.\d+)?)", text)
    if m:
        return max(float(m.group(1)), float(m.group(2)))
    m = re.search(r"(\d+(?:\.\d+)?)", text)
    if m:
        return float(m.group(1))
    return None


def _is_clean_int(s):
    return bool(re.fullmatch(r"\d+", (s or "").strip()))


def parse_md_criteria(text):
    """解析 markdown 表格形式的评分表（窄量规/宽矩阵/合并单元格续行）。返回 (criteria, ok)。"""
    blocks = _split_md_blocks(text)
    # 全局 role：取第一个含「序号」的表头
    role = None
    for b in blocks:
        for row in b:
            cells = _md_cells(row)
            if _looks_header(cells):
                role = _detect_role(cells)
                break
        if role:
            break
    if role is None:
        return None, False

    criteria, cid, last_factor = [], 0, ""
    for b in blocks:
        hdr = None
        for row in b:
            cells = _md_cells(row)
            if _looks_header(cells):
                hdr = cells
                break
        use_role = _detect_role(hdr) if hdr else role
        for row in b:
            cells = _md_cells(row)
            if not cells or _is_sep(cells):
                continue
            if hdr and cells == hdr:
                continue
            if re.match(r"^(合计|总计|汇总|总得分)", cells[0] or ""):
                continue
            if all(not c for c in cells):
                continue
            n = len(cells)

            def getc(i):
                return cells[i] if (i is not None and i < n) else ""

            seq = getc(use_role["seq"])
            factor_raw = getc(use_role["factor"])
            type_raw = getc(use_role["type"])
            score_raw = getc(use_role["score"])
            standard = getc(use_role["standard"])
            # 标准列若因 bidder 列导致为空，末列兜底（窄表末列）
            if not standard:
                for c in reversed(cells):
                    if c:
                        standard = c
                        break
            # 跳过纯投标人得分行（无标准/无分值/无因素）
            if not standard and not score_raw and not factor_raw:
                continue
            # factor 推导（合并单元格续行：factor_raw 为空时继承）
            if factor_raw:
                factor = factor_raw
            elif type_raw:
                factor = f"{type_raw}·{seq}" if _is_clean_int(seq) else type_raw
            elif _is_clean_int(seq):
                factor = f"评分项{seq}"
            else:
                factor = last_factor
            # type 判定
            if "【客观分】" in standard:
                t = "objective"
            elif "【主观分】" in standard:
                t = "subjective"
            elif type_raw and "主观" in type_raw:
                t = "subjective"
            elif type_raw and "客观" in type_raw:
                t = "objective"
            else:
                t = classify_type(standard) if standard else "unknown"
            score = _parse_score_value(score_raw) if score_raw else None
            if score is None and standard:
                score = _parse_score_value(standard)
            if factor_raw or (factor and factor != last_factor):
                last_factor = factor
            criteria.append({
                "id": f"C{cid:02d}",
                "factor": factor.strip(),
                "score": score,
                "standard": standard.strip(),
                "type": t,
                "evidence_required": bool(re.search(r"提供|证书|复印件|承诺|证明|盖章|签字", standard or "")),
                "is_knockout": is_knockout(standard),
                "needs_review": (t == "unknown" or score is None
                                 or (not factor_raw and factor.startswith("评分项"))),
            })
            cid += 1
    if not criteria:
        return None, False
    return criteria, True


def parse_txt_criteria(path):
    text = open(path, encoding="utf-8", errors="ignore").read()
    # 优先尝试 markdown 表格（IMA/PDF 抽取的真实形态）
    md, ok = parse_md_criteria(text)
    if ok and md:
        return md, True
    # 兜底：散文式「因素(分): 标准」正则
    criteria = []
    cid = 0
    pat1 = re.compile(r"^(?:\d+[.\、]?)?\s*([^（(]{2,30}?)\s*[（(]\s*(\d+(?:\.\d+)?)\s*分\s*[）)]\s*[:：]?\s*(.*)$")
    for m in pat1.finditer(text):
        factor, score, standard = m.group(1).strip(), float(m.group(2)), m.group(3).strip()
        t = classify_type(standard) if standard else "unknown"
        criteria.append({
            "id": f"C{cid:02d}", "factor": factor, "score": score, "standard": standard,
            "type": t, "evidence_required": bool(re.search(r"提供|证书|复印件|承诺|证明", standard)),
            "is_knockout": is_knockout(standard),
            "needs_review": (t == "unknown"),
        })
        cid += 1
    if criteria:
        return criteria, True
    return [{
        "id": "C00", "factor": "(全文，需LLM按评标表结构拆分)", "score": None,
        "standard": text.strip()[:2000], "type": "unknown", "evidence_required": False,
        "is_knockout": is_knockout(text),
        "needs_review": True,
    }], False


def extract_response(path):
    """抽取响应正文为纯文本"""
    if path.lower().endswith(".docx"):
        import docx
        doc = docx.Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells if c.text.strip()))
        return "\n".join(parts)
    else:
        return open(path, encoding="utf-8", errors="ignore").read()


def keywords_of(text):
    """取候选关键词：先按标点断句（保留>=3字短语），长串再补 3/4 字滑动窗口。
    避开 2 字泛词（方案/投标/设计…），在中文无分词下兼顾召回与精度。"""
    segs = re.split(r"[、，。；：（）()\s\-—/｜|]", text or "")
    grams = set()
    for s in segs:
        s = s.strip()
        if len(s) >= 3:
            grams.add(s)
        if len(s) > 5:  # 无标点长串补滑动窗口，避免只产出一个超长 token
            for n in (3, 4):
                for i in range(len(s) - n + 1):
                    grams.add(s[i:i + n])
    return [g for g in grams if g not in STOP]


# 形式要件关键词（用于标记评分项是否需核验形式，提示 LLM 重点查 format-risk）
FORMAT_KEYWORDS = ["原件", "扫描", "扫描件", "公章", "盖章", "签字", "签名", "骑缝", "截图", "公证", "特定页面"]


def audit_coverage(criteria_path, response_path, out_path):
    """读 criteria.json + 响应正文，做关键词覆盖初判

    coverage_hint 仅为初判（关键词命中），最终状态须由 LLM 精修；
    format_sensitive=True 表示该评分项标准含形式要件要求（原件/公章/签字等），
    提示 LLM 在精修时重点核验「内容在但形式不符」的 format-risk 状态。
    """
    criteria = json.load(open(criteria_path, encoding="utf-8"))
    resp = extract_response(response_path)
    hints = []
    for c in criteria:
        kws = keywords_of(c.get("factor", "") + " " + c.get("standard", ""))
        matched = [k for k in kws if k in resp]
        # 覆盖初判：命中 >=2 个实义词 → covered；命中 1 → partial；0 → missing
        if len(matched) >= 2:
            hint = "covered"
        elif len(matched) == 1:
            hint = "partial"
        else:
            hint = "missing"
        # 形式要件敏感标记：标准文本含形式关键词即置位（不臆断响应是否达标）
        format_sensitive = any(k in (c.get("standard", "") or "") for k in FORMAT_KEYWORDS)
        hints.append({
            "id": c.get("id"),
            "factor": c.get("factor", ""),
            "matched_keywords": matched[:8],
            "coverage_hint": hint,
            "format_sensitive": format_sensitive,
        })
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(hints, f, ensure_ascii=False, indent=2)
    cnt = {"covered": 0, "partial": 0, "missing": 0}
    for h in hints:
        cnt[h["coverage_hint"]] += 1
    n_fmt = sum(1 for h in hints if h["format_sensitive"])
    print(f"覆盖初判 → {out_path}：covered={cnt['covered']} partial={cnt['partial']} missing={cnt['missing']}（初判，待 LLM 精修）；其中 {n_fmt} 项含形式要件（重点查 format-risk）")


def main():
    ap = argparse.ArgumentParser(description="评标表/响应文件解析")
    ap.add_argument("input", help="评标表/响应文件 docx/txt，或 audit 模式下的响应文件")
    ap.add_argument("--out", default="criteria.json", help="输出路径")
    ap.add_argument("--mode", choices=["table", "response", "audit"], default="table",
                    help="table=抽评分项；response=抽响应正文；audit=覆盖初判(需 --criteria)")
    ap.add_argument("--criteria", help="audit 模式：criteria.json 路径")
    args = ap.parse_args()

    if args.mode == "response":
        txt = extract_response(args.input)
        with open(args.out, "w", encoding="utf-8") as f:
            f.write(txt)
        print(f"响应正文已提取 → {args.out}（{len(txt)} 字符）")
        return

    if args.mode == "audit":
        if not args.criteria:
            print("audit 模式需提供 --criteria criteria.json", file=sys.stderr)
            sys.exit(2)
        audit_coverage(args.criteria, args.input, args.out)
        return

    is_docx = args.input.lower().endswith(".docx")
    if is_docx:
        criteria, used = parse_docx_criteria(args.input)
    else:
        criteria, used = parse_txt_criteria(args.input)

    with open(args.out, "w", encoding="utf-8") as f:
        json.dump(criteria, f, ensure_ascii=False, indent=2)
    ko = sum(1 for c in criteria if c.get("is_knockout"))
    flag = "⚠️ 部分项需人工核定" if any(c["needs_review"] for c in criteria) else "✅ 抽取较完整"
    print(f"评标表解析 → {args.out}：共 {len(criteria)} 项（否决项 {ko}）| {flag}")
    if not used:
        print("提示：未识别到标准评分表结构，已按全文返回，请由 LLM 拆分核定。")


if __name__ == "__main__":
    main()
