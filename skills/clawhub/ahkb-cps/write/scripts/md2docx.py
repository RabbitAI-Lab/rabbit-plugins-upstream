"""将 Markdown 文章转换为专业 Word (.docx) 文档

用法:
    python md2docx.py <input.md> <output.docx> [--title "标题"] [--subtitle "副标题"]

    python md2docx.py article.md output.docx
    python md2docx.py article.md output.docx --title "系统科学导论" --subtitle "从控制论到大系统观"
"""

import sys
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import date


def parse_args():
    """解析命令行参数。返回 (input, output, title, subtitle)。"""
    args = sys.argv[1:]
    if not args or args[0] in ("-h", "--help"):
        print(__doc__)
        sys.exit(0)

    input_file = args[0]
    output_file = args[1] if len(args) > 1 else input_file.replace(".md", ".docx")

    # 从文件名推断默认标题和副标题
    basename = input_file.rsplit("/", 1)[-1].rsplit("\\", 1)[-1]
    default_title = basename.replace(".md", "").replace("_", " ").strip()
    title = default_title
    subtitle = ""

    # 解析命名参数
    i = 2
    while i < len(args):
        if args[i] == "--title" and i + 1 < len(args):
            title = args[i + 1]
            i += 2
        elif args[i] == "--subtitle" and i + 1 < len(args):
            subtitle = args[i + 1]
            i += 2
        else:
            print(f"未知参数: {args[i]}", file=sys.stderr)
            i += 1

    return input_file, output_file, title, subtitle


def build_docx(input_md, output_docx, title, subtitle):
    """构建 .docx 文档。"""

    doc = Document()

    # ── 页面设置 ──
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ── 设置默认字体 ──
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    # ── 设置标题样式 ──
    for level, size, bold, color in [
        ('Heading 1', 18, True, RGBColor(0x1A, 0x3C, 0x6D)),
        ('Heading 2', 14, True, RGBColor(0x2B, 0x57, 0x8C)),
        ('Heading 3', 12, True, RGBColor(0x3A, 0x6E, 0xA5)),
    ]:
        s = doc.styles[level]
        s.font.name = '黑体'
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = color
        s.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)

    # ── 封面 ──
    for _ in range(6):
        doc.add_paragraph('')

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 副标题（可选）
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x4A, 0x6F, 0x9A)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 30)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)
    run.font.size = Pt(12)

    # 日期
    today = date.today()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{today.year}年{today.month}月')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ── 解析 Markdown ──
    with open(input_md, 'r', encoding='utf-8') as f:
        md_content = f.read()

    _parse_and_add(doc, md_content)

    # ── 保存 ──
    doc.save(output_docx)
    print(f"✅ Word 文档已生成：{output_docx}")


def _parse_and_add(doc, md_text):
    """解析 Markdown 并添加到文档"""
    lines = md_text.split('\n')
    i = 0
    in_blockquote = False
    blockquote_lines = []

    while i < len(lines):
        line = lines[i]

        # 跳过 YAML 前置元数据 (--- ... ---)
        if line.strip() == '---' and i == 0:
            i += 1
            while i < len(lines) and lines[i].strip() != '---':
                i += 1
            i += 1
            continue

        # 引用块
        if line.startswith('> '):
            blockquote_lines.append(line[2:])
            in_blockquote = True
            i += 1
            continue
        else:
            if in_blockquote:
                _flush_blockquote(doc, blockquote_lines)
                blockquote_lines = []
                in_blockquote = False

        # 标题
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # 列表项
        list_match = re.match(r'^(\s*)[*-]\s+(.+)$', line)
        if list_match:
            text = list_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(text)
            run.font.size = Pt(11)
            i += 1
            continue

        # 水平分隔线
        if re.match(r'^---$|^___$|^\*\*\*$', line.strip()):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('— ' * 15)
            run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            run.font.size = Pt(10)
            i += 1
            continue

        # 空行
        if line.strip() == '':
            i += 1
            continue

        # 普通正文
        _add_paragraph(doc, line.strip())
        i += 1

    # 如果最后还有未输出的引用块
    if in_blockquote and blockquote_lines:
        _flush_blockquote(doc, blockquote_lines)


def _add_paragraph(doc, text):
    """添加正文段落，支持 ** 加粗标记。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.font.bold = True
        else:
            run = p.add_run(part)
    p.paragraph_format.first_line_indent = Cm(0.74)


def _flush_blockquote(doc, lines):
    """输出引用块。"""
    text = '\n'.join(lines)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'"{text}"')
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = '楷体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


if __name__ == "__main__":
    input_file, output_file, title, subtitle = parse_args()
    if not input_file:
        print("用法: python md2docx.py <input.md> [output.docx] [--title TITLE] [--subtitle SUBTITLE]", file=sys.stderr)
        sys.exit(1)
    build_docx(input_file, output_file, title, subtitle)
