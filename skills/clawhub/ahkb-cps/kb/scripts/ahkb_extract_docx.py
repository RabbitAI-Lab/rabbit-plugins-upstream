"""
ahkb_extract_docx.py — DOCX 解析器
提取：段落文字（按标题分节）、嵌入图片、嵌入音视频
输出：chunks 嵌套结构，资源归属明确
"""
from pathlib import Path
import zipfile, os, hashlib
from xml.etree import ElementTree as ET

# ─── 媒体文件类型检测 ───

IMAGE_EXTS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.wmf', '.emf'}
VIDEO_EXTS = {'.mp4', '.avi', '.mov', '.wmv', '.m4v', '.mpg', '.mpeg', '.flv', '.webm'}
AUDIO_EXTS = {'.mp3', '.wav', '.wma', '.aac', '.ogg', '.flac', '.m4a'}
MEDIA_EXTS = IMAGE_EXTS | VIDEO_EXTS | AUDIO_EXTS


def _get_media_type(ext):
    ext = ext.lower()
    if ext in IMAGE_EXTS:
        return "image"
    elif ext in VIDEO_EXTS:
        return "video"
    elif ext in AUDIO_EXTS:
        return "audio"
    return "other"


def extract_docx(filepath, workspace):
    """Extract text, images, and media from DOCX. Returns structured dict with chunks."""
    from docx import Document

    doc = Document(filepath)
    base = Path(filepath).stem
    safe_base = "".join(c if c.isalnum() or c in '-_ ' else '_' for c in base)

    # 目录结构
    img_dir = Path(workspace) / "图片及其他资源" / "images"
    video_dir = Path(workspace) / "图片及其他资源" / "videos"
    audio_dir = Path(workspace) / "图片及其他资源" / "audios"
    other_dir = Path(workspace) / "图片及其他资源" / "others"
    for d in [img_dir, video_dir, audio_dir, other_dir]:
        d.mkdir(parents=True, exist_ok=True)

    result = {
        "file": str(filepath),
        "type": "docx",
        "metadata": {"paragraph_count": len(doc.paragraphs)},
        "chunks": [],
        "full_text": "",
        "resources_flat": [],
    }

    # ── 第一步：提取文字并按标题分节 ──
    sections_text = []
    current_heading = None
    current_content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            if current_heading:
                sections_text.append({
                    "level": _heading_level(para.style.name),
                    "heading": current_heading,
                    "text": "\n".join(current_content),
                })
            current_heading = text
            current_content = []
        else:
            current_content.append(text)

    if current_heading:
        sections_text.append({"level": 1, "heading": current_heading, "text": "\n".join(current_content)})
    elif current_content:
        sections_text.append({"level": 0, "heading": "(正文)", "text": "\n".join(current_content)})

    if not sections_text:
        sections_text.append({"level": 0, "heading": "(正文)", "text": ""})

    # ── 第二步：从 ZIP 中提取所有媒体文件（图片+视频+音频）──
    all_media = []  # [{filename, ext, bytes, rid, para_idx}]
    seen_blobs = set()

    try:
        with zipfile.ZipFile(filepath, 'r') as z:
            # 读取 document.xml 用于图片定位
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)

            ns = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            }

            # ── 建立 rId → media 文件名 的映射（从 .rels）──
            rels_map = {}
            try:
                with z.open("word/_rels/document.xml.rels") as f:
                    rels_tree = ET.parse(f)
                for rel in rels_tree.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                    rid = rel.get('Id')
                    target = rel.get('Target', '')
                    # 匹配 media/ 下的所有文件（图片、视频、音频）
                    if target.startswith('media/'):
                        rels_map[rid] = f"word/{target}"
            except Exception:
                pass

            # ── 遍历段落，找图片和媒体引用 ──
            # 收集: 段落索引 → rId 列表
            para_media_refs = {}  # rid → para_idx
            for para_idx, para_elem in enumerate(tree.findall('.//w:p', ns)):
                # 找图片引用：w:drawing → wp:inline → a:blip
                for drawing in para_elem.findall('.//w:drawing', ns):
                    for blip in drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed and embed in rels_map:
                            para_media_refs[embed] = para_idx

                # 找其他媒体引用（视频/音频等通过 w:object 或 w:oleObject）
                for ole in para_elem.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed', ns):
                    rid = ole.text if ole.text else ole.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rid and rid in rels_map:
                        para_media_refs[rid] = para_idx

            # ── 段落索引 → 节索引 ──
            para_to_section = {}
            section_idx = 0
            for para_idx, para_elem in enumerate(tree.findall('.//w:p', ns)):
                pPr = para_elem.find('.//w:pPr', ns)
                if pPr is not None:
                    pStyle = pPr.find('w:pStyle', ns)
                    if pStyle is not None:
                        style_val = pStyle.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
                        if style_val and style_val.lower().startswith('heading'):
                            section_idx += 1
                para_to_section[para_idx] = min(section_idx, len(sections_text) - 1) if sections_text else 0

            # ── 列出 word/media/ 下的所有文件 ──
            media_files_in_zip = {}
            for name in z.namelist():
                if name.startswith("word/media/"):
                    fname = Path(name).name
                    ext = Path(name).suffix.lower()
                    if ext in MEDIA_EXTS:
                        media_files_in_zip[name] = {
                            "filename": fname,
                            "ext": ext[1:],
                            "bytes": z.read(name),
                        }

            # ── 按 refs 提取并保存媒体文件 ──
            for rid, para_idx in para_media_refs.items():
                media_path = rels_map.get(rid)
                if not media_path or media_path not in media_files_in_zip:
                    continue
                minfo = media_files_in_zip[media_path]
                # 去重（基于内容哈希）
                blob_hash = hashlib.md5(minfo["bytes"]).hexdigest()[:16]
                if blob_hash in seen_blobs:
                    continue
                seen_blobs.add(blob_hash)

                mtype = _get_media_type(f".{minfo['ext']}")
                target_dir = {"image": img_dir, "video": video_dir, "audio": audio_dir, "other": other_dir}[mtype]

                # 生成文件名
                media_ext = minfo["ext"]
                count = len(all_media) + 1
                fname = f"{safe_base}-{mtype}{count:02d}.{media_ext}"
                with open(target_dir / fname, "wb") as f:
                    f.write(minfo["bytes"])

                # 获取上下文文本
                sec_idx = para_to_section.get(para_idx, 0)
                ctx_text = ""
                if sections_text and 0 <= sec_idx < len(sections_text):
                    ctx_text = sections_text[sec_idx].get("text", "")

                all_media.append({
                    "type": mtype,
                    "filename": fname,
                    "ext": media_ext,
                    "context_text": ctx_text,
                    "source_ref": f"paragraph {para_idx}, section '{sections_text[sec_idx].get('heading', '')}'" if sections_text else "",
                    "para_idx": para_idx,
                    "sec_idx": sec_idx,
                })

            # ── 提取未在 XML 中引用的媒体文件（如页眉页脚中的图片）──
            referenced_paths = {rels_map.get(rid) for rid in para_media_refs}
            for media_path, minfo in media_files_in_zip.items():
                if media_path not in referenced_paths:
                    blob_hash = hashlib.md5(minfo["bytes"]).hexdigest()[:16]
                    if blob_hash in seen_blobs:
                        continue
                    seen_blobs.add(blob_hash)

                    mtype = _get_media_type(f".{minfo['ext']}")
                    target_dir = {"image": img_dir, "video": video_dir, "audio": audio_dir, "other": other_dir}[mtype]

                    count = len(all_media) + 1
                    fname = f"{safe_base}-{mtype}{count:02d}.{minfo['ext']}"
                    with open(target_dir / fname, "wb") as f:
                        f.write(minfo["bytes"])

                    ctx = sections_text[0].get("text", "") if sections_text else ""
                    all_media.append({
                        "type": mtype,
                        "filename": fname,
                        "ext": minfo["ext"],
                        "context_text": ctx,
                        "source_ref": f"word/media/ (unreferenced)",
                        "para_idx": -1,
                        "sec_idx": 0,
                    })

    except Exception as e:
        # ZIP 解析失败时回退到简单方式
        try:
            with zipfile.ZipFile(filepath, 'r') as z:
                for name in z.namelist():
                    if name.startswith("word/media/"):
                        ext = Path(name).suffix.lower()
                        if ext not in MEDIA_EXTS:
                            continue
                        fname = Path(name).name
                        mtype = _get_media_type(ext)
                        target_dir = {"image": img_dir, "video": video_dir, "audio": audio_dir, "other": other_dir}[mtype]
                        img_bytes = z.read(name)
                        count = len(all_media) + 1
                        out_fname = f"{safe_base}-{mtype}{count:02d}{ext}"
                        with open(target_dir / out_fname, "wb") as f:
                            f.write(img_bytes)
                        ctx = sections_text[0].get("text", "") if sections_text else ""
                        all_media.append({
                            "type": mtype,
                            "filename": out_fname,
                            "ext": ext[1:],
                            "context_text": ctx,
                            "source_ref": f"word/media/ (fallback)",
                            "para_idx": -1,
                            "sec_idx": 0,
                        })
        except Exception:
            pass

    # ── 第三步：构建 chunks ──
    # 将 all_media 按 sec_idx 分组到对应的节
    media_by_section = {}
    for m in all_media:
        sec_idx = m["sec_idx"]
        if sec_idx not in media_by_section:
            media_by_section[sec_idx] = []
        media_by_section[sec_idx].append(m)

    for sec_idx, sec in enumerate(sections_text):
        chunk_resources = media_by_section.get(sec_idx, [])
        # 去掉 para_idx 和 sec_idx（内部字段，不输出）
        clean_resources = []
        for r in chunk_resources:
            clean_resources.append({
                "type": r["type"],
                "filename": r["filename"],
                "ext": r["ext"],
                "context_text": r["context_text"],
                "source_ref": r["source_ref"],
            })

        heading = sec["heading"]
        text = sec["text"]
        chunk = {
            "id": f"sec-{sec_idx+1:03d}",
            "heading": heading,
            "source_position": heading,
            "type": "section",
            "text": text,
            "resources": clean_resources,
        }
        result["chunks"].append(chunk)
        result["full_text"] += f"\n\n--- {heading} ---\n\n{text}"

    # ── 构建扁平资源列表 ──
    flat = []
    for chunk in result["chunks"]:
        for r in chunk["resources"]:
            r_copy = dict(r)
            r_copy["belongs_to_chunk"] = chunk["id"]
            r_copy["chunk_heading"] = chunk["heading"]
            r_copy["chunk_text"] = chunk["text"]
            flat.append(r_copy)
    result["resources_flat"] = flat

    return result


def _heading_level(style_name):
    """Convert heading style name to level number."""
    name = style_name.lower()
    for level in range(1, 10):
        if f"heading {level}" in name or f"headings{level}" in name:
            return level
    if "heading" in name or "标题" in name:
        return 1
    return 1
