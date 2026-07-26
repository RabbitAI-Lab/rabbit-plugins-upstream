#!/usr/bin/env python3
"""Render findings into Markdown + DOCX report with role-based templates.

Role-based rendering:
  --role bidder   → 投标人版：红线→投标策略汇总→报价影响→时限权利→逐条明细
  --role tenderer → 招标人版（待完善）：合规自检→称谓一致性→质疑风险预判→逐条明细

Each finding item may carry: clause_id, change_type, old_text, new_text,
numeric_delta, dimension, sentiment, severity, impact, basis, basis_source,
action, confidence, context, compliance_check, timeliness_warning,
final_severity, bid_impact (bidder), implicit_barrier (bidder),
objection_eligible (bidder), pricing_note/pricing_risk/pricing_action (bidder).
"""
import argparse
import json
from collections import Counter

try:
    from docx import Document
except Exception:
    Document = None

SENT = {'风险': '[风险]', '中性': '[中性]', '利好': '[利好]'}
BID_IMPACT_LABELS = {
    '报价需调整': '💰',
    '技术方案需改': '🔧',
    '资格/业绩需补': '📋',
    '时限权利受影响': '⏰',
    '仅关注不需动作': '👁',
    '无影响': '✅',
}


def load(path):
    with open(path, encoding='utf-8') as fh:
        data = json.load(fh)
    if isinstance(data, list):
        return data, {}
    return data.get('items', []), data.get('meta', {})


def _loc(it):
    ctx = it.get('context')
    return f"位置：{ctx}" if ctx else ''


def _severity(it):
    return it.get('final_severity') or it.get('severity', '')


def _flagged(items):
    return [i for i in items
            if i.get('is_redline')
            or i.get('sentiment') == '风险'
            or i.get('severity') == '高'
            or i.get('final_severity') == '高']


# ============================================================
# BIDDER REPORT TEMPLATES
# ============================================================

def _bidder_summary_section(items):
    """投标策略汇总 — 报告首屏核心决策信息。"""
    lines = []
    c = Counter(i.get('sentiment', '中性') for i in items)
    redline = sum(1 for i in items if i.get('is_redline'))
    ibarrier = sum(1 for i in items if i.get('implicit_barrier'))
    tw = any(i.get('timeliness_warning') for i in items)
    obj_ok = sum(1 for i in items if i.get('objection_eligible'))
    pricing_items = [i for i in items if i.get('bid_impact') == '报价需调整']
    tech_items = [i for i in items if i.get('bid_impact') == '技术方案需改']
    qual_items = [i for i in items if i.get('bid_impact') == '资格/业绩需补']

    lines.append('## 影响概览')
    for k in ['风险', '中性', '利好']:
        if c.get(k):
            lines.append(f"- {SENT.get(k, k)}：{c[k]} 条")
    if redline:
        lines.append(f"- 🔴 **[红线]**：{redline} 条（强制合规风险，建议优先处理）")
    if ibarrier:
        lines.append(f"- ⚠️ **[隐性门槛]**：{ibarrier} 条（准入范围可能收窄）")
    lines.append('')

    # 投标行动清单
    actions = []
    if redline or obj_ok:
        actions.append("存在合规红线或可质疑项 —— 评估是否提交书面异议/质疑")
    if pricing_items:
        actions.append(f"**报价需重新测算**（{len(pricing_items)} 条变更影响成本/收入项）")
    if tech_items:
        actions.append(f"**技术方案需调整**（{len(tech_items)} 条变更涉及技术响应内容）")
    if qual_items:
        actions.append(f"**资格/业绩需核查**（{len(qual_items)} 条涉及准入条件变化）")
    if tw:
        actions.append("**时限权利** —— 变更距截止<15日，有权要求顺延")

    if actions:
        lines.append('## 📋 投标行动清单')
        for idx, a in enumerate(actions, 1):
            lines.append(f"{idx}. {a}")
        lines.append('')

    # 报价影响专段
    if pricing_items:
        lines.append('## 💰 报价影响评估')
        for it in pricing_items:
            note = it.get('pricing_note') or it.get('impact', '')
            risk = it.get('pricing_risk', '')
            paction = it.get('pricing_action') or it.get('action', '')
            risk_tag = f"（风险等级：{risk}）" if risk else ''
            lines.append(f"- **{it.get('clause_id','')}**：{note} {risk_tag}")
            if paction:
                lines.append(f"  → {paction}")
        lines.append('')

    # 时限权利专段
    if tw:
        lines.append('## ⏰ 投标人时限权利')
        lines.append('- **顺延权**：本次变更可能影响投标文件编制且距截止不足15日，依法你有权要求顺延投标截止时间。')
        lines.append('- **依据**：87号令第二十七条 / 招标投标法实施条例第二十一条')
        lines.append('- **行动**：立即致函招标人确认新截止时间；保留要求顺延的证据。')
        lines.append('')

    return '\n'.join(lines)


def _bidder_detail(idx, it):
    """投标人版单条详情模板。"""
    out = []
    loc = _loc(it)
    tag = f"{it.get('dimension','')}/{it.get('sentiment','')}"
    if it.get('is_redline'):
        tag += " 🔴"
    if it.get('implicit_barrier'):
        tag += " ⚠️"

    bi = it.get('bid_impact', '')
    bi_icon = BID_IMPACT_LABELS.get(bi, '')
    head = f"### {idx}. [{tag}] {it.get('clause_id','')}"
    if bi_icon:
        head += f" {bi_icon}"
    if loc:
        head += f"\n> {loc}"
    out.append(head)

    # Meta line
    meta_parts = [
        f"变更类型：{it.get('change_type','')}",
        f"严重度：{_severity(it)}",
        f"置信度：{it.get('confidence','')}",
    ]
    if bi:
        meta_parts.insert(2, f"投标影响：{bi}")
    out.append(f"- {'　'.join(meta_parts)}")

    if it.get('numeric_delta'):
        out.append(f"- 数值变更：{it['numeric_delta']}")
    if it.get('old_text'):
        out.append(f"- 原版：{it['old_text'][:500]}")
    if it.get('new_text'):
        out.append(f"- 新版：{it['new_text'][:500]}")
    out.append(f"- 影响：{it.get('impact','')}")

    if it.get('pricing_note'):
        out.append(f"- 📊 **报价影响**：{it['pricing_note']}（风险：{it.get('pricing_risk','N/A')}）")
        if it.get('pricing_action'):
            out.append(f"  → {it['pricing_action']}")

    out.append(f"- 依据：{it.get('basis','')}（{it.get('basis_source','')}）")
    if it.get('compliance_check'):
        out.append(f"- 合规核查：{it['compliance_check']}")
    if it.get('timeliness_warning'):
        out.append('- ⏰ **时限预警**：你有权要求顺延投标截止时间（<15日 + 影响编制）')
    if it.get('objection_eligible') and not it.get('is_redline'):
        out.append('- ✉️ 该条存在隐性风险，若认为不合理可在质疑期内提出书面质疑')

    # Action as final actionable item
    action = it.get('action', '')
    if action:
        out.append(f"- **建议操作**：{action}")
    out.append('')
    return '\n'.join(out)


def markdown_bidder(items, meta):
    """投标人版完整 Markdown 报告。"""
    lines = ['# 招标文件版本比对报告 · 投标人版', '']
    lines.append(f"- 原版：{meta.get('old_file','')}（{meta.get('old_date','')}）")
    lines.append(f"- 新版：{meta.get('new_file','')}（{meta.get('new_date','')}）")
    lines.append(f"- 投标截止：{meta.get('bid_deadline','未提供')}")
    lines.append(f"- 差异条数：{len(items)}")
    lines.append('')

    # Section 1: Summary + Action Items
    lines.append(_bidder_summary_section(items))

    # Section 2: Redlines / High Risk first
    lines.append('## 🔴 红线与高风险明细（优先处理）')
    flagged = _flagged(items)
    if not flagged:
        lines.append('（无）\n')
    for idx, it in enumerate(flagged, 1):
        lines.append(_bidder_detail(idx, it))

    # Section 3: All details
    lines.append('## 全部差异明细')
    for idx, it in enumerate(items, 1):
        lines.append(_bidder_detail(idx, it))
    return '\n'.join(lines)


# ============================================================
# TENDERER REPORT TEMPLATE (placeholder)
# ============================================================

def markdown_tenderer(items, meta):
    """招标人版报告（待完善）。"""
    lines = ['# 招标文件版本自检报告 · 招标人版', '']
    lines.append('> ⚠️ 招标人版报告模板正在建设中，当前输出为简化视图。\n')
    lines.append(f"- 原版：{meta.get('old_file','')}")
    lines.append(f"- 新版：{meta.get('new_file','')}")
    lines.append(f"- 差异条数：{len(items)}")
    lines.append('')
    c = Counter(i.get('sentiment', '中性') for i in items)
    lines.append('## 变更摘要')
    for k in ['风险', '中性', '利好']:
        if c.get(k):
            lines.append(f"- {k}：{c[k]} 条")
    lines.append('')
    lines.append('## 逐条变更')
    for idx, it in enumerate(items, 1):
        lines.append(f"### {idx}. [{it.get('dimension','')}/{it.get('sentiment','')}] {it.get('clause_id','')}")
        if it.get('impact'):
            lines.append(f"- {it.get('impact','')}")
        else:
            lines.append(f"- {it.get('old_text','')[:200]} → {it.get('new_text','')[:200]}")
        lines.append('')
    return '\n'.join(lines)


# ============================================================
# DOCX RENDERER (shared structure, role-aware heading)
# ============================================================

def docx_render(items, meta, out, role='bidder'):
    if Document is None:
        raise RuntimeError('python-docx not installed')
    doc = Document()
    title = '招标文件版本比对报告'
    if role == 'tenderer':
        title += ' · 招标人版（待完善）'
    else:
        title += ' · 投标人版'
    doc.add_heading(title, 0)

    doc.add_paragraph(f"原版：{meta.get('old_file','')}（{meta.get('old_date','')}）")
    doc.add_paragraph(f"新版：{meta.get('new_file','')}（{meta.get('new_date','')}）")
    doc.add_paragraph(f"投标截止：{meta.get('bid_deadline','未提供')}")
    doc.add_paragraph(f"差异条数：{len(items)}")

    # Use markdown content to populate body (simple approach: parse sections)
    if role == 'tenderer':
        md = markdown_tenderer(items, meta)
    else:
        md = markdown_bidder(items, meta)

    # Render markdown sections into docx paragraphs
    for line in md.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph('')
        elif line.startswith('# '):
            doc.add_heading(line[2:], 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
        elif line.startswith('- '):
            # Bold markers within list items
            text = line[2:]
            p = doc.add_paragraph(text, style='List Bullet')
        else:
            doc.add_paragraph(line)

    doc.save(out)
    print(f'Report ({role}) -> {out}')


def main():
    ap = argparse.ArgumentParser(description='Build version-compare report (role-aware)')
    ap.add_argument('--findings', required=True, help='findings.json path')
    ap.add_argument('--out', required=True, help='output .docx path')
    ap.add_argument('--markdown', default='', help='optional output .md path')
    ap.add_argument('--role', default='bidder', choices=['bidder', 'tenderer'],
                    help='report role: bidder (default) or tenderer')
    args = ap.parse_args()

    items, meta = load(args.findings)

    if args.markdown:
        md_fn = markdown_bidder if args.role == 'bidder' else markdown_tenderer
        with open(args.markdown, 'w', encoding='utf-8') as fh:
            fh.write(md_fn(items, meta))
        print(f'Markdown -> {args.markdown}')

    docx_render(items, meta, args.out, role=args.role)


if __name__ == '__main__':
    main()
