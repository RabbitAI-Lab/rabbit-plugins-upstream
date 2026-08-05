#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
医疗器械培训活动报告生成脚本 v3.0
根据 JSON 数据生成符合模板格式的 docx 活动报告

v3.0 核心升级:
  - 支持三种会议模式：纯线下(offline)、线上+线下(hybrid)、纯线上(online)
  - 新增活动致辞部分
  - 理论授课支持讲座题目+讲课人+医院+PPT内容总结
  - 新增病例讨论环节
  - 全新总结格式：参会人数+满意度+各环节亮点
  - 纯线上支持根据文字版记录提炼每个环节3-5句话

用法:
    python generate_report.py input.json output.docx
"""

import sys
import os
import json
import glob as glob_mod
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ============================================================
# 格式常量
# ============================================================
FONT_BODY = "宋体"
FONT_HEADING = "微软雅黑"
FONT_SIZE_BODY = Pt(10.5)
FONT_SIZE_TITLE = Pt(18)
FONT_SIZE_SECTION = Pt(12)
FONT_SIZE_SUBSECTION = Pt(11)
FONT_SIZE_FIELD = Pt(10.5)
FONT_SIZE_TABLE_HEAD = Pt(10)

COLOR_GRAY_BG = "D9D9D9"
COLOR_LIGHT_BG = "F2F2F2"
COLOR_TITLE_BG = "4472C4"

IMAGE_EXTENSIONS = (".jpg", ".jpeg", ".png", ".bmp", ".gif")


# ============================================================
# Excel 解析函数 (v2.2 legacy, 保留)
# ============================================================

def parse_excel_scores(excel_path, sheet_name=None):
    try:
        import openpyxl
    except ImportError:
        print("错误：需要安装 openpyxl 库。运行：pip install openpyxl")
        return []

    if not os.path.exists(excel_path):
        print(f"错误：Excel 文件不存在: {excel_path}")
        return []

    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb[sheet_name] if sheet_name else wb.active

    rows = list(ws.iter_rows(min_row=1, values_only=True))
    if not rows:
        print("错误：Excel 工作表为空")
        return []

    header_row = rows[0]
    col_map = {}
    header_texts = [str(c).strip() if c else "" for c in header_row]

    key_map = {
        "seq": ["序号", "编号", "no.", "no"],
        "name": ["姓名", "学员姓名", "名字", "名称"],
        "province": ["省份", "省", "省/直辖市", "地区", "区域"],
        "hospital": ["医院", "单位", "医院名称", "工作单位", "机构"],
        "score": ["得分", "成绩", "分数", "操作得分", "score", "考核得分"],
    }

    for ci, text in enumerate(header_texts):
        text_lower = text.lower().replace(" ", "")
        for field, keywords in key_map.items():
            if field in col_map:
                continue
            for kw in keywords:
                if kw in text_lower or text_lower in kw:
                    col_map[field] = ci
                    break

    required = ["name", "hospital", "score"]
    missing = [r for r in required if r not in col_map]
    if missing:
        print(f"警告：Excel中未识别到以下列: {missing}")
        if "name" in missing or "score" in missing:
            return []

    scores = []
    for ri in range(1, len(rows)):
        row = rows[ri]
        if not row or all(c is None or str(c).strip() == "" for c in row):
            continue

        def get_val(field, default=""):
            if field in col_map:
                v = row[col_map[field]]
                return str(v).strip() if v is not None else default
            return default

        try:
            score_val = get_val("score", "0")
            try:
                score_num = float(score_val)
            except ValueError:
                score_num = 0

            scores.append({
                "seq": len(scores) + 1,
                "name": get_val("name"),
                "province": get_val("province"),
                "hospital": get_val("hospital"),
                "score": score_num,
            })
        except Exception as e:
            print(f"  跳过第{ri + 1}行: {e}")

    wb.close()

    if not scores:
        print("错误：未从 Excel 中解析到任何成绩数据")
        return []

    scores.sort(key=lambda x: x["score"], reverse=True)
    for i, s in enumerate(scores):
        s["seq"] = i + 1

    print(f"从 Excel 解析到 {len(scores)} 条学员成绩（已按得分降序排列）")
    return scores


def parse_excel_quiz(excel_path, sheet_name=None):
    try:
        import openpyxl
    except ImportError:
        print("错误：需要安装 openpyxl 库。运行：pip install openpyxl")
        return {}

    if not os.path.exists(excel_path):
        print(f"错误：Excel 文件不存在: {excel_path}")
        return {}

    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb[sheet_name] if sheet_name else wb.active
    rows = list(ws.iter_rows(min_row=1, values_only=True))
    if not rows:
        return {}

    header_row = rows[0]
    header_texts = [str(c).strip() if c else "" for c in header_row]
    header_texts_lower = [h.lower().replace(" ", "") for h in header_texts]

    result = {}
    data_rows = [r for r in rows[1:] if r and any(c is not None and str(c).strip() != "" for c in r)]

    summary_keywords = {
        "participants": ["参与人数", "答题人数", "quiz人数", "参与学员", "参与数"],
        "avg_score": ["平均分", "quiz平均分", "平均成绩", "均分"],
        "full_score": ["满分", "总分", "quiz满分"],
    }

    for field, keywords in summary_keywords.items():
        for ci, h in enumerate(header_texts_lower):
            for kw in keywords:
                if kw in h:
                    for dr in data_rows:
                        if ci < len(dr) and dr[ci] is not None:
                            val = str(dr[ci]).strip()
                            if val:
                                try:
                                    result[field] = float(val) if "." in val else int(float(val))
                                except ValueError:
                                    result[field] = val
                            break
                    break
            if field in result:
                break

    if "participants" not in result:
        score_keywords = ["quiz得分", "quiz成绩", "答题得分", "quiz", "线上测试得分", "理论得分", "测验得分"]
        name_keywords = ["姓名", "学员姓名", "名字"]
        score_col = None
        name_col = None
        for ci, h in enumerate(header_texts_lower):
            if score_col is None:
                for kw in score_keywords:
                    if kw in h:
                        score_col = ci
                        break
            if name_col is None:
                for kw in name_keywords:
                    if kw in h:
                        name_col = ci
                        break

        if score_col is not None:
            individual_scores = []
            for dr in data_rows:
                if score_col >= len(dr) or dr[score_col] is None:
                    continue
                try:
                    score_val = float(dr[score_col])
                    name_val = str(dr[name_col]).strip() if name_col is not None and name_col < len(dr) and dr[name_col] is not None else ""
                    individual_scores.append({"name": name_val, "score": score_val})
                except (ValueError, TypeError):
                    continue
            if individual_scores:
                result["participants"] = len(individual_scores)
                avg = sum(s["score"] for s in individual_scores) / len(individual_scores)
                result["avg_score"] = round(avg, 1)
                result["individual_scores"] = individual_scores
                if "full_score" not in result:
                    result["full_score"] = 5 if individual_scores and individual_scores[0]["score"] <= 5 else 100

    wb.close()
    if result.get("participants"):
        print(f"从 QUIZ Excel 解析到：{result.get('participants')} 人，平均分 {result.get('avg_score')}，满分 {result.get('full_score')}")
    return result


def parse_excel_feedback(excel_path, sheet_name=None):
    try:
        import openpyxl
    except ImportError:
        print("错误：需要安装 openpyxl 库。运行：pip install openpyxl")
        return {}

    if not os.path.exists(excel_path):
        print(f"错误：Excel 文件不存在: {excel_path}")
        return {}

    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb[sheet_name] if sheet_name else wb.active
    rows = list(ws.iter_rows(min_row=1, values_only=True))
    if not rows:
        return {}

    header_row = rows[0]
    header_texts = [str(c).strip() if c else "" for c in header_row]
    header_texts_lower = [h.lower().replace(" ", "") for h in header_texts]
    result = {}
    data_rows = [r for r in rows[1:] if r and any(c is not None and str(c).strip() != "" for c in r)]

    summary_kw = {
        "respondents": ["反馈人数", "填写人数", "参与反馈", "反馈学员", "回收问卷"],
        "satisfaction": ["满意度评分", "满意度", "平均满意度", "满意度平均分"],
        "full_score": ["满分", "满意度满分"],
    }

    for field, keywords in summary_kw.items():
        for ci, h in enumerate(header_texts_lower):
            for kw in keywords:
                if kw in h:
                    for dr in data_rows:
                        if ci < len(dr) and dr[ci] is not None:
                            val = str(dr[ci]).strip()
                            if val:
                                try:
                                    result[field] = float(val) if "." in val else int(float(val))
                                except ValueError:
                                    result[field] = val
                            break
                    break
            if field in result:
                break

    if "respondents" not in result:
        sat_keywords = ["满意度评分", "满意度", "评分", "打分", "satisfaction"]
        sat_col = None
        for ci, h in enumerate(header_texts_lower):
            if sat_col is None:
                for kw in sat_keywords:
                    if kw in h:
                        sat_col = ci
                        break

        if sat_col is not None:
            satisfaction_scores = []
            for dr in data_rows:
                if sat_col >= len(dr) or dr[sat_col] is None:
                    continue
                try:
                    score_val = float(dr[sat_col])
                    satisfaction_scores.append(score_val)
                except (ValueError, TypeError):
                    continue
            if satisfaction_scores:
                result["respondents"] = len(satisfaction_scores)
                avg = sum(satisfaction_scores) / len(satisfaction_scores)
                result["satisfaction"] = round(avg, 1)
                if "full_score" not in result:
                    result["full_score"] = 5 if satisfaction_scores and satisfaction_scores[0] <= 5 else 10

    wb.close()
    if result.get("respondents"):
        print(f"从反馈 Excel 解析到：{result.get('respondents')} 人反馈，满意度 {result.get('satisfaction')}，满分 {result.get('full_score')}")
    return result


# ============================================================
# 照片文件夹处理
# ============================================================

def find_photos(photo_path):
    photos = []
    if not photo_path or not os.path.exists(photo_path):
        return photos
    if os.path.isfile(photo_path):
        ext = os.path.splitext(photo_path)[1].lower()
        if ext in IMAGE_EXTENSIONS:
            photos = [photo_path]
    elif os.path.isdir(photo_path):
        for ext in IMAGE_EXTENSIONS:
            pattern = os.path.join(photo_path, f"*{ext}")
            photos.extend(glob_mod.glob(pattern))
        photos.sort(key=lambda p: os.path.basename(p).lower())
    if photos:
        print(f"发现 {len(photos)} 张照片")
    return photos


# ============================================================
# 辅助函数
# ============================================================

def set_cell_background(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_font(cell, font_name=FONT_BODY, font_size=FONT_SIZE_BODY, bold=False, color=None):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
            run.font.bold = bold
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            if color:
                run.font.color.rgb = color


def add_formatted_paragraph(container, text, font_name=FONT_BODY, font_size=FONT_SIZE_BODY,
                            bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6)):
    p = container.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    return p


def add_multiline_text(cell, text, font_name=FONT_BODY, font_size=FONT_SIZE_BODY, bold=False):
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.font.name = font_name
        run.font.size = font_size
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
    </w:tblBorders>'''
    tblPr.append(parse_xml(borders_xml))


def set_column_width(table, col_idx, width):
    for row in table.rows:
        cell = row.cells[col_idx]
        cell.width = width


# ============================================================
# v3.0 新增：活动致辞
# ============================================================

def create_activity_opening(container, opening_data):
    """添加活动致辞部分"""
    if not opening_data:
        return

    add_formatted_paragraph(container, "一、活动致辞", FONT_HEADING, FONT_SIZE_SECTION, True, space_after=Pt(8))

    expert = opening_data.get("expert", "")
    host = opening_data.get("host", "")

    # 主持人+致辞专家
    header_parts = []
    if host:
        header_parts.append(f"本次培训由{host}担任主持")
    if expert:
        header_parts.append(f"{expert}莅临致开场辞")

    if header_parts:
        header_text = "，".join(header_parts) + "。"
        add_formatted_paragraph(container, header_text, FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(6))

    # 致辞内容
    speech = opening_data.get("speech", "")
    if speech:
        add_formatted_paragraph(container, speech, FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(8))
    else:
        # 默认致辞
        default_speech = (
            "致辞专家对各位学员的到来表示热烈欢迎，指出消化内镜技术的规范化培训"
            "对于提高基层医疗机构的诊疗水平具有重要意义，希望学员们珍惜学习机会，"
            "通过理论学习和实践操作，切实提升自身技术水平，更好地服务于临床患者。"
        )
        add_formatted_paragraph(container, default_speech, FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(8))


# ============================================================
# 日程表
# ============================================================

def create_schedule_table(container, schedule_data):
    if not schedule_data:
        return

    add_formatted_paragraph(container, "二、活动日程", FONT_HEADING, FONT_SIZE_SECTION, True, space_after=Pt(6))

    table = container.add_table(rows=1 + len(schedule_data), cols=3)
    set_table_borders(table)

    headers = ["时间", "内容", "讲者/专家"]
    for ci, header in enumerate(headers):
        cell = table.rows[0].cells[ci]
        set_cell_background(cell, COLOR_GRAY_BG)
        add_multiline_text(cell, header, FONT_BODY, FONT_SIZE_TABLE_HEAD, True)

    for ri, item in enumerate(schedule_data):
        row = table.rows[1 + ri]
        time_val = item.get("time", "")
        content_val = item.get("content", "")
        speaker_val = item.get("speaker", "")

        if time_val and not content_val and not speaker_val:
            merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
            set_cell_background(merged, COLOR_LIGHT_BG)
            add_multiline_text(merged, time_val, FONT_HEADING, FONT_SIZE_TABLE_HEAD, True)
        else:
            add_multiline_text(row.cells[0], time_val)
            add_multiline_text(row.cells[1], content_val)
            add_multiline_text(row.cells[2], speaker_val)

    set_column_width(table, 0, Cm(3.5))
    set_column_width(table, 1, Cm(9))
    set_column_width(table, 2, Cm(4.5))


# ============================================================
# v3.0 新增：理论授课环节
# ============================================================

def create_lecture_section(container, lectures_data):
    """添加理论授课环节"""
    if not lectures_data:
        return

    # 每个讲座以子标题形式呈现
    for i, lec in enumerate(lectures_data):
        title = lec.get("title", "")
        speaker = lec.get("speaker", "")
        hospital = lec.get("hospital", "")
        content_summary = lec.get("content_summary", "")

        # 标题行：讲座题目 — 讲者（医院）
        title_line = f"《{title}》"
        if speaker or hospital:
            title_line += f" — {speaker}"
            if hospital:
                title_line += f"（{hospital}）"
        add_formatted_paragraph(container, title_line, FONT_BODY, FONT_SIZE_SUBSECTION, True, space_after=Pt(4))

        # 内容摘要
        if content_summary:
            add_formatted_paragraph(container, content_summary, FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(8))
        else:
            add_formatted_paragraph(container, f"{speaker if speaker else '讲师'}就本专题进行了系统讲解。", FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(8))


# ============================================================
# v3.0 新增：操作环节（含成绩表）
# ============================================================

def create_operation_section(container, operations_data):
    """添加模型操作/临床操作环节"""
    if not operations_data:
        return

    for op in operations_data:
        name = op.get("name", "操作培训")
        summary = op.get("summary", "")
        assessment = op.get("assessment", {})
        scores = op.get("scores", [])

        add_formatted_paragraph(container, name, FONT_BODY, FONT_SIZE_SUBSECTION, True, space_after=Pt(4))

        if summary:
            add_formatted_paragraph(container, summary, FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(4))

        # 考核概述
        if assessment:
            assess_parts = [f"共{assessment.get('participants', '')}名学员完成操作考核"]
            if assessment.get("full_score"):
                assess_parts.append(f"满分{assessment['full_score']}分")
            if assessment.get("avg_score"):
                assess_parts.append(f"平均得分{assessment['avg_score']}分")
            add_formatted_paragraph(container, "，".join(assess_parts) + "。", FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(4))

        # 成绩表
        if scores:
            course_name = op.get("course_name", "")
            display_name = course_name if course_name else "操作"
            headers = ["序号", "姓名", "省份", "医院", f"{display_name}得分"]
            table = container.add_table(rows=1 + len(scores), cols=len(headers))
            set_table_borders(table)

            for ci, header in enumerate(headers):
                cell = table.rows[0].cells[ci]
                set_cell_background(cell, COLOR_GRAY_BG)
                add_multiline_text(cell, header, FONT_BODY, FONT_SIZE_TABLE_HEAD, True)

            for ri, s in enumerate(scores):
                row = table.rows[1 + ri]
                score_str = str(s.get("score", ""))
                try:
                    sc = float(score_str)
                    if sc == int(sc):
                        score_str = str(int(sc))
                except ValueError:
                    pass
                values = [
                    str(s.get("seq", ri + 1)),
                    s.get("name", ""),
                    s.get("province", ""),
                    s.get("hospital", ""),
                    score_str,
                ]
                for ci, val in enumerate(values):
                    add_multiline_text(row.cells[ci], val)

            set_column_width(table, 0, Cm(1.2))
            set_column_width(table, 1, Cm(2.5))
            set_column_width(table, 2, Cm(2))
            set_column_width(table, 3, Cm(7))
            set_column_width(table, 4, Cm(3.5))


# ============================================================
# v3.0 新增：病例讨论
# ============================================================

def create_case_discussion_section(container, cases_data):
    """添加病例讨论环节"""
    if not cases_data:
        return

    for cd in cases_data:
        presenter = cd.get("presenter", "")
        hospital = cd.get("hospital", "")
        topic = cd.get("topic", "")
        summary = cd.get("summary", "")

        title_line = f"病例讨论"
        if presenter:
            title_line += f" — {presenter}"
        if hospital:
            title_line += f"（{hospital}）"
        if topic:
            title_line += f"：{topic}"

        add_formatted_paragraph(container, title_line, FONT_BODY, FONT_SIZE_SUBSECTION, True, space_after=Pt(4))

        if summary:
            add_formatted_paragraph(container, summary, FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(8))


# ============================================================
# v3.0 新增：纯线上会议环节总结
# ============================================================

def create_online_sessions(container, sessions_data):
    """添加纯线上会议各环节总结（每环节3-5句话）"""
    if not sessions_data:
        return

    for ses in sessions_data:
        title = ses.get("title", "")
        speaker = ses.get("speaker", "")
        summary = ses.get("summary", "")

        if speaker:
            title_line = f"{title}（{speaker}）"
        else:
            title_line = title

        add_formatted_paragraph(container, title_line, FONT_BODY, FONT_SIZE_SUBSECTION, True, space_after=Pt(4))

        if summary:
            add_formatted_paragraph(container, summary, FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(8))


# ============================================================
# v3.0 新增：全新总结格式
# ============================================================

def create_new_summary(container, summary_data):
    """添加新的活动总结"""
    if not summary_data:
        return

    add_formatted_paragraph(container, "四、总结", FONT_HEADING, FONT_SIZE_SECTION, True, space_after=Pt(8))

    # 参会人数
    total = summary_data.get("total_participants", "")
    if total:
        add_formatted_paragraph(container, f"本次培训共有{total}名学员参会。", FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(4))

    # 满意度
    fb_respondents = summary_data.get("feedback_respondents", "")
    fb_avg = summary_data.get("feedback_avg_score", "")
    if fb_respondents and fb_avg:
        add_formatted_paragraph(container, f"培训后共有{fb_respondents}名学员参与了满意度评价，满意度平均分为{fb_avg}分。", FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(6))

    # 各环节亮点
    highlights = summary_data.get("highlights", {})
    if highlights:
        add_formatted_paragraph(container, "培训亮点：", FONT_BODY, FONT_SIZE_BODY, True, space_after=Pt(4))

        lec_hl = highlights.get("lecture", "")
        if lec_hl:
            add_formatted_paragraph(container, f"授课环节：{lec_hl}", FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(4))

        op_hl = highlights.get("operation", "")
        if op_hl:
            add_formatted_paragraph(container, f"操作环节：{op_hl}", FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(4))

        case_hl = highlights.get("case_discussion", "")
        if case_hl:
            add_formatted_paragraph(container, f"病例讨论：{case_hl}", FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(4))

    # 改善建议（保留）
    improvements = summary_data.get("improvements", "")
    if improvements:
        add_formatted_paragraph(container, f"改进建议：{improvements}", FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(6))


# ============================================================
# v3.0 重构：在线部分（hybrid/online 共用）
# ============================================================

def create_online_part(container, online_data):
    """添加线上培训部分"""
    if not online_data:
        return

    # 线上课程列表
    lectures = online_data.get("lectures", [])
    if lectures:
        add_formatted_paragraph(container, "线上理论学习内容：", FONT_BODY, FONT_SIZE_BODY, True, space_after=Pt(4))
        for lec in lectures:
            title = lec.get("title", "")
            speaker = lec.get("speaker", "")
            line = f"  《{title}》"
            if speaker:
                line += f" — {speaker}"
            add_formatted_paragraph(container, line, FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(2))
        add_formatted_paragraph(container, "", space_after=Pt(4))

    # QUIZ
    quiz = online_data.get("quiz", {})
    if quiz:
        participants = quiz.get("participants", "")
        avg = quiz.get("avg_score", "")
        full = quiz.get("full_score", "")
        quiz_text = f"线上QUIZ答题情况：共{participants}名学员参与，平均分{avg}分（满分{full}分）。"
        add_formatted_paragraph(container, quiz_text, FONT_BODY, FONT_SIZE_BODY, False, space_after=Pt(6))


# ============================================================
# 基本信息表（保留）
# ============================================================

def create_basic_info_table(doc, data):
    basic = data.get("basic_info", {})

    fields = [
        ("会议名称", basic.get("meeting_name", "")),
        ("主办单位", basic.get("organizer", "")),
        ("会议时间", basic.get("meeting_time", "")),
        ("会议地点", basic.get("meeting_location", "")),
        ("参会专家", basic.get("experts", "")),
        ("会议规模", basic.get("scale", "")),
        ("参加员工", basic.get("staff", "")),
        ("项目介绍", basic.get("project_intro", "")),
    ]

    table = doc.add_table(rows=2 + len(fields), cols=2)
    set_table_borders(table)
    set_column_width(table, 0, Cm(4))
    set_column_width(table, 1, Cm(13))

    title_row = table.rows[0]
    title_cell = title_row.cells[0].merge(title_row.cells[1])
    set_cell_background(title_cell, COLOR_TITLE_BG)
    add_multiline_text(title_cell, "活动报告书", FONT_HEADING, FONT_SIZE_TITLE, True)
    for p in title_cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    section_row = table.rows[1]
    section_cell = section_row.cells[0].merge(section_row.cells[1])
    set_cell_background(section_cell, COLOR_GRAY_BG)
    add_multiline_text(section_cell, "会议基本信息", FONT_HEADING, FONT_SIZE_SECTION, True)

    for i, (field_name, field_value) in enumerate(fields):
        row = table.rows[2 + i]
        left_cell = row.cells[0]
        set_cell_background(left_cell, COLOR_LIGHT_BG)
        add_multiline_text(left_cell, field_name, FONT_BODY, FONT_SIZE_FIELD, True)

        right_cell = row.cells[1]
        add_multiline_text(right_cell, field_value, FONT_BODY, FONT_SIZE_BODY, False)

    return table


# ============================================================
# v3.0 核心：培训内容表（按模式分发）
# ============================================================

def create_training_content_table(doc, data):
    """根据 meeting_mode 生成不同结构的培训内容"""
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table)
    cell = table.rows[0].cells[0]
    cell.width = Cm(17)
    cell.paragraphs[0].text = ""

    meeting_mode = data.get("meeting_mode", "hybrid")

    # 1. 活动致辞
    opening = data.get("activity_opening", {})
    create_activity_opening(cell, opening)

    # 2. 活动日程
    schedule = data.get("schedule", [])
    create_schedule_table(cell, schedule)

    # 3. 会议详细内容
    add_formatted_paragraph(cell, "三、会议详细内容", FONT_HEADING, FONT_SIZE_SECTION, True, space_after=Pt(8))

    detail = data.get("detailed_content", {})

    if meeting_mode == "offline":
        # === 纯线下 ===
        lectures = detail.get("lectures", [])
        if lectures:
            add_formatted_paragraph(cell, "（一）理论授课", FONT_HEADING, FONT_SIZE_SUBSECTION, True, space_after=Pt(6))
            create_lecture_section(cell, lectures)

        operations = detail.get("operations", [])
        if operations:
            add_formatted_paragraph(cell, "（二）模型操作/临床操作", FONT_HEADING, FONT_SIZE_SUBSECTION, True, space_after=Pt(6))
            create_operation_section(cell, operations)

        cases = detail.get("case_discussions", [])
        if cases:
            add_formatted_paragraph(cell, "（三）病例讨论", FONT_HEADING, FONT_SIZE_SUBSECTION, True, space_after=Pt(6))
            create_case_discussion_section(cell, cases)

    elif meeting_mode == "hybrid":
        # === 线上 + 线下 ===
        online = detail.get("online", {})
        if online:
            add_formatted_paragraph(cell, "（一）线上理论学习", FONT_HEADING, FONT_SIZE_SUBSECTION, True, space_after=Pt(6))
            create_online_part(cell, online)

        add_formatted_paragraph(cell, "（二）线下培训活动", FONT_HEADING, FONT_SIZE_SUBSECTION, True, space_after=Pt(6))

        lectures = detail.get("lectures", [])
        if lectures:
            add_formatted_paragraph(cell, "1. 理论授课", FONT_BODY, FONT_SIZE_SUBSECTION, True, space_after=Pt(4))
            create_lecture_section(cell, lectures)

        operations = detail.get("operations", [])
        if operations:
            add_formatted_paragraph(cell, "2. 模型操作/临床操作", FONT_BODY, FONT_SIZE_SUBSECTION, True, space_after=Pt(4))
            create_operation_section(cell, operations)

        cases = detail.get("case_discussions", [])
        if cases:
            add_formatted_paragraph(cell, "3. 病例讨论", FONT_BODY, FONT_SIZE_SUBSECTION, True, space_after=Pt(4))
            create_case_discussion_section(cell, cases)

    elif meeting_mode == "online":
        # === 纯线上 ===
        sessions = detail.get("online_sessions", [])
        if sessions:
            create_online_sessions(cell, sessions)

    # 4. 总结
    summary = data.get("summary", {})
    create_new_summary(cell, summary)

    # 5. 合影
    group_photo_path = data.get("group_photo_path", "")
    group_photo_folder = data.get("group_photo_folder", "")
    photos = []
    if group_photo_path:
        photos = find_photos(group_photo_path)
    if group_photo_folder:
        photos.extend(find_photos(group_photo_folder))
    seen = set()
    unique_photos = []
    for p in photos:
        ap = os.path.abspath(p)
        if ap not in seen:
            seen.add(ap)
            unique_photos.append(p)

    if unique_photos:
        add_formatted_paragraph(cell, "", space_after=Pt(6))
        for photo_path in unique_photos:
            try:
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(photo_path, width=Cm(14))
                add_formatted_paragraph(cell, "", FONT_BODY, Pt(4), False, WD_ALIGN_PARAGRAPH.CENTER, Pt(2))
            except Exception as e:
                print(f"  警告：无法嵌入照片 '{photo_path}': {e}")
        add_formatted_paragraph(cell, "合影留念", FONT_BODY, Pt(9), False, WD_ALIGN_PARAGRAPH.CENTER, Pt(6))

    # 6. 落款
    signature = data.get("signature", "[培训部门]")
    add_formatted_paragraph(cell, signature, FONT_BODY, Pt(10), False, WD_ALIGN_PARAGRAPH.RIGHT, Pt(6))

    return table


# ============================================================
# 数据预处理器
# ============================================================

def preprocess_data(data):
    """预处理数据：Excel解析、成绩排序、平均分计算"""
    detail = data.get("detailed_content", {})

    # 处理操作环节的 Excel 和成绩
    operations = detail.get("operations", [])
    for op in operations:
        excel_path = op.get("excel_path", "")
        excel_sheet = op.get("excel_sheet", None)
        if excel_path and os.path.exists(excel_path):
            scores = parse_excel_scores(excel_path, excel_sheet)
            if scores:
                op["scores"] = scores

        scores = op.get("scores", [])
        if scores:
            scores.sort(key=lambda x: x.get("score", 0), reverse=True)
            for i, s in enumerate(scores):
                s["seq"] = i + 1

        assessment = op.get("assessment", {})
        if scores and not assessment.get("avg_score"):
            avg = sum(s.get("score", 0) for s in scores) / len(scores)
            assessment["avg_score"] = f"{avg:.1f}"
        if scores and not assessment.get("participants"):
            assessment["participants"] = str(len(scores))

    # 处理线上部分的 QUIZ Excel
    online = detail.get("online", {})
    if online:
        quiz = online.get("quiz", {})
        quiz_excel = quiz.get("excel_path", "")
        if quiz_excel and os.path.exists(quiz_excel):
            quiz_data = parse_excel_quiz(quiz_excel, quiz.get("excel_sheet", None))
            if quiz_data:
                if not quiz.get("participants"):
                    quiz["participants"] = str(quiz_data.get("participants", ""))
                if not quiz.get("avg_score"):
                    quiz["avg_score"] = str(quiz_data.get("avg_score", ""))
                if not quiz.get("full_score"):
                    quiz["full_score"] = str(quiz_data.get("full_score", ""))

    return data


# ============================================================
# 主函数
# ============================================================

def generate_report_from_dict(data, output_path):
    data = preprocess_data(data)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Emu(7560310)
    section.page_height = Emu(10692130)
    section.top_margin = Emu(900430)
    section.bottom_margin = Emu(900430)
    section.left_margin = Emu(540385)
    section.right_margin = Emu(540385)

    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = FONT_SIZE_BODY
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    create_basic_info_table(doc, data)
    doc.add_paragraph()
    create_training_content_table(doc, data)

    doc.save(output_path)
    print(f"报告已生成: {output_path}")


def generate_report(json_path, output_path):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    generate_report_from_dict(data, output_path)


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="生成医疗器械培训活动报告 v3.0")
    parser.add_argument("json_file", nargs="?", help="JSON 数据文件")
    parser.add_argument("output", nargs="?", help="输出 docx 文件路径")
    args = parser.parse_args()
    if not args.json_file or not args.output:
        parser.print_help()
        sys.exit(1)
    generate_report(args.json_file, args.output)
