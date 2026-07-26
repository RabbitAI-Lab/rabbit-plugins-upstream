import docx
import json
import sys

doc = docx.Document('tests/design.docx')

# Extract paragraphs
print("=== PARAGRAPHS ===")
for p in doc.paragraphs:
    if p.text.strip():
        print(p.text)

# Extract tables
print("\n=== TABLES ===")
for i, table in enumerate(doc.tables):
    print(f"\n--- Table {i+1} ---")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cells))
