#!/usr/bin/env python3
"""
Word 文档文本提取工具 — 供 SKILL.md 阶段一/二使用。
从 .docx 文件中提取所有文本（含表格），输出结构化 Markdown 供 AI 分析。

用法:
    python scripts/extract_docx.py <path/to/file.docx> [-o output.txt]
    python scripts/extract_docx.py <path/to/file.docx>  # 输出到 stdout

依赖: python-docx (pip install python-docx)
"""
import sys, argparse
from pathlib import Path


def extract_docx(path: str) -> str:
    """使用 python-docx 提取文本和表格"""
    from docx import Document
    doc = Document(path)
    parts = []

    # 提取段落
    para_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else ""
            if "Heading" in style:
                level = style.replace("Heading ", "").replace("heading ", "")
                try:
                    parts.append(f"{'#' * int(level)} {text}")
                except ValueError:
                    parts.append(f"**{text}**")
            else:
                parts.append(text)
            para_count += 1

    # 提取表格
    table_count = 0
    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)

        if not rows_data:
            continue

        # 构建 Markdown 表格
        table_lines = ["", f"### 表格 {table_count + 1}"]
        table_lines.append("| " + " | ".join(rows_data[0]) + " |")
        table_lines.append("| " + " | ".join("---" for _ in rows_data[0]) + " |")
        for row in rows_data[1:]:
            table_lines.append("| " + " | ".join(row) + " |")
        table_lines.append("")
        parts.extend(table_lines)
        table_count += 1

    header = (
        f"# 提取自: {Path(path).name}\n"
        f"# 段落数: {para_count}, 表格数: {table_count}\n"
        f"# 总字符: {sum(len(p) for p in parts)}\n"
    )
    return header + "\n".join(parts)


def main():
    parser = argparse.ArgumentParser(description="从 Word 文档提取文本和表格")
    parser.add_argument("input", help="Word .docx 文件路径")
    parser.add_argument("-o", "--output", help="输出文件路径（默认输出到 stdout）")
    args = parser.parse_args()

    src = Path(args.input)
    if not src.exists():
        print(f"错误: 文件不存在 - {src}", file=sys.stderr)
        sys.exit(1)
    if src.suffix.lower() not in (".docx",):
        print(f"警告: 输入文件不是 .docx 格式: {src}", file=sys.stderr)

    try:
        text = extract_docx(str(src))
    except ImportError:
        print("错误: 需要安装 python-docx (pip install python-docx)", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"错误: 提取失败 - {e}", file=sys.stderr)
        sys.exit(1)

    if args.output:
        dst = Path(args.output)
        dst.parent.mkdir(parents=True, exist_ok=True)
        dst.write_text(text, encoding="utf-8")
        print(f"文本已提取到: {dst}", file=sys.stdout)
    else:
        sys.stdout.write(text)


if __name__ == "__main__":
    main()
