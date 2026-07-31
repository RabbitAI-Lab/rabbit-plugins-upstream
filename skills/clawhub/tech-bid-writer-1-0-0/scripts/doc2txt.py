#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
doc2txt.py — 招标文件格式转换工具（技术标写作专家配套）

支持格式：
  .docx  → python-docx 段落/表格提取 → 同名 _doc2txt.txt
  .doc   → olefile + UTF-16LE 文本块扫描（尽力而为）→ 同名 _doc2txt.txt
  .pdf   → pymupdf 文本层提取 → 同名 _doc2txt.txt
           扫描版（无文字层）PDF 返回提示，不强行解析
  .xlsx  → openpyxl 逐工作表转 Markdown 表格 → 同名 _doc2txt.md

用法：
  python doc2txt.py <file1> [file2 ...] [-o OUTPUT_DIR]

依赖（纯 pip，零系统依赖）：
  pip install python-docx olefile pymupdf openpyxl
"""

import sys
import os
import re
import argparse


def convert_docx(path):
    import docx
    document = docx.Document(path)
    lines = []
    for p in document.paragraphs:
        if p.text.strip():
            lines.append(p.text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def convert_doc(path):
    """尽力而为：扫描 WordDocument 流中的 UTF-16LE 可视文本。"""
    import olefile
    if not olefile.isOleFile(path):
        raise ValueError("不是有效的 .doc OLE 文件: %s" % path)
    with olefile.OleFileIO(path) as ole:
        if not ole.exists("WordDocument"):
            raise ValueError(".doc 中未找到 WordDocument 流")
        data = ole.openstream("WordDocument").read()
    text = data.decode("utf-16-le", errors="ignore")
    # 提取连续可打印字符（含 CJK），按换行/控制符断行
    lines = []
    buf = []
    for ch in text:
        o = ord(ch)
        if ch in "\r\n":
            if buf:
                lines.append("".join(buf))
                buf = []
        elif o < 32 and ch not in "\t":
            if buf:
                lines.append("".join(buf))
                buf = []
        else:
            buf.append(ch)
    if buf:
        lines.append("".join(buf))
    # 去掉仅含空白/孤立标点的噪点行
    cleaned = [ln.strip() for ln in lines if ln.strip()]
    return "\n".join(cleaned)


def convert_pdf(path):
    """返回 (text, is_scanned)。is_scanned=True 表示无文字层。"""
    import fitz  # pymupdf
    doc = fitz.open(path)
    pages = []
    has_text = False
    for page in doc:
        t = page.get_text()
        if t.strip():
            has_text = True
        pages.append(t)
    doc.close()
    full = "\n".join(pages)
    return full, (not has_text)


def convert_xlsx(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append("# %s" % ws.title)
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            out.append("| " + " | ".join(cells) + " |")
        out.append("")
    return "\n".join(out)


def process_file(path, out_dir):
    ext = os.path.splitext(path)[1].lower()
    stem = os.path.splitext(os.path.basename(path))[0]
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
        base = out_dir
    else:
        base = os.path.dirname(os.path.abspath(path))

    if ext == ".docx":
        text = convert_docx(path)
        out_path = os.path.join(base, "%s_doc2txt.txt" % stem)
        _write(out_path, text)
        return "OK  -> %s" % out_path

    if ext == ".doc":
        text = convert_doc(path)
        out_path = os.path.join(base, "%s_doc2txt.txt" % stem)
        _write(out_path, text)
        return "OK  -> %s" % out_path

    if ext == ".pdf":
        text, scanned = convert_pdf(path)
        if scanned:
            return ("SKIP -> %s 是扫描版 PDF（无文字层），请先用 OCR 识别后再转换。" % path)
        out_path = os.path.join(base, "%s_doc2txt.txt" % stem)
        _write(out_path, text)
        return "OK  -> %s" % out_path

    if ext == ".xlsx":
        text = convert_xlsx(path)
        out_path = os.path.join(base, "%s_doc2txt.md" % stem)
        _write(out_path, text)
        return "OK  -> %s" % out_path

    return "SKIP -> 不支持的格式: %s" % path


def _write(path, text):
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)


def main(argv=None):
    parser = argparse.ArgumentParser(description="招标文件格式转换工具")
    parser.add_argument("files", nargs="+", help="输入文件（.doc/.docx/.pdf/.xlsx）")
    parser.add_argument("-o", "--out", default=None, help="输出目录（默认源文件同目录）")
    args = parser.parse_args(argv)

    print("=" * 60)
    print("doc2txt — 招标文件格式转换")
    print("=" * 60)
    for f in args.files:
        if not os.path.exists(f):
            print("MISS -> 文件不存在: %s" % f)
            continue
        try:
            print(process_file(f, args.out))
        except Exception as e:
            print("ERR  -> %s : %s" % (f, e))
    print("=" * 60)
    print("完成。")


if __name__ == "__main__":
    main()
