#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
convert_to_word.py — Markdown → Word 转换器（技术标写作专家配套）

读取 Markdown 顶部的 `<!-- doc-format -->` 注释块，输出合规版式 Word 文档。
适用于技术标交付：仿宋、首行缩进、三线表、标题层级、行距与页边距可控。

doc-format 注释块示例（放在 Markdown 最前面）：
  <!-- doc-format
  font: 仿宋
  body-size: 12pt
  title-level: 16pt
  sub-level: 15pt
  line-spacing: 1.5
  margins: 3.7,3.5,2.8,2.6
  first-line-indent: 0.85cm
  -->
  （说明：正文仿宋小四；一级标题仿宋三号=16pt，其他标题仿宋小三=15pt；
    line-spacing 为倍数时写纯数字如 1.5，为固定磅值写如 18pt；
    margins 可写单值 2.5 或 上,下,左,右 四值。）

用法：
  python convert_to_word.py <input.md> [--out OUTPUT_DIR] [--name NAME.docx]

依赖：pip install python-docx

说明：
  - 仅转换内容，不硬编码任何输出路径；输出目录由 --out 或输入文件同目录决定。
  - Mermaid/代码块以等宽文本保留（图表需外部渲染，本工具不联网）。
  - 图片（![]()）跳过。
  - 表格输出为带边框表格，表头浅底。
"""

import sys
import os
import re
import argparse

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- 默认版式参数 ----------
DEFAULTS = {
    "font": "仿宋",
    "body_size": 12.0,            # 正文 小四
    "title_level": 16.0,         # 一级标题 三号
    "sub_level": 15.0,           # 其他标题 小三
    "line_spacing": 1.5,         # 行距（倍数）
    "line_spacing_mode": "mul",  # mul=倍数 / pt=固定磅值
    "margins_top": 3.7,          # 上页边距 cm（公文标准）
    "margins_bottom": 3.5,       # 下
    "margins_left": 2.8,         # 左
    "margins_right": 2.6,        # 右
    "first_line_indent": 0.85,   # 首行缩进 cm ≈ 2 字符（小四）
}


def _parse_doc_format(md_text):
    """提取并解析 `<!-- doc-format ... -->` 块，返回参数 dict（覆盖默认值）。"""
    params = dict(DEFAULTS)
    m = re.search(r"<!--\s*doc-format(.*?)-->", md_text, re.DOTALL)
    if not m:
        return params
    block = m.group(1)
    for line in block.splitlines():
        line = line.strip()
        if not line or ":" not in line:
            continue
        key, _, val = line.partition(":")
        key = key.strip().lower()
        val = val.strip()
        if key in ("font",):
            params["font"] = val
        elif key in ("body-size", "body_size"):
            params["body_size"] = _to_pt(val)
        elif key in ("title-level", "title_level"):
            params["title_level"] = _to_pt(val)
        elif key in ("sub-level", "sub_level"):
            params["sub_level"] = _to_pt(val)
        elif key in ("line-spacing", "line_spacing"):
            if val.lower().endswith("pt"):
                params["line_spacing"] = _to_pt(val)
                params["line_spacing_mode"] = "pt"
            else:
                try:
                    params["line_spacing"] = float(val)
                    params["line_spacing_mode"] = "mul"
                except ValueError:
                    pass
        elif key in ("margins",):
            if "," in val:
                parts = [float(x) for x in val.split(",") if x.strip() != ""]
                if len(parts) == 4:
                    params["margins_top"], params["margins_bottom"], params["margins_left"], params["margins_right"] = parts
                elif len(parts) == 1:
                    v = parts[0]
                    params["margins_top"] = params["margins_bottom"] = params["margins_left"] = params["margins_right"] = v
            else:
                v = _to_cm(val)
                params["margins_top"] = params["margins_bottom"] = params["margins_left"] = params["margins_right"] = v
        elif key in ("first-line-indent", "first_line_indent"):
            params["first_line_indent"] = _to_cm(val)
    return params


def _to_pt(val):
    val = val.lower().replace("pt", "").strip()
    try:
        return float(val)
    except ValueError:
        return DEFAULTS.get("body_size", 12.0)


def _to_cm(val):
    val = val.lower().replace("cm", "").strip()
    try:
        return float(val)
    except ValueError:
        return 2.0


def _set_run_font(run, font, size=None, bold=False):
    run.font.name = font
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)


def _shade_cell(cell, fill="D9E2F3"):
    """给单元格加底纹（表头用）。"""
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def _strip_doc_format(md_text):
    return re.sub(r"<!--\s*doc-format.*?-->", "", md_text, flags=re.DOTALL).strip()


def _add_inline_runs(paragraph, text, params, base_size=None, base_bold=False):
    """处理 **bold** 行内格式，写入段落。"""
    size = base_size if base_size is not None else params["body_size"]
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            _set_run_font(r, params["font"], size, bold=True)
        else:
            r = paragraph.add_run(part)
            _set_run_font(r, params["font"], size, bold=base_bold)


def _is_table_row(line):
    line = line.strip()
    return line.startswith("|") and line.endswith("|") and line.count("|") >= 2


def _parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def _is_separator_row(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def build_doc(md_text, params):
    doc = Document()

    # 默认字体（中英）
    normal = doc.styles["Normal"]
    normal.font.name = params["font"]
    normal.font.size = Pt(params["body_size"])
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), params["font"])

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(params["margins_top"])
        section.bottom_margin = Cm(params["margins_bottom"])
        section.left_margin = Cm(params["margins_left"])
        section.right_margin = Cm(params["margins_right"])

    lines = md_text.splitlines()
    i = 0
    n = len(lines)

    while i < n:
        raw = lines[i]
        line = raw.rstrip()

        # 空行
        if not line.strip():
            i += 1
            continue

        # 代码块 ```
        if line.strip().startswith("```"):
            lang = line.strip().strip("`").strip()
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结尾 ```
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.left_indent = Cm(0.5)
            label = "（图表-Mermaid，需外部渲染）" if lang.lower() == "mermaid" else "（代码块）"
            r0 = p.add_run(label + "\n" + "\n".join(code_lines))
            _set_run_font(r0, "Consolas", params["body_size"], bold=False)
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            p = doc.add_paragraph()
            if level == 1:
                size = params["title_level"]   # 一级标题 三号
            else:
                size = params["sub_level"]     # 其他标题 小三
            _add_inline_runs(p, text, params, base_size=size, base_bold=True)
            i += 1
            continue

        # 表格
        if _is_table_row(line):
            # 收集连续表格行
            tbl_lines = []
            while i < n and _is_table_row(lines[i]):
                tbl_lines.append(lines[i])
                i += 1
            # 过滤掉分隔行
            data_rows = []
            for tl in tbl_lines:
                if _is_separator_row(tl):
                    continue
                data_rows.append(_parse_table_row(tl))
            if not data_rows:
                continue
            cols = max(len(r) for r in data_rows)
            table = doc.add_table(rows=len(data_rows), cols=cols)
            table.style = "Table Grid"
            for ri, row in enumerate(data_rows):
                for ci in range(cols):
                    cell_text = row[ci] if ci < len(row) else ""
                    cell = table.cell(ri, ci)
                    cell.text = ""
                    cp = cell.paragraphs[0]
                    _add_inline_runs(cp, cell_text, params)
                    if ri == 0:
                        _shade_cell(cell)
            # 表格后空一行
            doc.add_paragraph()
            continue

        # 图片（跳过）
        if line.strip().startswith("!["):
            i += 1
            continue

        # 列表（- 或 1. ）
        m = re.match(r"^(\s*[-*]\s+)(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, m.group(2).strip(), params)
            i += 1
            continue
        m = re.match(r"^(\s*\d+\.\s+)(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, m.group(2).strip(), params)
            i += 1
            continue

        # 引用 >
        if line.strip().startswith(">"):
            stripped = line.strip()[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            _add_inline_runs(p, stripped, params)
            i += 1
            continue

        # 普通段落（可能跨多行直至空行）
        para_lines = [line.strip()]
        i += 1
        while i < n and lines[i].strip() and not lines[i].strip().startswith(("#", "|", ">", "```", "![")) \
                and not re.match(r"^\s*[-*]\s+", lines[i]) and not re.match(r"^\s*\d+\.\s+", lines[i]):
            para_lines.append(lines[i].strip())
            i += 1
        text = " ".join(para_lines)
        p = doc.add_paragraph()
        pf = p.paragraph_format
        if params.get("line_spacing_mode", "pt") == "pt":
            pf.line_spacing = Pt(params["line_spacing"])
        else:
            pf.line_spacing = params["line_spacing"]   # 倍数（如 1.5）
        pf.first_line_indent = Cm(params["first_line_indent"])
        _add_inline_runs(p, text, params)

    return doc


def main(argv=None):
    parser = argparse.ArgumentParser(description="Markdown → Word 转换器")
    parser.add_argument("md", help="输入 Markdown 文件")
    parser.add_argument("--out", default=None, help="输出目录（默认与输入同目录）")
    parser.add_argument("--name", default=None, help="输出文件名（默认 <输入名>.docx）")
    args = parser.parse_args(argv)

    if not os.path.exists(args.md):
        print("ERR -> 文件不存在: %s" % args.md)
        return 1

    with open(args.md, "r", encoding="utf-8") as f:
        md_text = f.read()

    params = _parse_doc_format(md_text)
    md_body = _strip_doc_format(md_text)
    doc = build_doc(md_body, params)

    out_dir = args.out if args.out else os.path.dirname(os.path.abspath(args.md))
    os.makedirs(out_dir, exist_ok=True)
    stem = os.path.splitext(os.path.basename(args.md))[0]
    out_name = args.name if args.name else ("%s.docx" % stem)
    if not out_name.lower().endswith(".docx"):
        out_name += ".docx"
    out_path = os.path.join(out_dir, out_name)
    doc.save(out_path)

    print("=" * 60)
    print("convert_to_word — Markdown → Word")
    print("=" * 60)
    print("输入 : %s" % args.md)
    print("输出 : %s" % out_path)
    ls = ("%.2f倍" % params["line_spacing"]) if params.get("line_spacing_mode", "pt") != "pt" else ("%.0fpt" % params["line_spacing"])
    print("版式 : 字体=%s 正文=%.0fpt 行距=%s 页边距=上%.1f/下%.1f/左%.1f/右%.1fcm 首行缩进=%.2fcm"
          % (params["font"], params["body_size"], ls,
             params["margins_top"], params["margins_bottom"], params["margins_left"], params["margins_right"],
             params["first_line_indent"]))
    print("=" * 60)
    print("完成。")
    return 0


if __name__ == "__main__":
    sys.exit(main())
