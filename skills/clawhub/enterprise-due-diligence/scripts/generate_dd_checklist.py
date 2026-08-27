#!/usr/bin/env python3
"""
企业尽调 - 检查清单生成模块
生成 Word 格式的尽职调查检查清单（30项，四色状态标记）。

使用方式:
    python generate_dd_checklist.py --data checklist_data.json --company "企业名称" --output "检查清单.docx"
"""

import json
import os
import sys
import argparse
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("请先安装 python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ============================================================
# 颜色常量
# ============================================================
HEADING_COLOR = RGBColor(55, 96, 146)  # 标题颜色 #376092（匹配模板）
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(255, 255, 255)

DARK_BG = "0F3460"     # 表头深色背景
GREEN_BG = "C6EFCE"    # 已确认 绿色
YELLOW_BG = "FFEB9C"   # 需关注 黄色
RED_BG = "FFC7CE"      # 高风险 红色
GRAY_BG = "D9D9D9"    # 待核实 灰色
ZEBRA_BG = "F2F2F2"    # 隔行浅色

STATUS_MAP = {
    "confirmed": ("✅ 已确认", GREEN_BG, "已确认"),
    "watch": ("⚠️ 需关注", YELLOW_BG, "需关注"),
    "high_risk": ("❌ 高风险", RED_BG, "高风险"),
    "pending": ("⬜ 待核实", GRAY_BG, "待核实"),
}

STATUS_NAMES = ["confirmed", "watch", "high_risk", "pending"]


# ============================================================
# 默认检查清单模板
# ============================================================

LEGAL_ITEMS = [
    {"id": "1.1", "item": "公司注册信息完整准确", "source": "公开数据"},
    {"id": "1.2", "item": "注册资本实缴情况", "source": "公开数据"},
    {"id": "1.3", "item": "股权结构清晰无争议", "source": "公开数据"},
    {"id": "1.4", "item": "实际控制人明确", "source": "公开数据"},
    {"id": "1.5", "item": "无重大未决诉讼", "source": "公开数据"},
    {"id": "1.6", "item": "无被执行信息", "source": "公开数据"},
    {"id": "1.7", "item": "无失信/限高记录", "source": "公开数据"},
    {"id": "1.8", "item": "核心专利状态有效", "source": "公开数据"},
    {"id": "1.9", "item": "商标注册完整", "source": "公开数据"},
    {"id": "1.10", "item": "经营许可/资质齐全", "source": "公开数据/用户提供"},
    {"id": "1.11", "item": "无重大行政处罚", "source": "公开数据"},
    {"id": "1.12", "item": "无知识产权纠纷", "source": "公开数据"},
]

FINANCIAL_ITEMS = [
    {"id": "2.1", "item": "财务报表完整可用", "source": "用户提供"},
    {"id": "2.2", "item": "资产负债率合理", "source": "用户提供"},
    {"id": "2.3", "item": "流动/速动比率合理", "source": "用户提供"},
    {"id": "2.4", "item": "盈利能力正常", "source": "用户提供"},
    {"id": "2.5", "item": "现金流状况健康", "source": "用户提供"},
    {"id": "2.6", "item": "税务合规无重大欠税", "source": "公开数据/用户提供"},
    {"id": "2.7", "item": "无异常关联交易", "source": "用户提供"},
    {"id": "2.8", "item": "应收账款周转正常", "source": "用户提供"},
    {"id": "2.9", "item": "存货周转正常", "source": "用户提供"},
    {"id": "2.10", "item": "无重大表外负债", "source": "用户提供"},
]

BUSINESS_ITEMS = [
    {"id": "3.1", "item": "核心产品/服务清晰", "source": "公开数据/用户提供"},
    {"id": "3.2", "item": "商业模式可持续", "source": "公开数据/用户提供"},
    {"id": "3.3", "item": "核心团队完整", "source": "公开数据"},
    {"id": "3.4", "item": "市场地位明确", "source": "公开数据"},
    {"id": "3.5", "item": "客户集中度可控", "source": "用户提供"},
    {"id": "3.6", "item": "供应商集中度可控", "source": "用户提供"},
    {"id": "3.7", "item": "行业政策风险可控", "source": "公开数据"},
    {"id": "3.8", "item": "无重大业务依赖风险", "source": "用户提供"},
]


# ============================================================
# 辅助函数
# ============================================================

def _set_run_font(run, size=10, bold=False, color=BLACK, font_name="微软雅黑"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def _set_cell_shading(cell, color_hex):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text, size=10, bold=False, color=BLACK, align="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(str(text) if text else "")
    _set_run_font(run, size=size, bold=bold, color=color)


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_run_font(run, color=HEADING_COLOR)
    return h


def _add_body(doc, text, bold=False, color=BLACK):
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    _set_run_font(run, size=10, bold=bold, color=color)
    return p


# ============================================================
# 检查清单生成
# ============================================================

def _build_checklist_table(doc, items, title, start_id="1"):
    """
    生成某个维度的检查清单表格。
    items: [{"id": "1.1", "item": "...", "source": "...", "status": "...", "note": "..."}]
    """
    if not items:
        return

    _add_heading(doc, title, level=2)

    table = doc.add_table(rows=1 + len(items), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    headers = ["序号", "检查项", "来源", "状态", "备注"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, h, size=10, bold=True, color=WHITE)
        _set_cell_shading(cell, DARK_BG)

    # 数据行
    for r_idx, item in enumerate(items):
        status = item.get("status", "pending")
        status_label, status_color, _ = STATUS_MAP.get(status, STATUS_MAP["pending"])

        row_data = [item.get("id", f"{start_id}.{r_idx+1}"),
                    item.get("item", ""),
                    item.get("source", ""),
                    status_label,
                    item.get("note", "")]

        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            _set_cell_text(cell, val, size=9, color=BLACK, align="left" if c_idx == 1 else "center")

            # 设置状态列底色
            if c_idx == 3:
                _set_cell_shading(cell, status_color)
            elif r_idx % 2 == 1:
                _set_cell_shading(cell, ZEBRA_BG)

        # 高风险行整体标红底色
        if status == "high_risk":
            for c_idx in range(5):
                _set_cell_shading(table.rows[r_idx + 1].cells[c_idx], RED_BG)
        elif status == "watch":
            for c_idx in range(5):
                _set_cell_shading(table.rows[r_idx + 1].cells[c_idx], YELLOW_BG)

    doc.add_paragraph()


def _build_summary_table(doc, all_items):
    """生成风险汇总表。"""
    _add_heading(doc, "风险汇总", level=2)

    total = len(all_items)
    counts = {"confirmed": 0, "watch": 0, "high_risk": 0, "pending": 0}
    status_labels = {"confirmed": "✅ 已确认", "watch": "⚠️ 需关注",
                     "high_risk": "❌ 高风险", "pending": "⬜ 待核实"}

    for item in all_items:
        s = item.get("status", "pending")
        if s in counts:
            counts[s] += 1

    high_risk_items = [item.get("item", "") for item in all_items if item.get("status") == "high_risk"]
    watch_items = [item.get("item", "") for item in all_items if item.get("status") == "watch"]

    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for i, h in enumerate(["风险级别", "数量", "涉及检查项"]):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, h, size=10, bold=True, color=WHITE)
        _set_cell_shading(cell, DARK_BG)

    # 汇总行
    summary_rows = [
        ("✅ 已确认", counts["confirmed"], ""),
        ("⚠️ 需关注", counts["watch"], "；".join(watch_items[:3])),
        ("❌ 高风险", counts["high_risk"], "；".join(high_risk_items[:3])),
        ("⬜ 待核实", counts["pending"], ""),
    ]

    for r_idx, (label, count, items_str) in enumerate(summary_rows):
        _set_cell_text(table.rows[r_idx + 1].cells[0], label, size=10, color=BLACK)
        _set_cell_text(table.rows[r_idx + 1].cells[1], str(count), size=10, color=BLACK)
        _set_cell_text(table.rows[r_idx + 1].cells[2], items_str, size=9, color=BLACK,
                       align="left")

        # 行底色
        bg = {"confirmed": GREEN_BG, "watch": YELLOW_BG,
              "high_risk": RED_BG, "pending": GRAY_BG}
        for c_idx in range(3):
            _set_cell_shading(table.rows[r_idx + 1].cells[c_idx],
                              bg[list(bg.keys())[r_idx]])

    doc.add_paragraph()

    # 总数
    _add_body(doc, f"合计：{total} 项检查项", bold=True, color=BLACK)


def generate_dd_checklist(checklist_data, company_name, output_path):
    """
    生成检查清单 Word 文档。
    """
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{company_name} — 尽职调查检查清单")
    _set_run_font(run, size=16, bold=True, color=BLACK)

    # 报告信息
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run(f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
    _set_run_font(info_run, size=10, color=BLACK)

    doc.add_paragraph()

    # 状态说明
    _add_heading(doc, "状态说明", level=2)
    legend_text = (
        "✅ 已确认 — 数据核实无误，无异常\n"
        "⚠️ 需关注 — 存在一定风险或异常，建议进一步核查\n"
        "❌ 高风险 — 存在重大风险，可能影响交易决策\n"
        "⬜ 待核实 — 数据不可得或未验证，需要补充信息"
    )
    _add_body(doc, legend_text, color=BLACK)

    doc.add_paragraph()

    # 从数据中获取各维度清单，合并默认+用户数据
    legal_items = _merge_items(LEGAL_ITEMS, checklist_data.get("legal", []))
    financial_items = _merge_items(FINANCIAL_ITEMS, checklist_data.get("financial", []))
    business_items = _merge_items(BUSINESS_ITEMS, checklist_data.get("business", []))
    all_items = legal_items + financial_items + business_items

    # 生成各维度表格
    _build_checklist_table(doc, legal_items, "法律维度检查项")
    _build_checklist_table(doc, financial_items, "财务维度检查项")
    _build_checklist_table(doc, business_items, "业务维度检查项")

    # 风险汇总
    _build_summary_table(doc, all_items)

    doc.save(output_path)
    return output_path


def _merge_items(defaults, overrides):
    """
    合并默认清单和用户数据中的检查项。
    overrides 中的项覆盖 defaults 中同 id 的项。
    """
    override_map = {item.get("id"): item for item in overrides}
    merged = []
    for default in defaults:
        item = dict(default)
        override = override_map.get(default["id"])
        if override:
            if "status" in override:
                item["status"] = override["status"]
            if "note" in override:
                item["note"] = override["note"]
        else:
            item.setdefault("status", "pending")
            item.setdefault("note", "")
        merged.append(item)
    return merged


def main():
    parser = argparse.ArgumentParser(description="企业尽调 - 检查清单生成")
    parser.add_argument("--data", "-d", required=True, help="检查清单数据 JSON 文件路径")
    parser.add_argument("--company", "-c", required=True, help="企业名称")
    parser.add_argument("--output", "-o", required=True, help="输出 Word 文件路径")
    args = parser.parse_args()

    with open(args.data, "r", encoding="utf-8") as f:
        checklist_data = json.load(f)

    output = generate_dd_checklist(checklist_data, args.company, args.output)
    print(f"✅ 检查清单已生成: {output}", file=sys.stderr)


if __name__ == "__main__":
    main()
