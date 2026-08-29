#!/usr/bin/env python3
"""
企业尽调 - 报告生成模块（双模板）
- lite（轻量版，默认）: 5 章公开信息初筛报告
- standard（基础版）: 7 章完整尽调报告（法律/财务/业务三维度，需 --template standard 显式指定）

使用方式:
    python generate_dd_report.py --data dd_data.json --company "企业名称" --output "报告.docx" [--template lite|standard]
"""

import json
import os
import sys
import argparse
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("请先安装 python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ============================================================
# 公共颜色/字号常量
# ============================================================
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(255, 255, 255)
DARK_BG = RGBColor(15, 52, 96)      # 表头深色背景 #0F3460
LIGHT_BG = RGBColor(242, 242, 242)  # 隔行浅色背景 #F2F2F2
GRAY_NOTE = RGBColor(128, 128, 128)  # 待核实/免责灰

SIZE_COVER_SUBTITLE = 20
SIZE_H1 = 14
SIZE_H2 = 13
SIZE_BODY = 11
SIZE_TABLE = 9.5

# lite（轻量版）专属色
LITE_TITLE = RGBColor(0x1A, 0x1A, 0x2E)   # 章/节标题
LITE_BLUE = RGBColor(0x00, 0x72, 0xC6)    # 融资要点标题
LITE_RED = RGBColor(0xEE, 0x00, 0x00)     # 风险/待验证标注
LITE_GRAY = RGBColor(0x44, 0x44, 0x44)    # 正文说明
LITE_PRODUCT = RGBColor(0x0F, 0x34, 0x60)  # 产品方向标题
LITE_CAP = RGBColor(0x4F, 0x81, 0xBD)     # 附录来源
LITE_DIS = RGBColor(0x88, 0x88, 0x88)     # 免责声明


# ============================================================
# 模板配置（standard 基础版 / lite 轻量版）
# ============================================================
STANDARD_PROFILE = {
    "h1_color": RGBColor(55, 96, 146),    # #376092
    "h2_color": RGBColor(79, 129, 189),   # #4F81BD
    "cover_title_size": 36,
    "cover_title_color": BLACK,
    "cover_title_text": "尽职调查报告",
    "zebra": True,                        # 表格隔行变色
}

LITE_PROFILE = {
    "h1_color": LITE_TITLE,               # #1A1A2E
    "h2_color": LITE_TITLE,
    "cover_title_size": 42,
    "cover_title_color": LITE_TITLE,
    "cover_title_text": "企业尽职调查报告",
    "zebra": False,
}

CURRENT_PROFILE = STANDARD_PROFILE


def _profile():
    return CURRENT_PROFILE


# ============================================================
# 辅助函数
# ============================================================

def _set_run_font(run, size=SIZE_BODY, bold=False, color=BLACK, font_name="微软雅黑"):
    """设置 run 的字体属性（含中文字体）。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def _set_cell_shading(cell, color_hex):
    """设置单元格底色。"""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text, size=SIZE_TABLE, bold=False, color=BLACK, align="left"):
    """设置单元格文字。"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(str(text) if text else "")
    _set_run_font(run, size=size, bold=bold, color=color)


def _add_heading(doc, text, level=1):
    """添加标题段落（内置 Heading 样式驱动导航窗格 + run 级直接格式覆盖配色）。"""
    h = doc.add_heading(text, level=level)
    pf = _profile()
    hc = pf["h1_color"] if level == 1 else pf["h2_color"]
    sz = SIZE_H1 if level == 1 else SIZE_H2
    for run in h.runs:
        _set_run_font(run, size=sz, bold=True, color=hc)
    return h


def _add_body(doc, text, bold=False, color=BLACK, size=SIZE_BODY):
    """添加正文段落。"""
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    _set_run_font(run, size=size, bold=bold, color=color)
    return p


def _add_data_table(doc, headers, rows, col_widths=None, zebra=True, font_size=10):
    """
    添加数据表格。
    headers: 表头列表；rows: 行数据（列表套列表）；col_widths: 列宽比。
    """
    if not rows:
        return None
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    if col_widths:
        total = sum(col_widths)
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width * 15 / total)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, h, size=font_size, bold=True, color=WHITE, align="center")
        _set_cell_shading(cell, "0F3460")
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            _set_cell_text(cell, str(cell_text) if cell_text else "", size=font_size, color=BLACK)
            if zebra and r_idx % 2 == 1:
                _set_cell_shading(cell, "F2F2F2")
    return table


def _add_pending_note(doc, field_name):
    """添加待核实标记。"""
    p = doc.add_paragraph()
    run = p.add_run(f"【待核实】{field_name} — 数据不可得，需用户补充。")
    _set_run_font(run, size=10, color=GRAY_NOTE, bold=True)
    return p


def _add_lite_red(doc, text):
    """lite 版红色风险/待验证标注。"""
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    _set_run_font(run, size=10.5, bold=True, color=LITE_RED)
    return p


def _add_lite_body(doc, text):
    """lite 版正文说明（10.5pt #444444）。"""
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    _set_run_font(run, size=10.5, color=LITE_GRAY)
    return p


def _add_lite_product_title(doc, text):
    """lite 版产品方向标题（13pt #0F3460）。"""
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    _set_run_font(run, size=13, bold=True, color=LITE_PRODUCT)
    return p


# ============================================================
# standard 版：封面 + 7 章 + 附录（基础版，默认）
# ============================================================

def build_cover_page(doc, company, report_type="投资尽调"):
    """standard 封面页。"""
    pf = _profile()
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{company}")
    _set_run_font(run, size=pf["cover_title_size"], bold=True, color=pf["cover_title_color"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(pf["cover_title_text"])
    _set_run_font(run, size=SIZE_COVER_SUBTITLE, bold=True, color=pf["h1_color"])
    doc.add_paragraph()
    info_items = [
        f"报告日期：{datetime.now().strftime('%Y年%m月%d日')}",
        f"尽调类型：{report_type}",
        "尽调范围：法律/财务/业务综合尽调",
    ]
    for item in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        _set_run_font(run, size=11, color=BLACK)
    doc.add_page_break()


def build_section1_basic_info(doc, data):
    """第一章：公司基本信息与股权架构。"""
    _add_lite_chapter(doc, "一、公司基本信息与股权架构")
    _add_heading(doc, "1.1 公司基础信息", level=2)
    basic = data.get("basic_info", {})
    if basic.get("source") != "未获取" and basic.get("credit_code"):
        headers = ["项目", "内容"]
        rows = [
            ["统一社会信用代码", basic.get("credit_code", "")],
            ["法定代表人", basic.get("legal_person", "")],
            ["成立日期", basic.get("established_date", "")],
            ["注册资本", f"{basic.get('reg_capital', '')}{basic.get('reg_capital_unit', '万元')}"],
            ["注册地址", basic.get("address", "")],
            ["企业类型", basic.get("company_type", "")],
            ["经营状态", basic.get("status", "")],
            ["经营范围", basic.get("business_scope", "")],
        ]
        _add_data_table(doc, headers, rows, col_widths=[4, 10], zebra=_profile()["zebra"])
    else:
        _add_pending_note(doc, "公司基础信息（公开数据不可得）")
    _add_heading(doc, "1.2 股权架构分析", level=2)
    equity = data.get("equity", [])
    if equity:
        headers = ["股东名称", "持股比例", "股东类型", "认缴出资额"]
        rows = [[e.get("shareholder", ""), f"{e.get('ratio', '')}%",
                 e.get("type", ""), e.get("amount", "")] for e in equity]
        _add_data_table(doc, headers, rows, col_widths=[5, 2, 3, 4], zebra=_profile()["zebra"])
    else:
        _add_pending_note(doc, "股权架构信息（公开数据不可得）")
    controller = data.get("controller", "")
    if controller:
        _add_body(doc, f"实际控制人：{controller}")
    else:
        _add_pending_note(doc, "实际控制人信息")
    changes = data.get("change_records", [])
    if changes:
        _add_heading(doc, "1.3 历史沿革", level=2)
        headers = ["变更日期", "变更事项", "变更前", "变更后"]
        rows = [[c.get("date", ""), c.get("item", ""),
                 c.get("before", ""), c.get("after", "")] for c in changes]
        _add_data_table(doc, headers, rows, col_widths=[3, 3, 4, 4], zebra=_profile()["zebra"])
    doc.add_paragraph()


def build_section2_team(doc, data):
    """第二章：核心团队画像。"""
    _add_lite_chapter(doc, "二、核心团队画像")
    management = data.get("management", [])
    if management:
        _add_heading(doc, "2.1 核心高管背景", level=2)
        headers = ["姓名", "职务", "背景简介", "任职时间"]
        rows = [[m.get("name", ""), m.get("position", ""),
                 m.get("background", ""), m.get("tenure", "")] for m in management]
        _add_data_table(doc, headers, rows, col_widths=[3, 3, 8, 2], zebra=_profile()["zebra"])
    else:
        _add_pending_note(doc, "核心团队信息（公开数据不可得）")
    team_analysis = data.get("team_analysis", "")
    if team_analysis:
        _add_heading(doc, "2.2 团队评估意见", level=2)
        _add_body(doc, team_analysis, color=BLACK)
    doc.add_paragraph()


def build_section3_ip(doc, data):
    """第三章：知识产权与专利。"""
    _add_lite_chapter(doc, "三、知识产权与专利")
    ip = data.get("ip", {})
    patents = ip.get("patents", [])
    if patents:
        _add_heading(doc, "3.1 专利分析", level=2)
        _add_body(doc, f"共计 {ip.get('patent_count', len(patents))} 项专利（含申请中）")
        headers = ["专利名称", "类型", "申请号", "状态"]
        rows = [[p.get("name", ""), p.get("type", ""),
                 p.get("app_number", ""), p.get("status", "")] for p in patents]
        _add_data_table(doc, headers, rows, col_widths=[6, 2, 4, 2], zebra=_profile()["zebra"])
    else:
        _add_pending_note(doc, "专利信息（公开数据不可得）")
    trademarks = ip.get("trademarks", [])
    if trademarks:
        _add_heading(doc, "3.2 商标分析", level=2)
        headers = ["商标名称", "注册类别", "注册号", "状态"]
        rows = [[t.get("name", ""), t.get("class", ""),
                 t.get("reg_number", ""), t.get("status", "")] for t in trademarks]
        _add_data_table(doc, headers, rows, col_widths=[4, 3, 4, 3], zebra=_profile()["zebra"])
    else:
        _add_pending_note(doc, "商标信息（公开数据不可得）")
    copyrights = ip.get("copyrights", [])
    if copyrights:
        _add_heading(doc, "3.3 软件著作权", level=2)
        _add_body(doc, f"共计 {ip.get('copyright_count', len(copyrights))} 项软件著作权")
    tech_analysis = data.get("tech_analysis", "")
    if tech_analysis:
        _add_heading(doc, "3.4 核心技术评估", level=2)
        _add_body(doc, tech_analysis, color=BLACK)
    doc.add_paragraph()


def build_section4_finance(doc, data):
    """第四章：财务分析与税务情况。"""
    _add_heading(doc, "四、财务分析与税务情况", level=1)
    financial = data.get("financial_analysis", {})
    if financial.get("data_available"):
        _add_heading(doc, "4.1 财务报表摘要", level=2)
        bs = financial.get("balance_sheet", {})
        if bs:
            headers = ["项目", "期末余额（万元）"]
            rows = [[k, str(v)] for k, v in bs.items()]
            _add_data_table(doc, headers, rows, col_widths=[6, 8], zebra=_profile()["zebra"])
        ratios = financial.get("ratios", [])
        if ratios:
            _add_heading(doc, "4.2 关键财务比率", level=2)
            headers = ["比率名称", "数值", "参考值", "评价"]
            rows = [[r.get("name", ""), r.get("value", ""),
                     r.get("ref", ""), r.get("assessment", "")] for r in ratios]
            _add_data_table(doc, headers, rows, col_widths=[4, 2, 2, 6], zebra=_profile()["zebra"])
    else:
        _add_body(doc, "财务数据通常来自用户提供的财务报表。企查查等公开数据源一般不提供详细财务数据。")
        _add_pending_note(doc, "财务数据（需要用户提供财务报表）")
    _add_heading(doc, "4.3 税务情况", level=2)
    tax_data = financial.get("tax", {})
    if tax_data:
        _add_body(doc, f"税务状态：{tax_data.get('status', '待核实')}")
    else:
        _add_pending_note(doc, "税务信息（公开数据不可得）")
    _add_heading(doc, "4.4 关联交易分析", level=2)
    related = financial.get("related_transactions", [])
    if related:
        headers = ["交易对手", "交易类型", "金额（万元）", "公允性"]
        rows = [[r.get("party", ""), r.get("type", ""),
                 r.get("amount", ""), r.get("fairness", "")] for r in related]
        _add_data_table(doc, headers, rows, col_widths=[4, 3, 3, 4], zebra=_profile()["zebra"])
    else:
        _add_pending_note(doc, "关联交易信息（需要用户提供财务报表）")
    doc.add_paragraph()


def build_section5_business(doc, data):
    """第五章：业务与市场分析。"""
    _add_heading(doc, "五、业务与市场分析", level=1)
    products = data.get("products", [])
    if products:
        _add_heading(doc, "5.1 核心产品/服务分析", level=2)
        headers = ["产品名称", "产品类别", "功能定位", "关键信息"]
        rows = [[p.get("name", ""), p.get("category", ""),
                 p.get("description", ""), p.get("note", "")] for p in products]
        _add_data_table(doc, headers, rows, col_widths=[3, 2, 5, 4], zebra=_profile()["zebra"])
        doc.add_paragraph()
        service_target = data.get("service_target", "")
        if service_target:
            _add_body(doc, f"服务对象：{service_target}", color=BLACK)
        usage_scenario = data.get("usage_scenario", "")
        if usage_scenario:
            _add_body(doc, f"使用场景：{usage_scenario}", color=BLACK)
    biz_model = data.get("business_model", "")
    if biz_model:
        _add_heading(doc, "5.2 商业模式分析", level=2)
        _add_body(doc, biz_model, color=BLACK)
    market = data.get("market_analysis", "")
    if market:
        _add_heading(doc, "5.3 市场与竞争分析", level=2)
        _add_body(doc, market, color=BLACK)
    _add_heading(doc, "5.4 客户与供应商分析", level=2)
    customer_data = data.get("customer_analysis", {})
    if customer_data:
        if "top_customers" in customer_data:
            _add_body(doc, f"前5大客户占比：{customer_data.get('top_customers', '待核实')}%")
        if "top_suppliers" in customer_data:
            _add_body(doc, f"前5大供应商占比：{customer_data.get('top_suppliers', '待核实')}%")
    else:
        _add_pending_note(doc, "客户与供应商信息（需要用户提供）")
    doc.add_paragraph()


def build_section6_financing(doc, data):
    """第六章：融资历史与估值。"""
    _add_heading(doc, "六、融资历史与资本结构", level=1)
    financing = data.get("financing", [])
    if financing:
        _add_heading(doc, "6.1 融资历程", level=2)
        headers = ["融资轮次", "融资金额", "投资方", "融资时间"]
        rows = [[f.get("round", ""), f.get("amount", ""),
                 f.get("investors", ""), f.get("date", "")] for f in financing]
        _add_data_table(doc, headers, rows, col_widths=[3, 3, 5, 3], zebra=_profile()["zebra"])
    else:
        _add_pending_note(doc, "融资历史信息（公开数据不可得）")
    capital = data.get("capital_structure", "")
    if capital:
        _add_heading(doc, "6.2 资本结构分析", level=2)
        _add_body(doc, capital, color=BLACK)
    doc.add_paragraph()


def build_section7_risk(doc, data):
    """第七章：风险评估与结论。"""
    _add_heading(doc, "七、风险评估与结论", level=1)
    risk_matrix = data.get("risk_matrix", {})
    _add_heading(doc, "7.1 综合风险评估", level=2)
    p0 = risk_matrix.get("p0", [])
    p1 = risk_matrix.get("p1", [])
    p2 = risk_matrix.get("p2", [])
    summary_headers = ["风险等级", "数量"]
    summary_rows = [
        ["P0 - Deal Breaker（致命风险）", str(len(p0))],
        ["P1 - Material（重大事项）", str(len(p1))],
        ["P2 - Informational（参考信息）", str(len(p2))],
    ]
    _add_data_table(doc, summary_headers, summary_rows, col_widths=[10, 4], zebra=_profile()["zebra"])
    doc.add_paragraph()
    if p0:
        _add_heading(doc, "P0 - Deal Breaker 级风险", level=2)
        p0_headers = ["序号", "风险事项", "当前状态", "建议行动"]
        p0_rows = [[str(i+1), item.get("risk", ""), item.get("status", ""),
                    item.get("action", "")] for i, item in enumerate(p0)]
        _add_data_table(doc, p0_headers, p0_rows, col_widths=[1, 5, 3, 5], zebra=_profile()["zebra"])
    if p1:
        _add_heading(doc, "P1 - Material 级风险", level=2)
        p1_headers = ["序号", "风险事项", "当前状态", "建议行动"]
        p1_rows = [[str(i+1), item.get("risk", ""), item.get("status", ""),
                    item.get("action", "")] for i, item in enumerate(p1)]
        _add_data_table(doc, p1_headers, p1_rows, col_widths=[1, 5, 3, 5], zebra=_profile()["zebra"])
    if p2:
        _add_heading(doc, "P2 - Informational 级风险", level=2)
        p2_headers = ["序号", "事项", "说明"]
        p2_rows = [[str(i+1), item.get("risk", ""), item.get("note", "")]
                   for i, item in enumerate(p2)]
        _add_data_table(doc, p2_headers, p2_rows, col_widths=[1, 5, 8], zebra=_profile()["zebra"])
    _add_heading(doc, "7.2 各维度风险评级", level=2)
    dim_headers = ["尽调维度", "风险等级", "说明"]
    dim_rows = [
        ["法律尽调", data.get("legal_analysis", {}).get("risk_level", "待核实"),
         data.get("legal_analysis", {}).get("notes", "")],
        ["财务尽调", data.get("financial_analysis", {}).get("risk_level", "待核实"),
         data.get("financial_analysis", {}).get("notes", "")],
        ["业务尽调", data.get("business_analysis", {}).get("risk_level", "待核实"),
         data.get("business_analysis", {}).get("notes", "")],
    ]
    _add_data_table(doc, dim_headers, dim_rows, col_widths=[3, 3, 8], zebra=_profile()["zebra"])
    doc.add_paragraph()
    _add_heading(doc, "7.3 尽调结论", level=2)
    overall = data.get("overall_rating", "数据不完整")
    _add_body(doc, f"综合风险评级：{overall}", bold=True, color=BLACK)
    _add_body(doc, "", color=BLACK)
    conclusion = data.get("conclusion", "")
    if conclusion:
        _add_body(doc, conclusion, color=BLACK)
    _add_heading(doc, "7.4 建议", level=2)
    recommendations = data.get("recommendations", "")
    if recommendations:
        _add_body(doc, recommendations, color=BLACK)
    doc.add_paragraph()


def build_appendix(doc, data):
    """standard 附录。"""
    _add_heading(doc, "附录", level=1)
    _add_heading(doc, "数据来源清单", level=2)
    sources = data.get("data_sources", [])
    if sources:
        for s in sources:
            _add_body(doc, f"- {s}", color=BLACK)
    else:
        _add_pending_note(doc, "数据来源信息")
    doc.add_paragraph()
    _add_heading(doc, "免责声明", level=2)
    disclaimer = (
        "本报告基于公开可获取信息及用户提供的文档自动生成，不构成投资建议、"
        "证券推荐或任何形式的承诺。报告中的分析结论仅供参考，具体交易决策请以"
        "一手尽职调查、专业顾问意见和官方文件为准。所有标注为'待核实'的内容，"
        "必须在正式尽调中通过一手来源核实。"
    )
    _add_body(doc, disclaimer, color=GRAY_NOTE)


# ============================================================
# lite 版：封面 + 目录 + 5 章 + 附录（轻量版）

def _add_lite_chapter(doc, text):
    """lite 章节标题：Heading 1 + 段前分页（每个板块从新页开始）。"""
    h = _add_heading(doc, text, level=1)
    h.paragraph_format.page_break_before = True
    return h

# ============================================================

def build_lite_cover(doc, company, report_type="投资/合作背景调查"):
    """lite 封面（42pt #1A1A2E 公司名）。"""
    pf = _profile()
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{company}")
    _set_run_font(run, size=pf["cover_title_size"], bold=True, color=pf["cover_title_color"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(pf["cover_title_text"])
    _set_run_font(run, size=SIZE_COVER_SUBTITLE, bold=False, color=pf["cover_title_color"])
    doc.add_paragraph()
    info_items = [
        f"报告日期：{datetime.now().strftime('%Y年%m月%d日')}",
        f"尽调类型：{report_type}",
        f"信息基准日：{datetime.now().strftime('%Y年%m月%d日')}（公开信息检索日）",
    ]
    for item in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        _set_run_font(run, size=11, color=LITE_GRAY)
    doc.add_page_break()


def build_lite_toc(doc):
    """lite 手写静态目录（五章 + 附录）。"""
    _add_heading(doc, "目录", level=1)
    toc_items = [
        "一、公司基本信息与股权架构",
        "二、核心团队画像",
        "三、知识产权与专利",
        "四、融资历史与资本结构",
        "五、核心产品",
        "附录：数据来源清单",
    ]
    for it in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(it)
        _set_run_font(run, size=12, color=BLACK)
def _lite_kv_table(doc, rows, widths=None):
    """lite 两列「维度-信息」表。"""
    if not rows:
        return None
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    for i, h in enumerate(["维度", "信息"]):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, h, size=9.5, bold=True, color=WHITE, align="center")
        _set_cell_shading(cell, "0F3460")
    for k, v in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], str(k), size=9.5)
        _set_cell_text(cells[1], str(v), size=9.5)
    return table


def build_lite_s1(doc, data):
    """lite 第一章：公司基本信息与股权架构。"""
    _add_lite_chapter(doc, "一、公司基本信息与股权架构")
    _add_heading(doc, "1.1 公司基础信息", level=2)
    basic = data.get("basic_info", {})
    lite_rows = data.get("lite_basic_rows")
    if lite_rows:
        rows = lite_rows
    else:
        rows = [
            ["统一社会信用代码", basic.get("credit_code", "")],
            ["法定代表人", basic.get("legal_person", "")],
            ["成立日期", basic.get("established_date", "")],
            ["注册资本", f"{basic.get('reg_capital', '')}{basic.get('reg_capital_unit', '万元')}"],
            ["实缴资本", data.get("paid_in_capital", "待核实")],
            ["企业类型", basic.get("company_type", "")],
            ["注册地址", basic.get("address", "")],
            ["经营状态", basic.get("status", "")],
            ["营业期限", basic.get("business_term", "待核实")],
            ["登记机关", basic.get("reg_authority", "待核实")],
            ["行业分类", basic.get("industry", "待核实")],
            ["员工参保人数", data.get("employee_count", "待核实")],
            ["企业规模/资质", data.get("company_qualifications", "待核实")],
            ["官网", data.get("website", "")],
            ["联系电话", data.get("contact_phone", "")],
            ["联系邮箱", data.get("contact_email", "")],
        ]
    _lite_kv_table(doc, rows, widths=[4.2, 11.6])
    note = data.get("lite_basic_note", "")
    if note:
        p = doc.add_paragraph()
        run = p.add_run(note)
        _set_run_font(run, size=9, color=LITE_DIS)

    _add_heading(doc, "1.2 经营范围", level=2)
    scope = data.get("lite_business_scope") or basic.get("business_scope", "待核实")
    _add_lite_body(doc, scope)
    scope_note = data.get("lite_scope_note", "")
    if scope_note:
        p = doc.add_paragraph()
        run = p.add_run(scope_note)
        _set_run_font(run, size=9, color=LITE_DIS)

    _add_heading(doc, "1.3 股权结构", level=2)
    equity = data.get("equity", [])
    if equity:
        headers = ["序号", "股东名称", "股东类型", "持股比例", "认缴出资(万元)"]
        rows = [[str(i+1), e.get("shareholder", ""), e.get("type", ""),
                 f"{e.get('ratio', '')}%", e.get("amount", "")] for i, e in enumerate(equity)]
        _add_data_table(doc, headers, rows, col_widths=[1.2, 6.0, 4.4, 2.6, 2.8],
                        zebra=_profile()["zebra"], font_size=9.5)
    else:
        _add_pending_note(doc, "股权结构信息（公开数据不可得）")
    controller = data.get("lite_controller_note") or (f"实际控制人：{data.get('controller', '')}" if data.get("controller") else "")
    if controller:
        p = doc.add_paragraph()
        run = p.add_run(controller)
        _set_run_font(run, size=10.5)
    comment = data.get("lite_equity_comment", "")
    if comment:
        p = doc.add_paragraph()
        run = p.add_run(comment)
        _set_run_font(run, size=10.5)

    _add_heading(doc, "1.4 主要人员", level=2)
    management = data.get("management", [])
    if management:
        headers = ["姓名", "职务", "持股/备注"]
        rows = [[m.get("name", ""), m.get("position", ""),
                 m.get("background", "") or m.get("tenure", "")] for m in management]
        _add_data_table(doc, headers, rows, col_widths=[2.6, 7.2, 6.0],
                        zebra=_profile()["zebra"], font_size=9.5)
    else:
        _add_pending_note(doc, "主要人员信息（公开数据不可得）")
    mgmt_note = data.get("lite_mgmt_note", "")
    if mgmt_note:
        _add_lite_red(doc, mgmt_note)
def build_lite_s2(doc, data):
    """lite 第二章：核心团队画像（维度-信息表）。"""
    _add_lite_chapter(doc, "二、核心团队画像")
    profiles = data.get("team_profiles", [])
    if not profiles:
        _add_pending_note(doc, "核心团队信息（公开数据不可得）")
    for idx, prof in enumerate(profiles, 1):
        _add_heading(doc, f"2.{idx} {prof.get('title', '核心成员')}", level=2)
        dims = prof.get("dims", [])
        if dims:
            _lite_kv_table(doc, dims, widths=[3.2, 12.6])
    team_analysis = data.get("team_analysis", "")
    if team_analysis:
        p = doc.add_paragraph()
        run = p.add_run(team_analysis)
        _set_run_font(run, size=10.5)
    team_note = data.get("lite_team_note", "")
    if team_note:
        _add_lite_red(doc, team_note)
def build_lite_s3(doc, data):
    """lite 第三章：知识产权与专利。"""
    _add_lite_chapter(doc, "三、知识产权与专利")
    ip = data.get("ip", {})
    _add_heading(doc, "3.1 知识产权总览", level=2)
    patents = data.get("lite_patents") or ip.get("patents", [])
    if patents:
        headers = ["名称", "类型", "法律状态", "公开/授权时间"]
        rows = [[p.get("name", ""), p.get("type", ""),
                 p.get("status", ""), p.get("date", "")] for p in patents]
        _add_data_table(doc, headers, rows, col_widths=[6.4, 2.2, 3.2, 3.2],
                        zebra=_profile()["zebra"], font_size=9.5)
    else:
        _add_pending_note(doc, "专利信息（公开数据不可得）")
    patent_note = data.get("lite_patent_note", "")
    if patent_note:
        p = doc.add_paragraph()
        run = p.add_run(patent_note)
        _set_run_font(run, size=10.5)
    patent_risk = data.get("lite_patent_risk", "")
    if patent_risk:
        _add_lite_red(doc, patent_risk)

    _add_heading(doc, "3.2 商标与资质", level=2)
    ip_rows = data.get("lite_ip_rows")
    if not ip_rows:
        trademarks = ip.get("trademarks", [])
        ip_rows = [[t.get("name", ""), t.get("class", ""), t.get("status", "待核实")]
                   for t in trademarks]
    if ip_rows:
        _add_data_table(doc, ["项目", "内容", "状态"], ip_rows,
                        col_widths=[4.2, 8.8, 3.0], zebra=_profile()["zebra"], font_size=9.5)
    else:
        _add_pending_note(doc, "商标/资质信息（公开数据不可得）")
    qual_risk = data.get("lite_qual_risk", "")
    if qual_risk:
        _add_lite_red(doc, qual_risk)
def build_lite_s4(doc, data):
    """lite 第四章：融资历史与资本结构。"""
    _add_lite_chapter(doc, "四、融资历史与资本结构")
    _add_heading(doc, "4.1 融资时间线", level=2)
    financing = data.get("financing", [])
    if financing:
        headers = ["时间", "事件", "金额", "来源"]
        rows = [[f.get("date", ""), f.get("event", ""),
                 f.get("amount", ""), f.get("source", "")] for f in financing]
        _add_data_table(doc, headers, rows, col_widths=[3.0, 8.6, 2.0, 2.4],
                        zebra=_profile()["zebra"], font_size=9.5)
    else:
        _add_pending_note(doc, "融资历史信息（公开数据不可得）")
    _add_heading(doc, "4.2 资本结构与融资要点", level=2)
    p = doc.add_paragraph()
    run = p.add_run("融资要点")
    _set_run_font(run, size=14, bold=True, color=LITE_BLUE)
    fin_notes = data.get("financing_notes", [])
    if fin_notes:
        for note in fin_notes:
            _add_lite_red(doc, note)
    else:
        _add_lite_red(doc, "融资要点：暂无公开披露信息，需在正式尽调中通过一手资料核实。")
def build_lite_s5(doc, data):
    """lite 第五章：核心产品。"""
    _add_lite_chapter(doc, "五、核心产品")
    _add_heading(doc, "5.1 核心产品", level=2)
    groups = data.get("product_groups", [])
    if groups:
        for g in groups:
            _add_lite_product_title(doc, g.get("title", ""))
            headers = g.get("headers", [])
            rows = g.get("rows", [])
            if headers and rows:
                widths = g.get("widths")
                _add_data_table(doc, headers, rows, col_widths=widths,
                                zebra=_profile()["zebra"], font_size=9.5)
    else:
        products = data.get("products", [])
        if products:
            headers = ["产品名称", "产品类别", "功能定位", "关键信息"]
            rows = [[p.get("name", ""), p.get("category", ""),
                     p.get("description", ""), p.get("note", "")] for p in products]
            _add_data_table(doc, headers, rows, col_widths=[3, 2, 5, 4],
                            zebra=_profile()["zebra"], font_size=9.5)
        else:
            _add_pending_note(doc, "产品信息（公开数据不可得）")
    _add_heading(doc, "5.2 核心技术", level=2)
    core_tech = data.get("core_tech", [])
    if core_tech:
        for t in core_tech:
            _add_lite_body(doc, t)
    else:
        _add_pending_note(doc, "核心技术信息（公开数据不可得）")
    tech_note = data.get("lite_tech_note", "")
    if tech_note:
        _add_lite_red(doc, tech_note)
    _add_heading(doc, "5.3 行业与市场", level=2)
    market = data.get("market_texts", [])
    if not market and data.get("market_analysis"):
        market = [data["market_analysis"]]
    if market:
        for m in market:
            _add_lite_body(doc, m)
    plan_note = data.get("lite_plan_note", "")
    if plan_note:
        _add_lite_red(doc, plan_note)
def build_lite_appendix(doc, data):
    """lite 附录：数据来源清单 + 免责声明。"""
    _add_lite_chapter(doc, "附录：数据来源清单")
    sources = data.get("data_sources", [])
    if sources:
        for i, s in enumerate(sources, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {s}")
            _set_run_font(run, size=9, color=LITE_CAP)
    else:
        _add_pending_note(doc, "数据来源信息")
    doc.add_paragraph()
    disclaimer = data.get("disclaimer") or (
        "免责声明：本报告基于公开可获取信息整理，不构成投资建议、证券推荐或任何机构的承诺；"
        "具体交易决策请以一手尽职调查、专业顾问意见和官方文件为准。所有标有「待验证」「待核实」"
        "「公开渠道未见」的内容，必须在正式尽调中通过一手来源（工商内档、药监局注册查询、"
        "审计报告、银行流水、访谈记录等）核实。"
    )
    p = doc.add_paragraph()
    run = p.add_run(disclaimer)
    _set_run_font(run, size=9, color=LITE_DIS)


# ============================================================
# 主入口
# ============================================================

def _add_static_toc(doc):
    """基础版静态目录：按章节级条目列出（小节因数据可得性会跳号，不列小节避免错位）。

    同时插入 TOC 域：Word/WPS 打开后右键「更新域」可生成带页码的自动目录，
    预览器中直接显示静态条目。
    """
    chapters = [
        "一、公司基本信息与股权架构",
        "二、核心团队画像",
        "三、知识产权与专利",
        "四、财务分析与税务情况",
        "五、业务与市场分析",
        "六、融资历史与资本结构",
        "七、风险评估与结论",
        "附录：数据来源清单与免责声明",
    ]
    for ch in chapters:
        p = doc.add_paragraph()
        run = p.add_run(ch)
        _set_run_font(run, size=11, bold=True, color=BLACK)
    note = doc.add_paragraph()
    note_run = note.add_run("（说明：各章小节因数据可得性可能存在编号跳项；Word/WPS 中可通过导航窗格按标题跳转。）")
    _set_run_font(note_run, size=9, color=RGBColor(0x88, 0x88, 0x88))


def generate_dd_report(dd_data, company_name, output_path, report_type="投资尽调",
                       template="lite"):
    """
    生成尽调报告 Word 文档。
    template: lite（默认，轻量版）| standard（基础版）
    """
    global CURRENT_PROFILE
    doc = Document()

    if template == "lite":
        CURRENT_PROFILE = LITE_PROFILE
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = section.bottom_margin = Cm(2.5)
        section.left_margin = section.right_margin = Cm(2.5)
        build_lite_cover(doc, company_name, report_type)
        build_lite_toc(doc)
        build_lite_s1(doc, dd_data)
        build_lite_s2(doc, dd_data)
        build_lite_s3(doc, dd_data)
        build_lite_s4(doc, dd_data)
        build_lite_s5(doc, dd_data)
        build_lite_appendix(doc, dd_data)
    else:
        CURRENT_PROFILE = STANDARD_PROFILE
        section = doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        build_cover_page(doc, company_name, report_type)
        _add_heading(doc, "目录", level=1)
        _add_static_toc(doc)
        doc.add_page_break()
        build_section1_basic_info(doc, dd_data)
        build_section2_team(doc, dd_data)
        build_section3_ip(doc, dd_data)
        build_section4_finance(doc, dd_data)
        build_section5_business(doc, dd_data)
        build_section6_financing(doc, dd_data)
        build_section7_risk(doc, dd_data)
        build_appendix(doc, dd_data)

    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="企业尽调 - 报告生成")
    parser.add_argument("--data", "-d", required=True, help="尽调数据 JSON 文件路径")
    parser.add_argument("--company", "-c", required=True, help="企业名称")
    parser.add_argument("--output", "-o", required=True, help="输出 Word 文件路径")
    parser.add_argument("--type", "-t", default="投资尽调", help="尽调类型")
    parser.add_argument("--template", "-tpl", default="lite",
                        choices=["lite", "standard"],
                        help="报告模板：lite（轻量版，默认）| standard（基础版）")
    args = parser.parse_args()

    with open(args.data, "r", encoding="utf-8") as f:
        dd_data = json.load(f)

    output = generate_dd_report(dd_data, args.company, args.output, args.type, args.template)
    print(f"✅ 尽调报告已生成: {output}（模板: {args.template}）", file=sys.stderr)


if __name__ == "__main__":
    main()
