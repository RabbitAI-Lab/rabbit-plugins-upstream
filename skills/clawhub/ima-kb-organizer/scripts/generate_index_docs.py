#!/usr/bin/env python3
"""
IMA 知识库分类索引文档生成器（Skill 版）
读取 config.json 和 tracker.json，为每个有效分类生成 Word + Markdown 索引文档。
支持配置驱动、多格式输出、自动复制到用户目录。

用法：
  python generate_index_docs.py --config <config.json> [--tracker <tracker.json>]

如果不指定 --tracker，默认使用 config 中 tracker_directory 下的 tracker.json。
"""

import json
import sys
import os
import shutil
import argparse
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def set_cell_font(cell, font_name="微软雅黑", font_size=9, bold=False, color=None):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_font(paragraph, font_name="微软雅黑", font_size=11, bold=False, color=None):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def generate_word_doc(category_name, description, items, kb_names, output_path, scan_time):
    """生成 Word 索引文档"""
    if not HAS_DOCX:
        print(f"警告：python-docx 未安装，跳过 Word 生成：{category_name}")
        return False

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    title = doc.add_heading(f"{category_name} - 资料索引", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    desc_para = doc.add_paragraph()
    desc_para.add_run(f"分类说明：{description or ''}")
    set_paragraph_font(desc_para, font_size=10, color=RGBColor(0x66, 0x66, 0x66))

    meta_para = doc.add_paragraph()
    meta_para.add_run(f"知识库：{kb_names}　|　资料数量：{len(items)} 篇　|　最后更新：{scan_time}")
    set_paragraph_font(meta_para, font_size=9, color=RGBColor(0x99, 0x99, 0x99))

    doc.add_paragraph("")

    if items:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_cells = table.rows[0].cells
        headers = ["序号", "资料标题", "来源", "导入知识库日期"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            set_cell_font(header_cells[i], font_size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            shading = header_cells[i]._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "4472C4"
            })
            shading.append(shading_elm)

        for idx, item in enumerate(items, 1):
            row = table.add_row()
            cells = row.cells
            cells[0].text = str(idx)
            set_cell_font(cells[0], font_size=9)
            cells[1].text = item.get("title", "（无标题）")
            set_cell_font(cells[1], font_size=9)
            cells[2].text = item.get("source", "")
            set_cell_font(cells[2], font_size=9)
            first_seen = item.get("first_seen", "")
            date_str = first_seen[:10] if len(first_seen) >= 10 else first_seen
            cells[3].text = date_str
            set_cell_font(cells[3], font_size=9)

        for row in table.rows:
            row.cells[0].width = Inches(0.5)
            row.cells[1].width = Inches(4.5)
            row.cells[2].width = Inches(1.5)
            row.cells[3].width = Inches(1.2)
    else:
        empty_para = doc.add_paragraph()
        empty_para.add_run("（暂无资料）")
        set_paragraph_font(empty_para, font_size=11, color=RGBColor(0x99, 0x99, 0x99))
        empty_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("")
    usage_heading = doc.add_paragraph()
    usage_heading.add_run("使用说明")
    set_paragraph_font(usage_heading, font_size=10, bold=True, color=RGBColor(0x33, 0x33, 0x33))
    usage_para = doc.add_paragraph()
    usage_para.add_run(
        "1. 撰写文档时，告诉 AI 助手使用本分类的资料作为素材来源。\n"
        "2. AI 会根据本索引中的资料，从 IMA 知识库拉取全文内容，进行 RAG 检索增强生成。\n"
        "3. 新资料请放入 IMA 知识库根目录，AI 会定期扫描并自动分类更新本文档。\n"
        "4. 如发现分类有误，请告知 AI 助手调整。"
    )
    set_paragraph_font(usage_para, font_size=9, color=RGBColor(0x66, 0x66, 0x66))

    doc.save(output_path)
    return True


def generate_markdown_doc(category_name, description, items, kb_names, output_path, scan_time):
    """生成 Markdown 索引文档"""
    lines = [
        f"# {category_name} - 资料索引\n",
        f"> {description or ''}\n",
        f"- **知识库**：{kb_names}",
        f"- **资料数量**：{len(items)} 篇",
        f"- **最后更新**：{scan_time}\n",
    ]

    if items:
        lines.append("| 序号 | 资料标题 | 来源 | 导入知识库日期 |")
        lines.append("|------|----------|------|----------------|")
        for idx, item in enumerate(items, 1):
            title = item.get("title", "（无标题）").replace("|", "\\|")
            source = item.get("source", "").replace("|", "\\|")
            first_seen = item.get("first_seen", "")
            date_str = first_seen[:10] if len(first_seen) >= 10 else first_seen
            lines.append(f"| {idx} | {title} | {source} | {date_str} |")
    else:
        lines.append("*（暂无资料）*\n")

    lines.append("\n---\n")
    lines.append("## 使用说明\n")
    lines.append("1. 撰写文档时，告诉 AI 助手使用本分类的资料作为素材来源。")
    lines.append("2. AI 会根据本索引中的资料，从 IMA 知识库拉取全文内容，进行 RAG 检索增强生成。")
    lines.append("3. 新资料请放入 IMA 知识库根目录，AI 会定期扫描并自动分类更新本文档。")
    lines.append("4. 如发现分类有误，请告知 AI 助手调整。")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines) + "\n")


def main():
    parser = argparse.ArgumentParser(description="IMA 知识库分类索引文档生成器")
    parser.add_argument("--config", required=True, help="配置文件路径 (config.json)")
    parser.add_argument("--tracker", help="追踪文件路径 (tracker.json)，默认从 config 推导")
    args = parser.parse_args()

    # 读取配置
    with open(args.config, "r", encoding="utf-8") as f:
        config = json.load(f)

    # 确定 tracker 路径
    tracker_path = args.tracker
    if not tracker_path:
        tracker_dir = config.get("tracker_directory", ".workbuddy/ima-tracker")
        tracker_path = os.path.join(tracker_dir, "tracker.json")

    if not os.path.exists(tracker_path):
        print(f"错误：追踪文件不存在：{tracker_path}")
        sys.exit(1)

    with open(tracker_path, "r", encoding="utf-8") as f:
        tracker = json.load(f)

    scan_time = tracker.get("last_scan_time", datetime.now().isoformat())
    formats = config.get("formats", ["docx"])
    output_dir = config.get("output_directory", os.path.dirname(tracker_path))

    # 合并 tracker 目录下的 category_docs 子目录（工作副本）
    work_output_dir = os.path.join(os.path.dirname(tracker_path), "category_docs")
    os.makedirs(work_output_dir, exist_ok=True)

    # 用户指定输出目录
    os.makedirs(output_dir, exist_ok=True)

    # 获取有效分类名称
    valid_categories = [c["name"] for c in config.get("categories", [])]
    category_desc = {c["name"]: c.get("description", "") for c in config.get("categories", [])}

    # 汇总所有知识库中各分类的内容
    all_categories = {}
    for kb_id, kb_info in tracker.get("knowledge_bases", {}).items():
        kb_name = kb_info.get("name", "")
        for cat_name, cat_data in kb_info.get("categories", {}).items():
            if cat_name not in valid_categories:
                continue
            if cat_name not in all_categories:
                all_categories[cat_name] = {"items": [], "kb_names": set()}
            all_categories[cat_name]["items"].extend(cat_data.get("items", []))
            all_categories[cat_name]["kb_names"].add(kb_name)

    generated = []
    for cat_name in valid_categories:
        cat_info = all_categories.get(cat_name, {"items": [], "kb_names": set()})
        items = cat_info["items"]
        kb_names = "、".join(sorted(cat_info["kb_names"])) if cat_info["kb_names"] else "（无）"
        desc = category_desc.get(cat_name, "")
        safe_name = cat_name.replace("+", "与").replace(" ", "")

        # 生成到工作目录
        if "docx" in formats and HAS_DOCX:
            word_path = os.path.join(work_output_dir, f"{safe_name}_资料索引.docx")
            generate_word_doc(cat_name, desc, items, kb_names, word_path, scan_time)
            # 复制到用户目录
            dst = os.path.join(output_dir, f"{safe_name}_资料索引.docx")
            shutil.copy2(word_path, dst)
            generated.append({"category": cat_name, "file": dst, "format": "docx", "count": len(items)})
            print(f"已生成 Word：{dst}（{len(items)} 篇）")

        if "markdown" in formats:
            md_path = os.path.join(work_output_dir, f"{safe_name}_资料索引.md")
            generate_markdown_doc(cat_name, desc, items, kb_names, md_path, scan_time)
            dst = os.path.join(output_dir, f"{safe_name}_资料索引.md")
            shutil.copy2(md_path, dst)
            generated.append({"category": cat_name, "file": dst, "format": "markdown", "count": len(items)})
            print(f"已生成 Markdown：{dst}（{len(items)} 篇）")

    # 摘要
    summary = {
        "generated_at": datetime.now().isoformat(),
        "scan_time": scan_time,
        "output_directory": output_dir,
        "documents": generated,
        "total_items": sum(g["count"] for g in generated),
    }
    summary_path = os.path.join(work_output_dir, "_生成摘要.json")
    with open(summary_path, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)

    print(f"\n生成完毕！共 {len(generated)} 个文档，{summary['total_items']} 篇资料。")
    print(f"输出目录：{output_dir}")


if __name__ == "__main__":
    main()
