# -*- coding: utf-8 -*-
"""
生成《IMA知识库智能整理：方法论与实践复盘》Word 文档
作者: sus-yugaohe
许可: CC BY-NC 4.0
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# 辅助函数
# ============================================================

def set_cell_font(cell, font_name="微软雅黑", font_size=10, bold=False, color=None):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_custom(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

def add_para(doc, text, font_size=11, bold=False, italic=False, alignment=None, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_bullet(doc, text, level=0, font_size=11):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    p.paragraph_format.line_spacing = 1.5
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_font(cell, font_size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    # 数据行
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
            set_cell_font(row.cells[i], font_size=10)
    # 列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table

def add_copyright_page(doc, title, subtitle, author="sus-yugaohe", version="v1.0"):
    """添加版权页"""
    doc.add_page_break()
    add_para(doc, "版权声明", font_size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    add_para(doc, f"作品名称：{title}", font_size=12, space_after=6)
    add_para(doc, f"副标题：{subtitle}", font_size=12, space_after=6)
    add_para(doc, f"作者：{author}", font_size=12, space_after=6)
    add_para(doc, f"版本：{version}", font_size=12, space_after=6)
    add_para(doc, f"发布日期：2026年7月", font_size=12, space_after=20)

    add_para(doc, "—" * 40, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_para(doc, f"\u00a9 2026 {author} 保留所有权利", font_size=12, bold=True, space_after=12)

    add_para(doc, "本作品采用知识共享署名-非商业性使用 4.0 国际许可协议（CC BY-NC 4.0）进行许可。", font_size=11, space_after=8)

    add_para(doc, "许可内容：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "署名（BY）：使用者必须在适当位置注明原作者 sus-yugaohe 及作品名称")
    add_bullet(doc, "非商业性使用（NC）：本作品及其衍生作品不得用于商业目的")
    add_bullet(doc, "共享：使用者可以复制、分发、传播本作品")
    add_bullet(doc, "改编：使用者可以创作基于本作品的衍生作品（须同样署名且非商业）")

    add_para(doc, "使用限制：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "未经作者书面授权，禁止将本作品用于任何商业用途，包括但不限于销售、付费课程、商业培训")
    add_bullet(doc, "引用本作品时，必须注明出处：\"sus-yugaohe, IMA知识库智能整理, 2026\"")
    add_bullet(doc, "改编作品须以相同或兼容的许可协议发布")
    add_bullet(doc, "作者不对本作品的适用性作任何担保，使用者自行承担使用风险")

    add_para(doc, "完整许可协议文本请访问：", font_size=11, space_after=4)
    add_para(doc, "https://creativecommons.org/licenses/by-nc/4.0/legalcode", font_size=10, italic=True, space_after=12)

    add_para(doc, "—" * 40, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_para(doc, "联系方式与反馈：通过 WorkBuddy 平台联系作者 sus-yugaohe", font_size=10, italic=True, color=RGBColor(0x80, 0x80, 0x80))

    doc.add_page_break()

def add_title_page(doc, title, subtitle, author="sus-yugaohe"):
    """添加封面"""
    # 空行留白
    for _ in range(6):
        add_para(doc, "", font_size=12, space_after=6)

    add_para(doc, title, font_size=28, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, color=RGBColor(0x1a, 0x56, 0xc4))
    add_para(doc, subtitle, font_size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30, color=RGBColor(0x40, 0x40, 0x40))

    for _ in range(8):
        add_para(doc, "", font_size=12, space_after=6)

    add_para(doc, f"作者：{author}", font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "版本：v1.0", font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "2026年7月", font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "\u00a9 2026 sus-yugaohe | CC BY-NC 4.0", font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x80, 0x80, 0x80))

    doc.add_page_break()

# ============================================================
# 文档生成
# ============================================================

def generate_document():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # ===== 封面 =====
    add_title_page(
        doc,
        "IMA知识库智能整理",
        "方法论与实践复盘\n—— 从痛点发现到技能封装的完整历程"
    )

    # ===== 版权页 =====
    add_copyright_page(
        doc,
        "IMA知识库智能整理：方法论与实践复盘",
        "从痛点发现到技能封装的完整历程"
    )

    # ===== 目录提示 =====
    add_heading_custom(doc, "目录", level=1, color=RGBColor(0x1a, 0x56, 0xc4))
    add_para(doc, "本文档整合三个层次的内容，读者可按需选择阅读：", font_size=11, space_after=12)
    add_para(doc, "第一部分　方法论总结 —— 适合想理解核心思路和设计原理的读者", font_size=11, bold=True, space_after=4)
    add_para(doc, "第二部分　完整沟通过程还原 —— 适合想了解每个决策细节的读者", font_size=11, bold=True, space_after=4)
    add_para(doc, "第三部分　简明步骤手册 —— 适合想快速复刻整套方案的读者", font_size=11, bold=True, space_after=4)
    add_para(doc, "附录 —— 文件清单与技术架构", font_size=11, bold=True, space_after=4)
    doc.add_page_break()

    # ================================================================
    # 第一部分：方法论总结
    # ================================================================
    add_heading_custom(doc, "第一部分　方法论总结", level=1, color=RGBColor(0x1a, 0x56, 0xc4))
    add_para(doc, "本部分提炼核心设计思路和关键决策，适合想快速理解方案原理的读者。", font_size=11, italic=True, color=RGBColor(0x80, 0x80, 0x80), space_after=12)

    # --- 1. 问题背景 ---
    add_heading_custom(doc, "一、问题背景", level=2)
    add_para(doc, "在使用腾讯 IMA 知识库的过程中，随着资料不断积累，逐渐出现以下痛点：", font_size=11, space_after=8)

    add_table(doc,
        ["痛点", "具体表现", "影响"],
        [
            ["内容杂乱无章", "37篇资料全部堆在根目录，无分类、无标签、无文件夹", "查找困难，无法按主题定位"],
            ["重复与无关内容", "同一文件被不同公众号转载导致重复；菜谱图片等与主题无关", "干扰检索，降低知识库质量"],
            ["无法精准检索", "撰写文档时无法按主题筛选素材，全量阅读效率极低", "知识库\u201c有料用不上\u201d"],
            ["API 限制", "IMA API 不支持删除、移动、建文件夹、打标签", "无法通过程序自动整理"],
        ],
        col_widths=[4, 7, 5]
    )
    add_para(doc, "", space_after=8)

    # --- 2. 核心创新 ---
    add_heading_custom(doc, "二、核心创新：索引代替物理文件夹", level=2)
    add_para(doc, "面对 IMA API 不支持文件移动和文件夹创建的限制，传统\u201c移动文件到文件夹\u201d的整理方式无法实现。本方案的核心创新在于：", font_size=11, space_after=8)

    add_para(doc, "用分类索引文档（Word/Markdown）记录每篇资料的分类归属，不移动文件本身。", font_size=12, bold=True, space_after=8)

    add_para(doc, "对比两种方案：", font_size=11, bold=True, space_after=6)
    add_table(doc,
        ["维度", "传统方案（物理移动）", "创新方案（索引追踪）"],
        [
            ["实现方式", "在 IMA 中创建文件夹，移动文件", "本地维护分类索引文档，记录 media_id 与分类的映射"],
            ["API 依赖", "需要移动/建文件夹 API（IMA 不支持）", "只需列表 API + 内容获取 API（IMA 支持）"],
            ["自动化程度", "需要手动移动（API 限制）", "全自动扫描、分类、更新索引"],
            ["可追溯性", "文件移动后历史信息丢失", "索引文档记录完整分类历史"],
            ["RAG 兼容性", "需要按文件夹逐个查询", "按分类索引精准定位，直接拉取全文"],
            ["维护成本", "每次新增都要手动归档", "每周自动扫描，零手动"],
        ],
        col_widths=[3, 6.5, 6.5]
    )
    add_para(doc, "", space_after=8)

    # --- 3. RAG 链路 ---
    add_heading_custom(doc, "三、RAG 检索增强生成链路设计", level=2)
    add_para(doc, "整套方案最终服务于一个目标：基于知识库素材，生成符合个性化需求的文档。RAG 链路如下：", font_size=11, space_after=8)

    add_para(doc, "链路流程：", font_size=11, bold=True, space_after=4)
    add_para(doc, "\u2460 用户指定分类（如\u201c用 AI+体育 类别的资料帮我写文档\u201d）", font_size=11, space_after=4)
    add_para(doc, "② 读取该分类的索引文档，获取所有资料的 media_id 列表", font_size=11, space_after=4)
    add_para(doc, "③ 逐篇调用 fetch_media_content 拉取全文内容", font_size=11, space_after=4)
    add_para(doc, "④ 以全文作为上下文，生成符合用户需求的文档", font_size=11, space_after=8)

    add_para(doc, "两种 RAG 模式：", font_size=11, bold=True, space_after=6)
    add_table(doc,
        ["模式", "适用场景", "操作方式", "优势"],
        [
            ["全量拉取", "分类下资料较少（\u22645篇）", "拉取该分类所有资料全文", "上下文完整，生成质量高"],
            ["搜索优先", "分类下资料较多（>5篇）", "先用 search_knowledge 搜索关键词，取 Top 5 再拉全文", "减少 API 调用，聚焦最相关内容"],
        ],
        col_widths=[3, 4, 6, 3]
    )
    add_para(doc, "", space_after=8)

    # --- 4. 三个关键决策 ---
    add_heading_custom(doc, "四、三个关键决策转折点", level=2)

    add_para(doc, "决策一：放弃物理整理，改用索引追踪", font_size=12, bold=True, color=RGBColor(0x1a, 0x56, 0xc4), space_after=4)
    add_para(doc, "在发现 IMA API 不支持删除、移动、建文件夹后，没有放弃整理目标，而是转换思路：既然不能移动物理文件，就用逻辑分类代替。这是整套方案的基石。", font_size=11, space_after=8)

    add_para(doc, "决策二：选择 Word 文档作为索引载体", font_size=12, bold=True, color=RGBColor(0x1a, 0x56, 0xc4), space_after=4)
    add_para(doc, "在 Word 和 txt 之间选择 Word，原因：学术/教育场景下 Word 更专业、可打印、可批注；表格形式清晰展示标题/来源/日期；同时保留 JSON 作为程序后台。后来在 skill 化时又增加了 Markdown 格式输出，兼顾程序化处理需求。", font_size=11, space_after=8)

    add_para(doc, "决策三：打包成可复用 Skill", font_size=12, bold=True, color=RGBColor(0x1a, 0x56, 0xc4), space_after=4)
    add_para(doc, "在验证方案可行后，主动分析 6 个不足（全量拉取效率低、仅靠标题分类、不支持多标签、无内容缓存、配置硬编码、只输出 Word），将其改进后打包成标准 skill，让其他人可以直接使用。这使方案从个人工具升级为可复用的通用方案。", font_size=11, space_after=8)

    # --- 5. 通用原则 ---
    add_heading_custom(doc, "五、通用原则（适用于其他知识库场景）", level=2)
    add_para(doc, "以下原则不仅适用于 IMA 知识库，也适用于其他知识管理场景：", font_size=11, space_after=8)

    principles = [
        ("先扫描分析，再设计分类", "在制定分类规则前，先全面扫描现有内容，用数据驱动分类设计，而非凭空想象分类结构。确保分类规则覆盖所有实际内容。"),
        ("分类规则要 MECE", "分类之间互斥（一篇资料只归入一个最核心类别），合起来穷尽（所有内容都有归属）。无法归类的内容设\u201c待确认\u201d类别，而非强行归入。"),
        ("逻辑分类优于物理整理", "当工具 API 有限时，用索引/标记/追踪文件实现逻辑分类，不依赖物理移动。更灵活、可追溯、可自动化。"),
        ("自动化扫描 + 人工归档", "机器负责扫描、分类、生成报告；人负责最终确认和特殊情况处理。既高效又可控。"),
        ("索引文档即虚拟文件夹", "用文档记录归属关系，检索时按索引定位再拉取全文。索引文档本身就是 RAG 的检索层。"),
    ]
    for i, (title, desc) in enumerate(principles, 1):
        add_para(doc, f"{i}. {title}", font_size=11, bold=True, space_after=4)
        add_para(doc, f"   {desc}", font_size=11, space_after=8)

    doc.add_page_break()

    # ================================================================
    # 第二部分：完整沟通过程还原
    # ================================================================
    add_heading_custom(doc, "第二部分　完整沟通过程还原", level=1, color=RGBColor(0x1a, 0x56, 0xc4))
    add_para(doc, "本部分还原从需求提出到方案成型的完整沟通过程，适合想了解每个决策细节的读者。", font_size=11, italic=True, color=RGBColor(0x80, 0x80, 0x80), space_after=12)

    # --- 第一轮 ---
    add_heading_custom(doc, "第一轮：需求发现与现状盘点", level=2)

    add_para(doc, "用户需求：", font_size=11, bold=True, space_after=4)
    add_para(doc, "\"你可以整理 IMA 知识库么\"——一个看似简单的问题，开启了整套方案的探索。", font_size=11, space_after=8)

    add_para(doc, "AI 分析与行动：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "加载 IMA MCP 工具，确认可用能力：列表查询、搜索、添加知识、导入URL、创建media、获取内容")
    add_bullet(doc, "拉取知识库列表：发现 2 个个人知识库")
    add_bullet(doc, "逐个扫描内容：共 37 篇资料")
    add_bullet(doc, "按主题归类：AI+体育（10篇）、AI+教育研究（7篇）、AI教学工具实操（4篇）、教育政策文件（5篇）、AI见解与培训（3篇）、无关内容（3篇）")
    add_bullet(doc, "生成可视化盘点报告（SVG 图表）")

    add_para(doc, "关键发现：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "重复内容：复旦AI教学指引1.0 有两篇（不同公众号转载同一文件）")
    add_bullet(doc, "内容混杂：菜谱图片和留学生新闻跟知识库主题完全不相关")
    add_bullet(doc, "零分类：37 篇全部堆在根目录，没有标签、没有文件夹")
    add_bullet(doc, "API 限制：IMA 接口不支持删除、移动、建文件夹、打标签")

    add_para(doc, "产出物：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "知识库内容盘点可视化报告")
    add_bullet(doc, "问题诊断清单（3个问题）")

    doc.add_page_break()

    # --- 第二轮 ---
    add_heading_custom(doc, "第二轮：方案设计与体系搭建", level=2)

    add_para(doc, "用户需求：", font_size=11, bold=True, space_after=4)
    add_para(doc, "用户提出了更明确的目标：定期整理知识库，自己只负责添加资料，AI 负责分类整理，最终用于 RAG 检索增强生成，确保文档质量。", font_size=11, space_after=8)

    add_para(doc, "AI 分析与行动：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "检查 IMA 是否支持创建文件夹和笔记（发现不支持）")
    add_bullet(doc, "设计三层方案：分类追踪文件 + 整理指南 + 定期自动化任务")
    add_bullet(doc, "创建 tracker.json 基线文件，记录全部 37 篇内容的 media_id、标题、分类、首次发现时间")
    add_bullet(doc, "生成 IMA知识库整理指南 HTML 文档，列出每篇内容应归入的文件夹或建议删除")
    add_bullet(doc, "创建每周日 10:00 的自动化任务，自动扫描新增内容、分类、生成报告")

    add_para(doc, "关键决策：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "分类规则设计：6 个类别（5个有效类别 + 1个待清理），基于关键词匹配")
    add_bullet(doc, "追踪文件格式选择 JSON（程序友好），而非 CSV 或 Excel")
    add_bullet(doc, "自动化频率选择每周一次（周日 10:00），平衡及时性和打扰频率")

    add_para(doc, "产出物：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "tracker.json 分类追踪文件（37篇基线数据）")
    add_bullet(doc, "IMA知识库整理指南.html")
    add_bullet(doc, "automation-1785299014871 自动化任务（每周日 10:00）")
    add_bullet(doc, "reports/ 报告目录")

    doc.add_page_break()

    # --- 第三轮 ---
    add_heading_custom(doc, "第三轮：方案调整——索引代替物理移动", level=2)

    add_para(doc, "用户需求：", font_size=11, bold=True, space_after=4)
    add_para(doc, "用户提出了关键思路转变：既然不能物理移动文件，是否可以用文档记录归属？将资料按子类别分开生成列表文档，定期更新，撰写文档时指定使用哪个类别的资料。", font_size=11, space_after=8)

    add_para(doc, "AI 分析：", font_size=11, bold=True, space_after=4)
    add_para(doc, "AI 确认这个方案完全可行，而且比物理移动更聪明。原因：IMA 的 fetch_media_content 可以按 media_id 拉取任意一篇资料的全文，所以 RAG 链路完整：指定分类 → 读取索引获取 media_id → 拉取全文 → 生成文档。这比 IMA 自带语义搜索更精准，因为是先按分类过滤再取全文，避免了跨主题噪音。", font_size=11, space_after=8)

    add_para(doc, "AI 行动：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "安装 python-docx 库")
    add_bullet(doc, "编写 generate_index_docs.py 脚本，读取 tracker.json 生成 5 个 Word 索引文档")
    add_bullet(doc, "每个文档包含表格：序号、标题、来源、导入知识库日期")
    add_bullet(doc, "更新自动化任务，增加步骤6：扫描分类后自动运行脚本重新生成文档")
    add_bullet(doc, "用户指定输出目录后，更新脚本增加自动复制功能")

    add_para(doc, "关键决策：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "文件格式选择 Word（学术场景更专业、可打印、可批注），同时保留 JSON 作为程序后台")
    add_bullet(doc, "3 篇无关内容和 1 篇重复内容不纳入索引文档（归在\u201c待清理\u201d类别，不影响 RAG 质量）")
    add_bullet(doc, "脚本增加自动复制到用户指定目录功能，每周自动更新无需手动操作")

    add_para(doc, "产出物：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "5 个分类索引 Word 文档（AI与体育 10篇、AI与教育研究 9篇、AI教学工具实操 4篇、教育政策文件 5篇、AI见解与培训 3篇）")
    add_bullet(doc, "generate_index_docs.py 文档生成脚本（含自动复制功能）")
    add_bullet(doc, "更新后的自动化任务（增加文档生成步骤）")

    doc.add_page_break()

    # --- 第四轮 ---
    add_heading_custom(doc, "第四轮：改进与 Skill 化", level=2)

    add_para(doc, "用户需求：", font_size=11, bold=True, space_after=4)
    add_para(doc, "用户询问改进建议，并问是否可以做成供其他人直接调用的 skill。", font_size=11, space_after=8)

    add_para(doc, "AI 分析的 6 个不足：", font_size=11, bold=True, space_after=4)
    add_table(doc,
        ["#", "问题", "改进方向"],
        [
            ["1", "RAG 是全量拉取而非检索后拉取", "先用 search_knowledge 按主题搜索，只拉最相关的几篇"],
            ["2", "分类仅靠标题关键词", "先 fetch_media_content 取摘要，再做 LLM 分类"],
            ["3", "不支持多标签", "允许一篇资料同时出现在多个分类索引中"],
            ["4", "无内容缓存", "本地缓存已拉取的全文内容，避免重复 API 调用"],
            ["5", "配置硬编码", "抽成 config.json，其他用户只需改配置"],
            ["6", "只输出 Word", "同时输出 Markdown 版本，方便程序化处理"],
        ],
        col_widths=[1, 6, 9]
    )
    add_para(doc, "", space_after=8)

    add_para(doc, "AI 行动：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "加载 skill-creator 技能，学习标准 skill 创建流程")
    add_bullet(doc, "使用 init_skill.py 初始化 skill 骨架")
    add_bullet(doc, "创建 6 个核心文件：SKILL.md + generate_index_docs.py（配置驱动版）+ 3 份参考文档 + config 模板")
    add_bullet(doc, "使用 quick_validate.py 验证 skill 结构")
    add_bullet(doc, "使用 package_skill.py 打包为 ima-kb-organizer.zip")

    add_para(doc, "Skill 包含内容：", font_size=11, bold=True, space_after=4)
    add_table(doc,
        ["文件", "作用"],
        [
            ["SKILL.md", "3 个工作流：首次设置 / 定期扫描 / RAG 生成"],
            ["scripts/generate_index_docs.py", "配置驱动的文档生成器（Word + Markdown）"],
            ["references/ima_api_reference.md", "IMA 全部工具能力 + 限制说明"],
            ["references/category_rules.md", "分类规则设计指南（MECE、颗粒度、多标签）"],
            ["references/rag_workflow.md", "RAG 两种模式选择策略 + 缓存方案"],
            ["assets/config_template.json", "配置模板，新用户填空即可"],
        ],
        col_widths=[7, 9]
    )
    add_para(doc, "", space_after=8)

    add_para(doc, "产出物：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "ima-kb-organizer skill（已安装到用户级 skills 目录）")
    add_bullet(doc, "ima-kb-organizer.zip（可分发的打包文件）")

    doc.add_page_break()

    # ================================================================
    # 第三部分：简明步骤手册
    # ================================================================
    add_heading_custom(doc, "第三部分　简明步骤手册", level=1, color=RGBColor(0x1a, 0x56, 0xc4))
    add_para(doc, "本部分提供 10 步快速复刻清单，适合想直接动手的读者。每步含操作要点和注意事项。", font_size=11, italic=True, color=RGBColor(0x80, 0x80, 0x80), space_after=12)

    steps = [
        ("步骤1：连接 IMA 知识库",
         "在 WorkBuddy 中连接 ima-mcp 连接器，确保 IMA 知识库可访问。",
         "确认连接器状态为 connected。需要至少一个可访问的知识库。"),
        ("步骤2：安装 Skill",
         "安装 ima-kb-organizer skill（双击 zip 文件或通过 WorkBuddy 安装）。",
         "安装后会自动出现在技能管理面板中。"),
        ("步骤3：首次扫描",
         "对 AI 说\"帮我整理 IMA 知识库\"，触发 skill 的首次设置工作流。",
         "AI 会扫描全部内容，生成内容盘点。耐心等待扫描完成。"),
        ("步骤4：确认分类规则",
         "与 AI 讨论确认分类规则。AI 会根据你的内容推荐分类，你可以调整。",
         "分类规则要 MECE：互斥且穷尽。关键词要覆盖你的领域。"),
        ("步骤5：生成配置文件",
         "AI 根据确认的分类规则生成 config.json，包含 KB ID、输出目录、分类定义。",
         "检查 KB ID 和输出目录是否正确。输出目录需提前创建。"),
        ("步骤6：生成索引文档",
         "AI 运行 generate_index_docs.py，为每个分类生成 Word + Markdown 索引文档。",
         "确认每个文档中的资料数量和分类是否正确。"),
        ("步骤7：创建自动化任务",
         "AI 创建每周定期扫描任务，自动检测新增内容、分类、更新索引文档。",
         "确认任务频率和时间合适。可以在 WorkBuddy 设置中查看和管理。"),
        ("步骤8：日常使用——添加资料",
         "看到好文章，直接添加到 IMA 知识库根目录。不需要手动分类。",
         "把知识库根目录当作\"收件箱\"，只管往里扔。"),
        ("步骤9：查看整理报告",
         "每周自动化任务运行后，查看生成的整理报告，了解新增内容的分类情况。",
         "如果分类有误，可以告知 AI 调整规则或手动修正。"),
        ("步骤10：RAG 文档生成",
         "需要撰写文档时，告诉 AI\"用 [分类名] 的资料帮我写关于 [主题] 的文档\"。",
         "AI 会自动拉取该分类下的资料全文作为素材，生成符合需求的文档。资料多时自动切换为搜索优先模式。"),
    ]

    for title, desc, note in steps:
        add_para(doc, title, font_size=12, bold=True, color=RGBColor(0x1a, 0x56, 0xc4), space_after=4)
        add_para(doc, f"操作：{desc}", font_size=11, space_after=4)
        add_para(doc, f"注意：{note}", font_size=10, italic=True, color=RGBColor(0x80, 0x80, 0x80), space_after=10)

    doc.add_page_break()

    # ================================================================
    # 附录
    # ================================================================
    add_heading_custom(doc, "附录", level=1, color=RGBColor(0x1a, 0x56, 0xc4))

    add_heading_custom(doc, "附录 A：文件清单", level=2)
    add_table(doc,
        ["文件", "路径", "用途"],
        [
            ["SKILL.md", "~/.workbuddy/skills/ima-kb-organizer/", "Skill 主文件，定义工作流"],
            ["generate_index_docs.py", "scripts/", "文档生成脚本（配置驱动）"],
            ["config_template.json", "assets/", "配置模板"],
            ["ima_api_reference.md", "references/", "IMA API 能力与限制"],
            ["category_rules.md", "references/", "分类规则设计指南"],
            ["rag_workflow.md", "references/", "RAG 工作流说明"],
            ["tracker.json", ".workbuddy/ima-tracker/", "分类追踪文件（基线数据）"],
            ["category_docs/*.docx", ".workbuddy/ima-tracker/", "分类索引 Word 文档"],
            ["reports/*.md", ".workbuddy/ima-tracker/", "整理报告"],
        ],
        col_widths=[5, 6, 5]
    )
    add_para(doc, "", space_after=8)

    add_heading_custom(doc, "附录 B：技术架构", level=2)
    add_para(doc, "数据流：", font_size=11, bold=True, space_after=4)
    add_para(doc, "IMA 知识库 → get_knowledge_list（扫描）→ tracker.json（追踪）→ generate_index_docs.py（生成）→ Word/Markdown 索引文档（输出）", font_size=11, space_after=8)
    add_para(doc, "RAG 流：", font_size=11, bold=True, space_after=4)
    add_para(doc, "用户指定分类 → 读取索引文档 → fetch_media_content（拉全文）→ AI 生成文档", font_size=11, space_after=8)
    add_para(doc, "自动化流：", font_size=11, bold=True, space_after=4)
    add_para(doc, "每周日 10:00 → 扫描新增内容 → 分类 → 更新 tracker.json → 重新生成索引文档 → 生成报告", font_size=11, space_after=8)

    add_heading_custom(doc, "附录 C：版本历史", level=2)
    add_table(doc,
        ["版本", "日期", "变更内容"],
        [
            ["v1.0", "2026-07-29", "初始版本：分类追踪 + 索引文档 + 自动化任务 + Skill 封装"],
        ],
        col_widths=[2, 3, 11]
    )
    add_para(doc, "", space_after=8)

    # ===== 末页 =====
    add_para(doc, "", space_after=30)
    add_para(doc, "\u2014\u2014 文档结束 \u2014\u2014", font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x80, 0x80, 0x80))
    add_para(doc, f"\u00a9 2026 sus-yugaohe | CC BY-NC 4.0", font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x80, 0x80, 0x80))

    # 保存
    output_path = os.path.join(OUTPUT_DIR, "IMA知识库智能整理_方法论与实践复盘.docx")
    doc.save(output_path)
    print(f"文档已生成：{output_path}")
    return output_path


if __name__ == "__main__":
    generate_document()
