# -*- coding: utf-8 -*-
"""
生成《IMA知识库整理技能使用说明》Word 文档
作者: sus-yugaohe
许可: CC BY-NC 4.0
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# 辅助函数（与第一份文档相同）
# ============================================================

def set_cell_font(cell, font_name="微软雅黑", font_size=10, bold=False, color=None):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_custom(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

def add_para(doc, text, font_size=11, bold=False, italic=False, alignment=None, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_bullet(doc, text, level=0, font_size=11):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    p.paragraph_format.line_spacing = 1.5
    return p

def add_numbered(doc, text, font_size=11):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    p.paragraph_format.line_spacing = 1.5
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_font(cell, font_size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
            set_cell_font(row.cells[i], font_size=10)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table

def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Consolas")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    # 添加灰色底色
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F0F0F0')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    return p

def add_title_page(doc, title, subtitle, author="sus-yugaohe"):
    for _ in range(6):
        add_para(doc, "", font_size=12, space_after=6)
    add_para(doc, title, font_size=28, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, color=RGBColor(0x1a, 0x56, 0xc4))
    add_para(doc, subtitle, font_size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30, color=RGBColor(0x40, 0x40, 0x40))
    for _ in range(8):
        add_para(doc, "", font_size=12, space_after=6)
    add_para(doc, f"作者：{author}", font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "版本：v1.0", font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "2026年7月", font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "\u00a9 2026 sus-yugaohe | CC BY-NC 4.0", font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x80, 0x80, 0x80))
    doc.add_page_break()

def add_copyright_page(doc, title, author="sus-yugaohe", version="v1.0"):
    doc.add_page_break()
    add_para(doc, "版权声明", font_size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_para(doc, f"作品名称：{title}", font_size=12, space_after=6)
    add_para(doc, f"作者：{author}", font_size=12, space_after=6)
    add_para(doc, f"版本：{version}", font_size=12, space_after=6)
    add_para(doc, f"发布日期：2026年7月", font_size=12, space_after=20)
    add_para(doc, "—" * 40, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(doc, f"\u00a9 2026 {author} 保留所有权利", font_size=12, bold=True, space_after=12)
    add_para(doc, "本作品采用知识共享署名-非商业性使用 4.0 国际许可协议（CC BY-NC 4.0）进行许可。", font_size=11, space_after=8)
    add_para(doc, "许可内容：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "署名（BY）：使用者必须在适当位置注明原作者 sus-yugaohe 及作品名称")
    add_bullet(doc, "非商业性使用（NC）：本作品及其衍生作品不得用于商业目的")
    add_bullet(doc, "共享：使用者可以复制、分发、传播本作品")
    add_bullet(doc, "改编：使用者可以创作基于本作品的衍生作品（须同样署名且非商业）")
    add_para(doc, "使用限制：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "未经作者书面授权，禁止将本作品用于任何商业用途，包括但不限于销售、付费课程、商业培训")
    add_bullet(doc, "引用本作品时，必须注明出处：\"sus-yugaohe, IMA知识库整理技能使用说明, 2026\"")
    add_bullet(doc, "改编作品须以相同或兼容的许可协议发布")
    add_bullet(doc, "作者不对本作品的适用性作任何担保，使用者自行承担使用风险")
    add_para(doc, "完整许可协议文本请访问：", font_size=11, space_after=4)
    add_para(doc, "https://creativecommons.org/licenses/by-nc/4.0/legalcode", font_size=10, italic=True, space_after=12)
    add_para(doc, "—" * 40, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(doc, "联系方式与反馈：通过 WorkBuddy 平台联系作者 sus-yugaohe", font_size=10, italic=True, color=RGBColor(0x80, 0x80, 0x80))
    doc.add_page_break()


# ============================================================
# 文档生成
# ============================================================

def generate_document():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # ===== 封面 =====
    add_title_page(doc, "IMA知识库整理技能", "使用说明\n—— 安装、配置与日常使用指南")

    # ===== 版权页 =====
    add_copyright_page(doc, "IMA知识库整理技能使用说明")

    # ===== 目录 =====
    add_heading_custom(doc, "目录", level=1, color=RGBColor(0x1a, 0x56, 0xc4))
    chapters = [
        "第一章　概述",
        "第二章　前置条件",
        "第三章　安装步骤",
        "第四章　配置说明",
        "第五章　日常工作流",
        "第六章　RAG 文档生成",
        "第七章　定期任务管理",
        "第八章　常见问题",
        "第九章　版权与许可",
    ]
    for ch in chapters:
        add_para(doc, ch, font_size=11, space_after=4)
    doc.add_page_break()

    # ===== 第一章 概述 =====
    add_heading_custom(doc, "第一章　概述", level=1, color=RGBColor(0x1a, 0x56, 0xc4))

    add_heading_custom(doc, "1.1 功能简介", level=2)
    add_para(doc, "IMA知识库整理技能（ima-kb-organizer）是一套基于 WorkBuddy 平台的自动化知识库管理方案，解决腾讯 IMA 知识库\u201c只进不出、无法分类整理\u201d的痛点。", font_size=11, space_after=8)
    add_para(doc, "核心功能：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "自动扫描 IMA 知识库内容，按自定义规则分类")
    add_bullet(doc, "生成分类索引文档（Word + Markdown），替代物理文件夹")
    add_bullet(doc, "定期自动整理，检测新增内容并分类归档")
    add_bullet(doc, "基于分类索引的精准 RAG 文档生成")
    add_bullet(doc, "支持多知识库、多标签分类")

    add_heading_custom(doc, "1.2 适用场景", level=2)
    add_table(doc,
        ["场景", "说明"],
        [
            ["学术研究", "管理论文、报告、政策文件，按研究方向分类检索"],
            ["教学备课", "收集教学素材，按课程/章节分类，快速生成教案"],
            ["知识管理", "个人/团队知识库整理，定期归档新增内容"],
            ["文档撰写", "基于指定分类的资料库，AI 辅助生成文档"],
            ["项目资料", "项目文档分类管理，按阶段/模块组织检索"],
        ],
        col_widths=[4, 12]
    )
    add_para(doc, "", space_after=8)

    add_heading_custom(doc, "1.3 技术架构", level=2)
    add_para(doc, "本技能采用\u201c逻辑分类\u201d架构，不依赖物理文件夹：", font_size=11, space_after=8)
    add_code_block(doc, "IMA 知识库（数据源）\n    |\n    v\n扫描分类（get_knowledge_list + 关键词匹配）\n    |\n    v\ntracker.json（分类追踪文件）\n    |\n    v\ngenerate_index_docs.py（文档生成器）\n    |\n    v\nWord + Markdown 索引文档（分类索引）\n    |\n    v\nRAG 生成（fetch_media_content → 全文 → AI 生成文档）")
    add_para(doc, "", space_after=8)

    doc.add_page_break()

    # ===== 第二章 前置条件 =====
    add_heading_custom(doc, "第二章　前置条件", level=1, color=RGBColor(0x1a, 0x56, 0xc4))

    add_heading_custom(doc, "2.1 环境要求", level=2)
    add_table(doc,
        ["项目", "要求", "说明"],
        [
            ["WorkBuddy", "最新版本", "AI 助手平台，提供 Skill 运行环境"],
            ["IMA 知识库连接器", "已连接", "在 WorkBuddy 连接器管理中连接 ima-mcp"],
            ["Python", "3.13+（推荐）", "用于运行文档生成脚本"],
            ["python-docx", "已安装", "Word 文档生成依赖库"],
            ["IMA 知识库", "至少 1 个", "需要有可访问的个人知识库"],
        ],
        col_widths=[4, 4, 8]
    )
    add_para(doc, "", space_after=8)

    add_heading_custom(doc, "2.2 连接器配置", level=2)
    add_para(doc, "在 WorkBuddy 中连接 IMA 知识库：", font_size=11, space_after=6)
    add_numbered(doc, "打开 WorkBuddy 设置 → 连接器管理")
    add_numbered(doc, "找到\"ima 知识库\"连接器，点击连接")
    add_numbered(doc, "授权后，连接器状态显示为\"已连接\"")
    add_numbered(doc, "在对话中可以验证：让 AI 帮你列出知识库内容")
    add_para(doc, "注意：连接器需要使用腾讯账号登录 IMA 后才能正常工作。", font_size=10, italic=True, color=RGBColor(0x80, 0x80, 0x80), space_after=8)

    add_heading_custom(doc, "2.3 Python 依赖安装", level=2)
    add_para(doc, "如果使用 WorkBuddy 内置 Python，执行以下命令安装依赖：", font_size=11, space_after=4)
    add_code_block(doc, 'pip install python-docx')
    add_para(doc, "如果使用系统 Python，确保 pip 可用后执行相同命令。", font_size=10, italic=True, color=RGBColor(0x80, 0x80, 0x80), space_after=8)

    doc.add_page_break()

    # ===== 第三章 安装步骤 =====
    add_heading_custom(doc, "第三章　安装步骤", level=1, color=RGBColor(0x1a, 0x56, 0xc4))

    add_para(doc, "从拿到 zip 文件到开始使用，共 5 步：", font_size=11, space_after=12)

    add_heading_custom(doc, "3.1 安装 Skill", level=2)
    add_para(doc, "方式一：通过 WorkBuddy 推荐市场安装（如果已上线）", font_size=11, bold=True, space_after=4)
    add_para(doc, "在 WorkBuddy 中对 AI 说\"帮我安装 ima-kb-organizer 技能\"，AI 会从推荐市场搜索并安装。", font_size=11, space_after=8)

    add_para(doc, "方式二：通过 zip 文件安装", font_size=11, bold=True, space_after=4)
    add_numbered(doc, "获取 ima-kb-organizer.zip 文件")
    add_numbered(doc, "在 WorkBuddy 中打开技能管理面板")
    add_numbered(doc, "选择\"导入 Skill\"，选择 zip 文件")
    add_numbered(doc, "等待安装完成，确认技能出现在列表中")

    add_heading_custom(doc, "3.2 触发首次扫描", level=2)
    add_para(doc, "安装完成后，在 WorkBuddy 对话中对 AI 说：", font_size=11, space_after=4)
    add_code_block(doc, "帮我整理 IMA 知识库")
    add_para(doc, "AI 将触发 ima-kb-organizer 技能，执行首次设置工作流。", font_size=11, space_after=8)

    add_heading_custom(doc, "3.3 确认分类规则", level=2)
    add_para(doc, "AI 会根据你的知识库内容推荐分类规则。你可以：", font_size=11, space_after=6)
    add_bullet(doc, "接受 AI 推荐的分类（适合快速上手）")
    add_bullet(doc, "调整分类名称和关键词（适合有特定需求）")
    add_bullet(doc, "增加或删除分类（适合自定义场景）")
    add_para(doc, "分类规则确认后，AI 会生成 config.json 配置文件。", font_size=11, space_after=8)

    add_heading_custom(doc, "3.4 生成索引文档", level=2)
    add_para(doc, "AI 自动运行文档生成脚本，为每个分类生成 Word 和 Markdown 索引文档。文档包含：", font_size=11, space_after=6)
    add_bullet(doc, "序号")
    add_bullet(doc, "资料标题")
    add_bullet(doc, "来源（公众号/网站名）")
    add_bullet(doc, "导入知识库日期")
    add_para(doc, "文档默认保存到 .workbuddy/ima-tracker/category_docs/ 目录，可配置自定义输出目录。", font_size=11, space_after=8)

    add_heading_custom(doc, "3.5 创建自动化任务", level=2)
    add_para(doc, "AI 会自动创建每周定期扫描任务。默认每周日 10:00 运行，可调整。", font_size=11, space_after=8)

    doc.add_page_break()

    # ===== 第四章 配置说明 =====
    add_heading_custom(doc, "第四章　配置说明", level=1, color=RGBColor(0x1a, 0x56, 0xc4))

    add_para(doc, "所有配置集中在 config.json 文件中。以下是完整字段说明：", font_size=11, space_after=8)

    add_heading_custom(doc, "4.1 配置文件示例", level=2)
    add_code_block(doc, """{
  "knowledge_bases": [
    {
      "id": "0019ed21e8806f81",
      "name": "我的知识库",
      "is_primary": true
    }
  ],
  "output_dir": ".workbuddy/ima-tracker/category_docs",
  "extra_output_dirs": [
    "D:/my-docs/ima-index"
  ],
  "categories": [
    {
      "name": "AI+体育",
      "keywords": ["体育", "运动", "穿戴", "传感器", "体测"],
      "description": "AI 在体育领域的应用"
    }
  ],
  "scan_frequency": "WEEKLY",
  "scan_day": "SU",
  "scan_time": "10:00",
  "rag_mode": "auto",
  "enable_cache": true,
  "cache_dir": ".workbuddy/ima-tracker/cache"
}""")

    add_heading_custom(doc, "4.2 字段说明", level=2)
    add_table(doc,
        ["字段", "类型", "说明"],
        [
            ["knowledge_bases", "数组", "知识库列表，每个含 id、name、is_primary"],
            ["knowledge_bases[].id", "字符串", "IMA 知识库 ID（必填）"],
            ["knowledge_bases[].name", "字符串", "知识库显示名称"],
            ["knowledge_bases[].is_primary", "布尔", "是否为主知识库（用作收件箱）"],
            ["output_dir", "字符串", "索引文档默认输出目录"],
            ["extra_output_dirs", "数组", "额外输出目录（自动复制到这些目录）"],
            ["categories", "数组", "分类规则列表"],
            ["categories[].name", "字符串", "分类名称"],
            ["categories[].keywords", "数组", "分类关键词（用于自动匹配）"],
            ["categories[].description", "字符串", "分类描述"],
            ["scan_frequency", "字符串", "扫描频率：WEEKLY / DAILY"],
            ["scan_day", "字符串", "扫描日：MO/TU/WE/TH/FR/SA/SU"],
            ["scan_time", "字符串", "扫描时间（24小时制）"],
            ["rag_mode", "字符串", "RAG 模式：auto / full / search_first"],
            ["enable_cache", "布尔", "是否启用内容缓存"],
            ["cache_dir", "字符串", "缓存目录路径"],
        ],
        col_widths=[5, 2, 9]
    )
    add_para(doc, "", space_after=8)

    add_heading_custom(doc, "4.3 RAG 模式说明", level=2)
    add_table(doc,
        ["模式", "说明", "适用场景"],
        [
            ["auto", "自动选择：分类下 ≤5 篇用全量拉取，>5 篇用搜索优先", "推荐默认值"],
            ["full", "始终全量拉取该分类所有资料全文", "资料少、需要完整上下文"],
            ["search_first", "始终先搜索关键词，取 Top 5 再拉全文", "资料多、需要聚焦"],
        ],
        col_widths=[3, 8, 5]
    )
    add_para(doc, "", space_after=8)

    doc.add_page_break()

    # ===== 第五章 日常工作流 =====
    add_heading_custom(doc, "第五章　日常工作流", level=1, color=RGBColor(0x1a, 0x56, 0xc4))

    add_heading_custom(doc, "5.1 资料收集", level=2)
    add_para(doc, "日常看到优秀的文章、报告、资料时：", font_size=11, space_after=6)
    add_numbered(doc, "在 IMA 客户端或网页版中，将资料添加到你的主知识库（收件箱）")
    add_numbered(doc, "不需要手动分类、不需要选择文件夹")
    add_numbered(doc, "资料会自动被纳入下一轮扫描")
    add_para(doc, "提示：主知识库就是 config.json 中 is_primary 为 true 的知识库。", font_size=10, italic=True, color=RGBColor(0x80, 0x80, 0x80), space_after=8)

    add_heading_custom(doc, "5.2 自动扫描", level=2)
    add_para(doc, "每周（默认周日 10:00），自动化任务会自动执行：", font_size=11, space_after=6)
    add_bullet(doc, "扫描知识库全部内容")
    add_bullet(doc, "与本地追踪文件比对，找出新增内容")
    add_bullet(doc, "按分类规则对新增内容自动分类")
    add_bullet(doc, "更新 tracker.json 追踪文件")
    add_bullet(doc, "重新生成所有分类索引文档（Word + Markdown）")
    add_bullet(doc, "自动复制到配置的输出目录")
    add_bullet(doc, "生成整理报告（Markdown 格式）")

    add_heading_custom(doc, "5.3 查看整理报告", level=2)
    add_para(doc, "每次扫描后，报告保存在：", font_size=11, space_after=4)
    add_code_block(doc, ".workbuddy/ima-tracker/reports/report-YYYY-MM-DD.md")
    add_para(doc, "报告包含：", font_size=11, space_after=6)
    add_bullet(doc, "扫描时间和新增内容数量")
    add_bullet(doc, "每篇新增内容的标题、来源、分类结果")
    add_bullet(doc, "疑似重复内容提醒")
    add_bullet(doc, "无关内容\u201c建议删除\u201d标记")
    add_para(doc, "如果分类有误，可以在对话中告知 AI 调整分类规则。", font_size=11, space_after=8)

    add_heading_custom(doc, "5.4 撰写文档", level=2)
    add_para(doc, "需要撰写文档时，在 WorkBuddy 对话中指定分类和主题：", font_size=11, space_after=4)
    add_code_block(doc, "用 AI+教育研究 分类的资料，帮我写一篇关于\"教育智能体在个性化学习中的应用\"的文档")
    add_para(doc, "AI 会自动：", font_size=11, space_after=6)
    add_bullet(doc, "读取该分类的索引文档，获取所有 media_id")
    add_bullet(doc, "按 RAG 模式拉取全文（全量或搜索优先）")
    add_bullet(doc, "以全文作为素材上下文生成文档")
    add_bullet(doc, "生成的文档会引用素材来源")

    doc.add_page_break()

    # ===== 第六章 RAG 文档生成 =====
    add_heading_custom(doc, "第六章　RAG 文档生成", level=1, color=RGBColor(0x1a, 0x56, 0xc4))

    add_heading_custom(doc, "6.1 基本用法", level=2)
    add_para(doc, "在 WorkBuddy 对话中，使用以下格式触发 RAG 生成：", font_size=11, space_after=8)
    add_code_block(doc, "用 [分类名] 的资料，帮我写 [文档主题/要求]")
    add_para(doc, "示例：", font_size=11, bold=True, space_after=4)
    add_bullet(doc, "\"用 AI+体育 的资料，帮我写一篇关于柔性穿戴传感器在运动姿态分析中应用的综述\"")
    add_bullet(doc, "\"用 教育政策文件 的资料，帮我整理一份 AI 教育政策要点摘要\"")
    add_bullet(doc, "\"用 AI教学工具实操 和 AI见解与培训 的资料，帮我写一份 AI 工具教学指南\"")
    add_para(doc, "可以指定多个分类，AI 会合并多个分类的资料作为素材。", font_size=10, italic=True, color=RGBColor(0x80, 0x80, 0x80), space_after=8)

    add_heading_custom(doc, "6.2 RAG 模式选择", level=2)
    add_para(doc, "两种模式的工作方式：", font_size=11, space_after=8)

    add_para(doc, "全量拉取模式（full）：", font_size=11, bold=True, space_after=4)
    add_para(doc, "读取该分类下所有资料的 media_id → 逐篇调用 fetch_media_content → 全部全文作为上下文 → 生成文档", font_size=11, space_after=4)
    add_para(doc, "优势：上下文最完整，生成质量最高。劣势：API 调用多，上下文可能超长。", font_size=11, space_after=8)

    add_para(doc, "搜索优先模式（search_first）：", font_size=11, bold=True, space_after=4)
    add_para(doc, "用文档主题作为关键词调用 search_knowledge → 取相关度最高的 Top 5 → 只拉取这 5 篇的全文 → 生成文档", font_size=11, space_after=4)
    add_para(doc, "优势：API 调用少，上下文聚焦。劣势：可能遗漏相关但未被搜索命中的资料。", font_size=11, space_after=8)

    add_para(doc, "自动模式（auto，推荐）：", font_size=11, bold=True, space_after=4)
    add_para(doc, "分类下 ≤5 篇自动用全量拉取，>5 篇自动用搜索优先。兼顾质量和效率。", font_size=11, space_after=8)

    add_heading_custom(doc, "6.3 内容缓存", level=2)
    add_para(doc, "启用缓存后（config.json 中 enable_cache: true），已拉取的全文内容会缓存到本地。", font_size=11, space_after=6)
    add_bullet(doc, "同一篇资料不会重复调用 fetch_media_content")
    add_bullet(doc, "缓存文件保存在 cache_dir 指定目录")
    add_bullet(doc, "如果知识库中的资料被更新，可在对话中让 AI\"清除缓存\"重新拉取")

    add_heading_custom(doc, "6.4 最佳实践", level=2)
    add_bullet(doc, "撰写前先查看索引文档，了解该分类有哪些资料")
    add_bullet(doc, "文档主题尽量具体，帮助搜索优先模式更精准匹配")
    add_bullet(doc, "如果生成结果不理想，可以尝试切换 RAG 模式")
    add_bullet(doc, "多分类组合使用时，注意主题的相关性")
    add_bullet(doc, "生成后可以让 AI 基于反馈修改，不需要重新拉取全文")

    doc.add_page_break()

    # ===== 第七章 定期任务管理 =====
    add_heading_custom(doc, "第七章　定期任务管理", level=1, color=RGBColor(0x1a, 0x56, 0xc4))

    add_heading_custom(doc, "7.1 查看任务", level=2)
    add_para(doc, "在 WorkBuddy 对话中对 AI 说：", font_size=11, space_after=4)
    add_code_block(doc, "查看我的自动化任务")
    add_para(doc, "AI 会列出所有自动化任务，包括名称、频率、状态。", font_size=11, space_after=8)

    add_heading_custom(doc, "7.2 暂停/恢复", level=2)
    add_para(doc, "暂停任务：", font_size=11, space_after=4)
    add_code_block(doc, "暂停 IMA 知识库定期整理任务")
    add_para(doc, "恢复任务：", font_size=11, space_after=4)
    add_code_block(doc, "恢复 IMA 知识库定期整理任务")

    add_heading_custom(doc, "7.3 修改频率", level=2)
    add_para(doc, "在对话中直接告诉 AI 新的频率：", font_size=11, space_after=4)
    add_bullet(doc, "\"把知识库整理任务改为每周一早上 8 点\"")
    add_bullet(doc, "\"把知识库整理任务改为每天扫描一次\"")
    add_para(doc, "AI 会自动更新自动化任务的调度规则。", font_size=11, space_after=8)

    add_heading_custom(doc, "7.4 手动触发", level=2)
    add_para(doc, "不想等定期任务，可以随时手动触发：", font_size=11, space_after=4)
    add_code_block(doc, "现在帮我扫描整理一下 IMA 知识库")
    add_para(doc, "AI 会立即执行完整的扫描分类流程。", font_size=11, space_after=8)

    add_heading_custom(doc, "7.5 删除任务", level=2)
    add_para(doc, "如果不再需要：", font_size=11, space_after=4)
    add_code_block(doc, "删除 IMA 知识库定期整理自动化任务")
    add_para(doc, "注意：删除任务后，已生成的索引文档和追踪文件仍然保留，但不会再自动更新。", font_size=10, italic=True, color=RGBColor(0x80, 0x80, 0x80), space_after=8)

    doc.add_page_break()

    # ===== 第八章 常见问题 =====
    add_heading_custom(doc, "第八章　常见问题", level=1, color=RGBColor(0x1a, 0x56, 0xc4))

    faqs = [
        ("Q: 分类不准确怎么办？",
         "A: 在对话中告知 AI 哪篇资料分错了，AI 会调整分类规则关键词。也可以直接编辑 config.json 中的 categories 部分修改关键词列表，然后手动触发一次扫描重新分类。"),
        ("Q: 新增的资料没有被扫描到？",
         "A: 检查：1) 资料是否添加到了主知识库（is_primary 为 true 的那个）；2) 自动化任务是否处于 ACTIVE 状态；3) 可以手动触发扫描确认。"),
        ("Q: 索引文档没有自动更新？",
         "A: 检查自动化任务是否正常运行。查看 .workbuddy/ima-tracker/reports/ 目录下是否有最新报告。如果脚本执行失败，可能是 python-docx 未安装或路径问题。"),
        ("Q: RAG 生成时提示资料太多？",
         "A: 将 config.json 中的 rag_mode 改为 search_first，AI 会先用关键词搜索筛选最相关的资料，再拉取全文。"),
        ("Q: 可以同时管理多个知识库吗？",
         "A: 可以。在 config.json 的 knowledge_bases 数组中添加多个知识库。标记一个 is_primary 为 true 作为收件箱，其他作为辅助知识库。"),
        ("Q: 缓存太大怎么清理？",
         "A: 在对话中说\"清除 IMA 知识库内容缓存\"，或直接删除 cache_dir 目录下的文件。下次 RAG 生成时会重新拉取。"),
        ("Q: 支持 IMA 之外的知识库吗？",
         "A: 当前版本仅支持 IMA 知识库。如果有其他知识库的 MCP 连接器，可以参照 SKILL.md 中的架构自行适配。"),
        ("Q: 商业使用授权如何申请？",
         "A: 本技能采用 CC BY-NC 4.0 协议，禁止商业使用。如需商业授权，请通过 WorkBuddy 平台联系作者 sus-yugaohe。"),
    ]

    for q, a in faqs:
        add_para(doc, q, font_size=11, bold=True, color=RGBColor(0x1a, 0x56, 0xc4), space_after=4)
        add_para(doc, a, font_size=11, space_after=10)

    doc.add_page_break()

    # ===== 第九章 版权与许可 =====
    add_heading_custom(doc, "第九章　版权与许可", level=1, color=RGBColor(0x1a, 0x56, 0xc4))

    add_para(doc, f"\u00a9 2026 sus-yugaohe 保留所有权利", font_size=14, bold=True, space_after=12)

    add_para(doc, "本技能及其配套文档采用知识共享署名-非商业性使用 4.0 国际许可协议（CC BY-NC 4.0）进行许可。", font_size=12, space_after=12)

    add_heading_custom(doc, "9.1 您可以自由地：", level=2)
    add_bullet(doc, "共享 — 在任何媒介以任何形式复制、发行本作品")
    add_bullet(doc, "改编 — 修改、转换或以本作品为基础进行创作")

    add_heading_custom(doc, "9.2 但须遵守以下条件：", level=2)
    add_bullet(doc, "署名 — 必须注明原作者 sus-yugaohe，并提供许可协议链接")
    add_bullet(doc, "非商业性使用 — 不得将本作品用于商业目的")
    add_bullet(doc, "相同方式共享 — 衍生作品须以相同许可协议发布")

    add_heading_custom(doc, "9.3 禁止行为", level=2)
    add_bullet(doc, "未经授权将本技能用于付费课程、商业培训、销售产品")
    add_bullet(doc, "移除或篡改版权声明和作者信息")
    add_bullet(doc, "以本技能名义进行虚假宣传")

    add_heading_custom(doc, "9.4 免责声明", level=2)
    add_para(doc, "本技能按\"原样\"提供，作者不对技能的适用性、准确性或完整性作任何担保。使用者自行承担使用风险。作者不对使用本技能造成的任何直接或间接损失负责。", font_size=11, space_after=12)

    add_heading_custom(doc, "9.5 商业授权", level=2)
    add_para(doc, "如需商业使用授权，请通过 WorkBuddy 平台联系作者 sus-yugaohe 协商授权条件。", font_size=11, space_after=12)

    add_para(doc, "完整许可协议文本：https://creativecommons.org/licenses/by-nc/4.0/legalcode", font_size=10, italic=True, space_after=12)

    # ===== 末页 =====
    add_para(doc, "", space_after=30)
    add_para(doc, "\u2014\u2014 文档结束 \u2014\u2014", font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x80, 0x80, 0x80))
    add_para(doc, f"\u00a9 2026 sus-yugaohe | CC BY-NC 4.0", font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x80, 0x80, 0x80))

    # 保存
    output_path = os.path.join(OUTPUT_DIR, "IMA知识库整理技能_使用说明.docx")
    doc.save(output_path)
    print(f"文档已生成：{output_path}")
    return output_path


if __name__ == "__main__":
    generate_document()
