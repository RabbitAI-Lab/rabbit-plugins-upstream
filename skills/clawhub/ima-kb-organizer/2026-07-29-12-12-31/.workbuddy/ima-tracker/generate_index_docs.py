#!/usr/bin/env python3
"""
IMA 知识库分类索引文档生成器
读取 tracker.json，为每个有效分类生成一个 Word 索引文档。
可被自动化任务或手动调用。
"""

import json
import sys
import os
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# 配置路径
TRACKER_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "tracker.json"
)
OUTPUT_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "category_docs"
)

# 用户指定目录：生成后自动复制到此目录
USER_DIR = r"D:\tuixiu 20240906备份\tuixiu\tui xiu\教学材料\2026-2027（1）数智素养与工具应用\workbuddy\ima知识库检索文档"

# 需要生成文档的分类（排除"待清理-无关内容"）
VALID_CATEGORIES = [
    "AI+体育",
    "AI+教育研究",
    "AI教学工具实操",
    "教育政策文件",
    "AI见解与培训",
]

# 分类描述映射
CATEGORY_DESC = {
    "AI+体育": "AI在体育领域的应用研究：运动捕捉、穿戴传感器、体测评估、赛事分析、竞技训练、生物力学等",
    "AI+教育研究": "AI在教育领域的学术研究：教育智能体、学习投入度、作业评估、学生画像、教师角色转型、协作学习等",
    "AI教学工具实操": "AI教学工具的实操教程：交互式网页制作、部署上线、数据分析、防作弊、数据回收等",
    "教育政策文件": "教育领域的人工智能相关政策文件、指导意见、实施方案、职称评定等官方文件及解读",
    "AI见解与培训": "AI领域的观点见解、技术分析、教师培训资源、提示词技巧等非学术研究类内容",
}


def set_cell_font(cell, font_name="微软雅黑", font_size=9, bold=False, color=None):
    """设置单元格字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color
            # 设置中文字体
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_font(paragraph, font_name="微软雅黑", font_size=11, bold=False, color=None):
    """设置段落字体"""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def generate_category_doc(category_name, items, kb_name, output_path, scan_time):
    """为单个分类生成 Word 文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # === 标题 ===
    title = doc.add_heading(f"{category_name} - 资料索引", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # === 分类描述 ===
    desc_para = doc.add_paragraph()
    desc_run = desc_para.add_run(f"分类说明：{CATEGORY_DESC.get(category_name, '')}")
    set_paragraph_font(desc_para, font_size=10, color=RGBColor(0x66, 0x66, 0x66))

    # === 元信息 ===
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(
        f"知识库：{kb_name}　|　资料数量：{len(items)} 篇　|　"
        f"最后更新：{scan_time}"
    )
    set_paragraph_font(meta_para, font_size=9, color=RGBColor(0x99, 0x99, 0x99))

    doc.add_paragraph("")  # 空行

    # === 资料列表表格 ===
    if items:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        header_cells = table.rows[0].cells
        headers = ["序号", "资料标题", "来源", "导入知识库日期"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            set_cell_font(header_cells[i], font_size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            # 设置表头背景色
            shading = header_cells[i]._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "4472C4"
            })
            shading.append(shading_elm)

        # 数据行
        for idx, item in enumerate(items, 1):
            row = table.add_row()
            cells = row.cells

            # 序号
            cells[0].text = str(idx)
            set_cell_font(cells[0], font_size=9)

            # 标题
            cells[1].text = item.get("title", "（无标题）")
            set_cell_font(cells[1], font_size=9)

            # 来源
            cells[2].text = item.get("source", "")
            set_cell_font(cells[2], font_size=9)

            # 日期（从 first_seen 提取日期部分）
            first_seen = item.get("first_seen", "")
            if first_seen:
                try:
                    dt = datetime.fromisoformat(first_seen)
                    date_str = dt.strftime("%Y-%m-%d")
                except (ValueError, TypeError):
                    date_str = first_seen[:10] if len(first_seen) >= 10 else first_seen
            else:
                date_str = ""
            cells[3].text = date_str
            set_cell_font(cells[3], font_size=9)

        # 设置列宽
        for row in table.rows:
            row.cells[0].width = Inches(0.5)
            row.cells[1].width = Inches(4.5)
            row.cells[2].width = Inches(1.5)
            row.cells[3].width = Inches(1.2)

    else:
        empty_para = doc.add_paragraph()
        empty_run = empty_para.add_run("（暂无资料）")
        set_paragraph_font(empty_para, font_size=11, color=RGBColor(0x99, 0x99, 0x99))
        empty_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # === 使用说明 ===
    doc.add_paragraph("")
    doc.add_paragraph("")
    usage_heading = doc.add_paragraph()
    usage_run = usage_heading.add_run("使用说明")
    set_paragraph_font(usage_heading, font_size=10, bold=True, color=RGBColor(0x33, 0x33, 0x33))

    usage_text = (
        "1. 撰写文档时，告诉 AI 助手使用本分类的资料作为素材来源。\n"
        "2. AI 会根据本索引中的资料，从 IMA 知识库拉取全文内容，进行 RAG 检索增强生成。\n"
        "3. 新资料请放入 IMA 知识库根目录，AI 会定期扫描并自动分类更新本文档。\n"
        "4. 如发现分类有误，请告知 AI 助手调整。"
    )
    usage_para = doc.add_paragraph()
    usage_para.add_run(usage_text)
    set_paragraph_font(usage_para, font_size=9, color=RGBColor(0x66, 0x66, 0x66))

    # 保存
    doc.save(output_path)


def main():
    # 读取 tracker
    if not os.path.exists(TRACKER_PATH):
        print(f"错误：追踪文件不存在：{TRACKER_PATH}")
        sys.exit(1)

    with open(TRACKER_PATH, "r", encoding="utf-8") as f:
        tracker = json.load(f)

    scan_time = tracker.get("last_scan_time", datetime.now().isoformat())

    # 创建输出目录
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # 汇总所有知识库中各分类的内容
    all_categories = {}
    for kb_id, kb_info in tracker.get("knowledge_bases", {}).items():
        kb_name = kb_info.get("name", "")
        for cat_name, cat_data in kb_info.get("categories", {}).items():
            if cat_name not in VALID_CATEGORIES:
                continue
            if cat_name not in all_categories:
                all_categories[cat_name] = {"items": [], "kb_names": set()}
            all_categories[cat_name]["items"].extend(cat_data.get("items", []))
            all_categories[cat_name]["kb_names"].add(kb_name)

    # 为每个分类生成文档
    generated = []
    for cat_name in VALID_CATEGORIES:
        cat_info = all_categories.get(cat_name, {"items": [], "kb_names": set()})
        items = cat_info["items"]
        kb_names = "、".join(sorted(cat_info["kb_names"])) if cat_info["kb_names"] else "（无）"

        # 文件名中替换特殊字符
        safe_name = cat_name.replace("+", "与").replace(" ", "")
        output_path = os.path.join(OUTPUT_DIR, f"{safe_name}_资料索引.docx")

        generate_category_doc(cat_name, items, kb_names, output_path, scan_time)
        generated.append({
            "category": cat_name,
            "file": output_path,
            "item_count": len(items),
        })
        print(f"已生成：{output_path}（{len(items)} 篇资料）")

    # 生成汇总信息
    summary_path = os.path.join(OUTPUT_DIR, "_生成摘要.json")
    summary = {
        "generated_at": datetime.now().isoformat(),
        "scan_time": scan_time,
        "documents": generated,
        "total_items": sum(g["item_count"] for g in generated),
    }
    with open(summary_path, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)

    print(f"\n生成完毕！共 {len(generated)} 个文档，{summary['total_items']} 篇资料。")
    print(f"摘要文件：{summary_path}")

    # 复制到用户指定目录
    if os.path.isdir(USER_DIR):
        copied = 0
        for g in generated:
            src_file = g["file"]
            dst_file = os.path.join(USER_DIR, os.path.basename(src_file))
            try:
                shutil.copy2(src_file, dst_file)
                copied += 1
            except Exception as e:
                print(f"警告：复制失败 {os.path.basename(src_file)}：{e}")
        print(f"已复制 {copied}/{len(generated)} 个文档到用户目录：{USER_DIR}")
    else:
        print(f"警告：用户目录不存在，跳过复制：{USER_DIR}")


if __name__ == "__main__":
    main()
