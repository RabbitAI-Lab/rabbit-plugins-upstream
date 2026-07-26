# -*- coding: utf-8 -*-
"""
公文格式Word文档生成脚本
按 GB/T 9704-2012《党政机关公文格式》标准生成Word文档

用法:
  python generate_gongwen.py <config.json>

config.json 结构:
{
  "output_path": "输出路径.docx",
  "title": "公文标题",
  "content": [
    {"type": "body", "text": "正文段落"},
    {"type": "h1", "text": "一、一级标题"},
    {"type": "h2", "text": "（一）二级标题"},
    {"type": "h3", "text": "1.三级标题"},
    {"type": "blank", "text": ""}
  ]
}
"""
import sys
import json
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


# ============ 字体配置 ============
# 系统自带字体
FANGSONG = '仿宋'        # 正文：3号仿宋
HEITI = '黑体'           # 一级标题：3号黑体
KAITI = '楷体'           # 二级标题：3号楷体
# 公文标题标准字体为"方正小标宋简体"，需安装该字体
# 如未安装，Word会自动用默认字体渲染；安装后即严格对标GB/T 9704
TITLE_FONT = '方正小标宋简体'  # 大标题：2号小标宋

# 字号换算：磅值
# 初号=42pt, 小初=36pt, 一号=26pt, 小一=24pt, 二号=22pt, 小二=18pt
# 三号=16pt, 小三=15pt, 四号=14pt, 小四=12pt, 五号=10.5pt
SIZE_TITLE = 22   # 2号
SIZE_BODY = 16    # 3号
LINE_SPACING = 29 # 固定行距29pt（每面约22行，每行约28字）


def set_run_font(run, font_name, size_pt, bold=False, color=(0, 0, 0)):
    """设置run的字体、字号、粗体、颜色，含东亚字体设置"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def set_paragraph_format(para, indent_chars=0, line_spacing_pt=LINE_SPACING,
                         space_after=0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """设置段落格式：首行缩进、固定行距、段后距、对齐方式"""
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.line_spacing = Pt(line_spacing_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    # 首行缩进：3号字16pt，2字符=32pt
    pf.first_line_indent = Pt(SIZE_BODY * indent_chars) if indent_chars > 0 else Pt(0)
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)


def add_title(doc, text):
    """公文标题：2号方正小标宋简体，居中"""
    para = doc.add_paragraph()
    set_paragraph_format(para, indent_chars=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = para.add_run(text)
    set_run_font(run, TITLE_FONT, SIZE_TITLE, bold=False)
    return para


def add_h1(doc, text):
    """一级标题：3号黑体，首行缩进2字。对应"一、""二、""三、"""
    para = doc.add_paragraph()
    set_paragraph_format(para, indent_chars=2)
    run = para.add_run(text)
    set_run_font(run, HEITI, SIZE_BODY)
    return para


def add_h2(doc, text):
    """二级标题：3号楷体，首行缩进2字。对应"（一）""（二）""（三）"""
    para = doc.add_paragraph()
    set_paragraph_format(para, indent_chars=2)
    run = para.add_run(text)
    set_run_font(run, KAITI, SIZE_BODY)
    return para


def add_h3(doc, text):
    """三级标题：3号仿宋，首行缩进2字。对应"1.""2.""3."
    注意：四级标题"（1）（2）"同样用仿宋，调用本函数即可"""
    para = doc.add_paragraph()
    set_paragraph_format(para, indent_chars=2)
    run = para.add_run(text)
    set_run_font(run, FANGSONG, SIZE_BODY)
    return para


def add_body(doc, text):
    """正文段落：3号仿宋，首行缩进2字，回行顶格"""
    para = doc.add_paragraph()
    set_paragraph_format(para, indent_chars=2)
    run = para.add_run(text)
    set_run_font(run, FANGSONG, SIZE_BODY)
    return para


def add_blank_line(doc):
    """空行"""
    para = doc.add_paragraph()
    set_paragraph_format(para, indent_chars=0)
    run = para.add_run('')
    set_run_font(run, FANGSONG, SIZE_BODY)
    return para


# 内容类型到处理函数的映射
TYPE_HANDLERS = {
    'title': add_title,
    'h1': add_h1,
    'h2': add_h2,
    'h3': add_h3,
    'body': add_body,
    'blank': lambda doc, text: add_blank_line(doc),
}


def generate_document(config):
    """根据JSON配置生成公文格式Word文档"""
    doc = Document()

    # 页面设置：A4，上37mm，左28mm，版心156×225mm
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)      # 天头（上白边）37mm±1mm
    section.bottom_margin = Mm(35)   # 297-37-225=35
    section.left_margin = Mm(28)     # 订口（左白边）28mm±1mm
    section.right_margin = Mm(26)     # 210-28-156=26

    # 添加标题
    if 'title' in config and config['title']:
        add_title(doc, config['title'])
        add_blank_line(doc)

    # 添加正文内容
    for item in config.get('content', []):
        item_type = item.get('type', 'body')
        item_text = item.get('text', '')
        handler = TYPE_HANDLERS.get(item_type, add_body)
        handler(doc, item_text)

    # 保存
    output_path = config['output_path']
    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) < 2:
        print("用法: python generate_gongwen.py <config.json>")
        sys.exit(1)

    config_path = sys.argv[1]
    with open(config_path, 'r', encoding='utf-8') as f:
        config = json.load(f)

    output_path = generate_document(config)
    print(f"文档已保存：{output_path}")


if __name__ == '__main__':
    main()
